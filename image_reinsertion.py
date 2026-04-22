"""DOCX image reinsertion and variant resolution.

Replaces image placeholders in generated DOCX files with actual image bytes,
handles synthetic multi-variant image blocks, and resolves final image selection.
"""

import logging
import re
from dataclasses import dataclass
from copy import deepcopy
from io import BytesIO
from typing import cast

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import lxml.etree as etree

from document import (
    IMAGE_PLACEHOLDER_PATTERN,
    _extract_run_text,
    _find_child_element,
    _xml_local_name,
)
from logger import log_event
from models import ImageAsset, ImageDeliveryPayload, get_image_variant_bytes


def resolve_final_image_bytes(asset: ImageAsset) -> bytes:
    delivery_payload = _resolve_delivery_payload(asset)
    if isinstance(delivery_payload.final_bytes, (bytes, bytearray)):
        return bytes(delivery_payload.final_bytes)

    if asset.selected_compare_variant:
        if asset.selected_compare_variant == "original":
            return asset.original_bytes
        selected_variant = asset.comparison_variants.get(asset.selected_compare_variant)
        selected_bytes = get_image_variant_bytes(selected_variant)
        if selected_bytes:
            return selected_bytes
    if asset.final_variant == "redrawn" and asset.redrawn_bytes:
        return asset.redrawn_bytes
    if asset.final_variant == "safe" and asset.safe_bytes:
        return asset.safe_bytes
    return asset.original_bytes


def resolve_image_insertions(asset: ImageAsset) -> list[tuple[str | None, bytes]]:
    delivery_payload = _resolve_delivery_payload(asset)
    if delivery_payload.insertions:
        return [(insertion.label, bytes(insertion.bytes)) for insertion in delivery_payload.insertions]

    final_bytes = resolve_final_image_bytes(asset)
    if not final_bytes:
        return []
    return [(None, final_bytes)]


def _resolve_delivery_payload(asset: ImageAsset) -> ImageDeliveryPayload:
    return asset.resolved_delivery_payload()


def reinsert_inline_images(docx_bytes: bytes, image_assets: list[ImageAsset]) -> bytes:
    if not docx_bytes or not image_assets:
        return docx_bytes

    source_stream = BytesIO(docx_bytes)
    document = Document(source_stream)
    asset_map = {asset.placeholder: asset for asset in image_assets}
    insertion_cache: dict[str, list[tuple[str | None, bytes]]] = {}

    for paragraph in _iter_reinsertion_paragraphs(document):
        paragraph_text = paragraph.text
        placeholders = _find_known_placeholders(paragraph_text, asset_map)
        if not placeholders:
            continue

        if _replace_multi_variant_placeholders_with_blocks(paragraph, asset_map, insertion_cache):
            continue

        if _replace_run_level_placeholders(paragraph, placeholders, asset_map, insertion_cache):
            continue

        if _replace_multi_run_placeholders(paragraph, asset_map, insertion_cache):
            continue

        if _replace_paragraph_placeholders_fallback(paragraph, paragraph_text, asset_map, insertion_cache):
            continue

        log_event(
            logging.WARNING,
            "image_reinsertion_placeholder_unhandled",
            "Не удалось безопасно заменить image placeholder ни одной стратегией reinsertion; placeholder оставлен как текст.",
            placeholder_count=len(placeholders),
            placeholders=placeholders,
            paragraph_text_preview=_paragraph_preview(paragraph_text),
        )

    output_stream = BytesIO()
    document.save(output_stream)
    return output_stream.getvalue()


# ---------------------------------------------------------------------------
# Paragraph iteration helpers
# ---------------------------------------------------------------------------


def _iter_reinsertion_paragraphs(document):
    visited_paragraph_elements: set[object] = set()

    yield from _iter_container_paragraphs(
        document,
        _visited_cell_elements=set(),
        _visited_paragraph_elements=visited_paragraph_elements,
    )

    for story_container in _iter_section_story_containers(document):
        yield from _iter_container_paragraphs(
            story_container,
            _visited_cell_elements=set(),
            _visited_paragraph_elements=visited_paragraph_elements,
        )


def _iter_section_story_containers(document):
    for section in document.sections:
        for attribute_name in (
            "header",
            "first_page_header",
            "even_page_header",
            "footer",
            "first_page_footer",
            "even_page_footer",
        ):
            story_container = getattr(section, attribute_name, None)
            if story_container is None:
                continue
            yield story_container


def _iter_container_paragraphs(
    container,
    *,
    _visited_cell_elements: set[object] | None = None,
    _visited_paragraph_elements: set[object] | None = None,
):
    if _visited_cell_elements is None:
        _visited_cell_elements = set()
    if _visited_paragraph_elements is None:
        _visited_paragraph_elements = set()

    for paragraph in getattr(container, "paragraphs", []):
        paragraph_element = paragraph._element
        if paragraph_element not in _visited_paragraph_elements:
            _visited_paragraph_elements.add(paragraph_element)
            yield paragraph

        yield from _iter_textbox_paragraphs(paragraph, _visited_paragraph_elements)

    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                cell_element = cell._tc
                if cell_element in _visited_cell_elements:
                    continue
                _visited_cell_elements.add(cell_element)
                yield from _iter_container_paragraphs(
                    cell,
                    _visited_cell_elements=_visited_cell_elements,
                    _visited_paragraph_elements=_visited_paragraph_elements,
                )


def _iter_textbox_paragraphs(paragraph, visited_paragraph_elements: set[object]):
    for textbox_paragraph_element in paragraph._element.xpath(".//w:txbxContent//w:p"):
        if textbox_paragraph_element in visited_paragraph_elements:
            continue

        visited_paragraph_elements.add(textbox_paragraph_element)
        yield Paragraph(textbox_paragraph_element, paragraph._parent)


# ---------------------------------------------------------------------------
# Placeholder replacement strategies
# ---------------------------------------------------------------------------


def _find_known_placeholders(text: str, asset_map: dict[str, ImageAsset]) -> list[str]:
    return [token for token in IMAGE_PLACEHOLDER_PATTERN.findall(text) if token in asset_map]


def _paragraph_preview(text: str, *, limit: int = 120) -> str:
    normalized = re.sub(r"\s+", " ", text).strip()
    if len(normalized) <= limit:
        return normalized
    return normalized[: limit - 3].rstrip() + "..."


def _resolve_cached_insertions(
    asset: ImageAsset,
    insertion_cache: dict[str, list[tuple[str | None, bytes]]],
) -> list[tuple[str | None, bytes]]:
    placeholder = asset.placeholder.strip()
    if placeholder in insertion_cache:
        return insertion_cache[placeholder]

    insertions = resolve_image_insertions(asset)
    insertion_cache[placeholder] = insertions
    return insertions


@dataclass(frozen=True)
class _MultiVariantBlockFragment:
    asset: ImageAsset
    insertions: list[tuple[str | None, bytes]]


def _replace_run_level_placeholders(
    paragraph,
    placeholders: list[str],
    asset_map: dict[str, ImageAsset],
    insertion_cache: dict[str, list[tuple[str | None, bytes]]],
) -> bool:
    runs = list(paragraph.runs)
    run_placeholders: list[str] = []
    for run in runs:
        run_placeholders.extend(_find_known_placeholders(run.text, asset_map))

    if len(run_placeholders) != len(placeholders):
        return False

    for run in runs:
        run_text = run.text
        run_tokens = _find_known_placeholders(run_text, asset_map)
        if not run_tokens:
            continue

        replacement_elements = _build_run_replacement_elements(
            paragraph,
            run._element,
            run_text,
            asset_map,
            insertion_cache,
        )
        _replace_xml_element_with_sequence(run._element, replacement_elements)
    return True


def _replace_multi_run_placeholders(
    paragraph,
    asset_map: dict[str, ImageAsset],
    insertion_cache: dict[str, list[tuple[str | None, bytes]]],
) -> bool:
    paragraph_children = [child for child in list(paragraph._element) if _xml_local_name(child.tag) in {"r", "hyperlink"}]
    if not paragraph_children:
        return False

    child_texts = [_extract_paragraph_child_text(child) for child in paragraph_children]
    full_text = "".join(child_texts)
    placeholder_matches = [
        match
        for match in IMAGE_PLACEHOLDER_PATTERN.finditer(full_text)
        if match.group(0) in asset_map
    ]
    if not placeholder_matches:
        return False

    replacement_elements: list[etree._Element] = []
    child_ranges: list[tuple[etree._Element, str, int, int]] = []
    cursor = 0
    for child, child_text in zip(paragraph_children, child_texts):
        next_cursor = cursor + len(child_text)
        child_ranges.append((child, child_text, cursor, next_cursor))
        cursor = next_cursor

    match_index = 0
    current_match = placeholder_matches[match_index] if placeholder_matches else None

    for child, child_text, child_start, child_end in child_ranges:
        position = child_start
        while position < child_end:
            while current_match is not None and current_match.end() <= position:
                match_index += 1
                current_match = placeholder_matches[match_index] if match_index < len(placeholder_matches) else None

            if current_match is not None and position >= current_match.start() and position < current_match.end():
                if position == current_match.start():
                    placeholder_text = current_match.group(0)
                    if _xml_local_name(child.tag) == "hyperlink":
                        return False
                    replacement_elements.extend(
                        _build_insertion_run_elements(
                            paragraph,
                            child,
                            asset_map[placeholder_text],
                            insertion_cache,
                            placeholder_text=placeholder_text,
                        )
                    )
                position = min(child_end, current_match.end())
                continue

            segment_end = child_end
            if current_match is not None and position < current_match.start():
                segment_end = min(segment_end, current_match.start())

            if segment_end > position:
                if _xml_local_name(child.tag) == "hyperlink":
                    if position != child_start or segment_end != child_end:
                        return False
                    replacement_elements.append(deepcopy(child))
                else:
                    replacement_elements.append(
                        _build_text_run_element(paragraph, child, full_text[position:segment_end])
                    )
            position = segment_end

    if not replacement_elements:
        return False

    _clear_paragraph_runs(paragraph)
    for element in replacement_elements:
        paragraph._element.append(element)
    return True


def _replace_paragraph_placeholders_fallback(
    paragraph,
    paragraph_text: str,
    asset_map: dict[str, ImageAsset],
    insertion_cache: dict[str, list[tuple[str | None, bytes]]],
) -> bool:
    if any(_xml_local_name(child.tag) == "hyperlink" for child in list(paragraph._element)):
        return False

    parts = re.split(f"({IMAGE_PLACEHOLDER_PATTERN.pattern})", paragraph_text)
    _clear_paragraph_runs(paragraph)
    for part in parts:
        if not part:
            continue
        asset = asset_map.get(part)
        if asset is None:
            paragraph.add_run(part)
            continue
        _append_image_insertions_to_paragraph(
            paragraph,
            asset,
            insertion_cache,
            placeholder_text=part,
        )
    return True


def _replace_multi_variant_placeholders_with_blocks(
    paragraph,
    asset_map: dict[str, ImageAsset],
    insertion_cache: dict[str, list[tuple[str | None, bytes]]],
) -> bool:
    paragraph_children = [child for child in list(paragraph._element) if _xml_local_name(child.tag) in {"r", "hyperlink"}]
    if not paragraph_children:
        return False

    child_texts = [_extract_paragraph_child_text(child) for child in paragraph_children]
    full_text = "".join(child_texts)
    placeholder_matches = [
        match
        for match in IMAGE_PLACEHOLDER_PATTERN.finditer(full_text)
        if match.group(0) in asset_map
    ]
    if not placeholder_matches:
        return False

    multi_variant_placeholders = [
        match.group(0)
        for match in placeholder_matches
        if len(_resolve_cached_insertions(asset_map[match.group(0)], insertion_cache)) > 1
    ]
    if not multi_variant_placeholders:
        return False

    fragments: list[etree._Element | _MultiVariantBlockFragment] = []
    child_ranges: list[tuple[etree._Element, str, int, int]] = []
    cursor = 0
    for child, child_text in zip(paragraph_children, child_texts):
        next_cursor = cursor + len(child_text)
        child_ranges.append((child, child_text, cursor, next_cursor))
        cursor = next_cursor

    match_index = 0
    current_match = placeholder_matches[match_index]

    for child, child_text, child_start, child_end in child_ranges:
        position = child_start
        while position < child_end:
            while current_match is not None and current_match.end() <= position:
                match_index += 1
                current_match = placeholder_matches[match_index] if match_index < len(placeholder_matches) else None

            if current_match is not None and position >= current_match.start() and position < current_match.end():
                if position == current_match.start():
                    placeholder_text = current_match.group(0)
                    asset = asset_map[placeholder_text]
                    insertions = _resolve_cached_insertions(asset, insertion_cache)
                    if len(insertions) > 1:
                        if _xml_local_name(child.tag) != "r":
                            _log_multi_variant_block_warning(
                                "image_reinsertion_multi_variant_block_fallback_to_text",
                                paragraph,
                                [placeholder_text],
                                reason="multi_variant_placeholder_inside_hyperlink_or_non_run_child",
                            )
                            fragments.append(deepcopy(child))
                            position = min(child_end, current_match.end())
                            continue
                        fragments.append(_MultiVariantBlockFragment(asset=asset, insertions=insertions))
                    elif _xml_local_name(child.tag) == "r":
                        fragments.extend(
                            _build_insertion_run_elements(
                                paragraph,
                                child,
                                asset,
                                insertion_cache,
                                placeholder_text=placeholder_text,
                            )
                        )
                    else:
                        fragments.append(deepcopy(child))
                position = min(child_end, current_match.end())
                continue

            segment_end = child_end
            if current_match is not None and position < current_match.start():
                segment_end = min(segment_end, current_match.start())

            if segment_end > position:
                if _xml_local_name(child.tag) == "hyperlink":
                    if position != child_start or segment_end != child_end:
                        return False
                    fragments.append(deepcopy(child))
                else:
                    fragments.append(_build_text_run_element(paragraph, child, full_text[position:segment_end]))
            position = segment_end

    replacement_blocks = _build_replacement_blocks_from_fragments(paragraph, fragments)
    if not replacement_blocks:
        _log_multi_variant_block_warning(
            "image_reinsertion_multi_variant_block_unresolved",
            paragraph,
            multi_variant_placeholders,
            reason="multi_variant_block_builder_returned_no_output",
        )
        return False

    anchor = cast(etree._Element, paragraph._element)
    for block in replacement_blocks:
        anchor.addnext(block)
        anchor = block
    paragraph._element.getparent().remove(paragraph._element)
    return True


# ---------------------------------------------------------------------------
# Run / element building helpers
# ---------------------------------------------------------------------------


def _build_run_replacement_elements(
    paragraph,
    template_run_element,
    run_text: str,
    asset_map: dict[str, ImageAsset],
    insertion_cache: dict[str, list[tuple[str | None, bytes]]],
):
    replacement_elements = []
    parts = re.split(f"({IMAGE_PLACEHOLDER_PATTERN.pattern})", run_text)
    for part in parts:
        if not part:
            continue
        asset = asset_map.get(part)
        if asset is None:
            replacement_elements.append(_build_text_run_element(paragraph, template_run_element, part))
            continue
        replacement_elements.extend(
            _build_insertion_run_elements(
                paragraph,
                template_run_element,
                asset,
                insertion_cache,
                placeholder_text=part,
            )
        )
    return replacement_elements


def _append_image_insertions_to_paragraph(
    paragraph,
    asset: ImageAsset,
    insertion_cache: dict[str, list[tuple[str | None, bytes]]],
    *,
    placeholder_text: str,
) -> None:
    insertions = _resolve_cached_insertions(asset, insertion_cache)
    if not insertions:
        paragraph.add_run(placeholder_text)
        return

    if len(insertions) > 1:
        _log_multi_variant_block_warning(
            "image_reinsertion_multi_variant_block_fallback_to_text",
            paragraph,
            [placeholder_text],
            reason="single_paragraph_fallback_supports_only_one_image",
        )
        paragraph.add_run(placeholder_text)
        return

    add_picture_kwargs = _build_picture_size_kwargs(asset)
    run = paragraph.add_run()
    run.add_picture(BytesIO(insertions[0][1]), **add_picture_kwargs)
    if insertions[0][0]:
        _set_picture_description(run._element, insertions[0][0])
    _apply_source_geometry_metadata(run._element, asset, preferred_description=insertions[0][0])


def _build_insertion_run_elements(
    paragraph,
    template_run_element,
    asset: ImageAsset,
    insertion_cache: dict[str, list[tuple[str | None, bytes]]],
    *,
    placeholder_text: str,
):
    insertions = _resolve_cached_insertions(asset, insertion_cache)
    if not insertions:
        return [_build_text_run_element(paragraph, template_run_element, placeholder_text)]

    if len(insertions) > 1:
        _log_multi_variant_block_warning(
            "image_reinsertion_multi_variant_block_fallback_to_text",
            paragraph,
            [placeholder_text],
            reason="run_level_replacement_supports_only_one_image",
        )
        return [_build_text_run_element(paragraph, template_run_element, placeholder_text)]

    add_picture_kwargs = _build_picture_size_kwargs(asset)
    return [
        _build_picture_run_element(
            paragraph,
            template_run_element,
            asset,
            insertions[0][1],
            add_picture_kwargs,
            description=insertions[0][0],
        )
    ]


def _build_text_run_element(paragraph, template_run_element, text: str):
    run_element = OxmlElement("w:r")
    _copy_run_properties(template_run_element, run_element)
    Run(run_element, paragraph).text = text.replace("<br/>", "\n")
    return run_element


def _build_picture_run_element(
    paragraph,
    template_run_element,
    asset: ImageAsset,
    image_bytes: bytes,
    add_picture_kwargs: dict[str, Emu],
    *,
    description: str | None = None,
):
    run_element = OxmlElement("w:r")
    _copy_run_properties(template_run_element, run_element)
    Run(run_element, paragraph).add_picture(BytesIO(image_bytes), **add_picture_kwargs)
    if description:
        _set_picture_description(run_element, description)
    _apply_source_geometry_metadata(run_element, asset, preferred_description=description)
    return run_element


def _copy_run_properties(template_run_element, target_run_element) -> None:
    run_properties = _find_child_element(template_run_element, "rPr")
    if run_properties is not None:
        target_run_element.append(deepcopy(run_properties))


def _replace_xml_element_with_sequence(element, replacements) -> None:
    if not replacements:
        return
    parent = element.getparent()
    if parent is None:
        return

    anchor = element
    for replacement in replacements:
        anchor.addnext(replacement)
        anchor = replacement
    parent.remove(element)


# ---------------------------------------------------------------------------
# Multi-variant block building
# ---------------------------------------------------------------------------


def _build_replacement_blocks_from_fragments(
    paragraph,
    fragments: list[etree._Element | _MultiVariantBlockFragment],
) -> list[etree._Element]:
    replacement_blocks: list[etree._Element] = []
    current_paragraph = None

    def flush_current_paragraph() -> None:
        nonlocal current_paragraph
        if current_paragraph is not None and _paragraph_element_has_content(current_paragraph):
            replacement_blocks.append(current_paragraph)
        current_paragraph = None

    for fragment in fragments:
        if isinstance(fragment, _MultiVariantBlockFragment):
            flush_current_paragraph()
            replacement_blocks.extend(_build_variant_block_elements(paragraph, fragment.asset, insertions=fragment.insertions))
            continue

        if current_paragraph is None:
            current_paragraph = _clone_paragraph_element(paragraph)
        current_paragraph.append(fragment)

    flush_current_paragraph()
    return replacement_blocks


def _clone_paragraph_element(paragraph):
    paragraph_element = OxmlElement("w:p")
    paragraph_properties = _find_child_element(paragraph._element, "pPr")
    if paragraph_properties is not None:
        paragraph_element.append(deepcopy(paragraph_properties))
    return paragraph_element


def _paragraph_element_has_content(paragraph_element) -> bool:
    return any(_xml_local_name(child.tag) != "pPr" for child in paragraph_element)


def _extract_paragraph_child_text(child) -> str:
    if _xml_local_name(child.tag) == "hyperlink":
        return "".join(_extract_run_text(run_element) for run_element in child.xpath("./w:r"))
    return _extract_run_text(child)


def _build_variant_block_elements(
    paragraph,
    asset: ImageAsset,
    *,
    insertions: list[tuple[str | None, bytes]] | None = None,
) -> list[etree._Element]:
    insertions = resolve_image_insertions(asset) if insertions is None else insertions
    if not insertions:
        return []

    add_picture_kwargs = _build_picture_size_kwargs(asset)
    blocks: list[etree._Element] = []
    for label, image_bytes in insertions:
        image_paragraph = _build_synthetic_image_block_paragraph()

        image_run = OxmlElement("w:r")
        Run(image_run, paragraph).add_picture(BytesIO(image_bytes), **add_picture_kwargs)
        if label:
            _set_picture_description(image_run, label)
        _apply_source_geometry_metadata(image_run, asset, preferred_description=label)

        image_paragraph.append(image_run)
        blocks.append(image_paragraph)
    return blocks


def _build_synthetic_image_block_paragraph():
    paragraph_element = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")
    alignment = OxmlElement("w:jc")
    alignment.set(qn("w:val"), "center")
    paragraph_properties.append(alignment)
    paragraph_element.append(paragraph_properties)
    return paragraph_element


def _set_picture_description(run_element, description: str) -> None:
    if not description:
        return

    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    wordprocessing_drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

    for drawing in run_element.iterchildren(tag=f"{{{word_ns}}}drawing"):
        for container_local_name in ("inline", "anchor"):
            container = drawing.find(f"{{{wordprocessing_drawing_ns}}}{container_local_name}")
            if container is None:
                continue
            doc_properties = container.find(f"{{{wordprocessing_drawing_ns}}}docPr")
            if doc_properties is not None:
                doc_properties.set("descr", description)
                return


def _apply_source_geometry_metadata(
    run_element,
    asset: ImageAsset | None,
    *,
    source_forensics: dict[str, object] | None = None,
    preferred_description: str | None = None,
) -> None:
    resolved_forensics = source_forensics
    if resolved_forensics is None and asset is not None:
        resolved_forensics = asset.source_forensics
    if not isinstance(resolved_forensics, dict) or not resolved_forensics:
        return

    _apply_source_drawing_container(run_element, resolved_forensics)
    _apply_picture_source_rect(run_element, resolved_forensics.get("source_rect"))
    _apply_picture_doc_properties(
        run_element,
        resolved_forensics.get("doc_properties"),
        preferred_description=preferred_description,
    )


def _apply_source_drawing_container(run_element, source_forensics: dict[str, object]) -> None:
    if source_forensics.get("drawing_container") != "anchor":
        return

    container_xml = source_forensics.get("drawing_container_xml")
    if not isinstance(container_xml, str) or not container_xml.strip():
        return

    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    wordprocessing_drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    drawingml_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"

    drawings = list(run_element.iterchildren(tag=f"{{{word_ns}}}drawing"))
    if not drawings:
        return
    drawing = drawings[0]

    current_container = None
    for container_local_name in ("inline", "anchor"):
        current_container = drawing.find(f"{{{wordprocessing_drawing_ns}}}{container_local_name}")
        if current_container is not None:
            break
    if current_container is None:
        return

    try:
        source_container = etree.fromstring(container_xml.encode("utf-8"))
    except etree.XMLSyntaxError:
        return

    if _is_page_relative_anchor_container(source_container):
        return

    current_graphic = current_container.find(f"{{{drawingml_ns}}}graphic")
    source_graphic = source_container.find(f"{{{drawingml_ns}}}graphic")
    if current_graphic is not None:
        if source_graphic is not None:
            source_graphic.getparent().replace(source_graphic, deepcopy(current_graphic))
        else:
            source_container.append(deepcopy(current_graphic))

    current_extent = current_container.find(f"{{{wordprocessing_drawing_ns}}}extent")
    source_extent = source_container.find(f"{{{wordprocessing_drawing_ns}}}extent")
    if current_extent is not None:
        if source_extent is None:
            source_container.insert(0, deepcopy(current_extent))
        else:
            for attribute_name in ("cx", "cy"):
                attribute_value = current_extent.get(attribute_name)
                if attribute_value:
                    source_extent.set(attribute_name, attribute_value)

    current_doc_properties = current_container.find(f"{{{wordprocessing_drawing_ns}}}docPr")
    source_doc_properties = source_container.find(f"{{{wordprocessing_drawing_ns}}}docPr")
    if current_doc_properties is not None and source_doc_properties is not None:
        current_id = current_doc_properties.get("id")
        if current_id:
            source_doc_properties.set("id", current_id)

    drawing.replace(current_container, source_container)


def _is_page_relative_anchor_container(source_container) -> bool:
    wordprocessing_drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    for position_local_name in ("positionH", "positionV"):
        position_element = source_container.find(f"{{{wordprocessing_drawing_ns}}}{position_local_name}")
        if position_element is not None and position_element.get("relativeFrom") == "page":
            return True
    return False


def _apply_picture_source_rect(run_element, source_rect_payload: object) -> None:
    source_rect = _normalize_source_rect(source_rect_payload)
    if source_rect is None:
        return

    drawingml_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    blip_fills = run_element.xpath(".//pic:blipFill")
    if not blip_fills:
        return

    blip_fill = blip_fills[0]
    existing_source_rects = blip_fill.xpath("./a:srcRect")
    if existing_source_rects:
        source_rect_element = existing_source_rects[0]
    else:
        source_rect_element = OxmlElement("a:srcRect")
        blip = blip_fill.find(f"{{{drawingml_ns}}}blip")
        if blip is not None:
            blip.addnext(source_rect_element)
        else:
            blip_fill.insert(0, source_rect_element)

    for key, value in source_rect.items():
        source_rect_element.set(key, str(value))


def _apply_picture_doc_properties(
    run_element,
    doc_properties_payload: object,
    *,
    preferred_description: str | None = None,
) -> None:
    if not isinstance(doc_properties_payload, dict):
        if preferred_description:
            _set_picture_description(run_element, preferred_description)
        return

    wordprocessing_drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    description_value = preferred_description or doc_properties_payload.get("descr")
    for drawing in run_element.iterchildren(tag=f"{{{word_ns}}}drawing"):
        for container_local_name in ("inline", "anchor"):
            container = drawing.find(f"{{{wordprocessing_drawing_ns}}}{container_local_name}")
            if container is None:
                continue
            doc_properties = container.find(f"{{{wordprocessing_drawing_ns}}}docPr")
            if doc_properties is None:
                continue
            if description_value:
                doc_properties.set("descr", str(description_value))
            title_value = doc_properties_payload.get("title")
            if title_value:
                doc_properties.set("title", str(title_value))
            name_value = doc_properties_payload.get("name")
            if name_value:
                doc_properties.set("name", str(name_value))
            return


def _normalize_source_rect(source_rect_payload: object) -> dict[str, int] | None:
    if not isinstance(source_rect_payload, dict):
        return None
    normalized: dict[str, int] = {}
    for key in ("l", "t", "r", "b"):
        raw_value = source_rect_payload.get(key)
        if raw_value is None:
            continue
        try:
            normalized[key] = int(raw_value)
        except (TypeError, ValueError):
            continue
    return normalized or None


def _log_multi_variant_block_warning(event_name: str, paragraph, placeholders: list[str], *, reason: str) -> None:
    log_event(
        logging.WARNING,
        event_name,
        "Multi-variant image placeholder не удалось безопасно развернуть в synthetic blocks; placeholder оставлен как текст.",
        placeholder_count=len(placeholders),
        placeholders=placeholders,
        paragraph_text_preview=_paragraph_preview(paragraph.text),
        reason=reason,
    )


# ---------------------------------------------------------------------------
# Shared low-level helpers
# ---------------------------------------------------------------------------


def _clear_paragraph_runs(paragraph) -> None:
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if _xml_local_name(child.tag) in {"r", "hyperlink"}:
            paragraph_element.remove(child)


def _build_picture_size_kwargs(asset: ImageAsset) -> dict[str, Emu]:
    size_kwargs: dict[str, Emu] = {}
    if isinstance(asset.width_emu, int) and asset.width_emu > 0:
        size_kwargs["width"] = Emu(asset.width_emu)
    if isinstance(asset.height_emu, int) and asset.height_emu > 0:
        size_kwargs["height"] = Emu(asset.height_emu)
    return size_kwargs
