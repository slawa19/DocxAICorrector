from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import Callable
import zipfile

from docx import Document
from docxaicorrector.document.extraction import (
    extract_document_content_from_docx,
    validate_docx_source_bytes,
)
from docxaicorrector.processing.processing_runtime import (
    _convert_pdf_text_layer_to_docx,
    _convert_pdf_to_docx,
    build_in_memory_uploaded_file,
)
from docxaicorrector.pdf_import.text_layer_quality import (
    build_text_layer_quality_report,
    extract_pdf_text_spans_with_pdfminer,
    unsupported_quality_report,
)


_PAGE_NUMBER_TEXT_PATTERN = re.compile(r"^(?:\d{1,4}|[ivxlcdmIVXLCDM]{1,12})$")


def build_pdf_import_backend_comparison(
    *,
    input_pdf: Path,
    converters: dict[str, Callable[[str, bytes], tuple[bytes, str]]] | None = None,
) -> dict[str, object]:
    source_bytes = input_pdf.read_bytes()
    converter_map = converters or {
        "libreoffice": lambda filename, content: _convert_pdf_to_docx(
            filename=filename,
            source_bytes=content,
        ),
        "pdf_text_layer": lambda filename, content: _convert_pdf_text_layer_to_docx(
            filename=filename,
            source_bytes=content,
        ),
    }
    return {
        "input_pdf": str(input_pdf),
        "source_bytes": len(source_bytes),
        "text_layer_quality": _build_text_layer_quality(input_pdf),
        "pdf_image_objects": _count_pdf_image_objects(input_pdf),
        "backends": {
            name: _run_backend(name=name, filename=input_pdf.name, source_bytes=source_bytes, converter=converter)
            for name, converter in converter_map.items()
        },
    }


def _build_text_layer_quality(input_pdf: Path) -> dict[str, object]:
    try:
        spans = extract_pdf_text_spans_with_pdfminer(input_pdf)
    except Exception as exc:
        return unsupported_quality_report(str(exc)).to_dict()
    return build_text_layer_quality_report(spans).to_dict()


def _count_pdf_image_objects(input_pdf: Path) -> dict[str, object]:
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTContainer, LTFigure, LTImage
    except ImportError as exc:
        return {"status": "unsupported", "reason": str(exc), "lt_image_count": 0, "lt_figure_count": 0}
    try:
        lt_image_count = 0
        lt_figure_count = 0

        def walk(item) -> None:
            nonlocal lt_image_count, lt_figure_count
            if isinstance(item, LTImage):
                lt_image_count += 1
            if isinstance(item, LTFigure):
                lt_figure_count += 1
            if isinstance(item, LTContainer):
                for child in item:
                    walk(child)

        for page in extract_pages(str(input_pdf)):
            walk(page)
    except Exception as exc:
        return {"status": "error", "reason": str(exc), "lt_image_count": 0, "lt_figure_count": 0}
    return {
        "status": "ok",
        "lt_image_count": lt_image_count,
        "lt_figure_count": lt_figure_count,
    }


def _run_backend(
    *,
    name: str,
    filename: str,
    source_bytes: bytes,
    converter: Callable[[str, bytes], tuple[bytes, str]],
) -> dict[str, object]:
    try:
        docx_bytes, conversion_backend = converter(filename, source_bytes)
        validate_docx_source_bytes(docx_bytes)
        paragraphs, image_assets = extract_document_content_from_docx(
            build_in_memory_uploaded_file(
                source_name=Path(filename).with_suffix(".docx").name,
                source_bytes=docx_bytes,
            )
        )
    except Exception as exc:
        return {
            "status": "error",
            "backend": name,
            "error_type": type(exc).__name__,
            "error": str(exc),
        }
    return {
        "status": "ok",
        "backend": name,
        "conversion_backend": conversion_backend,
        "docx_bytes": len(docx_bytes),
        **summarize_docx_formatting(docx_bytes),
        **summarize_paragraph_units(paragraphs, image_assets=image_assets),
    }


def summarize_docx_formatting(docx_bytes: bytes) -> dict[str, object]:
    document = Document(BytesIO(docx_bytes))
    paragraph_style_counts = Counter(paragraph.style.name if paragraph.style else "" for paragraph in document.paragraphs)
    direct_bold_run_count = sum(1 for paragraph in document.paragraphs for run in paragraph.runs if run.bold is True)
    direct_italic_run_count = sum(1 for paragraph in document.paragraphs for run in paragraph.runs if run.italic is True)
    media_count = 0
    drawing_count = 0
    anchored_shape_count = 0
    inline_shape_count = len(document.inline_shapes)
    try:
        with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
            names = archive.namelist()
            media_count = sum(1 for name in names if name.startswith("word/media/"))
            xml = "".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in names
                if name.startswith("word/") and name.endswith(".xml")
            )
            drawing_count = xml.count("<w:drawing")
            anchored_shape_count = xml.count("wp:anchor")
    except zipfile.BadZipFile:
        pass
    return {
        "docx_paragraph_count": len(document.paragraphs),
        "docx_style_counts": dict(paragraph_style_counts.most_common(8)),
        "direct_bold_run_count": direct_bold_run_count,
        "direct_italic_run_count": direct_italic_run_count,
        "docx_media_count": media_count,
        "docx_drawing_count": drawing_count,
        "docx_inline_shape_count": inline_shape_count,
        "docx_anchored_shape_count": anchored_shape_count,
    }


def summarize_paragraph_units(paragraphs, *, image_assets) -> dict[str, object]:
    texts = [str(paragraph.text or "").strip() for paragraph in paragraphs if str(paragraph.text or "").strip()]
    role_counts = Counter(str(getattr(paragraph, "role", "") or "unknown") for paragraph in paragraphs)
    repeated_short_texts = Counter(_normalize_text(text) for text in texts if len(_normalize_text(text)) <= 80)
    repeated_short_text_count = sum(1 for count in repeated_short_texts.values() if count >= 3)
    page_number_like_count = sum(1 for text in texts if _PAGE_NUMBER_TEXT_PATTERN.match(_normalize_text(text)))
    markdown_emphasis_marker_count = sum(
        1 for text in texts if text.startswith(("**", "*")) or text.endswith(("**", "*"))
    )
    return {
        "paragraph_count": len(paragraphs),
        "nonempty_paragraph_count": len(texts),
        "char_count": sum(len(text) for text in texts),
        "role_counts": dict(sorted(role_counts.items())),
        "heading_count": role_counts.get("heading", 0),
        "list_count": role_counts.get("list", 0),
        "bold_paragraph_count": sum(1 for paragraph in paragraphs if bool(getattr(paragraph, "is_bold", False))),
        "italic_paragraph_count": sum(1 for paragraph in paragraphs if bool(getattr(paragraph, "is_italic", False))),
        "page_number_like_paragraph_count": page_number_like_count,
        "markdown_emphasis_marker_count": markdown_emphasis_marker_count,
        "repeated_short_text_count": repeated_short_text_count,
        "extracted_image_asset_count": len(image_assets),
        "first_paragraph_previews": texts[:8],
    }


def _normalize_text(value: str) -> str:
    return " ".join(value.split())


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Compare LibreOffice PDF import with the feature-flagged text-layer bridge."
    )
    parser.add_argument("--input-pdf", type=Path, required=True)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args(argv)

    comparison = build_pdf_import_backend_comparison(input_pdf=args.input_pdf)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(json.dumps(comparison, ensure_ascii=False, indent=2), encoding="utf-8")
    else:
        json.dump(comparison, sys.stdout, ensure_ascii=False, indent=2)
        sys.stdout.write("\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
