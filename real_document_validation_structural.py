from __future__ import annotations

from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from difflib import SequenceMatcher
from io import BytesIO
import json
from pathlib import Path
import re
from typing import Any, cast

from docx import Document
from docx.oxml.ns import qn

import processing_runtime
from config import load_app_config
from document import (
    build_document_text,
    extract_document_content_from_docx,
    extract_document_content_with_normalization_reports,
    extract_document_content_with_boundary_report,
    inspect_placeholder_integrity,
    summarize_boundary_normalization_metrics,
)
from formatting_transfer import normalize_semantic_output_docx, preserve_source_paragraph_properties
from generation import convert_markdown_to_docx_bytes, ensure_pandoc_available
from image_reinsertion import reinsert_inline_images
from models import ParagraphBoundaryNormalizationReport
from processing_service import clone_processing_service
from real_document_validation_common import build_validation_event_logger, build_validation_runtime_config
from real_document_validation_profiles import DocumentProfile, RunProfile, apply_runtime_resolution_to_app_config, resolve_runtime_resolution

PROJECT_ROOT = Path(__file__).resolve().parent
FORMATTING_DIAGNOSTICS_DIR = PROJECT_ROOT / ".run" / "formatting_diagnostics"


@dataclass
class UploadedFileStub:
    name: str
    content: bytes
    position: int = 0

    def read(self, size: int = -1) -> bytes:
        if size < 0:
            data = self.content[self.position :]
            self.position = len(self.content)
            return data
        start = self.position
        end = min(len(self.content), start + size)
        self.position = end
        return self.content[start:end]

    def getvalue(self) -> bytes:
        return self.content

    def seek(self, offset: int, whence: int = 0) -> int:
        if whence == 0:
            self.position = max(0, offset)
        elif whence == 1:
            self.position = max(0, self.position + offset)
        elif whence == 2:
            self.position = max(0, len(self.content) + offset)
        else:
            raise ValueError(f"Unsupported whence: {whence}")
        return self.position


def evaluate_extraction_profile(document_profile: DocumentProfile) -> dict[str, object]:
    source_path = document_profile.resolved_source_path(PROJECT_ROOT)
    source_bytes = source_path.read_bytes()
    normalized_source = processing_runtime.normalize_uploaded_document(filename=source_path.name, source_bytes=source_bytes)
    paragraphs, image_assets, normalization_report, _, relation_report = extract_document_content_with_normalization_reports(
        BytesIO(normalized_source.content_bytes)
    )
    metrics = _build_structural_metrics(
        paragraphs=paragraphs,
        image_assets=image_assets,
        normalization_report=normalization_report,
        relation_report=relation_report,
    )
    checks = _build_extraction_checks(document_profile, metrics)
    return _build_validation_result(
        document_profile=document_profile,
        run_profile=None,
        tier="extraction",
        source_path=source_path,
        result="succeeded",
        metrics=metrics,
        checks=checks,
        runtime_config=None,
        output_artifacts=None,
        formatting_diagnostics=[],
    )


def run_structural_passthrough_validation(
    document_profile: DocumentProfile,
    run_profile: RunProfile,
) -> dict[str, object]:
    source_path = document_profile.resolved_source_path(PROJECT_ROOT)
    source_bytes = source_path.read_bytes()
    app_config = load_app_config()
    runtime_resolution = resolve_runtime_resolution(app_config, run_profile)
    runtime_config = apply_runtime_resolution_to_app_config(app_config, runtime_resolution)
    runtime = _build_runtime_capture()
    event_log: list[dict[str, object]] = []
    formatting_before = _snapshot_formatting_diagnostics_paths()
    result, prepared = _build_validation_processing_service(event_log).run_prepared_background_document(
        uploaded_file=UploadedFileStub(source_path.name, source_bytes),
        chunk_size=int(runtime_resolution.effective.chunk_size),
        image_mode=str(runtime_resolution.effective.image_mode),
        keep_all_image_variants=bool(runtime_resolution.effective.keep_all_image_variants),
        app_config=runtime_config,
        model=str(runtime_resolution.effective.model),
        max_retries=int(runtime_resolution.effective.max_retries),
        job_mutator=_build_passthrough_job,
        progress_callback=None,
        runtime=runtime,
    )

    formatting_after = _snapshot_formatting_diagnostics_paths()
    formatting_paths = _collect_new_formatting_diagnostics_paths(formatting_before, formatting_after)
    formatting_diagnostics = _load_formatting_diagnostics_payloads(formatting_paths)

    runtime_state = _runtime_state(runtime)
    latest_docx_bytes = runtime_state.get("latest_docx_bytes")
    latest_markdown = str(runtime_state.get("latest_markdown") or "")
    if not isinstance(latest_docx_bytes, (bytes, bytearray)):
        latest_docx_bytes = b""
    output_artifacts = _build_output_artifacts(bytes(latest_docx_bytes), latest_markdown)

    (
        source_paragraphs,
        source_image_assets,
        source_normalization_report,
        _,
        source_relation_report,
    ) = extract_document_content_with_normalization_reports(BytesIO(prepared.uploaded_file_bytes))
    output_paragraphs = []
    output_image_assets = []
    if output_artifacts["output_docx_openable"]:
        output_paragraphs, output_image_assets = extract_document_content_from_docx(BytesIO(bytes(latest_docx_bytes)))
    metrics = _build_structural_metrics(
        paragraphs=source_paragraphs,
        image_assets=source_image_assets,
        normalization_report=source_normalization_report,
    )
    metrics.update(
        {
            "output_paragraph_count": len(output_paragraphs),
            "output_heading_count": sum(1 for paragraph in output_paragraphs if paragraph.role == "heading"),
            "output_numbered_item_count": sum(
                1 for paragraph in output_paragraphs if paragraph.role == "list" and paragraph.list_kind == "ordered"
            ),
            "output_image_count": len(output_image_assets),
            "output_table_count": sum(1 for paragraph in output_paragraphs if paragraph.role == "table"),
            "formatting_diagnostics_count": len(formatting_diagnostics),
            "max_unmapped_source_paragraphs": _max_payload_length(formatting_diagnostics, "unmapped_source_ids"),
            "max_unmapped_target_paragraphs": _max_payload_length(formatting_diagnostics, "unmapped_target_indexes"),
            "accepted_merged_sources_count": _count_payload_items(formatting_diagnostics, "accepted_merged_sources"),
            "max_accepted_merged_sources": _max_accepted_merged_sources(formatting_diagnostics),
            "relation_count": source_relation_report.total_relations,
            "rejected_relation_candidate_count": source_relation_report.rejected_candidate_count,
            "relation_counts": dict(source_relation_report.relation_counts),
            "text_similarity": _calculate_text_similarity(source_paragraphs, output_paragraphs),
            "heading_level_drift": _calculate_heading_level_drift(source_paragraphs, output_paragraphs),
            "heading_only_output_detected": _is_heading_only_markdown(latest_markdown),
            "output_docx_openable": bool(output_artifacts["output_docx_openable"]),
        }
    )
    checks = _build_extraction_checks(document_profile, metrics)
    checks.extend(
        _build_structural_checks(
            document_profile=document_profile,
            result=result,
            metrics=metrics,
            output_artifacts=output_artifacts,
        )
    )
    return _build_validation_result(
        document_profile=document_profile,
        run_profile=run_profile,
        tier="structural",
        source_path=source_path,
        result=result,
        metrics=metrics,
        checks=checks,
        runtime_config=build_validation_runtime_config(runtime_resolution),
        output_artifacts=output_artifacts,
        formatting_diagnostics=formatting_diagnostics,
        event_log=event_log,
    )


def _build_validation_result(
    *,
    document_profile: DocumentProfile,
    run_profile: RunProfile | None,
    tier: str,
    source_path: Path,
    result: str,
    metrics: dict[str, object],
    checks: list[dict[str, object]],
    runtime_config: dict[str, object] | None,
    output_artifacts: dict[str, object] | None,
    formatting_diagnostics: list[dict[str, object]],
    event_log: list[dict[str, object]] | None = None,
) -> dict[str, object]:
    failed_checks = [str(check["name"]) for check in checks if not bool(check["passed"])]
    return {
        "document_profile_id": document_profile.id,
        "run_profile_id": None if run_profile is None else run_profile.id,
        "validation_tier": tier,
        "source_document_path": str(source_path.relative_to(PROJECT_ROOT)).replace("\\", "/"),
        "result": result,
        "passed": not failed_checks,
        "failed_checks": failed_checks,
        "metrics": metrics,
        "checks": checks,
        "runtime_config": runtime_config,
        "output_artifacts": output_artifacts,
        "formatting_diagnostics": formatting_diagnostics,
        "event_log": event_log or [],
    }


def _build_passthrough_job(job: Mapping[str, object]) -> dict[str, object]:
    cloned = dict(job)
    cloned["job_kind"] = "passthrough"
    return cloned


def _build_validation_processing_service(event_log: list[dict[str, object]]):
    def _run_document_processing_impl(**kwargs: object) -> str:
        from document_pipeline import run_document_processing

        return cast(str, run_document_processing(**kwargs))

    return clone_processing_service(
        get_client_fn=lambda: object(),
        load_system_prompt_fn=lambda: "",
        ensure_pandoc_available_fn=ensure_pandoc_available,
        generate_markdown_block_fn=lambda **kwargs: cast(str, kwargs.get("target_text") or ""),
        convert_markdown_to_docx_bytes_fn=convert_markdown_to_docx_bytes,
        process_document_images_impl_fn=lambda **kwargs: list(kwargs.get("image_assets") or []),
        analyze_image_fn=lambda *args, **kwargs: None,
        generate_image_candidate_fn=lambda *args, **kwargs: b"",
        validate_redraw_result_fn=lambda *args, **kwargs: None,
        detect_image_mime_type_fn=lambda *args, **kwargs: None,
        inspect_placeholder_integrity_fn=_inspect_placeholder_integrity_adapter,
        preserve_source_paragraph_properties_fn=_preserve_source_paragraph_properties_adapter,
        normalize_semantic_output_docx_fn=_normalize_semantic_output_docx_adapter,
        reinsert_inline_images_fn=_reinsert_inline_images_adapter,
        run_document_processing_impl_fn=_run_document_processing_impl,
        present_error_fn=lambda code, exc, title, **context: f"{title}: {exc}",
        log_event_fn=build_validation_event_logger(event_log),
        emit_state_fn=_emit_state,
        emit_finalize_fn=_emit_finalize,
        emit_activity_fn=_emit_activity,
        emit_log_fn=_emit_log,
        emit_status_fn=_emit_status,
        emit_image_log_fn=lambda runtime, **payload: None,
        emit_image_reset_fn=lambda runtime: None,
        should_stop_processing_fn=lambda runtime: False,
        resolve_uploaded_filename_fn=processing_runtime.resolve_uploaded_filename,
        image_model_call_budget_cls=object,
        image_model_call_budget_exceeded_cls=RuntimeError,
    )


def _build_structural_metrics(
    *,
    paragraphs: Sequence[object],
    image_assets: Sequence[object],
    normalization_report: ParagraphBoundaryNormalizationReport | None = None,
    relation_report=None,
) -> dict[str, object]:
    paragraph_units = list(paragraphs)
    return {
        "paragraph_count": len(paragraph_units),
        "heading_count": sum(1 for paragraph in paragraph_units if getattr(paragraph, "role", None) == "heading"),
        "numbered_item_count": sum(
            1
            for paragraph in paragraph_units
            if getattr(paragraph, "role", None) == "list" and getattr(paragraph, "list_kind", None) == "ordered"
        ),
        "image_count": len(list(image_assets)),
        "table_count": sum(1 for paragraph in paragraph_units if getattr(paragraph, "role", None) == "table"),
        "raw_paragraph_count": 0 if normalization_report is None else normalization_report.total_raw_paragraphs,
        "logical_paragraph_count": 0 if normalization_report is None else normalization_report.total_logical_paragraphs,
        "merged_group_count": 0 if normalization_report is None else normalization_report.merged_group_count,
        "merged_raw_paragraph_count": 0 if normalization_report is None else normalization_report.merged_raw_paragraph_count,
        **summarize_boundary_normalization_metrics(normalization_report),
        "relation_count": 0 if relation_report is None else relation_report.total_relations,
        "rejected_relation_candidate_count": 0 if relation_report is None else relation_report.rejected_candidate_count,
        "relation_counts": {} if relation_report is None else dict(relation_report.relation_counts),
    }


def _build_extraction_checks(document_profile: DocumentProfile, metrics: Mapping[str, object]) -> list[dict[str, object]]:
    checks = [
        _check_minimum("paragraph_count_minimum", _as_int(metrics, "paragraph_count"), document_profile.min_paragraphs),
    ]
    if document_profile.min_merged_groups > 0:
        checks.append(
            _check_minimum(
                "merged_group_count_minimum",
                _as_int(metrics, "merged_group_count"),
                document_profile.min_merged_groups,
            )
        )
    if document_profile.min_merged_raw_paragraphs > 0:
        checks.append(
            _check_minimum(
                "merged_raw_paragraph_count_minimum",
                _as_int(metrics, "merged_raw_paragraph_count"),
                document_profile.min_merged_raw_paragraphs,
            )
        )
    if document_profile.has_headings:
        checks.append(_check_minimum("heading_count_minimum", _as_int(metrics, "heading_count"), document_profile.min_headings))
    if document_profile.has_numbered_lists:
        checks.append(
            _check_minimum(
                "numbered_item_count_minimum",
                _as_int(metrics, "numbered_item_count"),
                document_profile.min_numbered_items,
            )
        )
    if document_profile.has_images:
        checks.append(_check_minimum("image_count_minimum", _as_int(metrics, "image_count"), document_profile.min_images))
    if document_profile.has_tables:
        checks.append(_check_minimum("table_count_minimum", _as_int(metrics, "table_count"), document_profile.min_tables))
    return checks


def _build_structural_checks(
    *,
    document_profile: DocumentProfile,
    result: str,
    metrics: Mapping[str, object],
    output_artifacts: Mapping[str, object],
) -> list[dict[str, object]]:
    checks = [
        {"name": "pipeline_succeeded", "passed": result == "succeeded", "result": result},
        {
            "name": "output_docx_openable",
            "passed": bool(output_artifacts.get("output_docx_openable")),
            "output_docx_openable": output_artifacts.get("output_docx_openable"),
        },
        {
            "name": "formatting_diagnostics_threshold",
            "passed": _as_int(metrics, "formatting_diagnostics_count") <= document_profile.max_formatting_diagnostics,
            "actual": metrics["formatting_diagnostics_count"],
            "allowed": document_profile.max_formatting_diagnostics,
        },
        {
            "name": "unmapped_source_threshold",
            "passed": _as_int(metrics, "max_unmapped_source_paragraphs") <= document_profile.max_unmapped_source_paragraphs,
            "actual": metrics["max_unmapped_source_paragraphs"],
            "allowed": document_profile.max_unmapped_source_paragraphs,
        },
        {
            "name": "unmapped_target_threshold",
            "passed": _as_int(metrics, "max_unmapped_target_paragraphs") <= document_profile.max_unmapped_target_paragraphs,
            "actual": metrics["max_unmapped_target_paragraphs"],
            "allowed": document_profile.max_unmapped_target_paragraphs,
        },
        {
            "name": "heading_level_drift_threshold",
            "passed": _as_int(metrics, "heading_level_drift") <= document_profile.max_heading_level_drift,
            "actual": metrics["heading_level_drift"],
            "allowed": document_profile.max_heading_level_drift,
        },
        {
            "name": "text_similarity_threshold",
            "passed": _as_float(metrics, "text_similarity") >= document_profile.min_text_similarity,
            "actual": metrics["text_similarity"],
            "required": document_profile.min_text_similarity,
        },
    ]
    if document_profile.require_numbered_lists_preserved:
        checks.append(
            {
                "name": "numbered_lists_preserved",
                "passed": _as_int(metrics, "output_numbered_item_count") >= _as_int(metrics, "numbered_item_count"),
                "source": metrics["numbered_item_count"],
                "output": metrics["output_numbered_item_count"],
            }
        )
    if document_profile.require_nonempty_output:
        checks.append(
            {
                "name": "nonempty_output_required",
                "passed": bool(output_artifacts.get("output_visible_text_chars")),
                "output_visible_text_chars": output_artifacts.get("output_visible_text_chars"),
            }
        )
    if document_profile.forbid_heading_only_collapse:
        checks.append(
            {
                "name": "forbid_heading_only_collapse",
                "passed": not bool(metrics["heading_only_output_detected"]),
                "heading_only_output_detected": metrics["heading_only_output_detected"],
            }
        )
    return checks


def _check_minimum(name: str, actual: int, minimum: int) -> dict[str, object]:
    return {"name": name, "passed": actual >= minimum, "actual": actual, "minimum": minimum}


def _snapshot_formatting_diagnostics_paths() -> set[str]:
    if not FORMATTING_DIAGNOSTICS_DIR.exists():
        return set()
    return {str(path.resolve()) for path in FORMATTING_DIAGNOSTICS_DIR.glob("*.json") if path.is_file()}


def _collect_new_formatting_diagnostics_paths(before: set[str], after: set[str]) -> list[str]:
    new_paths = [Path(path) for path in after - before]
    return [str(path) for path in sorted(new_paths, key=lambda candidate: (candidate.stat().st_mtime, str(candidate)))]


def _load_formatting_diagnostics_payloads(artifact_paths: Sequence[str]) -> list[dict[str, object]]:
    payloads: list[dict[str, object]] = []
    for artifact_path in artifact_paths:
        try:
            payload = json.loads(Path(artifact_path).read_text(encoding="utf-8"))
        except Exception:
            continue
        if isinstance(payload, dict):
            payloads.append(payload)
    return payloads


def _max_payload_length(payloads: Sequence[Mapping[str, object]], key: str) -> int:
    maximum = 0
    for payload in payloads:
        values = payload.get(key) or []
        if isinstance(values, Sequence) and not isinstance(values, (str, bytes, bytearray)):
            maximum = max(maximum, len(values))
    return maximum


def _count_payload_items(payloads: Sequence[Mapping[str, object]], key: str) -> int:
    total = 0
    for payload in payloads:
        values = payload.get(key) or []
        if isinstance(values, Sequence) and not isinstance(values, (str, bytes, bytearray)):
            total += len(values)
    return total


def _max_accepted_merged_sources(payloads: Sequence[Mapping[str, object]]) -> int:
    maximum = 0
    for payload in payloads:
        explicit_value = payload.get("max_accepted_merged_sources")
        if explicit_value is not None:
            try:
                maximum = max(maximum, int(explicit_value))
                continue
            except (TypeError, ValueError):
                pass

        accepted_sources = payload.get("accepted_merged_sources") or []
        if not isinstance(accepted_sources, Sequence) or isinstance(accepted_sources, (str, bytes, bytearray)):
            continue
        for entry in accepted_sources:
            if not isinstance(entry, Mapping):
                continue
            explicit_count = entry.get("accepted_merged_sources_count")
            if explicit_count is not None:
                try:
                    maximum = max(maximum, int(explicit_count))
                    continue
                except (TypeError, ValueError):
                    pass
            raw_indexes = entry.get("origin_raw_indexes") or []
            if isinstance(raw_indexes, Sequence) and not isinstance(raw_indexes, (str, bytes, bytearray)):
                maximum = max(maximum, len(raw_indexes))
    return maximum


def _calculate_text_similarity(source_paragraphs: Sequence[object], output_paragraphs: Sequence[object]) -> float:
    source_text = _normalize_text(build_document_text(cast(list[Any], list(source_paragraphs))))
    output_text = _normalize_text(build_document_text(cast(list[Any], list(output_paragraphs))))
    if not source_text and not output_text:
        return 1.0
    return round(SequenceMatcher(None, source_text, output_text).ratio(), 4)


def _calculate_heading_level_drift(source_paragraphs: Sequence[object], output_paragraphs: Sequence[object]) -> int:
    output_levels: dict[str, int] = {}
    for paragraph in output_paragraphs:
        if getattr(paragraph, "role", None) != "heading":
            continue
        normalized = _normalize_text(str(getattr(paragraph, "text", "")))
        if not normalized:
            continue
        output_levels[normalized] = int(getattr(paragraph, "heading_level", 0) or 0)

    max_drift = 0
    for paragraph in source_paragraphs:
        if getattr(paragraph, "role", None) != "heading":
            continue
        normalized = _normalize_text(str(getattr(paragraph, "text", "")))
        if not normalized or normalized not in output_levels:
            continue
        source_level = int(getattr(paragraph, "heading_level", 0) or 0)
        max_drift = max(max_drift, abs(source_level - output_levels[normalized]))
    return max_drift


def _normalize_text(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text).strip().lower()
    normalized = re.sub(r"^#{1,6}\s+", "", normalized)
    return normalized


def _is_heading_only_markdown(text: str) -> bool:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return bool(lines) and all(line.startswith("#") and len(line.split()) >= 2 for line in lines)


def _build_output_artifacts(docx_bytes: bytes, markdown_text: str) -> dict[str, object]:
    openable = False
    output_paragraphs = 0
    output_inline_shapes = 0
    output_visible_text_chars = 0
    output_contains_placeholder_markup = False
    if docx_bytes:
        try:
            document = Document(BytesIO(docx_bytes))
        except Exception:
            document = None
        if document is not None:
            openable = True
            output_paragraphs = len(document.paragraphs)
            output_inline_shapes = len(document.inline_shapes)
            output_visible_text_chars = len("\n".join(paragraph.text for paragraph in document.paragraphs))
            output_contains_placeholder_markup = "[[DOCX_IMAGE_" in document._element.xml
    return {
        "output_docx_openable": openable,
        "output_paragraphs": output_paragraphs,
        "output_inline_shapes": output_inline_shapes,
        "output_visible_text_chars": output_visible_text_chars,
        "output_contains_placeholder_markup": output_contains_placeholder_markup,
        "markdown_chars": len(markdown_text),
    }


def _build_runtime_capture() -> dict[str, object]:
    return {"state": {}, "finalize": [], "activity": [], "log": [], "status": []}


def _runtime_mapping(runtime: object) -> dict[str, object]:
    return cast(dict[str, object], runtime)


def _runtime_state(runtime: object) -> dict[str, object]:
    return cast(dict[str, object], _runtime_mapping(runtime).setdefault("state", {}))


def _emit_state(runtime: object, **values: object) -> None:
    _runtime_state(runtime).update(values)


def _emit_finalize(
    runtime: object,
    stage: str,
    detail: str,
    progress: float,
    terminal_kind: str | None = None,
) -> None:
    cast(list[object], _runtime_mapping(runtime).setdefault("finalize", [])).append(
        (stage, detail, progress, terminal_kind)
    )


def _emit_activity(runtime: object, message: str) -> None:
    cast(list[object], _runtime_mapping(runtime).setdefault("activity", [])).append(message)


def _emit_log(runtime: object, **payload: object) -> None:
    cast(list[object], _runtime_mapping(runtime).setdefault("log", [])).append(payload)


def _emit_status(runtime: object, **payload: object) -> None:
    cast(list[object], _runtime_mapping(runtime).setdefault("status", [])).append(payload)


def _inspect_placeholder_integrity_adapter(markdown_text: str, image_assets: Sequence[object]) -> Mapping[str, str]:
    return inspect_placeholder_integrity(markdown_text, cast(list[Any], list(image_assets)))


def _preserve_source_paragraph_properties_adapter(
    docx_bytes: bytes,
    paragraphs: Sequence[object],
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
) -> bytes:
    return preserve_source_paragraph_properties(
        docx_bytes,
        cast(list[Any], list(paragraphs)),
        generated_paragraph_registry=generated_paragraph_registry,
    )


def _normalize_semantic_output_docx_adapter(
    docx_bytes: bytes,
    paragraphs: Sequence[object],
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
) -> bytes:
    return normalize_semantic_output_docx(
        docx_bytes,
        cast(list[Any], list(paragraphs)),
        generated_paragraph_registry=generated_paragraph_registry,
    )


def _reinsert_inline_images_adapter(docx_bytes: bytes, image_assets: Sequence[object]) -> bytes:
    return reinsert_inline_images(docx_bytes, cast(list[Any], list(image_assets)))


def _as_int(metrics: Mapping[str, object], key: str) -> int:
    return int(cast(int, metrics[key]))


def _as_float(metrics: Mapping[str, object], key: str) -> float:
    return float(cast(float, metrics[key]))
