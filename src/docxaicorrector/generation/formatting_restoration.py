"""Minimal DOCX formatting restoration (docx-apply half of spec 033).

Extracted verbatim from ``generation/formatting_transfer.py`` (spec 033, Step 2). Holds
the docx-apply functions — minimal image/caption normalization, split-heading styling,
list-numbering restoration, pPr/XML helpers, TOC paragraph/run restoration, and
direct-alignment / semantic-quote restoration — plus the CENTER_/quote constants they
use.

The dependency direction is one-way: this module imports the shared low-level helpers
from ``formatting_mapping`` (the leaf mapper cluster); ``formatting_mapping`` never
imports this module, so there is no import cycle. The facade
(``formatting_transfer``) re-exports these names so ``formatting_transfer.<name>`` and
``from ...formatting_transfer import <name>`` keep resolving for callers and tests.
"""

import logging
import re
from typing import Mapping, Sequence, cast

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from docxaicorrector.core.logger import log_event
from docxaicorrector.core.models import ParagraphUnit
from docxaicorrector.document.extraction import IMAGE_PLACEHOLDER_PATTERN
from docxaicorrector.document.roles import (
    find_child_element,
    get_xml_attribute,
    infer_heuristic_heading_level,
    is_image_only_text,
    is_likely_caption_text,
    xml_local_name,
)
from docxaicorrector.generation.formatting_mapping import (
    _build_generated_registry_by_paragraph_id,
    _extract_target_heading_level,
    _generated_registry_text,
    _is_heading_like_source_paragraph,
    _normalize_text_for_mapping,
    _paragraph_preview,
    _target_paragraph_has_heading_style,
    _target_paragraph_style_name,
)

# Spec TOC/minimal-formatting 2026-04-21: centered direct alignment is allowed
# only for narrow non-heading cases, with an explicit short-paragraph heuristic.
CENTER_SHORT_NON_HEADING_MAX_CHARS = 90
CENTER_SHORT_NON_HEADING_MAX_WORDS = 12
ALLOWED_CENTERED_QUOTE_STRUCTURAL_ROLES = {"epigraph", "attribution", "dedication"}
DISALLOWED_CENTER_SHORT_STRUCTURAL_ROLES = {"toc_header", "toc_entry", "heading", "caption"}


def _apply_minimal_image_formatting(document) -> None:
    for paragraph in document.paragraphs:
        if is_image_only_text(paragraph.text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _apply_minimal_caption_formatting(
    document,
    source_paragraphs: Sequence[ParagraphUnit],
    *,
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
) -> None:
    if not document.paragraphs:
        return

    generated_registry_by_id = _build_generated_registry_by_paragraph_id(generated_paragraph_registry)
    source_caption_texts = {
        _normalize_text_for_mapping(paragraph.text)
        for paragraph in source_paragraphs
        if paragraph.role == "caption" and paragraph.text.strip()
    }
    source_caption_texts.discard("")
    generated_caption_texts = {
        _normalize_text_for_mapping(_generated_registry_text(generated_registry_by_id.get(paragraph.paragraph_id or "")))
        for paragraph in source_paragraphs
        if paragraph.role == "caption" and paragraph.paragraph_id
    }
    generated_caption_texts.discard("")

    for paragraph in document.paragraphs:
        if not _is_caption_candidate(
            document,
            paragraph,
            source_caption_texts=source_caption_texts,
            generated_caption_texts=generated_caption_texts,
        ):
            continue
        if _style_exists(document, "Caption"):
            paragraph.style = document.styles["Caption"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _is_caption_candidate(
    document,
    paragraph,
    *,
    source_caption_texts: set[str],
    generated_caption_texts: set[str],
) -> bool:
    text = paragraph.text.strip()
    if not text or IMAGE_PLACEHOLDER_PATTERN.search(text):
        return False

    normalized_text = _normalize_text_for_mapping(text)
    if not normalized_text:
        return False

    has_anchor_context = _has_caption_anchor_context(document, paragraph)
    if not has_anchor_context:
        return False

    if normalized_text in source_caption_texts or normalized_text in generated_caption_texts:
        return True
    return is_likely_caption_text(text)


def _has_caption_anchor_context(document, paragraph) -> bool:
    body_children = list(document._element.body.iterchildren())
    paragraph_element = paragraph._element

    for index, child in enumerate(body_children):
        if child != paragraph_element:
            continue
        previous_child = body_children[index - 1] if index > 0 else None
        next_child = body_children[index + 1] if index + 1 < len(body_children) else None
        return _is_caption_anchor_block(previous_child) or _is_caption_anchor_block(next_child)

    return False


def _is_caption_anchor_block(block_element) -> bool:
    if block_element is None:
        return False

    local_name = xml_local_name(block_element.tag)
    if local_name == "tbl":
        return True
    if local_name != "p":
        return False

    text_content = "".join(block_element.itertext())
    return is_image_only_text(text_content)


# ---------------------------------------------------------------------------
# Paragraph-level normalization and list numbering restoration
# ---------------------------------------------------------------------------


def _apply_accepted_split_heading_styles(
    document,
    target_paragraphs: Sequence[Paragraph],
    accepted_split_targets: Sequence[Mapping[str, object]],
    source_paragraphs: Sequence[ParagraphUnit],
) -> None:
    for accepted_target in accepted_split_targets:
        target_index_value = accepted_target.get("target_index")
        source_index_value = accepted_target.get("derived_from_source_index")
        try:
            target_index = int(cast(int | str, target_index_value))
            source_index = int(cast(int | str, source_index_value))
        except (TypeError, ValueError):
            continue
        if target_index < 0 or target_index >= len(target_paragraphs):
            continue
        if source_index < 0 or source_index >= len(source_paragraphs):
            continue

        paragraph = target_paragraphs[target_index]
        source_paragraph = source_paragraphs[source_index]
        normalized_target = _normalize_text_for_mapping(paragraph.text)
        normalized_source = _normalize_text_for_mapping(source_paragraph.text)
        if not normalized_target or not normalized_source.startswith(normalized_target):
            continue
        inferred_level = infer_heuristic_heading_level(paragraph.text)
        heading_style = f"Heading {min(max(inferred_level, 1), 6)}"
        if _style_exists(document, heading_style):
            paragraph.style = document.styles[heading_style]


def _extract_paragraph_num_id(paragraph) -> str | None:
    paragraph_properties = find_child_element(paragraph._element, "pPr")
    num_pr = find_child_element(paragraph_properties, "numPr")
    num_id = find_child_element(num_pr, "numId")
    return get_xml_attribute(num_id, "val") if num_id is not None else None


def _restore_list_numbering_for_mapped_paragraphs(document, mapping_pairs: list[tuple[ParagraphUnit, Paragraph]]) -> list[dict[str, object]]:
    numbering_root = _get_target_numbering_root(document)
    decisions: list[dict[str, object]] = []
    if numbering_root is None:
        for source_paragraph, _ in mapping_pairs:
            if source_paragraph.role != "list":
                continue
            decisions.append(
                {
                    "paragraph_id": source_paragraph.paragraph_id,
                    "text_preview": _paragraph_preview(source_paragraph.text),
                    "action": "target_numbering_part_missing",
                }
            )
        return decisions

    numbering_id_map: dict[tuple[str, str, str, str], tuple[int, int]] = {}
    next_abstract_num_id = _next_numbering_identifier(numbering_root, "abstractNum", "abstractNumId")
    next_num_id = _next_numbering_identifier(numbering_root, "num", "numId")

    for source_paragraph, target_paragraph in mapping_pairs:
        if source_paragraph.role != "list":
            continue
        existing_target_num_id = _extract_paragraph_num_id(target_paragraph)
        if existing_target_num_id is not None:
            decisions.append(
                {
                    "paragraph_id": source_paragraph.paragraph_id,
                    "text_preview": _paragraph_preview(source_paragraph.text),
                    "action": "kept_existing_target_numbering",
                    "list_kind": source_paragraph.list_kind,
                    "list_level": source_paragraph.list_level,
                    "target_num_id": existing_target_num_id,
                }
            )
            continue
        if not source_paragraph.list_num_xml or not source_paragraph.list_abstract_num_xml:
            decisions.append(
                {
                    "paragraph_id": source_paragraph.paragraph_id,
                    "text_preview": _paragraph_preview(source_paragraph.text),
                    "action": "missing_source_numbering_xml",
                    "list_kind": source_paragraph.list_kind,
                    "list_level": source_paragraph.list_level,
                }
            )
            continue

        mapping_key = (
            source_paragraph.list_num_id or "",
            source_paragraph.list_abstract_num_id or "",
            source_paragraph.list_num_xml,
            source_paragraph.list_abstract_num_xml,
        )
        mapped_ids = numbering_id_map.get(mapping_key)
        if mapped_ids is None:
            abstract_num_id = next_abstract_num_id
            num_id = next_num_id
            next_abstract_num_id += 1
            next_num_id += 1
            if not _append_numbering_definition(numbering_root, source_paragraph, abstract_num_id, num_id):
                decisions.append(
                    {
                        "paragraph_id": source_paragraph.paragraph_id,
                        "text_preview": _paragraph_preview(source_paragraph.text),
                        "action": "append_definition_failed",
                        "list_kind": source_paragraph.list_kind,
                        "list_level": source_paragraph.list_level,
                    }
                )
                continue
            mapped_ids = (abstract_num_id, num_id)
            numbering_id_map[mapping_key] = mapped_ids

        _apply_list_numbering_to_paragraph(target_paragraph, list_level=source_paragraph.list_level, num_id=mapped_ids[1])
        decisions.append(
            {
                "paragraph_id": source_paragraph.paragraph_id,
                "text_preview": _paragraph_preview(source_paragraph.text),
                "action": "restored",
                "list_kind": source_paragraph.list_kind,
                "list_level": source_paragraph.list_level,
                "target_num_id": mapped_ids[1],
                "target_abstract_num_id": mapped_ids[0],
            }
        )

    return decisions


def _get_target_numbering_root(document):
    numbering_part = getattr(document.part, "numbering_part", None)
    return getattr(numbering_part, "element", None)


def _next_numbering_identifier(numbering_root, tag_name: str, attribute_name: str) -> int:
    max_identifier = 0
    for child in numbering_root:
        if xml_local_name(child.tag) != tag_name:
            continue
        value = get_xml_attribute(child, attribute_name)
        try:
            max_identifier = max(max_identifier, int(value or "0"))
        except ValueError:
            continue
    return max_identifier + 1


def _append_numbering_definition(numbering_root, source_paragraph: ParagraphUnit, abstract_num_id: int, num_id: int) -> bool:
    try:
        abstract_num_element = parse_xml(source_paragraph.list_abstract_num_xml or "")
        num_element = parse_xml(source_paragraph.list_num_xml or "")
    except Exception:
        return False

    abstract_num_element.set(qn("w:abstractNumId"), str(abstract_num_id))
    num_element.set(qn("w:numId"), str(num_id))

    abstract_num_id_element = find_child_element(num_element, "abstractNumId")
    if abstract_num_id_element is None:
        abstract_num_id_element = OxmlElement("w:abstractNumId")
        num_element.insert(0, abstract_num_id_element)
    abstract_num_id_element.set(qn("w:val"), str(abstract_num_id))

    numbering_root.append(abstract_num_element)
    numbering_root.append(num_element)
    return True


def _apply_list_numbering_to_paragraph(paragraph, *, list_level: int, num_id: int) -> None:
    paragraph_properties = _ensure_paragraph_properties(paragraph)
    existing_num_pr = find_child_element(paragraph_properties, "numPr")
    if existing_num_pr is not None:
        paragraph_properties.remove(existing_num_pr)

    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(max(0, list_level)))
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)
    paragraph_properties.append(num_pr)


def _style_exists(document, style_name: str) -> bool:
    try:
        document.styles[style_name]
    except KeyError:
        return False
    return True


def _replace_paragraph_properties_from_xml(paragraph, paragraph_properties_xml: str) -> bool:
    if not paragraph_properties_xml.strip():
        return False

    paragraph_properties = parse_xml(paragraph_properties_xml)
    # Defensive: a paragraph-level ``sectPr`` is a section break carrying page geometry
    # (incl. multi-column layout). It must never ride into the output via a pPr copy —
    # that grafts source scan columns onto the delivered DOCX. Strip it at the single
    # graft point regardless of which caller supplied the pPr.
    for child in list(paragraph_properties):
        if xml_local_name(child.tag) == "sectPr":
            paragraph_properties.remove(child)
    existing_properties = find_child_element(paragraph._element, "pPr")
    if existing_properties is not None:
        paragraph._element.remove(existing_properties)
    paragraph._element.insert(0, paragraph_properties)
    return True


def _sanitize_toc_paragraph_properties_xml(paragraph_properties_xml: str) -> str:
    if not paragraph_properties_xml.strip():
        return ""

    paragraph_properties = parse_xml(paragraph_properties_xml)
    # ``sectPr`` carries page geometry (size/margins/**columns**). A source scanned in
    # two columns (FineReader) stores a continuous 2-column section break inside a
    # paragraph's pPr; copying it onto the target paragraph grafts a 2-column section
    # into the output (the narrow-column front-matter defect). Geometry is dropped by
    # the minimal-formatting contract, so strip ``sectPr`` here too.
    unsafe_geometry_names = {"ind", "tabs", "spacing", "jc", "pStyle", "sectPr"}
    for child in list(paragraph_properties):
        if xml_local_name(child.tag) in unsafe_geometry_names:
            paragraph_properties.remove(child)

    return paragraph_properties.xml if len(paragraph_properties) > 0 else ""


def _restore_toc_paragraph_properties_for_mapped_pairs(document, mapping_pairs: list[tuple[ParagraphUnit, Paragraph]]) -> list[dict[str, object]]:
    decisions: list[dict[str, object]] = []
    for source_paragraph, target_paragraph in mapping_pairs:
        structural_role = str(getattr(source_paragraph, "structural_role", "") or "").strip().lower()
        if structural_role not in {"toc_header", "toc_entry"}:
            continue

        decision = {
            "paragraph_id": source_paragraph.paragraph_id,
            "structural_role": structural_role,
            "source_style_name": str(getattr(source_paragraph, "style_name", "") or "").strip() or None,
            "target_style_name": _target_paragraph_style_name(target_paragraph),
            "source_alignment": str(getattr(source_paragraph, "paragraph_alignment", "") or "").strip().lower() or None,
            "source_preview": _paragraph_preview(source_paragraph.text),
        }

        paragraph_properties_xml = str(getattr(source_paragraph, "paragraph_properties_xml", "") or "")
        if not paragraph_properties_xml:
            decisions.append({**decision, "action": "skipped", "reason": "missing_source_paragraph_properties"})
            continue

        sanitized_paragraph_properties_xml = _sanitize_toc_paragraph_properties_xml(paragraph_properties_xml)
        if not sanitized_paragraph_properties_xml:
            decisions.append({**decision, "action": "skipped", "reason": "source_toc_properties_only_contained_unsafe_geometry"})
            continue

        original_target_style_name = _target_paragraph_style_name(target_paragraph)
        restored = _replace_paragraph_properties_from_xml(target_paragraph, sanitized_paragraph_properties_xml)
        if not restored:
            decisions.append({**decision, "action": "skipped", "reason": "invalid_source_paragraph_properties"})
            continue

        if original_target_style_name and _style_exists(document, original_target_style_name):
            target_paragraph.style = document.styles[original_target_style_name]

        decisions.append({**decision, "action": "restored", "reason": "copied_sanitized_source_toc_paragraph_properties"})
    return decisions


def _restore_toc_run_formatting_for_mapped_pairs(mapping_pairs: list[tuple[ParagraphUnit, Paragraph]]) -> list[dict[str, object]]:
    decisions: list[dict[str, object]] = []
    for source_paragraph, target_paragraph in mapping_pairs:
        structural_role = str(getattr(source_paragraph, "structural_role", "") or "").strip().lower()
        if structural_role not in {"toc_header", "toc_entry"}:
            continue

        font_size_pt = getattr(source_paragraph, "font_size_pt", None)
        should_apply_bold = bool(getattr(source_paragraph, "is_bold", False))
        should_apply_italic = bool(getattr(source_paragraph, "is_italic", False))
        nonempty_runs = [run for run in target_paragraph.runs if run.text and run.text.strip()]
        if not nonempty_runs:
            decisions.append(
                {
                    "paragraph_id": source_paragraph.paragraph_id,
                    "structural_role": structural_role,
                    "action": "skipped",
                    "reason": "no_nonempty_target_runs",
                }
            )
            continue

        if font_size_pt is None and not should_apply_bold and not should_apply_italic:
            decisions.append(
                {
                    "paragraph_id": source_paragraph.paragraph_id,
                    "structural_role": structural_role,
                    "action": "skipped",
                    "reason": "no_safe_toc_run_formatting_to_restore",
                }
            )
            continue

        for run in nonempty_runs:
            if font_size_pt is not None:
                run.font.size = Pt(font_size_pt)
            if should_apply_bold:
                run.bold = True
            if should_apply_italic:
                run.italic = True

        decisions.append(
            {
                "paragraph_id": source_paragraph.paragraph_id,
                "structural_role": structural_role,
                "action": "restored",
                "reason": "copied_safe_toc_run_formatting",
                "font_size_pt": font_size_pt,
                "applied_bold": should_apply_bold,
                "applied_italic": should_apply_italic,
            }
        )
    return decisions


def _set_direct_paragraph_alignment(paragraph, alignment_value: str | None) -> None:
    paragraph_properties = _ensure_paragraph_properties(paragraph)
    existing_alignment = find_child_element(paragraph_properties, "jc")
    if alignment_value is None:
        if existing_alignment is not None:
            paragraph_properties.remove(existing_alignment)
        return
    if existing_alignment is None:
        existing_alignment = OxmlElement("w:jc")
        paragraph_properties.append(existing_alignment)
    existing_alignment.set(qn("w:val"), alignment_value)


def _is_allowlisted_centered_quote_paragraph(source_paragraph: ParagraphUnit) -> bool:
    return str(getattr(source_paragraph, "structural_role", "") or "").strip().lower() in ALLOWED_CENTERED_QUOTE_STRUCTURAL_ROLES


def _is_short_non_heading_paragraph(source_paragraph: ParagraphUnit) -> bool:
    role = str(getattr(source_paragraph, "role", "") or "").strip().lower()
    structural_role = str(getattr(source_paragraph, "structural_role", "") or "").strip().lower()
    if role in {"heading", "list", "caption"}:
        return False
    if structural_role in DISALLOWED_CENTER_SHORT_STRUCTURAL_ROLES:
        return False
    if getattr(source_paragraph, "heading_level", None) is not None or getattr(source_paragraph, "heading_source", None):
        return False
    normalized = re.sub(r"\s+", " ", str(getattr(source_paragraph, "text", "") or "")).strip()
    if not normalized:
        return False
    if len(normalized) > CENTER_SHORT_NON_HEADING_MAX_CHARS:
        return False
    words = [token for token in normalized.split(" ") if token]
    return len(words) <= CENTER_SHORT_NON_HEADING_MAX_WORDS


def _is_short_centered_caption_paragraph(source_paragraph: ParagraphUnit) -> bool:
    role = str(getattr(source_paragraph, "role", "") or "").strip().lower()
    if role != "caption":
        return False
    normalized = re.sub(r"\s+", " ", str(getattr(source_paragraph, "text", "") or "")).strip()
    if not normalized:
        return False
    if len(normalized) > CENTER_SHORT_NON_HEADING_MAX_CHARS:
        return False
    words = [token for token in normalized.split(" ") if token]
    return len(words) <= CENTER_SHORT_NON_HEADING_MAX_WORDS


def _is_short_centered_attribution_paragraph(source_paragraph: ParagraphUnit) -> bool:
    role = str(getattr(source_paragraph, "role", "") or "").strip().lower()
    structural_role = str(getattr(source_paragraph, "structural_role", "") or "").strip().lower()
    if structural_role in ALLOWED_CENTERED_QUOTE_STRUCTURAL_ROLES:
        return True
    if role not in {"body", "quote"}:
        return False
    normalized = re.sub(r"\s+", " ", str(getattr(source_paragraph, "text", "") or "")).strip()
    if not normalized:
        return False
    if len(normalized) > 60:
        return False
    words = [token for token in normalized.split(" ") if token]
    if len(words) > 4:
        return False
    letters_only = re.sub(r"[^A-Za-zА-Яа-яЁё]", "", normalized)
    return bool(letters_only) and letters_only.upper() == letters_only


def _resolve_direct_alignment_restoration_decision(
    source_paragraph: ParagraphUnit,
    target_paragraph: Paragraph,
) -> dict[str, object] | None:
    alignment_value = str(source_paragraph.paragraph_alignment or "").strip().lower()
    if not alignment_value:
        return None
    decision = {
        "paragraph_id": source_paragraph.paragraph_id,
        "source_alignment": alignment_value,
        "target_style_name": _target_paragraph_style_name(target_paragraph),
        "target_heading_level": _extract_target_heading_level(target_paragraph),
        "source_role": str(getattr(source_paragraph, "role", "") or "").strip().lower(),
        "source_structural_role": str(getattr(source_paragraph, "structural_role", "") or "").strip().lower(),
        "source_preview": _paragraph_preview(source_paragraph.text),
    }
    if _is_heading_like_source_paragraph(source_paragraph):
        return {**decision, "action": "skipped", "reason": "source_heading_semantics"}
    if _target_paragraph_has_heading_style(target_paragraph):
        return {**decision, "action": "skipped", "reason": "target_heading_style"}
    if alignment_value != "center":
        return {**decision, "action": "skipped", "reason": "unsupported_alignment_value"}
    if str(getattr(source_paragraph, "role", "") or "").strip().lower() == "list":
        return {**decision, "action": "skipped", "reason": "list_item_not_allowlisted"}
    if _is_short_centered_caption_paragraph(source_paragraph):
        return {**decision, "action": "restored", "reason": "center_caption_allowlisted"}
    if _is_short_centered_attribution_paragraph(source_paragraph):
        return {**decision, "action": "restored", "reason": "center_attribution_allowlisted"}
    if _is_allowlisted_centered_quote_paragraph(source_paragraph):
        return {**decision, "action": "restored", "reason": "center_quote_allowlisted"}
    if not _is_short_non_heading_paragraph(source_paragraph):
        return {**decision, "action": "skipped", "reason": "not_short_non_heading_paragraph"}
    return {**decision, "action": "restored", "reason": "center_allowlisted"}


def _restore_direct_paragraph_alignment_for_mapped_pairs(mapping_pairs: list[tuple[ParagraphUnit, Paragraph]]) -> list[dict[str, object]]:
    decisions: list[dict[str, object]] = []
    for source_paragraph, target_paragraph in mapping_pairs:
        decision = _resolve_direct_alignment_restoration_decision(source_paragraph, target_paragraph)
        if decision is None:
            continue
        decisions.append(decision)
        if decision["action"] != "restored":
            log_event(
                logging.INFO,
                "alignment_restoration_skipped",
                "Пропущено восстановление direct alignment для абзаца.",
                paragraph_id=decision.get("paragraph_id"),
                source_alignment=decision.get("source_alignment"),
                skip_reason=decision.get("reason"),
                target_style_name=decision.get("target_style_name"),
                target_heading_level=decision.get("target_heading_level"),
                source_preview=decision.get("source_preview"),
            )
            continue
        _set_direct_paragraph_alignment(target_paragraph, source_paragraph.paragraph_alignment)
    return decisions


def _restore_semantic_quote_formatting_for_mapped_pairs(mapping_pairs: list[tuple[ParagraphUnit, Paragraph]]) -> None:
    for source_paragraph, target_paragraph in mapping_pairs:
        if source_paragraph.structural_role not in {"epigraph", "attribution", "dedication"}:
            continue
        if source_paragraph.is_italic:
            for run in target_paragraph.runs:
                if run.text and run.text.strip():
                    run.italic = True


def _ensure_paragraph_properties(paragraph):
    paragraph_properties = find_child_element(paragraph._element, "pPr")
    if paragraph_properties is not None:
        return paragraph_properties

    paragraph_properties = OxmlElement("w:pPr")
    paragraph._element.insert(0, paragraph_properties)
    return paragraph_properties
