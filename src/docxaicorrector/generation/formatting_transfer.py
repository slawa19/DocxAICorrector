"""Formatting-transfer facade: paragraph mapping + minimal DOCX formatting restoration.

Thin public entry point for the formatting-transfer wave (spec 033). The mapper cluster
lives in ``formatting_mapping`` (Step 1) and the docx-apply restoration cluster lives in
``formatting_restoration`` (Step 2); both are re-exported here so
``formatting_transfer.<name>`` and ``from ...formatting_transfer import <name>`` keep
resolving for external callers and tests. The dependency direction is one-way
(facade/restoration -> mapping; mapping never imports this module), so there is no cycle.
"""

import logging
from io import BytesIO
from typing import Mapping, Sequence, cast

from docx import Document
from docx.text.paragraph import Paragraph

from docxaicorrector.document.extraction import IMAGE_PLACEHOLDER_PATTERN
from docxaicorrector.generation.formatting_diagnostics_retention import (
    get_formatting_diagnostics_dir,
    write_formatting_diagnostics_artifact,
)
# The mapper cluster lives in formatting_mapping (spec 033, Step 1). These names are
# re-exported so ``formatting_transfer.<name>`` and ``from ...formatting_transfer
# import <name>`` keep resolving for the restoration functions below, for external
# callers, and for the mapper golden — the dependency direction is one-way
# (restoration/facade -> mapping; mapping never imports this module).
from docxaicorrector.generation.formatting_mapping import (  # noqa: F401
    MARKDOWN_HEADING_LINE_PATTERN,
    TEXT_VERIFIED_MAPPING_STRATEGIES,
    _EMPHASIS_ASTERISK_ITALIC_SPAN_PATTERN,
    _EMPHASIS_BOLD_SPAN_PATTERN,
    _EMPHASIS_TRIPLE_SPAN_PATTERN,
    _EMPHASIS_UNDERSCORE_ITALIC_SPAN_PATTERN,
    _ROLE_UNRESOLVED,
    _SYMBOL_ONLY_CARRYOVER_MARKER_PATTERN,
    _TargetRoleResolver,
    _build_caption_heading_conflicts,
    _build_emphasis_coverage_diagnostics,
    _build_generated_registry_by_paragraph_id,
    _build_generated_registry_candidates,
    _build_mapping_text_quality_diagnostics,
    _build_rebuild_key_mapping_quality_diagnostics,
    _build_source_registry_entry,
    _build_target_registry_entry,
    _build_unmapped_source_residual_diagnostics,
    _build_unmapped_source_role_counts,
    _build_unmapped_source_samples,
    _build_unmapped_target_residual_diagnostics,
    _build_unmapped_target_samples,
    _classify_effective_formatting_coverage,
    _classify_residual_closability,
    _classify_unmapped_source_residual,
    _collect_accepted_aggregated_sources,
    _collect_accepted_merged_sources,
    _collect_accepted_split_targets,
    _count_inline_emphasis_spans,
    _count_mapping_strategies,
    _count_output_emphasis,
    _count_relation_id_population,
    _count_source_emphasis,
    _emphasis_retention_ratio,
    _extract_target_heading_level,
    _generated_registry_merged_ids,
    _generated_registry_target_indexes,
    _generated_registry_target_text_compatible,
    _generated_registry_text,
    _is_heading_like_source_paragraph,
    _is_list_source_paragraph,
    _is_toc_source_paragraph,
    _map_source_target_paragraphs,
    _mapping_similarity_score,
    _mapping_text_floor_is_bad,
    _mapping_text_floor_quality,
    _neighbor_candidate_evidence,
    _normalize_text_for_mapping,
    _note_marker_key,
    _occupied_candidate_evidence,
    _paragraph_preview,
    _project_target_index_from_mapped_neighbors,
    _projected_registry_text_floor_satisfied,
    _register_mapping,
    _register_repeated_note_sequence_mappings,
    _registry_candidate_mapping_evidence,
    _registry_mapping_role_compatible,
    _source_format_role,
    _source_paragraph_aggregation_kind,
    _strip_markdown_list_prefixes_for_mapping,
    _target_format_role,
    _target_has_heading_format,
    _target_indexes_containing_any_candidate,
    _target_indexes_containing_candidate,
    _target_paragraph_has_heading_style,
    _target_paragraph_style_name,
    _text_coverage_evidence,
    _token_set,
    _try_register_bounded_registry_mapping,
    _try_register_local_gap_fallback,
    _try_register_projected_registry_mapping,
    _try_register_unique_registry_text_floor_mapping,
)
# The docx-apply restoration cluster lives in formatting_restoration (spec 033, Step 2).
# These names are re-exported so ``apply_output_formatting`` below can orchestrate them
# by bare name, and so ``formatting_transfer.<name>`` / ``from ...formatting_transfer
# import <name>`` keep resolving for external callers and tests. The dependency
# direction is one-way (restoration -> mapping; mapping never imports restoration).
from docxaicorrector.generation.formatting_restoration import (  # noqa: F401
    ALLOWED_CENTERED_QUOTE_STRUCTURAL_ROLES,
    CENTER_SHORT_NON_HEADING_MAX_CHARS,
    CENTER_SHORT_NON_HEADING_MAX_WORDS,
    DISALLOWED_CENTER_SHORT_STRUCTURAL_ROLES,
    _append_numbering_definition,
    _apply_accepted_split_heading_styles,
    _apply_list_numbering_to_paragraph,
    _apply_minimal_caption_formatting,
    _apply_minimal_image_formatting,
    _ensure_paragraph_properties,
    _extract_paragraph_num_id,
    _get_target_numbering_root,
    _has_caption_anchor_context,
    _is_allowlisted_centered_quote_paragraph,
    _is_caption_anchor_block,
    _is_caption_candidate,
    _is_short_centered_attribution_paragraph,
    _is_short_centered_caption_paragraph,
    _is_short_non_heading_paragraph,
    _next_numbering_identifier,
    _replace_paragraph_properties_from_xml,
    _resolve_direct_alignment_restoration_decision,
    _restore_direct_paragraph_alignment_for_mapped_pairs,
    _restore_list_numbering_for_mapped_paragraphs,
    _restore_semantic_quote_formatting_for_mapped_pairs,
    _restore_toc_paragraph_properties_for_mapped_pairs,
    _restore_toc_run_formatting_for_mapped_pairs,
    _sanitize_toc_paragraph_properties_xml,
    _set_direct_paragraph_alignment,
    _style_exists,
)
from docxaicorrector.core.logger import log_event
from docxaicorrector.core.models import ParagraphUnit

FORMATTING_DIAGNOSTICS_DIR = get_formatting_diagnostics_dir()


def _collect_target_paragraphs(document) -> list[Paragraph]:
    return [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() or IMAGE_PLACEHOLDER_PATTERN.search(paragraph.text)
    ]


def _write_formatting_diagnostics_artifact(
    stage: str,
    diagnostics: dict[str, object],
    *,
    run_id: str | None = None,
    source_token: str | None = None,
) -> str | None:
    # Round-11 F1: live ownership needs BOTH identities present and non-blank. A blank
    # identity ("" is not None) used to select "live", which then raised inside the
    # writer and fail-opened WITHOUT writing any artifact at all. Fall back to "offline"
    # so the diagnostic is still retained for explicit replay instead of destroyed.
    has_live_identity = bool((run_id or "").strip()) and bool((source_token or "").strip())
    return write_formatting_diagnostics_artifact(
        stage=stage,
        diagnostics=diagnostics,
        diagnostics_dir=FORMATTING_DIAGNOSTICS_DIR,
        scope="live" if has_live_identity else "offline",
        run_id=run_id,
        source_token=source_token,
    )


# ---------------------------------------------------------------------------
# Formatting preservation and semantic normalization
# ---------------------------------------------------------------------------


def restore_source_formatting(
    docx_bytes: bytes,
    paragraphs: list[ParagraphUnit],
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
    *,
    run_id: str | None = None,
    source_token: str | None = None,
) -> bytes:
    return _restore_source_formatting_impl(
        docx_bytes,
        paragraphs,
        generated_paragraph_registry=generated_paragraph_registry,
        run_id=run_id,
        source_token=source_token,
        mismatch_event_name="paragraph_count_mismatch_restore",
        mismatch_log_message=(
            "Число source/target абзацев не совпадает при unified formatting restore; "
            "применяю только консервативно сопоставленные абзацы."
        ),
    )


def preserve_source_paragraph_properties(
    docx_bytes: bytes,
    paragraphs: list[ParagraphUnit],
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
    *,
    run_id: str | None = None,
    source_token: str | None = None,
) -> bytes:
    """Canonical public formatting entry point for the current transition wave."""
    return apply_output_formatting(
        docx_bytes,
        paragraphs,
        generated_paragraph_registry=generated_paragraph_registry,
        run_id=run_id,
        source_token=source_token,
        mismatch_event_name="paragraph_count_mismatch_preserve",
        mismatch_log_message=(
            "Число source/target абзацев не совпадает при переносе свойств форматирования; "
            "применяю только консервативно сопоставленные абзацы."
        ),
    )


def apply_output_formatting(
    docx_bytes: bytes,
    paragraphs: list[ParagraphUnit],
    *,
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
    run_id: str | None = None,
    source_token: str | None = None,
    mismatch_event_name: str,
    mismatch_log_message: str,
) -> bytes:
    if not docx_bytes or not paragraphs:
        return docx_bytes

    document = Document(BytesIO(docx_bytes))
    target_paragraphs = _collect_target_paragraphs(document)
    relevant_source_paragraphs = [paragraph for paragraph in paragraphs if paragraph.role != "table"]
    mapping_pairs, diagnostics = _map_source_target_paragraphs(
        list(relevant_source_paragraphs),
        list(target_paragraphs),
        generated_paragraph_registry=generated_paragraph_registry,
    )
    unmapped_source_ids = cast(list[str], diagnostics["unmapped_source_ids"])
    unmapped_target_indexes = cast(list[int], diagnostics["unmapped_target_indexes"])

    _apply_minimal_image_formatting(document)
    _apply_minimal_caption_formatting(
        document,
        paragraphs,
        generated_paragraph_registry=generated_paragraph_registry,
    )
    diagnostics["toc_format_restoration_decisions"] = _restore_toc_paragraph_properties_for_mapped_pairs(document, mapping_pairs)
    diagnostics["toc_run_format_restoration_decisions"] = _restore_toc_run_formatting_for_mapped_pairs(mapping_pairs)
    diagnostics["alignment_restoration_decisions"] = _restore_direct_paragraph_alignment_for_mapped_pairs(mapping_pairs)
    _restore_semantic_quote_formatting_for_mapped_pairs(mapping_pairs)

    mismatch_detected = bool(unmapped_source_ids or unmapped_target_indexes)
    if not mismatch_detected:
        _apply_accepted_split_heading_styles(
            document,
            target_paragraphs,
            cast(list[dict[str, object]], diagnostics.get("accepted_split_targets", [])),
            list(relevant_source_paragraphs),
        )

    # Always restore list numbering for successfully mapped paragraph pairs, even when
    # there are unmapped paragraphs.  The mapping_pairs list already contains only the
    # paragraphs that were matched, so applying list formatting to them is always safe.
    # Skipping this entirely on any mismatch was the root cause of lists disappearing
    # whenever the AI added or removed even one paragraph in its output.
    diagnostics["list_restoration_decisions"] = _restore_list_numbering_for_mapped_paragraphs(document, mapping_pairs)

    artifact_path = _write_formatting_diagnostics_artifact(
        "restore",
        diagnostics,
        run_id=run_id,
        source_token=source_token,
    )
    if mismatch_detected:
        log_event(
            logging.WARNING,
            mismatch_event_name,
            mismatch_log_message,
            source_count=diagnostics["source_count"],
            target_count=len(target_paragraphs),
            mapped_count=diagnostics["mapped_count"],
            unmapped_source_count=len(unmapped_source_ids),
            unmapped_target_count=len(unmapped_target_indexes),
            artifact_path=artifact_path,
        )

    if _style_exists(document, "Table Grid"):
        for table in document.tables:
            table.style = "Table Grid"

    output_stream = BytesIO()
    document.save(output_stream)
    return output_stream.getvalue()


def _restore_source_formatting_impl(
    docx_bytes: bytes,
    paragraphs: list[ParagraphUnit],
    *,
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
    run_id: str | None = None,
    source_token: str | None = None,
    mismatch_event_name: str,
    mismatch_log_message: str,
) -> bytes:
    return apply_output_formatting(
        docx_bytes,
        paragraphs,
        generated_paragraph_registry=generated_paragraph_registry,
        run_id=run_id,
        source_token=source_token,
        mismatch_event_name=mismatch_event_name,
        mismatch_log_message=mismatch_log_message,
    )


def _build_output_formatting_diagnostics(
    source_paragraphs: Sequence[ParagraphUnit],
    target_paragraphs: Sequence[Paragraph],
    *,
    document=None,
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
) -> dict[str, object]:
    relevant_source_paragraphs = [paragraph for paragraph in source_paragraphs if paragraph.role != "table"]
    _, diagnostics = _map_source_target_paragraphs(
        list(relevant_source_paragraphs),
        list(target_paragraphs),
        generated_paragraph_registry=generated_paragraph_registry,
    )
    return diagnostics
