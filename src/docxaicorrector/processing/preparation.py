from collections import Counter, OrderedDict
from copy import deepcopy
from dataclasses import asdict, dataclass, field, replace
import hashlib
import inspect
import json
import logging
import os
from io import BytesIO
from pathlib import Path
from threading import Event, Lock
from collections.abc import Callable, Mapping, Sequence
from typing import Any, cast

from docx import Document as DocxDocument

from docxaicorrector.core.config import (
    get_client,
    get_client_for_model_selector,
    get_model_role_value,
    get_provider_config,
    load_app_config,
    load_project_dotenv,
    resolve_model_selector,
)
from docxaicorrector.core.constants import RUN_DIR
from docxaicorrector.document.boundaries import summarize_boundary_normalization_metrics
from docxaicorrector.document.boundary_review import resolve_paragraph_boundary_ai_review_settings
from docxaicorrector.document.extraction import (
    build_document_text,
    extract_document_content_with_normalization_reports,
)
from docxaicorrector.document.semantic_blocks import build_editing_jobs, build_semantic_blocks
from docxaicorrector.core.logger import log_event
from docxaicorrector.core.models import LayoutArtifactCleanupReport, ParagraphBoundaryNormalizationReport, ParagraphRelation, RelationNormalizationReport
from docxaicorrector.core.models import PARAGRAPH_BOUNDARY_AI_REVIEW_MODE_VALUES
from docxaicorrector.core.models import StructureRepairReport
from docxaicorrector.core.models import clone_prepared_image_asset
from docxaicorrector.processing.upload_ports import FrozenUploadPayload, HeartbeatBeacon, build_in_memory_uploaded_file
from docxaicorrector.document.segments import (
    CHAPTER_SEGMENTS_DETECTOR_VERSION,
    DocumentContextProfile,
    DocumentSegment,
    GlossaryTerm,
    SegmentDetectionReport,
    SegmentOutlineEntry,
    build_segment_to_job_mapping,
    detect_document_segments,
    resolve_segment_hard_boundary_paragraph_ids,
    validate_segment_coverage,
)
from docxaicorrector.document.provenance import resolve_scan_origin_config
from docxaicorrector.document.structure_authority import get_effective_structural_role
from docxaicorrector.text.translation_domains import build_terminology_plan, build_translation_domain_instructions


_REASON_LABELS: dict[str, str] = {
    # Preparation-stage first-block composition gate (application_flow humanization).
    "first_block_mixed_toc_and_epigraph": "первый блок смешивает элементы оглавления и эпиграфа",
    "first_block_mixed_toc_and_body_start": "первый блок смешивает элементы оглавления и начало основного текста",
    # Post-translation document-level quality gate (late_phases humanization).
    "untranslated_structural_text_review_required": "структурные элементы остались на исходном языке",
    "untranslated_body_text_review_required": "фрагменты основного текста остались на исходном языке",
    "untranslated_body_text_above_threshold": "слишком большой объём основного текста остался на исходном языке",
    # spec 042 P1-B: caption→heading structural conflict blocks delivery (fatal gate).
    "caption_heading_conflict": "подпись к рисунку или таблице превратилась в заголовок",
}


def humanize_quality_gate_reason(reason: str) -> str:
    normalized = str(reason or "").strip()
    return _REASON_LABELS.get(normalized, normalized.replace("_", " "))


def humanize_quality_gate_reasons(reasons) -> list[str]:
    return [humanize_quality_gate_reason(str(reason).strip()) for reason in reasons or () if str(reason).strip()]


@dataclass
class PreparedDocumentData:
    source_text: str
    paragraphs: list
    image_assets: list
    relations: list[ParagraphRelation]
    jobs: list[dict[str, Any]]
    prepared_source_key: str
    segments: list[DocumentSegment] | None = None
    segment_diagnostics: SegmentDetectionReport = field(default_factory=SegmentDetectionReport)
    structure_fingerprint: str = ""
    detector_version: str = CHAPTER_SEGMENTS_DETECTOR_VERSION
    segment_to_job: dict[str, tuple[int, ...]] | None = None
    source_format: str = "docx"
    conversion_backend: str | None = None
    normalization_report: ParagraphBoundaryNormalizationReport | None = None
    relation_report: RelationNormalizationReport | None = None
    cleanup_report: LayoutArtifactCleanupReport | None = None
    structure_repair_report: StructureRepairReport | None = None
    quality_gate_status: str = "pass"
    quality_gate_reasons: tuple[str, ...] = ()
    translation_domain: str = "general"
    translation_domain_instructions: str = ""
    document_context_profile: DocumentContextProfile = field(default_factory=DocumentContextProfile)
    cached: bool = False


def _build_normalization_metrics(
    normalization_report: ParagraphBoundaryNormalizationReport | None,
    relation_report: RelationNormalizationReport | None = None,
    cleanup_report: LayoutArtifactCleanupReport | None = None,
    structure_repair_report: StructureRepairReport | None = None,
) -> dict[str, int]:
    metrics: dict[str, int] = {}
    if normalization_report is not None:
        metrics.update(
            {
                "raw_paragraph_count": normalization_report.total_raw_paragraphs,
                "logical_paragraph_count": normalization_report.total_logical_paragraphs,
                "merged_group_count": normalization_report.merged_group_count,
                "merged_raw_paragraph_count": normalization_report.merged_raw_paragraph_count,
            }
        )
        metrics.update(summarize_boundary_normalization_metrics(normalization_report))
    if relation_report is not None:
        metrics.update(
            {
                "relation_count": relation_report.total_relations,
                "rejected_relation_candidate_count": relation_report.rejected_candidate_count,
            }
        )
        for relation_kind, count in relation_report.relation_counts.items():
            metrics[f"relation_{relation_kind}_count"] = count
    if cleanup_report is not None:
        metrics.update(flatten_layout_cleanup_metrics(cleanup_report))
    if structure_repair_report is not None:
        metrics.update(flatten_structure_repair_metrics(structure_repair_report))
    return metrics


def flatten_layout_cleanup_metrics(cleanup_report) -> dict[str, int]:
    if cleanup_report is None:
        return {}
    cleanup_mode = str(getattr(cleanup_report, "cleanup_mode", "remove") or "remove").strip().lower()
    if cleanup_mode == "flag":
        return {
            "layout_cleanup_removed_count": int(getattr(cleanup_report, "flagged_page_number_count", 0) or 0)
            + int(getattr(cleanup_report, "flagged_repeated_artifact_count", 0) or 0)
            + int(getattr(cleanup_report, "flagged_empty_or_whitespace_count", 0) or 0),
            "layout_cleanup_page_number_count": int(getattr(cleanup_report, "flagged_page_number_count", 0) or 0),
            "layout_cleanup_repeated_artifact_count": int(
                getattr(cleanup_report, "flagged_repeated_artifact_count", 0) or 0
            ),
            "layout_cleanup_empty_or_whitespace_count": int(
                getattr(cleanup_report, "flagged_empty_or_whitespace_count", 0) or 0
            ),
        }
    return {
        "layout_cleanup_removed_count": int(getattr(cleanup_report, "removed_paragraph_count", 0) or 0),
        "layout_cleanup_page_number_count": int(getattr(cleanup_report, "removed_page_number_count", 0) or 0),
        "layout_cleanup_repeated_artifact_count": int(getattr(cleanup_report, "removed_repeated_artifact_count", 0) or 0),
        "layout_cleanup_empty_or_whitespace_count": int(
            getattr(cleanup_report, "removed_empty_or_whitespace_count", 0) or 0
        ),
    }


def flatten_structure_repair_metrics(structure_repair_report) -> dict[str, int]:
    if structure_repair_report is None:
        return {}
    return {
        "structure_repair_bullet_items": int(getattr(structure_repair_report, "repaired_bullet_items", 0) or 0),
        "structure_repair_numbered_items": int(getattr(structure_repair_report, "repaired_numbered_items", 0) or 0),
        "structure_repair_bounded_toc_regions": int(getattr(structure_repair_report, "bounded_toc_regions", 0) or 0),
        "structure_repair_toc_body_boundary_repairs": int(
            getattr(structure_repair_report, "toc_body_boundary_repairs", 0) or 0
        ),
        "structure_repair_heading_candidates_from_toc": int(
            getattr(structure_repair_report, "heading_candidates_from_toc", 0) or 0
        ),
        "structure_repair_remaining_isolated_markers": int(
            getattr(structure_repair_report, "remaining_isolated_marker_count", 0) or 0
        ),
    }


def _build_preparation_stage_metrics(
    *,
    paragraph_count: int,
    image_count: int,
    normalization_report: ParagraphBoundaryNormalizationReport | None,
    relation_report: RelationNormalizationReport | None,
    cleanup_report: LayoutArtifactCleanupReport | None = None,
    structure_repair_report: StructureRepairReport | None = None,
    source_text: str | None = None,
    block_count: int | None = None,
) -> dict[str, int]:
    metrics = {
        "paragraph_count": paragraph_count,
        "image_count": image_count,
        **_build_normalization_metrics(normalization_report, relation_report, cleanup_report, structure_repair_report),
    }
    if source_text is not None:
        metrics["source_chars"] = len(source_text)
    if block_count is not None:
        metrics["block_count"] = block_count
    return metrics


# Bump whenever paragraph-import / text-layer segmentation logic changes (footnote-marker
# separation, line-fill paragraph boundaries, heading/list segmentation, etc.). It is folded
# into the prepared-source cache key so a full-pipeline run on a previously prepared book is
# invalidated and re-imports with the new logic instead of reusing stale cached structure.
PDF_IMPORT_PARAGRAPH_LOGIC_VERSION = 2


def emit_preparation_progress(progress_callback, *, stage: str, detail: str, progress: float, metrics: dict[str, Any] | None = None) -> None:
    if progress_callback is None:
        return
    progress_callback(stage=stage, detail=detail, progress=progress, metrics=metrics or {})


def _build_source_import_progress(*, source_format: str) -> tuple[str, str]:
    normalized = str(source_format or "docx").strip().lower()
    if normalized == "pdf":
        return (
            "Разбор DOCX (из PDF)",
            "Извлекаю абзацы, встроенные изображения и структуру из сконвертированного DOCX.",
        )
    if normalized == "doc":
        return (
            "Разбор DOCX (из DOC)",
            "Извлекаю абзацы, встроенные изображения и структуру из сконвертированного DOCX.",
        )
    return ("Разбор DOCX", "Извлекаю абзацы и встроенные изображения.")


def _resolve_layout_cleanup_cache_key(app_config: Mapping[str, Any]) -> str:
    if not bool(app_config.get("layout_artifact_cleanup_enabled", True)):
        return "off"
    min_repeat_count = max(2, int(app_config.get("layout_artifact_cleanup_min_repeat_count", 3) or 3))
    max_repeated_text_chars = max(1, int(app_config.get("layout_artifact_cleanup_max_repeated_text_chars", 80) or 80))
    cleanup_mode = str(app_config.get("layout_artifact_cleanup_mode", "flag") or "flag").strip().lower() or "flag"
    return f"1:{min_repeat_count}:{max_repeated_text_chars}:{cleanup_mode}"


def _resolve_scan_origin_cache_key(app_config: Mapping[str, Any] | None) -> str:
    # F2: the document-level scan-origin (OCR) thresholds change which tables are
    # flattened into linear body paragraphs vs. preserved (see
    # ``resolve_scan_origin_config`` / ``classify_document_scan_origin``), so they
    # shape the prepared/cached structure. Resolve them through the SAME resolver
    # the extraction stage uses so the key's scan-origin fingerprint is authoritative
    # and two runs that differ only by a threshold do not share a prepared entry.
    config = resolve_scan_origin_config(app_config)
    return f"{config.multi_column_absolute_min}:{config.multi_column_ratio_min}:{config.authored_uniform_grid_max_ratio}"


_DEFAULT_SCAN_ORIGIN_CACHE_KEY = _resolve_scan_origin_cache_key(None)


def _apply_first_block_composition_quality_gate(
    *,
    blocks: list,
    processing_operation: str,
    quality_gate_status: str,
    quality_gate_reasons: tuple[str, ...],
    structure_phase: str = "post_ai_final",
) -> tuple[str, tuple[str, ...]]:
    if processing_operation != "translate" or not blocks:
        return quality_gate_status, quality_gate_reasons
    first_block_paragraphs = list(getattr(blocks[0], "paragraphs", ()) or ())
    if not first_block_paragraphs:
        return quality_gate_status, quality_gate_reasons

    additional_reasons: list[str] = []
    if _block_has_toc_roles(first_block_paragraphs, structure_phase=structure_phase):
        if _block_has_epigraph_roles(first_block_paragraphs, structure_phase=structure_phase):
            additional_reasons.append("first_block_mixed_toc_and_epigraph")
        if _block_has_body_start_roles(first_block_paragraphs, structure_phase=structure_phase):
            additional_reasons.append("first_block_mixed_toc_and_body_start")
    if not additional_reasons:
        return quality_gate_status, quality_gate_reasons

    merged_reasons = tuple(dict.fromkeys([*quality_gate_reasons, *additional_reasons]))
    return "warning", merged_reasons


def _block_has_toc_roles(paragraphs: list, *, structure_phase: str) -> bool:
    return any(get_effective_structural_role(paragraph, phase=structure_phase) in {"toc_header", "toc_entry"} for paragraph in paragraphs)


def _block_has_epigraph_roles(paragraphs: list, *, structure_phase: str) -> bool:
    return any(
        get_effective_structural_role(paragraph, phase=structure_phase) in {"epigraph", "attribution", "dedication"}
        for paragraph in paragraphs
    )


def _block_has_body_start_roles(paragraphs: list, *, structure_phase: str) -> bool:
    for paragraph in paragraphs:
        role = str(getattr(paragraph, "role", "") or "").strip().lower()
        structural_role = get_effective_structural_role(paragraph, phase=structure_phase)
        if role in {"heading", "body", "list"} and structural_role not in {"toc_header", "toc_entry", "epigraph", "attribution", "dedication"}:
            return True
    return False


def build_prepared_source_key(
    uploaded_file_token: str,
    chunk_size: int,
    *,
    processing_operation: str = "edit",
    paragraph_boundary_normalization_mode: str = "high_only",
    paragraph_boundary_ai_review_mode: str = "off",
    paragraph_boundary_ai_review_model: str = "",
    paragraph_boundary_ai_review_candidate_limit: int = 0,
    paragraph_boundary_ai_review_timeout_seconds: int = 0,
    paragraph_boundary_ai_review_max_tokens_per_candidate: int = 0,
    relation_normalization_key: str = "phase2_default:epigraph_attribution,image_caption,table_caption,toc_region",
    layout_artifact_cleanup_key: str = "1:3:80",
    source_language: str = "en",
    target_language: str = "ru",
    translation_domain: str = "general",
    structure_recovery_enabled: bool = False,
    structure_recovery_mode: str = "legacy",
    scan_origin_key: str = _DEFAULT_SCAN_ORIGIN_CACHE_KEY,
    client_identity: str = "",
) -> str:
    resolved_operation = str(processing_operation or "edit").strip().lower() or "edit"
    operation_suffix = "" if resolved_operation == "edit" else f":op={resolved_operation}"
    normalized_source_language = str(source_language or "en").strip().lower() or "en"
    normalized_target_language = str(target_language or "ru").strip().lower() or "ru"
    normalized_translation_domain = str(translation_domain or "general").strip().lower() or "general"
    normalized_structure_recovery_mode = str(structure_recovery_mode or "legacy").strip().lower() or "legacy"
    structure_recovery_flag = "1" if bool(structure_recovery_enabled) else "0"
    # F10: the boundary AI-review artifact is shaped by MORE than its mode — the
    # structure-recognition MODEL, the candidate limit, the per-candidate timeout, and the
    # per-candidate token budget all change what recommendations get cached (see
    # ``resolve_paragraph_boundary_ai_review_settings``). Fold them into the key so two runs
    # that differ only by AI-review model (or any limit) do not share a prepared entry. They
    # only influence the result when AI review is actually running, so when the mode is "off"
    # a single stable ``ar=off`` token is used to avoid needless cross-invalidation.
    normalized_ai_review_mode = str(paragraph_boundary_ai_review_mode or "off").strip().lower() or "off"
    if normalized_ai_review_mode == "off":
        ai_review_fingerprint = ":ar=off"
    else:
        normalized_ai_review_model = str(paragraph_boundary_ai_review_model or "").strip().lower()
        ai_review_fingerprint = (
            f":arm={normalized_ai_review_model}"
            f":arcl={int(paragraph_boundary_ai_review_candidate_limit or 0)}"
            f":arts={int(paragraph_boundary_ai_review_timeout_seconds or 0)}"
            f":armt={int(paragraph_boundary_ai_review_max_tokens_per_candidate or 0)}"
        )
    # F14: fold every setting that influences the prepared/cached result (languages,
    # translation domain, structure-recovery knobs) into the key so run B cannot serve
    # run A's glossary/context. `pk` is an explicit key-format version tag: bumping it
    # invalidates stale entries whenever this fingerprint layout changes.
    normalized_scan_origin_key = str(scan_origin_key or "").strip() or _DEFAULT_SCAN_ORIGIN_CACHE_KEY
    context_fingerprint = (
        f":pk=4"
        f":sl={normalized_source_language}"
        f":tl={normalized_target_language}"
        f":td={normalized_translation_domain}"
        f":sr={structure_recovery_flag}"
        f":srm={normalized_structure_recovery_mode}"
        f":so={normalized_scan_origin_key}"
        f"{ai_review_fingerprint}"
    )
    # Spec 040: fold a secret-safe client/credential fingerprint into the key so a
    # client-dependent prepared document (AI boundary review) is never served across
    # differing credentials. The segment is appended ONLY when non-empty; when empty the
    # key is byte-identical to the pre-040 output (single-tenant / review-off sharing
    # preserved, existing cache entries not invalidated). ``cid`` is a fresh segment name
    # that cannot collide with any existing token above.
    normalized_client_identity = str(client_identity or "").strip()
    client_identity_suffix = f":cid={normalized_client_identity}" if normalized_client_identity else ""
    return (
        f"{uploaded_file_token}:{chunk_size}:{paragraph_boundary_normalization_mode}:"
        f"{paragraph_boundary_ai_review_mode}:{relation_normalization_key}:lc={layout_artifact_cleanup_key}"
        f"{operation_suffix}"
        f":pv={PDF_IMPORT_PARAGRAPH_LOGIC_VERSION}"
        f"{context_fingerprint}"
        f"{client_identity_suffix}"
    )


def _resolve_prepared_cache_client_identity(
    *,
    resolved_config: Any,
    ai_review_effective_enabled: bool,
    ai_review_model: str,
) -> str:
    # Spec 040: derive a stable, secret-safe fingerprint of the AI-boundary-review client
    # (the ONE client whose output is folded into the prepared/cached document). Returns
    # "" unless AI review is effectively enabled — in every other case the caller must get
    # a byte-identical key so single-tenant / review-off sharing is preserved. The raw
    # api-key VALUE never appears in the output: only sha256(value). Fail-open: any
    # resolution error (misconfig, disabled/unknown provider, missing registry) returns ""
    # rather than raising — a cache-key builder must never blow up the request.
    if not ai_review_effective_enabled:
        return ""
    try:
        resolved_selector = resolve_model_selector(
            str(ai_review_model or "").strip(),
            config_like=resolved_config,
            source_name="paragraph_boundary_ai_review_model",
        )
        provider_config = get_provider_config(resolved_selector.provider, resolved_config)
        canonical_selector = str(resolved_selector.canonical_selector or "")
        base_url = str(provider_config.base_url or "")
        api_key_env = str(provider_config.api_key_env or "")
        # Load project dotenv BEFORE reading the secret, mirroring the client's own secret
        # resolution (core.config._fingerprint_provider_secret): at cache-key time the
        # AI-review client is not built yet (it is created on the cache MISS path), so a
        # bare os.environ read would be empty when the credential lives only in a
        # not-yet-loaded .env — and would then fail to discriminate two tenants (both hash
        # ""). Loading dotenv makes the fingerprint track the credential the client actually
        # resolves. Only sha256(value) enters the key; the raw secret never does.
        load_project_dotenv()
        secret_fingerprint = hashlib.sha256(
            (os.environ.get(api_key_env, "").strip() or "").encode()
        ).hexdigest()
        return hashlib.sha256(
            "\x1f".join([canonical_selector, base_url, api_key_env, secret_fingerprint]).encode()
        ).hexdigest()[:16]
    except Exception:
        return ""


def resolve_prepared_cache_client_identity(app_config: Mapping[str, object] | None) -> str:
    """Public, secret-safe fingerprint of the AI-boundary-review client an injected tenant
    ``client_factory`` resolves for ``app_config`` (spec 041 P1-1).

    The injecting callers (``ProcessingService`` / UI preparation) compute this and pass it
    as ``client_cache_identity=`` (or tag it onto their factory as ``prepared_cache_identity``)
    so the shared preparation cache isolates tenants that share an ``app_config`` but differ
    by credential/endpoint. Returns "" when AI boundary review is effectively disabled (the
    client does not shape the artifact, so shared caching is safe) or on any resolution error
    (fail-open) — an empty identity then drives the safe shared-cache bypass in
    ``prepare_document_for_processing`` instead of a cross-tenant collision. This reuses the
    spec-040 fingerprint logic (provider selector + base_url + api_key_env + sha256(secret)).
    """
    resolved_config = load_app_config() if app_config is None else app_config
    (
        ai_review_effective_enabled,
        _ai_review_mode,
        _ai_review_candidate_limit,
        _ai_review_timeout_seconds,
        _ai_review_max_tokens_per_candidate,
        ai_review_model,
    ) = resolve_paragraph_boundary_ai_review_settings(
        allowed_modes=PARAGRAPH_BOUNDARY_AI_REVIEW_MODE_VALUES,
        app_config=resolved_config,
    )
    return _resolve_prepared_cache_client_identity(
        resolved_config=resolved_config,
        ai_review_effective_enabled=ai_review_effective_enabled,
        ai_review_model=ai_review_model,
    )


def _attach_prepared_job_ids(jobs: Sequence[Mapping[str, Any]]) -> list[dict[str, Any]]:
    prepared_jobs: list[dict[str, Any]] = []
    for index, job in enumerate(jobs):
        normalized_job = dict(job)
        normalized_job["job_id"] = str(normalized_job.get("job_id", "") or "").strip() or f"job_{index:04d}"
        prepared_jobs.append(normalized_job)
    return prepared_jobs


def _build_editing_jobs_with_optional_operation(*, blocks, max_chars: int, processing_operation: str, structure_phase: str):
    signature = inspect.signature(build_editing_jobs)
    accepts_processing_operation = "processing_operation" in signature.parameters
    accepts_structure_phase = "structure_phase" in signature.parameters
    if not accepts_processing_operation and not accepts_structure_phase:
        if str(getattr(build_editing_jobs, "__module__", "")) not in {"document", "document_semantic_blocks"}:
            return build_editing_jobs(blocks, max_chars=max_chars)
        raise RuntimeError("build_editing_jobs must accept processing_operation or structure_phase")
    # F25: kwargs are already gated by inspect.signature, so call exactly once. A
    # genuine internal TypeError must propagate rather than be misread as a signature
    # mismatch (which would silently re-run the target and double its side effects).
    kwargs: dict[str, Any] = {"max_chars": max_chars}
    if accepts_processing_operation:
        kwargs["processing_operation"] = processing_operation
    if accepts_structure_phase:
        kwargs["structure_phase"] = structure_phase
    return build_editing_jobs(blocks, **kwargs)


def _build_semantic_blocks_with_optional_boundaries(*, paragraphs, max_chars: int, relations, hard_boundary_paragraph_ids: set[str], structure_phase: str):
    signature = inspect.signature(build_semantic_blocks)
    accepts_hard_boundaries = "hard_boundary_paragraph_ids" in signature.parameters
    accepts_structure_phase = "structure_phase" in signature.parameters
    # F25: signature-gated kwargs, called exactly once; internal TypeErrors propagate.
    kwargs: dict[str, Any] = {"max_chars": max_chars, "relations": relations}
    if accepts_hard_boundaries:
        kwargs["hard_boundary_paragraph_ids"] = hard_boundary_paragraph_ids
    if accepts_structure_phase:
        kwargs["structure_phase"] = structure_phase
    return build_semantic_blocks(paragraphs, **kwargs)


def _detect_document_segments_with_optional_phase(*, paragraphs, source_content_hash16: str, chunk_size: int, structure_phase: str):
    signature = inspect.signature(detect_document_segments)
    # F25: signature-gated kwargs, called exactly once; internal TypeErrors propagate.
    kwargs: dict[str, Any] = {
        "source_content_hash16": source_content_hash16,
        "chunk_size": chunk_size,
    }
    if "structure_phase" in signature.parameters:
        kwargs["structure_phase"] = structure_phase
    return detect_document_segments(paragraphs, **kwargs)


def _build_document_context_glossary_terms(*, translation_domain: str, source_text: str) -> tuple[GlossaryTerm, ...]:
    terminology_plan = build_terminology_plan(source_text=source_text, translation_domain=translation_domain)
    if not terminology_plan:
        return ()

    glossary_terms: list[GlossaryTerm] = []
    for line in terminology_plan.splitlines():
        normalized_line = str(line or "").strip()
        if not normalized_line or "->" not in normalized_line:
            continue
        source_term, target_term = normalized_line.split("->", 1)
        source_term = source_term.strip()
        target_term = target_term.strip()
        if not source_term or not target_term:
            continue
        glossary_terms.append(
            GlossaryTerm(
                source_term=source_term,
                target_term=target_term,
                confidence="medium",
            )
        )
    return tuple(glossary_terms)


def _extract_docx_detected_author(*, source_bytes: bytes, source_format: str) -> str | None:
    if str(source_format or "").strip().lower() != "docx" or not source_bytes:
        return None
    try:
        document = DocxDocument(BytesIO(source_bytes))
    except Exception:
        return None
    author = str(getattr(document.core_properties, "author", "") or "").strip()
    return author or None


def _build_document_context_profile(
    *,
    segments: Sequence[DocumentSegment],
    translation_domain: str,
    translation_domain_instructions: str,
    source_text: str,
    source_token: str,
    source_title: str,
    detected_author: str | None,
    structure_fingerprint: str,
    source_language: str,
    target_language: str,
) -> DocumentContextProfile:
    outline_entries = tuple(
        SegmentOutlineEntry(
            segment_id=str(getattr(segment, "segment_id", "") or "").strip(),
            title=str(getattr(segment, "title", "") or "").strip(),
            level=max(1, int(getattr(segment, "level", 1) or 1)),
            structural_role=str(getattr(segment, "structural_role", "body_range") or "body_range").strip() or "body_range",
        )
        for segment in segments
        if str(getattr(segment, "segment_id", "") or "").strip() and str(getattr(segment, "title", "") or "").strip()
    )
    glossary_terms = _build_document_context_glossary_terms(
        translation_domain=translation_domain,
        source_text=source_text,
    )
    return DocumentContextProfile(
        source_token=source_token,
        structure_fingerprint=structure_fingerprint,
        source_title=source_title,
        detected_author=detected_author,
        source_language=source_language,
        target_language=target_language,
        translation_domain=translation_domain,
        style_instructions=translation_domain_instructions,
        outline_entries=outline_entries,
        glossary_terms=glossary_terms,
    )


def _supports_segment_detection(paragraphs: Sequence[Any]) -> bool:
    return all(hasattr(paragraph, "text") and hasattr(paragraph, "role") for paragraph in paragraphs)


def _extract_document_content_with_optional_app_config(
    *,
    uploaded_file,
    app_config: Mapping[str, Any],
    client_factory: Callable[[str], object] | None = None,
):
    signature = inspect.signature(extract_document_content_with_normalization_reports)
    parameters = signature.parameters
    accepts_kwargs = any(parameter.kind == inspect.Parameter.VAR_KEYWORD for parameter in parameters.values())
    accepts_app_config = accepts_kwargs or "app_config" in parameters
    # F3: thread the tenant client factory (when supplied) into extraction exactly
    # like app_config is gated here, so a per-tenant factory reaches the boundary-review
    # stage instead of falling back to the global provider client. Injected test doubles
    # that predate the parameter keep working because it is only forwarded when accepted.
    forwards_client_factory = client_factory is not None and (accepts_kwargs or "client_factory" in parameters)
    if forwards_client_factory:
        if accepts_app_config:
            return extract_document_content_with_normalization_reports(
                uploaded_file, app_config=app_config, client_factory=client_factory
            )
        return extract_document_content_with_normalization_reports(uploaded_file, client_factory=client_factory)
    if accepts_app_config:
        return extract_document_content_with_normalization_reports(uploaded_file, app_config=app_config)
    return extract_document_content_with_normalization_reports(uploaded_file)


def _prepare_document_for_processing(
    source_name: str,
    source_bytes: bytes,
    chunk_size: int,
    *,
    source_token: str = "",
    source_format: str = "docx",
    conversion_backend: str | None = None,
    app_config: Mapping[str, Any],
    processing_operation: str = "edit",
    get_client_fn,
    client_factory: Callable[[str], object] | None = None,
    progress_callback=None,
):
    initial_stage, initial_detail = _build_source_import_progress(source_format=source_format)
    emit_preparation_progress(
        progress_callback,
        stage=initial_stage,
        detail=initial_detail,
        progress=0.2,
        metrics={
            "source_format": source_format,
            "conversion_backend": conversion_backend,
        },
    )
    uploaded_file = build_in_memory_uploaded_file(source_name=source_name, source_bytes=source_bytes)
    with HeartbeatBeacon(
        progress_callback,
        stage=initial_stage,
        detail_template=(
            initial_detail
            + " ({elapsed} сек идёт чтение DOCX-архива и извлечение абзацев/изображений.)"
        ),
        progress=0.22,
        metrics={"source_format": source_format, "conversion_backend": conversion_backend},
        interval_seconds=2.0,
    ):
        extraction_result = _extract_document_content_with_optional_app_config(
            uploaded_file=uploaded_file, app_config=app_config, client_factory=client_factory
        )
    paragraphs, image_assets, normalization_report, relations, relation_report, cleanup_report = extraction_result[:6]
    structure_repair_report = extraction_result[6] if len(extraction_result) > 6 else None
    emit_preparation_progress(
        progress_callback,
        stage="Структура извлечена",
        detail="Документ прочитан, собираю текст для анализа.",
        progress=0.3,
        metrics={
            "paragraph_count": len(paragraphs),
            "image_count": len(image_assets),
            "source_format": source_format,
            "conversion_backend": conversion_backend,
            **_build_normalization_metrics(normalization_report, relation_report, cleanup_report, structure_repair_report),
        },
    )
    # Structure recognition (#2) has been removed: preparation flows straight from the
    # importer-provided roles to planning. The downstream structure phase is the deterministic
    # diagnostic phase that consumes importer advisory hints (the former production path with
    # structure recognition disabled).
    quality_gate_status: str = "pass"
    quality_gate_reasons: tuple[str, ...] = ()
    downstream_structure_phase = "pre_ai_diagnostic"
    if _supports_segment_detection(paragraphs):
        source_content_hash16 = hashlib.sha256(source_bytes).hexdigest()[:16]
        segments, segment_diagnostics, structure_fingerprint = _detect_document_segments_with_optional_phase(
            paragraphs=paragraphs,
            source_content_hash16=source_content_hash16,
            chunk_size=chunk_size,
            structure_phase=downstream_structure_phase,
        )
        for segment in segments:
            for index in range(segment.start_paragraph_index, segment.end_paragraph_index + 1):
                paragraph = paragraphs[index]
                paragraph.segment_id = segment.segment_id
                paragraph.segment_level = segment.level
                if index == segment.start_paragraph_index:
                    paragraph.segment_boundary_before = segment.ordinal > 1
    else:
        segments = []
        segment_diagnostics = SegmentDetectionReport()
        structure_fingerprint = ""
    source_text = build_document_text(paragraphs)
    translation_domain = str(app_config.get("translation_domain_default", "general") or "general").strip().lower() or "general"
    translation_domain_instructions = build_translation_domain_instructions(
        translation_domain=translation_domain,
        source_text=source_text,
    )
    detected_author = _extract_docx_detected_author(source_bytes=source_bytes, source_format=source_format)
    document_context_profile = _build_document_context_profile(
        segments=segments,
        translation_domain=translation_domain,
        translation_domain_instructions=translation_domain_instructions,
        source_text=source_text,
        source_token=str(source_token or "").strip(),
        source_title=Path(str(source_name or "")).stem,
        detected_author=detected_author,
        structure_fingerprint=structure_fingerprint,
        source_language=str(app_config.get("source_language", app_config.get("source_language_default", "en")) or "en").strip().lower() or "en",
        target_language=str(app_config.get("target_language", app_config.get("target_language_default", "ru")) or "ru").strip().lower() or "ru",
    )
    emit_preparation_progress(
        progress_callback,
        stage="Текст собран",
        detail="Формирую цельный текст документа и считаю объём.",
        progress=0.6,
        metrics={
            **_build_preparation_stage_metrics(
                paragraph_count=len(paragraphs),
                image_count=len(image_assets),
                normalization_report=normalization_report,
                relation_report=relation_report,
                cleanup_report=cleanup_report,
                structure_repair_report=structure_repair_report,
                source_text=source_text,
            ),
            "source_format": source_format,
            "conversion_backend": conversion_backend,
        },
    )
    downstream_relations = relations
    hard_boundary_paragraph_ids = resolve_segment_hard_boundary_paragraph_ids(segments)
    blocks = _build_semantic_blocks_with_optional_boundaries(
        paragraphs=paragraphs,
        max_chars=chunk_size,
        relations=downstream_relations,
        hard_boundary_paragraph_ids=hard_boundary_paragraph_ids,
        structure_phase=downstream_structure_phase,
    )
    quality_gate_status, quality_gate_reasons = _apply_first_block_composition_quality_gate(
        blocks=blocks,
        processing_operation=processing_operation,
        quality_gate_status=quality_gate_status,
        quality_gate_reasons=quality_gate_reasons,
        structure_phase=downstream_structure_phase,
    )
    log_event(
        logging.INFO,
        "preparation_outcome",
        "Определён итог подготовки документа (роли из импорта).",
        quality_gate_status=quality_gate_status,
        quality_gate_reasons=list(quality_gate_reasons),
        first_block_has_toc=(
            _block_has_toc_roles(
                list(getattr(blocks[0], "paragraphs", ()) or []),
                structure_phase=downstream_structure_phase,
            )
            if blocks
            else False
        ),
        first_block_has_epigraph=(
            _block_has_epigraph_roles(
                list(getattr(blocks[0], "paragraphs", ()) or []),
                structure_phase=downstream_structure_phase,
            )
            if blocks
            else False
        ),
        first_block_has_body_start=(
            _block_has_body_start_roles(
                list(getattr(blocks[0], "paragraphs", ()) or []),
                structure_phase=downstream_structure_phase,
            )
            if blocks
            else False
        ),
        **flatten_layout_cleanup_metrics(cleanup_report),
        **flatten_structure_repair_metrics(structure_repair_report),
    )
    emit_preparation_progress(
        progress_callback,
        stage="Смысловые блоки",
        detail="Группирую абзацы в блоки для модели.",
        progress=0.75,
        metrics={
            **_build_preparation_stage_metrics(
                paragraph_count=len(paragraphs),
                image_count=len(image_assets),
                normalization_report=normalization_report,
                relation_report=relation_report,
                cleanup_report=cleanup_report,
                structure_repair_report=structure_repair_report,
                source_text=source_text,
                block_count=len(blocks),
            ),
            "source_format": source_format,
            "conversion_backend": conversion_backend,
        },
    )
    jobs = _build_editing_jobs_with_optional_operation(
        blocks=blocks,
        max_chars=chunk_size,
        processing_operation=processing_operation,
        structure_phase=downstream_structure_phase,
    )
    jobs = _attach_prepared_job_ids(jobs)
    if _supports_segment_detection(paragraphs):
        segment_to_job = build_segment_to_job_mapping(segments, jobs)
        coverage_warnings = validate_segment_coverage(
            paragraphs=paragraphs,
            segments=segments,
            jobs=jobs,
            segment_to_job=segment_to_job,
        )
        if coverage_warnings:
            segment_diagnostics = replace(
                segment_diagnostics,
                warnings=tuple(dict.fromkeys((*segment_diagnostics.warnings, *coverage_warnings))),
            )
    else:
        segment_to_job = {}
    emit_preparation_progress(
        progress_callback,
        stage="Задания собраны",
        detail="Готовлю финальный набор задач для обработки.",
        progress=0.9,
        metrics={
            **_build_preparation_stage_metrics(
                paragraph_count=len(paragraphs),
                image_count=len(image_assets),
                normalization_report=normalization_report,
                relation_report=relation_report,
                cleanup_report=cleanup_report,
                structure_repair_report=structure_repair_report,
                source_text=source_text,
                block_count=len(jobs),
            ),
            "source_format": source_format,
            "conversion_backend": conversion_backend,
        },
    )
    return PreparedDocumentData(
        source_text=source_text,
        paragraphs=paragraphs,
        image_assets=image_assets,
        relations=relations,
        jobs=jobs,
        segments=segments,
        segment_diagnostics=segment_diagnostics,
        structure_fingerprint=structure_fingerprint,
        detector_version=CHAPTER_SEGMENTS_DETECTOR_VERSION,
        segment_to_job=segment_to_job,
        prepared_source_key="",
        source_format=source_format,
        conversion_backend=conversion_backend,
        normalization_report=normalization_report,
        relation_report=relation_report,
        cleanup_report=cleanup_report,
        structure_repair_report=structure_repair_report,
        quality_gate_status=quality_gate_status,
        quality_gate_reasons=quality_gate_reasons,
        translation_domain=translation_domain,
        translation_domain_instructions=translation_domain_instructions,
        document_context_profile=document_context_profile,
        cached=False,
    )


PREPARATION_CACHE_LIMIT = 2
_shared_preparation_cache: OrderedDict[str, PreparedDocumentData] = OrderedDict()
_shared_preparation_cache_lock = Lock()
_shared_preparation_inflight: dict[str, Event] = {}


def _get_preparation_cache(session_state) -> dict[str, PreparedDocumentData]:
    if session_state is None:
        return {}
    cache = session_state.get("preparation_cache")
    if not isinstance(cache, dict):
        cache = {}
        session_state["preparation_cache"] = cache
    return cache


def _touch_cache_entry(cache: dict[str, PreparedDocumentData], prepared_source_key: str, prepared_document: PreparedDocumentData) -> None:
    cache.pop(prepared_source_key, None)
    cache[prepared_source_key] = prepared_document


def _trim_cache(cache: dict[str, PreparedDocumentData]) -> None:
    while len(cache) > PREPARATION_CACHE_LIMIT:
        oldest_key = next(iter(cache))
        cache.pop(oldest_key, None)


def _read_cache_entry(cache: dict[str, PreparedDocumentData], prepared_source_key: str):
    cached = cache.get(prepared_source_key)
    if cached is None:
        return None
    _touch_cache_entry(cache, prepared_source_key, cached)
    return cached


def _clone_prepared_document(data: PreparedDocumentData, prepared_source_key: str, *, cached: bool) -> PreparedDocumentData:
    return PreparedDocumentData(
        source_text=data.source_text,
        paragraphs=deepcopy(data.paragraphs),
        image_assets=[clone_prepared_image_asset(asset) for asset in data.image_assets],
        relations=deepcopy(data.relations),
        jobs=[dict(job) for job in data.jobs],
        segments=deepcopy(data.segments),
        segment_diagnostics=deepcopy(data.segment_diagnostics),
        structure_fingerprint=data.structure_fingerprint,
        detector_version=data.detector_version,
        segment_to_job=deepcopy(data.segment_to_job),
        prepared_source_key=prepared_source_key,
        normalization_report=deepcopy(data.normalization_report),
        relation_report=deepcopy(data.relation_report),
        cleanup_report=deepcopy(data.cleanup_report),
        structure_repair_report=deepcopy(data.structure_repair_report),
        quality_gate_status=data.quality_gate_status,
        quality_gate_reasons=tuple(data.quality_gate_reasons),
        source_format=data.source_format,
        conversion_backend=data.conversion_backend,
        translation_domain=data.translation_domain,
        translation_domain_instructions=data.translation_domain_instructions,
        document_context_profile=data.document_context_profile,
        cached=cached,
    )


def _read_or_reserve_cached_prepared_document(*, session_state, prepared_source_key: str, allow_shared_cache: bool = True):
    # Session cache is only touched from the Streamlit rerun thread. Background preparation
    # workers always pass session_state=None and only participate in the shared cache path.
    session_cache = _get_preparation_cache(session_state) if session_state is not None else None
    if session_cache is not None:
        cached = _read_cache_entry(session_cache, prepared_source_key)
        if cached is not None:
            return _clone_prepared_document(cached, prepared_source_key, cached=True), None, "session"

    if not allow_shared_cache:
        # Spec 041 P1-1: the shared (process-global) tier is bypassed for this run — a
        # client-dependent artifact must not cross an unknown-identity boundary. The
        # per-session tier read above stays active (session-scoped, tenant-safe). Returning
        # no in-flight reservation means the caller rebuilds and (per the matching store
        # guard) does not publish to the shared tier.
        return None, None, None

    while True:
        with _shared_preparation_cache_lock:
            cached = _read_cache_entry(_shared_preparation_cache, prepared_source_key)
            if cached is not None:
                if session_cache is not None:
                    _touch_cache_entry(session_cache, prepared_source_key, cached)
                    _trim_cache(session_cache)
                return _clone_prepared_document(cached, prepared_source_key, cached=True), None, "shared"

            in_flight = _shared_preparation_inflight.get(prepared_source_key)
            if in_flight is None:
                in_flight = Event()
                _shared_preparation_inflight[prepared_source_key] = in_flight
                return None, in_flight, None

        in_flight.wait()


def _release_shared_preparation(prepared_source_key: str) -> None:
    with _shared_preparation_cache_lock:
        in_flight = _shared_preparation_inflight.pop(prepared_source_key, None)
    if in_flight is not None:
        in_flight.set()


def _store_cached_prepared_document(*, session_state, prepared_source_key: str, prepared_document: PreparedDocumentData, allow_shared_cache: bool = True) -> None:
    prepared_document.prepared_source_key = ""
    prepared_document.cached = False
    if session_state is not None:
        cache = _get_preparation_cache(session_state)
        _touch_cache_entry(cache, prepared_source_key, prepared_document)
        _trim_cache(cache)

    if not allow_shared_cache:
        # Spec 041 P1-1: mirror the read guard — only the shared (process-global) tier is
        # skipped; the per-session store above still runs so a session keeps serving its own
        # prepared document across reruns.
        return
    with _shared_preparation_cache_lock:
        _touch_cache_entry(_shared_preparation_cache, prepared_source_key, prepared_document)
        _trim_cache(_shared_preparation_cache)


def clear_preparation_cache(*, session_state=None, clear_shared: bool = False) -> None:
    if session_state is not None:
        session_state["preparation_cache"] = {}
    if clear_shared:
        with _shared_preparation_cache_lock:
            _shared_preparation_cache.clear()


def prepare_document_for_processing(
    *,
    uploaded_payload: FrozenUploadPayload,
    chunk_size: int,
    app_config: dict[str, Any] | None = None,
    processing_operation: str | None = None,
    session_state=None,
    get_client_fn=None,
    client_factory: Callable[[str], object] | None = None,
    client_cache_identity: str | None = None,
    progress_callback=None,
) -> PreparedDocumentData:
    resolved_config = load_app_config() if app_config is None else app_config
    resolved_get_client_fn = get_client if get_client_fn is None else get_client_fn
    resolved_processing_operation = str(
        processing_operation if processing_operation is not None else resolved_config.get("processing_operation", "edit")
    ).strip().lower() or "edit"
    normalization_mode = (
        str(resolved_config["paragraph_boundary_normalization_mode"])
        if bool(resolved_config["paragraph_boundary_normalization_enabled"])
        else "off"
    )
    # F10: resolve the AI-review settings through the SAME resolver the extraction stage
    # uses, so the key's model/limit fingerprint is authoritative (mode gated to "off"
    # when disabled; model/limits are the exact values that shape the cached artifact).
    (
        ai_review_effective_enabled,
        ai_review_mode,
        ai_review_candidate_limit,
        ai_review_timeout_seconds,
        ai_review_max_tokens_per_candidate,
        ai_review_model,
    ) = resolve_paragraph_boundary_ai_review_settings(
        allowed_modes=PARAGRAPH_BOUNDARY_AI_REVIEW_MODE_VALUES,
        app_config=resolved_config,
    )
    relation_normalization_key = "off"
    if bool(resolved_config.get("relation_normalization_enabled", True)):
        relation_profile = str(resolved_config.get("relation_normalization_profile", "phase2_default"))
        configured_relation_kinds = resolved_config.get("relation_normalization_enabled_relation_kinds", ())
        if not isinstance(configured_relation_kinds, (list, tuple, set)):
            configured_relation_kinds = ()
        enabled_relation_kinds = ",".join(
            sorted(str(kind) for kind in configured_relation_kinds)
        )
        relation_normalization_key = f"{relation_profile}:{enabled_relation_kinds}"
    layout_cleanup_key = _resolve_layout_cleanup_cache_key(resolved_config)
    key_source_language = str(
        resolved_config.get("source_language", resolved_config.get("source_language_default", "en")) or "en"
    ).strip().lower() or "en"
    key_target_language = str(
        resolved_config.get("target_language", resolved_config.get("target_language_default", "ru")) or "ru"
    ).strip().lower() or "ru"
    key_translation_domain = str(
        resolved_config.get("translation_domain_default", "general") or "general"
    ).strip().lower() or "general"
    key_structure_recovery_enabled = bool(resolved_config.get("structure_recovery_enabled", False))
    key_structure_recovery_mode = str(
        resolved_config.get("structure_recovery_mode", "legacy") or "legacy"
    ).strip().lower() or "legacy"
    key_scan_origin = _resolve_scan_origin_cache_key(resolved_config)
    # Spec 041 P1-1: the shared (process-global) preparation cache must isolate the tenant
    # client_factory that is ACTUALLY injected, not just the config-derived client — two
    # callers with the same app_config but different factories/credentials must never share
    # one client-dependent (AI-boundary-review) entry. When no explicit identity is supplied,
    # fall back to one carried on the factory object: the UI background path threads its
    # factory (not a param) through intermediaries this change does not touch, so it tags the
    # factory with `.prepared_cache_identity`.
    if client_cache_identity is None and client_factory is not None:
        factory_identity = getattr(client_factory, "prepared_cache_identity", None)
        client_cache_identity = str(factory_identity) if factory_identity is not None else None
    if client_factory is None:
        # Config-default path: the AI-boundary-review client is config-derived, so the
        # config-derived fingerprint (spec 040) is authoritative. Byte-identical key and
        # shared caching are preserved exactly.
        prepared_cache_client_identity = _resolve_prepared_cache_client_identity(
            resolved_config=resolved_config,
            ai_review_effective_enabled=ai_review_effective_enabled,
            ai_review_model=ai_review_model,
        )
        allow_shared_cache = True
    elif not ai_review_effective_enabled:
        # An injected factory does not shape the prepared artifact when AI boundary review is
        # off, so the shared cache is safe and the client identity stays "" (unchanged key).
        prepared_cache_client_identity = ""
        allow_shared_cache = True
    else:
        # Injected factory + AI review ON: fold a caller-supplied, secret-safe identity into
        # the key so distinct tenants get distinct shared entries. With NO usable identity,
        # never serve a client-dependent artifact across an unknown boundary — bypass the
        # shared tier for this run (the per-session tier stays available and tenant-safe).
        resolved_client_identity = str(client_cache_identity or "").strip()
        if resolved_client_identity:
            prepared_cache_client_identity = resolved_client_identity
            allow_shared_cache = True
        else:
            prepared_cache_client_identity = ""
            allow_shared_cache = False
    prepared_source_key = build_prepared_source_key(
        uploaded_payload.file_token,
        chunk_size,
        processing_operation=resolved_processing_operation,
        paragraph_boundary_normalization_mode=normalization_mode,
        paragraph_boundary_ai_review_mode=ai_review_mode,
        paragraph_boundary_ai_review_model=ai_review_model,
        paragraph_boundary_ai_review_candidate_limit=ai_review_candidate_limit,
        paragraph_boundary_ai_review_timeout_seconds=ai_review_timeout_seconds,
        paragraph_boundary_ai_review_max_tokens_per_candidate=ai_review_max_tokens_per_candidate,
        relation_normalization_key=relation_normalization_key,
        layout_artifact_cleanup_key=layout_cleanup_key,
        source_language=key_source_language,
        target_language=key_target_language,
        translation_domain=key_translation_domain,
        structure_recovery_enabled=key_structure_recovery_enabled,
        structure_recovery_mode=key_structure_recovery_mode,
        scan_origin_key=key_scan_origin,
        client_identity=prepared_cache_client_identity,
    )
    cached, in_flight, cache_level = _read_or_reserve_cached_prepared_document(
        session_state=session_state,
        prepared_source_key=prepared_source_key,
        allow_shared_cache=allow_shared_cache,
    )
    if cached is not None:
        log_event(
            logging.INFO,
            "preparation_cache_hit",
            "Использован кэш подготовки документа.",
            prepared_source_key=prepared_source_key,
            cache_level=cache_level,
        )
        emit_preparation_progress(
            progress_callback,
            stage="Подготовка документа",
            detail="Использую кэш подготовки для текущего файла.",
            progress=0.95,
            metrics={
                "paragraph_count": len(cached.paragraphs),
                "image_count": len(cached.image_assets),
                "source_chars": len(cached.source_text),
                "block_count": len(cached.jobs),
                "cached": cached.cached,
                "source_format": cached.source_format,
                "conversion_backend": cached.conversion_backend,
                **_build_normalization_metrics(
                    cached.normalization_report,
                    cached.relation_report,
                    cached.cleanup_report,
                    cached.structure_repair_report,
                ),
            },
        )
        return cached

    log_event(
        logging.INFO,
        "preparation_cache_miss",
        "Подготовка документа выполняется без готового cache-hit.",
        prepared_source_key=prepared_source_key,
    )

    try:
        prepared_document = _prepare_document_for_processing(
            uploaded_payload.filename,
            uploaded_payload.content_bytes,
            chunk_size,
            source_token=str(uploaded_payload.file_token or "").strip(),
            source_format=str(getattr(uploaded_payload, "source_format", "docx") or "docx"),
            conversion_backend=getattr(uploaded_payload, "conversion_backend", None),
            app_config=resolved_config,
            processing_operation=resolved_processing_operation,
            get_client_fn=resolved_get_client_fn,
            client_factory=client_factory,
            progress_callback=progress_callback,
        )
        _store_cached_prepared_document(
            session_state=session_state,
            prepared_source_key=prepared_source_key,
            prepared_document=prepared_document,
            allow_shared_cache=allow_shared_cache,
        )
    except Exception:
        if in_flight is not None:
            _release_shared_preparation(prepared_source_key)
        raise

    if in_flight is not None:
        _release_shared_preparation(prepared_source_key)
    return _clone_prepared_document(prepared_document, prepared_source_key, cached=False)
