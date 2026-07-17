import hashlib
import logging
import multiprocessing
import os
import pickle
import queue
import signal
import shutil
import subprocess
import tempfile
import threading
import time
from collections import OrderedDict
from dataclasses import dataclass, replace
from io import BytesIO
from pathlib import Path
from collections.abc import Callable, Mapping
from typing import cast
from uuid import uuid4

import streamlit as st

from docxaicorrector.core.logger import log_event
from docxaicorrector.processing.restart_store import clear_restart_source, load_restart_source_bytes, store_completed_source, store_restart_source
from docxaicorrector.runtime.state import (
    apply_processing_start,
    apply_preparation_complete,
    apply_preparation_failure,
    apply_preparation_stop,
    apply_processing_completion,
    get_latest_audiobook_postprocess_enabled,
    get_processing_event_queue,
    get_latest_processing_operation,
    get_processing_stop_event,
    get_processing_worker,
    get_preparation_event_queue,
    get_preparation_worker,
    get_latest_narration_text,
    get_latest_source_name,
    get_latest_source_token,
    get_restart_source,
    mark_preparation_started,
    request_processing_stop as request_processing_stop_via_state,
    reset_image_state,
    set_preparation_runtime,
    set_restart_source,
)
from docxaicorrector.runtime.events import (
    AppendImageLogEvent,
    AppendLogEvent,
    FinalizeProcessingStatusEvent,
    PreparationCompleteEvent,
    PreparationFailedEvent,
    PreparationStoppedEvent,
    ProcessingEvent,
    PushActivityEvent,
    ResetImageStateEvent,
    SetProcessingStatusEvent,
    SetStateEvent,
    WorkerCompleteEvent,
)
from docxaicorrector.runtime.workflow_state import ProcessingOutcome
from docxaicorrector.processing.upload_ports import (  # noqa: F401
    _DEFAULT_UPLOADED_FILENAME,
    FrozenUploadPayload,
    HeartbeatBeacon,
    InMemoryUploadedFile,
    UploadedFileLike,
    _looks_like_runtime_object_repr,
    build_in_memory_uploaded_file,
    read_uploaded_file_bytes,
    resolve_uploaded_filename,
)
from docxaicorrector.processing.service_ports import (  # noqa: F401
    normalize_background_error,
    should_stop_processing,
)


MAX_COMPLETED_SOURCE_BYTES = 8 * 1024 * 1024
_DOCX_ZIP_MAGIC = b"PK\x03\x04"
_LEGACY_DOC_MAGIC = bytes.fromhex("D0CF11E0A1B11AE1")
_PDF_MAGIC = b"%PDF-"
_DOC_CONVERSION_TIMEOUT_SECONDS = 120
_MATERIALIZED_UPLOAD_CACHE_LIMIT = 4

# ---------------------------------------------------------------------------
# F7: conservative resource budgets for the in-process pdfminer PDF parse.
# Normal documents are well within these limits; over-budget documents fail
# fast with a typed ``pdf_import_over_budget:*`` error (never silently
# truncated) so a pathological upload cannot exhaust RAM/CPU in-process.
# ---------------------------------------------------------------------------
_MAX_PDF_IMPORT_FILE_BYTES = 64 * 1024 * 1024  # 64 MiB
_MAX_PDF_IMPORT_PAGE_COUNT = 2000
# Object-count caps (F7): reject a pathological document whose extracted spans /
# images would exhaust RAM. These are deliberately generous (a real 2000-page
# book stays well under them); over-cap uploads fail fast with a typed
# ``pdf_import_over_budget:*`` error and an ERROR log (never silent truncation).
_MAX_PDF_IMPORT_SPAN_COUNT = 2_000_000
_MAX_PDF_IMPORT_IMAGE_COUNT = 20_000
# Single unified wall-clock deadline covering ALL heavy pdfminer/OCR stages
# (page count + span extract + optional OCR reparse + image extract). In the
# production path this bounds a killable child process; the in-process test
# fallback bounds a daemon thread (which cannot be force-killed).
_PDF_PARSE_WALLCLOCK_BUDGET_SECONDS = 300  # 5 minutes around the whole parse

# F7 test seam: the production parse runs in a spawned child process so an
# over-budget document can be genuinely TERMINATED. Existing tests monkeypatch
# in-process pdfminer functions, which a child process cannot see, so they set
# this flag (or ``DOCXAI_PDF_PARSE_IN_PROCESS=1``) to run the SAME worker logic
# in-process under the thread guard. The in-process path is test-only: it cannot
# force-kill an overrunning parse.
_PDF_PARSE_IN_PROCESS = False


def _pdf_parse_in_process_enabled() -> bool:
    if _PDF_PARSE_IN_PROCESS:
        return True
    return os.getenv("DOCXAI_PDF_PARSE_IN_PROCESS", "").strip().lower() in {"1", "true", "yes", "on"}

# ---------------------------------------------------------------------------
# F27: process-wide admission gate. Each background processing worker acquires a
# slot before doing real work so N concurrent Streamlit sessions cannot multiply
# PDF RAM / subprocess / API cost without bound. Single-session behaviour is
# unchanged (default limit >= 1); override via the env var.
# ---------------------------------------------------------------------------
_PROCESSING_ADMISSION_LIMIT_ENV = "DOCXAI_MAX_CONCURRENT_PROCESSING"
_DEFAULT_PROCESSING_ADMISSION_LIMIT = 2


def _resolve_processing_admission_limit() -> int:
    raw_value = os.getenv(_PROCESSING_ADMISSION_LIMIT_ENV, "")
    try:
        limit = int(str(raw_value).strip())
    except (TypeError, ValueError):
        return _DEFAULT_PROCESSING_ADMISSION_LIMIT
    return limit if limit >= 1 else _DEFAULT_PROCESSING_ADMISSION_LIMIT


def _build_processing_admission_gate(limit: int) -> threading.BoundedSemaphore:
    return threading.BoundedSemaphore(max(1, int(limit)))


def _acquire_admission_slot_cancellable(
    gate: threading.Semaphore,
    stop_event: threading.Event | None,
    *,
    poll_seconds: float = 0.1,
) -> bool:
    """Acquire an admission slot while honouring ``stop_event`` (F27).

    Blocks until a slot is free, but polls so a stopped upload does not wait on
    a slot forever. Returns ``True`` once a slot is held, ``False`` if the wait
    was cancelled (caller must NOT release in that case).
    """

    while True:
        if stop_event is not None and stop_event.is_set():
            return False
        if gate.acquire(timeout=poll_seconds):
            return True


_PROCESSING_ADMISSION_GATE = _build_processing_admission_gate(_resolve_processing_admission_limit())
_ALLOWED_SET_STATE_EVENT_KEYS = {
    "final_generated_paragraph_registry",
    "image_assets",
    "last_background_error",
    "last_error",
    "latest_docx_bytes",
    "latest_markdown",
    "latest_narration_text",
    "latest_marker_diagnostics_artifact",
    "latest_quality_warning",
    "latest_result_notice",
    "processed_block_markdowns",
    "processed_paragraph_registry",
}

_shared_materialized_upload_cache: OrderedDict[str, "FrozenUploadPayload"] = OrderedDict()
_shared_materialized_upload_cache_lock = threading.Lock()
_shared_materialized_upload_inflight: dict[str, threading.Event] = {}

__all__ = [
    "BackgroundRuntime",
    "FrozenUploadPayload",
    "NormalizedUploadedDocument",
    "UploadSourceIdentity",
    "ResolvedUploadContract",
    "InMemoryUploadedFile",
    "RuntimeEventEmitterDependencies",
    "RuntimeEventEmitters",
    "legacy_doc_conversion_available",
    "read_uploaded_file_bytes",
    "normalize_uploaded_document",
    "resolve_upload_contract",
    "freeze_resolved_upload",
    "freeze_uploaded_file",
    "freeze_uploaded_file_lightweight",
    "materialize_uploaded_payload",
    "HeartbeatBeacon",
    "build_uploaded_file_token",
    "build_in_memory_uploaded_file",
    "build_uploaded_file_selection_marker",
    "build_preparation_request_marker",
    "normalize_background_error",
    "build_result_bundle",
    "should_cache_completed_source",
    "get_current_result_bundle",
    "emit_or_apply_state",
    "emit_or_apply_image_reset",
    "emit_or_apply_status",
    "emit_or_apply_finalize",
    "emit_or_apply_activity",
    "emit_or_apply_log",
    "emit_or_apply_image_log",
    "build_runtime_event_emitters",
    "should_stop_processing",
    "resolve_uploaded_filename",
    "drain_processing_events",
    "drain_preparation_events",
    "preparation_worker_is_active",
    "processing_worker_is_active",
    "request_processing_stop",
    "start_background_processing",
    "start_background_preparation",
]


class BackgroundRuntime:
    def __init__(self, event_queue, stop_event, source_token: str = ""):
        self._event_queue = event_queue
        self._stop_event = stop_event
        self._source_token = str(source_token or "")

    def emit(self, event: ProcessingEvent) -> None:
        event_source_token = getattr(event, "source_token", None)
        if self._source_token and isinstance(event_source_token, str) and not event_source_token:
            event = replace(event, source_token=self._source_token)
        self._event_queue.put(event)

    def should_stop(self) -> bool:
        return self._stop_event.is_set()


@dataclass(frozen=True)
class NormalizedUploadedDocument:
    original_filename: str
    filename: str
    content_bytes: bytes
    source_format: str
    conversion_backend: str | None


@dataclass(frozen=True)
class UploadSourceIdentity:
    original_filename: str
    source_bytes: bytes
    token_size: int
    token_hash: str


@dataclass(frozen=True)
class ResolvedUploadContract:
    source_identity: UploadSourceIdentity
    normalized_document: NormalizedUploadedDocument

    @property
    def filename(self) -> str:
        return self.normalized_document.filename

    @property
    def content_bytes(self) -> bytes:
        return self.normalized_document.content_bytes

    @property
    def file_size(self) -> int:
        return len(self.normalized_document.content_bytes)

    @property
    def content_hash(self) -> str:
        return hashlib.sha256(self.normalized_document.content_bytes).hexdigest()[:16]

    @property
    def file_token(self) -> str:
        return f"{self.filename}:{self.source_identity.token_size}:{self.source_identity.token_hash}"


def _detect_uploaded_document_format(*, filename: str, source_bytes: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if source_bytes.startswith(_DOCX_ZIP_MAGIC):
        return "docx"
    if source_bytes.startswith(_LEGACY_DOC_MAGIC):
        return "doc" if suffix == ".doc" else "unknown"
    if source_bytes.startswith(_PDF_MAGIC):
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix == ".doc":
        return "doc"
    if suffix == ".pdf":
        return "pdf"
    return "unknown"


def _build_normalized_docx_filename(filename: str) -> str:
    path = Path(filename or "document")
    if path.suffix.lower() == ".docx":
        return path.name or "document.docx"
    if path.suffix:
        return str(path.with_suffix(".docx"))
    if path.name:
        return f"{path.name}.docx"
    return "document.docx"


def _run_completed_process(
    command: list[str],
    *,
    error_message: str,
    text: bool = True,
    timeout_seconds: int = _DOC_CONVERSION_TIMEOUT_SECONDS,
    cleanup_process_group: bool = False,
):
    if cleanup_process_group:
        process = None
        try:
            popen_kwargs = {
                "stdout": subprocess.PIPE,
                "stderr": subprocess.PIPE,
                "text": text,
            }
            if text:
                popen_kwargs["encoding"] = "utf-8"
                popen_kwargs["errors"] = "replace"
            if os.name == "nt":
                popen_kwargs["creationflags"] = subprocess.CREATE_NEW_PROCESS_GROUP
            else:
                popen_kwargs["start_new_session"] = True
            process = subprocess.Popen(command, **popen_kwargs)
            stdout, stderr = process.communicate(timeout=timeout_seconds)
        except subprocess.TimeoutExpired as exc:
            if process is not None:
                _terminate_process_tree(process)
                process.communicate()
            raise RuntimeError(
                f"{error_message} Превышено время ожидания {timeout_seconds} сек."
            ) from exc
        except OSError as exc:
            raise RuntimeError(error_message) from exc

        if process is None:
            raise RuntimeError(error_message)
        if process.returncode != 0:
            stderr_text = (stderr or "").strip() if text else ""
            stdout_text = (stdout or "").strip() if text else ""
            detail = stderr_text or stdout_text or f"exit_code={process.returncode}"
            raise RuntimeError(f"{error_message} {detail}")
        return subprocess.CompletedProcess(command, process.returncode, stdout, stderr)

    try:
        result = subprocess.run(
            command,
            capture_output=True,
            check=False,
            text=text,
            timeout=timeout_seconds,
            encoding="utf-8" if text else None,
            errors="replace" if text else None,
        )
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError(
            f"{error_message} Превышено время ожидания {timeout_seconds} сек."
        ) from exc
    except OSError as exc:
        raise RuntimeError(error_message) from exc

    if result.returncode != 0:
        stderr = (result.stderr or "").strip()
        stdout = (result.stdout or "").strip() if text else ""
        detail = stderr or stdout or f"exit_code={result.returncode}"
        raise RuntimeError(f"{error_message} {detail}")
    return result


def _terminate_process_tree(process: subprocess.Popen) -> None:
    try:
        if os.name == "nt":
            process.kill()
            return
        os.killpg(process.pid, signal.SIGTERM)
    except ProcessLookupError:
        return
    except OSError:
        process.kill()
        return

    try:
        process.wait(timeout=5)
    except subprocess.TimeoutExpired:
        try:
            os.killpg(process.pid, signal.SIGKILL)
        except ProcessLookupError:
            return
        except OSError:
            process.kill()


def _convert_legacy_doc_with_soffice(*, soffice_path: str, filename: str, source_bytes: bytes) -> bytes:
    normalized_filename = _build_normalized_docx_filename(filename)
    with tempfile.TemporaryDirectory(prefix="docxaicorrector_doc_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        input_path = temp_dir / (Path(filename).name or "document.doc")
        output_path = temp_dir / Path(normalized_filename).name
        input_path.write_bytes(source_bytes)
        _run_completed_process(
            [
                soffice_path,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(temp_dir),
                str(input_path),
            ],
            error_message="Не удалось конвертировать legacy DOC через LibreOffice.",
            cleanup_process_group=True,
        )
        output_path = _resolve_converted_docx_output(
            temp_dir=temp_dir,
            expected_output_path=output_path,
            missing_output_message="Не удалось конвертировать legacy DOC через LibreOffice: выходной DOCX не создан.",
        )
        output_bytes = output_path.read_bytes()
        if not output_bytes:
            raise RuntimeError("Не удалось конвертировать legacy DOC через LibreOffice: выходной DOCX пуст.")
        return output_bytes


def _resolve_converted_docx_output(*, temp_dir: Path, expected_output_path: Path, missing_output_message: str) -> Path:
    if expected_output_path.exists():
        return expected_output_path
    fallback_matches = sorted(temp_dir.glob("*.docx"))
    if len(fallback_matches) == 1:
        return fallback_matches[0]
    raise RuntimeError(missing_output_message)


def _convert_pdf_to_docx(*, filename: str, source_bytes: bytes) -> tuple[bytes, str]:
    soffice_path = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice_path:
        raise RuntimeError(
            "Загружен PDF-файл, но автоконвертация недоступна. "
            "Установите LibreOffice (`soffice`) внутри WSL."
        )

    normalized_filename = _build_normalized_docx_filename(filename)
    with tempfile.TemporaryDirectory(prefix="docxaicorrector_pdf_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        input_path = temp_dir / (Path(filename).name or "document.pdf")
        output_path = temp_dir / Path(normalized_filename).name
        input_path.write_bytes(source_bytes)
        _run_completed_process(
            [
                soffice_path,
                "--headless",
                "--infilter=writer_pdf_import",
                "--convert-to",
                "docx",
                "--outdir",
                str(temp_dir),
                str(input_path),
            ],
            error_message="Не удалось конвертировать PDF через LibreOffice.",
            cleanup_process_group=True,
        )
        output_path = _resolve_converted_docx_output(
            temp_dir=temp_dir,
            expected_output_path=output_path,
            missing_output_message="Не удалось конвертировать PDF через LibreOffice: выходной DOCX не создан.",
        )
        output_bytes = output_path.read_bytes()
        if not output_bytes:
            raise RuntimeError("Не удалось конвертировать PDF через LibreOffice: выходной DOCX пуст.")
        return output_bytes, "libreoffice"


def _pdf_text_layer_import_enabled() -> bool:
    return True


def _pdf_ocr_import_enabled() -> bool:
    value = os.getenv("DOCXAI_PDF_OCR_IMPORT_ENABLED", "")
    return value.strip().lower() in {"1", "true", "yes", "on"}


def _pdf_ocr_languages() -> str:
    return (os.getenv("DOCXAI_PDF_OCR_LANGUAGES", "eng+rus") or "eng+rus").strip()


def _run_pdf_ocr_to_text_layer_pdf(*, input_path: Path, output_path: Path) -> None:
    ocrmypdf_path = shutil.which("ocrmypdf")
    if not ocrmypdf_path:
        raise RuntimeError("pdf_ocr_import_unavailable:ocrmypdf")
    if not shutil.which("tesseract"):
        raise RuntimeError("pdf_ocr_import_unavailable:tesseract")
    languages = _pdf_ocr_languages()
    command = [
        ocrmypdf_path,
        "--force-ocr",
        "-l",
        languages,
        str(input_path),
        str(output_path),
    ]
    _run_completed_process(
        command,
        error_message="Не удалось выполнить OCR PDF через OCRmyPDF/Tesseract.",
        text=True,
        timeout_seconds=_DOC_CONVERSION_TIMEOUT_SECONDS,
        cleanup_process_group=True,
    )
    if not output_path.exists() or output_path.stat().st_size <= 0:
        raise RuntimeError("pdf_ocr_import_empty_output")


def _count_pdf_pages_for_budget(pdf_path: Path, *, cap: int) -> int | None:
    """Best-effort page count that short-circuits once ``cap`` is exceeded.

    Returns ``None`` when the page tree cannot be read (e.g. a stub/broken PDF),
    in which case callers skip the page-count budget rather than reject a
    document we could not measure. The scan stops at ``cap + 1`` so a huge page
    tree is never fully walked.
    """

    try:
        from pdfminer.pdfpage import PDFPage
    except ImportError:  # pragma: no cover - depends on optional env
        return None
    try:
        with open(pdf_path, "rb") as handle:
            count = 0
            for _ in PDFPage.get_pages(handle):
                count += 1
                if count > cap:
                    break
            return count
    except Exception:
        return None


def _enforce_pdf_import_file_size_budget(*, filename: str, source_bytes: bytes) -> None:
    """Reject an oversize PDF before any parse work (F7).

    File size is an instant check on already-materialized bytes, so it runs in
    the parent (outside the wall-clock deadline) before we spawn the parse.
    """

    max_bytes = _MAX_PDF_IMPORT_FILE_BYTES
    file_size = len(source_bytes)
    if file_size > max_bytes:
        log_event(
            logging.ERROR,
            "pdf_import_over_budget",
            "PDF отклонён: превышен лимит размера файла для импорта.",
            filename=filename,
            limit="file_size",
            file_size_bytes=file_size,
            max_file_size_bytes=max_bytes,
        )
        raise RuntimeError(f"pdf_import_over_budget:file_size:{file_size}>{max_bytes}")


@dataclass
class _PdfParseStagesResult:
    """Serializable payload produced by the PDF parse worker (F7).

    Every field is picklable (dataclasses of primitives / bytes) so the result
    can be written to a file in the temp dir by a spawned child process and read
    back by the parent.
    """

    spans: list
    image_objects: list
    quality_decision: str
    quality_decision_reasons: tuple[str, ...]
    body_text_ratio: float
    ocr_used: bool


def _run_pdf_parse_stages(
    *,
    input_path: Path,
    ocr_output_path: Path,
    filename: str,
    ocr_import_enabled: bool,
    max_page_count: int,
    max_span_count: int,
    max_image_count: int,
) -> _PdfParseStagesResult:
    """Run ALL heavy pdfminer/OCR stages for a PDF import under one deadline (F7).

    Covers the page-count budget, span extraction, the optional OCR reparse, and
    image extraction. Object-count caps reject a pathological document (never a
    silent truncation). This is a module-level, picklable function so it can run
    either in-process (test seam) or in a spawned child process (production).
    """

    from docxaicorrector.pdf_import.images import extract_pdf_images_with_pdfminer
    from docxaicorrector.pdf_import.text_layer_quality import (
        build_text_layer_quality_report,
        extract_pdf_text_spans_with_pdfminer,
    )

    page_count = _count_pdf_pages_for_budget(input_path, cap=max_page_count)
    if page_count is not None and page_count > max_page_count:
        log_event(
            logging.ERROR,
            "pdf_import_over_budget",
            "PDF отклонён: превышен лимит числа страниц для импорта.",
            filename=filename,
            limit="page_count",
            page_count=page_count,
            max_page_count=max_page_count,
        )
        raise RuntimeError(f"pdf_import_over_budget:page_count:{page_count}>{max_page_count}")

    spans = list(extract_pdf_text_spans_with_pdfminer(input_path))
    _enforce_pdf_span_count_cap(spans, filename=filename, max_span_count=max_span_count)
    quality_report = build_text_layer_quality_report(spans)
    ocr_used = False

    if quality_report.decision != "promising":
        if not ocr_import_enabled:
            # Not promising and OCR disabled: return early so the parent raises
            # the typed not-promising error (no image extraction needed).
            return _PdfParseStagesResult(
                spans=spans,
                image_objects=[],
                quality_decision=str(quality_report.decision),
                quality_decision_reasons=tuple(str(r) for r in quality_report.decision_reasons),
                body_text_ratio=float(getattr(quality_report, "body_text_ratio", 0.0) or 0.0),
                ocr_used=False,
            )
        _run_pdf_ocr_to_text_layer_pdf(input_path=input_path, output_path=ocr_output_path)
        spans = list(extract_pdf_text_spans_with_pdfminer(ocr_output_path))
        _enforce_pdf_span_count_cap(spans, filename=filename, max_span_count=max_span_count)
        quality_report = build_text_layer_quality_report(spans)
        ocr_used = True
        if quality_report.decision != "promising":
            return _PdfParseStagesResult(
                spans=spans,
                image_objects=[],
                quality_decision=str(quality_report.decision),
                quality_decision_reasons=tuple(str(r) for r in quality_report.decision_reasons),
                body_text_ratio=float(getattr(quality_report, "body_text_ratio", 0.0) or 0.0),
                ocr_used=True,
            )

    try:
        image_objects = list(extract_pdf_images_with_pdfminer(input_path))
    except Exception as exc:
        log_event(
            logging.WARNING,
            "pdf_text_layer_image_extraction_failed",
            "Text-layer PDF import will continue without embedded images.",
            filename=filename,
            reason=str(exc),
        )
        image_objects = []
    else:
        _enforce_pdf_image_count_cap(image_objects, filename=filename, max_image_count=max_image_count)

    return _PdfParseStagesResult(
        spans=spans,
        image_objects=image_objects,
        quality_decision=str(quality_report.decision),
        quality_decision_reasons=tuple(str(r) for r in quality_report.decision_reasons),
        body_text_ratio=float(getattr(quality_report, "body_text_ratio", 0.0) or 0.0),
        ocr_used=ocr_used,
    )


def _enforce_pdf_span_count_cap(spans, *, filename: str, max_span_count: int) -> None:
    span_count = len(spans)
    if span_count > max_span_count:
        log_event(
            logging.ERROR,
            "pdf_import_over_budget",
            "PDF отклонён: превышен лимит числа text-span объектов при импорте.",
            filename=filename,
            limit="span_count",
            span_count=span_count,
            max_span_count=max_span_count,
        )
        raise RuntimeError(f"pdf_import_over_budget:span_count:{span_count}>{max_span_count}")


def _enforce_pdf_image_count_cap(image_objects, *, filename: str, max_image_count: int) -> None:
    image_count = len(image_objects)
    if image_count > max_image_count:
        log_event(
            logging.ERROR,
            "pdf_import_over_budget",
            "PDF отклонён: превышен лимит числа image объектов при импорте.",
            filename=filename,
            limit="image_count",
            image_count=image_count,
            max_image_count=max_image_count,
        )
        raise RuntimeError(f"pdf_import_over_budget:image_count:{image_count}>{max_image_count}")


def _pdf_parse_subprocess_entry(
    result_path: str,
    input_path: str,
    ocr_output_path: str,
    filename: str,
    ocr_import_enabled: bool,
    max_page_count: int,
    max_span_count: int,
    max_image_count: int,
) -> None:
    """Child-process entrypoint (F7): run the parse stages and pickle the result.

    Errors are serialized as a plain message string (all parse-stage errors are
    ``RuntimeError`` with a typed message) so the parent can re-raise them even
    if the original exception type is not picklable.
    """

    try:
        result = _run_pdf_parse_stages(
            input_path=Path(input_path),
            ocr_output_path=Path(ocr_output_path),
            filename=filename,
            ocr_import_enabled=ocr_import_enabled,
            max_page_count=max_page_count,
            max_span_count=max_span_count,
            max_image_count=max_image_count,
        )
        payload: dict[str, object] = {"ok": True, "result": result}
    except BaseException as exc:  # noqa: BLE001 - marshalled to the parent
        payload = {"ok": False, "error_message": str(exc) or exc.__class__.__name__}
    try:
        with open(result_path, "wb") as handle:
            pickle.dump(payload, handle, protocol=pickle.HIGHEST_PROTOCOL)
    except Exception:
        with open(result_path, "wb") as handle:
            pickle.dump({"ok": False, "error_message": "pdf_parse_result_unserializable"}, handle)


# Test seam: the parent uses this callable as the spawned-process target. Tests
# monkeypatch it to a module-level worker that overruns the deadline so the
# real terminate()/join() kill path can be exercised.
_PDF_PARSE_SUBPROCESS_ENTRY: Callable[..., None] = _pdf_parse_subprocess_entry

# Optional observer invoked with the spawned process right after start; tests set
# it to capture the child and assert it was terminated. Kept ``None`` in prod.
_pdf_parse_process_observer: Callable[[object], None] | None = None


def _pdf_parse_sleep_forever_entry(result_path: str, *args: object, **kwargs: object) -> None:
    """Test-only subprocess target that never finishes, to exercise termination."""

    while True:  # pragma: no cover - runs in a child process that is killed
        time.sleep(0.05)


def _pdf_parse_canned_result_entry(result_path: str, *args: object, **kwargs: object) -> None:
    """Test-only subprocess target: write a valid result to prove the parent
    deserializes a spawned child's payload across the process boundary."""

    result = _PdfParseStagesResult(
        spans=[],
        image_objects=[],
        quality_decision="promising",
        quality_decision_reasons=(),
        body_text_ratio=1.0,
        ocr_used=False,
    )
    with open(result_path, "wb") as handle:
        pickle.dump({"ok": True, "result": result}, handle, protocol=pickle.HIGHEST_PROTOCOL)


def _run_pdf_parse_stages_with_deadline(
    *,
    temp_dir: Path,
    input_path: Path,
    ocr_output_path: Path,
    filename: str,
) -> _PdfParseStagesResult:
    """Run the parse stages under the single unified wall-clock deadline (F7).

    Production spawns a killable child process (``spawn`` context) that writes a
    serialized result into the temp dir; on overrun the parent ``terminate()``s
    and ``join()``s it (Python cannot kill a thread, only a process), then raises
    ``pdf_import_over_budget:deadline``. The in-process fallback (test seam) runs
    the same worker in a daemon thread and cannot force-kill it — test-only.
    """

    budget = _PDF_PARSE_WALLCLOCK_BUDGET_SECONDS
    ocr_import_enabled = _pdf_ocr_import_enabled()
    max_page_count = _MAX_PDF_IMPORT_PAGE_COUNT
    max_span_count = _MAX_PDF_IMPORT_SPAN_COUNT
    max_image_count = _MAX_PDF_IMPORT_IMAGE_COUNT

    if _pdf_parse_in_process_enabled():
        return _run_pdf_parse_stages_in_process_guarded(
            budget=budget,
            input_path=input_path,
            ocr_output_path=ocr_output_path,
            filename=filename,
            ocr_import_enabled=ocr_import_enabled,
            max_page_count=max_page_count,
            max_span_count=max_span_count,
            max_image_count=max_image_count,
        )

    result_path = temp_dir / "pdf_parse_result.pickle"
    context = multiprocessing.get_context("spawn")
    process = context.Process(
        target=_PDF_PARSE_SUBPROCESS_ENTRY,
        args=(
            str(result_path),
            str(input_path),
            str(ocr_output_path),
            filename,
            ocr_import_enabled,
            max_page_count,
            max_span_count,
            max_image_count,
        ),
        daemon=True,
        name="pdf-parse-worker",
    )
    started_at = time.monotonic()
    process.start()
    if _pdf_parse_process_observer is not None:
        _pdf_parse_process_observer(process)
    process.join(timeout=budget)
    if process.is_alive():
        # Python cannot kill a thread, but a child process CAN be terminated.
        process.terminate()
        process.join(timeout=10)
        if process.is_alive():  # pragma: no cover - SIGKILL escalation
            process.kill()
            process.join(timeout=10)
        elapsed = time.monotonic() - started_at
        log_event(
            logging.ERROR,
            "pdf_import_over_budget",
            "PDF отклонён: разбор PDF превысил единый лимит времени; дочерний процесс завершён.",
            filename=filename,
            limit="deadline",
            elapsed_seconds=round(elapsed, 3),
            max_parse_seconds=budget,
        )
        raise RuntimeError(f"pdf_import_over_budget:deadline:{budget}")

    try:
        with open(result_path, "rb") as handle:
            payload = pickle.load(handle)
    except (OSError, EOFError, pickle.UnpicklingError) as exc:
        raise RuntimeError(f"pdf_parse_subprocess_no_result:exitcode={process.exitcode}") from exc

    if not payload.get("ok"):
        raise RuntimeError(str(payload.get("error_message") or "pdf_parse_subprocess_failed"))
    result = payload.get("result")
    if not isinstance(result, _PdfParseStagesResult):
        raise RuntimeError("pdf_parse_subprocess_malformed_result")
    return result


def _run_pdf_parse_stages_in_process_guarded(
    *,
    budget: float,
    input_path: Path,
    ocr_output_path: Path,
    filename: str,
    ocr_import_enabled: bool,
    max_page_count: int,
    max_span_count: int,
    max_image_count: int,
) -> _PdfParseStagesResult:
    """In-process fallback (TEST-ONLY) under the unified deadline.

    Runs the worker in a daemon thread and fails fast on overrun. NOTE: unlike
    the production child process, an in-process daemon thread CANNOT be
    force-killed — an overrunning parse keeps running until the interpreter
    exits. This path exists only so tests that monkeypatch in-process pdfminer
    functions keep working; production uses the killable subprocess path.
    """

    result_box: dict[str, _PdfParseStagesResult] = {}
    error_box: dict[str, BaseException] = {}

    def _runner() -> None:
        try:
            result_box["value"] = _run_pdf_parse_stages(
                input_path=input_path,
                ocr_output_path=ocr_output_path,
                filename=filename,
                ocr_import_enabled=ocr_import_enabled,
                max_page_count=max_page_count,
                max_span_count=max_span_count,
                max_image_count=max_image_count,
            )
        except BaseException as exc:  # noqa: BLE001 - re-raised in the caller thread
            error_box["error"] = exc

    parse_thread = threading.Thread(target=_runner, daemon=True, name="pdf-parse-inproc")
    started_at = time.monotonic()
    parse_thread.start()
    parse_thread.join(timeout=budget)
    if parse_thread.is_alive():
        elapsed = time.monotonic() - started_at
        log_event(
            logging.ERROR,
            "pdf_import_over_budget",
            "PDF отклонён: разбор PDF превысил лимит времени (in-process fallback).",
            filename=filename,
            limit="parse_wallclock",
            elapsed_seconds=round(elapsed, 3),
            max_parse_seconds=budget,
        )
        raise RuntimeError(f"pdf_import_over_budget:parse_wallclock:{budget}")
    if "error" in error_box:
        raise error_box["error"]
    return result_box["value"]


def _convert_pdf_text_layer_to_docx(*, filename: str, source_bytes: bytes) -> tuple[bytes, str]:
    from docx import Document

    from docxaicorrector.pdf_import.logical_import import build_paragraph_units_from_text_spans

    _enforce_pdf_import_file_size_budget(filename=filename, source_bytes=source_bytes)

    with tempfile.TemporaryDirectory(prefix="docxaicorrector_pdf_text_layer_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        input_path = temp_dir / (Path(filename).name or "document.pdf")
        output_path = temp_dir / Path(_build_normalized_docx_filename(filename)).name
        ocr_output_path = temp_dir / f"{input_path.stem}.ocr.pdf"
        input_path.write_bytes(source_bytes)

        # All heavy pdfminer/OCR stages (page count + span extract + optional OCR
        # reparse + image extract) run under a SINGLE unified deadline (F7). In
        # production this is a genuinely killable child process; the in-process
        # test seam runs the same worker under a thread guard.
        parse_result = _run_pdf_parse_stages_with_deadline(
            temp_dir=temp_dir,
            input_path=input_path,
            ocr_output_path=ocr_output_path,
            filename=filename,
        )
        spans = parse_result.spans
        image_objects = parse_result.image_objects
        if parse_result.quality_decision != "promising":
            prefix = (
                "pdf_ocr_text_layer_import_not_promising:"
                if parse_result.ocr_used
                else "pdf_text_layer_import_not_promising:"
            )
            raise RuntimeError(prefix + ",".join(parse_result.quality_decision_reasons))

        import_result = build_paragraph_units_from_text_spans(spans)
        if not import_result.paragraphs:
            raise RuntimeError("pdf_text_layer_import_empty_document")

        document = Document()
        content_items = [
            ("paragraph", int(getattr(paragraph, "source_index", 0)), index, paragraph)
            for index, paragraph in enumerate(import_result.paragraphs)
        ]
        content_items.extend(
            ("image", int(getattr(image, "source_index", 0)), index, image)
            for index, image in enumerate(image_objects)
        )
        images_emitted = 0
        images_dropped = 0
        for item_type, _source_index, _index, payload in sorted(content_items, key=lambda item: (item[1], item[0] == "paragraph", item[2])):
            if item_type == "image":
                if _append_pdf_image_to_docx(document, payload):
                    images_emitted += 1
                else:
                    images_dropped += 1
                continue
            paragraph = payload
            _append_pdf_text_paragraph_to_docx(document, paragraph)
        document.save(output_path)
        output_bytes = output_path.read_bytes() if output_path.exists() else b""
        if not output_bytes:
            raise RuntimeError("pdf_text_layer_import_empty_docx")
        log_event(
            logging.INFO,
            "pdf_text_layer_import_succeeded",
            "PDF импортирован через text-layer importer в generated DOCX.",
            filename=filename,
            span_count=len(spans),
            paragraph_count=len(import_result.paragraphs),
            image_count=len(image_objects),
            images_emitted=images_emitted,
            images_dropped=images_dropped,
            skipped_page_number_count=import_result.report.skipped_page_number_count,
            skipped_repeated_page_furniture_count=import_result.report.skipped_repeated_page_furniture_count,
            skipped_blank_page_notice_count=import_result.report.skipped_blank_page_notice_count,
            body_text_ratio=parse_result.body_text_ratio,
        )
        return output_bytes, "pdf-text-layer"


def _append_pdf_text_paragraph_to_docx(document, paragraph) -> None:
    style = _pdf_text_layer_docx_style(paragraph.role, paragraph.heading_level)
    docx_paragraph = document.add_paragraph(style=style)

    # ``pdf_emphasis_runs`` carries the fully-built paragraph text (with de-hyphenation
    # and inline footnote markers already applied) split into character-level emphasis
    # runs; its concatenation equals ``paragraph.text``. Emit each run with its own
    # bold/italic so sub-line emphasis survives into the DOCX.
    emphasis_runs = [
        (str(run_text), bool(is_bold), bool(is_italic))
        for run_text, is_bold, is_italic in (getattr(paragraph, "pdf_emphasis_runs", None) or [])
        if run_text
    ]

    if paragraph.role == "heading":
        # A heading already conveys weight via its style, so a *uniformly*
        # bold/italic heading must not encode false character emphasis. Only
        # genuinely mixed intra-heading emphasis (F24) is preserved run-by-run;
        # a uniform (or empty) run set collapses to a single plain run.
        distinct_states = {(is_bold, is_italic) for _text, is_bold, is_italic in emphasis_runs}
        if len(distinct_states) > 1:
            for run_text, is_bold, is_italic in emphasis_runs:
                run = docx_paragraph.add_run(run_text)
                if is_bold:
                    run.bold = True
                if is_italic:
                    run.italic = True
            return
        docx_paragraph.add_run(paragraph.text)
        return

    if emphasis_runs:
        for run_text, is_bold, is_italic in emphasis_runs:
            run = docx_paragraph.add_run(run_text)
            if is_bold:
                run.bold = True
            if is_italic:
                run.italic = True
        return

    run = docx_paragraph.add_run(paragraph.text)
    if bool(getattr(paragraph, "is_bold", False)):
        run.bold = True
    if bool(getattr(paragraph, "is_italic", False)):
        run.italic = True


def _append_pdf_image_to_docx(document, image_object) -> bool:
    image_bytes = getattr(image_object, "image_bytes", None)
    if not isinstance(image_bytes, bytes) or not image_bytes:
        log_event(
            logging.WARNING,
            "pdf_text_layer_image_render_dropped",
            "PDF image could not be emitted into DOCX.",
            source_index=getattr(image_object, "source_index", None),
            mime_type=getattr(image_object, "mime_type", None),
            reason="empty_image_bytes",
        )
        return False
    try:
        _append_image_bytes_to_docx(document, image_bytes)
        return True
    except Exception as primary_exc:
        coerced_image_bytes = _coerce_image_bytes_for_docx(image_bytes)
        if coerced_image_bytes is None:
            log_event(
                logging.WARNING,
                "pdf_text_layer_image_render_dropped",
                "PDF image could not be emitted into DOCX.",
                source_index=getattr(image_object, "source_index", None),
                mime_type=getattr(image_object, "mime_type", None),
                reason="docx_render_failed",
                error_message=str(primary_exc),
            )
            return False
    try:
        _append_image_bytes_to_docx(document, coerced_image_bytes)
        return True
    except Exception as coerced_exc:
        log_event(
            logging.WARNING,
            "pdf_text_layer_image_render_dropped",
            "PDF image could not be emitted into DOCX.",
            source_index=getattr(image_object, "source_index", None),
            mime_type=getattr(image_object, "mime_type", None),
            reason="coerced_docx_render_failed",
            error_message=str(coerced_exc),
        )
        return False


def _append_image_bytes_to_docx(document, image_bytes: bytes) -> None:
    paragraph = document.add_paragraph()
    try:
        paragraph.add_run().add_picture(BytesIO(image_bytes))
    except Exception:
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)
        raise


def _coerce_image_bytes_for_docx(image_bytes: bytes) -> bytes | None:
    try:
        from PIL import Image
    except ImportError:  # pragma: no cover - optional transcode path
        return None
    try:
        with Image.open(BytesIO(image_bytes)) as image:
            converted = image.convert("RGBA") if image.mode not in {"RGB", "RGBA"} else image
            output = BytesIO()
            converted.save(output, format="PNG")
            return output.getvalue()
    except Exception:
        return None


def _pdf_text_layer_docx_style(role: str, heading_level: int | None) -> str | None:
    if role == "heading":
        level = min(max(int(heading_level or 2), 1), 6)
        return f"Heading {level}"
    if role == "list":
        return "List Bullet"
    return None


def _convert_pdf_to_docx_with_optional_text_layer(*, filename: str, source_bytes: bytes) -> tuple[bytes, str]:
    return _convert_pdf_text_layer_to_docx(filename=filename, source_bytes=source_bytes)


def _convert_legacy_doc_with_antiword(*, antiword_path: str, filename: str, source_bytes: bytes) -> bytes:
    from docxaicorrector.generation._generation import ensure_pandoc_available
    import pypandoc

    ensure_pandoc_available()
    normalized_filename = _build_normalized_docx_filename(filename)
    with tempfile.TemporaryDirectory(prefix="docxaicorrector_doc_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        input_path = temp_dir / (Path(filename).name or "document.doc")
        output_path = temp_dir / Path(normalized_filename).name
        input_path.write_bytes(source_bytes)
        antiword_result = _run_completed_process(
            [antiword_path, "-x", "db", str(input_path)],
            error_message="Не удалось извлечь legacy DOC через antiword.",
        )
        docbook_xml = (antiword_result.stdout or "").strip()
        if not docbook_xml:
            raise RuntimeError("Не удалось извлечь legacy DOC через antiword: получен пустой DocBook XML.")
        try:
            pypandoc.convert_text(docbook_xml, to="docx", format="docbook", outputfile=str(output_path))
        except (OSError, RuntimeError) as exc:
            raise RuntimeError("Не удалось собрать DOCX из legacy DOC через Pandoc.") from exc
        output_bytes = output_path.read_bytes() if output_path.exists() else b""
        if not output_bytes:
            raise RuntimeError("Не удалось собрать DOCX из legacy DOC через Pandoc: выходной DOCX пуст.")
        return output_bytes


def _convert_legacy_doc_to_docx(*, filename: str, source_bytes: bytes) -> tuple[bytes, str]:
    soffice_path = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice_path:
        try:
            return _convert_legacy_doc_with_soffice(
                soffice_path=soffice_path,
                filename=filename,
                source_bytes=source_bytes,
            ), "libreoffice"
        except RuntimeError as soffice_exc:
            antiword_path = shutil.which("antiword")
            if antiword_path:
                try:
                    return _convert_legacy_doc_with_antiword(
                        antiword_path=antiword_path,
                        filename=filename,
                        source_bytes=source_bytes,
                    ), "antiword+pandoc"
                except RuntimeError as antiword_exc:
                    raise RuntimeError(
                        "Не удалось конвертировать legacy DOC ни через LibreOffice, ни через antiword+pandoc. "
                        f"LibreOffice: {soffice_exc} antiword+pandoc: {antiword_exc}"
                    ) from antiword_exc
            raise soffice_exc

    antiword_path = shutil.which("antiword")
    if antiword_path:
        return _convert_legacy_doc_with_antiword(
            antiword_path=antiword_path,
            filename=filename,
            source_bytes=source_bytes,
        ), "antiword+pandoc"

    raise RuntimeError(
        "Загружен legacy DOC-файл, но автоконвертация недоступна. "
        "Установите в WSL LibreOffice (`soffice`) или связку `antiword` + `pandoc`."
    )


def legacy_doc_conversion_available() -> bool:
    if shutil.which("soffice") or shutil.which("libreoffice"):
        return True

    if not shutil.which("antiword"):
        return False

    try:
        import pypandoc

        pypandoc.get_pandoc_version()
    except OSError:
        return False

    return True


def normalize_uploaded_document(*, filename: str, source_bytes: bytes) -> NormalizedUploadedDocument:
    source_format = _detect_uploaded_document_format(filename=filename, source_bytes=source_bytes)
    normalized_filename = _build_normalized_docx_filename(filename) if source_format in {"doc", "docx", "pdf"} else filename

    if source_format == "doc":
        converted_bytes, conversion_backend = _convert_legacy_doc_to_docx(
            filename=filename,
            source_bytes=source_bytes,
        )
        return NormalizedUploadedDocument(
            original_filename=filename,
            filename=normalized_filename,
            content_bytes=converted_bytes,
            source_format=source_format,
            conversion_backend=conversion_backend,
        )

    if source_format == "pdf":
        converted_bytes, conversion_backend = _convert_pdf_to_docx_with_optional_text_layer(
            filename=filename,
            source_bytes=source_bytes,
        )
        return NormalizedUploadedDocument(
            original_filename=filename,
            filename=normalized_filename,
            content_bytes=converted_bytes,
            source_format=source_format,
            conversion_backend=conversion_backend,
        )

    return NormalizedUploadedDocument(
        original_filename=filename,
        filename=normalized_filename if source_format == "docx" else filename,
        content_bytes=bytes(source_bytes),
        source_format=source_format,
        conversion_backend=None,
    )


def _build_uploaded_file_token_components(*, normalized_document: NormalizedUploadedDocument, source_bytes: bytes) -> tuple[int, str]:
    identity_bytes = source_bytes if normalized_document.source_format in {"doc", "pdf"} else normalized_document.content_bytes
    identity_hash = hashlib.sha256(identity_bytes).hexdigest()[:16]
    return len(identity_bytes), identity_hash


def resolve_upload_contract(*, filename: str, source_bytes: bytes) -> ResolvedUploadContract:
    normalized_document = normalize_uploaded_document(filename=filename, source_bytes=source_bytes)
    token_size, token_hash = _build_uploaded_file_token_components(
        normalized_document=normalized_document,
        source_bytes=source_bytes,
    )
    return ResolvedUploadContract(
        source_identity=UploadSourceIdentity(
            original_filename=filename,
            source_bytes=bytes(source_bytes),
            token_size=token_size,
            token_hash=token_hash,
        ),
        normalized_document=normalized_document,
    )


def freeze_resolved_upload(contract: ResolvedUploadContract) -> FrozenUploadPayload:
    return FrozenUploadPayload(
        filename=contract.filename,
        content_bytes=contract.content_bytes,
        file_size=contract.file_size,
        content_hash=contract.content_hash,
        file_token=contract.file_token,
        source_format=contract.normalized_document.source_format,
        conversion_backend=contract.normalized_document.conversion_backend,
    )


def freeze_uploaded_file(uploaded_file: UploadedFileLike | BytesIO) -> FrozenUploadPayload:
    source_bytes = read_uploaded_file_bytes(uploaded_file)
    filename = getattr(uploaded_file, "name", "") or _DEFAULT_UPLOADED_FILENAME
    return freeze_resolved_upload(resolve_upload_contract(filename=filename, source_bytes=source_bytes))


def freeze_uploaded_file_lightweight(uploaded_file: UploadedFileLike | BytesIO) -> FrozenUploadPayload:
    """Cheap, main-thread-safe payload freeze that defers heavy format conversions.

    For PDF/DOC sources this returns a payload whose ``content_bytes`` are still
    raw input bytes and ``conversion_backend`` is ``None``. The actual conversion
    must run inside the preparation worker via :func:`materialize_uploaded_payload`.
    For DOCX sources behavior is identical to :func:`freeze_uploaded_file` — no
    conversion is needed and ``content_bytes`` already point at the DOCX bytes.

    The ``file_token`` is computed from raw input bytes for PDF/DOC, matching
    :func:`_build_uploaded_file_token_components`, so a lightweight payload and
    its later materialized counterpart share the same token.
    """

    source_bytes = read_uploaded_file_bytes(uploaded_file)
    filename = getattr(uploaded_file, "name", "") or _DEFAULT_UPLOADED_FILENAME
    source_format = _detect_uploaded_document_format(filename=filename, source_bytes=source_bytes)
    raw = bytes(source_bytes)
    identity_hash = hashlib.sha256(raw).hexdigest()[:16]
    if source_format == "docx":
        normalized_filename = _build_normalized_docx_filename(filename)
        token = f"{normalized_filename}:{len(raw)}:{identity_hash}"
        return FrozenUploadPayload(
            filename=normalized_filename,
            content_bytes=raw,
            file_size=len(raw),
            content_hash=identity_hash,
            file_token=token,
            source_format="docx",
            conversion_backend=None,
        )
    if source_format in {"pdf", "doc"}:
        normalized_filename = _build_normalized_docx_filename(filename)
        token = f"{normalized_filename}:{len(raw)}:{identity_hash}"
        return FrozenUploadPayload(
            filename=filename,  # keep original filename until materialization
            content_bytes=raw,
            file_size=len(raw),
            content_hash=identity_hash,
            file_token=token,
            source_format=source_format,
            conversion_backend=None,
        )
    # Unknown formats fall back to the eager path so we still get a consistent error.
    return freeze_resolved_upload(resolve_upload_contract(filename=filename, source_bytes=raw))


def _touch_materialized_upload_cache_entry(
    cache: OrderedDict[str, FrozenUploadPayload],
    cache_key: str,
    payload: FrozenUploadPayload,
) -> None:
    cache.pop(cache_key, None)
    cache[cache_key] = payload


def _trim_materialized_upload_cache(cache: OrderedDict[str, FrozenUploadPayload]) -> None:
    while len(cache) > _MATERIALIZED_UPLOAD_CACHE_LIMIT:
        cache.popitem(last=False)


def _read_or_reserve_materialized_upload(cache_key: str) -> tuple[FrozenUploadPayload | None, threading.Event | None]:
    while True:
        with _shared_materialized_upload_cache_lock:
            cached = _shared_materialized_upload_cache.get(cache_key)
            if cached is not None:
                _touch_materialized_upload_cache_entry(_shared_materialized_upload_cache, cache_key, cached)
                return cached, None

            in_flight = _shared_materialized_upload_inflight.get(cache_key)
            if in_flight is None:
                in_flight = threading.Event()
                _shared_materialized_upload_inflight[cache_key] = in_flight
                return None, in_flight

        in_flight.wait()


def _store_materialized_upload(payload: FrozenUploadPayload, *, cache_key: str | None = None) -> None:
    with _shared_materialized_upload_cache_lock:
        _touch_materialized_upload_cache_entry(
            _shared_materialized_upload_cache,
            cache_key or payload.file_token,
            payload,
        )
        _trim_materialized_upload_cache(_shared_materialized_upload_cache)


def _release_materialized_upload_reservation(cache_key: str) -> None:
    with _shared_materialized_upload_cache_lock:
        in_flight = _shared_materialized_upload_inflight.pop(cache_key, None)
    if in_flight is not None:
        in_flight.set()


def _materialized_upload_cache_key(payload: FrozenUploadPayload) -> str:
    fmt = (payload.source_format or "").lower()
    if fmt == "pdf" and _pdf_text_layer_import_enabled():
        return f"{payload.file_token}:pdf-text-layer"
    return payload.file_token


def materialize_uploaded_payload(
    payload: FrozenUploadPayload,
    *,
    progress_callback=None,
) -> FrozenUploadPayload:
    """Run any deferred format conversion (PDF/DOC -> DOCX) and return a payload
    whose ``content_bytes`` are DOCX bytes ready for downstream extraction.

    Already-materialized payloads (``conversion_backend`` set, or DOCX source)
    are returned unchanged. Long-running conversions emit periodic heartbeat
    progress events through ``progress_callback`` so the UI never stalls.
    The returned payload preserves ``file_token`` from the input so cache
    keys stay stable.
    """

    if payload.conversion_backend or (payload.source_format or "").lower() == "docx":
        return payload

    fmt = (payload.source_format or "").lower()
    if fmt in {"pdf", "doc"}:
        cache_key = _materialized_upload_cache_key(payload)
        cached_payload, reservation = _read_or_reserve_materialized_upload(cache_key)
        if cached_payload is not None:
            if progress_callback is not None:
                progress_callback(
                    stage="DOCX готов",
                    detail="Использую уже сконвертированную копию DOCX. Повторная конвертация не нужна.",
                    progress=0.18,
                    metrics={
                        "source_format": fmt,
                        "conversion_backend": cached_payload.conversion_backend,
                        "file_size_bytes": cached_payload.file_size,
                        "conversion_reused": True,
                    },
                )
            log_event(
                logging.INFO,
                "materialized_upload_cache_hit",
                "Использована уже сконвертированная DOCX-копия исходного файла.",
                source_format=fmt,
                file_token=payload.file_token,
                conversion_backend=cached_payload.conversion_backend,
            )
            return cached_payload
    else:
        cache_key = payload.file_token
        reservation = None

    if fmt == "pdf":
        try:
            if progress_callback is not None:
                pdf_import_detail = (
                    "Запускаю text-layer импорт PDF в DOCX…"
                    if _pdf_text_layer_import_enabled()
                    else "Запускаю конвертацию PDF в DOCX через LibreOffice…"
                )
                progress_callback(
                    stage="Импорт PDF",
                    detail=pdf_import_detail,
                    progress=0.05,
                    metrics={
                        "source_format": "pdf",
                        "file_size_bytes": payload.file_size,
                        "conversion_reused": False,
                    },
                )
            with HeartbeatBeacon(
                progress_callback,
                stage="Импорт PDF",
                detail_template=(
                    "Text-layer importer собирает PDF в DOCX… ({elapsed} сек)."
                    if _pdf_text_layer_import_enabled()
                    else "LibreOffice конвертирует PDF в DOCX… ({elapsed} сек). Для крупных книг это может занять 30–120 сек."
                ),
                progress=0.10,
                metrics={"source_format": "pdf", "file_size_bytes": payload.file_size, "conversion_reused": False},
                interval_seconds=2.0,
            ):
                converted_bytes, conversion_backend = _convert_pdf_to_docx_with_optional_text_layer(
                    filename=payload.filename,
                    source_bytes=payload.content_bytes,
                )
            if progress_callback is not None:
                progress_callback(
                    stage="DOCX готов",
                    detail="PDF сконвертирован, начинаю разбор содержимого.",
                    progress=0.18,
                    metrics={
                        "source_format": "pdf",
                        "conversion_backend": conversion_backend,
                        "file_size_bytes": len(converted_bytes),
                        "conversion_reused": False,
                    },
                )
            normalized_filename = _build_normalized_docx_filename(payload.filename)
            materialized_payload = FrozenUploadPayload(
                filename=normalized_filename,
                content_bytes=converted_bytes,
                file_size=len(converted_bytes),
                content_hash=hashlib.sha256(converted_bytes).hexdigest()[:16],
                file_token=payload.file_token,
                source_format="pdf",
                conversion_backend=conversion_backend,
            )
            _store_materialized_upload(materialized_payload, cache_key=cache_key)
            return materialized_payload
        finally:
            if reservation is not None:
                _release_materialized_upload_reservation(cache_key)

    if fmt == "doc":
        try:
            if progress_callback is not None:
                progress_callback(
                    stage="Импорт DOC",
                    detail="Запускаю конвертацию legacy DOC в DOCX…",
                    progress=0.05,
                    metrics={
                        "source_format": "doc",
                        "file_size_bytes": payload.file_size,
                        "conversion_reused": False,
                    },
                )
            with HeartbeatBeacon(
                progress_callback,
                stage="Импорт DOC",
                detail_template="Конвертирую legacy DOC в DOCX… ({elapsed} сек).",
                progress=0.10,
                metrics={"source_format": "doc", "file_size_bytes": payload.file_size, "conversion_reused": False},
                interval_seconds=2.0,
            ):
                converted_bytes, conversion_backend = _convert_legacy_doc_to_docx(
                    filename=payload.filename,
                    source_bytes=payload.content_bytes,
                )
            if progress_callback is not None:
                progress_callback(
                    stage="DOCX готов",
                    detail="DOC сконвертирован, начинаю разбор содержимого.",
                    progress=0.18,
                    metrics={
                        "source_format": "doc",
                        "conversion_backend": conversion_backend,
                        "file_size_bytes": len(converted_bytes),
                        "conversion_reused": False,
                    },
                )
            normalized_filename = _build_normalized_docx_filename(payload.filename)
            materialized_payload = FrozenUploadPayload(
                filename=normalized_filename,
                content_bytes=converted_bytes,
                file_size=len(converted_bytes),
                content_hash=hashlib.sha256(converted_bytes).hexdigest()[:16],
                file_token=payload.file_token,
                source_format="doc",
                conversion_backend=conversion_backend,
            )
            _store_materialized_upload(materialized_payload, cache_key=cache_key)
            return materialized_payload
        finally:
            if reservation is not None:
                _release_materialized_upload_reservation(cache_key)

    return payload


def build_uploaded_file_token(uploaded_file: UploadedFileLike | BytesIO | None = None, *, source_name: str | None = None, source_bytes: bytes | None = None) -> str:
    if isinstance(uploaded_file, FrozenUploadPayload):
        return uploaded_file.file_token
    if source_bytes is None:
        if uploaded_file is None:
            raise ValueError("Для построения токена нужен uploaded_file или source_bytes.")
        source_bytes = read_uploaded_file_bytes(uploaded_file)
    file_name = source_name if source_name is not None else (getattr(uploaded_file, "name", "") or _DEFAULT_UPLOADED_FILENAME)
    return resolve_upload_contract(filename=file_name, source_bytes=bytes(source_bytes)).file_token


def _coerce_metric_to_int(metrics: dict[str, object], key: str) -> int:
    value = metrics.get(key, 0)
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return int(value)
    if isinstance(value, str):
        try:
            return int(value)
        except ValueError:
            return 0
    return 0


def build_uploaded_file_selection_marker(uploaded_file) -> str:
    return build_uploaded_file_token(uploaded_file)


def build_preparation_request_marker(uploaded_file, *, chunk_size: int, processing_operation: str = "edit") -> str:
    resolved_operation = str(processing_operation or "edit").strip().lower() or "edit"
    operation_suffix = "" if resolved_operation == "edit" else f":op={resolved_operation}"
    return f"{build_uploaded_file_selection_marker(uploaded_file)}:{chunk_size}{operation_suffix}"


def build_result_bundle(
    *,
    source_name: str,
    source_token: str,
    docx_bytes: bytes | None,
    markdown_text: str,
    narration_text: str | None = None,
    processing_operation: str = "edit",
    audiobook_postprocess_enabled: bool = False,
    quality_warning: Mapping[str, object] | None = None,
) -> dict[str, object]:
    return {
        "source_name": source_name,
        "source_token": source_token,
        "docx_bytes": docx_bytes,
        "markdown_text": markdown_text,
        "narration_text": narration_text,
        "processing_operation": processing_operation,
        "audiobook_postprocess_enabled": audiobook_postprocess_enabled,
        "quality_warning": quality_warning,
    }


def should_cache_completed_source(*, source_bytes: bytes) -> bool:
    return len(source_bytes) <= MAX_COMPLETED_SOURCE_BYTES


def get_current_result_bundle() -> dict[str, object] | None:
    latest_docx_bytes = st.session_state.get("latest_docx_bytes")
    latest_narration_text = get_latest_narration_text()
    processing_operation = get_latest_processing_operation()
    if latest_docx_bytes is None:
        if latest_narration_text is None:
            return None
        if processing_operation == "audiobook":
            return None
    if not latest_docx_bytes and latest_narration_text is None:
        return None
    return build_result_bundle(
        source_name=get_latest_source_name(),
        source_token=get_latest_source_token(),
        docx_bytes=latest_docx_bytes,
        markdown_text=st.session_state.get("latest_markdown", ""),
        narration_text=latest_narration_text,
        processing_operation=processing_operation,
        audiobook_postprocess_enabled=get_latest_audiobook_postprocess_enabled(),
        quality_warning=cast(Mapping[str, object] | None, st.session_state.get("latest_quality_warning")),
    )


def set_session_values(**values) -> None:
    for key, value in values.items():
        st.session_state[key] = value


def _filter_allowed_set_state_values(values: dict[str, object]) -> dict[str, object]:
    unknown_keys = sorted(key for key in values if key not in _ALLOWED_SET_STATE_EVENT_KEYS)
    if unknown_keys:
        log_event(
            logging.WARNING,
            "state_event_unknown_keys",
            "SetStateEvent попытался записать ключи вне разрешённого allowlist.",
            unknown_keys=unknown_keys,
        )
    return {key: value for key, value in values.items() if key in _ALLOWED_SET_STATE_EVENT_KEYS}


def emit_or_apply_state(runtime: BackgroundRuntime | None, **values) -> None:
    if runtime is None:
        set_session_values(**values)
        return
    runtime.emit(SetStateEvent(values=values))


def emit_or_apply_image_reset(runtime: BackgroundRuntime | None) -> None:
    if runtime is None:
        reset_image_state()
        return
    runtime.emit(ResetImageStateEvent())


def emit_or_apply_status(runtime: BackgroundRuntime | None, *, set_processing_status, **payload) -> None:
    if runtime is None:
        set_processing_status(**payload)
        return
    runtime.emit(SetProcessingStatusEvent(payload=payload))


def emit_or_apply_finalize(runtime: BackgroundRuntime | None, *, finalize_processing_status, stage: str, detail: str, progress: float, terminal_kind: str | None = None) -> None:
    if runtime is None:
        finalize_processing_status(stage, detail, progress, terminal_kind)
        return
    runtime.emit(FinalizeProcessingStatusEvent(stage=stage, detail=detail, progress=progress, terminal_kind=terminal_kind))


def emit_or_apply_activity(runtime: BackgroundRuntime | None, *, push_activity, message: str) -> None:
    if runtime is None:
        push_activity(message)
        return
    runtime.emit(PushActivityEvent(message=message))


def emit_or_apply_log(runtime: BackgroundRuntime | None, *, append_log, **payload) -> None:
    if runtime is None:
        append_log(**payload)
        return
    runtime.emit(AppendLogEvent(payload=payload))


def emit_or_apply_image_log(runtime: BackgroundRuntime | None, *, append_image_log, **payload) -> None:
    if runtime is None:
        append_image_log(**payload)
        return
    runtime.emit(AppendImageLogEvent(payload=payload))


@dataclass(frozen=True)
class RuntimeEventEmitterDependencies:
    set_processing_status: Callable[..., None]
    finalize_processing_status: Callable[..., None]
    push_activity: Callable[[str], None]
    append_log: Callable[..., None]
    append_image_log: Callable[..., None]


@dataclass(frozen=True)
class RuntimeEventEmitters:
    emit_state: Callable[..., None]
    emit_image_reset: Callable[..., None]
    emit_status: Callable[..., None]
    emit_finalize: Callable[..., None]
    emit_activity: Callable[..., None]
    emit_log: Callable[..., None]
    emit_image_log: Callable[..., None]


def build_runtime_event_emitters(*, dependencies: RuntimeEventEmitterDependencies) -> RuntimeEventEmitters:
    return RuntimeEventEmitters(
        emit_state=emit_or_apply_state,
        emit_image_reset=emit_or_apply_image_reset,
        emit_status=lambda runtime, **payload: emit_or_apply_status(
            runtime,
            set_processing_status=dependencies.set_processing_status,
            **payload,
        ),
        emit_finalize=lambda runtime, stage, detail, progress, terminal_kind=None: emit_or_apply_finalize(
            runtime,
            finalize_processing_status=dependencies.finalize_processing_status,
            stage=stage,
            detail=detail,
            progress=progress,
            terminal_kind=terminal_kind,
        ),
        emit_activity=lambda runtime, message: emit_or_apply_activity(
            runtime,
            push_activity=dependencies.push_activity,
            message=message,
        ),
        emit_log=lambda runtime, **payload: emit_or_apply_log(
            runtime,
            append_log=dependencies.append_log,
            **payload,
        ),
        emit_image_log=lambda runtime, **payload: emit_or_apply_image_log(
            runtime,
            append_image_log=dependencies.append_image_log,
            **payload,
        ),
    )


def _is_stale_processing_event(event: ProcessingEvent) -> bool:
    event_source_token = str(getattr(event, "source_token", "") or "")
    if not event_source_token:
        return False
    latest_source_token = str(get_latest_source_token() or "")
    return event_source_token != latest_source_token


def drain_processing_events(*, set_processing_status, finalize_processing_status, push_activity, append_log, append_image_log) -> None:
    event_queue = get_processing_event_queue()
    if event_queue is None:
        return
    while True:
        try:
            event = event_queue.get_nowait()
        except queue.Empty:
            break

        if _is_stale_processing_event(event):
            continue

        if isinstance(event, SetStateEvent):
            allowed_values = _filter_allowed_set_state_values(event.values)
            if allowed_values:
                set_session_values(**allowed_values)
        elif isinstance(event, ResetImageStateEvent):
            reset_image_state()
        elif isinstance(event, SetProcessingStatusEvent):
            set_processing_status(**event.payload)
        elif isinstance(event, FinalizeProcessingStatusEvent):
            finalize_processing_status(event.stage, event.detail, event.progress, event.terminal_kind)
        elif isinstance(event, PushActivityEvent):
            push_activity(event.message)
        elif isinstance(event, AppendLogEvent):
            append_log(**event.payload)
        elif isinstance(event, AppendImageLogEvent):
            append_image_log(**event.payload)
        elif isinstance(event, WorkerCompleteEvent):
            apply_processing_completion(
                outcome=event.outcome,
                push_activity=push_activity,
                load_restart_source_bytes_fn=load_restart_source_bytes,
                clear_restart_source_fn=clear_restart_source,
                store_completed_source_fn=store_completed_source,
                should_cache_completed_source_fn=should_cache_completed_source,
                log_event_fn=log_event,
            )


def drain_preparation_events(*, reset_run_state, set_processing_status, finalize_processing_status, push_activity) -> None:
    event_queue = get_preparation_event_queue()
    if event_queue is None:
        return
    while True:
        try:
            event = event_queue.get_nowait()
        except queue.Empty:
            break

        active_upload_marker = str(st.session_state.get("preparation_input_marker", "") or "")
        if isinstance(event, (PreparationCompleteEvent, PreparationFailedEvent, PreparationStoppedEvent)) and active_upload_marker and event.upload_marker != active_upload_marker:
            continue

        if isinstance(event, SetProcessingStatusEvent):
            set_processing_status(**event.payload)
        elif isinstance(event, FinalizeProcessingStatusEvent):
            finalize_processing_status(event.stage, event.detail, event.progress, event.terminal_kind)
        elif isinstance(event, PushActivityEvent):
            push_activity(event.message)
        elif isinstance(event, PreparationCompleteEvent):
            apply_preparation_complete(
                prepared_run_context=event.prepared_run_context,
                upload_marker=event.upload_marker,
                reset_run_state_fn=reset_run_state,
            )
            finalize_processing_status(
                "Документ подготовлен",
                "",
                1.0,
                "completed",
            )
        elif isinstance(event, PreparationFailedEvent):
            apply_preparation_failure(
                upload_marker=event.upload_marker,
                error_message=event.error_message,
                error_details=event.error_details,
            )
            finalize_processing_status(
                "Ошибка подготовки",
                event.error_message,
                1.0,
                "error",
            )
            push_activity("Не удалось прочитать и проанализировать документ.")
        elif isinstance(event, PreparationStoppedEvent):
            apply_preparation_stop(upload_marker=event.upload_marker)
            finalize_processing_status(
                "Подготовка остановлена",
                "",
                1.0,
                "stopped",
            )
            push_activity("Подготовка документа остановлена.")


def preparation_worker_is_active() -> bool:
    worker = get_preparation_worker()
    return worker is not None and worker.is_alive()


def processing_worker_is_active() -> bool:
    worker = get_processing_worker()
    return worker is not None and worker.is_alive()


def request_processing_stop() -> None:
    request_processing_stop_via_state()


def start_background_processing(
    *,
    worker_target,
    reset_run_state,
    push_activity,
    set_processing_status,
    uploaded_filename: str,
    uploaded_token: str,
    source_bytes: bytes,
    prepared_source_key: str | None = None,
    structure_fingerprint: str | None = None,
    jobs: list[dict[str, str | int]],
    selected_segment_ids: list[str] | None = None,
    document_segments: list | None = None,
    output_mode: str | None = None,
    include_front_matter: bool = False,
    include_toc: bool = False,
    source_paragraphs: list | None = None,
    image_assets: list,
    image_mode: str,
    app_config: dict[str, object],
    model: str,
    max_retries: int,
    processing_operation: str = "edit",
    source_language: str = "en",
    target_language: str = "ru",
) -> None:
    previous_restart_source = get_restart_source()
    restart_session_id = str(st.session_state.get("restart_session_id", ""))
    reset_run_state(preserve_preparation=True)
    try:
        set_restart_source(store_restart_source(
            session_id=restart_session_id,
            source_name=uploaded_filename,
            source_token=uploaded_token,
            source_bytes=source_bytes,
            previous_restart_source=previous_restart_source,
        ))
    except OSError as exc:
        set_restart_source(None)
        log_event(
            logging.WARNING,
            "restart_source_store_failed",
            "Не удалось сохранить временный restart source; продолжаю обработку без возможности restart без повторной загрузки.",
            filename=uploaded_filename,
            source_token=uploaded_token,
            error_message=str(exc),
        )
        push_activity("Не удалось сохранить временный файл для restart. Повторный запуск без загрузки файла будет недоступен.")

    processing_events = queue.Queue()
    stop_event = threading.Event()
    runtime = BackgroundRuntime(processing_events, stop_event, source_token=uploaded_token)
    run_id = uuid4().hex

    push_activity("Запуск обработки документа.")
    set_processing_status(
        stage="Инициализация",
        detail="Проверяю доступность OpenAI, Pandoc и системного промпта.",
        current_block=0,
        block_count=len(jobs),
        progress=0.0,
        is_running=True,
    )

    def _admission_guarded_worker_target(**worker_kwargs) -> None:
        # F27: acquire a process-wide admission slot before doing real work so
        # concurrent sessions cannot multiply PDF RAM / subprocess / API cost.
        # Spec 041 P1-2: the wait is cancellable (mirrors the preparation path)
        # so a Stop while the gate is full cancels the queued run instead of
        # letting it run once a slot frees. ``stop_event`` and ``runtime`` are
        # this run's own event/runtime captured from the enclosing scope.
        if not _acquire_admission_slot_cancellable(_PROCESSING_ADMISSION_GATE, stop_event):
            # Cancelled during the wait: nothing was acquired, so do NOT release.
            # Surface the same stopped completion a normally-stopped run emits.
            runtime.emit(WorkerCompleteEvent(outcome="stopped"))
            return
        try:
            worker_target(**worker_kwargs)
        finally:
            # A slot was acquired above, so release it on every exit path
            # (completion, stop_event, or error).
            _PROCESSING_ADMISSION_GATE.release()

    worker = threading.Thread(
        target=_admission_guarded_worker_target,
        kwargs={
            "runtime": runtime,
            "uploaded_filename": uploaded_filename,
            "source_token": uploaded_token,
            "run_id": run_id,
            "prepared_source_key": prepared_source_key,
            "structure_fingerprint": structure_fingerprint,
            "jobs": jobs,
            "selected_segment_ids": selected_segment_ids,
            "document_segments": document_segments,
            "output_mode": output_mode,
            "include_front_matter": include_front_matter,
            "include_toc": include_toc,
            "source_paragraphs": source_paragraphs,
            "image_assets": image_assets,
            "image_mode": image_mode,
            "app_config": app_config,
            "model": model,
            "max_retries": max_retries,
            "processing_operation": processing_operation,
            "source_language": source_language,
            "target_language": target_language,
        },
        daemon=True,
    )
    apply_processing_start(
        uploaded_filename=uploaded_filename,
        uploaded_token=uploaded_token,
        image_mode=image_mode,
        processing_operation=processing_operation,
        audiobook_postprocess_enabled=bool(app_config.get("audiobook_postprocess_enabled", False)),
        worker=worker,
        event_queue=processing_events,
        stop_event=stop_event,
    )
    worker.start()


def start_background_preparation(
    *,
    worker_target,
    reset_run_state,
    push_activity,
    set_processing_status,
    uploaded_payload,
    upload_marker: str,
    chunk_size: int,
    image_mode: str,
    keep_all_image_variants: bool,
    processing_operation: str = "edit",
    app_config: dict[str, object] | None = None,
    client_factory=None,
) -> None:
    reset_run_state(keep_restart_source=False)
    mark_preparation_started(upload_marker)

    preparation_events = queue.Queue()
    preparation_stop_event = threading.Event()
    runtime = BackgroundRuntime(preparation_events, preparation_stop_event)

    push_activity("Файл получен сервером. Запускаю анализ документа.")
    set_processing_status(
        stage="Файл получен",
        detail="Файл передан на сервер. Запускаю анализ документа.",
        progress=0.02,
        is_running=True,
        phase="preparing",
        source_format=str(getattr(uploaded_payload, "source_format", "docx") or "docx"),
        conversion_backend=getattr(uploaded_payload, "conversion_backend", None),
        conversion_reused=False,
    )

    last_reported_stage = {"value": ""}
    last_activity = {"key": "", "at": 0.0}

    def report_progress(*, stage: str, detail: str, progress: float, metrics: dict[str, object]) -> None:
        runtime.emit(
            SetProcessingStatusEvent(
                payload={
                    "stage": stage,
                    "detail": detail,
                    "progress": progress,
                    "is_running": True,
                    "phase": "preparing",
                    "block_count": _coerce_metric_to_int(metrics, "block_count"),
                    "file_size_bytes": _coerce_metric_to_int(metrics, "file_size_bytes"),
                    "paragraph_count": _coerce_metric_to_int(metrics, "paragraph_count"),
                    "image_count": _coerce_metric_to_int(metrics, "image_count"),
                    "source_chars": _coerce_metric_to_int(metrics, "source_chars"),
                    "cached": bool(metrics.get("cached", False)),
                    "conversion_reused": bool(metrics.get("conversion_reused", False)),
                    "source_format": str(metrics.get("source_format") or "docx"),
                    "conversion_backend": str(metrics.get("conversion_backend") or "") or None,
                }
            )
        )
        # Activity feed: emit on stage change, and additionally throttle per-detail
        # heartbeats so users see live progress instead of stage-only entries.
        now = time.monotonic()
        if stage and stage != last_reported_stage["value"]:
            runtime.emit(PushActivityEvent(message=f"[Анализ] {stage}: {detail}"))
            last_reported_stage["value"] = stage
            last_activity["key"] = f"{stage}|{detail}"
            last_activity["at"] = now
            return
        activity_key = f"{stage}|{detail}"
        if (
            stage
            and detail
            and activity_key != last_activity["key"]
            and (now - float(last_activity["at"] or 0.0)) >= 3.0
        ):
            runtime.emit(PushActivityEvent(message=f"[Анализ] {detail}"))
            last_activity["key"] = activity_key
            last_activity["at"] = now

    def _run_preparation_stages() -> None:
        try:
            materialized_payload = materialize_uploaded_payload(
                uploaded_payload,
                progress_callback=report_progress,
            )
        except Exception as exc:
            error_details = normalize_background_error(
                stage="preparation",
                exc=exc,
                user_message=str(exc),
            )
            runtime.emit(
                PreparationFailedEvent(
                    upload_marker=upload_marker,
                    error_message=str(error_details["user_message"]),
                    error_details=error_details,
                )
            )
            return
        if runtime.should_stop():
            runtime.emit(PreparationStoppedEvent(upload_marker=upload_marker))
            return
        try:
            worker_kwargs: dict[str, object] = {
                "uploaded_payload": materialized_payload,
                "chunk_size": chunk_size,
                "image_mode": image_mode,
                "keep_all_image_variants": keep_all_image_variants,
                "processing_operation": processing_operation,
                "app_config": app_config,
                "progress_callback": report_progress,
            }
            # Forward the tenant client_factory (spec 039 part B) so the UI
            # preparation path honors per-tenant endpoint/credentials, mirroring
            # ProcessingService. None keeps the worker call byte-compatible.
            if client_factory is not None:
                worker_kwargs["client_factory"] = client_factory
            prepared_run_context = worker_target(**worker_kwargs)
        except Exception as exc:
            error_details = normalize_background_error(
                stage="preparation",
                exc=exc,
                user_message=str(exc),
            )
            runtime.emit(
                PreparationFailedEvent(
                    upload_marker=upload_marker,
                    error_message=str(error_details["user_message"]),
                    error_details=error_details,
                )
            )
            return
        runtime.emit(PreparationCompleteEvent(prepared_run_context=prepared_run_context, upload_marker=upload_marker))

    def run_preparation() -> None:
        # F27: PDF materialization/parse is the costliest stage, so preparation
        # must hold a process-wide admission slot too — not just main processing.
        # The wait is cancellable so a stopped upload never blocks a slot forever;
        # the slot is released on every exit path (complete, stop, failure).
        if not _acquire_admission_slot_cancellable(_PROCESSING_ADMISSION_GATE, preparation_stop_event):
            runtime.emit(PreparationStoppedEvent(upload_marker=upload_marker))
            return
        try:
            _run_preparation_stages()
        finally:
            _PROCESSING_ADMISSION_GATE.release()

    worker = threading.Thread(target=run_preparation, daemon=True, name="docx-preparation-worker")
    set_preparation_runtime(worker=worker, event_queue=preparation_events, stop_event=preparation_stop_event)
    worker.start()
