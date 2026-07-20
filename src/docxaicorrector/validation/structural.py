from __future__ import annotations

import argparse
from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from io import BytesIO
import json
from pathlib import Path
from typing import Any, cast

from docx import Document
from docx.oxml.ns import qn

import docxaicorrector.processing.processing_runtime as processing_runtime
from docxaicorrector.core.config import get_client, get_client_for_model_selector, get_provider_client, load_app_config, resolve_model_selector
from docxaicorrector.document.extraction import (
    extract_document_content_from_docx,
    extract_document_content_with_normalization_reports,
    extract_document_content_with_boundary_report,
    inspect_placeholder_integrity,
)
from docxaicorrector.document.semantic_blocks import build_semantic_blocks
from docxaicorrector.generation.formatting_diagnostics_retention import write_formatting_diagnostics_artifact
from docxaicorrector.generation.formatting_transfer import preserve_source_paragraph_properties
from docxaicorrector.generation._generation import convert_markdown_to_docx_bytes, ensure_pandoc_available
from docxaicorrector.image.reinsertion import reinsert_inline_images
from docxaicorrector.processing.processing_service import clone_processing_service
from docxaicorrector.validation.common import build_validation_event_logger, build_validation_runtime_config
from docxaicorrector.validation.formatting_coverage import (
    resolve_role_aware_formatting_unmapped_source_summary,
)
from docxaicorrector.validation.profiles import (
    DocumentProfile,
    RunProfile,
    apply_runtime_resolution_to_app_config,
    load_validation_registry,
    resolve_runtime_resolution,
)
from docxaicorrector.processing.preparation import flatten_structure_repair_metrics
from docxaicorrector.validation.structural_metrics_common import (  # noqa: F401
    _as_float,
    _as_int,
)
from docxaicorrector.validation.structural_event_log import (  # noqa: F401
    _extract_event_context,
    _extract_event_context_bool,
    _extract_event_context_float,
    _extract_event_context_int,
    _extract_event_context_int_list,
    _extract_event_context_list,
    _extract_event_context_value,
)
from docxaicorrector.validation.structural_text_metrics import (  # noqa: F401
    _calculate_heading_level_drift,
    _calculate_text_similarity,
    _count_bullet_headings,
    _has_toc_body_concat_markdown,
    _has_toc_structural_roles,
    _is_heading_only_markdown,
    _normalize_text,
    _relation_count,
)
from docxaicorrector.validation.structural_toc_signals import (  # noqa: F401
    _count_authoritative_topology_operations,
    _count_authoritative_topology_toc_entry_units,
    _count_document_map_anchor_roles,
    _count_effective_toc_regions_from_source,
    _count_high_confidence_compound_toc_split_hints,
    _count_topology_operations,
    _count_topology_toc_entry_units,
    _derive_toc_body_concat_gate_fields,
    _document_map_has_high_confidence_outline_inside_toc_region,
    _has_high_confidence_bounded_document_map_toc_region,
    _is_authoritative_topology_signal,
    _normalized_structural_role,
    _projection_has_heading_inside_toc_region,
    _projection_has_toc_entry_outside_toc_region,
    _projection_has_units_or_operations,
    _projection_supports_toc_body_concat_gate,
    _resolve_bounded_toc_region_range,
    has_toc_body_concat_structure,
)
from docxaicorrector.validation.structural_unit_alignment import (  # noqa: F401
    _align_target_indexes_from_generated_registry,
    _align_target_indexes_to_unit_keys,
    _build_generated_registry_text_by_paragraph_id,
    _build_source_paragraph_unit_membership,
    _build_target_alignments_from_source_registry,
    _collect_target_alignment_preview_trace,
    _derive_unit_aware_unmapped_fields,
    _generated_paragraph_spans_empty_body_interval_target,
    _infer_target_alignment_unit_keys_from_source_intervals,
    _logical_index_for_unit_accounting,
    _merge_target_alignment_unit_keys,
    _normalize_registry_preview_for_unit_alignment,
    _normalize_registry_text_for_unit_alignment,
    _paragraph_id_for_unit_accounting,
    _projection_units_for_logical_index,
    _registry_entry_relation_ids,
    _registry_entry_unit_keys,
    _registry_text_matches_target_preview,
    _serialize_compact_target_alignment_trace_entry,
    _truncate_target_alignment_trace_preview,
)
from docxaicorrector.validation.structural_checks import (  # noqa: F401
    _apply_metric_snapshot_fields,
    _build_extraction_checks,
    _build_markdown_quality_metrics,
    _build_sentinel_threshold_checks,
    _build_structural_checks,
    _build_structural_metrics,
    _check_minimum,
    _is_effectively_infinite_threshold,
)
from docxaicorrector.validation.structural_prep_snapshot_helpers import (  # noqa: F401
    _HUMANIZED_QUALITY_GATE_REASON_TO_CODE,
    _apply_prepared_metric_fields,
    _apply_prepared_snapshot_fields,
    _apply_preparation_error_snapshot_fallback,
    _apply_quality_gate_readiness_fallback,
    _apply_structure_summary_snapshot_fields,
    _apply_structure_validation_snapshot_fields,
    _block_has_body_start,
    _block_has_epigraph,
    _block_has_isolated_marker,
    _block_has_toc,
    _build_layout_signals_snapshot_context,
    _extract_first_block_target_chars,
    _extract_quality_gate_reasons_from_error,
    _infer_readiness_status_from_quality_gate_reasons,
    _is_isolated_marker_text,
    _maybe_backfill_layout_signals_snapshot,
    _normalize_snapshot_or_metric_statuses,
)
from docxaicorrector.structure.validation import validate_structure_quality

PROJECT_ROOT = Path(__file__).resolve().parents[3]


def _extract_quality_report_artifact_path(event_log: Sequence[Mapping[str, object]]) -> str | None:
    for event in reversed(event_log):
        if str(event.get("event_id") or "").strip() != "quality_report_saved":
            continue
        context = event.get("context")
        if not isinstance(context, Mapping):
            continue
        artifact_path = context.get("artifact_path")
        if isinstance(artifact_path, str) and artifact_path.strip():
            return artifact_path.strip()
    return None


def _load_translation_quality_report(event_log: Sequence[Mapping[str, object]]) -> tuple[dict[str, object] | None, str | None]:
    artifact_path = _extract_quality_report_artifact_path(event_log)
    if not artifact_path:
        return None, None
    candidate = Path(artifact_path)
    if not candidate.is_absolute():
        candidate = (PROJECT_ROOT / candidate).resolve()
    try:
        payload = json.loads(candidate.read_text(encoding="utf-8"))
    except Exception:
        return None, str(candidate)
    if not isinstance(payload, dict):
        return None, str(candidate)
    return payload, str(candidate)


def _merge_translation_quality_report_metrics(
    metrics: dict[str, object],
    translation_quality_report: Mapping[str, object],
) -> None:
    for key in (
        "unmapped_source_count_basis",
        "worst_unmapped_source_count",
        "raw_unmapped_source_paragraph_count",
        "filtered_unmapped_source_count",
        "format_neutral_creditable_count",
        "effective_unmapped_source_count",
        "role_aware_formatting_coverage_note",
        "bullet_heading_count",
        "bullet_heading_gate_source",
        "bullet_heading_classification",
        "raw_bullet_heading_count",
        "false_fragment_heading_count",
        "false_fragment_heading_gate_source",
        "raw_false_fragment_heading_count",
        "page_placeholder_heading_concat_count",
        "page_placeholder_heading_concat_source",
        "page_placeholder_heading_concat_classification",
        "raw_page_placeholder_heading_concat_count",
        "residual_bullet_glyph_count",
        "residual_bullet_glyph_gate_source",
        "residual_bullet_glyph_classification",
        "raw_residual_bullet_glyph_count",
        "scripture_reference_heading_count",
        "suspicious_heading_repetition_count",
        "list_fragment_regression_count",
        "list_fragment_regression_gate_source",
        "raw_list_fragment_regression_count",
        "mixed_script_term_count",
        "mixed_script_term_gate_source",
        "mixed_script_term_classification",
        "raw_mixed_script_term_count",
        "theology_style_deterministic_issue_count",
        "theology_style_deterministic_issue_source",
        "theology_style_deterministic_issue_classification",
        "raw_theology_style_deterministic_issue_count",
    ):
        if key in translation_quality_report:
            metrics[key] = translation_quality_report[key]


def _emit_target_alignment_trace_artifact(
    *,
    source_paragraphs: Sequence[object],
    topology_projection: object | None,
    formatting_payload: Mapping[str, object] | None,
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None,
) -> str | None:
    if formatting_payload is None or not generated_paragraph_registry:
        return None

    raw_unmapped_target_indexes = formatting_payload.get("unmapped_target_indexes")
    if not isinstance(raw_unmapped_target_indexes, list):
        return None

    unmapped_target_indexes: list[int] = []
    for value in raw_unmapped_target_indexes:
        try:
            unmapped_target_indexes.append(int(cast(Any, value)))
        except (TypeError, ValueError):
            continue
    if not unmapped_target_indexes:
        return None

    paragraph_unit_keys, _ = _build_source_paragraph_unit_membership(source_paragraphs, topology_projection)
    generated_registry_alignments, _ = _align_target_indexes_from_generated_registry(
        formatting_payload,
        generated_paragraph_registry=generated_paragraph_registry,
        paragraph_unit_keys=paragraph_unit_keys,
        trace_target_indexes=unmapped_target_indexes,
    )
    full_alignments = _align_target_indexes_to_unit_keys(
        formatting_payload,
        generated_paragraph_registry=generated_paragraph_registry,
        paragraph_unit_keys=paragraph_unit_keys,
    ) or {}
    compact_trace = [
        _serialize_compact_target_alignment_trace_entry(entry)
        for entry in _collect_target_alignment_preview_trace(
            formatting_payload,
            generated_paragraph_registry=generated_paragraph_registry,
            paragraph_unit_keys=paragraph_unit_keys,
            target_indexes=unmapped_target_indexes,
        )
    ]
    aligned_target_unit_keys_snapshot = {
        str(target_index): sorted(full_alignments.get(target_index, frozenset())) or None
        for target_index in unmapped_target_indexes
    }
    aligned_via_generated_registry = sum(
        1 for target_index in unmapped_target_indexes if generated_registry_alignments.get(target_index)
    )
    aligned_via_full_inference = sum(1 for target_index in unmapped_target_indexes if full_alignments.get(target_index))

    return write_formatting_diagnostics_artifact(
        stage="target_alignment_trace",
        filename_prefix="target_alignment_trace",
        scope="offline",
        diagnostics={
            "unmapped_target_indexes": unmapped_target_indexes,
            "generated_registry_alignment_trace": compact_trace,
            "aligned_target_unit_keys_snapshot": aligned_target_unit_keys_snapshot,
            "alignment_coverage_summary": {
                "total_unmapped": len(unmapped_target_indexes),
                "aligned_via_generated_registry": aligned_via_generated_registry,
                "aligned_via_interval_inference": aligned_via_full_inference,
                "still_unaligned": max(0, len(unmapped_target_indexes) - aligned_via_full_inference),
            },
        },
    )


def _unpack_extraction_result(extraction_result):
    paragraphs, image_assets, normalization_report, relations, relation_report, cleanup_report = extraction_result[:6]
    structure_repair_report = extraction_result[6] if len(extraction_result) > 6 else None
    return (
        paragraphs,
        image_assets,
        normalization_report,
        relations,
        relation_report,
        cleanup_report,
        structure_repair_report,
    )


@dataclass
class UploadedFileStub:
    name: str
    content: bytes
    position: int = 0

    def read(self, size: int = -1) -> bytes:
        if size < 0:
            data = self.content[self.position :]
            self.position = len(self.content)
            return data
        start = self.position
        end = min(len(self.content), start + size)
        self.position = end
        return self.content[start:end]

    def getvalue(self) -> bytes:
        return self.content

    def seek(self, offset: int, whence: int = 0) -> int:
        if whence == 0:
            self.position = max(0, offset)
        elif whence == 1:
            self.position = max(0, self.position + offset)
        elif whence == 2:
            self.position = max(0, len(self.content) + offset)
        else:
            raise ValueError(f"Unsupported whence: {whence}")
        return self.position


def evaluate_extraction_profile(document_profile: DocumentProfile) -> dict[str, object]:
    source_path = document_profile.resolved_source_path(PROJECT_ROOT)
    source_bytes = source_path.read_bytes()
    normalized_source = processing_runtime.normalize_uploaded_document(filename=source_path.name, source_bytes=source_bytes)
    (
        paragraphs,
        image_assets,
        normalization_report,
        _,
        relation_report,
        cleanup_report,
        _,
    ) = _unpack_extraction_result(extract_document_content_with_normalization_reports(BytesIO(normalized_source.content_bytes)))
    metrics = _build_structural_metrics(
        paragraphs=paragraphs,
        image_assets=image_assets,
        normalization_report=normalization_report,
        relation_report=relation_report,
        cleanup_report=cleanup_report,
    )
    checks = _build_extraction_checks(document_profile, metrics)
    return _build_validation_result(
        document_profile=document_profile,
        run_profile=None,
        tier="extraction",
        source_path=source_path,
        result="succeeded",
        metrics=metrics,
        checks=checks,
        runtime_config=None,
        output_artifacts=None,
        formatting_diagnostics=[],
    )


def run_structural_passthrough_validation(
    document_profile: DocumentProfile,
    run_profile: RunProfile,
) -> dict[str, object]:
    source_path = document_profile.resolved_source_path(PROJECT_ROOT)
    source_bytes = source_path.read_bytes()
    app_config = load_app_config()
    runtime_resolution = resolve_runtime_resolution(app_config, run_profile)
    runtime_config = apply_runtime_resolution_to_app_config(app_config, runtime_resolution)
    runtime = _build_runtime_capture()
    event_log: list[dict[str, object]] = []
    try:
        result, prepared = _build_validation_processing_service(event_log).run_prepared_background_document(
            uploaded_file=UploadedFileStub(source_path.name, source_bytes),
            chunk_size=int(runtime_resolution.effective.chunk_size),
            image_mode=str(runtime_resolution.effective.image_mode),
            keep_all_image_variants=bool(runtime_resolution.effective.keep_all_image_variants),
            app_config=runtime_config,
            model=str(runtime_resolution.effective.model),
            max_retries=int(runtime_resolution.effective.max_retries),
            job_mutator=_build_passthrough_job,
            progress_callback=None,
            runtime=runtime,
        )
    except Exception as exc:
        checks = [
            {
                "name": "preparation_quality_gate_blocked",
                "passed": False,
                "error": str(exc),
            }
        ]
        preparation_diagnostic_snapshot = _build_preparation_diagnostic_snapshot_from_source(
            source_path=source_path,
            source_bytes=source_bytes,
            chunk_size=int(runtime_resolution.effective.chunk_size),
            event_log=event_log,
        )
        _apply_preparation_error_snapshot_fallback(preparation_diagnostic_snapshot, str(exc))
        return _build_validation_result(
            document_profile=document_profile,
            run_profile=run_profile,
            tier="structural",
            source_path=source_path,
            result="failed",
            metrics={
                "preparation_error": str(exc),
                "quality_gate_status": "",
                "quality_gate_reasons": [],
                "readiness_status": "",
            },
            checks=checks,
            runtime_config=build_validation_runtime_config(runtime_resolution),
            output_artifacts=None,
            formatting_diagnostics=[],
            event_log=event_log,
            preparation_diagnostic_snapshot=preparation_diagnostic_snapshot,
            validation_execution_mode="passthrough",
        )

    formatting_paths = _extract_run_formatting_diagnostics_paths(event_log)
    formatting_diagnostics = _load_formatting_diagnostics_payloads(formatting_paths)
    canonical_formatting_diagnostics = _select_canonical_formatting_diagnostics_payload(formatting_diagnostics)
    canonical_formatting_payloads = [] if canonical_formatting_diagnostics is None else [canonical_formatting_diagnostics]

    runtime_state = _runtime_state(runtime)
    latest_docx_bytes = runtime_state.get("latest_docx_bytes")
    latest_markdown = str(runtime_state.get("latest_markdown") or "")
    processed_block_markdowns = cast(Sequence[object], runtime_state.get("processed_block_markdowns") or [])
    raw_structural_markdown = "\n\n".join(
        str(item) for item in processed_block_markdowns if isinstance(item, str) and item.strip()
    )
    raw_markdown = raw_structural_markdown
    if not raw_markdown:
        raw_markdown = latest_markdown
    if not isinstance(latest_docx_bytes, (bytes, bytearray)):
        latest_docx_bytes = b""
    output_artifacts = _build_output_artifacts(bytes(latest_docx_bytes), latest_markdown)

    (
        source_paragraphs,
        source_image_assets,
        source_normalization_report,
        source_relations,
        source_relation_report,
        source_cleanup_report,
        source_structure_repair_report,
    ) = _unpack_extraction_result(extract_document_content_with_normalization_reports(BytesIO(prepared.uploaded_file_bytes)))
    preparation_diagnostic_snapshot = build_preparation_diagnostic_snapshot(
        paragraphs=source_paragraphs,
        relations=source_relations,
        structure_repair_report=source_structure_repair_report,
        chunk_size=int(runtime_resolution.effective.chunk_size),
        event_log=event_log,
    )
    _apply_prepared_snapshot_fields(
        preparation_diagnostic_snapshot,
        prepared,
        app_config=runtime_config,
    )
    output_paragraphs = []
    output_image_assets = []
    if output_artifacts["output_docx_openable"]:
        output_paragraphs, output_image_assets = extract_document_content_from_docx(BytesIO(bytes(latest_docx_bytes)))
    metrics = _build_structural_metrics(
        paragraphs=source_paragraphs,
        image_assets=source_image_assets,
        normalization_report=source_normalization_report,
        relation_report=source_relation_report,
        cleanup_report=source_cleanup_report,
    )
    metrics.update(flatten_structure_repair_metrics(source_structure_repair_report))
    metrics.update(
        {
            "output_paragraph_count": len(output_paragraphs),
            "output_heading_count": sum(1 for paragraph in output_paragraphs if paragraph.role == "heading"),
            "output_numbered_item_count": sum(
                1 for paragraph in output_paragraphs if paragraph.role == "list" and paragraph.list_kind == "ordered"
            ),
            "output_image_count": len(output_image_assets),
            "output_table_count": sum(1 for paragraph in output_paragraphs if paragraph.role == "table"),
            "formatting_diagnostics_count": len(canonical_formatting_payloads),
            "max_unmapped_source_paragraphs": _max_payload_length(canonical_formatting_payloads, "unmapped_source_ids"),
            "max_unmapped_target_paragraphs": _max_payload_length(canonical_formatting_payloads, "unmapped_target_indexes"),
            "accepted_merged_sources_count": _count_payload_items(canonical_formatting_payloads, "accepted_merged_sources"),
            "max_accepted_merged_sources": _max_accepted_merged_sources(canonical_formatting_payloads),
            "relation_count": source_relation_report.total_relations,
            "rejected_relation_candidate_count": source_relation_report.rejected_candidate_count,
            "relation_counts": dict(source_relation_report.relation_counts),
            "text_similarity": _calculate_text_similarity(source_paragraphs, output_paragraphs),
            "heading_level_drift": _calculate_heading_level_drift(source_paragraphs, output_paragraphs),
            "heading_only_output_detected": _is_heading_only_markdown(latest_markdown),
            "output_docx_openable": bool(output_artifacts["output_docx_openable"]),
            "source_toc_detected": _has_toc_structural_roles(source_paragraphs),
            "output_toc_detected": _has_toc_structural_roles(output_paragraphs),
            "source_toc_region_count": _relation_count(source_relation_report, "toc_region"),
            "effective_source_toc_region_count": _count_effective_toc_regions_from_source(source_paragraphs),
            "bullet_heading_count": _count_bullet_headings(latest_markdown),
            "bullet_heading_gate_source": "legacy_markdown",
            "bullet_heading_classification": "markdown_gate",
            "raw_bullet_heading_count": _count_bullet_headings(latest_markdown),
            "toc_body_concat_detected": _has_toc_body_concat_markdown(latest_markdown),
            "toc_body_concat_markdown_detected": _has_toc_body_concat_markdown(latest_markdown),
            "require_pdf_conversion_satisfied": source_path.suffix.lower() == ".pdf",
            "runtime_translation_domain": str(runtime_resolution.effective.translation_domain),
            "quality_gate_status": _extract_event_context_value(event_log, "structure_processing_outcome", "quality_gate_status"),
            "quality_gate_reasons": _extract_event_context_list(event_log, "structure_processing_outcome", "quality_gate_reasons"),
            "readiness_status": _extract_event_context_value(event_log, "structure_processing_outcome", "readiness_status"),
            "block_count": _extract_event_context_int(event_log, "block_plan_summary", "block_count"),
            "llm_block_count": _extract_event_context_int(event_log, "block_plan_summary", "llm_block_count"),
            "passthrough_block_count": _extract_event_context_int(event_log, "block_plan_summary", "passthrough_block_count"),
            "first_block_target_chars": _extract_event_context_int_list(
                event_log,
                "block_plan_summary",
                "first_block_target_chars",
            ),
        }
    )
    role_aware_summary = resolve_role_aware_formatting_unmapped_source_summary(canonical_formatting_payloads)
    if role_aware_summary is not None:
        metrics.update(role_aware_summary)
    metrics.update(
        _build_markdown_quality_metrics(
            latest_markdown=latest_markdown,
            raw_markdown=raw_markdown,
            raw_structural_markdown=raw_structural_markdown,
            translation_domain=str(runtime_resolution.effective.translation_domain),
        )
    )
    translation_quality_report, translation_quality_report_path = _load_translation_quality_report(event_log)
    if translation_quality_report is not None:
        _merge_translation_quality_report_metrics(metrics, translation_quality_report)
    metrics["pdf_blank_page_marker_leakage_threshold"] = getattr(document_profile, "max_pdf_blank_page_marker_leakage", None)
    metrics["inline_page_furniture_leakage_threshold"] = getattr(document_profile, "max_inline_page_furniture_leakage", None)
    metrics["adjacent_h1_without_body_threshold"] = getattr(document_profile, "max_adjacent_h1_without_body", None)
    metrics["heading_body_concat_detected_threshold"] = getattr(document_profile, "max_heading_body_concat_detected", None)
    metrics["h1_epigraph_attribution_pattern_threshold"] = getattr(document_profile, "max_h1_epigraph_attribution_pattern", None)
    if translation_quality_report_path:
        metrics["translation_quality_report_path"] = translation_quality_report_path
    generated_paragraph_registry = runtime_state.get("generated_paragraph_registry")
    if not isinstance(generated_paragraph_registry, list):
        generated_paragraph_registry = None
    _apply_prepared_metric_fields(
        metrics,
        prepared,
        source_paragraphs=source_paragraphs,
        formatting_payload=canonical_formatting_diagnostics,
        generated_paragraph_registry=cast(Sequence[Mapping[str, object]] | None, generated_paragraph_registry),
    )
    target_alignment_trace_path = _emit_target_alignment_trace_artifact(
        source_paragraphs=source_paragraphs,
        topology_projection=getattr(prepared, "document_topology_projection", None),
        formatting_payload=canonical_formatting_diagnostics,
        generated_paragraph_registry=cast(Sequence[Mapping[str, object]] | None, generated_paragraph_registry),
    )
    if target_alignment_trace_path:
        formatting_diagnostics.extend(
            _load_formatting_diagnostics_payloads([target_alignment_trace_path])
        )
    _apply_metric_snapshot_fields(preparation_diagnostic_snapshot, metrics)
    _normalize_snapshot_or_metric_statuses(metrics)
    checks = _build_extraction_checks(document_profile, metrics)
    checks.extend(
        _build_structural_checks(
            document_profile=document_profile,
            result=result,
            metrics=metrics,
            output_artifacts=output_artifacts,
        )
    )
    return _build_validation_result(
        document_profile=document_profile,
        run_profile=run_profile,
        tier="structural",
        source_path=source_path,
        result=result,
        metrics=metrics,
        checks=checks,
        runtime_config=build_validation_runtime_config(runtime_resolution),
        output_artifacts=output_artifacts,
        formatting_diagnostics=formatting_diagnostics,
        event_log=event_log,
        preparation_diagnostic_snapshot=preparation_diagnostic_snapshot,
        validation_execution_mode="passthrough",
    )


def evaluate_structural_preparation_diagnostic(
    document_profile: DocumentProfile,
    run_profile: RunProfile,
) -> dict[str, object]:
    result = run_structural_passthrough_validation(document_profile, run_profile)
    metrics = cast(dict[str, object], result.get("metrics") or {})
    return {
        "document_profile_id": document_profile.id,
        "run_profile_id": run_profile.id,
        "validation_tier": result.get("validation_tier"),
        "validation_execution_mode": result.get("validation_execution_mode"),
        "passed": bool(result.get("passed")),
        "failed_checks": list(cast(list[str], result.get("failed_checks") or [])),
        "preparation_error": metrics.get("preparation_error"),
        "preparation_diagnostic_snapshot": result.get("preparation_diagnostic_snapshot"),
    }


def _build_validation_result(
    *,
    document_profile: DocumentProfile,
    run_profile: RunProfile | None,
    tier: str,
    source_path: Path,
    result: str,
    metrics: dict[str, object],
    checks: list[dict[str, object]],
    runtime_config: dict[str, object] | None,
    output_artifacts: dict[str, object] | None,
    formatting_diagnostics: list[dict[str, object]],
    event_log: list[dict[str, object]] | None = None,
    preparation_diagnostic_snapshot: dict[str, object] | None = None,
    validation_execution_mode: str | None = None,
) -> dict[str, object]:
    failed_checks = [str(check["name"]) for check in checks if not bool(check["passed"])]
    return {
        "document_profile_id": document_profile.id,
        "run_profile_id": None if run_profile is None else run_profile.id,
        "validation_tier": tier,
        "validation_execution_mode": validation_execution_mode,
        "source_document_path": str(source_path.relative_to(PROJECT_ROOT)).replace("\\", "/"),
        "result": result,
        "passed": not failed_checks,
        "failed_checks": failed_checks,
        "metrics": metrics,
        "checks": checks,
        "runtime_config": runtime_config,
        "output_artifacts": output_artifacts,
        "formatting_diagnostics": formatting_diagnostics,
        "event_log": event_log or [],
        "preparation_diagnostic_snapshot": preparation_diagnostic_snapshot,
    }


def _build_preparation_diagnostic_snapshot_from_source(
    *,
    source_path: Path,
    source_bytes: bytes,
    chunk_size: int,
    event_log: Sequence[Mapping[str, object]],
) -> dict[str, object]:
    try:
        normalized_source = processing_runtime.normalize_uploaded_document(filename=source_path.name, source_bytes=source_bytes)
        (
            paragraphs,
            _,
            _,
            relations,
            _,
            _,
            structure_repair_report,
        ) = _unpack_extraction_result(
            extract_document_content_with_normalization_reports(BytesIO(normalized_source.content_bytes))
        )
    except Exception as exc:
        snapshot = _build_preparation_diagnostic_defaults(event_log)
        snapshot["snapshot_error"] = str(exc)
        return snapshot
    snapshot = build_preparation_diagnostic_snapshot(
        paragraphs=paragraphs,
        relations=relations,
        structure_repair_report=structure_repair_report,
        chunk_size=chunk_size,
        event_log=event_log,
    )
    app_config = load_app_config()
    if isinstance(app_config, Mapping):
        app_config_mapping = cast(Mapping[str, Any], app_config)
    elif hasattr(app_config, "to_dict") and callable(getattr(app_config, "to_dict")):
        app_config_mapping = cast(Mapping[str, Any], app_config.to_dict())
    else:
        app_config_mapping = {}
    structure_validation_report = validate_structure_quality(
        paragraphs=cast(Sequence[Any], paragraphs),
        app_config=app_config_mapping,
        structure_repair_report=structure_repair_report,
    )
    _apply_structure_validation_snapshot_fields(snapshot, structure_validation_report)
    return snapshot


def build_preparation_diagnostic_snapshot(
    *,
    paragraphs: Sequence[object],
    relations: Sequence[object] | None,
    structure_repair_report: object | None,
    chunk_size: int,
    event_log: Sequence[Mapping[str, object]],
) -> dict[str, object]:
    snapshot = _build_preparation_diagnostic_defaults(event_log)
    paragraph_units = list(paragraphs)
    semantic_blocks = build_semantic_blocks(
        cast(list[Any], paragraph_units),
        max_chars=chunk_size,
        relations=None if relations is None else cast(list[Any], list(relations)),
    )
    first_block_paragraphs = [] if not semantic_blocks else list(semantic_blocks[0].paragraphs)
    first_block_target_chars = _extract_first_block_target_chars(
        event_log=event_log,
        semantic_blocks=semantic_blocks,
    )
    snapshot.update(
        {
            "paragraph_count": len(paragraph_units),
            "heading_count": sum(1 for paragraph in paragraph_units if getattr(paragraph, "role", None) == "heading"),
            "toc_header_count": sum(
                1
                for paragraph in paragraph_units
                if str(getattr(paragraph, "structural_role", "") or "").strip().lower() == "toc_header"
            ),
            "toc_entry_count": sum(
                1
                for paragraph in paragraph_units
                if str(getattr(paragraph, "structural_role", "") or "").strip().lower() == "toc_entry"
            ),
            "bounded_toc_region_count": int(getattr(structure_repair_report, "bounded_toc_regions", 0) or 0),
            "repaired_bullet_items": int(getattr(structure_repair_report, "repaired_bullet_items", 0) or 0),
            "repaired_numbered_items": int(getattr(structure_repair_report, "repaired_numbered_items", 0) or 0),
            "toc_body_boundary_repairs": int(getattr(structure_repair_report, "toc_body_boundary_repairs", 0) or 0),
            "remaining_isolated_marker_count": int(
                getattr(structure_repair_report, "remaining_isolated_marker_count", 0) or 0
            ),
            "semantic_block_count": len(semantic_blocks),
            "first_block_target_chars": first_block_target_chars,
            "first_block_has_toc": _block_has_toc(first_block_paragraphs),
            "first_block_has_epigraph": _block_has_epigraph(first_block_paragraphs),
            "first_block_has_body_start": _block_has_body_start(first_block_paragraphs),
            "first_block_has_isolated_marker": _block_has_isolated_marker(first_block_paragraphs),
        }
    )
    return snapshot


def _build_preparation_diagnostic_defaults(event_log: Sequence[Mapping[str, object]]) -> dict[str, object]:
    outline_coverage_ratio = _extract_event_context_float(event_log, "structure_processing_outcome", "outline_coverage_ratio")
    if outline_coverage_ratio is None:
        outline_coverage_ratio = _extract_event_context_float(event_log, "reconciliation_report_saved", "outline_coverage_ratio")
    document_map_status = str(
        _extract_event_context_value(event_log, "structure_processing_outcome", "document_map_status") or ""
    ).strip()
    document_map_status_reason = str(
        _extract_event_context_value(event_log, "structure_processing_outcome", "document_map_status_reason") or ""
    ).strip()
    topology_status = str(
        _extract_event_context_value(event_log, "structure_processing_outcome", "document_topology_projection_status") or ""
    ).strip()
    topology_status_reason = str(
        _extract_event_context_value(event_log, "structure_processing_outcome", "document_topology_projection_status_reason") or ""
    ).strip()
    topology_artifact_path = ""
    for event in reversed(list(event_log)):
        event_id = str(event.get("event_id") or "").strip()
        if event_id not in {"document_topology_projection_built", "document_topology_projection_skipped"}:
            continue
        context = cast(Mapping[str, object], event.get("context") or {})
        if not topology_status:
            if event_id == "document_topology_projection_built":
                topology_status = "built"
            else:
                reason = str(context.get("reason") or "").strip()
                topology_status = "no_operations" if reason == "no_operations" else "skipped"
                if not topology_status_reason:
                    topology_status_reason = reason
        topology_artifact_path = str(context.get("artifact_path") or "").strip()
        if topology_artifact_path:
            break
    topology_projection = None
    if topology_artifact_path:
        artifact_path = Path(topology_artifact_path)
        if not artifact_path.is_absolute():
            artifact_path = (PROJECT_ROOT / artifact_path).resolve()
        if artifact_path.exists():
            try:
                topology_projection = json.loads(artifact_path.read_text(encoding="utf-8"))
            except (OSError, ValueError, json.JSONDecodeError):
                topology_projection = None
    layout_signals_context = dict(_extract_event_context(event_log, "document_topology_layout_signals_built"))
    snapshot = {
        "paragraph_count": 0,
        "heading_count": 0,
        "toc_header_count": 0,
        "toc_entry_count": 0,
        "bounded_toc_region_count": 0,
        "repaired_bullet_items": 0,
        "repaired_numbered_items": 0,
        "toc_body_boundary_repairs": 0,
        "remaining_isolated_marker_count": 0,
        "readiness_status": _extract_event_context_value(event_log, "structure_processing_outcome", "readiness_status"),
        "readiness_reasons": _extract_event_context_list(event_log, "structure_processing_outcome", "readiness_reasons"),
        "document_map_present": _extract_event_context_bool(event_log, "structure_processing_outcome", "document_map_present"),
        "outline_coverage_ratio": outline_coverage_ratio,
        "document_topology_projection_status": topology_status or "not_requested",
        "document_topology_projection_status_reason": topology_status_reason,
        "document_topology_projection": topology_projection,
        "document_topology_layout_signals": None if not layout_signals_context else layout_signals_context,
        "front_matter_leaks": _extract_event_context_int_list(event_log, "reconciliation_report_saved", "front_matter_leaks"),
        "front_matter_body_advisories": _extract_event_context_int_list(
            event_log,
            "reconciliation_report_saved",
            "front_matter_body_advisories",
        ),
        "targeted_recall_invoked": _extract_event_context_bool(event_log, "reconciliation_report_saved", "targeted_recall_invoked"),
        "quality_gate_status": _extract_event_context_value(event_log, "structure_processing_outcome", "quality_gate_status"),
        "quality_gate_reasons": _extract_event_context_list(event_log, "structure_processing_outcome", "quality_gate_reasons"),
        "structure_ai_attempted": _extract_event_context_bool(event_log, "structure_processing_outcome", "structure_ai_attempted"),
        "ai_first_degraded": _extract_event_context_bool(event_log, "structure_processing_outcome", "ai_first_degraded"),
        "fallback_stage": _extract_event_context_value(event_log, "structure_processing_outcome", "fallback_stage"),
        "fallback_reason": _extract_event_context_value(event_log, "structure_processing_outcome", "fallback_reason"),
        "document_map_status": document_map_status or "not_requested",
        "document_map_status_reason": document_map_status_reason,
        "ai_classified_count": _extract_event_context_int(event_log, "structure_processing_outcome", "ai_classified_count"),
        "ai_heading_count": _extract_event_context_int(event_log, "structure_processing_outcome", "ai_heading_count"),
        "structure_window_split_count": _extract_event_context_int(
            event_log,
            "structure_processing_outcome",
            "structure_window_split_count",
        ),
        "structure_max_fallback_depth": _extract_event_context_int(
            event_log,
            "structure_processing_outcome",
            "structure_max_fallback_depth",
        ),
        "structure_split_fallback_descriptor_count": _extract_event_context_int(
            event_log,
            "structure_processing_outcome",
            "structure_split_fallback_descriptor_count",
        ),
        "structure_timeout_retry_count": _extract_event_context_int(
            event_log,
            "structure_processing_outcome",
            "structure_timeout_retry_count",
        ),
        "structure_timeout_retry_succeeded_count": _extract_event_context_int(
            event_log,
            "structure_processing_outcome",
            "structure_timeout_retry_succeeded_count",
        ),
        "structure_timeout_retry_failed_count": _extract_event_context_int(
            event_log,
            "structure_processing_outcome",
            "structure_timeout_retry_failed_count",
        ),
        "structure_split_fallback_capped_descriptor_count": _extract_event_context_int(
            event_log,
            "structure_processing_outcome",
            "structure_split_fallback_capped_descriptor_count",
        ),
        "semantic_block_count": 0,
        "first_block_target_chars": 0,
        "first_block_has_toc": False,
        "first_block_has_epigraph": False,
        "first_block_has_body_start": False,
        "first_block_has_isolated_marker": False,
    }
    return snapshot


def _build_passthrough_job(job: Mapping[str, object]) -> dict[str, object]:
    cloned = dict(job)
    cloned["job_kind"] = "passthrough"
    return cloned


def _build_validation_processing_service(event_log: list[dict[str, object]]):
    def _run_document_processing_impl(**kwargs: Any) -> str:
        from docxaicorrector.pipeline._pipeline import run_document_processing

        return cast(str, run_document_processing(**kwargs))

    return clone_processing_service(
        get_client_fn=get_client,
        get_provider_client_fn=get_provider_client,
        get_client_for_model_selector_fn=get_client_for_model_selector,
        resolve_model_selector_fn=resolve_model_selector,
        load_system_prompt_fn=lambda: "",
        ensure_pandoc_available_fn=ensure_pandoc_available,
        generate_markdown_block_fn=lambda **kwargs: cast(str, kwargs.get("target_text") or ""),
        convert_markdown_to_docx_bytes_fn=convert_markdown_to_docx_bytes,
        process_document_images_impl_fn=lambda **kwargs: list(kwargs.get("image_assets") or []),
        analyze_image_fn=lambda *args, **kwargs: None,
        generate_image_candidate_fn=lambda *args, **kwargs: b"",
        validate_redraw_result_fn=lambda *args, **kwargs: None,
        detect_image_mime_type_fn=lambda *args, **kwargs: None,
        inspect_placeholder_integrity_fn=_inspect_placeholder_integrity_adapter,
        preserve_source_paragraph_properties_fn=_preserve_source_paragraph_properties_adapter,
        reinsert_inline_images_fn=_reinsert_inline_images_adapter,
        run_document_processing_impl_fn=_run_document_processing_impl,
        present_error_fn=lambda code, exc, title, **context: f"{title}: {exc}",
        log_event_fn=build_validation_event_logger(event_log),
        emit_state_fn=_emit_state,
        emit_finalize_fn=_emit_finalize,
        emit_activity_fn=_emit_activity,
        emit_log_fn=_emit_log,
        emit_status_fn=_emit_status,
        emit_image_log_fn=lambda runtime, **payload: None,
        emit_image_reset_fn=lambda runtime: None,
        should_stop_processing_fn=lambda runtime: False,
        resolve_uploaded_filename_fn=processing_runtime.resolve_uploaded_filename,
        image_model_call_budget_cls=object,
        image_model_call_budget_exceeded_cls=RuntimeError,
    )


def _extract_run_formatting_diagnostics_paths(
    event_log: Sequence[Mapping[str, object]],
) -> list[str]:
    for event in reversed(event_log):
        if str(event.get("event_id") or "") != "formatting_diagnostics_artifacts_detected":
            continue
        context = event.get("context")
        if not isinstance(context, Mapping):
            continue
        artifact_paths = context.get("artifact_paths")
        if not isinstance(artifact_paths, Sequence) or isinstance(
            artifact_paths, (str, bytes, bytearray)
        ):
            continue
        return [str(path) for path in artifact_paths if isinstance(path, str) and path]
    return []


def _load_formatting_diagnostics_payloads(artifact_paths: Sequence[str]) -> list[dict[str, object]]:
    payloads: list[dict[str, object]] = []
    for artifact_path in artifact_paths:
        try:
            payload = json.loads(Path(artifact_path).read_text(encoding="utf-8"))
        except Exception:
            continue
        if isinstance(payload, dict):
            payloads.append(payload)
    return payloads


def _select_canonical_formatting_diagnostics_payload(
    payloads: Sequence[Mapping[str, object]],
) -> Mapping[str, object] | None:
    if not payloads:
        return None
    return payloads[-1]


def _max_payload_length(payloads: Sequence[Mapping[str, object]], key: str) -> int:
    maximum = 0
    for payload in payloads:
        values = payload.get(key) or []
        if isinstance(values, Sequence) and not isinstance(values, (str, bytes, bytearray)):
            maximum = max(maximum, len(values))
    return maximum


def _count_payload_items(payloads: Sequence[Mapping[str, object]], key: str) -> int:
    total = 0
    for payload in payloads:
        values = payload.get(key) or []
        if isinstance(values, Sequence) and not isinstance(values, (str, bytes, bytearray)):
            total += len(values)
    return total


def _max_accepted_merged_sources(payloads: Sequence[Mapping[str, object]]) -> int:
    maximum = 0
    for payload in payloads:
        explicit_value = payload.get("max_accepted_merged_sources")
        if explicit_value is not None:
            try:
                maximum = max(maximum, int(cast(Any, explicit_value)))
                continue
            except (TypeError, ValueError):
                pass

        accepted_sources = payload.get("accepted_merged_sources") or []
        if not isinstance(accepted_sources, Sequence) or isinstance(accepted_sources, (str, bytes, bytearray)):
            continue
        for entry in accepted_sources:
            if not isinstance(entry, Mapping):
                continue
            explicit_count = entry.get("accepted_merged_sources_count")
            if explicit_count is not None:
                try:
                    maximum = max(maximum, int(explicit_count))
                    continue
                except (TypeError, ValueError):
                    pass
            raw_indexes = entry.get("origin_raw_indexes") or []
            if isinstance(raw_indexes, Sequence) and not isinstance(raw_indexes, (str, bytes, bytearray)):
                maximum = max(maximum, len(raw_indexes))
    return maximum


def _build_output_artifacts(docx_bytes: bytes, markdown_text: str) -> dict[str, object]:
    openable = False
    output_paragraphs = 0
    output_inline_shapes = 0
    output_visible_text_chars = 0
    output_contains_placeholder_markup = False
    if docx_bytes:
        try:
            document = Document(BytesIO(docx_bytes))
        except Exception:
            document = None
        if document is not None:
            openable = True
            output_paragraphs = len(document.paragraphs)
            output_inline_shapes = len(document.inline_shapes)
            output_visible_text_chars = len("\n".join(paragraph.text for paragraph in document.paragraphs))
            output_contains_placeholder_markup = "[[DOCX_IMAGE_" in document._element.xml
    return {
        "output_docx_openable": openable,
        "output_paragraphs": output_paragraphs,
        "output_inline_shapes": output_inline_shapes,
        "output_visible_text_chars": output_visible_text_chars,
        "output_contains_placeholder_markup": output_contains_placeholder_markup,
        "markdown_chars": len(markdown_text),
    }


def _build_runtime_capture() -> dict[str, object]:
    return {"state": {}, "finalize": [], "activity": [], "log": [], "status": []}


def _runtime_mapping(runtime: object) -> dict[str, object]:
    return cast(dict[str, object], runtime)


def _runtime_state(runtime: object) -> dict[str, object]:
    return cast(dict[str, object], _runtime_mapping(runtime).setdefault("state", {}))


def _emit_state(runtime: object, **values: object) -> None:
    _runtime_state(runtime).update(values)


def _emit_finalize(
    runtime: object,
    stage: str,
    detail: str,
    progress: float,
    terminal_kind: str | None = None,
) -> None:
    cast(list[object], _runtime_mapping(runtime).setdefault("finalize", [])).append(
        (stage, detail, progress, terminal_kind)
    )


def _emit_activity(runtime: object, message: str) -> None:
    cast(list[object], _runtime_mapping(runtime).setdefault("activity", [])).append(message)


def _emit_log(runtime: object, **payload: object) -> None:
    cast(list[object], _runtime_mapping(runtime).setdefault("log", [])).append(payload)


def _emit_status(runtime: object, **payload: object) -> None:
    cast(list[object], _runtime_mapping(runtime).setdefault("status", [])).append(payload)


def _inspect_placeholder_integrity_adapter(markdown_text: str, image_assets: Sequence[object]) -> Mapping[str, str]:
    return inspect_placeholder_integrity(markdown_text, cast(list[Any], list(image_assets)))


def _parse_cli_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Print a structural preparation diagnostic snapshot for a validation profile.")
    parser.add_argument("document_profile_id", help="Validation document profile id, for example lietaer-pdf-full-benchmark.")
    parser.add_argument("--run-profile-id", dest="run_profile_id", help="Optional run profile override.")
    return parser.parse_args(list(argv) if argv is not None else None)


def main(argv: Sequence[str] | None = None) -> int:
    args = _parse_cli_args(argv)
    registry = load_validation_registry()
    document_profile = registry.get_document_profile(str(args.document_profile_id))
    run_profile_id = str(args.run_profile_id or getattr(document_profile, "structural_run_profile", "") or "").strip()
    run_profile = registry.get_run_profile(run_profile_id) if run_profile_id else registry.get_run_profile("ui-parity-translate-benchmark-advisory")
    payload = evaluate_structural_preparation_diagnostic(document_profile, run_profile)
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 0


def _preserve_source_paragraph_properties_adapter(
    docx_bytes: bytes,
    paragraphs: Sequence[object],
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
) -> bytes:
    return preserve_source_paragraph_properties(
        docx_bytes,
        cast(list[Any], list(paragraphs)),
        generated_paragraph_registry=generated_paragraph_registry,
    )


def _reinsert_inline_images_adapter(docx_bytes: bytes, image_assets: Sequence[object]) -> bytes:
    return reinsert_inline_images(docx_bytes, cast(list[Any], list(image_assets)))


if __name__ == "__main__":
    raise SystemExit(main())
