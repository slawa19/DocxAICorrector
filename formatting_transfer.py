"""Paragraph mapping and minimal DOCX formatting restoration.

Handles source-to-target paragraph alignment, conservative diagnostics,
minimal caption/image/table normalization, split-heading normalization,
and list numbering restoration.
"""

import logging
import re
from difflib import SequenceMatcher
from io import BytesIO
from typing import Mapping, Sequence, cast

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from document import (
    HEADING_STYLE_PATTERN,
    IMAGE_PLACEHOLDER_PATTERN,
    INLINE_HTML_TAG_PATTERN,
    MARKDOWN_LINK_PATTERN,
    build_paragraph_relations,
    detect_explicit_list_kind,
    find_child_element,
    get_xml_attribute,
    infer_heuristic_heading_level,
    is_image_only_text,
    is_likely_caption_text,
    resolve_effective_relation_kinds,
    resolve_paragraph_outline_level,
    xml_local_name,
)
from logger import log_event
from models import ParagraphUnit
from formatting_diagnostics_retention import get_formatting_diagnostics_dir, write_formatting_diagnostics_artifact

FORMATTING_DIAGNOSTICS_DIR = get_formatting_diagnostics_dir()
MARKDOWN_HEADING_LINE_PATTERN = re.compile(r"^#{1,6}\s+\S")

# Spec TOC/minimal-formatting 2026-04-21: centered direct alignment is allowed
# only for narrow non-heading cases, with an explicit short-paragraph heuristic.
CENTER_SHORT_NON_HEADING_MAX_CHARS = 90
CENTER_SHORT_NON_HEADING_MAX_WORDS = 12
ALLOWED_CENTERED_QUOTE_STRUCTURAL_ROLES = {"epigraph", "attribution", "dedication"}
DISALLOWED_CENTER_SHORT_STRUCTURAL_ROLES = {"toc_header", "toc_entry", "heading", "caption"}


def _paragraph_preview(text: str, *, limit: int = 120) -> str:
    normalized = re.sub(r"\s+", " ", text).strip()
    if len(normalized) <= limit:
        return normalized
    return normalized[: limit - 1].rstrip() + "…"


def _normalize_text_for_mapping(text: str) -> str:
    normalized = text.strip()
    normalized = re.sub(r"^#{1,6}\s+", "", normalized)
    normalized = MARKDOWN_LINK_PATTERN.sub(r"\1", normalized)
    normalized = INLINE_HTML_TAG_PATTERN.sub("", normalized)
    normalized = normalized.replace("***", "").replace("**", "").replace("*", "")
    normalized = normalized.replace("<br/>", " ")
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip().lower()


def _build_generated_registry_candidates(source_paragraph: ParagraphUnit, generated_text: str) -> list[str]:
    candidates: list[str] = []

    def add_candidate(text: str) -> None:
        normalized = _normalize_text_for_mapping(text)
        if normalized and normalized not in candidates:
            candidates.append(normalized)

    add_candidate(generated_text)

    lines = [line.strip() for line in generated_text.splitlines() if line.strip()]
    if not lines:
        return candidates

    non_heading_lines = [line for line in lines if not MARKDOWN_HEADING_LINE_PATTERN.match(line)]
    if source_paragraph.role == "body" and non_heading_lines:
        add_candidate(" ".join(non_heading_lines))
        for line in non_heading_lines:
            add_candidate(line)
        return candidates

    for line in lines:
        add_candidate(line)
    return candidates


def _collect_target_paragraphs(document) -> list[Paragraph]:
    return [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() or IMAGE_PLACEHOLDER_PATTERN.search(paragraph.text)
    ]


def _build_source_registry_entry(
    paragraph: ParagraphUnit,
    fallback_index: int,
    *,
    mapped_target_index: int | None,
    strategy: str | None,
    relation_ids: Sequence[str],
) -> dict[str, object]:
    source_index = paragraph.source_index if paragraph.source_index >= 0 else fallback_index
    return {
        "paragraph_id": paragraph.paragraph_id or f"p{source_index:04d}",
        "source_index": source_index,
        "role": paragraph.role,
        "structural_role": paragraph.structural_role,
        "role_confidence": paragraph.role_confidence,
        "asset_id": paragraph.asset_id,
        "attached_to_asset_id": paragraph.attached_to_asset_id,
        "heading_level": paragraph.heading_level,
        "list_kind": paragraph.list_kind,
        "list_level": paragraph.list_level,
        "list_numbering_format": paragraph.list_numbering_format,
        "list_num_id": paragraph.list_num_id,
        "list_abstract_num_id": paragraph.list_abstract_num_id,
        "origin_raw_indexes": list(paragraph.origin_raw_indexes),
        "origin_raw_text_count": len(paragraph.origin_raw_texts),
        "boundary_source": paragraph.boundary_source,
        "boundary_confidence": paragraph.boundary_confidence,
        "boundary_rationale": paragraph.boundary_rationale,
        "relation_ids": list(relation_ids),
        "mapped_target_index": mapped_target_index,
        "mapping_strategy": strategy,
        "text_preview": _paragraph_preview(_normalize_text_for_mapping(paragraph.text) or paragraph.text),
    }


def _build_target_registry_entry(paragraph, target_index: int, *, mapped: bool) -> dict[str, object]:
    style = getattr(paragraph, "style", None)
    style_name = getattr(style, "name", None)
    return {
        "target_index": target_index,
        "mapped": mapped,
        "style_name": style_name,
        "heading_level": _extract_target_heading_level(paragraph),
        "text_preview": _paragraph_preview(_normalize_text_for_mapping(paragraph.text) or paragraph.text),
    }


def _extract_target_heading_level(paragraph) -> int | None:
    style = getattr(paragraph, "style", None)
    style_name = getattr(style, "name", "") or ""
    normalized_style = style_name.strip().lower()
    style_match = HEADING_STYLE_PATTERN.match(normalized_style)
    if style_match is not None:
        level_text = style_match.group(1)
        if level_text:
            try:
                return max(1, min(int(level_text), 6))
            except ValueError:
                return 1
        return 1

    outline_level = resolve_paragraph_outline_level(paragraph)
    if outline_level is not None:
        return outline_level
    return None


def _build_caption_heading_conflicts(
    source_paragraphs: list[ParagraphUnit],
    target_paragraphs: list[Paragraph],
    mapped_target_by_source: Mapping[int, int],
    strategy_by_source: Mapping[int, str],
) -> list[dict[str, object]]:
    conflicts: list[dict[str, object]] = []
    for source_index, target_index in sorted(mapped_target_by_source.items()):
        source_paragraph = source_paragraphs[source_index]
        if source_paragraph.role != "caption":
            continue

        target_paragraph = target_paragraphs[target_index]
        target_heading_level = _extract_target_heading_level(target_paragraph)
        if target_heading_level is None:
            continue

        conflicts.append(
            {
                "paragraph_id": source_paragraph.paragraph_id or f"p{source_index:04d}",
                "source_index": source_paragraph.source_index if source_paragraph.source_index >= 0 else source_index,
                "mapped_target_index": target_index,
                "mapping_strategy": strategy_by_source.get(source_index),
                "source_role": source_paragraph.role,
                "source_role_confidence": source_paragraph.role_confidence,
                "attached_to_asset_id": source_paragraph.attached_to_asset_id,
                "target_style_name": getattr(getattr(target_paragraph, "style", None), "name", None),
                "target_heading_level": target_heading_level,
                "source_text_preview": _paragraph_preview(source_paragraph.text),
                "target_text_preview": _paragraph_preview(target_paragraph.text),
            }
        )
    return conflicts


def _mapping_similarity_score(source_paragraph: ParagraphUnit, target_text: str) -> float:
    source_text = _normalize_text_for_mapping(source_paragraph.text)
    normalized_target = _normalize_text_for_mapping(target_text)
    if not source_text or not normalized_target:
        return 0.0

    score = SequenceMatcher(None, source_text, normalized_target).ratio()
    if source_paragraph.role == "caption" and is_likely_caption_text(target_text):
        score += 0.08
    if source_paragraph.role == "list":
        target_list_kind = detect_explicit_list_kind(target_text)
        if target_list_kind is not None and target_list_kind == source_paragraph.list_kind:
            score += 0.05
    if source_paragraph.role == "heading" and len(target_text.split()) <= 18:
        score += 0.03
    return min(score, 1.0)


def _register_mapping(
    source_index: int,
    target_index: int,
    strategy: str,
    *,
    mapped_target_by_source: dict[int, int],
    strategy_by_source: dict[int, str],
    available_target_indexes: set[int],
) -> None:
    mapped_target_by_source[source_index] = target_index
    strategy_by_source[source_index] = strategy
    available_target_indexes.discard(target_index)


def _build_generated_registry_by_paragraph_id(
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None,
) -> dict[str, str]:
    registry_by_id: dict[str, str] = {}
    if not generated_paragraph_registry:
        return registry_by_id

    for entry in generated_paragraph_registry:
        paragraph_id = entry.get("paragraph_id")
        text = entry.get("text")
        if isinstance(paragraph_id, str) and paragraph_id and isinstance(text, str) and text.strip():
            registry_by_id[paragraph_id] = text
    return registry_by_id


def _collect_accepted_split_targets(
    source_paragraphs: list[ParagraphUnit],
    target_paragraphs: list[Paragraph],
    mapped_target_by_source: Mapping[int, int],
    generated_registry_by_id: Mapping[str, str],
) -> list[dict[str, object]]:
    accepted_targets: list[dict[str, object]] = []
    accepted_target_indexes: set[int] = set()

    for source_index, target_index in sorted(mapped_target_by_source.items()):
        if target_index <= 0:
            continue

        source_paragraph = source_paragraphs[source_index]
        paragraph_id = source_paragraph.paragraph_id
        if not paragraph_id:
            continue

        generated_text = generated_registry_by_id.get(paragraph_id)
        if not generated_text:
            continue

        generated_lines = [line.strip() for line in generated_text.splitlines() if line.strip()]
        if len(generated_lines) < 2:
            continue

        heading_line = generated_lines[0]
        if not MARKDOWN_HEADING_LINE_PATTERN.match(heading_line):
            continue

        body_lines = [line for line in generated_lines[1:] if not MARKDOWN_HEADING_LINE_PATTERN.match(line)]
        if not body_lines:
            continue

        candidate_target_index = target_index - 1
        if candidate_target_index in accepted_target_indexes:
            continue

        heading_target = target_paragraphs[candidate_target_index]
        if _extract_target_heading_level(heading_target) is None:
            continue

        body_target = target_paragraphs[target_index]
        normalized_heading_line = _normalize_text_for_mapping(heading_line)
        normalized_heading_target = _normalize_text_for_mapping(heading_target.text)
        normalized_body_text = _normalize_text_for_mapping(" ".join(body_lines))
        normalized_body_target = _normalize_text_for_mapping(body_target.text)

        if not normalized_heading_line or normalized_heading_line != normalized_heading_target:
            continue
        if not normalized_body_text or normalized_body_text != normalized_body_target:
            continue

        accepted_target_indexes.add(candidate_target_index)
        accepted_targets.append(
            {
                "target_index": candidate_target_index,
                "derived_from_source_index": source_index,
                "kind": "split_heading_prefix",
                "heading_level": _extract_target_heading_level(heading_target),
                "target_text_preview": _paragraph_preview(heading_target.text),
                "source_text_preview": _paragraph_preview(source_paragraph.text),
            }
        )

    return accepted_targets


def _collect_accepted_merged_sources(
    source_paragraphs: list[ParagraphUnit],
    target_paragraphs: list[Paragraph],
    mapped_target_by_source: Mapping[int, int],
) -> list[dict[str, object]]:
    accepted_sources: list[dict[str, object]] = []
    for source_index, target_index in sorted(mapped_target_by_source.items()):
        source_paragraph = source_paragraphs[source_index]
        if len(source_paragraph.origin_raw_indexes) <= 1:
            continue
        accepted_sources.append(
            {
                "logical_paragraph_id": source_paragraph.paragraph_id or f"p{source_index:04d}",
                "origin_raw_indexes": list(source_paragraph.origin_raw_indexes),
                "accepted_merged_sources_count": len(source_paragraph.origin_raw_indexes),
                "dominant_raw_index": source_paragraph.origin_raw_indexes[0],
                "kind": source_paragraph.boundary_source or "false_paragraph_boundary_merge",
                "boundary_confidence": source_paragraph.boundary_confidence,
                "boundary_decision_class": "medium_accepted" if source_paragraph.boundary_confidence == "medium" else "high",
                "boundary_rationale": source_paragraph.boundary_rationale,
                "target_index": target_index,
                "target_text_preview": _paragraph_preview(target_paragraphs[target_index].text),
                "source_text_preview": _paragraph_preview(source_paragraph.text),
            }
        )
    return accepted_sources


def _map_source_target_paragraphs(
    source_paragraphs: list[ParagraphUnit],
    target_paragraphs: list[Paragraph],
    *,
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
):
    mapped_target_by_source: dict[int, int] = {}
    strategy_by_source: dict[int, str] = {}
    generated_registry_by_id = _build_generated_registry_by_paragraph_id(generated_paragraph_registry)

    available_target_indexes = set(range(len(target_paragraphs)))
    target_indexes_by_exact_text: dict[str, list[int]] = {}
    target_indexes_by_normalized_text: dict[str, list[int]] = {}

    for target_index, target_paragraph in enumerate(target_paragraphs):
        exact_text = target_paragraph.text.strip()
        target_indexes_by_exact_text.setdefault(exact_text, []).append(target_index)
        normalized_text = _normalize_text_for_mapping(target_paragraph.text)
        if normalized_text:
            target_indexes_by_normalized_text.setdefault(normalized_text, []).append(target_index)

    for source_index, source_paragraph in enumerate(source_paragraphs):
        paragraph_id = source_paragraph.paragraph_id
        if not paragraph_id or source_index in mapped_target_by_source:
            continue
        generated_text = generated_registry_by_id.get(paragraph_id)
        if not generated_text:
            continue
        matching_target_indexes: set[int] = set()
        for normalized_generated_text in _build_generated_registry_candidates(source_paragraph, generated_text):
            matching_target_indexes.update(
                target_index
                for target_index in target_indexes_by_normalized_text.get(normalized_generated_text, [])
                if target_index in available_target_indexes
            )
        if len(matching_target_indexes) == 1:
            _register_mapping(
                source_index,
                next(iter(matching_target_indexes)),
                "paragraph_id_registry",
                mapped_target_by_source=mapped_target_by_source,
                strategy_by_source=strategy_by_source,
                available_target_indexes=available_target_indexes,
            )

    for source_index, source_paragraph in enumerate(source_paragraphs):
        paragraph_id = source_paragraph.paragraph_id
        if not paragraph_id or source_index in mapped_target_by_source or source_paragraph.role == "image":
            continue
        generated_text = generated_registry_by_id.get(paragraph_id)
        if not generated_text:
            continue

        scored_candidates: list[tuple[float, int]] = []
        registry_candidates = _build_generated_registry_candidates(source_paragraph, generated_text)
        if not registry_candidates:
            continue

        for target_index in sorted(available_target_indexes):
            if abs(target_index - source_index) > 3:
                continue
            score = max(
                SequenceMatcher(None, candidate_text, _normalize_text_for_mapping(target_paragraphs[target_index].text)).ratio()
                for candidate_text in registry_candidates
            )
            if score >= 0.75:
                scored_candidates.append((score, target_index))

        if not scored_candidates:
            continue

        scored_candidates.sort(reverse=True)
        best_score, best_target_index = scored_candidates[0]
        next_best_score = scored_candidates[1][0] if len(scored_candidates) > 1 else 0.0
        if best_score - next_best_score < 0.05:
            continue

        _register_mapping(
            source_index,
            best_target_index,
            "paragraph_id_registry_similarity",
            mapped_target_by_source=mapped_target_by_source,
            strategy_by_source=strategy_by_source,
            available_target_indexes=available_target_indexes,
        )

    for source_index, source_paragraph in enumerate(source_paragraphs):
        if source_index in mapped_target_by_source or source_paragraph.role != "image":
            continue
        candidates = [
            target_index
            for target_index in target_indexes_by_exact_text.get(source_paragraph.text.strip(), [])
            if target_index in available_target_indexes
        ]
        if len(candidates) == 1:
            _register_mapping(
                source_index,
                candidates[0],
                "image_anchor",
                mapped_target_by_source=mapped_target_by_source,
                strategy_by_source=strategy_by_source,
                available_target_indexes=available_target_indexes,
            )

    for source_index, source_paragraph in enumerate(source_paragraphs):
        if source_index == 0 or source_index in mapped_target_by_source or source_paragraph.role != "caption":
            continue
        previous_paragraph = source_paragraphs[source_index - 1]
        previous_target_index = mapped_target_by_source.get(source_index - 1)
        if previous_paragraph.role not in {"image", "table"} or previous_target_index is None:
            continue
        candidate_index = previous_target_index + 1
        if candidate_index not in available_target_indexes or candidate_index >= len(target_paragraphs):
            continue
        candidate_text = target_paragraphs[candidate_index].text.strip()
        if candidate_text and (
            is_likely_caption_text(candidate_text)
            or _normalize_text_for_mapping(source_paragraph.text) == _normalize_text_for_mapping(candidate_text)
        ):
            _register_mapping(
                source_index,
                candidate_index,
                "adjacent_caption",
                mapped_target_by_source=mapped_target_by_source,
                strategy_by_source=strategy_by_source,
                available_target_indexes=available_target_indexes,
            )

    for source_index, source_paragraph in enumerate(source_paragraphs):
        if source_index in mapped_target_by_source:
            continue
        normalized_text = _normalize_text_for_mapping(source_paragraph.text)
        if not normalized_text:
            continue
        candidates = [
            target_index
            for target_index in target_indexes_by_normalized_text.get(normalized_text, [])
            if target_index in available_target_indexes
        ]
        if len(candidates) == 1:
            _register_mapping(
                source_index,
                candidates[0],
                "exact_text",
                mapped_target_by_source=mapped_target_by_source,
                strategy_by_source=strategy_by_source,
                available_target_indexes=available_target_indexes,
            )

    for source_index, source_paragraph in enumerate(source_paragraphs):
        if source_index in mapped_target_by_source or source_paragraph.role == "image":
            continue

        scored_candidates: list[tuple[float, int]] = []
        for target_index in sorted(available_target_indexes):
            score = _mapping_similarity_score(source_paragraph, target_paragraphs[target_index].text)
            if score >= 0.9:
                scored_candidates.append((score, target_index))

        if not scored_candidates:
            continue

        scored_candidates.sort(reverse=True)
        best_score, best_target_index = scored_candidates[0]
        next_best_score = scored_candidates[1][0] if len(scored_candidates) > 1 else 0.0
        if best_score - next_best_score < 0.05:
            continue

        _register_mapping(
            source_index,
            best_target_index,
            "similarity",
            mapped_target_by_source=mapped_target_by_source,
            strategy_by_source=strategy_by_source,
            available_target_indexes=available_target_indexes,
        )

    if len(source_paragraphs) == len(target_paragraphs):
        for source_index, source_paragraph in enumerate(source_paragraphs):
            if source_index in mapped_target_by_source or source_index not in available_target_indexes:
                continue

            target_paragraph = target_paragraphs[source_index]
            source_exact = source_paragraph.text.strip()
            target_exact = target_paragraph.text.strip()
            source_normalized = _normalize_text_for_mapping(source_paragraph.text)
            target_normalized = _normalize_text_for_mapping(target_paragraph.text)

            if (source_exact and source_exact == target_exact) or (
                source_normalized and source_normalized == target_normalized
            ):
                _register_mapping(
                    source_index,
                    source_index,
                    "positional_exact",
                    mapped_target_by_source=mapped_target_by_source,
                    strategy_by_source=strategy_by_source,
                    available_target_indexes=available_target_indexes,
                )

    mapping_pairs = [
        (source_paragraphs[source_index], target_paragraphs[target_index])
        for source_index, target_index in sorted(mapped_target_by_source.items())
    ]

    accepted_split_targets = _collect_accepted_split_targets(
        source_paragraphs,
        target_paragraphs,
        mapped_target_by_source,
        generated_registry_by_id,
    )
    accepted_merged_sources = _collect_accepted_merged_sources(
        source_paragraphs,
        target_paragraphs,
        mapped_target_by_source,
    )
    accepted_relations, relation_report = build_paragraph_relations(
        source_paragraphs,
        enabled_relation_kinds=resolve_effective_relation_kinds(),
    )
    relation_ids_by_paragraph: dict[str, list[str]] = {}
    for relation in accepted_relations:
        for paragraph_id in relation.member_paragraph_ids:
            relation_ids_by_paragraph.setdefault(paragraph_id, []).append(relation.relation_id)
    accepted_split_target_indexes = {entry["target_index"] for entry in accepted_split_targets}

    diagnostics = {
        "source_count": len(source_paragraphs),
        "target_count": len(target_paragraphs),
        "mapped_count": len(mapping_pairs),
        "unmapped_source_ids": [
            source_paragraphs[source_index].paragraph_id or f"p{source_index:04d}"
            for source_index in range(len(source_paragraphs))
            if source_index not in mapped_target_by_source
        ],
        "unmapped_target_indexes": [
            target_index
            for target_index in range(len(target_paragraphs))
            if target_index not in mapped_target_by_source.values() and target_index not in accepted_split_target_indexes
        ],
        "source_registry": [
            _build_source_registry_entry(
                paragraph,
                index,
                mapped_target_index=mapped_target_by_source.get(index),
                strategy=strategy_by_source.get(index),
                relation_ids=relation_ids_by_paragraph.get(paragraph.paragraph_id or f"p{index:04d}", []),
            )
            for index, paragraph in enumerate(source_paragraphs)
        ],
        "target_registry": [
            _build_target_registry_entry(
                paragraph,
                index,
                mapped=index in mapped_target_by_source.values() or index in accepted_split_target_indexes,
            )
            for index, paragraph in enumerate(target_paragraphs)
        ],
        "accepted_split_targets": accepted_split_targets,
        "accepted_merged_sources": accepted_merged_sources,
        "accepted_merged_sources_count": len(accepted_merged_sources),
        "max_accepted_merged_sources": max(
            (int(cast(int, entry.get("accepted_merged_sources_count", 0)) or 0) for entry in accepted_merged_sources),
            default=0,
        ),
        "high_confidence_merge_count": sum(
            1 for entry in accepted_merged_sources if entry.get("boundary_decision_class") == "high"
        ),
        "medium_accepted_merge_count": sum(
            1 for entry in accepted_merged_sources if entry.get("boundary_decision_class") == "medium_accepted"
        ),
        "accepted_relations": [
            {
                "relation_id": relation.relation_id,
                "relation_kind": relation.relation_kind,
                "member_paragraph_ids": list(relation.member_paragraph_ids),
                "anchor_asset_id": relation.anchor_asset_id,
                "confidence": relation.confidence,
                "rationale": list(relation.rationale),
            }
            for relation in accepted_relations
        ],
        "relation_decisions": [
            {
                "relation_kind": decision.relation_kind,
                "decision": decision.decision,
                "member_paragraph_ids": list(decision.member_paragraph_ids),
                "anchor_asset_id": decision.anchor_asset_id,
                "reasons": list(decision.reasons),
            }
            for decision in relation_report.decisions
        ],
        "relation_count": relation_report.total_relations,
        "relation_counts": dict(relation_report.relation_counts),
        "rejected_relation_candidate_count": relation_report.rejected_candidate_count,
        "caption_heading_conflicts": _build_caption_heading_conflicts(
            source_paragraphs,
            target_paragraphs,
            mapped_target_by_source,
            strategy_by_source,
        ),
        "list_restoration_decisions": [],
    }
    return mapping_pairs, diagnostics


def _write_formatting_diagnostics_artifact(stage: str, diagnostics: dict[str, object]) -> str | None:
    return write_formatting_diagnostics_artifact(
        stage=stage,
        diagnostics=diagnostics,
        diagnostics_dir=FORMATTING_DIAGNOSTICS_DIR,
    )


# ---------------------------------------------------------------------------
# Formatting preservation and semantic normalization
# ---------------------------------------------------------------------------


def restore_source_formatting(
    docx_bytes: bytes,
    paragraphs: list[ParagraphUnit],
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
) -> bytes:
    return _restore_source_formatting_impl(
        docx_bytes,
        paragraphs,
        generated_paragraph_registry=generated_paragraph_registry,
        mismatch_event_name="paragraph_count_mismatch_restore",
        mismatch_log_message=(
            "Число source/target абзацев не совпадает при unified formatting restore; "
            "применяю только консервативно сопоставленные абзацы."
        ),
    )


def preserve_source_paragraph_properties(
    docx_bytes: bytes,
    paragraphs: list[ParagraphUnit],
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
) -> bytes:
    """Canonical public formatting entry point for the current transition wave."""
    return apply_output_formatting(
        docx_bytes,
        paragraphs,
        generated_paragraph_registry=generated_paragraph_registry,
        mismatch_event_name="paragraph_count_mismatch_preserve",
        mismatch_log_message=(
            "Число source/target абзацев не совпадает при переносе свойств форматирования; "
            "применяю только консервативно сопоставленные абзацы."
        ),
    )


def apply_output_formatting(
    docx_bytes: bytes,
    paragraphs: list[ParagraphUnit],
    *,
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
    mismatch_event_name: str,
    mismatch_log_message: str,
) -> bytes:
    if not docx_bytes or not paragraphs:
        return docx_bytes

    document = Document(BytesIO(docx_bytes))
    target_paragraphs = _collect_target_paragraphs(document)
    relevant_source_paragraphs = [paragraph for paragraph in paragraphs if paragraph.role != "table"]
    mapping_pairs, diagnostics = _map_source_target_paragraphs(
        list(relevant_source_paragraphs),
        list(target_paragraphs),
        generated_paragraph_registry=generated_paragraph_registry,
    )
    unmapped_source_ids = cast(list[str], diagnostics["unmapped_source_ids"])
    unmapped_target_indexes = cast(list[int], diagnostics["unmapped_target_indexes"])

    _apply_minimal_image_formatting(document)
    _apply_minimal_caption_formatting(
        document,
        paragraphs,
        generated_paragraph_registry=generated_paragraph_registry,
    )
    diagnostics["toc_format_restoration_decisions"] = _restore_toc_paragraph_properties_for_mapped_pairs(document, mapping_pairs)
    diagnostics["alignment_restoration_decisions"] = _restore_direct_paragraph_alignment_for_mapped_pairs(mapping_pairs)
    _restore_semantic_quote_formatting_for_mapped_pairs(mapping_pairs)

    mismatch_detected = bool(unmapped_source_ids or unmapped_target_indexes)
    if not mismatch_detected:
        _apply_accepted_split_heading_styles(
            document,
            target_paragraphs,
            cast(list[dict[str, object]], diagnostics.get("accepted_split_targets", [])),
            list(relevant_source_paragraphs),
        )

    # Always restore list numbering for successfully mapped paragraph pairs, even when
    # there are unmapped paragraphs.  The mapping_pairs list already contains only the
    # paragraphs that were matched, so applying list formatting to them is always safe.
    # Skipping this entirely on any mismatch was the root cause of lists disappearing
    # whenever the AI added or removed even one paragraph in its output.
    diagnostics["list_restoration_decisions"] = _restore_list_numbering_for_mapped_paragraphs(document, mapping_pairs)

    # Persisted diagnostics artifacts remain mismatch-only by contract; on the
    # happy path, alignment decisions are still available in runtime logs.
    if mismatch_detected:
        artifact_path = _write_formatting_diagnostics_artifact("restore", diagnostics)
        log_event(
            logging.WARNING,
            mismatch_event_name,
            mismatch_log_message,
            source_count=diagnostics["source_count"],
            target_count=len(target_paragraphs),
            mapped_count=diagnostics["mapped_count"],
            unmapped_source_count=len(unmapped_source_ids),
            unmapped_target_count=len(unmapped_target_indexes),
            artifact_path=artifact_path,
        )

    if _style_exists(document, "Table Grid"):
        for table in document.tables:
            table.style = "Table Grid"

    output_stream = BytesIO()
    document.save(output_stream)
    return output_stream.getvalue()


def _restore_source_formatting_impl(
    docx_bytes: bytes,
    paragraphs: list[ParagraphUnit],
    *,
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
    mismatch_event_name: str,
    mismatch_log_message: str,
) -> bytes:
    return apply_output_formatting(
        docx_bytes,
        paragraphs,
        generated_paragraph_registry=generated_paragraph_registry,
        mismatch_event_name=mismatch_event_name,
        mismatch_log_message=mismatch_log_message,
    )


def _build_output_formatting_diagnostics(
    source_paragraphs: Sequence[ParagraphUnit],
    target_paragraphs: Sequence[Paragraph],
    *,
    document=None,
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
) -> dict[str, object]:
    relevant_source_paragraphs = [paragraph for paragraph in source_paragraphs if paragraph.role != "table"]
    _, diagnostics = _map_source_target_paragraphs(
        list(relevant_source_paragraphs),
        list(target_paragraphs),
        generated_paragraph_registry=generated_paragraph_registry,
    )
    return diagnostics


def _apply_minimal_image_formatting(document) -> None:
    for paragraph in document.paragraphs:
        if is_image_only_text(paragraph.text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _apply_minimal_caption_formatting(
    document,
    source_paragraphs: Sequence[ParagraphUnit],
    *,
    generated_paragraph_registry: Sequence[Mapping[str, object]] | None = None,
) -> None:
    if not document.paragraphs:
        return

    generated_registry_by_id = _build_generated_registry_by_paragraph_id(generated_paragraph_registry)
    source_caption_texts = {
        _normalize_text_for_mapping(paragraph.text)
        for paragraph in source_paragraphs
        if paragraph.role == "caption" and paragraph.text.strip()
    }
    source_caption_texts.discard("")
    generated_caption_texts = {
        _normalize_text_for_mapping(generated_registry_by_id.get(paragraph.paragraph_id or "", ""))
        for paragraph in source_paragraphs
        if paragraph.role == "caption" and paragraph.paragraph_id
    }
    generated_caption_texts.discard("")

    for paragraph in document.paragraphs:
        if not _is_caption_candidate(
            document,
            paragraph,
            source_caption_texts=source_caption_texts,
            generated_caption_texts=generated_caption_texts,
        ):
            continue
        if _style_exists(document, "Caption"):
            paragraph.style = document.styles["Caption"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _is_caption_candidate(
    document,
    paragraph,
    *,
    source_caption_texts: set[str],
    generated_caption_texts: set[str],
) -> bool:
    text = paragraph.text.strip()
    if not text or IMAGE_PLACEHOLDER_PATTERN.search(text):
        return False

    normalized_text = _normalize_text_for_mapping(text)
    if not normalized_text:
        return False

    has_anchor_context = _has_caption_anchor_context(document, paragraph)
    if not has_anchor_context:
        return False

    if normalized_text in source_caption_texts or normalized_text in generated_caption_texts:
        return True
    return is_likely_caption_text(text)


def _has_caption_anchor_context(document, paragraph) -> bool:
    body_children = list(document._element.body.iterchildren())
    paragraph_element = paragraph._element

    for index, child in enumerate(body_children):
        if child != paragraph_element:
            continue
        previous_child = body_children[index - 1] if index > 0 else None
        next_child = body_children[index + 1] if index + 1 < len(body_children) else None
        return _is_caption_anchor_block(previous_child) or _is_caption_anchor_block(next_child)

    return False


def _is_caption_anchor_block(block_element) -> bool:
    if block_element is None:
        return False

    local_name = xml_local_name(block_element.tag)
    if local_name == "tbl":
        return True
    if local_name != "p":
        return False

    text_content = "".join(block_element.itertext())
    return is_image_only_text(text_content)


# ---------------------------------------------------------------------------
# Paragraph-level normalization and list numbering restoration
# ---------------------------------------------------------------------------


def _apply_accepted_split_heading_styles(
    document,
    target_paragraphs: Sequence[Paragraph],
    accepted_split_targets: Sequence[Mapping[str, object]],
    source_paragraphs: Sequence[ParagraphUnit],
) -> None:
    for accepted_target in accepted_split_targets:
        target_index_value = accepted_target.get("target_index")
        source_index_value = accepted_target.get("derived_from_source_index")
        try:
            target_index = int(cast(int | str, target_index_value))
            source_index = int(cast(int | str, source_index_value))
        except (TypeError, ValueError):
            continue
        if target_index < 0 or target_index >= len(target_paragraphs):
            continue
        if source_index < 0 or source_index >= len(source_paragraphs):
            continue

        paragraph = target_paragraphs[target_index]
        source_paragraph = source_paragraphs[source_index]
        normalized_target = _normalize_text_for_mapping(paragraph.text)
        normalized_source = _normalize_text_for_mapping(source_paragraph.text)
        if not normalized_target or not normalized_source.startswith(normalized_target):
            continue
        inferred_level = infer_heuristic_heading_level(paragraph.text)
        heading_style = f"Heading {min(max(inferred_level, 1), 6)}"
        if _style_exists(document, heading_style):
            paragraph.style = document.styles[heading_style]


def _extract_paragraph_num_id(paragraph) -> str | None:
    paragraph_properties = find_child_element(paragraph._element, "pPr")
    num_pr = find_child_element(paragraph_properties, "numPr")
    num_id = find_child_element(num_pr, "numId")
    return get_xml_attribute(num_id, "val") if num_id is not None else None


def _restore_list_numbering_for_mapped_paragraphs(document, mapping_pairs: list[tuple[ParagraphUnit, Paragraph]]) -> list[dict[str, object]]:
    numbering_root = _get_target_numbering_root(document)
    decisions: list[dict[str, object]] = []
    if numbering_root is None:
        for source_paragraph, _ in mapping_pairs:
            if source_paragraph.role != "list":
                continue
            decisions.append(
                {
                    "paragraph_id": source_paragraph.paragraph_id,
                    "text_preview": _paragraph_preview(source_paragraph.text),
                    "action": "target_numbering_part_missing",
                }
            )
        return decisions

    numbering_id_map: dict[tuple[str, str, str, str], tuple[int, int]] = {}
    next_abstract_num_id = _next_numbering_identifier(numbering_root, "abstractNum", "abstractNumId")
    next_num_id = _next_numbering_identifier(numbering_root, "num", "numId")

    for source_paragraph, target_paragraph in mapping_pairs:
        if source_paragraph.role != "list":
            continue
        existing_target_num_id = _extract_paragraph_num_id(target_paragraph)
        if existing_target_num_id is not None:
            decisions.append(
                {
                    "paragraph_id": source_paragraph.paragraph_id,
                    "text_preview": _paragraph_preview(source_paragraph.text),
                    "action": "kept_existing_target_numbering",
                    "list_kind": source_paragraph.list_kind,
                    "list_level": source_paragraph.list_level,
                    "target_num_id": existing_target_num_id,
                }
            )
            continue
        if not source_paragraph.list_num_xml or not source_paragraph.list_abstract_num_xml:
            decisions.append(
                {
                    "paragraph_id": source_paragraph.paragraph_id,
                    "text_preview": _paragraph_preview(source_paragraph.text),
                    "action": "missing_source_numbering_xml",
                    "list_kind": source_paragraph.list_kind,
                    "list_level": source_paragraph.list_level,
                }
            )
            continue

        mapping_key = (
            source_paragraph.list_num_id or "",
            source_paragraph.list_abstract_num_id or "",
            source_paragraph.list_num_xml,
            source_paragraph.list_abstract_num_xml,
        )
        mapped_ids = numbering_id_map.get(mapping_key)
        if mapped_ids is None:
            abstract_num_id = next_abstract_num_id
            num_id = next_num_id
            next_abstract_num_id += 1
            next_num_id += 1
            if not _append_numbering_definition(numbering_root, source_paragraph, abstract_num_id, num_id):
                decisions.append(
                    {
                        "paragraph_id": source_paragraph.paragraph_id,
                        "text_preview": _paragraph_preview(source_paragraph.text),
                        "action": "append_definition_failed",
                        "list_kind": source_paragraph.list_kind,
                        "list_level": source_paragraph.list_level,
                    }
                )
                continue
            mapped_ids = (abstract_num_id, num_id)
            numbering_id_map[mapping_key] = mapped_ids

        _apply_list_numbering_to_paragraph(target_paragraph, list_level=source_paragraph.list_level, num_id=mapped_ids[1])
        decisions.append(
            {
                "paragraph_id": source_paragraph.paragraph_id,
                "text_preview": _paragraph_preview(source_paragraph.text),
                "action": "restored",
                "list_kind": source_paragraph.list_kind,
                "list_level": source_paragraph.list_level,
                "target_num_id": mapped_ids[1],
                "target_abstract_num_id": mapped_ids[0],
            }
        )

    return decisions


def _get_target_numbering_root(document):
    numbering_part = getattr(document.part, "numbering_part", None)
    return getattr(numbering_part, "element", None)


def _next_numbering_identifier(numbering_root, tag_name: str, attribute_name: str) -> int:
    max_identifier = 0
    for child in numbering_root:
        if xml_local_name(child.tag) != tag_name:
            continue
        value = get_xml_attribute(child, attribute_name)
        try:
            max_identifier = max(max_identifier, int(value or "0"))
        except ValueError:
            continue
    return max_identifier + 1


def _append_numbering_definition(numbering_root, source_paragraph: ParagraphUnit, abstract_num_id: int, num_id: int) -> bool:
    try:
        abstract_num_element = parse_xml(source_paragraph.list_abstract_num_xml or "")
        num_element = parse_xml(source_paragraph.list_num_xml or "")
    except Exception:
        return False

    abstract_num_element.set(qn("w:abstractNumId"), str(abstract_num_id))
    num_element.set(qn("w:numId"), str(num_id))

    abstract_num_id_element = find_child_element(num_element, "abstractNumId")
    if abstract_num_id_element is None:
        abstract_num_id_element = OxmlElement("w:abstractNumId")
        num_element.insert(0, abstract_num_id_element)
    abstract_num_id_element.set(qn("w:val"), str(abstract_num_id))

    numbering_root.append(abstract_num_element)
    numbering_root.append(num_element)
    return True


def _apply_list_numbering_to_paragraph(paragraph, *, list_level: int, num_id: int) -> None:
    paragraph_properties = _ensure_paragraph_properties(paragraph)
    existing_num_pr = find_child_element(paragraph_properties, "numPr")
    if existing_num_pr is not None:
        paragraph_properties.remove(existing_num_pr)

    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(max(0, list_level)))
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)
    paragraph_properties.append(num_pr)


def _style_exists(document, style_name: str) -> bool:
    try:
        document.styles[style_name]
    except KeyError:
        return False
    return True


def _replace_paragraph_properties_from_xml(paragraph, paragraph_properties_xml: str) -> bool:
    if not paragraph_properties_xml.strip():
        return False

    paragraph_properties = parse_xml(paragraph_properties_xml)
    existing_properties = find_child_element(paragraph._element, "pPr")
    if existing_properties is not None:
        paragraph._element.remove(existing_properties)
    paragraph._element.insert(0, paragraph_properties)
    return True


def _restore_toc_paragraph_properties_for_mapped_pairs(document, mapping_pairs: list[tuple[ParagraphUnit, Paragraph]]) -> list[dict[str, object]]:
    decisions: list[dict[str, object]] = []
    for source_paragraph, target_paragraph in mapping_pairs:
        structural_role = str(getattr(source_paragraph, "structural_role", "") or "").strip().lower()
        if structural_role not in {"toc_header", "toc_entry"}:
            continue

        decision = {
            "paragraph_id": source_paragraph.paragraph_id,
            "structural_role": structural_role,
            "source_style_name": str(getattr(source_paragraph, "style_name", "") or "").strip() or None,
            "target_style_name": _target_paragraph_style_name(target_paragraph),
            "source_alignment": str(getattr(source_paragraph, "paragraph_alignment", "") or "").strip().lower() or None,
            "source_preview": _paragraph_preview(source_paragraph.text),
        }

        paragraph_properties_xml = str(getattr(source_paragraph, "paragraph_properties_xml", "") or "")
        if not paragraph_properties_xml:
            decisions.append({**decision, "action": "skipped", "reason": "missing_source_paragraph_properties"})
            continue

        restored = _replace_paragraph_properties_from_xml(target_paragraph, paragraph_properties_xml)
        if not restored:
            decisions.append({**decision, "action": "skipped", "reason": "invalid_source_paragraph_properties"})
            continue

        if decision["source_style_name"] and _style_exists(document, cast(str, decision["source_style_name"])):
            target_paragraph.style = document.styles[cast(str, decision["source_style_name"])]
        decisions.append({**decision, "action": "restored", "reason": "copied_source_toc_paragraph_properties"})
    return decisions


def _set_direct_paragraph_alignment(paragraph, alignment_value: str | None) -> None:
    paragraph_properties = _ensure_paragraph_properties(paragraph)
    existing_alignment = find_child_element(paragraph_properties, "jc")
    if alignment_value is None:
        if existing_alignment is not None:
            paragraph_properties.remove(existing_alignment)
        return
    if existing_alignment is None:
        existing_alignment = OxmlElement("w:jc")
        paragraph_properties.append(existing_alignment)
    existing_alignment.set(qn("w:val"), alignment_value)


def _target_paragraph_has_heading_style(paragraph) -> bool:
    style = getattr(paragraph, "style", None)
    style_name = str(getattr(style, "name", "") or "").strip().lower()
    return style_name.startswith("heading ")


def _target_paragraph_style_name(paragraph) -> str | None:
    style = getattr(paragraph, "style", None)
    style_name = str(getattr(style, "name", "") or "").strip()
    return style_name or None


def _is_heading_like_source_paragraph(source_paragraph: ParagraphUnit) -> bool:
    source_role = str(getattr(source_paragraph, "role", "") or "").strip().lower()
    structural_role = str(getattr(source_paragraph, "structural_role", "") or "").strip().lower()
    return (
        source_role == "heading"
        or structural_role == "heading"
        or getattr(source_paragraph, "heading_level", None) is not None
        or bool(getattr(source_paragraph, "heading_source", None))
    )


def _is_allowlisted_centered_quote_paragraph(source_paragraph: ParagraphUnit) -> bool:
    return str(getattr(source_paragraph, "structural_role", "") or "").strip().lower() in ALLOWED_CENTERED_QUOTE_STRUCTURAL_ROLES


def _is_short_non_heading_paragraph(source_paragraph: ParagraphUnit) -> bool:
    role = str(getattr(source_paragraph, "role", "") or "").strip().lower()
    structural_role = str(getattr(source_paragraph, "structural_role", "") or "").strip().lower()
    if role in {"heading", "list", "caption"}:
        return False
    if structural_role in DISALLOWED_CENTER_SHORT_STRUCTURAL_ROLES:
        return False
    if getattr(source_paragraph, "heading_level", None) is not None or getattr(source_paragraph, "heading_source", None):
        return False
    normalized = re.sub(r"\s+", " ", str(getattr(source_paragraph, "text", "") or "")).strip()
    if not normalized:
        return False
    if len(normalized) > CENTER_SHORT_NON_HEADING_MAX_CHARS:
        return False
    words = [token for token in normalized.split(" ") if token]
    return len(words) <= CENTER_SHORT_NON_HEADING_MAX_WORDS


def _resolve_direct_alignment_restoration_decision(
    source_paragraph: ParagraphUnit,
    target_paragraph: Paragraph,
) -> dict[str, object] | None:
    alignment_value = str(source_paragraph.paragraph_alignment or "").strip().lower()
    if not alignment_value:
        return None
    decision = {
        "paragraph_id": source_paragraph.paragraph_id,
        "source_alignment": alignment_value,
        "target_style_name": _target_paragraph_style_name(target_paragraph),
        "target_heading_level": _extract_target_heading_level(target_paragraph),
        "source_role": str(getattr(source_paragraph, "role", "") or "").strip().lower(),
        "source_structural_role": str(getattr(source_paragraph, "structural_role", "") or "").strip().lower(),
        "source_preview": _paragraph_preview(source_paragraph.text),
    }
    if _is_heading_like_source_paragraph(source_paragraph):
        return {**decision, "action": "skipped", "reason": "source_heading_semantics"}
    if _target_paragraph_has_heading_style(target_paragraph):
        return {**decision, "action": "skipped", "reason": "target_heading_style"}
    if alignment_value != "center":
        return {**decision, "action": "skipped", "reason": "unsupported_alignment_value"}
    if str(getattr(source_paragraph, "role", "") or "").strip().lower() == "list":
        return {**decision, "action": "skipped", "reason": "list_item_not_allowlisted"}
    if _is_allowlisted_centered_quote_paragraph(source_paragraph):
        return {**decision, "action": "restored", "reason": "center_quote_allowlisted"}
    if not _is_short_non_heading_paragraph(source_paragraph):
        return {**decision, "action": "skipped", "reason": "not_short_non_heading_paragraph"}
    return {**decision, "action": "restored", "reason": "center_allowlisted"}


def _restore_direct_paragraph_alignment_for_mapped_pairs(mapping_pairs: list[tuple[ParagraphUnit, Paragraph]]) -> list[dict[str, object]]:
    decisions: list[dict[str, object]] = []
    for source_paragraph, target_paragraph in mapping_pairs:
        decision = _resolve_direct_alignment_restoration_decision(source_paragraph, target_paragraph)
        if decision is None:
            continue
        decisions.append(decision)
        if decision["action"] != "restored":
            log_event(
                logging.INFO,
                "alignment_restoration_skipped",
                "Пропущено восстановление direct alignment для абзаца.",
                paragraph_id=decision.get("paragraph_id"),
                source_alignment=decision.get("source_alignment"),
                skip_reason=decision.get("reason"),
                target_style_name=decision.get("target_style_name"),
                target_heading_level=decision.get("target_heading_level"),
                source_preview=decision.get("source_preview"),
            )
            continue
        _set_direct_paragraph_alignment(target_paragraph, source_paragraph.paragraph_alignment)
    return decisions


def _restore_semantic_quote_formatting_for_mapped_pairs(mapping_pairs: list[tuple[ParagraphUnit, Paragraph]]) -> None:
    for source_paragraph, target_paragraph in mapping_pairs:
        if source_paragraph.structural_role not in {"epigraph", "attribution", "dedication"}:
            continue
        if source_paragraph.is_italic:
            for run in target_paragraph.runs:
                if run.text and run.text.strip():
                    run.italic = True


def _ensure_paragraph_properties(paragraph):
    paragraph_properties = find_child_element(paragraph._element, "pPr")
    if paragraph_properties is not None:
        return paragraph_properties

    paragraph_properties = OxmlElement("w:pPr")
    paragraph._element.insert(0, paragraph_properties)
    return paragraph_properties
