from __future__ import annotations

import json
import subprocess
import sys
from io import BytesIO
from pathlib import Path

from docx import Document

from docxaicorrector.validation.formatting_replay import replay_formatting_diagnostics_from_report


def _build_docx_bytes(paragraphs: list[str]) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_replay_formatting_diagnostics_from_report_recomputes_from_saved_final_docx(tmp_path: Path):
    target_docx_path = tmp_path / "target.docx"
    target_docx_path.write_bytes(_build_docx_bytes(["Heading translated", "Body translated"]))

    report_path = tmp_path / "report.json"
    report_payload = {
        "output_artifacts": {"docx_path": str(target_docx_path)},
        "formatting_diagnostics": [
            {
                "stage": "post_cleanup",
                "mapped_count": 0,
                "unmapped_source_ids": ["p0001", "p0002"],
                "unmapped_target_indexes": [0, 1],
                "source_registry": [
                    {
                        "paragraph_id": "p0001",
                        "source_index": 0,
                        "text_preview": "Heading translated",
                        "role": "heading",
                        "structural_role": "heading",
                        "heading_level": 1,
                    },
                    {
                        "paragraph_id": "p0002",
                        "source_index": 1,
                        "text_preview": "Body translated",
                        "role": "body",
                        "structural_role": "body",
                    },
                ],
            }
        ],
    }
    report_path.write_text(json.dumps(report_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    replayed = replay_formatting_diagnostics_from_report(report_path=report_path)

    assert replayed["source_reconstruction_basis"] == "saved_source_registry_preview"
    assert replayed["replay_scope"] == "restore_diagnostics_from_saved_final_docx"
    assert replayed["replay_mode"] == "no_llm_current_mapping_code"
    assert replayed["replay_fidelity"] == "matched_saved_source_count"
    assert replayed["saved_mapped_count"] == 0
    replayed_diagnostics = replayed["replayed_diagnostics"]
    assert replayed_diagnostics["mapped_count"] == 2
    assert replayed["replayed_summary"]["unmapped_source_count"] == 0
    role_aware_summary = replayed["replayed_summary"]["role_aware_summary"]
    assert role_aware_summary is None or role_aware_summary["effective_unmapped_source_count"] == 0


def test_replay_formatting_diagnostics_marks_count_only_fidelity_for_source_language_registry_preview(tmp_path: Path):
    target_docx_path = tmp_path / "target.docx"
    target_docx_path.write_bytes(_build_docx_bytes(["СОДЕРЖАНИЕ", "Предисловие"]))

    report_path = tmp_path / "report.json"
    report_payload = {
        "output_artifacts": {"docx_path": str(target_docx_path)},
        "formatting_diagnostics": [
            {
                "stage": "post_cleanup",
                "mapped_count": 0,
                "unmapped_source_ids": ["p0001", "p0002"],
                "unmapped_target_indexes": [0, 1],
                "source_registry": [
                    {
                        "paragraph_id": "p0001",
                        "source_index": 0,
                        "text_preview": "contents",
                        "role": "heading",
                        "structural_role": "heading",
                        "heading_level": 1,
                    },
                    {
                        "paragraph_id": "p0002",
                        "source_index": 1,
                        "text_preview": "foreword",
                        "role": "body",
                        "structural_role": "body",
                    },
                ],
                "target_registry": [
                    {"target_index": 0, "text_preview": "СОДЕРЖАНИЕ", "heading_level": 1},
                    {"target_index": 1, "text_preview": "Предисловие"},
                ],
            }
        ],
    }
    report_path.write_text(json.dumps(report_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    replayed = replay_formatting_diagnostics_from_report(report_path=report_path)

    assert replayed["source_reconstruction_basis"] == "saved_source_registry_preview"
    assert replayed["saved_source_count"] == 2
    assert replayed["replayed_source_count"] == 2
    assert replayed["replay_fidelity"] == "count_parity_only_source_language_preview"


def test_replay_formatting_diagnostics_from_report_prefers_saved_registry_over_discovered_source_docx(tmp_path: Path):
    target_docx_path = tmp_path / "target.docx"
    target_docx_path.write_bytes(_build_docx_bytes(["Registry heading", "Registry body"]))

    source_docx_path = tmp_path / "source.docx"
    source_docx_path.write_bytes(_build_docx_bytes(["Wrong source only"]))

    report_path = tmp_path / "report.json"
    report_payload = {
        "preparation": {"uploaded_filename": source_docx_path.name},
        "output_artifacts": {"docx_path": str(target_docx_path)},
        "formatting_diagnostics": [
            {
                "stage": "post_cleanup",
                "mapped_count": 0,
                "unmapped_source_ids": ["p0001", "p0002"],
                "unmapped_target_indexes": [0, 1],
                "source_registry": [
                    {
                        "paragraph_id": "p0001",
                        "source_index": 0,
                        "text_preview": "Registry heading",
                        "role": "heading",
                        "structural_role": "heading",
                        "heading_level": 1,
                    },
                    {
                        "paragraph_id": "p0002",
                        "source_index": 1,
                        "text_preview": "Registry body",
                        "role": "body",
                        "structural_role": "body",
                    },
                ],
            }
        ],
    }
    report_path.write_text(json.dumps(report_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    replayed = replay_formatting_diagnostics_from_report(report_path=report_path)

    assert replayed["source_reconstruction_basis"] == "saved_source_registry_preview"
    assert replayed["source_docx_path"] is None
    assert replayed["saved_source_count"] == 2
    assert replayed["replayed_source_count"] == 2
    assert replayed["replay_fidelity"] == "matched_saved_source_count"
    assert replayed["replayed_diagnostics"]["mapped_count"] == 2


def test_replay_formatting_diagnostics_from_report_uses_explicit_source_docx_override_when_requested(tmp_path: Path):
    target_docx_path = tmp_path / "target.docx"
    target_docx_path.write_bytes(_build_docx_bytes(["Override heading"]))

    source_docx_path = tmp_path / "source.docx"
    source_docx_path.write_bytes(_build_docx_bytes(["Override heading"]))

    report_path = tmp_path / "report.json"
    report_payload = {
        "output_artifacts": {"docx_path": str(target_docx_path)},
        "formatting_diagnostics": [
            {
                "stage": "post_cleanup",
                "mapped_count": 0,
                "unmapped_source_ids": ["p0001", "p0002"],
                "unmapped_target_indexes": [0],
                "source_registry": [
                    {
                        "paragraph_id": "p0001",
                        "source_index": 0,
                        "text_preview": "Registry heading",
                        "role": "heading",
                        "structural_role": "heading",
                        "heading_level": 1,
                    },
                    {
                        "paragraph_id": "p0002",
                        "source_index": 1,
                        "text_preview": "Registry body",
                        "role": "body",
                        "structural_role": "body",
                    },
                ],
            }
        ],
    }
    report_path.write_text(json.dumps(report_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    replayed = replay_formatting_diagnostics_from_report(
        report_path=report_path,
        source_docx_path=source_docx_path,
    )

    assert replayed["source_reconstruction_basis"] == "source_docx_override"
    assert replayed["source_docx_path"] == str(source_docx_path)
    assert replayed["saved_source_count"] == 2
    assert replayed["replayed_source_count"] == 1
    assert replayed["replay_fidelity"] == "source_count_mismatch_vs_saved_report"


def test_replay_formatting_diagnostics_from_report_reconstructs_from_final_generated_registry(tmp_path: Path):
    target_docx_path = tmp_path / "target.docx"
    document = Document()
    document.add_paragraph("Глава восьмая", style="Heading 2")
    document.add_paragraph("Body translated")
    buffer = BytesIO()
    document.save(buffer)
    target_docx_path.write_bytes(buffer.getvalue())

    report_path = tmp_path / "report.json"
    report_payload = {
        "output_artifacts": {"docx_path": str(target_docx_path)},
        "formatting_diagnostics": [],
        "runtime": {
            "state": {
                "final_generated_paragraph_registry": [
                    {"paragraph_id": "p0001", "text": "## Глава восьмая", "target_paragraph_indexes": [0]},
                    {"paragraph_id": "p0002", "text": "Body translated", "target_paragraph_indexes": [1]},
                ]
            }
        },
    }
    report_path.write_text(json.dumps(report_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    replayed = replay_formatting_diagnostics_from_report(report_path=report_path)

    assert replayed["source_reconstruction_basis"] == "final_generated_paragraph_registry"
    assert replayed["replay_fidelity"] == "current_output_registry_replay"
    assert replayed["saved_source_count"] == 2
    assert replayed["replayed_source_count"] == 2
    assert replayed["saved_mapped_count"] is None
    assert replayed["replayed_diagnostics"]["mapped_count"] == 2
    assert replayed["replayed_summary"]["unmapped_source_count"] == 0
    role_aware_summary = replayed["replayed_summary"]["role_aware_summary"]
    assert role_aware_summary is None or role_aware_summary["effective_unmapped_source_count"] == 0


def test_replay_formatting_diagnostics_prefers_final_registry_over_source_language_preview(tmp_path: Path):
    target_docx_path = tmp_path / "target.docx"
    document = Document()
    document.add_paragraph("СОДЕРЖАНИЕ", style="Heading 1")
    document.add_paragraph("Предисловие")
    buffer = BytesIO()
    document.save(buffer)
    target_docx_path.write_bytes(buffer.getvalue())

    report_path = tmp_path / "report.json"
    report_payload = {
        "output_artifacts": {"docx_path": str(target_docx_path)},
        "formatting_diagnostics": [
            {
                "stage": "restore",
                "mapped_count": 2,
                "unmapped_source_ids": [],
                "unmapped_target_indexes": [],
                "source_registry": [
                    {
                        "paragraph_id": "p0001",
                        "source_index": 0,
                        "text_preview": "contents",
                        "role": "heading",
                        "structural_role": "heading",
                        "heading_level": 1,
                    },
                    {
                        "paragraph_id": "p0002",
                        "source_index": 1,
                        "text_preview": "foreword",
                        "role": "body",
                        "structural_role": "body",
                    },
                ],
                "target_registry": [
                    {"target_index": 0, "text_preview": "СОДЕРЖАНИЕ", "heading_level": 1},
                    {"target_index": 1, "text_preview": "Предисловие"},
                ],
            }
        ],
        "runtime": {
            "state": {
                "final_generated_paragraph_registry": [
                    {"paragraph_id": "p0001", "text": "# СОДЕРЖАНИЕ", "target_paragraph_indexes": [0]},
                    {"paragraph_id": "p0002", "text": "Предисловие", "target_paragraph_indexes": [1]},
                ]
            }
        },
    }
    report_path.write_text(json.dumps(report_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    replayed = replay_formatting_diagnostics_from_report(report_path=report_path)

    assert replayed["source_reconstruction_basis"] == "final_generated_paragraph_registry"
    assert replayed["replay_fidelity"] == "current_output_registry_replay"
    assert replayed["replayed_diagnostics"]["mapped_count"] == 2
    assert replayed["replayed_summary"]["unmapped_source_count"] == 0


def test_classify_formatting_residuals_replay_mode_reports_replayed_restore_diagnostics(tmp_path: Path):
    target_docx_path = tmp_path / "target.docx"
    target_docx_path.write_bytes(_build_docx_bytes(["Heading translated", "Body translated"]))

    report_path = tmp_path / "report.json"
    report_payload = {
        "output_artifacts": {"docx_path": str(target_docx_path)},
        "formatting_diagnostics": [
            {
                "stage": "post_cleanup",
                "mapped_count": 0,
                "unmapped_source_ids": ["p0001", "p0002"],
                "unmapped_target_indexes": [0, 1],
                "source_registry": [
                    {
                        "paragraph_id": "p0001",
                        "source_index": 0,
                        "text_preview": "Heading translated",
                        "role": "heading",
                        "structural_role": "heading",
                        "heading_level": 1,
                    },
                    {
                        "paragraph_id": "p0002",
                        "source_index": 1,
                        "text_preview": "Body translated",
                        "role": "body",
                        "structural_role": "body",
                    },
                ],
            }
        ],
    }
    report_path.write_text(json.dumps(report_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    completed = subprocess.run(
        [
            sys.executable,
            "scripts/classify-formatting-residuals.py",
            str(report_path),
            "--replay",
        ],
        cwd=Path(__file__).resolve().parents[1],
        capture_output=True,
        text=True,
        check=True,
    )
    payload = json.loads(completed.stdout)

    assert payload["classification_basis"] == "replayed_restore_diagnostics"
    assert payload["replay_scope"] == "restore_diagnostics_from_saved_final_docx"
    assert payload["replay_mode"] == "no_llm_current_mapping_code"
    assert payload["source_reconstruction_basis"] == "saved_source_registry_preview"
    assert payload["replay_fidelity"] == "matched_saved_source_count"
    assert payload["saved_mapped_count"] == 0
    assert payload["mapped_count"] == 2
    assert payload["unmapped_source_count"] == 0
