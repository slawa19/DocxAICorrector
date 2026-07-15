"""Characterization gate for the formatting-transfer restoration/facade path (spec 033).

Whereas ``tests/test_formatting_mapper_golden.py`` freezes the mapper's blob, this
gate freezes the *restoration* half of ``generation/formatting_transfer.py`` — the
docx-apply functions (image/caption/list/TOC/alignment/quote/split-heading) plus the
facade entry point ``preserve_source_paragraph_properties`` / ``apply_output_formatting``
and the diagnostics artifact it writes.

A single in-memory document exercises, in one pass, every restoration seam:

* a numbered **list item** (numbering restore -> ``word/numbering.xml`` gains a definition),
* a **TOC header + entry** (pPr graft + safe run formatting incl. italic),
* a **centered caption** (Caption style promotion via image anchor + alignment allowlist),
* an **image-only paragraph** (centered by minimal image formatting),
* a **centered attribution** and an **epigraph** (alignment allowlist + semantic-quote italic),
* a **split-heading** source whose registry text splits into a heading + body target
  (accepted by the mapper; the accepted-split decision is snapshotted).

Both goldens are deterministic (fixed inputs, no RNG / time). To regenerate after an
intentional, reviewed behaviour change, run with::

    UPDATE_FORMATTING_RESTORATION_GOLDEN=1 <run this test>

Golden A — the restored docx: canonicalized (pretty, attribute-sorted) ``word/document.xml``
and ``word/numbering.xml`` from ``preserve_source_paragraph_properties(...)``.
Golden B — the written diagnostics artifact's restoration-decision payload
(``toc_format_restoration_decisions``, ``toc_run_format_restoration_decisions``,
``alignment_restoration_decisions``, ``list_restoration_decisions``,
``caption_heading_conflicts``, plus ``accepted_split_targets``).
"""

from __future__ import annotations

import json
import os
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

import docxaicorrector.generation.formatting_transfer as formatting_transfer
from docxaicorrector.core.models import ParagraphUnit

# Builds an in-memory docx and drives the full facade path — an integration-grade
# characterization gate, not a fast unit test.
pytestmark = pytest.mark.integration_local

_REPO_ROOT = Path(__file__).resolve().parent.parent
_FIXTURE_DIR = _REPO_ROOT / "tests" / "fixtures" / "formatting_restoration_golden"

_UPDATE_ENV = "UPDATE_FORMATTING_RESTORATION_GOLDEN"

_TOC_HEADER_PPR = (
    '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:pStyle w:val="Normal"/>'
    '<w:spacing w:before="153" w:after="10" w:line="272" w:lineRule="exact"/>'
    '<w:ind w:start="1519" w:end="0" w:hanging="0"/>'
    '<w:tabs><w:tab w:val="right" w:leader="dot" w:pos="9350"/></w:tabs>'
    '<w:jc w:val="end"/>'
    '<w:keepNext/>'
    '</w:pPr>'
)
_TOC_ENTRY_PPR = (
    '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:pStyle w:val="Normal"/>'
    '<w:spacing w:before="281" w:after="10" w:line="272" w:lineRule="exact"/>'
    '<w:ind w:start="1588" w:end="4878" w:hanging="0"/>'
    '<w:tabs><w:tab w:val="right" w:leader="dot" w:pos="9350"/></w:tabs>'
    '<w:jc w:val="end"/>'
    '<w:keepLines/>'
    '</w:pPr>'
)
_ABSTRACT_NUM_XML = (
    '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'w:abstractNumId="7">'
    '<w:lvl w:ilvl="0"><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/></w:lvl>'
    '</w:abstractNum>'
)
_NUM_XML = (
    '<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="9">'
    '<w:abstractNumId w:val="7"/></w:num>'
)

# The restoration-decision keys the artifact carries (spec 033 Golden B). The first
# four are grafted onto the diagnostics by ``apply_output_formatting``; the last two
# are produced by the mapper and forwarded through the artifact unchanged.
_ARTIFACT_DECISION_KEYS = (
    "toc_format_restoration_decisions",
    "toc_run_format_restoration_decisions",
    "alignment_restoration_decisions",
    "list_restoration_decisions",
    "caption_heading_conflicts",
    "accepted_split_targets",
)


def _build_case() -> tuple[list[ParagraphUnit], bytes, list[dict[str, object]]]:
    """Deterministically synthesize (source paragraphs, target docx bytes, registry)."""
    source_paragraphs = [
        ParagraphUnit(
            paragraph_id="p0000",
            text="Содержание",
            role="body",
            structural_role="toc_header",
            style_name="Normal",
            paragraph_alignment="end",
            font_size_pt=15.0,
            paragraph_properties_xml=_TOC_HEADER_PPR,
        ),
        ParagraphUnit(
            paragraph_id="p0001",
            text="Введение в тему",
            role="body",
            structural_role="toc_entry",
            style_name="Normal",
            is_italic=True,
            font_size_pt=11.0,
            paragraph_properties_xml=_TOC_ENTRY_PPR,
        ),
        ParagraphUnit(
            paragraph_id="p0002",
            text="Первый пункт списка",
            role="list",
            structural_role="list",
            list_kind="ordered",
            list_level=0,
            list_num_id="9",
            list_abstract_num_id="7",
            list_num_xml=_NUM_XML,
            list_abstract_num_xml=_ABSTRACT_NUM_XML,
        ),
        ParagraphUnit(
            paragraph_id="p0003",
            text="Рис. 1. Схема устройства",
            role="caption",
            structural_role="caption",
            paragraph_alignment="center",
        ),
        ParagraphUnit(
            paragraph_id="p0004",
            text="— И. ИВАНОВ",
            role="body",
            structural_role="attribution",
            paragraph_alignment="center",
            is_italic=True,
        ),
        ParagraphUnit(
            paragraph_id="p0005",
            text="Знание сила",
            role="body",
            structural_role="epigraph",
            paragraph_alignment="center",
            is_italic=True,
        ),
        ParagraphUnit(
            paragraph_id="p0006",
            text="Заголовок главы Первый абзац основного текста.",
            role="body",
            structural_role="body",
        ),
    ]

    registry: list[dict[str, object]] = [
        {"paragraph_id": "p0000", "text": "Содержание"},
        {"paragraph_id": "p0001", "text": "Введение в тему"},
        {"paragraph_id": "p0002", "text": "Первый пункт списка"},
        {"paragraph_id": "p0003", "text": "Рис. 1. Схема устройства"},
        {"paragraph_id": "p0004", "text": "— И. ИВАНОВ"},
        {"paragraph_id": "p0005", "text": "Знание сила"},
        {"paragraph_id": "p0006", "text": "### Заголовок главы\nПервый абзац основного текста."},
    ]

    target = Document()
    target.add_paragraph("Содержание", style="Body Text")
    target.add_paragraph("Введение в тему", style="Body Text")
    target.add_paragraph("Первый пункт списка")
    target.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target.add_paragraph("Рис. 1. Схема устройства")
    target.add_paragraph("— И. ИВАНОВ")
    target.add_paragraph("Знание сила")
    target.add_paragraph("Заголовок главы", style="Heading 2")
    target.add_paragraph("Первый абзац основного текста.")
    buffer = BytesIO()
    target.save(buffer)
    return source_paragraphs, buffer.getvalue(), registry


def _canonicalize_xml(xml_bytes: bytes) -> str:
    """Pretty, attribute-sorted canonical form for a stable, diffable snapshot."""
    parser = etree.XMLParser(remove_blank_text=True)
    root = etree.fromstring(xml_bytes, parser=parser)
    for element in root.iter():
        if len(element.attrib) > 1:
            sorted_items = sorted(element.attrib.items())
            element.attrib.clear()
            for key, value in sorted_items:
                element.set(key, value)
    return etree.tostring(root, pretty_print=True, encoding="unicode")


def _restored_document_bytes(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> bytes:
    # Redirect the diagnostics artifact away from the repo working tree.
    monkeypatch.setattr(formatting_transfer, "FORMATTING_DIAGNOSTICS_DIR", tmp_path)
    source_paragraphs, target_bytes, registry = _build_case()
    return formatting_transfer.preserve_source_paragraph_properties(
        target_bytes,
        source_paragraphs,
        generated_paragraph_registry=registry,
    )


def _extract_part(docx_bytes: bytes, part_name: str) -> bytes:
    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        return archive.read(part_name)


def _check_or_update(fixture_name: str, actual: str) -> None:
    fixture_path = _FIXTURE_DIR / fixture_name
    if os.environ.get(_UPDATE_ENV):
        fixture_path.parent.mkdir(parents=True, exist_ok=True)
        fixture_path.write_text(actual, encoding="utf-8")
        pytest.skip(f"regenerated golden fixture: {fixture_path.name}")
    assert fixture_path.exists(), (
        f"missing golden fixture {fixture_path}; regenerate with {_UPDATE_ENV}=1"
    )
    expected = fixture_path.read_text(encoding="utf-8")
    assert actual == expected, (
        f"formatting restoration output diverged from golden {fixture_path.name}; "
        f"regenerate with {_UPDATE_ENV}=1 only after a reviewed behaviour change"
    )


def test_restored_document_xml_matches_golden(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    restored_bytes = _restored_document_bytes(monkeypatch, tmp_path)
    document_xml = _canonicalize_xml(_extract_part(restored_bytes, "word/document.xml"))
    _check_or_update("restored_document.xml", document_xml)


def test_restored_numbering_xml_matches_golden(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    restored_bytes = _restored_document_bytes(monkeypatch, tmp_path)
    numbering_xml = _canonicalize_xml(_extract_part(restored_bytes, "word/numbering.xml"))
    _check_or_update("restored_numbering.xml", numbering_xml)


def test_restoration_artifact_decisions_match_golden(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    monkeypatch.setattr(formatting_transfer, "FORMATTING_DIAGNOSTICS_DIR", tmp_path)
    source_paragraphs, target_bytes, registry = _build_case()
    formatting_transfer.preserve_source_paragraph_properties(
        target_bytes,
        source_paragraphs,
        generated_paragraph_registry=registry,
    )

    artifacts = sorted(tmp_path.glob("*.json"))
    assert len(artifacts) == 1, f"expected exactly one diagnostics artifact, found {artifacts}"
    payload = json.loads(artifacts[0].read_text(encoding="utf-8"))

    decisions = {key: payload.get(key) for key in _ARTIFACT_DECISION_KEYS}
    blob = json.dumps(decisions, ensure_ascii=False, sort_keys=True, indent=2) + "\n"
    _check_or_update("restored_artifact_decisions.json", blob)
