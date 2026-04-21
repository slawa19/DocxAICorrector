import base64
from io import BytesIO

import image_reinsertion
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches

from image_reinsertion import (
    _build_variant_block_elements,
    _replace_xml_element_with_sequence,
    resolve_final_image_bytes,
    resolve_image_insertions,
    reinsert_inline_images,
)
from models import ImageAsset, ImageVariantCandidate


PNG_BYTES = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aK3cAAAAASUVORK5CYII=")


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_textbox_with_text(paragraph, text: str) -> None:
    paragraph._p.append(
        parse_xml(
            f"""
            <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                 xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                 xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                 xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <w:drawing>
                    <wp:inline>
                        <wp:extent cx="914400" cy="914400"/>
                        <wp:docPr id="1" name="TextBox 1"/>
                        <a:graphic>
                            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                                <wps:wsp>
                                    <wps:txbx>
                                        <w:txbxContent>
                                            <w:p>
                                                <w:r>
                                                    <w:t>{text}</w:t>
                                                </w:r>
                                            </w:p>
                                        </w:txbxContent>
                                    </wps:txbx>
                                    <wps:bodyPr/>
                                </wps:wsp>
                            </a:graphicData>
                        </a:graphic>
                    </wp:inline>
                </w:drawing>
            </w:r>
            """
        )
    )


def _extract_docpr_descriptions(element) -> list[str]:
    return [doc_pr.get("descr") for doc_pr in element.xpath(".//wp:docPr") if doc_pr.get("descr")]


def _extract_source_rects(element) -> list[dict[str, str]]:
    return [
        {key: src_rect.get(key) for key in ("l", "t", "r", "b") if src_rect.get(key) is not None}
        for src_rect in element.xpath(".//a:srcRect")
    ]


def _append_decimal_numbering_definition(document, *, num_id: str, abstract_num_id: str) -> None:
    numbering_root = document.part.numbering_part.element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_num_id)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    abstract_num.append(lvl)
    numbering_root.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_num_ref = OxmlElement("w:abstractNumId")
    abstract_num_ref.set(qn("w:val"), abstract_num_id)
    num.append(abstract_num_ref)
    numbering_root.append(num)


def _attach_numbering(paragraph, *, num_id: str, ilvl: str) -> None:
    paragraph_properties = paragraph._element.get_or_add_pPr()
    num_pr = paragraph_properties.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        paragraph_properties.append(num_pr)

    ilvl_element = num_pr.find(qn("w:ilvl"))
    if ilvl_element is None:
        ilvl_element = OxmlElement("w:ilvl")
        num_pr.append(ilvl_element)
    ilvl_element.set(qn("w:val"), ilvl)

    num_id_element = num_pr.find(qn("w:numId"))
    if num_id_element is None:
        num_id_element = OxmlElement("w:numId")
        num_pr.append(num_id_element)
    num_id_element.set(qn("w:val"), num_id)


def test_replace_xml_element_with_sequence_empty_replacements_is_noop():
    parent = parse_xml("<root><child>text</child></root>")
    child = parent[0]

    _replace_xml_element_with_sequence(child, [])

    assert len(parent) == 1
    assert parent[0].text == "text"


def test_resolve_image_insertions_keeps_safe_and_candidates_for_manual_review():
    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        safe_bytes=PNG_BYTES,
        attempt_variants=[
            ImageVariantCandidate(mode="candidate1", bytes=PNG_BYTES, mime_type="image/png"),
            ImageVariantCandidate(mode="candidate2", bytes=PNG_BYTES, mime_type="image/png"),
        ],
    )
    asset.update_pipeline_metadata(preserve_all_variants_in_docx=True)

    assert resolve_image_insertions(asset) == [
        ("safe", PNG_BYTES),
        ("candidate1", PNG_BYTES),
        ("candidate2", PNG_BYTES),
    ]


def test_reinsert_inline_images_labels_manual_review_variants():
    doc = Document()
    doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        safe_bytes=PNG_BYTES,
        attempt_variants=[
            ImageVariantCandidate(mode="candidate1", bytes=PNG_BYTES, mime_type="image/png"),
            ImageVariantCandidate(mode="candidate2", bytes=PNG_BYTES, mime_type="image/png"),
        ],
    )
    asset.update_pipeline_metadata(preserve_all_variants_in_docx=True)

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [asset])
    updated_doc = Document(BytesIO(updated_bytes))
    visible_text = "\n".join(paragraph.text for paragraph in updated_doc.paragraphs)
    visible_text += "\n" + "\n".join(
        paragraph.text
        for table in updated_doc.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )

    assert len(updated_doc.tables) == 0
    assert len(updated_doc.inline_shapes) == 3
    assert len(updated_doc.paragraphs) == 3
    assert "candidate1" not in visible_text
    assert "candidate2" not in visible_text
    assert _extract_docpr_descriptions(updated_doc._element) == ["safe", "candidate1", "candidate2"]


def test_reinsert_inline_images_replaces_placeholder_with_picture():
    doc = Document()
    doc.add_paragraph("До")
    image_paragraph = doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("После")
    buffer = BytesIO()
    doc.save(buffer)

    expected_indent = image_paragraph.paragraph_format.left_indent

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                final_variant="original",
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[1].text == ""
    assert updated_doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert updated_doc.paragraphs[1].paragraph_format.left_indent == expected_indent
    assert len(updated_doc.inline_shapes) == 1
    assert updated_doc.inline_shapes[0].width == 914400
    assert updated_doc.inline_shapes[0].height == 914400


def test_reinsert_inline_images_preserves_formatted_text_around_placeholder_in_same_paragraph():
    doc = Document()
    paragraph = doc.add_paragraph()
    before_run = paragraph.add_run("До ")
    before_run.bold = True
    paragraph.add_run("[[DOCX_IMAGE_img_001]]")
    after_run = paragraph.add_run(" после")
    after_run.italic = True
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                final_variant="original",
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))
    updated_paragraph = updated_doc.paragraphs[0]

    assert updated_paragraph.text == "До  после"
    assert updated_paragraph.runs[0].text == "До "
    assert updated_paragraph.runs[0].bold is True
    assert updated_paragraph.runs[-1].text == " после"
    assert updated_paragraph.runs[-1].italic is True
    assert len(updated_doc.inline_shapes) == 1


def test_reinsert_inline_images_preserves_hyperlink_xml_when_placeholder_is_in_same_paragraph():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("До ")
    _add_hyperlink(paragraph, "ссылка", "https://example.com")
    paragraph.add_run(" [[DOCX_")
    paragraph.add_run("IMAGE_img_001]] после")
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                final_variant="original",
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))
    updated_paragraph = updated_doc.paragraphs[0]

    assert "[[DOCX_IMAGE_img_001]]" not in updated_paragraph.text
    assert "ссылка" in updated_paragraph.text
    assert "после" in updated_paragraph.text
    assert len(updated_doc.inline_shapes) == 1
    assert len(updated_paragraph._element.xpath("./w:hyperlink")) == 1


def test_reinsert_inline_images_uses_shared_block_layout_for_multi_variant_placeholder_inside_paragraph():
    doc = Document()
    paragraph = doc.add_paragraph()
    before_run = paragraph.add_run("До ")
    before_run.bold = True
    paragraph.add_run("[[DOCX_IMAGE_img_001]]")
    after_run = paragraph.add_run(" после")
    after_run.italic = True
    buffer = BytesIO()
    doc.save(buffer)

    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        safe_bytes=PNG_BYTES,
        attempt_variants=[
            ImageVariantCandidate(mode="candidate1", bytes=PNG_BYTES, mime_type="image/png"),
            ImageVariantCandidate(mode="candidate2", bytes=PNG_BYTES, mime_type="image/png"),
        ],
    )
    asset.update_pipeline_metadata(preserve_all_variants_in_docx=True)

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [asset])
    updated_doc = Document(BytesIO(updated_bytes))

    assert [paragraph.text for paragraph in updated_doc.paragraphs] == ["До ", "", "", "", " после"]
    assert updated_doc.paragraphs[0].runs[0].bold is True
    assert updated_doc.paragraphs[-1].runs[0].italic is True
    assert len(updated_doc.tables) == 0
    assert len(updated_doc.inline_shapes) == 3
    assert _extract_docpr_descriptions(updated_doc._element) == ["safe", "candidate1", "candidate2"]


def test_reinsert_inline_images_preserves_hyperlink_when_multi_variant_blocks_are_inserted_nearby():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("До ")
    _add_hyperlink(paragraph, "ссылка", "https://example.com")
    paragraph.add_run(" [[DOCX_IMAGE_img_001]] после")
    buffer = BytesIO()
    doc.save(buffer)

    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        safe_bytes=PNG_BYTES,
        attempt_variants=[
            ImageVariantCandidate(mode="candidate1", bytes=PNG_BYTES, mime_type="image/png"),
            ImageVariantCandidate(mode="candidate2", bytes=PNG_BYTES, mime_type="image/png"),
        ],
    )
    asset.update_pipeline_metadata(preserve_all_variants_in_docx=True)

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [asset])
    updated_doc = Document(BytesIO(updated_bytes))
    visible_text = "\n".join(paragraph.text for paragraph in updated_doc.paragraphs)

    assert "ссылка" in visible_text
    assert "после" in visible_text
    assert len(updated_doc.tables) == 0
    assert len(updated_doc.inline_shapes) == 3
    assert len(updated_doc._element.xpath(".//w:hyperlink")) == 1
    assert _extract_docpr_descriptions(updated_doc._element) == ["safe", "candidate1", "candidate2"]


def test_reinsert_inline_images_logs_warning_when_all_replacement_strategies_fail(monkeypatch):
    doc = Document()
    doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
    )

    events = []
    monkeypatch.setattr(
        image_reinsertion,
        "_replace_multi_variant_placeholders_with_blocks",
        lambda paragraph, asset_map, insertion_cache: False,
    )
    monkeypatch.setattr(
        image_reinsertion,
        "_replace_run_level_placeholders",
        lambda paragraph, placeholders, asset_map, insertion_cache: False,
    )
    monkeypatch.setattr(
        image_reinsertion,
        "_replace_multi_run_placeholders",
        lambda paragraph, asset_map, insertion_cache: False,
    )
    monkeypatch.setattr(
        image_reinsertion,
        "_replace_paragraph_placeholders_fallback",
        lambda paragraph, paragraph_text, asset_map, insertion_cache: False,
    )
    monkeypatch.setattr(image_reinsertion, "log_event", lambda level, event, message, **context: events.append((event, context)))

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [asset])
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].text == "[[DOCX_IMAGE_img_001]]"
    assert events == [
        (
            "image_reinsertion_placeholder_unhandled",
            {
                "placeholder_count": 1,
                "placeholders": ["[[DOCX_IMAGE_img_001]]"],
                "paragraph_text_preview": "[[DOCX_IMAGE_img_001]]",
            },
        )
    ]


def test_build_variant_block_elements_returns_empty_for_empty_insertions(monkeypatch):
    doc = Document()
    paragraph = doc.add_paragraph("placeholder")
    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
    )

    monkeypatch.setattr(image_reinsertion, "resolve_image_insertions", lambda current_asset: [])

    assert _build_variant_block_elements(paragraph, asset) == []


def test_reinsert_inline_images_keeps_placeholder_text_when_no_image_bytes_resolved():
    doc = Document()
    doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=b"",
                mime_type="image/png",
                position_index=0,
                final_variant="original",
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].text == "[[DOCX_IMAGE_img_001]]"
    assert len(updated_doc.inline_shapes) == 0


def test_reinsert_inline_images_replaces_placeholder_with_picture_inside_table_cell():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                final_variant="original",
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))
    cell_paragraph = updated_doc.tables[0].cell(0, 0).paragraphs[0]

    assert cell_paragraph.text == ""
    assert len(updated_doc.inline_shapes) == 1
    assert updated_doc.inline_shapes[0].width == 914400
    assert updated_doc.inline_shapes[0].height == 914400


def test_reinsert_inline_images_reapplies_source_rect_and_doc_properties_for_original_asset():
    doc = Document()
    doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                final_variant="original",
                source_forensics={
                    "source_rect": {"l": 1250, "t": 2500, "r": 3750, "b": 5000},
                    "doc_properties": {
                        "descr": "Исходное описание",
                        "title": "Исходный title",
                        "name": "Исходное имя",
                    },
                },
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert _extract_source_rects(updated_doc._element) == [{"l": "1250", "t": "2500", "r": "3750", "b": "5000"}]
    doc_pr = updated_doc._element.xpath(".//wp:docPr")[0]
    assert doc_pr.get("descr") == "Исходное описание"
    assert doc_pr.get("title") == "Исходный title"
    assert doc_pr.get("name") == "Исходное имя"


def test_reinsert_inline_images_restores_anchor_container_from_source_forensics():
    doc = Document()
    doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    source_anchor_xml = """
    <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
               simplePos="0" relativeHeight="0" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>
        <wp:extent cx="914400" cy="914400"/>
        <wp:wrapNone/>
        <wp:docPr id="7" name="Source Anchor" descr="Исходный anchor"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic/>
            </a:graphicData>
        </a:graphic>
    </wp:anchor>
    """

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                final_variant="original",
                source_forensics={
                    "drawing_container": "anchor",
                    "drawing_container_xml": source_anchor_xml,
                    "doc_properties": {"descr": "Исходный anchor", "name": "Source Anchor"},
                },
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert len(updated_doc._element.xpath(".//wp:anchor")) == 1
    assert len(updated_doc._element.xpath(".//wp:inline")) == 0
    assert updated_doc._element.xpath(".//wp:docPr")[0].get("descr") == "Исходный anchor"


def test_reinsert_inline_images_replaces_placeholder_with_picture_inside_nested_table_cell():
    doc = Document()
    outer_table = doc.add_table(rows=1, cols=1)
    nested_table = outer_table.cell(0, 0).add_table(rows=1, cols=1)
    nested_table.cell(0, 0).paragraphs[0].add_run("[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                final_variant="original",
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))
    nested_cell_paragraph = updated_doc.tables[0].cell(0, 0).tables[0].cell(0, 0).paragraphs[0]

    assert nested_cell_paragraph.text == ""
    assert len(updated_doc.inline_shapes) == 1
    assert updated_doc.inline_shapes[0].width == 914400
    assert updated_doc.inline_shapes[0].height == 914400


def test_reinsert_inline_images_replaces_placeholder_split_across_runs_without_plain_text_fallback():
    doc = Document()
    paragraph = doc.add_paragraph()
    first_run = paragraph.add_run("До [[DOCX_")
    first_run.bold = True
    second_run = paragraph.add_run("IMAGE_img_001]] после")
    second_run.italic = True
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                final_variant="original",
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))
    updated_paragraph = updated_doc.paragraphs[0]

    assert updated_paragraph.text == "До  после"
    assert updated_paragraph.runs[0].text == "До "
    assert updated_paragraph.runs[0].bold is True
    assert updated_paragraph.runs[-1].text == " после"
    assert updated_paragraph.runs[-1].italic is True
    assert len(updated_doc.inline_shapes) == 1


def test_reinsert_inline_images_uses_shared_block_layout_for_split_run_multi_variant_placeholder():
    doc = Document()
    paragraph = doc.add_paragraph()
    first_run = paragraph.add_run("До [[DOCX_")
    first_run.bold = True
    second_run = paragraph.add_run("IMAGE_img_001]] после")
    second_run.italic = True
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                validation_status="compared",
                comparison_variants={
                    "safe": {"bytes": PNG_BYTES},
                    "semantic_redraw_direct": {"bytes": PNG_BYTES},
                    "semantic_redraw_structured": {"bytes": PNG_BYTES},
                },
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert [paragraph.text for paragraph in updated_doc.paragraphs] == ["До ", "", "", "", " после"]
    assert updated_doc.paragraphs[0].runs[0].bold is True
    assert updated_doc.paragraphs[-1].runs[0].italic is True
    assert len(updated_doc.tables) == 0
    assert len(updated_doc.inline_shapes) == 3
    assert _extract_docpr_descriptions(updated_doc._element) == [
        "Вариант 1: Просто улучшить",
        "Вариант 2: Креативная AI-перерисовка",
        "Вариант 3: Структурная AI-перерисовка",
    ]


def test_reinsert_inline_images_replaces_placeholder_in_header_and_footer():
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].add_run("[[DOCX_IMAGE_img_001]]")
    section.footer.paragraphs[0].add_run("[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                final_variant="original",
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    header_xml = updated_doc.sections[0].header._element.xml
    footer_xml = updated_doc.sections[0].footer._element.xml

    assert "[[DOCX_IMAGE_img_001]]" not in header_xml
    assert "[[DOCX_IMAGE_img_001]]" not in footer_xml
    assert "a:blip" in header_xml
    assert "a:blip" in footer_xml


def test_reinsert_inline_images_replaces_placeholder_inside_textbox():
    doc = Document()
    host_paragraph = doc.add_paragraph("Перед textbox")
    _append_textbox_with_text(host_paragraph, "[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=PNG_BYTES,
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                final_variant="original",
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))
    document_xml = updated_doc._element.xml

    assert "[[DOCX_IMAGE_img_001]]" not in document_xml
    assert "w:txbxContent" in document_xml
    assert "a:blip" in document_xml


def test_reinsert_inline_images_in_compare_all_mode_inserts_all_generated_variants():
    doc = Document()
    doc.add_paragraph("До")
    doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    doc.add_paragraph("После")
    buffer = BytesIO()
    doc.save(buffer)

    updated_bytes = reinsert_inline_images(
        buffer.getvalue(),
        [
            ImageAsset(
                image_id="img_001",
                placeholder="[[DOCX_IMAGE_img_001]]",
                original_bytes=b"original",
                mime_type="image/png",
                position_index=0,
                width_emu=914400,
                height_emu=914400,
                validation_status="compared",
                comparison_variants={
                    "safe": {"bytes": PNG_BYTES},
                    "semantic_redraw_direct": {"bytes": PNG_BYTES},
                    "semantic_redraw_structured": {"bytes": PNG_BYTES},
                },
            )
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))
    visible_text = "\n".join(paragraph.text for paragraph in updated_doc.paragraphs)
    visible_text += "\n" + "\n".join(
        paragraph.text
        for table in updated_doc.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )

    assert len(updated_doc.inline_shapes) == 3
    assert len(updated_doc.tables) == 0
    assert [paragraph.text for paragraph in updated_doc.paragraphs] == ["До", "", "", "", "После"]
    assert "Вариант 1: Просто улучшить" not in visible_text
    assert "Вариант 2: Креативная AI-перерисовка" not in visible_text
    assert "Вариант 3: Структурная AI-перерисовка" not in visible_text
    assert _extract_docpr_descriptions(updated_doc._element) == [
        "Вариант 1: Просто улучшить",
        "Вариант 2: Креативная AI-перерисовка",
        "Вариант 3: Структурная AI-перерисовка",
    ]


def test_reinsert_inline_images_resolves_multi_variant_insertions_once_per_placeholder(monkeypatch):
    doc = Document()
    doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
    )
    call_count = 0

    def fake_resolve_image_insertions(current_asset):
        nonlocal call_count
        call_count += 1
        return [
            ("safe", PNG_BYTES),
            ("candidate1", PNG_BYTES),
            ("candidate2", PNG_BYTES),
        ]

    monkeypatch.setattr(image_reinsertion, "resolve_image_insertions", fake_resolve_image_insertions)

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [asset])
    updated_doc = Document(BytesIO(updated_bytes))

    assert call_count == 1
    assert len(updated_doc.inline_shapes) == 3


def test_reinsert_inline_images_resolves_reused_placeholder_once_per_pass_across_paragraphs(monkeypatch):
    doc = Document()
    doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    doc.add_paragraph("Before [[DOCX_IMAGE_img_001]] after")
    buffer = BytesIO()
    doc.save(buffer)

    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
    )
    call_count = 0

    def fake_resolve_image_insertions(current_asset):
        nonlocal call_count
        call_count += 1
        assert current_asset.placeholder == "[[DOCX_IMAGE_img_001]]"
        return [(None, PNG_BYTES)]

    monkeypatch.setattr(image_reinsertion, "resolve_image_insertions", fake_resolve_image_insertions)

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [asset])
    updated_doc = Document(BytesIO(updated_bytes))

    assert call_count == 1
    assert len(updated_doc.inline_shapes) == 2
    assert updated_doc.paragraphs[0].text == ""
    assert updated_doc.paragraphs[1].text == "Before  after"


def test_reinsert_inline_images_different_placeholders_keep_separate_cache_entries(monkeypatch):
    doc = Document()
    doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    doc.add_paragraph("[[DOCX_IMAGE_img_002]]")
    buffer = BytesIO()
    doc.save(buffer)

    first_asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        width_emu=914400,
        height_emu=914400,
        final_variant="original",
    )
    second_asset = ImageAsset(
        image_id="img_002",
        placeholder="[[DOCX_IMAGE_img_002]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=1,
        width_emu=1828800,
        height_emu=1828800,
        final_variant="original",
    )
    seen_placeholders = []

    def fake_resolve_image_insertions(current_asset):
        seen_placeholders.append(current_asset.placeholder)
        return [(None, current_asset.original_bytes)]

    monkeypatch.setattr(image_reinsertion, "resolve_image_insertions", fake_resolve_image_insertions)

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [first_asset, second_asset])
    updated_doc = Document(BytesIO(updated_bytes))

    assert seen_placeholders == ["[[DOCX_IMAGE_img_001]]", "[[DOCX_IMAGE_img_002]]"]
    assert len(updated_doc.inline_shapes) == 2
    assert updated_doc.inline_shapes[0].width == 914400
    assert updated_doc.inline_shapes[1].width == 1828800


def test_reinsert_inline_images_multi_variant_blocks_drop_list_indent_and_keep_next_formatting():
    doc = Document()
    _append_decimal_numbering_definition(doc, num_id="77", abstract_num_id="700")
    paragraph = doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    _attach_numbering(paragraph, num_id="77", ilvl="0")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph_properties = paragraph._element.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    paragraph_properties.append(keep_next)
    buffer = BytesIO()
    doc.save(buffer)

    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        safe_bytes=PNG_BYTES,
        attempt_variants=[
            ImageVariantCandidate(mode="candidate1", bytes=PNG_BYTES, mime_type="image/png"),
            ImageVariantCandidate(mode="candidate2", bytes=PNG_BYTES, mime_type="image/png"),
        ],
    )
    asset.update_pipeline_metadata(preserve_all_variants_in_docx=True)

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [asset])
    updated_doc = Document(BytesIO(updated_bytes))

    assert len(updated_doc.paragraphs) == 3
    for updated_paragraph in updated_doc.paragraphs:
        paragraph_properties = updated_paragraph._element.pPr
        assert paragraph_properties is not None
        alignment = paragraph_properties.find(qn("w:jc"))
        assert alignment is not None
        assert alignment.get(qn("w:val")) == "center"
        assert paragraph_properties.find(qn("w:numPr")) is None
        assert paragraph_properties.find(qn("w:ind")) is None
        assert paragraph_properties.find(qn("w:keepNext")) is None


def test_reinsert_inline_images_multi_variant_blocks_drop_heading_style_from_source_paragraph():
    doc = Document()
    paragraph = doc.add_paragraph("[[DOCX_IMAGE_img_001]]", style="Heading 1")
    paragraph_properties = paragraph._element.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    paragraph_properties.append(keep_lines)
    buffer = BytesIO()
    doc.save(buffer)

    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        safe_bytes=PNG_BYTES,
        attempt_variants=[
            ImageVariantCandidate(mode="candidate1", bytes=PNG_BYTES, mime_type="image/png"),
            ImageVariantCandidate(mode="candidate2", bytes=PNG_BYTES, mime_type="image/png"),
        ],
    )
    asset.update_pipeline_metadata(preserve_all_variants_in_docx=True)

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [asset])
    updated_doc = Document(BytesIO(updated_bytes))

    assert len(updated_doc.paragraphs) == 3
    for updated_paragraph in updated_doc.paragraphs:
        paragraph_properties = updated_paragraph._element.pPr
        assert paragraph_properties is not None
        assert paragraph_properties.find(qn("w:pStyle")) is None
        assert paragraph_properties.find(qn("w:outlineLvl")) is None
        assert paragraph_properties.find(qn("w:keepLines")) is None
        alignment = paragraph_properties.find(qn("w:jc"))
        assert alignment is not None
        assert alignment.get(qn("w:val")) == "center"


def test_reinsert_inline_images_keeps_single_variant_placeholder_text_and_logs_when_inside_hyperlink_in_mixed_multi_variant_paragraph(monkeypatch):
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Lead ")
    _add_hyperlink(paragraph, "[[DOCX_IMAGE_img_001]]", "https://example.com")
    paragraph.add_run(" middle [[DOCX_IMAGE_img_002]] tail")
    buffer = BytesIO()
    doc.save(buffer)

    first_asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        width_emu=914400,
        height_emu=914400,
        final_variant="original",
    )
    second_asset = ImageAsset(
        image_id="img_002",
        placeholder="[[DOCX_IMAGE_img_002]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=1,
        safe_bytes=PNG_BYTES,
        attempt_variants=[
            ImageVariantCandidate(mode="candidate1", bytes=PNG_BYTES, mime_type="image/png"),
            ImageVariantCandidate(mode="candidate2", bytes=PNG_BYTES, mime_type="image/png"),
        ],
    )
    second_asset.update_pipeline_metadata(preserve_all_variants_in_docx=True)

    events = []
    monkeypatch.setattr(image_reinsertion, "log_event", lambda level, event, message, **context: events.append((event, context)))

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [first_asset, second_asset])
    updated_doc = Document(BytesIO(updated_bytes))
    visible_text = "\n".join(paragraph.text for paragraph in updated_doc.paragraphs)

    assert "[[DOCX_IMAGE_img_001]]" in visible_text
    assert "tail" in visible_text
    assert len(updated_doc.inline_shapes) == 3
    assert len(updated_doc._element.xpath(".//w:hyperlink")) == 1
    assert not any(
        event == "image_reinsertion_multi_variant_block_fallback_to_text"
        and context.get("reason") == "multi_variant_placeholder_inside_hyperlink_or_non_run_child"
        for event, context in events
    )


def test_reinsert_inline_images_logs_multi_variant_specific_warning_when_block_build_fails(monkeypatch):
    doc = Document()
    doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    buffer = BytesIO()
    doc.save(buffer)

    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        safe_bytes=PNG_BYTES,
        attempt_variants=[
            ImageVariantCandidate(mode="candidate1", bytes=PNG_BYTES, mime_type="image/png"),
            ImageVariantCandidate(mode="candidate2", bytes=PNG_BYTES, mime_type="image/png"),
        ],
    )
    asset.update_pipeline_metadata(preserve_all_variants_in_docx=True)

    events = []
    monkeypatch.setattr(image_reinsertion, "_build_replacement_blocks_from_fragments", lambda paragraph, fragments: [])
    monkeypatch.setattr(image_reinsertion, "log_event", lambda level, event, message, **context: events.append((event, context)))

    updated_bytes = reinsert_inline_images(buffer.getvalue(), [asset])
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].text == "[[DOCX_IMAGE_img_001]]"
    assert any(
        event == "image_reinsertion_multi_variant_block_unresolved"
        and context.get("reason") == "multi_variant_block_builder_returned_no_output"
        for event, context in events
    )


def test_resolve_final_image_bytes_prefers_selected_compare_variant():
    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=b"original",
        mime_type="image/png",
        position_index=0,
        safe_bytes=b"safe",
        redrawn_bytes=b"redrawn",
        final_variant="redrawn",
        comparison_variants={
            "semantic_redraw_direct": {"bytes": b"chosen"},
        },
        selected_compare_variant="semantic_redraw_direct",
    )

    assert resolve_final_image_bytes(asset) == b"chosen"
    assert asset.resolved_delivery_payload().selected_variant == "semantic_redraw_direct"
    assert asset.resolved_delivery_payload().final_bytes == b"chosen"


def test_resolve_final_image_bytes_returns_original_for_explicit_original_compare_choice():
    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=b"original",
        mime_type="image/png",
        position_index=0,
        safe_bytes=b"safe",
        redrawn_bytes=b"redrawn",
        final_variant="redrawn",
        comparison_variants={"semantic_redraw_direct": {"bytes": b"chosen"}},
        selected_compare_variant="original",
    )

    assert resolve_final_image_bytes(asset) == b"original"


def test_resolve_image_insertions_returns_all_compare_all_variants_before_single_final_choice():
    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=b"original",
        mime_type="image/png",
        position_index=0,
        validation_status="compared",
        comparison_variants={
            "safe": {"bytes": b"safe"},
            "semantic_redraw_direct": {"bytes": b"direct"},
            "semantic_redraw_structured": {"bytes": b"structured"},
        },
    )

    assert resolve_image_insertions(asset) == [
        ("Вариант 1: Просто улучшить", b"safe"),
        ("Вариант 2: Креативная AI-перерисовка", b"direct"),
        ("Вариант 3: Структурная AI-перерисовка", b"structured"),
    ]
    assert asset.resolved_delivery_payload().delivery_kind == "compare_all_variants"
    assert [insertion.variant_key for insertion in asset.resolved_delivery_payload().insertions] == [
        "safe",
        "semantic_redraw_direct",
        "semantic_redraw_structured",
    ]


def test_resolved_delivery_payload_uses_manual_review_insertions_when_enabled():
    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=b"original",
        mime_type="image/png",
        position_index=0,
        safe_bytes=b"safe",
        attempt_variants=[
            ImageVariantCandidate(mode="candidate1", bytes=b"candidate-1", mime_type="image/png"),
            ImageVariantCandidate(mode="candidate2", bytes=b"candidate-2", mime_type="image/png"),
        ],
    )
    asset.update_pipeline_metadata(preserve_all_variants_in_docx=True)

    payload = asset.resolved_delivery_payload()

    assert payload.delivery_kind == "manual_review_variants"
    assert [(insertion.label, insertion.bytes) for insertion in payload.insertions] == [
        ("safe", b"safe"),
        ("candidate1", b"candidate-1"),
        ("candidate2", b"candidate-2"),
    ]