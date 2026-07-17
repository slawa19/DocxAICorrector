from io import BytesIO
import json
from threading import Event, Thread
from types import SimpleNamespace
from typing import Any, cast

from docx import Document
import pytest

import docxaicorrector.document.segments as document_segments
import docxaicorrector.processing.preparation as preparation
import docxaicorrector.validation.structural as structural_validation
from docxaicorrector.core.config import ModelRegistry, TextModelConfig
from docxaicorrector.core.models import DocumentBlock
from docxaicorrector.core.models import ImageAsset, ImageVariantCandidate
from docxaicorrector.core.models import LayoutArtifactCleanupReport, ParagraphBoundaryNormalizationReport
from docxaicorrector.core.models import ParagraphRelation, ParagraphRelationDecision, ParagraphUnit, RelationNormalizationReport
from docxaicorrector.core.models import StructureRepairReport
from docxaicorrector.document.segments import CHAPTER_SEGMENTS_DETECTOR_VERSION
from docxaicorrector.processing.processing_runtime import FrozenUploadPayload, build_in_memory_uploaded_file, build_preparation_request_marker


def setup_function():
    preparation.clear_preparation_cache(clear_shared=True)


def _build_report(*, raw=0, logical=0, merged_groups=0, merged_raw=0):
    return ParagraphBoundaryNormalizationReport(
        total_raw_paragraphs=raw,
        total_logical_paragraphs=logical,
        merged_group_count=merged_groups,
        merged_raw_paragraph_count=merged_raw,
    )


def _build_cleanup_report(*, original=0, cleaned=0, removed=0, page_numbers=0, repeated=0):
    return LayoutArtifactCleanupReport(
        original_paragraph_count=original,
        cleaned_paragraph_count=cleaned,
        removed_paragraph_count=removed,
        removed_page_number_count=page_numbers,
        removed_repeated_artifact_count=repeated,
        removed_empty_or_whitespace_count=0,
        cleanup_applied=True,
    )


def _build_default_prepared_source_key(file_token: str, chunk_size: int = 6000) -> str:
    app_config = preparation.load_app_config()
    configured_relation_kinds = app_config.get("relation_normalization_enabled_relation_kinds", ())
    if not isinstance(configured_relation_kinds, (list, tuple, set)):
        configured_relation_kinds = ()
    relation_key = (
        f"{str(app_config.get('relation_normalization_profile', 'phase2_default') or 'phase2_default')}:"
        f"{','.join(sorted(str(kind) for kind in configured_relation_kinds))}"
    )
    return preparation.build_prepared_source_key(
        file_token,
        chunk_size,
        paragraph_boundary_normalization_mode=str(
            app_config.get("paragraph_boundary_normalization_mode", "high_only") or "high_only"
        ),
        paragraph_boundary_ai_review_mode=str(app_config.get("paragraph_boundary_ai_review_mode", "off") or "off"),
        relation_normalization_key=relation_key,
        layout_artifact_cleanup_key=preparation._resolve_layout_cleanup_cache_key(app_config),
    )


def test_flatten_layout_cleanup_metrics_includes_empty_paragraphs():
    metrics = preparation.flatten_layout_cleanup_metrics(
        LayoutArtifactCleanupReport(
            original_paragraph_count=5,
            cleaned_paragraph_count=2,
            removed_paragraph_count=3,
            removed_page_number_count=1,
            removed_repeated_artifact_count=1,
            removed_empty_or_whitespace_count=1,
            cleanup_applied=True,
        )
    )

    assert metrics == {
        "layout_cleanup_removed_count": 3,
        "layout_cleanup_page_number_count": 1,
        "layout_cleanup_repeated_artifact_count": 1,
        "layout_cleanup_empty_or_whitespace_count": 1,
    }


def test_flatten_layout_cleanup_metrics_uses_flagged_counts_for_signal_mode():
    metrics = preparation.flatten_layout_cleanup_metrics(
        LayoutArtifactCleanupReport(
            original_paragraph_count=5,
            cleaned_paragraph_count=5,
            removed_paragraph_count=0,
            removed_page_number_count=0,
            removed_repeated_artifact_count=0,
            removed_empty_or_whitespace_count=0,
            cleanup_applied=True,
            cleanup_mode="flag",
            flagged_page_number_count=1,
            flagged_repeated_artifact_count=1,
            flagged_empty_or_whitespace_count=1,
        )
    )

    assert metrics == {
        "layout_cleanup_removed_count": 3,
        "layout_cleanup_page_number_count": 1,
        "layout_cleanup_repeated_artifact_count": 1,
        "layout_cleanup_empty_or_whitespace_count": 1,
    }


def _build_extract_result(paragraphs, image_assets, report, relations=None, relation_report=None, cleanup_report=None):
    if cleanup_report is None:
        cleanup_report = _build_cleanup_report(original=len(paragraphs), cleaned=len(paragraphs))
    return (
        paragraphs,
        image_assets,
        report,
        ([] if relations is None else relations),
        relation_report,
        cleanup_report,
        StructureRepairReport(
            applied=False,
            repaired_bullet_items=0,
            repaired_numbered_items=0,
            bounded_toc_regions=0,
            toc_body_boundary_repairs=0,
            heading_candidates_from_toc=0,
            remaining_isolated_marker_count=0,
        ),
    )


def _build_docx_bytes(paragraphs: list[str]) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_uploaded_payload(
    filename: str,
    content_bytes: bytes,
    file_token: str,
    *,
    source_format: str = "docx",
    conversion_backend: str | None = None,
) -> FrozenUploadPayload:
    return FrozenUploadPayload(
        filename=filename,
        content_bytes=content_bytes,
        file_size=len(content_bytes),
        content_hash="test-hash",
        file_token=file_token,
        source_format=source_format,
        conversion_backend=conversion_backend,
    )


def _build_paragraph(*, source_index: int, text: str, role: str = "body", **kwargs) -> ParagraphUnit:
    return ParagraphUnit(text=text, role=role, source_index=source_index, logical_index=source_index, **kwargs)


def _build_runtime_model_registry(*, structure_recognition_model: str = "gpt-4o-mini") -> ModelRegistry:
    return ModelRegistry(
        text=TextModelConfig(default="gpt-5.4-mini", options=("gpt-5.4-mini",)),
        structure_recognition=structure_recognition_model,
        image_analysis="gpt-5.4-mini",
        image_validation="gpt-5.4-mini",
        image_reconstruction="gpt-5.4-mini",
        image_generation="gpt-image-1.5",
        image_edit="gpt-image-1.5",
        image_generation_vision="gpt-5.4-mini",
    )


def _make_ai_first_config(
    *,
    structure_recognition_model: str = "gpt-4o-mini",
    relation_kinds: tuple[str, ...] = ("epigraph_attribution", "image_caption", "table_caption", "toc_region"),
    **overrides,
):
    config = {
        "paragraph_boundary_normalization_enabled": True,
        "paragraph_boundary_normalization_mode": "high_only",
        "paragraph_boundary_ai_review_enabled": False,
        "paragraph_boundary_ai_review_mode": "off",
        "relation_normalization_enabled": True,
        "relation_normalization_profile": "phase2_default",
        "relation_normalization_enabled_relation_kinds": relation_kinds,
        "structure_recognition_mode": "always",
        "structure_recognition_enabled": True,
        "structure_recognition_model": structure_recognition_model,
        "structure_recognition_max_window_paragraphs": 1800,
        "structure_recognition_overlap_paragraphs": 50,
        "structure_recognition_timeout_seconds": 60,
        "structure_recognition_split_fallback_max_depth": 3,
        "structure_recognition_split_fallback_max_expansions": 8,
        "structure_recognition_min_confidence": "medium",
        "structure_recognition_cache_enabled": False,
        "structure_recognition_save_debug_artifacts": False,
        "structure_validation_enabled": True,
        "structure_validation_save_debug_artifacts": False,
        "structure_recovery_enabled": True,
        "structure_recovery_mode": "ai_first",
        "structure_recovery_document_map_enabled": True,
        "structure_recovery_document_map_model": "openrouter:test/document-map",
        "structure_recovery_document_map_timeout_seconds": 45,
        "structure_recovery_document_map_max_input_paragraphs": 200,
        "structure_recovery_document_map_max_input_tokens": 180000,
        "structure_recovery_document_map_preview_chars": 120,
        "structure_recovery_document_map_cache_enabled": False,
        "structure_recovery_document_map_save_debug_artifacts": False,
        "structure_recovery_anchored_classification_max_window_paragraphs": 3000,
        "structure_recovery_anchored_classification_overlap_paragraphs": 0,
        "structure_recovery_anchored_classification_preview_chars": 1500,
        "structure_recovery_anchored_classification_target_input_tokens": 180000,
        "structure_recovery_anchored_classification_min_confidence": "high",
        "structure_recovery_topology_projection_enabled": False,
        "structure_recovery_topology_projection_save_debug_artifacts": True,
        "structure_recovery_topology_projection_binding_splits_enabled": False,
        "structure_recovery_topology_projection_layout_signals_enabled": False,
        "structure_recovery_topology_projection_layout_signals_heading_ratio": 1.15,
        "structure_recovery_topology_projection_layout_signals_short_line_chars": 80,
        "structure_recovery_topology_projection_layout_signals_baseline_tolerance_pt": 0.25,
        "structure_recovery_topology_projection_layout_signals_min_tier_population": 2,
        "models": _build_runtime_model_registry(structure_recognition_model=structure_recognition_model),
    }
    config.update(overrides)
    if "models" not in overrides:
        resolved_model = str(config.get("structure_recognition_model", structure_recognition_model) or structure_recognition_model)
        config["models"] = _build_runtime_model_registry(structure_recognition_model=resolved_model)
    return config


def _stub_single_block_preparation_builders(monkeypatch, *, source_text: str, job_text: str = "block"):
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: source_text)

    def _fake_build_semantic_blocks(
        paragraphs,
        max_chars,
        relations=None,
        hard_boundary_paragraph_ids=None,
        structure_phase="post_ai_final",
    ):
        return ["block"]

    def _fake_build_editing_jobs(blocks, max_chars, processing_operation="edit", structure_phase="post_ai_final"):
        return [{"target_text": job_text, "target_chars": len(job_text), "context_chars": 0}]

    monkeypatch.setattr(preparation, "build_semantic_blocks", _fake_build_semantic_blocks)
    monkeypatch.setattr(preparation, "build_editing_jobs", _fake_build_editing_jobs)


def test_prepare_document_for_processing_uses_cache_for_identical_inputs(monkeypatch):
    calls = {"count": 0}
    session_state = {"preparation_cache": {}}
    progress_events = []

    def fake_extract(uploaded_file):
        calls["count"] += 1
        return _build_extract_result([], [], None)

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", fake_extract)
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: [])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [])

    source_bytes = b"docx-bytes"
    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", source_bytes, "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
        progress_callback=lambda **payload: progress_events.append(payload),
    )
    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", source_bytes, "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
        progress_callback=lambda **payload: progress_events.append(payload),
    )

    assert calls["count"] == 1
    assert progress_events[-1]["metrics"]["cached"] is True


def test_prepare_document_for_processing_passes_processing_operation_to_job_builder(monkeypatch):
    captured = {}

    monkeypatch.setattr(
        preparation,
        "extract_document_content_with_normalization_reports",
        lambda uploaded_file: _build_extract_result([_build_paragraph(source_index=0, text="Contents")], [], None),
    )
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "Contents")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block"])

    def fake_build_editing_jobs(blocks, max_chars, processing_operation="edit"):
        captured["processing_operation"] = processing_operation
        return [{"target_text": "Contents", "target_chars": 8, "context_chars": 0}]

    monkeypatch.setattr(preparation, "build_editing_jobs", fake_build_editing_jobs)

    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        processing_operation="translate",
        session_state={"preparation_cache": {}},
    )

    assert captured["processing_operation"] == "translate"


def test_prepare_document_for_processing_assigns_stable_job_ids(monkeypatch):
    monkeypatch.setattr(
        preparation,
        "extract_document_content_with_normalization_reports",
        lambda uploaded_file: _build_extract_result([_build_paragraph(source_index=0, text="Contents")], [], None),
    )
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "Contents")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block-1", "block-2"])
    monkeypatch.setattr(
        preparation,
        "build_editing_jobs",
        lambda blocks, max_chars, processing_operation="edit": [
            {"target_text": "First block", "target_chars": 11, "context_chars": 0},
            {"target_text": "Second block", "target_chars": 12, "context_chars": 0},
        ],
    )

    result = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state={"preparation_cache": {}},
    )

    assert [job["job_id"] for job in result.jobs] == ["job_0000", "job_0001"]


def test_prepare_document_for_processing_keeps_toc_jobs_operation_aware_across_cache_entries(monkeypatch):
    calls = {"count": 0}
    session_state = {"preparation_cache": {}}
    toc_paragraphs = [
        ParagraphUnit(text="Contents", role="body", structural_role="toc_header", source_index=0),
        ParagraphUnit(text="Chapter 1........ 12", role="body", structural_role="toc_entry", source_index=1),
    ]

    def fake_extract(uploaded_file):
        calls["count"] += 1
        return _build_extract_result(toc_paragraphs, [], None)

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", fake_extract)
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "\n\n".join(paragraph.text for paragraph in paragraphs))
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: [DocumentBlock(paragraphs=list(paragraphs))])

    edit_result = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        processing_operation="edit",
        session_state=session_state,
    )
    translate_result = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        processing_operation="translate",
        session_state=session_state,
    )

    assert calls["count"] == 2
    assert edit_result.jobs[0]["job_kind"] == "passthrough"
    assert translate_result.jobs[0]["job_kind"] == "llm"
    assert edit_result.jobs[0]["narration_include"] is False
    assert translate_result.jobs[0]["narration_include"] is False


def test_detect_document_segments_splits_oversized_heading_ranges_into_parent_and_children():
    paragraphs = [
        ParagraphUnit(text="Chapter 1", role="heading", heading_level=1, heading_source="explicit", paragraph_id="p0000", source_index=0),
        ParagraphUnit(text="A" * 12000, role="body", paragraph_id="p0001", source_index=1),
        ParagraphUnit(text="B" * 12000, role="body", paragraph_id="p0002", source_index=2),
        ParagraphUnit(text="C" * 12000, role="body", paragraph_id="p0003", source_index=3),
    ]

    segments, diagnostics, structure_fingerprint = preparation.detect_document_segments(
        paragraphs,
        source_content_hash16="abcd1234ef567890",
        chunk_size=6000,
    )

    assert structure_fingerprint
    assert [segment.title for segment in segments] == ["Chapter 1", "Chapter 1 Part 1", "Chapter 1 Part 2"]
    assert [segment.level for segment in segments] == [1, 2, 2]
    assert segments[0].parent_segment_id is None
    assert segments[1].parent_segment_id == segments[0].segment_id
    assert segments[2].parent_segment_id == segments[0].segment_id
    assert segments[0].start_paragraph_index == 0 and segments[0].end_paragraph_index == 0
    assert segments[1].start_paragraph_index == 1 and segments[1].end_paragraph_index == 2
    assert segments[2].start_paragraph_index == 3 and segments[2].end_paragraph_index == 3
    assert diagnostics.segment_count == 3
    assert diagnostics.low_confidence_count == 2
    assert "low_confidence_segments_present" in diagnostics.warnings


def test_detect_document_segments_adds_all_caps_typography_evidence_for_heading_candidates():
    paragraphs = [
        ParagraphUnit(text="APPENDIX", role="body", paragraph_id="p0000", source_index=0),
        ParagraphUnit(text="Body paragraph", role="body", paragraph_id="p0001", source_index=1),
    ]

    segments, diagnostics, structure_fingerprint = preparation.detect_document_segments(
        paragraphs,
        source_content_hash16="abcd1234ef567890",
        chunk_size=6000,
        structure_phase="pre_ai_diagnostic",
    )

    assert structure_fingerprint
    assert diagnostics.segment_count >= 1
    typography_evidence = [
        evidence
        for evidence in segments[0].boundary_evidence
        if evidence.source == "typography_fallback"
    ]
    assert typography_evidence
    assert typography_evidence[0].details["is_all_caps"] is True
    assert typography_evidence[0].details["is_bold"] is False


def test_detect_document_segments_post_ai_final_ignores_typography_only_heading_fallbacks():
    paragraphs = [
        ParagraphUnit(text="APPENDIX", role="body", paragraph_id="p0000", source_index=0),
        ParagraphUnit(text="Body paragraph", role="body", paragraph_id="p0001", source_index=1),
    ]

    segments, diagnostics, structure_fingerprint = preparation.detect_document_segments(
        paragraphs,
        source_content_hash16="abcd1234ef567890",
        chunk_size=6000,
        structure_phase="post_ai_final",
    )

    assert structure_fingerprint
    assert diagnostics.segment_count == 1
    assert segments[0].title == "Body Range 1"


def test_detect_document_segments_is_deterministic_for_same_input():
    paragraphs = [
        ParagraphUnit(text="Contents", role="body", structural_role="toc_header", paragraph_id="p0000", source_index=0),
        ParagraphUnit(text="1 Chapter 1", role="body", structural_role="toc_entry", paragraph_id="p0001", source_index=1),
        ParagraphUnit(text="Chapter 1", role="heading", heading_level=1, heading_source="explicit", paragraph_id="p0002", source_index=2),
        ParagraphUnit(text="First chapter body", role="body", paragraph_id="p0003", source_index=3),
        ParagraphUnit(text="Chapter 2", role="heading", heading_level=1, heading_source="explicit", paragraph_id="p0004", source_index=4),
        ParagraphUnit(text="Second chapter body", role="body", paragraph_id="p0005", source_index=5),
    ]

    first_segments, first_diagnostics, first_structure_fingerprint = preparation.detect_document_segments(
        paragraphs,
        source_content_hash16="abcd1234ef567890",
        chunk_size=6000,
    )
    second_segments, second_diagnostics, second_structure_fingerprint = preparation.detect_document_segments(
        paragraphs,
        source_content_hash16="abcd1234ef567890",
        chunk_size=6000,
    )

    assert first_structure_fingerprint == second_structure_fingerprint
    assert first_diagnostics == second_diagnostics
    assert [
        (
            segment.segment_id,
            segment.parent_segment_id,
            segment.boundary_fingerprint,
            segment.start_paragraph_id,
            segment.end_paragraph_id,
            segment.title,
        )
        for segment in first_segments
    ] == [
        (
            segment.segment_id,
            segment.parent_segment_id,
            segment.boundary_fingerprint,
            segment.start_paragraph_id,
            segment.end_paragraph_id,
            segment.title,
        )
        for segment in second_segments
    ]


def test_detect_document_segments_counts_hinted_toc_entries():
    paragraphs = [
        ParagraphUnit(text="Contents", role="body", structural_role="body", paragraph_id="p0000", source_index=0, heuristic_structural_role_hint="toc_header"),
        ParagraphUnit(text="1 Chapter 1", role="body", structural_role="body", paragraph_id="p0001", source_index=1, heuristic_structural_role_hint="toc_entry"),
        ParagraphUnit(text="Chapter 1", role="heading", heading_level=1, heading_source="explicit", paragraph_id="p0002", source_index=2),
        ParagraphUnit(text="First chapter body", role="body", paragraph_id="p0003", source_index=3),
    ]

    _segments, diagnostics, _structure_fingerprint = preparation.detect_document_segments(
        paragraphs,
        source_content_hash16="abcd1234ef567890",
        chunk_size=6000,
        structure_phase="pre_ai_diagnostic",
    )

    assert diagnostics.toc_entry_count == 2


def test_detect_document_segments_ignores_advisory_toc_hints_in_post_ai_final_phase():
    paragraphs = [
        ParagraphUnit(text="Contents", role="body", structural_role="body", paragraph_id="p0000", source_index=0, heuristic_structural_role_hint="toc_header"),
        ParagraphUnit(text="1 Chapter 1", role="body", structural_role="body", paragraph_id="p0001", source_index=1, heuristic_structural_role_hint="toc_entry"),
        ParagraphUnit(text="Chapter 1", role="heading", heading_level=1, heading_source="explicit", paragraph_id="p0002", source_index=2),
        ParagraphUnit(text="First chapter body", role="body", paragraph_id="p0003", source_index=3),
    ]

    _segments, diagnostics, _structure_fingerprint = preparation.detect_document_segments(
        paragraphs,
        source_content_hash16="abcd1234ef567890",
        chunk_size=6000,
        structure_phase="post_ai_final",
    )

    assert diagnostics.toc_entry_count == 0


def test_prepare_document_for_processing_returns_independent_copies():
    session_state = {"preparation_cache": {}}
    source_bytes = _build_docx_bytes(["Первый абзац.", "Второй абзац."])

    first = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", source_bytes, "report.docx:token"),
        chunk_size=6000,
        session_state=session_state,
    )
    second = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", source_bytes, "report.docx:token"),
        chunk_size=6000,
        session_state=session_state,
    )

    first.paragraphs[0].text = "Изменено"
    first.jobs[0]["target_text"] = "Изменено"

    assert second.paragraphs[0].text == "Первый абзац.\n\nВторой абзац." or second.paragraphs[0].text == "Первый абзац."
    assert str(second.jobs[0]["target_text"]).strip() != "Изменено"


def test_prepare_document_for_processing_clones_attempt_variants_independently(monkeypatch):
    session_state = {"preparation_cache": {}}
    source_bytes = b"docx-bytes"

    asset = ImageAsset(
        image_id="img-1",
        placeholder="[[DOCX_IMAGE_img-1]]",
        original_bytes=b"orig",
        mime_type="image/png",
        position_index=0,
        attempt_variants=[
            ImageVariantCandidate(
                mode="safe",
                validation_result={"score": 0.5},
                final_reason="initial",
            )
        ],
    )

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", lambda uploaded_file: _build_extract_result([{"text": "p"}], [asset], None))
    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", lambda uploaded_file: _build_extract_result([{"text": "p"}], [asset], None))
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "text")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block"])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [{"target_text": "text", "target_chars": 4, "context_chars": 0}])

    first = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", source_bytes, "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
    )
    second = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", source_bytes, "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
    )

    first.image_assets[0].attempt_variants[0].final_reason = "mutated"
    first.image_assets[0].attempt_variants[0].validation_result["score"] = 0.1

    assert len(second.image_assets[0].attempt_variants) == 1
    assert second.image_assets[0].attempt_variants[0].final_reason == "initial"
    assert second.image_assets[0].attempt_variants[0].validation_result == {"score": 0.5}


def test_prepare_document_for_processing_clones_comparison_variants_independently(monkeypatch):
    session_state = {"preparation_cache": {}}
    source_bytes = b"docx-bytes"

    asset = ImageAsset(
        image_id="img-1",
        placeholder="[[DOCX_IMAGE_img-1]]",
        original_bytes=b"orig",
        mime_type="image/png",
        position_index=0,
        comparison_variants={
            "safe": {
                "mode": "safe",
                "validation_result": {"score": 0.7},
                "final_reason": "initial",
            }
        },
    )

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", lambda uploaded_file: _build_extract_result([{"text": "p"}], [asset], None))
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "text")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block"])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [{"target_text": "text", "target_chars": 4, "context_chars": 0}])

    first = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", source_bytes, "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
    )
    second = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", source_bytes, "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
    )

    first.image_assets[0].comparison_variants["safe"]["final_reason"] = "mutated"
    first.image_assets[0].comparison_variants["safe"]["validation_result"]["score"] = 0.2

    assert second.image_assets[0].comparison_variants["safe"]["final_reason"] == "initial"
    assert second.image_assets[0].comparison_variants["safe"]["validation_result"] == {"score": 0.7}


def test_prepare_document_for_processing_limits_session_cache_size(monkeypatch):
    session_state = {"preparation_cache": {}}
    expected_keys = [
        _build_default_prepared_source_key("two:3:b"),
        _build_default_prepared_source_key("three:5:c"),
    ]

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", lambda uploaded_file: _build_extract_result([], [], None))
    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", lambda uploaded_file: _build_extract_result([], [], None))
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: [])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [])

    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("one.docx", b"one", "one:3:a"),
        chunk_size=6000,
        session_state=session_state,
    )
    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("two.docx", b"two", "two:3:b"),
        chunk_size=6000,
        session_state=session_state,
    )
    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("three.docx", b"three", "three:5:c"),
        chunk_size=6000,
        session_state=session_state,
    )

    assert list(session_state["preparation_cache"].keys()) == expected_keys


def test_prepare_document_for_processing_uses_shared_cache_without_session_state(monkeypatch):
    calls = {"count": 0}
    progress_events = []

    def fake_extract(uploaded_file):
        calls["count"] += 1
        return _build_extract_result([], [], None)

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", fake_extract)
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: [])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [])

    source_bytes = b"docx-bytes"
    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", source_bytes, "report.docx:10:hash"),
        chunk_size=6000,
        session_state=None,
        progress_callback=lambda **payload: progress_events.append(payload),
    )
    second = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", source_bytes, "report.docx:10:hash"),
        chunk_size=6000,
        session_state=None,
        progress_callback=lambda **payload: progress_events.append(payload),
    )

    assert calls["count"] == 1
    assert second.cached is True
    assert progress_events[-1]["metrics"]["cached"] is True


def test_prepare_document_for_processing_uses_single_flight_for_shared_cache(monkeypatch):
    calls = {"count": 0}
    extract_started = Event()
    release_extract = Event()
    second_finished = Event()
    results = {}

    def fake_extract(uploaded_file):
        calls["count"] += 1
        extract_started.set()
        assert release_extract.wait(timeout=5)
        return _build_extract_result([], [], None)

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", fake_extract)
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: [])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [])

    def worker(result_key: str, finished_event: Event):
        results[result_key] = preparation.prepare_document_for_processing(
            uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
            chunk_size=6000,
            session_state=None,
        )
        finished_event.set()

    first_finished = Event()
    first_thread = Thread(target=worker, args=("first", first_finished))
    second_thread = Thread(target=worker, args=("second", second_finished))

    first_thread.start()
    assert extract_started.wait(timeout=5)
    second_thread.start()

    assert not second_finished.wait(timeout=0.2)

    release_extract.set()
    first_thread.join(timeout=5)
    second_thread.join(timeout=5)

    assert not first_thread.is_alive()
    assert not second_thread.is_alive()
    assert calls["count"] == 1
    assert results["first"].cached is False
    assert results["second"].cached is True


def test_clear_preparation_cache_requires_explicit_shared_clear(monkeypatch):
    calls = {"count": 0}

    def fake_extract(uploaded_file):
        calls["count"] += 1
        return _build_extract_result([], [], None)

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", fake_extract)
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: [])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [])

    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state=None,
    )

    session_state = {"preparation_cache": {"stale": object()}}
    preparation.clear_preparation_cache(session_state=session_state)

    assert session_state["preparation_cache"] == {}

    cached_again = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state=None,
    )

    preparation.clear_preparation_cache(clear_shared=True)

    uncached_after_explicit_clear = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state=None,
    )

    assert calls["count"] == 2
    assert cached_again.cached is True
    assert uncached_after_explicit_clear.cached is False


def test_prepare_document_for_processing_miss_returns_clone_separate_from_cached_entry(monkeypatch):
    session_state = {"preparation_cache": {}}
    expected_key = _build_default_prepared_source_key("report.docx:10:hash")

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", lambda uploaded_file: _build_extract_result([{"text": "p"}], [{"image": b"x"}], None))
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "text")
    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", lambda uploaded_file: _build_extract_result([{"text": "p"}], [{"image": b"x"}], None))
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block"])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [{"target_text": "text", "target_chars": 4, "context_chars": 0}])

    result = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
    )

    cached_entry = session_state["preparation_cache"][expected_key]

    assert result is not cached_entry
    assert result.paragraphs is not cached_entry.paragraphs
    assert result.image_assets is not cached_entry.image_assets
    assert result.jobs is not cached_entry.jobs
    assert result.jobs[0] is not cached_entry.jobs[0]


def test_prepare_document_for_processing_reports_pdf_import_stage(monkeypatch):
    session_state = {"preparation_cache": {}}
    events = []

    monkeypatch.setattr(
        preparation,
        "extract_document_content_with_normalization_reports",
        lambda uploaded_file: _build_extract_result(["p1"], [], _build_report(raw=1, logical=1)),
    )
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "text-value")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block-a"])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [{"target_text": "a", "target_chars": 1, "context_chars": 0}])

    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload(
            "report.docx",
            b"docx-bytes",
            "report.pdf:10:hash",
            source_format="pdf",
            conversion_backend="libreoffice",
        ),
        chunk_size=6000,
        session_state=session_state,
        progress_callback=lambda **payload: events.append(payload),
    )

    assert events[0]["stage"] == "Разбор DOCX (из PDF)"
    assert events[0]["detail"] == "Извлекаю абзацы, встроенные изображения и структуру из сконвертированного DOCX."
    assert events[0]["metrics"]["source_format"] == "pdf"


def test_apply_first_block_composition_quality_gate_uses_hint_toc_only_in_pre_ai_diagnostic():
    first_block = DocumentBlock(
        paragraphs=[
            ParagraphUnit(text="Contents", role="body", structural_role="body", heuristic_structural_role_hint="toc_header", source_index=0),
            ParagraphUnit(text="Chapter 1........ 12", role="body", structural_role="body", heuristic_structural_role_hint="toc_entry", source_index=1),
            ParagraphUnit(text="Epigraph line", role="body", structural_role="epigraph", source_index=2),
            ParagraphUnit(text="Introduction", role="heading", structural_role="body", heading_source="heuristic", source_index=3),
        ]
    )

    diagnostic_status, diagnostic_reasons = preparation._apply_first_block_composition_quality_gate(
        blocks=[first_block],
        processing_operation="translate",
        quality_gate_status="pass",
        quality_gate_reasons=(),
        structure_phase="pre_ai_diagnostic",
    )
    final_status, final_reasons = preparation._apply_first_block_composition_quality_gate(
        blocks=[first_block],
        processing_operation="translate",
        quality_gate_status="pass",
        quality_gate_reasons=(),
        structure_phase="post_ai_final",
    )

    assert diagnostic_status == "warning"
    assert diagnostic_reasons == (
        "first_block_mixed_toc_and_epigraph",
        "first_block_mixed_toc_and_body_start",
    )
    assert final_status == "pass"
    assert final_reasons == ()


def test_apply_first_block_composition_quality_gate_preserves_final_binding_mixed_block_warnings():
    first_block = DocumentBlock(
        paragraphs=[
            ParagraphUnit(text="Contents", role="body", structural_role="toc_header", source_index=0),
            ParagraphUnit(text="Chapter 1........ 12", role="body", structural_role="toc_entry", source_index=1),
            ParagraphUnit(text="Epigraph line", role="body", structural_role="epigraph", source_index=2),
            ParagraphUnit(text="Introduction", role="heading", structural_role="body", heading_source="heuristic", source_index=3),
        ]
    )

    status, reasons = preparation._apply_first_block_composition_quality_gate(
        blocks=[first_block],
        processing_operation="translate",
        quality_gate_status="pass",
        quality_gate_reasons=(),
        structure_phase="post_ai_final",
    )

    assert status == "warning"
    assert reasons == (
        "first_block_mixed_toc_and_epigraph",
        "first_block_mixed_toc_and_body_start",
    )


def test_prepare_document_for_processing_retains_normalization_report_on_fresh_preparation(monkeypatch):
    session_state = {"preparation_cache": {}}
    report = _build_report(raw=4, logical=3, merged_groups=1, merged_raw=2)

    monkeypatch.setattr(
        preparation,
        "extract_document_content_with_normalization_reports",
        lambda uploaded_file: _build_extract_result([{"text": "p"}], [], report),
    )
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "text")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block"])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [{"target_text": "text", "target_chars": 4, "context_chars": 0}])

    result = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
    )

    assert result.normalization_report == report
    assert result.normalization_report is not report


def test_prepare_document_for_processing_retains_normalization_report_across_cache_hit_clone(monkeypatch):
    session_state = {"preparation_cache": {}}
    extract_calls = {"count": 0}

    def fake_extract(uploaded_file):
        extract_calls["count"] += 1
        return _build_extract_result([{"text": "p"}], [], _build_report(raw=4, logical=3, merged_groups=1, merged_raw=2))

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", fake_extract)
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "text")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block"])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [{"target_text": "text", "target_chars": 4, "context_chars": 0}])

    first = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
    )
    second = preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
    )

    assert first.normalization_report is not None
    first.normalization_report.merged_group_count = 99

    assert extract_calls["count"] == 1
    assert second.cached is True
    assert second.normalization_report is not None
    assert second.normalization_report.merged_group_count == 1
    assert second.normalization_report is not first.normalization_report


def test_prepare_document_for_processing_reports_cache_hit_normalization_metrics(monkeypatch):
    session_state = {"preparation_cache": {}}
    events = []

    monkeypatch.setattr(
        preparation,
        "extract_document_content_with_normalization_reports",
        lambda uploaded_file: _build_extract_result(["p1", "p2"], ["img"], _build_report(raw=3, logical=2, merged_groups=1, merged_raw=2)),
    )
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "text-value")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block-a", "block-b"])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [{"target_text": "a", "target_chars": 1, "context_chars": 0}, {"target_text": "b", "target_chars": 1, "context_chars": 0}])

    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
    )
    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
        progress_callback=lambda **payload: events.append(payload),
    )

    assert len(events) == 1
    assert events[0]["stage"] == "Подготовка документа"
    assert events[0]["metrics"]["cached"] is True
    assert events[0]["metrics"]["raw_paragraph_count"] == 3
    assert events[0]["metrics"]["logical_paragraph_count"] == 2
    assert events[0]["metrics"]["merged_group_count"] == 1
    assert events[0]["metrics"]["merged_raw_paragraph_count"] == 2


def test_prepare_document_for_processing_emits_heartbeat_during_extraction(monkeypatch):
    """Live-progress contract: while the blocking DOCX extraction call runs,
    `HeartbeatBeacon` must continue to emit progress events so the UI activity
    feed and progress bar do not freeze.
    """
    import time as _time

    import docxaicorrector.processing.processing_runtime as _processing_runtime

    events: list[dict] = []
    session_state = {"preparation_cache": {}}

    class _FastBeacon(_processing_runtime.HeartbeatBeacon):
        def __init__(self, *args, **kwargs):
            kwargs["interval_seconds"] = 0.05
            super().__init__(*args, **kwargs)

    def _slow_extract(uploaded_file, *, app_config=None):
        _time.sleep(0.3)
        return _build_extract_result(["p1"], [], _build_report(raw=1, logical=1))

    monkeypatch.setattr(preparation, "HeartbeatBeacon", _FastBeacon)
    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", _slow_extract)
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "text-value")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block-a"])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [{"target_text": "a", "target_chars": 1, "context_chars": 0}])

    preparation.prepare_document_for_processing(
        uploaded_payload=_build_uploaded_payload("report.docx", b"docx-bytes", "report.docx:10:hash"),
        chunk_size=6000,
        session_state=session_state,
        progress_callback=lambda **payload: events.append(payload),
    )

    initial_stage = "Разбор DOCX"
    heartbeat_events = [e for e in events if e.get("stage") == initial_stage and "сек идёт чтение" in (e.get("detail") or "")]
    # Extraction takes ~0.3s with 0.05s heartbeat interval — expect at least 2 ticks.
    assert len(heartbeat_events) >= 2, [e.get("detail") for e in events]
    # Progress value during heartbeat should match the wired-in 0.22 anchor.
    assert all(abs(float(e["progress"]) - 0.22) < 1e-6 for e in heartbeat_events)


def test_build_prepared_source_key_distinguishes_target_language():
    # F14: two keys differing ONLY by target_language must be distinct so a run with a
    # different target language cannot serve another run's cached glossary/context.
    key_ru = preparation.build_prepared_source_key(
        "token",
        6000,
        target_language="ru",
        processing_operation="edit",
        paragraph_boundary_normalization_mode="high_only",
        paragraph_boundary_ai_review_mode="off",
        source_language="en",
        translation_domain="general",
        structure_recovery_enabled=False,
        structure_recovery_mode="legacy",
    )
    key_de = preparation.build_prepared_source_key(
        "token",
        6000,
        target_language="de",
        processing_operation="edit",
        paragraph_boundary_normalization_mode="high_only",
        paragraph_boundary_ai_review_mode="off",
        source_language="en",
        translation_domain="general",
        structure_recovery_enabled=False,
        structure_recovery_mode="legacy",
    )
    assert key_ru != key_de


def test_build_prepared_source_key_distinguishes_translation_domain():
    # F14: keys differing ONLY by translation_domain must be distinct.
    key_general = preparation.build_prepared_source_key(
        "token", 6000, translation_domain="general", source_language="en", target_language="ru"
    )
    key_legal = preparation.build_prepared_source_key(
        "token", 6000, translation_domain="legal", source_language="en", target_language="ru"
    )
    assert key_general != key_legal


def test_build_prepared_source_key_distinguishes_structure_recovery_mode():
    # F14: keys differing ONLY by structure-recovery mode must be distinct.
    key_legacy = preparation.build_prepared_source_key(
        "token", 6000, structure_recovery_enabled=True, structure_recovery_mode="legacy"
    )
    key_ai_first = preparation.build_prepared_source_key(
        "token", 6000, structure_recovery_enabled=True, structure_recovery_mode="ai_first"
    )
    assert key_legacy != key_ai_first


def test_build_prepared_source_key_identical_settings_collide():
    # F14: identical settings must still produce the SAME key (cache hit preserved).
    key_a = preparation.build_prepared_source_key(
        "token",
        6000,
        target_language="ru",
        source_language="en",
        translation_domain="general",
        structure_recovery_enabled=True,
        structure_recovery_mode="ai_first",
    )
    key_b = preparation.build_prepared_source_key(
        "token",
        6000,
        target_language="ru",
        source_language="en",
        translation_domain="general",
        structure_recovery_enabled=True,
        structure_recovery_mode="ai_first",
    )
    assert key_a == key_b


def test_build_prepared_source_key_distinguishes_ai_review_model():
    # F10: with AI review ON, two keys differing ONLY by the AI-review model must be
    # distinct — a different structure-recognition model shapes a different cached
    # AI-review artifact, so it must not be served from another model's prepared entry.
    key_model_a = preparation.build_prepared_source_key(
        "token",
        6000,
        paragraph_boundary_ai_review_mode="review_only",
        paragraph_boundary_ai_review_model="gpt-model-a",
        paragraph_boundary_ai_review_candidate_limit=200,
        paragraph_boundary_ai_review_timeout_seconds=30,
        paragraph_boundary_ai_review_max_tokens_per_candidate=120,
    )
    key_model_b = preparation.build_prepared_source_key(
        "token",
        6000,
        paragraph_boundary_ai_review_mode="review_only",
        paragraph_boundary_ai_review_model="gpt-model-b",
        paragraph_boundary_ai_review_candidate_limit=200,
        paragraph_boundary_ai_review_timeout_seconds=30,
        paragraph_boundary_ai_review_max_tokens_per_candidate=120,
    )
    assert key_model_a != key_model_b


def test_build_prepared_source_key_distinguishes_ai_review_candidate_limit():
    # F10: with AI review ON, two keys differing ONLY by candidate-limit must be distinct.
    key_limit_200 = preparation.build_prepared_source_key(
        "token",
        6000,
        paragraph_boundary_ai_review_mode="review_only",
        paragraph_boundary_ai_review_model="gpt-model-a",
        paragraph_boundary_ai_review_candidate_limit=200,
        paragraph_boundary_ai_review_timeout_seconds=30,
        paragraph_boundary_ai_review_max_tokens_per_candidate=120,
    )
    key_limit_500 = preparation.build_prepared_source_key(
        "token",
        6000,
        paragraph_boundary_ai_review_mode="review_only",
        paragraph_boundary_ai_review_model="gpt-model-a",
        paragraph_boundary_ai_review_candidate_limit=500,
        paragraph_boundary_ai_review_timeout_seconds=30,
        paragraph_boundary_ai_review_max_tokens_per_candidate=120,
    )
    assert key_limit_200 != key_limit_500


def test_build_prepared_source_key_ai_review_off_ignores_model_and_limits():
    # F10: when AI review is OFF the model/limits do not shape any cached artifact, so
    # keys that differ only by those knobs must still COLLIDE (no needless invalidation).
    key_off_a = preparation.build_prepared_source_key(
        "token",
        6000,
        paragraph_boundary_ai_review_mode="off",
        paragraph_boundary_ai_review_model="gpt-model-a",
        paragraph_boundary_ai_review_candidate_limit=200,
        paragraph_boundary_ai_review_timeout_seconds=30,
        paragraph_boundary_ai_review_max_tokens_per_candidate=120,
    )
    key_off_b = preparation.build_prepared_source_key(
        "token",
        6000,
        paragraph_boundary_ai_review_mode="off",
        paragraph_boundary_ai_review_model="gpt-model-b",
        paragraph_boundary_ai_review_candidate_limit=999,
        paragraph_boundary_ai_review_timeout_seconds=99,
        paragraph_boundary_ai_review_max_tokens_per_candidate=999,
    )
    assert key_off_a == key_off_b


def test_build_prepared_source_key_identical_ai_review_settings_collide():
    # F10: identical AI-review settings still produce the SAME key (cache hit preserved).
    key_a = preparation.build_prepared_source_key(
        "token",
        6000,
        paragraph_boundary_ai_review_mode="review_only",
        paragraph_boundary_ai_review_model="gpt-model-a",
        paragraph_boundary_ai_review_candidate_limit=200,
        paragraph_boundary_ai_review_timeout_seconds=30,
        paragraph_boundary_ai_review_max_tokens_per_candidate=120,
    )
    key_b = preparation.build_prepared_source_key(
        "token",
        6000,
        paragraph_boundary_ai_review_mode="review_only",
        paragraph_boundary_ai_review_model="gpt-model-a",
        paragraph_boundary_ai_review_candidate_limit=200,
        paragraph_boundary_ai_review_timeout_seconds=30,
        paragraph_boundary_ai_review_max_tokens_per_candidate=120,
    )
    assert key_a == key_b


def test_resolve_scan_origin_cache_key_defaults_and_folds_overrides():
    # F2: the default resolves to the general conservative thresholds, and an
    # app_config override changes the resolved fingerprint.
    default_key = preparation._resolve_scan_origin_cache_key(None)
    assert default_key == "10:0.1:1.5"
    assert preparation._DEFAULT_SCAN_ORIGIN_CACHE_KEY == default_key
    tuned_key = preparation._resolve_scan_origin_cache_key(
        {"scan_origin_multi_column_absolute_min": 25}
    )
    assert tuned_key == "25:0.1:1.5"
    assert tuned_key != default_key


def test_build_prepared_source_key_distinguishes_scan_origin_threshold():
    # F2: two keys differing ONLY by a scan-origin threshold must be distinct — the
    # thresholds change which tables are flattened, shaping different prepared structure.
    key_default = preparation.build_prepared_source_key(
        "token", 6000, scan_origin_key="10:0.1:1.5"
    )
    key_tuned = preparation.build_prepared_source_key(
        "token", 6000, scan_origin_key="20:0.1:1.5"
    )
    assert key_default != key_tuned


def test_build_prepared_source_key_identical_scan_origin_collide():
    # F2: identical scan-origin thresholds still produce the SAME key (cache hit preserved).
    key_a = preparation.build_prepared_source_key("token", 6000, scan_origin_key="10:0.1:1.5")
    key_b = preparation.build_prepared_source_key("token", 6000, scan_origin_key="10:0.1:1.5")
    assert key_a == key_b


def test_build_editing_jobs_adapter_propagates_internal_type_error_without_retry():
    # F25: a target that accepts the signature-checked kwargs but raises TypeError
    # DEEP inside must propagate and be invoked exactly once (no swallow/retry).
    calls = {"count": 0}

    def flaky_build_editing_jobs(blocks, *, max_chars, processing_operation="edit", structure_phase="pre_ai_diagnostic"):
        calls["count"] += 1
        raise TypeError("internal boom, not a signature mismatch")

    original = preparation.build_editing_jobs
    preparation.build_editing_jobs = flaky_build_editing_jobs
    try:
        with pytest.raises(TypeError, match="internal boom"):
            preparation._build_editing_jobs_with_optional_operation(
                blocks=["b"],
                max_chars=6000,
                processing_operation="edit",
                structure_phase="pre_ai_diagnostic",
            )
    finally:
        preparation.build_editing_jobs = original
    assert calls["count"] == 1


def test_build_semantic_blocks_adapter_propagates_internal_type_error_without_retry():
    # F25: same contract for the semantic-blocks adapter.
    calls = {"count": 0}

    def flaky_build_semantic_blocks(paragraphs, *, max_chars, relations, hard_boundary_paragraph_ids=None, structure_phase="pre_ai_diagnostic"):
        calls["count"] += 1
        raise TypeError("internal boom, not a signature mismatch")

    original = preparation.build_semantic_blocks
    preparation.build_semantic_blocks = flaky_build_semantic_blocks
    try:
        with pytest.raises(TypeError, match="internal boom"):
            preparation._build_semantic_blocks_with_optional_boundaries(
                paragraphs=["p"],
                max_chars=6000,
                relations=[],
                hard_boundary_paragraph_ids=set(),
                structure_phase="pre_ai_diagnostic",
            )
    finally:
        preparation.build_semantic_blocks = original
    assert calls["count"] == 1


def test_detect_segments_adapter_propagates_internal_type_error_without_retry():
    # F25: same contract for the segment-detection adapter.
    calls = {"count": 0}

    def flaky_detect_document_segments(paragraphs, *, source_content_hash16, chunk_size, structure_phase="pre_ai_diagnostic"):
        calls["count"] += 1
        raise TypeError("internal boom, not a signature mismatch")

    original = preparation.detect_document_segments
    preparation.detect_document_segments = flaky_detect_document_segments
    try:
        with pytest.raises(TypeError, match="internal boom"):
            preparation._detect_document_segments_with_optional_phase(
                paragraphs=["p"],
                source_content_hash16="abc123",
                chunk_size=6000,
                structure_phase="pre_ai_diagnostic",
            )
    finally:
        preparation.detect_document_segments = original
    assert calls["count"] == 1


# --- Spec 040: preparation cache client/credential (tenant) identity ---------------------

# Byte-for-byte snapshot of ``build_prepared_source_key('tok123', 1200)`` captured from the
# pre-040 implementation. The client-identity axis MUST be a no-op when empty: an empty
# identity may never change this string (no cache invalidation, single-tenant unchanged).
_PRE_040_REPRESENTATIVE_KEY = (
    "tok123:1200:high_only:off:phase2_default:"
    "epigraph_attribution,image_caption,table_caption,toc_region:lc=1:3:80:pv=2:pk=4:"
    "sl=en:tl=ru:td=general:sr=0:srm=legacy:so=10:0.1:1.5:ar=off"
)


def test_build_prepared_source_key_client_identity_empty_is_byte_identical():
    # Anti-regression: default (no client_identity) == explicit "" == the pre-040 output.
    default_key = preparation.build_prepared_source_key("tok123", 1200)
    empty_identity_key = preparation.build_prepared_source_key("tok123", 1200, client_identity="")
    assert default_key == empty_identity_key
    assert default_key == _PRE_040_REPRESENTATIVE_KEY
    assert ":cid=" not in default_key


def test_build_prepared_source_key_appends_cid_segment_when_identity_present():
    base_key = preparation.build_prepared_source_key("tok123", 1200)
    keyed = preparation.build_prepared_source_key("tok123", 1200, client_identity="abcd1234abcd1234")
    # Non-empty identity appends exactly one stable ``:cid=<identity>`` segment; the rest of
    # the key is untouched, so keyed == base + suffix.
    assert keyed == f"{base_key}:cid=abcd1234abcd1234"
    assert keyed != base_key
    # Different identity -> different key; same identity -> same key.
    other = preparation.build_prepared_source_key("tok123", 1200, client_identity="ffff0000ffff0000")
    assert other != keyed
    same = preparation.build_prepared_source_key("tok123", 1200, client_identity="abcd1234abcd1234")
    assert same == keyed


def test_resolve_prepared_cache_client_identity_off_returns_empty(monkeypatch):
    # AI review OFF -> the prepared document is NOT client-dependent -> identity "" no matter
    # what credentials are in the environment (sharing preserved, key byte-identical).
    monkeypatch.setenv("OPENAI_API_KEY", "sk-whatever-value")
    cfg = preparation.load_app_config()
    identity = preparation._resolve_prepared_cache_client_identity(
        resolved_config=cfg,
        ai_review_effective_enabled=False,
        ai_review_model="gpt-4o-mini",
    )
    assert identity == ""


def test_resolve_prepared_cache_client_identity_differs_by_env_secret(monkeypatch):
    # With AI review ON, two runs whose ONLY difference is os.environ[api_key_env] must
    # produce DIFFERENT 16-hex identities; an unchanged secret produces the SAME identity.
    cfg = preparation.load_app_config()

    def _identity() -> str:
        return preparation._resolve_prepared_cache_client_identity(
            resolved_config=cfg,
            ai_review_effective_enabled=True,
            ai_review_model="gpt-4o-mini",
        )

    monkeypatch.setenv("OPENAI_API_KEY", "sk-tenant-A")
    identity_a = _identity()
    monkeypatch.setenv("OPENAI_API_KEY", "sk-tenant-B")
    identity_b = _identity()
    monkeypatch.setenv("OPENAI_API_KEY", "sk-tenant-A")
    identity_a_again = _identity()

    assert len(identity_a) == 16 and all(c in "0123456789abcdef" for c in identity_a)
    assert len(identity_b) == 16 and all(c in "0123456789abcdef" for c in identity_b)
    assert identity_a != identity_b
    assert identity_a == identity_a_again

    # And the identity flows through into the full cache key.
    key_a = preparation.build_prepared_source_key(
        "token", 6000, paragraph_boundary_ai_review_mode="review_only", client_identity=identity_a
    )
    key_b = preparation.build_prepared_source_key(
        "token", 6000, paragraph_boundary_ai_review_mode="review_only", client_identity=identity_b
    )
    assert key_a != key_b


def test_resolve_prepared_cache_client_identity_never_leaks_secret(monkeypatch):
    # Secret-safety: the raw api-key value must NEVER appear in the identity or the key,
    # only its sha256. Also assert the env NAME's value cannot be recovered by substring.
    secret = "SUPER-SECRET-KEY-VALUE-1234567890"
    monkeypatch.setenv("OPENAI_API_KEY", secret)
    cfg = preparation.load_app_config()
    identity = preparation._resolve_prepared_cache_client_identity(
        resolved_config=cfg,
        ai_review_effective_enabled=True,
        ai_review_model="gpt-4o-mini",
    )
    assert identity != ""
    assert secret not in identity
    key = preparation.build_prepared_source_key(
        "token", 6000, paragraph_boundary_ai_review_mode="review_only", client_identity=identity
    )
    assert secret not in key


def test_resolve_prepared_cache_client_identity_fails_open_on_bad_selector(monkeypatch):
    # Fail-open: an unresolvable selector (unknown provider) must return "" rather than
    # raise — cache-key construction can never blow up the request.
    monkeypatch.setenv("OPENAI_API_KEY", "sk-anything")
    cfg = preparation.load_app_config()
    identity = preparation._resolve_prepared_cache_client_identity(
        resolved_config=cfg,
        ai_review_effective_enabled=True,
        ai_review_model="bogusprovider:some-model",
    )
    assert identity == ""


def _build_sentinel_prepared_document(marker: str) -> "preparation.PreparedDocumentData":
    return preparation.PreparedDocumentData(
        source_text=marker,
        paragraphs=[],
        image_assets=[],
        relations=[],
        jobs=[],
        prepared_source_key="",
    )


def test_prepared_cache_no_cross_credential_bleed(monkeypatch):
    # End-to-end (shared cache): prime the shared cache under identity A (review ON), then a
    # second run with the SAME token/settings but a DIFFERENT api-key value must NOT read
    # identity A's prepared document; a THIRD run back on identity A DOES hit the cache.
    # setup_function() already cleared the shared cache before this test.
    cfg = preparation.load_app_config()

    def _identity() -> str:
        return preparation._resolve_prepared_cache_client_identity(
            resolved_config=cfg,
            ai_review_effective_enabled=True,
            ai_review_model="gpt-4o-mini",
        )

    def _key(identity: str) -> str:
        return preparation.build_prepared_source_key(
            "token-shared",
            6000,
            paragraph_boundary_ai_review_mode="review_only",
            paragraph_boundary_ai_review_model="gpt-4o-mini",
            client_identity=identity,
        )

    # Identity A primes the shared cache.
    monkeypatch.setenv("OPENAI_API_KEY", "sk-tenant-A")
    identity_a = _identity()
    key_a = _key(identity_a)
    preparation._store_cached_prepared_document(
        session_state=None,
        prepared_source_key=key_a,
        prepared_document=_build_sentinel_prepared_document("tenant-A-document"),
    )

    # Identity B: different credential -> different key -> shared-cache MISS (reserves inflight).
    monkeypatch.setenv("OPENAI_API_KEY", "sk-tenant-B")
    identity_b = _identity()
    key_b = _key(identity_b)
    assert identity_b != identity_a
    assert key_b != key_a
    cached_b, in_flight_b, _level_b = preparation._read_or_reserve_cached_prepared_document(
        session_state=None,
        prepared_source_key=key_b,
    )
    assert cached_b is None  # no cross-credential bleed
    assert in_flight_b is not None
    preparation._release_shared_preparation(key_b)  # tidy up the reservation

    # Identity A again: same credential -> same key -> shared-cache HIT with A's document.
    monkeypatch.setenv("OPENAI_API_KEY", "sk-tenant-A")
    identity_a_again = _identity()
    assert identity_a_again == identity_a
    cached_a, in_flight_a, level_a = preparation._read_or_reserve_cached_prepared_document(
        session_state=None,
        prepared_source_key=_key(identity_a_again),
    )
    assert in_flight_a is None
    assert cached_a is not None
    assert level_a == "shared"
    assert cached_a.source_text == "tenant-A-document"


def _install_counting_prepared_builder(monkeypatch) -> dict[str, int]:
    # Spec 041 P1-1 tests drive prepare_document_for_processing itself but stub the heavy
    # pipeline: each real build increments the counter and returns a fresh sentinel so a cache
    # HIT (no build) is distinguishable from a MISS (rebuild) by the source_text marker.
    builds = {"count": 0}

    def _fake_build(*_args, **_kwargs):
        builds["count"] += 1
        return _build_sentinel_prepared_document(f"prepared-doc-{builds['count']}")

    monkeypatch.setattr(preparation, "_prepare_document_for_processing", _fake_build)
    return builds


def test_prepare_document_shared_cache_isolates_injected_tenant_identity(monkeypatch):
    # Injected factory + AI review ON + explicit client_cache_identity: distinct tenant
    # identities MUST NOT share one shared-cache entry even with the SAME token/app_config.
    preparation.clear_preparation_cache(clear_shared=True)
    config = _make_ai_first_config(
        paragraph_boundary_ai_review_enabled=True,
        paragraph_boundary_ai_review_mode="review_only",
    )
    builds = _install_counting_prepared_builder(monkeypatch)
    payload = _build_uploaded_payload("report.docx", b"docx-bytes", "shared-token")

    def factory_a(*_a, **_k):
        return object()

    def factory_b(*_a, **_k):
        return object()

    # Tenant A primes the shared cache.
    result_a1 = preparation.prepare_document_for_processing(
        uploaded_payload=payload,
        chunk_size=6000,
        app_config=config,
        session_state=None,
        client_factory=factory_a,
        client_cache_identity="idA",
    )
    assert builds["count"] == 1

    # Tenant B: different identity, SAME token/app_config -> MISS (must not receive A's doc).
    result_b1 = preparation.prepare_document_for_processing(
        uploaded_payload=payload,
        chunk_size=6000,
        app_config=config,
        session_state=None,
        client_factory=factory_b,
        client_cache_identity="idB",
    )
    assert builds["count"] == 2
    assert result_b1.source_text != result_a1.source_text

    # Tenant A again: same identity -> shared-cache HIT (no rebuild), serves A's document.
    result_a2 = preparation.prepare_document_for_processing(
        uploaded_payload=payload,
        chunk_size=6000,
        app_config=config,
        session_state=None,
        client_factory=factory_a,
        client_cache_identity="idA",
    )
    assert builds["count"] == 2
    assert result_a2.source_text == result_a1.source_text
    assert result_a2.cached is True


def test_prepare_document_bypasses_shared_cache_when_injected_identity_unknown(monkeypatch):
    # Injected factory + AI review ON + NO identity: the shared (process-global) tier must be
    # bypassed entirely, so a second identical run does not serve the first run's entry.
    preparation.clear_preparation_cache(clear_shared=True)
    config = _make_ai_first_config(
        paragraph_boundary_ai_review_enabled=True,
        paragraph_boundary_ai_review_mode="review_only",
    )
    builds = _install_counting_prepared_builder(monkeypatch)
    payload = _build_uploaded_payload("report.docx", b"docx-bytes", "shared-token")

    def factory(*_a, **_k):
        return object()

    first = preparation.prepare_document_for_processing(
        uploaded_payload=payload,
        chunk_size=6000,
        app_config=config,
        session_state=None,
        client_factory=factory,
        client_cache_identity=None,
    )
    assert builds["count"] == 1

    second = preparation.prepare_document_for_processing(
        uploaded_payload=payload,
        chunk_size=6000,
        app_config=config,
        session_state=None,
        client_factory=factory,
        client_cache_identity=None,
    )
    # No shared hit: each run rebuilds and nothing is published to the shared tier.
    assert builds["count"] == 2
    assert second.cached is False
    assert second.source_text != first.source_text
    assert len(preparation._shared_preparation_cache) == 0


def test_prepare_document_config_path_uses_shared_cache_with_config_identity(monkeypatch):
    # Regression: client_factory=None (config-default path) with AI review ON must keep using
    # the config-derived identity + shared cache exactly as spec 040 (byte-identical key).
    preparation.clear_preparation_cache(clear_shared=True)
    monkeypatch.setenv("OPENAI_API_KEY", "sk-config-tenant")
    config = dict(preparation.load_app_config())
    config["paragraph_boundary_ai_review_enabled"] = True
    config["paragraph_boundary_ai_review_mode"] = "review_only"
    builds = _install_counting_prepared_builder(monkeypatch)
    payload = _build_uploaded_payload("report.docx", b"docx-bytes", "config-token")

    first = preparation.prepare_document_for_processing(
        uploaded_payload=payload,
        chunk_size=6000,
        app_config=config,
        session_state=None,
        client_factory=None,
    )
    assert builds["count"] == 1

    second = preparation.prepare_document_for_processing(
        uploaded_payload=payload,
        chunk_size=6000,
        app_config=config,
        session_state=None,
        client_factory=None,
    )
    # Shared cache HIT: no rebuild, and the stored key folds the config-derived identity.
    assert builds["count"] == 1
    assert second.cached is True
    assert first.source_text == second.source_text

    (
        _enabled,
        _mode,
        _candidate_limit,
        _timeout_seconds,
        _max_tokens,
        ai_review_model,
    ) = preparation.resolve_paragraph_boundary_ai_review_settings(
        allowed_modes=preparation.PARAGRAPH_BOUNDARY_AI_REVIEW_MODE_VALUES,
        app_config=config,
    )
    config_identity = preparation._resolve_prepared_cache_client_identity(
        resolved_config=config,
        ai_review_effective_enabled=True,
        ai_review_model=ai_review_model,
    )
    assert config_identity != ""
    assert len(preparation._shared_preparation_cache) == 1
    stored_key = next(iter(preparation._shared_preparation_cache))
    assert f":cid={config_identity}" in stored_key
