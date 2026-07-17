"""Counter-proof tests for spec 020 — preserve authored tables; flatten scan-origin.

Two behaviors are pinned here:

1. Part 1 — a KEPT table is emitted as a Pandoc-markdown table and therefore
   survives the render as a real Word ``w:tbl`` (raw ``<table>`` HTML is dropped
   by Pandoc; markdown tables are not).
2. Part 2 — a document classified as scan-origin (OCR) has its tables flattened
   into linear body paragraphs (no ``w:tbl``), with cell text conserved, while
   authored documents keep their tables. The classifier is pinned on the
   measured provenance signals (RESISTANCE=129 multi-col vs authored 0-2).
"""

import re
import zipfile
from io import BytesIO
from pathlib import Path

import pytest

import docxaicorrector.generation._generation as generation
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

from docxaicorrector.core.models import RawParagraph, RawTable
from docxaicorrector.document.extraction import (
    _build_raw_document_blocks,
    build_document_text,
    extract_document_content_from_docx,
)
from docxaicorrector.document.provenance import (
    classify_document_scan_origin,
    classify_scan_origin_from_document_xml,
    table_has_authored_signals,
)


BOOK_SOURCES = Path(__file__).resolve().parent / "sources" / "book"
RESISTANCE = BOOK_SOURCES / "RESISTANCE FACTORS AND SPECIAL FORCES AREAS UKRAINE.docx"
MAZZUCATO = BOOK_SOURCES / "The Value of Everything. Making and Taking in the Global Economy by Mariana Mazzucato (z-lib.org).docx"
LIETAER = BOOK_SOURCES / "Rethinking-money_-How-new-currencies-turn-scarcity-into-prosperity-Bernard-Lietaer-Jacqui-Dunne.docx"


def _pandoc_available() -> bool:
    generation.ensure_pandoc_available.cache_clear()
    try:
        generation.ensure_pandoc_available()
        return True
    except Exception:
        return False
    finally:
        generation.ensure_pandoc_available.cache_clear()


def _count_word_tables(docx_bytes: bytes) -> int:
    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8", "ignore")
    return len(re.findall(r"<w:tbl\b", document_xml))


def _output_all_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _strip_ws(text: str) -> str:
    return re.sub(r"\s+", "", text)


def _synthetic_document_xml(*, num_sections: int, multi_column_sections: int) -> str:
    sections = []
    for index in range(num_sections):
        num = 2 if index < multi_column_sections else 1
        sections.append(f'<w:p><w:pPr><w:sectPr><w:cols w:num="{num}"/></w:sectPr></w:pPr></w:p>')
    body = "".join(sections)
    return f'<w:document xmlns:w="ns"><w:body>{body}<w:sectPr><w:cols w:num="1"/></w:sectPr></w:body></w:document>'


# --------------------------------------------------------------------------- #
# Required counter-proof 1: scan-origin classifier (threshold pinned on corpus) #
# --------------------------------------------------------------------------- #


def test_scan_origin_classifier_pins_corpus_provenance():
    resistance = classify_document_scan_origin(RESISTANCE.read_bytes())
    mazzucato = classify_document_scan_origin(MAZZUCATO.read_bytes())
    lietaer = classify_document_scan_origin(LIETAER.read_bytes())

    assert resistance.is_scan_origin is True
    assert resistance.multi_column_section_count == 129

    assert mazzucato.is_scan_origin is False
    assert mazzucato.multi_column_section_count == 0

    assert lietaer.is_scan_origin is False
    assert lietaer.multi_column_section_count == 2


def test_scan_origin_classifier_does_not_over_trigger_on_authored_two_column_layout():
    # Anti-vacuum: a normal authored document with a couple of two-column
    # sections must stay authored (bias to authored; don't flatten magazines).
    authored_xml = _synthetic_document_xml(num_sections=10, multi_column_sections=3)
    classification = classify_scan_origin_from_document_xml(authored_xml)

    assert classification.multi_column_section_count == 3
    assert classification.is_scan_origin is False


def test_scan_origin_classifier_flags_high_multi_column_density():
    scan_xml = _synthetic_document_xml(num_sections=200, multi_column_sections=120)
    classification = classify_scan_origin_from_document_xml(scan_xml)

    assert classification.multi_column_section_count == 120
    assert classification.is_scan_origin is True


# --------------------------------------------------------------------------- #
# Required counter-proof 2: scan tables flattened + conservation (RESISTANCE)  #
# --------------------------------------------------------------------------- #


def test_scan_origin_document_flattens_tables_and_conserves_cell_text():
    source_bytes = RESISTANCE.read_bytes()
    paragraphs, _ = extract_document_content_from_docx(BytesIO(source_bytes))

    # No table survives extraction for a scan-origin document.
    assert all(paragraph.role != "table" for paragraph in paragraphs)

    body_stream = _strip_ws(build_document_text(paragraphs))
    source_document = Document(BytesIO(source_bytes))

    missing: list[str] = []
    checked = 0
    for table in source_document.tables:
        for row in table.rows:
            seen: set[int] = set()
            for cell in row.cells:
                identity = id(cell._tc)
                if identity in seen:
                    continue
                seen.add(identity)
                cell_text = _strip_ws(cell.text)
                if not cell_text:
                    continue
                checked += 1
                if cell_text not in body_stream:
                    missing.append(cell.text)

    assert checked > 100  # sanity: the fixture really has substantial tabular data
    # Whitespace-insensitive conservation: at most a negligible tail of OCR cells
    # whose surrounding punctuation (wrapping quotes) reflowed; the token content
    # is preserved. Pin a hard, near-total conservation bound.
    assert len(missing) <= 1, f"table cell text lost on flatten: {missing[:10]}"


# --------------------------------------------------------------------------- #
# Required counter-proof 3: authored tables survive the render as real w:tbl   #
# --------------------------------------------------------------------------- #


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_raw_html_table_is_dropped_but_markdown_table_survives_render():
    raw_html_table = "<table>\n<tbody>\n<tr><td>A</td><td>B</td></tr>\n<tr><td>1</td><td>2</td></tr>\n</tbody>\n</table>"
    markdown_table = "|  |  |\n| --- | --- |\n| A | B |\n| 1 | 2 |"

    assert _count_word_tables(generation.convert_markdown_to_docx_bytes(raw_html_table)) == 0
    assert _count_word_tables(generation.convert_markdown_to_docx_bytes(markdown_table)) >= 1


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_authored_document_table_renders_as_real_word_table():
    document = Document()
    document.add_paragraph("Перед таблицей")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Колонка A"
    table.cell(0, 1).text = "Колонка B"
    table.cell(1, 0).text = "Значение 1"
    table.cell(1, 1).text = "Значение 2"
    document.add_paragraph("После таблицы")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)
    assert [paragraph.role for paragraph in paragraphs] == ["body", "table", "body"]

    markdown = build_document_text(paragraphs)
    output = generation.convert_markdown_to_docx_bytes(markdown)

    assert _count_word_tables(output) >= 1
    output_text = _output_all_text(output)
    for value in ("Колонка A", "Колонка B", "Значение 1", "Значение 2"):
        assert value in output_text


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_mazzucato_authored_tables_survive_render_as_word_tables():
    paragraphs, _ = extract_document_content_from_docx(BytesIO(MAZZUCATO.read_bytes()))
    table_paragraphs = [paragraph for paragraph in paragraphs if paragraph.role == "table"]
    assert len(table_paragraphs) == 3

    markdown = build_document_text(paragraphs)
    output = generation.convert_markdown_to_docx_bytes(markdown)

    assert _count_word_tables(output) >= 3


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_resistance_scan_tables_do_not_render_word_tables():
    paragraphs, _ = extract_document_content_from_docx(BytesIO(RESISTANCE.read_bytes()))
    markdown = build_document_text(paragraphs)
    output = generation.convert_markdown_to_docx_bytes(markdown)

    assert _count_word_tables(output) == 0


# --------------------------------------------------------------------------- #
# F13: per-table authored-signal override for the scan-origin prior           #
# --------------------------------------------------------------------------- #

_W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _tbl_xml(*, borders: bool, grid_widths: list[int], border_val: str = "single") -> object:
    border_edges = ""
    if borders:
        border_edges = (
            "<w:tblBorders>"
            + "".join(
                f'<w:{edge} w:val="{border_val}" w:sz="4" w:space="0" w:color="000000"/>'
                for edge in ("top", "left", "bottom", "right", "insideH", "insideV")
            )
            + "</w:tblBorders>"
        )
    grid = "".join(f'<w:gridCol w:w="{width}"/>' for width in grid_widths)
    return parse_xml(
        f"<w:tbl {_W_NS}><w:tblPr>{border_edges}</w:tblPr><w:tblGrid>{grid}</w:tblGrid></w:tbl>"
    )


def test_table_has_authored_signals_distinguishes_authored_from_scan_shapes():
    # Real table-level borders alone are an authored signal (even if columns vary).
    assert table_has_authored_signals(_tbl_xml(borders=True, grid_widths=[500, 2500, 900])) is True
    # A uniform multi-column grid alone is an authored signal (even without borders).
    assert table_has_authored_signals(_tbl_xml(borders=False, grid_widths=[2000, 2000])) is True
    # Borderless + irregular widths == the OCR column-region shape -> flatten.
    assert table_has_authored_signals(_tbl_xml(borders=False, grid_widths=[500, 2500, 900])) is False
    # A tblBorders element whose edges are all "nil" is NOT a real border.
    assert (
        table_has_authored_signals(
            _tbl_xml(borders=True, grid_widths=[500, 2500, 900], border_val="nil")
        )
        is False
    )


def test_table_has_authored_signals_treats_real_scan_tables_as_non_authored():
    # Anti-regression: the actual RESISTANCE scan tables must keep flattening.
    source_document = Document(BytesIO(RESISTANCE.read_bytes()))
    assert source_document.tables  # sanity
    for table in source_document.tables:
        assert table_has_authored_signals(table._tbl) is False


def _add_bordered_uniform_table(document, values: list[list[str]], *, widths: list[int]) -> None:
    table = document.add_table(rows=len(values), cols=len(values[0]))
    for row_index, row_values in enumerate(values):
        for col_index, value in enumerate(row_values):
            table.cell(row_index, col_index).text = value
    grid = table._tbl.find(qn("w:tblGrid"))
    for grid_col, width in zip(grid.findall(qn("w:gridCol")), widths):
        grid_col.set(qn("w:w"), str(width))
    table_properties = table._tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_element = OxmlElement(f"w:{edge}")
        edge_element.set(qn("w:val"), "single")
        edge_element.set(qn("w:sz"), "4")
        borders.append(edge_element)
    table_properties.append(borders)


def _add_borderless_irregular_table(document, values: list[list[str]], *, widths: list[int]) -> None:
    table = document.add_table(rows=len(values), cols=len(values[0]))
    for row_index, row_values in enumerate(values):
        for col_index, value in enumerate(row_values):
            table.cell(row_index, col_index).text = value
    grid = table._tbl.find(qn("w:tblGrid"))
    for grid_col, width in zip(grid.findall(qn("w:gridCol")), widths):
        grid_col.set(qn("w:w"), str(width))


def test_scan_origin_preserves_authored_table_but_flattens_scan_table():
    # Inside a document treated as scan-origin, the per-table override keeps a
    # genuinely authored table (real borders + uniform grid) as a real table
    # while still flattening a borderless, irregular scan-shape table.
    document = Document()
    document.add_paragraph("Перед таблицами")
    _add_bordered_uniform_table(
        document,
        [["Автор A", "Автор B"], ["Значение 1", "Значение 2"]],
        widths=[2000, 2000],
    )
    document.add_paragraph("Между таблицами")
    _add_borderless_irregular_table(
        document,
        [["Скан ячейка A", "Скан ячейка B", "Скан ячейка C"]],
        widths=[500, 2500, 900],
    )

    raw_blocks, _ = _build_raw_document_blocks(document, is_scan_origin=True)

    table_blocks = [block for block in raw_blocks if isinstance(block, RawTable)]
    flattened_blocks = [
        block
        for block in raw_blocks
        if isinstance(block, RawParagraph) and block.layout_origin == "table_flattened"
    ]

    # The authored table survives as one RawTable carrying its cell text.
    assert len(table_blocks) == 1
    assert "Автор A" in table_blocks[0].html_text
    assert "Значение 2" in table_blocks[0].html_text

    # The scan-shape table is flattened into linear body lines.
    flattened_text = "\n".join(block.text for block in flattened_blocks)
    for value in ("Скан ячейка A", "Скан ячейка B", "Скан ячейка C"):
        assert value in flattened_text
    # The authored cells are NOT in the flattened stream (they stayed a table).
    assert "Автор A" not in flattened_text
