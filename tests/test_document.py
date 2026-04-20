import base64
import json
import os
import zipfile
from io import BytesIO
from pathlib import Path

import pytest

import document
import formatting_transfer
import formatting_diagnostics_retention
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from document import (
    build_document_text,
    extract_document_content_from_docx,
)
from formatting_transfer import (
    normalize_semantic_output_docx,
    preserve_source_paragraph_properties,
)
from models import ParagraphUnit


PNG_BYTES = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aK3cAAAAASUVORK5CYII=")
def _extract_numbering_ids(paragraph) -> tuple[str | None, str | None]:
    paragraph_properties = paragraph._element.pPr
    if paragraph_properties is None:
        return None, None
    num_pr = paragraph_properties.find(qn("w:numPr"))
    if num_pr is None:
        return None, None
    ilvl = num_pr.find(qn("w:ilvl"))
    num_id = num_pr.find(qn("w:numId"))
    return (
        None if ilvl is None else ilvl.get(qn("w:val")),
        None if num_id is None else num_id.get(qn("w:val")),
    )


def _numbering_root_contains_num_id(document, num_id: str) -> bool:
    numbering_root = document.part.numbering_part.element
    for child in numbering_root:
        if child.tag == qn("w:num") and child.get(qn("w:numId")) == num_id:
            abstract_num = child.find(qn("w:abstractNumId"))
            if abstract_num is None:
                return False
            abstract_num_id = abstract_num.get(qn("w:val"))
            return any(
                candidate.tag == qn("w:abstractNum") and candidate.get(qn("w:abstractNumId")) == abstract_num_id
                for candidate in numbering_root
            )
    return False


def _append_decimal_numbering_definition(document, *, num_id: str, abstract_num_id: str) -> None:
    numbering_root = document.part.numbering_part.element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_num_id)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    abstract_num.append(lvl)
    numbering_root.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_num_ref = OxmlElement("w:abstractNumId")
    abstract_num_ref.set(qn("w:val"), abstract_num_id)
    num.append(abstract_num_ref)
    numbering_root.append(num)


def _attach_numbering(paragraph, *, num_id: str, ilvl: str) -> None:
    paragraph_properties = paragraph._element.get_or_add_pPr()
    num_pr = paragraph_properties.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        paragraph_properties.append(num_pr)

    ilvl_element = num_pr.find(qn("w:ilvl"))
    if ilvl_element is None:
        ilvl_element = OxmlElement("w:ilvl")
        num_pr.append(ilvl_element)
    ilvl_element.set(qn("w:val"), ilvl)

    num_id_element = num_pr.find(qn("w:numId"))
    if num_id_element is None:
        num_id_element = OxmlElement("w:numId")
        num_pr.append(num_id_element)
    num_id_element.set(qn("w:val"), num_id)


def test_preserve_source_paragraph_properties_keeps_existing_heading_semantics_in_target_docx():
    source_paragraphs = [ParagraphUnit(text="Заголовок", role="heading", heading_level=1)]

    target_doc = Document()
    target_doc.add_paragraph("Заголовок", style="Heading 2")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Heading 2"


def test_preserve_source_paragraph_properties_logs_mismatch_warning(monkeypatch):
    source_doc = Document()
    source_paragraph = source_doc.add_paragraph("Абзац")
    source_paragraph.paragraph_format.left_indent = Inches(0.5)
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("Абзац")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    events = []
    monkeypatch.setattr(formatting_transfer, "log_event", lambda level, event, message, **context: events.append((event, context)))

    preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)

    assert len(events) == 1
    event_name, context = events[0]
    assert event_name == "paragraph_count_mismatch_preserve"
    assert context["source_count"] == 1
    assert context["target_count"] == 2
    assert context["mapped_count"] == 1
    assert context["unmapped_source_count"] == 0
    assert context["unmapped_target_count"] == 1
    assert isinstance(context["artifact_path"], str)


def test_preserve_source_paragraph_properties_does_not_replay_raw_xml_on_mismatch():
    source_doc = Document()
    source_paragraph = source_doc.add_paragraph("Абзац")
    source_paragraph.paragraph_format.left_indent = Inches(0.5)
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("Абзац")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))
    first_paragraph_properties = updated_doc.paragraphs[0]._element.pPr
    first_indentation = None if first_paragraph_properties is None else first_paragraph_properties.find(qn("w:ind"))

    assert first_indentation is None


def test_preserve_source_paragraph_properties_artifact_records_caption_heading_conflict(tmp_path, monkeypatch):
    image_path = tmp_path / "artifact_caption_image.png"
    image_path.write_bytes(PNG_BYTES)

    source_doc = Document()
    source_doc.add_paragraph().add_run().add_picture(str(image_path))
    source_caption = source_doc.add_paragraph("Рис. 1. Подпись к изображению")
    source_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target_doc.add_paragraph("Рис. 1. Подпись к изображению", style="Heading 1")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    diagnostics_dir = tmp_path / "formatting_diagnostics"
    monkeypatch.setattr(formatting_transfer, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)

    preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)

    artifacts = sorted(diagnostics_dir.glob("*.json"))
    assert len(artifacts) == 1
    payload = json.loads(artifacts[0].read_text(encoding="utf-8"))
    assert len(payload["caption_heading_conflicts"]) == 1
    assert payload["caption_heading_conflicts"][0]["target_style_name"] == "Heading 1"
    assert payload["caption_heading_conflicts"][0]["target_heading_level"] == 1


def test_preserve_source_paragraph_properties_artifact_records_restored_list_decisions_during_mismatch(tmp_path, monkeypatch):
    source_doc = Document()
    source_doc.add_paragraph("Первый пункт", style="List Number")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("Первый пункт")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    diagnostics_dir = tmp_path / "formatting_diagnostics"
    monkeypatch.setattr(formatting_transfer, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)

    artifacts = sorted(diagnostics_dir.glob("*.json"))
    assert len(artifacts) == 1
    payload = json.loads(artifacts[0].read_text(encoding="utf-8"))
    assert len(payload["list_restoration_decisions"]) == 1
    assert payload["list_restoration_decisions"][0]["action"] == "restored"


def test_prune_formatting_diagnostics_removes_oldest_and_preserves_newest(tmp_path):
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    diagnostics_dir.mkdir()
    timestamps = [1000, 2000, 3000, 4000]
    paths = []
    for timestamp in timestamps:
        path = diagnostics_dir / f"restore_{timestamp}.json"
        path.write_text("{}", encoding="utf-8")
        paths.append(path)

    for offset, path in enumerate(paths, start=1):
        mtime = float(offset)
        path.touch()
        os.utime(path, (mtime, mtime))

    pruned = formatting_diagnostics_retention.prune_formatting_diagnostics(
        diagnostics_dir=diagnostics_dir,
        now_epoch_seconds=10.0,
        max_age_seconds=100,
        max_count=2,
    )

    remaining = sorted(path.name for path in diagnostics_dir.glob("*.json"))
    assert remaining == ["restore_3000.json", "restore_4000.json"]
    assert sorted(Path(path).name for path in pruned) == ["restore_1000.json", "restore_2000.json"]


def test_write_formatting_diagnostics_artifact_prunes_expired_runtime_files_only(tmp_path):
    runtime_dir = tmp_path / ".run" / "formatting_diagnostics"
    tests_dir = tmp_path / "tests" / "artifacts" / "formatting_diagnostics"
    runtime_dir.mkdir(parents=True)
    tests_dir.mkdir(parents=True)
    old_runtime = runtime_dir / "restore_old.json"
    old_runtime.write_text("{}", encoding="utf-8")
    test_artifact = tests_dir / "keep.json"
    test_artifact.write_text("{}", encoding="utf-8")

    import os
    os.utime(old_runtime, (1.0, 1.0))

    artifact_path = formatting_diagnostics_retention.write_formatting_diagnostics_artifact(
        stage="restore",
        diagnostics={"mapped_count": 1},
        diagnostics_dir=runtime_dir,
        now_epoch_ms=200_000,
    )

    assert artifact_path is not None
    assert sorted(path.name for path in runtime_dir.glob("*.json")) == [Path(artifact_path).name]
    assert test_artifact.exists()


def test_validate_docx_archive_rejects_zip_slip_paths():
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("../evil.txt", "boom")

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "подозрительные пути" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for zip-slip path")


def test_validate_docx_archive_rejects_oversized_archive_bytes(monkeypatch):
    monkeypatch.setattr(document, "MAX_DOCX_ARCHIVE_SIZE_BYTES", 5)

    try:
        document._validate_docx_archive(b"123456")
    except RuntimeError as exc:
        assert "превышает допустимый размер архива" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for oversized DOCX archive")


def test_validate_docx_archive_rejects_bad_zip_payload():
    try:
        document._validate_docx_archive(b"not-a-zip")
    except RuntimeError as exc:
        assert "поврежденный или неподдерживаемый DOCX-архив" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for invalid DOCX archive")


def test_validate_docx_archive_rejects_empty_archive():
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w"):
        pass

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "пустой DOCX-архив" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for empty DOCX archive")


def test_validate_docx_archive_rejects_too_many_entries(monkeypatch):
    monkeypatch.setattr(document, "MAX_DOCX_ENTRY_COUNT", 1)
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", "<w:document/>")

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "слишком много файлов" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for excessive DOCX entry count")


def test_validate_docx_archive_rejects_suspicious_compression_ratio(monkeypatch):
    monkeypatch.setattr(document, "MAX_DOCX_COMPRESSION_RATIO", 1)
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "A" * 2048)
        archive.writestr("word/document.xml", "B" * 2048)

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "подозрительно высокий коэффициент сжатия" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for suspicious DOCX compression ratio")


def test_validate_docx_archive_rejects_absolute_paths():
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("/word/document.xml", "boom")

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "абсолютные пути" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for absolute path in DOCX archive")


def test_validate_docx_archive_rejects_missing_content_types():
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("word/document.xml", "<w:document/>")

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "отсутствует [Content_Types].xml" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for missing [Content_Types].xml")


def test_preserve_source_paragraph_properties_applies_minimal_output_formatting():
    source_paragraphs = [
        ParagraphUnit(text="Глава", role="heading", heading_level=1),
        ParagraphUnit(text="[[DOCX_IMAGE_img_001]]", role="image"),
        ParagraphUnit(text="Рис. 1. Подпись", role="caption"),
        ParagraphUnit(text="Обычный текст", role="body"),
    ]
    target_doc = Document()
    target_doc.add_paragraph("Глава")
    target_doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target_doc.add_paragraph("Рис. 1. Подпись")
    target_doc.add_paragraph("Обычный текст")
    table = target_doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "A"
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[2].style is not None
    assert updated_doc.paragraphs[3].style is not None
    assert updated_doc.tables[0].style is not None

    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert updated_doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert updated_doc.paragraphs[2].style.name == "Caption"
    assert updated_doc.paragraphs[3].style.name in {"Body Text", "Normal"}
    assert updated_doc.tables[0].style.name == "Table Grid"


def test_preserve_source_paragraph_properties_applies_partial_transfer_on_semantic_mismatch():
    source_paragraphs = [ParagraphUnit(text="Заголовок", role="heading", heading_level=1)]
    target_doc = Document()
    target_doc.add_paragraph("Заголовок")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_uses_content_heuristics_for_captions_without_mapping():
    source_paragraphs = [ParagraphUnit(text="Рис. 1. Подпись к изображению", role="caption")]
    target_doc = Document()
    target_doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target_doc.add_paragraph("Рисунок 1 Подпись к изображению")
    target_doc.add_paragraph("Посторонний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Caption"
    assert updated_doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_preserve_source_paragraph_properties_restores_direct_center_alignment_for_mapped_paragraphs():
    source_paragraphs = [
        ParagraphUnit(text="ГЛАВА 1", role="heading", heading_level=1, paragraph_alignment="center", paragraph_id="p0000"),
        ParagraphUnit(text="Богатство заключается не в том, чтобы иметь много имущества.", role="body", paragraph_alignment="center", paragraph_id="p0001"),
    ]
    target_doc = Document()
    target_doc.add_paragraph("ГЛАВА 1")
    target_doc.add_paragraph("Богатство заключается не в том, чтобы иметь много имущества.")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert updated_doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_preserve_source_paragraph_properties_does_not_promote_generated_registry_text_to_heading():
    source_paragraphs = [ParagraphUnit(text="Старый заголовок", role="heading", heading_level=1, paragraph_id="p0000")]
    target_doc = Document()
    target_doc.add_paragraph("Совершенно новый заголовок")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(
        target_buffer.getvalue(),
        source_paragraphs,
        generated_paragraph_registry=[{"paragraph_id": "p0000", "text": "Совершенно новый заголовок"}],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_does_not_apply_body_formatting_via_generated_registry_similarity():
    source_paragraphs = [ParagraphUnit(text="Исходный абзац сильно отличается", role="body", paragraph_id="p0010")]
    target_doc = Document()
    target_doc.add_paragraph("Лишний абзац перед целью")
    target_doc.add_paragraph("Итоговый литературно отредактированный абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(
        target_buffer.getvalue(),
        source_paragraphs,
        generated_paragraph_registry=[
            {"paragraph_id": "p0010", "text": "Итоговый литературно отредактированный абзац"}
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_leaves_ambiguous_generated_registry_targets_unchanged():
    source_paragraphs = [
        ParagraphUnit(text="Исходный абзац сильно отличается", role="body", paragraph_id="p0011"),
    ]
    target_doc = Document()
    target_doc.add_paragraph("Отредактированный абзац с важной мыслью о богатстве и сообществе сегодня")
    target_doc.add_paragraph("Отредактированный абзац с важной мыслью о богатстве и сообществе завтра")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(
        target_buffer.getvalue(),
        source_paragraphs,
        generated_paragraph_registry=[
            {
                "paragraph_id": "p0011",
                "text": "Отредактированный абзац с важной мыслью о богатстве и сообществе",
            }
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_does_not_use_split_generated_registry_mapping_for_body():
    source_paragraphs = [ParagraphUnit(text="Старый слитый абзац", role="body", paragraph_id="p0056")]
    target_doc = Document()
    target_doc.add_paragraph("Новый заголовок", style="Heading 3")
    target_doc.add_paragraph("Текст после нового заголовка")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(
        target_buffer.getvalue(),
        source_paragraphs,
        generated_paragraph_registry=[
            {"paragraph_id": "p0056", "text": "### Новый заголовок\nТекст после нового заголовка"}
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Heading 3"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_does_not_restyle_reordered_paragraphs():
    source_paragraphs = [
        ParagraphUnit(text="Заголовок", role="heading", heading_level=1),
        ParagraphUnit(text="Обычный текст", role="body"),
    ]
    target_doc = Document()
    target_doc.add_paragraph("Обычный текст")
    target_doc.add_paragraph("Заголовок")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name in {"Body Text", "Normal"}
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_keeps_existing_numbered_list_semantics_without_injecting_source_numbering():
    source_doc = Document()
    source_doc.add_paragraph("Первый пункт", style="List Number")
    source_doc.add_paragraph("Второй пункт", style="List Number")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)

    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    _append_decimal_numbering_definition(target_doc, num_id="77", abstract_num_id="700")
    first = target_doc.add_paragraph("Первый пункт")
    second = target_doc.add_paragraph("Второй пункт")
    _attach_numbering(first, num_id="77", ilvl="0")
    _attach_numbering(second, num_id="77", ilvl="0")
    original_first = _extract_numbering_ids(first)
    original_second = _extract_numbering_ids(second)
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    first_ilvl, first_num_id = _extract_numbering_ids(updated_doc.paragraphs[0])
    second_ilvl, second_num_id = _extract_numbering_ids(updated_doc.paragraphs[1])

    assert (first_ilvl, first_num_id) == original_first
    assert (second_ilvl, second_num_id) == original_second
    assert first_num_id is not None
    assert _numbering_root_contains_num_id(updated_doc, first_num_id)


def test_preserve_source_paragraph_properties_restores_numbering_for_mapped_plain_target_paragraphs_despite_mismatch():
    source_doc = Document()
    source_doc.add_paragraph("Первый пункт", style="List Number")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)

    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("Первый пункт")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    ilvl, num_id = _extract_numbering_ids(updated_doc.paragraphs[0])

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert ilvl == "0"
    assert num_id is not None
    assert _numbering_root_contains_num_id(updated_doc, num_id)
    assert _extract_numbering_ids(updated_doc.paragraphs[1]) == (None, None)


def test_normalize_semantic_output_docx_remains_noop():
    target_doc = Document()
    paragraph = target_doc.add_paragraph("Обычный текст")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    docx_bytes = target_buffer.getvalue()
    assert normalize_semantic_output_docx(docx_bytes, [ParagraphUnit(text="Обычный текст", role="body")]) == docx_bytes


def test_caption_survives_extraction_markdown_and_preserve_after_image(tmp_path):
    image_path = tmp_path / "docx_caption_pipeline_image.png"
    image_path.write_bytes(PNG_BYTES)

    source_doc = Document()
    source_doc.add_paragraph().add_run().add_picture(str(image_path))
    caption = source_doc.add_paragraph(style="Caption")
    caption.add_run("Рисунок 1 Образец подписи").bold = True
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)

    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)
    markdown = build_document_text(source_paragraphs)

    target_doc = Document()
    target_doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target_doc.add_paragraph("Рисунок 1 Образец подписи")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert [paragraph.role for paragraph in source_paragraphs] == ["image", "caption"]
    assert markdown == "[[DOCX_IMAGE_img_001]]\n\n**Рисунок 1 Образец подписи**"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Caption"
    assert updated_doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_extract_document_content_from_docx_rejects_suspicious_uncompressed_archive(monkeypatch):
    monkeypatch.setattr(document, "MAX_DOCX_UNCOMPRESSED_SIZE_BYTES", 100)

    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "x" * 150)
        archive.writestr("word/document.xml", "<w:document />")
    buffer.seek(0)

    try:
        extract_document_content_from_docx(buffer)
    except RuntimeError as exc:
        assert "слишком велик после распаковки" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for suspiciously large uncompressed DOCX archive")


def test_read_uploaded_docx_bytes_preserves_original_cause(monkeypatch):
    failing_error = ValueError("bad upload")
    monkeypatch.setattr(document, "read_uploaded_file_bytes", lambda uploaded_file: (_ for _ in ()).throw(failing_error))

    try:
        document._read_uploaded_docx_bytes(object())
    except ValueError as exc:
        assert "Не удалось прочитать содержимое DOCX-файла" in str(exc)
        assert exc.__cause__ is failing_error
    else:
        raise AssertionError("Expected ValueError when uploaded DOCX bytes cannot be read")


def test_read_uploaded_docx_bytes_reuses_existing_docx_bytes_without_renormalizing(monkeypatch):
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w") as archive:
        archive.writestr("word/document.xml", "<w:document />")
    docx_bytes = buffer.getvalue()

    monkeypatch.setattr(document, "read_uploaded_file_bytes", lambda uploaded_file: docx_bytes)
    assert document._read_uploaded_docx_bytes(object()) == docx_bytes


def test_read_uploaded_docx_bytes_rejects_non_normalized_non_docx_input(monkeypatch):
    monkeypatch.setattr(document, "read_uploaded_file_bytes", lambda uploaded_file: b"legacy-binary")
    monkeypatch.setattr(document, "resolve_uploaded_filename", lambda uploaded_file: "legacy.doc")

    with pytest.raises(ValueError, match="Ожидался уже нормализованный DOCX-архив"):
        document._read_uploaded_docx_bytes(object())
