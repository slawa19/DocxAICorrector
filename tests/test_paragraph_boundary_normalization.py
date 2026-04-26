from io import BytesIO
import base64
import json

import config
from docx import Document as make_document
import document as document_module
import document_boundary_review as boundary_review_module
from models import (
    ParagraphBoundaryDecision,
    ParagraphBoundaryNormalizationReport,
    ParagraphRelationDecision,
    ParagraphUnit,
    RawParagraph,
    RelationNormalizationReport,
)

from document import (
    build_document_text,
    build_marker_wrapped_block_text,
    build_paragraph_relations,
    extract_document_content_with_normalization_reports,
    extract_document_content_with_boundary_report,
)


PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aK3cAAAAASUVORK5CYII="
)


def _save_document(document) -> BytesIO:
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def test_false_body_boundary_is_merged_into_one_logical_paragraph():
    document = make_document()
    document.add_paragraph("архетипами: повторяющимися моделями")
    document.add_paragraph("поведения во времени, наблюдаемыми в разных системах.")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert len(paragraphs) == 1
    assert paragraphs[0].text == "архетипами: повторяющимися моделями поведения во времени, наблюдаемыми в разных системах."
    assert paragraphs[0].origin_raw_indexes == [0, 1]
    assert paragraphs[0].origin_raw_texts == [
        "архетипами: повторяющимися моделями",
        "поведения во времени, наблюдаемыми в разных системах.",
    ]
    assert paragraphs[0].boundary_source == "normalized_merge"
    assert report.merged_group_count == 1
    assert report.merged_raw_paragraph_count == 2
    assert build_document_text(paragraphs) == paragraphs[0].text
    assert build_marker_wrapped_block_text(paragraphs) == f"[[DOCX_PARA_{paragraphs[0].paragraph_id}]]\n{paragraphs[0].text}"


def test_terminal_punctuation_preserves_real_paragraph_boundary():
    document = make_document()
    document.add_paragraph("Первый абзац закончен.")
    document.add_paragraph("Второй абзац начинается с новой мысли.")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert [paragraph.text for paragraph in paragraphs] == [
        "Первый абзац закончен.",
        "Второй абзац начинается с новой мысли.",
    ]
    assert report.merged_group_count == 0
    assert report.decisions[0].decision == "keep"
    assert report.decisions[0].confidence == "blocked"


def test_heading_body_boundary_is_never_merged():
    document = make_document()
    document.add_paragraph("Глава 1", style="Heading 1")
    document.add_paragraph("основной текст после заголовка")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert [paragraph.role for paragraph in paragraphs] == ["heading", "body"]
    assert report.merged_group_count == 0
    assert report.decisions[0].reasons == ("heading_boundary", "non_body_role", "style_transition")


def test_list_boundary_is_never_merged():
    document = make_document()
    document.add_paragraph("Вступление без финальной точки")
    document.add_paragraph("Первый пункт", style="List Number")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert [paragraph.role for paragraph in paragraphs] == ["body", "list"]
    assert report.merged_group_count == 0
    assert "list_metadata" in report.decisions[0].reasons


def test_merged_text_spacing_is_normalized_conservatively():
    document = make_document()
    document.add_paragraph("Это важное")
    document.add_paragraph("продолжение текста")

    paragraphs, _, _ = extract_document_content_with_boundary_report(_save_document(document))

    assert paragraphs[0].text == "Это важное продолжение текста"


def test_dash_led_continuation_after_superscript_note_is_merged():
    document = make_document()
    first = document.add_paragraph()
    first.add_run("Apple held cash worth $187 billion")
    note = first.add_run("4")
    note.font.superscript = True
    document.add_paragraph(
        "— about the same size as the Czech economy that year, which made the tax arrangement especially visible in public debate."
    )

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert len(paragraphs) == 1
    assert paragraphs[0].text == (
        "Apple held cash worth $187 billion<sup>4</sup> — about the same size as the Czech economy "
        "that year, which made the tax arrangement especially visible in public debate."
    )
    assert paragraphs[0].origin_raw_indexes == [0, 1]
    assert paragraphs[0].boundary_source == "normalized_merge"
    assert paragraphs[0].boundary_confidence == "high"
    assert report.decisions[0].decision == "merge"
    assert report.decisions[0].confidence == "high"
    assert "right_starts_continuation" in report.decisions[0].reasons


def test_high_confidence_merge_chain_collapses_three_adjacent_body_paragraphs():
    document = make_document()
    document.add_paragraph("Это особенно важно")
    document.add_paragraph("для устойчивой")
    document.add_paragraph("работы всей системы.")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert len(paragraphs) == 1
    assert paragraphs[0].text == "Это особенно важно для устойчивой работы всей системы."
    assert paragraphs[0].origin_raw_indexes == [0, 1, 2]
    assert paragraphs[0].boundary_source == "normalized_merge"
    assert paragraphs[0].boundary_confidence == "high"
    assert paragraphs[0].boundary_rationale == (
        "same_body_style, compatible_alignment, left_not_terminal, "
        "right_starts_continuation, left_incomplete, combined_sentence_plausible"
    )
    assert report.merged_group_count == 1
    assert report.merged_raw_paragraph_count == 3
    assert [(decision.left_raw_index, decision.right_raw_index, decision.decision, decision.confidence) for decision in report.decisions] == [
        (0, 1, "merge", "high"),
        (1, 2, "merge", "high"),
    ]


def test_high_only_promotes_medium_chain_lead_when_followed_by_mergeable_continuation(monkeypatch):
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": True,
            "paragraph_boundary_normalization_mode": "high_only",
            "paragraph_boundary_normalization_save_debug_artifacts": False,
        },
    )
    document = make_document()
    document.add_paragraph("Почему она")
    document.add_paragraph("Чаще всего просто считается")
    document.add_paragraph("менее эффективной версией частного сектора?")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert len(paragraphs) == 1
    assert paragraphs[0].text == "Почему она Чаще всего просто считается менее эффективной версией частного сектора?"
    assert paragraphs[0].origin_raw_indexes == [0, 1, 2]
    assert paragraphs[0].boundary_source == "normalized_merge"
    assert paragraphs[0].boundary_confidence == "medium"
    assert [(decision.left_raw_index, decision.right_raw_index, decision.decision, decision.confidence) for decision in report.decisions] == [
        (0, 1, "merge", "medium"),
        (1, 2, "merge", "high"),
    ]
    assert report.decisions[0].reasons[-1] == "chain_continuation_supported"


def test_toc_like_entries_are_not_merged_by_boundary_normalization():
    document = make_document()
    document.add_paragraph("Contents")
    document.add_paragraph("Meet the Production Boundary")
    document.add_paragraph("Why Value Theory Matters")
    document.add_paragraph("The Structure of the Book")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert [paragraph.text for paragraph in paragraphs] == [
        "Contents",
        "Meet the Production Boundary",
        "Why Value Theory Matters",
        "The Structure of the Book",
    ]
    assert [paragraph.structural_role for paragraph in paragraphs] == [
        "toc_header",
        "toc_entry",
        "toc_entry",
        "toc_entry",
    ]
    assert report.merged_group_count == 0
    assert report.decisions[1].decision == "keep"
    assert report.decisions[1].confidence == "blocked"
    assert "adjacent_toc_like_entries" in report.decisions[1].reasons


def test_debug_artifact_report_writes_expected_path_and_payload(monkeypatch, tmp_path):
    reports_dir = tmp_path / "paragraph-boundary-reports"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_REPORTS_DIR", reports_dir)
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": True,
            "paragraph_boundary_normalization_mode": "high_only",
            "paragraph_boundary_normalization_save_debug_artifacts": True,
        },
    )
    monkeypatch.setattr(document_module, "resolve_uploaded_filename", lambda uploaded_file: "debug sample.docx")

    debug_doc = make_document()
    debug_doc.add_paragraph("архетипами: повторяющимися моделями")
    debug_doc.add_paragraph("поведения во времени, наблюдаемыми в разных системах.")
    uploaded = _save_document(debug_doc)

    paragraphs, _, report = extract_document_content_with_boundary_report(uploaded)

    assert len(paragraphs) == 1
    artifact_files = sorted(reports_dir.glob("*.json"))
    assert len(artifact_files) == 1
    assert artifact_files[0].name.startswith("debug_sample.docx_")

    payload = json.loads(artifact_files[0].read_text(encoding="utf-8"))
    assert payload == {
        "version": 1,
        "source_file": "debug sample.docx",
        "source_hash": payload["source_hash"],
        "mode": "high_only",
        "total_raw_paragraphs": 2,
        "total_logical_paragraphs": 1,
        "merged_group_count": 1,
        "merged_raw_paragraph_count": 2,
        "high_confidence_merge_count": 1,
        "medium_accepted_merge_count": 0,
        "medium_rejected_candidate_count": 0,
        "decisions": [
            {
                "left_raw_index": 0,
                "right_raw_index": 1,
                "decision": "merge",
                "confidence": "high",
                "reasons": [
                    "same_body_style",
                    "compatible_alignment",
                    "left_not_terminal",
                    "right_starts_continuation",
                    "left_incomplete",
                    "combined_sentence_plausible",
                ],
            }
        ],
    }
    assert len(payload["source_hash"]) == 8


def test_debug_artifact_report_uses_default_filename_for_plain_bytesio(monkeypatch, tmp_path):
    reports_dir = tmp_path / "paragraph-boundary-reports"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_REPORTS_DIR", reports_dir)
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": True,
            "paragraph_boundary_normalization_mode": "high_only",
            "paragraph_boundary_normalization_save_debug_artifacts": True,
        },
    )

    debug_doc = make_document()
    debug_doc.add_paragraph("архетипами: повторяющимися моделями")
    debug_doc.add_paragraph("поведения во времени, наблюдаемыми в разных системах.")

    extract_document_content_with_boundary_report(_save_document(debug_doc))

    artifact_files = sorted(reports_dir.glob("*.json"))
    assert len(artifact_files) == 1
    assert artifact_files[0].name.startswith("document.docx_")


def test_relation_debug_artifact_writes_expected_payload(monkeypatch, tmp_path):
    relation_reports_dir = tmp_path / "relation-normalization-reports"
    monkeypatch.setattr(document_module, "RELATION_NORMALIZATION_REPORTS_DIR", relation_reports_dir)
    report = RelationNormalizationReport(
        total_relations=1,
        relation_counts={"toc_region": 1},
        rejected_candidate_count=0,
        decisions=[
            ParagraphRelationDecision(
                relation_kind="toc_region",
                decision="accept",
                member_paragraph_ids=("p0000", "p0001", "p0002"),
                reasons=("toc_header_with_entries",),
            )
        ],
    )

    artifact_path = document_module._write_relation_normalization_report_artifact(
        source_name="debug sample.docx",
        source_bytes=b"debug-source-bytes",
        profile="phase2_default",
        enabled_relation_kinds=("image_caption", "table_caption", "epigraph_attribution", "toc_region"),
        report=report,
    )

    assert artifact_path is not None
    artifact_files = sorted(relation_reports_dir.glob("*.json"))
    assert len(artifact_files) == 1

    payload = json.loads(artifact_files[0].read_text(encoding="utf-8"))
    assert payload == {
        "version": 1,
        "source_file": "debug sample.docx",
        "source_hash": payload["source_hash"],
        "profile": "phase2_default",
        "enabled_relation_kinds": ["image_caption", "table_caption", "epigraph_attribution", "toc_region"],
        "total_relations": 1,
        "relation_counts": {"toc_region": 1},
        "rejected_candidate_count": 0,
        "decisions": [
            {
                "relation_kind": "toc_region",
                "decision": "accept",
                "member_paragraph_ids": ["p0000", "p0001", "p0002"],
                "anchor_asset_id": None,
                "reasons": ["toc_header_with_entries"],
            }
        ],
    }
    assert len(payload["source_hash"]) == 8


def test_debug_artifact_report_is_not_written_when_saving_disabled(monkeypatch, tmp_path):
    reports_dir = tmp_path / "paragraph-boundary-reports"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_REPORTS_DIR", reports_dir)
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": True,
            "paragraph_boundary_normalization_mode": "high_only",
            "paragraph_boundary_normalization_save_debug_artifacts": False,
        },
    )
    monkeypatch.setattr(document_module, "resolve_uploaded_filename", lambda uploaded_file: "debug sample.docx")

    debug_doc = make_document()
    debug_doc.add_paragraph("архетипами: повторяющимися моделями")
    debug_doc.add_paragraph("поведения во времени, наблюдаемыми в разных системах.")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(debug_doc))

    assert len(paragraphs) == 1
    assert report.merged_group_count == 1
    assert not reports_dir.exists()


def test_debug_artifact_directory_creation_failure_is_non_fatal(monkeypatch, tmp_path):
    reports_dir = tmp_path / "paragraph-boundary-reports"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_REPORTS_DIR", reports_dir)
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": True,
            "paragraph_boundary_normalization_mode": "high_only",
            "paragraph_boundary_normalization_save_debug_artifacts": True,
        },
    )

    def fail_mkdir(self, *args, **kwargs):
        raise OSError("mkdir blocked")

    monkeypatch.setattr(document_module.Path, "mkdir", fail_mkdir)

    debug_doc = make_document()
    debug_doc.add_paragraph("архетипами: повторяющимися моделями")
    debug_doc.add_paragraph("поведения во времени, наблюдаемыми в разных системах.")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(debug_doc))

    assert len(paragraphs) == 1
    assert paragraphs[0].text == "архетипами: повторяющимися моделями поведения во времени, наблюдаемыми в разных системах."
    assert report.merged_group_count == 1
    assert not reports_dir.exists()


def test_debug_artifact_file_write_failure_is_non_fatal(monkeypatch, tmp_path):
    reports_dir = tmp_path / "paragraph-boundary-reports"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_REPORTS_DIR", reports_dir)
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": True,
            "paragraph_boundary_normalization_mode": "high_only",
            "paragraph_boundary_normalization_save_debug_artifacts": True,
        },
    )
    monkeypatch.setattr(document_module, "resolve_uploaded_filename", lambda uploaded_file: "debug sample.docx")

    original_write_text = document_module.Path.write_text

    def fail_report_write(self, *args, **kwargs):
        if self.suffix == ".json":
            raise OSError("write blocked")
        return original_write_text(self, *args, **kwargs)

    monkeypatch.setattr(document_module.Path, "write_text", fail_report_write)

    debug_doc = make_document()
    debug_doc.add_paragraph("архетипами: повторяющимися моделями")
    debug_doc.add_paragraph("поведения во времени, наблюдаемыми в разных системах.")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(debug_doc))

    assert len(paragraphs) == 1
    assert report.merged_group_count == 1
    assert reports_dir.exists()
    assert list(reports_dir.glob("*.json")) == []


def test_disabled_normalization_preserves_legacy_boundaries(monkeypatch):
    document = make_document()
    document.add_paragraph("архетипами: повторяющимися моделями")
    document.add_paragraph("поведения во времени, наблюдаемыми в разных системах.")
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": False,
            "paragraph_boundary_normalization_mode": "high_only",
            "paragraph_boundary_normalization_save_debug_artifacts": False,
        },
    )

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert [paragraph.text for paragraph in paragraphs] == [
        "архетипами: повторяющимися моделями",
        "поведения во времени, наблюдаемыми в разных системах.",
    ]
    assert report.merged_group_count == 0
    assert report.merged_raw_paragraph_count == 0


def test_medium_candidate_is_rejected_in_high_only_mode(monkeypatch):
    document = make_document()
    document.add_paragraph("Это важное наблюдение:")
    document.add_paragraph("Следующий шаг требует дополнительной проверки.")
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": True,
            "paragraph_boundary_normalization_mode": "high_only",
            "paragraph_boundary_normalization_save_debug_artifacts": False,
        },
    )

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert [paragraph.text for paragraph in paragraphs] == [
        "Это важное наблюдение:",
        "Следующий шаг требует дополнительной проверки.",
    ]
    assert report.merged_group_count == 0
    assert report.decisions[0].decision == "keep"
    assert report.decisions[0].confidence == "medium"
    assert "medium_mode_disabled" in report.decisions[0].reasons


def test_medium_diagnostics_are_distinguished_in_boundary_report_artifact(monkeypatch, tmp_path):
    reports_dir = tmp_path / "paragraph-boundary-reports"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_REPORTS_DIR", reports_dir)
    monkeypatch.setattr(document_module, "resolve_uploaded_filename", lambda uploaded_file: "medium sample.docx")
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": True,
            "paragraph_boundary_normalization_mode": "high_only",
            "paragraph_boundary_normalization_save_debug_artifacts": True,
        },
    )

    document = make_document()
    document.add_paragraph("Это важное наблюдение:")
    document.add_paragraph("Следующий шаг требует дополнительной проверки.")

    extract_document_content_with_boundary_report(_save_document(document))

    artifact_files = sorted(reports_dir.glob("*.json"))
    payload = json.loads(artifact_files[0].read_text(encoding="utf-8"))
    assert payload["high_confidence_merge_count"] == 0
    assert payload["medium_accepted_merge_count"] == 0
    assert payload["medium_rejected_candidate_count"] == 1


def test_medium_candidate_is_merged_in_high_and_medium_mode(monkeypatch):
    document = make_document()
    document.add_paragraph("Это важное наблюдение:")
    document.add_paragraph("Следующий шаг требует дополнительной проверки.")
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": True,
            "paragraph_boundary_normalization_mode": "high_and_medium",
            "paragraph_boundary_normalization_save_debug_artifacts": False,
        },
    )

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert len(paragraphs) == 1
    assert paragraphs[0].text == "Это важное наблюдение: Следующий шаг требует дополнительной проверки."
    assert paragraphs[0].boundary_confidence == "medium"
    assert report.merged_group_count == 1
    assert report.decisions[0].decision == "merge"
    assert report.decisions[0].confidence == "medium"
    assert "medium_mode_disabled" not in report.decisions[0].reasons


def test_blocked_boundary_stays_blocked_in_high_and_medium_mode(monkeypatch):
    document = make_document()
    document.add_paragraph("Первый абзац закончен.")
    document.add_paragraph("Второй абзац начинается с новой мысли.")
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "paragraph_boundary_normalization_enabled": True,
            "paragraph_boundary_normalization_mode": "high_and_medium",
            "paragraph_boundary_normalization_save_debug_artifacts": False,
        },
    )

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert len(paragraphs) == 2
    assert report.merged_group_count == 0
    assert report.decisions[0].decision == "keep"
    assert report.decisions[0].confidence == "blocked"


def test_caption_boundary_is_never_merged():
    document = make_document()
    document.add_paragraph("Вступление без завершающей точки")
    document.add_paragraph("Рисунок 1. Подпись", style="Caption")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert [paragraph.role for paragraph in paragraphs] == ["body", "caption"]
    assert report.merged_group_count == 0
    assert "caption_like_boundary" in report.decisions[0].reasons


def test_build_paragraph_relations_accepts_adjacent_caption_candidate_without_prior_reclassification():
    paragraphs = [
        ParagraphUnit(text="[[DOCX_IMAGE_img-1]]", role="image", asset_id="img-1", paragraph_id="p0000"),
        ParagraphUnit(text="Рисунок 1. Подпись", role="body", paragraph_id="p0001"),
    ]

    relations, report = build_paragraph_relations(
        paragraphs,
        enabled_relation_kinds=("image_caption",),
    )

    assert len(relations) == 1
    assert relations[0].relation_kind == "image_caption"
    assert relations[0].member_paragraph_ids == ("p0000", "p0001")
    assert relations[0].anchor_asset_id == "img-1"
    assert report.total_relations == 1


def test_adjacent_caption_candidate_is_attached_after_relation_side_effects_and_reclassification(tmp_path):
    image_path = tmp_path / "image.png"
    image_path.write_bytes(PNG_BYTES)
    document = make_document()
    document.add_paragraph().add_run().add_picture(str(image_path))
    document.add_paragraph("Рисунок 1. Подпись")

    paragraphs, _, _, relations, relation_report, _ = extract_document_content_with_normalization_reports(_save_document(document))

    assert [paragraph.role for paragraph in paragraphs] == ["image", "caption"]
    assert paragraphs[1].attached_to_asset_id == paragraphs[0].asset_id
    assert len(relations) == 1
    assert relations[0].relation_kind == "image_caption"
    assert relation_report.total_relations == 1


def test_image_boundary_is_never_merged(tmp_path):
    image_path = tmp_path / "image.png"
    image_path.write_bytes(PNG_BYTES)
    document = make_document()
    document.add_paragraph("Вступление без завершающей точки")
    document.add_paragraph().add_run().add_picture(str(image_path))
    document.add_paragraph("продолжение после изображения")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert [paragraph.role for paragraph in paragraphs] == ["body", "image", "body"]
    assert report.merged_group_count == 0
    assert report.decisions[0].decision == "keep"
    assert report.decisions[0].confidence == "blocked"
    assert "non_body_role" in report.decisions[0].reasons
    assert report.decisions[1].decision == "keep"
    assert report.decisions[1].confidence == "blocked"
    assert "non_body_role" in report.decisions[1].reasons


def test_table_boundary_is_never_merged():
    document = make_document()
    document.add_paragraph("Вступление без завершающей точки")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "ячейка"
    document.add_paragraph("продолжение после таблицы")

    paragraphs, _, report = extract_document_content_with_boundary_report(_save_document(document))

    assert [paragraph.role for paragraph in paragraphs] == ["body", "table", "body"]
    assert report.merged_group_count == 0


def test_ai_review_candidates_include_medium_boundaries_and_rejected_relations():
    candidates = boundary_review_module.build_ai_review_candidates(
        raw_blocks=[
            RawParagraph(raw_index=0, text="Это важное наблюдение:", style_name="Normal"),
            RawParagraph(raw_index=1, text="Следующий шаг требует проверки.", style_name="Normal"),
        ],
        paragraphs=[
            ParagraphUnit(text="Подпись 1", role="caption", paragraph_id="p0001"),
            ParagraphUnit(text="Подпись 2", role="caption", paragraph_id="p0002"),
        ],
        boundary_report=ParagraphBoundaryNormalizationReport(
            total_raw_paragraphs=2,
            total_logical_paragraphs=2,
            merged_group_count=0,
            merged_raw_paragraph_count=0,
            decisions=[
                ParagraphBoundaryDecision(
                    left_raw_index=0,
                    right_raw_index=1,
                    decision="keep",
                    confidence="medium",
                    reasons=("combined_sentence_plausible", "medium_mode_disabled"),
                )
            ],
        ),
        relation_report=RelationNormalizationReport(
            total_relations=0,
            relation_counts={},
            rejected_candidate_count=1,
            decisions=[
                ParagraphRelationDecision(
                    relation_kind="epigraph_attribution",
                    decision="reject",
                    member_paragraph_ids=("p0001", "p0002"),
                    reasons=("spacing_gap",),
                )
            ],
        ),
        candidate_limit=10,
    )

    assert [candidate["candidate_kind"] for candidate in candidates] == [
        "boundary_medium",
        "relation_rejected",
    ]
    assert candidates[0]["candidate_id"] == "0:1"
    assert candidates[0]["left_text"] == "Это важное наблюдение:"
    assert candidates[1]["candidate_id"] == "epigraph_attribution:p0001|p0002"
    assert candidates[1]["member_texts"] == ["Подпись 1", "Подпись 2"]


def test_ai_review_timeout_writes_non_authoritative_artifact(monkeypatch, tmp_path):
    review_dir = tmp_path / "paragraph-boundary-ai-review"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_AI_REVIEW_DIR", review_dir)
    monkeypatch.setattr(
        document_module,
        "_request_ai_review_recommendations",
        lambda **kwargs: (_ for _ in ()).throw(TimeoutError("review timeout")),
    )

    artifact_path = document_module._run_paragraph_boundary_ai_review(
        source_name="debug sample.docx",
        source_bytes=b"debug-source-bytes",
        mode="review_only",
        model="gpt-5.4",
        raw_blocks=[
            RawParagraph(raw_index=0, text="Это важное наблюдение:", style_name="Normal"),
            RawParagraph(raw_index=1, text="Следующий шаг требует проверки.", style_name="Normal"),
        ],
        paragraphs=[],
        boundary_report=ParagraphBoundaryNormalizationReport(
            total_raw_paragraphs=2,
            total_logical_paragraphs=2,
            merged_group_count=0,
            merged_raw_paragraph_count=0,
            decisions=[
                ParagraphBoundaryDecision(
                    left_raw_index=0,
                    right_raw_index=1,
                    decision="keep",
                    confidence="medium",
                    reasons=("combined_sentence_plausible", "medium_mode_disabled"),
                )
            ],
        ),
        relation_report=RelationNormalizationReport(
            total_relations=0,
            relation_counts={},
            rejected_candidate_count=0,
            decisions=[],
        ),
        candidate_limit=10,
        timeout_seconds=5,
        max_tokens_per_candidate=120,
    )

    assert artifact_path is not None
    payload = json.loads(next(review_dir.glob("*.json")).read_text(encoding="utf-8"))
    assert payload["mode"] == "review_only"
    assert payload["error_code"] == "timeout"
    assert payload["reviewed_candidate_count"] == 1
    assert payload["accepted_candidate_count"] == 0
    assert payload["decisions"][0]["final_decision"] == "keep"
    assert "ai_review_unavailable:timeout" in payload["decisions"][0]["reasons"]


def test_ai_review_disagreement_is_logged_but_deterministic_decision_is_retained(monkeypatch, tmp_path):
    review_dir = tmp_path / "paragraph-boundary-ai-review"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_AI_REVIEW_DIR", review_dir)
    monkeypatch.setattr(
        document_module,
        "_request_ai_review_recommendations",
        lambda **kwargs: {"0:1": {"recommendation": "merge", "reasons": ["continuation"]}},
    )

    artifact_path = document_module._run_paragraph_boundary_ai_review(
        source_name="debug sample.docx",
        source_bytes=b"debug-source-bytes",
        mode="review_only",
        model="gpt-5.4",
        raw_blocks=[
            RawParagraph(raw_index=0, text="Это важное наблюдение:", style_name="Normal"),
            RawParagraph(raw_index=1, text="Следующий шаг требует проверки.", style_name="Normal"),
        ],
        paragraphs=[],
        boundary_report=ParagraphBoundaryNormalizationReport(
            total_raw_paragraphs=2,
            total_logical_paragraphs=2,
            merged_group_count=0,
            merged_raw_paragraph_count=0,
            decisions=[
                ParagraphBoundaryDecision(
                    left_raw_index=0,
                    right_raw_index=1,
                    decision="keep",
                    confidence="medium",
                    reasons=("combined_sentence_plausible", "medium_mode_disabled"),
                )
            ],
        ),
        relation_report=RelationNormalizationReport(
            total_relations=0,
            relation_counts={},
            rejected_candidate_count=0,
            decisions=[],
        ),
        candidate_limit=10,
        timeout_seconds=5,
        max_tokens_per_candidate=120,
    )

    assert artifact_path is not None
    payload = json.loads(next(review_dir.glob("*.json")).read_text(encoding="utf-8"))
    assert payload["decisions"][0]["ai_recommendation"] == "merge"
    assert payload["decisions"][0]["final_decision"] == "keep"
    assert "deterministic_decision_retained" in payload["decisions"][0]["reasons"]


def test_ai_review_empty_response_falls_back_to_deterministic_artifact(monkeypatch, tmp_path):
    review_dir = tmp_path / "paragraph-boundary-ai-review"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_AI_REVIEW_DIR", review_dir)
    monkeypatch.setattr(
        document_module,
        "_request_ai_review_recommendations",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("AI review returned empty output.")),
    )

    artifact_path = document_module._run_paragraph_boundary_ai_review(
        source_name="debug sample.docx",
        source_bytes=b"debug-source-bytes",
        mode="review_only",
        model="gpt-5.4",
        raw_blocks=[
            RawParagraph(raw_index=0, text="Это важное наблюдение:", style_name="Normal"),
            RawParagraph(raw_index=1, text="Следующий шаг требует проверки.", style_name="Normal"),
        ],
        paragraphs=[],
        boundary_report=ParagraphBoundaryNormalizationReport(
            total_raw_paragraphs=2,
            total_logical_paragraphs=2,
            merged_group_count=0,
            merged_raw_paragraph_count=0,
            decisions=[
                ParagraphBoundaryDecision(
                    left_raw_index=0,
                    right_raw_index=1,
                    decision="keep",
                    confidence="medium",
                    reasons=("combined_sentence_plausible",),
                )
            ],
        ),
        relation_report=RelationNormalizationReport(
            total_relations=0,
            relation_counts={},
            rejected_candidate_count=0,
            decisions=[],
        ),
        candidate_limit=10,
        timeout_seconds=5,
        max_tokens_per_candidate=120,
    )

    assert artifact_path is not None
    payload = json.loads(next(review_dir.glob("*.json")).read_text(encoding="utf-8"))
    assert payload["error_code"] == "empty_response"
    assert payload["decisions"][0]["final_decision"] == "keep"
    assert "ai_review_unavailable:empty_response" in payload["decisions"][0]["reasons"]


def test_ai_review_malformed_response_falls_back_to_deterministic_artifact(monkeypatch, tmp_path):
    review_dir = tmp_path / "paragraph-boundary-ai-review"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_AI_REVIEW_DIR", review_dir)
    monkeypatch.setattr(
        document_module,
        "_request_ai_review_recommendations",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("AI review did not return JSON.")),
    )

    artifact_path = document_module._run_paragraph_boundary_ai_review(
        source_name="debug sample.docx",
        source_bytes=b"debug-source-bytes",
        mode="review_only",
        model="gpt-5.4",
        raw_blocks=[
            RawParagraph(raw_index=0, text="Это важное наблюдение:", style_name="Normal"),
            RawParagraph(raw_index=1, text="Следующий шаг требует проверки.", style_name="Normal"),
        ],
        paragraphs=[],
        boundary_report=ParagraphBoundaryNormalizationReport(
            total_raw_paragraphs=2,
            total_logical_paragraphs=2,
            merged_group_count=0,
            merged_raw_paragraph_count=0,
            decisions=[
                ParagraphBoundaryDecision(
                    left_raw_index=0,
                    right_raw_index=1,
                    decision="keep",
                    confidence="medium",
                    reasons=("combined_sentence_plausible",),
                )
            ],
        ),
        relation_report=RelationNormalizationReport(
            total_relations=0,
            relation_counts={},
            rejected_candidate_count=0,
            decisions=[],
        ),
        candidate_limit=10,
        timeout_seconds=5,
        max_tokens_per_candidate=120,
    )

    assert artifact_path is not None
    payload = json.loads(next(review_dir.glob("*.json")).read_text(encoding="utf-8"))
    assert payload["error_code"] == "no_json_object"
    assert payload["decisions"][0]["final_decision"] == "keep"
    assert "ai_review_unavailable:no_json_object" in payload["decisions"][0]["reasons"]


def test_ai_review_without_matching_recommendation_marks_no_recommendation(monkeypatch, tmp_path):
    review_dir = tmp_path / "paragraph-boundary-ai-review"
    monkeypatch.setattr(document_module, "PARAGRAPH_BOUNDARY_AI_REVIEW_DIR", review_dir)
    monkeypatch.setattr(
        document_module,
        "_request_ai_review_recommendations",
        lambda **kwargs: {"different-candidate": {"recommendation": "merge", "reasons": ["other"]}},
    )

    artifact_path = document_module._run_paragraph_boundary_ai_review(
        source_name="debug sample.docx",
        source_bytes=b"debug-source-bytes",
        mode="review_only",
        model="gpt-5.4",
        raw_blocks=[
            RawParagraph(raw_index=0, text="Это важное наблюдение:", style_name="Normal"),
            RawParagraph(raw_index=1, text="Следующий шаг требует проверки.", style_name="Normal"),
        ],
        paragraphs=[],
        boundary_report=ParagraphBoundaryNormalizationReport(
            total_raw_paragraphs=2,
            total_logical_paragraphs=2,
            merged_group_count=0,
            merged_raw_paragraph_count=0,
            decisions=[
                ParagraphBoundaryDecision(
                    left_raw_index=0,
                    right_raw_index=1,
                    decision="keep",
                    confidence="medium",
                    reasons=("combined_sentence_plausible",),
                )
            ],
        ),
        relation_report=RelationNormalizationReport(
            total_relations=0,
            relation_counts={},
            rejected_candidate_count=0,
            decisions=[],
        ),
        candidate_limit=10,
        timeout_seconds=5,
        max_tokens_per_candidate=120,
    )

    assert artifact_path is not None
    payload = json.loads(next(review_dir.glob("*.json")).read_text(encoding="utf-8"))
    assert "error_code" not in payload
    assert payload["decisions"][0]["ai_recommendation"] is None
    assert payload["decisions"][0]["final_decision"] == "keep"
    assert "ai_review_no_recommendation" in payload["decisions"][0]["reasons"]
