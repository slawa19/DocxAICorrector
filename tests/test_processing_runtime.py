import hashlib
import queue
import sys
import subprocess
import threading
import time
from types import SimpleNamespace
from pathlib import Path
from typing import Any, cast

import pytest

import docxaicorrector.processing.processing_runtime as processing_runtime
import docxaicorrector.processing.upload_ports as upload_ports
import docxaicorrector.runtime.state as state
from docxaicorrector.document.extraction import extract_document_content_from_docx, extract_paragraph_units_from_docx
from docxaicorrector.pdf_import.images import PdfImageObject
from docxaicorrector.pdf_import.text_layer_quality import PdfTextSpan
from docxaicorrector.runtime.events import (
    AppendImageLogEvent,
    AppendLogEvent,
    FinalizeProcessingStatusEvent,
    PreparationCompleteEvent,
    PreparationFailedEvent,
    PreparationStoppedEvent,
    PushActivityEvent,
    ResetImageStateEvent,
    SetProcessingStatusEvent,
    SetStateEvent,
    WorkerCompleteEvent,
)


class SessionState(dict):
    def __getattr__(self, name):
        try:
            return self[name]
        except KeyError as exc:
            raise AttributeError(name) from exc

    def __setattr__(self, name, value):
        self[name] = value


@pytest.fixture(autouse=True)
def _pdf_parse_uses_in_process_seam(monkeypatch):
    """F7 test seam: the production PDF parse runs in a spawned child process
    that cannot see in-process monkeypatches. These tests patch pdfminer in
    process, so default every test to the in-process worker. Tests that must
    exercise the real subprocess path opt out with an explicit setattr(False)."""

    monkeypatch.setattr(processing_runtime, "_PDF_PARSE_IN_PROCESS", True)


class UploadedFileStub:
    def __init__(self, name: str, content: bytes):
        self.name = name
        self.size = len(content)
        self._content = content
        self._position = 0

    def read(self, size: int = -1) -> bytes:
        if size < 0:
            data = self._content[self._position :]
            self._position = len(self._content)
            return data
        start = self._position
        end = min(len(self._content), start + size)
        self._position = end
        return self._content[start:end]

    def getvalue(self) -> bytes:
        return self._content

    def seek(self, offset: int, whence: int = 0) -> int:
        if whence == 0:
            self._position = max(0, offset)
        elif whence == 1:
            self._position = max(0, self._position + offset)
        elif whence == 2:
            self._position = max(0, len(self._content) + offset)
        else:
            raise ValueError("Unsupported whence")
        self._position = min(self._position, len(self._content))
        return self._position


def _clear_materialized_upload_cache() -> None:
    processing_runtime._shared_materialized_upload_cache.clear()
    processing_runtime._shared_materialized_upload_inflight.clear()


def test_build_uploaded_file_token_uses_name_size_and_content_hash():
    token = processing_runtime.build_uploaded_file_token(UploadedFileStub("report.docx", b"abc"))

    assert token == "report.docx:3:ba7816bf8f01cfea"


def test_build_preparation_request_marker_includes_chunk_size():
    marker = processing_runtime.build_preparation_request_marker(UploadedFileStub("report.docx", b"abc"), chunk_size=6000)

    assert marker == "report.docx:3:ba7816bf8f01cfea:6000"


def test_build_preparation_request_marker_includes_non_default_operation():
    marker = processing_runtime.build_preparation_request_marker(
        UploadedFileStub("report.docx", b"abc"),
        chunk_size=6000,
        processing_operation="audiobook",
    )

    assert marker == "report.docx:3:ba7816bf8f01cfea:6000:op=audiobook"


def test_build_preparation_request_marker_uses_content_hash_for_same_name_same_size_files():
    marker_one = processing_runtime.build_preparation_request_marker(UploadedFileStub("report.docx", b"abc"), chunk_size=6000)
    marker_two = processing_runtime.build_preparation_request_marker(UploadedFileStub("report.docx", b"xyz"), chunk_size=6000)

    assert marker_one != marker_two


def test_drain_processing_events_applies_typed_runtime_events(monkeypatch):
    session_state = SessionState(
        processing_event_queue=queue.Queue(),
        image_assets=["stale"],
        image_validation_failures=["stale"],
        image_processing_summary={"total_images": 3, "processed_images": 2, "validation_errors": ["boom"]},
        restart_source=None,
        processing_worker=object(),
        processing_stop_event=object(),
        processing_stop_requested=True,
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)

    calls = {
        "status": [],
        "finalize": [],
        "activity": [],
        "log": [],
        "image_log": [],
    }

    session_state.processing_event_queue.put(SetStateEvent(values={"last_error": "boom"}))
    session_state.processing_event_queue.put(ResetImageStateEvent())
    session_state.processing_event_queue.put(
        SetProcessingStatusEvent(
            payload={
                "stage": "run",
                "detail": "detail",
                "segment_status_by_id": {"seg_0001": "processing"},
                "segment_progress_by_id": {"seg_0001": 0.5},
                "active_segment_id": "seg_0001",
                "active_segment_title": "Chapter 1",
            }
        )
    )
    session_state.processing_event_queue.put(FinalizeProcessingStatusEvent(stage="done", detail="ok", progress=1.0, terminal_kind="completed"))
    session_state.processing_event_queue.put(PushActivityEvent(message="hello"))
    session_state.processing_event_queue.put(AppendLogEvent(payload={"status": "OK", "block_index": 1, "block_count": 2, "target_chars": 3, "context_chars": 4, "details": "done"}))
    session_state.processing_event_queue.put(AppendImageLogEvent(payload={"image_id": "img_1", "status": "validated", "decision": "accept", "confidence": 0.9}))
    session_state.processing_event_queue.put(WorkerCompleteEvent(outcome="succeeded"))

    processing_runtime.drain_processing_events(
        set_processing_status=lambda **payload: calls["status"].append(payload),
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: calls["finalize"].append((stage, detail, progress, terminal_kind)),
        push_activity=lambda message: calls["activity"].append(message),
        append_log=lambda **payload: calls["log"].append(payload),
        append_image_log=lambda **payload: calls["image_log"].append(payload),
    )

    assert session_state.last_error == "boom"
    assert session_state.image_assets == []
    assert session_state.image_validation_failures == []
    assert session_state.image_processing_summary == {
        "total_images": 0,
        "processed_images": 0,
        "images_validated": 0,
        "validation_passed": 0,
        "fallbacks_applied": 0,
        "validation_errors": [],
    }
    assert calls["status"] == [{
        "stage": "run",
        "detail": "detail",
        "segment_status_by_id": {"seg_0001": "processing"},
        "segment_progress_by_id": {"seg_0001": 0.5},
        "active_segment_id": "seg_0001",
        "active_segment_title": "Chapter 1",
    }]
    assert calls["finalize"] == [("done", "ok", 1.0, "completed")]
    assert calls["activity"] == ["hello"]
    assert calls["log"][0]["status"] == "OK"
    assert calls["image_log"][0]["image_id"] == "img_1"
    assert session_state.processing_outcome == "succeeded"
    assert session_state.processing_worker is None
    assert session_state.processing_event_queue is None
    assert session_state.processing_stop_event is None
    assert session_state.processing_stop_requested is False


def test_drain_processing_events_warns_and_ignores_unknown_set_state_keys(monkeypatch):
    session_state = SessionState(
        processing_event_queue=queue.Queue(),
        processing_worker=object(),
        processing_stop_event=object(),
        processing_stop_requested=True,
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    log_calls = []
    monkeypatch.setattr(processing_runtime, "log_event", lambda *args, **kwargs: log_calls.append((args, kwargs)))

    session_state.processing_event_queue.put(
        SetStateEvent(values={"last_error": "boom", "unexpected_key": "nope"})
    )
    session_state.processing_event_queue.put(WorkerCompleteEvent(outcome="failed"))

    processing_runtime.drain_processing_events(
        set_processing_status=lambda **payload: None,
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: None,
        push_activity=lambda message: None,
        append_log=lambda **payload: None,
        append_image_log=lambda **payload: None,
    )

    assert session_state.last_error == "boom"
    assert "unexpected_key" not in session_state
    assert len(log_calls) == 1
    assert log_calls[0][0][1] == "state_event_unknown_keys"
    assert log_calls[0][1]["unknown_keys"] == ["unexpected_key"]


def test_build_runtime_event_emitters_emits_typed_events_for_background_runtime():
    emitted_events = []

    class RuntimeStub:
        def emit(self, event):
            emitted_events.append(event)

    emitters = processing_runtime.build_runtime_event_emitters(
        dependencies=processing_runtime.RuntimeEventEmitterDependencies(
            set_processing_status=lambda **payload: None,
            finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: None,
            push_activity=lambda message: None,
            append_log=lambda **payload: None,
            append_image_log=lambda **payload: None,
        )
    )

    runtime = RuntimeStub()
    emitters.emit_state(runtime, last_error="boom")
    emitters.emit_image_reset(runtime)
    emitters.emit_status(
        runtime,
        stage="run",
        detail="detail",
        segment_status_by_id={"seg_0001": "processing"},
        segment_progress_by_id={"seg_0001": 0.5},
        active_segment_id="seg_0001",
        active_segment_title="Chapter 1",
    )
    emitters.emit_finalize(runtime, "done", "ok", 1.0, "completed")
    emitters.emit_activity(runtime, "hello")
    emitters.emit_log(runtime, status="OK", block_index=1, block_count=1, target_chars=2, context_chars=0, details="done")
    emitters.emit_image_log(runtime, image_id="img_1", status="validated", decision="accept", confidence=0.9)

    assert emitted_events == [
        SetStateEvent(values={"last_error": "boom"}),
        ResetImageStateEvent(),
        SetProcessingStatusEvent(
            payload={
                "stage": "run",
                "detail": "detail",
                "segment_status_by_id": {"seg_0001": "processing"},
                "segment_progress_by_id": {"seg_0001": 0.5},
                "active_segment_id": "seg_0001",
                "active_segment_title": "Chapter 1",
            }
        ),
        FinalizeProcessingStatusEvent(stage="done", detail="ok", progress=1.0, terminal_kind="completed"),
        PushActivityEvent(message="hello"),
        AppendLogEvent(payload={"status": "OK", "block_index": 1, "block_count": 1, "target_chars": 2, "context_chars": 0, "details": "done"}),
        AppendImageLogEvent(payload={"image_id": "img_1", "status": "validated", "decision": "accept", "confidence": 0.9}),
    ]


def test_background_runtime_tags_processing_events_with_source_token():
    event_queue = queue.Queue()
    runtime = processing_runtime.BackgroundRuntime(
        event_queue,
        threading.Event(),
        source_token="report.docx:3:abc",
    )

    runtime.emit(SetStateEvent(values={"last_error": "boom"}))
    runtime.emit(WorkerCompleteEvent(outcome="succeeded"))

    assert event_queue.get_nowait() == SetStateEvent(
        values={"last_error": "boom"},
        source_token="report.docx:3:abc",
    )
    assert event_queue.get_nowait() == WorkerCompleteEvent(
        outcome="succeeded",
        source_token="report.docx:3:abc",
    )


def test_drain_processing_events_ignores_stale_source_token_events(monkeypatch):
    session_state = SessionState(
        processing_event_queue=queue.Queue(),
        image_assets=["current-image"],
        latest_docx_bytes=b"current-docx",
        latest_markdown="current-md",
        processing_worker=object(),
        processing_stop_event=object(),
        processing_stop_requested=True,
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(processing_runtime, "get_latest_source_token", lambda: "new.docx:3:new")

    finalize_calls = []

    session_state.processing_event_queue.put(
        SetStateEvent(
            values={"latest_docx_bytes": b"old-docx", "latest_markdown": "old-md", "last_error": "old-error"},
            source_token="old.docx:3:old",
        )
    )
    session_state.processing_event_queue.put(ResetImageStateEvent(source_token="old.docx:3:old"))
    session_state.processing_event_queue.put(
        FinalizeProcessingStatusEvent(
            stage="done",
            detail="stale",
            progress=1.0,
            terminal_kind="completed",
            source_token="old.docx:3:old",
        )
    )
    session_state.processing_event_queue.put(WorkerCompleteEvent(outcome="succeeded", source_token="old.docx:3:old"))

    processing_runtime.drain_processing_events(
        set_processing_status=lambda **payload: None,
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: finalize_calls.append((stage, detail, progress, terminal_kind)),
        push_activity=lambda message: None,
        append_log=lambda **payload: None,
        append_image_log=lambda **payload: None,
    )

    assert session_state.image_assets == ["current-image"]
    assert session_state.latest_docx_bytes == b"current-docx"
    assert session_state.latest_markdown == "current-md"
    assert session_state.processing_worker is not None
    assert session_state.processing_event_queue is not None
    assert session_state.processing_stop_event is not None
    assert session_state.processing_stop_requested is True
    assert finalize_calls == []


def test_drain_preparation_events_stores_prepared_context(monkeypatch):
    prepared_run_context = type("PreparedRunContextStub", (), {
        "uploaded_file_token": "report.docx:3:abc",
        "prepared_source_key": "report.docx:3:abc:6000",
    })()
    session_state = SessionState(
        preparation_event_queue=queue.Queue(),
        preparation_worker=object(),
        selected_source_token="",
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    finalized = []

    session_state.preparation_event_queue.put(
        PreparationCompleteEvent(
            prepared_run_context=prepared_run_context,
            upload_marker="report.docx:3:ba7816bf8f01cfea:6000",
        )
    )

    processing_runtime.drain_preparation_events(
        reset_run_state=lambda **kwargs: None,
        set_processing_status=lambda **payload: None,
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: finalized.append((stage, detail, progress, terminal_kind)),
        push_activity=lambda message: None,
    )

    assert session_state.prepared_run_context is prepared_run_context
    assert session_state.preparation_input_marker == "report.docx:3:ba7816bf8f01cfea:6000"
    assert session_state.selected_source_token == "report.docx:3:abc"
    assert session_state.preparation_worker is None
    assert session_state.preparation_event_queue is None
    assert finalized == [("Документ подготовлен", "", 1.0, "completed")]


def test_drain_preparation_events_marks_failure(monkeypatch):
    session_state = SessionState(
        preparation_event_queue=queue.Queue(),
        preparation_worker=object(),
        last_error="",
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    finalized = []
    activities = []

    session_state.preparation_event_queue.put(
        PreparationFailedEvent(
            upload_marker="report.docx:3:ba7816bf8f01cfea:6000",
            error_message="boom",
            error_details={
                "stage": "preparation",
                "severity": "error",
                "user_message": "boom",
                "technical_message": "boom",
                "error_type": "RuntimeError",
                "recoverable": False,
            },
        )
    )

    processing_runtime.drain_preparation_events(
        reset_run_state=lambda **kwargs: None,
        set_processing_status=lambda **payload: None,
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: finalized.append((stage, detail, progress, terminal_kind)),
        push_activity=lambda message: activities.append(message),
    )

    assert session_state.prepared_run_context is None
    assert session_state.preparation_failed_marker == "report.docx:3:ba7816bf8f01cfea:6000"
    assert session_state.last_error == "boom"
    assert session_state.last_background_error["stage"] == "preparation"
    assert session_state.preparation_worker is None
    assert session_state.preparation_event_queue is None
    assert finalized == [("Ошибка подготовки", "boom", 1.0, "error")]
    assert activities == ["Не удалось прочитать и проанализировать документ."]


def test_drain_preparation_events_ignores_stale_completion_marker(monkeypatch):
    current_queue = queue.Queue()
    current_worker = object()
    active_prepared_context = object()
    session_state = SessionState(
        preparation_input_marker="new.docx:3:def:6000",
        preparation_failed_marker="",
        preparation_event_queue=current_queue,
        preparation_worker=current_worker,
        selected_source_token="new.docx:3:def",
        prepared_run_context=active_prepared_context,
        processing_outcome="running",
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    finalized = []

    current_queue.put(
        PreparationCompleteEvent(
            prepared_run_context=type("PreparedRunContextStub", (), {
                "uploaded_file_token": "old.docx:3:abc",
                "prepared_source_key": "old.docx:3:abc:6000",
            })(),
            upload_marker="old.docx:3:abc:6000",
        )
    )

    processing_runtime.drain_preparation_events(
        reset_run_state=lambda **kwargs: (_ for _ in ()).throw(AssertionError("stale completion must be ignored")),
        set_processing_status=lambda **payload: (_ for _ in ()).throw(AssertionError("stale completion must be ignored")),
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: finalized.append((stage, detail, progress, terminal_kind)),
        push_activity=lambda message: (_ for _ in ()).throw(AssertionError("stale completion must be ignored")),
    )

    assert session_state.preparation_input_marker == "new.docx:3:def:6000"
    assert session_state.preparation_failed_marker == ""
    assert session_state.preparation_worker is current_worker
    assert session_state.preparation_event_queue is current_queue
    assert session_state.selected_source_token == "new.docx:3:def"
    assert session_state.prepared_run_context is active_prepared_context
    assert session_state.processing_outcome == "running"
    assert finalized == []


def test_drain_preparation_events_ignores_stale_failure_marker(monkeypatch):
    current_queue = queue.Queue()
    current_worker = object()
    active_prepared_context = object()
    session_state = SessionState(
        preparation_input_marker="new.docx:3:def:6000",
        preparation_failed_marker="",
        preparation_event_queue=current_queue,
        preparation_worker=current_worker,
        selected_source_token="new.docx:3:def",
        prepared_run_context=active_prepared_context,
        last_error="",
        processing_outcome="running",
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    finalized = []
    activities = []

    current_queue.put(
        PreparationFailedEvent(
            upload_marker="old.docx:3:abc:6000",
            error_message="stale boom",
            error_details={
                "stage": "preparation",
                "severity": "error",
                "user_message": "stale boom",
                "technical_message": "stale boom",
                "error_type": "RuntimeError",
                "recoverable": False,
            },
        )
    )

    processing_runtime.drain_preparation_events(
        reset_run_state=lambda **kwargs: (_ for _ in ()).throw(AssertionError("stale failure must be ignored")),
        set_processing_status=lambda **payload: (_ for _ in ()).throw(AssertionError("stale failure must be ignored")),
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: finalized.append((stage, detail, progress, terminal_kind)),
        push_activity=lambda message: activities.append(message),
    )

    assert session_state.preparation_input_marker == "new.docx:3:def:6000"
    assert session_state.preparation_failed_marker == ""
    assert session_state.preparation_worker is current_worker
    assert session_state.preparation_event_queue is current_queue
    assert session_state.selected_source_token == "new.docx:3:def"
    assert session_state.prepared_run_context is active_prepared_context
    assert session_state.last_error == ""
    assert session_state.processing_outcome == "running"
    assert finalized == []
    assert activities == []


def test_drain_preparation_events_ignores_stale_stopped_marker(monkeypatch):
    current_queue = queue.Queue()
    current_worker = object()
    session_state = SessionState(
        preparation_input_marker="new.docx:3:def:6000",
        preparation_failed_marker="",
        preparation_event_queue=current_queue,
        preparation_worker=current_worker,
        processing_outcome="running",
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    finalized = []
    activities = []

    current_queue.put(
        PreparationStoppedEvent(upload_marker="old.docx:3:abc:6000")
    )

    processing_runtime.drain_preparation_events(
        reset_run_state=lambda **kwargs: (_ for _ in ()).throw(AssertionError("stale stopped marker must be ignored")),
        set_processing_status=lambda **payload: (_ for _ in ()).throw(AssertionError("stale stopped marker must be ignored")),
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: finalized.append((stage, detail, progress, terminal_kind)),
        push_activity=lambda message: activities.append(message),
    )

    assert session_state.preparation_input_marker == "new.docx:3:def:6000"
    assert session_state.preparation_worker is current_worker
    assert session_state.preparation_event_queue is current_queue
    assert session_state.processing_outcome == "running"
    assert finalized == []
    assert activities == []


def test_start_background_preparation_creates_worker_and_status(monkeypatch):
    session_state = SessionState()
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    statuses = []
    activities = []
    payloads = []

    uploaded_payload = processing_runtime.FrozenUploadPayload(
        filename="report.docx",
        content_bytes=b"abc",
        file_size=3,
        content_hash="hash",
        file_token="report.docx:3:ba7816bf8f01cfea",
        source_format="pdf",
        conversion_backend="libreoffice",
    )

    processing_runtime.start_background_preparation(
        worker_target=lambda **kwargs: payloads.append(kwargs),
        reset_run_state=lambda **kwargs: None,
        push_activity=lambda message: activities.append(message),
        set_processing_status=lambda **payload: statuses.append(payload),
        uploaded_payload=uploaded_payload,
        upload_marker="report.docx:3:ba7816bf8f01cfea:6000",
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        processing_operation="audiobook",
        app_config={"processing_operation": "audiobook"},
    )

    session_state.preparation_worker.join(timeout=5)

    assert session_state.preparation_input_marker == "report.docx:3:ba7816bf8f01cfea:6000"
    assert session_state.preparation_event_queue is not None
    assert session_state.preparation_worker is not None
    assert statuses[0]["phase"] == "preparing"
    assert statuses[0]["stage"] == "Файл получен"
    assert statuses[0]["source_format"] == "pdf"
    assert statuses[0]["conversion_backend"] == "libreoffice"
    assert activities == ["Файл получен сервером. Запускаю анализ документа."]
    assert payloads[0]["uploaded_payload"] == uploaded_payload
    assert payloads[0]["processing_operation"] == "audiobook"
    assert payloads[0]["app_config"] == {"processing_operation": "audiobook"}


def test_start_background_preparation_propagates_cached_flag(monkeypatch):
    session_state = SessionState()
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    statuses = []
    activities = []
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(source_name="report.docx", source_bytes=b"abc")
    uploaded_payload = processing_runtime.freeze_uploaded_file(uploaded_file)

    def worker_target(**kwargs):
        kwargs["progress_callback"](
            stage="Подготовка документа",
            detail="cache hit",
            progress=0.9,
            metrics={"cached": True, "block_count": 5, "paragraph_count": 10, "image_count": 1, "source_chars": 2000},
        )

    processing_runtime.start_background_preparation(
        worker_target=worker_target,
        reset_run_state=lambda **kwargs: None,
        push_activity=lambda message: activities.append(message),
        set_processing_status=lambda **payload: statuses.append(payload),
        uploaded_payload=uploaded_payload,
        upload_marker="report.docx:3:ba7816bf8f01cfea:6000",
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
    )

    session_state.preparation_worker.join(timeout=5)
    processing_runtime.drain_preparation_events(
        reset_run_state=lambda **kwargs: None,
        set_processing_status=lambda **payload: statuses.append(payload),
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: None,
        push_activity=lambda message: activities.append(message),
    )
    assert any(payload.get("cached") is True for payload in statuses)
    assert "[Анализ] Подготовка документа: cache hit" in activities


def test_drain_processing_events_moves_restart_source_to_completed_cache_on_success(monkeypatch):
    session_state = SessionState(
        processing_event_queue=queue.Queue(),
        restart_source={"filename": "report.docx", "token": "report.docx:3:abc", "storage_path": "restart.bin", "session_id": "session-a"},
        processing_worker=object(),
        processing_stop_event=object(),
        processing_stop_requested=True,
        restart_session_id="session-a",
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(processing_runtime, "load_restart_source_bytes", lambda restart_source: b"abc")
    cleared = []
    monkeypatch.setattr(processing_runtime, "clear_restart_source", lambda restart_source: cleared.append(restart_source))
    monkeypatch.setattr(
        processing_runtime,
        "store_completed_source",
        lambda **kwargs: {
            "filename": kwargs["source_name"],
            "token": kwargs["source_token"],
            "storage_path": "completed.bin",
            "size": len(kwargs["source_bytes"]),
            "session_id": kwargs["session_id"],
            "storage_kind": "completed",
        },
    )

    session_state.processing_event_queue.put(WorkerCompleteEvent(outcome="succeeded"))

    processing_runtime.drain_processing_events(
        set_processing_status=lambda **payload: None,
        finalize_processing_status=lambda stage, detail, progress: None,
        push_activity=lambda message: None,
        append_log=lambda **payload: None,
        append_image_log=lambda **payload: None,
    )

    assert session_state.completed_source == {
        "filename": "report.docx",
        "token": "report.docx:3:abc",
        "storage_path": "completed.bin",
        "size": 3,
        "session_id": "session-a",
        "storage_kind": "completed",
    }
    assert session_state.restart_source is None
    assert cleared == [{"filename": "report.docx", "token": "report.docx:3:abc", "storage_path": "restart.bin", "session_id": "session-a"}]


def test_drain_processing_events_skips_completed_cache_for_large_sources(monkeypatch):
    session_state = SessionState(
        processing_event_queue=queue.Queue(),
        restart_source={"filename": "report.docx", "token": "report.docx:12:abc", "storage_path": "restart.bin", "session_id": "session-a"},
        processing_worker=object(),
        processing_stop_event=object(),
        processing_stop_requested=True,
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(processing_runtime, "MAX_COMPLETED_SOURCE_BYTES", 4)
    monkeypatch.setattr(processing_runtime, "load_restart_source_bytes", lambda restart_source: b"abcdef")
    cleared = []
    activities = []
    monkeypatch.setattr(processing_runtime, "clear_restart_source", lambda restart_source: cleared.append(restart_source))

    session_state.processing_event_queue.put(WorkerCompleteEvent(outcome="succeeded"))

    processing_runtime.drain_processing_events(
        set_processing_status=lambda **payload: None,
        finalize_processing_status=lambda stage, detail, progress: None,
        push_activity=lambda message: activities.append(message),
        append_log=lambda **payload: None,
        append_image_log=lambda **payload: None,
    )

    assert session_state.completed_source is None
    assert session_state.restart_source is None
    assert len(activities) == 1
    assert "слишком большой" in activities[0].lower()
    assert cleared == [{"filename": "report.docx", "token": "report.docx:12:abc", "storage_path": "restart.bin", "session_id": "session-a"}]


def test_start_background_processing_degrades_gracefully_when_restart_store_fails(monkeypatch):
    session_state = SessionState(restart_session_id="session-a")
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    state_monkey_session = session_state
    monkeypatch.setattr(state.st, "session_state", state_monkey_session)
    state.init_session_state()
    session_state.restart_session_id = "session-a"

    activity_messages = []
    monkeypatch.setattr(processing_runtime, "store_restart_source", lambda **kwargs: (_ for _ in ()).throw(OSError("disk full")))
    log_events = []
    monkeypatch.setattr(processing_runtime, "log_event", lambda *args, **kwargs: log_events.append((args, kwargs)))

    processing_runtime.start_background_processing(
        worker_target=lambda **kwargs: None,
        reset_run_state=state.reset_run_state,
        push_activity=lambda message: activity_messages.append(message),
        set_processing_status=lambda **kwargs: None,
        uploaded_filename="report.docx",
        uploaded_token="report.docx:3:abc",
        source_bytes=b"abc",
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=["paragraph"],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
    )

    session_state.processing_worker.join(timeout=5)

    assert session_state.restart_source is None
    assert session_state.processing_worker is not None
    assert any("restart" in message.lower() for message in activity_messages)
    assert len(log_events) == 1


def test_start_background_processing_preserves_prepared_context(monkeypatch):
    prepared_run_context = object()
    session_state = SessionState(
        restart_session_id="session-a",
        prepared_run_context=prepared_run_context,
        latest_preparation_summary={"stage": "Документ подготовлен"},
        preparation_input_marker="report.docx:3:abc:6000",
        prepared_source_key="report.docx:3:abc:6000",
        preparation_cache={"report.docx:3:abc:6000": {"cached": True}},
    )
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(state.st, "session_state", session_state)
    state.init_session_state()
    session_state.restart_session_id = "session-a"
    monkeypatch.setattr(
        processing_runtime,
        "store_restart_source",
        lambda **kwargs: {
            "filename": kwargs["source_name"],
            "token": kwargs["source_token"],
            "storage_path": "restart.bin",
            "session_id": kwargs["session_id"],
        },
    )

    processing_runtime.start_background_processing(
        worker_target=lambda **kwargs: None,
        reset_run_state=state.reset_run_state,
        push_activity=lambda message: None,
        set_processing_status=lambda **kwargs: None,
        uploaded_filename="report.docx",
        uploaded_token="report.docx:3:abc",
        source_bytes=b"abc",
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=["paragraph"],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
    )

    session_state.processing_worker.join(timeout=5)

    assert session_state.prepared_run_context is prepared_run_context
    assert session_state.latest_preparation_summary == {"stage": "Документ подготовлен"}
    assert session_state.preparation_input_marker == "report.docx:3:abc:6000"
    assert session_state.prepared_source_key == "report.docx:3:abc:6000"
    assert session_state.preparation_cache == {"report.docx:3:abc:6000": {"cached": True}}
    assert session_state.processing_outcome == "running"


def test_start_background_processing_delegates_p1a_start_state_to_state_owner(monkeypatch):
    session_state = SessionState(restart_session_id="session-a")
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(state.st, "session_state", session_state)
    state.init_session_state()
    session_state.restart_session_id = "session-a"
    monkeypatch.setattr(
        processing_runtime,
        "store_restart_source",
        lambda **kwargs: {
            "filename": kwargs["source_name"],
            "token": kwargs["source_token"],
            "storage_path": "restart.bin",
            "session_id": kwargs["session_id"],
        },
    )
    start_calls = []

    original_apply_processing_start = processing_runtime.apply_processing_start

    def tracking_apply_processing_start(**kwargs):
        start_calls.append(kwargs)
        return original_apply_processing_start(**kwargs)

    monkeypatch.setattr(processing_runtime, "apply_processing_start", tracking_apply_processing_start)

    processing_runtime.start_background_processing(
        worker_target=lambda **kwargs: None,
        reset_run_state=state.reset_run_state,
        push_activity=lambda message: None,
        set_processing_status=lambda **kwargs: None,
        uploaded_filename="report.docx",
        uploaded_token="report.docx:3:abc",
        source_bytes=b"abc",
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=["paragraph"],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
    )

    session_state.processing_worker.join(timeout=5)

    assert len(start_calls) == 1
    assert start_calls[0]["uploaded_filename"] == "report.docx"
    assert start_calls[0]["uploaded_token"] == "report.docx:3:abc"
    assert start_calls[0]["image_mode"] == "safe"
    assert start_calls[0]["worker"] is session_state.processing_worker
    assert start_calls[0]["event_queue"] is session_state.processing_event_queue
    assert start_calls[0]["stop_event"] is session_state.processing_stop_event


def test_start_background_processing_passes_selected_segment_ids_to_worker(monkeypatch):
    session_state = SessionState(restart_session_id="session-a")
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(state.st, "session_state", session_state)
    state.init_session_state()
    session_state.restart_session_id = "session-a"
    monkeypatch.setattr(
        processing_runtime,
        "store_restart_source",
        lambda **kwargs: {
            "filename": kwargs["source_name"],
            "token": kwargs["source_token"],
            "storage_path": "restart.bin",
            "session_id": kwargs["session_id"],
        },
    )
    captured = {}

    def worker_target(**kwargs):
        captured.update(kwargs)

    processing_runtime.start_background_processing(
        worker_target=worker_target,
        reset_run_state=state.reset_run_state,
        push_activity=lambda message: None,
        set_processing_status=lambda **kwargs: None,
        uploaded_filename="report.docx",
        uploaded_token="report.docx:3:abc",
        source_bytes=b"abc",
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        selected_segment_ids=["seg_0001", "seg_0002"],
        source_paragraphs=["paragraph"],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
    )

    session_state.processing_worker.join(timeout=5)

    assert captured["selected_segment_ids"] == ["seg_0001", "seg_0002"]


def test_start_background_processing_passes_none_selected_segment_ids_to_worker(monkeypatch):
    session_state = SessionState(restart_session_id="session-a")
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(state.st, "session_state", session_state)
    state.init_session_state()
    session_state.restart_session_id = "session-a"
    monkeypatch.setattr(
        processing_runtime,
        "store_restart_source",
        lambda **kwargs: {
            "filename": kwargs["source_name"],
            "token": kwargs["source_token"],
            "storage_path": "restart.bin",
            "session_id": kwargs["session_id"],
        },
    )
    captured = {}

    def worker_target(**kwargs):
        captured.update(kwargs)

    processing_runtime.start_background_processing(
        worker_target=worker_target,
        reset_run_state=state.reset_run_state,
        push_activity=lambda message: None,
        set_processing_status=lambda **kwargs: None,
        uploaded_filename="report.docx",
        uploaded_token="report.docx:3:abc",
        source_bytes=b"abc",
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=["paragraph"],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
    )

    session_state.processing_worker.join(timeout=5)

    assert captured["selected_segment_ids"] is None


def test_request_processing_stop_delegates_to_state_owner(monkeypatch):
    calls = []
    monkeypatch.setattr(processing_runtime, "request_processing_stop_via_state", lambda: calls.append("called"))

    processing_runtime.request_processing_stop()

    assert calls == ["called"]


def test_get_current_result_bundle_reads_p1a_source_identity_via_state_helpers(monkeypatch):
    session_state = SessionState(latest_docx_bytes=b"docx", latest_markdown="md")
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(processing_runtime, "get_latest_source_name", lambda: "report.docx")
    monkeypatch.setattr(processing_runtime, "get_latest_source_token", lambda: "report.docx:3:abc")

    result = processing_runtime.get_current_result_bundle()

    assert result == {
        "source_name": "report.docx",
        "source_token": "report.docx:3:abc",
        "docx_bytes": b"docx",
        "markdown_text": "md",
        "narration_text": None,
        "processing_operation": "edit",
        "audiobook_postprocess_enabled": False,
        "quality_warning": None,
    }


def test_get_current_result_bundle_allows_narration_only_result(monkeypatch):
    session_state = SessionState(latest_docx_bytes=None, latest_markdown="md", latest_narration_text="[thoughtful] text")
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(processing_runtime, "get_latest_source_name", lambda: "report.docx")
    monkeypatch.setattr(processing_runtime, "get_latest_source_token", lambda: "report.docx:3:abc")

    result = processing_runtime.get_current_result_bundle()

    assert result == {
        "source_name": "report.docx",
        "source_token": "report.docx:3:abc",
        "docx_bytes": None,
        "markdown_text": "md",
        "narration_text": "[thoughtful] text",
        "processing_operation": "edit",
        "audiobook_postprocess_enabled": False,
        "quality_warning": None,
    }


def test_get_current_result_bundle_rejects_incomplete_standalone_audiobook_result(monkeypatch):
    session_state = SessionState(latest_docx_bytes=None, latest_markdown="md", latest_narration_text="[thoughtful] text")
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(processing_runtime, "get_latest_processing_operation", lambda: "audiobook")

    assert processing_runtime.get_current_result_bundle() is None


def test_build_result_bundle_preserves_explicit_mode_metadata():
    result = processing_runtime.build_result_bundle(
        source_name="report.docx",
        source_token="report.docx:3:abc",
        docx_bytes=b"docx",
        markdown_text="md",
        narration_text="[thoughtful] narration",
        processing_operation="translate",
        audiobook_postprocess_enabled=True,
    )

    assert result["processing_operation"] == "translate"
    assert result["audiobook_postprocess_enabled"] is True


def test_freeze_uploaded_file_normalizes_legacy_doc_payload(monkeypatch):
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="legacy.doc",
        source_bytes=bytes.fromhex("D0CF11E0A1B11AE1") + b"legacy-binary",
    )
    monkeypatch.setattr(
        processing_runtime,
        "_convert_legacy_doc_to_docx",
        lambda **kwargs: (b"converted-docx", "antiword+pandoc"),
    )

    payload = processing_runtime.freeze_uploaded_file(uploaded_file)

    assert payload.filename == "legacy.docx"
    assert payload.content_bytes == b"converted-docx"
    assert payload.file_size == len(b"converted-docx")
    assert payload.file_token.startswith("legacy.docx:")


def test_build_uploaded_file_token_renames_zip_payloads_with_docx_magic_to_docx_extension():
    token = processing_runtime.build_uploaded_file_token(
        source_name="misnamed.doc",
        source_bytes=b"PK\x03\x04not-really-a-full-docx",
    )

    assert token.startswith("misnamed.docx:")


def test_detect_uploaded_document_format_rejects_non_doc_ole2_suffix() -> None:
    detected = processing_runtime._detect_uploaded_document_format(
        filename="worksheet.xls",
        source_bytes=bytes.fromhex("D0CF11E0A1B11AE1") + b"ole2-payload",
    )

    assert detected == "unknown"


def test_run_completed_process_raises_timeout_error(monkeypatch):
    def run_stub(*args, **kwargs):
        raise subprocess.TimeoutExpired(cmd=args[0], timeout=kwargs["timeout"])

    monkeypatch.setattr(processing_runtime.subprocess, "run", run_stub)

    with pytest.raises(RuntimeError, match="Превышено время ожидания"):
        processing_runtime._run_completed_process(["soffice"], error_message="boom")


def test_run_completed_process_cleans_process_group_on_timeout(monkeypatch):
    process_instances = []
    terminated = []

    class PopenStub:
        pid = 1234
        returncode = None

        def __init__(self, command, **kwargs):
            self.command = command
            self.kwargs = kwargs
            self.communicate_calls = 0
            process_instances.append(self)

        def communicate(self, timeout=None):
            self.communicate_calls += 1
            if self.communicate_calls == 1:
                raise subprocess.TimeoutExpired(
                    cmd=self.command,
                    timeout=float(timeout if timeout is not None else 0),
                )
            self.returncode = -15
            return "", ""

    monkeypatch.setattr(processing_runtime.subprocess, "Popen", PopenStub)
    monkeypatch.setattr(processing_runtime, "_terminate_process_tree", lambda process: terminated.append(process))

    with pytest.raises(RuntimeError, match="Превышено время ожидания"):
        processing_runtime._run_completed_process(
            ["soffice"],
            error_message="boom",
            cleanup_process_group=True,
            timeout_seconds=1,
        )

    assert terminated == process_instances
    assert process_instances[0].communicate_calls == 2


def test_convert_legacy_doc_to_docx_falls_back_to_antiword_when_soffice_fails(monkeypatch):
    calls = []

    monkeypatch.setattr(
        processing_runtime.shutil,
        "which",
        lambda name: {
            "soffice": "/usr/bin/soffice",
            "libreoffice": None,
            "antiword": "/usr/bin/antiword",
        }.get(name),
    )

    def soffice_stub(**kwargs):
        calls.append("soffice")
        raise RuntimeError("soffice failed")

    def antiword_stub(**kwargs):
        calls.append("antiword")
        return b"converted-docx"

    monkeypatch.setattr(processing_runtime, "_convert_legacy_doc_with_soffice", soffice_stub)
    monkeypatch.setattr(processing_runtime, "_convert_legacy_doc_with_antiword", antiword_stub)

    converted_bytes, backend = processing_runtime._convert_legacy_doc_to_docx(
        filename="legacy.doc",
        source_bytes=bytes.fromhex("D0CF11E0A1B11AE1") + b"legacy",
    )

    assert converted_bytes == b"converted-docx"
    assert backend == "antiword+pandoc"
    assert calls == ["soffice", "antiword"]


def test_convert_legacy_doc_with_soffice_cleans_process_group(monkeypatch, tmp_path):
    docx_bytes = b"PK\x03\x04converted-docx"
    cleanup_flags = []

    def fake_run_completed_process(command, *, error_message, text=True, timeout_seconds=120, cleanup_process_group=False):
        cleanup_flags.append(cleanup_process_group)
        outdir = Path(command[command.index("--outdir") + 1])
        input_path = Path(command[-1])
        output_path = outdir / input_path.with_suffix(".docx").name
        output_path.write_bytes(docx_bytes)
        return object()

    monkeypatch.setattr(processing_runtime, "_run_completed_process", fake_run_completed_process)

    converted = processing_runtime._convert_legacy_doc_with_soffice(
        soffice_path="/usr/bin/soffice",
        filename="legacy.doc",
        source_bytes=bytes.fromhex("D0CF11E0A1B11AE1") + b"legacy",
    )

    assert converted == docx_bytes
    assert cleanup_flags == [True]


def test_legacy_doc_conversion_available_requires_pandoc_for_antiword_path(monkeypatch):
    monkeypatch.setattr(
        processing_runtime.shutil,
        "which",
        lambda name: {
            "soffice": None,
            "libreoffice": None,
            "antiword": "/usr/bin/antiword",
        }.get(name),
    )

    class _PandocStub:
        @staticmethod
        def get_pandoc_version():
            raise OSError("pandoc missing")

    monkeypatch.setitem(sys.modules, "pypandoc", _PandocStub)

    assert processing_runtime.legacy_doc_conversion_available() is False


def test_legacy_doc_conversion_available_accepts_soffice_without_antiword(monkeypatch):
    monkeypatch.setattr(
        processing_runtime.shutil,
        "which",
        lambda name: {
            "soffice": "/usr/bin/soffice",
            "libreoffice": None,
            "antiword": None,
        }.get(name),
    )

    assert processing_runtime.legacy_doc_conversion_available() is True


def test_build_uploaded_file_token_for_legacy_doc_is_stable_across_converter_outputs(monkeypatch):
    converted_outputs = [b"converted-docx-a", b"converted-docx-b"]

    def convert_stub(**kwargs):
        return converted_outputs.pop(0), "libreoffice"

    monkeypatch.setattr(processing_runtime, "_convert_legacy_doc_to_docx", convert_stub)

    source_bytes = bytes.fromhex("D0CF11E0A1B11AE1") + b"same-legacy-doc"
    first = processing_runtime.build_uploaded_file_token(source_name="legacy.doc", source_bytes=source_bytes)
    second = processing_runtime.build_uploaded_file_token(source_name="legacy.doc", source_bytes=source_bytes)

    assert first == second


def test_detect_uploaded_document_format_recognizes_pdf_magic_bytes() -> None:
    detected = processing_runtime._detect_uploaded_document_format(
        filename="source.bin",
        source_bytes=b"%PDF-1.7\ncontent",
    )

    assert detected == "pdf"


def test_detect_uploaded_document_format_recognizes_pdf_suffix_fallback() -> None:
    detected = processing_runtime._detect_uploaded_document_format(
        filename="source.pdf",
        source_bytes=b"not-really-a-pdf-header",
    )

    assert detected == "pdf"


def test_normalize_uploaded_pdf_converts_to_docx(monkeypatch):
    pdf_bytes = b"%PDF-1.7\ncontent"
    docx_bytes = b"PK\x03\x04converted-docx"
    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_text_layer_to_docx",
        lambda **kwargs: (docx_bytes, "pdf-text-layer"),
    )

    normalized = processing_runtime.normalize_uploaded_document(filename="source.pdf", source_bytes=pdf_bytes)

    assert normalized.original_filename == "source.pdf"
    assert normalized.filename == "source.docx"
    assert normalized.content_bytes == docx_bytes
    assert normalized.source_format == "pdf"
    assert normalized.conversion_backend == "pdf-text-layer"


def test_convert_pdf_to_docx_uses_writer_pdf_import_filter(monkeypatch):
    pdf_bytes = b"%PDF-1.7\ncontent"
    docx_bytes = b"PK\x03\x04converted-docx"
    commands = []

    monkeypatch.setattr(
        processing_runtime.shutil,
        "which",
        lambda name: "/usr/bin/soffice" if name == "soffice" else None,
    )

    def fake_run_completed_process(command, *, error_message, text=True, timeout_seconds=120, cleanup_process_group=False):
        commands.append(command)
        outdir = Path(command[command.index("--outdir") + 1])
        input_path = Path(command[-1])
        output_path = outdir / input_path.with_suffix(".docx").name
        output_path.write_bytes(docx_bytes)
        return object()

    monkeypatch.setattr(processing_runtime, "_run_completed_process", fake_run_completed_process)

    converted_bytes, backend = processing_runtime._convert_pdf_to_docx(filename="source.pdf", source_bytes=pdf_bytes)

    assert converted_bytes == docx_bytes
    assert backend == "libreoffice"
    assert commands[0][2] == "--infilter=writer_pdf_import"


def test_normalize_uploaded_pdf_surfaces_text_layer_import_error(monkeypatch):
    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_text_layer_to_docx",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("pdf_text_layer_import_not_promising:no_text_layer")),
    )
    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_to_docx",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("libreoffice should not run")),
    )

    with pytest.raises(RuntimeError, match="pdf_text_layer_import_not_promising:no_text_layer"):
        processing_runtime.normalize_uploaded_document(
            filename="source.pdf",
            source_bytes=b"%PDF-1.7\ncontent",
        )


def test_build_uploaded_file_token_for_pdf_is_stable_across_converter_outputs(monkeypatch):
    converted_outputs = [b"PK\x03\x04converted-docx-a", b"PK\x03\x04converted-docx-b"]

    def convert_stub(**kwargs):
        return converted_outputs.pop(0), "pdf-text-layer"

    monkeypatch.setattr(processing_runtime, "_convert_pdf_text_layer_to_docx", convert_stub)

    source_bytes = b"%PDF-1.7\nsame-pdf"
    first = processing_runtime.build_uploaded_file_token(source_name="source.pdf", source_bytes=source_bytes)
    second = processing_runtime.build_uploaded_file_token(source_name="source.pdf", source_bytes=source_bytes)

    assert first == second


def test_normalize_uploaded_pdf_uses_text_layer_import_by_default(monkeypatch):
    monkeypatch.delenv("DOCXAI_PDF_TEXT_LAYER_IMPORT_ENABLED", raising=False)
    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_text_layer_to_docx",
        lambda **kwargs: (b"text-layer-docx", "pdf-text-layer"),
    )
    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_to_docx",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("libreoffice should not run")),
    )

    normalized = processing_runtime.normalize_uploaded_document(
        filename="source.pdf",
        source_bytes=b"%PDF-1.7\nsource",
    )

    assert normalized.filename == "source.docx"
    assert normalized.content_bytes == b"text-layer-docx"
    assert normalized.source_format == "pdf"
    assert normalized.conversion_backend == "pdf-text-layer"


def test_normalize_uploaded_pdf_ignores_legacy_libreoffice_override(monkeypatch):
    monkeypatch.setenv("DOCXAI_PDF_TEXT_LAYER_IMPORT_ENABLED", "0")
    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_text_layer_to_docx",
        lambda **kwargs: (b"text-layer-docx", "pdf-text-layer"),
    )
    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_to_docx",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("libreoffice should not run")),
    )

    normalized = processing_runtime.normalize_uploaded_document(
        filename="source.pdf",
        source_bytes=b"%PDF-1.7\nsource",
    )

    assert normalized.conversion_backend == "pdf-text-layer"
    assert normalized.content_bytes == b"text-layer-docx"


def test_resolve_upload_contract_uses_original_pdf_bytes_for_source_identity(monkeypatch):
    source_bytes = b"%PDF-1.7\nsame-pdf"
    converted_bytes = b"PK\x03\x04converted-docx"
    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_text_layer_to_docx",
        lambda **kwargs: (converted_bytes, "pdf-text-layer"),
    )

    contract = processing_runtime.resolve_upload_contract(filename="source.pdf", source_bytes=source_bytes)

    assert contract.normalized_document.filename == "source.docx"
    assert contract.normalized_document.content_bytes == converted_bytes
    assert contract.source_identity.original_filename == "source.pdf"
    assert contract.source_identity.source_bytes == source_bytes
    assert contract.source_identity.token_size == len(source_bytes)
    assert contract.source_identity.token_hash == hashlib.sha256(source_bytes).hexdigest()[:16]


def test_convert_pdf_to_docx_falls_back_to_single_generated_docx(monkeypatch):
    pdf_bytes = b"%PDF-1.7\ncontent"
    generated_bytes = b"PK\x03\x04generated-docx"

    monkeypatch.setattr(
        processing_runtime.shutil,
        "which",
        lambda name: "/usr/bin/soffice" if name == "soffice" else None,
    )

    def fake_run_completed_process(command, *, error_message, text=True, timeout_seconds=120, cleanup_process_group=False):
        outdir = Path(command[command.index("--outdir") + 1])
        (outdir / "unexpected-name.docx").write_bytes(generated_bytes)
        return object()

    monkeypatch.setattr(processing_runtime, "_run_completed_process", fake_run_completed_process)

    converted_bytes, backend = processing_runtime._convert_pdf_to_docx(
        filename="source.pdf",
        source_bytes=pdf_bytes,
    )

    assert converted_bytes == generated_bytes
    assert backend == "libreoffice"


def test_convert_pdf_to_docx_surfaces_converter_failure(monkeypatch):
    monkeypatch.setattr(
        processing_runtime.shutil,
        "which",
        lambda name: "/usr/bin/soffice" if name == "soffice" else None,
    )
    monkeypatch.setattr(
        processing_runtime,
        "_run_completed_process",
        lambda *args, **kwargs: (_ for _ in ()).throw(RuntimeError("Не удалось конвертировать PDF через LibreOffice. broken pdf")),
    )

    with pytest.raises(RuntimeError, match="Не удалось конвертировать PDF через LibreOffice"):
        processing_runtime._convert_pdf_to_docx(
            filename="broken.pdf",
            source_bytes=b"%PDF-1.7\nbroken",
        )


def test_resolve_upload_contract_separates_source_identity_from_normalized_payload(monkeypatch):
    source_bytes = bytes.fromhex("D0CF11E0A1B11AE1") + b"same-legacy-doc"
    monkeypatch.setattr(
        processing_runtime,
        "_convert_legacy_doc_to_docx",
        lambda **kwargs: (b"converted-docx", "libreoffice"),
    )

    contract = processing_runtime.resolve_upload_contract(filename="legacy.doc", source_bytes=source_bytes)
    payload = processing_runtime.freeze_resolved_upload(contract)

    assert contract.source_identity.original_filename == "legacy.doc"
    assert contract.source_identity.source_bytes == source_bytes
    assert contract.normalized_document.filename == "legacy.docx"
    assert contract.normalized_document.content_bytes == b"converted-docx"
    assert payload.filename == "legacy.docx"
    assert payload.content_bytes == b"converted-docx"
    assert payload.file_token == contract.file_token
    assert payload.file_token.endswith(f":{contract.source_identity.token_hash}")


def test_freeze_uploaded_file_lightweight_for_pdf_defers_conversion(monkeypatch):
    """PDF freeze on the main thread MUST NOT call LibreOffice; it stays cheap."""

    converter_calls: list[dict[str, object]] = []

    def fail_convert(**kwargs):  # pragma: no cover - guard
        converter_calls.append(kwargs)
        raise AssertionError("PDF conversion must happen in the worker, not on freeze")

    monkeypatch.setattr(processing_runtime, "_convert_pdf_to_docx", fail_convert)
    pdf_bytes = b"%PDF-1.4\n%fakepdf\n"
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="book.pdf",
        source_bytes=pdf_bytes,
    )

    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)

    assert payload.source_format == "pdf"
    assert payload.conversion_backend is None
    assert payload.content_bytes == pdf_bytes
    assert payload.filename == "book.pdf"
    assert payload.file_token.startswith("book.docx:")
    assert converter_calls == []


def test_freeze_uploaded_file_lightweight_passthrough_for_docx():
    docx_bytes = b"PK\x03\x04docx-bytes"
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="report.docx",
        source_bytes=docx_bytes,
    )

    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)

    assert payload.source_format == "docx"
    assert payload.content_bytes == docx_bytes
    assert payload.filename == "report.docx"


def test_materialize_uploaded_payload_runs_pdf_conversion_with_progress(monkeypatch):
    _clear_materialized_upload_cache()
    pdf_bytes = b"%PDF-1.4\n%fakepdf\n"
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="book.pdf",
        source_bytes=pdf_bytes,
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)

    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_text_layer_to_docx",
        lambda **kwargs: (b"converted-docx-bytes", "pdf-text-layer"),
    )

    events: list[dict[str, object]] = []

    def progress_cb(**kwargs):
        events.append(kwargs)

    materialized = processing_runtime.materialize_uploaded_payload(
        payload, progress_callback=progress_cb
    )

    assert materialized.source_format == "pdf"
    assert materialized.conversion_backend == "pdf-text-layer"
    assert materialized.content_bytes == b"converted-docx-bytes"
    assert materialized.filename == "book.docx"
    # token MUST be preserved so prepared_source_key cache stays stable
    assert materialized.file_token == payload.file_token
    stages = [e["stage"] for e in events]
    assert "Импорт PDF" in stages
    assert "DOCX готов" in stages


def test_materialize_uploaded_payload_uses_text_layer_import_by_default(monkeypatch):
    _clear_materialized_upload_cache()
    monkeypatch.delenv("DOCXAI_PDF_TEXT_LAYER_IMPORT_ENABLED", raising=False)
    pdf_bytes = b"%PDF-1.4\n%text-layer\n"
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="book.pdf",
        source_bytes=pdf_bytes,
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)

    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_text_layer_to_docx",
        lambda **kwargs: (b"text-layer-docx", "pdf-text-layer"),
    )
    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_to_docx",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("libreoffice should not run")),
    )

    events: list[dict[str, object]] = []
    materialized = processing_runtime.materialize_uploaded_payload(
        payload,
        progress_callback=lambda **kwargs: events.append(kwargs),
    )

    assert materialized.conversion_backend == "pdf-text-layer"
    assert materialized.content_bytes == b"text-layer-docx"
    assert materialized.filename == "book.docx"
    assert materialized.file_token == payload.file_token
    assert any("text-layer" in str(event["detail"]) for event in events if event["stage"] == "Импорт PDF")


def test_materialize_uploaded_payload_surfaces_text_layer_rejection(monkeypatch):
    _clear_materialized_upload_cache()
    monkeypatch.delenv("DOCXAI_PDF_TEXT_LAYER_IMPORT_ENABLED", raising=False)
    pdf_bytes = b"%PDF-1.4\n%fallback\n"
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="fallback.pdf",
        source_bytes=pdf_bytes,
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)

    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_text_layer_to_docx",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("pdf_text_layer_import_not_promising")),
    )
    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_to_docx",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("libreoffice should not run")),
    )

    with pytest.raises(RuntimeError, match="pdf_text_layer_import_not_promising"):
        processing_runtime.materialize_uploaded_payload(payload, progress_callback=None)


def test_materialized_upload_cache_uses_text_layer_when_legacy_env_is_set(monkeypatch):
    _clear_materialized_upload_cache()
    pdf_bytes = b"%PDF-1.4\n%same-token\n"
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="same.pdf",
        source_bytes=pdf_bytes,
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)
    calls: list[str] = []

    def convert_text_layer(**kwargs):
        calls.append("text-layer")
        return b"text-layer-docx", "pdf-text-layer"

    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_to_docx",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("libreoffice should not run")),
    )
    monkeypatch.setattr(processing_runtime, "_convert_pdf_text_layer_to_docx", convert_text_layer)

    monkeypatch.setenv("DOCXAI_PDF_TEXT_LAYER_IMPORT_ENABLED", "0")
    first = processing_runtime.materialize_uploaded_payload(payload, progress_callback=None)
    monkeypatch.delenv("DOCXAI_PDF_TEXT_LAYER_IMPORT_ENABLED", raising=False)
    second = processing_runtime.materialize_uploaded_payload(payload, progress_callback=None)

    assert first.conversion_backend == "pdf-text-layer"
    assert second.conversion_backend == "pdf-text-layer"
    assert calls == ["text-layer"]


def test_pdf_text_layer_generated_docx_does_not_encode_false_bold_or_italic(monkeypatch):
    from docxaicorrector.pdf_import import text_layer_quality

    spans = [
        PdfTextSpan(
            page_number=1,
            text="CONTENTS",
            x0=50,
            top=70,
            x1=250,
            bottom=90,
            page_height=800,
            font_size=18,
            is_bold=True,
        ),
        PdfTextSpan(
            page_number=1,
            text="Plain body line",
            x0=50,
            top=100,
            x1=450,
            bottom=112,
            page_height=800,
            font_size=10,
        )
    ]
    monkeypatch.setattr(text_layer_quality, "extract_pdf_text_spans_with_pdfminer", lambda path: spans)
    monkeypatch.setattr(
        text_layer_quality,
        "build_text_layer_quality_report",
        lambda spans: SimpleNamespace(
            decision="promising",
            decision_reasons=(),
            body_text_ratio=1.0,
        ),
    )

    docx_bytes, backend = processing_runtime._convert_pdf_text_layer_to_docx(
        filename="plain.pdf",
        source_bytes=b"%PDF-1.4\n",
    )
    paragraphs = extract_paragraph_units_from_docx(
        processing_runtime.build_in_memory_uploaded_file(
            source_name="plain.docx",
            source_bytes=docx_bytes,
        )
    )

    assert backend == "pdf-text-layer"
    assert [paragraph.text for paragraph in paragraphs] == ["CONTENTS", "Plain body line"]


def test_pdf_text_layer_generated_docx_preserves_span_level_bold_and_italic(monkeypatch):
    from docxaicorrector.pdf_import import text_layer_quality

    spans = [
        PdfTextSpan(
            page_number=1,
            text="Normal line",
            x0=50,
            top=100,
            x1=450,
            bottom=112,
            page_height=800,
            font_size=10,
        ),
        PdfTextSpan(
            page_number=1,
            text="emphasis line",
            x0=50,
            top=114,
            x1=450,
            bottom=126,
            page_height=800,
            font_size=10,
            is_italic=True,
        ),
        PdfTextSpan(
            page_number=1,
            text="bold phrase with enough words to avoid heading candidate in this importer",
            x0=50,
            top=128,
            x1=450,
            bottom=140,
            page_height=800,
            font_size=10,
            is_bold=True,
        ),
    ]
    monkeypatch.setattr(text_layer_quality, "extract_pdf_text_spans_with_pdfminer", lambda path: spans)
    monkeypatch.setattr(
        text_layer_quality,
        "build_text_layer_quality_report",
        lambda spans: SimpleNamespace(
            decision="promising",
            decision_reasons=(),
            body_text_ratio=1.0,
        ),
    )

    docx_bytes, backend = processing_runtime._convert_pdf_text_layer_to_docx(
        filename="formatting.pdf",
        source_bytes=b"%PDF-1.4\n",
    )
    paragraphs = extract_paragraph_units_from_docx(
        processing_runtime.build_in_memory_uploaded_file(
            source_name="formatting.docx",
            source_bytes=docx_bytes,
        )
    )

    assert backend == "pdf-text-layer"
    assert [paragraph.text for paragraph in paragraphs] == [
        "Normal line *emphasis line* **bold phrase with enough words to avoid heading candidate in this importer**",
    ]


def test_pdf_text_layer_generated_docx_preserves_pdf_images_as_docx_placeholders(monkeypatch):
    from docxaicorrector.pdf_import import images as pdf_images
    from docxaicorrector.pdf_import import text_layer_quality

    spans = [
        PdfTextSpan(
            page_number=1,
            text="Text before image.",
            x0=50,
            top=100,
            x1=450,
            bottom=112,
            page_height=800,
            font_size=10,
        ),
        PdfTextSpan(
            page_number=1,
            text="Text after image.",
            x0=50,
            top=220,
            x1=450,
            bottom=232,
            page_height=800,
            font_size=10,
        ),
    ]
    image_bytes = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
        b"\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4"
        b"\x89\x00\x00\x00\nIDATx\x9cc\xf8\x0f\x00\x01\x01"
        b"\x01\x00\x18\xdd\x8d\xb0\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    image_objects = [
        PdfImageObject(
            page_number=1,
            x0=50,
            top=150,
            x1=150,
            bottom=210,
            page_height=800,
            image_bytes=image_bytes,
            mime_type="image/png",
            source_index=150,
        )
    ]
    monkeypatch.setattr(text_layer_quality, "extract_pdf_text_spans_with_pdfminer", lambda path: spans)
    monkeypatch.setattr(pdf_images, "extract_pdf_images_with_pdfminer", lambda path: image_objects)
    monkeypatch.setattr(
        text_layer_quality,
        "build_text_layer_quality_report",
        lambda spans: SimpleNamespace(
            decision="promising",
            decision_reasons=(),
            body_text_ratio=1.0,
        ),
    )

    docx_bytes, backend = processing_runtime._convert_pdf_text_layer_to_docx(
        filename="with-image.pdf",
        source_bytes=b"%PDF-1.4\n",
    )
    paragraphs, image_assets = extract_document_content_from_docx(
        processing_runtime.build_in_memory_uploaded_file(
            source_name="with-image.docx",
            source_bytes=docx_bytes,
        )
    )

    assert backend == "pdf-text-layer"
    assert [paragraph.text for paragraph in paragraphs] == [
        "Text before image.",
        "[[DOCX_IMAGE_img_001]]",
        "Text after image.",
    ]
    assert len(image_assets) == 1
    assert image_assets[0].mime_type == "image/png"
    assert image_assets[0].original_bytes == image_bytes


def test_pdf_text_layer_generated_docx_counts_unrenderable_images_as_dropped(monkeypatch):
    from docxaicorrector.pdf_import import images as pdf_images
    from docxaicorrector.pdf_import import text_layer_quality

    events: list[tuple[int, str, dict[str, object]]] = []
    spans = [
        PdfTextSpan(
            page_number=1,
            text="Text around failed image.",
            x0=50,
            top=100,
            x1=450,
            bottom=112,
            page_height=800,
            font_size=10,
        )
    ]
    image_objects = [
        PdfImageObject(
            page_number=1,
            x0=50,
            top=150,
            x1=150,
            bottom=210,
            page_height=800,
            image_bytes=b"not-a-real-png",
            mime_type="image/png",
            source_index=150,
        )
    ]
    monkeypatch.setattr(text_layer_quality, "extract_pdf_text_spans_with_pdfminer", lambda path: spans)
    monkeypatch.setattr(pdf_images, "extract_pdf_images_with_pdfminer", lambda path: image_objects)
    monkeypatch.setattr(
        text_layer_quality,
        "build_text_layer_quality_report",
        lambda spans: SimpleNamespace(
            decision="promising",
            decision_reasons=(),
            body_text_ratio=1.0,
        ),
    )
    monkeypatch.setattr(
        processing_runtime,
        "log_event",
        lambda level, event, message, **context: events.append((level, event, context)),
    )

    docx_bytes, backend = processing_runtime._convert_pdf_text_layer_to_docx(
        filename="with-broken-image.pdf",
        source_bytes=b"%PDF-1.4\n",
    )
    paragraphs, image_assets = extract_document_content_from_docx(
        processing_runtime.build_in_memory_uploaded_file(
            source_name="with-broken-image.docx",
            source_bytes=docx_bytes,
        )
    )

    drop_event = next(event for event in events if event[1] == "pdf_text_layer_image_render_dropped")
    success_event = next(event for event in events if event[1] == "pdf_text_layer_import_succeeded")

    assert backend == "pdf-text-layer"
    assert [paragraph.text for paragraph in paragraphs] == ["Text around failed image."]
    assert image_assets == []
    assert drop_event[2]["source_index"] == 150
    assert drop_event[2]["mime_type"] == "image/png"
    assert success_event[2]["image_count"] == 1
    assert success_event[2]["images_emitted"] == 0
    assert success_event[2]["images_dropped"] == 1


def test_pdf_text_layer_import_can_run_ocr_fallback_for_scanned_pdf(monkeypatch, tmp_path):
    from docxaicorrector.pdf_import import text_layer_quality

    calls: list[list[str]] = []
    input_spans: list[PdfTextSpan] = []
    ocr_spans = [
        PdfTextSpan(
            page_number=1,
            text="OCR body text.",
            x0=50,
            top=100,
            x1=450,
            bottom=112,
            page_height=800,
            font_size=10,
        )
    ]

    monkeypatch.setenv("DOCXAI_PDF_OCR_IMPORT_ENABLED", "1")
    monkeypatch.setenv("DOCXAI_PDF_OCR_LANGUAGES", "eng+deu")
    monkeypatch.setattr(processing_runtime.shutil, "which", lambda name: f"/usr/bin/{name}")
    monkeypatch.setattr(text_layer_quality, "extract_pdf_text_spans_with_pdfminer", lambda path: ocr_spans if str(path).endswith(".ocr.pdf") else input_spans)

    def fake_quality(spans):
        return SimpleNamespace(
            decision="promising" if spans else "scanned_or_unsupported",
            decision_reasons=("empty_text_layer",) if not spans else (),
            body_text_ratio=1.0 if spans else 0.0,
        )

    def fake_run_completed_process(command, **kwargs):
        calls.append([str(part) for part in command])
        Path(command[-1]).write_bytes(b"%PDF-1.4\n%ocr\n")
        return object()

    monkeypatch.setattr(text_layer_quality, "build_text_layer_quality_report", fake_quality)
    monkeypatch.setattr(processing_runtime, "_run_completed_process", fake_run_completed_process)

    docx_bytes, backend = processing_runtime._convert_pdf_text_layer_to_docx(
        filename="scan.pdf",
        source_bytes=b"%PDF-1.4\n%scan\n",
    )
    paragraphs = extract_paragraph_units_from_docx(
        processing_runtime.build_in_memory_uploaded_file(
            source_name="scan.docx",
            source_bytes=docx_bytes,
        )
    )

    assert backend == "pdf-text-layer"
    assert [paragraph.text for paragraph in paragraphs] == ["OCR body text."]
    assert calls
    assert "-l" in calls[0]
    assert "eng+deu" in calls[0]


def test_pdf_text_layer_import_reports_missing_ocr_tools_when_ocr_enabled(monkeypatch):
    from docxaicorrector.pdf_import import text_layer_quality

    monkeypatch.setenv("DOCXAI_PDF_OCR_IMPORT_ENABLED", "1")
    monkeypatch.setattr(processing_runtime.shutil, "which", lambda name: None)
    monkeypatch.setattr(text_layer_quality, "extract_pdf_text_spans_with_pdfminer", lambda path: [])
    monkeypatch.setattr(
        text_layer_quality,
        "build_text_layer_quality_report",
        lambda spans: SimpleNamespace(
            decision="scanned_or_unsupported",
            decision_reasons=("empty_text_layer",),
            body_text_ratio=0.0,
        ),
    )

    with pytest.raises(RuntimeError, match="pdf_ocr_import_unavailable:ocrmypdf"):
        processing_runtime._convert_pdf_text_layer_to_docx(
            filename="scan.pdf",
            source_bytes=b"%PDF-1.4\n%scan\n",
        )


def test_materialize_uploaded_payload_reuses_cached_pdf_conversion(monkeypatch):
    _clear_materialized_upload_cache()
    pdf_bytes = b"%PDF-1.4\n%reuse-cache\n"
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="book-reuse.pdf",
        source_bytes=pdf_bytes,
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)

    convert_calls: list[dict[str, object]] = []

    def convert_pdf(**kwargs):
        convert_calls.append(kwargs)
        return b"converted-docx-reuse", "pdf-text-layer"

    monkeypatch.setattr(processing_runtime, "_convert_pdf_text_layer_to_docx", convert_pdf)

    first_events: list[dict[str, Any]] = []
    second_events: list[dict[str, Any]] = []

    first = processing_runtime.materialize_uploaded_payload(payload, progress_callback=lambda **kwargs: first_events.append(kwargs))
    second = processing_runtime.materialize_uploaded_payload(payload, progress_callback=lambda **kwargs: second_events.append(kwargs))

    assert len(convert_calls) == 1
    assert first.content_bytes == b"converted-docx-reuse"
    assert second.content_bytes == b"converted-docx-reuse"
    assert second.conversion_backend == "pdf-text-layer"
    assert any(e["stage"] == "DOCX готов" for e in second_events)
    assert any("Использую уже сконвертированную копию DOCX" in str(e["detail"]) for e in second_events)
    assert any(bool(cast(dict[str, Any], e["metrics"]).get("conversion_reused")) for e in second_events)


def test_materialize_uploaded_payload_passthrough_for_docx():
    docx_bytes = b"PK\x03\x04docx-bytes"
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="report.docx",
        source_bytes=docx_bytes,
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)

    materialized = processing_runtime.materialize_uploaded_payload(payload, progress_callback=None)

    assert materialized is payload  # no allocations, no work


def test_heartbeat_beacon_emits_progress_during_blocking_call():
    import time as _time

    events: list[dict[str, object]] = []

    def progress_cb(**kwargs):
        events.append(kwargs)

    with processing_runtime.HeartbeatBeacon(
        progress_cb,
        stage="Long phase",
        detail_template="working… ({elapsed} сек)",
        progress=0.5,
        metrics={"k": 1},
        interval_seconds=0.1,
    ):
        _time.sleep(0.35)

    assert len(events) >= 2
    for event in events:
        assert event["stage"] == "Long phase"
        assert "working…" in str(event["detail"])
        assert event["progress"] == 0.5
        assert event["metrics"] == {"k": 1}


def test_heartbeat_beacon_is_noop_when_callback_is_none():
    import time as _time

    with processing_runtime.HeartbeatBeacon(
        None,
        stage="Stage",
        detail_template="x",
        progress=0.0,
        interval_seconds=0.05,
    ):
        _time.sleep(0.15)
    # nothing to assert beyond not raising


def test_heartbeat_beacon_logs_warning_once_when_callback_fails(monkeypatch):
    import time as _time

    log_calls: list[dict[str, Any]] = []

    def failing_progress_cb(**kwargs):
        raise RuntimeError("boom")

    # HeartbeatBeacon lives in the ui-free upload_ports module (re-exported from
    # processing_runtime), so its log_event binding is upload_ports.log_event.
    monkeypatch.setattr(
        upload_ports,
        "log_event",
        lambda level, event, message, **context: log_calls.append(
            {
                "level": level,
                "event": event,
                "message": message,
                "context": context,
            }
        ),
    )

    with processing_runtime.HeartbeatBeacon(
        failing_progress_cb,
        stage="Long phase",
        detail_template="working… ({elapsed} сек)",
        progress=0.5,
        metrics={"k": 1},
        interval_seconds=0.05,
    ):
        _time.sleep(0.2)

    assert len(log_calls) == 1
    assert log_calls[0]["event"] == "heartbeat_callback_failed"
    assert cast(dict[str, Any], log_calls[0]["context"])["stage"] == "Long phase"


def test_start_background_preparation_materializes_pdf_inside_worker(monkeypatch):
    _clear_materialized_upload_cache()
    session_state = SessionState()
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    statuses: list[dict[str, object]] = []
    activities: list[str] = []

    pdf_bytes = b"%PDF-1.4\n%fakepdf\n"
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="book.pdf",
        source_bytes=pdf_bytes,
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)
    assert payload.conversion_backend is None  # precondition: cheap freeze

    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_text_layer_to_docx",
        lambda **kwargs: (b"converted-docx", "pdf-text-layer"),
    )

    received_payload: dict[str, Any] = {}

    def worker_target(**kwargs):
        received_payload["payload"] = kwargs["uploaded_payload"]
        kwargs["progress_callback"](
            stage="Документ подготовлен",
            detail="",
            progress=1.0,
            metrics={"cached": False},
        )

    processing_runtime.start_background_preparation(
        worker_target=worker_target,
        reset_run_state=lambda **kwargs: None,
        push_activity=lambda message: activities.append(message),
        set_processing_status=lambda **kw: statuses.append(kw),
        uploaded_payload=payload,
        upload_marker=payload.file_token,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
    )

    session_state.preparation_worker.join(timeout=5)
    processing_runtime.drain_preparation_events(
        reset_run_state=lambda **kwargs: None,
        set_processing_status=lambda **kw: statuses.append(kw),
        finalize_processing_status=lambda *args, **kw: None,
        push_activity=lambda message: activities.append(message),
    )

    materialized = cast(processing_runtime.FrozenUploadPayload, received_payload["payload"])
    assert isinstance(materialized, processing_runtime.FrozenUploadPayload)
    assert materialized.conversion_backend == "pdf-text-layer"
    assert materialized.content_bytes == b"converted-docx"
    # progress events from materialization must reach the UI
    stages_seen = [str(s.get("stage")) for s in statuses]
    assert "Импорт PDF" in stages_seen
    assert "DOCX готов" in stages_seen


def test_start_background_preparation_initial_status_for_lightweight_pdf(monkeypatch):
    _clear_materialized_upload_cache()
    """Bootstrap status contract: when the UI hands a lightweight PDF payload to
    `start_background_preparation`, the very first `set_processing_status` call
    must already advertise `source_format="pdf"` so the live-status panel shows
    a PDF-aware stage immediately, before LibreOffice runs in the worker.
    The lightweight payload must NOT carry a conversion_backend yet.
    """
    session_state = SessionState()
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    statuses: list[dict[str, object]] = []
    activities: list[str] = []

    pdf_bytes = b"%PDF-1.4\n%lightweight\n"
    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="book.pdf",
        source_bytes=pdf_bytes,
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)
    assert payload.source_format == "pdf"
    assert payload.conversion_backend is None
    assert payload.file_token.startswith("book.docx:")

    # Block the worker from progressing past start so we capture the initial status.
    started = threading.Event()
    release = threading.Event()

    def worker_target(**kwargs):
        started.set()
        release.wait(timeout=5)

    monkeypatch.setattr(
        processing_runtime,
        "_convert_pdf_to_docx",
        lambda **kwargs: (b"converted-docx", "libreoffice"),
    )

    processing_runtime.start_background_preparation(
        worker_target=worker_target,
        reset_run_state=lambda **kwargs: None,
        push_activity=lambda message: activities.append(message),
        set_processing_status=lambda **kw: statuses.append(kw),
        uploaded_payload=payload,
        upload_marker=payload.file_token,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
    )

    try:
        # First status is emitted synchronously by start_background_preparation.
        assert statuses, "start_background_preparation must emit an initial status"
        first = statuses[0]
        assert first["phase"] == "preparing"
        assert first["stage"] == "Файл получен"
        assert first["source_format"] == "pdf"
        # Lightweight payload: backend not known yet — worker will materialize.
        assert first.get("conversion_backend") is None
        assert activities and activities[0].startswith("Файл получен")
    finally:
        release.set()
        if session_state.preparation_worker is not None:
            session_state.preparation_worker.join(timeout=5)


def test_start_background_preparation_reports_materialization_failure_before_worker_target(monkeypatch):
    _clear_materialized_upload_cache()
    session_state = SessionState()
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    statuses: list[dict[str, object]] = []
    activities: list[str] = []
    finalized = []
    worker_calls: list[dict[str, object]] = []

    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="book.pdf",
        source_bytes=b"%PDF-1.4\nsource\n",
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)
    assert payload.source_format == "pdf"
    assert payload.conversion_backend is None

    monkeypatch.setattr(
        processing_runtime,
        "materialize_uploaded_payload",
        lambda uploaded_payload, progress_callback=None: (_ for _ in ()).throw(RuntimeError("pdf converter missing")),
    )

    processing_runtime.start_background_preparation(
        worker_target=lambda **kwargs: worker_calls.append(kwargs),
        reset_run_state=lambda **kwargs: None,
        push_activity=lambda message: activities.append(message),
        set_processing_status=lambda **kw: statuses.append(kw),
        uploaded_payload=payload,
        upload_marker=payload.file_token,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
    )

    session_state.preparation_worker.join(timeout=5)
    processing_runtime.drain_preparation_events(
        reset_run_state=lambda **kwargs: None,
        set_processing_status=lambda **kw: statuses.append(kw),
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: finalized.append((stage, detail, progress, terminal_kind)),
        push_activity=lambda message: activities.append(message),
    )

    assert worker_calls == []
    assert statuses[0]["stage"] == "Файл получен"
    assert statuses[0]["source_format"] == "pdf"
    assert session_state.preparation_failed_marker == payload.file_token
    assert session_state.last_error == "pdf converter missing"
    assert session_state.last_background_error["stage"] == "preparation"
    assert finalized == [("Ошибка подготовки", "pdf converter missing", 1.0, "error")]
    assert activities[-1] == "Не удалось прочитать и проанализировать документ."


def test_start_background_preparation_reports_worker_failure_after_materialization(monkeypatch):
    _clear_materialized_upload_cache()
    session_state = SessionState()
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    statuses: list[dict[str, object]] = []
    activities: list[str] = []
    finalized = []
    worker_calls: list[dict[str, object]] = []

    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="book.pdf",
        source_bytes=b"%PDF-1.4\nsource\n",
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)

    materialized_payload = processing_runtime.FrozenUploadPayload(
        filename="book.docx",
        content_bytes=b"PK\x03\x04converted-docx",
        file_size=len(b"PK\x03\x04converted-docx"),
        content_hash="mockedhash",
        file_token="book.docx:18:mockedhash",
        source_format="pdf",
        conversion_backend="libreoffice",
    )
    monkeypatch.setattr(
        processing_runtime,
        "materialize_uploaded_payload",
        lambda uploaded_payload, progress_callback=None: materialized_payload,
    )

    def worker_target(**kwargs):
        worker_calls.append(kwargs)
        raise RuntimeError("worker boom")

    processing_runtime.start_background_preparation(
        worker_target=worker_target,
        reset_run_state=lambda **kwargs: None,
        push_activity=lambda message: activities.append(message),
        set_processing_status=lambda **kw: statuses.append(kw),
        uploaded_payload=payload,
        upload_marker=payload.file_token,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
    )

    session_state.preparation_worker.join(timeout=5)
    processing_runtime.drain_preparation_events(
        reset_run_state=lambda **kwargs: None,
        set_processing_status=lambda **kw: statuses.append(kw),
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: finalized.append((stage, detail, progress, terminal_kind)),
        push_activity=lambda message: activities.append(message),
    )

    assert len(worker_calls) == 1
    assert worker_calls[0]["uploaded_payload"] == materialized_payload
    assert statuses[0]["source_format"] == "pdf"
    assert session_state.preparation_failed_marker == payload.file_token
    assert session_state.last_error == "worker boom"
    assert session_state.last_background_error["stage"] == "preparation"
    assert finalized == [("Ошибка подготовки", "worker boom", 1.0, "error")]
    assert activities[-1] == "Не удалось прочитать и проанализировать документ."


def test_start_background_preparation_stops_after_materialization_before_worker_target(monkeypatch):
    _clear_materialized_upload_cache()
    session_state = SessionState()
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    statuses: list[dict[str, object]] = []
    activities: list[str] = []
    finalized = []
    worker_calls: list[dict[str, object]] = []

    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="book.pdf",
        source_bytes=b"%PDF-1.4\nsource\n",
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)

    materialized_payload = processing_runtime.FrozenUploadPayload(
        filename="book.docx",
        content_bytes=b"PK\x03\x04converted-docx",
        file_size=len(b"PK\x03\x04converted-docx"),
        content_hash="mockedhash",
        file_token="book.docx:18:mockedhash",
        source_format="pdf",
        conversion_backend="libreoffice",
    )

    def materialize_and_request_stop(uploaded_payload, progress_callback=None):
        session_state.preparation_stop_event.set()
        return materialized_payload

    monkeypatch.setattr(
        processing_runtime,
        "materialize_uploaded_payload",
        materialize_and_request_stop,
    )

    processing_runtime.start_background_preparation(
        worker_target=lambda **kwargs: worker_calls.append(kwargs),
        reset_run_state=lambda **kwargs: None,
        push_activity=lambda message: activities.append(message),
        set_processing_status=lambda **kw: statuses.append(kw),
        uploaded_payload=payload,
        upload_marker=payload.file_token,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
    )

    session_state.preparation_worker.join(timeout=5)
    processing_runtime.drain_preparation_events(
        reset_run_state=lambda **kwargs: None,
        set_processing_status=lambda **kw: statuses.append(kw),
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: finalized.append((stage, detail, progress, terminal_kind)),
        push_activity=lambda message: activities.append(message),
    )

    assert worker_calls == []
    assert session_state.prepared_run_context is None
    assert session_state.preparation_worker is None
    assert session_state.preparation_event_queue is None
    assert session_state.preparation_stop_event is None
    assert session_state.processing_outcome == "idle"
    assert finalized == [("Подготовка остановлена", "", 1.0, "stopped")]
    assert activities[-1] == "Подготовка документа остановлена."



def test_append_pdf_text_paragraph_emits_per_run_emphasis() -> None:
    from docx import Document

    from docxaicorrector.core.models import ParagraphUnit

    paragraph = ParagraphUnit(
        text="See life-threatening multifaceted risk.¹",
        role="body",
        structural_role="body",
        pdf_emphasis_runs=[
            ("See life-threatening ", False, False),
            ("multifaceted", False, True),
            (" risk.¹", False, False),
        ],
    )
    document = Document()

    processing_runtime._append_pdf_text_paragraph_to_docx(document, paragraph)

    emitted = document.paragraphs[-1]
    assert emitted.text == "See life-threatening multifaceted risk.¹"
    italic_runs = [run.text for run in emitted.runs if run.italic]
    assert italic_runs == ["multifaceted"]
    assert all(not run.bold for run in emitted.runs)


def test_append_pdf_text_paragraph_falls_back_without_runs() -> None:
    from docx import Document

    from docxaicorrector.core.models import ParagraphUnit

    paragraph = ParagraphUnit(
        text="Plain body line.",
        role="body",
        structural_role="body",
        is_italic=True,
    )
    document = Document()

    processing_runtime._append_pdf_text_paragraph_to_docx(document, paragraph)

    emitted = document.paragraphs[-1]
    assert emitted.text == "Plain body line."
    assert len(emitted.runs) == 1
    assert emitted.runs[0].italic is True


# --- F24: heading emphasis runs preserved by the DOCX materializer ---------


def test_append_pdf_heading_preserves_mixed_emphasis_runs() -> None:
    from docx import Document

    from docxaicorrector.core.models import ParagraphUnit

    paragraph = ParagraphUnit(
        text="Chapter One: The Beginning",
        role="heading",
        structural_role="heading",
        heading_level=1,
        pdf_emphasis_runs=[
            ("Chapter One: ", True, False),
            ("The Beginning", False, True),
        ],
    )
    document = Document()

    processing_runtime._append_pdf_text_paragraph_to_docx(document, paragraph)

    emitted = document.paragraphs[-1]
    assert emitted.text == "Chapter One: The Beginning"
    assert [run.text for run in emitted.runs if run.bold] == ["Chapter One: "]
    assert [run.text for run in emitted.runs if run.italic] == ["The Beginning"]


def test_append_pdf_heading_without_runs_stays_single_plain_run() -> None:
    from docx import Document

    from docxaicorrector.core.models import ParagraphUnit

    paragraph = ParagraphUnit(
        text="Plain Heading",
        role="heading",
        structural_role="heading",
        heading_level=2,
    )
    document = Document()

    processing_runtime._append_pdf_text_paragraph_to_docx(document, paragraph)

    emitted = document.paragraphs[-1]
    assert emitted.text == "Plain Heading"
    assert len(emitted.runs) == 1
    assert not emitted.runs[0].bold
    assert not emitted.runs[0].italic


# --- F7: PDF import resource budget ----------------------------------------


def _promising_report_stub(spans):
    return SimpleNamespace(decision="promising", decision_reasons=(), body_text_ratio=1.0)


def test_convert_pdf_rejects_oversize_input_file(monkeypatch) -> None:
    events: list[tuple[int, str, dict[str, object]]] = []
    monkeypatch.setattr(processing_runtime, "_MAX_PDF_IMPORT_FILE_BYTES", 16)
    monkeypatch.setattr(
        processing_runtime,
        "log_event",
        lambda level, event, message, **context: events.append((level, event, context)),
    )

    with pytest.raises(RuntimeError, match="pdf_import_over_budget:file_size"):
        processing_runtime._convert_pdf_text_layer_to_docx(
            filename="oversize.pdf",
            source_bytes=b"%PDF-1.4\n" + b"x" * 100,
        )

    assert any(
        event == "pdf_import_over_budget" and context.get("limit") == "file_size"
        for _level, event, context in events
    )


def test_convert_pdf_rejects_over_page_budget(monkeypatch) -> None:
    events: list[tuple[int, str, dict[str, object]]] = []
    # The page-count budget now runs inside the parse worker under the unified
    # deadline; point the worker at the in-process seam so this monkeypatch is seen.
    monkeypatch.setattr(processing_runtime, "_PDF_PARSE_IN_PROCESS", True)
    monkeypatch.setattr(processing_runtime, "_MAX_PDF_IMPORT_PAGE_COUNT", 3)
    monkeypatch.setattr(processing_runtime, "_count_pdf_pages_for_budget", lambda path, cap: 5000)
    monkeypatch.setattr(
        processing_runtime,
        "log_event",
        lambda level, event, message, **context: events.append((level, event, context)),
    )

    with pytest.raises(RuntimeError, match="pdf_import_over_budget:page_count"):
        processing_runtime._convert_pdf_text_layer_to_docx(
            filename="many-pages.pdf",
            source_bytes=b"%PDF-1.4\n",
        )

    assert any(
        event == "pdf_import_over_budget" and context.get("limit") == "page_count"
        for _level, event, context in events
    )


def test_convert_pdf_rejects_when_parse_exceeds_wallclock_budget(monkeypatch) -> None:
    from docxaicorrector.pdf_import import text_layer_quality

    events: list[tuple[int, str, dict[str, object]]] = []
    monkeypatch.setattr(processing_runtime, "_PDF_PARSE_IN_PROCESS", True)
    monkeypatch.setattr(processing_runtime, "_PDF_PARSE_WALLCLOCK_BUDGET_SECONDS", 0.2)

    def slow_parse(path):
        time.sleep(1.0)
        return []

    monkeypatch.setattr(text_layer_quality, "extract_pdf_text_spans_with_pdfminer", slow_parse)
    monkeypatch.setattr(
        processing_runtime,
        "log_event",
        lambda level, event, message, **context: events.append((level, event, context)),
    )

    with pytest.raises(RuntimeError, match="pdf_import_over_budget:parse_wallclock"):
        processing_runtime._convert_pdf_text_layer_to_docx(
            filename="slow.pdf",
            source_bytes=b"%PDF-1.4\n",
        )

    assert any(
        event == "pdf_import_over_budget" and context.get("limit") == "parse_wallclock"
        for _level, event, context in events
    )


def test_convert_pdf_normal_document_parses_within_budget(monkeypatch) -> None:
    from docxaicorrector.pdf_import import images as pdf_images
    from docxaicorrector.pdf_import import text_layer_quality

    spans = [
        PdfTextSpan(
            page_number=1,
            text="Ordinary body line.",
            x0=50,
            top=100,
            x1=450,
            bottom=112,
            page_height=800,
            font_size=10,
        )
    ]
    monkeypatch.setattr(processing_runtime, "_PDF_PARSE_IN_PROCESS", True)
    monkeypatch.setattr(text_layer_quality, "extract_pdf_text_spans_with_pdfminer", lambda path: spans)
    monkeypatch.setattr(text_layer_quality, "build_text_layer_quality_report", _promising_report_stub)
    monkeypatch.setattr(pdf_images, "extract_pdf_images_with_pdfminer", lambda path: [])

    docx_bytes, backend = processing_runtime._convert_pdf_text_layer_to_docx(
        filename="normal.pdf",
        source_bytes=b"%PDF-1.4\n",
    )

    assert backend == "pdf-text-layer"
    assert docx_bytes


def test_convert_pdf_rejects_over_span_count_cap(monkeypatch, tmp_path) -> None:
    from docxaicorrector.pdf_import import text_layer_quality

    events: list[tuple[int, str, dict[str, object]]] = []
    monkeypatch.setattr(processing_runtime, "_PDF_PARSE_IN_PROCESS", True)
    monkeypatch.setattr(processing_runtime, "_MAX_PDF_IMPORT_SPAN_COUNT", 1)

    def two_spans(path):
        return [
            PdfTextSpan(page_number=1, text="a", x0=0, top=0, x1=1, bottom=1),
            PdfTextSpan(page_number=1, text="b", x0=0, top=2, x1=1, bottom=3),
        ]

    monkeypatch.setattr(text_layer_quality, "extract_pdf_text_spans_with_pdfminer", two_spans)
    monkeypatch.setattr(
        processing_runtime,
        "log_event",
        lambda level, event, message, **context: events.append((level, event, context)),
    )

    with pytest.raises(RuntimeError, match="pdf_import_over_budget:span_count"):
        processing_runtime._convert_pdf_text_layer_to_docx(
            filename="spans.pdf",
            source_bytes=b"%PDF-1.4\n",
        )

    assert any(
        event == "pdf_import_over_budget" and context.get("limit") == "span_count"
        for _level, event, context in events
    )


def test_convert_pdf_rejects_over_image_count_cap(monkeypatch) -> None:
    from docxaicorrector.pdf_import import images as pdf_images
    from docxaicorrector.pdf_import import text_layer_quality

    events: list[tuple[int, str, dict[str, object]]] = []
    monkeypatch.setattr(processing_runtime, "_PDF_PARSE_IN_PROCESS", True)
    monkeypatch.setattr(processing_runtime, "_MAX_PDF_IMPORT_IMAGE_COUNT", 1)

    spans = [PdfTextSpan(page_number=1, text="Body.", x0=0, top=0, x1=100, bottom=12)]
    monkeypatch.setattr(text_layer_quality, "extract_pdf_text_spans_with_pdfminer", lambda path: spans)
    monkeypatch.setattr(text_layer_quality, "build_text_layer_quality_report", _promising_report_stub)

    def two_images(path):
        return [SimpleNamespace(source_index=0), SimpleNamespace(source_index=1)]

    monkeypatch.setattr(pdf_images, "extract_pdf_images_with_pdfminer", two_images)
    monkeypatch.setattr(
        processing_runtime,
        "log_event",
        lambda level, event, message, **context: events.append((level, event, context)),
    )

    with pytest.raises(RuntimeError, match="pdf_import_over_budget:image_count"):
        processing_runtime._convert_pdf_text_layer_to_docx(
            filename="images.pdf",
            source_bytes=b"%PDF-1.4\n",
        )

    assert any(
        event == "pdf_import_over_budget" and context.get("limit") == "image_count"
        for _level, event, context in events
    )


def test_pdf_parse_subprocess_is_terminated_on_deadline_overrun(monkeypatch, tmp_path) -> None:
    """F7: the REAL subprocess path force-kills a worker that overruns the deadline.

    Uses a module-level worker that sleeps forever as the spawned target; the
    parent must terminate()/join() it (not leave it alive) and raise the typed
    ``pdf_import_over_budget:deadline`` error.
    """

    # Opt out of the in-process seam: exercise the REAL killable subprocess path.
    monkeypatch.setattr(processing_runtime, "_PDF_PARSE_IN_PROCESS", False)
    monkeypatch.setattr(processing_runtime, "_PDF_PARSE_WALLCLOCK_BUDGET_SECONDS", 0.5)
    monkeypatch.setattr(
        processing_runtime,
        "_PDF_PARSE_SUBPROCESS_ENTRY",
        processing_runtime._pdf_parse_sleep_forever_entry,
    )

    captured: dict[str, object] = {}
    monkeypatch.setattr(
        processing_runtime,
        "_pdf_parse_process_observer",
        lambda proc: captured.__setitem__("proc", proc),
    )

    input_path = tmp_path / "in.pdf"
    input_path.write_bytes(b"%PDF-1.4\n")
    ocr_output_path = tmp_path / "in.ocr.pdf"

    with pytest.raises(RuntimeError, match="pdf_import_over_budget:deadline"):
        processing_runtime._run_pdf_parse_stages_with_deadline(
            temp_dir=tmp_path,
            input_path=input_path,
            ocr_output_path=ocr_output_path,
            filename="slow.pdf",
        )

    process = cast(Any, captured.get("proc"))
    assert process is not None, "the spawned child process was never observed"
    # The child must have been genuinely terminated, not left running.
    assert process.is_alive() is False
    assert process.exitcode is not None


def test_pdf_parse_subprocess_result_is_deserialized_from_child(monkeypatch, tmp_path) -> None:
    """F7: the parent reads back a real spawned child's serialized parse result."""

    monkeypatch.setattr(processing_runtime, "_PDF_PARSE_IN_PROCESS", False)
    monkeypatch.setattr(
        processing_runtime,
        "_PDF_PARSE_SUBPROCESS_ENTRY",
        processing_runtime._pdf_parse_canned_result_entry,
    )

    input_path = tmp_path / "in.pdf"
    input_path.write_bytes(b"%PDF-1.4\n")
    ocr_output_path = tmp_path / "in.ocr.pdf"

    result = processing_runtime._run_pdf_parse_stages_with_deadline(
        temp_dir=tmp_path,
        input_path=input_path,
        ocr_output_path=ocr_output_path,
        filename="ok.pdf",
    )

    assert isinstance(result, processing_runtime._PdfParseStagesResult)
    assert result.quality_decision == "promising"
    assert result.spans == []
    assert result.image_objects == []


# --- F27: process-wide admission gate --------------------------------------


def test_processing_admission_limit_env_override(monkeypatch) -> None:
    monkeypatch.setenv("DOCXAI_MAX_CONCURRENT_PROCESSING", "5")
    assert processing_runtime._resolve_processing_admission_limit() == 5

    monkeypatch.setenv("DOCXAI_MAX_CONCURRENT_PROCESSING", "0")
    assert (
        processing_runtime._resolve_processing_admission_limit()
        == processing_runtime._DEFAULT_PROCESSING_ADMISSION_LIMIT
    )

    monkeypatch.setenv("DOCXAI_MAX_CONCURRENT_PROCESSING", "not-a-number")
    assert (
        processing_runtime._resolve_processing_admission_limit()
        == processing_runtime._DEFAULT_PROCESSING_ADMISSION_LIMIT
    )

    monkeypatch.delenv("DOCXAI_MAX_CONCURRENT_PROCESSING", raising=False)
    assert (
        processing_runtime._resolve_processing_admission_limit()
        == processing_runtime._DEFAULT_PROCESSING_ADMISSION_LIMIT
    )


def test_processing_admission_gate_caps_concurrency() -> None:
    gate = processing_runtime._build_processing_admission_gate(2)

    assert gate.acquire(blocking=False) is True
    assert gate.acquire(blocking=False) is True
    # The gate is bounded: with both slots held the next acquire cannot proceed.
    assert gate.acquire(blocking=False) is False

    gate.release()
    assert gate.acquire(blocking=False) is True

    gate.release()
    gate.release()


def test_processing_admission_gate_floor_is_single_slot() -> None:
    gate = processing_runtime._build_processing_admission_gate(0)

    assert gate.acquire(blocking=False) is True
    assert gate.acquire(blocking=False) is False

    gate.release()


def test_admission_gate_covers_preparation_work(monkeypatch) -> None:
    """F27: preparation acquires the process-wide admission slot around its real
    work and releases it on completion (the costliest PDF stage runs here)."""

    _clear_materialized_upload_cache()
    session_state = SessionState()
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)

    events_log: list[str] = []
    real_gate = processing_runtime._build_processing_admission_gate(1)

    class SpyGate:
        def acquire(self, *args, **kwargs):
            acquired = real_gate.acquire(*args, **kwargs)
            if acquired:
                events_log.append("acquire")
            return acquired

        def release(self) -> None:
            events_log.append("release")
            real_gate.release()

    monkeypatch.setattr(processing_runtime, "_PROCESSING_ADMISSION_GATE", SpyGate())

    uploaded_file = processing_runtime.build_in_memory_uploaded_file(
        source_name="book.docx",
        source_bytes=b"PK\x03\x04docx-body",
    )
    payload = processing_runtime.freeze_uploaded_file_lightweight(uploaded_file)

    def worker_target(**kwargs):
        events_log.append("worker_target")
        kwargs["progress_callback"](stage="Готово", detail="", progress=1.0, metrics={})

    processing_runtime.start_background_preparation(
        worker_target=worker_target,
        reset_run_state=lambda **kwargs: None,
        push_activity=lambda message: None,
        set_processing_status=lambda **kw: None,
        uploaded_payload=payload,
        upload_marker=payload.file_token,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
    )

    session_state.preparation_worker.join(timeout=5)

    # The slot must wrap the real work: acquired first, worker runs, released last.
    assert events_log == ["acquire", "worker_target", "release"]
    # And the slot is fully returned to the gate after preparation finishes.
    assert real_gate.acquire(blocking=False) is True
    real_gate.release()


def test_admission_gate_preparation_wait_is_cancellable(monkeypatch) -> None:
    """F27: a stopped upload must not block on a full admission slot forever."""

    stop_event = threading.Event()
    stop_event.set()
    full_gate = processing_runtime._build_processing_admission_gate(1)
    assert full_gate.acquire(blocking=False) is True  # exhaust the only slot

    acquired = processing_runtime._acquire_admission_slot_cancellable(
        full_gate, stop_event, poll_seconds=0.01
    )

    assert acquired is False  # cancelled instead of blocking on the held slot
    full_gate.release()


def test_admission_gate_processing_wait_is_cancellable(monkeypatch):
    """Spec 041 P1-2: a Stop while the processing admission gate is full must
    cancel the queued run — ``worker_target`` never runs, the run surfaces
    ``stopped`` to the UI, and the (never-acquired) slot is not released."""

    session_state = SessionState(restart_session_id="session-a")
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(state.st, "session_state", session_state)
    state.init_session_state()
    session_state.restart_session_id = "session-a"
    monkeypatch.setattr(
        processing_runtime,
        "store_restart_source",
        lambda **kwargs: {
            "filename": kwargs["source_name"],
            "token": kwargs["source_token"],
            "storage_path": "restart.bin",
            "session_id": kwargs["session_id"],
        },
    )

    # Saturate the admission gate: its only slot is held for the whole test, so
    # _acquire_admission_slot_cancellable can NEVER acquire — the only exit is a
    # Stop. A spy records every acquire/release so we can assert no over-release.
    real_gate = processing_runtime._build_processing_admission_gate(1)
    assert real_gate.acquire(blocking=False) is True  # exhaust the only slot

    acquire_results: list[bool] = []
    release_calls: list[str] = []

    class SpyGate:
        def acquire(self, *args, **kwargs):
            acquired = real_gate.acquire(*args, **kwargs)
            acquire_results.append(acquired)
            return acquired

        def release(self) -> None:
            release_calls.append("release")
            real_gate.release()

    monkeypatch.setattr(processing_runtime, "_PROCESSING_ADMISSION_GATE", SpyGate())

    worker_calls: list[dict] = []

    def worker_target(**kwargs):
        worker_calls.append(kwargs)

    processing_runtime.start_background_processing(
        worker_target=worker_target,
        reset_run_state=state.reset_run_state,
        push_activity=lambda message: None,
        set_processing_status=lambda **kwargs: None,
        uploaded_filename="report.docx",
        uploaded_token="report.docx:3:abc",
        source_bytes=b"abc",
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=["paragraph"],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
    )

    # The gate is permanently saturated, so the worker is parked polling for a
    # slot; requesting Stop is the only way out. Setting the run's own stop_event
    # (the exact object threaded into the guarded target) cancels the wait.
    run_stop_event = session_state.processing_stop_event
    run_stop_event.set()

    session_state.processing_worker.join(timeout=5)
    assert not session_state.processing_worker.is_alive()

    # worker_target must never run — the run was cancelled during admission.
    assert worker_calls == []
    # The slot was never acquired, so it must never be released (no over-release;
    # a BoundedSemaphore would raise on a spurious release — it did not, and the
    # spy confirms release was never called on the gate).
    assert release_calls == []
    assert acquire_results and all(result is False for result in acquire_results)

    # The run surfaces the normal stopped completion to the UI event stream.
    outcomes: list[str] = []
    while True:
        try:
            event = session_state.processing_event_queue.get_nowait()
        except queue.Empty:
            break
        if isinstance(event, WorkerCompleteEvent):
            outcomes.append(event.outcome)
    assert outcomes == ["stopped"]

    # The real gate still holds exactly the test's slot: releasing it once
    # succeeds, and a second release raises (proving no phantom slot was added).
    real_gate.release()
    with pytest.raises(ValueError):
        real_gate.release()


# --- F12: PDF image discovery vs emission is counted and warned ------------


class _FailingImageStream:
    def get_rawdata(self):
        raise ValueError("stream_decode_failed")

    def get_data(self):
        raise ValueError("stream_decode_failed")


class _StaticImageStream:
    def __init__(self, data: bytes) -> None:
        self._data = data

    def get_rawdata(self):
        return self._data


class _FakeLTFigure(list):
    pass


class _FakePage(list):
    height = 800.0


class _FakeLTImage:
    def __init__(self, stream, bbox):
        self.stream = stream
        self.x0, self.y0, self.x1, self.y1 = bbox


def _patch_pdfminer_image_layout(monkeypatch, pages) -> None:
    import pdfminer.high_level as high_level
    import pdfminer.layout as layout

    monkeypatch.setattr(layout, "LTImage", _FakeLTImage)
    monkeypatch.setattr(layout, "LTFigure", _FakeLTFigure)
    monkeypatch.setattr(high_level, "extract_pages", lambda path: pages)


def test_extract_pdf_images_warns_when_discovered_exceeds_emitted(monkeypatch) -> None:
    from docxaicorrector.pdf_import import images as pdf_images

    failing_image = _FakeLTImage(_FailingImageStream(), (0.0, 0.0, 10.0, 10.0))
    page = _FakePage([failing_image])
    _patch_pdfminer_image_layout(monkeypatch, pages=[page])

    events: list[tuple[int, str, dict[str, object]]] = []
    monkeypatch.setattr(
        pdf_images,
        "log_event",
        lambda level, event, message, **context: events.append((level, event, context)),
    )

    result = pdf_images.extract_pdf_images_with_pdfminer("dummy.pdf")

    assert result == []
    dropped = next(
        context for _level, event, context in events if event == "pdf_image_extraction_dropped_images"
    )
    assert dropped["discovered"] == 1
    assert dropped["emitted"] == 0
    assert dropped["skipped_no_image_bytes"] == 1


def test_extract_pdf_images_emits_summary_without_dropping_valid_image(monkeypatch) -> None:
    from docxaicorrector.pdf_import import images as pdf_images

    png_bytes = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
        b"\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4"
        b"\x89\x00\x00\x00\nIDATx\x9cc\xf8\x0f\x00\x01\x01"
        b"\x01\x00\x18\xdd\x8d\xb0\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    good_image = _FakeLTImage(_StaticImageStream(png_bytes), (10.0, 700.0, 60.0, 760.0))
    page = _FakePage([good_image])
    _patch_pdfminer_image_layout(monkeypatch, pages=[page])

    events: list[tuple[int, str, dict[str, object]]] = []
    monkeypatch.setattr(
        pdf_images,
        "log_event",
        lambda level, event, message, **context: events.append((level, event, context)),
    )

    result = pdf_images.extract_pdf_images_with_pdfminer("dummy.pdf")

    assert len(result) == 1
    assert result[0].mime_type == "image/png"
    assert not any(event == "pdf_image_extraction_dropped_images" for _level, event, _context in events)
    summary = next(
        context for _level, event, context in events if event == "pdf_image_extraction_summary"
    )
    assert summary["discovered"] == 1
    assert summary["emitted"] == 1
