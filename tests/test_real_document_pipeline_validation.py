import base64
import importlib.util
import io
import json
import os
import time
from io import BytesIO
from pathlib import Path
from contextlib import redirect_stdout
from types import SimpleNamespace

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxaicorrector.core.models import LayoutArtifactCleanupDecision, LayoutArtifactCleanupReport
from docxaicorrector.core.models import ParagraphUnit
import docxaicorrector.pipeline._pipeline as document_pipeline
import docxaicorrector.pipeline.late_phases as document_pipeline_late_phases


PNG_BYTES = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aK3cAAAAASUVORK5CYII=")


def _load_validation_module():
    project_root = Path(__file__).resolve().parents[1]
    module_path = project_root / "tests" / "artifacts" / "real_document_pipeline" / "run_lietaer_validation.py"
    spec = importlib.util.spec_from_file_location("run_lietaer_validation", module_path)
    if spec is None or spec.loader is None:
        raise AssertionError("Unable to load run_lietaer_validation.py")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _docx_bytes(document: Document) -> bytes:  # type: ignore[reportGeneralTypeIssues]
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _append_numbering_level(level: str, fmt: str) -> OxmlElement:  # type: ignore[reportGeneralTypeIssues]
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), level)

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(
        qn("w:val"),
        "%1." if fmt in {"decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman"} else "•",
    )
    lvl.append(lvl_text)
    return lvl


def _append_multilevel_numbering_definition(document: Document, *, num_id: str, abstract_num_id: str) -> None:  # type: ignore[reportGeneralTypeIssues]
    numbering_root = document.part.numbering_part.element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_num_id)
    abstract_num.append(_append_numbering_level("0", "bullet"))
    abstract_num.append(_append_numbering_level("1", "decimal"))
    numbering_root.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_num_ref = OxmlElement("w:abstractNumId")
    abstract_num_ref.set(qn("w:val"), abstract_num_id)
    num.append(abstract_num_ref)
    numbering_root.append(num)


def _attach_numbering(paragraph, *, num_id: str, ilvl: str) -> None:
    paragraph_properties = paragraph._element.get_or_add_pPr()
    num_pr = paragraph_properties.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        paragraph_properties.append(num_pr)

    ilvl_element = num_pr.find(qn("w:ilvl"))
    if ilvl_element is None:
        ilvl_element = OxmlElement("w:ilvl")
        num_pr.append(ilvl_element)
    ilvl_element.set(qn("w:val"), ilvl)

    num_id_element = num_pr.find(qn("w:numId"))
    if num_id_element is None:
        num_id_element = OxmlElement("w:numId")
        num_pr.append(num_id_element)
    num_id_element.set(qn("w:val"), num_id)


def test_evaluate_lietaer_acceptance_detects_caption_heading_regression(tmp_path):
    validation = _load_validation_module()
    image_path = tmp_path / "caption_image.png"
    image_path.write_bytes(PNG_BYTES)

    source_doc = Document()
    source_doc.add_paragraph().add_run().add_picture(str(image_path))
    source_doc.add_paragraph("Рисунок 1. Подпись")

    output_doc = Document()
    output_doc.add_paragraph().add_run().add_picture(str(image_path))
    output_doc.add_paragraph("Рисунок 1. Подпись", style="Heading 1")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    assert acceptance["passed"] is False
    assert "captions_not_promoted_to_headings" in acceptance["failed_checks"]


def test_evaluate_lietaer_acceptance_detects_center_alignment_regression() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_centered = source_doc.add_paragraph("ЭПИКТЕТ")
    source_centered.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_doc = Document()
    output_doc.add_paragraph("ЭПИКТЕТ")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    assert acceptance["passed"] is False
    assert "centered_short_paragraphs_preserved" in acceptance["failed_checks"]


def test_evaluate_lietaer_acceptance_passes_for_clean_structural_output(tmp_path):
    validation = _load_validation_module()
    image_path = tmp_path / "clean_image.png"
    image_path.write_bytes(PNG_BYTES)

    source_doc = Document()
    centered_title = source_doc.add_paragraph("Глава 1", style="Heading 1")
    centered_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_doc.add_paragraph("Первый пункт", style="List Number")
    source_doc.add_paragraph("Второй пункт", style="List Number")
    source_doc.add_paragraph().add_run().add_picture(str(image_path))
    source_caption = source_doc.add_paragraph("Рисунок 1. Корректная подпись")
    source_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_doc = Document()
    centered_output_title = output_doc.add_paragraph("Глава 1", style="Heading 1")
    centered_output_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    output_doc.add_paragraph("Первый пункт", style="List Number")
    output_doc.add_paragraph("Второй пункт", style="List Number")
    output_doc.add_paragraph().add_run().add_picture(str(image_path))
    output_caption = output_doc.add_paragraph("Рисунок 1. Корректная подпись")
    output_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    assert acceptance["passed"] is True
    assert acceptance["failed_checks"] == []


def test_evaluate_lietaer_acceptance_fails_on_reader_cleanup_failed_stage() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Body")
    output_doc = Document()
    output_doc.add_paragraph("Body")
    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "reader_cleanup_evidence": {
            "stage_status": "failed",
            "failed_chunk_count": 75,
            "cleanup_chunk_count": 75,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )
    by_name = {check["name"]: check for check in acceptance["checks"]}

    assert acceptance["passed"] is False
    assert "reader_cleanup_stage_completed" in acceptance["failed_checks"]
    assert by_name["reader_cleanup_stage_completed"]["stage_status"] == "failed"
    assert by_name["reader_cleanup_stage_completed"]["failed_chunk_ratio"] == 1.0


def test_evaluate_lietaer_acceptance_fails_on_translation_quality_report_residual_defects() -> None:
    validation = _load_validation_module()

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "translation_quality_report": {
            "bullet_heading_count": 0,
            "bullet_heading_gate_source": "legacy_markdown",
            "bullet_heading_classification": "markdown_gate",
            "raw_bullet_heading_count": 0,
            "false_fragment_heading_count": 2,
            "residual_bullet_glyph_count": 1,
            "residual_bullet_glyph_gate_source": "legacy_markdown",
            "residual_bullet_glyph_classification": "display_hygiene",
            "raw_residual_bullet_glyph_count": 1,
            "list_fragment_regression_count": 1,
            "mixed_script_term_count": 1,
            "mixed_script_term_gate_source": "legacy_markdown",
            "mixed_script_term_classification": "non_structural_hygiene",
            "raw_mixed_script_term_count": 1,
            "theology_style_deterministic_issue_count": 3,
            "theology_style_deterministic_issue_source": "legacy_markdown",
            "theology_style_deterministic_issue_classification": "domain_style_advisory",
            "raw_theology_style_deterministic_issue_count": 3,
            "toc_body_concat_detected": True,
            "toc_body_concat_markdown_detected": True,
            "toc_body_concat_structure_detected": False,
            "toc_body_concat_gate_source": "legacy_markdown",
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(report)
    by_name = {check["name"]: check for check in acceptance["checks"]}

    assert acceptance["passed"] is False
    assert acceptance["failed_checks"] == [
        "false_fragment_headings_present",
        "residual_bullet_glyphs_present",
        "list_fragment_regressions_present",
        "mixed_script_terms_present",
        "structural_comparison_available",
    ]
    assert by_name["residual_bullet_glyphs_present"]["residual_bullet_glyph_gate_source"] == "legacy_markdown"
    assert by_name["residual_bullet_glyphs_present"]["residual_bullet_glyph_classification"] == "display_hygiene"
    assert by_name["residual_bullet_glyphs_present"]["raw_residual_bullet_glyph_count"] == 1
    assert by_name["mixed_script_terms_present"]["mixed_script_term_gate_source"] == "legacy_markdown"
    assert by_name["mixed_script_terms_present"]["mixed_script_term_classification"] == "non_structural_hygiene"
    assert by_name["mixed_script_terms_present"]["raw_mixed_script_term_count"] == 1
    assert by_name["theology_style_deterministic_issues_present"]["theology_style_deterministic_issue_source"] == "legacy_markdown"
    assert (
        by_name["theology_style_deterministic_issues_present"]["theology_style_deterministic_issue_classification"]
        == "domain_style_advisory"
    )
    assert by_name["theology_style_deterministic_issues_present"]["raw_theology_style_deterministic_issue_count"] == 3
    assert "toc_body_concatenation_detected" not in by_name


def test_build_source_cleanup_evidence_reports_artifact_summary(tmp_path) -> None:
    validation = _load_validation_module()
    artifact_path = tmp_path / "layout_cleanup.json"
    artifact_path.write_text("{}", encoding="utf-8")
    cleanup_report = LayoutArtifactCleanupReport(
        original_paragraph_count=6,
        cleaned_paragraph_count=3,
        removed_paragraph_count=3,
        removed_page_number_count=1,
        removed_repeated_artifact_count=2,
        removed_empty_or_whitespace_count=0,
        decisions=[
            LayoutArtifactCleanupDecision(
                original_source_index=0,
                original_paragraph_id="p0000",
                origin_raw_indexes=(0,),
                text_preview="1",
                action="remove",
                reason="page_number_pattern",
                confidence="high",
                normalized_text="1",
                repeat_count=1,
                page_number=1,
                layout_origin="paragraph",
            ),
            LayoutArtifactCleanupDecision(
                original_source_index=1,
                original_paragraph_id="p0001",
                origin_raw_indexes=(1,),
                text_preview="Introduction",
                action="keep",
                reason="uncertain_repeated_artifact",
                confidence="low",
                normalized_text="introduction",
                repeat_count=2,
                page_number=None,
                layout_origin="paragraph",
            ),
        ],
        cleanup_applied=True,
        cleanup_mode="remove",
        artifact_path=str(artifact_path),
    )

    evidence = validation._build_source_cleanup_evidence(cleanup_report)

    assert evidence is not None
    assert evidence["cleanup_mode"] == "remove"
    assert evidence["artifact_path"] == str(artifact_path).replace("\\", "/")
    assert evidence["removed_paragraph_count"] == 3
    assert evidence["reason_counts"]["page_number_pattern"] == 1
    assert evidence["reason_counts"]["uncertain_repeated_artifact"] == 1
    assert evidence["removed_samples"][0]["page_number"] == 1
    assert evidence["kept_uncertain_samples"][0]["reason"] == "uncertain_repeated_artifact"


def test_evaluate_lietaer_acceptance_uses_authoritative_structural_markdown_counts() -> None:
    validation = _load_validation_module()

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "translation_quality_report": {
            "bullet_heading_count": 0,
            "false_fragment_heading_count": 0,
            "false_fragment_heading_gate_source": "entry_assembly",
            "raw_false_fragment_heading_count": 2,
            "residual_bullet_glyph_count": 0,
            "residual_bullet_glyph_gate_source": "legacy_markdown",
            "raw_residual_bullet_glyph_count": 0,
            "list_fragment_regression_count": 0,
            "list_fragment_regression_gate_source": "topology_projection",
            "raw_list_fragment_regression_count": 1,
            "mixed_script_term_count": 0,
            "theology_style_deterministic_issue_count": 0,
            "toc_body_concat_detected": False,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(report)
    by_name = {check["name"]: check for check in acceptance["checks"]}

    assert "false_fragment_headings_present" not in acceptance["failed_checks"]
    assert "list_fragment_regressions_present" not in acceptance["failed_checks"]
    assert by_name["false_fragment_headings_present"]["false_fragment_heading_count"] == 0
    assert by_name["false_fragment_headings_present"]["false_fragment_heading_gate_source"] == "entry_assembly"
    assert by_name["false_fragment_headings_present"]["raw_false_fragment_heading_count"] == 2
    assert by_name["residual_bullet_glyphs_present"]["residual_bullet_glyph_count"] == 0
    assert by_name["residual_bullet_glyphs_present"]["residual_bullet_glyph_gate_source"] == "legacy_markdown"
    assert by_name["residual_bullet_glyphs_present"]["raw_residual_bullet_glyph_count"] == 0
    assert by_name["list_fragment_regressions_present"]["list_fragment_regression_count"] == 0
    assert by_name["list_fragment_regressions_present"]["list_fragment_regression_gate_source"] == "topology_projection"
    assert by_name["list_fragment_regressions_present"]["raw_list_fragment_regression_count"] == 1


def test_evaluate_lietaer_acceptance_tolerates_review_only_list_fragment_residue() -> None:
    validation = _load_validation_module()

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "translation_quality_report": {
            "bullet_heading_count": 0,
            "false_fragment_heading_count": 0,
            "residual_bullet_glyph_count": 0,
            "list_fragment_regression_count": 1,
            "list_fragment_regression_gate_source": "entry_assembly",
            "raw_list_fragment_regression_count": 47,
            "mixed_script_term_count": 0,
            "theology_style_deterministic_issue_count": 0,
            "toc_body_concat_detected": False,
            "quality_status": "warn",
            "gate_reasons": ["list_fragment_regressions_review_required"],
            "formatting_review_items": [
                {
                    "reason": "list_fragment_regressions_review_required",
                    "label": "Одиночный номер в сносках или библиографии",
                    "sample": {
                        "line": 42,
                        "text": "1489.",
                        "reason": "list_fragment_regressions_present",
                    },
                }
            ],
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(report)
    by_name = {check["name"]: check for check in acceptance["checks"]}

    assert "list_fragment_regressions_present" not in acceptance["failed_checks"]
    assert by_name["list_fragment_regressions_present"]["passed"] is True
    assert by_name["list_fragment_regressions_present"]["review_reason"] == "list_fragment_regressions_review_required"
    assert by_name["list_fragment_regressions_present"]["list_fragment_regression_count"] == 1


def test_evaluate_lietaer_acceptance_tolerates_review_only_legacy_hygiene_residue() -> None:
    validation = _load_validation_module()

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "translation_quality_report": {
            "quality_status": "warn",
            "gate_reasons": [
                "false_fragment_headings_review_required",
                "mixed_script_terms_review_required",
            ],
            "bullet_heading_count": 0,
            "false_fragment_heading_count": 1,
            "false_fragment_heading_gate_source": "legacy_markdown",
            "raw_false_fragment_heading_count": 1,
            "residual_bullet_glyph_count": 0,
            "list_fragment_regression_count": 0,
            "mixed_script_term_count": 1,
            "mixed_script_term_gate_source": "legacy_markdown",
            "mixed_script_term_classification": "non_structural_hygiene",
            "raw_mixed_script_term_count": 1,
            "theology_style_deterministic_issue_count": 0,
            "toc_body_concat_detected": False,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(report)

    assert "false_fragment_headings_present" not in acceptance["failed_checks"]
    assert "mixed_script_terms_present" not in acceptance["failed_checks"]


def test_evaluate_lietaer_acceptance_fails_failed_translation_quality_report() -> None:
    validation = _load_validation_module()

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "translation_quality_report": {
            "quality_status": "fail",
            "gate_reasons": ["role_loss_above_manual_review_threshold"],
            "bullet_heading_count": 0,
            "false_fragment_heading_count": 0,
            "residual_bullet_glyph_count": 0,
            "list_fragment_regression_count": 0,
            "mixed_script_term_count": 0,
            "theology_style_deterministic_issue_count": 0,
            "toc_body_concat_detected": False,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(report)
    by_name = {check["name"]: check for check in acceptance["checks"]}

    assert acceptance["passed"] is False
    assert "translation_quality_report_not_failed" in acceptance["failed_checks"]
    assert by_name["translation_quality_report_not_failed"]["quality_status"] == "fail"
    assert by_name["translation_quality_report_not_failed"]["gate_reasons"] == [
        "role_loss_above_manual_review_threshold"
    ]


def test_prod_pipeline_quality_report_matches_validation_harness_replay_basis(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    quality_dir = tmp_path / "quality_reports"
    diagnostics_dir.mkdir(parents=True, exist_ok=True)
    quality_dir.mkdir(parents=True, exist_ok=True)

    monkeypatch.setattr(document_pipeline, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    runtime = {"state": {}, "finalize": [], "activity": [], "log": [], "status": []}

    def _emit_state(runtime, **values):
        runtime.setdefault("state", {}).update(values)

    def _emit_finalize(runtime, stage, detail, progress, terminal_kind=None):
        runtime.setdefault("finalize", []).append((stage, detail, progress, terminal_kind))

    def _emit_activity(runtime, message):
        runtime.setdefault("activity", []).append(message)

    def _emit_log(runtime, **payload):
        runtime.setdefault("log", []).append(payload)

    def _emit_status(runtime, **payload):
        runtime.setdefault("status", []).append(payload)

    def _preserve_with_role_aware_diagnostics(docx_bytes, paragraphs, generated_paragraph_registry=None):
        (diagnostics_dir / "restore_replay.json").write_text(
            json.dumps(
                {
                    "stage": "restore",
                    "source_count": 2,
                    "target_count": 2,
                    "mapped_count": 1,
                    "unmapped_source_ids": ["p0001"],
                    "unmapped_target_indexes": [1],
                    "unmapped_source_residual_diagnostics": {
                        "effective_formatting_coverage_diagnostics": {
                            "counts": {
                                "format_neutral_body_residue": 1,
                            },
                            "format_neutral_creditable_count": 1,
                        },
                        "samples": [
                            {
                                "paragraph_id": "p0001",
                                "source_index": 1,
                                "role": "body",
                                "structural_role": "body",
                                "text_preview": "Small note",
                                "effective_formatting_coverage_class": "format_neutral_body_residue",
                            }
                        ],
                    },
                    "unmapped_target_residual_diagnostics": {
                        "split_accounting_creditable_count": 1,
                        "controlled_fallback_creditable_count": 1,
                        "counts": {"controlled_fallback_covered": 1},
                        "residual_rows": [
                            {
                                "target_index": 1,
                                "target_text_preview": "Fallback paragraph",
                                "residual_class": "controlled_fallback_covered",
                                "controlled_fallback_kind": "english_residual_output",
                                "controlled_fallback_block_index": 7,
                            }
                        ],
                    },
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        return docx_bytes

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[ParagraphUnit(text="Body", role="body", structural_role="body", paragraph_id="p0000")],
        image_assets=[],
        image_mode="safe",
        app_config={"translation_output_quality_gate_policy": "strict"},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="translate",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Processed block",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=lambda markdown_text, image_assets: {},
        convert_markdown_to_docx_bytes=lambda markdown_text: b"docx-bytes",
        preserve_source_paragraph_properties=_preserve_with_role_aware_diagnostics,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
    )

    quality_report_files = list(quality_dir.glob("*.json"))
    assert result == "succeeded"
    assert len(quality_report_files) == 1
    prod_quality_report = json.loads(quality_report_files[0].read_text(encoding="utf-8"))

    benchmark_report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [json.loads((diagnostics_dir / "restore_replay.json").read_text(encoding="utf-8"))],
        "translation_quality_report": prod_quality_report,
    }

    source_doc = Document()
    source_doc.add_paragraph("Body")
    output_doc = Document()
    output_doc.add_paragraph("Processed block")

    acceptance = validation.evaluate_lietaer_acceptance(
        benchmark_report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
        mismatch_threshold=0,
        unmapped_target_threshold=0,
    )
    by_name = {check["name"]: check for check in acceptance["checks"]}

    assert acceptance["passed"] is True
    assert by_name["translation_quality_report_not_failed"]["quality_status"] == prod_quality_report["quality_status"]
    assert by_name["translation_quality_report_not_failed"]["gate_reasons"] == prod_quality_report["gate_reasons"]
    assert by_name["unmapped_source_threshold"]["count_basis"] == prod_quality_report["unmapped_source_count_basis"]
    assert by_name["unmapped_target_threshold"]["count_basis"] == prod_quality_report["unmapped_target_count_basis"]
    assert prod_quality_report["quality_status"] == "warn"
    assert prod_quality_report["gate_reasons"] == ["controlled_fallback_blocks_review_required"]
    assert prod_quality_report["unmapped_source_count_basis"] == "role_aware_formatting_coverage"
    assert prod_quality_report["unmapped_target_count_basis"] == "role_aware_formatting_coverage"


def test_evaluate_lietaer_acceptance_ignores_centered_heading_alignment_for_minimal_formatter_contract() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_heading = source_doc.add_paragraph("Глава 1", style="Heading 1")
    source_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_attribution = source_doc.add_paragraph("ЭПИКТЕТ")
    source_attribution.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_doc = Document()
    output_doc.add_paragraph("Глава 1", style="Heading 1")
    output_doc.add_paragraph("ЭПИКТЕТ")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    centered_check = next(check for check in acceptance["checks"] if check["name"] == "centered_short_paragraphs_preserved")
    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")

    assert heading_check["passed"] is True
    assert heading_check["missing"] == []
    assert centered_check["passed"] is False
    assert centered_check["missing"] == ["эпиктет"]
    assert centered_check["source_centered_count"] == 1


def test_extract_allowlisted_centered_paragraph_texts_excludes_centered_chapter_marker_without_heading_style() -> None:
    validation = _load_validation_module()

    document = Document()
    chapter_marker = document.add_paragraph("ГЛАВА 1")
    chapter_marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    attribution = document.add_paragraph("ЭПИКТЕТ")
    attribution.alignment = WD_ALIGN_PARAGRAPH.CENTER

    assert sorted(validation._extract_allowlisted_centered_paragraph_texts(document)) == ["эпиктет"]


def test_evaluate_lietaer_acceptance_translate_mode_relaxes_source_language_heading_and_numbering_checks() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Introduction: Making versus Taking", style="Heading 1")
    source_doc.add_paragraph("First item", style="List Number")
    source_doc.add_paragraph("Second item", style="List Number")

    output_doc = Document()
    output_doc.add_paragraph("Введение: созидание и присвоение", style="Heading 1")
    output_doc.add_paragraph("Первый пункт")
    output_doc.add_paragraph("Второй пункт")

    report = {
        "result": "succeeded",
        "runtime_config": {"effective": {"processing_operation": "translate"}},
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")
    numbering_check = next(check for check in acceptance["checks"] if check["name"] == "word_numbering_preserved")

    assert heading_check["passed"] is True
    assert heading_check["missing"] == []
    assert numbering_check["passed"] is True
    assert numbering_check["processing_operation"] == "translate"


def test_evaluate_lietaer_acceptance_translate_mode_ignores_numeric_and_english_heading_renames() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Bowl Judgment #1 (Revelation 16:3):", style="Heading 2")
    source_doc.add_paragraph("Year 3 (sometime between 2028-2036:)", style="Heading 2")
    source_doc.add_paragraph("In Conclusion", style="Heading 2")

    output_doc = Document()
    output_doc.add_paragraph("Суд над чашей № 1 (Откровение 16:3):", style="Heading 2")
    output_doc.add_paragraph("Год 3 (где-то между 2028 и 2036)", style="Heading 2")
    output_doc.add_paragraph("В заключение", style="Heading 2")

    report = {
        "result": "succeeded",
        "runtime_config": {"effective": {"processing_operation": "translate"}},
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")

    assert heading_check["passed"] is True
    assert heading_check["missing"] == []


def test_evaluate_lietaer_acceptance_translate_mode_only_enforces_translation_stable_scripture_heading() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("(Matthew 24:36)", style="Heading 2")
    source_doc.add_paragraph("Year 3 (sometime between 2028-2036:)", style="Heading 2")
    source_doc.add_paragraph("A Secular Timeline for AI's Destruction of Humanity", style="Heading 2")

    output_doc = Document()
    output_doc.add_paragraph("(Матфея 24:36)", style="Heading 2")
    output_doc.add_paragraph("Год 3 (примерно между 2028 и 2036:)", style="Heading 2")
    output_doc.add_paragraph("Светская хронология уничтожения человечества ИИ", style="Heading 2")

    report = {
        "result": "succeeded",
        "runtime_config": {"effective": {"processing_operation": "translate"}},
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")

    assert heading_check["passed"] is True
    assert heading_check["source_heading_count"] == 1
    assert heading_check["missing"] == []


def test_evaluate_lietaer_acceptance_translate_mode_matches_scripture_heading_by_reference_anchor() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("(Matthew 24:36)", style="Heading 2")

    output_doc = Document()
    output_doc.add_paragraph("(Матфея 24:36)", style="Heading 2")

    report = {
        "result": "succeeded",
        "runtime_config": {"effective": {"processing_operation": "translate"}},
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")

    assert heading_check["passed"] is True
    assert heading_check["missing"] == []


def test_evaluate_lietaer_acceptance_translate_mode_ignores_numeric_body_marker_heading_token() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("11,12", style="Heading 2")

    output_doc = Document()
    output_doc.add_paragraph("Деньги против бартера", style="Heading 2")

    report = {
        "result": "succeeded",
        "runtime_config": {"effective": {"processing_operation": "translate"}},
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")

    assert heading_check["passed"] is True
    assert heading_check["missing"] == []
    assert heading_check["source_heading_count"] == 0


def test_evaluate_lietaer_acceptance_translate_mode_ignores_page_range_heading_token() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("179–180", style="Heading 2")

    output_doc = Document()
    output_doc.add_paragraph("Указатель", style="Heading 1")

    report = {
        "result": "succeeded",
        "runtime_config": {"effective": {"processing_operation": "translate"}},
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")

    assert heading_check["passed"] is True
    assert heading_check["missing"] == []
    assert heading_check["source_heading_count"] == 0


def test_evaluate_lietaer_acceptance_translate_mode_ignores_comma_plus_page_range_heading_token() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("182, 192–193", style="Heading 2")

    output_doc = Document()
    output_doc.add_paragraph("Указатель", style="Heading 1")

    report = {
        "result": "succeeded",
        "runtime_config": {"effective": {"processing_operation": "translate"}},
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")

    assert heading_check["passed"] is True
    assert heading_check["missing"] == []
    assert heading_check["source_heading_count"] == 0


def test_evaluate_lietaer_acceptance_translate_mode_still_requires_meaningful_heading() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Устойчивое изобилие", style="Heading 2")

    output_doc = Document()
    output_doc.add_paragraph("Другая тема", style="Heading 2")

    report = {
        "result": "succeeded",
        "runtime_config": {"effective": {"processing_operation": "translate"}},
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")

    assert heading_check["passed"] is False
    assert heading_check["missing"] == ["устойчивое изобилие"]
    assert heading_check["source_heading_count"] == 1


def test_evaluate_lietaer_acceptance_translate_mode_does_not_blanket_exclude_reference_heading() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Справочный указатель", style="Heading 1")

    output_doc = Document()
    output_doc.add_paragraph("Справочный указатель", style="Heading 1")

    report = {
        "result": "succeeded",
        "runtime_config": {"effective": {"processing_operation": "translate"}},
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")

    assert heading_check["passed"] is True
    assert heading_check["missing"] == []
    assert heading_check["source_heading_count"] == 1


def test_normalize_structural_text_strips_markdown_wrappers() -> None:
    validation = _load_validation_module()

    assert validation._normalize_structural_text("## **Религии безличного абсолюта**") == "религии безличного абсолюта"
    assert validation._normalize_structural_text("<u>Дао</u>") == "дао"


def test_evaluate_lietaer_acceptance_ignores_short_garbage_heading_and_markdown_wrapped_heading() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("ð¢", style="Heading 1")
    source_doc.add_paragraph("**Религии безличного абсолюта**", style="Heading 2")

    output_doc = Document()
    output_doc.add_paragraph("Т", style="Heading 1")
    output_doc.add_paragraph("Религии безличного абсолюта", style="Heading 2")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    heading_check = next(check for check in acceptance["checks"] if check["name"] == "key_headings_preserved")

    assert heading_check["passed"] is True
    assert heading_check["missing"] == []
    assert heading_check["source_heading_count"] == 1
    assert heading_check["output_heading_count"] == 1


def test_evaluate_lietaer_acceptance_detects_known_false_split_in_runtime_markdown() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Однако деньги — не единственное средство обмена.")
    output_doc = Document()
    output_doc.add_paragraph("Однако деньги — не единственное средство обмена.")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "runtime": {
            "state": {
                "latest_markdown": "Вы помогаете соседу установить\n\nустановить новую крышу.",
                "processed_block_markdowns": ["Вы помогаете соседу установить\n\nустановить новую крышу."],
            }
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    assert acceptance["passed"] is False
    assert "known_false_split_absent_in_final_markdown:lietaer_exchange_install_roof_split" in acceptance["failed_checks"]
    assert "known_false_split_absent_in_processed_markdown:lietaer_exchange_install_roof_split" in acceptance["failed_checks"]


def test_evaluate_lietaer_acceptance_classifies_placeholder_heading_concat_as_display_hygiene() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Однако деньги — не единственное средство обмена.")
    output_doc = Document()
    output_doc.add_paragraph("Однако деньги — не единственное средство обмена.")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "runtime": {
            "state": {
                "latest_markdown": "This page intentionally left blank\n\nChapter Nine STRATEGIES FOR NGO S",
                "processed_block_markdowns": ["This page intentionally left blank Chapter Nine STRATEGIES FOR NGO S"],
            }
        },
        "translation_quality_report": {
            "page_placeholder_heading_concat_count": 0,
            "page_placeholder_heading_concat_source": "legacy_markdown",
            "page_placeholder_heading_concat_classification": "display_hygiene",
            "raw_page_placeholder_heading_concat_count": 1,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )
    by_name = {check["name"]: check for check in acceptance["checks"]}

    assert "page_placeholder_heading_concat_hygiene_applied" not in acceptance["failed_checks"]
    assert by_name["page_placeholder_heading_concat_hygiene_applied"]["page_placeholder_heading_concat_count"] == 0
    assert by_name["page_placeholder_heading_concat_hygiene_applied"]["raw_page_placeholder_heading_concat_count"] == 1
    assert by_name["page_placeholder_heading_concat_hygiene_applied"]["page_placeholder_heading_concat_source"] == "legacy_markdown"
    assert by_name["page_placeholder_heading_concat_hygiene_applied"]["page_placeholder_heading_concat_classification"] == "display_hygiene"


def test_evaluate_lietaer_acceptance_preserves_richer_formatting_diagnostics_payload() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Это один логический абзац после нормализации.")
    output_doc = Document()
    output_doc.add_paragraph("Это один логический абзац после нормализации.")

    formatting_payload = {
        "accepted_merged_sources": [
            {
                "logical_paragraph_id": "p0012",
                "origin_raw_indexes": [12, 13, 14],
                "accepted_merged_sources_count": 3,
                "target_index": 0,
            }
        ],
        "accepted_merged_sources_count": 1,
        "max_accepted_merged_sources": 3,
        "unmapped_source_ids": [],
        "unmapped_target_indexes": [],
    }
    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [formatting_payload],
        "metrics": {
            "accepted_merged_sources_count": 1,
            "max_accepted_merged_sources": 3,
        },
        "runtime": {
            "state": {
                "latest_markdown": "Это один логический абзац после нормализации.",
                "processed_block_markdowns": ["Это один логический абзац после нормализации."],
            }
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    assert acceptance["passed"] is True
    assert report["formatting_diagnostics"] == [formatting_payload]
    assert report["metrics"]["accepted_merged_sources_count"] == 1
    assert report["metrics"]["max_accepted_merged_sources"] == 3


def test_evaluate_lietaer_acceptance_emits_explicit_unmapped_threshold_checks() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Один абзац")
    output_doc = Document()
    output_doc.add_paragraph("Один абзац")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [
            {
                "unmapped_source_ids": ["p0003", "p0004"],
                "unmapped_target_indexes": [12],
            }
        ],
        "translation_quality_report": {
            "worst_unmapped_source_count": 2,
            "unmapped_target_count": 1,
            "toc_body_concat_detected": False,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
        mismatch_threshold=2,
        unmapped_target_threshold=1,
    )

    by_name = {check["name"]: check for check in acceptance["checks"]}
    assert by_name["unmapped_source_threshold"]["passed"] is True
    assert by_name["unmapped_source_threshold"]["actual"] == 2
    assert by_name["unmapped_source_threshold"]["allowed"] == 2
    assert by_name["unmapped_target_threshold"]["passed"] is True
    assert by_name["unmapped_target_threshold"]["actual"] == 1
    assert by_name["unmapped_target_threshold"]["allowed"] == 1
    assert "structural_comparison_available" not in by_name


def test_evaluate_lietaer_acceptance_discounts_payload_backed_benign_unmapped_source_merges() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Один абзац")
    output_doc = Document()
    output_doc.add_paragraph("Один абзац")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [
            {
                "unmapped_source_ids": ["p0000", "p0002", "p0005", "p0006"],
                "unmapped_target_indexes": [],
                "source_registry": [
                    {
                        "paragraph_id": "p0000",
                        "source_index": 0,
                        "role": "body",
                        "structural_role": "body",
                        "mapped_target_index": None,
                        "text_preview": "глава 1",
                    },
                    {
                        "paragraph_id": "p0001",
                        "source_index": 1,
                        "role": "heading",
                        "structural_role": "heading",
                        "mapped_target_index": 0,
                        "text_preview": "что такое богатство?",
                    },
                    {
                        "paragraph_id": "p0002",
                        "source_index": 2,
                        "role": "image",
                        "structural_role": "image",
                        "asset_id": "img_001",
                        "mapped_target_index": None,
                        "text_preview": "[[docx_image_img_001]]",
                    },
                    {
                        "paragraph_id": "p0003",
                        "source_index": 3,
                        "role": "body",
                        "structural_role": "body",
                        "attached_to_asset_id": "img_001",
                        "mapped_target_index": 1,
                        "text_preview": "рисунок 1.1. Подпись к рисунку",
                    },
                    {
                        "paragraph_id": "p0004",
                        "source_index": 4,
                        "role": "body",
                        "structural_role": "body",
                        "mapped_target_index": 2,
                        "text_preview": "Предыдущий абзац без терминальной точки",
                    },
                    {
                        "paragraph_id": "p0005",
                        "source_index": 5,
                        "role": "body",
                        "structural_role": "body",
                        "mapped_target_index": None,
                        "text_preview": ", и эти усилия продолжаются в следующей строке.",
                    },
                    {
                        "paragraph_id": "p0006",
                        "source_index": 6,
                        "role": "body",
                        "structural_role": "body",
                        "mapped_target_index": None,
                        "text_preview": "эпиктет",
                    },
                ],
            }
        ],
        "translation_quality_report": {
            "worst_unmapped_source_count": 4,
            "unmapped_target_count": 0,
            "toc_body_concat_detected": False,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
        mismatch_threshold=1,
    )

    by_name = {check["name"]: check for check in acceptance["checks"]}
    assert by_name["formatting_diagnostics_threshold"]["passed"] is True
    assert by_name["formatting_diagnostics_threshold"]["actual"] == 1
    assert by_name["unmapped_source_threshold"]["passed"] is True
    assert by_name["unmapped_source_threshold"]["actual"] == 1


def test_evaluate_lietaer_acceptance_prefers_structure_unit_unmapped_basis_over_raw_formatting_counts() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Один абзац")
    output_doc = Document()
    output_doc.add_paragraph("Один абзац")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [
            {
                "unmapped_source_ids": ["p0003", "p0004"],
                "unmapped_target_indexes": [12, 13, 14],
                "unmapped_target_residual_diagnostics": {
                    "split_accounting_creditable_count": 2,
                },
            }
        ],
        "translation_quality_report": {
            "worst_unmapped_source_count": 1,
            "unmapped_source_count": 1,
            "raw_unmapped_source_paragraph_count": 2,
            "structure_unit_unmapped_source_count": 1,
            "unmapped_source_count_basis": "topology_unit",
            "unmapped_target_count": 1,
            "raw_unmapped_target_paragraph_count": 2,
            "structure_unit_unmapped_target_count": 1,
            "unmapped_target_count_basis": "topology_unit",
            "toc_body_concat_detected": False,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
        mismatch_threshold=1,
        unmapped_target_threshold=1,
    )

    by_name = {check["name"]: check for check in acceptance["checks"]}
    assert by_name["formatting_diagnostics_threshold"]["passed"] is True
    assert by_name["formatting_diagnostics_threshold"]["actual"] == 1
    assert by_name["formatting_diagnostics_threshold"]["raw_worst_unmapped_source_count"] == 2
    assert by_name["formatting_diagnostics_threshold"]["unmapped_source_count_basis"] == "topology_unit"
    assert by_name["unmapped_source_threshold"]["passed"] is True
    assert by_name["unmapped_source_threshold"]["actual"] == 1
    assert by_name["unmapped_target_threshold"]["passed"] is True
    assert by_name["unmapped_target_threshold"]["actual"] == 1
    assert by_name["unmapped_target_threshold"]["count_basis"] == "role_aware_formatting_coverage"
    assert by_name["unmapped_target_threshold"]["raw_unmapped_target_count"] == 3
    assert by_name["unmapped_target_threshold"]["role_aware_effective_unmapped_target_count"] == 1
    assert by_name["unmapped_target_threshold"]["target_split_accounting_creditable_count"] == 2
    assert by_name["unmapped_target_threshold"]["quality_unmapped_target_count"] == 1


def test_evaluate_lietaer_acceptance_prefers_accepted_aggregation_legacy_basis_over_raw_counts() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Один абзац")
    output_doc = Document()
    output_doc.add_paragraph("Один абзац")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [
            {
                "unmapped_source_ids": ["p0003", "p0004"],
                "unmapped_target_indexes": [12],
                "accepted_aggregated_sources": [
                    {"paragraph_id": "p0003", "target_index": 12},
                    {"paragraph_id": "p0004", "target_index": 12},
                ],
            }
        ],
        "translation_quality_report": {
            "worst_unmapped_source_count": 2,
            "unmapped_source_count": 2,
            "raw_unmapped_source_paragraph_count": 2,
            "structure_unit_unmapped_source_count": 0,
            "unmapped_source_count_basis": "accepted_aggregation_legacy",
            "unmapped_target_count": 1,
            "raw_unmapped_target_paragraph_count": 1,
            "structure_unit_unmapped_target_count": 0,
            "unmapped_target_count_basis": "accepted_aggregation_legacy",
            "accepted_aggregated_source_unit_count": 2,
            "accepted_aggregated_target_index_count": 1,
            "toc_body_concat_detected": False,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
        mismatch_threshold=0,
        unmapped_target_threshold=0,
    )

    by_name = {check["name"]: check for check in acceptance["checks"]}
    assert by_name["formatting_diagnostics_threshold"]["passed"] is True
    assert by_name["formatting_diagnostics_threshold"]["actual"] == 0
    assert by_name["formatting_diagnostics_threshold"]["raw_worst_unmapped_source_count"] == 2
    assert by_name["formatting_diagnostics_threshold"]["unmapped_source_count_basis"] == "accepted_aggregation_legacy"
    assert by_name["unmapped_source_threshold"]["passed"] is True
    assert by_name["unmapped_source_threshold"]["actual"] == 0
    assert by_name["unmapped_target_threshold"]["passed"] is True
    assert by_name["unmapped_target_threshold"]["actual"] == 0


def test_evaluate_lietaer_acceptance_uses_role_aware_effective_formatting_source_count() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Один абзац")
    output_doc = Document()
    output_doc.add_paragraph("Один абзац")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [
            {
                "unmapped_source_ids": ["p0001", "p0002", "p0003"],
                "unmapped_target_indexes": [],
                "unmapped_source_residual_diagnostics": {
                    "effective_formatting_coverage_diagnostics": {
                        "classification_basis": "full_unmapped_source_set",
                        "format_neutral_creditable_count": 2,
                    }
                },
            }
        ],
        "translation_quality_report": {
            "worst_unmapped_source_count": 3,
            "unmapped_source_count": 3,
            "unmapped_target_count": 0,
            "toc_body_concat_detected": False,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
        mismatch_threshold=1,
    )

    by_name = {check["name"]: check for check in acceptance["checks"]}
    assert by_name["formatting_diagnostics_threshold"]["passed"] is True
    assert by_name["formatting_diagnostics_threshold"]["actual"] == 1
    assert by_name["formatting_diagnostics_threshold"]["raw_worst_unmapped_source_count"] == 3
    assert by_name["formatting_diagnostics_threshold"]["role_aware_effective_unmapped_source_count"] == 1
    assert by_name["formatting_diagnostics_threshold"]["format_neutral_creditable_count"] == 2
    assert by_name["formatting_diagnostics_threshold"]["unmapped_source_count_basis"] == "role_aware_formatting_coverage"
    assert by_name["unmapped_source_threshold"]["passed"] is True
    assert by_name["unmapped_source_threshold"]["actual"] == 1
    assert by_name["unmapped_source_threshold"]["raw_worst_unmapped_source_count"] == 3
    assert by_name["unmapped_source_threshold"]["count_basis"] == "role_aware_formatting_coverage"
    assert by_name["unmapped_source_threshold"]["format_neutral_creditable_count"] == 2


def test_evaluate_lietaer_acceptance_uses_role_aware_effective_formatting_target_count() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Один абзац")
    output_doc = Document()
    output_doc.add_paragraph("Один абзац")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [
            {
                "unmapped_source_ids": [],
                "unmapped_target_indexes": [1, 2, 3],
                "unmapped_target_residual_diagnostics": {
                    "split_accounting_creditable_count": 2,
                },
            }
        ],
        "translation_quality_report": {
            "worst_unmapped_source_count": 0,
            "unmapped_source_count": 0,
            "unmapped_target_count": 3,
            "toc_body_concat_detected": False,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
        mismatch_threshold=0,
        unmapped_target_threshold=1,
    )

    by_name = {check["name"]: check for check in acceptance["checks"]}
    assert by_name["unmapped_target_threshold"]["passed"] is True
    assert by_name["unmapped_target_threshold"]["actual"] == 1
    assert by_name["unmapped_target_threshold"]["raw_unmapped_target_count"] == 3
    assert by_name["unmapped_target_threshold"]["role_aware_effective_unmapped_target_count"] == 1
    assert by_name["unmapped_target_threshold"]["target_split_accounting_creditable_count"] == 2
    assert by_name["unmapped_target_threshold"]["count_basis"] == "role_aware_formatting_coverage"


def test_evaluate_lietaer_acceptance_emits_required_no_toc_body_concat_check() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Содержание")
    output_doc = Document()
    output_doc.add_paragraph("Содержание")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "translation_quality_report": {
            "toc_body_concat_detected": False,
        },
        "preparation_diagnostic_snapshot": {
            "toc_body_concat_detected": False,
            "toc_body_concat_markdown_detected": False,
            "toc_body_concat_gate_source": "legacy_markdown",
            "effective_source_toc_region_count": 1,
            "document_map_toc_region_count": 1,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
        require_no_toc_body_concat=True,
    )

    by_name = {check["name"]: check for check in acceptance["checks"]}
    assert by_name["no_toc_body_concat_required"]["passed"] is True
    assert by_name["no_toc_body_concat_required"]["toc_body_concat_gate_source"] == "legacy_markdown"


def test_evaluate_lietaer_acceptance_prefers_topology_toc_body_signal_for_required_gate() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Содержание")
    output_doc = Document()
    output_doc.add_paragraph("Содержание")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "translation_quality_report": {
            "toc_body_concat_detected": False,
        },
        "preparation_diagnostic_snapshot": {
            "toc_body_concat_detected": False,
            "toc_body_concat_markdown_detected": True,
            "toc_body_concat_structure_detected": False,
            "toc_body_concat_gate_source": "topology_projection",
            "structure_repair_toc_body_boundary_repairs": 0,
            "effective_source_toc_region_count": 0,
            "document_map_toc_region_count": 1,
            "topology_toc_entry_count": 2,
            "topology_split_compound_toc_operation_count": 1,
            "document_map_compound_toc_split_hint_count": 1,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
        require_no_toc_body_concat=True,
    )

    by_name = {check["name"]: check for check in acceptance["checks"]}
    assert by_name["no_toc_body_concat_required"]["passed"] is True
    assert by_name["no_toc_body_concat_required"]["toc_body_concat_gate_source"] == "topology_projection"
    assert by_name["no_toc_body_concat_required"]["toc_body_concat_markdown_detected"] is True
    assert by_name["no_toc_body_concat_required"]["toc_body_concat_structure_detected"] is False


def test_evaluate_lietaer_acceptance_uses_explicit_report_markdown_field_for_legacy_toc_gate_fallback() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_doc.add_paragraph("Содержание")
    output_doc = Document()
    output_doc.add_paragraph("Содержание")

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
        "translation_quality_report": {
            "toc_body_concat_detected": False,
            "toc_body_concat_markdown_detected": True,
            "toc_body_concat_structure_detected": False,
            "toc_body_concat_gate_source": "legacy_markdown",
        },
        "preparation_diagnostic_snapshot": {
            "toc_body_concat_detected": False,
            "effective_source_toc_region_count": 1,
            "document_map_toc_region_count": 1,
        },
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
        require_no_toc_body_concat=True,
    )

    by_name = {check["name"]: check for check in acceptance["checks"]}
    assert by_name["no_toc_body_concat_required"]["passed"] is False
    assert by_name["no_toc_body_concat_required"]["toc_body_concat_gate_source"] == "legacy_markdown"
    assert by_name["no_toc_body_concat_required"]["toc_body_concat_detected"] is False
    assert by_name["no_toc_body_concat_required"]["toc_body_concat_markdown_detected"] is True
    assert by_name["no_toc_body_concat_required"]["toc_body_concat_structure_detected"] is False


def test_evaluate_lietaer_acceptance_allows_centered_text_edits_when_alignment_is_preserved() -> None:
    validation = _load_validation_module()

    source_doc = Document()
    source_quote = source_doc.add_paragraph(
        "Богатство заключается не в том, чтобы иметь много имущества, а в том, чтобы иметь мало желаний."
    )
    source_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_caption = source_doc.add_paragraph("Рисунок 1.2. Взаимосвязь между потребностями, активами и капиталом")
    source_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_doc = Document()
    output_quote = output_doc.add_paragraph("Богатство — не в обилии имущества, а в умении довольствоваться малым.")
    output_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    output_caption = output_doc.add_paragraph("Рисунок 1.2. Взаимосвязь потребностей, активов и капитала")
    output_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    report = {
        "result": "succeeded",
        "output_artifacts": {
            "output_docx_openable": True,
            "output_contains_placeholder_markup": False,
        },
        "formatting_diagnostics": [],
    }

    acceptance = validation.evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=_docx_bytes(source_doc),
        output_docx_bytes=_docx_bytes(output_doc),
    )

    centered_check = next(check for check in acceptance["checks"] if check["name"] == "centered_short_paragraphs_preserved")
    assert centered_check["passed"] is True
    assert len(centered_check["matches"]) == 2


def test_centered_quote_similarity_allows_anchored_paraphrase_but_not_unrelated_short_text() -> None:
    validation = _load_validation_module()

    paraphrase_score = validation._centered_quote_similarity(
        "богатство заключается не в том, чтобы иметь много имущества, а в том, чтобы иметь мало желаний.",
        "богатство не в обилии имущества, а в умении довольствоваться малым.",
    )
    unrelated_score = validation._centered_quote_similarity(
        "богатство заключается не в том, чтобы иметь много имущества, а в том, чтобы иметь мало желаний.",
        "эпиктет",
    )

    assert paraphrase_score >= 0.55
    assert unrelated_score < 0.55


def test_match_centered_structural_texts_matches_normalized_single_token_attribution() -> None:
    validation = _load_validation_module()

    missing, matches = validation._match_centered_structural_texts(["эпиктет"], ["эпиктет"])

    assert missing == []
    assert matches == [{"source": "эпиктет", "output": "эпиктет", "similarity": 1.0}]


def test_count_ordered_word_numbered_paragraphs_handles_multilevel_numbering() -> None:
    validation = _load_validation_module()

    document = Document()
    _append_multilevel_numbering_definition(document, num_id="9001", abstract_num_id="9000")

    bullet_paragraph = document.add_paragraph("Маркер")
    first_ordered_paragraph = document.add_paragraph("Первый пункт")
    second_ordered_paragraph = document.add_paragraph("Второй пункт")

    _attach_numbering(bullet_paragraph, num_id="9001", ilvl="0")
    _attach_numbering(first_ordered_paragraph, num_id="9001", ilvl="1")
    _attach_numbering(second_ordered_paragraph, num_id="9001", ilvl="1")

    assert validation._count_ordered_word_numbered_paragraphs(document) == 2


def test_summarize_repeat_runs_detects_intermittent_failures() -> None:
    validation = _load_validation_module()

    summary, acceptance, failure_classification = validation._summarize_repeat_runs(
        [
            {
                "repeat_index": 1,
                "run_id": "run-1",
                "result": "succeeded",
                "acceptance_passed": True,
                "failure_classification": None,
                "signals": {"heading_only_output_detected": False},
            },
            {
                "repeat_index": 2,
                "run_id": "run-2",
                "result": "failed",
                "acceptance_passed": False,
                "failure_classification": "heading_only_output",
                "signals": {"heading_only_output_detected": True},
            },
            {
                "repeat_index": 3,
                "run_id": "run-3",
                "result": "succeeded",
                "acceptance_passed": True,
                "failure_classification": None,
                "signals": {"heading_only_output_detected": False},
            },
        ]
    )

    assert summary["repeat_count"] == 3
    assert summary["pipeline_succeeded_count"] == 2
    assert summary["acceptance_passed_count"] == 2
    assert summary["intermittent_failure_detected"] is True
    assert summary["heading_only_output_detected_count"] == 1
    assert summary["failed_repeat_indexes"] == [2]
    assert acceptance["passed"] is False
    assert acceptance["failed_checks"] == ["all_repeat_runs_succeeded", "all_repeat_runs_acceptance_passed"]
    assert failure_classification == "intermittent_failure"


def test_build_validation_mode_payload_marks_comparison_only_runs() -> None:
    validation = _load_validation_module()

    comparison_only = validation._build_validation_mode_payload(
        validation.load_validation_registry().get_run_profile(
            "ui-parity-translate-simple-reader-cleanup-comparison-only"
        )
    )
    acceptance = validation._build_validation_mode_payload(
        validation.load_validation_registry().get_run_profile("ui-parity-default")
    )

    assert comparison_only == {
        "comparison_only_validation": True,
        "validation_run_type": "comparison_only",
        "acceptance_contract_active": False,
        "evidence_label": "comparison_only_non_acceptance",
        "success_criterion": "pipeline_result_and_artifacts",
    }
    assert acceptance == {
        "comparison_only_validation": False,
        "validation_run_type": "acceptance",
        "acceptance_contract_active": True,
        "evidence_label": "acceptance_contract",
        "success_criterion": "acceptance_passed",
    }


def test_resolve_validation_final_status_uses_comparison_only_success_criterion() -> None:
    validation = _load_validation_module()

    comparison_only = validation._build_validation_mode_payload(
        validation.load_validation_registry().get_run_profile(
            "ui-parity-translate-simple-reader-cleanup-comparison-only"
        )
    )
    acceptance = validation._build_validation_mode_payload(
        validation.load_validation_registry().get_run_profile("ui-parity-default")
    )

    assert (
        validation._resolve_validation_final_status(
            result="succeeded",
            acceptance_passed=False,
            validation_mode=comparison_only,
        )
        == "completed"
    )
    assert (
        validation._resolve_validation_final_status(
            result="succeeded",
            acceptance_passed=False,
            validation_mode=acceptance,
        )
        == "failed"
    )


def test_resolve_reader_verifier_config_is_off_by_default_for_proof_profiles() -> None:
    validation = _load_validation_module()

    config = validation._resolve_reader_verifier_config(
        validation_mode={
            "comparison_only_validation": True,
            "validation_run_type": "comparison_only",
        },
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        app_config=None,
        runtime_app_config={},
    )

    assert config["enabled"] is False
    assert config["model"] == "openrouter:google/gemini-3-flash-preview"


def test_resolve_reader_verifier_config_respects_explicit_opt_in_override() -> None:
    validation = _load_validation_module()

    config = validation._resolve_reader_verifier_config(
        validation_mode={
            "comparison_only_validation": True,
            "validation_run_type": "comparison_only",
        },
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        app_config=None,
        runtime_app_config={
            "reader_verifier_enabled": True,
            "reader_verifier_model": "anthropic:claude-sonnet-4-6",
        },
    )

    assert config == {
        "enabled": True,
        "model": "anthropic:claude-sonnet-4-6",
    }


def test_load_reader_cleanup_evidence_extracts_artifacts_and_delete_stats(tmp_path) -> None:
    validation = _load_validation_module()

    cleanup_report_path = tmp_path / "ui_results" / "chapter.reader_cleanup_report.json"
    cleanup_report_path.parent.mkdir(parents=True, exist_ok=True)
    cleanup_report_path.write_text(
        json.dumps(
            {
                "stage_status": "completed",
                "changed": True,
                "stats": {
                    "proposed_delete_block_count": 3,
                    "accepted_delete_block_count": 1,
                    "ignored_delete_block_count": 2,
                    "failed_chunk_count": 0,
                },
                "accepted_delete_blocks": [
                    {
                        "id": "b_000001",
                        "reason": "repeated_running_header",
                        "confidence": "high",
                        "raw_text_preview": "chapter eight",
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    event_log = [
        {
            "event_id": "ui_result_artifacts_saved",
            "context": {
                "artifact_paths": {
                    "markdown_path": str(tmp_path / "ui_results" / "chapter.result.md"),
                    "docx_path": str(tmp_path / "ui_results" / "chapter.result.docx"),
                    "reader_cleanup_raw_markdown_path": str(tmp_path / "ui_results" / "chapter.raw.result.md"),
                    "reader_cleanup_report_path": str(cleanup_report_path),
                }
            },
        }
    ]

    evidence = validation._load_reader_cleanup_evidence(event_log)

    assert evidence["cleaned_markdown_path"] == str(tmp_path / "ui_results" / "chapter.result.md")
    assert evidence["cleaned_docx_path"] == str(tmp_path / "ui_results" / "chapter.result.docx")
    assert evidence["raw_markdown_path"] == str(tmp_path / "ui_results" / "chapter.raw.result.md")
    assert evidence["reader_cleanup_report_path"] == str(cleanup_report_path)
    assert evidence["stage_status"] == "completed"
    assert evidence["changed"] is True
    assert evidence["accepted_delete_block_count"] == 1
    assert evidence["ignored_delete_block_count"] == 2
    assert evidence["rejected_delete_block_count"] == 0
    assert evidence["deleted_block_previews"] == [
        {
            "id": "b_000001",
            "reason": "repeated_running_header",
            "confidence": "high",
            "raw_text_preview": "chapter eight",
        }
    ]


def test_load_reader_cleanup_evidence_reports_runtime_anchor_repair_status(tmp_path) -> None:
    validation = _load_validation_module()

    cleanup_report_path = tmp_path / "ui_results" / "chapter.reader_cleanup_report.json"
    cleanup_report_path.parent.mkdir(parents=True, exist_ok=True)
    anchor_target = {
        "anchor_id": "anchor-1",
        "category": "heading_fused_with_body",
        "block_id": "b_000002",
        "line_ref": "cleaned_markdown:3",
        "snippet": "КАК ЭТО РАБОТАЕТ: Местные органы власти...",
    }
    cleanup_report_path.write_text(
        json.dumps(
            {
                "stage_status": "completed",
                "changed": True,
                "stats": {
                    "proposed_delete_block_count": 0,
                    "accepted_delete_block_count": 0,
                    "ignored_delete_block_count": 0,
                    "failed_chunk_count": 0,
                },
                "accepted_delete_blocks": [],
                "passes": {
                    "anchor_repair_pass": {
                        "selected_anchor_count": 1,
                        "selected_anchors": [anchor_target],
                        "stats": {
                            "accepted_cleanup_operation_count": 1,
                            "accepted_delete_block_count": 0,
                        },
                    }
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    evidence = validation._build_reader_cleanup_evidence_from_artifact_paths(
        {"reader_cleanup_report_path": str(cleanup_report_path)}
    )

    assert evidence["anchor_repair_status"] == "runtime_applied"
    assert evidence["recommended_anchor_targets"] == [anchor_target]
    assert evidence["recommended_anchor_target_count"] == 1

    cleanup_report_path.write_text(
        json.dumps(
            {
                "stage_status": "completed",
                "changed": False,
                "stats": {
                    "proposed_delete_block_count": 0,
                    "accepted_delete_block_count": 0,
                    "ignored_delete_block_count": 0,
                    "failed_chunk_count": 0,
                },
                "accepted_delete_blocks": [],
                "passes": {
                    "anchor_repair_pass": {
                        "selected_anchor_count": 1,
                        "selected_anchors": [anchor_target],
                        "stats": {
                            "accepted_cleanup_operation_count": 0,
                            "accepted_delete_block_count": 0,
                        },
                    }
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    evidence = validation._build_reader_cleanup_evidence_from_artifact_paths(
        {"reader_cleanup_report_path": str(cleanup_report_path)}
    )

    assert evidence["anchor_repair_status"] == "runtime_attempted_no_safe_ops"
    assert evidence["recommended_anchor_targets"] == [anchor_target]
    assert evidence["recommended_anchor_target_count"] == 1


def test_apply_repeat_count_override_ignores_invalid_value() -> None:
    validation = _load_validation_module()
    run_profile = validation.load_validation_registry().get_run_profile("ui-parity-default")
    buffer = io.StringIO()

    with redirect_stdout(buffer):
        updated = validation._apply_repeat_count_override(run_profile, "abc")

    assert updated.repeat_count == run_profile.repeat_count
    assert "invalid DOCXAI_REAL_DOCUMENT_REPEAT_COUNT_OVERRIDE" in buffer.getvalue()


def test_select_repeat_artifact_references_exposes_failing_and_successful_runs() -> None:
    validation = _load_validation_module()

    references = validation._select_repeat_artifact_references(
        [
            {
                "run_id": "run-1",
                "acceptance_passed": False,
                "report_path": "runs/run-1/report.json",
                "summary_path": "runs/run-1/summary.txt",
                "output_artifacts": {
                    "markdown_path": "runs/run-1/output.md",
                    "docx_path": "runs/run-1/output.docx",
                },
            },
            {
                "run_id": "run-2",
                "acceptance_passed": True,
                "report_path": "runs/run-2/report.json",
                "summary_path": "runs/run-2/summary.txt",
                "output_artifacts": {
                    "markdown_path": "runs/run-2/output.md",
                    "docx_path": "runs/run-2/output.docx",
                },
            },
        ]
    )

    assert references["first_failing_run_id"] == "run-1"
    assert references["first_failing_docx_path"] == "runs/run-1/output.docx"
    assert references["representative_success_run_id"] == "run-2"
    assert references["representative_success_markdown_path"] == "runs/run-2/output.md"


def test_write_latest_alias_artifacts_preserves_stable_manifest_schema(tmp_path) -> None:
    validation = _load_validation_module()

    report_path = tmp_path / "run_report.json"
    summary_path = tmp_path / "run_summary.txt"
    progress_path = tmp_path / "run_progress.json"
    markdown_path = tmp_path / "run_output.md"
    docx_path = tmp_path / "run_output.docx"
    latest_report_path = tmp_path / "latest_report.json"
    latest_summary_path = tmp_path / "latest_summary.txt"
    latest_markdown_path = tmp_path / "latest_output.md"
    latest_docx_path = tmp_path / "latest_output.docx"
    latest_manifest_path = tmp_path / "latest.json"

    report_path.write_text("{}", encoding="utf-8")
    summary_path.write_text("summary", encoding="utf-8")
    progress_path.write_text("{}", encoding="utf-8")
    markdown_path.write_text("markdown", encoding="utf-8")
    docx_path.write_bytes(b"PK")

    manifest_payload = {
        "run_id": "run-123",
        "document_profile_id": "lietaer-core",
        "run_profile_id": "ui-parity-default",
        "validation_tier": "full",
        "status": "completed",
        "report_json": validation._path_for_report(report_path),
        "summary_txt": validation._path_for_report(summary_path),
        "progress_json": validation._path_for_report(progress_path),
        "latest_progress_json": validation._path_for_report(progress_path),
        "acceptance_passed": True,
    }

    validation._write_latest_alias_artifacts(
        report_path=report_path,
        summary_path=summary_path,
        markdown_artifact=markdown_path,
        docx_artifact=docx_path,
        tts_artifact=None,
        latest_report_path=latest_report_path,
        latest_summary_path=latest_summary_path,
        latest_markdown_path=latest_markdown_path,
        latest_docx_path=latest_docx_path,
        latest_tts_path=None,
        latest_manifest_path=latest_manifest_path,
        run_id="run-123",
        run_dir=tmp_path,
        manifest_payload=manifest_payload,
    )

    latest_manifest = json.loads(latest_manifest_path.read_text(encoding="utf-8"))

    assert latest_manifest["status"] == "completed"
    assert latest_manifest["report_json"] == manifest_payload["report_json"]
    assert latest_manifest["summary_txt"] == manifest_payload["summary_txt"]
    assert latest_manifest["progress_json"] == manifest_payload["progress_json"]
    assert latest_manifest["latest_report"] == validation._path_for_report(latest_report_path)
    assert latest_manifest["latest_summary"] == validation._path_for_report(latest_summary_path)


def test_full_tier_runtime_contract_is_nested_only() -> None:
    validation = _load_validation_module()

    runtime_config = validation.build_validation_runtime_config(None)

    assert set(runtime_config.keys()) == {"effective", "ui_defaults", "overrides"}


def test_main_uses_processing_service_facade_and_runtime_config_only(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()
    from docxaicorrector.core.models import ImageAsset, ParagraphUnit

    source_path = tmp_path / "legacy.doc"
    source_bytes = bytes.fromhex("D0CF11E0A1B11AE1") + b"legacy-source"
    source_path.write_bytes(source_bytes)

    def _resolution_payload(**values):
        return SimpleNamespace(**values, to_dict=lambda: dict(values))

    document_profile = SimpleNamespace(
        id="lietaer-core",
        artifact_prefix="lietaer_validation",
        output_basename="lietaer_output",
        max_unmapped_source_paragraphs=0,
        max_unmapped_target_paragraphs=0,
        require_no_toc_body_concat=True,
        resolved_source_path=lambda project_root=None: source_path,
    )
    run_profile = SimpleNamespace(
        id="ui-parity-default",
        tier="full",
        repeat_count=1,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=False,
        model="gpt-5.4",
        max_retries=1,
    )
    registry = SimpleNamespace(
        get_document_profile=lambda profile_id: document_profile,
        resolve_run_profile=lambda profile, requested_run_profile_id: run_profile,
    )
    captured = {}
    quality_report_path = tmp_path / ".run" / "quality_reports" / "prepared_quality_report.json"
    quality_report_path.parent.mkdir(parents=True, exist_ok=True)
    quality_report_path.write_text(
        json.dumps(
            {
                "quality_status": "fail",
                "gate_reasons": [
                    "false_fragment_headings_present",
                    "residual_bullet_glyphs_present",
                ],
                "bullet_heading_count": 0,
                "bullet_heading_gate_source": "legacy_markdown",
                "bullet_heading_classification": "markdown_gate",
                "raw_bullet_heading_count": 0,
                "page_placeholder_heading_concat_count": 0,
                "page_placeholder_heading_concat_source": "legacy_markdown",
                "page_placeholder_heading_concat_classification": "display_hygiene",
                "raw_page_placeholder_heading_concat_count": 1,
                "false_fragment_heading_count": 2,
                "false_fragment_heading_gate_source": "entry_assembly",
                "raw_false_fragment_heading_count": 2,
                "residual_bullet_glyph_count": 1,
                "residual_bullet_glyph_gate_source": "legacy_markdown",
                "residual_bullet_glyph_classification": "display_hygiene",
                "raw_residual_bullet_glyph_count": 1,
                "list_fragment_regression_count": 0,
                "list_fragment_regression_gate_source": "topology_projection",
                "raw_list_fragment_regression_count": 1,
                "mixed_script_term_count": 2,
                "mixed_script_term_gate_source": "legacy_markdown",
                "mixed_script_term_classification": "non_structural_hygiene",
                "raw_mixed_script_term_count": 2,
                "theology_style_deterministic_issue_count": 0,
                "theology_style_deterministic_issue_source": "legacy_markdown",
                "theology_style_deterministic_issue_classification": "domain_style_advisory",
                "raw_theology_style_deterministic_issue_count": 0,
                "toc_body_concat_detected": False,
                "toc_body_concat_markdown_detected": False,
                "toc_body_concat_structure_detected": False,
                "toc_body_concat_gate_source": "legacy_markdown",
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    prepared_asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        width_emu=123,
        height_emu=456,
        source_forensics={"drawing_container": "inline", "source_rect": {"l": 10}},
    )
    processed_asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
        width_emu=123,
        height_emu=456,
        source_forensics={"drawing_container": "inline", "source_rect": {"l": 10}},
        final_decision="accept",
        final_variant="redrawn",
        final_reason="accepted",
    )
    processed_asset.update_runtime_attempt_state(validation_status="passed")
    prepared_paragraphs = [
        ParagraphUnit(text="Contents", role="body", structural_role="toc_header", source_index=0),
        ParagraphUnit(text="Chapter 1........ 12", role="body", structural_role="toc_entry", source_index=1),
        ParagraphUnit(text="-", role="body", structural_role="body", source_index=2),
        ParagraphUnit(text="Epigraph line", role="body", structural_role="epigraph", source_index=3),
        ParagraphUnit(text="Introduction", role="heading", structural_role="body", heading_level=1, source_index=4),
    ]
    structure_repair_report = SimpleNamespace(
        repaired_bullet_items=4,
        repaired_numbered_items=5,
        bounded_toc_regions=1,
        toc_body_boundary_repairs=1,
        heading_candidates_from_toc=7,
        remaining_isolated_marker_count=0,
    )

    class _ValidationServiceStub:
        def __init__(self, log_event_fn):
            self.log_event_fn = log_event_fn

        def run_prepared_background_document(self, **kwargs):
            uploaded_file = kwargs["uploaded_file"]
            captured["uploaded_filename"] = uploaded_file.name
            captured["uploaded_bytes"] = uploaded_file.getvalue()
            captured["processing_operation"] = kwargs["processing_operation"]
            captured["source_language"] = kwargs["source_language"]
            captured["target_language"] = kwargs["target_language"]
            self.log_event_fn(
                20,
                "structure_processing_outcome",
                "structure outcome",
                structure_ai_attempted=True,
                quality_gate_status="blocked",
                quality_gate_reasons=["structure_readiness_blocked_unsafe_best_effort_only"],
                readiness_status="blocked_unsafe_best_effort_only",
                readiness_reasons=["heading_count_far_below_toc_expectation"],
                ai_classified_count=7,
                ai_heading_count=3,
            )
            self.log_event_fn(
                20,
                "block_plan_summary",
                "block summary",
                block_count=3,
                llm_block_count=2,
                passthrough_block_count=1,
                first_block_target_chars=[3891, 946, 935],
            )
            self.log_event_fn(
                20,
                "quality_report_saved",
                "quality report saved",
                artifact_path=str(quality_report_path),
                quality_status="fail",
                gate_reasons=["false_fragment_headings_present", "residual_bullet_glyphs_present"],
            )
            kwargs["runtime"].emit(
                validation.SetStateEvent(
                    values={
                        "latest_markdown": "validated output",
                        "latest_docx_bytes": _docx_bytes(Document()),
                        "image_assets": [processed_asset],
                    }
                )
            )
            return "succeeded", SimpleNamespace(
                uploaded_filename="prepared.docx",
                uploaded_file_bytes=b"PK\x03\x04normalized-source",
                paragraphs=prepared_paragraphs,
                image_assets=[prepared_asset],
                jobs=[{"job_kind": "block"}],
                source_text="source text",
                preparation_cached=False,
                preparation_elapsed_seconds=0.1,
                ai_classified_count=7,
                ai_heading_count=3,
                ai_role_change_count=2,
                ai_heading_promotion_count=1,
                ai_heading_demotion_count=1,
                ai_structural_role_change_count=1,
                structure_repair_report=structure_repair_report,
                uploaded_file_token="token-1",
            )

    monkeypatch.setattr(validation, "PROJECT_ROOT", tmp_path)
    monkeypatch.setattr(validation, "REAL_DOCUMENT_ARTIFACT_ROOT", tmp_path / "artifacts")
    monkeypatch.setattr(validation, "load_validation_registry", lambda: registry)
    monkeypatch.setattr(validation, "load_app_config", lambda: object())
    monkeypatch.setattr(
        validation,
        "resolve_runtime_resolution",
        lambda app_config, run_profile: SimpleNamespace(
            effective=_resolution_payload(
                chunk_size=6000,
                image_mode="safe",
                keep_all_image_variants=False,
                model="gpt-5.4",
                max_retries=1,
                processing_operation="translate",
                source_language="en",
                target_language="ru",
            ),
            ui_defaults=_resolution_payload(
                chunk_size=6000,
                image_mode="safe",
                keep_all_image_variants=False,
                model="gpt-5.4",
                max_retries=1,
                processing_operation="edit",
                source_language="en",
                target_language="ru",
            ),
            overrides={},
        ),
    )
    monkeypatch.setattr(validation, "apply_runtime_resolution_to_app_config", lambda app_config, resolution: {"x": 1})
    monkeypatch.setattr(validation, "evaluate_lietaer_acceptance", lambda report, **kwargs: {"passed": True, "failed_checks": [], "checks": []})
    monkeypatch.setattr(validation, "_snapshot_formatting_diagnostics_paths", lambda: set())
    monkeypatch.setattr(validation, "_collect_new_formatting_diagnostics_paths", lambda before, after: [])
    monkeypatch.setattr(validation, "_extract_run_formatting_diagnostics_paths", lambda event_log: [])
    monkeypatch.setattr(validation, "_load_recent_formatting_diagnostics", lambda started_at: ([], []))
    monkeypatch.setattr(validation, "_load_formatting_diagnostics_payloads", lambda paths: [])
    monkeypatch.setattr(validation, "_print_terminal_completion_summary", lambda **kwargs: None)
    monkeypatch.setattr(validation.processing_service, "clone_processing_service", lambda **kwargs: _ValidationServiceStub(kwargs["log_event_fn"]))
    monkeypatch.setattr(
        validation.application_flow,
        "prepare_run_context_for_background",
        lambda **kwargs: (_ for _ in ()).throw(AssertionError("main should use service facade instead of direct prepare")),
    )
    monkeypatch.setattr(
        validation.document_pipeline,
        "run_document_processing",
        lambda **kwargs: (_ for _ in ()).throw(AssertionError("main should use service facade instead of direct pipeline orchestration")),
    )

    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_PROFILE", "lietaer-core")
    monkeypatch.delenv("DOCXAI_REAL_DOCUMENT_RUN_PROFILE", raising=False)
    monkeypatch.delenv("DOCXAI_REAL_DOCUMENT_REPEAT_COUNT_OVERRIDE", raising=False)
    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_FORCED_RUN_ID", "run-123")

    validation.main()

    report_path = tmp_path / "artifacts" / "runs" / "run-123" / "lietaer_validation_report.json"
    summary_path = tmp_path / "artifacts" / "runs" / "run-123" / "lietaer_validation_summary.txt"
    report = json.loads(report_path.read_text(encoding="utf-8"))
    summary_text = summary_path.read_text(encoding="utf-8")

    assert captured["uploaded_filename"] == "legacy.doc"
    assert captured["uploaded_bytes"] == source_bytes
    assert captured["processing_operation"] == "translate"
    assert captured["source_language"] == "en"
    assert captured["target_language"] == "ru"
    assert report["runtime_config"]["effective"]["image_mode"] == "safe"
    assert report["preparation"]["ai_classified_count"] == 7
    assert report["preparation"]["ai_heading_count"] == 3
    assert report["preparation"]["ai_role_change_count"] == 2
    assert report["preparation"]["ai_heading_promotion_count"] == 1
    assert report["preparation"]["ai_heading_demotion_count"] == 1
    assert report["preparation"]["ai_structural_role_change_count"] == 1
    snapshot = report["preparation_diagnostic_snapshot"]
    assert snapshot["paragraph_count"] == 5
    assert snapshot["heading_count"] == 1
    assert snapshot["toc_header_count"] == 1
    assert snapshot["toc_entry_count"] == 1
    assert snapshot["bounded_toc_region_count"] == 1
    assert snapshot["repaired_bullet_items"] == 4
    assert snapshot["repaired_numbered_items"] == 5
    assert snapshot["toc_body_boundary_repairs"] == 1
    assert snapshot["remaining_isolated_marker_count"] == 0
    assert snapshot["readiness_status"] == "blocked_unsafe_best_effort_only"
    assert snapshot["readiness_reasons"] == ["heading_count_far_below_toc_expectation"]
    assert snapshot["quality_gate_status"] == "blocked"
    assert snapshot["quality_gate_reasons"] == ["structure_readiness_blocked_unsafe_best_effort_only"]
    assert snapshot["structure_ai_attempted"] is True
    assert snapshot["ai_classified_count"] == 7
    assert snapshot["ai_heading_count"] == 3
    assert snapshot["semantic_block_count"] >= 1
    assert snapshot["first_block_target_chars"] == 3891
    assert report["image_forensics"]["prepared_assets"][0]["source"]["source_forensics"]["drawing_container"] == "inline"
    assert report["image_forensics"]["prepared_assets"][0]["source"]["source_sha256"]
    assert report["image_forensics"]["processed_assets"][0]["final_selection"]["final_variant"] == "redrawn"
    assert report["translation_quality_report_path"] == ".run/quality_reports/prepared_quality_report.json"
    assert report["translation_quality_report"]["quality_status"] == "fail"
    assert report["translation_quality_report"]["false_fragment_heading_count"] == 2
    assert 'preparation_diagnostic_snapshot={"ai_classified_count": 7' in summary_text
    assert '"readiness_status": "blocked_unsafe_best_effort_only"' in summary_text
    assert "translation_quality_status=fail" in summary_text
    assert "translation_quality_gate_reasons=false_fragment_headings_present,residual_bullet_glyphs_present" in summary_text
    assert "translation_quality_bullet_heading_gate_source=legacy_markdown" in summary_text
    assert "translation_quality_bullet_heading_classification=markdown_gate" in summary_text
    assert "translation_quality_raw_bullet_heading_count=0" in summary_text
    assert "translation_quality_page_placeholder_heading_concat_source=legacy_markdown" in summary_text
    assert "translation_quality_page_placeholder_heading_concat_classification=display_hygiene" in summary_text
    assert "translation_quality_raw_page_placeholder_heading_concat_count=1" in summary_text
    assert "translation_quality_false_fragment_heading_gate_source=entry_assembly" in summary_text
    assert "translation_quality_raw_false_fragment_heading_count=2" in summary_text
    assert "translation_quality_residual_bullet_glyph_gate_source=legacy_markdown" in summary_text
    assert "translation_quality_residual_bullet_glyph_classification=display_hygiene" in summary_text
    assert "translation_quality_raw_residual_bullet_glyph_count=1" in summary_text
    assert "translation_quality_list_fragment_regression_gate_source=topology_projection" in summary_text
    assert "translation_quality_raw_list_fragment_regression_count=1" in summary_text
    assert "translation_quality_mixed_script_term_gate_source=legacy_markdown" in summary_text
    assert "translation_quality_mixed_script_term_classification=non_structural_hygiene" in summary_text
    assert "translation_quality_raw_mixed_script_term_count=2" in summary_text
    assert "translation_quality_theology_style_deterministic_issue_source=legacy_markdown" in summary_text
    assert "translation_quality_theology_style_deterministic_issue_classification=domain_style_advisory" in summary_text
    assert "translation_quality_raw_theology_style_deterministic_issue_count=0" in summary_text
    assert "translation_quality_toc_body_concat_gate_source=legacy_markdown" in summary_text
    assert "source_file" not in report
    assert "runtime_configuration" not in report


def test_main_comparison_only_reader_cleanup_reports_non_acceptance_artifacts_for_chapter_region(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()
    from docxaicorrector.core.models import ParagraphUnit

    source_path = tmp_path / "Rethinking-money-chapter-region-pages-10-11-and-156-217.pdf"
    source_path.write_bytes(b"%PDF-1.4 chapter-region")
    ui_results_dir = tmp_path / ".run" / "ui_results"
    ui_results_dir.mkdir(parents=True, exist_ok=True)
    cleaned_markdown_path = ui_results_dir / "chapter_region.result.md"
    cleaned_docx_path = ui_results_dir / "chapter_region.result.docx"
    raw_markdown_path = ui_results_dir / "chapter_region.raw.result.md"
    cleanup_report_path = ui_results_dir / "chapter_region.reader_cleanup_report.json"

    cleaned_markdown_path.write_text(
        "12 CHAPTER HEADER\n\n• Currency menu entry\n\nPhoto: market square\n\nlowercase carryover after caption\n\n"
        "Repeated body fragment that is intentionally long enough to trigger duplicate detection.\n\n"
        "Repeated body fragment that is intentionally long enough to trigger duplicate detection.",
        encoding="utf-8",
    )
    cleaned_docx_path.write_bytes(_docx_bytes(Document()))
    raw_markdown_path.write_text("Raw markdown", encoding="utf-8")
    cleanup_report_path.write_text(
        json.dumps(
            {
                "stage_status": "completed",
                "changed": True,
                "stats": {
                    "proposed_delete_block_count": 2,
                    "accepted_delete_block_count": 1,
                    "ignored_delete_block_count": 1,
                    "failed_chunk_count": 0,
                },
                "accepted_delete_blocks": [
                    {
                        "id": "b_000007",
                        "reason": "repeated_running_header",
                        "confidence": "high",
                        "raw_text_preview": "rethinking money",
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    def _resolution_payload(**values):
        return SimpleNamespace(**values, to_dict=lambda: dict(values))

    document_profile = SimpleNamespace(
        id="lietaer-pdf-chapter-region-core",
        artifact_prefix="lietaer_pdf_chapter_region",
        output_basename="Rethinking_money_chapter_region",
        max_unmapped_source_paragraphs=12,
        max_unmapped_target_paragraphs=6,
        require_no_toc_body_concat=True,
        resolved_source_path=lambda project_root=None: source_path,
    )
    run_profile = SimpleNamespace(
        id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        tier="full",
        repeat_count=1,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=False,
        model="gpt-5.4",
        max_retries=1,
        comparison_only_validation=True,
        reader_verifier_enabled=True,
    )
    registry = SimpleNamespace(
        get_document_profile=lambda profile_id: document_profile,
        resolve_run_profile=lambda profile, requested_run_profile_id: run_profile,
    )

    prepared_paragraphs = [
        ParagraphUnit(text="Chapter 8", role="heading", structural_role="heading", heading_level=1, source_index=0)
    ]

    class _ValidationServiceStub:
        def __init__(self, log_event_fn):
            self.log_event_fn = log_event_fn

        def run_prepared_background_document(self, **kwargs):
            self.log_event_fn(
                20,
                "ui_result_artifacts_saved",
                "Сохранены итоговые UI-артефакты обработки.",
                filename="chapter-region.pdf",
                artifact_paths={
                    "markdown_path": str(cleaned_markdown_path),
                    "docx_path": str(cleaned_docx_path),
                    "reader_cleanup_raw_markdown_path": str(raw_markdown_path),
                    "reader_cleanup_report_path": str(cleanup_report_path),
                },
            )
            kwargs["runtime"].emit(
                validation.SetStateEvent(
                    values={
                        "latest_markdown": "Cleaned markdown",
                        "latest_docx_bytes": _docx_bytes(Document()),
                    }
                )
            )
            return "succeeded", SimpleNamespace(
                uploaded_filename="chapter-region.pdf",
                uploaded_file_bytes=b"%PDF-1.4 chapter-region",
                paragraphs=prepared_paragraphs,
                image_assets=[],
                jobs=[{"job_kind": "block"}],
                source_text="Source chapter region text",
                preparation_cached=False,
                preparation_elapsed_seconds=0.1,
                ai_classified_count=0,
                ai_heading_count=0,
                ai_role_change_count=0,
                ai_heading_promotion_count=0,
                ai_heading_demotion_count=0,
                ai_structural_role_change_count=0,
                structure_repair_report=None,
                uploaded_file_token="token-1",
            )

    monkeypatch.setattr(validation, "PROJECT_ROOT", tmp_path)
    monkeypatch.setattr(validation, "REAL_DOCUMENT_ARTIFACT_ROOT", tmp_path / "artifacts")
    monkeypatch.setattr(validation, "load_validation_registry", lambda: registry)
    monkeypatch.setattr(validation, "load_app_config", lambda: object())
    monkeypatch.setattr(
        validation,
        "resolve_runtime_resolution",
        lambda app_config, run_profile: SimpleNamespace(
            effective=_resolution_payload(
                chunk_size=6000,
                image_mode="safe",
                keep_all_image_variants=False,
                model="gpt-5.4",
                max_retries=1,
                processing_operation="translate",
                source_language="auto",
                target_language="ru",
            ),
            ui_defaults=_resolution_payload(
                chunk_size=6000,
                image_mode="safe",
                keep_all_image_variants=False,
                model="gpt-5.4",
                max_retries=1,
                processing_operation="translate",
                source_language="auto",
                target_language="ru",
            ),
            overrides={
                "reader_cleanup_enabled": True,
                "reader_cleanup_policy": "advisory",
                "structure_recognition_mode": "off",
            },
        ),
    )
    monkeypatch.setattr(validation, "apply_runtime_resolution_to_app_config", lambda app_config, resolution: {"x": 1})
    monkeypatch.setattr(
        validation,
        "evaluate_lietaer_acceptance",
        lambda report, **kwargs: {"passed": False, "failed_checks": ["false_fragment_headings_present"], "checks": []},
    )
    monkeypatch.setattr(validation, "_snapshot_formatting_diagnostics_paths", lambda: set())
    monkeypatch.setattr(validation, "_collect_new_formatting_diagnostics_paths", lambda before, after: [])
    monkeypatch.setattr(validation, "_extract_run_formatting_diagnostics_paths", lambda event_log: [])
    monkeypatch.setattr(validation, "_load_recent_formatting_diagnostics", lambda started_at: ([], []))
    monkeypatch.setattr(validation, "_load_formatting_diagnostics_payloads", lambda paths: [])
    monkeypatch.setattr(validation, "_load_translation_quality_report", lambda event_log: (None, None))
    monkeypatch.setattr(validation, "_print_terminal_completion_summary", lambda **kwargs: None)
    monkeypatch.setattr(
        validation.processing_service,
        "clone_processing_service",
        lambda **kwargs: _ValidationServiceStub(kwargs["log_event_fn"]),
    )

    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_PROFILE", "lietaer-pdf-chapter-region-core")
    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_RUN_PROFILE", "ui-parity-translate-simple-reader-cleanup-comparison-only")
    monkeypatch.delenv("DOCXAI_REAL_DOCUMENT_REPEAT_COUNT_OVERRIDE", raising=False)
    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_FORCED_RUN_ID", "run-compare")

    validation.main()

    report_path = tmp_path / "artifacts" / "runs" / "run-compare" / "lietaer_pdf_chapter_region_report.json"
    summary_path = tmp_path / "artifacts" / "runs" / "run-compare" / "lietaer_pdf_chapter_region_summary.txt"
    latest_manifest_path = tmp_path / "artifacts" / "lietaer_pdf_chapter_region_latest.json"

    report = json.loads(report_path.read_text(encoding="utf-8"))
    summary_text = summary_path.read_text(encoding="utf-8")
    latest_manifest = json.loads(latest_manifest_path.read_text(encoding="utf-8"))

    assert report["document_profile_id"] == "lietaer-pdf-chapter-region-core"
    assert report["run_profile_id"] == "ui-parity-translate-simple-reader-cleanup-comparison-only"
    assert report["validation_mode"]["evidence_label"] == "comparison_only_non_acceptance"
    assert report["validation_mode"]["success_criterion"] == "pipeline_result_and_artifacts"
    assert report["acceptance"]["passed"] is False
    assert report["reader_cleanup_evidence"] == {
        "artifacts_present": True,
        "cleaned_markdown_path": str(cleaned_markdown_path),
        "cleaned_docx_path": str(cleaned_docx_path),
        "raw_markdown_path": str(raw_markdown_path),
        "reader_cleanup_report_path": str(cleanup_report_path),
        "stage_status": "completed",
        "changed": True,
        "accepted_delete_block_count": 1,
        "ignored_delete_block_count": 1,
        "rejected_delete_block_count": 0,
            "failed_chunk_count": 0,
            "cleanup_chunk_count": None,
            "cleanup_settings": {},
            "anchor_repair_status": "diagnostic_only_not_applied",
            "recommended_anchor_targets": report["reader_cleanup_evidence"]["recommended_anchor_targets"],
        "recommended_anchor_target_count": report["reader_cleanup_evidence"]["recommended_anchor_target_count"],
        "verifier_recommended_anchor_targets": report["reader_cleanup_evidence"]["verifier_recommended_anchor_targets"],
        "verifier_recommended_anchor_target_count": report["reader_cleanup_evidence"][
            "verifier_recommended_anchor_target_count"
        ],
        "deleted_block_previews": [
            {
                "id": "b_000007",
                "reason": "repeated_running_header",
                "confidence": "high",
                "raw_text_preview": "rethinking money",
            }
        ],
    }
    assert report["reader_cleanup_evidence"]["recommended_anchor_target_count"] == len(
        report["reader_cleanup_evidence"]["recommended_anchor_targets"]
    )
    assert report["reader_cleanup_evidence"]["recommended_anchor_target_count"] >= 1
    assert report["reader_cleanup_evidence"]["verifier_recommended_anchor_target_count"] == len(
        report["reader_cleanup_evidence"]["verifier_recommended_anchor_targets"]
    )
    assert report["reader_cleanup_evidence"]["verifier_recommended_anchor_target_count"] >= 1
    assert report["output_artifacts"]["cleaned_markdown_path"] == str(cleaned_markdown_path)
    assert report["output_artifacts"]["cleaned_docx_path"] == str(cleaned_docx_path)
    assert report["output_artifacts"]["reader_cleanup_raw_markdown_path"] == str(raw_markdown_path)
    assert report["output_artifacts"]["reader_cleanup_report_path"] == str(cleanup_report_path)

    assert latest_manifest["status"] == "completed"
    assert latest_manifest["acceptance_passed"] is False
    assert "validation_evidence_label=comparison_only_non_acceptance" in summary_text
    assert "comparison_only_acceptance_diagnostic_only=True" in summary_text
    assert f"cleaned_markdown_path={cleaned_markdown_path}" in summary_text
    assert f"cleaned_docx_path={cleaned_docx_path}" in summary_text
    assert f"reader_cleanup_raw_markdown_path={raw_markdown_path}" in summary_text
    assert f"reader_cleanup_report_path={cleanup_report_path}" in summary_text
    assert "reader_cleanup_accepted_delete_block_count=1" in summary_text
    assert "reader_cleanup_ignored_delete_block_count=1" in summary_text
    assert "reader_mvp_status_anchor_repair_status=diagnostic_only_not_applied" in summary_text
    assert 'reader_cleanup_deleted_block_previews=[{"id": "b_000007", "reason": "repeated_running_header", "confidence": "high", "raw_text_preview": "rethinking money"}]' in summary_text


def _reader_verifier_test_evidence_payload(*, pre_audit_issue_counts: dict[str, int] | None = None) -> dict[str, object]:
    counts = {
        "page_furniture_inline": 0,
        "heading_fused_with_body": 0,
        "broken_list_marker": 0,
        "fragmented_paragraph": 0,
        "duplicate_fragment": 0,
        "orphan_caption": 0,
        "mixed_language_leak": 0,
        "quote_not_block_formatted": 0,
    }
    if pre_audit_issue_counts:
        counts.update(pre_audit_issue_counts)
    findings = []
    if pre_audit_issue_counts:
        for category, count in pre_audit_issue_counts.items():
            if count <= 0:
                continue
            for index in range(count):
                findings.append(
                    {
                        "category": category,
                        "artifact": "cleaned_markdown",
                        "line_ref": f"cleaned_markdown:{12 + index}",
                        "snippet": f"Synthetic {category} finding {index + 1} remains in cleaned output.",
                        "note": f"Deterministic pre-audit found {category} #{index + 1}.",
                    }
                )
    return {
        "artifact_paths": {
            "raw_markdown": "runs/run-1/raw.md",
            "cleaned_markdown": "runs/run-1/cleaned.md",
            "cleaned_docx": "runs/run-1/cleaned.docx",
            "reader_cleanup_report": "runs/run-1/cleanup.json",
        },
        "validator_boundary": {
            "observer_only": True,
            "runs_cleanup_repair": False,
            "runs_anchor_repair": False,
            "mutates_cleaned_markdown": False,
            "mutates_cleaned_docx": False,
            "rebuilds_docx": False,
        },
        "pre_audit_issue_counts": counts,
        "pre_audit_findings": findings,
        "mandatory_review_targets": findings,
        "toc_filtering_policy": {
            "mode": "evidence_only",
            "toc_out_of_review_scope": False,
            "policy_source": "toc_in_review_scope",
            "artifact_repair_applied": False,
        },
        "evidence_filtering_note": (
            "TOC-like review targets may be filtered from verifier evidence when profile policy marks TOC out of scope; "
            "this is evidence filtering only and does not repair, rewrite, or rebuild output artifacts."
        ),
        "filtered_toc_issue_count": 0,
    }


def _reader_verifier_test_response(**overrides: object) -> str:
    payload: dict[str, object] = {
        "overall_verdict": "cleaned_better",
        "cleaned_audit_verdict": "clean",
        "reader_quality_score_raw": 4,
        "reader_quality_score_cleaned": 7,
        "confidence": "medium",
        "noise_removed": ["Repeated running headers are gone."],
        "possible_false_deletions": [],
        "readability_regressions": [],
        "remaining_issues": [],
        "evidence_anchors": [
            {
                "kind": "improvement_seen",
                "artifact": "comparison",
                "line_ref": "comparison:1",
                "snippet": "Header noise was removed from the cleaned artifact.",
                "note": "Concrete readability improvement anchor.",
            }
        ],
        "recommended_next_changes": [
            {
                "change_type": "ai_operation_contract",
                "recommendation": "Ask the cleanup operation contract to remove standalone blank-page markers before verifier review.",
                "why": "These artifacts remain obvious reader-visible noise.",
            }
        ],
        "summary_for_human": "Cleaned output is easier to read and no major text loss was observed.",
        "simple_user_summary": "The cleaned version is easier to read than the raw translation.",
        "simple_user_risk_statement": "No major text loss was detected at current review confidence.",
        "simple_user_next_step": "Add one more deterministic cleanup rule for blank-page markers.",
    }
    payload.update(overrides)
    return json.dumps(payload, ensure_ascii=False)


def test_build_reader_verifier_evidence_payload_excludes_toc_like_pre_audit_findings(tmp_path) -> None:
    validation = _load_validation_module()

    raw_markdown_path = tmp_path / "raw.md"
    cleaned_markdown_path = tmp_path / "cleaned.md"
    cleanup_report_path = tmp_path / "cleanup.json"
    cleaned_docx_path = tmp_path / "cleaned.docx"
    source_document_path = tmp_path / "source.docx"

    cleaned_markdown_path.write_text(
        "СОДЕРЖАНИЕ Предисловие ix Введение: от дефицита к процветанию 1\n\n"
        "КАК ЭТО РАБОТАЕТ. Местные органы власти могут ежегодно запрашивать гражданский взнос.",
        encoding="utf-8",
    )
    raw_markdown_path.write_text(cleaned_markdown_path.read_text(encoding="utf-8"), encoding="utf-8")
    cleanup_report_path.write_text(json.dumps({"warnings": []}, ensure_ascii=False), encoding="utf-8")
    cleaned_docx_path.write_bytes(b"PK")
    source_document_path.write_bytes(b"PK")

    evidence = validation._build_reader_verifier_evidence_payload(
        run_id="run-toc-filter",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        source_document_path=source_document_path,
        source_text="source",
        reader_cleanup_evidence={
            "raw_markdown_path": str(raw_markdown_path),
            "cleaned_markdown_path": str(cleaned_markdown_path),
            "reader_cleanup_report_path": str(cleanup_report_path),
            "cleaned_docx_path": str(cleaned_docx_path),
            "stage_status": "completed",
            "changed": True,
        },
        runtime_app_config={"reader_cleanup_keep_toc": False},
    )

    assert evidence["toc_filtering_policy"]["toc_out_of_review_scope"] is True
    assert evidence["toc_filtering_policy"]["artifact_repair_applied"] is False
    assert evidence["filtered_toc_issue_count"] == 1
    assert evidence["filtered_toc_pre_audit_count"] == 1
    assert evidence["raw_pre_audit_issue_counts"]["heading_fused_with_body"] == 2
    assert evidence["pre_audit_issue_counts"]["heading_fused_with_body"] == 1
    assert evidence["validator_boundary"]["observer_only"] is True
    assert evidence["validator_boundary"]["rebuilds_docx"] is False
    assert evidence["filtered_toc_issue_previews"] == [
        {
            "source": "pre_audit",
            "category": "heading_fused_with_body",
            "artifact": "cleaned_markdown",
            "line_ref": "cleaned_markdown:1",
            "snippet": "СОДЕРЖАНИЕ Предисловие ix Введение: от дефицита к процветанию 1",
        }
    ]
    assert all("СОДЕРЖАНИЕ" not in finding["snippet"] for finding in evidence["pre_audit_findings"])
    assert all("СОДЕРЖАНИЕ" not in finding["snippet"] for finding in evidence["mandatory_review_targets"])
    assert len(evidence["mandatory_review_targets"]) == 1
    assert "КАК ЭТО РАБОТАЕТ" in evidence["mandatory_review_targets"][0]["snippet"]


def test_build_reader_verifier_evidence_payload_keeps_toc_like_pre_audit_findings_when_toc_in_scope(tmp_path) -> None:
    validation = _load_validation_module()

    raw_markdown_path = tmp_path / "raw.md"
    cleaned_markdown_path = tmp_path / "cleaned.md"
    cleanup_report_path = tmp_path / "cleanup.json"
    cleaned_docx_path = tmp_path / "cleaned.docx"
    source_document_path = tmp_path / "source.docx"

    cleaned_markdown_path.write_text(
        "СОДЕРЖАНИЕ Предисловие ix Введение: от дефицита к процветанию 1\n\n"
        "КАК ЭТО РАБОТАЕТ. Местные органы власти могут ежегодно запрашивать гражданский взнос.",
        encoding="utf-8",
    )
    raw_markdown_path.write_text(cleaned_markdown_path.read_text(encoding="utf-8"), encoding="utf-8")
    cleanup_report_path.write_text(json.dumps({"warnings": []}, ensure_ascii=False), encoding="utf-8")
    cleaned_docx_path.write_bytes(b"PK")
    source_document_path.write_bytes(b"PK")

    evidence = validation._build_reader_verifier_evidence_payload(
        run_id="run-toc-in-scope",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        source_document_path=source_document_path,
        source_text="source",
        reader_cleanup_evidence={
            "raw_markdown_path": str(raw_markdown_path),
            "cleaned_markdown_path": str(cleaned_markdown_path),
            "reader_cleanup_report_path": str(cleanup_report_path),
            "cleaned_docx_path": str(cleaned_docx_path),
            "stage_status": "completed",
            "changed": True,
        },
        runtime_app_config={"reader_cleanup_keep_toc": True},
    )

    assert evidence["toc_filtering_policy"]["toc_out_of_review_scope"] is False
    assert evidence["filtered_toc_issue_count"] == 0
    assert evidence["pre_audit_issue_counts"]["heading_fused_with_body"] == 2
    assert any("СОДЕРЖАНИЕ" in finding["snippet"] for finding in evidence["mandatory_review_targets"])


def test_parse_reader_verifier_completed_review_filters_toc_like_remaining_issues(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_payload = _reader_verifier_test_evidence_payload()
    evidence_payload["toc_filtering_policy"] = {
        "mode": "evidence_only",
        "toc_out_of_review_scope": True,
        "policy_source": "reader_cleanup_keep_toc_false",
        "artifact_repair_applied": False,
    }
    evidence_payload["cleaned_markdown"] = (
        "СОДЕРЖАНИЕ Предисловие ix Введение: от дефицита к процветанию 1\n\n"
        "КАК ЭТО РАБОТАЕТ. Местные органы власти могут ежегодно запрашивать гражданский взнос."
    )
    raw_response = _reader_verifier_test_response(
        cleaned_audit_verdict="improved_but_has_remaining_issues",
        remaining_issues=[
            {
                "category": "heading_fused_with_body",
                "severity": "high",
                "artifact": "cleaned_markdown",
                "line_ref": "cleaned_markdown:1",
                "snippet": "СОДЕРЖАНИЕ Предисловие ix Введение: от дефицита к процветанию 1",
                "why_reader_hurts": "TOC-like opening text is out of scope for the review.",
                "recommended_fix_type": "split_heading",
            },
            {
                "category": "heading_fused_with_body",
                "severity": "high",
                "artifact": "cleaned_markdown",
                "line_ref": "cleaned_markdown:3",
                "snippet": "КАК ЭТО РАБОТАЕТ. Местные органы власти могут ежегодно запрашивать гражданский взнос.",
                "why_reader_hurts": "The section heading is fused into body text.",
                "recommended_fix_type": "split_heading",
            },
        ],
        evidence_anchors=[
            {
                "kind": "improvement_seen",
                "artifact": "comparison",
                "line_ref": "comparison:1",
                "snippet": "Header noise was removed from the cleaned artifact.",
                "note": "Concrete readability improvement anchor.",
            },
            {
                "kind": "remaining_issue",
                "artifact": "cleaned_markdown",
                "line_ref": "cleaned_markdown:1",
                "snippet": "СОДЕРЖАНИЕ Предисловие ix Введение: от дефицита к процветанию 1",
                "note": "TOC-like issue.",
            },
            {
                "kind": "remaining_issue",
                "artifact": "cleaned_markdown",
                "line_ref": "cleaned_markdown:3",
                "snippet": "КАК ЭТО РАБОТАЕТ. Местные органы власти могут ежегодно запрашивать гражданский взнос.",
                "note": "Real remaining issue.",
            },
        ],
    )

    review = validation._parse_reader_verifier_completed_review(
        raw_response=raw_response,
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=tmp_path / "evidence.json",
        evidence_payload=evidence_payload,
        run_id="run-toc-filter",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
    )

    assert review["issue_summary_by_category"]["heading_fused_with_body"] == 1
    assert review["filtered_toc_issue_count"] == 1
    assert review["filtered_toc_pre_audit_count"] == 0
    assert review["filtered_toc_verifier_issue_count"] == 1
    assert review["filtered_toc_evidence_anchor_count"] == 1
    assert len(review["remaining_issues"]) == 1
    assert review["filtered_toc_issue_previews"] == [
        {
            "source": "verifier_remaining_issue",
            "category": "heading_fused_with_body",
            "artifact": "cleaned_markdown",
            "line_ref": "cleaned_markdown:1",
            "snippet": "СОДЕРЖАНИЕ Предисловие ix Введение: от дефицита к процветанию 1",
        }
    ]
    assert "СОДЕРЖАНИЕ" not in review["remaining_issues"][0]["snippet"]
    assert all("СОДЕРЖАНИЕ" not in anchor["snippet"] for anchor in review["evidence_anchors"])
    assert review["toc_filtering_policy"]["artifact_repair_applied"] is False
    assert "evidence filtering only" in review["evidence_filtering_note"]
    assert review["validator_boundary"]["observer_only"] is True


def test_build_reader_verifier_evidence_payload_emits_cleanup_ignored_reason_summary(tmp_path) -> None:
    validation = _load_validation_module()

    raw_markdown_path = tmp_path / "raw.md"
    cleaned_markdown_path = tmp_path / "cleaned.md"
    cleanup_report_path = tmp_path / "cleanup.json"
    cleaned_docx_path = tmp_path / "cleaned.docx"
    source_document_path = tmp_path / "source.docx"

    cleaned_markdown_path.write_text("Intro\n\nBody paragraph\n\nOutro", encoding="utf-8")
    raw_markdown_path.write_text(cleaned_markdown_path.read_text(encoding="utf-8"), encoding="utf-8")
    cleanup_report_path.write_text(
        json.dumps(
            {
                "warnings": [],
                "accepted_cleanup_operations": [
                    {
                        "operation": "remove_inline_noise",
                        "reason": "repeated_running_header",
                        "chunk_index": 1,
                    }
                ],
                "ignored_delete_blocks": [
                    {
                        "operation": "remove_inline_noise",
                        "reason": "repeated_running_header",
                        "ignored_reason": "prior_same_block_operation_not_applied",
                        "chunk_index": 1,
                        "raw_text_preview": "10",
                    },
                    {
                        "operation": "normalize_heading_boundary",
                        "reason": "heading_body_fusion",
                        "ignored_reason": "heading_boundary_unaccounted_text",
                        "chunk_index": 1,
                        "raw_text_preview": "HEADING Body text continues",
                    },
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    cleaned_docx_path.write_bytes(b"PK")
    source_document_path.write_bytes(b"PK")

    evidence = validation._build_reader_verifier_evidence_payload(
        run_id="run-cleanup-diagnostics",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        source_document_path=source_document_path,
        source_text="source",
        reader_cleanup_evidence={
            "raw_markdown_path": str(raw_markdown_path),
            "cleaned_markdown_path": str(cleaned_markdown_path),
            "reader_cleanup_report_path": str(cleanup_report_path),
            "cleaned_docx_path": str(cleaned_docx_path),
            "stage_status": "completed",
            "changed": True,
        },
        runtime_app_config={"reader_cleanup_keep_toc": False},
    )

    assert evidence["cleanup_diagnostics"]["accepted_operation_counts"] == {"remove_inline_noise": 1}
    assert evidence["cleanup_diagnostics"]["ignored_reason_counts"]["prior_same_block_operation_not_applied"] == 1
    assert evidence["cleanup_diagnostics"]["ignored_reason_counts"]["heading_boundary_unaccounted_text"] == 1
    assert evidence["cleanup_diagnostics"]["top_ignored_reasons"] == [
        {
            "ignored_reason": "prior_same_block_operation_not_applied",
            "count": 1,
            "examples": [
                {
                    "operation": "remove_inline_noise",
                    "reason": "repeated_running_header",
                    "chunk_index": 1,
                    "text_preview": "10",
                    "sequence_decision": "",
                }
            ],
        },
        {
            "ignored_reason": "heading_boundary_unaccounted_text",
            "count": 1,
            "examples": [
                {
                    "operation": "normalize_heading_boundary",
                    "reason": "heading_body_fusion",
                    "chunk_index": 1,
                    "text_preview": "HEADING Body text continues",
                    "sequence_decision": "",
                }
            ],
        },
    ]


def test_parse_reader_verifier_completed_review_accepts_structured_payload(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")
    evidence_payload = _reader_verifier_test_evidence_payload(pre_audit_issue_counts={"broken_list_marker": 1})
    raw_response = _reader_verifier_test_response(
        cleaned_audit_verdict="improved_but_has_remaining_issues",
        readability_regressions=["Some list spacing still looks uneven."],
        remaining_issues=[
            {
                "category": "broken_list_marker",
                "severity": "medium",
                "artifact": "cleaned_markdown",
                "line_ref": "cleaned_markdown:12",
                "snippet": "• Currency menu entry remains unnormalized.",
                "why_reader_hurts": "The residual bullet character breaks markdown list readability.",
                "recommended_fix_type": "normalize_list",
            }
        ],
        evidence_anchors=[
            {
                "kind": "improvement_seen",
                "artifact": "comparison",
                "line_ref": "comparison:1",
                "snippet": "Repeated running header removed in cleaned markdown.",
                "note": "Header noise no longer interrupts reading.",
            },
            {
                "kind": "remaining_issue",
                "artifact": "cleaned_markdown",
                "line_ref": "cleaned_markdown:12",
                "snippet": "• Currency menu entry remains unnormalized.",
                "note": "The list marker still looks raw in the cleaned artifact.",
            },
        ],
    )

    review = validation._parse_reader_verifier_completed_review(
        raw_response=raw_response,
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=evidence_payload,
    )

    assert review["review_mode"] == "development_only_non_acceptance"
    assert review["verifier_status"] == "completed"
    assert review["verifier_requested_selector"] == "openrouter:google/gemini-3-flash-preview"
    assert review["verifier_canonical_selector"] == "openrouter:google/gemini-3-flash-preview"
    assert review["verifier_provider"] == "openrouter"
    assert review["verifier_model_id"] == "google/gemini-3-flash-preview"
    assert review["overall_verdict"] == "cleaned_better"
    assert review["cleaned_audit_verdict"] == "improved_but_has_remaining_issues"
    assert review["pre_audit_issue_counts"]["broken_list_marker"] == 1
    assert review["issue_summary_by_category"]["broken_list_marker"] == 2
    assert review["remaining_issues"][0]["category"] == "broken_list_marker"
    assert any(
        issue["snippet"] == "Synthetic broken_list_marker finding 1 remains in cleaned output."
        for issue in review["remaining_issues"]
    )
    assert "current review confidence" in review["simple_user_summary"]
    assert "still remain in the cleaned output" in review["simple_user_summary"]
    assert "Review confidence is medium" in review["simple_user_risk_statement"]


def test_parse_reader_verifier_completed_review_dedupes_matching_pre_audit_issue(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")
    evidence_payload = _reader_verifier_test_evidence_payload()
    evidence_payload["pre_audit_issue_counts"] = {
        **evidence_payload["pre_audit_issue_counts"],
        "heading_fused_with_body": 1,
    }
    evidence_payload["pre_audit_findings"] = [
        {
            "category": "heading_fused_with_body",
            "artifact": "cleaned_markdown",
            "line_ref": "cleaned_markdown:25",
            "snippet": "СТРАТЕГИИ ДЛЯ ГОСУДАРСТВ Деньги — это рычаг власти.",
            "note": "Detected heading-like uppercase text fused into running body prose.",
        }
    ]
    evidence_payload["mandatory_review_targets"] = list(evidence_payload["pre_audit_findings"])
    raw_response = _reader_verifier_test_response(
        cleaned_audit_verdict="improved_but_has_remaining_issues",
        remaining_issues=[
            {
                "category": "heading_fused_with_body",
                "severity": "medium",
                "artifact": "cleaned_markdown",
                "line_ref": "25",
                "snippet": "СТРАТЕГИИ ДЛЯ ГОСУДАРСТВ Деньги — это рычаг власти.",
                "why_reader_hurts": "The chapter title is fused with the epigraph.",
                "recommended_fix_type": "split_heading",
            }
        ],
        evidence_anchors=[
            {
                "kind": "improvement_seen",
                "artifact": "comparison",
                "line_ref": "comparison:1",
                "snippet": "Repeated running headers were removed.",
                "note": "Concrete readability improvement anchor.",
            },
            {
                "kind": "remaining_issue",
                "artifact": "cleaned_markdown",
                "line_ref": "25",
                "snippet": "СТРАТЕГИИ ДЛЯ ГОСУДАРСТВ Деньги — это рычаг власти.",
                "note": "The heading remains fused with body text.",
            },
        ],
    )

    review = validation._parse_reader_verifier_completed_review(
        raw_response=raw_response,
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=evidence_payload,
    )

    assert len(review["remaining_issues"]) == 1
    assert review["issue_summary_by_category"]["heading_fused_with_body"] == 1
    assert review["pre_audit_issue_counts"]["heading_fused_with_body"] == 1


def test_parse_reader_verifier_completed_review_normalizes_object_findings_to_strings(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")
    evidence_payload = _reader_verifier_test_evidence_payload()
    raw_response = _reader_verifier_test_response(
        overall_verdict="mixed",
        cleaned_audit_verdict="unsafe_or_regressed",
        reader_quality_score_raw=5,
        reader_quality_score_cleaned=6,
        confidence="low",
        noise_removed=[{"text": "Repeated header removed."}],
        possible_false_deletions=[{"issue": "One short paragraph may be too aggressive."}],
        readability_regressions=[{"description": "Some spacing drift remains."}],
        recommended_next_changes=[
            {
                "change_type": "prompt",
                "recommendation": "Ask for flatter finding lists.",
                "why": "The reviewer currently drifts toward object-shaped findings.",
            }
        ],
        summary_for_human="The cleanup helps, but the evidence is still mixed.",
        simple_user_summary="The cleanup pass improved some noisy sections, but it also introduced enough risk or regression that the result is not clearly better yet.",
        simple_user_risk_statement="One short paragraph may be too aggressive.",
        simple_user_next_step="Tighten the prompt before the next comparison-only run.",
        evidence_anchors=[
            {
                "kind": "improvement_seen",
                "artifact": "comparison",
                "line_ref": "comparison:3",
                "snippet": "Repeated header removed.",
                "note": "Reader-visible header noise was reduced.",
            },
            {
                "kind": "possible_false_deletion",
                "artifact": "comparison",
                "line_ref": "comparison:5",
                "snippet": "One short paragraph may be too aggressive.",
                "note": "Possible deletion still needs review.",
            },
        ],
    )

    review = validation._parse_reader_verifier_completed_review(
        raw_response=raw_response,
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=evidence_payload,
    )

    assert review["noise_removed"] == ["Repeated header removed."]
    assert review["possible_false_deletions"] == ["One short paragraph may be too aggressive."]
    assert review["readability_regressions"] == ["Some spacing drift remains."]


def test_parse_reader_verifier_completed_review_ignores_negated_safety_summaries(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")
    evidence_payload = _reader_verifier_test_evidence_payload()
    raw_response = _reader_verifier_test_response(
        overall_verdict="cleaned_better",
        cleaned_audit_verdict="clean",
        reader_quality_score_raw=5,
        reader_quality_score_cleaned=6,
        confidence="high",
        noise_removed=["Image placeholders were removed from reader-facing Markdown."],
        possible_false_deletions=[
            "No content deletions were accepted per cleanup report; no false deletions detected."
        ],
        readability_regressions=["No regressions introduced by cleanup; structural improvements are net positive."],
        remaining_issues=[],
        evidence_anchors=[
            {
                "kind": "improvement_seen",
                "artifact": "comparison",
                "line_ref": "comparison:3",
                "snippet": "Image placeholders were removed from reader-facing Markdown.",
                "note": "Reader-visible internal markers were removed.",
            }
        ],
    )

    review = validation._parse_reader_verifier_completed_review(
        raw_response=raw_response,
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=evidence_payload,
    )

    assert review["possible_false_deletions"] == []
    assert review["readability_regressions"] == []


def test_parse_reader_verifier_completed_review_hardens_overconfident_simple_language(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")
    evidence_payload = _reader_verifier_test_evidence_payload(pre_audit_issue_counts={"broken_list_marker": 1})
    raw_response = _reader_verifier_test_response(
        cleaned_audit_verdict="improved_but_has_remaining_issues",
        readability_regressions=["Some list spacing still looks uneven."],
        remaining_issues=[
            {
                "category": "broken_list_marker",
                "severity": "medium",
                "artifact": "cleaned_markdown",
                "line_ref": "cleaned_markdown:15",
                "snippet": "• Currency glossary item is still not normalized.",
                "why_reader_hurts": "The raw bullet character interrupts the cleaned reading flow.",
                "recommended_fix_type": "normalize_list",
            }
        ],
        evidence_anchors=[
            {
                "kind": "improvement_seen",
                "artifact": "comparison",
                "line_ref": "comparison:1",
                "snippet": "Repeated running headers are gone.",
                "note": "The cleaned version removes repeated header noise.",
            },
            {
                "kind": "remaining_issue",
                "artifact": "cleaned_markdown",
                "line_ref": "cleaned_markdown:15",
                "snippet": "• Currency glossary item is still not normalized.",
                "note": "A broken list marker remains in the cleaned output.",
            },
        ],
        recommended_next_changes=[
            {
                "change_type": "cleanup_core",
                "recommendation": "Apply one narrow list-formatting follow-up through the cleanup operation applicator.",
                "why": "The current slice still has a reader-visible formatting inconsistency.",
            }
        ],
        simple_user_summary="The cleaned text is much easier to read and no actual story content was lost.",
        simple_user_risk_statement="The risk is very low.",
        simple_user_next_step="This is ready for production.",
    )

    review = validation._parse_reader_verifier_completed_review(
        raw_response=raw_response,
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=evidence_payload,
    )

    assert "story content" not in review["simple_user_summary"].lower()
    assert "no major text loss was detected at current review confidence" in review["simple_user_summary"].lower()
    assert "current review confidence" in review["simple_user_summary"].lower()
    assert "still remain in the cleaned output" in review["simple_user_summary"]
    assert "Review confidence is medium" in review["simple_user_risk_statement"]
    assert review["simple_user_next_step"] == "Use this as development-only comparison evidence and apply one narrow follow-up change before rerunning the same profile."


def test_parse_reader_verifier_completed_review_hardens_high_confidence_final_wording(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")
    evidence_payload = _reader_verifier_test_evidence_payload()
    raw_response = _reader_verifier_test_response(
        cleaned_audit_verdict="clean",
        reader_quality_score_raw=3,
        reader_quality_score_cleaned=5,
        confidence="high",
        noise_removed=["Removed distracting page numbers and headers."],
        recommended_next_changes=[
            {
                "change_type": "cleanup_core",
                "recommendation": "Standardize bullet point characters through cleanup-core safety checks.",
                "why": "The text still mixes bullet characters across sections.",
            }
        ],
        summary_for_human="The cleaned version significantly improves readability for this slice.",
        simple_user_summary="The cleaned version is easier to read because it removes stray numbers and page artifacts that were scattered throughout the text. No actual information was lost during this process.",
        simple_user_risk_statement="The risk of missing information is very low; only isolated page numbers and redundant markers were removed. The core text and all translated concepts remain fully intact.",
        simple_user_next_step="No further cleanup is required for this block; the text is ready for reading or further translation review.",
    )

    review = validation._parse_reader_verifier_completed_review(
        raw_response=raw_response,
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=evidence_payload,
    )

    assert review["simple_user_summary"] == (
        "The cleaned version is easier to read because it removes stray numbers and page artifacts that were "
        "scattered throughout the text. No major text loss was detected at current review confidence."
    )
    assert review["simple_user_risk_statement"] == (
        "No major text loss was detected at current review confidence. "
        "This remains development-only comparison evidence, not an acceptance result."
    )
    assert review["recommended_next_changes"] == [
        {
            "change_type": "safety_application",
            "recommendation": "Use the AI cleanup operation contract rather than a deterministic rule for this follow-up: Standardize bullet point characters through cleanup-core safety checks.",
            "why": "Legacy verifier change_type 'cleanup_core' was normalized to 'safety_application'. The text still mixes bullet characters across sections.",
        }
    ]
    assert review["simple_user_next_step"] == (
        "Use this as development-only comparison evidence. Next, apply one narrow safety_application change: "
        "Use the AI cleanup operation contract rather than a deterministic rule for this follow-up: "
        "Standardize bullet point characters through cleanup-core safety checks."
    )


def test_parse_reader_verifier_completed_review_rejects_clean_audit_with_remaining_issues(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")

    with pytest.raises(RuntimeError, match="reader_verifier_remaining_issues_forbid_cleaned_audit_clean"):
        validation._parse_reader_verifier_completed_review(
            raw_response=_reader_verifier_test_response(
                cleaned_audit_verdict="clean",
                remaining_issues=[
                    {
                        "category": "page_furniture_inline",
                        "severity": "high",
                        "artifact": "cleaned_markdown",
                        "line_ref": "cleaned_markdown:3",
                        "snippet": "151 CHAPTER TITLE",
                        "why_reader_hurts": "The reader still sees page furniture inside the body flow.",
                        "recommended_fix_type": "delete_noise",
                    }
                ],
                evidence_anchors=[
                    {
                        "kind": "improvement_seen",
                        "artifact": "comparison",
                        "line_ref": "comparison:1",
                        "snippet": "Some running headers were removed.",
                        "note": "One noise class improved.",
                    },
                    {
                        "kind": "remaining_issue",
                        "artifact": "cleaned_markdown",
                        "line_ref": "cleaned_markdown:3",
                        "snippet": "151 CHAPTER TITLE",
                        "note": "Inline page furniture still remains.",
                    },
                ],
            ),
            run_id="run-1",
            document_profile_id="lietaer-pdf-chapter-region-core",
            run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
            requested_selector="openrouter:google/gemini-3-flash-preview",
            canonical_selector="openrouter:google/gemini-3-flash-preview",
            provider="openrouter",
            model_id="google/gemini-3-flash-preview",
            evidence_path=evidence_path,
            evidence_payload=_reader_verifier_test_evidence_payload(pre_audit_issue_counts={"page_furniture_inline": 1}),
        )


def test_parse_reader_verifier_completed_review_allows_raw_better_without_improvement_anchor(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")

    review = validation._parse_reader_verifier_completed_review(
        raw_response=_reader_verifier_test_response(
            overall_verdict="raw_better",
            cleaned_audit_verdict="unsafe_or_regressed",
            noise_removed=[],
            possible_false_deletions=["Cleanup appears to have removed meaningful text."],
            readability_regressions=["A heading was fused into body prose in the cleaned output."],
            remaining_issues=[
                {
                    "category": "heading_fused_with_body",
                    "severity": "high",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:14",
                    "snippet": "200 ПЕРЕОСМЫСЛИВАЯ ДЕНЬГИ Наконец-то пришло осознание необходимости...",
                    "why_reader_hurts": "The cleaned artifact still contains a heading glued into running prose.",
                    "recommended_fix_type": "split_heading",
                }
            ],
            evidence_anchors=[
                {
                    "kind": "remaining_issue",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:14",
                    "snippet": "200 ПЕРЕОСМЫСЛИВАЯ ДЕНЬГИ Наконец-то пришло осознание необходимости...",
                    "note": "The cleaned artifact still contains a fused heading/body defect.",
                },
                {
                    "kind": "possible_false_deletion",
                    "artifact": "comparison",
                    "line_ref": "comparison:9",
                    "snippet": "One body sentence is missing in cleaned output.",
                    "note": "The raw output appears safer than the cleaned output for this region.",
                },
            ],
            summary_for_human="The cleaned output is not safer than the raw output for this slice.",
            simple_user_summary="The cleanup pass removed or damaged meaningful text more than it improved readability. The raw version is safer to keep until cleanup rules are fixed.",
            simple_user_risk_statement="The cleaned output still contains reader-visible defects and may have removed meaningful text.",
            simple_user_next_step="Tighten the cleanup rules for heading/body splits before rerunning the same comparison-only profile.",
        ),
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=_reader_verifier_test_evidence_payload(pre_audit_issue_counts={"heading_fused_with_body": 1}),
    )

    assert review["overall_verdict"] == "raw_better"
    assert review["cleaned_audit_verdict"] == "unsafe_or_regressed"
    assert review["remaining_issues"][0]["category"] == "heading_fused_with_body"


def test_parse_reader_verifier_completed_review_ignores_malformed_improvement_anchor(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")

    review = validation._parse_reader_verifier_completed_review(
        raw_response=_reader_verifier_test_response(
            overall_verdict="cleaned_better",
            cleaned_audit_verdict="clean",
            noise_removed=["Repeated running headers are gone."],
            remaining_issues=[],
            evidence_anchors=[
                {
                    "kind": "improvement_seen",
                    "artifact": "comparison",
                    "line_ref": "",
                    "snippet": "",
                    "note": "",
                }
            ],
        ),
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=_reader_verifier_test_evidence_payload(),
    )

    diagnostics = review["evidence_anchor_diagnostics"]
    assert review["overall_verdict"] == "cleaned_better"
    assert diagnostics["ignored_anchor_count"] == 1
    assert diagnostics["ignored_kind_counts"]["improvement_seen"] == 1
    assert diagnostics["repaired_anchor_counts"]["improvement_seen"] == 1
    assert diagnostics["warnings"] == [
        "reader_verifier_evidence_anchor_ignored_missing_required_text:index=0:kind=improvement_seen:artifact=comparison"
    ]
    assert any(anchor["kind"] == "improvement_seen" for anchor in review["evidence_anchors"])


def test_parse_reader_verifier_completed_review_ignores_malformed_remaining_issue_anchor_and_restores_targets(
    tmp_path,
) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")

    review = validation._parse_reader_verifier_completed_review(
        raw_response=_reader_verifier_test_response(
            overall_verdict="cleaned_better",
            cleaned_audit_verdict="improved_but_has_remaining_issues",
            remaining_issues=[
                {
                    "category": "heading_fused_with_body",
                    "severity": "high",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:14",
                    "snippet": "200 ПЕРЕОСМЫСЛИВАЯ ДЕНЬГИ Наконец-то пришло осознание необходимости...",
                    "why_reader_hurts": "The cleaned artifact still contains a heading glued into running prose.",
                    "recommended_fix_type": "split_heading",
                }
            ],
            evidence_anchors=[
                {
                    "kind": "improvement_seen",
                    "artifact": "comparison",
                    "line_ref": "comparison:1",
                    "snippet": "Repeated running header removed.",
                    "note": "One readability improvement is confirmed.",
                },
                {
                    "kind": "remaining_issue",
                    "artifact": "cleaned_markdown",
                    "line_ref": "",
                    "snippet": "",
                    "note": "",
                },
            ],
        ),
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=_reader_verifier_test_evidence_payload(pre_audit_issue_counts={"fragmented_paragraph": 1}),
    )

    diagnostics = review["evidence_anchor_diagnostics"]
    assert review["overall_verdict"] == "cleaned_better"
    assert diagnostics["ignored_anchor_count"] == 1
    assert diagnostics["ignored_kind_counts"]["remaining_issue"] == 1
    assert any(issue["category"] == "heading_fused_with_body" for issue in review["remaining_issues"])
    assert any(issue["category"] == "fragmented_paragraph" for issue in review["remaining_issues"])
    assert any(anchor["kind"] == "remaining_issue" for anchor in review["evidence_anchors"])


def test_parse_reader_verifier_completed_review_ignores_unknown_remaining_issue_note_field(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")

    review = validation._parse_reader_verifier_completed_review(
        raw_response=_reader_verifier_test_response(
            overall_verdict="cleaned_better",
            cleaned_audit_verdict="improved_but_has_remaining_issues",
            remaining_issues=[
                {
                    "category": "heading_fused_with_body",
                    "severity": "high",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:93",
                    "snippet": "КАК ЭТО РАБОТАЕТ. Местные органы власти ежегодно запрашивают...",
                    "why_reader_hurts": "The cleaned artifact still contains a heading glued into running prose.",
                    "recommended_fix_type": "split_heading",
                    "note": "Useful verifier explanation that is outside the strict issue schema.",
                }
            ],
            evidence_anchors=[
                {
                    "kind": "improvement_seen",
                    "artifact": "comparison",
                    "line_ref": "comparison:1",
                    "snippet": "Repeated running header removed.",
                    "note": "One readability improvement is confirmed.",
                },
                {
                    "kind": "remaining_issue",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:93",
                    "snippet": "КАК ЭТО РАБОТАЕТ. Местные органы власти ежегодно запрашивают...",
                    "note": "The heading/body boundary still needs cleanup.",
                },
            ],
        ),
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=_reader_verifier_test_evidence_payload(),
    )

    diagnostics = review["remaining_issue_diagnostics"]
    assert review["overall_verdict"] == "cleaned_better"
    assert review["cleaned_audit_verdict"] == "improved_but_has_remaining_issues"
    assert review["remaining_issues"][0]["category"] == "heading_fused_with_body"
    assert "note" not in review["remaining_issues"][0]
    assert diagnostics["ignored_unknown_field_counts"] == {"note": 1}
    assert diagnostics["warnings"] == ["reader_verifier_remaining_issue_ignored_unknown_fields:index=0:fields=note"]


def test_parse_reader_verifier_completed_review_downgrades_contradictory_removed_claim(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")

    review = validation._parse_reader_verifier_completed_review(
        raw_response=_reader_verifier_test_response(
            cleaned_audit_verdict="improved_but_has_remaining_issues",
            noise_removed=["Broken list markers were removed from the cleaned output."],
            remaining_issues=[
                {
                    "category": "broken_list_marker",
                    "severity": "medium",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:8",
                    "snippet": "• Residual list marker remains.",
                    "why_reader_hurts": "The bullet still looks raw to the reader.",
                    "recommended_fix_type": "normalize_list",
                }
            ],
            evidence_anchors=[
                {
                    "kind": "improvement_seen",
                    "artifact": "comparison",
                    "line_ref": "comparison:1",
                    "snippet": "Some list markers were normalized.",
                    "note": "The cleanup improved list readability in several places.",
                },
                {
                    "kind": "remaining_issue",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:8",
                    "snippet": "• Residual list marker remains.",
                    "note": "The same defect class still remains in the cleaned artifact.",
                },
            ],
        ),
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=_reader_verifier_test_evidence_payload(pre_audit_issue_counts={"broken_list_marker": 1}),
    )

    assert "broken_list_marker still has remaining review targets" in review["noise_removed"][0]


def test_parse_reader_verifier_completed_review_surfaces_empty_llm_issues_from_pre_audit(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")

    review = validation._parse_reader_verifier_completed_review(
        raw_response=_reader_verifier_test_response(
            overall_verdict="unclear",
            cleaned_audit_verdict="unclear",
            noise_removed=[],
            remaining_issues=[],
            evidence_anchors=[],
            summary_for_human="The verifier could not confidently classify the slice.",
            simple_user_summary="The evidence is unclear.",
            simple_user_risk_statement="The risk is unclear.",
            simple_user_next_step="Gather clearer comparison evidence before changing the profile.",
        ),
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=_reader_verifier_test_evidence_payload(pre_audit_issue_counts={"page_furniture_inline": 1}),
    )

    assert review["overall_verdict"] == "unclear"
    assert review["remaining_issues"][0]["category"] == "page_furniture_inline"
    assert review["issue_summary_by_category"]["page_furniture_inline"] == 1
    assert review["evidence_anchors"][0]["kind"] == "remaining_issue"


def test_parse_reader_verifier_completed_review_restores_missing_pre_audit_targets_when_llm_returns_partial_list(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")

    review = validation._parse_reader_verifier_completed_review(
        raw_response=_reader_verifier_test_response(
            cleaned_audit_verdict="improved_but_has_remaining_issues",
            remaining_issues=[
                {
                    "category": "heading_fused_with_body",
                    "severity": "high",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:14",
                    "snippet": "200 ПЕРЕОСМЫСЛИВАЯ ДЕНЬГИ Наконец-то пришло осознание необходимости...",
                    "why_reader_hurts": "The cleaned artifact still contains a fused heading/body defect.",
                    "recommended_fix_type": "split_heading",
                }
            ],
            evidence_anchors=[
                {
                    "kind": "improvement_seen",
                    "artifact": "comparison",
                    "line_ref": "comparison:1",
                    "snippet": "Repeated running header removed.",
                    "note": "One readability improvement is confirmed.",
                },
                {
                    "kind": "remaining_issue",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:14",
                    "snippet": "200 ПЕРЕОСМЫСЛИВАЯ ДЕНЬГИ Наконец-то пришло осознание необходимости...",
                    "note": "One fused heading/body issue is still present.",
                },
            ],
        ),
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=_reader_verifier_test_evidence_payload(
            pre_audit_issue_counts={"heading_fused_with_body": 2, "fragmented_paragraph": 1}
        ),
    )

    assert review["issue_summary_by_category"]["heading_fused_with_body"] >= 2
    assert review["issue_summary_by_category"]["fragmented_paragraph"] == 1
    assert any(issue["category"] == "fragmented_paragraph" for issue in review["remaining_issues"])
    assert sum(1 for anchor in review["evidence_anchors"] if anchor["kind"] == "remaining_issue") >= 3


def test_parse_reader_verifier_completed_review_restores_missing_pre_audit_targets_with_equal_category_count(tmp_path) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")

    review = validation._parse_reader_verifier_completed_review(
        raw_response=_reader_verifier_test_response(
            cleaned_audit_verdict="improved_but_has_remaining_issues",
            remaining_issues=[
                {
                    "category": "heading_fused_with_body",
                    "severity": "high",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:12",
                    "snippet": "Synthetic heading_fused_with_body finding 1 remains in cleaned output.",
                    "why_reader_hurts": "The first fused heading/body defect remains.",
                    "recommended_fix_type": "split_heading",
                },
                {
                    "category": "heading_fused_with_body",
                    "severity": "high",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:99",
                    "snippet": "LLM-only heading_fused_with_body finding outside mandatory targets.",
                    "why_reader_hurts": "The verifier found another fused heading/body defect.",
                    "recommended_fix_type": "split_heading",
                },
            ],
            evidence_anchors=[
                {
                    "kind": "improvement_seen",
                    "artifact": "comparison",
                    "line_ref": "comparison:1",
                    "snippet": "Repeated running header removed.",
                    "note": "One readability improvement is confirmed.",
                },
                {
                    "kind": "remaining_issue",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:12",
                    "snippet": "Synthetic heading_fused_with_body finding 1 remains in cleaned output.",
                    "note": "The first mandatory issue remains.",
                },
                {
                    "kind": "remaining_issue",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:99",
                    "snippet": "LLM-only heading_fused_with_body finding outside mandatory targets.",
                    "note": "The verifier found another same-category issue.",
                },
            ],
        ),
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        requested_selector="openrouter:google/gemini-3-flash-preview",
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
        evidence_path=evidence_path,
        evidence_payload=_reader_verifier_test_evidence_payload(pre_audit_issue_counts={"heading_fused_with_body": 2}),
    )

    assert review["issue_summary_by_category"]["heading_fused_with_body"] == 3
    assert any(
        issue["line_ref"] == "cleaned_markdown:13"
        and issue["snippet"] == "Synthetic heading_fused_with_body finding 2 remains in cleaned output."
        for issue in review["remaining_issues"]
    )


def test_run_reader_verifier_marks_model_resolution_failure_without_fallback(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")
    evidence_payload = {
        "base_artifacts_present": True,
        "artifact_paths": {
            "raw_markdown": "runs/run-1/raw.md",
            "cleaned_markdown": "runs/run-1/cleaned.md",
            "cleaned_docx": "runs/run-1/cleaned.docx",
            "reader_cleanup_report": "runs/run-1/cleanup.json",
        },
    }

    monkeypatch.setattr(
        validation,
        "describe_provider_availability",
        lambda selector, app_config=None: SimpleNamespace(
            enabled=True,
            has_api_key=True,
            error_message=None,
            selector=SimpleNamespace(
                canonical_selector="openrouter:google/gemini-3-flash-preview",
                provider="openrouter",
                model_id="google/gemini-3-flash-preview",
            ),
        ),
    )
    monkeypatch.setattr(validation, "resolve_model_selector", lambda *args, **kwargs: (_ for _ in ()).throw(RuntimeError("resolution boom")))
    monkeypatch.setattr(validation, "get_client_for_model_selector", lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("must not fallback to another client")))

    review = validation._run_reader_verifier(
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        app_config=object(),
        runtime_app_config={"reader_verifier_enabled": True},
        validation_mode={"comparison_only_validation": True},
        evidence_payload=evidence_payload,
        evidence_path=evidence_path,
        max_retries=1,
    )

    assert review["verifier_status"] == "not_run"
    assert review["verifier_reason"] == "model_resolution_failed"
    assert review["overall_verdict"] == "unclear"
    assert review["verifier_canonical_selector"] is None


def test_run_reader_verifier_failure_surfaces_pre_audit_remaining_issues(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")
    evidence_payload = {
        **_reader_verifier_test_evidence_payload(pre_audit_issue_counts={"page_furniture_inline": 1}),
        "base_artifacts_present": True,
    }

    monkeypatch.setattr(
        validation,
        "describe_provider_availability",
        lambda selector, app_config=None: SimpleNamespace(
            enabled=True,
            has_api_key=True,
            error_message=None,
            selector=SimpleNamespace(
                canonical_selector="openrouter:google/gemini-3-flash-preview",
                provider="openrouter",
                model_id="google/gemini-3-flash-preview",
            ),
        ),
    )
    monkeypatch.setattr(
        validation,
        "resolve_model_selector",
        lambda *args, **kwargs: SimpleNamespace(
            canonical_selector="openrouter:google/gemini-3-flash-preview",
            provider="openrouter",
            model_id="google/gemini-3-flash-preview",
        ),
    )
    monkeypatch.setattr(validation, "get_client_for_model_selector", lambda *args, **kwargs: object())
    monkeypatch.setattr(validation, "generate_markdown_block", lambda *args, **kwargs: (_ for _ in ()).throw(RuntimeError("llm boom")))

    review = validation._run_reader_verifier(
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        app_config=object(),
        runtime_app_config={"reader_verifier_enabled": True},
        validation_mode={"comparison_only_validation": True},
        evidence_payload=evidence_payload,
        evidence_path=evidence_path,
        max_retries=1,
    )

    assert review["verifier_status"] == "failed"
    assert review["verifier_reason"] == "execution_failed"
    assert review["remaining_issues"][0]["category"] == "page_furniture_inline"
    assert review["issue_summary_by_category"]["page_furniture_inline"] == 1
    assert review["verifier_canonical_selector"] == "openrouter:google/gemini-3-flash-preview"


def test_run_reader_verifier_timeout_surfaces_pre_audit_remaining_issues(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()

    evidence_path = tmp_path / "reader_quality_evidence.json"
    evidence_path.write_text("{}", encoding="utf-8")
    evidence_payload = {
        **_reader_verifier_test_evidence_payload(pre_audit_issue_counts={"heading_fused_with_body": 1}),
        "base_artifacts_present": True,
    }

    monkeypatch.setattr(
        validation,
        "describe_provider_availability",
        lambda selector, app_config=None: SimpleNamespace(
            enabled=True,
            has_api_key=True,
            error_message=None,
            selector=SimpleNamespace(
                canonical_selector="openrouter:google/gemini-3-flash-preview",
                provider="openrouter",
                model_id="google/gemini-3-flash-preview",
            ),
        ),
    )
    monkeypatch.setattr(
        validation,
        "resolve_model_selector",
        lambda *args, **kwargs: SimpleNamespace(
            canonical_selector="openrouter:google/gemini-3-flash-preview",
            provider="openrouter",
            model_id="google/gemini-3-flash-preview",
        ),
    )
    monkeypatch.setattr(validation, "get_client_for_model_selector", lambda *args, **kwargs: object())

    def _slow_generate_markdown_block(**kwargs):
        time.sleep(0.2)
        return "{}"

    monkeypatch.setattr(validation, "generate_markdown_block", _slow_generate_markdown_block)
    monkeypatch.setenv("DOCXAI_READER_VERIFIER_TIMEOUT_SECONDS", "0.01")

    review = validation._run_reader_verifier(
        run_id="run-1",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        app_config=object(),
        runtime_app_config={"reader_verifier_enabled": True},
        validation_mode={"comparison_only_validation": True},
        evidence_payload=evidence_payload,
        evidence_path=evidence_path,
        max_retries=1,
    )

    assert review["verifier_status"] == "failed"
    assert review["verifier_reason"] == "execution_timeout"
    assert review["remaining_issues"][0]["category"] == "heading_fused_with_body"
    assert review["issue_summary_by_category"]["heading_fused_with_body"] == 1
    assert review["verifier_canonical_selector"] == "openrouter:google/gemini-3-flash-preview"


def test_main_comparison_only_reader_verifier_writes_artifacts_and_metadata(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()
    from docxaicorrector.core.models import ParagraphUnit

    source_path = tmp_path / "Rethinking-money-chapter-region-pages-10-11-and-156-217.pdf"
    source_path.write_bytes(b"%PDF-1.4 chapter-region")
    ui_results_dir = tmp_path / ".run" / "ui_results"
    ui_results_dir.mkdir(parents=True, exist_ok=True)
    cleaned_markdown_path = ui_results_dir / "chapter_region.result.md"
    cleaned_docx_path = ui_results_dir / "chapter_region.result.docx"
    raw_markdown_path = ui_results_dir / "chapter_region.raw.result.md"
    cleanup_report_path = ui_results_dir / "chapter_region.reader_cleanup_report.json"

    cleaned_markdown_path.write_text(
        "12 CHAPTER HEADER\n\n• Currency menu entry\n\nPhoto: market square\n\nlowercase carryover after caption\n\n"
        "Repeated body fragment that is intentionally long enough to trigger duplicate detection.\n\n"
        "Repeated body fragment that is intentionally long enough to trigger duplicate detection.",
        encoding="utf-8",
    )
    cleaned_docx_path.write_bytes(_docx_bytes(Document()))
    raw_markdown_path.write_text("Header\n\nBody paragraph\n\nFooter", encoding="utf-8")
    cleanup_report_path.write_text(
        json.dumps(
            {
                "stage_status": "completed",
                "changed": True,
                "stats": {
                    "proposed_delete_block_count": 1,
                    "accepted_delete_block_count": 1,
                    "ignored_delete_block_count": 0,
                    "failed_chunk_count": 0,
                    "cleanup_chunk_count": 1,
                },
                "accepted_delete_blocks": [
                    {
                        "id": "b_000000",
                        "text_hash": "abc",
                        "reason": "repeated_running_header",
                        "confidence": "high",
                        "raw_text_preview": "Header",
                        "char_count": 6,
                        "kind": "heading",
                    }
                ],
                "ignored_delete_blocks": [
                    {
                        "operation": "remove_inline_noise",
                        "reason": "repeated_running_header",
                        "ignored_reason": "prior_same_block_operation_not_applied",
                        "chunk_index": 1,
                        "raw_text_preview": "12 CHAPTER HEADER",
                    }
                ],
                "warnings": [],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    def _resolution_payload(**values):
        return SimpleNamespace(**values, to_dict=lambda: dict(values))

    document_profile = SimpleNamespace(
        id="lietaer-pdf-chapter-region-core",
        artifact_prefix="lietaer_pdf_chapter_region",
        output_basename="Rethinking_money_chapter_region",
        max_unmapped_source_paragraphs=12,
        max_unmapped_target_paragraphs=6,
        require_no_toc_body_concat=True,
        resolved_source_path=lambda project_root=None: source_path,
    )
    run_profile = SimpleNamespace(
        id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        tier="full",
        repeat_count=1,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=False,
        model="gpt-5.4",
        max_retries=1,
        comparison_only_validation=True,
    )
    registry = SimpleNamespace(
        get_document_profile=lambda profile_id: document_profile,
        resolve_run_profile=lambda profile, requested_run_profile_id: run_profile,
    )
    prepared_paragraphs = [
        ParagraphUnit(text="Chapter 8", role="heading", structural_role="heading", heading_level=1, source_index=0)
    ]

    class _ValidationServiceStub:
        def __init__(self, log_event_fn):
            self.log_event_fn = log_event_fn

        def run_prepared_background_document(self, **kwargs):
            self.log_event_fn(
                20,
                "ui_result_artifacts_saved",
                "Сохранены итоговые UI-артефакты обработки.",
                filename="chapter-region.pdf",
                artifact_paths={
                    "markdown_path": str(cleaned_markdown_path),
                    "docx_path": str(cleaned_docx_path),
                    "reader_cleanup_raw_markdown_path": str(raw_markdown_path),
                    "reader_cleanup_report_path": str(cleanup_report_path),
                },
            )
            kwargs["runtime"].emit(
                validation.SetStateEvent(
                    values={
                        "latest_markdown": "Cleaned markdown",
                        "latest_docx_bytes": _docx_bytes(Document()),
                    }
                )
            )
            return "succeeded", SimpleNamespace(
                uploaded_filename="chapter-region.pdf",
                uploaded_file_bytes=b"%PDF-1.4 chapter-region",
                paragraphs=prepared_paragraphs,
                image_assets=[],
                jobs=[{"job_kind": "block"}],
                source_text="Source chapter region text.",
                preparation_cached=False,
                preparation_elapsed_seconds=0.1,
                ai_classified_count=0,
                ai_heading_count=0,
                ai_role_change_count=0,
                ai_heading_promotion_count=0,
                ai_heading_demotion_count=0,
                ai_structural_role_change_count=0,
                structure_repair_report=None,
                uploaded_file_token="token-1",
            )

    monkeypatch.setattr(validation, "PROJECT_ROOT", tmp_path)
    monkeypatch.setattr(validation, "REAL_DOCUMENT_ARTIFACT_ROOT", tmp_path / "artifacts")
    monkeypatch.setattr(validation, "load_validation_registry", lambda: registry)
    monkeypatch.setattr(validation, "load_app_config", lambda: object())
    monkeypatch.setattr(
        validation,
        "resolve_runtime_resolution",
        lambda app_config, run_profile: SimpleNamespace(
            effective=_resolution_payload(
                chunk_size=6000,
                image_mode="safe",
                keep_all_image_variants=False,
                model="gpt-5.4",
                max_retries=1,
                processing_operation="translate",
                source_language="auto",
                target_language="ru",
            ),
            ui_defaults=_resolution_payload(
                chunk_size=6000,
                image_mode="safe",
                keep_all_image_variants=False,
                model="gpt-5.4",
                max_retries=1,
                processing_operation="translate",
                source_language="auto",
                target_language="ru",
            ),
            overrides={
                "reader_cleanup_enabled": True,
                "reader_cleanup_policy": "advisory",
                "reader_verifier_enabled": True,
                "structure_recognition_mode": "off",
            },
        ),
    )
    monkeypatch.setattr(
        validation,
        "apply_runtime_resolution_to_app_config",
        lambda app_config, resolution: {"reader_verifier_enabled": True},
    )
    monkeypatch.setattr(
        validation,
        "evaluate_lietaer_acceptance",
        lambda report, **kwargs: {"passed": False, "failed_checks": ["false_fragment_headings_present"], "checks": []},
    )
    monkeypatch.setattr(validation, "_snapshot_formatting_diagnostics_paths", lambda: set())
    monkeypatch.setattr(validation, "_collect_new_formatting_diagnostics_paths", lambda before, after: [])
    monkeypatch.setattr(validation, "_extract_run_formatting_diagnostics_paths", lambda event_log: [])
    monkeypatch.setattr(validation, "_load_recent_formatting_diagnostics", lambda started_at: ([], []))
    monkeypatch.setattr(validation, "_load_formatting_diagnostics_payloads", lambda paths: [])
    monkeypatch.setattr(validation, "_load_translation_quality_report", lambda event_log: (None, None))
    monkeypatch.setattr(validation, "_print_terminal_completion_summary", lambda **kwargs: None)
    monkeypatch.setattr(
        validation.processing_service,
        "clone_processing_service",
        lambda **kwargs: _ValidationServiceStub(kwargs["log_event_fn"]),
    )
    selector_payload = SimpleNamespace(
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
    )
    monkeypatch.setattr(
        validation,
        "describe_provider_availability",
        lambda selector, app_config=None: SimpleNamespace(
            enabled=True,
            has_api_key=True,
            error_message=None,
            selector=selector_payload,
        ),
    )
    monkeypatch.setattr(validation, "resolve_model_selector", lambda *args, **kwargs: selector_payload)
    monkeypatch.setattr(validation, "get_client_for_model_selector", lambda *args, **kwargs: object())
    monkeypatch.setattr(
        validation,
        "generate_markdown_block",
        lambda **kwargs: json.dumps(
            {
                "overall_verdict": "cleaned_better",
                "cleaned_audit_verdict": "improved_but_has_remaining_issues",
                "reader_quality_score_raw": 4,
                "reader_quality_score_cleaned": 7,
                "confidence": "high",
                "noise_removed": ["Repeated running header removed."],
                "possible_false_deletions": [],
                "readability_regressions": ["Footnote markers still feel noisy in places."],
                "remaining_issues": [
                    {
                        "category": "broken_list_marker",
                        "severity": "medium",
                        "artifact": "cleaned_markdown",
                        "line_ref": "cleaned_markdown:3",
                        "snippet": "• Currency menu entry",
                        "why_reader_hurts": "The raw bullet marker still looks unfinished to the reader.",
                        "recommended_fix_type": "normalize_list",
                    }
                ],
                "evidence_anchors": [
                    {
                        "kind": "improvement_seen",
                        "artifact": "comparison",
                        "line_ref": "comparison:1",
                        "snippet": "Repeated running header removed.",
                        "note": "The cleaned artifact removes repeated page furniture.",
                    },
                    {
                        "kind": "remaining_issue",
                        "artifact": "cleaned_markdown",
                        "line_ref": "cleaned_markdown:3",
                        "snippet": "• Currency menu entry",
                        "note": "A broken list marker still remains in the cleaned markdown.",
                    },
                ],
                "recommended_next_changes": [
                    {
                        "change_type": "prompt",
                        "recommendation": "Tell the cleanup verifier to focus more on leftover orphan markers.",
                        "why": "The cleaned sample still has minor reader-visible noise.",
                    }
                ],
                "summary_for_human": "Cleaned output is easier to read than raw output for this slice.",
                "simple_user_summary": "The cleaned version is easier to read than the raw translation. Most of the benefit comes from removing repeated reader-visible noise, and no major text loss was detected at current review confidence.",
                "simple_user_risk_statement": "Some minor reader-visible noise remains, but no major semantic loss was identified.",
                "simple_user_next_step": "Tighten the cleanup prompt for leftover orphan markers.",
            },
            ensure_ascii=False,
        ),
    )

    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_PROFILE", "lietaer-pdf-chapter-region-core")
    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_RUN_PROFILE", "ui-parity-translate-simple-reader-cleanup-comparison-only")
    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_FORCED_RUN_ID", "run-verifier")

    validation.main()

    run_dir = tmp_path / "artifacts" / "runs" / "run-verifier"
    report_path = run_dir / "lietaer_pdf_chapter_region_report.json"
    summary_path = run_dir / "lietaer_pdf_chapter_region_summary.txt"
    evidence_path = run_dir / "lietaer_pdf_chapter_region_reader_quality_evidence.json"
    review_json_path = run_dir / "lietaer_pdf_chapter_region_reader_quality_review.json"
    review_md_path = run_dir / "lietaer_pdf_chapter_region_reader_quality_review.md"
    status_md_path = run_dir / "lietaer_pdf_chapter_region_reader_mvp_status.md"
    latest_manifest_path = tmp_path / "artifacts" / "lietaer_pdf_chapter_region_latest.json"

    report = json.loads(report_path.read_text(encoding="utf-8"))
    summary_text = summary_path.read_text(encoding="utf-8")
    review = json.loads(review_json_path.read_text(encoding="utf-8"))
    review_md = review_md_path.read_text(encoding="utf-8")
    status_md = status_md_path.read_text(encoding="utf-8")
    evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
    latest_manifest = json.loads(latest_manifest_path.read_text(encoding="utf-8"))

    assert evidence["evidence_mode"] == "full_selected_slice"
    assert evidence["base_artifacts_present"] is True
    assert evidence["pre_audit_issue_counts"]["page_furniture_inline"] >= 1
    assert evidence["pre_audit_issue_counts"]["broken_list_marker"] >= 1
    assert evidence["pre_audit_issue_counts"]["fragmented_paragraph"] >= 1
    assert evidence["pre_audit_issue_counts"]["duplicate_fragment"] >= 1
    assert evidence["pre_audit_findings"]
    assert review["review_mode"] == "development_only_non_acceptance"
    assert review["verifier_status"] == "completed"
    assert review["verifier_requested_selector"] == "openrouter:google/gemini-3-flash-preview"
    assert review["verifier_canonical_selector"] == "openrouter:google/gemini-3-flash-preview"
    assert review["verifier_provider"] == "openrouter"
    assert review["verifier_model_id"] == "google/gemini-3-flash-preview"
    assert review["overall_verdict"] == "cleaned_better"
    assert review["cleaned_audit_verdict"] == "improved_but_has_remaining_issues"
    assert review["pre_audit_issue_counts"]["broken_list_marker"] >= 1
    assert review["issue_summary_by_category"]["broken_list_marker"] == 1
    assert review["issue_summary_by_category"]["page_furniture_inline"] >= 1
    assert review["issue_summary_by_category"]["fragmented_paragraph"] >= 1
    assert review["issue_summary_by_category"]["duplicate_fragment"] >= 1
    assert review["simple_user_summary"].startswith("The cleaned version is easier to read")
    assert report["reader_verifier_evidence"]["artifact_paths"]["source_evidence_json"] == "artifacts/runs/run-verifier/lietaer_pdf_chapter_region_reader_quality_evidence.json"
    assert report["reader_verifier_evidence"]["artifact_paths"]["review_json"] == "artifacts/runs/run-verifier/lietaer_pdf_chapter_region_reader_quality_review.json"
    assert report["reader_verifier_evidence"]["artifact_paths"]["review_md"] == "artifacts/runs/run-verifier/lietaer_pdf_chapter_region_reader_quality_review.md"
    assert report["reader_verifier_evidence"]["artifact_paths"]["mvp_status_md"] == "artifacts/runs/run-verifier/lietaer_pdf_chapter_region_reader_mvp_status.md"
    assert report["output_artifacts"]["reader_mvp_status_md"] == "artifacts/runs/run-verifier/lietaer_pdf_chapter_region_reader_mvp_status.md"
    assert report["reader_mvp_status"]["status_label"] == "readable_draft_not_acceptance_ready"
    assert report["reader_mvp_status"]["comparison_only_acceptance_diagnostic"] is True
    assert report["reader_mvp_status"]["cleanup_score_delta"] == 3.0
    assert report["reader_mvp_status"]["no_false_deletions_reported"] is True
    assert report["reader_mvp_status"]["no_readability_regressions_reported"] is False
    assert report["reader_verifier_evidence"]["cleanup_diagnostics"]["ignored_reason_counts"]["prior_same_block_operation_not_applied"] == 1
    assert report["reader_verifier_evidence"]["validator_boundary"]["observer_only"] is True
    assert report["reader_mvp_status"]["blocker_groups"]["mapping_quality_gate_diagnostics"] == [
        "acceptance_diagnostic_checks=false_fragment_headings_present"
    ]
    assert report["reader_mvp_status"]["blocker_groups"]["cleanup_application_diagnostics"] == [
        "prior_same_block_operation_not_applied=1"
    ]
    assert latest_manifest["status"] == "completed"
    assert latest_manifest["reader_verifier_status"] == "completed"
    assert latest_manifest["reader_verifier_model_selector"] == "openrouter:google/gemini-3-flash-preview"
    assert latest_manifest["reader_verifier_model_id"] == "google/gemini-3-flash-preview"
    assert latest_manifest["reader_verifier_cleaned_audit_verdict"] == "improved_but_has_remaining_issues"
    assert latest_manifest["reader_verifier_cleanup_ignored_reason_counts"] == {
        "duplicate_operation_incompatible": 0,
        "heading_boundary_substrings_not_found": 0,
        "heading_boundary_unaccounted_text": 0,
        "noise_substring_not_found": 0,
        "prior_same_block_operation_not_applied": 1,
        "remove_inline_noise_not_exact_noise_pattern": 0,
    }
    assert latest_manifest["reader_verifier_remaining_issue_count"] == 4
    assert latest_manifest["reader_verifier_high_severity_issue_count"] == 2
    assert latest_manifest["reader_verifier_top_issue_categories"] == [
        "broken_list_marker",
        "duplicate_fragment",
        "fragmented_paragraph",
    ]
    assert latest_manifest["reader_mvp_status_label"] == "readable_draft_not_acceptance_ready"
    assert latest_manifest["reader_mvp_status_acceptance_diagnostic_only"] is True
    assert latest_manifest["reader_mvp_status_cleanup_score_delta"] == 3.0
    assert latest_manifest["reader_mvp_status_false_deletion_status"] == "none_reported"
    assert latest_manifest["reader_mvp_status_readability_regression_status"] == "reported"
    assert latest_manifest["reader_mvp_status_md"] == "artifacts/runs/run-verifier/lietaer_pdf_chapter_region_reader_mvp_status.md"
    assert "reader_verifier_status=completed" in summary_text
    assert "reader_verifier_overall_verdict=cleaned_better" in summary_text
    assert "reader_verifier_cleaned_audit_verdict=improved_but_has_remaining_issues" in summary_text
    assert "reader_verifier_remaining_issue_count=4" in summary_text
    assert "reader_verifier_high_severity_issue_count=2" in summary_text
    assert "reader_verifier_top_issue_categories=broken_list_marker,duplicate_fragment,fragmented_paragraph" in summary_text
    assert "reader_verifier_cleanup_ignored_reason_counts={\"duplicate_operation_incompatible\": 0, \"heading_boundary_substrings_not_found\": 0, \"heading_boundary_unaccounted_text\": 0, \"noise_substring_not_found\": 0, \"prior_same_block_operation_not_applied\": 1, \"remove_inline_noise_not_exact_noise_pattern\": 0}" in summary_text
    assert "reader_verifier_evidence_json=artifacts/runs/run-verifier/lietaer_pdf_chapter_region_reader_quality_evidence.json" in summary_text
    assert "reader_verifier_review_json=artifacts/runs/run-verifier/lietaer_pdf_chapter_region_reader_quality_review.json" in summary_text
    assert "reader_verifier_simple_user_next_step=Tighten the cleanup prompt for leftover orphan markers." in summary_text
    assert "reader_mvp_status_md=artifacts/runs/run-verifier/lietaer_pdf_chapter_region_reader_mvp_status.md" in summary_text
    assert "reader_mvp_status_acceptance_diagnostic_only=True" in summary_text
    assert "reader_mvp_status_cleanup_score_delta=3.0" in summary_text
    assert "reader_mvp_status_false_deletion_status=none_reported" in summary_text
    assert "reader_mvp_status_readability_regression_status=reported" in summary_text
    assert "reader_mvp_status_blocker_group_cleanup_application=prior_same_block_operation_not_applied=1" in summary_text
    assert "reader_mvp_status_blocker_group_quality_gate=acceptance_diagnostic_checks=false_fragment_headings_present" in summary_text
    assert "# Verdict" in review_md
    assert "# Out-of-Scope Filtered TOC Issues" in review_md
    assert "# Cleanup Application Diagnostics" in review_md
    assert "# Audit Verdict" in review_md
    assert "# Remaining Issues" in review_md
    assert "# Verifier Metadata" in review_md
    assert "# Статус MVP" in status_md
    assert "# Result Layers" in status_md
    assert "## Cleanup Application Diagnostics" in status_md
    assert "# Cleanup Diagnostic Examples" in status_md
    assert "comparison-only прогона" in status_md
    assert "Verifier не сообщил о false deletions." in status_md
    assert "Verifier сообщил о readability regressions: 1." in status_md


def test_build_reader_mvp_status_payload_groups_blockers_and_positive_signals() -> None:
    validation = _load_validation_module()

    status = validation._build_reader_mvp_status_payload(
        {
            "result": "succeeded",
            "validation_mode": {
                "comparison_only_validation": True,
                "acceptance_contract_active": False,
            },
            "runtime_config": {
                "target_language": "ru",
                "overrides": {"target_language": "ru"},
            },
            "acceptance": {
                "passed": False,
                "failed_checks": [
                    "formatting_diagnostics_threshold",
                    "unmapped_source_threshold",
                    "unmapped_target_threshold",
                ],
            },
            "reader_cleanup_evidence": {
                "stage_status": "completed",
                "failed_chunk_count": 1,
            },
            "reader_verifier_evidence": {
                "overall_verdict": "cleaned_better",
                "reader_quality_score_raw": 3,
                "reader_quality_score_cleaned": 5,
                "remaining_issues": [
                    {"severity": "high"},
                    {"severity": "high"},
                    {"severity": "medium"},
                ],
                "issue_summary_by_category": {
                    "heading_fused_with_body": 8,
                    "fragmented_paragraph": 2,
                    "page_furniture_inline": 1,
                },
                "possible_false_deletions": [],
                "readability_regressions": [],
                "filtered_toc_issue_count": 2,
                "filtered_toc_pre_audit_count": 1,
                "filtered_toc_verifier_issue_count": 1,
                "filtered_toc_evidence_anchor_count": 0,
                "cleanup_diagnostics": {
                    "ignored_reason_counts": {
                        "prior_same_block_operation_not_applied": 2,
                    },
                    "top_ignored_reasons": [
                        {
                            "ignored_reason": "prior_same_block_operation_not_applied",
                            "count": 2,
                            "examples": [
                                {
                                    "operation": "remove_inline_noise",
                                    "reason": "repeated_running_header",
                                    "text_preview": "10",
                                }
                            ],
                        }
                    ],
                },
            },
            "translation_quality_report": {
                "quality_status": "warn",
                "gate_reasons": ["unmapped_source_paragraphs_above_advisory_threshold"],
            },
        }
    )

    assert status["status_label"] == "readable_draft_not_acceptance_ready"
    assert status["comparison_only_acceptance_diagnostic"] is True
    assert status["cleanup_score_delta"] == 2.0
    assert status["no_false_deletions_reported"] is True
    assert status["no_readability_regressions_reported"] is True
    assert status["signal_layers"]["comparison_only_validation_status"] == "observer_only"
    assert status["signal_layers"]["cleaned_audit_verdict"] == "unclear"
    assert status["filtered_issue_counts"] == {
        "toc_total": 2,
        "toc_pre_audit": 1,
        "toc_verifier_issue": 1,
        "toc_evidence_anchor": 0,
    }
    assert status["blocker_groups"]["cleanup_contract"] == ["cleanup_chunk_failures=1"]
    assert status["blocker_groups"]["reader_visible_cleanup_defects"] == [
        "heading_fused_with_body=8",
        "fragmented_paragraph=2",
        "page_furniture_inline=1",
    ]
    assert status["blocker_groups"]["cleanup_application_diagnostics"] == [
        "prior_same_block_operation_not_applied=2"
    ]
    assert status["blocker_groups"]["mapping_quality_gate_diagnostics"] == [
        "translation_quality_status=warn",
        "translation_quality_gate_reasons=unmapped_source_paragraphs_above_advisory_threshold",
        "acceptance_diagnostic_checks=formatting_diagnostics_threshold,unmapped_source_threshold,unmapped_target_threshold",
    ]

    status_md = validation._render_reader_mvp_status_markdown(status)
    summary_lines = validation._build_reader_mvp_status_summary_lines(status)

    assert "# Статус MVP" in status_md
    assert "Есть полезное улучшение читабельности" in status_md
    assert "observer_only" in status_md
    assert "Out-of-scope filtered issue counts: total=2, pre_audit=1, verifier=1, anchors=0" in status_md
    assert "Verifier не сообщил о false deletions." in status_md
    assert "Verifier не сообщил о readability regressions." in status_md
    assert "prior_same_block_operation_not_applied=2" in status_md
    assert "reader_mvp_status_acceptance_diagnostic_only=True" in summary_lines
    assert "reader_mvp_status_cleanup_score_delta=2.0" in summary_lines
    assert "reader_mvp_status_false_deletion_status=none_reported" in summary_lines
    assert "reader_mvp_status_readability_regression_status=none_reported" in summary_lines
    assert 'reader_mvp_status_filtered_issue_counts={"toc_evidence_anchor": 0, "toc_pre_audit": 1, "toc_total": 2, "toc_verifier_issue": 1}' in summary_lines
    assert "reader_mvp_status_blocker_group_cleanup_application=prior_same_block_operation_not_applied=2" in summary_lines


def test_build_reader_mvp_status_payload_does_not_treat_diagnostic_anchor_targets_as_cleanup_contract_blocker() -> None:
    validation = _load_validation_module()

    status = validation._build_reader_mvp_status_payload(
        {
            "result": "succeeded",
            "validation_mode": {
                "comparison_only_validation": True,
                "acceptance_contract_active": False,
            },
            "reader_cleanup_evidence": {
                "stage_status": "completed",
                "failed_chunk_count": 0,
                "anchor_repair_status": "diagnostic_only_not_applied",
                "recommended_anchor_target_count": 2,
                "recommended_anchor_targets": [
                    {
                        "anchor_id": "diag-anchor-1",
                        "category": "heading_fused_with_body",
                        "block_id": "block-2",
                        "line_ref": "cleaned_markdown:3",
                        "snippet": "КАК ЭТО РАБОТАЕТ",
                    }
                ],
            },
            "reader_verifier_evidence": {
                "overall_verdict": "cleaned_better",
                "cleaned_audit_verdict": "improved_but_has_remaining_issues",
                "reader_quality_score_raw": 3,
                "reader_quality_score_cleaned": 5,
                "remaining_issues": [{"severity": "medium"}],
                "issue_summary_by_category": {"heading_fused_with_body": 1},
                "possible_false_deletions": [],
                "readability_regressions": [],
            },
            "acceptance": {"passed": False, "failed_checks": []},
        }
    )

    assert status["anchor_repair_status"] == "diagnostic_only_not_applied"
    assert status["recommended_anchor_target_count"] == 2
    assert "anchor_repair_status=diagnostic_only_not_applied" not in status["blocker_groups"]["cleanup_contract"]


def test_build_reader_mvp_status_payload_avoids_success_wording_for_raw_better_or_unclear() -> None:
    validation = _load_validation_module()

    raw_better_status = validation._build_reader_mvp_status_payload(
        {
            "result": "succeeded",
            "runtime_config": {
                "target_language": "ru",
                "overrides": {"target_language": "ru"},
            },
            "reader_verifier_evidence": {
                "overall_verdict": "raw_better",
                "cleaned_audit_verdict": "improved_but_has_remaining_issues",
                "reader_quality_score_raw": 6,
                "reader_quality_score_cleaned": 4,
                "remaining_issues": [],
                "issue_summary_by_category": {},
                "possible_false_deletions": [],
                "readability_regressions": [],
            },
            "acceptance": {"passed": True, "failed_checks": []},
        }
    )
    mixed_status = validation._build_reader_mvp_status_payload(
        {
            "result": "succeeded",
            "runtime_config": {
                "target_language": "en",
                "overrides": {"target_language": "en"},
            },
            "reader_verifier_evidence": {
                "overall_verdict": "mixed",
                "cleaned_audit_verdict": "unclear",
                "reader_quality_score_raw": 4,
                "reader_quality_score_cleaned": 4,
                "remaining_issues": [],
                "issue_summary_by_category": {},
                "possible_false_deletions": [],
                "readability_regressions": [],
            },
            "acceptance": {"passed": True, "failed_checks": []},
        }
    )

    assert raw_better_status["status_label"] == "cleanup_regressed"
    assert raw_better_status["status_label"] != "cleaned_better_diagnostic_evidence"
    assert "Стало лучше" not in raw_better_status["user_summary"]
    assert "cleaned output is easier" not in raw_better_status["user_summary"]
    assert "cleaned output is easier to read" not in raw_better_status["user_summary"]

    assert mixed_status["status_label"] == "mixed_or_unclear"
    assert mixed_status["status_label"] != "cleaned_better_diagnostic_evidence"
    assert "mixed or unclear" in mixed_status["user_summary"]
    assert "Стало лучше" not in mixed_status["user_summary"]
    assert "cleaned output is easier" not in mixed_status["user_summary"]
    assert "cleaned output is easier to read" not in mixed_status["user_summary"]


def test_build_reader_mvp_status_payload_avoids_success_wording_for_pipeline_failed() -> None:
    validation = _load_validation_module()

    status = validation._build_reader_mvp_status_payload(
        {
            "result": "failed",
            "runtime_config": {
                "target_language": "en",
                "overrides": {"target_language": "en"},
            },
            "reader_verifier_evidence": {
                "overall_verdict": "cleaned_better",
                "cleaned_audit_verdict": "clean",
                "reader_quality_score_raw": 3,
                "reader_quality_score_cleaned": 6,
                "remaining_issues": [],
                "issue_summary_by_category": {},
                "possible_false_deletions": [],
                "readability_regressions": [],
            },
            "acceptance": {"passed": False, "failed_checks": []},
        }
    )

    assert status["status_label"] == "pipeline_failed"
    assert "pipeline did not succeed" in status["pipeline_summary"]
    assert "pipeline did not succeed" in status["user_summary"]
    assert "cleaned output is easier" not in status["user_summary"]
    assert "cleaned output is easier to read" not in status["user_summary"]
    assert "cleanup quality" in status["risk_summary"]


def test_build_reader_mvp_status_payload_treats_false_deletions_or_readability_regressions_as_non_success() -> None:
    validation = _load_validation_module()

    status = validation._build_reader_mvp_status_payload(
        {
            "result": "succeeded",
            "runtime_config": {
                "target_language": "en",
                "overrides": {"target_language": "en"},
            },
            "reader_verifier_evidence": {
                "overall_verdict": "cleaned_better",
                "cleaned_audit_verdict": "improved_but_has_remaining_issues",
                "reader_quality_score_raw": 3,
                "reader_quality_score_cleaned": 6,
                "remaining_issues": [],
                "issue_summary_by_category": {},
                "possible_false_deletions": ["The verifier suspects one paragraph may have been removed."],
                "readability_regressions": ["One list is still harder to scan after cleanup."],
            },
            "acceptance": {"passed": True, "failed_checks": []},
        }
    )

    assert status["status_label"] == "readable_draft_not_acceptance_ready"
    assert status["status_label"] != "cleaned_better_diagnostic_evidence"
    assert status["no_false_deletions_reported"] is False
    assert status["no_readability_regressions_reported"] is False
    assert "not acceptance-ready" in status["user_summary"]
    assert "false deletions=1" in status["risk_summary"]
    assert "readability regressions=1" in status["risk_summary"]


def test_build_reader_cleanup_anchor_targets_ignores_raw_and_comparison_issues() -> None:
    validation = _load_validation_module()

    cleaned_markdown = "Intro\n\nКАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь.\n\nTail"
    raw_or_comparison_targets = validation._build_reader_cleanup_anchor_targets(
        review_payload={
            "remaining_issues": [
                {
                    "category": "heading_fused_with_body",
                    "artifact": "raw_markdown",
                    "line_ref": "raw_markdown:3",
                    "snippet": "КАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь.",
                },
                {
                    "category": "heading_fused_with_body",
                    "artifact": "comparison",
                    "line_ref": "comparison:8",
                    "snippet": "КАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь.",
                },
            ]
        },
        cleaned_markdown=cleaned_markdown,
    )
    cleaned_targets = validation._build_reader_cleanup_anchor_targets(
        review_payload={
            "remaining_issues": [
                {
                    "category": "heading_fused_with_body",
                    "artifact": "cleaned_markdown",
                    "line_ref": "cleaned_markdown:3",
                    "snippet": "КАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь.",
                }
            ]
        },
        cleaned_markdown=cleaned_markdown,
    )

    assert raw_or_comparison_targets == []
    assert len(cleaned_targets) == 1
    assert cleaned_targets[0]["category"] == "heading_fused_with_body"
    assert cleaned_targets[0]["line_ref"] == "cleaned_markdown:3"


def test_load_runtime_reader_cleanup_anchor_targets_from_env_accepts_review_payload(
    tmp_path,
    monkeypatch,
) -> None:
    validation = _load_validation_module()

    anchor_path = tmp_path / "reader_quality_review.json"
    heading_target = {
        "anchor_id": "anchor-heading",
        "category": "heading_fused_with_body",
        "block_id": "b_000002",
        "line_ref": "cleaned_markdown:3",
        "snippet": "КАК ЭТО РАБОТАЕТ: Местные органы власти...",
    }
    out_of_scope_target = {
        "anchor_id": "anchor-quote",
        "category": "quote_not_block_formatted",
        "block_id": "b_000009",
        "line_ref": "cleaned_markdown:18",
        "snippet": "Quote.",
    }
    anchor_path.write_text(
        json.dumps(
            {"recommended_anchor_targets": [heading_target, out_of_scope_target]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    monkeypatch.setenv("DOCXAI_READER_CLEANUP_ANCHOR_TARGETS_PATH", str(anchor_path))
    monkeypatch.delenv("DOCXAI_READER_CLEANUP_ANCHOR_TARGETS_JSON", raising=False)

    targets = validation._load_runtime_reader_cleanup_anchor_targets_from_env()

    assert targets == [heading_target]

    monkeypatch.setenv("DOCXAI_READER_CLEANUP_ANCHOR_TARGETS_JSON", json.dumps([heading_target], ensure_ascii=False))
    monkeypatch.delenv("DOCXAI_READER_CLEANUP_ANCHOR_TARGETS_PATH", raising=False)

    assert validation._load_runtime_reader_cleanup_anchor_targets_from_env() == [heading_target]


def test_load_runtime_reader_cleanup_anchor_targets_from_env_prefers_fresh_verifier_targets(
    tmp_path,
    monkeypatch,
) -> None:
    validation = _load_validation_module()

    anchor_path = tmp_path / "reader_quality_review.json"
    stale_runtime_target = {
        "anchor_id": "stale-page-furniture",
        "category": "page_furniture_inline",
        "block_id": "b_000278",
        "line_ref": "559",
        "snippet": "20 ДЕНЬГИ, КОТОРЫЕ ПАХНУТ?",
    }
    fresh_verifier_target = {
        "anchor_id": "fresh-page-furniture",
        "category": "page_furniture_inline",
        "block_id": "b_000228",
        "line_ref": "459",
        "snippet": "190 ПЕРЕОСМЫСЛЕНИЕ ДЕНЕГ Особый интерес представляет",
    }
    anchor_path.write_text(
        json.dumps(
            {
                "recommended_anchor_targets": [stale_runtime_target],
                "verifier_recommended_anchor_targets": [fresh_verifier_target],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    monkeypatch.setenv("DOCXAI_READER_CLEANUP_ANCHOR_TARGETS_PATH", str(anchor_path))
    monkeypatch.delenv("DOCXAI_READER_CLEANUP_ANCHOR_TARGETS_JSON", raising=False)

    assert validation._load_runtime_reader_cleanup_anchor_targets_from_env() == [fresh_verifier_target]


def test_write_reader_verifier_artifacts_keeps_anchor_repair_diagnostic_only(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()

    artifact_dir = tmp_path / "artifacts" / "runs" / "run-anchor"
    artifact_dir.mkdir(parents=True, exist_ok=True)
    ui_results_dir = tmp_path / ".run" / "ui_results"
    ui_results_dir.mkdir(parents=True, exist_ok=True)
    source_path = tmp_path / "chapter-region.pdf"
    source_path.write_bytes(b"%PDF-1.4 anchor")

    initial_cleaned_markdown = (
        "Intro\n\nPreface\n\nКАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь.\n\nTail\n\nOutro"
    )
    cleaned_markdown_path = ui_results_dir / "chapter_region.result.md"
    cleaned_docx_path = ui_results_dir / "chapter_region.result.docx"
    raw_markdown_path = ui_results_dir / "chapter_region.raw.result.md"
    cleanup_report_path = ui_results_dir / "chapter_region.reader_cleanup_report.json"
    cleaned_markdown_path.write_text(initial_cleaned_markdown, encoding="utf-8")
    cleaned_docx_path.write_bytes(_docx_bytes(Document()))
    raw_markdown_path.write_text("Raw\n\nBody", encoding="utf-8")
    cleanup_report_path.write_text(
        json.dumps(
            {
                "stage_status": "completed",
                "changed": True,
                "warnings": [],
                "global_plan": {
                    "repeated_noise_patterns": [],
                    "candidate_block_ids": [],
                    "document_specific_running_headers": [],
                    "examples_do_not_delete": [],
                    "likely_heading_body_patterns": [],
                    "likely_fragmentation_patterns": [],
                    "warnings": [],
                },
                "stats": {
                    "raw_block_count": 5,
                    "raw_char_count": len(initial_cleaned_markdown),
                    "cleanup_chunk_count": 1,
                    "failed_chunk_count": 0,
                    "proposed_cleanup_operation_count": 0,
                    "proposed_delete_block_count": 0,
                    "accepted_cleanup_operation_count": 0,
                    "accepted_delete_block_count": 0,
                    "ignored_cleanup_operation_count": 0,
                    "ignored_delete_block_count": 0,
                    "deleted_non_whitespace_char_count": 0,
                    "deleted_char_ratio": 0.0,
                },
                "accepted_cleanup_operations": [],
                "accepted_delete_blocks": [],
                "ignored_delete_blocks": [],
                "chunk_results": [
                    {
                        "chunk_index": 1,
                        "status": "completed",
                            "target_block_count": 5,
                        "target_chars": len(initial_cleaned_markdown),
                        "elapsed_ms": 1.0,
                        "proposed_cleanup_operation_count": 0,
                        "proposed_delete_block_count": 0,
                        "accepted_cleanup_operation_count": 0,
                        "accepted_delete_block_count": 0,
                        "ignored_cleanup_operation_count": 0,
                        "ignored_delete_block_count": 0,
                        "repair_attempted": False,
                        "repair_status": "not_attempted",
                        "schema_validation_error": "",
                        "repair_error": "",
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    selector_payload = SimpleNamespace(
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
    )
    original_docx_bytes = cleaned_docx_path.read_bytes()
    call_counts = {"verifier": 0}

    monkeypatch.setattr(
        validation,
        "describe_provider_availability",
        lambda selector, app_config=None: SimpleNamespace(
            enabled=True,
            has_api_key=True,
            error_message=None,
            selector=selector_payload,
        ),
    )
    monkeypatch.setattr(validation, "resolve_model_selector", lambda *args, **kwargs: selector_payload)
    monkeypatch.setattr(validation, "get_client_for_model_selector", lambda *args, **kwargs: object())
    def _fake_generate_markdown_block(**kwargs):
        call_counts["verifier"] += 1
        return json.dumps(
            {
                "overall_verdict": "cleaned_better",
                "cleaned_audit_verdict": "improved_but_has_remaining_issues",
                "reader_quality_score_raw": 4,
                "reader_quality_score_cleaned": 6,
                "confidence": "high",
                "noise_removed": ["The cleaned artifact is better than raw overall."],
                "possible_false_deletions": [],
                "readability_regressions": [],
                "remaining_issues": [
                    {
                        "category": "heading_fused_with_body",
                        "severity": "high",
                        "artifact": "cleaned_markdown",
                        "line_ref": "cleaned_markdown:3",
                        "snippet": "КАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь.",
                        "why_reader_hurts": "The heading is fused with the first sentence of the paragraph.",
                        "recommended_fix_type": "split_heading",
                    }
                ],
                "evidence_anchors": [
                    {
                        "kind": "improvement_seen",
                        "artifact": "comparison",
                        "line_ref": "comparison:1",
                        "snippet": "Some initial cleanup already happened.",
                        "note": "The first cleanup pass removed other noise.",
                    },
                    {
                        "kind": "remaining_issue",
                        "artifact": "cleaned_markdown",
                        "line_ref": "cleaned_markdown:3",
                        "snippet": "КАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь.",
                        "note": "The heading remains fused with body text.",
                    },
                ],
                "recommended_next_changes": [
                    {
                        "change_type": "operation_contract",
                        "recommendation": "Apply a bounded heading boundary repair to the remaining fused title.",
                        "why": "The exact heading/body split is still reader-visible.",
                    }
                ],
                "summary_for_human": "One reader-visible heading/body issue still remains.",
                "simple_user_summary": "The cleaned version is easier to read than the raw translation.",
                "simple_user_risk_statement": "One fused heading still remains in the cleaned artifact.",
                "simple_user_next_step": "Run a bounded anchor repair pass for the remaining heading boundary.",
            },
            ensure_ascii=False,
        )

    monkeypatch.setattr(validation, "generate_markdown_block", _fake_generate_markdown_block)

    reader_cleanup_evidence = {
        "artifacts_present": True,
        "cleaned_markdown_path": str(cleaned_markdown_path),
        "cleaned_docx_path": str(cleaned_docx_path),
        "raw_markdown_path": str(raw_markdown_path),
        "reader_cleanup_report_path": str(cleanup_report_path),
        "stage_status": "completed",
        "changed": True,
        "accepted_delete_block_count": 0,
        "ignored_delete_block_count": 0,
        "rejected_delete_block_count": 0,
        "failed_chunk_count": 0,
        "deleted_block_previews": [],
    }

    review = validation._write_reader_verifier_artifacts(
        run_id="run-anchor",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        artifact_prefix="lietaer_pdf_chapter_region",
        artifact_dir=artifact_dir,
        source_document_path=source_path,
        source_text="Synthetic source text.",
        reader_cleanup_evidence=reader_cleanup_evidence,
        app_config=object(),
        runtime_app_config={
            "reader_cleanup_enabled": True,
            "reader_cleanup_anchor_repair_enabled": True,
            "reader_cleanup_policy": "advisory",
            "reader_verifier_enabled": True,
            "model": "gpt-5.4",
        },
        validation_mode={
            "comparison_only_validation": True,
            "validation_run_type": "comparison_only",
        },
        max_retries=1,
    )

    updated_cleanup_report = json.loads(cleanup_report_path.read_text(encoding="utf-8"))
    updated_evidence = review["updated_reader_cleanup_evidence"]
    persisted_review = json.loads(
        (artifact_dir / "lietaer_pdf_chapter_region_reader_quality_review.json").read_text(encoding="utf-8")
    )
    persisted_evidence = json.loads(
        (artifact_dir / "lietaer_pdf_chapter_region_reader_quality_evidence.json").read_text(encoding="utf-8")
    )

    assert call_counts == {"verifier": 1}
    assert cleaned_markdown_path.read_text(encoding="utf-8") == initial_cleaned_markdown
    assert cleaned_docx_path.read_bytes() == original_docx_bytes
    assert "passes" not in updated_cleanup_report
    assert updated_evidence["reader_cleanup_report_path"] == str(cleanup_report_path)
    assert updated_evidence["anchor_repair_status"] == "diagnostic_only_not_applied"
    assert updated_evidence["recommended_anchor_target_count"] == len(updated_evidence["recommended_anchor_targets"])
    assert updated_evidence["recommended_anchor_target_count"] >= 1
    assert review["cleaned_audit_verdict"] == "improved_but_has_remaining_issues"
    assert review["anchor_repair_status"] == "diagnostic_only_not_applied"
    assert review["recommended_anchor_target_count"] == len(review["recommended_anchor_targets"])
    assert review["validator_boundary"]["observer_only"] is True
    assert review["validator_boundary"]["runs_anchor_repair"] is False
    assert any(target["category"] == "heading_fused_with_body" for target in review["recommended_anchor_targets"])
    assert persisted_review["anchor_repair_status"] == "diagnostic_only_not_applied"
    assert persisted_review["recommended_anchor_target_count"] == len(persisted_review["recommended_anchor_targets"])
    assert persisted_review["validator_boundary"]["mutates_cleaned_markdown"] is False
    assert persisted_review["validator_boundary"]["rebuilds_docx"] is False
    assert persisted_evidence["anchor_repair_status"] == "diagnostic_only_not_applied"
    assert persisted_evidence["recommended_anchor_target_count"] == len(persisted_evidence["recommended_anchor_targets"])
    assert persisted_evidence["validator_boundary"]["mutates_cleaned_docx"] is False


def test_write_reader_verifier_artifacts_preserves_runtime_applied_anchor_targets(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()

    artifact_dir = tmp_path / "artifacts" / "runs" / "run-anchor-runtime"
    artifact_dir.mkdir(parents=True, exist_ok=True)
    ui_results_dir = tmp_path / ".run" / "ui_results"
    ui_results_dir.mkdir(parents=True, exist_ok=True)
    source_path = tmp_path / "chapter-region-runtime.pdf"
    source_path.write_bytes(b"%PDF-1.4 runtime-anchor")

    initial_cleaned_markdown = (
        "Runtime intro\n\nКАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь.\n\nTail block"
    )
    cleaned_markdown_path = ui_results_dir / "chapter_region_runtime.result.md"
    cleaned_docx_path = ui_results_dir / "chapter_region_runtime.result.docx"
    raw_markdown_path = ui_results_dir / "chapter_region_runtime.raw.result.md"
    cleanup_report_path = ui_results_dir / "chapter_region_runtime.reader_cleanup_report.json"
    cleaned_markdown_path.write_text(initial_cleaned_markdown, encoding="utf-8")
    cleaned_docx_path.write_bytes(_docx_bytes(Document()))
    raw_markdown_path.write_text("Raw runtime\n\nBody", encoding="utf-8")
    cleanup_report_path.write_text(
        json.dumps(
            {
                "stage_status": "completed",
                "changed": True,
                "warnings": [],
                "global_plan": {
                    "repeated_noise_patterns": [],
                    "candidate_block_ids": [],
                    "document_specific_running_headers": [],
                    "examples_do_not_delete": [],
                    "likely_heading_body_patterns": [],
                    "likely_fragmentation_patterns": [],
                    "warnings": [],
                },
                "stats": {
                    "raw_block_count": 3,
                    "raw_char_count": len(initial_cleaned_markdown),
                    "cleanup_chunk_count": 1,
                    "failed_chunk_count": 0,
                    "proposed_cleanup_operation_count": 1,
                    "proposed_delete_block_count": 0,
                    "accepted_cleanup_operation_count": 1,
                    "accepted_delete_block_count": 0,
                    "ignored_cleanup_operation_count": 0,
                    "ignored_delete_block_count": 0,
                    "deleted_non_whitespace_char_count": 0,
                    "deleted_char_ratio": 0.0,
                },
                "accepted_cleanup_operations": [],
                "accepted_delete_blocks": [],
                "ignored_delete_blocks": [],
                "chunk_results": [
                    {
                        "chunk_index": 1,
                        "status": "completed",
                        "target_block_count": 3,
                        "target_chars": len(initial_cleaned_markdown),
                        "elapsed_ms": 1.0,
                        "proposed_cleanup_operation_count": 1,
                        "proposed_delete_block_count": 0,
                        "accepted_cleanup_operation_count": 1,
                        "accepted_delete_block_count": 0,
                        "ignored_cleanup_operation_count": 0,
                        "ignored_delete_block_count": 0,
                        "repair_attempted": True,
                        "repair_status": "applied",
                        "schema_validation_error": "",
                        "repair_error": "",
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    selector_payload = SimpleNamespace(
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
    )
    runtime_anchor_target = {
        "anchor_id": "runtime-anchor-1",
        "category": "page_furniture_inline",
        "block_id": "block_000001",
        "line_ref": "cleaned_markdown:1",
        "snippet": "Runtime intro",
    }

    monkeypatch.setattr(
        validation,
        "describe_provider_availability",
        lambda selector, app_config=None: SimpleNamespace(
            enabled=True,
            has_api_key=True,
            error_message=None,
            selector=selector_payload,
        ),
    )
    monkeypatch.setattr(validation, "resolve_model_selector", lambda *args, **kwargs: selector_payload)
    monkeypatch.setattr(validation, "get_client_for_model_selector", lambda *args, **kwargs: object())

    def _fake_generate_markdown_block(**kwargs):
        return json.dumps(
            {
                "overall_verdict": "cleaned_better",
                "cleaned_audit_verdict": "improved_but_has_remaining_issues",
                "reader_quality_score_raw": 4,
                "reader_quality_score_cleaned": 6,
                "confidence": "high",
                "noise_removed": ["The runtime pass already applied one safe anchor repair."],
                "possible_false_deletions": [],
                "readability_regressions": [],
                "remaining_issues": [
                    {
                        "category": "heading_fused_with_body",
                        "severity": "high",
                        "artifact": "cleaned_markdown",
                        "line_ref": "cleaned_markdown:3",
                        "snippet": "КАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь.",
                        "why_reader_hurts": "The heading remains fused with body text.",
                        "recommended_fix_type": "split_heading",
                    }
                ],
                "evidence_anchors": [
                    {
                        "kind": "remaining_issue",
                        "artifact": "cleaned_markdown",
                        "line_ref": "cleaned_markdown:3",
                        "snippet": "КАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь.",
                        "note": "A verifier-only diagnostic anchor still exists.",
                    }
                ],
                "recommended_next_changes": [
                    {
                        "change_type": "operation_contract",
                        "recommendation": "Keep the runtime-repaired anchor and inspect the remaining fused heading separately.",
                        "why": "Runtime repair evidence should remain intact while diagnostics stay observer-only.",
                    }
                ],
                "summary_for_human": "One remaining verifier-only heading issue is still visible.",
                "simple_user_summary": "The cleaned version is easier to read than the raw translation.",
                "simple_user_risk_statement": "One verifier-only heading issue remains visible.",
                "simple_user_next_step": "Inspect the remaining fused heading without mutating the cleaned artifact.",
            },
            ensure_ascii=False,
        )

    monkeypatch.setattr(validation, "generate_markdown_block", _fake_generate_markdown_block)

    reader_cleanup_evidence = {
        "artifacts_present": True,
        "cleaned_markdown_path": str(cleaned_markdown_path),
        "cleaned_docx_path": str(cleaned_docx_path),
        "raw_markdown_path": str(raw_markdown_path),
        "reader_cleanup_report_path": str(cleanup_report_path),
        "stage_status": "completed",
        "changed": True,
        "accepted_delete_block_count": 0,
        "ignored_delete_block_count": 0,
        "rejected_delete_block_count": 0,
        "failed_chunk_count": 0,
        "deleted_block_previews": [],
        "anchor_repair_status": "runtime_applied",
        "recommended_anchor_targets": [runtime_anchor_target],
        "recommended_anchor_target_count": 1,
    }

    review = validation._write_reader_verifier_artifacts(
        run_id="run-anchor-runtime",
        document_profile_id="lietaer-pdf-chapter-region-core",
        run_profile_id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        artifact_prefix="lietaer_pdf_chapter_region_runtime",
        artifact_dir=artifact_dir,
        source_document_path=source_path,
        source_text="Synthetic source text.",
        reader_cleanup_evidence=reader_cleanup_evidence,
        app_config=object(),
        runtime_app_config={
            "reader_cleanup_enabled": True,
            "reader_cleanup_anchor_repair_enabled": True,
            "reader_cleanup_policy": "advisory",
            "model": "gpt-5.4",
        },
        validation_mode={
            "comparison_only_validation": True,
            "validation_run_type": "comparison_only",
        },
        max_retries=1,
    )

    updated_evidence = review["updated_reader_cleanup_evidence"]
    persisted_review = json.loads(
        (artifact_dir / "lietaer_pdf_chapter_region_runtime_reader_quality_review.json").read_text(encoding="utf-8")
    )
    persisted_evidence = json.loads(
        (artifact_dir / "lietaer_pdf_chapter_region_runtime_reader_quality_evidence.json").read_text(encoding="utf-8")
    )

    assert updated_evidence["anchor_repair_status"] == "runtime_applied"
    assert updated_evidence["recommended_anchor_targets"] == [runtime_anchor_target]
    assert updated_evidence["recommended_anchor_target_count"] == 1
    assert review["anchor_repair_status"] == "runtime_applied"
    assert review["recommended_anchor_targets"] == [runtime_anchor_target]
    assert review["recommended_anchor_target_count"] == 1
    assert updated_evidence["verifier_recommended_anchor_target_count"] == 1
    assert updated_evidence["verifier_recommended_anchor_targets"][0]["category"] == "heading_fused_with_body"
    assert updated_evidence["verifier_recommended_anchor_targets"] != [runtime_anchor_target]
    assert persisted_review["recommended_anchor_targets"] == [runtime_anchor_target]
    assert persisted_review["verifier_recommended_anchor_target_count"] == 1
    assert persisted_evidence["recommended_anchor_targets"] == [runtime_anchor_target]
    assert persisted_evidence["verifier_recommended_anchor_target_count"] == 1


def test_main_comparison_only_reader_verifier_failure_is_non_blocking(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()
    from docxaicorrector.core.models import ParagraphUnit

    source_path = tmp_path / "Rethinking-money-chapter-region-pages-10-11-and-156-217.pdf"
    source_path.write_bytes(b"%PDF-1.4 chapter-region")
    ui_results_dir = tmp_path / ".run" / "ui_results"
    ui_results_dir.mkdir(parents=True, exist_ok=True)
    cleaned_markdown_path = ui_results_dir / "chapter_region.result.md"
    cleaned_docx_path = ui_results_dir / "chapter_region.result.docx"
    raw_markdown_path = ui_results_dir / "chapter_region.raw.result.md"
    cleanup_report_path = ui_results_dir / "chapter_region.reader_cleanup_report.json"

    cleaned_markdown_path.write_text("Cleaned markdown", encoding="utf-8")
    cleaned_docx_path.write_bytes(_docx_bytes(Document()))
    raw_markdown_path.write_text("Header\n\nBody paragraph", encoding="utf-8")
    cleanup_report_path.write_text(
        json.dumps(
            {
                "stage_status": "completed",
                "changed": True,
                "stats": {
                    "proposed_delete_block_count": 1,
                    "accepted_delete_block_count": 1,
                    "ignored_delete_block_count": 0,
                    "failed_chunk_count": 0,
                    "cleanup_chunk_count": 1,
                },
                "accepted_delete_blocks": [
                    {
                        "id": "b_000000",
                        "text_hash": "abc",
                        "reason": "repeated_running_header",
                        "confidence": "high",
                        "raw_text_preview": "Header",
                        "char_count": 6,
                        "kind": "heading",
                    }
                ],
                "warnings": [],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    def _resolution_payload(**values):
        return SimpleNamespace(**values, to_dict=lambda: dict(values))

    document_profile = SimpleNamespace(
        id="lietaer-pdf-chapter-region-core",
        artifact_prefix="lietaer_pdf_chapter_region",
        output_basename="Rethinking_money_chapter_region",
        max_unmapped_source_paragraphs=12,
        max_unmapped_target_paragraphs=6,
        require_no_toc_body_concat=True,
        resolved_source_path=lambda project_root=None: source_path,
    )
    run_profile = SimpleNamespace(
        id="ui-parity-translate-simple-reader-cleanup-comparison-only",
        tier="full",
        repeat_count=1,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=False,
        model="gpt-5.4",
        max_retries=1,
        comparison_only_validation=True,
        reader_verifier_enabled=True,
    )
    registry = SimpleNamespace(
        get_document_profile=lambda profile_id: document_profile,
        resolve_run_profile=lambda profile, requested_run_profile_id: run_profile,
    )
    prepared_paragraphs = [
        ParagraphUnit(text="Chapter 8", role="heading", structural_role="heading", heading_level=1, source_index=0)
    ]

    class _ValidationServiceStub:
        def __init__(self, log_event_fn):
            self.log_event_fn = log_event_fn

        def run_prepared_background_document(self, **kwargs):
            self.log_event_fn(
                20,
                "ui_result_artifacts_saved",
                "Сохранены итоговые UI-артефакты обработки.",
                filename="chapter-region.pdf",
                artifact_paths={
                    "markdown_path": str(cleaned_markdown_path),
                    "docx_path": str(cleaned_docx_path),
                    "reader_cleanup_raw_markdown_path": str(raw_markdown_path),
                    "reader_cleanup_report_path": str(cleanup_report_path),
                },
            )
            kwargs["runtime"].emit(
                validation.SetStateEvent(
                    values={
                        "latest_markdown": "Cleaned markdown",
                        "latest_docx_bytes": _docx_bytes(Document()),
                    }
                )
            )
            return "succeeded", SimpleNamespace(
                uploaded_filename="chapter-region.pdf",
                uploaded_file_bytes=b"%PDF-1.4 chapter-region",
                paragraphs=prepared_paragraphs,
                image_assets=[],
                jobs=[{"job_kind": "block"}],
                source_text="Source chapter region text.",
                preparation_cached=False,
                preparation_elapsed_seconds=0.1,
                ai_classified_count=0,
                ai_heading_count=0,
                ai_role_change_count=0,
                ai_heading_promotion_count=0,
                ai_heading_demotion_count=0,
                ai_structural_role_change_count=0,
                structure_repair_report=None,
                uploaded_file_token="token-1",
            )

    monkeypatch.setattr(validation, "PROJECT_ROOT", tmp_path)
    monkeypatch.setattr(validation, "REAL_DOCUMENT_ARTIFACT_ROOT", tmp_path / "artifacts")
    monkeypatch.setattr(validation, "load_validation_registry", lambda: registry)
    monkeypatch.setattr(validation, "load_app_config", lambda: object())
    monkeypatch.setattr(
        validation,
        "resolve_runtime_resolution",
        lambda app_config, run_profile: SimpleNamespace(
            effective=_resolution_payload(
                chunk_size=6000,
                image_mode="safe",
                keep_all_image_variants=False,
                model="gpt-5.4",
                max_retries=1,
                processing_operation="translate",
                source_language="auto",
                target_language="ru",
            ),
            ui_defaults=_resolution_payload(
                chunk_size=6000,
                image_mode="safe",
                keep_all_image_variants=False,
                model="gpt-5.4",
                max_retries=1,
                processing_operation="translate",
                source_language="auto",
                target_language="ru",
            ),
            overrides={
                "reader_cleanup_enabled": True,
                "reader_cleanup_policy": "advisory",
                "reader_verifier_enabled": True,
                "structure_recognition_mode": "off",
            },
        ),
    )
    monkeypatch.setattr(
        validation,
        "apply_runtime_resolution_to_app_config",
        lambda app_config, resolution: {"reader_verifier_enabled": True},
    )
    monkeypatch.setattr(
        validation,
        "evaluate_lietaer_acceptance",
        lambda report, **kwargs: {"passed": False, "failed_checks": ["false_fragment_headings_present"], "checks": []},
    )
    monkeypatch.setattr(validation, "_snapshot_formatting_diagnostics_paths", lambda: set())
    monkeypatch.setattr(validation, "_collect_new_formatting_diagnostics_paths", lambda before, after: [])
    monkeypatch.setattr(validation, "_extract_run_formatting_diagnostics_paths", lambda event_log: [])
    monkeypatch.setattr(validation, "_load_recent_formatting_diagnostics", lambda started_at: ([], []))
    monkeypatch.setattr(validation, "_load_formatting_diagnostics_payloads", lambda paths: [])
    monkeypatch.setattr(validation, "_load_translation_quality_report", lambda event_log: (None, None))
    monkeypatch.setattr(validation, "_print_terminal_completion_summary", lambda **kwargs: None)
    monkeypatch.setattr(
        validation.processing_service,
        "clone_processing_service",
        lambda **kwargs: _ValidationServiceStub(kwargs["log_event_fn"]),
    )
    selector_payload = SimpleNamespace(
        canonical_selector="openrouter:google/gemini-3-flash-preview",
        provider="openrouter",
        model_id="google/gemini-3-flash-preview",
    )
    monkeypatch.setattr(
        validation,
        "describe_provider_availability",
        lambda selector, app_config=None: SimpleNamespace(
            enabled=True,
            has_api_key=True,
            error_message=None,
            selector=selector_payload,
        ),
    )
    monkeypatch.setattr(validation, "resolve_model_selector", lambda *args, **kwargs: selector_payload)
    monkeypatch.setattr(validation, "get_client_for_model_selector", lambda *args, **kwargs: object())
    monkeypatch.setattr(validation, "generate_markdown_block", lambda **kwargs: "not json at all")

    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_PROFILE", "lietaer-pdf-chapter-region-core")
    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_RUN_PROFILE", "ui-parity-translate-simple-reader-cleanup-comparison-only")
    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_FORCED_RUN_ID", "run-verifier-failed")

    validation.main()

    run_dir = tmp_path / "artifacts" / "runs" / "run-verifier-failed"
    report_path = run_dir / "lietaer_pdf_chapter_region_report.json"
    review_json_path = run_dir / "lietaer_pdf_chapter_region_reader_quality_review.json"
    latest_manifest_path = tmp_path / "artifacts" / "lietaer_pdf_chapter_region_latest.json"

    report = json.loads(report_path.read_text(encoding="utf-8"))
    review = json.loads(review_json_path.read_text(encoding="utf-8"))
    latest_manifest = json.loads(latest_manifest_path.read_text(encoding="utf-8"))

    assert report["acceptance"]["passed"] is False
    assert report["reader_verifier_evidence"]["verifier_status"] == "failed"
    assert report["reader_verifier_evidence"]["verifier_reason"] == "execution_failed"
    assert report["reader_verifier_evidence"]["overall_verdict"] == "unclear"
    assert review["verifier_status"] == "failed"
    assert review["overall_verdict"] == "unclear"
    assert latest_manifest["status"] == "completed"
    assert latest_manifest["reader_verifier_status"] == "failed"


def test_main_falls_back_to_prepared_snapshot_statuses_when_event_log_lacks_structure_outcome(tmp_path, monkeypatch):
    validation = _load_validation_module()

    source_path = tmp_path / "sample.pdf"
    source_path.write_bytes(b"%PDF-1.4 sample")

    def _resolution_payload(**values):
        return SimpleNamespace(**values, to_dict=lambda: dict(values))

    document_profile = SimpleNamespace(
        id="stub-structural-profile",
        artifact_prefix="stub_structural_validation",
        output_basename="Stub_structural_validated",
        max_unmapped_source_paragraphs=0,
        max_unmapped_target_paragraphs=0,
        require_no_toc_body_concat=True,
        resolved_source_path=lambda project_root=None: source_path,
    )
    run_profile = SimpleNamespace(
        id="ui-parity-pdf-structural-recovery",
        tier="full",
        repeat_count=1,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=False,
        model="gpt-5.4",
        max_retries=1,
    )
    registry = SimpleNamespace(
        get_document_profile=lambda profile_id: document_profile,
        resolve_run_profile=lambda profile, requested_run_profile_id: run_profile,
    )

    prepared_paragraphs = [
        ParagraphUnit(text="Contents", role="body", structural_role="toc_header", source_index=0, paragraph_id="p0000"),
        ParagraphUnit(text="Chapter 1........ 12", role="body", structural_role="toc_entry", source_index=1, paragraph_id="p0001"),
        ParagraphUnit(text="Introduction", role="heading", structural_role="body", heading_level=1, source_index=2, paragraph_id="p0002"),
    ]

    class _ValidationServiceStub:
        def __init__(self, log_event_fn):
            self.log_event_fn = log_event_fn

        def run_prepared_background_document(self, **kwargs):
            kwargs["runtime"].emit(
                validation.SetStateEvent(
                    values={
                        "latest_markdown": "validated output",
                        "latest_docx_bytes": _docx_bytes(Document()),
                        "image_assets": [],
                    }
                )
            )
            return "succeeded", SimpleNamespace(
                uploaded_filename="prepared.docx",
                uploaded_file_bytes=b"PK\x03\x04normalized-source",
                paragraphs=prepared_paragraphs,
                image_assets=[],
                jobs=[{"job_kind": "block"}],
                source_text="source text",
                preparation_cached=False,
                preparation_elapsed_seconds=0.1,
                ai_classified_count=0,
                ai_heading_count=0,
                ai_role_change_count=0,
                ai_heading_promotion_count=0,
                ai_heading_demotion_count=0,
                ai_structural_role_change_count=0,
                structure_repair_report=SimpleNamespace(
                    repaired_bullet_items=0,
                    repaired_numbered_items=0,
                    bounded_toc_regions=1,
                    toc_body_boundary_repairs=1,
                    heading_candidates_from_toc=0,
                    remaining_isolated_marker_count=0,
                ),
                uploaded_file_token="token-1",
                quality_gate_status="pass",
                quality_gate_reasons=(),
                structure_validation_report=SimpleNamespace(readiness_status="ready", readiness_reasons=()),
                structure_ai_attempted=False,
            )

    monkeypatch.setattr(validation, "PROJECT_ROOT", tmp_path)
    monkeypatch.setattr(validation, "REAL_DOCUMENT_ARTIFACT_ROOT", tmp_path / "artifacts")
    monkeypatch.setattr(validation, "load_validation_registry", lambda: registry)
    monkeypatch.setattr(validation, "load_app_config", lambda: object())
    monkeypatch.setattr(
        validation,
        "resolve_runtime_resolution",
        lambda app_config, run_profile: SimpleNamespace(
            effective=_resolution_payload(
                chunk_size=6000,
                image_mode="safe",
                keep_all_image_variants=False,
                model="gpt-5.4",
                max_retries=1,
                processing_operation="translate",
                source_language="auto",
                target_language="ru",
                translation_domain="theology",
            ),
            ui_defaults=_resolution_payload(
                chunk_size=6000,
                image_mode="safe",
                keep_all_image_variants=False,
                model="gpt-5.4",
                max_retries=1,
                processing_operation="edit",
                source_language="auto",
                target_language="ru",
                translation_domain="general",
            ),
            overrides={},
        ),
    )
    monkeypatch.setattr(validation, "apply_runtime_resolution_to_app_config", lambda app_config, resolution: {"x": 1})
    monkeypatch.setattr(validation, "evaluate_lietaer_acceptance", lambda report, **kwargs: {"passed": True, "failed_checks": [], "checks": []})
    monkeypatch.setattr(validation, "_snapshot_formatting_diagnostics_paths", lambda: set())
    monkeypatch.setattr(validation, "_collect_new_formatting_diagnostics_paths", lambda before, after: [])
    monkeypatch.setattr(validation, "_extract_run_formatting_diagnostics_paths", lambda event_log: [])
    monkeypatch.setattr(validation, "_load_recent_formatting_diagnostics", lambda started_at: ([], []))
    monkeypatch.setattr(validation, "_load_formatting_diagnostics_payloads", lambda paths: [])
    monkeypatch.setattr(validation, "_print_terminal_completion_summary", lambda **kwargs: None)
    monkeypatch.setattr(validation.processing_service, "clone_processing_service", lambda **kwargs: _ValidationServiceStub(kwargs["log_event_fn"]))

    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_PROFILE", "stub-structural-profile")
    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_RUN_PROFILE", "ui-parity-pdf-structural-recovery")
    monkeypatch.delenv("DOCXAI_REAL_DOCUMENT_REPEAT_COUNT_OVERRIDE", raising=False)
    monkeypatch.setenv("DOCXAI_REAL_DOCUMENT_FORCED_RUN_ID", "run-structured-fallback")

    validation.main()

    report_path = tmp_path / "artifacts" / "runs" / "run-structured-fallback" / "stub_structural_validation_report.json"
    report = json.loads(report_path.read_text(encoding="utf-8"))

    assert report["preparation_diagnostic_snapshot"]["quality_gate_status"] == "pass"
    assert report["preparation_diagnostic_snapshot"]["readiness_status"] == "ready"
    assert report["output_artifacts"]["output_docx_openable"] is True
    assert report["output_artifacts"]["docx_path"] is not None


def test_classify_failure_detects_heading_only_output_from_block_rejection_event() -> None:
    validation = _load_validation_module()

    failure_classification = validation.classify_failure(
        {
            "result": "failed",
            "last_error": "",
            "event_log": [
                {
                    "event_id": "block_rejected",
                    "context": {
                        "output_classification": "heading_only_output",
                        "message": "Модель вернула только заголовок при наличии основного текста во входном блоке.",
                    },
                }
            ],
        }
    )

    assert failure_classification == "heading_only_output"


def test_extract_run_formatting_diagnostics_paths_prefers_current_run_artifacts() -> None:
    validation = _load_validation_module()

    event_log = [
        {
            "event_id": "formatting_diagnostics_artifacts_detected",
            "context": {"artifact_paths": [".run\\formatting_diagnostics\\older.json"]},
        },
        {
            "event_id": "formatting_diagnostics_artifacts_detected",
            "context": {
                "artifact_paths": [
                    ".run\\formatting_diagnostics\\normalize_current.json",
                    ".run\\formatting_diagnostics\\preserve_current.json",
                ]
            },
        },
    ]

    assert validation._extract_run_formatting_diagnostics_paths(event_log) == [
        ".run\\formatting_diagnostics\\normalize_current.json",
        ".run\\formatting_diagnostics\\preserve_current.json",
    ]


def test_collect_new_formatting_diagnostics_paths_returns_only_new_files_sorted(tmp_path) -> None:
    validation = _load_validation_module()

    old_path = tmp_path / "old.json"
    new_first = tmp_path / "new_first.json"
    new_second = tmp_path / "new_second.json"
    old_path.write_text("{}", encoding="utf-8")
    new_first.write_text("{}", encoding="utf-8")
    new_second.write_text("{}", encoding="utf-8")

    before = {str(old_path.resolve())}
    after = {str(old_path.resolve()), str(new_second.resolve()), str(new_first.resolve())}

    assert validation._collect_new_formatting_diagnostics_paths(before, after) == [
        str(new_first),
        str(new_second),
    ]


def test_build_environment_snapshot_reports_workspace_runtime(monkeypatch) -> None:
    validation = _load_validation_module()
    monkeypatch.setenv("PYTHONPATH", "src:.")
    monkeypatch.setenv("VIRTUAL_ENV", "/tmp/docxai/.venv")

    snapshot = validation._build_environment_snapshot()

    assert snapshot["project_root"] == "."
    assert isinstance(snapshot["python_executable"], str)
    assert snapshot["python_version"]
    assert snapshot["pythonpath"] == "src:."
    assert snapshot["virtual_env"] == "/tmp/docxai/.venv"
    assert "is_wsl" in snapshot


def test_format_terminal_progress_line_includes_phase_progress_and_metrics() -> None:
    validation = _load_validation_module()

    line = validation._format_terminal_progress_line(
        event_type="status",
        phase="process",
        stage="Ожидание ответа OpenAI",
        detail="Блок отправлен в модель.",
        progress=0.5,
        elapsed_seconds=12.34,
        metrics={"current_block": 3, "block_count": 6, "output_ratio": 0.91},
    )

    assert "[status]" in line
    assert "[process]" in line
    assert "50.0%" in line
    assert "current_block=3" in line
    assert "block_count=6" in line
    assert "output_ratio=0.91" in line


def test_normalize_terminal_detail_removes_nested_heartbeat_prefixes() -> None:
    validation = _load_validation_module()

    assert validation._normalize_terminal_detail("Heartbeat: Heartbeat: Блок 4 отправлен") == "Блок 4 отправлен"


def test_print_terminal_completion_summary_is_concise() -> None:
    validation = _load_validation_module()
    buffer = io.StringIO()

    with redirect_stdout(buffer):
        validation._print_terminal_completion_summary(
            final_status="failed",
            report={
                "result": "succeeded",
                "progress_path": "tests/artifacts/real_document_pipeline/runs/run-1/lietaer_validation_progress.json",
                "run": {"run_id": "run-1"},
                "output_artifacts": {
                    "report_json": "tests/artifacts/real_document_pipeline/runs/run-1/lietaer_validation_report.json",
                    "summary_txt": "tests/artifacts/real_document_pipeline/runs/run-1/lietaer_validation_summary.txt",
                },
                "acceptance": {
                    "passed": False,
                    "failed_checks": ["centered_short_paragraphs_preserved"],
                },
                "translation_quality_report": {
                    "quality_status": "fail",
                    "gate_reasons": ["false_fragment_headings_present"],
                    "page_placeholder_heading_concat_count": 0,
                    "page_placeholder_heading_concat_source": "legacy_markdown",
                    "page_placeholder_heading_concat_classification": "display_hygiene",
                    "raw_page_placeholder_heading_concat_count": 1,
                    "false_fragment_heading_count": 0,
                    "false_fragment_heading_gate_source": "entry_assembly",
                    "raw_false_fragment_heading_count": 2,
                    "residual_bullet_glyph_count": 1,
                    "residual_bullet_glyph_gate_source": "legacy_markdown",
                    "raw_residual_bullet_glyph_count": 1,
                    "list_fragment_regression_count": 0,
                    "list_fragment_regression_gate_source": "topology_projection",
                    "raw_list_fragment_regression_count": 1,
                },
            },
        )

    output = buffer.getvalue()

    assert "[summary]" in output
    assert "[artifacts]" in output
    assert "[translation_quality]" in output
    assert "page_placeholder_heading_concat_source=legacy_markdown" in output
    assert "false_fragment_heading_gate_source=entry_assembly" in output
    assert "residual_bullet_glyph_gate_source=legacy_markdown" in output
    assert "raw_list_fragment_regression_count=1" in output
    assert "[acceptance] failed_checks=centered_short_paragraphs_preserved" in output
    assert "latest_docx_bytes" not in output
    assert "processed_block_markdowns" not in output


def test_safe_terminal_print_disables_output_after_broken_pipe(monkeypatch) -> None:
    validation = _load_validation_module()
    calls: list[tuple[object, ...]] = []

    def broken_print(*args, **kwargs):
        calls.append(args)
        raise BrokenPipeError()

    monkeypatch.setattr(validation, "print", broken_print, raising=False)

    validation._safe_terminal_print("first", flush=True)
    validation._safe_terminal_print("second", flush=True)

    assert calls == [("first",)]
    assert validation._TERMINAL_OUTPUT_DISABLED is True


def test_validation_progress_tracker_writes_progress_and_manifest(tmp_path) -> None:
    validation = _load_validation_module()

    run_dir = tmp_path / "runs" / "run-1"
    artifact_root = tmp_path
    progress_path = run_dir / "lietaer_validation_progress.json"
    latest_progress_path = artifact_root / "lietaer_validation_progress.json"
    latest_manifest_path = artifact_root / "lietaer_validation_latest.json"
    report_path = run_dir / "lietaer_validation_report.json"
    summary_path = run_dir / "lietaer_validation_summary.txt"
    markdown_path = run_dir / "validated.md"
    docx_path = run_dir / "validated.docx"
    latest_report_path = artifact_root / "lietaer_validation_report.json"
    latest_summary_path = artifact_root / "lietaer_validation_summary.txt"
    latest_markdown_path = artifact_root / "validated.md"
    latest_docx_path = artifact_root / "validated.docx"

    tracker = validation.ValidationProgressTracker(
        run_id="run-1",
        source_path=tmp_path / "source.docx",
        run_dir=run_dir,
        artifact_root=artifact_root,
        progress_path=progress_path,
        latest_progress_path=latest_progress_path,
        latest_manifest_path=latest_manifest_path,
        report_path=report_path,
        summary_path=summary_path,
        markdown_path=markdown_path,
        docx_path=docx_path,
        latest_report_path=latest_report_path,
        latest_summary_path=latest_summary_path,
        latest_markdown_path=latest_markdown_path,
        latest_docx_path=latest_docx_path,
        started_at_utc=validation.datetime.now(validation.UTC),
    )

    tracker.emit(
        event_type="prepare",
        phase="prepare",
        stage="Разбор документа",
        detail="Построено 12 paragraph units.",
        progress=0.2,
        metrics={"job_count": 4},
        print_line=False,
    )
    tracker.finalize(
        status="completed",
        result="succeeded",
        acceptance_passed=True,
        failure_classification=None,
        last_error="",
        detail="Acceptance=passed",
    )

    progress_payload = json.loads(progress_path.read_text(encoding="utf-8"))
    latest_progress_payload = json.loads(latest_progress_path.read_text(encoding="utf-8"))
    latest_manifest = json.loads(latest_manifest_path.read_text(encoding="utf-8"))

    assert progress_payload["status"] == "completed"
    assert progress_payload["phase"] == "completed"
    assert progress_payload["acceptance_passed"] is True
    assert progress_payload["metrics"] == {"job_count": 4}
    assert latest_progress_payload["run_id"] == "run-1"
    assert latest_progress_payload["status"] == "completed"
    assert latest_manifest["run_id"] == "run-1"
    assert latest_manifest["status"] == "completed"
    assert latest_manifest["progress_json"] == str(progress_path).replace("\\", "/")


def test_write_json_atomic_uses_run_dir_temp_for_latest_alias_retry(tmp_path, monkeypatch) -> None:
    validation = _load_validation_module()

    run_dir = tmp_path / "runs" / "run-1"
    artifact_root = tmp_path / "artifacts"
    latest_progress_path = artifact_root / "lietaer_validation_progress.json"
    original_replace = Path.replace
    replace_calls: list[tuple[Path, Path]] = []

    def flaky_replace(self: Path, target: Path) -> Path:
        replace_calls.append((self, target))
        if len(replace_calls) == 1:
            raise PermissionError("locked progress alias")
        return original_replace(self, target)

    monkeypatch.setattr(validation.time, "sleep", lambda _seconds: None)
    monkeypatch.setattr(Path, "replace", flaky_replace)

    validation._write_json_atomic(latest_progress_path, {"status": "in_progress"}, temp_dir=run_dir)

    assert len(replace_calls) == 2
    assert replace_calls[0][0].parent == run_dir
    assert replace_calls[0][1] == latest_progress_path
    assert json.loads(latest_progress_path.read_text(encoding="utf-8")) == {"status": "in_progress"}
    assert not list(artifact_root.glob("*.tmp"))
