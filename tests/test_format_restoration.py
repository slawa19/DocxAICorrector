import base64
import json
import os
from io import BytesIO
from pathlib import Path
from typing import cast

import config
import formatting_diagnostics_retention
import formatting_transfer
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from document import build_document_text, extract_document_content_from_docx
from formatting_transfer import (
    _build_output_formatting_diagnostics,
    _map_source_target_paragraphs,
    _apply_minimal_caption_formatting,
    _apply_minimal_image_formatting,
    normalize_semantic_output_docx,
    preserve_source_paragraph_properties,
    restore_source_formatting,
)
from models import ParagraphUnit


PNG_BYTES = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aK3cAAAAASUVORK5CYII=")


def _set_raw_paragraph_alignment(paragraph, value: str) -> None:
    paragraph_properties = paragraph._element.get_or_add_pPr()
    alignment = paragraph_properties.find(qn("w:jc"))
    if alignment is None:
        alignment = OxmlElement("w:jc")
        paragraph_properties.append(alignment)
    alignment.set(qn("w:val"), value)


def _get_alignment_xml_value(paragraph) -> str | None:
    paragraph_properties = paragraph._element.pPr
    if paragraph_properties is None:
        return None
    alignment = paragraph_properties.find(qn("w:jc"))
    if alignment is None:
        return None
    return alignment.get(qn("w:val"))


def _extract_numbering_ids(paragraph) -> tuple[str | None, str | None]:
    paragraph_properties = paragraph._element.pPr
    if paragraph_properties is None:
        return None, None
    num_pr = paragraph_properties.find(qn("w:numPr"))
    if num_pr is None:
        return None, None
    ilvl = num_pr.find(qn("w:ilvl"))
    num_id = num_pr.find(qn("w:numId"))
    return (
        None if ilvl is None else ilvl.get(qn("w:val")),
        None if num_id is None else num_id.get(qn("w:val")),
    )


def _numbering_root_contains_num_id(document, num_id: str) -> bool:
    numbering_root = document.part.numbering_part.element
    for child in numbering_root:
        if child.tag == qn("w:num") and child.get(qn("w:numId")) == num_id:
            abstract_num = child.find(qn("w:abstractNumId"))
            if abstract_num is None:
                return False
            abstract_num_id = abstract_num.get(qn("w:val"))
            return any(
                candidate.tag == qn("w:abstractNum") and candidate.get(qn("w:abstractNumId")) == abstract_num_id
                for candidate in numbering_root
            )
    return False


def _append_decimal_numbering_definition(document, *, num_id: str, abstract_num_id: str) -> None:
    numbering_root = document.part.numbering_part.element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_num_id)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    abstract_num.append(lvl)
    numbering_root.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_num_ref = OxmlElement("w:abstractNumId")
    abstract_num_ref.set(qn("w:val"), abstract_num_id)
    num.append(abstract_num_ref)
    numbering_root.append(num)


def _attach_numbering(paragraph, *, num_id: str, ilvl: str) -> None:
    paragraph_properties = paragraph._element.get_or_add_pPr()
    num_pr = paragraph_properties.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        paragraph_properties.append(num_pr)

    ilvl_element = num_pr.find(qn("w:ilvl"))
    if ilvl_element is None:
        ilvl_element = OxmlElement("w:ilvl")
        num_pr.append(ilvl_element)
    ilvl_element.set(qn("w:val"), ilvl)

    num_id_element = num_pr.find(qn("w:numId"))
    if num_id_element is None:
        num_id_element = OxmlElement("w:numId")
        num_pr.append(num_id_element)
    num_id_element.set(qn("w:val"), num_id)


def test_restore_source_formatting_preserves_existing_heading_semantics():
    source_doc = Document()
    source_doc.add_paragraph("Что такое богатство?")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("Что такое богатство?", style="Heading 2")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_doc = Document(BytesIO(restore_source_formatting(target_buffer.getvalue(), source_paragraphs)))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Heading 2"


def test_preserve_source_paragraph_properties_restores_epigraph_italics_and_alignment_from_semantics():
    source_paragraphs = [
        ParagraphUnit(
            text="Богатство заключается в свободе желаний.",
            role="body",
            structural_role="epigraph",
            paragraph_id="p0001",
            paragraph_alignment="center",
            is_italic=True,
        ),
        ParagraphUnit(
            text="— Эпиктет",
            role="body",
            structural_role="attribution",
            paragraph_id="p0002",
            paragraph_alignment="center",
            is_italic=True,
        ),
    ]
    generated_registry = [
        {"paragraph_id": "p0001", "text": "Богатство заключается в свободе желаний."},
        {"paragraph_id": "p0002", "text": "— Эпиктет"},
    ]

    target_doc = Document()
    target_doc.add_paragraph("Богатство заключается в свободе желаний.")
    target_doc.add_paragraph("— Эпиктет")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_doc = Document(
        BytesIO(
            restore_source_formatting(
                target_buffer.getvalue(),
                source_paragraphs,
                generated_paragraph_registry=generated_registry,
            )
        )
    )

    assert updated_doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert updated_doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert all(run.italic for run in updated_doc.paragraphs[0].runs if run.text.strip())
    assert all(run.italic for run in updated_doc.paragraphs[1].runs if run.text.strip())

def test_restore_source_formatting_does_not_inject_source_numbering_xml():
    source_doc = Document()
    source_doc.add_paragraph("Первый пункт", style="List Number")
    source_doc.add_paragraph("Второй пункт", style="List Number")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    _append_decimal_numbering_definition(target_doc, num_id="77", abstract_num_id="700")
    first = target_doc.add_paragraph("Первый пункт")
    second = target_doc.add_paragraph("Второй пункт")
    _attach_numbering(first, num_id="77", ilvl="0")
    _attach_numbering(second, num_id="77", ilvl="0")
    first_ilvl, first_num_id = _extract_numbering_ids(first)
    second_ilvl, second_num_id = _extract_numbering_ids(second)
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_doc = Document(BytesIO(restore_source_formatting(target_buffer.getvalue(), source_paragraphs)))
    updated_first_ilvl, updated_first_num_id = _extract_numbering_ids(updated_doc.paragraphs[0])
    updated_second_ilvl, updated_second_num_id = _extract_numbering_ids(updated_doc.paragraphs[1])

    assert updated_first_ilvl == first_ilvl
    assert updated_second_ilvl == second_ilvl
    assert updated_first_num_id == first_num_id
    assert updated_second_num_id == second_num_id
    assert updated_first_num_id is not None
    assert _numbering_root_contains_num_id(updated_doc, updated_first_num_id)


def test_unified_restoration_preserves_heading_styles():
    source_paragraphs = [
        ParagraphUnit(text="Глава", role="heading", heading_level=1),
        ParagraphUnit(text="Раздел", role="heading", heading_level=2),
        ParagraphUnit(text="Краткое описание установки", role="heading", heading_level=2),
        ParagraphUnit(text="Обычный текст", role="body"),
    ]

    target_doc = Document()
    target_doc.add_paragraph(source_paragraphs[0].text, style="Heading 1")
    target_doc.add_paragraph(source_paragraphs[1].text, style="Heading 2")
    target_doc.add_paragraph(source_paragraphs[2].text, style="Heading 2")
    target_doc.add_paragraph(source_paragraphs[3].text)
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_doc = Document(BytesIO(restore_source_formatting(target_buffer.getvalue(), source_paragraphs)))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[2].style is not None
    assert updated_doc.paragraphs[3].style is not None
    assert updated_doc.paragraphs[0].style.name == "Heading 1"
    assert updated_doc.paragraphs[1].style.name == "Heading 2"
    assert updated_doc.paragraphs[2].style.name == "Heading 2"
    assert updated_doc.paragraphs[3].style.name in {"Body Text", "Normal"}


def test_normalize_semantic_output_docx_is_noop():
    document = Document()
    paragraph = document.add_paragraph("Обычный текст")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buffer = BytesIO()
    document.save(buffer)

    docx_bytes = buffer.getvalue()
    assert normalize_semantic_output_docx(docx_bytes, [ParagraphUnit(text="Обычный текст", role="body")]) == docx_bytes


def test_mapping_accepts_split_heading_target_for_merged_source_paragraph():
    source_paragraphs = [
        ParagraphUnit(
            paragraph_id="p0056",
            text=(
                "Миф (и потенциал) индивидуального богатства "
                "До сих пор мы затронули три вещи, которые определяют наше понимание богатства:"
            ),
            role="body",
            structural_role="body",
            role_confidence="heuristic",
        )
    ]
    generated_registry = [
        {
            "paragraph_id": "p0056",
            "text": (
                "### Миф (и потенциал) индивидуального богатства\n"
                "До сих пор мы затронули три вещи, которые определяют наше понимание богатства:"
            ),
        }
    ]

    target_doc = Document()
    target_doc.add_paragraph("Миф (и потенциал) индивидуального богатства", style="Heading 3")
    target_doc.add_paragraph("До сих пор мы затронули три вещи, которые определяют наше понимание богатства:")

    mapping_pairs, diagnostics = _map_source_target_paragraphs(
        source_paragraphs,
        target_doc.paragraphs,
        generated_paragraph_registry=generated_registry,
    )

    assert [(source.paragraph_id, target.text) for source, target in mapping_pairs] == [
        ("p0056", "До сих пор мы затронули три вещи, которые определяют наше понимание богатства:"),
    ]
    assert diagnostics["unmapped_source_ids"] == []
    assert diagnostics["unmapped_target_indexes"] == []
    assert len(list(diagnostics["accepted_split_targets"])) == 1
    accepted_target = diagnostics["accepted_split_targets"][0]
    assert accepted_target["target_index"] == 0
    assert accepted_target["derived_from_source_index"] == 0
    assert accepted_target["kind"] == "split_heading_prefix"
    assert accepted_target["heading_level"] == 3
    assert accepted_target["target_text_preview"] == "Миф (и потенциал) индивидуального богатства"
    assert accepted_target["source_text_preview"].startswith("Миф (и потенциал) индивидуального богатства")


def test_build_output_formatting_diagnostics_uses_real_mapping_instead_of_tail_count_mismatch():
    source_paragraphs = [
        ParagraphUnit(
            paragraph_id="p0056",
            text=(
                "Миф (и потенциал) индивидуального богатства "
                "До сих пор мы затронули три вещи, которые определяют наше понимание богатства:"
            ),
            role="body",
            structural_role="body",
            role_confidence="heuristic",
        )
    ]
    generated_registry = [
        {
            "paragraph_id": "p0056",
            "text": (
                "### Миф (и потенциал) индивидуального богатства\n"
                "До сих пор мы затронули три вещи, которые определяют наше понимание богатства:"
            ),
        }
    ]

    target_doc = Document()
    target_doc.add_paragraph("Миф (и потенциал) индивидуального богатства", style="Heading 3")
    target_doc.add_paragraph("До сих пор мы затронули три вещи, которые определяют наше понимание богатства:")

    diagnostics = _build_output_formatting_diagnostics(
        source_paragraphs,
        list(target_doc.paragraphs),
        document=target_doc,
        generated_paragraph_registry=generated_registry,
    )

    assert diagnostics["source_count"] == 1
    assert diagnostics["target_count"] == 2
    assert diagnostics["mapped_count"] == 1
    assert diagnostics["unmapped_source_ids"] == []
    assert diagnostics["unmapped_target_indexes"] == []
    assert len(cast(list[dict[str, object]], diagnostics["accepted_split_targets"])) == 1


def test_mapping_reports_accepted_merged_sources_in_diagnostics():
    source_paragraphs = [
        ParagraphUnit(
            paragraph_id="p0010",
            text="Это один логический абзац после нормализации.",
            role="body",
            structural_role="body",
            role_confidence="heuristic",
            origin_raw_indexes=[10, 11],
            origin_raw_texts=["Это один логический", "абзац после нормализации."],
            boundary_source="normalized_merge",
            boundary_confidence="high",
        )
    ]

    target_doc = Document()
    target_doc.add_paragraph("Это один логический абзац после нормализации.")

    _, diagnostics = _map_source_target_paragraphs(source_paragraphs, target_doc.paragraphs)

    assert diagnostics["unmapped_source_ids"] == []
    assert diagnostics["unmapped_target_indexes"] == []
    assert diagnostics["accepted_merged_sources_count"] == 1
    assert diagnostics["max_accepted_merged_sources"] == 2
    assert diagnostics["accepted_merged_sources"] == [
        {
            "logical_paragraph_id": "p0010",
            "origin_raw_indexes": [10, 11],
            "accepted_merged_sources_count": 2,
            "dominant_raw_index": 10,
            "kind": "normalized_merge",
            "boundary_confidence": "high",
            "boundary_decision_class": "high",
            "boundary_rationale": None,
            "target_index": 0,
            "target_text_preview": "Это один логический абзац после нормализации.",
            "source_text_preview": "Это один логический абзац после нормализации.",
        }
    ]
    assert diagnostics["high_confidence_merge_count"] == 1
    assert diagnostics["medium_accepted_merge_count"] == 0


def test_formatting_diagnostics_propagate_boundary_rationale_for_merged_source():
    source_paragraphs = [
        ParagraphUnit(
            paragraph_id="p0012",
            text="Это один логический абзац после нормализации.",
            role="body",
            structural_role="body",
            role_confidence="heuristic",
            origin_raw_indexes=[12, 13, 14],
            origin_raw_texts=["Это один", "логический абзац", "после нормализации."],
            boundary_source="normalized_merge",
            boundary_confidence="high",
            boundary_rationale=(
                "same_body_style, compatible_alignment, left_not_terminal, "
                "right_starts_continuation, left_incomplete, combined_sentence_plausible"
            ),
        )
    ]

    target_doc = Document()
    target_doc.add_paragraph("Это один логический абзац после нормализации.")

    diagnostics = _build_output_formatting_diagnostics(source_paragraphs, list(target_doc.paragraphs), document=target_doc)

    source_entry = cast(list[dict[str, object]], diagnostics["source_registry"])[0]
    assert source_entry["paragraph_id"] == "p0012"
    assert source_entry["origin_raw_indexes"] == [12, 13, 14]
    assert source_entry["origin_raw_text_count"] == 3
    assert source_entry["boundary_source"] == "normalized_merge"
    assert source_entry["boundary_confidence"] == "high"
    assert source_entry["boundary_rationale"] == (
        "same_body_style, compatible_alignment, left_not_terminal, "
        "right_starts_continuation, left_incomplete, combined_sentence_plausible"
    )
    assert source_entry["mapped_target_index"] == 0
    assert source_entry["mapping_strategy"] == "exact_text"

    accepted_merged_source = cast(list[dict[str, object]], diagnostics["accepted_merged_sources"])[0]
    assert accepted_merged_source == {
        "logical_paragraph_id": "p0012",
        "origin_raw_indexes": [12, 13, 14],
        "accepted_merged_sources_count": 3,
        "dominant_raw_index": 12,
        "kind": "normalized_merge",
        "boundary_confidence": "high",
        "boundary_decision_class": "high",
        "boundary_rationale": (
            "same_body_style, compatible_alignment, left_not_terminal, "
            "right_starts_continuation, left_incomplete, combined_sentence_plausible"
        ),
        "target_index": 0,
        "target_text_preview": "Это один логический абзац после нормализации.",
        "source_text_preview": "Это один логический абзац после нормализации.",
    }
    assert diagnostics["accepted_merged_sources_count"] == 1
    assert diagnostics["max_accepted_merged_sources"] == 3


def test_mapping_reports_medium_accepted_merged_sources_in_diagnostics():
    source_paragraphs = [
        ParagraphUnit(
            paragraph_id="p0015",
            text="Это важное наблюдение: Следующий шаг требует дополнительной проверки.",
            role="body",
            structural_role="body",
            role_confidence="heuristic",
            origin_raw_indexes=[15, 16],
            origin_raw_texts=["Это важное наблюдение:", "Следующий шаг требует дополнительной проверки."],
            boundary_source="normalized_merge",
            boundary_confidence="medium",
            boundary_rationale="same_body_style, compatible_alignment, left_not_terminal",
        )
    ]

    target_doc = Document()
    target_doc.add_paragraph("Это важное наблюдение: Следующий шаг требует дополнительной проверки.")

    _, diagnostics = _map_source_target_paragraphs(source_paragraphs, target_doc.paragraphs)

    assert diagnostics["high_confidence_merge_count"] == 0
    assert diagnostics["medium_accepted_merge_count"] == 1
    assert diagnostics["accepted_merged_sources"] == [
        {
            "logical_paragraph_id": "p0015",
            "origin_raw_indexes": [15, 16],
            "accepted_merged_sources_count": 2,
            "dominant_raw_index": 15,
            "kind": "normalized_merge",
            "boundary_confidence": "medium",
            "boundary_decision_class": "medium_accepted",
            "boundary_rationale": "same_body_style, compatible_alignment, left_not_terminal",
            "target_index": 0,
            "target_text_preview": "Это важное наблюдение: Следующий шаг требует дополнительной проверки.",
            "source_text_preview": "Это важное наблюдение: Следующий шаг требует дополнительной проверки.",
        }
    ]


def test_formatting_diagnostics_include_relation_metadata_and_registry_membership():
    source_paragraphs = [
        ParagraphUnit(
            paragraph_id="p0000",
            text="[[DOCX_IMAGE_img_001]]",
            role="image",
            structural_role="image",
            asset_id="img_001",
            role_confidence="explicit",
        ),
        ParagraphUnit(
            paragraph_id="p0001",
            text="Рис. 1. Подпись",
            role="caption",
            structural_role="caption",
            attached_to_asset_id="img_001",
            role_confidence="adjacent",
        ),
        ParagraphUnit(
            paragraph_id="p0002",
            text="Богатство заключается в свободе желаний.",
            role="body",
            structural_role="epigraph",
            paragraph_alignment="center",
        ),
        ParagraphUnit(
            paragraph_id="p0003",
            text="— Эпиктет",
            role="body",
            structural_role="attribution",
        ),
    ]

    target_doc = Document()
    target_doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target_doc.add_paragraph("Рис. 1. Подпись")
    target_doc.add_paragraph("Богатство заключается в свободе желаний.")
    target_doc.add_paragraph("— Эпиктет")

    diagnostics = _build_output_formatting_diagnostics(source_paragraphs, list(target_doc.paragraphs), document=target_doc)

    assert diagnostics["relation_count"] == 2
    assert diagnostics["relation_counts"] == {
        "image_caption": 1,
        "epigraph_attribution": 1,
    }
    assert [relation["relation_kind"] for relation in cast(list[dict[str, object]], diagnostics["accepted_relations"])] == [
        "image_caption",
        "epigraph_attribution",
    ]
    source_registry = cast(list[dict[str, object]], diagnostics["source_registry"])
    assert source_registry[0]["relation_ids"] == ["rel_0001"]
    assert source_registry[1]["relation_ids"] == ["rel_0001"]
    assert source_registry[2]["relation_ids"] == ["rel_0002"]
    assert source_registry[3]["relation_ids"] == ["rel_0002"]


def test_formatting_diagnostics_use_effective_relation_config(monkeypatch):
    monkeypatch.setattr(
        config,
        "load_app_config",
        lambda: {
            "relation_normalization_enabled": True,
            "relation_normalization_profile": "phase2_default",
            "relation_normalization_enabled_relation_kinds": ("image_caption",),
            "relation_normalization_save_debug_artifacts": True,
        },
    )
    source_paragraphs = [
        ParagraphUnit(
            paragraph_id="p0000",
            text="[[DOCX_IMAGE_img_001]]",
            role="image",
            structural_role="image",
            asset_id="img_001",
        ),
        ParagraphUnit(
            paragraph_id="p0001",
            text="Рис. 1. Подпись",
            role="caption",
            structural_role="caption",
            attached_to_asset_id="img_001",
        ),
        ParagraphUnit(
            paragraph_id="p0002",
            text="Богатство заключается в свободе желаний.",
            role="body",
            structural_role="epigraph",
            paragraph_alignment="center",
        ),
        ParagraphUnit(
            paragraph_id="p0003",
            text="— Эпиктет",
            role="body",
            structural_role="attribution",
        ),
    ]

    target_doc = Document()
    target_doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target_doc.add_paragraph("Рис. 1. Подпись")
    target_doc.add_paragraph("Богатство заключается в свободе желаний.")
    target_doc.add_paragraph("— Эпиктет")

    diagnostics = _build_output_formatting_diagnostics(source_paragraphs, list(target_doc.paragraphs), document=target_doc)

    assert diagnostics["relation_count"] == 1
    assert diagnostics["relation_counts"] == {"image_caption": 1}
    assert [relation["relation_kind"] for relation in cast(list[dict[str, object]], diagnostics["accepted_relations"])] == [
        "image_caption",
    ]
    source_registry = cast(list[dict[str, object]], diagnostics["source_registry"])
    assert source_registry[0]["relation_ids"] == ["rel_0001"]
    assert source_registry[1]["relation_ids"] == ["rel_0001"]
    assert source_registry[2]["relation_ids"] == []
    assert source_registry[3]["relation_ids"] == []

    rejected_kinds = [
        decision["relation_kind"]
        for decision in cast(list[dict[str, object]], diagnostics.get("relation_decisions", []))
        if decision.get("decision") == "reject"
    ]
    assert rejected_kinds == []


def test_restore_source_formatting_normalizes_split_heading_prefix_to_heading_2():
    source_paragraphs = [
        ParagraphUnit(
            paragraph_id="p0056",
            text=(
                "Миф (и потенциал) индивидуального богатства "
                "До сих пор мы затронули три вещи, которые определяют наше понимание богатства:"
            ),
            role="body",
            structural_role="body",
            role_confidence="heuristic",
        )
    ]
    generated_registry = [
        {
            "paragraph_id": "p0056",
            "text": (
                "### Миф (и потенциал) индивидуального богатства\n"
                "До сих пор мы затронули три вещи, которые определяют наше понимание богатства:"
            ),
        }
    ]

    target_doc = Document()
    target_doc.add_paragraph("Миф (и потенциал) индивидуального богатства", style="Heading 3")
    target_doc.add_paragraph("До сих пор мы затронули три вещи, которые определяют наше понимание богатства:")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_doc = Document(
        BytesIO(
            restore_source_formatting(
                target_buffer.getvalue(),
                source_paragraphs,
                generated_paragraph_registry=generated_registry,
            )
        )
    )

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Heading 2"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name in {"Body Text", "Normal"}


def test_restore_source_formatting_restores_list_numbering_when_target_loses_it():
    source_doc = Document()
    source_doc.add_paragraph("Первый пункт", style="List Number")
    source_doc.add_paragraph("Второй пункт", style="List Number")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("Первый пункт")
    target_doc.add_paragraph("Второй пункт")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_doc = Document(BytesIO(restore_source_formatting(target_buffer.getvalue(), source_paragraphs)))

    first_ilvl, first_num_id = _extract_numbering_ids(updated_doc.paragraphs[0])
    second_ilvl, second_num_id = _extract_numbering_ids(updated_doc.paragraphs[1])

    assert first_ilvl == "0"
    assert second_ilvl == "0"
    assert first_num_id is not None
    assert second_num_id == first_num_id
    assert _numbering_root_contains_num_id(updated_doc, first_num_id)


def test_apply_minimal_image_formatting_centers_only_image_only_paragraphs():
    document = Document()
    image_only = document.add_paragraph("[[DOCX_IMAGE_img_001]]")
    mixed = document.add_paragraph("Текст [[DOCX_IMAGE_img_002]] подпись")

    _apply_minimal_image_formatting(document)

    assert image_only.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert mixed.alignment is None


def test_apply_minimal_caption_formatting_marks_caption_after_image_anchor():
    document = Document()
    document.add_paragraph("[[DOCX_IMAGE_img_001]]")
    caption = document.add_paragraph("Рис. 1. Подпись к изображению")
    source_paragraphs = [
        ParagraphUnit(text="[[DOCX_IMAGE_img_001]]", role="image", paragraph_id="p0000"),
        ParagraphUnit(text="Рис. 1. Подпись к изображению", role="caption", paragraph_id="p0001"),
    ]

    _apply_minimal_caption_formatting(document, source_paragraphs)

    assert caption.style is not None
    assert caption.style.name == "Caption"
    assert caption.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_apply_minimal_caption_formatting_marks_caption_after_table_anchor():
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Данные"
    caption = document.add_paragraph("Table 1. Summary")
    source_paragraphs = [
        ParagraphUnit(text="<table><tbody><tr><td>Данные</td></tr></tbody></table>", role="table", paragraph_id="p0000"),
        ParagraphUnit(text="Table 1. Summary", role="caption", paragraph_id="p0001"),
    ]

    _apply_minimal_caption_formatting(document, source_paragraphs)

    assert caption.style is not None
    assert caption.style.name == "Caption"
    assert caption.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_apply_minimal_caption_formatting_does_not_promote_body_caption_like_text_without_anchor():
    document = Document()
    body = document.add_paragraph("Рис. 5. Это обычный абзац основного текста")
    source_paragraphs = [
        ParagraphUnit(text="Рис. 5. Это обычный абзац основного текста", role="body", paragraph_id="p0000"),
    ]

    _apply_minimal_caption_formatting(document, source_paragraphs)

    assert body.style is None or body.style.name != "Caption"
    assert body.alignment is None


def test_apply_minimal_caption_formatting_keeps_exact_caption_match_gated_by_anchor_context():
    document = Document()
    document.add_paragraph("Вводный абзац")
    body = document.add_paragraph("Table 3. Generated caption text")
    source_paragraphs = [
        ParagraphUnit(text="Table 3. Generated caption text", role="caption", paragraph_id="p0007"),
    ]
    generated_registry = [{"paragraph_id": "p0007", "text": "Table 3. Generated caption text"}]

    _apply_minimal_caption_formatting(
        document,
        source_paragraphs,
        generated_paragraph_registry=generated_registry,
    )

    assert body.style is None or body.style.name != "Caption"
    assert body.alignment is None


def test_preserve_source_paragraph_properties_keeps_existing_heading_semantics_in_target_docx():
    source_paragraphs = [ParagraphUnit(text="Заголовок", role="heading", heading_level=1)]

    target_doc = Document()
    target_doc.add_paragraph("Заголовок", style="Heading 2")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Heading 2"


def test_preserve_source_paragraph_properties_logs_mismatch_warning(monkeypatch):
    source_doc = Document()
    source_paragraph = source_doc.add_paragraph("Абзац")
    source_paragraph.paragraph_format.left_indent = Inches(0.5)
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("Абзац")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    events = []
    monkeypatch.setattr(formatting_transfer, "log_event", lambda level, event, message, **context: events.append((event, context)))

    preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)

    assert len(events) == 1
    event_name, context = events[0]
    assert event_name == "paragraph_count_mismatch_preserve"
    assert context["source_count"] == 1
    assert context["target_count"] == 2
    assert context["mapped_count"] == 1
    assert context["unmapped_source_count"] == 0
    assert context["unmapped_target_count"] == 1
    assert isinstance(context["artifact_path"], str)


def test_preserve_source_paragraph_properties_does_not_replay_raw_xml_on_mismatch():
    source_doc = Document()
    source_paragraph = source_doc.add_paragraph("Абзац")
    source_paragraph.paragraph_format.left_indent = Inches(0.5)
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("Абзац")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))
    first_paragraph_properties = updated_doc.paragraphs[0]._element.pPr
    first_indentation = None if first_paragraph_properties is None else first_paragraph_properties.find(qn("w:ind"))

    assert first_indentation is None


def test_preserve_source_paragraph_properties_artifact_records_caption_heading_conflict(tmp_path, monkeypatch):
    image_path = tmp_path / "artifact_caption_image.png"
    image_path.write_bytes(PNG_BYTES)

    source_doc = Document()
    source_doc.add_paragraph().add_run().add_picture(str(image_path))
    source_caption = source_doc.add_paragraph("Рис. 1. Подпись к изображению")
    source_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target_doc.add_paragraph("Рис. 1. Подпись к изображению", style="Heading 1")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    diagnostics_dir = tmp_path / "formatting_diagnostics"
    monkeypatch.setattr(formatting_transfer, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)

    preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)

    artifacts = sorted(diagnostics_dir.glob("*.json"))
    assert len(artifacts) == 1
    payload = json.loads(artifacts[0].read_text(encoding="utf-8"))
    assert len(payload["caption_heading_conflicts"]) == 1
    assert payload["caption_heading_conflicts"][0]["target_style_name"] == "Heading 1"
    assert payload["caption_heading_conflicts"][0]["target_heading_level"] == 1


def test_preserve_source_paragraph_properties_artifact_records_restored_list_decisions_during_mismatch(tmp_path, monkeypatch):
    source_doc = Document()
    source_doc.add_paragraph("Первый пункт", style="List Number")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)
    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("Первый пункт")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    diagnostics_dir = tmp_path / "formatting_diagnostics"
    monkeypatch.setattr(formatting_transfer, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)

    preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)

    artifacts = sorted(diagnostics_dir.glob("*.json"))
    assert len(artifacts) == 1
    payload = json.loads(artifacts[0].read_text(encoding="utf-8"))
    assert len(payload["list_restoration_decisions"]) == 1
    assert payload["list_restoration_decisions"][0]["action"] == "restored"


def test_prune_formatting_diagnostics_removes_oldest_and_preserves_newest(tmp_path):
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    diagnostics_dir.mkdir()
    timestamps = [1000, 2000, 3000, 4000]
    paths = []
    for timestamp in timestamps:
        path = diagnostics_dir / f"restore_{timestamp}.json"
        path.write_text("{}", encoding="utf-8")
        paths.append(path)

    for offset, path in enumerate(paths, start=1):
        mtime = float(offset)
        path.touch()
        os.utime(path, (mtime, mtime))

    pruned = formatting_diagnostics_retention.prune_formatting_diagnostics(
        diagnostics_dir=diagnostics_dir,
        now_epoch_seconds=10.0,
        max_age_seconds=100,
        max_count=2,
    )

    remaining = sorted(path.name for path in diagnostics_dir.glob("*.json"))
    assert remaining == ["restore_3000.json", "restore_4000.json"]
    assert sorted(Path(path).name for path in pruned) == ["restore_1000.json", "restore_2000.json"]


def test_write_formatting_diagnostics_artifact_prunes_expired_runtime_files_only(tmp_path):
    runtime_dir = tmp_path / ".run" / "formatting_diagnostics"
    tests_dir = tmp_path / "tests" / "artifacts" / "formatting_diagnostics"
    runtime_dir.mkdir(parents=True)
    tests_dir.mkdir(parents=True)
    old_runtime = runtime_dir / "restore_old.json"
    old_runtime.write_text("{}", encoding="utf-8")
    test_artifact = tests_dir / "keep.json"
    test_artifact.write_text("{}", encoding="utf-8")

    os.utime(old_runtime, (1.0, 1.0))

    artifact_path = formatting_diagnostics_retention.write_formatting_diagnostics_artifact(
        stage="restore",
        diagnostics={"mapped_count": 1},
        diagnostics_dir=runtime_dir,
        now_epoch_ms=200_000,
    )

    assert artifact_path is not None
    assert sorted(path.name for path in runtime_dir.glob("*.json")) == [Path(artifact_path).name]
    assert test_artifact.exists()


def test_preserve_source_paragraph_properties_applies_minimal_output_formatting():
    source_paragraphs = [
        ParagraphUnit(text="Глава", role="heading", heading_level=1),
        ParagraphUnit(text="[[DOCX_IMAGE_img_001]]", role="image"),
        ParagraphUnit(text="Рис. 1. Подпись", role="caption"),
        ParagraphUnit(text="Обычный текст", role="body"),
    ]
    target_doc = Document()
    target_doc.add_paragraph("Глава")
    target_doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target_doc.add_paragraph("Рис. 1. Подпись")
    target_doc.add_paragraph("Обычный текст")
    table = target_doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "A"
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[2].style is not None
    assert updated_doc.paragraphs[3].style is not None
    assert updated_doc.tables[0].style is not None

    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert updated_doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert updated_doc.paragraphs[2].style.name == "Caption"
    assert updated_doc.paragraphs[3].style.name in {"Body Text", "Normal"}
    assert updated_doc.tables[0].style.name == "Table Grid"


def test_preserve_source_paragraph_properties_applies_partial_transfer_on_semantic_mismatch():
    source_paragraphs = [ParagraphUnit(text="Заголовок", role="heading", heading_level=1)]
    target_doc = Document()
    target_doc.add_paragraph("Заголовок")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_uses_content_heuristics_for_captions_without_mapping():
    source_paragraphs = [ParagraphUnit(text="Рис. 1. Подпись к изображению", role="caption")]
    target_doc = Document()
    target_doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target_doc.add_paragraph("Рисунок 1 Подпись к изображению")
    target_doc.add_paragraph("Посторонний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Caption"
    assert updated_doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_preserve_source_paragraph_properties_restores_direct_center_alignment_for_mapped_paragraphs():
    source_paragraphs = [
        ParagraphUnit(text="ГЛАВА 1", role="heading", heading_level=1, paragraph_alignment="center", paragraph_id="p0000"),
        ParagraphUnit(text="Богатство заключается не в том, чтобы иметь много имущества.", role="body", paragraph_alignment="center", paragraph_id="p0001"),
    ]
    target_doc = Document()
    target_doc.add_paragraph("ГЛАВА 1")
    target_doc.add_paragraph("Богатство заключается не в том, чтобы иметь много имущества.")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert updated_doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_preserve_source_paragraph_properties_does_not_promote_generated_registry_text_to_heading():
    source_paragraphs = [ParagraphUnit(text="Старый заголовок", role="heading", heading_level=1, paragraph_id="p0000")]
    target_doc = Document()
    target_doc.add_paragraph("Совершенно новый заголовок")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(
        target_buffer.getvalue(),
        source_paragraphs,
        generated_paragraph_registry=[{"paragraph_id": "p0000", "text": "Совершенно новый заголовок"}],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_does_not_apply_body_formatting_via_generated_registry_similarity():
    source_paragraphs = [ParagraphUnit(text="Исходный абзац сильно отличается", role="body", paragraph_id="p0010")]
    target_doc = Document()
    target_doc.add_paragraph("Лишний абзац перед целью")
    target_doc.add_paragraph("Итоговый литературно отредактированный абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(
        target_buffer.getvalue(),
        source_paragraphs,
        generated_paragraph_registry=[
            {"paragraph_id": "p0010", "text": "Итоговый литературно отредактированный абзац"}
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_leaves_ambiguous_generated_registry_targets_unchanged():
    source_paragraphs = [
        ParagraphUnit(text="Исходный абзац сильно отличается", role="body", paragraph_id="p0011"),
    ]
    target_doc = Document()
    target_doc.add_paragraph("Отредактированный абзац с важной мыслью о богатстве и сообществе сегодня")
    target_doc.add_paragraph("Отредактированный абзац с важной мыслью о богатстве и сообществе завтра")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(
        target_buffer.getvalue(),
        source_paragraphs,
        generated_paragraph_registry=[
            {
                "paragraph_id": "p0011",
                "text": "Отредактированный абзац с важной мыслью о богатстве и сообществе",
            }
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_does_not_use_split_generated_registry_mapping_for_body():
    source_paragraphs = [ParagraphUnit(text="Старый слитый абзац", role="body", paragraph_id="p0056")]
    target_doc = Document()
    target_doc.add_paragraph("Новый заголовок", style="Heading 3")
    target_doc.add_paragraph("Текст после нового заголовка")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(
        target_buffer.getvalue(),
        source_paragraphs,
        generated_paragraph_registry=[
            {"paragraph_id": "p0056", "text": "### Новый заголовок\nТекст после нового заголовка"}
        ],
    )
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Heading 3"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_does_not_restyle_reordered_paragraphs():
    source_paragraphs = [
        ParagraphUnit(text="Заголовок", role="heading", heading_level=1),
        ParagraphUnit(text="Обычный текст", role="body"),
    ]
    target_doc = Document()
    target_doc.add_paragraph("Обычный текст")
    target_doc.add_paragraph("Заголовок")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name in {"Body Text", "Normal"}
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Normal"


def test_preserve_source_paragraph_properties_keeps_existing_numbered_list_semantics_without_injecting_source_numbering():
    source_doc = Document()
    source_doc.add_paragraph("Первый пункт", style="List Number")
    source_doc.add_paragraph("Второй пункт", style="List Number")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)

    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    _append_decimal_numbering_definition(target_doc, num_id="77", abstract_num_id="700")
    first = target_doc.add_paragraph("Первый пункт")
    second = target_doc.add_paragraph("Второй пункт")
    _attach_numbering(first, num_id="77", ilvl="0")
    _attach_numbering(second, num_id="77", ilvl="0")
    original_first = _extract_numbering_ids(first)
    original_second = _extract_numbering_ids(second)
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    first_ilvl, first_num_id = _extract_numbering_ids(updated_doc.paragraphs[0])
    second_ilvl, second_num_id = _extract_numbering_ids(updated_doc.paragraphs[1])

    assert (first_ilvl, first_num_id) == original_first
    assert (second_ilvl, second_num_id) == original_second
    assert first_num_id is not None
    assert _numbering_root_contains_num_id(updated_doc, first_num_id)


def test_preserve_source_paragraph_properties_restores_numbering_for_mapped_plain_target_paragraphs_despite_mismatch():
    source_doc = Document()
    source_doc.add_paragraph("Первый пункт", style="List Number")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)

    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)

    target_doc = Document()
    target_doc.add_paragraph("Первый пункт")
    target_doc.add_paragraph("Лишний абзац")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    ilvl, num_id = _extract_numbering_ids(updated_doc.paragraphs[0])

    assert updated_doc.paragraphs[0].style is not None
    assert updated_doc.paragraphs[0].style.name == "Normal"
    assert ilvl == "0"
    assert num_id is not None
    assert _numbering_root_contains_num_id(updated_doc, num_id)
    assert _extract_numbering_ids(updated_doc.paragraphs[1]) == (None, None)


def test_normalize_semantic_output_docx_remains_noop():
    target_doc = Document()
    paragraph = target_doc.add_paragraph("Обычный текст")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    docx_bytes = target_buffer.getvalue()
    assert normalize_semantic_output_docx(docx_bytes, [ParagraphUnit(text="Обычный текст", role="body")]) == docx_bytes


def test_caption_survives_extraction_markdown_and_preserve_after_image(tmp_path):
    image_path = tmp_path / "docx_caption_pipeline_image.png"
    image_path.write_bytes(PNG_BYTES)

    source_doc = Document()
    source_doc.add_paragraph().add_run().add_picture(str(image_path))
    caption = source_doc.add_paragraph(style="Caption")
    caption.add_run("Рисунок 1 Образец подписи").bold = True
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)

    source_paragraphs, _ = extract_document_content_from_docx(source_buffer)
    markdown = build_document_text(source_paragraphs)

    target_doc = Document()
    target_doc.add_paragraph("[[DOCX_IMAGE_img_001]]")
    target_doc.add_paragraph("Рисунок 1 Образец подписи")
    target_buffer = BytesIO()
    target_doc.save(target_buffer)

    updated_bytes = preserve_source_paragraph_properties(target_buffer.getvalue(), source_paragraphs)
    updated_doc = Document(BytesIO(updated_bytes))

    assert [paragraph.role for paragraph in source_paragraphs] == ["image", "caption"]
    assert markdown == "[[DOCX_IMAGE_img_001]]\n\n**Рисунок 1 Образец подписи**"
    assert updated_doc.paragraphs[1].style is not None
    assert updated_doc.paragraphs[1].style.name == "Caption"
    assert updated_doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
