import queue
import threading
import hashlib
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from typing import Any, cast

import pytest

import docxaicorrector.processing.preparation as preparation
import docxaicorrector.processing.processing_service as processing_service
import docxaicorrector.processing.processing_runtime as processing_runtime
import docxaicorrector.processing.restart_store as restart_store
import docxaicorrector.runtime.artifacts as runtime_artifacts
import docxaicorrector.runtime.state as state
import docxaicorrector.ui.application_flow as application_flow
import docxaicorrector.processing.application_flow as flow_core
from conftest import SessionState as SessionState  # noqa: F811
from docx import Document
from docxaicorrector.document.segments import DocumentContextProfile, DocumentSegment, SegmentBoundaryEvidence, SegmentDetectionReport, SegmentOutlineEntry
from docxaicorrector.runtime.events import FinalizeProcessingStatusEvent, PreparationCompleteEvent, PushActivityEvent, SetProcessingStatusEvent, SetStateEvent


@pytest.fixture(autouse=True)
def _session_state_factory(make_session_state):
    globals()["SessionState"] = make_session_state


class UploadedFileStub:
    def __init__(self, name: str, content: bytes):
        self.name = name
        self._content = content
        self.size = len(content)
        self._cursor = 0

    def getvalue(self):
        return self._content

    def read(self, size: int = -1) -> bytes:
        if size < 0:
            chunk = self._content[self._cursor :]
            self._cursor = len(self._content)
            return chunk
        end = min(len(self._content), self._cursor + size)
        chunk = self._content[self._cursor : end]
        self._cursor = end
        return chunk

    def seek(self, offset: int, whence: int = 0) -> int:
        if whence == 0:
            self._cursor = max(0, offset)
        elif whence == 1:
            self._cursor = max(0, self._cursor + offset)
        elif whence == 2:
            self._cursor = max(0, len(self._content) + offset)
        else:
            raise ValueError(f"Unsupported whence: {whence}")
        self._cursor = min(self._cursor, len(self._content))
        return self._cursor


def _freeze_uploaded_file(name: str, content: bytes):
    return processing_runtime.freeze_uploaded_file(UploadedFileStub(name, content))


def test_prepare_run_context_updates_selected_token_and_prepared_key(monkeypatch):
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    session_state = SessionState(
        selected_source_token="",
        prepared_source_key="",
        completed_source={"filename": "report.docx", "token": "report.docx:3:ba7816bf8f01cfea", "storage_path": "completed.bin"},
    )
    monkeypatch.setattr(state.st, "session_state", session_state)
    logged = []
    progress_events = []

    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=["img"],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="report.docx:hash:6000",
        normalization_report=SimpleNamespace(
            total_raw_paragraphs=3,
            total_logical_paragraphs=2,
            merged_group_count=1,
            merged_raw_paragraph_count=2,
        ),
        cached=False,
    )

    result = application_flow.prepare_run_context(
        uploaded_file=UploadedFileStub("report.docx", b"abc"),
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        session_state=session_state,
        reset_run_state_fn=lambda **kwargs: None,
        fail_critical_fn=lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("unexpected critical error")),
        log_event_fn=lambda *args, **kwargs: logged.append((args, kwargs)),
        prepare_document_for_processing_fn=lambda **kwargs: prepared_document,
        progress_callback=lambda **payload: progress_events.append(payload),
    )

    assert result.uploaded_filename == "report.docx"
    assert result.uploaded_file_bytes == b"abc"
    assert result.uploaded_file_token.startswith("report.docx:3:")
    assert result.jobs == prepared_document.jobs
    assert result.preparation_stage == "Документ подготовлен"
    assert result.preparation_detail == ""
    assert result.preparation_cached is False
    assert result.preparation_elapsed_seconds >= 0.0
    assert result.normalization_report is prepared_document.normalization_report
    assert session_state.selected_source_token == result.uploaded_file_token
    assert session_state.prepared_source_key == "report.docx:hash:6000"
    assert session_state.completed_source is None
    assert len(logged) == 1
    assert progress_events[0]["stage"] == "Чтение файла"
    assert progress_events[1]["metrics"]["file_size_bytes"] == 3
    assert progress_events[-1]["stage"] == "Документ подготовлен"
    assert progress_events[-1]["metrics"]["block_count"] == 1
    assert progress_events[-1]["metrics"]["raw_paragraph_count"] == 3
    assert progress_events[-1]["metrics"]["logical_paragraph_count"] == 2
    assert progress_events[-1]["metrics"]["merged_group_count"] == 1
    assert progress_events[-1]["metrics"]["merged_raw_paragraph_count"] == 2
    assert logged[0][1]["raw_paragraph_count"] == 3
    assert logged[0][1]["logical_paragraph_count"] == 2
    assert logged[0][1]["merged_group_count"] == 1
    assert logged[0][1]["merged_raw_paragraph_count"] == 2


def test_sync_selected_file_context_delegates_selected_token_write_to_state(monkeypatch):
    session_state = SessionState(selected_source_token="")
    monkeypatch.setattr(state.st, "session_state", session_state)
    delegated_tokens = []

    def record_selected_token(token, session_state=None):
        delegated_tokens.append(token)
        if session_state is not None:
            session_state.selected_source_token = token

    monkeypatch.setattr(flow_core, "set_selected_source_token", record_selected_token)

    application_flow.sync_selected_file_context(
        session_state=session_state,
        reset_run_state_fn=lambda **kwargs: None,
        uploaded_file_token="report.docx:3:abc",
    )

    assert delegated_tokens == ["report.docx:3:abc"]
    assert session_state.selected_source_token == "report.docx:3:abc"


def test_prepare_run_context_raises_on_empty_job_target(monkeypatch):
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    failures = []

    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "   ", "target_chars": 0, "context_chars": 0}],
        prepared_source_key="prepared-key",
        cached=False,
    )

    def fail_critical_stub(event, message, **context):
        failures.append((event, message, context))
        raise RuntimeError(message)

    try:
        application_flow.prepare_run_context(
            uploaded_file=UploadedFileStub("report.docx", b"abc"),
            chunk_size=6000,
            image_mode="safe",
            keep_all_image_variants=True,
            session_state=session_state,
            reset_run_state_fn=lambda **kwargs: None,
            fail_critical_fn=fail_critical_stub,
            log_event_fn=lambda *args, **kwargs: None,
            prepare_document_for_processing_fn=lambda **kwargs: prepared_document,
        )
    except RuntimeError:
        pass
    else:
        raise AssertionError("prepare_run_context must fail on empty target_text")

    assert failures[0][0] == "empty_target_block"


def test_prepare_run_context_raises_on_none_job_target(monkeypatch):
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    failures = []

    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": None, "target_chars": 0, "context_chars": 0}],
        prepared_source_key="prepared-key",
        cached=False,
    )

    def fail_critical_stub(event, message, **context):
        failures.append((event, message, context))
        raise RuntimeError(message)

    try:
        application_flow.prepare_run_context(
            uploaded_file=UploadedFileStub("report.docx", b"abc"),
            chunk_size=6000,
            image_mode="safe",
            keep_all_image_variants=True,
            session_state=session_state,
            reset_run_state_fn=lambda **kwargs: None,
            fail_critical_fn=fail_critical_stub,
            log_event_fn=lambda *args, **kwargs: None,
            prepare_document_for_processing_fn=lambda **kwargs: prepared_document,
        )
    except RuntimeError:
        pass
    else:
        raise AssertionError("prepare_run_context must fail on None target_text")

    assert failures[0][0] == "empty_target_block"


def test_prepare_run_context_keeps_best_effort_warning_when_quality_gate_warns(monkeypatch):
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    failures = []

    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="prepared-key",
        quality_gate_status="warning",
        quality_gate_reasons=("first_block_mixed_toc_and_epigraph", "first_block_mixed_toc_and_body_start"),
        cached=False,
        source_format="pdf",
        conversion_backend="libreoffice",
    )

    prepared_run_context = application_flow.prepare_run_context(
        uploaded_file=UploadedFileStub("report.docx", b"abc"),
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        session_state=session_state,
        reset_run_state_fn=lambda **kwargs: None,
        fail_critical_fn=lambda *args, **kwargs: failures.append((args, kwargs)),
        log_event_fn=lambda *args, **kwargs: None,
        prepare_document_for_processing_fn=lambda **kwargs: prepared_document,
    )

    assert failures == []
    assert prepared_run_context.quality_gate_status == "warning"
    assert prepared_run_context.source_format == "pdf"
    assert prepared_run_context.conversion_backend == "libreoffice"
    assert prepared_run_context.preparation_detail.endswith(
        "Причины: первый блок смешивает элементы оглавления и эпиграфа, первый блок смешивает элементы оглавления и начало основного текста"
    )


def test_prepare_run_context_copies_segment_fields_from_prepared_document(monkeypatch):
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    document_context_profile = DocumentContextProfile(
        outline_entries=(SegmentOutlineEntry(segment_id="seg_0001_abcd1234", title="Chapter 1", level=1),),
    )

    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="prepared-key",
        cached=False,
        segments=[SimpleNamespace(segment_id="seg_0001_abcd1234", title="Chapter 1")],
        segment_diagnostics=SimpleNamespace(segment_count=1, toc_matched_count=0),
        structure_fingerprint="abc123def456",
        detector_version="chapter_segments_v1",
        segment_to_job={"seg_0001_abcd1234": (0,)},
        document_context_profile=document_context_profile,
    )

    prepared_run_context = application_flow.prepare_run_context(
        uploaded_file=UploadedFileStub("report.docx", b"abc"),
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        session_state=session_state,
        reset_run_state_fn=lambda **kwargs: None,
        fail_critical_fn=lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("unexpected critical error")),
        log_event_fn=lambda *args, **kwargs: None,
        prepare_document_for_processing_fn=lambda **kwargs: prepared_document,
    )

    assert prepared_run_context.segments == prepared_document.segments
    assert prepared_run_context.segment_diagnostics is prepared_document.segment_diagnostics
    assert prepared_run_context.structure_fingerprint == "abc123def456"
    assert prepared_run_context.detector_version == "chapter_segments_v1"
    assert prepared_run_context.segment_to_job == {"seg_0001_abcd1234": (0,)}
    assert prepared_run_context.document_context_profile == document_context_profile


def test_prepare_run_context_keeps_other_completed_source_tokens(monkeypatch):
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    session_state = SessionState(
        selected_source_token="",
        prepared_source_key="",
        completed_source={"filename": "other.docx", "token": "other.docx:3:def", "storage_path": "other.bin"},
    )

    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=["img"],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="report.docx:hash:6000",
        cached=False,
    )

    application_flow.prepare_run_context(
        uploaded_file=UploadedFileStub("report.docx", b"abc"),
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        session_state=session_state,
        reset_run_state_fn=lambda **kwargs: None,
        fail_critical_fn=lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("unexpected critical error")),
        log_event_fn=lambda *args, **kwargs: None,
        prepare_document_for_processing_fn=lambda **kwargs: prepared_document,
    )

    assert session_state.completed_source == {"filename": "other.docx", "token": "other.docx:3:def", "storage_path": "other.bin"}


def test_get_cached_completed_file_loads_bytes_from_store():
    session_state = SessionState(completed_source={"filename": "report.docx", "token": "report.docx:3:abc", "storage_path": "completed.bin", "size": 3, "payload_sha256": hashlib.sha256(b"abc").hexdigest(), "source_format": "docx", "conversion_backend": None})

    uploaded_file = application_flow.get_cached_completed_file(
        session_state=cast(application_flow.SessionStateLike, session_state),
        load_completed_source_bytes_fn=lambda source: b"abc",
    )

    assert uploaded_file is not None
    assert uploaded_file.filename == "report.docx"
    assert uploaded_file.content_bytes == b"abc"
    assert uploaded_file.file_token == "report.docx:3:abc"


def test_consume_completed_source_if_used_clears_persisted_file(monkeypatch):
    session_state = SessionState(completed_source={"filename": "report.docx", "token": "report.docx:3:abc", "storage_path": "completed.bin"})
    cleared = []
    monkeypatch.setattr(application_flow, "clear_restart_source", lambda source: cleared.append(source))

    application_flow.consume_completed_source_if_used(session_state=session_state, uploaded_file_token="report.docx:3:abc")

    assert session_state.completed_source is None
    assert cleared == [{"filename": "report.docx", "token": "report.docx:3:abc", "storage_path": "completed.bin"}]


def test_prepare_run_context_reports_invalid_archive_via_fail_critical(monkeypatch):
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    validated = []
    failures = []

    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: validated.append(source_bytes) or (_ for _ in ()).throw(RuntimeError("bad archive")))

    def fail_critical_stub(event, message, **context):
        failures.append((event, message, context))
        raise RuntimeError(message)

    with pytest.raises(RuntimeError, match="bad archive"):
        application_flow.prepare_run_context(
            uploaded_file=UploadedFileStub("report.docx", b"abc"),
            chunk_size=6000,
            image_mode="safe",
            keep_all_image_variants=True,
            session_state=session_state,
            reset_run_state_fn=lambda **kwargs: None,
            fail_critical_fn=fail_critical_stub,
            log_event_fn=lambda *args, **kwargs: None,
            prepare_document_for_processing_fn=lambda **kwargs: (_ for _ in ()).throw(AssertionError("prepare_document_for_processing must not run")),
        )

    assert validated == [b"abc"]
    assert failures == [("doc_validation_failed", "bad archive", {"filename": "report.docx"})]


def test_prepare_run_context_reports_broken_relationships_via_fail_critical(monkeypatch):
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    validated = []
    failures = []

    monkeypatch.setattr(
        flow_core,
        "validate_docx_source_bytes",
        lambda source_bytes: validated.append(source_bytes)
        or (_ for _ in ()).throw(RuntimeError("broken relationship target: rId5")),
    )

    def fail_critical_stub(event, message, **context):
        failures.append((event, message, context))
        raise RuntimeError(message)

    with pytest.raises(RuntimeError, match="broken relationship target: rId5"):
        application_flow.prepare_run_context(
            uploaded_file=UploadedFileStub("report.docx", b"abc"),
            chunk_size=6000,
            image_mode="safe",
            keep_all_image_variants=True,
            session_state=session_state,
            reset_run_state_fn=lambda **kwargs: None,
            fail_critical_fn=fail_critical_stub,
            log_event_fn=lambda *args, **kwargs: None,
            prepare_document_for_processing_fn=lambda **kwargs: (_ for _ in ()).throw(AssertionError("prepare_document_for_processing must not run")),
        )

    assert validated == [b"abc"]
    assert failures == [
        ("doc_validation_failed", "broken relationship target: rId5", {"filename": "report.docx"})
    ]


def test_prepare_run_context_reports_suspicious_uncompressed_archive_via_fail_critical(monkeypatch):
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    validated = []
    failures = []

    monkeypatch.setattr(
        flow_core,
        "validate_docx_source_bytes",
        lambda source_bytes: validated.append(source_bytes)
        or (_ for _ in ()).throw(RuntimeError("DOCX-архив слишком велик после распаковки и отклонен из соображений безопасности.")),
    )

    def fail_critical_stub(event, message, **context):
        failures.append((event, message, context))
        raise RuntimeError(message)

    with pytest.raises(RuntimeError, match="слишком велик после распаковки"):
        application_flow.prepare_run_context(
            uploaded_file=UploadedFileStub("report.docx", b"abc"),
            chunk_size=6000,
            image_mode="safe",
            keep_all_image_variants=True,
            session_state=session_state,
            reset_run_state_fn=lambda **kwargs: None,
            fail_critical_fn=fail_critical_stub,
            log_event_fn=lambda *args, **kwargs: None,
            prepare_document_for_processing_fn=lambda **kwargs: (_ for _ in ()).throw(AssertionError("prepare_document_for_processing must not run")),
        )

    assert validated == [b"abc"]
    assert failures == [
        (
            "doc_validation_failed",
            "DOCX-архив слишком велик после распаковки и отклонен из соображений безопасности.",
            {"filename": "report.docx"},
        )
    ]


def test_prepare_run_context_reports_absolute_archive_paths_via_fail_critical(monkeypatch):
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    validated = []
    failures = []

    monkeypatch.setattr(
        flow_core,
        "validate_docx_source_bytes",
        lambda source_bytes: validated.append(source_bytes)
        or (_ for _ in ()).throw(RuntimeError("DOCX-архив содержит абсолютные пути и отклонён из соображений безопасности.")),
    )

    def fail_critical_stub(event, message, **context):
        failures.append((event, message, context))
        raise RuntimeError(message)

    with pytest.raises(RuntimeError, match="абсолютные пути"):
        application_flow.prepare_run_context(
            uploaded_file=UploadedFileStub("report.docx", b"abc"),
            chunk_size=6000,
            image_mode="safe",
            keep_all_image_variants=True,
            session_state=session_state,
            reset_run_state_fn=lambda **kwargs: None,
            fail_critical_fn=fail_critical_stub,
            log_event_fn=lambda *args, **kwargs: None,
            prepare_document_for_processing_fn=lambda **kwargs: (_ for _ in ()).throw(AssertionError("prepare_document_for_processing must not run")),
        )

    assert validated == [b"abc"]
    assert failures == [
        (
            "doc_validation_failed",
            "DOCX-архив содержит абсолютные пути и отклонён из соображений безопасности.",
            {"filename": "report.docx"},
        )
    ]


def test_prepare_run_context_reports_encrypted_or_protected_input_via_fail_critical(monkeypatch):
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    validated = []
    failures = []

    monkeypatch.setattr(
        flow_core,
        "validate_docx_source_bytes",
        lambda source_bytes: validated.append(source_bytes)
        or (_ for _ in ()).throw(RuntimeError("Документ защищён паролем или шифрованием и не может быть обработан.")),
    )

    def fail_critical_stub(event, message, **context):
        failures.append((event, message, context))
        raise RuntimeError(message)

    with pytest.raises(RuntimeError, match="защищён паролем"):
        application_flow.prepare_run_context(
            uploaded_file=UploadedFileStub("report.docx", b"abc"),
            chunk_size=6000,
            image_mode="safe",
            keep_all_image_variants=True,
            session_state=session_state,
            reset_run_state_fn=lambda **kwargs: None,
            fail_critical_fn=fail_critical_stub,
            log_event_fn=lambda *args, **kwargs: None,
            prepare_document_for_processing_fn=lambda **kwargs: (_ for _ in ()).throw(AssertionError("prepare_document_for_processing must not run")),
        )

    assert validated == [b"abc"]
    assert failures == [
        (
            "doc_validation_failed",
            "Документ защищён паролем или шифрованием и не может быть обработан.",
            {"filename": "report.docx"},
        )
    ]


def test_prepare_run_context_normalizes_legacy_doc_before_validation(monkeypatch):
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    validated = []
    received = {}
    freeze_calls = []

    monkeypatch.setattr(
        flow_core,
        "freeze_uploaded_file",
        lambda uploaded_file: freeze_calls.append(uploaded_file.name)
        or SimpleNamespace(filename="legacy.docx", content_bytes=b"converted-docx", file_token="legacy.docx:6:mocked"),
    )
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: validated.append(source_bytes) or None)

    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="legacy.docx:hash:6000",
        cached=False,
    )

    def prepare_document_for_processing_stub(**kwargs):
        received.update(kwargs)
        return prepared_document

    result = application_flow.prepare_run_context(
        uploaded_file=UploadedFileStub("legacy.doc", b"legacy"),
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        session_state=session_state,
        reset_run_state_fn=lambda **kwargs: None,
        fail_critical_fn=lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("unexpected critical error")),
        log_event_fn=lambda *args, **kwargs: None,
        prepare_document_for_processing_fn=prepare_document_for_processing_stub,
    )

    assert freeze_calls == ["legacy.doc"]
    assert validated == [b"converted-docx"]
    assert received["uploaded_payload"].filename == "legacy.docx"
    assert received["uploaded_payload"].content_bytes == b"converted-docx"
    assert received["uploaded_payload"].file_token == "legacy.docx:6:mocked"
    assert result.uploaded_filename == "legacy.docx"
    assert result.uploaded_file_bytes == b"converted-docx"
    assert result.uploaded_file_token == "legacy.docx:6:mocked"


def test_prepare_run_context_normalizes_pdf_before_validation(monkeypatch):
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    validated = []
    received = {}
    freeze_calls = []

    monkeypatch.setattr(
        flow_core,
        "freeze_uploaded_file",
        lambda uploaded_file: freeze_calls.append(uploaded_file.name)
        or SimpleNamespace(filename="source.docx", content_bytes=b"PK\x03\x04converted-docx", file_token="source.docx:16:mocked"),
    )
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: validated.append(source_bytes) or None)

    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="source.docx:hash:6000",
        cached=False,
    )

    def prepare_document_for_processing_stub(**kwargs):
        received.update(kwargs)
        return prepared_document

    result = application_flow.prepare_run_context(
        uploaded_file=UploadedFileStub("source.pdf", b"%PDF-1.7\nsource"),
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        session_state=session_state,
        reset_run_state_fn=lambda **kwargs: None,
        fail_critical_fn=lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("unexpected critical error")),
        log_event_fn=lambda *args, **kwargs: None,
        prepare_document_for_processing_fn=prepare_document_for_processing_stub,
    )

    assert freeze_calls == ["source.pdf"]
    assert validated == [b"PK\x03\x04converted-docx"]
    assert received["uploaded_payload"].filename == "source.docx"
    assert received["uploaded_payload"].content_bytes == b"PK\x03\x04converted-docx"
    assert received["uploaded_payload"].file_token == "source.docx:16:mocked"
    assert result.uploaded_filename == "source.docx"
    assert result.uploaded_file_bytes == b"PK\x03\x04converted-docx"
    assert result.uploaded_file_token == "source.docx:16:mocked"


def test_prepare_run_context_reports_doc_conversion_failure_via_fail_critical(monkeypatch):
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    failures = []

    monkeypatch.setattr(
        flow_core,
        "freeze_uploaded_file",
        lambda uploaded_file: (_ for _ in ()).throw(RuntimeError("converter missing")),
    )

    def fail_critical_stub(event, message, **context):
        failures.append((event, message, context))
        raise RuntimeError(message)

    with pytest.raises(RuntimeError, match="converter missing"):
        application_flow.prepare_run_context(
            uploaded_file=UploadedFileStub("legacy.doc", b"legacy"),
            chunk_size=6000,
            image_mode="safe",
            keep_all_image_variants=True,
            session_state=session_state,
            reset_run_state_fn=lambda **kwargs: None,
            fail_critical_fn=fail_critical_stub,
            log_event_fn=lambda *args, **kwargs: None,
            prepare_document_for_processing_fn=lambda **kwargs: (_ for _ in ()).throw(AssertionError("unexpected")),
        )

    assert failures == [("doc_conversion_failed", "converter missing", {"filename": "legacy.doc"})]


def test_prepare_run_context_reports_pdf_conversion_failure_via_fail_critical(monkeypatch):
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    failures = []

    monkeypatch.setattr(
        flow_core,
        "freeze_uploaded_file",
        lambda uploaded_file: (_ for _ in ()).throw(RuntimeError("pdf converter missing")),
    )

    def fail_critical_stub(event, message, **context):
        failures.append((event, message, context))
        raise RuntimeError(message)

    with pytest.raises(RuntimeError, match="pdf converter missing"):
        application_flow.prepare_run_context(
            uploaded_file=UploadedFileStub("source.pdf", b"%PDF-1.7\nsource"),
            chunk_size=6000,
            image_mode="safe",
            keep_all_image_variants=True,
            session_state=session_state,
            reset_run_state_fn=lambda **kwargs: None,
            fail_critical_fn=fail_critical_stub,
            log_event_fn=lambda *args, **kwargs: None,
            prepare_document_for_processing_fn=lambda **kwargs: (_ for _ in ()).throw(AssertionError("unexpected")),
        )

    assert failures == [("doc_conversion_failed", "pdf converter missing", {"filename": "source.pdf"})]


def test_prepare_run_context_sync_path_freezes_upload_once(monkeypatch):
    session_state = SessionState(selected_source_token="", prepared_source_key="")
    freeze_calls = []
    validate_calls = []

    monkeypatch.setattr(
        flow_core,
        "freeze_uploaded_file",
        lambda uploaded_file: freeze_calls.append(uploaded_file.name)
        or SimpleNamespace(filename="legacy.docx", content_bytes=b"converted-docx", file_token="legacy.docx:token"),
    )
    monkeypatch.setattr(
        flow_core,
        "validate_docx_source_bytes",
        lambda source_bytes: validate_calls.append(source_bytes),
    )

    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="prepared-key",
        cached=False,
    )

    application_flow.prepare_run_context(
        uploaded_file=UploadedFileStub("legacy.doc", b"legacy"),
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        session_state=session_state,
        reset_run_state_fn=lambda **kwargs: None,
        fail_critical_fn=lambda *args, **kwargs: None,
        log_event_fn=lambda *args, **kwargs: None,
        prepare_document_for_processing_fn=lambda **kwargs: prepared_document,
    )

    assert freeze_calls == ["legacy.doc"]
    assert validate_calls == [b"converted-docx"]


def test_restart_flow_restores_uploaded_file_from_run_store_and_cleans_up(tmp_path, monkeypatch):
    session_state = SessionState()
    monkeypatch.setattr(state.st, "session_state", session_state)
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    monkeypatch.setattr(restart_store, "RUN_DIR", tmp_path)

    state.init_session_state()
    session_state.restart_session_id = "session-a"

    processing_runtime.start_background_processing(
        worker_target=lambda **kwargs: None,
        reset_run_state=state.reset_run_state,
        push_activity=state.push_activity,
        set_processing_status=state.set_processing_status,
        uploaded_filename="report.docx",
        uploaded_token="report.docx:3:abc",
        source_bytes=b"abc",
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
    )

    session_state.processing_worker.join(timeout=5)
    session_state.processing_worker = None
    session_state.processing_outcome = "stopped"

    restored_file = application_flow.resolve_effective_uploaded_file(
        uploaded_file=None,
        current_result=None,
        session_state=session_state,
    )

    restart_path = session_state.restart_source["storage_path"]

    assert restored_file is not None
    assert restored_file.filename == "report.docx"
    assert restored_file.content_bytes == b"abc"
    assert restored_file.file_token == "report.docx:3:abc"

    state.reset_run_state(keep_restart_source=False)


def test_sync_selected_file_context_resets_run_state_for_new_file(monkeypatch):
    session_state = SessionState(
        selected_source_token="old.docx:10",
        sidebar_text_operation="Перевод",
        sidebar_source_language="Авто",
        sidebar_target_language="Русский",
        previous_result=None,
        latest_docx_bytes=b"docx",
        latest_source_name="old.docx",
        latest_source_token="old.docx:10",
        latest_markdown="markdown",
        run_log=[{"status": "STOP"}],
        activity_feed=[{"time": "10:00:00", "message": "stale"}],
        processed_block_markdowns=["partial"],
        last_error="",
        image_assets=[],
        image_validation_failures=[],
        image_processing_summary={},
        processing_status={},
        processing_stop_requested=False,
        processing_worker=None,
        processing_event_queue=None,
        processing_stop_event=None,
        processing_outcome="idle",
        prepared_source_key="old.docx:10:6000",
        restart_source={"filename": "old.docx", "storage_path": "old.bin"},
    )
    reset_calls = []
    monkeypatch.setattr(state.st, "session_state", session_state)

    application_flow.sync_selected_file_context(
        session_state=session_state,
        reset_run_state_fn=lambda **kwargs: (reset_calls.append(kwargs), session_state.update(run_log=[], activity_feed=[], restart_source=None)),
        uploaded_file_token="new.docx:20",
    )

    assert reset_calls == [{"keep_restart_source": False}]
    assert session_state.selected_source_token == "new.docx:20"
    assert session_state.restart_source is None
    assert session_state.run_log == []
    assert session_state.activity_feed == []
    assert "sidebar_text_operation" not in session_state
    assert "sidebar_source_language" not in session_state
    assert "sidebar_target_language" not in session_state


def test_has_resettable_state_depends_on_restartable_source(tmp_path):
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    session_state = SessionState(processing_outcome="stopped", restart_source=_valid_restart_record(restart_path))

    assert application_flow.has_resettable_state(current_result=None, session_state=session_state) is True  # type: ignore[arg-type]

    session_state.processing_outcome = "idle"

    assert application_flow.has_resettable_state(current_result=None, session_state=session_state) is False  # type: ignore[arg-type]


def test_derive_idle_view_state_covers_idle_paths(tmp_path):
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    session_state = SessionState(processing_outcome="stopped", restart_source=_valid_restart_record(restart_path))

    assert application_flow.derive_app_idle_view_state(current_result=None, uploaded_file=object(), session_state=session_state) == "file_selected"
    assert application_flow.derive_app_idle_view_state(current_result={"docx_bytes": b"x"}, uploaded_file=None, session_state=session_state) == "completed"
    assert application_flow.derive_app_idle_view_state(current_result=None, uploaded_file=None, session_state=session_state) == "restartable"

    session_state.processing_outcome = "idle"

    assert application_flow.derive_app_idle_view_state(current_result=None, uploaded_file=None, session_state=session_state) == "empty"


def test_get_cached_restart_file_returns_none_when_storage_missing(monkeypatch):
    session_state = SessionState(restart_source={"filename": "report.docx", "storage_path": "missing.bin"})
    monkeypatch.setattr(
        application_flow,
        "load_persisted_source_bytes_with_reason",
        lambda restart_source: (None, "unreadable_payload"),
    )

    assert application_flow.get_cached_restart_file(session_state=session_state) is None  # type: ignore[arg-type]


def test_resolve_effective_uploaded_file_uses_completed_source_after_success():
    session_state = SessionState(
        completed_source={"filename": "report.docx", "storage_path": "completed.bin", "token": "report.docx:3:abc", "size": 3, "payload_sha256": hashlib.sha256(b"abc").hexdigest(), "source_format": "docx", "conversion_backend": None}
    )

    uploaded_file = application_flow.resolve_effective_uploaded_file(
        uploaded_file=None,
        current_result={"docx_bytes": b"done"},
        session_state=session_state,
        load_restart_source_bytes_fn=lambda source: b"abc",
    )

    assert uploaded_file is not None
    assert uploaded_file.filename == "report.docx"
    assert uploaded_file.content_bytes == b"abc"
    assert uploaded_file.file_token == "report.docx:3:abc"


@pytest.mark.parametrize("source_format", ["pdf", "doc"])
def test_cached_normalized_source_restores_authoritative_original_token(source_format):
    source_bytes = b"normalized-docx"
    source_token = f"report.{source_format}:123:original"
    record = {
        "filename": "report.docx",
        "storage_path": "completed.bin",
        "token": source_token,
        "size": len(source_bytes),
        "payload_sha256": hashlib.sha256(source_bytes).hexdigest(),
        "source_format": source_format,
        "conversion_backend": "libreoffice",
    }

    restored = application_flow.get_cached_completed_file(
        session_state=cast(application_flow.SessionStateLike, SessionState(completed_source=record)),
        load_completed_source_bytes_fn=lambda source: source_bytes,
    )

    assert restored is not None
    assert restored.file_token == source_token
    assert restored.source_format == source_format
    assert restored.conversion_backend == "libreoffice"


def test_resolve_preparation_upload_passes_through_frozen_payload_without_reread():
    # spec-045 P0 seam: a restored persisted source arrives as a FrozenUploadPayload.
    # It must be passed through as-is (NOT re-frozen, which would crash because
    # FrozenUploadPayload has no .read/.seek/.getvalue), and its authoritative
    # SOURCE-derived token must survive rather than being recomputed from the
    # normalized DOCX bytes. If the isinstance() branch is reverted this test fails.
    payload = flow_core.FrozenUploadPayload(
        filename="book.pdf",
        content_bytes=b"PK\x03\x04normalized-docx-ish-bytes",
        file_size=len(b"PK\x03\x04normalized-docx-ish-bytes"),
        content_hash="deadbeef",
        file_token="book.pdf:100:deadbeef",
        source_format="pdf",
        conversion_backend="libreoffice",
    )

    resolved = flow_core._resolve_preparation_upload(uploaded_file=payload, uploaded_payload=None)

    assert resolved.needs_read_stage is False
    assert resolved.uploaded_file_token == payload.file_token
    assert resolved.uploaded_file_bytes == payload.content_bytes
    assert resolved.uploaded_filename == payload.filename


def test_fresh_upload_wins_when_persisted_source_is_unverifiable():
    fresh_upload = UploadedFileStub("fresh.pdf", b"fresh")
    session_state = SessionState(
        processing_outcome="stopped",
        restart_source={"filename": "report.docx", "storage_path": "missing.bin"},
    )

    assert application_flow.resolve_effective_uploaded_file(
        uploaded_file=fresh_upload,
        current_result=None,
        session_state=session_state,
        load_restart_source_bytes_fn=lambda source: None,
    ) is fresh_upload


def test_has_restartable_source_does_not_materialize_restart_bytes(tmp_path, monkeypatch):
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    session_state = SessionState(processing_outcome="stopped", restart_source=_valid_restart_record(restart_path))
    load_calls = []
    monkeypatch.setattr(
        application_flow,
        "load_persisted_source_bytes_with_reason",
        lambda restart_source: (load_calls.append(restart_source) or b"abc", None),
    )

    assert application_flow.has_restartable_source(session_state=session_state) is True  # type: ignore[arg-type]
    assert load_calls == []


def test_has_restartable_source_returns_false_when_restart_file_was_removed(tmp_path):
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    session_state = SessionState(processing_outcome="stopped", restart_source=_valid_restart_record(restart_path))
    restart_path.unlink()

    assert application_flow.has_restartable_source(session_state=session_state) is False  # type: ignore[arg-type]


def _valid_restart_record(restart_path, source_bytes: bytes = b"abc") -> dict[str, object]:
    return {
        "filename": "report.docx",
        "token": "report.docx:3:abc",
        "storage_path": str(restart_path),
        "size": len(source_bytes),
        "payload_sha256": hashlib.sha256(source_bytes).hexdigest(),
        "source_format": "docx",
        "conversion_backend": None,
        "storage_kind": "restart",
    }


def test_blocked_result_falls_back_to_restart_source_for_reprocess(tmp_path):
    # Round-11 Fix A: a BLOCKED run is terminal-failed, so completed-source caching
    # (SUCCEEDED-only) never ran, yet it DOES produce a current_result bundle. Before
    # the fix the restart branch was skipped and the blocked view offered only Reset.
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    session_state = SessionState(
        processing_outcome="failed",
        restart_source=_valid_restart_record(restart_path),
    )

    restored = application_flow.resolve_effective_uploaded_file(
        uploaded_file=None,
        current_result={"docx_bytes": None, "delivery_disposition": {"status": "blocked", "explanation": "gate"}},
        session_state=session_state,
        load_restart_source_bytes_fn=lambda source: b"abc",
    )

    assert restored is not None
    assert restored.filename == "report.docx"
    assert restored.file_token == "report.docx:3:abc"
    assert restored.content_bytes == b"abc"


def test_accepted_succeeded_result_never_falls_back_to_restart_source(tmp_path):
    # Anti-regression for Fix A: the fall-through is keyed on restart ELIGIBILITY
    # (a stopped/failed outcome that kept its record), never on byte presence. A
    # SUCCEEDED run — the only outcome that pairs with an accepted delivery in
    # practice — must behave exactly as before and yield no restart file, even if a
    # stale record is still lying around.
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    session_state = SessionState(
        processing_outcome="succeeded",
        restart_source=_valid_restart_record(restart_path),
    )

    assert application_flow.resolve_effective_uploaded_file(
        uploaded_file=None,
        current_result={"docx_bytes": b"done", "delivery_disposition": {"status": "accepted"}},
        session_state=session_state,
        load_restart_source_bytes_fn=lambda source: b"abc",
    ) is None


def test_stopped_after_publication_still_offers_restart_source(tmp_path):
    # Round-11 R1: a stop observed AFTER the result was published leaves an accepted
    # bundle plus a retained restart record (completed-source caching is SUCCEEDED-only).
    # Keying the fall-through on the blocked disposition alone missed this class, so the
    # run rendered as COMPLETED with no reprocess control while valid bytes sat on disk.
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    session_state = SessionState(
        processing_outcome="stopped",
        restart_source=_valid_restart_record(restart_path),
    )

    restored = application_flow.resolve_effective_uploaded_file(
        uploaded_file=None,
        current_result={"docx_bytes": b"done", "delivery_disposition": {"status": "accepted"}},
        session_state=session_state,
        load_restart_source_bytes_fn=lambda source: b"abc",
    )

    assert restored is not None
    assert restored.file_token == "report.docx:3:abc"


def test_missing_result_bundle_still_uses_restart_source(tmp_path):
    # Anti-regression for Fix A: the no-bundle path is untouched.
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    session_state = SessionState(
        processing_outcome="failed",
        restart_source=_valid_restart_record(restart_path),
    )

    restored = application_flow.resolve_effective_uploaded_file(
        uploaded_file=None,
        current_result=None,
        session_state=session_state,
        load_restart_source_bytes_fn=lambda source: b"abc",
    )

    assert restored is not None
    assert restored.file_token == "report.docx:3:abc"


def test_has_restartable_source_rejects_record_without_spec045_metadata(tmp_path):
    # Round-11 Fix B: a record missing payload_sha256/source_format is rejected by
    # load_restart_source_bytes as invalid_metadata on EVERY rerun, so the RESTARTABLE
    # offer could never work. It must not reach the view at all.
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    legacy_record = {"filename": "report.docx", "storage_path": str(restart_path)}
    session_state = SessionState(processing_outcome="stopped", restart_source=legacy_record)

    assert application_flow.has_restartable_source(session_state=session_state) is False  # type: ignore[arg-type]
    assert application_flow.derive_app_idle_view_state(
        current_result=None, uploaded_file=None, session_state=session_state
    ) == "empty"


def test_has_restartable_source_rejects_record_with_malformed_payload_digest(tmp_path):
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    record = _valid_restart_record(restart_path)
    record["payload_sha256"] = "not-a-digest"
    session_state = SessionState(processing_outcome="stopped", restart_source=record)

    assert application_flow.has_restartable_source(session_state=session_state) is False  # type: ignore[arg-type]


def test_has_restartable_source_still_offers_and_restores_a_valid_record(tmp_path):
    # ANTI-VACUUM for Fix B: the added structural gate must not swallow good records.
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    record = _valid_restart_record(restart_path)
    session_state = SessionState(processing_outcome="stopped", restart_source=record)

    assert application_flow.has_restartable_source(session_state=session_state) is True  # type: ignore[arg-type]
    assert application_flow.derive_app_idle_view_state(
        current_result=None, uploaded_file=None, session_state=session_state
    ) == "restartable"

    restored = application_flow.get_cached_restart_file(
        session_state=cast(application_flow.SessionStateLike, session_state),
        load_restart_source_bytes_fn=lambda source: b"abc",
    )

    assert restored is not None
    assert restored.content_bytes == b"abc"


def test_transient_unreadable_payload_keeps_restart_record_and_file(tmp_path):
    # Fix B must stay non-destructive: a momentarily unreadable payload (e.g. a locked
    # file) is NOT a permanently bad record, so neither the session record nor the file
    # may be dropped.
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    record = _valid_restart_record(restart_path)
    session_state = SessionState(processing_outcome="stopped", restart_source=record)

    assert application_flow.get_cached_restart_file(
        session_state=cast(application_flow.SessionStateLike, session_state),
        load_restart_source_bytes_fn=lambda source: None,
    ) is None
    assert session_state.restart_source == record
    assert restart_path.is_file()
    assert application_flow.has_restartable_source(session_state=session_state) is True  # type: ignore[arg-type]


def _confined_restart_record(run_dir: Path, source_bytes: bytes = b"abc") -> dict[str, object]:
    """A record whose file satisfies restart_store's confined-path rule (inside RUN_DIR,
    ``restart_`` prefix), so the real loader and the real clear helper both engage."""
    restart_path = run_dir / "restart_session_token.docx"
    restart_path.write_bytes(source_bytes)
    return {
        "filename": "report.docx",
        "token": "report.docx:3:abc",
        "storage_path": str(restart_path),
        "size": len(source_bytes),
        "payload_sha256": hashlib.sha256(source_bytes).hexdigest(),
        "source_format": "docx",
        "conversion_backend": None,
        "storage_kind": "restart",
    }


def test_restart_record_with_permanently_bad_metadata_self_heals(tmp_path, monkeypatch):
    # Round-11 Fix 2: invalid_metadata is a DETERMINISTIC verdict — every rerun re-reads
    # and re-rejects the same record. The offer must heal itself: session record cleared
    # and the file removed through the confined helper.
    monkeypatch.setattr(restart_store, "RUN_DIR", tmp_path)
    record = _confined_restart_record(tmp_path)
    record.pop("token")
    restart_path = Path(str(record["storage_path"]))
    session_state = SessionState(processing_outcome="stopped", restart_source=record)

    assert application_flow.get_cached_restart_file(
        session_state=cast(application_flow.SessionStateLike, session_state)
    ) is None
    assert session_state.restart_source is None
    assert not restart_path.exists()
    assert application_flow.has_restartable_source(session_state=session_state) is False  # type: ignore[arg-type]


def test_restart_record_with_integrity_mismatch_self_heals(tmp_path, monkeypatch):
    # Same for integrity_mismatch: the stored digest can never match the bytes on disk.
    monkeypatch.setattr(restart_store, "RUN_DIR", tmp_path)
    record = _confined_restart_record(tmp_path)
    restart_path = Path(str(record["storage_path"]))
    restart_path.write_bytes(b"tampered")
    session_state = SessionState(processing_outcome="stopped", restart_source=record)

    assert application_flow.get_cached_restart_file(
        session_state=cast(application_flow.SessionStateLike, session_state)
    ) is None
    assert session_state.restart_source is None
    assert not restart_path.exists()
    assert application_flow.has_restartable_source(session_state=session_state) is False  # type: ignore[arg-type]


def test_unconfined_restart_record_is_dropped_without_deleting_the_outside_file(tmp_path, monkeypatch):
    # unconfined_path is deterministic too, so the session record goes; but deletion stays
    # inside the confined helper, which refuses a path outside RUN_DIR. Self-healing must
    # never widen the deletion scope.
    run_dir = tmp_path / "run"
    run_dir.mkdir()
    monkeypatch.setattr(restart_store, "RUN_DIR", run_dir)
    outside_path = tmp_path / "restart_outside.docx"
    outside_path.write_bytes(b"abc")
    record = _confined_restart_record(run_dir)
    record["storage_path"] = str(outside_path)
    session_state = SessionState(processing_outcome="stopped", restart_source=record)

    assert application_flow.get_cached_restart_file(
        session_state=cast(application_flow.SessionStateLike, session_state)
    ) is None
    assert session_state.restart_source is None
    assert outside_path.is_file()


def test_unreadable_restart_payload_preserves_record_and_file(tmp_path, monkeypatch):
    # COUNTER-PROOF that the self-heal is scoped, not indiscriminate: unreadable_payload
    # can be transient (a momentarily locked file), so both the session record and the
    # stored file survive and the record stays on offer.
    monkeypatch.setattr(restart_store, "RUN_DIR", tmp_path)
    record = _confined_restart_record(tmp_path)
    restart_path = Path(str(record["storage_path"]))
    restart_path.unlink()
    restart_path.mkdir()  # a path that exists but raises OSError on read_bytes
    session_state = SessionState(processing_outcome="stopped", restart_source=record)

    assert application_flow.get_cached_restart_file(
        session_state=cast(application_flow.SessionStateLike, session_state)
    ) is None
    assert session_state.restart_source == record
    assert restart_path.exists()


def test_valid_restart_record_is_restored_and_left_untouched(tmp_path, monkeypatch):
    # ANTI-VACUUM for Fix 2: self-healing must not touch a record that loads fine.
    monkeypatch.setattr(restart_store, "RUN_DIR", tmp_path)
    record = _confined_restart_record(tmp_path)
    restart_path = Path(str(record["storage_path"]))
    session_state = SessionState(processing_outcome="stopped", restart_source=record)

    restored = application_flow.get_cached_restart_file(
        session_state=cast(application_flow.SessionStateLike, session_state)
    )

    assert restored is not None
    assert restored.content_bytes == b"abc"
    assert session_state.restart_source == record
    assert restart_path.is_file()


def test_completed_record_with_permanently_bad_metadata_self_heals(tmp_path, monkeypatch):
    # The completed-source cache heals through its own existing confined helper.
    monkeypatch.setattr(restart_store, "RUN_DIR", tmp_path)
    record = _confined_restart_record(tmp_path)
    record["payload_sha256"] = "not-a-digest"
    completed_path = Path(str(record["storage_path"]))
    session_state = SessionState(completed_source=record)

    assert application_flow.get_cached_completed_file(
        session_state=cast(application_flow.SessionStateLike, session_state)
    ) is None
    assert session_state.completed_source is None
    assert not completed_path.exists()


def test_has_restartable_source_rejects_pdf_record_without_conversion_backend(tmp_path):
    # Round-11 Fix 1 (the DRIFT case): the gate re-implemented only a subset of the
    # loader's rule and omitted conversion_backend, so this record was offered in the
    # RESTARTABLE view and then always failed to restore.
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    record = _valid_restart_record(restart_path)
    record["source_format"] = "pdf"
    record["conversion_backend"] = None
    session_state = SessionState(processing_outcome="stopped", restart_source=record)

    assert application_flow.has_restartable_source(session_state=session_state) is False  # type: ignore[arg-type]

    record["conversion_backend"] = "   "

    assert application_flow.has_restartable_source(session_state=session_state) is False  # type: ignore[arg-type]


def test_has_restartable_source_still_offers_pdf_record_with_conversion_backend(tmp_path):
    # ANTI-VACUUM for Fix 1: the shared helper must not start rejecting good records.
    restart_path = tmp_path / "restart.bin"
    restart_path.write_bytes(b"abc")
    record = _valid_restart_record(restart_path)
    record["source_format"] = "pdf"
    record["conversion_backend"] = "libreoffice-writer-pdf-import"
    session_state = SessionState(processing_outcome="stopped", restart_source=record)

    assert application_flow.has_restartable_source(session_state=session_state) is True  # type: ignore[arg-type]


def test_prepare_run_context_for_background_uses_frozen_upload_payload(monkeypatch):
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    captured = {}
    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="prepared-key",
        cached=False,
    )
    payload = _freeze_uploaded_file("report.docx", b"abc")

    result = application_flow.prepare_run_context_for_background(
        uploaded_payload=payload,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        processing_operation="audiobook",
        app_config={"translation_domain_default": "general", "reader_cleanup_default": True},
        prepare_document_for_processing_fn=lambda **kwargs: (captured.setdefault("prepare", kwargs), prepared_document)[1],
    )

    assert result.uploaded_filename == "report.docx"
    assert result.uploaded_file_bytes == b"abc"
    assert result.uploaded_file_token == payload.file_token
    assert captured["prepare"]["app_config"] == {"translation_domain_default": "general", "reader_cleanup_default": True}
    assert captured["prepare"]["processing_operation"] == "audiobook"


def test_prepare_run_context_for_background_forwards_tenant_client_factory(monkeypatch):
    """Spec 039 (B): a tenant client_factory injected into the UI preparation
    entry must reach prepare_document_for_processing (as both get_client_fn and
    client_factory, the SAME object), and the None-default path must stay
    byte-compatible (no factory kwargs)."""
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="prepared-key",
        cached=False,
    )
    payload = _freeze_uploaded_file("report.docx", b"abc")

    def _sentinel_factory(selector=None, required_capability="responses_text", *, config_like=None):
        return object()

    with_factory: dict[str, Any] = {}
    application_flow.prepare_run_context_for_background(
        uploaded_payload=payload,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        app_config={},
        prepare_document_for_processing_fn=lambda **kwargs: (with_factory.setdefault("kwargs", kwargs), prepared_document)[1],
        client_factory=_sentinel_factory,
    )
    assert with_factory["kwargs"]["client_factory"] is _sentinel_factory
    assert with_factory["kwargs"]["get_client_fn"] is _sentinel_factory

    without_factory: dict[str, Any] = {}
    application_flow.prepare_run_context_for_background(
        uploaded_payload=payload,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        app_config={},
        prepare_document_for_processing_fn=lambda **kwargs: (without_factory.setdefault("kwargs", kwargs), prepared_document)[1],
    )
    assert "client_factory" not in without_factory["kwargs"]
    assert "get_client_fn" not in without_factory["kwargs"]


def test_prepare_run_context_sync_path_forwards_tenant_client_factory(monkeypatch):
    """Spec 039 (B) follow-up (post-verification-rigor): the SYNC restart/reuse entry
    (ui/_app.py:894 -> application_flow.prepare_run_context) must ALSO forward a tenant
    client_factory to prepare_document_for_processing (both get_client_fn and
    client_factory, the SAME object); the None-default path stays byte-compatible. The
    background path was already covered; this closes the sync fallback the reviewer found
    unfixed after the first pass."""
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    session_state = SessionState(
        selected_source_token="",
        prepared_source_key="",
        completed_source={"filename": "report.docx", "token": "report.docx:3:ba7816bf8f01cfea", "storage_path": "completed.bin"},
    )
    monkeypatch.setattr(state.st, "session_state", session_state)
    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="prepared-key",
        cached=False,
    )

    def _sentinel_factory(selector=None, required_capability="responses_text", *, config_like=None):
        return object()

    with_factory: dict[str, Any] = {}
    application_flow.prepare_run_context(
        uploaded_file=UploadedFileStub("report.docx", b"abc"),
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        app_config={},
        session_state=session_state,
        reset_run_state_fn=lambda **kw: None,
        fail_critical_fn=lambda *a, **kw: (_ for _ in ()).throw(AssertionError("unexpected critical error")),
        log_event_fn=lambda *a, **kw: None,
        prepare_document_for_processing_fn=lambda **kwargs: (with_factory.setdefault("kwargs", kwargs), prepared_document)[1],
        client_factory=_sentinel_factory,
    )
    assert with_factory["kwargs"]["client_factory"] is _sentinel_factory
    assert with_factory["kwargs"]["get_client_fn"] is _sentinel_factory

    without_factory: dict[str, Any] = {}
    application_flow.prepare_run_context(
        uploaded_file=UploadedFileStub("report.docx", b"abc"),
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        app_config={},
        session_state=session_state,
        reset_run_state_fn=lambda **kw: None,
        fail_critical_fn=lambda *a, **kw: (_ for _ in ()).throw(AssertionError("unexpected critical error")),
        log_event_fn=lambda *a, **kw: None,
        prepare_document_for_processing_fn=lambda **kwargs: (without_factory.setdefault("kwargs", kwargs), prepared_document)[1],
    )
    assert "client_factory" not in without_factory["kwargs"]
    assert "get_client_fn" not in without_factory["kwargs"]


def test_prepare_run_context_for_background_uses_real_cache(monkeypatch):
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    preparation.clear_preparation_cache(clear_shared=True)
    calls = {"extract": 0}
    progress_events = []

    def fake_extract(uploaded_file):
        calls["extract"] += 1
        return ["paragraph"], [], None, [], None, None

    monkeypatch.setattr(preparation, "extract_document_content_with_normalization_reports", fake_extract)
    monkeypatch.setattr(preparation, "build_document_text", lambda paragraphs: "text")
    monkeypatch.setattr(preparation, "build_semantic_blocks", lambda paragraphs, max_chars, relations=None: ["block"])
    monkeypatch.setattr(preparation, "build_editing_jobs", lambda blocks, max_chars: [{"target_text": "block", "target_chars": 5, "context_chars": 0}])

    uploaded_payload = _freeze_uploaded_file("report.docx", b"abc")
    application_flow.prepare_run_context_for_background(
        uploaded_payload=uploaded_payload,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        progress_callback=lambda **payload: progress_events.append(payload),
    )
    second = application_flow.prepare_run_context_for_background(
        uploaded_payload=uploaded_payload,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        progress_callback=lambda **payload: progress_events.append(payload),
    )

    assert calls["extract"] == 1
    assert second.prepared_source_key.endswith(
        ":6000:high_only:off:phase2_default:epigraph_attribution,image_caption,table_caption,toc_region:lc=1:3:80:flag:pv=2"
        ":pk=4:sl=en:tl=ru:td=general:sr=0:srm=legacy:so=10:0.1:1.5:ar=off"
    )
    assert second.preparation_cached is True
    assert second.preparation_stage == "Документ подготовлен"
    assert progress_events[-1]["stage"] == "Документ подготовлен"
    assert progress_events[-1]["metrics"]["cached"] is True


def test_prepare_run_context_for_background_processes_real_docx_without_mocks():
    preparation.clear_preparation_cache(clear_shared=True)
    source_doc = Document()
    source_doc.add_heading("Глава 1", level=1)
    source_doc.add_paragraph(
        "Первый абзац документа содержит достаточно длинный связный текст, чтобы не считаться подозрительно коротким body-блоком."
    )
    source_doc.add_paragraph(
        "Второй абзац документа тоже содержит несколько обычных слов и завершенное предложение для стабильной подготовки."
    )
    source_buffer = BytesIO()
    source_doc.save(source_buffer)

    payload = _freeze_uploaded_file("report.docx", source_buffer.getvalue())
    progress_events = []

    result = application_flow.prepare_run_context_for_background(
        uploaded_payload=payload,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        progress_callback=lambda **payload: progress_events.append(payload),
    )

    assert result.uploaded_filename == "report.docx"
    assert result.uploaded_file_bytes == source_buffer.getvalue()
    assert result.source_text
    assert len(result.paragraphs) == 3
    assert len(result.jobs) >= 1
    assert any("Глава 1" in str(job.get("target_text", "")) for job in result.jobs)
    assert any(event["stage"] == "Разбор DOCX" for event in progress_events)
    assert progress_events[-1]["stage"] == "Документ подготовлен"


def test_background_handoff_persists_result_bundle_and_ui_artifacts(tmp_path, monkeypatch):
    session_state = SessionState()
    monkeypatch.setattr(state.st, "session_state", session_state)
    monkeypatch.setattr(processing_runtime.st, "session_state", session_state)
    state.init_session_state()
    preparation.clear_preparation_cache(clear_shared=True)

    source_doc = Document()
    source_doc.add_heading("Глава 1", level=1)
    source_doc.add_paragraph(
        "Первый абзац документа содержит достаточно длинный связный текст, чтобы подготовка прошла без fallback-поведения."
    )
    source_doc.add_paragraph(
        "Второй абзац нужен, чтобы downstream handoff собрал непустой result bundle и сохранил итоговые артефакты."
    )
    source_buffer = BytesIO()
    source_doc.save(source_buffer)

    uploaded_payload = _freeze_uploaded_file("report.docx", source_buffer.getvalue())
    prepared = application_flow.prepare_run_context_for_background(
        uploaded_payload=uploaded_payload,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
    )

    preparation_events: queue.Queue[object] = queue.Queue()
    session_state.preparation_event_queue = preparation_events
    session_state.preparation_worker = object()
    preparation_events.put(
        PreparationCompleteEvent(
            prepared_run_context=prepared,
            upload_marker=f"{prepared.uploaded_file_token}:6000",
        )
    )
    preparation_finalizations = []

    processing_runtime.drain_preparation_events(
        reset_run_state=state.reset_run_state,
        set_processing_status=state.set_processing_status,
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: preparation_finalizations.append(
            (stage, detail, progress, terminal_kind)
        ),
        push_activity=state.push_activity,
    )

    assert session_state.prepared_run_context is prepared
    assert preparation_finalizations[-1] == ("Документ подготовлен", "", 1.0, "completed")

    processing_events: queue.Queue[object] = queue.Queue()
    stop_event = threading.Event()
    runtime = processing_runtime.BackgroundRuntime(processing_events, stop_event)
    state.apply_processing_start(
        uploaded_filename=prepared.uploaded_filename,
        uploaded_token=prepared.uploaded_file_token,
        image_mode="safe",
        processing_operation="edit",
        audiobook_postprocess_enabled=False,
        worker=object(),
        event_queue=processing_events,
        stop_event=stop_event,
    )

    fake_client = object()
    artifact_paths: dict[str, str] = {}
    markdown_text = "# Результат\n\nИтоговый markdown собран из background processing handoff."
    docx_bytes = b"result-docx-bytes"

    def _fake_run_document_processing(**kwargs):
        nonlocal artifact_paths
        artifact_paths = runtime_artifacts.write_ui_result_artifacts(
            source_name=str(kwargs["uploaded_file"]),
            markdown_text=markdown_text,
            docx_bytes=docx_bytes,
            output_dir=tmp_path,
            created_at=1_766_636_465.0,
        )
        emitted_runtime = kwargs["runtime"]
        emitted_runtime.emit(
            SetProcessingStatusEvent(
                payload={"stage": "Сборка результата", "detail": "Сохраняю финальные артефакты", "progress": 0.95}
            )
        )
        emitted_runtime.emit(SetStateEvent(values={"latest_markdown": markdown_text, "latest_docx_bytes": docx_bytes}))
        emitted_runtime.emit(PushActivityEvent(message="Финальные UI-артефакты сохранены."))
        emitted_runtime.emit(
            FinalizeProcessingStatusEvent(stage="Готово", detail="", progress=1.0, terminal_kind="completed")
        )
        return "succeeded"

    service = processing_service.clone_processing_service(
        get_client_fn=lambda: fake_client,
        run_document_processing_impl_fn=_fake_run_document_processing,
    )
    service.run_processing_worker(
        runtime=runtime,
        uploaded_filename=prepared.uploaded_filename,
        source_token=prepared.uploaded_file_token,
        prepared_source_key=prepared.prepared_source_key,
        structure_fingerprint=getattr(prepared, "structure_fingerprint", None),
        jobs=prepared.jobs,
        selected_segment_ids=getattr(prepared, "selected_segment_ids", None),
        document_segments=getattr(prepared, "segments", None),
        output_mode=getattr(prepared, "output_mode", None),
        include_front_matter=bool(getattr(prepared, "include_front_matter", False)),
        include_toc=bool(getattr(prepared, "include_toc", False)),
        source_paragraphs=prepared.paragraphs,
        image_assets=prepared.image_assets,
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="edit",
        document_context_prompt="",
    )

    processing_finalizations = []

    processing_runtime.drain_processing_events(
        set_processing_status=state.set_processing_status,
        finalize_processing_status=lambda stage, detail, progress, terminal_kind=None: processing_finalizations.append(
            (stage, detail, progress, terminal_kind)
        ),
        push_activity=state.push_activity,
        append_log=state.append_log,
        append_image_log=state.append_image_log,
    )

    result_bundle = processing_runtime.get_current_result_bundle()
    markdown_path = Path(artifact_paths["markdown_path"])
    docx_path = Path(artifact_paths["docx_path"])

    assert result_bundle is not None
    assert result_bundle["source_name"] == "report.docx"
    assert result_bundle["source_token"] == prepared.uploaded_file_token
    assert result_bundle["markdown_text"] == markdown_text
    assert result_bundle["docx_bytes"] == docx_bytes
    assert session_state.processing_outcome == "succeeded"
    assert session_state.processing_worker is None
    assert session_state.processing_event_queue is None
    assert session_state.latest_markdown == markdown_text
    assert session_state.latest_docx_bytes == docx_bytes
    assert session_state.activity_feed[-1]["message"] == "Финальные UI-артефакты сохранены."
    assert processing_finalizations[-1] == ("Готово", "", 1.0, "completed")
    assert markdown_path.exists()
    assert markdown_path.read_text(encoding="utf-8") == markdown_text
    assert docx_path.exists()
    assert docx_path.read_bytes() == docx_bytes


def test_prepare_run_context_for_background_skips_renormalization_for_frozen_payload(monkeypatch):
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    monkeypatch.setattr(
        flow_core,
        "freeze_uploaded_file",
        lambda uploaded_file: (_ for _ in ()).throw(AssertionError("unexpected refreeze")),
    )
    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="prepared-key",
        cached=False,
    )
    payload = _freeze_uploaded_file("report.docx", b"abc")

    result = application_flow.prepare_run_context_for_background(
        uploaded_payload=payload,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        prepare_document_for_processing_fn=lambda **kwargs: prepared_document,
    )

    assert result.uploaded_filename == payload.filename
    assert result.uploaded_file_bytes == payload.content_bytes
    assert result.uploaded_file_token == payload.file_token


def test_prepare_run_context_sync_and_background_share_same_upload_contract(monkeypatch):
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="prepared-key",
        cached=False,
    )
    sync_calls = []
    background_calls = []
    payload = _freeze_uploaded_file("report.docx", b"abc")

    def prepare_sync(**kwargs):
        sync_calls.append(kwargs)
        return prepared_document

    def prepare_background(**kwargs):
        background_calls.append(kwargs)
        return prepared_document

    sync_result = application_flow.prepare_run_context(
        uploaded_file=UploadedFileStub("report.docx", b"abc"),
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        session_state=SessionState(selected_source_token="", prepared_source_key=""),
        reset_run_state_fn=lambda **kwargs: None,
        fail_critical_fn=lambda *args, **kwargs: None,
        log_event_fn=lambda *args, **kwargs: None,
        prepare_document_for_processing_fn=prepare_sync,
    )
    background_result = application_flow.prepare_run_context_for_background(
        uploaded_payload=payload,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        prepare_document_for_processing_fn=prepare_background,
    )

    assert sync_calls[0]["uploaded_payload"].filename == background_calls[0]["uploaded_payload"].filename == "report.docx"
    assert sync_calls[0]["uploaded_payload"].content_bytes == background_calls[0]["uploaded_payload"].content_bytes == b"abc"
    assert sync_calls[0]["uploaded_payload"].file_token == background_calls[0]["uploaded_payload"].file_token == payload.file_token
    assert sync_result.uploaded_filename == background_result.uploaded_filename
    assert sync_result.uploaded_file_bytes == background_result.uploaded_file_bytes
    assert sync_result.uploaded_file_token == background_result.uploaded_file_token


@pytest.mark.parametrize(
    ("restore_kind", "source_format", "conversion_backend"),
    [
        ("completed", "pdf", "pdf-text-layer"),
        ("restart", "doc", "libreoffice"),
        ("completed", "docx", None),
    ],
)
def test_restored_frozen_payload_enters_sync_preparation_without_refreeze_or_reconversion(
    restore_kind,
    source_format,
    conversion_backend,
    monkeypatch,
):
    normalized_bytes = f"normalized-{source_format}".encode()
    source_token = f"report.{source_format}:{len(normalized_bytes)}:original"
    source_record = {
        "filename": "report.docx",
        "token": source_token,
        "storage_path": f"{restore_kind}.bin",
        "size": len(normalized_bytes),
        "payload_sha256": hashlib.sha256(normalized_bytes).hexdigest(),
        "source_format": source_format,
        "conversion_backend": conversion_backend,
    }
    session_state = SessionState(
        selected_source_token="",
        prepared_source_key="",
        **{f"{restore_kind}_source": source_record},
    )
    load_calls = []
    restore_fn = (
        application_flow.get_cached_completed_file
        if restore_kind == "completed"
        else application_flow.get_cached_restart_file
    )
    load_kwarg = (
        {"load_completed_source_bytes_fn": lambda record: load_calls.append(record) or normalized_bytes}
        if restore_kind == "completed"
        else {"load_restart_source_bytes_fn": lambda record: load_calls.append(record) or normalized_bytes}
    )
    restored_payload = restore_fn(
        session_state=cast(application_flow.SessionStateLike, session_state),
        **load_kwarg,
    )
    assert isinstance(restored_payload, processing_runtime.FrozenUploadPayload)

    freeze_calls = []
    conversion_calls = []
    monkeypatch.setattr(flow_core, "validate_docx_source_bytes", lambda source_bytes: None)
    monkeypatch.setattr(
        flow_core,
        "freeze_uploaded_file",
        lambda uploaded_file: freeze_calls.append(uploaded_file),
    )
    monkeypatch.setattr(
        processing_runtime,
        "normalize_uploaded_document",
        lambda **kwargs: conversion_calls.append(kwargs),
    )
    prepared_document = SimpleNamespace(
        source_text="text",
        paragraphs=["p1"],
        image_assets=[],
        jobs=[{"target_text": "block", "target_chars": 5, "context_chars": 0}],
        prepared_source_key="prepared-key",
        cached=False,
    )
    prepare_calls = []

    result = application_flow.prepare_run_context(
        uploaded_file=restored_payload,
        chunk_size=6000,
        image_mode="safe",
        keep_all_image_variants=True,
        session_state=session_state,
        reset_run_state_fn=lambda **kwargs: None,
        fail_critical_fn=lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("unexpected critical error")),
        log_event_fn=lambda *args, **kwargs: None,
        prepare_document_for_processing_fn=lambda **kwargs: prepare_calls.append(kwargs) or prepared_document,
    )

    assert load_calls == [source_record]
    assert freeze_calls == []
    assert conversion_calls == []
    assert prepare_calls[0]["uploaded_payload"] is restored_payload
    assert prepare_calls[0]["uploaded_payload"].source_format == source_format
    assert prepare_calls[0]["uploaded_payload"].conversion_backend == conversion_backend
    assert result.uploaded_file_bytes == normalized_bytes
    assert result.uploaded_file_token == source_token
