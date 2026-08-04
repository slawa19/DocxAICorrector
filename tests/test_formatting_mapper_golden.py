"""Full-blob characterization gate for the formatting-transfer mapper (spec 029).

This test snapshots the ENTIRE output of ``_map_source_target_paragraphs`` — the
``mapping_pairs`` (serialized as ``[paragraph_id, target_index]``) PLUS the complete
``diagnostics`` dict — to canonical sorted JSON, and asserts byte-for-byte equality
against a committed golden fixture.

It is the correctness gate for the spec-029 performance refactor: every optimization
lever (E/A/B/C) MUST leave this blob byte-identical. Inputs are built fully offline
(no LLM/API): the five books under ``tests/sources/book`` are extracted, a realistic
identity target ``Document`` is synthesized (headings styled ``Heading N``, everything
else ``Normal``), deterministic target perturbations force residuals into the fuzzy
passes (9/10/11) and the global similarity pass (13), and a synthesized
``generated_paragraph_registry`` carries deliberate SPLIT / MERGED / RENAMED / rebuild-key
entries so the registry-driven passes are exercised too.

Determinism is guaranteed by index arithmetic and content hashing only (no RNG / time). To
regenerate the fixtures after an intentional, reviewed behavior change, run::

    UPDATE_FORMATTING_MAPPER_GOLDEN=1 <run this test>

REGENERATION EVIDENCE CONTRACT — read before setting that variable. This gate is the only
place ``UPDATE_FORMATTING_MAPPER_GOLDEN`` is documented, so it would otherwise authorise its
own regeneration. Regenerating is allowed (spec 029 scopes byte-identity to ITS optimization
levers, not to a permanent freeze), but the review that blesses it MUST be argued **per source
paragraph, keyed on the paragraph's TEXT — never on the blob's counters**. ``mapped_count``,
``bad_pair_count`` and the per-strategy tallies are aggregates over a synthetic problem whose
size and composition move with the document, so they can drop, rise or stay flat on the very
same change and prove nothing in any direction. The claim a regeneration has to support is:
*of every source paragraph whose text is unchanged, here is what changed in its mapping
outcome, and here is why that is the intended behavior change.* Attach that comparison to the
spec entry; "the counters look fine" is not a review.
"""

from __future__ import annotations

import glob
import hashlib
import json
import os
import re
from pathlib import Path

import pytest
from docx import Document

from docxaicorrector.document.extraction import extract_document_content_from_docx
from docxaicorrector.generation.formatting_transfer import _map_source_target_paragraphs

# Cap paragraphs per book: the gate exercises a broad spread of passes at this size (9-10
# distinct mapping strategies per book, 14 across the corpus) while staying a few seconds per
# book. Full-book scaling is the job of the offline profiling harness, not this correctness gate.
# Which strategies fire is a PROPERTY OF THE SYNTHETIC INPUT, not a contract: `bounded_registry_fuzzy`
# fired 4 times corpus-wide before the F16 key change and 0 after, and is covered directly by
# tests/test_format_restoration.py::test_mapping_uses_bounded_registry_fuzzy_pass_for_translated_body_target.
_PARAGRAPH_CAP = 500

# Loads the real book fixtures under tests/sources/book/ (offline, no API) — an
# integration-grade characterization gate, not a fast unit test.
pytestmark = pytest.mark.integration_local

_REPO_ROOT = Path(__file__).resolve().parent.parent
_BOOK_GLOB = str(_REPO_ROOT / "tests" / "sources" / "book" / "*.docx")
_FIXTURE_DIR = _REPO_ROOT / "tests" / "fixtures" / "formatting_mapper_golden"


def _slug(path: str) -> str:
    stem = Path(path).stem
    return re.sub(r"[^a-z0-9]+", "_", stem.lower()).strip("_")


def _stable_perturb_key(source_paragraph) -> int:
    """Content-addressed perturbation selector: a pure function of the paragraph's own TEXT.

    The bucket is ``sha1`` over the paragraph's whitespace-normalized text — the same normalized
    word sequence the perturbation branches in ``_build_case`` themselves read (``stripped`` /
    ``words``) — and over nothing else.

    It deliberately does NOT use ``paragraph_id``. ``paragraph_id`` is *positional*:
    ``_assign_paragraph_identity`` (``document/extraction.py``) hands out ``f"p{logical_index:04d}"``
    over the final paragraph list, so inserting, deleting or MERGING a single paragraph renumbers
    every later one. Keying on it therefore did the exact opposite of what the previous docstring
    here claimed: the mapper was handed a wholly different synthetic problem for the entire tail of
    the book, and the resulting fixture churn looked like a mapper change when nothing about those
    paragraphs had changed. Measured on the five book fixtures, merging one paragraph at source
    index 5 re-rolled the perturbation of 245-273 of the 498 text-identical survivors per book
    under the id key, and of 0 under this text key.
    ``test_perturbation_selection_is_stable_under_paragraph_merge`` below is the standing guard.

    The text fallback the id key carried was dead code in every real run: extraction always
    assigns a non-empty ``paragraph_id``, so the ``or`` never reached the text.

    Two consequences, accepted deliberately:

    * Two source paragraphs with identical text get the identical perturbation. Real books do
      repeat short boilerplate lines. A collision only means the mapper is handed the same
      synthetic sub-problem twice, which is itself realistic; the only way to disambiguate
      duplicates is an occurrence ordinal, and that re-introduces precisely the positional coupling
      this key exists to remove. Measured: 5-22 duplicate paragraphs per capped 500.
    * Empty-text paragraphs (image-only, etc.) all share the ``sha1("")`` bucket, for the same
      reason. Measured: 0 of them inside the capped window on all five books today.

    Determinism is unaffected — ``sha1`` of a fixed string, no RNG and no time.
    """
    identity = " ".join(str(getattr(source_paragraph, "text", "") or "").split())
    return int(hashlib.sha1(identity.encode("utf-8")).hexdigest()[:8], 16)


def _build_case(source_paragraphs):
    """Deterministically synthesize (target Document, registry) from source paragraphs.

    Perturbation SELECTION is keyed on the paragraph's own text content (F16,
    ``_stable_perturb_key``); only the operations that are genuinely about adjacency/ordering (the
    MERGED next-source fold, the target index a rebuild-key hint points at) still reference the
    index. The blob stays fully reproducible (hash of fixed text, no RNG / time).
    """
    doc = Document()
    registry: list[dict[str, object]] = []
    n = len(source_paragraphs)
    for i, sp in enumerate(source_paragraphs):
        base = sp.text
        stripped = base.strip()
        words = stripped.split()
        key = _stable_perturb_key(sp)

        # --- target text: mostly identity, ~14% small edit (high ratio -> fuzzy/13
        # accept), ~9% large edit (low ratio -> residual, exercises reject paths).
        ttext = base
        if stripped and key % 7 == 2 and len(words) >= 2:
            ttext = base + " " + words[-1]
        elif stripped and key % 11 == 0 and len(words) >= 4:
            ttext = " ".join(words[: max(1, len(words) // 2)])

        # Realistic styling so role resolution mirrors a real generated docx.
        if sp.role == "heading":
            level = min(max(sp.heading_level or 1, 1), 6)
            doc.add_paragraph(ttext, style=f"Heading {level}")
        else:
            doc.add_paragraph(ttext, style="Normal")

        # --- registry entry (the "generated" text the model produced).
        pid = sp.paragraph_id
        if not pid:
            continue
        if key % 23 == 5:
            # RENAMED: drop the registry entry so this source must be recovered by
            # the exact-text / global-similarity passes instead of a registry pass.
            continue
        entry: dict[str, object] = {"paragraph_id": pid, "text": base}
        if key % 17 == 0 and len(words) >= 3:
            # SPLIT: generated markdown split a heading off the front of the body.
            entry["text"] = "### " + " ".join(words[:2]) + "\n" + " ".join(words[2:])
        elif (
            key % 19 == 0
            and i + 1 < n
            and source_paragraphs[i + 1].paragraph_id
            and source_paragraphs[i + 1].text.strip()
        ):
            # MERGED: generated text folded the next source paragraph in (genuinely
            # positional — the fold target is the immediate document neighbour).
            entry["text"] = base + "\n" + source_paragraphs[i + 1].text
            entry["merged_paragraph_ids"] = [source_paragraphs[i + 1].paragraph_id]
        if key % 31 == 3:
            # Rebuild-key hint (drives the paragraph_id_rebuild_key pass). The hint VALUE is
            # the paragraph's own target index (positional by definition).
            entry["target_paragraph_indexes"] = [i]
        registry.append(entry)
    return doc, registry


def _serialize(mapping_pairs, diagnostics, target_paragraphs) -> str:
    index_by_target = {id(paragraph): index for index, paragraph in enumerate(target_paragraphs)}
    pairs = [
        [source_paragraph.paragraph_id, index_by_target[id(target_paragraph)]]
        for source_paragraph, target_paragraph in mapping_pairs
    ]
    blob = {"mapping_pairs": pairs, "diagnostics": diagnostics}
    return json.dumps(blob, ensure_ascii=False, sort_keys=True, indent=2) + "\n"


def _compute_blob(book_path: str) -> str:
    with open(book_path, "rb") as handle:
        source_paragraphs, _ = extract_document_content_from_docx(handle)
    source_paragraphs = source_paragraphs[:_PARAGRAPH_CAP]
    document, registry = _build_case(source_paragraphs)
    target_paragraphs = document.paragraphs
    mapping_pairs, diagnostics = _map_source_target_paragraphs(
        source_paragraphs,
        target_paragraphs,
        generated_paragraph_registry=registry,
    )
    return _serialize(mapping_pairs, diagnostics, target_paragraphs)


_BOOKS = sorted(glob.glob(_BOOK_GLOB))


@pytest.mark.parametrize("book_path", _BOOKS, ids=[_slug(book) for book in _BOOKS])
def test_formatting_mapper_output_matches_golden(book_path: str) -> None:
    assert _BOOKS, f"no book fixtures found under {_BOOK_GLOB}"
    blob = _compute_blob(book_path)
    fixture_path = _FIXTURE_DIR / f"{_slug(book_path)}.json"

    if os.environ.get("UPDATE_FORMATTING_MAPPER_GOLDEN"):
        fixture_path.parent.mkdir(parents=True, exist_ok=True)
        fixture_path.write_text(blob, encoding="utf-8")
        pytest.skip(f"regenerated golden fixture: {fixture_path.name}")

    assert fixture_path.exists(), (
        f"missing golden fixture {fixture_path}; regenerate with "
        f"UPDATE_FORMATTING_MAPPER_GOLDEN=1"
    )
    expected = fixture_path.read_text(encoding="utf-8")
    assert blob == expected, (
        f"formatting mapper output diverged from golden for {Path(book_path).name}; "
        f"a spec-029 lever must be byte-identical. If the divergence is an INTENDED, reviewed "
        f"behavior change, regenerate with UPDATE_FORMATTING_MAPPER_GOLDEN=1 — but the review "
        f"MUST be argued per source paragraph keyed on the paragraph's TEXT (which unchanged-text "
        f"paragraphs changed mapping outcome, and why that is intended), NOT on this blob's "
        f"counters: mapped_count / bad_pair_count / strategy tallies are aggregates over a "
        f"synthetic problem and move in both directions on the same change. See this module's "
        f"docstring, 'REGENERATION EVIDENCE CONTRACT'."
    )


class _FakeSourceParagraph:
    """Minimal source-paragraph stand-in exposing exactly the attributes _build_case reads."""

    def __init__(self, paragraph_id: str, text: str, role: str = "body", heading_level: int | None = None) -> None:
        self.paragraph_id = paragraph_id
        self.text = text
        self.role = role
        self.heading_level = heading_level


def test_perturbation_selection_is_stable_under_unrelated_insertion() -> None:
    """F16: a source paragraph's perturbation must depend only on its own content, so
    inserting an unrelated paragraph elsewhere does not reshuffle it.

    Builds the same paragraphs twice — once plain, once with an unrelated paragraph inserted
    at the FRONT (shifting every positional index by one) — and asserts each shared paragraph
    lands the SAME target text. This covers the ``enumerate``-index selection that predated F16;
    it does NOT renumber ``paragraph_id``, so it alone would also have passed on the id-keyed
    version. ``test_perturbation_selection_is_stable_under_paragraph_merge`` below is the test
    that pins the real-world case, where extraction re-assigns every later ``paragraph_id``.
    """
    shared = [
        _FakeSourceParagraph(f"p{index:03d}", f"Source paragraph number {index} with several words here.")
        for index in range(40)
    ]
    baseline_doc, _ = _build_case(list(shared))
    inserted = [_FakeSourceParagraph("x_inserted", "An unrelated recovered image paragraph."), *shared]
    inserted_doc, _ = _build_case(inserted)

    baseline_text_by_pid = {sp.paragraph_id: para.text for sp, para in zip(shared, baseline_doc.paragraphs)}
    # The inserted doc has one extra leading paragraph; the rest align to `shared` by offset 1.
    inserted_text_by_pid = {
        sp.paragraph_id: para.text for sp, para in zip(shared, inserted_doc.paragraphs[1:])
    }

    assert baseline_text_by_pid == inserted_text_by_pid
    # Guard against a degenerate all-identity map: the selection must actually perturb some
    # paragraphs, otherwise the invariance above would be vacuous.
    assert any(para.text != sp.text for sp, para in zip(shared, baseline_doc.paragraphs))

    # Stability is per-content: the key is a pure function of the paragraph's own text, and the
    # positional paragraph_id has no influence at all.
    assert _stable_perturb_key(shared[5]) == _stable_perturb_key(
        _FakeSourceParagraph("some_other_id", shared[5].text)
    )
    assert _stable_perturb_key(shared[5]) != _stable_perturb_key(
        _FakeSourceParagraph(shared[5].paragraph_id, "different text")
    )


def test_perturbation_selection_is_stable_under_paragraph_merge() -> None:
    """F16 (the case the id-keyed version got wrong): MERGING two source paragraphs must move
    only the merged paragraph's own perturbation, not every later paragraph's.

    A merge is what the import stage does by construction, and it renumbers ``paragraph_id`` for
    the whole tail (``p{logical_index:04d}``). Under the old ``sha1(paragraph_id)`` key that handed
    the mapper a different synthetic problem for part of the tail: 6 of the 38 survivors here, and
    245-273 of 498 on the real book fixtures. Under the text key every paragraph whose own text
    survives unchanged must keep its exact target text — 0 may move.
    """
    shared = [
        _FakeSourceParagraph(f"p{index:03d}", f"Source paragraph number {index} with several words here.")
        for index in range(40)
    ]
    baseline_doc, _ = _build_case(list(shared))
    baseline_by_text = {sp.text: para.text for sp, para in zip(shared, baseline_doc.paragraphs)}

    # Merge #10 into #11 and renumber positionally, exactly as extraction re-assigns identities.
    merged_sources = [*shared[:10], _FakeSourceParagraph("", shared[10].text + " " + shared[11].text), *shared[12:]]
    for index, sp in enumerate(merged_sources):
        sp.paragraph_id = f"p{index:03d}"
    merged_doc, _ = _build_case(merged_sources)
    merged_by_text = {sp.text: para.text for sp, para in zip(merged_sources, merged_doc.paragraphs)}

    surviving = [text for text in baseline_by_text if text in merged_by_text]
    assert len(surviving) == len(shared) - 2, "only the two merged paragraphs should lose their text"
    changed = [text for text in surviving if baseline_by_text[text] != merged_by_text[text]]
    assert changed == [], f"{len(changed)} unchanged-text paragraphs had their perturbation re-rolled by a merge"
