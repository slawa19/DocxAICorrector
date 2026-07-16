import base64
import zipfile
from io import BytesIO
from typing import Any, cast

import pytest

import docxaicorrector.document.extraction as document
import docxaicorrector.document.extraction as document_extraction
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from docxaicorrector.core.models import ImageAsset, ParagraphUnit
from docxaicorrector.document.extraction import (
    build_document_text,
    extract_document_content_from_docx,
    inspect_placeholder_integrity,
)
from docxaicorrector.document.roles import (
    paragraph_has_strong_heading_format,
    resolve_effective_paragraph_font_size,
)
from docxaicorrector.document.semantic_blocks import build_marker_wrapped_block_text
from docxaicorrector.document.extraction import extract_document_content_with_normalization_reports


PNG_BYTES = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aK3cAAAAASUVORK5CYII=")


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_tab(run) -> None:
    run._element.append(OxmlElement("w:tab"))


def _append_textbox_with_paragraphs(paragraph, texts: list[str]) -> None:
    textbox_paragraphs = "".join(
        f"""
        <w:p>
            <w:r>
                <w:t>{text}</w:t>
            </w:r>
        </w:p>
        """
        for text in texts
    )
    paragraph._p.append(
        parse_xml(
            f"""
            <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                 xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                 xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                 xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <w:drawing>
                    <wp:inline>
                        <wp:extent cx="914400" cy="914400"/>
                        <wp:docPr id="1" name="TextBox 1"/>
                        <a:graphic>
                            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                                <wps:wsp>
                                    <wps:txbx>
                                        <w:txbxContent>
                                            {textbox_paragraphs}
                                        </w:txbxContent>
                                    </wps:txbx>
                                    <wps:bodyPr/>
                                </wps:wsp>
                            </a:graphicData>
                        </a:graphic>
                    </wp:inline>
                </w:drawing>
            </w:r>
            """
        )
    )


def _detach_inline_drawing(doc, image_path):
    """Build a valid inline-image ``w:drawing`` bound to the document part, detached."""
    throwaway = doc.add_paragraph()
    run = throwaway.add_run()
    run.add_picture(str(image_path))
    drawing = run._element.find(qn("w:drawing"))
    run._element.remove(drawing)
    throwaway._p.getparent().remove(throwaway._p)
    return drawing


def _append_textbox_with_interior_drawing(paragraph, text: str, drawing_element) -> None:
    """Append a textbox whose interior holds a text paragraph AND an image drawing."""
    textbox_run = parse_xml(
        """
        <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
             xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
             xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <w:drawing>
                <wp:inline>
                    <wp:extent cx="914400" cy="914400"/>
                    <wp:docPr id="2" name="TextBox 2"/>
                    <a:graphic>
                        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                            <wps:wsp>
                                <wps:txbx>
                                    <w:txbxContent/>
                                </wps:txbx>
                                <wps:bodyPr/>
                            </wps:wsp>
                        </a:graphicData>
                    </a:graphic>
                </wp:inline>
            </w:drawing>
        </w:r>
        """
    )
    txbx_content = textbox_run.find(".//" + qn("w:txbxContent"))
    text_paragraph = parse_xml(
        f"""
        <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:r><w:t>{text}</w:t></w:r>
        </w:p>
        """
    )
    txbx_content.append(text_paragraph)
    image_paragraph = OxmlElement("w:p")
    image_run = OxmlElement("w:r")
    image_run.append(drawing_element)
    image_paragraph.append(image_run)
    txbx_content.append(image_paragraph)
    paragraph._p.append(textbox_run)


def _nested_textbox_run(inner_run_xml: str) -> str:
    """A textbox run whose interior paragraph holds ``inner_run_xml`` (another run)."""
    return f"""
        <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
             xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
             xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <w:drawing>
                <wp:inline>
                    <wp:extent cx="914400" cy="914400"/>
                    <wp:docPr id="3" name="OuterTextBox"/>
                    <a:graphic>
                        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                            <wps:wsp>
                                <wps:txbx>
                                    <w:txbxContent>
                                        <w:p>{inner_run_xml}</w:p>
                                    </w:txbxContent>
                                </wps:txbx>
                                <wps:bodyPr/>
                            </wps:wsp>
                        </a:graphicData>
                    </a:graphic>
                </wp:inline>
            </w:drawing>
        </w:r>
        """


def _append_nested_textbox_with_interior_drawing(paragraph, drawing_element) -> None:
    """Append a NESTED textbox (a txbxContent inside a txbxContent) whose deepest
    interior paragraph holds exactly one image drawing."""
    outer_run = parse_xml(_nested_textbox_run(_nested_textbox_run("")))
    txbx_contents = outer_run.findall(".//" + qn("w:txbxContent"))
    innermost = txbx_contents[-1]
    image_paragraph = OxmlElement("w:p")
    image_run = OxmlElement("w:r")
    image_run.append(drawing_element)
    image_paragraph.append(image_run)
    innermost.append(image_paragraph)
    paragraph._p.append(outer_run)


def _set_raw_paragraph_alignment(paragraph, value: str) -> None:
    paragraph_properties = paragraph._element.get_or_add_pPr()
    alignment = paragraph_properties.find(qn("w:jc"))
    if alignment is None:
        alignment = OxmlElement("w:jc")
        paragraph_properties.append(alignment)
    alignment.set(qn("w:val"), value)


def _set_outline_level(paragraph, value: int) -> None:
    paragraph_properties = paragraph._element.get_or_add_pPr()
    outline_level = paragraph_properties.find(qn("w:outlineLvl"))
    if outline_level is None:
        outline_level = OxmlElement("w:outlineLvl")
        paragraph_properties.append(outline_level)
    outline_level.set(qn("w:val"), str(value))


def _set_style_outline_level(style, value: int) -> None:
    style_properties = style._element.get_or_add_pPr()
    outline_level = style_properties.find(qn("w:outlineLvl"))
    if outline_level is None:
        outline_level = OxmlElement("w:outlineLvl")
        style_properties.append(outline_level)
    outline_level.set(qn("w:val"), str(value))


def _set_style_alignment(style, value: str) -> None:
    style_properties = style._element.get_or_add_pPr()
    alignment = style_properties.find(qn("w:jc"))
    if alignment is None:
        alignment = OxmlElement("w:jc")
        style_properties.append(alignment)
    alignment.set(qn("w:val"), value)


def test_extraction_cleanup_removes_textbox_artifacts_and_reassigns_identity(tmp_path, monkeypatch):
    monkeypatch.setattr(document_extraction, "_resolve_paragraph_boundary_normalization_settings", lambda *a, **k: ("off", False))
    document_obj = Document()
    document_obj.add_heading("Synthetic Title", level=1)
    for page_number in range(1, 5):
        document_obj.add_paragraph(f"Body content {page_number}.")
        document_obj.add_paragraph(str(page_number))
        holder = document_obj.add_paragraph()
        _append_textbox_with_paragraphs(holder, ["www.example.com"])

    source_path = tmp_path / "textbox-artifacts.docx"
    document_obj.save(source_path)

    with source_path.open("rb") as source_file:
        paragraphs, _, _, _, _, cleanup_report, _ = extract_document_content_with_normalization_reports(source_file)

    texts = [paragraph.text for paragraph in paragraphs]
    assert texts.count("www.example.com") == 4
    assert sum(1 for text in texts if text in {"1", "2", "3", "4"}) == 4
    assert "Synthetic Title" in texts
    assert "Body content 1." in texts
    assert cleanup_report.cleanup_mode == "flag"
    assert cleanup_report.flagged_page_number_count == 4
    assert cleanup_report.flagged_repeated_artifact_count == 4
    assert cleanup_report.removed_page_number_count == 0
    assert cleanup_report.removed_repeated_artifact_count == 0
    assert [paragraph.logical_index for paragraph in paragraphs] == list(range(len(paragraphs)))
    assert [paragraph.paragraph_id for paragraph in paragraphs] == [f"p{index:04d}" for index in range(len(paragraphs))]


def test_extraction_passes_legacy_structure_recovery_mode_to_helpers(tmp_path, monkeypatch):
    document_obj = Document()
    document_obj.add_paragraph("Title")
    source_path = tmp_path / "legacy-structure-recovery.docx"
    document_obj.save(source_path)

    monkeypatch.setattr(document_extraction, "_resolve_paragraph_boundary_normalization_settings", lambda *a, **k: ("off", False))
    monkeypatch.setattr(document_extraction, "_resolve_layout_artifact_cleanup_settings", lambda app_config=None: (True, 3, 80, False, "legacy"))
    monkeypatch.setattr(document_extraction, "_resolve_relation_normalization_settings", lambda *a, **k: (False, "phase2_default", (), False))
    monkeypatch.setattr(document_extraction, "_resolve_paragraph_boundary_ai_review_settings", lambda *a, **k: (False, "off", 0, 0, 0, ""))
    monkeypatch.setattr(document_extraction, "build_paragraph_relations", lambda paragraphs, enabled_relation_kinds=(): ([], None))
    monkeypatch.setattr(document_extraction, "apply_relation_side_effects", lambda paragraphs, relations: None)
    def fake_reclassify_adjacent_captions(
        paragraphs,
        *,
        structure_recovery_enabled=False,
        structure_recovery_mode="legacy",
    ):
        captured["caption"] = (structure_recovery_enabled, structure_recovery_mode)

    monkeypatch.setattr(document_extraction, "reclassify_adjacent_captions", fake_reclassify_adjacent_captions)

    captured = {}

    def fake_promote(paragraphs, *, structure_recovery_enabled=False, structure_recovery_mode="legacy"):
        captured["promote"] = (structure_recovery_enabled, structure_recovery_mode)

    def fake_normalize(paragraphs, *, structure_recovery_enabled=False, structure_recovery_mode="legacy"):
        captured["normalize"] = (structure_recovery_enabled, structure_recovery_mode)

    def fake_cleanup(
        paragraphs,
        *,
        enabled=True,
        min_repeat_count=3,
        max_repeated_text_chars=80,
        cleanup_mode="legacy",
        structure_recovery_enabled=False,
        structure_recovery_mode="legacy",
    ):
        captured["cleanup"] = (structure_recovery_enabled, structure_recovery_mode)
        return paragraphs, document_extraction.LayoutArtifactCleanupReport(
            original_paragraph_count=len(paragraphs),
            cleaned_paragraph_count=len(paragraphs),
            removed_paragraph_count=0,
            removed_page_number_count=0,
            removed_repeated_artifact_count=0,
            removed_empty_or_whitespace_count=0,
            cleanup_applied=True,
        )

    def fake_repair(
        paragraphs,
        *,
        app_config=None,
        structure_recovery_enabled=False,
        structure_recovery_mode="legacy",
    ):
        captured["repair"] = (structure_recovery_enabled, structure_recovery_mode)
        return paragraphs, document_extraction.StructureRepairReport(
            applied=False,
            repaired_bullet_items=0,
            repaired_numbered_items=0,
            bounded_toc_regions=0,
            toc_body_boundary_repairs=0,
            heading_candidates_from_toc=0,
            remaining_isolated_marker_count=0,
        )

    monkeypatch.setattr(document_extraction, "promote_short_standalone_headings", fake_promote)
    monkeypatch.setattr(document_extraction, "normalize_front_matter_display_title", fake_normalize)
    monkeypatch.setattr(document_extraction, "clean_paragraph_layout_artifacts", fake_cleanup)
    monkeypatch.setattr(document_extraction, "repair_pdf_derived_structure", fake_repair)

    with source_path.open("rb") as source_file:
        extract_document_content_with_normalization_reports(
            source_file,
            app_config={
                "structure_recovery_enabled": False,
                "structure_recovery_mode": "ai_first",
            },
        )

    assert captured == {
        "promote": (False, "legacy"),
        "normalize": (False, "legacy"),
        "cleanup": (False, "legacy"),
        "repair": (False, "legacy"),
        "caption": (False, "legacy"),
    }


def _extract_source_rects(element) -> list[dict[str, str]]:
    return [
        {key: src_rect.get(key) for key in ("l", "t", "r", "b") if src_rect.get(key) is not None}
        for src_rect in element.xpath(".//a:srcRect")
    ]


def _make_docx_with_emdash_bullet_numbering(texts: list[str]) -> BytesIO:
    """Create a DOCX where 'List Paragraph' paragraphs use em-dash (U+2014) as bullet char."""
    doc = Document()
    doc.add_paragraph("Обычный текст перед списком.")

    numbering_part = doc.part.numbering_part
    numbering_root = numbering_part._element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), "900")
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2014")
    lvl.append(lvl_text)
    abstract_num.append(lvl)
    numbering_root.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "900")
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), "900")
    num.append(abstract_ref)
    numbering_root.append(num)

    for text in texts:
        para = doc.add_paragraph(text, style="List Paragraph")
        paragraph_properties = para._element.find(qn("w:pPr"))
        if paragraph_properties is None:
            paragraph_properties = OxmlElement("w:pPr")
            para._element.insert(0, paragraph_properties)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "900")
        num_pr.append(ilvl)
        num_pr.append(num_id)
        paragraph_properties.append(num_pr)

    doc.add_paragraph("Обычный текст после.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def test_extract_document_content_from_docx_merges_false_body_boundary_in_public_api():
    doc = Document()
    doc.add_paragraph("архетипами: повторяющимися моделями")
    doc.add_paragraph("поведения во времени, наблюдаемыми в разных системах.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].text == "архетипами: повторяющимися моделями поведения во времени, наблюдаемыми в разных системах."
    assert paragraphs[0].origin_raw_indexes == [0, 1]
    assert paragraphs[0].boundary_normalization_applied is True
    assert build_marker_wrapped_block_text(paragraphs) == (
        f"[[DOCX_PARA_{paragraphs[0].paragraph_id}]]\n"
        "архетипами: повторяющимися моделями поведения во времени, наблюдаемыми в разных системах."
    )


def test_extract_document_content_from_docx_flattens_inline_break_wrapped_prose():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("For centuries, economists and policymakers")
    paragraph.add_run().add_break()
    paragraph.add_run("divided activities by whether they created value.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].text == (
        "For centuries, economists and policymakers divided activities by whether they created value."
    )


def test_extract_document_content_from_docx_drops_break_only_spacer_paragraphs():
    doc = Document()
    doc.add_paragraph("Before")
    spacer = doc.add_paragraph()
    spacer.add_run().add_break()
    doc.add_paragraph("After")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.text for paragraph in paragraphs] == ["Before", "After"]


def test_extract_document_content_from_docx_extracts_textbox_paragraphs():
    doc = Document()
    doc.add_paragraph("Before")
    host = doc.add_paragraph()
    _append_textbox_with_paragraphs(host, ["Inside text box", "Second textbox paragraph"])
    doc.add_paragraph("Middle")
    doc.add_paragraph("After")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, image_assets = extract_document_content_from_docx(buffer)

    assert [paragraph.text for paragraph in paragraphs] == [
        "Before Inside text box Second textbox paragraph Middle",
        "After",
    ]
    assert image_assets == []


def test_extract_document_content_from_docx_deduplicates_adjacent_textbox_paragraphs():
    doc = Document()
    host = doc.add_paragraph()
    _append_textbox_with_paragraphs(host, ["Duplicated line", "Duplicated line", "Next line"])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.text for paragraph in paragraphs] == ["Duplicated line", "Next line"]


def test_extract_document_content_from_docx_splits_toc_like_inline_break_cluster_and_marks_toc_roles():
    doc = Document()
    doc.add_paragraph("Contents")
    paragraph = doc.add_paragraph()
    paragraph.add_run("Common Critiques of Value Extraction")
    paragraph.add_run().add_break()
    paragraph.add_run("What is Value?")
    paragraph.add_run().add_break()
    paragraph.add_run("Meet the Production Boundary")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _, _, _, _, _, _ = document.extract_document_content_with_normalization_reports(
        buffer,
        app_config={"structure_recovery_enabled": True, "structure_recovery_mode": "ai_first"},
    )

    assert [paragraph.text for paragraph in paragraphs] == [
        "Contents",
        "Common Critiques of Value Extraction",
        "What is Value?",
        "Meet the Production Boundary",
    ]
    assert [paragraph.structural_role for paragraph in paragraphs] == [
        "body",
        "body",
        "body",
        "body",
    ]
    assert [paragraph.heuristic_structural_role_hint for paragraph in paragraphs] == [
        "toc_header",
        "toc_entry",
        "toc_entry",
        "toc_entry",
    ]


def test_extract_document_content_from_docx_preserves_source_index_as_provenance_when_inline_break_cluster_splits():
    doc = Document()
    doc.add_paragraph("Contents")
    paragraph = doc.add_paragraph()
    paragraph.add_run("Common Critiques of Value Extraction")
    paragraph.add_run().add_break()
    paragraph.add_run("What is Value?")
    paragraph.add_run().add_break()
    paragraph.add_run("Meet the Production Boundary")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.source_index for paragraph in paragraphs] == [0, 1, 1, 1]
    assert [paragraph.logical_index for paragraph in paragraphs] == [0, 1, 2, 3]
    assert [paragraph.paragraph_id for paragraph in paragraphs] == ["p0000", "p0001", "p0002", "p0003"]
    assert [paragraph.toc_pattern_hint for paragraph in paragraphs] == [True, True, True, True]


def test_extract_document_content_with_normalization_reports_populates_page_number_stage0_signals(tmp_path, monkeypatch):
    monkeypatch.setattr(document_extraction, "_resolve_paragraph_boundary_normalization_settings", lambda *a, **k: ("off", False))
    document_obj = Document()
    document_obj.add_paragraph("Body content")
    document_obj.add_paragraph("12")

    source_path = tmp_path / "page-number-signals.docx"
    document_obj.save(source_path)

    with source_path.open("rb") as source_file:
        paragraphs, _, _, _, _, cleanup_report, _ = extract_document_content_with_normalization_reports(source_file)

    page_number_paragraph = next(paragraph for paragraph in paragraphs if paragraph.text == "12")

    assert cleanup_report.flagged_page_number_count == 1
    assert page_number_paragraph.is_likely_page_number is True
    assert page_number_paragraph.page_number == 12
    assert page_number_paragraph.position_fraction == 1.0


def test_extract_document_content_from_docx_splits_compact_toc_run_clusters_without_explicit_breaks():
    doc = Document()
    doc.add_paragraph("Contents")
    paragraph = doc.add_paragraph()
    paragraph.add_run("Banks and Financial Markets Become Allies")
    paragraph.add_run(" ")
    paragraph.add_run("The Banking Problem")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _, _, _, _, _, _ = document.extract_document_content_with_normalization_reports(
        buffer,
        app_config={"structure_recovery_enabled": True, "structure_recovery_mode": "ai_first"},
    )

    assert [paragraph.text for paragraph in paragraphs] == [
        "Contents",
        "Banks and Financial Markets Become Allies",
        "The Banking Problem",
    ]
    assert [paragraph.structural_role for paragraph in paragraphs] == [
        "body",
        "body",
        "body",
    ]
    assert [paragraph.heuristic_structural_role_hint for paragraph in paragraphs] == [
        "toc_header",
        "toc_entry",
        "toc_entry",
    ]


def test_extract_document_content_from_docx_splits_long_two_entry_compact_toc_run_clusters():
    doc = Document()
    doc.add_paragraph("Contents")
    paragraph = doc.add_paragraph()
    paragraph.add_run("Something Odd About the National Accounts: GDP Facit Saltus!")
    paragraph.add_run(" ")
    paragraph.add_run("Patching Up the National Accounts isn't Enough")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _, _, _, _, _, _ = document.extract_document_content_with_normalization_reports(
        buffer,
        app_config={"structure_recovery_enabled": True, "structure_recovery_mode": "ai_first"},
    )

    assert [paragraph.text for paragraph in paragraphs] == [
        "Contents",
        "Something Odd About the National Accounts: GDP Facit Saltus!",
        "Patching Up the National Accounts isn't Enough",
    ]
    assert [paragraph.structural_role for paragraph in paragraphs] == [
        "body",
        "body",
        "body",
    ]
    assert [paragraph.heuristic_structural_role_hint for paragraph in paragraphs] == [
        "toc_header",
        "toc_entry",
        "toc_entry",
    ]


def test_extract_document_content_with_normalization_reports_legacy_projects_toc_roles_from_inline_break_cluster(tmp_path):
    doc = Document()
    doc.add_paragraph("Contents")
    paragraph = doc.add_paragraph()
    paragraph.add_run("Common Critiques of Value Extraction")
    paragraph.add_run().add_break()
    paragraph.add_run("What is Value?")
    paragraph.add_run().add_break()
    paragraph.add_run("Meet the Production Boundary")

    source_path = tmp_path / "legacy-inline-break-toc.docx"
    doc.save(source_path)

    with source_path.open("rb") as source_file:
        paragraphs, _, _, _, _, _, _ = extract_document_content_with_normalization_reports(
            source_file,
            app_config={"structure_recovery_enabled": True, "structure_recovery_mode": "legacy"},
        )

    assert [paragraph.structural_role for paragraph in paragraphs] == [
        "toc_header",
        "toc_entry",
        "toc_entry",
        "toc_entry",
    ]


def test_extract_document_content_from_docx_keeps_regular_body_run_clusters_as_one_paragraph():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Market failures")
    paragraph.add_run(" ")
    paragraph.add_run("matter here")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.text for paragraph in paragraphs] == ["Market failures matter here"]


def test_extract_document_content_from_docx_inserts_image_placeholders(tmp_path):
    image_path = tmp_path / "image.png"
    image_path.write_bytes(PNG_BYTES)
    doc = Document()
    doc.add_paragraph("Вступление")
    doc.add_paragraph().add_run().add_picture(str(image_path))
    doc.add_paragraph("Завершение")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, image_assets = extract_document_content_from_docx(buffer)

    assert [paragraph.text for paragraph in paragraphs] == [
        "Вступление",
        "[[DOCX_IMAGE_img_001]]",
        "Завершение",
    ]
    assert len(image_assets) == 1
    assert image_assets[0].image_id == "img_001"
    assert image_assets[0].placeholder == "[[DOCX_IMAGE_img_001]]"
    assert image_assets[0].width_emu is not None
    assert image_assets[0].height_emu is not None
    assert paragraphs[1].asset_id == "img_001"
    assert [paragraph.paragraph_id for paragraph in paragraphs] == ["p0000", "p0001", "p0002"]
    assert [paragraph.logical_index for paragraph in paragraphs] == [0, 1, 2]
    assert [paragraph.structural_role for paragraph in paragraphs] == ["body", "image", "body"]
    assert inspect_placeholder_integrity("\n\n".join(paragraph.text for paragraph in paragraphs), image_assets) == {
        "img_001": "ok"
    }


def test_extract_document_content_from_docx_populates_image_asset_payload_fields(tmp_path):
    image_path = tmp_path / "image.png"
    image_path.write_bytes(PNG_BYTES)
    doc = Document()
    doc.add_paragraph().add_run().add_picture(str(image_path), width=Inches(1.25))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, image_assets = extract_document_content_from_docx(buffer)

    assert [paragraph.text for paragraph in paragraphs] == ["[[DOCX_IMAGE_img_001]]"]
    assert len(image_assets) == 1
    asset = image_assets[0]
    assert asset.original_bytes == PNG_BYTES
    assert asset.position_index == 0
    assert asset.placeholder == "[[DOCX_IMAGE_img_001]]"
    assert asset.image_id == "img_001"
    assert asset.width_emu is not None
    assert asset.height_emu is not None


def test_direct_inline_image_survives_when_paragraph_also_has_textbox(tmp_path):
    # F11: a direct (non-textbox) inline image sharing a paragraph with a textbox
    # must not be dropped. Previously the whole paragraph's images were suppressed
    # whenever a textbox was present.
    image_path = tmp_path / "image.png"
    image_path.write_bytes(PNG_BYTES)
    doc = Document()
    host = doc.add_paragraph("До картинки ")
    host.add_run().add_picture(str(image_path))  # direct inline image
    _append_textbox_with_paragraphs(host, ["Текст во врезке"])  # text-only textbox
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, image_assets = extract_document_content_from_docx(buffer)

    assert len(image_assets) == 1
    asset = image_assets[0]
    assert asset.image_id == "img_001"
    assert asset.original_bytes == PNG_BYTES

    joined = "\n".join(paragraph.text for paragraph in paragraphs)
    assert asset.placeholder in joined  # the placeholder survives in the body
    assert "Текст во врезке" in joined  # textbox text is still restored


def test_textbox_interior_image_is_not_double_counted_with_direct_image(tmp_path):
    # F11 guard: a direct image and a textbox-INTERIOR image are each captured
    # exactly once — the direct pass takes the direct image, the restore pass
    # takes the interior one; neither is double-counted.
    image_path = tmp_path / "image.png"
    image_path.write_bytes(PNG_BYTES)
    doc = Document()
    host = doc.add_paragraph()
    host.add_run().add_picture(str(image_path))  # direct inline image
    interior_drawing = _detach_inline_drawing(doc, image_path)
    _append_textbox_with_interior_drawing(host, "Врезка", interior_drawing)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, image_assets = extract_document_content_from_docx(buffer)

    assert len(image_assets) == 2
    assert [asset.image_id for asset in image_assets] == ["img_001", "img_002"]
    joined = "\n".join(paragraph.text for paragraph in paragraphs)
    assert "[[DOCX_IMAGE_img_001]]" in joined
    assert "[[DOCX_IMAGE_img_002]]" in joined
    assert "Врезка" in joined


def test_nested_textbox_interior_image_is_extracted_exactly_once(tmp_path):
    # F17: a nested textbox (txbxContent inside txbxContent) containing a single
    # image must yield exactly one image asset. Previously the outer restore
    # paragraph and the inner restore paragraph both captured the deep blip
    # (the "inside any textbox" test could not distinguish nesting levels), so
    # the same image was emitted twice.
    image_path = tmp_path / "image.png"
    image_path.write_bytes(PNG_BYTES)
    doc = Document()
    host = doc.add_paragraph()
    interior_drawing = _detach_inline_drawing(doc, image_path)
    _append_nested_textbox_with_interior_drawing(host, interior_drawing)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, image_assets = extract_document_content_from_docx(buffer)

    assert len(image_assets) == 1
    assert image_assets[0].image_id == "img_001"
    joined = "\n".join(paragraph.text for paragraph in paragraphs)
    assert joined.count("[[DOCX_IMAGE_img_001]]") == 1


def test_extract_document_content_from_docx_captures_source_rect_forensics(tmp_path):
    image_path = tmp_path / "cropped-image.png"
    image_path.write_bytes(PNG_BYTES)

    doc = Document()
    run = doc.add_paragraph().add_run()
    run.add_picture(str(image_path), width=Inches(1.25))
    blip_fill = run._element.xpath(".//pic:blipFill")[0]
    source_rect = OxmlElement("a:srcRect")
    source_rect.set("l", "1250")
    source_rect.set("t", "2500")
    source_rect.set("r", "3750")
    source_rect.set("b", "5000")
    blip = blip_fill.xpath("./a:blip")[0]
    blip.addnext(source_rect)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    _, image_assets = extract_document_content_from_docx(buffer)

    assert image_assets[0].source_forensics["source_rect"] == {
        "l": 1250,
        "t": 2500,
        "r": 3750,
        "b": 5000,
    }
    assert "<wp:inline" in str(image_assets[0].source_forensics["drawing_container_xml"])


def test_extract_document_content_from_docx_tolerates_malformed_drawing_extent_metadata(tmp_path):
    image_path = tmp_path / "bad-extent-image.png"
    image_path.write_bytes(PNG_BYTES)

    doc = Document()
    run = doc.add_paragraph().add_run()
    run.add_picture(str(image_path), width=Inches(1.25))
    extent = run._element.xpath(".//wp:extent")[0]
    extent.set("cx", "oops")
    extent.set("cy", "-15")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    _, image_assets = extract_document_content_from_docx(buffer)

    assert len(image_assets) == 1
    assert image_assets[0].width_emu is None
    assert image_assets[0].height_emu is None


def test_extract_document_content_from_docx_ignores_non_numeric_source_rect_metadata(tmp_path):
    image_path = tmp_path / "partial-source-rect-image.png"
    image_path.write_bytes(PNG_BYTES)

    doc = Document()
    run = doc.add_paragraph().add_run()
    run.add_picture(str(image_path), width=Inches(1.25))
    blip_fill = run._element.xpath(".//pic:blipFill")[0]
    source_rect = OxmlElement("a:srcRect")
    source_rect.set("l", "1250")
    source_rect.set("t", "oops")
    source_rect.set("r", "")
    blip = blip_fill.xpath("./a:blip")[0]
    blip.addnext(source_rect)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    _, image_assets = extract_document_content_from_docx(buffer)

    assert len(image_assets) == 1
    assert image_assets[0].source_forensics["source_rect"] == {"l": 1250}


def test_inspect_placeholder_integrity_reports_unexpected_placeholders():
    asset = ImageAsset(
        image_id="img_001",
        placeholder="[[DOCX_IMAGE_img_001]]",
        original_bytes=PNG_BYTES,
        mime_type="image/png",
        position_index=0,
    )

    status_map = inspect_placeholder_integrity(
        "[[DOCX_IMAGE_img_001]]\n\n[[DOCX_IMAGE_img_999]]",
        [asset],
    )

    assert status_map == {
        "img_001": "ok",
        "unexpected:[[DOCX_IMAGE_img_999]]": "unexpected",
    }


def test_build_document_text_renders_word_numbered_and_bulleted_lists_as_markdown():
    doc = Document()
    doc.add_paragraph("Вступление")
    doc.add_paragraph("Первый пункт", style="List Number")
    doc.add_paragraph("Второй пункт", style="List Number")
    doc.add_paragraph("Подпункт", style="List Bullet 2")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["body", "list", "list", "list"]
    assert paragraphs[1].list_kind == "ordered"
    assert paragraphs[2].list_kind == "ordered"
    assert paragraphs[3].list_kind == "unordered"
    assert paragraphs[3].list_level == 1
    assert paragraphs[1].list_numbering_format is not None
    assert paragraphs[1].list_num_id is not None
    assert paragraphs[1].list_abstract_num_id is not None
    assert paragraphs[1].list_num_xml is not None
    assert paragraphs[1].list_abstract_num_xml is not None
    assert build_document_text(paragraphs) == (
        "Вступление\n\n"
        "1. Первый пункт\n\n"
        "1. Второй пункт\n\n"
        "  - Подпункт"
    )


def test_build_document_text_renders_nested_ordered_lists_with_markdown_safe_indent():
    nested_ordered = ParagraphUnit(text="Вложенный пункт", role="list", list_kind="ordered", list_level=1)
    deeper_unordered = ParagraphUnit(text="Глубже", role="list", list_kind="unordered", list_level=2)

    assert nested_ordered.rendered_text == "  1. Вложенный пункт"
    assert deeper_unordered.rendered_text == "    - Глубже"


def test_build_document_text_does_not_duplicate_existing_list_markers():
    paragraphs = [
        ParagraphUnit(text="1. Уже размеченный пункт", role="list", list_kind="ordered"),
        ParagraphUnit(text="- Уже размеченный маркер", role="list", list_kind="unordered"),
    ]

    assert build_document_text(paragraphs) == "1. Уже размеченный пункт\n\n- Уже размеченный маркер"


def test_public_paragraph_helper_exports_match_heading_and_font_detection():
    doc = Document()
    paragraph = doc.add_paragraph("Ключевой раздел")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.runs[0].font.size = Pt(15)

    assert paragraph_has_strong_heading_format(paragraph) is True
    assert resolve_effective_paragraph_font_size(paragraph) == 15.0


def test_build_document_text_renders_epigraph_and_attribution_as_blockquotes():
    paragraphs = [
        ParagraphUnit(text="Богатство заключается в свободе желаний.", role="body", structural_role="epigraph"),
        ParagraphUnit(text="— Эпиктет", role="body", structural_role="attribution"),
        ParagraphUnit(text="Обычный абзац.", role="body", structural_role="body"),
    ]

    assert build_document_text(paragraphs) == (
        "> Богатство заключается в свободе желаний.\n\n"
        "> — Эпиктет\n\n"
        "Обычный абзац."
    )


def test_build_document_text_does_not_duplicate_existing_blockquote_prefixes_for_epigraph_roles():
    paragraphs = [
        ParagraphUnit(text="> Уже оформленная цитата", role="body", structural_role="epigraph"),
        ParagraphUnit(text="> — Уже оформленный автор", role="body", structural_role="attribution"),
    ]

    assert build_document_text(paragraphs) == "> Уже оформленная цитата\n\n> — Уже оформленный автор"


def test_build_marker_wrapped_block_text_preserves_blockquote_rendering_for_epigraph_roles():
    paragraphs = [
        ParagraphUnit(text="Богатство заключается в свободе желаний.", role="body", structural_role="epigraph", paragraph_id="p0001"),
        ParagraphUnit(text="— Эпиктет", role="body", structural_role="attribution", paragraph_id="p0002"),
    ]

    assert build_marker_wrapped_block_text(paragraphs) == (
        "[[DOCX_PARA_p0001]]\n> Богатство заключается в свободе желаний.\n\n"
        "[[DOCX_PARA_p0002]]\n> — Эпиктет"
    )


def test_emdash_bullet_paragraphs_are_not_classified_as_list():
    """Em-dash (—) bullet in OOXML numbering is Russian typographic convention, not a real list."""
    buffer = _make_docx_with_emdash_bullet_numbering([
        "Американская торговая палата тратит на лоббизм больше всех.",
        "Эти многоплановые усилия — прерогатива местных сообществ.",
    ])

    paragraphs, _ = extract_document_content_from_docx(buffer)

    emdash_paras = [p for p in paragraphs if "торговая палата" in p.text or "многоплановые" in p.text]
    assert len(emdash_paras) == 2
    for paragraph in emdash_paras:
        assert paragraph.role == "body", f"Expected role='body', got '{paragraph.role}' for: {paragraph.text[:60]}"
        assert paragraph.list_kind is None
        assert paragraph.list_num_xml is None
        assert paragraph.list_abstract_num_xml is None


def test_emdash_bullet_paragraphs_render_without_list_markers():
    """Paragraphs demoted from em-dash bullet should render as plain text, no '- ' prefix."""
    buffer = _make_docx_with_emdash_bullet_numbering(["Цитата из книги."])

    paragraphs, _ = extract_document_content_from_docx(buffer)

    quote_para = [p for p in paragraphs if "Цитата" in p.text][0]
    assert quote_para.role == "body"
    text = build_document_text([quote_para])
    assert not text.startswith("- ")
    assert not text.startswith("— ")


def test_classify_paragraph_role_does_not_treat_emdash_prefix_as_list():
    """Text starting with '— ' should not be classified as list by text pattern."""
    from docxaicorrector.document.roles import classify_paragraph_role

    assert classify_paragraph_role("— Это прямая речь", "Body Text") == "body"
    assert classify_paragraph_role("— Цитата из книги", "Normal") == "body"
    assert classify_paragraph_role("- Пункт списка", "Body Text") == "list"
    assert classify_paragraph_role("• Маркированный пункт", "Normal") == "list"


def test_extract_document_content_from_docx_renders_title_and_outline_levels_as_markdown_headings():
    doc = Document()
    doc.add_paragraph("Название главы", style="Title")
    subheading = doc.add_paragraph("Подзаголовок")
    _set_outline_level(subheading, 1)
    doc.add_paragraph("Основной текст.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["heading", "heading", "body"]
    assert paragraphs[0].heading_level == 1
    assert paragraphs[1].heading_level == 2
    assert build_document_text(paragraphs) == "# Название главы\n\n## Подзаголовок\n\nОсновной текст."


def test_extract_document_content_from_docx_uses_inherited_outline_level_from_base_style():
    doc = Document()
    base_style = doc.styles.add_style("Base Outline Heading", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_outline_level(base_style, 2)
    derived_style = cast(Any, doc.styles.add_style("Derived Outline Heading", WD_STYLE_TYPE.PARAGRAPH))
    derived_style.base_style = base_style
    doc.add_paragraph("Наследуемый заголовок", style="Derived Outline Heading")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "heading"
    assert paragraphs[0].heading_level == 3
    assert build_document_text(paragraphs) == "### Наследуемый заголовок"


def test_extract_document_content_from_docx_recognizes_russian_heading_alias_style():
    doc = Document()
    doc.styles.add_style("Заголовок 3", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Русский заголовок", style="Заголовок 3")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "heading"
    assert paragraphs[0].heading_level == 3
    assert build_document_text(paragraphs) == "### Русский заголовок"


def test_extract_document_content_from_docx_keeps_tables_in_document_order():
    doc = Document()
    doc.add_paragraph("Перед таблицей")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Колонка A"
    table.cell(0, 1).text = "Колонка B"
    table.cell(1, 0).text = "Значение 1"
    table.cell(1, 1).text = "Значение 2"
    doc.add_paragraph("После таблицы")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["body", "table", "body"]
    # Kept tables are emitted as Pandoc-markdown (pipe) tables so the render
    # produces a real Word ``w:tbl`` (raw ``<table>`` HTML would be dropped).
    assert paragraphs[1].text.startswith("| Колонка A | Колонка B |")
    assert "Колонка A" in paragraphs[1].text
    assert build_document_text(paragraphs).startswith("Перед таблицей\n\n| Колонка A | Колонка B |")


def test_extract_document_content_from_docx_marks_caption_after_image(tmp_path):
    image_path = tmp_path / "docx_caption_image.png"
    image_path.write_bytes(PNG_BYTES)

    doc = Document()
    doc.add_paragraph().add_run().add_picture(str(image_path))
    doc.add_paragraph("Рис. 1. Подпись к изображению")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["image", "caption"]
    assert [paragraph.role_confidence for paragraph in paragraphs] == ["explicit", "adjacent"]
    assert paragraphs[0].asset_id == "img_001"
    assert paragraphs[1].attached_to_asset_id == "img_001"


def test_extract_document_content_from_docx_ai_first_hints_caption_after_image_without_mutating_role(tmp_path):
    image_path = tmp_path / "docx_caption_image_ai_first.png"
    image_path.write_bytes(PNG_BYTES)

    doc = Document()
    doc.add_paragraph().add_run().add_picture(str(image_path))
    doc.add_paragraph("Рис. 1. Подпись к изображению")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _, _, _, _, _, _ = document.extract_document_content_with_normalization_reports(
        buffer,
        app_config={"structure_recovery_enabled": True, "structure_recovery_mode": "ai_first"},
    )

    assert [paragraph.role for paragraph in paragraphs] == ["image", "body"]
    assert [paragraph.structural_role for paragraph in paragraphs] == ["image", "body"]
    assert paragraphs[0].asset_id == "img_001"
    assert paragraphs[1].attached_to_asset_id == "img_001"
    assert paragraphs[1].heuristic_role_hint == "caption"
    assert paragraphs[1].heuristic_heading_level_hint is None
    assert paragraphs[1].role_confidence == "heuristic"


def test_extract_document_content_from_docx_keeps_caption_style_after_image_even_when_format_looks_like_heading(tmp_path):
    image_path = tmp_path / "docx_caption_headingish_image.png"
    image_path.write_bytes(PNG_BYTES)

    doc = Document()
    doc.add_paragraph().add_run().add_picture(str(image_path))
    caption = doc.add_paragraph(style="Caption")
    caption_run = caption.add_run("Рисунок 1 Образец подписи")
    caption_run.bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["image", "caption"]
    assert paragraphs[1].role_confidence == "explicit"
    assert paragraphs[1].attached_to_asset_id == "img_001"
    assert paragraphs[1].heading_level is None
    assert build_document_text(paragraphs) == "[[DOCX_IMAGE_img_001]]\n\n**Рисунок 1 Образец подписи**"


def test_validate_docx_archive_rejects_zip_slip_paths():
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("../evil.txt", "boom")

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "подозрительные пути" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for zip-slip path")


def test_validate_docx_archive_rejects_oversized_archive_bytes(monkeypatch):
    monkeypatch.setattr(document, "MAX_DOCX_ARCHIVE_SIZE_BYTES", 5)

    try:
        document._validate_docx_archive(b"123456")
    except RuntimeError as exc:
        assert "превышает допустимый размер архива" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for oversized DOCX archive")


def test_validate_docx_archive_rejects_bad_zip_payload():
    try:
        document._validate_docx_archive(b"not-a-zip")
    except RuntimeError as exc:
        assert "поврежденный или неподдерживаемый DOCX-архив" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for invalid DOCX archive")


def test_validate_docx_archive_rejects_empty_archive():
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w"):
        pass

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "пустой DOCX-архив" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for empty DOCX archive")


def test_validate_docx_archive_rejects_too_many_entries(monkeypatch):
    monkeypatch.setattr(document, "MAX_DOCX_ENTRY_COUNT", 1)
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", "<w:document/>")

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "слишком много файлов" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for excessive DOCX entry count")


def test_validate_docx_archive_rejects_suspicious_compression_ratio(monkeypatch):
    monkeypatch.setattr(document, "MAX_DOCX_COMPRESSION_RATIO", 1)
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "A" * 2048)
        archive.writestr("word/document.xml", "B" * 2048)

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "подозрительно высокий коэффициент сжатия" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for suspicious DOCX compression ratio")


def test_validate_docx_archive_rejects_absolute_paths():
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("/word/document.xml", "boom")

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "абсолютные пути" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for absolute path in DOCX archive")


def test_validate_docx_archive_rejects_missing_content_types():
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("word/document.xml", "<w:document/>")

    try:
        document._validate_docx_archive(buffer.getvalue())
    except RuntimeError as exc:
        assert "отсутствует [Content_Types].xml" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for missing [Content_Types].xml")


def test_extract_document_content_from_docx_rejects_suspicious_uncompressed_archive(monkeypatch):
    monkeypatch.setattr(document, "MAX_DOCX_UNCOMPRESSED_SIZE_BYTES", 100)

    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "x" * 150)
        archive.writestr("word/document.xml", "<w:document />")
    buffer.seek(0)

    try:
        extract_document_content_from_docx(buffer)
    except RuntimeError as exc:
        assert "слишком велик после распаковки" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for suspiciously large uncompressed DOCX archive")


def test_read_uploaded_docx_bytes_preserves_original_cause(monkeypatch):
    failing_error = ValueError("bad upload")
    monkeypatch.setattr(document, "read_uploaded_file_bytes", lambda uploaded_file: (_ for _ in ()).throw(failing_error))

    try:
        document._read_uploaded_docx_bytes(object())
    except ValueError as exc:
        assert "Не удалось прочитать содержимое DOCX-файла" in str(exc)
        assert exc.__cause__ is failing_error
    else:
        raise AssertionError("Expected ValueError when uploaded DOCX bytes cannot be read")


def test_read_uploaded_docx_bytes_reuses_existing_docx_bytes_without_renormalizing(monkeypatch):
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w") as archive:
        archive.writestr("word/document.xml", "<w:document />")
    docx_bytes = buffer.getvalue()

    monkeypatch.setattr(document, "read_uploaded_file_bytes", lambda uploaded_file: docx_bytes)
    assert document._read_uploaded_docx_bytes(object()) == docx_bytes


def test_read_uploaded_docx_bytes_rejects_non_normalized_non_docx_input(monkeypatch):
    monkeypatch.setattr(document, "read_uploaded_file_bytes", lambda uploaded_file: b"legacy-binary")
    monkeypatch.setattr(document, "resolve_uploaded_filename", lambda uploaded_file: "legacy.doc")

    with pytest.raises(ValueError, match="Ожидался уже нормализованный DOCX-архив"):
        document._read_uploaded_docx_bytes(object())


def test_extract_document_content_from_docx_marks_caption_after_table():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Ячейка"
    doc.add_paragraph("Таблица 1. Подпись к таблице")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["table", "caption"]
    assert paragraphs[0].asset_id == "table_001"
    assert paragraphs[1].attached_to_asset_id == "table_001"


def test_extract_document_content_from_docx_reclassifies_heading_like_caption_after_table():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Ячейка"
    caption = doc.add_paragraph("Таблица 1 Итоговые показатели")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["table", "caption"]
    assert paragraphs[1].attached_to_asset_id == "table_001"
    assert paragraphs[1].heading_level is None


def test_extract_document_content_from_docx_does_not_treat_justified_body_text_as_heading():
    doc = Document()
    paragraph = doc.add_paragraph("Это обычный выровненный по ширине абзац без признаков заголовка")
    _set_raw_paragraph_alignment(paragraph, "both")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "body"
    assert paragraphs[0].heading_level is None


def test_extract_document_content_from_docx_recovers_mixed_format_heading_in_normal_style():
    doc = Document()
    paragraph = doc.add_paragraph(style="Normal")
    first_run = paragraph.add_run("Раздел 1:")
    first_run.bold = True
    paragraph.add_run(" Основные результаты")
    second_run = paragraph.add_run(" исследования")
    second_run.bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "heading"
    assert paragraphs[0].heading_level == 2
    assert build_document_text(paragraphs) == "## **Раздел 1:** Основные результаты** исследования**"


def test_extract_document_content_from_docx_detects_heading_from_inherited_style_alignment_with_text_signal():
    doc = Document()
    centered_style = doc.styles.add_style("Centered Heading Candidate", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_alignment(centered_style, "center")
    paragraph = doc.add_paragraph("Раздел 3 Основные результаты", style="Centered Heading Candidate")
    run = paragraph.runs[0]
    run.bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "heading"
    assert paragraphs[0].heading_source == "heuristic"
    assert paragraphs[0].heading_level == 2


def test_extract_document_content_from_docx_preserves_numbered_section_heading_with_list_metadata():
    doc = Document()
    paragraph = doc.add_paragraph("2. Systemic crises: frequency, types and geographical spread", style="List Bullet")
    paragraph.runs[0].bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "heading"
    assert paragraphs[0].heading_source == "heuristic"
    assert paragraphs[0].heading_level == 2
    assert paragraphs[0].list_kind is not None
    assert build_document_text(paragraphs) == "## **2. Systemic crises: frequency, types and geographical spread**"


def test_extract_document_content_from_docx_recovers_markdown_bold_numbered_heading_with_list_metadata():
    doc = Document()
    paragraph = doc.add_paragraph("**3. The Sovereign Debt Squeeze**", style="List Number")
    paragraph.runs[0].bold = False
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "heading"
    assert paragraphs[0].heading_source == "heuristic"
    assert paragraphs[0].heading_level == 2
    assert paragraphs[0].list_kind == "ordered"
    assert build_document_text(paragraphs).startswith("## ")
    assert "3. The Sovereign Debt Squeeze" in build_document_text(paragraphs)


def test_extract_document_content_from_docx_keeps_plain_numbered_list_item_as_list():
    doc = Document()
    doc.add_paragraph("This paragraph introduces an ordinary operational checklist.")
    paragraph = doc.add_paragraph("Pay invoices within thirty days after each delivery.", style="List Number")
    paragraph.runs[0].bold = False
    doc.add_paragraph("The checklist continues with other ordinary administrative steps.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [item.role for item in paragraphs] == ["body", "list", "body"]
    assert paragraphs[1].heading_level is None
    assert paragraphs[1].list_kind == "ordered"


def test_extract_document_content_from_docx_detects_heading_from_base_style_chain_alignment_with_text_signal():
    doc = Document()
    base_style = doc.styles.add_style("Centered Heading Base", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_alignment(base_style, "center")
    derived_style = cast(Any, doc.styles.add_style("Centered Heading Derived", WD_STYLE_TYPE.PARAGRAPH))
    derived_style.base_style = base_style
    paragraph = doc.add_paragraph("Глава 2 Методика", style="Centered Heading Derived")
    run = paragraph.runs[0]
    run.bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "heading"
    assert paragraphs[0].heading_source == "heuristic"
    assert paragraphs[0].heading_level == 1


def test_extract_document_content_from_docx_promotes_front_matter_display_title_to_h1():
    doc = Document()
    author = doc.add_paragraph("Mariana Mazzucato", style="Heading 1")
    author.runs[0].font.size = Pt(28)

    title = doc.add_paragraph("T H E VALUE O F E V E RY T H I NG")
    title.runs[0].font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Making and Taking in the Global Economy")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(18)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert paragraphs[0].text == "Mariana Mazzucato"
    assert paragraphs[0].role == "body"
    assert paragraphs[0].heading_level is None
    assert paragraphs[1].text == "T H E VALUE O F E V E RY T H I NG"
    assert paragraphs[1].role == "heading"
    assert paragraphs[1].heading_source == "heuristic"
    assert paragraphs[1].heading_level == 1
    assert paragraphs[2].role == "body"
    assert paragraphs[2].is_italic is True


def test_extract_document_content_from_docx_ai_first_projects_heuristic_heading_to_advisory_hint():
    doc = Document()
    base_style = doc.styles.add_style("Centered Heading Base AI First", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_alignment(base_style, "center")
    derived_style = cast(Any, doc.styles.add_style("Centered Heading Derived AI First", WD_STYLE_TYPE.PARAGRAPH))
    derived_style.base_style = base_style
    paragraph = doc.add_paragraph("Глава 2 Методика", style="Centered Heading Derived AI First")
    paragraph.runs[0].bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _, _, _, _, _, _ = document.extract_document_content_with_normalization_reports(
        buffer,
        app_config={"structure_recovery_enabled": True, "structure_recovery_mode": "ai_first"},
    )

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "body"
    assert paragraphs[0].structural_role == "body"
    assert paragraphs[0].heading_level is None
    assert paragraphs[0].heading_source is None
    assert paragraphs[0].heuristic_role_hint == "heading"
    assert paragraphs[0].heuristic_heading_level_hint == 1


def test_extract_document_content_from_docx_ai_first_keeps_explicit_heading_binding():
    doc = Document()
    heading = doc.add_paragraph("Chapter 1 Value", style="Heading 1")
    heading.runs[0].font.size = Pt(20)
    body = doc.add_paragraph("This opening paragraph provides ordinary narrative context after the heading.")
    body.runs[0].font.size = Pt(11)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _, _, _, _, _, _ = document.extract_document_content_with_normalization_reports(
        buffer,
        app_config={"structure_recovery_enabled": True, "structure_recovery_mode": "ai_first"},
    )

    assert paragraphs[0].role == "heading"
    assert paragraphs[0].structural_role == "heading"
    assert paragraphs[0].heading_level == 1
    assert paragraphs[0].heading_source == "explicit"
    assert paragraphs[0].heuristic_role_hint is None
    assert paragraphs[0].heuristic_heading_level_hint is None


def test_normalize_front_matter_display_title_ai_first_sets_hint_without_mutating_role():
    paragraphs = [
        ParagraphUnit(text="Mariana Mazzucato", role="body", structural_role="body", font_size_pt=28, is_bold=True),
        ParagraphUnit(text="T H E VALUE O F E V E RY T H I NG", role="body", structural_role="body", font_size_pt=28, is_bold=True),
        ParagraphUnit(text="Making and Taking in the Global Economy", role="body", structural_role="body", font_size_pt=18, is_italic=True),
    ]

    document_extraction.normalize_front_matter_display_title(
        paragraphs,
        structure_recovery_enabled=True,
        structure_recovery_mode="ai_first",
    )

    assert paragraphs[1].role == "body"
    assert paragraphs[1].structural_role == "body"
    assert paragraphs[1].heading_level is None
    assert paragraphs[1].heading_source is None
    assert paragraphs[1].heuristic_role_hint == "heading"
    assert paragraphs[1].heuristic_heading_level_hint == 1


def test_normalize_front_matter_display_title_ai_first_does_not_demote_existing_heading_siblings():
    paragraphs = [
        ParagraphUnit(text="Mariana Mazzucato", role="heading", structural_role="heading", heading_level=2, heading_source="heuristic", font_size_pt=28, is_bold=True),
        ParagraphUnit(text="T H E VALUE O F E V E RY T H I NG", role="body", structural_role="body", font_size_pt=28, is_bold=True),
        ParagraphUnit(text="Making and Taking in the Global Economy", role="body", structural_role="body", font_size_pt=18, is_italic=True),
    ]

    document_extraction.normalize_front_matter_display_title(
        paragraphs,
        structure_recovery_enabled=True,
        structure_recovery_mode="ai_first",
    )

    assert paragraphs[0].role == "heading"
    assert paragraphs[0].structural_role == "heading"
    assert paragraphs[0].heading_level == 2
    assert paragraphs[0].heading_source == "heuristic"
    assert paragraphs[1].heuristic_role_hint == "heading"
    assert paragraphs[1].heuristic_heading_level_hint == 1


def test_extract_document_content_from_docx_keeps_true_structural_h1_when_no_cover_title_exists():
    doc = Document()
    heading = doc.add_paragraph("Chapter 1 Value", style="Heading 1")
    heading.runs[0].font.size = Pt(20)
    body = doc.add_paragraph("This opening paragraph provides ordinary narrative context after the heading.")
    body.runs[0].font.size = Pt(11)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert paragraphs[0].role == "heading"
    assert paragraphs[0].heading_level == 1
    assert paragraphs[0].text == "Chapter 1 Value"


def test_extract_document_content_from_docx_paragraph_alignment_override_beats_inherited_center_for_heading_detection():
    doc = Document()
    base_style = doc.styles.add_style("Centered Heading Base Override", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_alignment(base_style, "center")
    derived_style = cast(Any, doc.styles.add_style("Centered Heading Derived Override", WD_STYLE_TYPE.PARAGRAPH))
    derived_style.base_style = base_style
    paragraph = doc.add_paragraph("Краткое описание установки", style="Centered Heading Derived Override")
    _set_raw_paragraph_alignment(paragraph, "left")
    run = paragraph.runs[0]
    run.bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "body"
    assert paragraphs[0].heading_level is None
    assert paragraphs[0].role_confidence == "heuristic"


def test_extract_document_content_from_docx_does_not_promote_centered_bold_body_without_text_signal():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Краткое описание установки")
    run.bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "body"
    assert paragraphs[0].heading_level is None
    assert paragraphs[0].role_confidence == "heuristic"


def test_extract_document_content_from_docx_does_not_promote_inherited_centered_body_without_text_signal():
    doc = Document()
    centered_style = doc.styles.add_style("Centered Body Candidate", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_alignment(centered_style, "center")
    paragraph = doc.add_paragraph("Краткое описание установки", style="Centered Body Candidate")
    run = paragraph.runs[0]
    run.bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].role == "body"
    assert paragraphs[0].heading_level is None
    assert paragraphs[0].role_confidence == "heuristic"


def test_extract_document_content_from_docx_promotes_short_larger_subheading_between_body_paragraphs():
    doc = Document()

    first_paragraph = doc.add_paragraph(
        "Богатство может означать деньги, свободу выбора, устойчивость и доступ к возможностям, "
        "которые человек иначе не получил бы."
    )
    first_paragraph.runs[0].font.size = Pt(11)

    heading_paragraph = doc.add_paragraph("Переосмысление богатства")
    heading_paragraph.runs[0].font.size = Pt(14)

    third_paragraph = doc.add_paragraph(
        "Богатство - это не только владение активами, но и способность направлять время, внимание "
        "и отношения к осмысленным целям."
    )
    third_paragraph.runs[0].font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["body", "heading", "body"]
    assert paragraphs[1].heading_source == "heuristic"
    assert paragraphs[1].heading_level == 2
    assert build_document_text(paragraphs) == (
        "Богатство может означать деньги, свободу выбора, устойчивость и доступ к возможностям, "
        "которые человек иначе не получил бы.\n\n"
        "## Переосмысление богатства\n\n"
        "Богатство - это не только владение активами, но и способность направлять время, внимание "
        "и отношения к осмысленным целям."
    )


def test_extract_document_content_from_docx_promotes_very_short_subheading_between_body_paragraphs_without_larger_font():
    doc = Document()

    doc.add_paragraph(
        "Привлекательность лотерейных билетов с крупными призами отчасти объясняется мечтами о переменах и доступе к новым "
        "возможностям."
    )
    doc.add_paragraph("Переосмысление богатства")
    doc.add_paragraph(
        "Богатство - это то, чего мы все хотим, но его значение зависит не только от денег, а еще и от устойчивости, "
        "свободы выбора и качества связей."
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["body", "heading", "body"]
    assert paragraphs[1].heading_source == "heuristic"
    assert paragraphs[1].heading_level == 2


def test_promote_short_standalone_headings_ai_first_sets_hint_without_mutating_role():
    paragraphs = [
        ParagraphUnit(
            text="Привлекательность лотерейных билетов с крупными призами отчасти объясняется мечтами о переменах и доступе к новым возможностям.",
            role="body",
            structural_role="body",
            font_size_pt=11,
        ),
        ParagraphUnit(
            text="Переосмысление богатства",
            role="body",
            structural_role="body",
            font_size_pt=11,
        ),
        ParagraphUnit(
            text="Богатство - это то, чего мы все хотим, но его значение зависит не только от денег, а еще и от устойчивости, свободы выбора и качества связей.",
            role="body",
            structural_role="body",
            font_size_pt=11,
        ),
    ]

    document.promote_short_standalone_headings(
        paragraphs,
        structure_recovery_enabled=True,
        structure_recovery_mode="ai_first",
    )

    assert [paragraph.role for paragraph in paragraphs] == ["body", "body", "body"]
    assert paragraphs[1].structural_role == "body"
    assert paragraphs[1].heading_level is None
    assert paragraphs[1].heading_source is None
    assert paragraphs[1].heuristic_role_hint == "heading"
    assert paragraphs[1].heuristic_heading_level_hint == 2


def test_extract_document_content_from_docx_does_not_promote_very_short_sentence_with_terminal_period():
    doc = Document()

    doc.add_paragraph(
        "Привлекательность лотерейных билетов с крупными призами отчасти объясняется мечтами о переменах и доступе к новым "
        "возможностям."
    )
    doc.add_paragraph("Новое богатство.")
    doc.add_paragraph(
        "Богатство - это то, чего мы все хотим, но его значение зависит не только от денег, а еще и от устойчивости, "
        "свободы выбора и качества связей."
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["body", "body", "body"]


def test_promote_short_standalone_headings_does_not_override_ai_classified_body_heading_candidate():
    paragraphs = [
        ParagraphUnit(
            text="Привлекательность лотерейных билетов с крупными призами объясняется мечтами о переменах и доступе к новым возможностям.",
            role="body",
            source_index=0,
            font_size_pt=11.0,
        ),
        ParagraphUnit(
            text="Переосмысление богатства",
            role="body",
            structural_role="body",
            role_confidence="ai",
            source_index=1,
            font_size_pt=14.0,
        ),
        ParagraphUnit(
            text="Богатство зависит не только от денег, но и от устойчивости, свободы выбора и качества связей.",
            role="body",
            source_index=2,
            font_size_pt=11.0,
        ),
    ]

    document.promote_short_standalone_headings(paragraphs)

    assert [paragraph.role for paragraph in paragraphs] == ["body", "body", "body"]
    assert paragraphs[1].role_confidence == "ai"
    assert paragraphs[1].heading_source is None
    assert paragraphs[1].heading_level is None


def test_promote_short_standalone_headings_does_not_override_ai_structural_attribution():
    paragraphs = [
        ParagraphUnit(
            text="Богатство может означать деньги, свободу выбора и доступ к возможностям.",
            role="body",
            source_index=0,
            font_size_pt=11.0,
        ),
        ParagraphUnit(
            text="ЭПИКТЕТ",
            role="body",
            structural_role="attribution",
            role_confidence="ai",
            source_index=1,
            font_size_pt=16.0,
        ),
        ParagraphUnit(
            text="Следующий абзац продолжает мысль и даёт обычный текстовый контекст для эвристического паттерна.",
            role="body",
            source_index=2,
            font_size_pt=11.0,
        ),
    ]

    document.promote_short_standalone_headings(paragraphs)

    assert paragraphs[1].role == "body"
    assert paragraphs[1].structural_role == "attribution"
    assert paragraphs[1].role_confidence == "ai"
    assert paragraphs[1].heading_source is None


def test_promote_short_standalone_headings_does_not_promote_centered_all_caps_attribution_after_italic_quote():
    paragraphs = [
        ParagraphUnit(
            text="Богатство заключается не в том, чтобы иметь много имущества, а в том, чтобы иметь мало желаний.",
            role="body",
            source_index=0,
            is_italic=True,
            font_size_pt=11.0,
        ),
        ParagraphUnit(
            text="ЭПИКТЕТ",
            role="body",
            source_index=1,
            paragraph_alignment="center",
            font_size_pt=16.0,
        ),
        ParagraphUnit(
            text="Следующий абзац продолжает обычный текст и задаёт body-контекст для эвристики.",
            role="body",
            source_index=2,
            font_size_pt=11.0,
        ),
    ]

    document.promote_short_standalone_headings(paragraphs)

    assert paragraphs[1].role == "body"
    assert paragraphs[1].structural_role == "body"
    assert paragraphs[1].heading_source is None


def test_promote_short_standalone_headings_still_promotes_legitimate_centered_heading_after_italic_context():
    paragraphs = [
        ParagraphUnit(
            text="Богатство заключается не только в накоплении средств, но и в умении выстраивать устойчивые связи и долгосрочные цели.",
            role="body",
            source_index=0,
            is_italic=True,
            font_size_pt=11.0,
        ),
        ParagraphUnit(
            text="Переосмысление богатства",
            role="body",
            source_index=1,
            paragraph_alignment="center",
            font_size_pt=16.0,
        ),
        ParagraphUnit(
            text="Следующий абзац продолжает обычный основной текст и даёт достаточный body-контекст для эвристического промоушена.",
            role="body",
            source_index=2,
            font_size_pt=11.0,
        ),
    ]

    document.promote_short_standalone_headings(paragraphs)

    assert paragraphs[1].role == "heading"
    assert paragraphs[1].heading_source == "heuristic"
    assert paragraphs[1].heading_level == 2


def test_extract_document_content_from_docx_keeps_inherited_centered_caption_after_table():
    doc = Document()
    centered_style = doc.styles.add_style("Centered Caption Candidate", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_alignment(centered_style, "center")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Ячейка"
    paragraph = doc.add_paragraph("Таблица 1 Итоговые показатели", style="Centered Caption Candidate")
    run = paragraph.runs[0]
    run.bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["table", "caption"]
    assert paragraphs[1].attached_to_asset_id == "table_001"
    assert paragraphs[1].heading_level is None


def test_extract_document_content_from_docx_preserves_hyperlinks_tabs_and_inline_emphasis():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("До ")
    bold_run = paragraph.add_run("важно")
    bold_run.bold = True
    paragraph.add_run(" и ")
    italic_run = paragraph.add_run("курсив")
    italic_run.italic = True
    tab_run = paragraph.add_run()
    _append_tab(tab_run)
    _add_hyperlink(paragraph, "ссылка", "https://example.com")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].text == "До **важно** и *курсив*\t[ссылка](https://example.com)"


def _make_fld_char(char_type: str):
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), char_type)
    return fld_char


def test_extract_document_content_from_docx_renders_underline_superscript_and_subscript():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Вода ")
    underline_run = paragraph.add_run("важно")
    underline_run.underline = True
    paragraph.add_run(", формула H")
    subscript_run = paragraph.add_run("2")
    subscript_run.font.subscript = True
    paragraph.add_run("O и E=mc")
    superscript_run = paragraph.add_run("2")
    superscript_run.font.superscript = True
    paragraph.add_run(".")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].text == "Вода <u>важно</u>, формула H<sub>2</sub>O и E=mc<sup>2</sup>."


def test_apply_run_markdown_combines_bold_and_italic():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("сильно")
    run.bold = True
    run.italic = True

    assert document_extraction._apply_run_markdown("сильно", run._element) == "***сильно***"


def test_apply_run_markdown_ignores_explicitly_disabled_toggle_properties():
    """<w:b w:val="0"/>, <w:i w:val="0"/>, <w:u w:val="none"/> mean formatting OFF."""
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("обычный")
    run.bold = False
    run.italic = False
    run.underline = False

    assert document_extraction._apply_run_markdown("обычный", run._element) == "обычный"


def test_extract_document_content_from_docx_does_not_emphasize_disabled_toggle_runs():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Совершенно обычный абзац без выделения символов.")
    run.bold = False
    run.italic = False
    run.underline = False

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert paragraphs[0].text == "Совершенно обычный абзац без выделения символов."
    assert "**" not in paragraphs[0].text
    assert "<u>" not in paragraphs[0].text


def test_extract_run_text_excludes_field_instruction_codes():
    """<w:instrText> field codes (HYPERLINK/PAGEREF/TOC) must not leak into body text."""
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Смотрите ")
    begin = paragraph.add_run()
    begin._element.append(_make_fld_char("begin"))
    instruction = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.text = ' HYPERLINK "http://example.com" '
    instruction._element.append(instr_text)
    separate = paragraph.add_run()
    separate._element.append(_make_fld_char("separate"))
    paragraph.add_run("сайт примера")
    end = paragraph.add_run()
    end._element.append(_make_fld_char("end"))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert "HYPERLINK" not in paragraphs[0].text
    assert "http://example.com" not in paragraphs[0].text
    assert "сайт примера" in paragraphs[0].text


def test_extract_run_text_excludes_tracked_change_deletions():
    """<w:delText> tracked-change deletions must not resurface in the extracted text."""
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Оставленный текст ")
    deletion = paragraph.add_run()
    del_text = OxmlElement("w:delText")
    del_text.text = "удалённый фрагмент"
    deletion._element.append(del_text)
    paragraph.add_run("и хвост.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert len(paragraphs) == 1
    assert "удалённый фрагмент" not in paragraphs[0].text
    assert "Оставленный текст" in paragraphs[0].text
    assert "и хвост." in paragraphs[0].text


def test_extract_document_content_from_docx_handles_empty_document_without_crashing():
    doc = Document()
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, image_assets = extract_document_content_from_docx(buffer)

    assert paragraphs == []
    assert image_assets == []


def test_extract_document_content_from_docx_skips_blank_and_whitespace_only_paragraphs():
    doc = Document()
    doc.add_paragraph("Начало")
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("Конец")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.text for paragraph in paragraphs] == ["Начало", "Конец"]


def test_extract_document_content_from_docx_tolerates_nested_tables():
    doc = Document()
    outer_table = doc.add_table(rows=1, cols=1)
    outer_cell = outer_table.cell(0, 0)
    outer_cell.paragraphs[0].add_run("Внешняя ячейка")
    nested_table = outer_cell.add_table(rows=1, cols=1)
    nested_table.cell(0, 0).text = "Вложенная ячейка"
    doc.add_paragraph("После таблицы")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    table_paragraphs = [paragraph for paragraph in paragraphs if paragraph.role == "table"]
    assert len(table_paragraphs) == 1
    assert "Внешняя ячейка" in table_paragraphs[0].text
    assert any(paragraph.text == "После таблицы" for paragraph in paragraphs)


def test_extract_document_content_from_docx_renders_nested_list_levels():
    doc = Document()
    doc.add_paragraph("Верхний пункт", style="List Bullet")
    doc.add_paragraph("Вложенный пункт", style="List Bullet 2")
    doc.add_paragraph("Глубокий пункт", style="List Bullet 3")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    paragraphs, _ = extract_document_content_from_docx(buffer)

    assert [paragraph.role for paragraph in paragraphs] == ["list", "list", "list"]
    assert [paragraph.list_level for paragraph in paragraphs] == [0, 1, 2]
    assert all(paragraph.list_kind == "unordered" for paragraph in paragraphs)


def test_boundary_normalization_resolver_uses_passed_app_config(monkeypatch):
    # F15: the boundary-normalization resolver must honour an explicitly passed
    # app_config (the one threaded from extraction), not the GLOBAL load_app_config().
    import docxaicorrector.core.config as config

    global_config = {
        "paragraph_boundary_normalization_enabled": True,
        "paragraph_boundary_normalization_mode": "high_only",
        "paragraph_boundary_normalization_save_debug_artifacts": True,
    }
    monkeypatch.setattr(config, "load_app_config", lambda: global_config)

    # No app_config => falls back to the global (proves the global path still works).
    global_mode, _ = document_extraction._resolve_paragraph_boundary_normalization_settings()
    assert global_mode == "high_only"

    # Passed app_config overrides the global (disabled => "off").
    override_config = {"paragraph_boundary_normalization_enabled": False}
    override_mode, _ = document_extraction._resolve_paragraph_boundary_normalization_settings(override_config)
    assert override_mode == "off"
    assert override_mode != global_mode


def test_relation_normalization_resolver_uses_passed_app_config(monkeypatch):
    # F15: relation-normalization resolver honours the passed app_config.
    import docxaicorrector.core.config as config

    global_config = {"relation_normalization_enabled": True}
    monkeypatch.setattr(config, "load_app_config", lambda: global_config)

    global_enabled = document_extraction._resolve_relation_normalization_settings()[0]
    assert global_enabled is True

    override_enabled = document_extraction._resolve_relation_normalization_settings(
        {"relation_normalization_enabled": False}
    )[0]
    assert override_enabled is False
    assert override_enabled != global_enabled


def test_ai_review_resolver_uses_passed_app_config(monkeypatch):
    # F15: paragraph-boundary AI-review resolver honours the passed app_config.
    import docxaicorrector.core.config as config

    global_config = {
        "paragraph_boundary_ai_review_enabled": True,
        "paragraph_boundary_ai_review_mode": "review_only",
    }
    monkeypatch.setattr(config, "load_app_config", lambda: global_config)
    monkeypatch.setattr(config, "get_model_role_value", lambda app_config, role: "openai:gpt-test")

    global_mode = document_extraction._resolve_paragraph_boundary_ai_review_settings()[1]
    assert global_mode == "review_only"

    override_mode = document_extraction._resolve_paragraph_boundary_ai_review_settings(
        {"paragraph_boundary_ai_review_enabled": False}
    )[1]
    assert override_mode == "off"
    assert override_mode != global_mode
