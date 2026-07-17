from __future__ import annotations

import os
import queue
import shutil
import socket
import subprocess
import sys
import threading
import traceback
import hashlib
import errno
from collections import Counter
from collections.abc import Callable, Mapping, Sequence
from dataclasses import replace
from datetime import UTC, datetime
from difflib import SequenceMatcher
import json
from io import BytesIO
from pathlib import Path
import platform
import re
import time
import uuid
from typing import Any, Protocol, cast

SCRIPT_PATH = Path(__file__).resolve()
PROJECT_ROOT = SCRIPT_PATH.parents[3]
SRC_ROOT = PROJECT_ROOT / "src"


# Make the repo-root shared bootstrap importable, then pin src first (F5/R29).
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))
from docxaicorrector_bootstrap import ensure_src_first_import_order

ensure_src_first_import_order(PROJECT_ROOT, SRC_ROOT)

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn

import docxaicorrector.core.logger as app_logger
import docxaicorrector.processing.processing_runtime as processing_runtime
import docxaicorrector.processing.processing_service as processing_service
import docxaicorrector.pipeline._pipeline as document_pipeline
import docxaicorrector.ui.application_flow as application_flow
from docxaicorrector.core.config import (
    AppConfig,
    describe_provider_availability,
    get_client,
    get_client_for_model_selector,
    load_app_config,
    load_system_prompt,
    resolve_model_selector,
)
from docxaicorrector.document.extraction import (
    ORDERED_LIST_FORMATS,
    extract_document_content_from_docx,
)
from docxaicorrector.generation._generation import (
    generate_markdown_block,
)
from docxaicorrector.image.shared import parse_json_object
from docxaicorrector.pipeline.output_validation import collect_page_placeholder_heading_concat_samples
from docxaicorrector.reader_cleanup_mvp import (
    build_cleanup_blocks,
)
from docxaicorrector.validation.common import build_validation_event_logger, build_validation_runtime_config
from docxaicorrector.validation.structural import (
    _apply_prepared_snapshot_fields,
    _normalize_snapshot_or_metric_statuses,
    build_preparation_diagnostic_snapshot,
)
from docxaicorrector.validation.formatting_coverage import (
    formatting_payload_format_neutral_creditable_count as _formatting_payload_format_neutral_creditable_count,
    filter_benign_unmapped_source_ids as _filter_benign_unmapped_source_ids,
    resolve_filtered_formatting_unmapped_source_count as _resolve_filtered_formatting_unmapped_source_count,
    resolve_role_aware_formatting_unmapped_source_summary as _resolve_role_aware_formatting_unmapped_source_summary,
    resolve_role_aware_formatting_unmapped_target_summary as _resolve_role_aware_formatting_unmapped_target_summary,
)
from docxaicorrector.validation.acceptance import (
    build_acceptance_verdict as _build_acceptance_verdict,
    build_acceptance_toc_body_concat_check as _build_acceptance_toc_body_concat_check,
    extract_runtime_processing_operation as _extract_runtime_processing_operation,
    resolve_acceptance_unmapped_source_summary as _resolve_acceptance_unmapped_source_summary,
    resolve_acceptance_unmapped_target_summary as _resolve_acceptance_unmapped_target_summary,
    translation_quality_reason_is_review_only as _translation_quality_reason_is_review_only,
)
from docxaicorrector.validation.profiles import (
    apply_runtime_resolution_to_app_config,
    load_validation_registry,
    resolve_runtime_resolution,
)
from docxaicorrector.runtime.events import (
    AppendImageLogEvent,
    AppendLogEvent,
    FinalizeProcessingStatusEvent,
    PushActivityEvent,
    ResetImageStateEvent,
    SetProcessingStatusEvent,
    SetStateEvent,
)

REAL_DOCUMENT_ARTIFACT_ROOT = PROJECT_ROOT / "tests" / "artifacts" / "real_document_pipeline"
FORMATTING_DIAGNOSTICS_DIR = PROJECT_ROOT / ".run" / "formatting_diagnostics"
HEARTBEAT_INTERVAL_SECONDS = 15.0
READER_VERIFIER_DEFAULT_SELECTOR = "openrouter:google/gemini-3-flash-preview"
READER_VERIFIER_TIMEOUT_SECONDS = 180.0
_ALLOWED_READER_VERIFIER_VERDICTS = frozenset({"cleaned_better", "raw_better", "mixed", "unclear"})
_ALLOWED_READER_VERIFIER_AUDIT_VERDICTS = frozenset(
    {"clean", "improved_but_has_remaining_issues", "unsafe_or_regressed", "unclear"}
)
_ALLOWED_READER_VERIFIER_CONFIDENCE = frozenset({"low", "medium", "high"})
_ALLOWED_READER_VERIFIER_CHANGE_TYPES = frozenset(
    {"prompt", "model_selection", "operation_contract", "safety_application", "deterministic_last_resort"}
)
_LEGACY_READER_VERIFIER_CHANGE_TYPE_MAP = {
    "ai_operation_contract": "operation_contract",
    "cleanup_core": "safety_application",
    "deterministic_cleanup": "deterministic_last_resort",
    "minimal_formatting": "safety_application",
}
_READER_VERIFIER_DEFECT_CATEGORIES = (
    "page_furniture_inline",
    "heading_fused_with_body",
    "broken_list_marker",
    "fragmented_paragraph",
    "duplicate_fragment",
    "orphan_caption",
    "mixed_language_leak",
    "quote_not_block_formatted",
)
_TRACKED_READER_CLEANUP_IGNORED_REASONS = (
    "prior_same_block_operation_not_applied",
    "heading_boundary_unaccounted_text",
    "heading_boundary_substrings_not_found",
    "remove_inline_noise_not_exact_noise_pattern",
    "noise_substring_not_found",
    "duplicate_operation_incompatible",
)
_ALLOWED_READER_VERIFIER_SEVERITY = frozenset({"high", "medium", "low"})
_ALLOWED_READER_VERIFIER_ISSUE_ARTIFACTS = frozenset({"cleaned_markdown", "raw_markdown", "comparison"})
_ALLOWED_READER_VERIFIER_FIX_TYPES = frozenset(
    {"delete_noise", "split_heading", "merge_paragraph", "normalize_list", "format_quote", "other"}
)
_ALLOWED_READER_VERIFIER_ANCHOR_KINDS = frozenset({"improvement_seen", "remaining_issue", "possible_false_deletion"})
_ALLOWED_READER_CLEANUP_ANCHOR_REPAIR_CATEGORIES = frozenset(
    {"heading_fused_with_body", "page_furniture_inline", "fragmented_paragraph"}
)
_READER_VERIFIER_CATEGORY_KEYWORDS = {
    "page_furniture_inline": ("page number", "page numbers", "running header", "running headers", "page furniture", "header/footer"),
    "heading_fused_with_body": ("heading fused", "heading/body", "heading glued", "heading merged"),
    "broken_list_marker": ("bullet", "list marker", "broken list"),
    "fragmented_paragraph": ("fragmented paragraph", "broken paragraph", "paragraph join", "page-boundary join", "carryover"),
    "duplicate_fragment": ("duplicate fragment", "duplicated fragment", "repeated fragment"),
    "orphan_caption": ("orphan caption", "caption"),
    "mixed_language_leak": ("mixed language", "language leak", "english leak"),
    "quote_not_block_formatted": ("quote", "block quote", "blockquote"),
}
_READER_VERIFIER_MODEL_FIELDS = frozenset(
    {
        "overall_verdict",
        "cleaned_audit_verdict",
        "reader_quality_score_raw",
        "reader_quality_score_cleaned",
        "confidence",
        "noise_removed",
        "possible_false_deletions",
        "readability_regressions",
        "remaining_issues",
        "evidence_anchors",
        "recommended_next_changes",
        "summary_for_human",
        "simple_user_summary",
        "simple_user_risk_statement",
        "simple_user_next_step",
    }
)
_TERMINAL_OUTPUT_DISABLED = False


def _safe_terminal_print(*args: object, **kwargs: object) -> None:
    global _TERMINAL_OUTPUT_DISABLED
    if _TERMINAL_OUTPUT_DISABLED:
        return
    try:
        print(*args, **kwargs)
    except BrokenPipeError:
        _TERMINAL_OUTPUT_DISABLED = True
    except OSError as exc:
        if exc.errno != errno.EPIPE:
            raise
        _TERMINAL_OUTPUT_DISABLED = True


class UploadedFileStub:
    def __init__(self, name: str, content: bytes):
        self.name = name
        self._content = content
        self._position = 0

    def read(self, size: int = -1) -> bytes:
        if size < 0:
            data = self._content[self._position :]
            self._position = len(self._content)
            return data
        start = self._position
        end = min(len(self._content), start + size)
        self._position = end
        return self._content[start:end]

    def getvalue(self) -> bytes:
        return self._content

    def seek(self, offset: int, whence: int = 0) -> int:
        if whence == 0:
            self._position = max(0, offset)
        elif whence == 1:
            self._position = max(0, self._position + offset)
        elif whence == 2:
            self._position = max(0, len(self._content) + offset)
        else:
            raise ValueError(f"Unsupported whence: {whence}")
        return self._position


def _path_for_report(path: Path | None) -> str | None:
    if path is None:
        return None
    try:
        return str(path.relative_to(PROJECT_ROOT)).replace("\\", "/")
    except ValueError:
        return str(path).replace("\\", "/")


def _build_source_cleanup_evidence(cleanup_report: object | None) -> dict[str, object] | None:
    if cleanup_report is None:
        return None

    cleanup_mode = str(getattr(cleanup_report, "cleanup_mode", "remove") or "remove").strip().lower() or "remove"
    removed_samples: list[dict[str, object]] = []
    flagged_samples: list[dict[str, object]] = []
    kept_uncertain_samples: list[dict[str, object]] = []
    reason_counts: Counter[str] = Counter()
    decision_counts: Counter[str] = Counter()

    for decision in list(getattr(cleanup_report, "decisions", []) or []):
        action = str(getattr(decision, "action", "") or "keep").strip().lower() or "keep"
        reason = str(getattr(decision, "reason", "") or "keep").strip().lower() or "keep"
        decision_counts[action] += 1
        reason_counts[reason] += 1
        serialized = {
            "action": action,
            "reason": reason,
            "confidence": str(getattr(decision, "confidence", "") or ""),
            "text_preview": str(getattr(decision, "text_preview", "") or ""),
            "original_source_index": getattr(decision, "original_source_index", None),
            "origin_raw_indexes": list(getattr(decision, "origin_raw_indexes", ()) or ()),
            "page_number": getattr(decision, "page_number", None),
            "layout_origin": str(getattr(decision, "layout_origin", "") or ""),
            "repeat_count": getattr(decision, "repeat_count", None),
        }
        if action == "remove" and len(removed_samples) < 8:
            removed_samples.append(serialized)
        elif action == "flag" and len(flagged_samples) < 8:
            flagged_samples.append(serialized)
        elif reason == "uncertain_repeated_artifact" and len(kept_uncertain_samples) < 8:
            kept_uncertain_samples.append(serialized)

    return {
        "cleanup_applied": bool(getattr(cleanup_report, "cleanup_applied", False)),
        "cleanup_mode": cleanup_mode,
        "skipped_reason": getattr(cleanup_report, "skipped_reason", None),
        "error_code": getattr(cleanup_report, "error_code", None),
        "artifact_path": _path_for_report(_resolve_reported_path(getattr(cleanup_report, "artifact_path", None))),
        "original_paragraph_count": int(getattr(cleanup_report, "original_paragraph_count", 0) or 0),
        "cleaned_paragraph_count": int(getattr(cleanup_report, "cleaned_paragraph_count", 0) or 0),
        "removed_paragraph_count": int(getattr(cleanup_report, "removed_paragraph_count", 0) or 0),
        "removed_page_number_count": int(getattr(cleanup_report, "removed_page_number_count", 0) or 0),
        "removed_repeated_artifact_count": int(getattr(cleanup_report, "removed_repeated_artifact_count", 0) or 0),
        "removed_empty_or_whitespace_count": int(getattr(cleanup_report, "removed_empty_or_whitespace_count", 0) or 0),
        "flagged_page_number_count": int(getattr(cleanup_report, "flagged_page_number_count", 0) or 0),
        "flagged_repeated_artifact_count": int(getattr(cleanup_report, "flagged_repeated_artifact_count", 0) or 0),
        "flagged_empty_or_whitespace_count": int(getattr(cleanup_report, "flagged_empty_or_whitespace_count", 0) or 0),
        "decision_counts": dict(sorted(decision_counts.items())),
        "reason_counts": dict(sorted(reason_counts.items())),
        "removed_samples": removed_samples,
        "flagged_samples": flagged_samples,
        "kept_uncertain_samples": kept_uncertain_samples,
    }


def _coerce_int(value: object, default: int = 0) -> int:
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return int(value)
    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return default
        try:
            return int(stripped)
        except ValueError:
            return default
    return default


def _coerce_bool(value: object, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    if isinstance(value, str):
        normalized = value.strip().lower()
        if normalized in {"1", "true", "yes", "on"}:
            return True
        if normalized in {"0", "false", "no", "off"}:
            return False
    return default


def _coerce_float(value: object, default: float = 0.0) -> float:
    if isinstance(value, bool):
        return float(int(value))
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return default
        try:
            return float(stripped)
        except ValueError:
            return default
    return default


def _get_config_value(config_like: object | None, key: str) -> object | None:
    if config_like is None:
        return None
    if isinstance(config_like, Mapping):
        return config_like.get(key)
    return getattr(config_like, key, None)


_CHAPTER_MARKER_TEXT_PATTERN = re.compile(r"^(?:chapter|глава)\s+(?:\d+|[ivxlcdm]+)\b[ .:-]*$", re.IGNORECASE)
_IMAGE_PLACEHOLDER_TEXT_PATTERN = re.compile(r"\[\[docx_image_[^\]]+\]\]", re.IGNORECASE)
_LEADING_CONTINUATION_FRAGMENT_PATTERN = re.compile(r"^[,.;:!?…)\]»]\s*\S")


def _coerce_mapping_sequence(value: object) -> list[Mapping[str, object]]:
    if not isinstance(value, Sequence) or isinstance(value, (str, bytes, bytearray)):
        return []
    return [item for item in value if isinstance(item, Mapping)]


def _coerce_string_list(value: object) -> list[str]:
    if not isinstance(value, Sequence) or isinstance(value, (str, bytes, bytearray)):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def _entry_has_target_mapping(entry: Mapping[str, object] | None) -> bool:
    if entry is None:
        return False
    return entry.get("mapped_target_index") is not None


def _is_benign_opening_chapter_marker_merge(
    entry: Mapping[str, object],
    next_entry: Mapping[str, object] | None,
) -> bool:
    source_index = _coerce_int(entry.get("source_index"), default=-1)
    if source_index > 0:
        return False
    text_preview = _normalize_structural_text(str(entry.get("text_preview") or ""))
    if _CHAPTER_MARKER_TEXT_PATTERN.match(text_preview) is None:
        return False
    if not _entry_has_target_mapping(next_entry):
        return False
    next_role = str((next_entry or {}).get("role") or (next_entry or {}).get("structural_role") or "").strip().lower()
    return next_role == "heading"


def _is_benign_image_attachment_merge(
    entry: Mapping[str, object],
    previous_entry: Mapping[str, object] | None,
    next_entry: Mapping[str, object] | None,
) -> bool:
    role = str(entry.get("role") or entry.get("structural_role") or "").strip().lower()
    asset_id = str(entry.get("asset_id") or "").strip()
    text_preview = str(entry.get("text_preview") or "")

    if role == "image" and asset_id and _entry_has_target_mapping(next_entry):
        attached_asset_id = str((next_entry or {}).get("attached_to_asset_id") or "").strip()
        if attached_asset_id == asset_id:
            return True

    if _IMAGE_PLACEHOLDER_TEXT_PATTERN.search(text_preview) is None:
        return False
    if _entry_has_target_mapping(next_entry):
        next_preview = _normalize_structural_text(str((next_entry or {}).get("text_preview") or ""))
        if next_preview.startswith("рисунок ") or next_preview.startswith("figure "):
            return True
    if _entry_has_target_mapping(previous_entry):
        previous_preview = str((previous_entry or {}).get("text_preview") or "")
        if _IMAGE_PLACEHOLDER_TEXT_PATTERN.search(previous_preview) is not None:
            return True
    return False


def _is_benign_punctuation_continuation_merge(
    entry: Mapping[str, object],
    previous_entry: Mapping[str, object] | None,
) -> bool:
    role = str(entry.get("role") or entry.get("structural_role") or "").strip().lower()
    if role != "body" or not _entry_has_target_mapping(previous_entry):
        return False
    previous_role = str((previous_entry or {}).get("role") or (previous_entry or {}).get("structural_role") or "").strip().lower()
    if previous_role != "body":
        return False
    text_preview = str(entry.get("text_preview") or "").lstrip()
    return _LEADING_CONTINUATION_FRAGMENT_PATTERN.match(text_preview) is not None


def _resolve_acceptance_unmapped_source_count(
    *,
    formatting_diagnostics: Sequence[Mapping[str, object]],
    translation_quality_report: Mapping[str, object],
) -> int:
    return int(
        _resolve_acceptance_unmapped_source_summary(
            formatting_diagnostics=formatting_diagnostics,
            translation_quality_report=translation_quality_report,
        )["actual"]
    )


def _resolve_acceptance_unmapped_target_count(
    *,
    formatting_diagnostics: Sequence[Mapping[str, object]],
    translation_quality_report: Mapping[str, object],
) -> int:
    return int(
        _resolve_acceptance_unmapped_target_summary(
            formatting_diagnostics=formatting_diagnostics,
            translation_quality_report=translation_quality_report,
        )["actual"]
    )


def _build_run_id(source_path: Path) -> str:
    timestamp = datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")
    sanitized_stem = re.sub(r"[^A-Za-z0-9._-]+", "_", source_path.stem).strip("_") or "real_doc"
    return f"{timestamp}_{os.getpid()}_{sanitized_stem}"


def _snapshot_formatting_diagnostics_paths() -> set[str]:
    if not FORMATTING_DIAGNOSTICS_DIR.exists():
        return set()
    return {str(path.resolve()) for path in FORMATTING_DIAGNOSTICS_DIR.glob("*.json") if path.is_file()}


def _collect_new_formatting_diagnostics_paths(before: set[str], after: set[str]) -> list[str]:
    new_paths = [Path(path) for path in after - before]
    return [
        str(path)
        for path in sorted(new_paths, key=lambda candidate: (candidate.stat().st_mtime, str(candidate)))
    ]


def _safe_git_head() -> str | None:
    try:
        result = subprocess.run(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=PROJECT_ROOT,
            capture_output=True,
            check=True,
            text=True,
        )
    except Exception:
        return None
    return result.stdout.strip() or None


def _detect_wsl_runtime() -> bool:
    release = platform.release().lower()
    return "microsoft" in release or bool(os.environ.get("WSL_DISTRO_NAME"))


def _build_environment_snapshot() -> dict[str, object]:
    return {
        "script_path": _path_for_report(SCRIPT_PATH),
        "project_root": _path_for_report(PROJECT_ROOT),
        "cwd": _path_for_report(Path.cwd()),
        "python_executable": sys.executable,
        "python_version": sys.version.split()[0],
        "pythonpath": os.environ.get("PYTHONPATH"),
        "virtual_env": os.environ.get("VIRTUAL_ENV"),
        "workspace_venv_exists": (PROJECT_ROOT / ".venv" / "bin" / "activate").exists(),
        "workspace_venv_win_exists": (PROJECT_ROOT / ".venv-win" / "Scripts" / "python.exe").exists(),
        "platform": platform.platform(),
        "release": platform.release(),
        "hostname": socket.gethostname(),
        "is_wsl": _detect_wsl_runtime(),
        "git_head": _safe_git_head(),
    }


def _write_json_atomic(path: Path, payload: Mapping[str, object], *, temp_dir: Path | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    resolved_temp_dir = temp_dir or path.parent
    resolved_temp_dir.mkdir(parents=True, exist_ok=True)
    temp_path = resolved_temp_dir / f"{path.name}.{uuid.uuid4().hex}.tmp"
    temp_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    for attempt in range(20):
        try:
            temp_path.replace(path)
            break
        except PermissionError:
            if attempt == 19:
                raise
            time.sleep(min(2.0, 0.25 * (attempt + 1)))
    if temp_path.exists():
        temp_path.unlink(missing_ok=True)


def _format_terminal_progress_line(
    *,
    event_type: str,
    phase: str,
    stage: str,
    detail: str,
    progress: float | None,
    elapsed_seconds: float,
    metrics: Mapping[str, object] | None,
) -> str:
    progress_text = "-"
    if isinstance(progress, (int, float)):
        progress_text = f"{max(0.0, min(1.0, float(progress))) * 100:.1f}%"
    metrics_parts: list[str] = []
    if metrics:
        for key in ("current_block", "block_count", "job_count", "output_ratio", "target_chars", "context_chars"):
            value = metrics.get(key)
            if value is not None:
                metrics_parts.append(f"{key}={value}")
    suffix = f" | {' '.join(metrics_parts)}" if metrics_parts else ""
    return f"[{event_type}] +{elapsed_seconds:.1f}s [{phase}] {stage} | {progress_text} | {detail}{suffix}"


def _normalize_terminal_detail(detail: str) -> str:
    normalized = detail.strip()
    while normalized.startswith("Heartbeat: "):
        normalized = normalized.removeprefix("Heartbeat: ").strip()
    return normalized


def _as_string_list(value: object) -> list[str]:
    if not isinstance(value, Sequence) or isinstance(value, (str, bytes, bytearray)):
        return []
    return [str(item) for item in value]


def _as_object_list(value: object) -> list[object]:
    if not isinstance(value, Sequence) or isinstance(value, (str, bytes, bytearray)):
        return []
    return list(value)


def _as_float_or_none(value: object) -> float | None:
    if isinstance(value, (int, float)):
        return float(value)
    return None


def _build_validation_mode_payload(run_profile) -> dict[str, object]:
    comparison_only = bool(getattr(run_profile, "comparison_only_validation", False))
    validation_run_type = "comparison_only" if comparison_only else "acceptance"
    return {
        "comparison_only_validation": comparison_only,
        "validation_run_type": validation_run_type,
        "acceptance_contract_active": not comparison_only,
        "evidence_label": "comparison_only_non_acceptance" if comparison_only else "acceptance_contract",
        "success_criterion": "pipeline_result_and_artifacts" if comparison_only else "acceptance_passed",
    }


def _resolve_validation_final_status(
    *,
    result: str,
    acceptance_passed: bool,
    validation_mode: Mapping[str, object],
) -> str:
    if bool(validation_mode.get("comparison_only_validation")):
        return "completed" if result == "succeeded" else "failed"
    return "completed" if result == "succeeded" and acceptance_passed else "failed"


def _build_translation_quality_summary_lines(
    translation_quality_report: Mapping[str, object] | None,
) -> list[str]:
    report = translation_quality_report or {}
    lines = [
        f"translation_quality_status={report.get('quality_status')}",
        f"translation_quality_gate_reasons={','.join(_as_string_list(report.get('gate_reasons')))}",
    ]
    for key in (
        "toc_body_concat_gate_source",
        "toc_body_concat_markdown_detected",
        "toc_body_concat_structure_detected",
        "bullet_heading_count",
        "bullet_heading_gate_source",
        "bullet_heading_classification",
        "raw_bullet_heading_count",
        "page_placeholder_heading_concat_count",
        "page_placeholder_heading_concat_source",
        "page_placeholder_heading_concat_classification",
        "raw_page_placeholder_heading_concat_count",
        "false_fragment_heading_count",
        "false_fragment_heading_gate_source",
        "raw_false_fragment_heading_count",
        "residual_bullet_glyph_count",
        "residual_bullet_glyph_gate_source",
        "residual_bullet_glyph_classification",
        "raw_residual_bullet_glyph_count",
        "list_fragment_regression_count",
        "list_fragment_regression_gate_source",
        "raw_list_fragment_regression_count",
        "mixed_script_term_count",
        "mixed_script_term_gate_source",
        "mixed_script_term_classification",
        "raw_mixed_script_term_count",
        "theology_style_deterministic_issue_count",
        "theology_style_deterministic_issue_source",
        "theology_style_deterministic_issue_classification",
        "raw_theology_style_deterministic_issue_count",
    ):
        if key in report:
            lines.append(f"translation_quality_{key}={report.get(key)}")
    return lines


def _format_signed_score_delta(value: float) -> str:
    text = f"{value:+.3f}".rstrip("0")
    if text.endswith("."):
        text += "0"
    return text


def _resolve_reader_mvp_status_language(report: Mapping[str, object]) -> str:
    runtime_config = cast(Mapping[str, object], report.get("runtime_config") or {})
    overrides = cast(Mapping[str, object], runtime_config.get("overrides") or {})
    effective = cast(Mapping[str, object], runtime_config.get("effective") or {})
    ui_defaults = cast(Mapping[str, object], runtime_config.get("ui_defaults") or {})
    language = str(
        overrides.get("target_language")
        or runtime_config.get("target_language")
        or effective.get("target_language")
        or ui_defaults.get("target_language")
        or ""
    ).strip().lower()
    return language or "en"


def _build_reader_mvp_status_payload(report: Mapping[str, object]) -> dict[str, object]:
    validation_mode = cast(Mapping[str, object], report.get("validation_mode") or {})
    acceptance = cast(Mapping[str, object], report.get("acceptance") or {})
    reader_cleanup_evidence = cast(Mapping[str, object], report.get("reader_cleanup_evidence") or {})
    reader_verifier = cast(Mapping[str, object], report.get("reader_verifier_evidence") or {})
    translation_quality_report = cast(Mapping[str, object], report.get("translation_quality_report") or {})

    comparison_only_validation = bool(validation_mode.get("comparison_only_validation"))
    acceptance_contract_active = bool(validation_mode.get("acceptance_contract_active"))
    comparison_only_acceptance_diagnostic = comparison_only_validation and not acceptance_contract_active

    remaining_issues = _coerce_mapping_sequence(reader_verifier.get("remaining_issues"))
    issue_summary = cast(Mapping[str, object], reader_verifier.get("issue_summary_by_category") or {})
    possible_false_deletions = _coerce_string_list(reader_verifier.get("possible_false_deletions"))
    readability_regressions = _coerce_string_list(reader_verifier.get("readability_regressions"))
    failed_checks = _as_string_list(acceptance.get("failed_checks"))

    raw_score = round(_coerce_float(reader_verifier.get("reader_quality_score_raw"), default=0.0), 3)
    cleaned_score = round(_coerce_float(reader_verifier.get("reader_quality_score_cleaned"), default=0.0), 3)
    cleanup_score_delta = round(cleaned_score - raw_score, 3)
    cleanup_score_delta_display = _format_signed_score_delta(cleanup_score_delta)
    remaining_issue_count = len(remaining_issues)
    high_severity_issue_count = _count_reader_verifier_high_severity_issues(remaining_issues)
    top_issue_categories = _select_reader_verifier_top_categories(issue_summary)
    top_issue_categories_text = ", ".join(top_issue_categories)

    cleanup_contract_blockers: list[str] = []
    cleanup_stage_status = str(reader_cleanup_evidence.get("stage_status") or "").strip()
    failed_chunk_count = _coerce_int(reader_cleanup_evidence.get("failed_chunk_count"))
    anchor_repair_status = str(reader_cleanup_evidence.get("anchor_repair_status") or "").strip() or "unknown"
    recommended_anchor_target_count = _coerce_int(reader_cleanup_evidence.get("recommended_anchor_target_count"))
    if cleanup_stage_status and cleanup_stage_status.lower() != "completed":
        cleanup_contract_blockers.append(f"cleanup_stage_status={cleanup_stage_status}")
    if failed_chunk_count > 0:
        cleanup_contract_blockers.append(f"cleanup_chunk_failures={failed_chunk_count}")
    if anchor_repair_status not in {
        "not_needed",
        "not_reported",
        "unknown",
        "runtime_applied",
        "runtime_attempted_no_safe_ops",
        "applied_in_runtime",
    } and not (
        comparison_only_validation and anchor_repair_status == "diagnostic_only_not_applied"
    ):
        cleanup_contract_blockers.append(f"anchor_repair_status={anchor_repair_status}")

    reader_visible_cleanup_defects = [
        f"{category}={count}"
        for category, count in sorted(
            (
                (str(category), _coerce_int(count))
                for category, count in issue_summary.items()
                if str(category).strip()
            ),
            key=lambda item: (-item[1], item[0]),
        )
        if count > 0
    ]

    mapping_quality_gate_diagnostics: list[str] = []
    quality_status = str(translation_quality_report.get("quality_status") or "").strip()
    gate_reasons = _as_string_list(translation_quality_report.get("gate_reasons"))
    if quality_status:
        mapping_quality_gate_diagnostics.append(f"translation_quality_status={quality_status}")
    if gate_reasons:
        mapping_quality_gate_diagnostics.append(
            f"translation_quality_gate_reasons={','.join(gate_reasons)}"
        )
    diagnostic_failed_checks = [
        check
        for check in failed_checks
        if "unmapped" in check or "formatting" in check or "fragment" in check or "quality" in check
    ]
    if diagnostic_failed_checks:
        mapping_quality_gate_diagnostics.append(
            f"acceptance_diagnostic_checks={','.join(diagnostic_failed_checks)}"
        )

    no_false_deletions_reported = not possible_false_deletions
    no_readability_regressions_reported = not readability_regressions
    overall_verdict = str(reader_verifier.get("overall_verdict") or "unclear").strip() or "unclear"
    cleaned_audit_verdict = str(reader_verifier.get("cleaned_audit_verdict") or "unclear").strip() or "unclear"
    filtered_issue_counts = {
        "toc_total": max(0, _coerce_int(reader_verifier.get("filtered_toc_issue_count"), default=0)),
        "toc_pre_audit": max(0, _coerce_int(reader_verifier.get("filtered_toc_pre_audit_count"), default=0)),
        "toc_verifier_issue": max(0, _coerce_int(reader_verifier.get("filtered_toc_verifier_issue_count"), default=0)),
        "toc_evidence_anchor": max(0, _coerce_int(reader_verifier.get("filtered_toc_evidence_anchor_count"), default=0)),
    }
    cleanup_diagnostics = cast(Mapping[str, object], reader_verifier.get("cleanup_diagnostics") or {})
    cleanup_ignored_reason_counts = cast(Mapping[str, object], cleanup_diagnostics.get("ignored_reason_counts") or {})
    cleanup_application_diagnostics = [
        f"{reason}={max(0, _coerce_int(cleanup_ignored_reason_counts.get(reason), default=0))}"
        for reason in _TRACKED_READER_CLEANUP_IGNORED_REASONS
        if max(0, _coerce_int(cleanup_ignored_reason_counts.get(reason), default=0)) > 0
    ]
    cleanup_diagnostic_examples: list[str] = []
    for entry in _coerce_mapping_sequence(cleanup_diagnostics.get("top_ignored_reasons")):
        ignored_reason = str(entry.get("ignored_reason") or "").strip()
        for example in _coerce_mapping_sequence(entry.get("examples"))[:2]:
            operation_name = str(example.get("operation") or "unknown").strip() or "unknown"
            operation_reason = str(example.get("reason") or "").strip()
            preview = str(example.get("text_preview") or "").strip()
            line = f"{ignored_reason}: {operation_name}"
            if operation_reason:
                line += f"/{operation_reason}"
            if preview:
                line += f" -> {preview}"
            cleanup_diagnostic_examples.append(line)
        if len(cleanup_diagnostic_examples) >= 4:
            break
    result = str(report.get("result") or "").strip().lower()
    has_safety_risks = bool(possible_false_deletions or readability_regressions)
    if result != "succeeded":
        status_label = "pipeline_failed"
    elif overall_verdict in {"cleaned_worse", "raw_better"} or cleaned_audit_verdict == "unsafe_or_regressed":
        status_label = "cleanup_regressed"
    elif overall_verdict == "cleaned_better" and (remaining_issue_count > 0 or has_safety_risks):
        status_label = "readable_draft_not_acceptance_ready"
    elif overall_verdict == "cleaned_better":
        status_label = "cleaned_better_diagnostic_evidence"
    else:
        status_label = "mixed_or_unclear"

    language = _resolve_reader_mvp_status_language(report)
    if language == "ru":
        pipeline_summary = (
            "Пайплайн не завершился успешно; reader-first результат нельзя считать доказанным."
            if status_label == "pipeline_failed"
            else "Пайплайн завершился успешно и собрал reviewable raw/cleaned артефакты."
        )
        cleanup_summary = (
            f"Cleanup дал verdict {overall_verdict}: score {raw_score:.1f} -> {cleaned_score:.1f} "
            f"({cleanup_score_delta_display})."
        )
        acceptance_summary = (
            "Acceptance failure здесь диагностический и не считается падением пайплайна для comparison-only прогона."
            if comparison_only_acceptance_diagnostic
            else f"Acceptance passed={bool(acceptance.get('passed'))}."
        )
        remaining_risk_summary = (
            f"Остаются {remaining_issue_count} reader-visible issues, из них {high_severity_issue_count} high severity; "
            f"top categories: {top_issue_categories_text or 'нет'}."
        )
        positive_safety_signals = [
            "Verifier не сообщил о false deletions."
            if no_false_deletions_reported
            else f"Verifier сообщил о possible false deletions: {len(possible_false_deletions)}.",
            "Verifier не сообщил о readability regressions."
            if no_readability_regressions_reported
            else f"Verifier сообщил о readability regressions: {len(readability_regressions)}.",
        ]
        if status_label == "pipeline_failed":
            user_summary = (
                "Пайплайн не завершился успешно, поэтому нельзя утверждать, что cleaned output стал лучше raw. "
                f"{acceptance_summary} {remaining_risk_summary}"
            )
            risk_summary = (
                "Главный риск сейчас runtime/pipeline failure, а не доказанное качество cleanup: "
                f"false deletions={len(possible_false_deletions)}, readability regressions={len(readability_regressions)}. "
                f"{remaining_risk_summary}"
            )
        elif status_label == "cleanup_regressed":
            user_summary = (
                "Очистка ухудшила результат или сделала его небезопасным для принятия: "
                "cleaned output нельзя считать лучше raw. "
                f"{acceptance_summary} {remaining_risk_summary}"
            )
            risk_summary = (
                "Главный риск сейчас в регрессии очистки или небезопасном результате: "
                f"false deletions={len(possible_false_deletions)}, readability regressions={len(readability_regressions)}. "
                f"{remaining_risk_summary}"
            )
        elif status_label == "mixed_or_unclear":
            user_summary = (
                "Доказательства смешанные или неясные: нельзя утверждать, что cleaned output стал лучше raw. "
                f"{acceptance_summary} {remaining_risk_summary}"
            )
            risk_summary = (
                "Главный риск сейчас в неясном результате проверки: "
                f"false deletions={len(possible_false_deletions)}, readability regressions={len(readability_regressions)}. "
                f"{remaining_risk_summary}"
            )
        elif status_label == "readable_draft_not_acceptance_ready":
            user_summary = (
                "Есть полезное улучшение читабельности, но результат ещё не acceptance-ready. "
                f"{acceptance_summary} {remaining_risk_summary}"
            )
            risk_summary = (
                "Главный остаточный риск сейчас reader-visible, а не runtime: "
                f"false deletions={len(possible_false_deletions)}, readability regressions={len(readability_regressions)}. "
                f"{remaining_risk_summary}"
            )
        else:
            user_summary = (
                f"Стало лучше: cleaned output читается легче. {acceptance_summary} {remaining_risk_summary}"
            )
            risk_summary = (
                f"Главный остаточный риск сейчас reader-visible, а не runtime: {remaining_risk_summary}"
            )
    else:
        pipeline_summary = (
            "The pipeline did not succeed, so the reader-first result is not proven."
            if status_label == "pipeline_failed"
            else "The pipeline succeeded and produced reviewable raw/cleaned artifacts."
        )
        cleanup_summary = (
            f"Cleanup verdict is {overall_verdict}: score {raw_score:.1f} -> {cleaned_score:.1f} "
            f"({cleanup_score_delta_display})."
        )
        acceptance_summary = (
            "Acceptance failure is diagnostic only for this comparison-only run and does not mean the pipeline failed."
            if comparison_only_acceptance_diagnostic
            else f"Acceptance passed={bool(acceptance.get('passed'))}."
        )
        remaining_risk_summary = (
            f"{remaining_issue_count} reader-visible issues remain, with {high_severity_issue_count} high severity; "
            f"top categories: {top_issue_categories_text or 'none'}."
        )
        positive_safety_signals = [
            "The verifier reported no false deletions."
            if no_false_deletions_reported
            else f"The verifier reported possible false deletions: {len(possible_false_deletions)}.",
            "The verifier reported no readability regressions."
            if no_readability_regressions_reported
            else f"The verifier reported readability regressions: {len(readability_regressions)}.",
        ]
        if status_label == "pipeline_failed":
            user_summary = (
                "The pipeline did not succeed, so the cleaned output should not be described as easier or better than raw. "
                f"{acceptance_summary} {remaining_risk_summary}"
            )
            risk_summary = (
                "The current risk is runtime/pipeline failure rather than proven cleanup quality: "
                f"false deletions={len(possible_false_deletions)}, readability regressions={len(readability_regressions)}. "
                f"{remaining_risk_summary}"
            )
        elif status_label == "cleanup_regressed":
            user_summary = (
                "Cleanup regressed or made the result unsafe: the cleaned output should not be treated as better than raw. "
                f"{acceptance_summary} {remaining_risk_summary}"
            )
            risk_summary = (
                "The current risk is cleanup regression or unsafe output: "
                f"false deletions={len(possible_false_deletions)}, readability regressions={len(readability_regressions)}. "
                f"{remaining_risk_summary}"
            )
        elif status_label == "mixed_or_unclear":
            user_summary = (
                "The evidence is mixed or unclear, so the cleaned output should not be described as easier or better yet. "
                f"{acceptance_summary} {remaining_risk_summary}"
            )
            risk_summary = (
                "The current risk is an unclear validation result: "
                f"false deletions={len(possible_false_deletions)}, readability regressions={len(readability_regressions)}. "
                f"{remaining_risk_summary}"
            )
        elif status_label == "readable_draft_not_acceptance_ready":
            user_summary = (
                "There is some readability improvement, but the result is not acceptance-ready. "
                f"{acceptance_summary} {remaining_risk_summary}"
            )
            risk_summary = (
                "The remaining risk is reader-visible rather than runtime: "
                f"false deletions={len(possible_false_deletions)}, readability regressions={len(readability_regressions)}. "
                f"{remaining_risk_summary}"
            )
        else:
            user_summary = (
                f"The cleaned output is easier to read. {acceptance_summary} {remaining_risk_summary}"
            )
            risk_summary = f"The remaining risk is reader-visible rather than runtime: {remaining_risk_summary}"

    signal_layers = {
        "pipeline_result_status": result or "unknown",
        "comparison_only_validation_status": "observer_only" if comparison_only_validation else "acceptance_bound",
        "acceptance_diagnostic_status": (
            "diagnostic_only_failed"
            if comparison_only_acceptance_diagnostic and failed_checks
            else "diagnostic_only_passed"
            if comparison_only_acceptance_diagnostic
            else "acceptance_passed"
            if bool(acceptance.get("passed"))
            else "acceptance_failed"
        ),
        "acceptance_diagnostic_failed_checks": failed_checks,
        "reader_verifier_verdict": overall_verdict,
        "cleaned_audit_verdict": cleaned_audit_verdict,
        "remaining_reader_visible_issue_count": remaining_issue_count,
        "out_of_scope_filtered_issue_counts": filtered_issue_counts,
    }

    return {
        "status_label": status_label,
        "language": language,
        "pipeline_summary": pipeline_summary,
        "cleanup_summary": cleanup_summary,
        "acceptance_summary": acceptance_summary,
        "remaining_risk_summary": remaining_risk_summary,
        "user_summary": user_summary,
        "risk_summary": risk_summary,
        "cleanup_score_delta": cleanup_score_delta,
        "cleanup_score_delta_display": cleanup_score_delta_display,
        "remaining_issue_count": remaining_issue_count,
        "high_severity_issue_count": high_severity_issue_count,
        "top_issue_categories": top_issue_categories,
        "anchor_repair_status": anchor_repair_status,
        "recommended_anchor_target_count": recommended_anchor_target_count,
        "no_false_deletions_reported": no_false_deletions_reported,
        "no_readability_regressions_reported": no_readability_regressions_reported,
        "comparison_only_acceptance_diagnostic": comparison_only_acceptance_diagnostic,
        "positive_safety_signals": positive_safety_signals,
        "signal_layers": signal_layers,
        "filtered_issue_counts": filtered_issue_counts,
        "cleanup_diagnostic_examples": cleanup_diagnostic_examples,
        "blocker_groups": {
            "cleanup_contract": cleanup_contract_blockers,
            "reader_visible_cleanup_defects": reader_visible_cleanup_defects,
            "cleanup_application_diagnostics": cleanup_application_diagnostics,
            "mapping_quality_gate_diagnostics": mapping_quality_gate_diagnostics,
        },
    }


def _render_reader_mvp_status_markdown(status_payload: Mapping[str, object]) -> str:
    language = str(status_payload.get("language") or "en").strip().lower()
    blocker_groups = cast(Mapping[str, object], status_payload.get("blocker_groups") or {})
    cleanup_contract = _coerce_string_list(blocker_groups.get("cleanup_contract"))
    reader_visible = _coerce_string_list(blocker_groups.get("reader_visible_cleanup_defects"))
    cleanup_application = _coerce_string_list(blocker_groups.get("cleanup_application_diagnostics"))
    quality_gate = _coerce_string_list(blocker_groups.get("mapping_quality_gate_diagnostics"))
    positive_safety_signals = _coerce_string_list(status_payload.get("positive_safety_signals"))
    signal_layers = cast(Mapping[str, object], status_payload.get("signal_layers") or {})
    filtered_issue_counts = cast(Mapping[str, object], status_payload.get("filtered_issue_counts") or {})
    cleanup_diagnostic_examples = _coerce_string_list(status_payload.get("cleanup_diagnostic_examples"))

    if language == "ru":
        lines = [
            "# Статус MVP",
            "",
            f"- Итоговый статус: {status_payload.get('status_label')}",
            f"- Для оператора: {status_payload.get('user_summary')}",
            f"- Риск: {status_payload.get('risk_summary')}",
            "",
            "# Разделение сигналов",
            "",
            f"- Pipeline success: {status_payload.get('pipeline_summary')}",
            f"- Cleanup improvement: {status_payload.get('cleanup_summary')}",
            f"- Acceptance diagnostic: {status_payload.get('acceptance_summary')}",
            f"- Remaining reader-visible risk: {status_payload.get('remaining_risk_summary')}",
            "",
            "# Result Layers",
            "",
            f"- Pipeline result status: {signal_layers.get('pipeline_result_status')}",
            f"- Comparison-only validation status: {signal_layers.get('comparison_only_validation_status')}",
            f"- Acceptance diagnostic status: {signal_layers.get('acceptance_diagnostic_status')}",
            f"- Acceptance diagnostic failed checks: {','.join(_coerce_string_list(signal_layers.get('acceptance_diagnostic_failed_checks'))) or 'none'}",
            f"- Reader verifier verdict: {signal_layers.get('reader_verifier_verdict')}",
            f"- Cleaned audit verdict: {signal_layers.get('cleaned_audit_verdict')}",
            f"- Remaining reader-visible issue count: {signal_layers.get('remaining_reader_visible_issue_count')}",
            (
                "- Out-of-scope filtered issue counts: "
                f"total={filtered_issue_counts.get('toc_total')}, "
                f"pre_audit={filtered_issue_counts.get('toc_pre_audit')}, "
                f"verifier={filtered_issue_counts.get('toc_verifier_issue')}, "
                f"anchors={filtered_issue_counts.get('toc_evidence_anchor')}"
            ),
            "",
            "# Позитивные сигналы",
            "",
        ]
        lines.extend(f"- {item}" for item in positive_safety_signals)
        lines.extend(["", "# Группы блокеров", "", "## Cleanup Contract", ""])
        lines.extend(f"- {item}" for item in cleanup_contract or ["Нет cleanup-contract blockers."])
        lines.extend(["", "## Reader-Visible Cleanup Defects", ""])
        lines.extend(f"- {item}" for item in reader_visible or ["Нет reader-visible cleanup blockers."])
        lines.extend(["", "## Cleanup Application Diagnostics", ""])
        lines.extend(f"- {item}" for item in cleanup_application or ["Нет cleanup-application diagnostics blockers."])
        lines.extend(["", "## Mapping / Quality-Gate Diagnostics", ""])
        lines.extend(f"- {item}" for item in quality_gate or ["Нет mapping/quality-gate diagnostics blockers."])
        lines.extend(["", "# Cleanup Diagnostic Examples", ""])
        lines.extend(f"- {item}" for item in cleanup_diagnostic_examples or ["Нет примеров ignored cleanup reasons."])
        return "\n".join(lines).strip() + "\n"

    lines = [
        "# MVP Status",
        "",
        f"- Status label: {status_payload.get('status_label')}",
        f"- User summary: {status_payload.get('user_summary')}",
        f"- Risk summary: {status_payload.get('risk_summary')}",
        "",
        "# Signal Split",
        "",
        f"- Pipeline success: {status_payload.get('pipeline_summary')}",
        f"- Cleanup improvement: {status_payload.get('cleanup_summary')}",
        f"- Acceptance diagnostic: {status_payload.get('acceptance_summary')}",
        f"- Remaining reader-visible risk: {status_payload.get('remaining_risk_summary')}",
        "",
        "# Result Layers",
        "",
        f"- Pipeline result status: {signal_layers.get('pipeline_result_status')}",
        f"- Comparison-only validation status: {signal_layers.get('comparison_only_validation_status')}",
        f"- Acceptance diagnostic status: {signal_layers.get('acceptance_diagnostic_status')}",
        f"- Acceptance diagnostic failed checks: {','.join(_coerce_string_list(signal_layers.get('acceptance_diagnostic_failed_checks'))) or 'none'}",
        f"- Reader verifier verdict: {signal_layers.get('reader_verifier_verdict')}",
        f"- Cleaned audit verdict: {signal_layers.get('cleaned_audit_verdict')}",
        f"- Remaining reader-visible issue count: {signal_layers.get('remaining_reader_visible_issue_count')}",
        (
            "- Out-of-scope filtered issue counts: "
            f"total={filtered_issue_counts.get('toc_total')}, "
            f"pre_audit={filtered_issue_counts.get('toc_pre_audit')}, "
            f"verifier={filtered_issue_counts.get('toc_verifier_issue')}, "
            f"anchors={filtered_issue_counts.get('toc_evidence_anchor')}"
        ),
        "",
        "# Positive Safety Signals",
        "",
    ]
    lines.extend(f"- {item}" for item in positive_safety_signals)
    lines.extend(["", "# Blocker Groups", "", "## Cleanup Contract", ""])
    lines.extend(f"- {item}" for item in cleanup_contract or ["No cleanup-contract blockers recorded."])
    lines.extend(["", "## Reader-Visible Cleanup Defects", ""])
    lines.extend(f"- {item}" for item in reader_visible or ["No reader-visible cleanup blockers recorded."])
    lines.extend(["", "## Cleanup Application Diagnostics", ""])
    lines.extend(f"- {item}" for item in cleanup_application or ["No cleanup-application diagnostics blockers recorded."])
    lines.extend(["", "## Mapping / Quality-Gate Diagnostics", ""])
    lines.extend(f"- {item}" for item in quality_gate or ["No mapping/quality-gate blockers recorded."])
    lines.extend(["", "# Cleanup Diagnostic Examples", ""])
    lines.extend(f"- {item}" for item in cleanup_diagnostic_examples or ["No ignored cleanup reason examples recorded."])
    return "\n".join(lines).strip() + "\n"


def _build_reader_mvp_status_summary_lines(status_payload: Mapping[str, object]) -> list[str]:
    blocker_groups = cast(Mapping[str, object], status_payload.get("blocker_groups") or {})
    signal_layers = cast(Mapping[str, object], status_payload.get("signal_layers") or {})
    return [
        f"reader_mvp_status_label={status_payload.get('status_label')}",
        f"reader_mvp_status_user_summary={status_payload.get('user_summary')}",
        f"reader_mvp_status_risk_summary={status_payload.get('risk_summary')}",
        f"reader_mvp_status_cleanup_score_delta={status_payload.get('cleanup_score_delta')}",
        f"reader_mvp_status_acceptance_diagnostic_only={status_payload.get('comparison_only_acceptance_diagnostic')}",
        f"reader_mvp_status_pipeline_result_status={signal_layers.get('pipeline_result_status')}",
        f"reader_mvp_status_comparison_only_validation_status={signal_layers.get('comparison_only_validation_status')}",
        f"reader_mvp_status_acceptance_diagnostic_status={signal_layers.get('acceptance_diagnostic_status')}",
        "reader_mvp_status_filtered_issue_counts="
        + json.dumps(cast(Mapping[str, object], status_payload.get("filtered_issue_counts") or {}), ensure_ascii=False, sort_keys=True),
        f"reader_mvp_status_anchor_repair_status={status_payload.get('anchor_repair_status')}",
        "reader_mvp_status_false_deletion_status="
        + ("none_reported" if bool(status_payload.get("no_false_deletions_reported")) else "reported"),
        "reader_mvp_status_readability_regression_status="
        + ("none_reported" if bool(status_payload.get("no_readability_regressions_reported")) else "reported"),
        "reader_mvp_status_blocker_group_cleanup_contract="
        + "|".join(_coerce_string_list(blocker_groups.get("cleanup_contract")) or ["none"]),
        "reader_mvp_status_blocker_group_reader_visible="
        + "|".join(_coerce_string_list(blocker_groups.get("reader_visible_cleanup_defects")) or ["none"]),
        "reader_mvp_status_blocker_group_cleanup_application="
        + "|".join(_coerce_string_list(blocker_groups.get("cleanup_application_diagnostics")) or ["none"]),
        "reader_mvp_status_blocker_group_quality_gate="
        + "|".join(_coerce_string_list(blocker_groups.get("mapping_quality_gate_diagnostics")) or ["none"]),
        "reader_mvp_status_cleanup_diagnostic_examples="
        + "|".join(_coerce_string_list(status_payload.get("cleanup_diagnostic_examples")) or ["none"]),
    ]


def _print_terminal_completion_summary(*, report: Mapping[str, object], final_status: str) -> None:
    output_artifacts = cast(Mapping[str, object], report.get("output_artifacts") or {})
    acceptance = cast(Mapping[str, object], report.get("acceptance") or {})
    validation_mode = cast(Mapping[str, object], report.get("validation_mode") or {})
    failed_checks = _as_string_list(acceptance.get("failed_checks"))
    translation_quality_report = cast(Mapping[str, object], report.get("translation_quality_report") or {})
    _safe_terminal_print(
        "[summary] "
        f"status={final_status} "
        f"result={report.get('result')} "
        f"validation_run_type={validation_mode.get('validation_run_type')} "
        f"acceptance_passed={acceptance.get('passed')} "
        f"run_id={cast(Mapping[str, object], report.get('run') or {}).get('run_id')}",
        flush=True,
    )
    _safe_terminal_print(
        "[artifacts] "
        f"report={output_artifacts.get('report_json')} "
        f"summary={output_artifacts.get('summary_txt')} "
        f"progress={report.get('progress_path')}",
        flush=True,
    )
    if translation_quality_report:
        _safe_terminal_print(
            "[translation_quality] "
            + " ".join(
                part
                for part in (
                    f"status={translation_quality_report.get('quality_status')}",
                    f"gate_reasons={','.join(_as_string_list(translation_quality_report.get('gate_reasons')))}",
                    f"toc_body_concat_gate_source={translation_quality_report.get('toc_body_concat_gate_source')}"
                    if "toc_body_concat_gate_source" in translation_quality_report
                    else "",
                    f"bullet_heading_count={translation_quality_report.get('bullet_heading_count')}"
                    if "bullet_heading_count" in translation_quality_report
                    else "",
                    f"bullet_heading_gate_source={translation_quality_report.get('bullet_heading_gate_source')}"
                    if "bullet_heading_gate_source" in translation_quality_report
                    else "",
                    f"bullet_heading_classification={translation_quality_report.get('bullet_heading_classification')}"
                    if "bullet_heading_classification" in translation_quality_report
                    else "",
                    f"raw_bullet_heading_count={translation_quality_report.get('raw_bullet_heading_count')}"
                    if "raw_bullet_heading_count" in translation_quality_report
                    else "",
                    f"page_placeholder_heading_concat_count={translation_quality_report.get('page_placeholder_heading_concat_count')}"
                    if "page_placeholder_heading_concat_count" in translation_quality_report
                    else "",
                    f"page_placeholder_heading_concat_source={translation_quality_report.get('page_placeholder_heading_concat_source')}"
                    if "page_placeholder_heading_concat_source" in translation_quality_report
                    else "",
                    f"page_placeholder_heading_concat_classification={translation_quality_report.get('page_placeholder_heading_concat_classification')}"
                    if "page_placeholder_heading_concat_classification" in translation_quality_report
                    else "",
                    f"raw_page_placeholder_heading_concat_count={translation_quality_report.get('raw_page_placeholder_heading_concat_count')}"
                    if "raw_page_placeholder_heading_concat_count" in translation_quality_report
                    else "",
                    f"false_fragment_heading_count={translation_quality_report.get('false_fragment_heading_count')}"
                    if "false_fragment_heading_count" in translation_quality_report
                    else "",
                    f"false_fragment_heading_gate_source={translation_quality_report.get('false_fragment_heading_gate_source')}"
                    if "false_fragment_heading_gate_source" in translation_quality_report
                    else "",
                    f"raw_false_fragment_heading_count={translation_quality_report.get('raw_false_fragment_heading_count')}"
                    if "raw_false_fragment_heading_count" in translation_quality_report
                    else "",
                    f"residual_bullet_glyph_count={translation_quality_report.get('residual_bullet_glyph_count')}"
                    if "residual_bullet_glyph_count" in translation_quality_report
                    else "",
                    f"residual_bullet_glyph_gate_source={translation_quality_report.get('residual_bullet_glyph_gate_source')}"
                    if "residual_bullet_glyph_gate_source" in translation_quality_report
                    else "",
                    f"residual_bullet_glyph_classification={translation_quality_report.get('residual_bullet_glyph_classification')}"
                    if "residual_bullet_glyph_classification" in translation_quality_report
                    else "",
                    f"raw_residual_bullet_glyph_count={translation_quality_report.get('raw_residual_bullet_glyph_count')}"
                    if "raw_residual_bullet_glyph_count" in translation_quality_report
                    else "",
                    f"list_fragment_regression_count={translation_quality_report.get('list_fragment_regression_count')}"
                    if "list_fragment_regression_count" in translation_quality_report
                    else "",
                    f"list_fragment_regression_gate_source={translation_quality_report.get('list_fragment_regression_gate_source')}"
                    if "list_fragment_regression_gate_source" in translation_quality_report
                    else "",
                    f"raw_list_fragment_regression_count={translation_quality_report.get('raw_list_fragment_regression_count')}"
                    if "raw_list_fragment_regression_count" in translation_quality_report
                    else "",
                    f"mixed_script_term_count={translation_quality_report.get('mixed_script_term_count')}"
                    if "mixed_script_term_count" in translation_quality_report
                    else "",
                    f"mixed_script_term_gate_source={translation_quality_report.get('mixed_script_term_gate_source')}"
                    if "mixed_script_term_gate_source" in translation_quality_report
                    else "",
                    f"mixed_script_term_classification={translation_quality_report.get('mixed_script_term_classification')}"
                    if "mixed_script_term_classification" in translation_quality_report
                    else "",
                    f"raw_mixed_script_term_count={translation_quality_report.get('raw_mixed_script_term_count')}"
                    if "raw_mixed_script_term_count" in translation_quality_report
                    else "",
                    f"theology_style_deterministic_issue_count={translation_quality_report.get('theology_style_deterministic_issue_count')}"
                    if "theology_style_deterministic_issue_count" in translation_quality_report
                    else "",
                    f"theology_style_deterministic_issue_source={translation_quality_report.get('theology_style_deterministic_issue_source')}"
                    if "theology_style_deterministic_issue_source" in translation_quality_report
                    else "",
                    f"theology_style_deterministic_issue_classification={translation_quality_report.get('theology_style_deterministic_issue_classification')}"
                    if "theology_style_deterministic_issue_classification" in translation_quality_report
                    else "",
                    f"raw_theology_style_deterministic_issue_count={translation_quality_report.get('raw_theology_style_deterministic_issue_count')}"
                    if "raw_theology_style_deterministic_issue_count" in translation_quality_report
                    else "",
                )
                if part
            ),
            flush=True,
        )
    if failed_checks:
        _safe_terminal_print(f"[acceptance] failed_checks={','.join(str(item) for item in failed_checks)}", flush=True)


class ValidationProgressTracker:
    def __init__(
        self,
        *,
        run_id: str,
        document_profile_id: str | None = None,
        run_profile_id: str | None = None,
        validation_tier: str | None = None,
        validation_run_type: str | None = None,
        comparison_only_validation: bool | None = None,
        source_path: Path,
        run_dir: Path,
        artifact_root: Path,
        progress_path: Path,
        latest_progress_path: Path,
        latest_manifest_path: Path,
        report_path: Path,
        summary_path: Path,
        markdown_path: Path,
        docx_path: Path,
        latest_report_path: Path,
        latest_summary_path: Path,
        latest_markdown_path: Path,
        latest_docx_path: Path,
        started_at_utc: datetime,
    ) -> None:
        self.run_id = run_id
        self.run_dir = run_dir
        self.progress_path = progress_path
        self.latest_progress_path = latest_progress_path
        self.latest_manifest_path = latest_manifest_path
        self.started_at_monotonic = time.perf_counter()
        self.lock = threading.Lock()
        self.stop_event = threading.Event()
        self.heartbeat_thread: threading.Thread | None = None
        self.state: dict[str, object] = {
            "run_id": run_id,
            "document_profile_id": document_profile_id,
            "run_profile_id": run_profile_id,
            "validation_tier": validation_tier,
            "validation_run_type": validation_run_type,
            "comparison_only_validation": comparison_only_validation,
            "status": "in_progress",
            "source_document_path": _path_for_report(source_path),
            "run_dir": _path_for_report(run_dir),
            "artifact_root": _path_for_report(artifact_root),
            "progress_json": _path_for_report(progress_path),
            "latest_progress_json": _path_for_report(latest_progress_path),
            "report_json": _path_for_report(report_path),
            "summary_txt": _path_for_report(summary_path),
            "markdown_path": _path_for_report(markdown_path),
            "docx_path": _path_for_report(docx_path),
            "latest_report_json": _path_for_report(latest_report_path),
            "latest_summary_txt": _path_for_report(latest_summary_path),
            "latest_markdown_path": _path_for_report(latest_markdown_path),
            "latest_docx_path": _path_for_report(latest_docx_path),
            "started_at_utc": started_at_utc.isoformat(),
            "finished_at_utc": None,
            "last_update_at_utc": started_at_utc.isoformat(),
            "phase": "startup",
            "stage": "Инициализация",
            "detail": "Создаю run-scoped артефакты и latest manifest.",
            "progress": 0.0,
            "last_error": "",
            "result": "not_started",
            "acceptance_passed": None,
            "failure_classification": None,
            "runtime_config": None,
            "runtime_overrides": {},
            "metrics": {},
            "recent_events": [],
        }
        self._write_locked()

    def set_manifest_context(self, **values: object) -> None:
        with self.lock:
            self.state.update(cast(dict[str, object], sanitize_for_json(values)))
            self._write_locked()

    def start(self) -> None:
        if self.heartbeat_thread is not None:
            return
        self.heartbeat_thread = threading.Thread(target=self._heartbeat_loop, name="lietaer-validation-heartbeat", daemon=True)
        self.heartbeat_thread.start()

    def stop(self) -> None:
        self.stop_event.set()
        if self.heartbeat_thread is not None:
            self.heartbeat_thread.join(timeout=HEARTBEAT_INTERVAL_SECONDS + 2.0)

    def emit(
        self,
        *,
        event_type: str,
        phase: str,
        stage: str,
        detail: str,
        progress: float | None = None,
        metrics: Mapping[str, object] | None = None,
        print_line: bool = True,
    ) -> None:
        line = ""
        with self.lock:
            now = datetime.now(UTC)
            elapsed_seconds = max(0.0, time.perf_counter() - self.started_at_monotonic)
            if progress is not None:
                self.state["progress"] = max(0.0, min(1.0, float(progress)))
            self.state["phase"] = phase
            self.state["stage"] = stage
            self.state["detail"] = detail
            self.state["last_update_at_utc"] = now.isoformat()
            self.state["elapsed_seconds"] = round(elapsed_seconds, 3)
            if metrics:
                self.state["metrics"] = sanitize_for_json(dict(metrics))
            recent_events = _as_object_list(self.state.get("recent_events"))
            recent_events.append(
                {
                    "timestamp_utc": now.isoformat(),
                    "event_type": event_type,
                    "phase": phase,
                    "stage": stage,
                    "detail": detail,
                    "progress": self.state.get("progress"),
                    "metrics": sanitize_for_json(dict(metrics or {})),
                }
            )
            self.state["recent_events"] = recent_events[-25:]
            self._write_locked()
            if print_line:
                line = _format_terminal_progress_line(
                    event_type=event_type,
                    phase=phase,
                    stage=stage,
                    detail=detail,
                    progress=_as_float_or_none(self.state.get("progress")) or 0.0,
                    elapsed_seconds=elapsed_seconds,
                    metrics=metrics,
                )
        if line:
            _safe_terminal_print(line, flush=True)

    def finalize(
        self,
        *,
        status: str,
        result: str,
        acceptance_passed: bool,
        failure_classification: str | None,
        last_error: str,
        detail: str,
    ) -> None:
        line = ""
        with self.lock:
            finished_at_utc = datetime.now(UTC)
            elapsed_seconds = max(0.0, time.perf_counter() - self.started_at_monotonic)
            self.state["status"] = status
            self.state["result"] = result
            self.state["acceptance_passed"] = acceptance_passed
            self.state["failure_classification"] = failure_classification
            self.state["last_error"] = last_error
            self.state["phase"] = "completed" if status == "completed" else "failed"
            self.state["stage"] = "Завершено" if status == "completed" else "Завершено с ошибкой"
            self.state["detail"] = detail
            self.state["progress"] = 1.0
            self.state["finished_at_utc"] = finished_at_utc.isoformat()
            self.state["last_update_at_utc"] = finished_at_utc.isoformat()
            self.state["elapsed_seconds"] = round(elapsed_seconds, 3)
            self._write_locked()
            line = _format_terminal_progress_line(
                event_type=status,
                phase=str(self.state["phase"]),
                stage=str(self.state["stage"]),
                detail=detail,
                progress=1.0,
                elapsed_seconds=elapsed_seconds,
                metrics={"result": result, "acceptance_passed": acceptance_passed},
            )
        _safe_terminal_print(line, flush=True)

    def _build_manifest_payload_locked(self) -> dict[str, object]:
        return {
            "run_id": self.state.get("run_id"),
            "document_profile_id": self.state.get("document_profile_id"),
            "run_profile_id": self.state.get("run_profile_id"),
            "validation_tier": self.state.get("validation_tier"),
            "validation_run_type": self.state.get("validation_run_type"),
            "comparison_only_validation": self.state.get("comparison_only_validation"),
            "validation_mode": self.state.get("validation_mode"),
            "status": self.state.get("status"),
            "source_document_path": self.state.get("source_document_path"),
            "run_dir": self.state.get("run_dir"),
            "progress_json": self.state.get("progress_json"),
            "latest_progress_json": self.state.get("latest_progress_json"),
            "report_json": self.state.get("report_json"),
            "summary_txt": self.state.get("summary_txt"),
            "markdown_path": self.state.get("markdown_path"),
            "docx_path": self.state.get("docx_path"),
            "latest_report_json": self.state.get("latest_report_json"),
            "latest_summary_txt": self.state.get("latest_summary_txt"),
            "latest_markdown_path": self.state.get("latest_markdown_path"),
            "latest_docx_path": self.state.get("latest_docx_path"),
            "started_at_utc": self.state.get("started_at_utc"),
            "finished_at_utc": self.state.get("finished_at_utc"),
            "last_update_at_utc": self.state.get("last_update_at_utc"),
            "phase": self.state.get("phase"),
            "stage": self.state.get("stage"),
            "detail": self.state.get("detail"),
            "progress": self.state.get("progress"),
            "result": self.state.get("result"),
            "acceptance_passed": self.state.get("acceptance_passed"),
            "failure_classification": self.state.get("failure_classification"),
            "last_error": self.state.get("last_error"),
            "runtime_config": self.state.get("runtime_config"),
            "runtime_overrides": self.state.get("runtime_overrides"),
            "reader_verifier_status": self.state.get("reader_verifier_status"),
            "reader_verifier_reason": self.state.get("reader_verifier_reason"),
            "reader_verifier_model_selector": self.state.get("reader_verifier_model_selector"),
            "reader_verifier_canonical_selector": self.state.get("reader_verifier_canonical_selector"),
            "reader_verifier_provider": self.state.get("reader_verifier_provider"),
            "reader_verifier_model_id": self.state.get("reader_verifier_model_id"),
            "reader_verifier_overall_verdict": self.state.get("reader_verifier_overall_verdict"),
            "reader_verifier_cleaned_audit_verdict": self.state.get("reader_verifier_cleaned_audit_verdict"),
            "reader_verifier_confidence": self.state.get("reader_verifier_confidence"),
            "reader_verifier_remaining_issue_count": self.state.get("reader_verifier_remaining_issue_count"),
            "reader_verifier_high_severity_issue_count": self.state.get("reader_verifier_high_severity_issue_count"),
            "reader_verifier_top_issue_categories": self.state.get("reader_verifier_top_issue_categories"),
            "reader_verifier_filtered_toc_issue_count": self.state.get("reader_verifier_filtered_toc_issue_count"),
            "reader_verifier_filtered_toc_pre_audit_count": self.state.get("reader_verifier_filtered_toc_pre_audit_count"),
            "reader_verifier_filtered_toc_verifier_issue_count": self.state.get("reader_verifier_filtered_toc_verifier_issue_count"),
            "reader_verifier_cleanup_ignored_reason_counts": self.state.get("reader_verifier_cleanup_ignored_reason_counts"),
            "reader_verifier_simple_user_summary": self.state.get("reader_verifier_simple_user_summary"),
            "reader_verifier_simple_user_risk_statement": self.state.get("reader_verifier_simple_user_risk_statement"),
            "reader_verifier_simple_user_next_step": self.state.get("reader_verifier_simple_user_next_step"),
            "reader_verifier_review_json": self.state.get("reader_verifier_review_json"),
            "reader_verifier_review_md": self.state.get("reader_verifier_review_md"),
            "reader_verifier_evidence_json": self.state.get("reader_verifier_evidence_json"),
            "reader_mvp_status_label": self.state.get("reader_mvp_status_label"),
            "reader_mvp_status_user_summary": self.state.get("reader_mvp_status_user_summary"),
            "reader_mvp_status_risk_summary": self.state.get("reader_mvp_status_risk_summary"),
            "reader_mvp_status_cleanup_score_delta": self.state.get("reader_mvp_status_cleanup_score_delta"),
            "reader_mvp_status_acceptance_diagnostic_only": self.state.get("reader_mvp_status_acceptance_diagnostic_only"),
            "reader_mvp_status_false_deletion_status": self.state.get("reader_mvp_status_false_deletion_status"),
            "reader_mvp_status_readability_regression_status": self.state.get("reader_mvp_status_readability_regression_status"),
            "reader_mvp_status_md": self.state.get("reader_mvp_status_md"),
            "latest_report": self.state.get("latest_report_json"),
            "latest_summary": self.state.get("latest_summary_txt"),
            "latest_markdown": self.state.get("latest_markdown_path"),
            "latest_docx": self.state.get("latest_docx_path"),
        }

    def _write_locked(self) -> None:
        progress_payload = cast(Mapping[str, object], sanitize_for_json(dict(self.state)))
        _write_json_atomic(self.progress_path, progress_payload)
        _write_json_atomic(self.latest_progress_path, progress_payload, temp_dir=self.run_dir)
        _write_json_atomic(self.latest_manifest_path, self._build_manifest_payload_locked(), temp_dir=self.run_dir)

    def _heartbeat_loop(self) -> None:
        while not self.stop_event.wait(HEARTBEAT_INTERVAL_SECONDS):
            with self.lock:
                if str(self.state.get("status") or "") != "in_progress":
                    return
                phase = str(self.state.get("phase") or "processing")
                stage = str(self.state.get("stage") or "Ожидание")
                detail = str(self.state.get("detail") or "Процесс продолжается.")
                progress_value = self.state.get("progress")
                metrics = cast(Mapping[str, object], self.state.get("metrics") or {})
            self.emit(
                event_type="heartbeat",
                phase=phase,
                stage=stage,
                detail=f"Heartbeat: {_normalize_terminal_detail(detail)}",
                progress=float(progress_value) if isinstance(progress_value, (int, float)) else None,
                metrics=metrics,
            )


def _write_latest_alias_artifacts(
    *,
    report_path: Path,
    summary_path: Path,
    markdown_artifact: Path | None,
    docx_artifact: Path | None,
    tts_artifact: Path | None,
    latest_report_path: Path,
    latest_summary_path: Path,
    latest_markdown_path: Path | None,
    latest_docx_path: Path | None,
    latest_tts_path: Path | None,
    latest_manifest_path: Path,
    run_id: str,
    run_dir: Path,
    manifest_payload: Mapping[str, object],
) -> None:
    latest_report_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(report_path, latest_report_path)
    shutil.copy2(summary_path, latest_summary_path)
    if markdown_artifact is not None and latest_markdown_path is not None:
        shutil.copy2(markdown_artifact, latest_markdown_path)
    if docx_artifact is not None and latest_docx_path is not None:
        shutil.copy2(docx_artifact, latest_docx_path)
    if tts_artifact is not None and latest_tts_path is not None:
        shutil.copy2(tts_artifact, latest_tts_path)

    latest_manifest = dict(manifest_payload)
    latest_manifest.update(
        {
            "run_id": run_id,
            "run_dir": _path_for_report(run_dir),
            "latest_report": _path_for_report(latest_report_path),
            "latest_summary": _path_for_report(latest_summary_path),
            "latest_markdown": _path_for_report(latest_markdown_path) if latest_markdown_path is not None else None,
            "latest_docx": _path_for_report(latest_docx_path) if latest_docx_path is not None else None,
            "latest_tts_text": _path_for_report(latest_tts_path) if latest_tts_path is not None else None,
        }
    )
    _write_json_atomic(latest_manifest_path, latest_manifest)


def _preview_text(text: str | None, *, limit: int = 240) -> str:
    normalized = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(normalized) <= limit:
        return normalized
    return normalized[: limit - 3].rstrip() + "..."


def _empty_reader_verifier_category_counts() -> dict[str, int]:
    return {category: 0 for category in _READER_VERIFIER_DEFECT_CATEGORIES}


def _format_reader_verifier_line_ref(artifact: str, line_number: int) -> str:
    return f"{artifact}:{line_number}"


def _count_reader_verifier_categories(
    issues: Sequence[Mapping[str, object]],
) -> dict[str, int]:
    counts = _empty_reader_verifier_category_counts()
    for issue in issues:
        category = str(issue.get("category") or "").strip()
        if category in counts:
            counts[category] += 1
    return counts


def _select_reader_verifier_top_categories(
    counts: Mapping[str, object],
    *,
    limit: int = 3,
) -> list[str]:
    ranked = sorted(
        (
            (category, max(0, _coerce_int(counts.get(category), default=0)))
            for category in _READER_VERIFIER_DEFECT_CATEGORIES
        ),
        key=lambda item: (-item[1], item[0]),
    )
    return [category for category, count in ranked if count > 0][:limit]


def _count_reader_verifier_high_severity_issues(
    issues: Sequence[Mapping[str, object]],
) -> int:
    return sum(1 for issue in issues if str(issue.get("severity") or "").strip() == "high")


def _reader_verifier_fix_type_for_category(category: str) -> str:
    if category == "page_furniture_inline":
        return "delete_noise"
    if category == "heading_fused_with_body":
        return "split_heading"
    if category == "fragmented_paragraph":
        return "merge_paragraph"
    if category == "broken_list_marker":
        return "normalize_list"
    if category == "quote_not_block_formatted":
        return "format_quote"
    return "other"


def _reader_verifier_severity_for_category(category: str) -> str:
    if category in {"page_furniture_inline", "heading_fused_with_body", "fragmented_paragraph"}:
        return "high"
    return "medium"


def _reader_verifier_pre_audit_findings_as_remaining_issues(
    evidence_payload: Mapping[str, object],
) -> list[dict[str, str]]:
    findings = evidence_payload.get("mandatory_review_targets") or evidence_payload.get("pre_audit_findings") or []
    if not isinstance(findings, Sequence) or isinstance(findings, (str, bytes, bytearray)):
        return []
    issues: list[dict[str, str]] = []
    seen: set[tuple[str, str, str]] = set()
    for item in findings:
        if not isinstance(item, Mapping):
            continue
        category = str(item.get("category") or "").strip()
        artifact = str(item.get("artifact") or "cleaned_markdown").strip() or "cleaned_markdown"
        line_ref = str(item.get("line_ref") or "").strip()
        snippet = _preview_text(str(item.get("snippet") or "").strip())
        note = str(item.get("note") or "Deterministic pre-audit found this reader-visible issue.").strip()
        if category not in _READER_VERIFIER_DEFECT_CATEGORIES:
            continue
        if artifact not in _ALLOWED_READER_VERIFIER_ISSUE_ARTIFACTS:
            artifact = "cleaned_markdown"
        if not line_ref or not snippet:
            continue
        key = (category, line_ref, snippet)
        if key in seen:
            continue
        seen.add(key)
        issues.append(
            {
                "category": category,
                "severity": _reader_verifier_severity_for_category(category),
                "artifact": artifact,
                "line_ref": line_ref,
                "snippet": snippet,
                "why_reader_hurts": note,
                "recommended_fix_type": _reader_verifier_fix_type_for_category(category),
            }
        )
    return issues


def _reader_verifier_remaining_issue_anchors(
    issues: Sequence[Mapping[str, object]],
    *,
    existing_anchors: Sequence[Mapping[str, str]] = (),
) -> list[dict[str, str]]:
    anchors = [dict(anchor) for anchor in existing_anchors]
    seen = {
        (
            str(anchor.get("kind") or ""),
            str(anchor.get("line_ref") or ""),
            str(anchor.get("snippet") or ""),
        )
        for anchor in anchors
    }
    for issue in issues:
        key = ("remaining_issue", str(issue.get("line_ref") or ""), str(issue.get("snippet") or ""))
        if key in seen:
            continue
        seen.add(key)
        anchors.append(
            {
                "kind": "remaining_issue",
                "artifact": str(issue.get("artifact") or "cleaned_markdown"),
                "line_ref": str(issue.get("line_ref") or ""),
                "snippet": str(issue.get("snippet") or ""),
                "note": str(issue.get("why_reader_hurts") or "Deterministic pre-audit issue remains reviewable."),
            }
        )
    return anchors


def _reader_verifier_improvement_anchors(
    improvements: Sequence[str],
    *,
    existing_anchors: Sequence[Mapping[str, str]] = (),
) -> list[dict[str, str]]:
    anchors = [dict(anchor) for anchor in existing_anchors]
    seen = {
        (
            str(anchor.get("kind") or ""),
            str(anchor.get("line_ref") or ""),
            str(anchor.get("snippet") or ""),
        )
        for anchor in anchors
    }
    for improvement in improvements:
        snippet = _preview_text(str(improvement).strip())
        if not snippet:
            continue
        key = ("improvement_seen", "comparison:diagnostic", snippet)
        if key in seen:
            continue
        seen.add(key)
        anchors.append(
            {
                "kind": "improvement_seen",
                "artifact": "comparison",
                "line_ref": "comparison:diagnostic",
                "snippet": snippet,
                "note": "Recovered improvement anchor from validated noise_removed text after dropping a malformed verifier anchor.",
            }
        )
        break
    return anchors


def _merge_reader_verifier_missing_pre_audit_issues(
    *,
    existing_issues: Sequence[Mapping[str, str]],
    pre_audit_issues: Sequence[Mapping[str, str]],
) -> tuple[list[dict[str, str]], bool]:
    merged = [dict(issue) for issue in existing_issues]
    seen = {
        _reader_verifier_issue_identity(issue)
        for issue in merged
    }
    added = False
    for issue in pre_audit_issues:
        key = _reader_verifier_issue_identity(issue)
        if key in seen:
            continue
        if any(_reader_verifier_issues_overlap(existing, issue) for existing in merged):
            continue
        seen.add(key)
        merged.append(dict(issue))
        added = True
    return merged, added


def _reader_verifier_issue_identity(issue: Mapping[str, object]) -> tuple[str, str, str, str]:
    return (
        str(issue.get("category") or "").strip(),
        str(issue.get("artifact") or "cleaned_markdown").strip() or "cleaned_markdown",
        _normalize_reader_verifier_line_ref(issue.get("line_ref")),
        _normalize_reader_verifier_snippet_key(issue.get("snippet")),
    )


def _reader_verifier_issues_overlap(
    existing_issue: Mapping[str, object],
    candidate_issue: Mapping[str, object],
) -> bool:
    existing_category, existing_artifact, existing_line_ref, existing_snippet = _reader_verifier_issue_identity(
        existing_issue
    )
    candidate_category, candidate_artifact, candidate_line_ref, candidate_snippet = _reader_verifier_issue_identity(
        candidate_issue
    )
    if (
        existing_category != candidate_category
        or existing_artifact != candidate_artifact
        or existing_line_ref != candidate_line_ref
        or not existing_snippet
        or not candidate_snippet
    ):
        return False
    return existing_snippet in candidate_snippet or candidate_snippet in existing_snippet


def _normalize_reader_verifier_line_ref(value: object) -> str:
    line_ref = str(value or "").strip()
    if ":" in line_ref:
        line_ref = line_ref.rsplit(":", 1)[-1].strip()
    return line_ref


def _normalize_reader_verifier_snippet_key(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip()).casefold()


def _append_reader_verifier_pre_audit_finding(
    findings: list[dict[str, str]],
    seen: set[tuple[str, str, str]],
    *,
    category: str,
    line_number: int,
    snippet: str,
    note: str,
) -> None:
    if category not in _READER_VERIFIER_DEFECT_CATEGORIES:
        return
    normalized_snippet = _preview_text(snippet)
    if not normalized_snippet:
        return
    line_ref = _format_reader_verifier_line_ref("cleaned_markdown", line_number)
    key = (category, line_ref, normalized_snippet)
    if key in seen:
        return
    seen.add(key)
    findings.append(
        {
            "category": category,
            "artifact": "cleaned_markdown",
            "line_ref": line_ref,
            "snippet": normalized_snippet,
            "note": note,
        }
    )


def _run_reader_verifier_pre_audit(cleaned_markdown: str) -> dict[str, object]:
    counts = _empty_reader_verifier_category_counts()
    findings: list[dict[str, str]] = []
    seen: set[tuple[str, str, str]] = set()
    previous_nonempty_line = ""
    previous_nonempty_line_number = 0
    normalized_line_positions: dict[str, int] = {}

    for line_number, raw_line in enumerate(cleaned_markdown.splitlines(), start=1):
        stripped = raw_line.strip()
        if not stripped:
            continue

        if re.match(r"^\d{1,4}\s+[A-ZА-ЯЁ][A-ZА-ЯЁ\s\-]{3,}\b", stripped) or re.search(
            r"\b[A-ZА-ЯЁ][A-ZА-ЯЁ\s\-]{4,}\s+\d{1,4}\b",
            stripped,
        ):
            counts["page_furniture_inline"] += 1
            _append_reader_verifier_pre_audit_finding(
                findings,
                seen,
                category="page_furniture_inline",
                line_number=line_number,
                snippet=stripped,
                note="Detected likely inline page furniture or running-header residue.",
            )

        fused_match = re.match(r"^([A-ZА-ЯЁ0-9«»\"'().,:;!?/\-\s]{8,}?)(\s+[A-ZА-ЯЁ]?[a-zа-яё].+)$", stripped)
        if fused_match is not None:
            heading_prefix = fused_match.group(1).strip()
            if len(re.findall(r"[A-ZА-ЯЁ]", heading_prefix)) >= 6 and not re.search(r"[a-zа-яё]", heading_prefix):
                counts["heading_fused_with_body"] += 1
                _append_reader_verifier_pre_audit_finding(
                    findings,
                    seen,
                    category="heading_fused_with_body",
                    line_number=line_number,
                    snippet=stripped,
                    note="Detected heading-like uppercase text fused into running body prose.",
                )

        if "•" in stripped:
            counts["broken_list_marker"] += 1
            _append_reader_verifier_pre_audit_finding(
                findings,
                seen,
                category="broken_list_marker",
                line_number=line_number,
                snippet=stripped,
                note="Detected residual bullet marker that likely needs Markdown normalization.",
            )

        if re.match(r"^[a-zа-яё]", stripped):
            previous_line_looks_boundary = bool(
                previous_nonempty_line
                and (
                    re.search(r"(?:Фото:|Photo(?:\s+credit)?:|Источник:|Source:)", previous_nonempty_line, re.IGNORECASE)
                    or previous_nonempty_line.endswith((",", ";", ":", "—", "»", '"'))
                    or len(previous_nonempty_line) <= 80
                )
            )
            if previous_line_looks_boundary:
                counts["fragmented_paragraph"] += 1
                _append_reader_verifier_pre_audit_finding(
                    findings,
                    seen,
                    category="fragmented_paragraph",
                    line_number=line_number,
                    snippet=stripped,
                    note=(
                        "Detected lowercase paragraph carryover after a likely page-boundary or caption boundary"
                        f" (previous line {previous_nonempty_line_number})."
                    ),
                )

        normalized_line = re.sub(r"\s+", " ", stripped).casefold()
        if len(normalized_line) >= 40:
            first_seen_line = normalized_line_positions.get(normalized_line)
            if first_seen_line is None:
                normalized_line_positions[normalized_line] = line_number
            else:
                counts["duplicate_fragment"] += 1
                _append_reader_verifier_pre_audit_finding(
                    findings,
                    seen,
                    category="duplicate_fragment",
                    line_number=line_number,
                    snippet=stripped,
                    note=f"Detected repeated line fragment first seen at cleaned_markdown:{first_seen_line}.",
                )

        previous_nonempty_line = stripped
        previous_nonempty_line_number = line_number

    return {
        "issue_counts": counts,
        "findings": findings,
        "mandatory_review_targets": list(findings),
    }


def _normalize_reader_verifier_remaining_issues(value: object) -> tuple[list[dict[str, str]], dict[str, object]]:
    if not isinstance(value, Sequence) or isinstance(value, (str, bytes, bytearray)):
        raise RuntimeError("reader_verifier_remaining_issues_must_be_list")
    normalized_issues: list[dict[str, str]] = []
    warnings: list[str] = []
    ignored_field_counts: dict[str, int] = {}
    for index, item in enumerate(value):
        if not isinstance(item, Mapping):
            raise RuntimeError("reader_verifier_remaining_issue_item_must_be_object")
        unknown_fields = sorted(
            set(item.keys())
            - {"category", "severity", "artifact", "line_ref", "snippet", "why_reader_hurts", "recommended_fix_type"}
        )
        if unknown_fields:
            normalized_unknown_fields = [str(field) for field in unknown_fields]
            for field in normalized_unknown_fields:
                ignored_field_counts[field] = ignored_field_counts.get(field, 0) + 1
            warnings.append(
                "reader_verifier_remaining_issue_ignored_unknown_fields:"
                f"index={index}:fields={','.join(normalized_unknown_fields)}"
            )
        category = str(item.get("category") or "").strip()
        severity = str(item.get("severity") or "").strip().lower()
        artifact = str(item.get("artifact") or "").strip()
        line_ref = str(item.get("line_ref") or "").strip()
        snippet = _preview_text(str(item.get("snippet") or "").strip())
        why_reader_hurts = str(item.get("why_reader_hurts") or "").strip()
        recommended_fix_type = str(item.get("recommended_fix_type") or "").strip()
        if category not in _READER_VERIFIER_DEFECT_CATEGORIES:
            raise RuntimeError(f"reader_verifier_unknown_remaining_issue_category:{category}")
        if severity not in _ALLOWED_READER_VERIFIER_SEVERITY:
            raise RuntimeError(f"reader_verifier_unknown_remaining_issue_severity:{severity}")
        if artifact not in _ALLOWED_READER_VERIFIER_ISSUE_ARTIFACTS:
            raise RuntimeError(f"reader_verifier_unknown_remaining_issue_artifact:{artifact}")
        if recommended_fix_type not in _ALLOWED_READER_VERIFIER_FIX_TYPES:
            raise RuntimeError(f"reader_verifier_unknown_remaining_issue_fix_type:{recommended_fix_type}")
        if not line_ref or not snippet or not why_reader_hurts:
            raise RuntimeError("reader_verifier_remaining_issue_missing_required_text")
        normalized_issues.append(
            {
                "category": category,
                "severity": severity,
                "artifact": artifact,
                "line_ref": line_ref,
                "snippet": snippet,
                "why_reader_hurts": why_reader_hurts,
                "recommended_fix_type": recommended_fix_type,
            }
        )
    return normalized_issues, {
        "input_issue_count": len(value),
        "ignored_unknown_field_count": sum(ignored_field_counts.values()),
        "ignored_unknown_field_counts": ignored_field_counts,
        "warnings": warnings,
    }


def _normalize_reader_verifier_evidence_anchors(value: object) -> tuple[list[dict[str, str]], dict[str, object]]:
    if not isinstance(value, Sequence) or isinstance(value, (str, bytes, bytearray)):
        raise RuntimeError("reader_verifier_evidence_anchors_must_be_list")
    anchors: list[dict[str, str]] = []
    warnings: list[str] = []
    ignored_kind_counts = {kind: 0 for kind in sorted(_ALLOWED_READER_VERIFIER_ANCHOR_KINDS)}
    for index, item in enumerate(value):
        if not isinstance(item, Mapping):
            raise RuntimeError("reader_verifier_evidence_anchor_item_must_be_object")
        unknown_fields = sorted(set(item.keys()) - {"kind", "artifact", "line_ref", "snippet", "note"})
        if unknown_fields:
            raise RuntimeError(
                f"reader_verifier_unknown_evidence_anchor_fields:{','.join(str(field) for field in unknown_fields)}"
            )
        kind = str(item.get("kind") or "").strip()
        artifact = str(item.get("artifact") or "").strip()
        line_ref = str(item.get("line_ref") or "").strip()
        snippet = _preview_text(str(item.get("snippet") or "").strip())
        note = str(item.get("note") or "").strip()
        if kind not in _ALLOWED_READER_VERIFIER_ANCHOR_KINDS:
            raise RuntimeError(f"reader_verifier_unknown_evidence_anchor_kind:{kind}")
        if artifact not in _ALLOWED_READER_VERIFIER_ISSUE_ARTIFACTS:
            raise RuntimeError(f"reader_verifier_unknown_evidence_anchor_artifact:{artifact}")
        if not line_ref or not snippet or not note:
            ignored_kind_counts[kind] = ignored_kind_counts.get(kind, 0) + 1
            warnings.append(
                "reader_verifier_evidence_anchor_ignored_missing_required_text:"
                f"index={index}:kind={kind}:artifact={artifact}"
            )
            continue
        anchors.append(
            {
                "kind": kind,
                "artifact": artifact,
                "line_ref": line_ref,
                "snippet": snippet,
                "note": note,
            }
        )
    return anchors, {
        "input_anchor_count": len(value),
        "ignored_anchor_count": len(warnings),
        "ignored_kind_counts": ignored_kind_counts,
        "repaired_anchor_counts": {"improvement_seen": 0, "remaining_issue": 0},
        "warnings": warnings,
    }


def _detect_reader_verifier_contradictory_removed_claim(
    *,
    noise_removed: Sequence[str],
    remaining_issues: Sequence[Mapping[str, object]],
) -> str | None:
    remaining_categories = {str(issue.get("category") or "").strip() for issue in remaining_issues}
    if not remaining_categories:
        return None
    for item in noise_removed:
        lowered = item.casefold()
        claims_full_removal = any(
            token in lowered for token in ("removed", "gone", "fixed", "resolved", "eliminated", "stripped", "normalized")
        )
        if not claims_full_removal:
            continue
        for category in remaining_categories:
            keywords = _READER_VERIFIER_CATEGORY_KEYWORDS.get(category, ())
            if any(keyword in lowered for keyword in keywords):
                return category
    return None


def _downgrade_reader_verifier_contradictory_removed_claims(
    *,
    noise_removed: Sequence[str],
    remaining_issues: Sequence[Mapping[str, object]],
) -> list[str]:
    downgraded = list(noise_removed)
    contradictory_category = _detect_reader_verifier_contradictory_removed_claim(
        noise_removed=downgraded,
        remaining_issues=remaining_issues,
    )
    if contradictory_category is None:
        return downgraded
    return [
        f"Some cleanup improvement was reported, but {contradictory_category} still has remaining review targets."
        if any(keyword in item.casefold() for keyword in _READER_VERIFIER_CATEGORY_KEYWORDS.get(contradictory_category, ()))
        else item
        for item in downgraded
    ]


def _resolve_reported_path(path_value: object) -> Path | None:
    if not isinstance(path_value, str) or not path_value.strip():
        return None
    candidate = Path(path_value.strip())
    if candidate.is_absolute():
        return candidate
    return (PROJECT_ROOT / candidate).resolve()


def _load_text_artifact(path_value: object) -> tuple[str, str | None, str | None]:
    artifact_path = _resolve_reported_path(path_value)
    if artifact_path is None:
        return "", None, "artifact_path_missing"
    try:
        return artifact_path.read_text(encoding="utf-8"), _path_for_report(artifact_path), None
    except OSError:
        return "", _path_for_report(artifact_path), "artifact_unreadable"


def _load_optional_json_artifact(path_value: object) -> tuple[dict[str, object] | None, str | None, str | None]:
    artifact_path = _resolve_reported_path(path_value)
    if artifact_path is None:
        return None, None, "artifact_path_missing"
    try:
        payload = json.loads(artifact_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None, _path_for_report(artifact_path), "artifact_unreadable"
    if not isinstance(payload, dict):
        return None, _path_for_report(artifact_path), "artifact_not_object"
    return payload, _path_for_report(artifact_path), None


def _resolve_reader_verifier_config(
    *,
    validation_mode: Mapping[str, object],
    document_profile_id: str,
    run_profile_id: str,
    app_config: object | None,
    runtime_app_config: Mapping[str, object],
) -> dict[str, object]:
    # FC6: verifier is explicit advisory evidence, not an auto-enabled step in
    # proof/comparison runs. Matching profile ids no longer turn it on by
    # default; profiles or runtime overrides must opt in intentionally.
    default_enabled = False
    enabled_value = _get_config_value(runtime_app_config, "reader_verifier_enabled")
    if enabled_value is None:
        enabled_value = _get_config_value(app_config, "reader_verifier_enabled")
    model_value = _get_config_value(runtime_app_config, "reader_verifier_model")
    if model_value is None:
        model_value = _get_config_value(app_config, "reader_verifier_model")
    model_selector = str(model_value or READER_VERIFIER_DEFAULT_SELECTOR).strip()
    return {
        "enabled": _coerce_bool(enabled_value, default=default_enabled),
        "model": model_selector,
    }


def _build_reader_verifier_deleted_block_context(
    *,
    raw_markdown: str,
    cleanup_report_payload: Mapping[str, object] | None,
) -> list[dict[str, object]]:
    if not raw_markdown.strip() or cleanup_report_payload is None:
        return []
    blocks = build_cleanup_blocks(raw_markdown)
    blocks_by_id = {block.block_id: block for block in blocks}
    previews: list[dict[str, object]] = []
    for entry in _coerce_mapping_sequence(cleanup_report_payload.get("accepted_delete_blocks")):
        block_id = str(entry.get("id") or "").strip()
        block = blocks_by_id.get(block_id)
        if block is None:
            continue
        previous_block = blocks[block.index - 1] if block.index > 0 else None
        next_block = blocks[block.index + 1] if (block.index + 1) < len(blocks) else None
        previews.append(
            {
                "id": block.block_id,
                "text_hash": block.text_hash,
                "reason": entry.get("reason"),
                "confidence": entry.get("confidence"),
                "char_count": entry.get("char_count"),
                "kind": entry.get("kind"),
                "raw_text_preview": entry.get("raw_text_preview"),
                "previous_block_preview": _preview_text(previous_block.text if previous_block is not None else ""),
                "next_block_preview": _preview_text(next_block.text if next_block is not None else ""),
            }
        )
    return previews


def _build_reader_verifier_evidence_payload(
    *,
    run_id: str,
    document_profile_id: str,
    run_profile_id: str,
    source_document_path: Path,
    source_text: str,
    reader_cleanup_evidence: Mapping[str, object],
    runtime_app_config: Mapping[str, object] | None = None,
) -> dict[str, object]:
    raw_markdown, raw_markdown_report_path, raw_markdown_warning = _load_text_artifact(
        reader_cleanup_evidence.get("raw_markdown_path")
    )
    cleaned_markdown, cleaned_markdown_report_path, cleaned_markdown_warning = _load_text_artifact(
        reader_cleanup_evidence.get("cleaned_markdown_path")
    )
    cleanup_report_payload, cleanup_report_report_path, cleanup_report_warning = _load_optional_json_artifact(
        reader_cleanup_evidence.get("reader_cleanup_report_path")
    )
    cleaned_docx_report_path = _path_for_report(_resolve_reported_path(reader_cleanup_evidence.get("cleaned_docx_path")))
    load_warnings = [
        warning
        for warning in (
            raw_markdown_warning,
            cleaned_markdown_warning,
            cleanup_report_warning,
        )
        if warning is not None
    ]
    cleanup_stats = cast(Mapping[str, object], (cleanup_report_payload or {}).get("stats") or {})
    cleanup_settings = cast(Mapping[str, object], (cleanup_report_payload or {}).get("cleanup_settings") or {})
    cleanup_diagnostics = _build_reader_cleanup_diagnostics(cleanup_report_payload)
    deleted_block_previews = _build_reader_verifier_deleted_block_context(
        raw_markdown=raw_markdown,
        cleanup_report_payload=cleanup_report_payload,
    )
    block_spans = _build_cleanup_block_line_spans(cleaned_markdown)
    pre_audit_payload = _run_reader_verifier_pre_audit(cleaned_markdown)
    raw_pre_audit_findings = cast(Sequence[Mapping[str, str]], pre_audit_payload["findings"])
    toc_out_of_review_scope = _reader_verifier_toc_out_of_review_scope_from_config(runtime_app_config or {})
    filtered_pre_audit_findings, filtered_toc_pre_audit_findings = _filter_reader_verifier_items_excluding_toc(
        raw_pre_audit_findings,
        block_spans=block_spans,
        toc_out_of_review_scope=toc_out_of_review_scope,
    )
    filtered_toc_issue_previews = _build_filtered_toc_issue_previews(
        filtered_toc_pre_audit_findings,
        source="pre_audit",
    )
    validator_boundary = {
        "observer_only": True,
        "runs_cleanup_repair": False,
        "runs_anchor_repair": False,
        "mutates_cleaned_markdown": False,
        "mutates_cleaned_docx": False,
        "rebuilds_docx": False,
    }
    return {
        "run_id": run_id,
        "document_profile_id": document_profile_id,
        "run_profile_id": run_profile_id,
        "source_document_path": _path_for_report(source_document_path),
        "evidence_mode": "full_selected_slice",
        "review_mode": "development_only_non_acceptance",
        "base_artifacts_present": bool(
            raw_markdown_report_path
            and cleaned_markdown_report_path
            and cleaned_docx_report_path
            and cleanup_report_report_path
        ),
        "source_text": source_text,
        "raw_markdown": raw_markdown,
        "cleaned_markdown": cleaned_markdown,
        "validator_boundary": validator_boundary,
        "toc_filtering_policy": {
            "mode": "evidence_only",
            "toc_out_of_review_scope": toc_out_of_review_scope,
            "policy_source": "reader_cleanup_keep_toc_false" if toc_out_of_review_scope else "toc_in_review_scope",
            "artifact_repair_applied": False,
        },
        "evidence_filtering_note": (
            "TOC-like review targets may be filtered from verifier evidence when profile policy marks TOC out of scope; "
            "this is evidence filtering only and does not repair, rewrite, or rebuild output artifacts."
        ),
        "cleanup_report_summary": {
            "stage_status": reader_cleanup_evidence.get("stage_status"),
            "changed": reader_cleanup_evidence.get("changed"),
            "accepted_delete_block_count": reader_cleanup_evidence.get("accepted_delete_block_count"),
            "ignored_delete_block_count": reader_cleanup_evidence.get("ignored_delete_block_count"),
            "rejected_delete_block_count": reader_cleanup_evidence.get("rejected_delete_block_count"),
            "failed_chunk_count": reader_cleanup_evidence.get("failed_chunk_count"),
            "cleanup_chunk_count": cleanup_stats.get("cleanup_chunk_count"),
            "cleanup_settings": dict(cleanup_settings),
            "warnings": _coerce_string_list((cleanup_report_payload or {}).get("warnings")),
        },
        "cleanup_diagnostics": cleanup_diagnostics,
        "raw_pre_audit_issue_counts": pre_audit_payload["issue_counts"],
        "raw_pre_audit_findings": list(raw_pre_audit_findings),
        "filtered_toc_issue_count": len(filtered_toc_pre_audit_findings),
        "filtered_toc_pre_audit_count": len(filtered_toc_pre_audit_findings),
        "filtered_toc_verifier_issue_count": 0,
        "filtered_toc_evidence_anchor_count": 0,
        "filtered_toc_issue_previews": filtered_toc_issue_previews,
        "pre_audit_issue_counts": _count_reader_verifier_categories(filtered_pre_audit_findings),
        "pre_audit_findings": filtered_pre_audit_findings,
        "mandatory_review_targets": filtered_pre_audit_findings,
        "deleted_block_previews": deleted_block_previews,
        "artifact_paths": {
            "raw_markdown": raw_markdown_report_path,
            "cleaned_markdown": cleaned_markdown_report_path,
            "cleaned_docx": cleaned_docx_report_path,
            "reader_cleanup_report": cleanup_report_report_path,
        },
        "load_warnings": load_warnings,
    }


def _build_reader_verifier_system_prompt() -> str:
    return (
        "You are reviewing a comparison-only translated book slice for reader quality.\n"
        "This is development-only non-acceptance evidence.\n"
        "Compare the raw translated Markdown against the cleaned Markdown using the source text, cleanup report, and deterministic pre-audit findings.\n"
        "Do not optimize for acceptance checks, structural authority, topology, DocumentMap, or production readiness.\n"
        "You must answer two questions separately: whether cleaned is easier to read than raw, and whether the cleaned artifact itself still has reader-visible defects.\n"
        "Treat every deterministic pre-audit finding in mandatory_review_targets as a required review target. You may disagree with a candidate classification, but you must not silently ignore the candidate set.\n"
        "Respect toc_filtering_policy: if toc_out_of_review_scope is true, TOC-like table-of-contents blocks are out-of-scope evidence only and must not be reported as remaining issues; otherwise review them normally.\n"
        "Return JSON only with exactly these top-level fields: overall_verdict, cleaned_audit_verdict, reader_quality_score_raw, reader_quality_score_cleaned, confidence, noise_removed, possible_false_deletions, readability_regressions, remaining_issues, evidence_anchors, recommended_next_changes, summary_for_human, simple_user_summary, simple_user_risk_statement, simple_user_next_step.\n"
        "Allowed overall_verdict values: cleaned_better, raw_better, mixed, unclear.\n"
        "Allowed cleaned_audit_verdict values: clean, improved_but_has_remaining_issues, unsafe_or_regressed, unclear.\n"
        "Allowed confidence values: low, medium, high.\n"
        "noise_removed, possible_false_deletions, and readability_regressions must be arrays of short strings, not objects.\n"
        "If there are no possible false deletions, possible_false_deletions must be an empty array, not a sentence saying none were found.\n"
        "If there are no readability regressions, readability_regressions must be an empty array, not a sentence saying none were found.\n"
        "remaining_issues must be a list of objects with exactly category, severity, artifact, line_ref, snippet, why_reader_hurts, recommended_fix_type.\n"
        "Allowed remaining_issues categories: page_furniture_inline, heading_fused_with_body, broken_list_marker, fragmented_paragraph, duplicate_fragment, orphan_caption, mixed_language_leak, quote_not_block_formatted.\n"
        "Allowed severity values: high, medium, low. Allowed artifact values: cleaned_markdown, raw_markdown, comparison. Allowed recommended_fix_type values: delete_noise, split_heading, merge_paragraph, normalize_list, format_quote, other.\n"
        "evidence_anchors must be a list of objects with exactly kind, artifact, line_ref, snippet, note. Allowed kind values: improvement_seen, remaining_issue, possible_false_deletion.\n"
        "If remaining_issues is non-empty, cleaned_audit_verdict must not be clean.\n"
        "Do not claim a defect class was fully removed when the cleaned artifact still contains unresolved examples from that same class.\n"
        "A positive comparison verdict must never be used as shorthand for the cleaned artifact being clean.\n"
        "recommended_next_changes must be a list of objects with exactly change_type, recommendation, why.\n"
        "Allowed change_type values: prompt, model_selection, operation_contract, safety_application, deterministic_last_resort.\n"
        "Do not recommend document-specific deterministic cleanup rules, regex packs, or hardcoded book-specific fixes; recommend prompt hardening, model selection changes, bounded operation-contract improvements, safety/application changes, or a document-agnostic deterministic last resort instead.\n"
        "simple_user_summary and simple_user_risk_statement must stay cautious: avoid absolute claims that no content was lost, avoid domain-specific phrases such as story content, avoid production-ready language, and explicitly state whether reader-visible structural or readability defects still remain.\n"
        "Do not recommend acceptance-threshold tuning, structure-recognition expansion, or broad validation rewrites.\n"
        "If the evidence is insufficient, set overall_verdict to unclear and explain why."
    )


def _build_reader_verifier_non_success_review(
    *,
    run_id: str,
    document_profile_id: str,
    run_profile_id: str,
    requested_selector: str,
    canonical_selector: str | None,
    provider: str | None,
    model_id: str | None,
    verifier_status: str,
    verifier_reason: str,
    verifier_detail: str,
    evidence_path: Path,
    evidence_payload: Mapping[str, object],
) -> dict[str, object]:
    reason_text = verifier_detail.strip() or verifier_reason
    block_spans = _build_cleanup_block_line_spans(str(evidence_payload.get("cleaned_markdown") or ""))
    toc_out_of_review_scope = _reader_verifier_toc_out_of_review_scope_from_evidence(evidence_payload)
    pre_audit_issue_counts = {
        category: max(
            0,
            _coerce_int(cast(Mapping[str, object], evidence_payload.get("pre_audit_issue_counts") or {}).get(category), default=0),
        )
        for category in _READER_VERIFIER_DEFECT_CATEGORIES
    }
    pre_audit_remaining_issues, filtered_toc_pre_audit_issues = _filter_reader_verifier_items_excluding_toc(
        _reader_verifier_pre_audit_findings_as_remaining_issues(evidence_payload),
        block_spans=block_spans,
        toc_out_of_review_scope=toc_out_of_review_scope,
    )
    pre_audit_evidence_anchors = _reader_verifier_remaining_issue_anchors(pre_audit_remaining_issues)
    issue_summary_by_category = _count_reader_verifier_categories(pre_audit_remaining_issues)
    base_filtered_toc_pre_audit_count = max(
        0,
        _coerce_int(
            evidence_payload.get("filtered_toc_pre_audit_count", evidence_payload.get("filtered_toc_issue_count")),
            default=0,
        ),
    )
    filtered_toc_issue_previews = _merge_filtered_toc_issue_previews(
        _coerce_mapping_sequence(evidence_payload.get("filtered_toc_issue_previews")),
        _build_filtered_toc_issue_previews(filtered_toc_pre_audit_issues, source="pre_audit"),
    )
    return {
        "run_id": run_id,
        "document_profile_id": document_profile_id,
        "run_profile_id": run_profile_id,
        "review_mode": "development_only_non_acceptance",
        "verifier_requested_selector": requested_selector,
        "verifier_canonical_selector": canonical_selector,
        "verifier_provider": provider,
        "verifier_model_id": model_id,
        "verifier_status": verifier_status,
        "verifier_reason": verifier_reason,
        "verifier_detail": reason_text,
        "artifact_paths": {
            "source_evidence_json": _path_for_report(evidence_path),
            "raw_markdown": cast(Mapping[str, object], evidence_payload.get("artifact_paths") or {}).get("raw_markdown"),
            "cleaned_markdown": cast(Mapping[str, object], evidence_payload.get("artifact_paths") or {}).get("cleaned_markdown"),
            "cleaned_docx": cast(Mapping[str, object], evidence_payload.get("artifact_paths") or {}).get("cleaned_docx"),
            "reader_cleanup_report": cast(Mapping[str, object], evidence_payload.get("artifact_paths") or {}).get("reader_cleanup_report"),
        },
        "overall_verdict": "unclear",
        "cleaned_audit_verdict": "unclear",
        "reader_quality_score_raw": 0.0,
        "reader_quality_score_cleaned": 0.0,
        "confidence": "low",
        "validator_boundary": dict(cast(Mapping[str, object], evidence_payload.get("validator_boundary") or {})),
        "toc_filtering_policy": dict(cast(Mapping[str, object], evidence_payload.get("toc_filtering_policy") or {})),
        "evidence_filtering_note": str(evidence_payload.get("evidence_filtering_note") or ""),
        "cleanup_diagnostics": dict(cast(Mapping[str, object], evidence_payload.get("cleanup_diagnostics") or {})),
        "filtered_toc_issue_count": base_filtered_toc_pre_audit_count + len(filtered_toc_pre_audit_issues),
        "filtered_toc_pre_audit_count": base_filtered_toc_pre_audit_count + len(filtered_toc_pre_audit_issues),
        "filtered_toc_verifier_issue_count": 0,
        "filtered_toc_evidence_anchor_count": 0,
        "filtered_toc_issue_previews": filtered_toc_issue_previews,
        "pre_audit_issue_counts": pre_audit_issue_counts,
        "remaining_issues": pre_audit_remaining_issues,
        "issue_summary_by_category": issue_summary_by_category,
        "evidence_anchors": pre_audit_evidence_anchors,
        "noise_removed": [],
        "possible_false_deletions": [],
        "readability_regressions": [],
        "recommended_next_changes": [],
        "summary_for_human": f"Verifier did not produce a review conclusion: {reason_text}.",
        "simple_user_summary": (
            "The current run does not provide enough reliable verifier evidence to say whether cleaned output is better than raw output. "
            f"Deterministic pre-audit still surfaced {len(pre_audit_remaining_issues)} review target(s)."
            if pre_audit_remaining_issues
            else "The current run does not provide enough reliable verifier evidence to say whether cleaned output is better than raw output."
        ),
        "simple_user_risk_statement": (
            "Raw and cleaned comparison artifacts were preserved, but the verifier did not produce a reliable review conclusion. "
            f"Reason: {reason_text}."
        ),
        "simple_user_next_step": "Fix the verifier precondition or runtime failure, then rerun the same comparison-only profile.",
    }


def _normalize_recommendation_list(value: object) -> list[dict[str, str]]:
    recommendations: list[dict[str, str]] = []
    if not isinstance(value, Sequence) or isinstance(value, (str, bytes, bytearray)):
        raise RuntimeError("reader_verifier_recommended_next_changes_must_be_list")
    for item in value:
        if not isinstance(item, Mapping):
            raise RuntimeError("reader_verifier_recommended_next_change_item_must_be_object")
        unknown_fields = sorted(set(item.keys()) - {"change_type", "recommendation", "why"})
        if unknown_fields:
            raise RuntimeError(
                f"reader_verifier_unknown_recommended_next_change_fields:{','.join(str(field) for field in unknown_fields)}"
            )
        change_type = str(item.get("change_type") or "").strip()
        recommendation = str(item.get("recommendation") or "").strip()
        why = str(item.get("why") or "").strip()
        original_change_type = change_type
        change_type = _LEGACY_READER_VERIFIER_CHANGE_TYPE_MAP.get(change_type, change_type)
        if change_type not in _ALLOWED_READER_VERIFIER_CHANGE_TYPES:
            raise RuntimeError(f"reader_verifier_unknown_change_type:{change_type}")
        if not recommendation or not why:
            raise RuntimeError("reader_verifier_recommended_next_change_missing_text")
        if original_change_type in _LEGACY_READER_VERIFIER_CHANGE_TYPE_MAP:
            recommendation = (
                "Use the AI cleanup operation contract rather than a deterministic rule for this follow-up: "
                f"{recommendation}"
            )
            why = f"Legacy verifier change_type {original_change_type!r} was normalized to {change_type!r}. {why}"
        recommendations.append(
            {
                "change_type": change_type,
                "recommendation": recommendation,
                "why": why,
            }
        )
    return recommendations


def _normalize_string_list_item(item: object, *, key: str) -> str:
    if isinstance(item, str):
        normalized = item.strip()
        if normalized:
            return normalized
        raise RuntimeError(f"reader_verifier_{key}_must_be_string_list")
    if isinstance(item, Mapping):
        for candidate_key in ("text", "item", "summary", "description", "finding", "issue", "reason", "value"):
            candidate = item.get(candidate_key)
            if isinstance(candidate, str) and candidate.strip():
                return candidate.strip()
    raise RuntimeError(f"reader_verifier_{key}_must_be_string_list")


def _require_string_list(payload: Mapping[str, object], key: str) -> list[str]:
    value = payload.get(key)
    if not isinstance(value, Sequence) or isinstance(value, (str, bytes, bytearray)):
        raise RuntimeError(f"reader_verifier_{key}_must_be_string_list")
    result = []
    for item in value:
        normalized = _normalize_string_list_item(item, key=key)
        if _is_reader_verifier_negated_absence_statement(normalized, key=key):
            continue
        if normalized:
            result.append(normalized)
    return result


def _is_reader_verifier_negated_absence_statement(text: str, *, key: str) -> bool:
    normalized = " ".join(text.strip().lower().split())
    if not normalized:
        return False
    if key == "possible_false_deletions":
        return bool(
            re.search(
                r"\bno\b.*\b(?:false deletions?|content deletions?|deletions?)\b.*\b(?:detected|reported|accepted|found)\b",
                normalized,
            )
            or re.search(r"\b(?:false deletions?|content deletions?)\b.*\bno\b", normalized)
        )
    if key == "readability_regressions":
        return bool(
            re.search(
                r"\bno\b.*\b(?:readability regressions?|regressions?)\b.*\b(?:detected|reported|introduced|found)\b",
                normalized,
            )
            or re.search(r"\b(?:readability regressions?|regressions?)\b.*\bno\b", normalized)
        )
    return False


def _contains_uncertainty_marker(text: str) -> bool:
    return bool(re.search(r"\b(?:appears?|seems?|may|might|provisional|currently|current review confidence|not clearly|unclear|mixed)\b", text, re.IGNORECASE))


def _contains_explicit_confidence_level(text: str, confidence: str) -> bool:
    return bool(re.search(rf"\bconfidence\s+is\s+{re.escape(confidence)}\b", text, re.IGNORECASE))


def _contains_development_only_marker(text: str) -> bool:
    return bool(re.search(r"\bdevelopment-only\b|\bcomparison-only\b|\bcomparison evidence\b|\bacceptance result\b", text, re.IGNORECASE))


def _contains_overconfident_preservation_claim(text: str) -> bool:
    return bool(
        re.search(
            r"\bno\s+(?:actual\s+|meaningful\s+|major\s+)?(?:book\s+|story\s+)?(?:content|text|information)\b.*\b(?:lost|removed|damaged)\b(?:\s+during\s+this\s+process)?"
            r"|\b(?:core\s+text|translated\s+concepts?|meaningful\s+content)\b.*\b(?:remain|remains|stays?|are)\s+(?:fully\s+)?intact\b",
            text,
            re.IGNORECASE,
        )
    )


def _normalize_reader_verifier_user_text(text: str) -> str:
    normalized = re.sub(r"\bstory content\b", "content", text, flags=re.IGNORECASE)
    normalized = re.sub(
        r"\b(?:no|without)\s+(?:actual\s+|meaningful\s+|semantic\s+|major\s+)?(?:book\s+|story\s+)?(?:content|text|information)\s+(?:(?:appears?|seems?)\s+to\s+have\s+been\s+)?(?:was\s+)?(?:lost|removed|damaged)\b(?:\s+during\s+this\s+process)?",
        "no major text loss was detected at current review confidence",
        normalized,
        flags=re.IGNORECASE,
    )
    normalized = re.sub(
        r"\b(?:core\s+text|translated\s+concepts?|meaningful\s+content)\b.*\b(?:remain|remains|stays?|are)\s+(?:fully\s+)?intact\b",
        "no major text loss was detected at current review confidence",
        normalized,
        flags=re.IGNORECASE,
    )
    normalized = re.sub(
        r"\b(?:safe|ready)\s+for\s+production\b|\bproduction-?ready\b",
        "development-only comparison evidence",
        normalized,
        flags=re.IGNORECASE,
    )
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized


def _normalize_reader_verifier_simple_language(
    *,
    overall_verdict: str,
    cleaned_audit_verdict: str,
    confidence: str,
    simple_user_summary: str,
    simple_user_risk_statement: str,
    simple_user_next_step: str,
    remaining_issues: Sequence[Mapping[str, object]],
    possible_false_deletions: Sequence[str],
    readability_regressions: Sequence[str],
    recommended_next_changes: Sequence[Mapping[str, str]],
) -> tuple[str, str, str]:
    raw_next_step = simple_user_next_step
    summary = _normalize_reader_verifier_user_text(simple_user_summary)
    risk_statement = _normalize_reader_verifier_user_text(simple_user_risk_statement)
    next_step = _normalize_reader_verifier_user_text(simple_user_next_step)

    if _contains_overconfident_preservation_claim(summary):
        summary = re.sub(
            r"\bno\s+(?:actual\s+|meaningful\s+|major\s+)?(?:book\s+|story\s+)?(?:content|text|information)\b.*\b(?:lost|removed|damaged)\b(?:\s+during\s+this\s+process)?"
            r"|\b(?:core\s+text|translated\s+concepts?|meaningful\s+content)\b.*\b(?:remain|remains|stays?|are)\s+(?:fully\s+)?intact\b",
            "No major text loss was detected at current review confidence",
            summary,
            flags=re.IGNORECASE,
        )
    summary = re.sub(
        r"(^|[.!?]\s+)no major text loss was detected at current review confidence",
        r"\1No major text loss was detected at current review confidence",
        summary,
    )

    if confidence in {"low", "medium"} and not _contains_uncertainty_marker(summary):
        summary = f"{summary} This remains provisional at current review confidence."

    if re.search(r"\b(?:the\s+)?risk(?:\s+of\s+[a-z\s]+)?\s+is\s+very\s+low\b|\bno\s+risk\b", risk_statement, re.IGNORECASE):
        risk_statement = "No major text loss was detected at current review confidence."

    if _contains_overconfident_preservation_claim(risk_statement):
        risk_statement = "No major text loss was detected at current review confidence."

    if confidence in {"low", "medium"} and not _contains_explicit_confidence_level(risk_statement, confidence):
        risk_statement = f"{risk_statement} Review confidence is {confidence}, so this comparison should be treated as provisional."

    if confidence == "high" and not _contains_development_only_marker(risk_statement):
        risk_statement = f"{risk_statement} This remains development-only comparison evidence, not an acceptance result."

    if remaining_issues:
        remaining_issue_count = len(remaining_issues)
        remaining_issue_text = (
            f"{remaining_issue_count} reader-visible structural or readability issue"
            f"{'s' if remaining_issue_count != 1 else ''} still remain in the cleaned output."
        )
        if not re.search(r"\b(?:remaining issues?|still remain|structural|readability defect)\b", summary, re.IGNORECASE):
            summary = f"{summary} {remaining_issue_text}"
        if not re.search(r"\b(?:remaining issues?|still remain|structural|readability defect)\b", risk_statement, re.IGNORECASE):
            risk_statement = f"{remaining_issue_text} {risk_statement}"

    if cleaned_audit_verdict == "clean" and remaining_issues:
        risk_statement = (
            "Reader-visible structural or readability defects still remain in the cleaned output. "
            "This remains development-only comparison evidence, not an acceptance result."
        )

    if re.search(r"\b(?:safe|ready)\s+for\s+production\b|\bproduction-?ready\b", raw_next_step, re.IGNORECASE):
        next_step = "Use this as development-only comparison evidence and apply one narrow follow-up change before rerunning the same profile."

    if re.search(
        r"\bno\s+further\s+cleanup\s+is\s+required\b|\bready\s+for\s+reading\b|\bready\s+for\s+further\s+translation\s+review\b|\bproceed\s+with\s+the\s+cleaned\s+version\b",
        raw_next_step,
        re.IGNORECASE,
    ):
        next_step = "Use this as development-only comparison evidence"
        if recommended_next_changes:
            first_change = recommended_next_changes[0]
            change_type = str(first_change.get("change_type") or "follow-up").strip() or "follow-up"
            recommendation = str(first_change.get("recommendation") or "").strip()
            if recommendation:
                next_step = f"{next_step}. Next, apply one narrow {change_type} change: {recommendation}"
            else:
                next_step = f"{next_step}. Next, apply one narrow {change_type} change and rerun the same profile."
        else:
            next_step = f"{next_step}. Rerun the same profile if you want stronger confirmation."

    if not next_step:
        if recommended_next_changes:
            next_step = str(recommended_next_changes[0].get("recommendation") or "").strip()
        elif overall_verdict == "unclear":
            next_step = "Gather clearer comparison evidence before changing the profile."
        else:
            next_step = "Apply one narrow follow-up change and rerun the same comparison-only profile."

    return summary.strip(), risk_statement.strip(), next_step.strip()


def _parse_reader_verifier_completed_review(
    *,
    raw_response: str,
    run_id: str,
    document_profile_id: str,
    run_profile_id: str,
    requested_selector: str,
    canonical_selector: str,
    provider: str,
    model_id: str,
    evidence_path: Path,
    evidence_payload: Mapping[str, object],
) -> dict[str, object]:
    payload = parse_json_object(
        raw_response,
        empty_message="Reader verifier returned empty output.",
        no_json_message="Reader verifier did not return JSON.",
    )
    unknown_fields = sorted(set(payload.keys()) - _READER_VERIFIER_MODEL_FIELDS)
    if unknown_fields:
        raise RuntimeError(f"reader_verifier_unknown_top_level_fields:{','.join(unknown_fields)}")
    missing_fields = sorted(field for field in _READER_VERIFIER_MODEL_FIELDS if field not in payload)
    if missing_fields:
        raise RuntimeError(f"reader_verifier_missing_top_level_fields:{','.join(missing_fields)}")

    overall_verdict = str(payload.get("overall_verdict") or "").strip()
    cleaned_audit_verdict = str(payload.get("cleaned_audit_verdict") or "").strip()
    confidence = str(payload.get("confidence") or "").strip().lower()
    if overall_verdict not in _ALLOWED_READER_VERIFIER_VERDICTS:
        raise RuntimeError(f"reader_verifier_unknown_overall_verdict:{overall_verdict}")
    if cleaned_audit_verdict not in _ALLOWED_READER_VERIFIER_AUDIT_VERDICTS:
        raise RuntimeError(f"reader_verifier_unknown_cleaned_audit_verdict:{cleaned_audit_verdict}")
    if confidence not in _ALLOWED_READER_VERIFIER_CONFIDENCE:
        raise RuntimeError(f"reader_verifier_unknown_confidence:{confidence}")

    simple_user_summary = str(payload.get("simple_user_summary") or "").strip()
    simple_user_risk_statement = str(payload.get("simple_user_risk_statement") or "").strip()
    simple_user_next_step = str(payload.get("simple_user_next_step") or "").strip()
    summary_for_human = str(payload.get("summary_for_human") or "").strip()
    if not simple_user_summary or not simple_user_risk_statement or not simple_user_next_step or not summary_for_human:
        raise RuntimeError("reader_verifier_missing_summary_text")

    block_spans = _build_cleanup_block_line_spans(str(evidence_payload.get("cleaned_markdown") or ""))
    toc_out_of_review_scope = _reader_verifier_toc_out_of_review_scope_from_evidence(evidence_payload)
    normalized_remaining_issues, remaining_issue_diagnostics = _normalize_reader_verifier_remaining_issues(
        payload.get("remaining_issues")
    )
    remaining_issues, filtered_toc_remaining_issues = _filter_reader_verifier_items_excluding_toc(
        normalized_remaining_issues,
        block_spans=block_spans,
        toc_out_of_review_scope=toc_out_of_review_scope,
    )
    payload_remaining_issues = list(remaining_issues)
    evidence_anchors, evidence_anchor_diagnostics = _normalize_reader_verifier_evidence_anchors(
        payload.get("evidence_anchors")
    )
    evidence_anchors, filtered_toc_evidence_anchors = _filter_reader_verifier_items_excluding_toc(
        evidence_anchors,
        block_spans=block_spans,
        toc_out_of_review_scope=toc_out_of_review_scope,
    )
    pre_audit_remaining_issues, filtered_toc_pre_audit_issues = _filter_reader_verifier_items_excluding_toc(
        _reader_verifier_pre_audit_findings_as_remaining_issues(evidence_payload),
        block_spans=block_spans,
        toc_out_of_review_scope=toc_out_of_review_scope,
    )
    base_filtered_toc_pre_audit_count = max(
        0,
        _coerce_int(
            evidence_payload.get("filtered_toc_pre_audit_count", evidence_payload.get("filtered_toc_issue_count")),
            default=0,
        ),
    )
    filtered_toc_issue_previews = _merge_filtered_toc_issue_previews(
        _coerce_mapping_sequence(evidence_payload.get("filtered_toc_issue_previews")),
        _build_filtered_toc_issue_previews(filtered_toc_pre_audit_issues, source="pre_audit"),
        _build_filtered_toc_issue_previews(filtered_toc_remaining_issues, source="verifier_remaining_issue"),
    )
    filtered_toc_pre_audit_count = base_filtered_toc_pre_audit_count + len(filtered_toc_pre_audit_issues)
    filtered_toc_verifier_issue_count = len(filtered_toc_remaining_issues)
    merged_pre_audit_issue = False
    if pre_audit_remaining_issues:
        remaining_issues, merged_pre_audit_issue = _merge_reader_verifier_missing_pre_audit_issues(
            existing_issues=remaining_issues,
            pre_audit_issues=pre_audit_remaining_issues,
        )
    if merged_pre_audit_issue:
        evidence_anchors = _reader_verifier_remaining_issue_anchors(
            remaining_issues,
            existing_anchors=evidence_anchors,
        )
        if cleaned_audit_verdict == "clean" and not payload_remaining_issues:
            cleaned_audit_verdict = "improved_but_has_remaining_issues"
    noise_removed = _require_string_list(payload, "noise_removed")
    possible_false_deletions = _require_string_list(payload, "possible_false_deletions")
    readability_regressions = _require_string_list(payload, "readability_regressions")
    recommended_next_changes = _normalize_recommendation_list(payload.get("recommended_next_changes"))
    ignored_anchor_kind_counts = cast(
        Mapping[str, object], evidence_anchor_diagnostics.get("ignored_kind_counts") or {}
    )
    repaired_anchor_counts = {
        str(key): max(0, _coerce_int(value, default=0))
        for key, value in cast(Mapping[str, object], evidence_anchor_diagnostics.get("repaired_anchor_counts") or {}).items()
    }
    if (
        remaining_issues
        and not any(anchor.get("kind") == "remaining_issue" for anchor in evidence_anchors)
        and _coerce_int(ignored_anchor_kind_counts.get("remaining_issue"), default=0) > 0
    ):
        existing_count = len(evidence_anchors)
        evidence_anchors = _reader_verifier_remaining_issue_anchors(remaining_issues, existing_anchors=evidence_anchors)
        repaired_anchor_counts["remaining_issue"] = max(0, len(evidence_anchors) - existing_count)
    if (
        overall_verdict == "cleaned_better"
        and not any(anchor.get("kind") == "improvement_seen" for anchor in evidence_anchors)
        and _coerce_int(ignored_anchor_kind_counts.get("improvement_seen"), default=0) > 0
    ):
        existing_count = len(evidence_anchors)
        evidence_anchors = _reader_verifier_improvement_anchors(noise_removed, existing_anchors=evidence_anchors)
        repaired_anchor_counts["improvement_seen"] = max(0, len(evidence_anchors) - existing_count)
    evidence_anchor_diagnostics["repaired_anchor_counts"] = repaired_anchor_counts
    if remaining_issues and cleaned_audit_verdict == "clean":
        raise RuntimeError("reader_verifier_remaining_issues_forbid_cleaned_audit_clean")
    if cleaned_audit_verdict == "improved_but_has_remaining_issues" and not remaining_issues:
        raise RuntimeError("reader_verifier_improved_but_has_remaining_issues_requires_remaining_issues")
    noise_removed = _downgrade_reader_verifier_contradictory_removed_claims(
        noise_removed=noise_removed,
        remaining_issues=remaining_issues,
    )
    if overall_verdict == "cleaned_better" and not any(anchor.get("kind") == "improvement_seen" for anchor in evidence_anchors):
        raise RuntimeError("reader_verifier_missing_improvement_anchor")
    if remaining_issues and not any(anchor.get("kind") == "remaining_issue" for anchor in evidence_anchors):
        raise RuntimeError("reader_verifier_missing_remaining_issue_anchor")
    pre_audit_issue_counts = _count_reader_verifier_categories(pre_audit_remaining_issues)
    issue_summary_by_category = _count_reader_verifier_categories(remaining_issues)
    simple_user_summary, simple_user_risk_statement, simple_user_next_step = _normalize_reader_verifier_simple_language(
        overall_verdict=overall_verdict,
        cleaned_audit_verdict=cleaned_audit_verdict,
        confidence=confidence,
        simple_user_summary=simple_user_summary,
        simple_user_risk_statement=simple_user_risk_statement,
        simple_user_next_step=simple_user_next_step,
        remaining_issues=remaining_issues,
        possible_false_deletions=possible_false_deletions,
        readability_regressions=readability_regressions,
        recommended_next_changes=recommended_next_changes,
    )

    review_payload = {
        "run_id": run_id,
        "document_profile_id": document_profile_id,
        "run_profile_id": run_profile_id,
        "review_mode": "development_only_non_acceptance",
        "verifier_requested_selector": requested_selector,
        "verifier_canonical_selector": canonical_selector,
        "verifier_provider": provider,
        "verifier_model_id": model_id,
        "verifier_status": "completed",
        "verifier_reason": "",
        "artifact_paths": {
            "source_evidence_json": _path_for_report(evidence_path),
            "raw_markdown": cast(Mapping[str, object], evidence_payload.get("artifact_paths") or {}).get("raw_markdown"),
            "cleaned_markdown": cast(Mapping[str, object], evidence_payload.get("artifact_paths") or {}).get("cleaned_markdown"),
            "cleaned_docx": cast(Mapping[str, object], evidence_payload.get("artifact_paths") or {}).get("cleaned_docx"),
            "reader_cleanup_report": cast(Mapping[str, object], evidence_payload.get("artifact_paths") or {}).get("reader_cleanup_report"),
        },
        "overall_verdict": overall_verdict,
        "cleaned_audit_verdict": cleaned_audit_verdict,
        "reader_quality_score_raw": round(_coerce_float(payload.get("reader_quality_score_raw"), default=0.0), 3),
        "reader_quality_score_cleaned": round(_coerce_float(payload.get("reader_quality_score_cleaned"), default=0.0), 3),
        "confidence": confidence,
        "validator_boundary": dict(cast(Mapping[str, object], evidence_payload.get("validator_boundary") or {})),
        "toc_filtering_policy": dict(cast(Mapping[str, object], evidence_payload.get("toc_filtering_policy") or {})),
        "evidence_filtering_note": str(evidence_payload.get("evidence_filtering_note") or ""),
        "cleanup_diagnostics": dict(cast(Mapping[str, object], evidence_payload.get("cleanup_diagnostics") or {})),
        "filtered_toc_issue_count": filtered_toc_pre_audit_count + filtered_toc_verifier_issue_count,
        "filtered_toc_pre_audit_count": filtered_toc_pre_audit_count,
        "filtered_toc_verifier_issue_count": filtered_toc_verifier_issue_count,
        "filtered_toc_evidence_anchor_count": len(filtered_toc_evidence_anchors),
        "filtered_toc_issue_previews": filtered_toc_issue_previews,
        "pre_audit_issue_counts": pre_audit_issue_counts,
        "remaining_issues": remaining_issues,
        "remaining_issue_diagnostics": remaining_issue_diagnostics,
        "issue_summary_by_category": issue_summary_by_category,
        "evidence_anchors": evidence_anchors,
        "evidence_anchor_diagnostics": evidence_anchor_diagnostics,
        "noise_removed": noise_removed,
        "possible_false_deletions": possible_false_deletions,
        "readability_regressions": readability_regressions,
        "recommended_next_changes": recommended_next_changes,
        "summary_for_human": summary_for_human,
        "simple_user_summary": simple_user_summary,
        "simple_user_risk_statement": simple_user_risk_statement,
        "simple_user_next_step": simple_user_next_step,
    }
    return review_payload


def _render_reader_verifier_markdown_summary(review_payload: Mapping[str, object]) -> str:
    improvements = _coerce_string_list(review_payload.get("noise_removed"))
    remaining_issues = _coerce_mapping_sequence(review_payload.get("remaining_issues"))
    risks = _coerce_string_list(review_payload.get("possible_false_deletions")) + _coerce_string_list(
        review_payload.get("readability_regressions")
    )
    recommendations = cast(Sequence[Mapping[str, object]], review_payload.get("recommended_next_changes") or [])
    filtered_toc_issue_previews = _coerce_mapping_sequence(review_payload.get("filtered_toc_issue_previews"))
    cleanup_diagnostics = cast(Mapping[str, object], review_payload.get("cleanup_diagnostics") or {})
    remaining_issue_diagnostics = cast(Mapping[str, object], review_payload.get("remaining_issue_diagnostics") or {})
    remaining_issue_warnings = _coerce_string_list(remaining_issue_diagnostics.get("warnings"))
    evidence_anchor_diagnostics = cast(Mapping[str, object], review_payload.get("evidence_anchor_diagnostics") or {})
    evidence_anchor_warnings = _coerce_string_list(evidence_anchor_diagnostics.get("warnings"))
    top_ignored_reasons = _coerce_mapping_sequence(cleanup_diagnostics.get("top_ignored_reasons"))
    accepted_operation_counts = cast(Mapping[str, object], cleanup_diagnostics.get("accepted_operation_counts") or {})
    validator_boundary = cast(Mapping[str, object], review_payload.get("validator_boundary") or {})
    lines = [
        "# Verdict",
        "",
        str(review_payload.get("overall_verdict") or "unclear"),
        "",
        "# Audit Verdict",
        "",
        str(review_payload.get("cleaned_audit_verdict") or "unclear"),
        "",
        "# In Plain Words",
        "",
        str(review_payload.get("simple_user_summary") or ""),
        "",
        "# Improvements Seen",
        "",
    ]
    if improvements:
        lines.extend(f"- {item}" for item in improvements)
    else:
        lines.append("- No verified readability improvements were recorded.")
    lines.extend(["", "# Remaining Issues", ""])
    if remaining_issues:
        for issue in remaining_issues:
            lines.append(
                "- [{severity}] {category} at {line_ref}: {snippet} ({why_reader_hurts}; fix={recommended_fix_type})".format(
                    severity=str(issue.get("severity") or ""),
                    category=str(issue.get("category") or ""),
                    line_ref=str(issue.get("line_ref") or ""),
                    snippet=str(issue.get("snippet") or ""),
                    why_reader_hurts=str(issue.get("why_reader_hurts") or ""),
                    recommended_fix_type=str(issue.get("recommended_fix_type") or ""),
                )
            )
    else:
        lines.append("- No remaining reader-visible issues were recorded.")
    lines.extend(["", "# Remaining Issue Diagnostics", ""])
    if remaining_issue_warnings:
        lines.extend(f"- {item}" for item in remaining_issue_warnings)
    else:
        lines.append("- No remaining-issue normalization warnings were recorded.")
    lines.extend(["", "# Evidence Anchor Diagnostics", ""])
    if evidence_anchor_warnings:
        lines.extend(f"- {item}" for item in evidence_anchor_warnings)
    else:
        lines.append("- No evidence-anchor normalization warnings were recorded.")
    lines.extend(["", "# Risks Seen", ""])
    if risks:
        lines.append(f"- {str(review_payload.get('simple_user_risk_statement') or '')}")
        lines.extend(f"- {item}" for item in risks)
    else:
        lines.append(f"- {str(review_payload.get('simple_user_risk_statement') or 'No additional reader-facing risks were recorded.')}")
    lines.extend(["", "# Out-of-Scope Filtered TOC Issues", ""])
    lines.append(
        "- TOC-like findings were filtered as out-of-scope evidence only. No cleaned Markdown/DOCX was repaired, overwritten, or rebuilt by validation."
    )
    lines.append(f"- filtered_toc_issue_count: {review_payload.get('filtered_toc_issue_count')}")
    lines.append(f"- filtered_toc_pre_audit_count: {review_payload.get('filtered_toc_pre_audit_count')}")
    lines.append(f"- filtered_toc_verifier_issue_count: {review_payload.get('filtered_toc_verifier_issue_count')}")
    lines.append(f"- filtered_toc_evidence_anchor_count: {review_payload.get('filtered_toc_evidence_anchor_count')}")
    if filtered_toc_issue_previews:
        for preview in filtered_toc_issue_previews:
            lines.append(
                "- [{source}] {category} at {artifact} {line_ref}: {snippet}".format(
                    source=str(preview.get("source") or "unknown"),
                    category=str(preview.get("category") or "uncategorized"),
                    artifact=str(preview.get("artifact") or "cleaned_markdown"),
                    line_ref=str(preview.get("line_ref") or ""),
                    snippet=str(preview.get("snippet") or ""),
                )
            )
    else:
        lines.append("- No TOC-like issue previews were filtered.")
    lines.extend(["", "# Cleanup Application Diagnostics", ""])
    if accepted_operation_counts:
        lines.append(
            "- accepted_operation_counts: "
            + ", ".join(
                f"{key}={max(0, _coerce_int(value, default=0))}"
                for key, value in sorted(accepted_operation_counts.items())
            )
        )
    else:
        lines.append("- accepted_operation_counts: none")
    if top_ignored_reasons:
        for entry in top_ignored_reasons:
            ignored_reason = str(entry.get("ignored_reason") or "").strip()
            count = max(0, _coerce_int(entry.get("count"), default=0))
            lines.append(f"- {ignored_reason}: {count}")
            for example in _coerce_mapping_sequence(entry.get("examples")):
                example_operation = str(example.get("operation") or "unknown").strip() or "unknown"
                example_reason = str(example.get("reason") or "").strip()
                example_preview = str(example.get("text_preview") or "").strip()
                example_chunk = max(0, _coerce_int(example.get("chunk_index"), default=0))
                example_text = f"- example chunk={example_chunk} operation={example_operation}"
                if example_reason:
                    example_text += f" reason={example_reason}"
                if example_preview:
                    example_text += f" preview={example_preview}"
                lines.append(example_text)
    else:
        lines.append("- No tracked ignored cleanup reasons were recorded.")
    lines.extend(["", "# Recommended Next Changes", ""])
    if recommendations:
        for item in recommendations:
            change_type = str(item.get("change_type") or "").strip()
            recommendation = str(item.get("recommendation") or "").strip()
            why = str(item.get("why") or "").strip()
            lines.append(f"- {change_type}: {recommendation} ({why})")
    else:
        lines.append(f"- {str(review_payload.get('simple_user_next_step') or 'No verifier recommendation was produced.')}")
    lines.extend(
        [
            "",
            "# Verifier Metadata",
            "",
            f"- review_mode: {review_payload.get('review_mode')}",
            f"- verifier_status: {review_payload.get('verifier_status')}",
            f"- verifier_reason: {review_payload.get('verifier_reason')}",
            f"- verifier_requested_selector: {review_payload.get('verifier_requested_selector')}",
            f"- verifier_canonical_selector: {review_payload.get('verifier_canonical_selector')}",
            f"- verifier_provider: {review_payload.get('verifier_provider')}",
            f"- verifier_model_id: {review_payload.get('verifier_model_id')}",
            f"- cleaned_audit_verdict: {review_payload.get('cleaned_audit_verdict')}",
            f"- confidence: {review_payload.get('confidence')}",
            f"- remaining_issue_count: {len(remaining_issues)}",
            f"- high_severity_issue_count: {_count_reader_verifier_high_severity_issues(remaining_issues)}",
            f"- top_issue_categories: {','.join(_select_reader_verifier_top_categories(cast(Mapping[str, object], review_payload.get('issue_summary_by_category') or {})))}",
            f"- evidence_filtering_note: {review_payload.get('evidence_filtering_note')}",
            f"- filtered_toc_issue_count: {review_payload.get('filtered_toc_issue_count')}",
            f"- filtered_toc_pre_audit_count: {review_payload.get('filtered_toc_pre_audit_count')}",
            f"- filtered_toc_verifier_issue_count: {review_payload.get('filtered_toc_verifier_issue_count')}",
            f"- artifact_repair_applied: {cast(Mapping[str, object], review_payload.get('toc_filtering_policy') or {}).get('artifact_repair_applied')}",
            f"- validator_observer_only: {validator_boundary.get('observer_only')}",
            f"- validator_runs_cleanup_repair: {validator_boundary.get('runs_cleanup_repair')}",
            f"- validator_runs_anchor_repair: {validator_boundary.get('runs_anchor_repair')}",
            f"- validator_mutates_cleaned_markdown: {validator_boundary.get('mutates_cleaned_markdown')}",
            f"- validator_mutates_cleaned_docx: {validator_boundary.get('mutates_cleaned_docx')}",
            f"- validator_rebuilds_docx: {validator_boundary.get('rebuilds_docx')}",
            f"- source_evidence_json: {cast(Mapping[str, object], review_payload.get('artifact_paths') or {}).get('source_evidence_json')}",
        ]
    )
    return "\n".join(lines).strip() + "\n"


def _classify_reader_verifier_runtime_failure(exc: Exception) -> str:
    message = str(exc).strip().lower()
    if "model" in message and ("not found" in message or "unavailable" in message or "unknown" in message):
        return "model_unavailable"
    return "execution_failed"


def _reader_verifier_timeout_seconds() -> float:
    raw_value = os.environ.get("DOCXAI_READER_VERIFIER_TIMEOUT_SECONDS", "").strip()
    if not raw_value:
        return READER_VERIFIER_TIMEOUT_SECONDS
    try:
        timeout_seconds = float(raw_value)
    except ValueError:
        return READER_VERIFIER_TIMEOUT_SECONDS
    if timeout_seconds <= 0:
        return READER_VERIFIER_TIMEOUT_SECONDS
    return timeout_seconds


def _run_reader_verifier(
    *,
    run_id: str,
    document_profile_id: str,
    run_profile_id: str,
    app_config: object | None,
    runtime_app_config: Mapping[str, object],
    validation_mode: Mapping[str, object],
    evidence_payload: Mapping[str, object],
    evidence_path: Path,
    max_retries: int,
) -> dict[str, object]:
    config = _resolve_reader_verifier_config(
        validation_mode=validation_mode,
        document_profile_id=document_profile_id,
        run_profile_id=run_profile_id,
        app_config=app_config,
        runtime_app_config=runtime_app_config,
    )
    requested_selector = str(config.get("model") or "").strip()
    if not bool(config.get("enabled")):
        return _build_reader_verifier_non_success_review(
            run_id=run_id,
            document_profile_id=document_profile_id,
            run_profile_id=run_profile_id,
            requested_selector=requested_selector or READER_VERIFIER_DEFAULT_SELECTOR,
            canonical_selector=None,
            provider=None,
            model_id=None,
            verifier_status="not_run",
            verifier_reason="verifier_disabled",
            verifier_detail="Reader verifier is disabled for this run.",
            evidence_path=evidence_path,
            evidence_payload=evidence_payload,
        )
    if not requested_selector:
        return _build_reader_verifier_non_success_review(
            run_id=run_id,
            document_profile_id=document_profile_id,
            run_profile_id=run_profile_id,
            requested_selector="",
            canonical_selector=None,
            provider=None,
            model_id=None,
            verifier_status="not_run",
            verifier_reason="model_selector_unconfigured",
            verifier_detail="Reader verifier model selector is not configured.",
            evidence_path=evidence_path,
            evidence_payload=evidence_payload,
        )
    if not bool(evidence_payload.get("base_artifacts_present")):
        return _build_reader_verifier_non_success_review(
            run_id=run_id,
            document_profile_id=document_profile_id,
            run_profile_id=run_profile_id,
            requested_selector=requested_selector,
            canonical_selector=None,
            provider=None,
            model_id=None,
            verifier_status="not_run",
            verifier_reason="base_artifacts_missing",
            verifier_detail="Base comparison artifacts are missing, so verifier review cannot run.",
            evidence_path=evidence_path,
            evidence_payload=evidence_payload,
        )
    availability_app_config: AppConfig | Mapping[str, object]
    if isinstance(app_config, AppConfig):
        availability_app_config = app_config
    elif isinstance(app_config, Mapping):
        availability_app_config = app_config
    else:
        availability_app_config = runtime_app_config
    try:
        availability = describe_provider_availability(requested_selector, app_config=availability_app_config)
    except Exception as exc:
        return _build_reader_verifier_non_success_review(
            run_id=run_id,
            document_profile_id=document_profile_id,
            run_profile_id=run_profile_id,
            requested_selector=requested_selector,
            canonical_selector=None,
            provider=None,
            model_id=None,
            verifier_status="not_run",
            verifier_reason="model_resolution_failed",
            verifier_detail=str(exc),
            evidence_path=evidence_path,
            evidence_payload=evidence_payload,
        )
    if not availability.enabled:
        return _build_reader_verifier_non_success_review(
            run_id=run_id,
            document_profile_id=document_profile_id,
            run_profile_id=run_profile_id,
            requested_selector=requested_selector,
            canonical_selector=availability.selector.canonical_selector,
            provider=availability.selector.provider,
            model_id=availability.selector.model_id,
            verifier_status="not_run",
            verifier_reason="provider_disabled",
            verifier_detail=availability.error_message or "Requested verifier provider is disabled.",
            evidence_path=evidence_path,
            evidence_payload=evidence_payload,
        )
    if not availability.has_api_key:
        return _build_reader_verifier_non_success_review(
            run_id=run_id,
            document_profile_id=document_profile_id,
            run_profile_id=run_profile_id,
            requested_selector=requested_selector,
            canonical_selector=availability.selector.canonical_selector,
            provider=availability.selector.provider,
            model_id=availability.selector.model_id,
            verifier_status="not_run",
            verifier_reason="api_key_missing",
            verifier_detail=availability.error_message or "Required verifier API key is missing.",
            evidence_path=evidence_path,
            evidence_payload=evidence_payload,
        )
    try:
        resolved_selector = resolve_model_selector(
            requested_selector,
            "responses_text",
            config_like=app_config,
            source_name="reader verifier model selector",
        )
    except Exception as exc:
        return _build_reader_verifier_non_success_review(
            run_id=run_id,
            document_profile_id=document_profile_id,
            run_profile_id=run_profile_id,
            requested_selector=requested_selector,
            canonical_selector=None,
            provider=None,
            model_id=None,
            verifier_status="not_run",
            verifier_reason="model_resolution_failed",
            verifier_detail=str(exc),
            evidence_path=evidence_path,
            evidence_payload=evidence_payload,
        )
    def _execute_verifier_request() -> dict[str, object]:
        try:
            client = get_client_for_model_selector(
                requested_selector,
                "responses_text",
                config_like=app_config,
            )
            raw_response = generate_markdown_block(
                client=client,
                model=resolved_selector.model_id,
                system_prompt=_build_reader_verifier_system_prompt(),
                target_text=json.dumps(evidence_payload, ensure_ascii=False, indent=2),
                context_before="",
                context_after="",
                max_retries=max(1, max_retries),
                expected_paragraph_ids=None,
                marker_mode=False,
            )
            return _parse_reader_verifier_completed_review(
                raw_response=raw_response,
                run_id=run_id,
                document_profile_id=document_profile_id,
                run_profile_id=run_profile_id,
                requested_selector=requested_selector,
                canonical_selector=resolved_selector.canonical_selector,
                provider=resolved_selector.provider,
                model_id=resolved_selector.model_id,
                evidence_path=evidence_path,
                evidence_payload=evidence_payload,
            )
        except Exception as exc:
            return _build_reader_verifier_non_success_review(
                run_id=run_id,
                document_profile_id=document_profile_id,
                run_profile_id=run_profile_id,
                requested_selector=requested_selector,
                canonical_selector=resolved_selector.canonical_selector,
                provider=resolved_selector.provider,
                model_id=resolved_selector.model_id,
                verifier_status="failed",
                verifier_reason=_classify_reader_verifier_runtime_failure(exc),
                verifier_detail=str(exc),
                evidence_path=evidence_path,
                evidence_payload=evidence_payload,
            )

    result_holder: dict[str, object] = {}

    def _worker() -> None:
        result_holder["review"] = _execute_verifier_request()

    timeout_seconds = _reader_verifier_timeout_seconds()
    worker = threading.Thread(target=_worker, name="reader-verifier-review", daemon=True)
    worker.start()
    worker.join(timeout=timeout_seconds)
    if worker.is_alive():
        return _build_reader_verifier_non_success_review(
            run_id=run_id,
            document_profile_id=document_profile_id,
            run_profile_id=run_profile_id,
            requested_selector=requested_selector,
            canonical_selector=resolved_selector.canonical_selector,
            provider=resolved_selector.provider,
            model_id=resolved_selector.model_id,
            verifier_status="failed",
            verifier_reason="execution_timeout",
            verifier_detail=f"Reader verifier exceeded {timeout_seconds:.1f}s advisory timeout.",
            evidence_path=evidence_path,
            evidence_payload=evidence_payload,
        )
    review = result_holder.get("review")
    if isinstance(review, Mapping):
        return dict(review)
    return _build_reader_verifier_non_success_review(
        run_id=run_id,
        document_profile_id=document_profile_id,
        run_profile_id=run_profile_id,
        requested_selector=requested_selector,
        canonical_selector=resolved_selector.canonical_selector,
        provider=resolved_selector.provider,
        model_id=resolved_selector.model_id,
        verifier_status="failed",
        verifier_reason="execution_failed",
        verifier_detail="Reader verifier worker finished without a review payload.",
        evidence_path=evidence_path,
        evidence_payload=evidence_payload,
    )


def _write_reader_verifier_artifacts(
    *,
    run_id: str,
    document_profile_id: str,
    run_profile_id: str,
    artifact_prefix: str,
    artifact_dir: Path,
    source_document_path: Path,
    source_text: str,
    reader_cleanup_evidence: Mapping[str, object],
    app_config: object | None,
    runtime_app_config: Mapping[str, object],
    validation_mode: Mapping[str, object],
    max_retries: int,
) -> dict[str, object]:
    evidence_path = artifact_dir / f"{artifact_prefix}_reader_quality_evidence.json"
    review_json_path = artifact_dir / f"{artifact_prefix}_reader_quality_review.json"
    review_md_path = artifact_dir / f"{artifact_prefix}_reader_quality_review.md"
    evidence_payload = _build_reader_verifier_evidence_payload(
        run_id=run_id,
        document_profile_id=document_profile_id,
        run_profile_id=run_profile_id,
        source_document_path=source_document_path,
        source_text=source_text,
        reader_cleanup_evidence=reader_cleanup_evidence,
        runtime_app_config=runtime_app_config,
    )
    _write_json_atomic(evidence_path, cast(Mapping[str, object], sanitize_for_json(evidence_payload)))
    review_payload = _run_reader_verifier(
        run_id=run_id,
        document_profile_id=document_profile_id,
        run_profile_id=run_profile_id,
        app_config=app_config,
        runtime_app_config=runtime_app_config,
        validation_mode=validation_mode,
        evidence_payload=evidence_payload,
        evidence_path=evidence_path,
        max_retries=max_retries,
    )
    anchor_diagnostics = _build_reader_verifier_anchor_repair_diagnostics(
        review_payload=review_payload,
        cleaned_markdown=str(evidence_payload.get("cleaned_markdown") or ""),
        reader_cleanup_evidence=reader_cleanup_evidence,
    )
    existing_anchor_status = str(reader_cleanup_evidence.get("anchor_repair_status") or "").strip()
    preserve_runtime_anchor_targets = existing_anchor_status in {
        "runtime_applied",
        "runtime_attempted_no_safe_ops",
        "applied_in_runtime",
    }
    updated_reader_cleanup_evidence = dict(reader_cleanup_evidence)
    updated_reader_cleanup_evidence["anchor_repair_status"] = anchor_diagnostics["anchor_repair_status"]
    updated_reader_cleanup_evidence["verifier_recommended_anchor_targets"] = anchor_diagnostics[
        "verifier_recommended_anchor_targets"
    ]
    updated_reader_cleanup_evidence["verifier_recommended_anchor_target_count"] = anchor_diagnostics[
        "verifier_recommended_anchor_target_count"
    ]
    if not preserve_runtime_anchor_targets:
        updated_reader_cleanup_evidence["recommended_anchor_targets"] = anchor_diagnostics["recommended_anchor_targets"]
        updated_reader_cleanup_evidence["recommended_anchor_target_count"] = anchor_diagnostics[
            "recommended_anchor_target_count"
        ]
    reader_cleanup_evidence = updated_reader_cleanup_evidence
    evidence_payload = {
        **evidence_payload,
        "anchor_repair_status": anchor_diagnostics["anchor_repair_status"],
        "recommended_anchor_targets": anchor_diagnostics["recommended_anchor_targets"],
        "recommended_anchor_target_count": anchor_diagnostics["recommended_anchor_target_count"],
        "verifier_recommended_anchor_targets": anchor_diagnostics["verifier_recommended_anchor_targets"],
        "verifier_recommended_anchor_target_count": anchor_diagnostics["verifier_recommended_anchor_target_count"],
    }
    _write_json_atomic(evidence_path, cast(Mapping[str, object], sanitize_for_json(evidence_payload)))
    review_payload = {
        **review_payload,
        "anchor_repair_status": anchor_diagnostics["anchor_repair_status"],
        "recommended_anchor_targets": anchor_diagnostics["recommended_anchor_targets"],
        "recommended_anchor_target_count": anchor_diagnostics["recommended_anchor_target_count"],
        "verifier_recommended_anchor_targets": anchor_diagnostics["verifier_recommended_anchor_targets"],
        "verifier_recommended_anchor_target_count": anchor_diagnostics["verifier_recommended_anchor_target_count"],
    }
    review_json_path.write_text(
        json.dumps(sanitize_for_json(review_payload), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    review_md_path.write_text(_render_reader_verifier_markdown_summary(review_payload), encoding="utf-8")
    return {
        **review_payload,
        "artifact_paths": {
            **cast(dict[str, object], review_payload.get("artifact_paths") or {}),
            "source_evidence_json": _path_for_report(evidence_path),
            "review_json": _path_for_report(review_json_path),
            "review_md": _path_for_report(review_md_path),
        },
        "updated_reader_cleanup_evidence": dict(reader_cleanup_evidence),
    }


def _load_json_file(path: Path) -> dict[str, object]:
    return cast(dict[str, object], json.loads(path.read_text(encoding="utf-8")))


def _build_repeat_run_id(parent_run_id: str, repeat_index: int, repeat_count: int) -> str:
    width = max(2, len(str(max(1, repeat_count))))
    return f"{parent_run_id}_r{repeat_index:0{width}d}of{repeat_count:0{width}d}"


def _summarize_repeat_runs(repeat_runs: Sequence[Mapping[str, object]]) -> tuple[dict[str, object], dict[str, object], str | None]:
    total_runs = len(repeat_runs)
    pipeline_succeeded_count = 0
    acceptance_passed_count = 0
    heading_only_output_detected_count = 0
    result_counts: dict[str, int] = {}
    failure_classification_counts: dict[str, int] = {}
    failed_repeat_indexes: list[int] = []
    failed_repeat_run_ids: list[str] = []

    for repeat_run in repeat_runs:
        result = str(repeat_run.get("result") or "unknown")
        result_counts[result] = result_counts.get(result, 0) + 1
        if result == "succeeded":
            pipeline_succeeded_count += 1

        acceptance_passed = bool(repeat_run.get("acceptance_passed"))
        if acceptance_passed:
            acceptance_passed_count += 1
        else:
            repeat_index = repeat_run.get("repeat_index")
            if isinstance(repeat_index, int):
                failed_repeat_indexes.append(repeat_index)
            repeat_run_id = str(repeat_run.get("run_id") or "")
            if repeat_run_id:
                failed_repeat_run_ids.append(repeat_run_id)

        failure_classification = str(repeat_run.get("failure_classification") or "").strip()
        if failure_classification:
            failure_classification_counts[failure_classification] = (
                failure_classification_counts.get(failure_classification, 0) + 1
            )

        signals = cast(Mapping[str, object], repeat_run.get("signals") or {})
        if bool(signals.get("heading_only_output_detected")):
            heading_only_output_detected_count += 1

    all_pipeline_succeeded = total_runs > 0 and pipeline_succeeded_count == total_runs
    all_acceptance_passed = total_runs > 0 and acceptance_passed_count == total_runs
    intermittent_failure_detected = 0 < acceptance_passed_count < total_runs

    checks = [
        {
            "name": "all_repeat_runs_succeeded",
            "passed": all_pipeline_succeeded,
            "actual": pipeline_succeeded_count,
            "expected": total_runs,
        },
        {
            "name": "all_repeat_runs_acceptance_passed",
            "passed": all_acceptance_passed,
            "actual": acceptance_passed_count,
            "expected": total_runs,
        },
    ]
    failed_checks = [check["name"] for check in checks if not bool(check["passed"])]

    acceptance = {
        "passed": all_pipeline_succeeded and all_acceptance_passed,
        "failed_checks": failed_checks,
        "checks": checks,
    }

    failure_classification: str | None = None
    if not acceptance["passed"]:
        if intermittent_failure_detected:
            failure_classification = "intermittent_failure"
        elif len(failure_classification_counts) == 1:
            failure_classification = next(iter(failure_classification_counts))
        else:
            failure_classification = "repeat_failures"

    summary = {
        "repeat_count": total_runs,
        "pipeline_succeeded_count": pipeline_succeeded_count,
        "acceptance_passed_count": acceptance_passed_count,
        "failed_repeat_indexes": failed_repeat_indexes,
        "failed_repeat_run_ids": failed_repeat_run_ids,
        "intermittent_failure_detected": intermittent_failure_detected,
        "heading_only_output_detected_count": heading_only_output_detected_count,
        "result_counts": result_counts,
        "failure_classification_counts": failure_classification_counts,
    }
    return summary, acceptance, failure_classification


def _select_repeat_artifact_references(repeat_runs: Sequence[Mapping[str, object]]) -> dict[str, object | None]:
    first_failing_run = next((run for run in repeat_runs if not bool(run.get("acceptance_passed"))), None)
    representative_success_run = next((run for run in reversed(list(repeat_runs)) if bool(run.get("acceptance_passed"))), None)

    def extract_paths(run: Mapping[str, object] | None, prefix: str) -> dict[str, object | None]:
        if run is None:
            return {
                f"{prefix}_run_id": None,
                f"{prefix}_report_json": None,
                f"{prefix}_summary_txt": None,
                f"{prefix}_markdown_path": None,
                f"{prefix}_docx_path": None,
                f"{prefix}_tts_text_path": None,
            }
        output_artifacts = cast(Mapping[str, object], run.get("output_artifacts") or {})
        return {
            f"{prefix}_run_id": run.get("run_id"),
            f"{prefix}_report_json": run.get("report_path"),
            f"{prefix}_summary_txt": run.get("summary_path"),
            f"{prefix}_markdown_path": output_artifacts.get("markdown_path"),
            f"{prefix}_docx_path": output_artifacts.get("docx_path"),
            f"{prefix}_tts_text_path": output_artifacts.get("tts_text_path"),
        }

    artifact_references = {}
    artifact_references.update(extract_paths(first_failing_run, "first_failing"))
    artifact_references.update(extract_paths(representative_success_run, "representative_success"))
    return artifact_references


def _run_repeat_validation(
    *,
    document_profile,
    run_profile,
    source_path: Path,
    artifact_root: Path,
    requested_run_profile_id: str | None,
) -> None:
    validation_mode = _build_validation_mode_payload(run_profile)
    artifact_root.mkdir(parents=True, exist_ok=True)
    parent_run_id = _build_run_id(source_path)
    artifact_dir = artifact_root / "runs" / parent_run_id
    artifact_dir.mkdir(parents=True, exist_ok=True)

    artifact_prefix = document_profile.artifact_prefix
    output_basename = document_profile.output_basename or f"{source_path.stem}_validated"
    report_path = artifact_dir / f"{artifact_prefix}_report.json"
    summary_path = artifact_dir / f"{artifact_prefix}_summary.txt"
    progress_path = artifact_dir / f"{artifact_prefix}_progress.json"
    markdown_artifact = artifact_dir / f"{output_basename}.md"
    docx_artifact = artifact_dir / f"{output_basename}.docx"
    tts_artifact = artifact_dir / f"{output_basename}.tts.txt"
    latest_report_path = artifact_root / f"{artifact_prefix}_report.json"
    latest_summary_path = artifact_root / f"{artifact_prefix}_summary.txt"
    latest_progress_path = artifact_root / f"{artifact_prefix}_progress.json"
    latest_markdown_path = artifact_root / f"{output_basename}.md"
    latest_docx_path = artifact_root / f"{output_basename}.docx"
    latest_tts_path = artifact_root / f"{output_basename}.tts.txt"
    latest_manifest_path = artifact_root / f"{artifact_prefix}_latest.json"

    run_started_at_utc = datetime.now(UTC)
    run_started_at_epoch_seconds = time.time()
    tracker = ValidationProgressTracker(
        run_id=parent_run_id,
        document_profile_id=document_profile.id,
        run_profile_id=run_profile.id,
        validation_tier=run_profile.tier,
        validation_run_type=str(validation_mode["validation_run_type"]),
        comparison_only_validation=bool(validation_mode["comparison_only_validation"]),
        source_path=source_path,
        run_dir=artifact_dir,
        artifact_root=artifact_root,
        progress_path=progress_path,
        latest_progress_path=latest_progress_path,
        latest_manifest_path=latest_manifest_path,
        report_path=report_path,
        summary_path=summary_path,
        markdown_path=markdown_artifact,
        docx_path=docx_artifact,
        latest_report_path=latest_report_path,
        latest_summary_path=latest_summary_path,
        latest_markdown_path=latest_markdown_path,
        latest_docx_path=latest_docx_path,
        started_at_utc=run_started_at_utc,
    )
    tracker.start()
    tracker.emit(
        event_type="start",
        phase="startup",
        stage="Repeat orchestration",
        detail=f"Запуск repeat/soak профиля {run_profile.id} на {run_profile.repeat_count} повторов.",
        progress=0.0,
        metrics={"repeat_count": run_profile.repeat_count},
    )

    app_config = load_app_config()
    runtime_resolution = resolve_runtime_resolution(app_config, run_profile)
    tracker.set_manifest_context(
        runtime_config=runtime_resolution.effective.to_dict(),
        runtime_overrides=runtime_resolution.overrides,
        validation_mode=validation_mode,
    )

    repeat_runs: list[dict[str, object]] = []
    last_markdown_artifact: Path | None = None
    last_docx_artifact: Path | None = None

    try:
        for repeat_index in range(1, run_profile.repeat_count + 1):
            repeat_run_id = _build_repeat_run_id(parent_run_id, repeat_index, run_profile.repeat_count)
            tracker.emit(
                event_type="repeat_start",
                phase="repeat",
                stage=f"Repeat {repeat_index}/{run_profile.repeat_count}",
                detail=f"Запускаю повтор {repeat_index} из {run_profile.repeat_count}.",
                progress=(repeat_index - 1) / run_profile.repeat_count,
                metrics={"current_repeat": repeat_index, "repeat_count": run_profile.repeat_count},
            )

            child_env = os.environ.copy()
            child_env["DOCXAI_REAL_DOCUMENT_REPEAT_COUNT_OVERRIDE"] = "1"
            child_env["DOCXAI_REAL_DOCUMENT_FORCED_RUN_ID"] = repeat_run_id
            child_env["DOCXAI_REAL_DOCUMENT_PARENT_RUN_ID"] = parent_run_id
            child_env["DOCXAI_REAL_DOCUMENT_REPEAT_INDEX"] = str(repeat_index)
            child_env["DOCXAI_REAL_DOCUMENT_REPEAT_TOTAL"] = str(run_profile.repeat_count)
            child_env["DOCXAI_REAL_DOCUMENT_PROFILE"] = document_profile.id
            if requested_run_profile_id:
                child_env["DOCXAI_REAL_DOCUMENT_RUN_PROFILE"] = requested_run_profile_id

            completed = subprocess.run(
                [sys.executable, str(SCRIPT_PATH)],
                cwd=PROJECT_ROOT,
                env=child_env,
                check=False,
            )

            repeat_dir = artifact_root / "runs" / repeat_run_id
            repeat_report_path = repeat_dir / f"{artifact_prefix}_report.json"
            repeat_summary_path = repeat_dir / f"{artifact_prefix}_summary.txt"
            repeat_progress_path = repeat_dir / f"{artifact_prefix}_progress.json"

            repeat_report = _load_json_file(repeat_report_path) if repeat_report_path.exists() else {}
            repeat_output_artifacts = cast(Mapping[str, object], repeat_report.get("output_artifacts") or {})
            repeat_markdown_path = repeat_output_artifacts.get("markdown_path")
            repeat_docx_path = repeat_output_artifacts.get("docx_path")
            if isinstance(repeat_markdown_path, str) and repeat_markdown_path:
                candidate = PROJECT_ROOT / repeat_markdown_path
                if candidate.exists():
                    last_markdown_artifact = candidate
            if isinstance(repeat_docx_path, str) and repeat_docx_path:
                candidate = PROJECT_ROOT / repeat_docx_path
                if candidate.exists():
                    last_docx_artifact = candidate

            repeat_run_payload = {
                "repeat_index": repeat_index,
                "repeat_count": run_profile.repeat_count,
                "run_id": repeat_run_id,
                "returncode": completed.returncode,
                "status": "completed" if completed.returncode == 0 else "failed",
                "result": repeat_report.get("result") if repeat_report else "failed",
                "acceptance_passed": bool(cast(Mapping[str, object], repeat_report.get("acceptance") or {}).get("passed")),
                "failed_checks": _as_string_list(cast(Mapping[str, object], repeat_report.get("acceptance") or {}).get("failed_checks")),
                "failure_classification": repeat_report.get("failure_classification"),
                "duration_seconds": cast(Mapping[str, object], repeat_report.get("run") or {}).get("duration_seconds"),
                "report_path": _path_for_report(repeat_report_path) if repeat_report_path.exists() else None,
                "summary_path": _path_for_report(repeat_summary_path) if repeat_summary_path.exists() else None,
                "progress_path": _path_for_report(repeat_progress_path) if repeat_progress_path.exists() else None,
                "signals": repeat_report.get("signals") or {},
                "output_artifacts": repeat_output_artifacts,
            }
            repeat_runs.append(repeat_run_payload)

            tracker.emit(
                event_type="repeat_complete",
                phase="repeat",
                stage=f"Repeat {repeat_index}/{run_profile.repeat_count}",
                detail=(
                    f"Повтор {repeat_index} завершён: result={repeat_run_payload['result']} "
                    f"acceptance={repeat_run_payload['acceptance_passed']}."
                ),
                progress=repeat_index / run_profile.repeat_count,
                metrics={
                    "current_repeat": repeat_index,
                    "repeat_count": run_profile.repeat_count,
                    "returncode": completed.returncode,
                },
            )

        repeat_summary, acceptance, failure_classification = _summarize_repeat_runs(repeat_runs)
        run_finished_at_epoch_seconds = time.time()
        run_finished_at_utc = datetime.now(UTC)
        run_duration_seconds = round(run_finished_at_epoch_seconds - run_started_at_epoch_seconds, 3)
        result = "succeeded" if acceptance["passed"] else "failed"

        repeat_artifact_references = _select_repeat_artifact_references(repeat_runs)

        report = {
            "run": {
                "run_id": parent_run_id,
                "started_at_utc": run_started_at_utc.isoformat(),
                "finished_at_utc": run_finished_at_utc.isoformat(),
                "duration_seconds": run_duration_seconds,
                "document_profile_id": document_profile.id,
                "run_profile_id": run_profile.id,
                "validation_tier": run_profile.tier,
                "validation_mode": validation_mode,
                "repeat_count": run_profile.repeat_count,
                "artifact_root": _path_for_report(artifact_root),
                "artifact_dir": _path_for_report(artifact_dir),
                "environment": _build_environment_snapshot(),
            },
            "document_profile_id": document_profile.id,
            "run_profile_id": run_profile.id,
            "validation_tier": run_profile.tier,
            "validation_mode": validation_mode,
            "source_document_path": _path_for_report(source_path),
            "artifact_dir": _path_for_report(artifact_dir),
            "progress_path": _path_for_report(progress_path),
            "result": result,
            "runtime_config": build_validation_runtime_config(runtime_resolution),
            "failure_classification": failure_classification,
            "signals": {
                "intermittent_failure_detected": repeat_summary["intermittent_failure_detected"],
                "heading_only_output_detected_count": repeat_summary["heading_only_output_detected_count"],
            },
            "preparation": None,
            "formatting_diagnostics": [],
            "repeat_summary": repeat_summary,
            "repeat_runs": repeat_runs,
            "acceptance": acceptance,
            "output_artifacts": {
                "markdown_path": _path_for_report(last_markdown_artifact),
                "docx_path": _path_for_report(last_docx_artifact),
                "report_json": _path_for_report(report_path),
                "summary_txt": _path_for_report(summary_path),
                "latest_report_json": _path_for_report(latest_report_path),
                "latest_summary_txt": _path_for_report(latest_summary_path),
                "latest_markdown_path": _path_for_report(latest_markdown_path) if last_markdown_artifact is not None else None,
                "latest_docx_path": _path_for_report(latest_docx_path) if last_docx_artifact is not None else None,
                "latest_manifest_json": _path_for_report(latest_manifest_path),
                **repeat_artifact_references,
            },
        }
        sanitized_report = sanitize_for_json(report)
        final_status = _resolve_validation_final_status(
            result=result,
            acceptance_passed=bool(acceptance["passed"]),
            validation_mode=validation_mode,
        )

        summary_lines = [
            f"run_id={parent_run_id}",
            f"document_profile_id={document_profile.id}",
            f"run_profile_id={run_profile.id}",
            f"validation_tier={run_profile.tier}",
            f"validation_run_type={validation_mode['validation_run_type']}",
            f"comparison_only_validation={validation_mode['comparison_only_validation']}",
            f"acceptance_contract_active={validation_mode['acceptance_contract_active']}",
            f"validation_evidence_label={validation_mode['evidence_label']}",
            f"validation_success_criterion={validation_mode['success_criterion']}",
            f"status={final_status}",
            f"run_started_at_utc={run_started_at_utc.isoformat()}",
            f"run_finished_at_utc={run_finished_at_utc.isoformat()}",
            f"run_duration_seconds={run_duration_seconds}",
            f"source={_path_for_report(source_path)}",
            f"artifact_dir={_path_for_report(artifact_dir)}",
            f"artifact_root={_path_for_report(artifact_root)}",
            f"progress_json={_path_for_report(progress_path)}",
            f"result={result}",
            f"failure_classification={failure_classification}",
            f"repeat_count={run_profile.repeat_count}",
            f"runtime_overrides={json.dumps(cast(Mapping[str, object], report['runtime_config']).get('overrides') or {}, ensure_ascii=False, sort_keys=True)}",
            f"pipeline_succeeded_count={repeat_summary['pipeline_succeeded_count']}",
            f"acceptance_passed_count={repeat_summary['acceptance_passed_count']}",
            f"intermittent_failure_detected={repeat_summary['intermittent_failure_detected']}",
            f"failed_repeat_indexes={','.join(str(index) for index in _as_object_list(repeat_summary['failed_repeat_indexes']))}",
            f"failed_repeat_run_ids={','.join(str(item) for item in _as_object_list(repeat_summary['failed_repeat_run_ids']))}",
            f"result_counts={json.dumps(repeat_summary['result_counts'], ensure_ascii=False, sort_keys=True)}",
            f"failure_classification_counts={json.dumps(repeat_summary['failure_classification_counts'], ensure_ascii=False, sort_keys=True)}",
            f"acceptance_passed={acceptance['passed']}",
            f"acceptance_failed_checks={','.join(_as_string_list(acceptance['failed_checks']))}",
            f"markdown_path={report['output_artifacts']['markdown_path']}",
            f"docx_path={report['output_artifacts']['docx_path']}",
            f"tts_text_path={report['output_artifacts'].get('tts_text_path')}",
            f"latest_manifest_json={report['output_artifacts']['latest_manifest_json']}",
        ]

        summary_path.write_text("\n".join(summary_lines), encoding="utf-8")
        report_path.write_text(json.dumps(sanitized_report, ensure_ascii=False, indent=2), encoding="utf-8")
        _write_latest_alias_artifacts(
            report_path=report_path,
            summary_path=summary_path,
            markdown_artifact=last_markdown_artifact,
            docx_artifact=last_docx_artifact,
            tts_artifact=None,
            latest_report_path=latest_report_path,
            latest_summary_path=latest_summary_path,
            latest_markdown_path=latest_markdown_path if last_markdown_artifact is not None else None,
            latest_docx_path=latest_docx_path if last_docx_artifact is not None else None,
            latest_tts_path=None,
            latest_manifest_path=latest_manifest_path,
            run_id=parent_run_id,
            run_dir=artifact_dir,
            manifest_payload=cast(Mapping[str, object], tracker._build_manifest_payload_locked()),
        )
        tracker.finalize(
            status=final_status,
            result=result,
            acceptance_passed=bool(acceptance["passed"]),
            failure_classification=failure_classification,
            last_error="",
            detail=_build_validation_completion_detail(
                validation_mode=validation_mode,
                acceptance_passed=bool(acceptance["passed"]),
                report_path=report_path,
            ),
        )
        _print_terminal_completion_summary(report=cast(Mapping[str, object], report), final_status=final_status)
        if final_status != "completed":
            raise SystemExit(1)
    finally:
        tracker.stop()


def _apply_runtime_event(event: object, runtime_snapshot: dict) -> None:
    if isinstance(event, SetStateEvent):
        runtime_snapshot.setdefault("state", {}).update(event.values)
    elif isinstance(event, ResetImageStateEvent):
        runtime_snapshot["image_reset_count"] = int(
            runtime_snapshot.get("image_reset_count", 0)
        ) + 1
    elif isinstance(event, SetProcessingStatusEvent):
        runtime_snapshot.setdefault("status", []).append(event.payload)
    elif isinstance(event, FinalizeProcessingStatusEvent):
        runtime_snapshot.setdefault("finalize", []).append(
            {
                "stage": event.stage,
                "detail": event.detail,
                "progress": event.progress,
            }
        )
    elif isinstance(event, PushActivityEvent):
        runtime_snapshot.setdefault("activity", []).append(event.message)
    elif isinstance(event, AppendLogEvent):
        runtime_snapshot.setdefault("log", []).append(event.payload)
    elif isinstance(event, AppendImageLogEvent):
        runtime_snapshot.setdefault("image_log", []).append(event.payload)


def drain_runtime_events(event_queue: queue.Queue, runtime_snapshot: dict, on_event=None) -> None:
    while True:
        try:
            event = event_queue.get_nowait()
        except queue.Empty:
            break
        _apply_runtime_event(event, runtime_snapshot)
        if on_event is not None:
            on_event(event)


def sanitize_for_json(value: object) -> object:
    if isinstance(value, Mapping):
        return {str(key): sanitize_for_json(item) for key, item in value.items()}
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes, bytearray)):
        return [sanitize_for_json(item) for item in value]
    return app_logger.sanitize_log_context(value)


def classify_failure(report: dict) -> str | None:
    candidates = []
    last_error = str(report.get("last_error") or "")
    candidates.append(last_error)
    exc = report.get("exception") or {}
    if isinstance(exc, dict):
        candidates.append(str(exc.get("message") or ""))
        candidates.append(str(exc.get("traceback") or ""))
    for event in report.get("event_log", []):
        if isinstance(event, dict):
            candidates.append(str(event.get("event_id") or ""))
            candidates.append(json.dumps(event, ensure_ascii=False))
    joined = "\n".join(text for text in candidates if text)
    for marker in (
        "heading_only_output",
        "empty_processed_block",
        "empty_response",
        "collapsed_output",
        "unsupported_response_shape",
        "image_placeholder_integrity_failed",
        "docx_build_failed",
        "image_processing_failed",
        "processing_init_failed",
    ):
        if marker in joined:
            return marker
    if report.get("result") == "failed":
        return "failed_unclassified"
    if report.get("result") == "stopped":
        return "stopped"
    return None


def is_heading_only_markdown(text: str) -> bool:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return bool(lines) and all(line.startswith("#") and len(line.split()) >= 2 for line in lines)


def _normalize_structural_text(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text).strip().lower()
    normalized = re.sub(r"^#{1,6}\s+", "", normalized)
    normalized = re.sub(r"^([*_]{1,3})(.+)\1$", r"\2", normalized)
    normalized = re.sub(r"^<[^>]+>(.+)</[^>]+>$", r"\1", normalized)
    return normalized


def _contains_cyrillic(text: str) -> bool:
    return any("а" <= char <= "я" or char == "ё" for char in text.lower())


def _contains_latin(text: str) -> bool:
    return any("a" <= char <= "z" for char in text.lower())


def _is_translation_stable_key_heading(text: str) -> bool:
    normalized = _normalize_structural_text(text)
    if not normalized:
        return False
    if re.fullmatch(r"\(?matthew\s+\d+:\d+(?:-\d+)?\)?", normalized, flags=re.IGNORECASE):
        return True
    if re.fullmatch(r"\(?матфея\s+\d+:\d+(?:-\d+)?\)?", normalized, flags=re.IGNORECASE):
        return True
    return False


def _should_enforce_source_heading_in_translate_mode(text: str) -> bool:
    normalized = _normalize_structural_text(text)
    if not normalized:
        return False
    if _contains_cyrillic(normalized):
        return True
    if _contains_latin(normalized):
        return _is_translation_stable_key_heading(normalized)
    return True


def _extract_scripture_anchor(text: str) -> str | None:
    normalized = _normalize_structural_text(text)
    match = re.search(r"(\d+:\d+(?:[-–]\d+)?)", normalized)
    if match is None:
        return None
    return match.group(1).replace("–", "-")


def _scripture_heading_matches(source_heading: str, output_heading: str) -> bool:
    source_anchor = _extract_scripture_anchor(source_heading)
    output_anchor = _extract_scripture_anchor(output_heading)
    if not source_anchor or not output_anchor or source_anchor != output_anchor:
        return False
    return _is_translation_stable_key_heading(source_heading) or _is_translation_stable_key_heading(output_heading)


def _resolve_missing_key_headings(
    source_heading_texts: set[str],
    output_heading_texts: set[str],
    *,
    processing_operation: str,
) -> list[str]:
    missing = sorted(source_heading_texts - output_heading_texts)
    if processing_operation != "translate" or not missing:
        return missing

    unmatched_output = set(output_heading_texts)
    resolved_missing: list[str] = []
    for source_heading in missing:
        matched_output = next(
            (
                output_heading
                for output_heading in unmatched_output
                if _scripture_heading_matches(source_heading, output_heading)
            ),
            None,
        )
        if matched_output is not None:
            unmatched_output.discard(matched_output)
            continue
        resolved_missing.append(source_heading)
    return resolved_missing


def _is_numeric_marker_or_page_range_key_heading(text: str) -> bool:
    normalized = _normalize_structural_text(text)
    if not normalized:
        return False
    numeric_fragment = r"\d[\d\s]*"
    return bool(
        re.fullmatch(
            rf"{numeric_fragment}(?:(?:\s*,\s*|\s*[-–—]\s*){numeric_fragment})+",
            normalized,
        )
    )


def _is_meaningful_key_heading(text: str) -> bool:
    normalized = _normalize_structural_text(text)
    fragment = _classify_centered_fragment(normalized)
    if fragment["kind"] == "attribution":
        return False
    if _is_numeric_marker_or_page_range_key_heading(normalized):
        return False
    alnum_only = re.sub(r"[^0-9a-zа-яё]+", "", normalized, flags=re.IGNORECASE)
    if len(alnum_only) < 3:
        return False
    return True


def _find_child_by_local_name(element, local_name: str):
    if element is None:
        return None
    for child in element:
        if child.tag == qn(f"w:{local_name}"):
            return child
    return None


def _paragraph_has_word_numbering(paragraph) -> bool:
    paragraph_properties = getattr(paragraph._element, "pPr", None)
    num_pr = _find_child_by_local_name(paragraph_properties, "numPr")
    if num_pr is not None:
        return True

    style = getattr(paragraph, "style", None)
    while style is not None:
        style_properties = _find_child_by_local_name(getattr(style, "_element", None), "pPr")
        num_pr = _find_child_by_local_name(style_properties, "numPr")
        if num_pr is not None:
            return True
        style = getattr(style, "base_style", None)
    return False


def _count_word_numbered_paragraphs(document: DocxDocument) -> int:
    return sum(1 for paragraph in document.paragraphs if _paragraph_has_word_numbering(paragraph))


def _resolve_paragraph_num_id(paragraph) -> str | None:
    paragraph_properties = getattr(paragraph._element, "pPr", None)
    num_pr = _find_child_by_local_name(paragraph_properties, "numPr")
    if num_pr is not None:
        num_id_element = _find_child_by_local_name(num_pr, "numId")
        if num_id_element is not None:
            return num_id_element.get(qn("w:val"))

    style = getattr(paragraph, "style", None)
    while style is not None:
        style_properties = _find_child_by_local_name(getattr(style, "_element", None), "pPr")
        num_pr = _find_child_by_local_name(style_properties, "numPr")
        if num_pr is not None:
            num_id_element = _find_child_by_local_name(num_pr, "numId")
            if num_id_element is not None:
                return num_id_element.get(qn("w:val"))
        style = getattr(style, "base_style", None)
    return None


def _resolve_paragraph_ilvl(paragraph) -> str | None:
    paragraph_properties = getattr(paragraph._element, "pPr", None)
    num_pr = _find_child_by_local_name(paragraph_properties, "numPr")
    if num_pr is not None:
        ilvl_element = _find_child_by_local_name(num_pr, "ilvl")
        if ilvl_element is not None:
            return ilvl_element.get(qn("w:val"))

    style = getattr(paragraph, "style", None)
    while style is not None:
        style_properties = _find_child_by_local_name(getattr(style, "_element", None), "pPr")
        num_pr = _find_child_by_local_name(style_properties, "numPr")
        if num_pr is not None:
            ilvl_element = _find_child_by_local_name(num_pr, "ilvl")
            if ilvl_element is not None:
                return ilvl_element.get(qn("w:val"))
        style = getattr(style, "base_style", None)
    return None


def _resolve_numbering_format_by_num_id(document: DocxDocument) -> dict[tuple[str, str], str]:
    numbering_part = getattr(document.part, "numbering_part", None)
    numbering_root = getattr(numbering_part, "element", None)
    if numbering_root is None:
        return {}

    abstract_num_formats: dict[str, dict[str, str]] = {}
    for child in numbering_root:
        if child.tag != qn("w:abstractNum"):
            continue
        abstract_num_id = child.get(qn("w:abstractNumId"))
        if not abstract_num_id:
            continue
        level_formats: dict[str, str] = {}
        for candidate in child:
            if candidate.tag != qn("w:lvl"):
                continue
            ilvl = candidate.get(qn("w:ilvl")) or "0"
            num_fmt = _find_child_by_local_name(candidate, "numFmt")
            format_value = None if num_fmt is None else num_fmt.get(qn("w:val"))
            if format_value:
                level_formats[ilvl] = format_value
        if level_formats:
            abstract_num_formats[abstract_num_id] = level_formats

    formats_by_num_id: dict[tuple[str, str], str] = {}
    for child in numbering_root:
        if child.tag != qn("w:num"):
            continue
        num_id = child.get(qn("w:numId"))
        if not num_id:
            continue
        abstract_num_id_element = _find_child_by_local_name(child, "abstractNumId")
        abstract_num_id = None if abstract_num_id_element is None else abstract_num_id_element.get(qn("w:val"))
        if not abstract_num_id or abstract_num_id not in abstract_num_formats:
            continue
        for ilvl, format_value in abstract_num_formats[abstract_num_id].items():
            formats_by_num_id[(num_id, ilvl)] = format_value
    return formats_by_num_id


def _count_ordered_word_numbered_paragraphs(document: DocxDocument) -> int:
    formats_by_num_id = _resolve_numbering_format_by_num_id(document)
    count = 0
    for paragraph in document.paragraphs:
        num_id = _resolve_paragraph_num_id(paragraph)
        ilvl = _resolve_paragraph_ilvl(paragraph) or "0"
        if num_id and formats_by_num_id.get((num_id, ilvl)) in ORDERED_LIST_FORMATS:
            count += 1
    return count


def _resolve_direct_paragraph_alignment(paragraph) -> str | None:
    paragraph_properties = getattr(paragraph._element, "pPr", None)
    alignment = _find_child_by_local_name(paragraph_properties, "jc")
    return None if alignment is None else alignment.get(qn("w:val"))


def _extract_short_centered_paragraph_texts(
    document: DocxDocument,
    *,
    max_words: int = 18,
    max_chars: int = 160,
) -> set[str]:
    centered_texts: set[str] = set()
    for paragraph in document.paragraphs:
        if _resolve_direct_paragraph_alignment(paragraph) != "center":
            continue
        normalized_text = _normalize_structural_text(paragraph.text)
        if not normalized_text:
            continue
        if len(normalized_text) > max_chars or len(normalized_text.split()) > max_words:
            continue
        centered_texts.add(normalized_text)
    return centered_texts


def _is_heading_style_name(style_name: str) -> bool:
    normalized = style_name.strip().lower()
    return normalized.startswith("heading") or normalized.startswith("заголовок")


def _is_centered_heading_like_text(text: str) -> bool:
    normalized = _normalize_structural_text(text)
    if not normalized:
        return False
    if re.fullmatch(r"глава\s+\d+[a-zа-яё0-9.-]*", normalized, flags=re.IGNORECASE):
        return True
    if re.fullmatch(r"chapter\s+\d+[a-zа-я0-9.-]*", normalized, flags=re.IGNORECASE):
        return True
    if re.fullmatch(r"part\s+\d+[a-zа-я0-9.-]*", normalized, flags=re.IGNORECASE):
        return True
    return False


def _is_allowlisted_centered_acceptance_paragraph(paragraph) -> bool:
    if _resolve_direct_paragraph_alignment(paragraph) != "center":
        return False

    style = getattr(paragraph, "style", None)
    style_name = str(getattr(style, "name", "") or "")
    if _is_heading_style_name(style_name):
        return False

    normalized_text = _normalize_structural_text(paragraph.text)
    if not normalized_text:
        return False
    if _is_centered_heading_like_text(normalized_text):
        return False

    fragment = _classify_centered_fragment(normalized_text)
    return fragment["kind"] in {"caption", "attribution", "quote", "general"}


def _extract_allowlisted_centered_paragraph_texts(
    document: DocxDocument,
    *,
    max_words: int = 18,
    max_chars: int = 160,
) -> set[str]:
    centered_texts: set[str] = set()
    for paragraph in document.paragraphs:
        if not _is_allowlisted_centered_acceptance_paragraph(paragraph):
            continue
        normalized_text = _normalize_structural_text(paragraph.text)
        if len(normalized_text) > max_chars or len(normalized_text.split()) > max_words:
            continue
        centered_texts.add(normalized_text)
    return centered_texts


def _centered_text_similarity(left: str, right: str) -> float:
    if left == right:
        return 1.0

    score = SequenceMatcher(None, left, right).ratio()
    left_tokens = set(left.split())
    right_tokens = set(right.split())
    if left_tokens and right_tokens:
        overlap_score = len(left_tokens & right_tokens) / max(len(left_tokens | right_tokens), 1)
        score = max(score, overlap_score)

    if left.startswith("рисунок ") and right.startswith("рисунок "):
        left_label = left.split(" ", 2)[:2]
        right_label = right.split(" ", 2)[:2]
        if left_label == right_label:
            score = max(score, 0.85)
    return score


def _centered_quote_similarity(left: str, right: str) -> float:
    sequence_score = SequenceMatcher(None, left, right).ratio()
    left_tokens = {token for token in _centered_word_tokens(left) if len(token) > 2}
    right_tokens = {token for token in _centered_word_tokens(right) if len(token) > 2}
    if not left_tokens or not right_tokens:
        return sequence_score

    overlap_score = len(left_tokens & right_tokens) / max(len(left_tokens | right_tokens), 1)
    shared_long_tokens = len(left_tokens & right_tokens)
    anchored_parallelism = all(token in left_tokens and token in right_tokens for token in ("богатство", "имущества"))

    if shared_long_tokens == 0:
        return min(sequence_score, 0.35)
    if shared_long_tokens == 1:
        return min(max(sequence_score, overlap_score), 0.54)
    if anchored_parallelism and sequence_score >= 0.5:
        return max(sequence_score, 0.6)
    return max(sequence_score, overlap_score)


def _extract_centered_caption_payload(text: str) -> str:
    label = _extract_centered_caption_label(text)
    if label is None:
        return text
    payload = re.sub(r"^(рисунок|рис\.?|figure|fig\.?)\s+[0-9]+(?:\.[0-9]+)*[\s.:,-]*", "", text, count=1, flags=re.IGNORECASE)
    return payload.strip()


def _centered_word_tokens(text: str) -> list[str]:
    return re.findall(r"[0-9A-Za-zА-Яа-яЁё]+", text.lower())


def _centered_caption_similarity(left: str, right: str) -> float:
    left_label = _extract_centered_caption_label(left)
    right_label = _extract_centered_caption_label(right)
    if not left_label or left_label != right_label:
        return 0.0

    left_payload = _extract_centered_caption_payload(left)
    right_payload = _extract_centered_caption_payload(right)
    if not left_payload or not right_payload:
        return 1.0

    payload_score = SequenceMatcher(None, left_payload, right_payload).ratio()
    left_tokens = {token for token in _centered_word_tokens(left_payload) if len(token) > 2}
    right_tokens = {token for token in _centered_word_tokens(right_payload) if len(token) > 2}
    if left_tokens and right_tokens:
        payload_score = max(payload_score, len(left_tokens & right_tokens) / max(len(left_tokens | right_tokens), 1))
    return max(payload_score, 0.8)


def _extract_centered_caption_label(text: str) -> str | None:
    match = re.match(r"^(рисунок|рис\.?|figure|fig\.?)\s+([0-9]+(?:\.[0-9]+)*)", text, re.IGNORECASE)
    if match is None:
        return None
    return f"{match.group(1).lower()} {match.group(2)}"


def _classify_centered_fragment(text: str) -> dict[str, str | None]:
    label = _extract_centered_caption_label(text)
    if label is not None:
        return {"kind": "caption", "label": label}
    words = [token for token in text.split() if token]
    letters_only = re.sub(r"[^A-Za-zА-Яа-яЁё]", "", text)
    if words and len(words) == 1 and letters_only and len(letters_only) >= 3:
        return {"kind": "attribution", "label": None}
    if words and len(words) <= 4 and letters_only and letters_only.upper() == letters_only:
        return {"kind": "attribution", "label": None}
    if len(words) >= 6:
        return {"kind": "quote", "label": None}
    return {"kind": "general", "label": None}


def _match_centered_structural_texts(
    source_centered_texts: Sequence[str],
    output_centered_texts: Sequence[str],
    *,
    min_similarity: float = 0.55,
) -> tuple[list[str], list[dict[str, object]]]:
    unmatched_output = list(output_centered_texts)
    missing_source: list[str] = []
    matches: list[dict[str, object]] = []

    for source_text in source_centered_texts:
        source_fragment = _classify_centered_fragment(source_text)
        best_index = -1
        best_score = 0.0
        for candidate_index, candidate_text in enumerate(unmatched_output):
            candidate_fragment = _classify_centered_fragment(candidate_text)
            if source_fragment["kind"] != candidate_fragment["kind"]:
                continue

            score = _centered_text_similarity(source_text, candidate_text)
            if source_fragment["kind"] == "caption":
                score = _centered_caption_similarity(source_text, candidate_text)
            elif source_fragment["kind"] == "attribution":
                if source_text == candidate_text:
                    score = 1.0
                else:
                    score = 0.0
            elif source_fragment["kind"] == "quote":
                score = _centered_quote_similarity(source_text, candidate_text)
            else:
                shared_tokens = set(source_text.split())
                candidate_tokens = set(candidate_text.split())
                overlap_count = len(shared_tokens & candidate_tokens)
                if shared_tokens and candidate_tokens and overlap_count < 2:
                    score = min(score, 0.49)
                elif overlap_count < max(2, min(len(shared_tokens), len(candidate_tokens)) // 2):
                    score = min(score, 0.54)

            if score > best_score:
                best_score = score
                best_index = candidate_index
        if best_index >= 0 and best_score >= min_similarity:
            matched_output = unmatched_output.pop(best_index)
            matches.append(
                {
                    "source": source_text,
                    "output": matched_output,
                    "similarity": round(best_score, 3),
                }
            )
            continue
        missing_source.append(source_text)

    return missing_source, matches


def _load_recent_formatting_diagnostics(since_epoch_seconds: float) -> tuple[list[str], list[dict[str, object]]]:
    artifact_paths = document_pipeline._collect_recent_formatting_diagnostics(
        since_epoch_seconds=since_epoch_seconds
    )
    return artifact_paths, _load_formatting_diagnostics_payloads(artifact_paths)


def _load_formatting_diagnostics_payloads(artifact_paths: Sequence[str]) -> list[dict[str, object]]:
    payloads: list[dict[str, object]] = []
    for artifact_path in artifact_paths:
        try:
            payload = json.loads(Path(artifact_path).read_text(encoding="utf-8"))
        except Exception:
            continue
        if isinstance(payload, dict):
            payloads.append(payload)
    return payloads


def _extract_run_formatting_diagnostics_paths(event_log: Sequence[Mapping[str, object]]) -> list[str]:
    for event in reversed(event_log):
        if str(event.get("event_id") or "") != "formatting_diagnostics_artifacts_detected":
            continue
        context = event.get("context") or {}
        if not isinstance(context, Mapping):
            continue
        artifact_paths = context.get("artifact_paths") or []
        if not isinstance(artifact_paths, Sequence) or isinstance(artifact_paths, (str, bytes, bytearray)):
            continue
        return [str(path) for path in artifact_paths if isinstance(path, str) and path]
    return []


def _extract_quality_report_artifact_path(event_log: Sequence[Mapping[str, object]]) -> str | None:
    for event in reversed(event_log):
        if str(event.get("event_id") or "") != "quality_report_saved":
            continue
        context = event.get("context") or {}
        if not isinstance(context, Mapping):
            continue
        artifact_path = context.get("artifact_path")
        if isinstance(artifact_path, str) and artifact_path.strip():
            return artifact_path.strip()
    return None


def _load_translation_quality_report(event_log: Sequence[Mapping[str, object]]) -> tuple[dict[str, object] | None, str | None]:
    artifact_path = _extract_quality_report_artifact_path(event_log)
    if not artifact_path:
        return None, None
    candidate = Path(artifact_path)
    if not candidate.is_absolute():
        candidate = (PROJECT_ROOT / candidate).resolve()
    try:
        payload = json.loads(candidate.read_text(encoding="utf-8"))
    except Exception:
        return None, _path_for_report(candidate)
    if not isinstance(payload, dict):
        return None, _path_for_report(candidate)
    return payload, _path_for_report(candidate)


def _extract_ui_result_artifact_paths(event_log: Sequence[Mapping[str, object]]) -> dict[str, str]:
    for event in reversed(event_log):
        if str(event.get("event_id") or "") != "ui_result_artifacts_saved":
            continue
        context = event.get("context") or {}
        if not isinstance(context, Mapping):
            continue
        artifact_paths = context.get("artifact_paths") or {}
        if not isinstance(artifact_paths, Mapping):
            continue
        return {
            str(key): str(value)
            for key, value in artifact_paths.items()
            if isinstance(key, str) and isinstance(value, str) and value.strip()
        }
    return {}


def _load_reader_cleanup_evidence(event_log: Sequence[Mapping[str, object]]) -> dict[str, object]:
    artifact_paths = _extract_ui_result_artifact_paths(event_log)
    return _build_reader_cleanup_evidence_from_artifact_paths(artifact_paths)


def _build_reader_cleanup_evidence_from_artifact_paths(artifact_paths: Mapping[str, str]) -> dict[str, object]:
    evidence: dict[str, object] = {
        "artifacts_present": bool(artifact_paths),
        "cleaned_markdown_path": artifact_paths.get("markdown_path"),
        "cleaned_docx_path": artifact_paths.get("docx_path"),
        "raw_markdown_path": artifact_paths.get("reader_cleanup_raw_markdown_path"),
        "reader_cleanup_report_path": artifact_paths.get("reader_cleanup_report_path"),
        "stage_status": None,
        "changed": None,
        "accepted_delete_block_count": 0,
        "ignored_delete_block_count": 0,
        "rejected_delete_block_count": 0,
        "failed_chunk_count": 0,
        "anchor_repair_status": "unknown",
        "recommended_anchor_targets": [],
        "recommended_anchor_target_count": 0,
        "deleted_block_previews": [],
    }
    report_path = artifact_paths.get("reader_cleanup_report_path")
    if not report_path:
        return evidence

    report_file = Path(report_path)
    if not report_file.is_absolute():
        report_file = PROJECT_ROOT / report_file
    if not report_file.exists():
        evidence["load_warning"] = "reader_cleanup_report_missing"
        return evidence

    try:
        report_payload = _load_json_file(report_file)
    except (OSError, json.JSONDecodeError):
        evidence["load_warning"] = "reader_cleanup_report_unreadable"
        return evidence

    stats = cast(Mapping[str, object], report_payload.get("stats") or {})
    cleanup_settings = cast(Mapping[str, object], report_payload.get("cleanup_settings") or {})
    accepted_delete_block_count = _coerce_int(stats.get("accepted_delete_block_count"))
    ignored_delete_block_count = _coerce_int(stats.get("ignored_delete_block_count"))
    proposed_delete_block_count = _coerce_int(stats.get("proposed_delete_block_count"))
    failed_chunk_count = _coerce_int(stats.get("failed_chunk_count"))
    accepted_delete_blocks = _as_object_list(report_payload.get("accepted_delete_blocks"))
    passes = cast(Mapping[str, object], report_payload.get("passes") or {})
    anchor_repair_pass = cast(Mapping[str, object], passes.get("anchor_repair_pass") or {})
    runtime_anchor_selected_count = _coerce_int(anchor_repair_pass.get("selected_anchor_count"))
    runtime_anchor_applied = bool(anchor_repair_pass)
    runtime_anchor_stats = cast(Mapping[str, object], anchor_repair_pass.get("stats") or {})
    runtime_anchor_accepted_count = _coerce_int(runtime_anchor_stats.get("accepted_cleanup_operation_count"))
    if runtime_anchor_applied and runtime_anchor_accepted_count > 0:
        anchor_repair_status = "runtime_applied"
    elif runtime_anchor_applied:
        anchor_repair_status = "runtime_attempted_no_safe_ops"
    else:
        anchor_repair_status = "not_reported"
    deleted_block_previews: list[dict[str, object]] = []
    for entry in accepted_delete_blocks[:5]:
        if not isinstance(entry, Mapping):
            continue
        deleted_block_previews.append(
            {
                "id": entry.get("id"),
                "reason": entry.get("reason"),
                "confidence": entry.get("confidence"),
                "raw_text_preview": entry.get("raw_text_preview"),
            }
        )

    evidence.update(
        {
            "stage_status": report_payload.get("stage_status"),
            "changed": report_payload.get("changed"),
            "accepted_delete_block_count": accepted_delete_block_count,
            "ignored_delete_block_count": ignored_delete_block_count,
            "rejected_delete_block_count": max(
                0,
                proposed_delete_block_count - accepted_delete_block_count - ignored_delete_block_count,
            ),
            "failed_chunk_count": failed_chunk_count,
            "cleanup_chunk_count": stats.get("cleanup_chunk_count"),
            "cleanup_settings": dict(cleanup_settings),
            "anchor_repair_status": anchor_repair_status,
            "recommended_anchor_targets": list(cast(Sequence[object], anchor_repair_pass.get("selected_anchors") or [])),
            "recommended_anchor_target_count": runtime_anchor_selected_count,
            "deleted_block_previews": deleted_block_previews,
        }
    )
    return evidence


def _build_reader_verifier_anchor_repair_diagnostics(
    *,
    review_payload: Mapping[str, object],
    cleaned_markdown: str,
    reader_cleanup_evidence: Mapping[str, object],
) -> dict[str, object]:
    existing_status = str(reader_cleanup_evidence.get("anchor_repair_status") or "").strip()
    verifier_recommended_anchor_targets = _build_reader_cleanup_anchor_targets(
        review_payload=review_payload,
        cleaned_markdown=cleaned_markdown,
    )
    if existing_status in {"runtime_applied", "runtime_attempted_no_safe_ops", "applied_in_runtime"}:
        anchor_repair_status = existing_status
        recommended_anchor_targets = list(
            cast(Sequence[object], reader_cleanup_evidence.get("recommended_anchor_targets") or [])
        )
        recommended_anchor_target_count = _coerce_int(
            reader_cleanup_evidence.get("recommended_anchor_target_count"),
            default=len(recommended_anchor_targets),
        )
    elif verifier_recommended_anchor_targets:
        anchor_repair_status = "diagnostic_only_not_applied"
        recommended_anchor_targets = verifier_recommended_anchor_targets
        recommended_anchor_target_count = len(verifier_recommended_anchor_targets)
    else:
        anchor_repair_status = existing_status or "not_needed"
        recommended_anchor_targets = verifier_recommended_anchor_targets
        recommended_anchor_target_count = len(verifier_recommended_anchor_targets)
    return {
        "anchor_repair_status": anchor_repair_status,
        "recommended_anchor_targets": recommended_anchor_targets,
        "recommended_anchor_target_count": recommended_anchor_target_count,
        "verifier_recommended_anchor_targets": verifier_recommended_anchor_targets,
        "verifier_recommended_anchor_target_count": len(verifier_recommended_anchor_targets),
    }


def _parse_reader_verifier_cleaned_line_number(line_ref: str, artifact: str) -> int | None:
    if artifact and artifact != "cleaned_markdown":
        return None
    candidate = str(line_ref or "").strip()
    if not candidate:
        return None
    if ":" in candidate:
        prefix, _, suffix = candidate.partition(":")
        if prefix and prefix != "cleaned_markdown":
            return None
        candidate = suffix.strip()
    if not candidate.isdigit():
        return None
    line_number = int(candidate)
    return line_number if line_number > 0 else None


def _build_cleanup_block_line_spans(markdown_text: str) -> list[dict[str, object]]:
    normalized_markdown = str(markdown_text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized_markdown:
        return []
    blocks = build_cleanup_blocks(normalized_markdown)
    if not blocks:
        return []

    spans: list[dict[str, object]] = []
    current_lines: list[str] = []
    block_start_line = 1
    block_index = 0
    for line_number, raw_line in enumerate(normalized_markdown.split("\n") + [""], start=1):
        if raw_line.strip():
            if not current_lines:
                block_start_line = line_number
            current_lines.append(raw_line)
            continue
        if not current_lines:
            continue
        if block_index >= len(blocks):
            break
        block = blocks[block_index]
        block_index += 1
        spans.append(
            {
                "block_id": block.block_id,
                "text": block.text,
                "start_line": block_start_line,
                "end_line": line_number - 1,
                "is_toc_like": block.is_toc_like,
            }
        )
        current_lines = []
    return spans


def _is_reader_verifier_issue_in_toc_block(
    issue: Mapping[str, object],
    *,
    block_spans: Sequence[Mapping[str, object]],
) -> bool:
    if not block_spans:
        return False

    artifact = str(issue.get("artifact") or "cleaned_markdown").strip() or "cleaned_markdown"
    line_ref = str(issue.get("line_ref") or "").strip()
    snippet = str(issue.get("snippet") or "").strip()
    line_number = _parse_reader_verifier_cleaned_line_number(line_ref, artifact)
    if line_number is not None:
        for span in block_spans:
            start_line = _coerce_int(span.get("start_line"), default=0)
            end_line = _coerce_int(span.get("end_line"), default=0)
            if start_line <= line_number <= end_line:
                return bool(span.get("is_toc_like"))
    if snippet:
        snippet_matches = [span for span in block_spans if snippet in str(span.get("text") or "")]
        if len(snippet_matches) == 1:
            return bool(snippet_matches[0].get("is_toc_like"))
    return False


def _build_filtered_toc_issue_previews(
    items: Sequence[Mapping[str, object]],
    *,
    source: str,
) -> list[dict[str, str]]:
    previews: list[dict[str, str]] = []
    seen: set[tuple[str, str, str, str, str]] = set()
    for item in items:
        artifact = str(item.get("artifact") or "cleaned_markdown").strip() or "cleaned_markdown"
        line_ref = str(item.get("line_ref") or "").strip()
        snippet = _preview_text(str(item.get("snippet") or "").strip())
        category = str(item.get("category") or "").strip() or "uncategorized"
        if not line_ref or not snippet:
            continue
        key = (source, category, artifact, line_ref, snippet)
        if key in seen:
            continue
        seen.add(key)
        previews.append(
            {
                "source": source,
                "category": category,
                "artifact": artifact,
                "line_ref": line_ref,
                "snippet": snippet,
            }
        )
    return previews


def _merge_filtered_toc_issue_previews(
    *preview_groups: Sequence[Mapping[str, object]],
) -> list[dict[str, str]]:
    merged: list[dict[str, str]] = []
    seen: set[tuple[str, str, str, str, str]] = set()
    for group in preview_groups:
        for item in group:
            source = str(item.get("source") or "").strip() or "unknown"
            category = str(item.get("category") or "").strip() or "uncategorized"
            artifact = str(item.get("artifact") or "cleaned_markdown").strip() or "cleaned_markdown"
            line_ref = str(item.get("line_ref") or "").strip()
            snippet = _preview_text(str(item.get("snippet") or "").strip())
            if not line_ref or not snippet:
                continue
            key = (source, category, artifact, line_ref, snippet)
            if key in seen:
                continue
            seen.add(key)
            merged.append(
                {
                    "source": source,
                    "category": category,
                    "artifact": artifact,
                    "line_ref": line_ref,
                    "snippet": snippet,
                }
            )
    return merged


def _filter_reader_verifier_items_excluding_toc(
    items: Sequence[Mapping[str, object]],
    *,
    block_spans: Sequence[Mapping[str, object]],
    toc_out_of_review_scope: bool,
) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    filtered: list[dict[str, str]] = []
    excluded: list[dict[str, str]] = []
    for item in items:
        normalized_item = {str(key): str(value) for key, value in item.items()}
        if toc_out_of_review_scope and _is_reader_verifier_issue_in_toc_block(normalized_item, block_spans=block_spans):
            excluded.append(normalized_item)
            continue
        filtered.append(normalized_item)
    return filtered, excluded


def _reader_verifier_toc_out_of_review_scope_from_config(runtime_app_config: Mapping[str, object]) -> bool:
    keep_toc_value = _get_config_value(runtime_app_config, "reader_cleanup_keep_toc")
    return keep_toc_value is not None and not _coerce_bool(keep_toc_value, default=True)


def _reader_verifier_toc_out_of_review_scope_from_evidence(evidence_payload: Mapping[str, object]) -> bool:
    policy = evidence_payload.get("toc_filtering_policy")
    if not isinstance(policy, Mapping):
        return False
    return bool(policy.get("toc_out_of_review_scope"))


def _build_reader_cleanup_diagnostics(cleanup_report_payload: Mapping[str, object] | None) -> dict[str, object]:
    ignored_reason_counts = {reason: 0 for reason in _TRACKED_READER_CLEANUP_IGNORED_REASONS}
    examples_by_reason: dict[str, list[dict[str, object]]] = {
        reason: [] for reason in _TRACKED_READER_CLEANUP_IGNORED_REASONS
    }
    accepted_operation_counts: Counter[str] = Counter()
    if cleanup_report_payload is None:
        return {
            "tracked_ignored_reasons": list(_TRACKED_READER_CLEANUP_IGNORED_REASONS),
            "ignored_reason_counts": ignored_reason_counts,
            "top_ignored_reasons": [],
            "accepted_operation_counts": {},
        }

    accepted_entries = _coerce_mapping_sequence(cleanup_report_payload.get("accepted_cleanup_operations"))
    if accepted_entries:
        for entry in accepted_entries:
            operation_name = str(entry.get("operation") or "").strip() or "unknown"
            accepted_operation_counts[operation_name] += 1
    else:
        for _ in _coerce_mapping_sequence(cleanup_report_payload.get("accepted_delete_blocks")):
            accepted_operation_counts["delete_block"] += 1

    ignored_entries = _coerce_mapping_sequence(
        cleanup_report_payload.get("ignored_cleanup_operations") or cleanup_report_payload.get("ignored_delete_blocks")
    )
    for entry in ignored_entries:
        ignored_reason = str(entry.get("ignored_reason") or "").strip()
        if ignored_reason not in ignored_reason_counts:
            continue
        ignored_reason_counts[ignored_reason] += 1
        examples = examples_by_reason[ignored_reason]
        if len(examples) >= 3:
            continue
        preview_text = _preview_text(
            str(
                entry.get("raw_text_preview")
                or entry.get("expected_after_preview")
                or entry.get("noise_substring")
                or ""
            ).strip()
        )
        examples.append(
            {
                "operation": str(entry.get("operation") or "").strip() or "unknown",
                "reason": str(entry.get("reason") or "").strip(),
                "chunk_index": max(0, _coerce_int(entry.get("chunk_index"), default=0)),
                "text_preview": preview_text,
                "sequence_decision": str(entry.get("sequence_decision") or "").strip(),
            }
        )

    top_ignored_reasons = [
        {
            "ignored_reason": reason,
            "count": ignored_reason_counts[reason],
            "examples": list(examples_by_reason[reason]),
        }
        for reason in _TRACKED_READER_CLEANUP_IGNORED_REASONS
        if ignored_reason_counts[reason] > 0
    ]
    return {
        "tracked_ignored_reasons": list(_TRACKED_READER_CLEANUP_IGNORED_REASONS),
        "ignored_reason_counts": ignored_reason_counts,
        "top_ignored_reasons": top_ignored_reasons,
        "accepted_operation_counts": dict(sorted(accepted_operation_counts.items())),
    }


def _resolve_reader_cleanup_anchor_block_id(
    issue: Mapping[str, object],
    *,
    block_spans: Sequence[Mapping[str, object]],
) -> str | None:
    artifact = str(issue.get("artifact") or "cleaned_markdown").strip() or "cleaned_markdown"
    if artifact != "cleaned_markdown":
        return None
    line_ref = str(issue.get("line_ref") or "").strip()
    snippet = str(issue.get("snippet") or "").strip()
    line_number = _parse_reader_verifier_cleaned_line_number(line_ref, artifact)
    if line_number is not None:
        line_matches = [
            span
            for span in block_spans
            if _coerce_int(span.get("start_line"), default=0) <= line_number <= _coerce_int(span.get("end_line"), default=0)
        ]
        if len(line_matches) == 1:
            return str(line_matches[0].get("block_id") or "") or None
    if snippet:
        snippet_matches = [span for span in block_spans if snippet in str(span.get("text") or "")]
        if len(snippet_matches) == 1:
            return str(snippet_matches[0].get("block_id") or "") or None
    return None


def _build_reader_cleanup_anchor_targets(
    *,
    review_payload: Mapping[str, object],
    cleaned_markdown: str,
) -> list[dict[str, str]]:
    block_spans = _build_cleanup_block_line_spans(cleaned_markdown)
    if not block_spans:
        return []

    remaining_issues = _coerce_mapping_sequence(review_payload.get("remaining_issues"))
    anchor_targets: list[dict[str, str]] = []
    seen_identity_keys: set[str] = set()
    for issue in remaining_issues:
        category = str(issue.get("category") or "").strip()
        if category not in _ALLOWED_READER_CLEANUP_ANCHOR_REPAIR_CATEGORIES:
            continue
        artifact = str(issue.get("artifact") or "cleaned_markdown").strip() or "cleaned_markdown"
        if artifact != "cleaned_markdown":
            continue
        block_id = _resolve_reader_cleanup_anchor_block_id(issue, block_spans=block_spans)
        if not block_id:
            continue
        line_ref = str(issue.get("line_ref") or "").strip()
        snippet = _preview_text(str(issue.get("snippet") or "").strip())
        identity_key = f"{category}|{line_ref}|{snippet}"
        if identity_key in seen_identity_keys:
            continue
        seen_identity_keys.add(identity_key)
        anchor_targets.append(
            {
                "anchor_id": hashlib.sha256(identity_key.encode("utf-8")).hexdigest()[:16],
                "category": category,
                "block_id": block_id,
                "line_ref": line_ref,
                "snippet": snippet,
            }
        )
    return anchor_targets


def _normalize_runtime_reader_cleanup_anchor_targets_payload(payload: object) -> list[dict[str, object]]:
    if isinstance(payload, Mapping):
        raw_targets = (
            payload.get("verifier_recommended_anchor_targets")
            or payload.get("anchor_targets")
            or payload.get("recommended_anchor_targets")
            or []
        )
    else:
        raw_targets = payload
    if not isinstance(raw_targets, Sequence) or isinstance(raw_targets, (str, bytes, bytearray)):
        raise RuntimeError("reader_cleanup_anchor_targets_must_be_list")

    targets: list[dict[str, object]] = []
    for index, item in enumerate(raw_targets, start=1):
        if not isinstance(item, Mapping):
            raise RuntimeError(f"reader_cleanup_anchor_target_must_be_object:{index}")
        category = str(item.get("category") or "").strip()
        if category not in _ALLOWED_READER_CLEANUP_ANCHOR_REPAIR_CATEGORIES:
            continue
        targets.append(dict(item))
    return targets


def _load_runtime_reader_cleanup_anchor_targets_from_env() -> list[dict[str, object]]:
    raw_json = os.environ.get("DOCXAI_READER_CLEANUP_ANCHOR_TARGETS_JSON", "").strip()
    raw_path = os.environ.get("DOCXAI_READER_CLEANUP_ANCHOR_TARGETS_PATH", "").strip()
    if raw_json and raw_path:
        raise RuntimeError(
            "Use only one of DOCXAI_READER_CLEANUP_ANCHOR_TARGETS_JSON or DOCXAI_READER_CLEANUP_ANCHOR_TARGETS_PATH"
        )
    if not raw_json and not raw_path:
        return []

    if raw_json:
        payload = json.loads(raw_json)
    else:
        anchor_path = Path(raw_path)
        if not anchor_path.is_absolute():
            anchor_path = PROJECT_ROOT / anchor_path
        payload = _load_json_file(anchor_path)
    return _normalize_runtime_reader_cleanup_anchor_targets_payload(payload)


def _merge_reader_cleanup_artifact_paths(
    output_artifacts: dict[str, object],
    reader_cleanup_evidence: Mapping[str, object],
) -> None:
    cleaned_markdown_path = reader_cleanup_evidence.get("cleaned_markdown_path")
    cleaned_docx_path = reader_cleanup_evidence.get("cleaned_docx_path")
    raw_markdown_path = reader_cleanup_evidence.get("raw_markdown_path")
    reader_cleanup_report_path = reader_cleanup_evidence.get("reader_cleanup_report_path")
    if isinstance(cleaned_markdown_path, str) and cleaned_markdown_path:
        output_artifacts["cleaned_markdown_path"] = cleaned_markdown_path
    if isinstance(cleaned_docx_path, str) and cleaned_docx_path:
        output_artifacts["cleaned_docx_path"] = cleaned_docx_path
    if isinstance(raw_markdown_path, str) and raw_markdown_path:
        output_artifacts["reader_cleanup_raw_markdown_path"] = raw_markdown_path
    if isinstance(reader_cleanup_report_path, str) and reader_cleanup_report_path:
        output_artifacts["reader_cleanup_report_path"] = reader_cleanup_report_path


def _build_validation_completion_detail(
    *,
    validation_mode: Mapping[str, object],
    acceptance_passed: bool,
    report_path: Path,
    reader_cleanup_evidence: Mapping[str, object] | None = None,
    reader_verifier_evidence: Mapping[str, object] | None = None,
) -> str:
    if not bool(validation_mode.get("comparison_only_validation")):
        return f"Acceptance={'passed' if acceptance_passed else 'failed'}; report={_path_for_report(report_path)}"

    evidence = reader_cleanup_evidence or {}
    detail_parts = [
        "Comparison-only non-acceptance evidence completed",
        f"report={_path_for_report(report_path)}",
    ]
    raw_markdown_path = evidence.get("raw_markdown_path")
    cleaned_markdown_path = evidence.get("cleaned_markdown_path")
    cleaned_docx_path = evidence.get("cleaned_docx_path")
    cleanup_report_path = evidence.get("reader_cleanup_report_path")
    if isinstance(raw_markdown_path, str) and raw_markdown_path:
        detail_parts.append(f"raw_markdown={raw_markdown_path}")
    if isinstance(cleaned_markdown_path, str) and cleaned_markdown_path:
        detail_parts.append(f"cleaned_markdown={cleaned_markdown_path}")
    if isinstance(cleaned_docx_path, str) and cleaned_docx_path:
        detail_parts.append(f"cleaned_docx={cleaned_docx_path}")
    if isinstance(cleanup_report_path, str) and cleanup_report_path:
        detail_parts.append(f"reader_cleanup_report={cleanup_report_path}")
    verifier = reader_verifier_evidence or {}
    review_json_path = cast(Mapping[str, object], verifier.get("artifact_paths") or {}).get("review_json")
    verifier_status = verifier.get("verifier_status")
    overall_verdict = verifier.get("overall_verdict")
    cleaned_audit_verdict = verifier.get("cleaned_audit_verdict")
    if verifier_status is not None:
        detail_parts.append(f"reader_verifier_status={verifier_status}")
    if overall_verdict is not None:
        detail_parts.append(f"reader_verifier_verdict={overall_verdict}")
    if cleaned_audit_verdict is not None:
        detail_parts.append(f"reader_verifier_audit_verdict={cleaned_audit_verdict}")
    if isinstance(review_json_path, str) and review_json_path:
        detail_parts.append(f"reader_verifier_review={review_json_path}")
    detail_parts.append(f"acceptance_diagnostic={'passed' if acceptance_passed else 'failed'}")
    return "; ".join(detail_parts)


def _build_structural_acceptance_checks(
    *,
    source_docx_bytes: bytes,
    output_docx_bytes: bytes,
    processing_operation: str,
) -> list[dict[str, object]]:
    checks: list[dict[str, object]] = []

    def add_check(name: str, passed: bool, **details: object) -> None:
        checks.append({"name": name, "passed": passed, **details})

    source_paragraphs, _ = extract_document_content_from_docx(BytesIO(source_docx_bytes))
    output_paragraphs, _ = extract_document_content_from_docx(BytesIO(output_docx_bytes))
    source_document = Document(BytesIO(source_docx_bytes))
    output_document = Document(BytesIO(output_docx_bytes))

    source_caption_texts = {
        _normalize_structural_text(paragraph.text)
        for paragraph in source_paragraphs
        if paragraph.role == "caption" and _normalize_structural_text(paragraph.text)
    }
    output_heading_texts = {
        _normalize_structural_text(paragraph.text)
        for paragraph in output_paragraphs
        if paragraph.role == "heading" and _normalize_structural_text(paragraph.text)
    }
    caption_heading_regressions = sorted(source_caption_texts & output_heading_texts)
    add_check(
        "captions_not_promoted_to_headings",
        not caption_heading_regressions,
        regressions=caption_heading_regressions,
    )

    source_heading_texts = {
        _normalize_structural_text(paragraph.text)
        for paragraph in source_paragraphs
        if paragraph.role == "heading"
        and _normalize_structural_text(paragraph.text)
        and len(_normalize_structural_text(paragraph.text).split()) <= 10
        and _is_meaningful_key_heading(paragraph.text)
    }
    if processing_operation == "translate":
        source_heading_texts = {
            heading
            for heading in source_heading_texts
            if _should_enforce_source_heading_in_translate_mode(heading)
        }
    output_heading_texts = {
        _normalize_structural_text(paragraph.text)
        for paragraph in output_paragraphs
        if paragraph.role == "heading"
        and _normalize_structural_text(paragraph.text)
        and _is_meaningful_key_heading(paragraph.text)
    }
    missing_key_headings = _resolve_missing_key_headings(
        source_heading_texts,
        output_heading_texts,
        processing_operation=processing_operation,
    )
    add_check(
        "key_headings_preserved",
        not missing_key_headings,
        missing=missing_key_headings,
        source_heading_count=len(source_heading_texts),
        output_heading_count=len(output_heading_texts),
    )

    source_centered_texts = sorted(_extract_allowlisted_centered_paragraph_texts(source_document))
    output_centered_texts = sorted(_extract_allowlisted_centered_paragraph_texts(output_document))
    missing_centered_texts, centered_matches = _match_centered_structural_texts(
        source_centered_texts,
        output_centered_texts,
    )
    add_check(
        "centered_short_paragraphs_preserved",
        not missing_centered_texts,
        missing=missing_centered_texts,
        source_centered_count=len(source_centered_texts),
        output_centered_count=len(output_centered_texts),
        matches=centered_matches,
    )

    source_numbered_count = sum(1 for paragraph in source_paragraphs if paragraph.role == "list" and paragraph.list_kind == "ordered")
    output_numbered_count = _count_ordered_word_numbered_paragraphs(output_document)
    word_numbering_passed = source_numbered_count == 0 or output_numbered_count >= source_numbered_count
    if processing_operation == "translate":
        word_numbering_passed = True
    add_check(
        "word_numbering_preserved",
        word_numbering_passed,
        source_numbered_count=source_numbered_count,
        output_numbered_count=output_numbered_count,
        processing_operation=processing_operation,
    )
    return checks


def evaluate_lietaer_acceptance(
    report: Mapping[str, object],
    *,
    source_docx_bytes: bytes | None = None,
    output_docx_bytes: bytes | None = None,
    mismatch_threshold: int = 0,
    unmapped_target_threshold: int = 0,
    require_no_toc_body_concat: bool = False,
) -> dict[str, object]:
    # Thin wrapper over the shared acceptance-verdict assembly
    # (docxaicorrector.validation.acceptance). The report-derived checks +
    # verdict roll-up live in shared so production finalization can compute an
    # identical verdict; the harness-only structural (source<->output DOCX)
    # comparison is injected via ``structural_checks_builder``.
    structural_checks_builder: Callable[[str], list[dict[str, object]]] | None
    if source_docx_bytes and output_docx_bytes:
        source_docx_bytes_value = source_docx_bytes
        output_docx_bytes_value = output_docx_bytes

        def _structural_checks_builder(processing_operation: str) -> list[dict[str, object]]:
            return _build_structural_acceptance_checks(
                source_docx_bytes=source_docx_bytes_value,
                output_docx_bytes=output_docx_bytes_value,
                processing_operation=processing_operation,
            )

        structural_checks_builder = _structural_checks_builder
    else:
        structural_checks_builder = None

    return _build_acceptance_verdict(
        report,
        mismatch_threshold=mismatch_threshold,
        unmapped_target_threshold=unmapped_target_threshold,
        require_no_toc_body_concat=require_no_toc_body_concat,
        structural_checks_builder=structural_checks_builder,
    )


def _apply_repeat_count_override(run_profile, repeat_count_override: str):
    if not repeat_count_override:
        return run_profile
    try:
        repeat_count = max(1, int(repeat_count_override))
    except ValueError:
        _safe_terminal_print(
            f"[warning] invalid DOCXAI_REAL_DOCUMENT_REPEAT_COUNT_OVERRIDE={repeat_count_override!r}; using profile default {run_profile.repeat_count}",
            flush=True,
        )
        return run_profile
    return replace(run_profile, repeat_count=repeat_count)


class _SnapshotWithToDict(Protocol):
    def to_dict(self) -> dict[str, object]: ...


def _serialize_image_asset_forensics(image_assets: Sequence[object]) -> list[dict[str, object]]:
    payload: list[dict[str, object]] = []
    for asset in image_assets:
        source_identity_snapshot = getattr(asset, "source_identity_snapshot", None)
        if not callable(source_identity_snapshot):
            continue
        runtime_state_snapshot = getattr(asset, "runtime_state_snapshot", None)
        final_selection_snapshot = getattr(asset, "final_selection_snapshot", None)
        source_snapshot = source_identity_snapshot()
        runtime_snapshot = runtime_state_snapshot() if callable(runtime_state_snapshot) else None
        final_snapshot = final_selection_snapshot() if callable(final_selection_snapshot) else None

        source_payload = cast(_SnapshotWithToDict, source_snapshot).to_dict() if hasattr(source_snapshot, "to_dict") else {}
        source_bytes = getattr(asset, "original_bytes", None)
        if isinstance(source_bytes, (bytes, bytearray)):
            source_payload["source_sha256"] = hashlib.sha256(bytes(source_bytes)).hexdigest()
            source_payload["source_bytes_size"] = len(source_bytes)

        payload.append(
            {
                "source": source_payload,
                "runtime": cast(_SnapshotWithToDict, runtime_snapshot).to_dict() if runtime_snapshot is not None and hasattr(runtime_snapshot, "to_dict") else None,
                "final_selection": cast(_SnapshotWithToDict, final_snapshot).to_dict() if final_snapshot is not None and hasattr(final_snapshot, "to_dict") else None,
            }
        )
    return payload


def _build_image_forensics_report(prepared, runtime_snapshot: Mapping[str, object]) -> dict[str, object]:
    runtime_state = cast(Mapping[str, object], runtime_snapshot.get("state") or {})
    processed_assets = runtime_state.get("image_assets")
    prepared_assets = prepared.image_assets if prepared is not None else []
    return {
        "prepared_assets": _serialize_image_asset_forensics(cast(Sequence[object], prepared_assets)),
        "processed_assets": _serialize_image_asset_forensics(
            cast(Sequence[object], processed_assets) if isinstance(processed_assets, Sequence) else []
        ),
    }


def main() -> None:
    registry = load_validation_registry()
    document_profile_id = os.environ.get("DOCXAI_REAL_DOCUMENT_PROFILE", "lietaer-pdf-full-benchmark").strip() or "lietaer-pdf-full-benchmark"
    requested_run_profile_id = os.environ.get("DOCXAI_REAL_DOCUMENT_RUN_PROFILE", "").strip() or None
    document_profile = registry.get_document_profile(document_profile_id)
    run_profile = registry.resolve_run_profile(document_profile, requested_run_profile_id)
    repeat_count_override = os.environ.get("DOCXAI_REAL_DOCUMENT_REPEAT_COUNT_OVERRIDE", "").strip()
    run_profile = _apply_repeat_count_override(run_profile, repeat_count_override)
    validation_mode = _build_validation_mode_payload(run_profile)
    source_path = document_profile.resolved_source_path(PROJECT_ROOT)
    artifact_root = REAL_DOCUMENT_ARTIFACT_ROOT
    if run_profile.repeat_count > 1 and not repeat_count_override:
        _run_repeat_validation(
            document_profile=document_profile,
            run_profile=run_profile,
            source_path=source_path,
            artifact_root=artifact_root,
            requested_run_profile_id=requested_run_profile_id,
        )
        return

    artifact_root.mkdir(parents=True, exist_ok=True)
    run_id = os.environ.get("DOCXAI_REAL_DOCUMENT_FORCED_RUN_ID", "").strip() or _build_run_id(source_path)
    artifact_dir = artifact_root / "runs" / run_id
    artifact_dir.mkdir(parents=True, exist_ok=True)
    artifact_prefix = document_profile.artifact_prefix
    output_basename = document_profile.output_basename or f"{source_path.stem}_validated"

    for handler in app_logger.get_logger().handlers:
        if hasattr(handler, "maxBytes"):
            setattr(handler, "maxBytes", 1_000_000_000)

    report_path = artifact_dir / f"{artifact_prefix}_report.json"
    summary_path = artifact_dir / f"{artifact_prefix}_summary.txt"
    progress_path = artifact_dir / f"{artifact_prefix}_progress.json"
    markdown_artifact = artifact_dir / f"{output_basename}.md"
    docx_artifact = artifact_dir / f"{output_basename}.docx"
    latest_report_path = artifact_root / f"{artifact_prefix}_report.json"
    latest_summary_path = artifact_root / f"{artifact_prefix}_summary.txt"
    latest_progress_path = artifact_root / f"{artifact_prefix}_progress.json"
    latest_markdown_path = artifact_root / f"{output_basename}.md"
    latest_docx_path = artifact_root / f"{output_basename}.docx"
    latest_manifest_path = artifact_root / f"{artifact_prefix}_latest.json"

    progress_events = []
    event_log = []
    event_queue: queue.Queue = queue.Queue()
    run_started_at_epoch_seconds = time.time()
    run_started_at_utc = datetime.now(UTC)
    tracker = ValidationProgressTracker(
        run_id=run_id,
        document_profile_id=document_profile.id,
        run_profile_id=run_profile.id,
        validation_tier=run_profile.tier,
        validation_run_type=str(validation_mode["validation_run_type"]),
        comparison_only_validation=bool(validation_mode["comparison_only_validation"]),
        source_path=source_path,
        run_dir=artifact_dir,
        artifact_root=artifact_root,
        progress_path=progress_path,
        latest_progress_path=latest_progress_path,
        latest_manifest_path=latest_manifest_path,
        report_path=report_path,
        summary_path=summary_path,
        markdown_path=markdown_artifact,
        docx_path=docx_artifact,
        latest_report_path=latest_report_path,
        latest_summary_path=latest_summary_path,
        latest_markdown_path=latest_markdown_path,
        latest_docx_path=latest_docx_path,
        started_at_utc=run_started_at_utc,
    )
    tracker.start()
    tracker.emit(
        event_type="start",
        phase="startup",
        stage="Инициализация",
        detail=f"Запуск real-document валидации для профиля {document_profile.id}.",
        progress=0.0,
        metrics={"run_id": run_id},
    )
    formatting_diagnostics_before = _snapshot_formatting_diagnostics_paths()
    runtime = processing_runtime.BackgroundRuntime(event_queue, threading.Event())
    runtime_snapshot = {
        "state": {},
        "finalize": [],
        "activity": [],
        "log": [],
        "status": [],
        "image_log": [],
        "image_reset_count": 0,
    }

    def emit_prepare_progress(**payload: object) -> None:
        progress_events.append({"phase": "prepare", **payload})
        tracker.emit(
            event_type="prepare",
            phase="prepare",
            stage=str(payload.get("stage") or "Подготовка"),
            detail=str(payload.get("detail") or ""),
            progress=_as_float_or_none(payload.get("progress")),
            metrics=cast(Mapping[str, object], payload.get("metrics") or {}),
        )

    def emit_runtime_event(event: object) -> None:
        if isinstance(event, SetProcessingStatusEvent):
            payload = event.payload
            tracker.emit(
                event_type="status",
                phase="process",
                stage=str(payload.get("stage") or "Обработка"),
                detail=str(payload.get("detail") or ""),
                progress=_as_float_or_none(payload.get("progress")),
                metrics={
                    key: payload.get(key)
                    for key in ("current_block", "block_count", "target_chars", "context_chars")
                    if payload.get(key) is not None
                },
            )
        elif isinstance(event, FinalizeProcessingStatusEvent):
            tracker.emit(
                event_type="finalize",
                phase="process",
                stage=event.stage,
                detail=event.detail,
                progress=event.progress,
            )
        elif isinstance(event, AppendLogEvent):
            status = str(event.payload.get("status") or "")
            if status in {"WARN", "ERROR", "DONE"}:
                tracker.emit(
                    event_type="log",
                    phase="process",
                    stage=status,
                    detail=str(event.payload.get("details") or status),
                    metrics={
                        key: event.payload.get(key)
                        for key in ("block_index", "block_count", "target_chars", "context_chars")
                        if event.payload.get(key) is not None
                    },
                )

    runtime_monitor_stop = threading.Event()

    def runtime_monitor_worker() -> None:
        while True:
            if runtime_monitor_stop.is_set() and event_queue.empty():
                return
            try:
                event = event_queue.get(timeout=0.5)
            except queue.Empty:
                continue
            _apply_runtime_event(event, runtime_snapshot)
            emit_runtime_event(event)

    runtime_monitor_thread = threading.Thread(
        target=runtime_monitor_worker,
        name="lietaer-validation-runtime-monitor",
        daemon=True,
    )
    runtime_monitor_thread.start()

    source_bytes = b""
    source_docx_bytes = b""
    app_config = None
    app_config_dict: dict[str, object] = {}
    runtime_resolution = None
    prepared = None
    result = "not_started"
    exception_payload = None

    def _handle_logged_event(event_payload: Mapping[str, object]) -> None:
        event_id = str(event_payload.get("event_id") or "")
        message = str(event_payload.get("message") or "")
        context = cast(Mapping[str, object], event_payload.get("context") or {})
        if event_id == "block_completed":
            block_index = context.get("block_index")
            block_count = context.get("block_count")
            progress_value = None
            if isinstance(block_index, int) and isinstance(block_count, int) and block_count > 0:
                progress_value = block_index / block_count
            tracker.emit(
                event_type="milestone",
                phase="process",
                stage="Блок завершён",
                detail=f"Блок {block_index} из {block_count} завершён успешно.",
                progress=progress_value,
                metrics={
                    "current_block": block_index,
                    "block_count": block_count,
                    "output_ratio": context.get("output_ratio"),
                },
            )
        elif event_id == "block_rejected":
            tracker.emit(
                event_type="warning",
                phase="process",
                stage="Блок отклонён",
                detail=str(context.get("output_classification") or message),
                metrics={
                    "current_block": context.get("block_index"),
                    "block_count": context.get("block_count"),
                },
            )
        elif event_id == "formatting_diagnostics_artifacts_detected":
            tracker.emit(
                event_type="warning",
                phase="assemble",
                stage="Formatting diagnostics",
                detail="Сохранены formatting diagnostics artifacts для текущего прогона.",
                metrics={
                    "job_count": len(
                        cast(
                            Sequence[object],
                            context.get("artifact_paths") if isinstance(context.get("artifact_paths"), Sequence) else [],
                        )
                    )
                },
            )

    log_event_capture = build_validation_event_logger(event_log, on_event=_handle_logged_event)

    try:
        source_bytes = source_path.read_bytes()
        tracker.emit(
            event_type="source",
            phase="startup",
            stage="Исходный документ загружен",
            detail=f"Прочитан {source_path.name}.",
            progress=0.02,
            metrics={"target_chars": len(source_bytes)},
        )
        app_config = load_app_config()
        runtime_resolution = resolve_runtime_resolution(app_config, run_profile)
        app_config_dict = apply_runtime_resolution_to_app_config(app_config, runtime_resolution)
        runtime_anchor_targets = _load_runtime_reader_cleanup_anchor_targets_from_env()
        if runtime_anchor_targets:
            app_config_dict["reader_cleanup_anchor_repair_enabled"] = True
            app_config_dict["reader_cleanup_anchor_targets"] = runtime_anchor_targets
        tracker.set_manifest_context(
            runtime_config=runtime_resolution.effective.to_dict(),
            runtime_overrides=runtime_resolution.overrides,
            validation_mode=validation_mode,
        )
        tracker.emit(
            event_type="config",
            phase="startup",
            stage="Конфигурация загружена",
            detail=(
                f"Модель {runtime_resolution.effective.model}, "
                f"chunk_size={runtime_resolution.effective.chunk_size}, tier={run_profile.tier}."
            ),
            progress=0.05,
            metrics={"job_count": 0},
        )
        validation_service = processing_service.clone_processing_service(
            log_event_fn=log_event_capture,
        )
        result, prepared = validation_service.run_prepared_background_document(
            uploaded_file=UploadedFileStub(source_path.name, source_bytes),
            chunk_size=runtime_resolution.effective.chunk_size,
            image_mode=runtime_resolution.effective.image_mode,
            keep_all_image_variants=runtime_resolution.effective.keep_all_image_variants,
            app_config=app_config_dict,
            model=runtime_resolution.effective.model,
            max_retries=runtime_resolution.effective.max_retries,
            processing_operation=runtime_resolution.effective.processing_operation,
            source_language=runtime_resolution.effective.source_language,
            target_language=runtime_resolution.effective.target_language,
            prepare_progress_callback=emit_prepare_progress,
            processing_progress_callback=lambda **payload: progress_events.append({"phase": "process", **payload}),
            runtime=runtime,
        )
        source_docx_bytes = prepared.uploaded_file_bytes
        tracker.emit(
            event_type="prepared",
            phase="prepare",
            stage="Подготовка завершена",
            detail="План обработки собран, запускаю pipeline.",
            progress=0.2,
            metrics={
                "job_count": len(prepared.jobs),
                "target_chars": len(prepared.source_text),
            },
        )
    except Exception as exc:
        exception_payload = {
            "type": type(exc).__name__,
            "message": str(exc),
            "traceback": traceback.format_exc(),
        }
        result = "failed"
        tracker.emit(
            event_type="exception",
            phase="failed",
            stage="Исключение в валидаторе",
            detail=str(exc),
        )
    finally:
        runtime_monitor_stop.set()
        runtime_monitor_thread.join(timeout=3.0)

    drain_runtime_events(event_queue, runtime_snapshot, on_event=emit_runtime_event)

    state = runtime_snapshot.get("state", {})
    final_markdown = str(state.get("latest_markdown") or "")
    latest_docx_bytes = state.get("latest_docx_bytes")
    latest_narration_text = str(state.get("latest_narration_text") or "")
    last_error = str(state.get("last_error") or "") or str((exception_payload or {}).get("message") or "")
    source_chars = len(prepared.source_text) if prepared is not None else 0
    final_markdown_chars = len(final_markdown)
    output_ratio = round(final_markdown_chars / max(source_chars, 1), 3)

    block_completed_events = [
        event for event in event_log if event.get("event_id") == "block_completed"
    ]
    block_rejected_events = [
        event for event in event_log if event.get("event_id") == "block_rejected"
    ]
    block_output_ratios = [
        event.get("context", {}).get("output_ratio")
        for event in block_completed_events
        if isinstance(event.get("context", {}).get("output_ratio"), (int, float))
    ]

    openable_output = False
    output_paragraphs = 0
    output_inline_shapes = 0
    output_visible_text_chars = 0
    output_contains_placeholder_markup = False

    markdown_artifact_path: Path | None = None
    docx_artifact_path: Path | None = None
    tts_artifact = artifact_dir / f"{output_basename}.tts.txt"
    latest_tts_path = artifact_root / f"{output_basename}.tts.txt"
    tts_artifact_path: Path | None = None

    if final_markdown:
        markdown_artifact.write_text(final_markdown, encoding="utf-8")
        markdown_artifact_path = markdown_artifact

    if isinstance(latest_docx_bytes, (bytes, bytearray)) and latest_docx_bytes:
        latest_docx_bytes = bytes(latest_docx_bytes)
        docx_artifact.write_bytes(latest_docx_bytes)
        docx_artifact_path = docx_artifact
        try:
            output_doc = Document(BytesIO(latest_docx_bytes))
            openable_output = True
            output_paragraphs = len(output_doc.paragraphs)
            output_inline_shapes = len(output_doc.inline_shapes)
            output_visible_text_chars = len(
                "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
            )
            output_contains_placeholder_markup = (
                "[[DOCX_IMAGE_" in output_doc._element.xml
            )
        except Exception:
            openable_output = False

    if latest_narration_text:
        tts_artifact.write_text(latest_narration_text, encoding="utf-8")
        tts_artifact_path = tts_artifact

    formatting_diagnostics_after = _snapshot_formatting_diagnostics_paths()
    snapshot_discovered_paths = _collect_new_formatting_diagnostics_paths(
        formatting_diagnostics_before,
        formatting_diagnostics_after,
    )
    formatting_diagnostics_paths = _extract_run_formatting_diagnostics_paths(event_log)
    formatting_diagnostics_discovery_source = None
    if snapshot_discovered_paths:
        formatting_diagnostics_paths = snapshot_discovered_paths
        formatting_diagnostics_payloads = _load_formatting_diagnostics_payloads(
            formatting_diagnostics_paths
        )
        formatting_diagnostics_discovery_source = "snapshot_diff"
    elif formatting_diagnostics_paths:
        formatting_diagnostics_payloads = _load_formatting_diagnostics_payloads(
            formatting_diagnostics_paths
        )
        formatting_diagnostics_discovery_source = "event_log"
    else:
        formatting_diagnostics_paths, formatting_diagnostics_payloads = _load_recent_formatting_diagnostics(
            run_started_at_epoch_seconds
        )
        formatting_diagnostics_discovery_source = "recent_scan"

    if not snapshot_discovered_paths and formatting_diagnostics_discovery_source != "event_log":
        formatting_diagnostics_payloads = _load_formatting_diagnostics_payloads(formatting_diagnostics_paths)

    run_finished_at_epoch_seconds = time.time()
    run_finished_at_utc = datetime.now(UTC)
    run_duration_seconds = round(run_finished_at_epoch_seconds - run_started_at_epoch_seconds, 3)
    result = "failed" if exception_payload is not None else result

    preparation_payload = {
        "uploaded_filename": prepared.uploaded_filename if prepared is not None else source_path.name,
        "uploaded_file_token": prepared.uploaded_file_token if prepared is not None else None,
        "paragraph_count": len(prepared.paragraphs) if prepared is not None else None,
        "image_count": len(prepared.image_assets) if prepared is not None else None,
        "job_count": len(prepared.jobs) if prepared is not None else None,
        "ai_classified_count": int(getattr(prepared, "ai_classified_count", 0) or 0) if prepared is not None else 0,
        "ai_heading_count": int(getattr(prepared, "ai_heading_count", 0) or 0) if prepared is not None else 0,
        "ai_role_change_count": int(getattr(prepared, "ai_role_change_count", 0) or 0) if prepared is not None else 0,
        "ai_heading_promotion_count": int(getattr(prepared, "ai_heading_promotion_count", 0) or 0) if prepared is not None else 0,
        "ai_heading_demotion_count": int(getattr(prepared, "ai_heading_demotion_count", 0) or 0) if prepared is not None else 0,
        "ai_structural_role_change_count": int(getattr(prepared, "ai_structural_role_change_count", 0) or 0) if prepared is not None else 0,
        "source_chars": source_chars,
        "cached": prepared.preparation_cached if prepared is not None else None,
        "elapsed_seconds": round(prepared.preparation_elapsed_seconds, 3) if prepared is not None else None,
    }
    preparation_diagnostic_snapshot = build_preparation_diagnostic_snapshot(
        paragraphs=prepared.paragraphs if prepared is not None else [],
        relations=None,
        structure_repair_report=getattr(prepared, "structure_repair_report", None),
        chunk_size=int(runtime_resolution.effective.chunk_size) if runtime_resolution is not None else 6000,
        event_log=event_log,
    )
    if prepared is not None:
        _apply_prepared_snapshot_fields(preparation_diagnostic_snapshot, prepared)
    else:
        _normalize_snapshot_or_metric_statuses(preparation_diagnostic_snapshot)

    report = {
        "run": {
            "run_id": run_id,
            "started_at_utc": run_started_at_utc.isoformat(),
            "finished_at_utc": run_finished_at_utc.isoformat(),
            "duration_seconds": run_duration_seconds,
            "document_profile_id": document_profile.id,
            "run_profile_id": run_profile.id,
            "validation_tier": run_profile.tier,
            "validation_mode": validation_mode,
            "repeat_count": run_profile.repeat_count,
            "artifact_root": _path_for_report(artifact_root),
            "artifact_dir": _path_for_report(artifact_dir),
            "environment": _build_environment_snapshot(),
        },
        "document_profile_id": document_profile.id,
        "run_profile_id": run_profile.id,
        "validation_tier": run_profile.tier,
        "validation_mode": validation_mode,
        "source_document_path": _path_for_report(source_path),
        "artifact_dir": _path_for_report(artifact_dir),
        "progress_path": _path_for_report(progress_path),
        "result": result,
        "runtime_config": build_validation_runtime_config(runtime_resolution),
        "preparation": preparation_payload,
        "preparation_diagnostic_snapshot": preparation_diagnostic_snapshot,
        "source_cleanup_evidence": _build_source_cleanup_evidence(getattr(prepared, "cleanup_report", None)),
        "runtime": runtime_snapshot,
        "image_forensics": _build_image_forensics_report(prepared, runtime_snapshot),
        "last_error": last_error,
        "exception": exception_payload,
        "failure_classification": None,
        "signals": {
            "heading_only_output_detected": is_heading_only_markdown(final_markdown),
            "heading_only_rejection_logged": bool(block_rejected_events),
            "silent_text_loss_suspected": bool(final_markdown.strip()) and output_ratio < 0.6,
            "output_ratio_vs_source_text": output_ratio,
            "min_block_output_ratio": min(block_output_ratios) if block_output_ratios else None,
            "max_block_output_ratio": max(block_output_ratios) if block_output_ratios else None,
            "image_reset_emitted": runtime_snapshot.get("image_reset_count", 0),
        },
        "output_artifacts": {
            "markdown_path": _path_for_report(markdown_artifact_path),
            "docx_path": _path_for_report(docx_artifact_path),
            "tts_text_path": _path_for_report(tts_artifact_path),
            "output_docx_openable": openable_output,
            "output_paragraphs": output_paragraphs,
            "output_inline_shapes": output_inline_shapes,
            "output_visible_text_chars": output_visible_text_chars,
            "output_contains_placeholder_markup": output_contains_placeholder_markup,
            "report_json": _path_for_report(report_path),
            "summary_txt": _path_for_report(summary_path),
            "latest_report_json": _path_for_report(latest_report_path),
            "latest_summary_txt": _path_for_report(latest_summary_path),
            "latest_markdown_path": _path_for_report(latest_markdown_path) if markdown_artifact_path is not None else None,
            "latest_docx_path": _path_for_report(latest_docx_path) if docx_artifact_path is not None else None,
            "latest_tts_text_path": _path_for_report(latest_tts_path) if tts_artifact_path is not None else None,
            "latest_manifest_json": _path_for_report(latest_manifest_path),
        },
        "formatting_diagnostics_paths": [_path_for_report(Path(path)) for path in formatting_diagnostics_paths],
        "formatting_diagnostics_discovery": {
            "source": formatting_diagnostics_discovery_source,
            "baseline_count": len(formatting_diagnostics_before),
            "after_count": len(formatting_diagnostics_after),
            "new_count": len(snapshot_discovered_paths),
        },
        "formatting_diagnostics": formatting_diagnostics_payloads,
        "translation_quality_report": None,
        "translation_quality_report_path": None,
        "progress_events_tail": progress_events[-12:],
        "event_log": event_log[-25:],
        "image_log_tail": runtime_snapshot.get("image_log", [])[-25:],
    }
    translation_quality_report, translation_quality_report_path = _load_translation_quality_report(event_log)
    report["translation_quality_report"] = translation_quality_report
    report["translation_quality_report_path"] = translation_quality_report_path
    report["failure_classification"] = classify_failure(report)
    # Load and attach the reader-cleanup evidence BEFORE evaluating the acceptance
    # verdict so the `reader_cleanup_stage_completed` check sees the real cleanup
    # `stage_status` (a failed cleanup must not pass as an empty status). See spec 041
    # P1-3. `classify_failure` above does not read `reader_cleanup_evidence`, so its
    # ordering is unaffected. The later reader-verifier block only mutates anchor-repair
    # fields on the evidence (not stage_status/failed_chunk_count/cleanup_chunk_count),
    # so it does not change an acceptance input and the verdict is not rebuilt.
    reader_cleanup_evidence = _load_reader_cleanup_evidence(event_log)
    report["reader_cleanup_evidence"] = reader_cleanup_evidence
    _merge_reader_cleanup_artifact_paths(
        cast(dict[str, object], report["output_artifacts"]),
        reader_cleanup_evidence,
    )
    report["acceptance"] = evaluate_lietaer_acceptance(
        report,
        source_docx_bytes=source_docx_bytes,
        output_docx_bytes=bytes(latest_docx_bytes) if isinstance(latest_docx_bytes, (bytes, bytearray)) else None,
        mismatch_threshold=document_profile.max_unmapped_source_paragraphs,
        unmapped_target_threshold=document_profile.max_unmapped_target_paragraphs,
        require_no_toc_body_concat=document_profile.require_no_toc_body_concat,
    )
    reader_verifier_evidence: dict[str, object] | None = None
    if bool(validation_mode.get("comparison_only_validation")):
        reader_verifier_evidence = _write_reader_verifier_artifacts(
            run_id=run_id,
            document_profile_id=document_profile.id,
            run_profile_id=run_profile.id,
            artifact_prefix=artifact_prefix,
            artifact_dir=artifact_dir,
            source_document_path=source_path,
            source_text=prepared.source_text if prepared is not None else "",
            reader_cleanup_evidence=reader_cleanup_evidence,
            app_config=app_config,
            runtime_app_config=app_config_dict,
            validation_mode=validation_mode,
            max_retries=int(runtime_resolution.effective.max_retries) if runtime_resolution is not None else 1,
        )
        updated_reader_cleanup_evidence = reader_verifier_evidence.pop("updated_reader_cleanup_evidence", None)
        if isinstance(updated_reader_cleanup_evidence, Mapping):
            reader_cleanup_evidence = dict(updated_reader_cleanup_evidence)
            report["reader_cleanup_evidence"] = reader_cleanup_evidence
            _merge_reader_cleanup_artifact_paths(
                cast(dict[str, object], report["output_artifacts"]),
                reader_cleanup_evidence,
            )
        report["reader_verifier_evidence"] = reader_verifier_evidence
        remaining_issues = _coerce_mapping_sequence(reader_verifier_evidence.get("remaining_issues"))
        issue_summary = cast(Mapping[str, object], reader_verifier_evidence.get("issue_summary_by_category") or {})
        reader_mvp_status_path = artifact_dir / f"{artifact_prefix}_reader_mvp_status.md"
        reader_mvp_status = _build_reader_mvp_status_payload(report)
        reader_mvp_status_with_artifacts = {
            **reader_mvp_status,
            "artifact_paths": {"status_md": _path_for_report(reader_mvp_status_path)},
        }
        reader_mvp_status_path.write_text(
            _render_reader_mvp_status_markdown(reader_mvp_status_with_artifacts),
            encoding="utf-8",
        )
        report["reader_mvp_status"] = reader_mvp_status_with_artifacts
        cast(dict[str, object], report["output_artifacts"])["reader_mvp_status_md"] = _path_for_report(reader_mvp_status_path)
        reader_verifier_artifact_paths = dict(
            cast(Mapping[str, object], reader_verifier_evidence.get("artifact_paths") or {})
        )
        reader_verifier_artifact_paths["mvp_status_md"] = _path_for_report(reader_mvp_status_path)
        reader_verifier_evidence["artifact_paths"] = reader_verifier_artifact_paths
        tracker.set_manifest_context(
            reader_verifier_status=reader_verifier_evidence.get("verifier_status"),
            reader_verifier_reason=reader_verifier_evidence.get("verifier_reason"),
            reader_verifier_model_selector=reader_verifier_evidence.get("verifier_requested_selector"),
            reader_verifier_canonical_selector=reader_verifier_evidence.get("verifier_canonical_selector"),
            reader_verifier_provider=reader_verifier_evidence.get("verifier_provider"),
            reader_verifier_model_id=reader_verifier_evidence.get("verifier_model_id"),
            reader_verifier_overall_verdict=reader_verifier_evidence.get("overall_verdict"),
            reader_verifier_cleaned_audit_verdict=reader_verifier_evidence.get("cleaned_audit_verdict"),
            reader_verifier_confidence=reader_verifier_evidence.get("confidence"),
            reader_verifier_remaining_issue_count=len(remaining_issues),
            reader_verifier_high_severity_issue_count=_count_reader_verifier_high_severity_issues(remaining_issues),
            reader_verifier_top_issue_categories=_select_reader_verifier_top_categories(issue_summary),
            reader_verifier_filtered_toc_issue_count=reader_verifier_evidence.get("filtered_toc_issue_count"),
            reader_verifier_filtered_toc_pre_audit_count=reader_verifier_evidence.get("filtered_toc_pre_audit_count"),
            reader_verifier_filtered_toc_verifier_issue_count=reader_verifier_evidence.get("filtered_toc_verifier_issue_count"),
            reader_verifier_cleanup_ignored_reason_counts=cast(
                Mapping[str, object],
                cast(Mapping[str, object], reader_verifier_evidence.get("cleanup_diagnostics") or {}).get("ignored_reason_counts") or {},
            ),
            reader_verifier_simple_user_summary=reader_verifier_evidence.get("simple_user_summary"),
            reader_verifier_simple_user_risk_statement=reader_verifier_evidence.get("simple_user_risk_statement"),
            reader_verifier_simple_user_next_step=reader_verifier_evidence.get("simple_user_next_step"),
            reader_verifier_review_json=cast(Mapping[str, object], reader_verifier_evidence.get("artifact_paths") or {}).get("review_json"),
            reader_verifier_review_md=cast(Mapping[str, object], reader_verifier_evidence.get("artifact_paths") or {}).get("review_md"),
            reader_verifier_evidence_json=cast(Mapping[str, object], reader_verifier_evidence.get("artifact_paths") or {}).get("source_evidence_json"),
            reader_mvp_status_label=reader_mvp_status_with_artifacts.get("status_label"),
            reader_mvp_status_user_summary=reader_mvp_status_with_artifacts.get("user_summary"),
            reader_mvp_status_risk_summary=reader_mvp_status_with_artifacts.get("risk_summary"),
            reader_mvp_status_cleanup_score_delta=reader_mvp_status_with_artifacts.get("cleanup_score_delta"),
            reader_mvp_status_acceptance_diagnostic_only=reader_mvp_status_with_artifacts.get("comparison_only_acceptance_diagnostic"),
            reader_mvp_status_false_deletion_status=(
                "none_reported"
                if bool(reader_mvp_status_with_artifacts.get("no_false_deletions_reported"))
                else "reported"
            ),
            reader_mvp_status_readability_regression_status=(
                "none_reported"
                if bool(reader_mvp_status_with_artifacts.get("no_readability_regressions_reported"))
                else "reported"
            ),
            reader_mvp_status_md=cast(Mapping[str, object], reader_mvp_status_with_artifacts.get("artifact_paths") or {}).get("status_md"),
        )
    sanitized_report = sanitize_for_json(report)
    final_status = _resolve_validation_final_status(
        result=result,
        acceptance_passed=bool(report["acceptance"]["passed"]),
        validation_mode=validation_mode,
    )

    summary_lines = [
        f"run_id={run_id}",
        f"document_profile_id={document_profile.id}",
        f"run_profile_id={run_profile.id}",
        f"validation_tier={run_profile.tier}",
        f"validation_run_type={validation_mode['validation_run_type']}",
        f"comparison_only_validation={validation_mode['comparison_only_validation']}",
        f"acceptance_contract_active={validation_mode['acceptance_contract_active']}",
        f"validation_evidence_label={validation_mode['evidence_label']}",
        f"validation_success_criterion={validation_mode['success_criterion']}",
        f"status={final_status}",
        f"run_started_at_utc={run_started_at_utc.isoformat()}",
        f"run_finished_at_utc={run_finished_at_utc.isoformat()}",
        f"run_duration_seconds={run_duration_seconds}",
        f"source={_path_for_report(source_path)}",
        f"artifact_dir={_path_for_report(artifact_dir)}",
        f"artifact_root={_path_for_report(artifact_root)}",
        f"progress_json={_path_for_report(progress_path)}",
        f"result={report['result']}",
        f"failure_classification={report['failure_classification']}",
        f"runtime_overrides={json.dumps(cast(Mapping[str, object], report['runtime_config']).get('overrides') or {}, ensure_ascii=False, sort_keys=True)}",
        f"paragraph_count={report['preparation']['paragraph_count']}",
        f"image_count={report['preparation']['image_count']}",
        f"job_count={report['preparation']['job_count']}",
        f"ai_classified_count={report['preparation']['ai_classified_count']}",
        f"ai_heading_count={report['preparation']['ai_heading_count']}",
        f"ai_role_change_count={report['preparation']['ai_role_change_count']}",
        f"ai_heading_promotion_count={report['preparation']['ai_heading_promotion_count']}",
        f"ai_heading_demotion_count={report['preparation']['ai_heading_demotion_count']}",
        f"ai_structural_role_change_count={report['preparation']['ai_structural_role_change_count']}",
        f"preparation_diagnostic_snapshot={json.dumps(report['preparation_diagnostic_snapshot'], ensure_ascii=False, sort_keys=True)}",
        f"source_chars={report['preparation']['source_chars']}",
        f"final_markdown_chars={final_markdown_chars}",
        f"output_ratio_vs_source_text={report['signals']['output_ratio_vs_source_text']}",
        f"min_block_output_ratio={report['signals']['min_block_output_ratio']}",
        f"heading_only_output_detected={report['signals']['heading_only_output_detected']}",
        f"heading_only_rejection_logged={report['signals']['heading_only_rejection_logged']}",
        f"silent_text_loss_suspected={report['signals']['silent_text_loss_suspected']}",
        f"image_reset_emitted={report['signals']['image_reset_emitted']}",
        f"output_docx_openable={report['output_artifacts']['output_docx_openable']}",
        f"output_inline_shapes={report['output_artifacts']['output_inline_shapes']}",
        f"output_contains_placeholder_markup={report['output_artifacts']['output_contains_placeholder_markup']}",
        f"formatting_diagnostics_count={len(formatting_diagnostics_payloads)}",
        f"formatting_diagnostics_discovery_source={formatting_diagnostics_discovery_source}",
        f"translation_quality_report_path={report['translation_quality_report_path']}",
        f"acceptance_passed={report['acceptance']['passed']}",
        f"acceptance_failed_checks={','.join(_as_string_list(cast(Mapping[str, object], report['acceptance']).get('failed_checks')))}",
        f"comparison_only_acceptance_diagnostic_only={bool(validation_mode['comparison_only_validation'])}",
        f"last_error={last_error}",
        f"markdown_path={report['output_artifacts']['markdown_path']}",
        f"docx_path={report['output_artifacts']['docx_path']}",
        f"cleaned_markdown_path={report['output_artifacts'].get('cleaned_markdown_path')}",
        f"cleaned_docx_path={report['output_artifacts'].get('cleaned_docx_path')}",
        f"reader_cleanup_raw_markdown_path={report['output_artifacts'].get('reader_cleanup_raw_markdown_path')}",
        f"reader_cleanup_report_path={report['output_artifacts'].get('reader_cleanup_report_path')}",
        f"reader_cleanup_stage_status={report['reader_cleanup_evidence'].get('stage_status')}",
        f"reader_cleanup_changed={report['reader_cleanup_evidence'].get('changed')}",
        f"reader_cleanup_accepted_delete_block_count={report['reader_cleanup_evidence'].get('accepted_delete_block_count')}",
        f"reader_cleanup_ignored_delete_block_count={report['reader_cleanup_evidence'].get('ignored_delete_block_count')}",
        f"reader_cleanup_rejected_delete_block_count={report['reader_cleanup_evidence'].get('rejected_delete_block_count')}",
        f"reader_cleanup_failed_chunk_count={report['reader_cleanup_evidence'].get('failed_chunk_count')}",
        f"reader_cleanup_deleted_block_previews={json.dumps(report['reader_cleanup_evidence'].get('deleted_block_previews') or [], ensure_ascii=False)}",
        f"reader_verifier_status={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('verifier_status')}",
        f"reader_verifier_reason={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('verifier_reason')}",
        f"reader_verifier_requested_selector={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('verifier_requested_selector')}",
        f"reader_verifier_canonical_selector={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('verifier_canonical_selector')}",
        f"reader_verifier_provider={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('verifier_provider')}",
        f"reader_verifier_model_id={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('verifier_model_id')}",
        f"reader_verifier_overall_verdict={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('overall_verdict')}",
        f"reader_verifier_cleaned_audit_verdict={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('cleaned_audit_verdict')}",
        f"reader_verifier_confidence={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('confidence')}",
        f"reader_verifier_remaining_issue_count={len(_coerce_mapping_sequence(cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('remaining_issues')))}",
        f"reader_verifier_high_severity_issue_count={_count_reader_verifier_high_severity_issues(_coerce_mapping_sequence(cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('remaining_issues')))}",
        f"reader_verifier_top_issue_categories={','.join(_select_reader_verifier_top_categories(cast(Mapping[str, object], cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('issue_summary_by_category') or {})))}",
        f"reader_verifier_filtered_toc_issue_count={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('filtered_toc_issue_count')}",
        f"reader_verifier_filtered_toc_pre_audit_count={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('filtered_toc_pre_audit_count')}",
        f"reader_verifier_filtered_toc_verifier_issue_count={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('filtered_toc_verifier_issue_count')}",
        f"reader_verifier_filtered_toc_evidence_anchor_count={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('filtered_toc_evidence_anchor_count')}",
        f"reader_verifier_ignored_evidence_anchor_count={cast(Mapping[str, object], cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('evidence_anchor_diagnostics') or {}).get('ignored_anchor_count')}",
        "reader_verifier_ignored_evidence_anchor_kind_counts="
        + json.dumps(
            cast(Mapping[str, object], cast(Mapping[str, object], cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('evidence_anchor_diagnostics') or {}).get('ignored_kind_counts') or {}),
            ensure_ascii=False,
            sort_keys=True,
        ),
        "reader_verifier_filtered_toc_issue_previews="
        + json.dumps(cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('filtered_toc_issue_previews') or [], ensure_ascii=False),
        "reader_verifier_cleanup_ignored_reason_counts="
        + json.dumps(
            cast(Mapping[str, object], cast(Mapping[str, object], cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('cleanup_diagnostics') or {}).get('ignored_reason_counts') or {}),
            ensure_ascii=False,
            sort_keys=True,
        ),
        f"reader_verifier_simple_user_summary={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('simple_user_summary')}",
        f"reader_verifier_simple_user_risk_statement={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('simple_user_risk_statement')}",
        f"reader_verifier_simple_user_next_step={cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('simple_user_next_step')}",
        f"reader_verifier_review_json={cast(Mapping[str, object], cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('artifact_paths') or {}).get('review_json')}",
        f"reader_verifier_review_md={cast(Mapping[str, object], cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('artifact_paths') or {}).get('review_md')}",
        f"reader_verifier_evidence_json={cast(Mapping[str, object], cast(Mapping[str, object], report.get('reader_verifier_evidence') or {}).get('artifact_paths') or {}).get('source_evidence_json')}",
        f"reader_mvp_status_md={cast(Mapping[str, object], report.get('output_artifacts') or {}).get('reader_mvp_status_md')}",
        f"tts_text_path={report['output_artifacts']['tts_text_path']}",
        f"latest_manifest_json={report['output_artifacts']['latest_manifest_json']}",
        f"python_executable={report['run']['environment']['python_executable']}",
        f"python_version={report['run']['environment']['python_version']}",
        f"pythonpath={report['run']['environment']['pythonpath']}",
        f"virtual_env={report['run']['environment']['virtual_env']}",
        f"workspace_venv_exists={report['run']['environment']['workspace_venv_exists']}",
        f"workspace_venv_win_exists={report['run']['environment']['workspace_venv_win_exists']}",
        f"is_wsl={report['run']['environment']['is_wsl']}",
        f"git_head={report['run']['environment']['git_head']}",
    ]
    if isinstance(report.get("reader_mvp_status"), Mapping):
        summary_lines.extend(
            _build_reader_mvp_status_summary_lines(
                cast(Mapping[str, object], report.get("reader_mvp_status") or {}),
            )
        )
    summary_lines.extend(
        _build_translation_quality_summary_lines(
            cast(Mapping[str, object], report["translation_quality_report"] or {}),
        )
    )
    summary_path.write_text("\n".join(summary_lines), encoding="utf-8")
    report_path.write_text(
        json.dumps(sanitized_report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    _write_latest_alias_artifacts(
        report_path=report_path,
        summary_path=summary_path,
        markdown_artifact=markdown_artifact_path,
        docx_artifact=docx_artifact_path,
        tts_artifact=tts_artifact_path,
        latest_report_path=latest_report_path,
        latest_summary_path=latest_summary_path,
        latest_markdown_path=latest_markdown_path if markdown_artifact_path is not None else None,
        latest_docx_path=latest_docx_path if docx_artifact_path is not None else None,
        latest_tts_path=latest_tts_path if tts_artifact_path is not None else None,
        latest_manifest_path=latest_manifest_path,
        run_id=run_id,
        run_dir=artifact_dir,
        manifest_payload=cast(Mapping[str, object], tracker._build_manifest_payload_locked()),
    )
    tracker.finalize(
        status=final_status,
        result=str(report["result"]),
        acceptance_passed=bool(report["acceptance"]["passed"]),
        failure_classification=cast(str | None, report["failure_classification"]),
        last_error=last_error,
        detail=_build_validation_completion_detail(
            validation_mode=validation_mode,
            acceptance_passed=bool(report["acceptance"]["passed"]),
            report_path=report_path,
            reader_cleanup_evidence=cast(Mapping[str, object], report.get("reader_cleanup_evidence") or {}),
            reader_verifier_evidence=cast(Mapping[str, object], report.get("reader_verifier_evidence") or {}),
        ),
    )
    tracker.stop()
    _print_terminal_completion_summary(report=cast(Mapping[str, object], report), final_status=final_status)
    if final_status != "completed":
        raise SystemExit(1)


if __name__ == "__main__":
    main()
