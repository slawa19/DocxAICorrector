import base64
import json
import logging
import subprocess
import sys
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import pytest
from typing import Any, cast

import docxaicorrector.generation._generation as generation
import docxaicorrector.pipeline._pipeline as document_pipeline
import docxaicorrector.pipeline.late_phases as document_pipeline_late_phases
import docxaicorrector.pipeline.output_validation as document_pipeline_output_validation
import docxaicorrector.pipeline.reassembly as document_pipeline_reassembly
from docx import Document

from docxaicorrector.core.models import ParagraphUnit
from docxaicorrector.document._document import extract_document_content_from_docx
from docxaicorrector.reader_cleanup_mvp import build_cleanup_blocks


PNG_BYTES = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aK3cAAAAASUVORK5CYII=")


class AssetStub:
    def __init__(self, image_id: str):
        self.image_id = image_id
        self.placeholder_status = None

    def update_pipeline_metadata(self, **values):
        self.placeholder_status = values.get("placeholder_status")


class PlannedJobs:
    def __init__(self, jobs, *, planned_len: int):
        self._jobs = list(jobs)
        self._planned_len = planned_len

    def __iter__(self):
        return iter(self._jobs)

    def __len__(self):
        return self._planned_len


class ParagraphStub:
    role = "body"


def _build_runtime_capture():
    return {"state": {}, "finalize": [], "activity": [], "log": [], "status": []}


def _emit_state(runtime, **values):
    runtime.setdefault("state", {}).update(values)


def _emit_finalize(runtime, stage, detail, progress, terminal_kind=None):
    runtime.setdefault("finalize", []).append((stage, detail, progress, terminal_kind))


def _emit_activity(runtime, message):
    runtime.setdefault("activity", []).append(message)


def _emit_log(runtime, **payload):
    runtime.setdefault("log", []).append(payload)


def _emit_status(runtime, **payload):
    runtime.setdefault("status", []).append(payload)


def _inspect_placeholder_integrity(markdown_text, image_assets):
    return {asset.image_id: "ok" for asset in image_assets}


def _convert_markdown_to_docx_bytes(markdown_text):
    return b"docx-bytes"


def _reinsert_inline_images(docx_bytes, image_assets):
    return docx_bytes


def _run_processing(runtime, **overrides):
    params = {
        "uploaded_file": "report.docx",
        "source_token": "",
        "run_id": "",
        "jobs": [{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        "output_mode": None,
        "source_paragraphs": [],
        "image_assets": [],
        "image_mode": "safe",
        "app_config": {},
        "model": "gpt-5.4",
        "max_retries": 1,
        "on_progress": lambda **kwargs: None,
        "runtime": runtime,
        "resolve_uploaded_filename": lambda uploaded_file: str(uploaded_file),
        "get_client": lambda: object(),
        "ensure_pandoc_available": lambda: None,
        "load_system_prompt": lambda **_kw: "system",
        "log_event": lambda *args, **kwargs: None,
        "present_error": lambda code, exc, title, **kwargs: f"{title}: {exc}",
        "emit_state": _emit_state,
        "emit_finalize": _emit_finalize,
        "emit_activity": _emit_activity,
        "emit_log": _emit_log,
        "emit_status": _emit_status,
        "should_stop_processing": lambda runtime: False,
        "generate_markdown_block": lambda **kwargs: "Обработанный блок",
        "process_document_images": lambda **kwargs: [],
        "inspect_placeholder_integrity": _inspect_placeholder_integrity,
        "convert_markdown_to_docx_bytes": _convert_markdown_to_docx_bytes,
        "preserve_source_paragraph_properties": lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        "reinsert_inline_images": _reinsert_inline_images,
        "write_ui_result_artifacts": lambda **kwargs: {"markdown_path": "/tmp/final.result.md", "docx_path": "/tmp/final.result.docx"},
    }
    params.update(overrides)
    return document_pipeline.run_document_processing(**params)


def _run_processing_and_read_quality_report(tmp_path, runtime, **overrides):
    quality_dir = tmp_path / "quality_reports"
    overrides.setdefault("app_config", {"translation_output_quality_gate_policy": "strict"})
    overrides.setdefault("processing_operation", "translate")
    overrides.setdefault("generate_markdown_block", lambda **kwargs: "Обработанный блок")
    result = _run_processing(runtime, **overrides)
    report_files = list(quality_dir.glob("*.json"))
    payload = json.loads(report_files[0].read_text(encoding="utf-8")) if report_files else None
    return result, payload


def _capture_log_events():
    captured = []

    def log_event(level, event_id, message, **context):
        captured.append({"level": level, "event_id": event_id, "message": message, "context": context})

    return captured, log_event


def test_run_document_processing_happy_path_updates_runtime_state():
    runtime = _build_runtime_capture()
    progress_calls = []
    image_assets = [AssetStub("img_001")]

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[],
        image_assets=image_assets,
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: progress_calls.append(kwargs),
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Обработанный блок",
        process_document_images=lambda **kwargs: image_assets,
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_docx_bytes"] == b"final-docx"
    assert runtime["state"]["latest_markdown"] == "Обработанный блок"
    assert runtime["finalize"][-1] == ("Обработка завершена", runtime["finalize"][-1][1], 1.0, "completed")
    assert runtime["log"][-1]["status"] == "DONE"
    assert len(progress_calls) == 3


def test_run_document_processing_builds_docx_once_when_reader_cleanup_disabled():
    runtime = _build_runtime_capture()
    converted_markdown_inputs = []

    def convert_markdown_to_docx_bytes(markdown_text):
        converted_markdown_inputs.append(markdown_text)
        return markdown_text.encode("utf-8")

    result = _run_processing(
        runtime,
        app_config={"reader_cleanup_enabled": False},
        generate_markdown_block=lambda **kwargs: "final block",
        convert_markdown_to_docx_bytes=convert_markdown_to_docx_bytes,
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == "final block"
    assert runtime["state"]["latest_docx_bytes"] == b"final block"
    assert converted_markdown_inputs == ["final block"]


def test_pre_cleanup_formatting_baseline_is_diagnostic_only_rebuild_identity_snapshot():
    baseline = document_pipeline_late_phases._build_pre_cleanup_formatting_baseline(
        markdown_text="Intro translated\n\nBody translated",
        generated_paragraph_registry=[
            {"paragraph_id": "p0001", "text": "Intro translated"},
            {"paragraph_id": "p0002", "text": "Missing source"},
        ],
    )

    assert baseline == {
        "stage": "pre_reader_cleanup_rebuild_identity",
        "classification": "diagnostic_only",
        "mapping_basis": "ordered_exact_text_rebuild_sidecar",
        "metric_scope": "sidecar_only_proxy",
        "status": "computed",
        "source_count": 2,
        "target_count": 2,
        "mapped_count": 1,
        "unmapped_source_count": 1,
        "unmapped_target_count": 1,
        "unmapped_source_ids": ["p0002"],
        "unmapped_target_indexes": [1],
    }


def test_run_document_processing_passes_text_transform_context_to_system_prompt_loader():
    runtime = _build_runtime_capture()
    captured = {}

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={"editorial_intensity_default": "conservative"},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="de",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: captured.setdefault("prompt", dict(kwargs)) or "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Обработанный блок",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "succeeded"
    assert captured["prompt"] == {
        "operation": "translate",
        "source_language": "en",
        "target_language": "de",
        "editorial_intensity": "conservative",
        "prompt_variant": "default",
        "translation_domain": "general",
        "source_text": "",
    }


def test_run_document_processing_routes_provider_aware_text_and_image_clients():
    runtime = _build_runtime_capture()
    text_calls = []
    image_calls = []
    openrouter_client = object()
    openai_client = object()

    def resolve_model_selector(selector, required_capability=None):
        if selector == "openrouter:google/gemini-3.1-flash-lite-preview":
            return SimpleNamespace(
                raw_selector=selector,
                canonical_selector=selector,
                provider="openrouter",
                model_id="google/gemini-3.1-flash-lite-preview",
            )
        return SimpleNamespace(
            raw_selector=selector,
            canonical_selector=f"openai:{selector}",
            provider="openai",
            model_id=selector.removeprefix("openai:"),
        )

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[],
        image_assets=[AssetStub("img_001")],
        image_mode="semantic_redraw_direct",
        app_config={},
        model="openrouter:google/gemini-3.1-flash-lite-preview",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: openai_client,
        get_provider_client=lambda provider_name: openai_client if provider_name == "openai" else openrouter_client,
        get_client_for_model_selector=lambda selector, required_capability: openrouter_client,
        resolve_model_selector=resolve_model_selector,
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: text_calls.append(kwargs) or "Обработанный блок",
        process_document_images=lambda **kwargs: image_calls.append(kwargs) or kwargs["image_assets"],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "succeeded"
    assert text_calls[0]["client"] is openrouter_client
    assert text_calls[0]["model"] == "google/gemini-3.1-flash-lite-preview"
    assert image_calls[0]["client"] is openai_client


def test_run_document_processing_does_not_fallback_to_text_client_for_openai_image_phase():
    runtime = _build_runtime_capture()
    openrouter_client = object()
    image_calls = []

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[],
        image_assets=[AssetStub("img_001")],
        image_mode="semantic_redraw_direct",
        app_config={},
        model="openrouter:google/gemini-3.1-flash-lite-preview",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: (_ for _ in ()).throw(AssertionError("legacy get_client fallback must not run")),
        get_provider_client=lambda provider_name: (_ for _ in ()).throw(RuntimeError("missing OpenAI client")),
        get_client_for_model_selector=lambda selector, required_capability: openrouter_client,
        resolve_model_selector=lambda selector, required_capability=None: SimpleNamespace(
            raw_selector=selector,
            canonical_selector=selector,
            provider="openrouter",
            model_id="google/gemini-3.1-flash-lite-preview",
        ),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Обработанный блок",
        process_document_images=lambda **kwargs: image_calls.append(kwargs) or kwargs["image_assets"],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "failed"
    assert image_calls == []


def test_run_document_processing_persists_final_ui_result_artifacts_and_logs_paths():
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()
    captured = {}

    def write_ui_result_artifacts(**kwargs):
        captured["artifact_kwargs"] = kwargs
        return {
            "markdown_path": "/tmp/mariana.result.md",
            "docx_path": "/tmp/mariana.result.docx",
        }

    result = _run_processing(
        runtime,
        log_event=log_event,
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert captured["artifact_kwargs"] == {
        "source_name": "report.docx",
        "markdown_text": "Обработанный блок",
        "docx_bytes": b"docx-bytes",
        "assembly_mode": "full_document",
        "result_manifest": {
            "schema_version": 1,
            "source_name": "report.docx",
            "assembly_mode": "full_document",
            "output_mode": "legacy_full_document",
            "selected_segment_count": 0,
            "included_segment_count": 0,
            "included_segment_ids": [],
            "coverage": {
                "segment_ids": [],
                "paragraph_ranges": [],
            },
            "segments": [],
        },
    }
    info_events = [event for event in events if event["level"] == logging.INFO]
    saved_event = next(event for event in info_events if event["event_id"] == "ui_result_artifacts_saved")
    assert saved_event["context"]["artifact_paths"] == {
        "markdown_path": "/tmp/mariana.result.md",
        "docx_path": "/tmp/mariana.result.docx",
    }


def test_run_document_processing_builds_segment_aware_manifest_for_full_document_output():
    runtime = _build_runtime_capture()
    captured = {}
    source_paragraphs = [
        SimpleNamespace(role="body", segment_id="seg_0001", text="Chapter 1"),
        SimpleNamespace(role="body", segment_id="seg_0002", text="Chapter 2"),
    ]

    def write_ui_result_artifacts(**kwargs):
        captured["artifact_kwargs"] = kwargs
        return {
            "markdown_path": "/tmp/mariana.result.md",
            "docx_path": "/tmp/mariana.result.docx",
            "manifest_path": "/tmp/mariana.result.manifest.json",
        }

    result = _run_processing(
        runtime,
        source_token="report.docx:3:token",
        run_id="run-123",
        jobs=[
            {"target_text": "block 1", "context_before": "", "context_after": "", "target_chars": 7, "context_chars": 0, "segment_id": "seg_0001"},
            {"target_text": "block 2", "context_before": "", "context_after": "", "target_chars": 7, "context_chars": 0, "segment_id": "seg_0002"},
            {"target_text": "block 3", "context_before": "", "context_after": "", "target_chars": 7, "context_chars": 0, "segment_id": "seg_0002"},
        ],
        source_paragraphs=source_paragraphs,
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert captured["artifact_kwargs"]["assembly_mode"] == "full_document"
    assert "selected_segment_count" not in captured["artifact_kwargs"]
    assert captured["artifact_kwargs"]["result_manifest"] == {
        "schema_version": 1,
        "source_name": "report.docx",
        "source_token": "report.docx:3:token",
        "run_id": "run-123",
        "assembly_mode": "full_document",
        "output_mode": "legacy_full_document",
        "selected_segment_count": 0,
        "included_segment_count": 2,
        "included_segment_ids": ["seg_0001", "seg_0002"],
        "coverage": {
            "segment_ids": ["seg_0001", "seg_0002"],
            "paragraph_ranges": [
                {"segment_id": "seg_0001", "start_paragraph_index": 0, "end_paragraph_index": 0, "paragraph_count": 1},
                {"segment_id": "seg_0002", "start_paragraph_index": 1, "end_paragraph_index": 1, "paragraph_count": 1},
            ],
        },
        "segments": [
            {"segment_id": "seg_0001", "job_count": 1, "selected": False},
            {"segment_id": "seg_0002", "job_count": 2, "selected": False},
        ],
    }


def test_full_document_translate_run_preserves_segment_focus_and_document_context():
    """Anti-regression (spec 016 removal of partial translation).

    Segmentation is LOAD-BEARING for full-document TRANSLATE runs: ``document_segments``
    must still reach the worker so the per-block segment-focus prompt
    (``_build_block_segment_focus_prompt``, gated ``operation == "translate"``) and the
    document-context prompt (``build_document_context_prompt``) are built and threaded into
    the LLM system prompt for every block. Edit-mode tests alone would NOT catch a
    regression here, so this drives the translate path explicitly.
    """
    runtime = _build_runtime_capture()
    captured_system_prompts: list[str] = []
    captured_prompt_source_texts: list[str] = []

    document_segments = [
        SimpleNamespace(segment_id="seg_0001", title="Chapter 1", level=1, ordinal=1, structural_role="chapter"),
        SimpleNamespace(segment_id="seg_0002", title="Chapter 2", level=1, ordinal=2, structural_role="chapter"),
    ]

    def capture_generate(**kwargs):
        captured_system_prompts.append(str(kwargs.get("system_prompt", "")))
        return "Обработанный блок"

    def capture_load_system_prompt(**kwargs):
        captured_prompt_source_texts.append(str(kwargs.get("source_text", "")))
        return "system"

    result = _run_processing(
        runtime,
        jobs=[
            {"target_text": "block 1", "context_before": "", "context_after": "", "target_chars": 7, "context_chars": 0, "segment_id": "seg_0001"},
            {"target_text": "block 2", "context_before": "", "context_after": "", "target_chars": 7, "context_chars": 0, "segment_id": "seg_0002"},
        ],
        source_paragraphs=[
            SimpleNamespace(role="body", segment_id="seg_0001", text="Chapter 1"),
            SimpleNamespace(role="body", segment_id="seg_0002", text="Chapter 2"),
        ],
        document_segments=document_segments,
        output_mode="legacy_full_document",
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        document_context_prompt="ГЛОССАРИЙ: Great Tribulation -> Великая скорбь",
        load_system_prompt=capture_load_system_prompt,
        generate_markdown_block=capture_generate,
    )

    assert result == "succeeded"
    combined_system_prompt = "\n\n".join(captured_system_prompts)
    # document_segments reached the worker and the per-block segment-focus prompt was built:
    assert "ТЕКУЩИЙ БЛОК ДОКУМЕНТА" in combined_system_prompt
    assert "Chapter 1" in combined_system_prompt
    assert "Chapter 2" in combined_system_prompt
    # the document-context prompt (build_document_context_prompt output) still flows to blocks:
    assert any(
        "ГЛОССАРИЙ: Great Tribulation -> Великая скорбь" in source_text
        for source_text in captured_prompt_source_texts
    )


def test_full_document_edit_run_omits_translate_only_segment_focus_prompt():
    """Companion guard: the segment-focus prompt stays gated to translate runs, so an
    edit run over the same segment-tagged jobs must NOT emit the segment-focus block."""
    runtime = _build_runtime_capture()
    captured_system_prompts: list[str] = []

    def capture_generate(**kwargs):
        captured_system_prompts.append(str(kwargs.get("system_prompt", "")))
        return "Обработанный блок"

    result = _run_processing(
        runtime,
        jobs=[
            {"target_text": "block 1", "context_before": "", "context_after": "", "target_chars": 7, "context_chars": 0, "segment_id": "seg_0001"},
        ],
        document_segments=[
            SimpleNamespace(segment_id="seg_0001", title="Chapter 1", level=1, ordinal=1, structural_role="chapter"),
        ],
        processing_operation="edit",
        generate_markdown_block=capture_generate,
    )

    assert result == "succeeded"
    assert "ТЕКУЩИЙ БЛОК ДОКУМЕНТА" not in "\n\n".join(captured_system_prompts)


def test_run_document_processing_manifest_omits_noncanonical_segment_titles():
    runtime = _build_runtime_capture()
    captured = {}

    def write_ui_result_artifacts(**kwargs):
        captured["artifact_kwargs"] = kwargs
        return {
            "markdown_path": "/tmp/mariana.result.md",
            "docx_path": "/tmp/mariana.result.docx",
            "manifest_path": "/tmp/mariana.result.manifest.json",
        }

    result = _run_processing(
        runtime,
        jobs=[
            {"target_text": "block 1", "context_before": "", "context_after": "", "target_chars": 7, "context_chars": 0, "segment_id": "seg_0001"},
        ],
        source_paragraphs=[SimpleNamespace(role="body", segment_id="seg_0001", text="This is body text, not a canonical heading")],
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert captured["artifact_kwargs"]["result_manifest"]["segments"] == [
        {"segment_id": "seg_0001", "job_count": 1, "selected": False}
    ]


def test_run_document_processing_passes_machine_readable_quality_warning_to_artifact_writer_with_diagnostics(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    captured = {}

    def write_ui_result_artifacts(**kwargs):
        captured["artifact_kwargs"] = kwargs
        return {
            "markdown_path": "/tmp/report.result.md",
            "docx_path": "/tmp/report.result.docx",
            "manifest_path": "/tmp/report.result.manifest.json",
        }

    result = _run_processing(
        runtime,
        output_mode="legacy_full_document",
        jobs=[
            {
                "segment_id": "seg_0001",
                "target_text": "block",
                "context_before": "",
                "context_after": "",
                "target_chars": 5,
                "context_chars": 0,
            }
        ],
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert captured["artifact_kwargs"]["result_manifest"]["output_mode"] == "legacy_full_document"


def test_run_document_processing_passes_machine_readable_quality_warning_to_artifact_writer(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    quality_dir = tmp_path / "quality_reports"
    artifact_calls = {}

    monkeypatch.setattr(document_pipeline, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    diagnostics_dir.mkdir(parents=True, exist_ok=True)
    quality_dir.mkdir(parents=True, exist_ok=True)

    def preserve_with_unmapped_artifact(docx_bytes, paragraphs, generated_paragraph_registry=None):
        (diagnostics_dir / "preserve_001.json").write_text(
            json.dumps(
                {
                    "stage": "restore",
                    "source_count": 50,
                    "target_count": 48,
                    "mapped_count": 48,
                    "unmapped_source_ids": ["p0048", "p0049"],
                    "unmapped_target_indexes": [],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        return docx_bytes

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[ParagraphStub()],
        image_assets=[],
        image_mode="safe",
        app_config={"translation_output_quality_gate_policy": "advisory"},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="translate",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Обработанный блок",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: b"docx-bytes",
        preserve_source_paragraph_properties=preserve_with_unmapped_artifact,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
        write_ui_result_artifacts=lambda **kwargs: artifact_calls.setdefault("kwargs", dict(kwargs)) or {"markdown_path": "/tmp/report.result.md", "docx_path": "/tmp/report.result.docx", "metadata_path": "/tmp/report.result.meta.json"},
    )

    assert result == "succeeded"
    # Policy-independent discrepancy emission (GATE_TRUSTWORTHINESS Task B):
    # under advisory the pass/fail severity stays policy-scaled (warn +
    # unmapped_source_paragraphs_above_advisory_threshold) but the review-item
    # DATA is now emitted so the UI is not blind.
    assert artifact_calls["kwargs"]["quality_warning"] == {
        "kind": "translation_quality_gate",
        "quality_status": "warn",
        "gate_reasons": ["unmapped_source_paragraphs_above_advisory_threshold"],
        "message": (
            "Перевод завершён. Документ готов к использованию, но требует ручной "
            "проверки оформления: 2 абзаца с замечаниями. "
            "Подробности — в отчёте проверки (formatting_review.txt)."
        ),
        "formatting_review_items": [
            {
                "reason": "unmapped_source_paragraphs_review_required",
                "label": "Абзацы без явного соответствия оригиналу",
                "count": 2,
                "severity": "review",
            }
        ],
        "formatting_review_required_count": 2,
    }


def test_run_document_processing_warns_and_delivers_large_role_loss_with_formatting_review_items(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    quality_dir = tmp_path / "quality_reports"
    role_loss_ids = [f"p{index:04d}" for index in range(11)]

    monkeypatch.setattr(document_pipeline, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    diagnostics_dir.mkdir(parents=True, exist_ok=True)
    quality_dir.mkdir(parents=True, exist_ok=True)

    def preserve_with_role_loss_artifact(docx_bytes, paragraphs, generated_paragraph_registry=None):
        (diagnostics_dir / "restore_role_loss.json").write_text(
            json.dumps(
                {
                    "stage": "restore",
                    "unmapped_source_ids": role_loss_ids,
                    "unmapped_target_indexes": [],
                    "source_count": 1000,
                    "target_count": 989,
                    "source_registry": [
                        {
                            "paragraph_id": paragraph_id,
                            "source_index": index,
                            "role": "heading",
                            "structural_role": "heading",
                            "text_preview": f"Chapter {index}",
                        }
                        for index, paragraph_id in enumerate(role_loss_ids)
                    ],
                    "unmapped_source_residual_diagnostics": {
                        "effective_formatting_coverage_diagnostics": {
                            "counts": {
                                "content_survived_but_format_role_lost": 11,
                            },
                            "format_neutral_creditable_count": 0,
                        },
                        "samples": [
                            {
                                "paragraph_id": paragraph_id,
                                "source_index": index,
                                "role": "heading",
                                "structural_role": "heading",
                                "text_preview": f"Chapter {index}",
                                "effective_formatting_coverage_class": "content_survived_but_format_role_lost",
                            }
                            for index, paragraph_id in enumerate(role_loss_ids)
                        ],
                    },
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        return docx_bytes

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[
            ParagraphUnit(text=f"Chapter {index}", role="heading", structural_role="heading", paragraph_id=paragraph_id)
            for index, paragraph_id in enumerate(role_loss_ids)
        ],
        image_assets=[],
        image_mode="safe",
        app_config={"translation_output_quality_gate_policy": "strict", "enable_paragraph_markers": True},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="translate",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Many headings collapsed into body text",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: b"docx-bytes",
        preserve_source_paragraph_properties=preserve_with_role_loss_artifact,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
    )

    report_files = list(quality_dir.glob("*.json"))
    # spec 018: a large role_loss set is review-grade (fix severity) — the document is
    # fully usable, so the run is DELIVERED with a yellow "завершён, требует проверки"
    # notice instead of a blocking red failure. The gate_reasons + review items survive.
    assert result == "succeeded"
    assert runtime["finalize"][-1][3] != "error"
    assert runtime["state"]["latest_result_notice"] == {
        "level": "warning",
        "message": (
            "Перевод завершён. Документ готов к использованию, но требует ручной "
            "проверки оформления: 11 абзацев с замечаниями. "
            "Подробности — в отчёте проверки (formatting_review.txt)."
        ),
    }
    notice_message = runtime["state"]["latest_result_notice"]["message"]
    assert "translation_quality_gate_failed" not in notice_message
    assert "заблокирован" not in notice_message
    assert "критическая" not in notice_message.lower()
    assert len(report_files) == 1

    quality_report = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert quality_report["quality_status"] == "warn"
    assert quality_report["gate_reasons"] == ["role_loss_above_manual_review_threshold"]
    assert quality_report["formatting_review_required_count"] == 11
    review_items = quality_report["formatting_review_items"]
    assert len(review_items) == 8
    assert review_items[0]["aggregate_count"] == 11
    assert all(item["severity"] == "fix" for item in review_items)
    assert all(item["label"] == "Структурный абзац стал обычным текстом" for item in review_items)


def test_run_document_processing_persists_completed_job_result_records(monkeypatch):
    runtime = _build_runtime_capture()
    captured = {}

    def write_job_result_registry(*, records):
        captured["job_records"] = list(records)
        return {"job_0000": "/tmp/job-results/job_0000.job-result.json"}

    monkeypatch.setattr(document_pipeline, "write_job_result_registry_impl", write_job_result_registry)

    result = _run_processing(
        runtime,
        prepared_source_key="prep:report:1234",
        structure_fingerprint="struct-abc",
        jobs=[
            {
                "job_id": "job_0000",
                "segment_id": "seg_0001",
                "target_text": "block",
                "context_before": "",
                "context_after": "",
                "target_chars": 5,
                "context_chars": 0,
            }
        ],
    )

    assert result == "succeeded"
    assert len(captured["job_records"]) == 1
    completed_record = dict(captured["job_records"][0])
    assert isinstance(completed_record.pop("updated_at", ""), str)
    assert completed_record == {
        "schema_version": 1,
        "prepared_source_key": "prep:report:1234",
        "structure_fingerprint": "struct-abc",
        "job_id": "job_0000",
        "segment_id": "seg_0001",
        "status": "completed",
        "block_index": 1,
        "target_chars": 5,
        "context_chars": 0,
    }


def test_load_segment_result_records_returns_latest_record_per_segment(tmp_path):
    target_dir = tmp_path / "prep_report_1234" / "struct-abc"
    target_dir.mkdir(parents=True, exist_ok=True)
    (target_dir / "seg_0001.segment-result.json").write_text(
        json.dumps(
            {
                "segment_id": "seg_0001",
                "translated_markdown": "Stored translation",
                "prepared_source_key": "prep:report:1234",
                "structure_fingerprint": "struct-abc",
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    records = document_pipeline_reassembly.load_segment_result_records(
        prepared_source_key="prep:report:1234",
        structure_fingerprint="struct-abc",
        input_dir=tmp_path,
    )

    assert records == {
        "seg_0001": {
            "segment_id": "seg_0001",
            "translated_markdown": "Stored translation",
            "prepared_source_key": "prep:report:1234",
            "structure_fingerprint": "struct-abc",
        }
    }


def test_run_document_processing_builds_standalone_audiobook_artifact_and_coerces_image_mode():
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()
    captured = {}

    def write_ui_result_artifacts(**kwargs):
        captured["artifact_kwargs"] = kwargs
        return {
            "markdown_path": "/tmp/report.result.md",
            "docx_path": "/tmp/report.result.docx",
            "tts_text_path": "/tmp/report.result.tts.txt",
        }

    processed_image_modes = []

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "target_text": "# Title\n\n[thoughtful] Body text",
            "context_before": "",
            "context_after": "",
            "target_chars": 29,
            "context_chars": 0,
            "narration_include": True,
        }],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="audiobook",
        source_language="auto",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=log_event,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "# Title\n\n[thoughtful] Body text",
        process_document_images=lambda **kwargs: processed_image_modes.append(kwargs["image_mode"]) or [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert processed_image_modes == ["no_change"]
    assert runtime["state"]["latest_narration_text"] == "Title\n\n[thoughtful] Body text"
    assert captured["artifact_kwargs"]["narration_text"] == "Title\n\n[thoughtful] Body text"
    info_events = [event for event in events if event["level"] == logging.INFO]
    saved_event = next(event for event in info_events if event["event_id"] == "ui_audiobook_artifact_saved")
    assert saved_event["context"]["mode"] == "standalone"
    assert saved_event["context"]["filename"] == "report.docx"
    assert saved_event["context"]["artifact_paths"] == {
        "markdown_path": "/tmp/report.result.md",
        "docx_path": "/tmp/report.result.docx",
        "tts_text_path": "/tmp/report.result.tts.txt",
    }
    assert saved_event["context"]["tts_text_path"] == "/tmp/report.result.tts.txt"


def test_run_document_processing_runs_audiobook_postprocess_without_mutating_base_docx_branch():
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()
    generated_calls = []
    artifact_calls = {}

    def generate_markdown_block(**kwargs):
        generated_calls.append(dict(kwargs))
        if kwargs["system_prompt"] == "system:translate":
            return f"TRANSLATED::{kwargs['target_text']}"
        return f"# Narration\n\n[thoughtful] {kwargs['target_text']}"

    def write_ui_result_artifacts(**kwargs):
        artifact_calls["kwargs"] = dict(kwargs)
        return {
            "markdown_path": "/tmp/report.result.md",
            "docx_path": "/tmp/report.result.docx",
            "tts_text_path": "/tmp/report.result.tts.txt",
        }

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {
                "target_text": "Chapter one text",
                "context_before": "",
                "context_after": "Chapter two text",
                "target_chars": 16,
                "context_chars": 16,
                "narration_include": True,
            },
            {
                "target_text": "Index tail",
                "context_before": "Chapter two text",
                "context_after": "",
                "target_chars": 10,
                "context_chars": 16,
                "narration_include": False,
            },
            {
                "target_text": "Chapter two text",
                "context_before": "Chapter one text",
                "context_after": "",
                "target_chars": 16,
                "context_chars": 16,
                "narration_include": True,
            },
        ],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={
            "audiobook_postprocess_enabled": True,
            "audiobook_model": "gpt-5.4-audio",
            "chunk_size": 20,
            "editorial_intensity_default": "literary",
        },
        model="gpt-5.4-translate",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=log_event,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == "TRANSLATED::Chapter one text\n\nTRANSLATED::Index tail\n\nTRANSLATED::Chapter two text"
    assert runtime["state"]["latest_docx_bytes"] == runtime["state"]["latest_markdown"].encode("utf-8")
    narration_text = runtime["state"]["latest_narration_text"]
    assert narration_text is not None
    assert "TRANSLATED::Chapter one text" in narration_text
    assert "TRANSLATED::Chapter two text" in narration_text
    assert "TRANSLATED::Index tail" not in narration_text
    assert narration_text.count("[thoughtful]") == 1
    assert artifact_calls["kwargs"]["markdown_text"] == runtime["state"]["latest_markdown"]
    assert artifact_calls["kwargs"]["docx_bytes"] == runtime["state"]["latest_docx_bytes"]
    assert artifact_calls["kwargs"]["narration_text"] == narration_text

    assert [call["model"] for call in generated_calls] == [
        "gpt-5.4-translate",
        "gpt-5.4-translate",
        "gpt-5.4-translate",
        "gpt-5.4-audio",
    ]
    assert [call["target_text"] for call in generated_calls[3:]] == [
        "TRANSLATED::Chapter one text\n\nTRANSLATED::Chapter two text",
    ]
    assert generated_calls[3]["context_before"] == ""
    assert generated_calls[3]["context_after"] == ""

    info_events = [event for event in events if event["level"] == logging.INFO]
    saved_event = next(event for event in info_events if event["event_id"] == "ui_audiobook_artifact_saved")
    assert saved_event["context"]["mode"] == "postprocess"
    assert saved_event["context"]["filename"] == "report.docx"
    assert saved_event["context"]["artifact_paths"] == {
        "markdown_path": "/tmp/report.result.md",
        "docx_path": "/tmp/report.result.docx",
        "tts_text_path": "/tmp/report.result.tts.txt",
    }
    postprocess_started = [event for event in info_events if event["event_id"] == "audiobook_postprocess_chunk_started"]
    postprocess_completed = [event for event in info_events if event["event_id"] == "audiobook_postprocess_chunk_completed"]
    assert len(postprocess_started) == 1
    assert len(postprocess_completed) == 1
    assert all(event["context"]["operation"] == "audiobook" for event in postprocess_started)
    assert all(event["context"]["pass"] == "postprocess" for event in postprocess_started)
    completed_event = next(event for event in info_events if event["event_id"] == "processing_completed")
    assert completed_event["context"]["audiobook_postprocess_enabled"] is True


def test_run_document_processing_preserves_base_result_when_audiobook_postprocess_fails():
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()
    artifact_calls = {}

    def generate_markdown_block(**kwargs):
        if kwargs["system_prompt"] == "system:audiobook":
            raise RuntimeError("postprocess exploded")
        return "edited"

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0, "narration_include": True}],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={"audiobook_postprocess_enabled": True, "audiobook_model": "gpt-5.4-audio"},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="edit",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=log_event,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
        write_ui_result_artifacts=lambda **kwargs: artifact_calls.setdefault("kwargs", dict(kwargs)) or {"markdown_path": "/tmp/report.result.md", "docx_path": "/tmp/report.result.docx"},
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == "edited"
    assert runtime["state"]["latest_docx_bytes"] == b"edited"
    assert runtime["state"]["latest_narration_text"] is None
    assert "narration_text" not in artifact_calls["kwargs"]
    warning_events = [event for event in events if event["level"] == logging.WARNING]
    assert any(event["event_id"] == "audiobook_postprocess_failed_base_result_preserved" for event in warning_events)


def test_run_document_processing_applies_reader_cleanup_and_saves_raw_markdown_report_artifacts(tmp_path: Path):
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()
    artifact_calls = {}
    converted_markdown_inputs = []

    def generate_markdown_block(**kwargs):
        if kwargs["system_prompt"].startswith("You are cleaning translated book Markdown"):
            payload = json.loads(kwargs["target_text"])
            cleanup_operations = []
            for block in payload["blocks"]:
                if block["text"] == "Header":
                    cleanup_operations.append(
                        {
                            "id": block["id"],
                            "text_hash": block["text_hash"],
                            "operation": "delete_block",
                            "reason": "repeated_running_header",
                            "confidence": "high",
                            "evidence_before": block["text"],
                            "expected_after_preview": "",
                            "safety_note": "Test fixture deletes only the repeated running header block.",
                        }
                    )
            return json.dumps({"cleanup_operations": cleanup_operations, "warnings": []}, ensure_ascii=False)
        return kwargs["target_text"]

    def write_ui_result_artifacts(**kwargs):
        artifact_calls["kwargs"] = dict(kwargs)
        return {
            "markdown_path": str(tmp_path / "final.result.md"),
            "docx_path": str(tmp_path / "final.result.docx"),
        }

    def convert_markdown_to_docx_bytes(markdown_text):
        converted_markdown_inputs.append(markdown_text)
        return markdown_text.encode("utf-8")

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {"target_text": "Intro", "context_before": "", "context_after": "Header", "target_chars": 5, "context_chars": 6, "narration_include": True},
            {"target_text": "Header", "context_before": "Intro", "context_after": "Body paragraph", "target_chars": 6, "context_chars": 19, "narration_include": True},
            {"target_text": "Body paragraph", "context_before": "Header", "context_after": "Header", "target_chars": 14, "context_chars": 12, "narration_include": True},
            {"target_text": "Header", "context_before": "Body paragraph", "context_after": "Outro", "target_chars": 6, "context_chars": 19, "narration_include": True},
            {"target_text": "Outro", "context_before": "Header", "context_after": "", "target_chars": 5, "context_chars": 6, "narration_include": True},
        ],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={
            "reader_cleanup_enabled": True,
            "reader_cleanup_policy": "advisory",
            "reader_cleanup_chunk_size": 50,
            "reader_cleanup_keep_toc": True,
            "reader_cleanup_max_delete_block_ratio": 0.8,
            "reader_cleanup_max_delete_char_ratio": 0.8,
        },
        model="gpt-5.4-translate",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=log_event,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == "Intro\n\nBody paragraph\n\nOutro"
    assert runtime["state"]["latest_docx_bytes"] == b"Intro\n\nBody paragraph\n\nOutro"
    assert converted_markdown_inputs == ["Intro\n\nBody paragraph\n\nOutro"]
    assert artifact_calls["kwargs"]["markdown_text"] == "Intro\n\nBody paragraph\n\nOutro"
    info_events = [event for event in events if event["level"] == logging.INFO]
    saved_event = next(event for event in info_events if event["event_id"] == "ui_result_artifacts_saved")
    assert "reader_cleanup_raw_markdown_path" in saved_event["context"]["artifact_paths"]
    assert "reader_cleanup_report_path" in saved_event["context"]["artifact_paths"]
    cleanup_event = next(event for event in info_events if event["event_id"] == "reader_cleanup_applied")
    assert cleanup_event["context"]["accepted_delete_block_count"] == 2


def test_run_document_processing_records_pre_cleanup_formatting_baseline_without_extra_docx_build(
    tmp_path: Path,
    monkeypatch,
):
    runtime = _build_runtime_capture()
    converted_markdown_inputs = []
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    def generate_markdown_block(**kwargs):
        if kwargs["system_prompt"].startswith("You are cleaning translated book Markdown"):
            payload = json.loads(kwargs["target_text"])
            return json.dumps(
                {
                    "cleanup_operations": [
                        {
                            "id": block["id"],
                            "text_hash": block["text_hash"],
                            "operation": "delete_block",
                            "reason": "repeated_running_header",
                            "confidence": "high",
                            "evidence_before": "Header",
                            "expected_after_preview": "",
                            "safety_note": "Test fixture deletes only the repeated running header block.",
                        }
                        for block in payload["blocks"]
                        if block["text"] == "Header"
                    ],
                    "warnings": [],
                },
                ensure_ascii=False,
            )
        return kwargs["target_text"]

    def convert_markdown_to_docx_bytes(markdown_text):
        converted_markdown_inputs.append(markdown_text)
        return markdown_text.encode("utf-8")

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {"target_text": "Intro", "context_before": "", "context_after": "Header", "target_chars": 5, "context_chars": 6, "narration_include": True},
            {"target_text": "Header", "context_before": "Intro", "context_after": "Body", "target_chars": 6, "context_chars": 4, "narration_include": True},
            {"target_text": "Body", "context_before": "Header", "context_after": "Header", "target_chars": 4, "context_chars": 12, "narration_include": True},
            {"target_text": "Header", "context_before": "Body", "context_after": "Outro", "target_chars": 6, "context_chars": 9, "narration_include": True},
            {"target_text": "Outro", "context_before": "Header", "context_after": "", "target_chars": 5, "context_chars": 6, "narration_include": True},
        ],
        source_paragraphs=[
            ParagraphUnit(text="Intro", role="body", paragraph_id="p0001"),
            ParagraphUnit(text="Header", role="body", paragraph_id="p0002"),
            ParagraphUnit(text="Body", role="body", paragraph_id="p0003"),
            ParagraphUnit(text="Header", role="body", paragraph_id="p0004"),
            ParagraphUnit(text="Outro", role="body", paragraph_id="p0005"),
        ],
        image_assets=[],
        image_mode="safe",
        app_config={
            "reader_cleanup_enabled": True,
            "reader_cleanup_policy": "advisory",
            "reader_cleanup_chunk_size": 50,
            "reader_cleanup_max_delete_block_ratio": 0.8,
            "reader_cleanup_max_delete_char_ratio": 0.8,
        },
        model="gpt-5.4-translate",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
        write_ui_result_artifacts=lambda **kwargs: {
            "markdown_path": str(tmp_path / "final.result.md"),
            "docx_path": str(tmp_path / "final.result.docx"),
        },
    )

    assert result == "succeeded"
    assert converted_markdown_inputs == ["Intro\n\nBody\n\nOutro"]
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    quality_payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert quality_payload["formatting_diagnostics_artifact_count"] == 0
    assert quality_payload["pre_cleanup_formatting_baseline"] == {
        "stage": "pre_reader_cleanup_rebuild_identity",
        "classification": "diagnostic_only",
        "mapping_basis": "ordered_exact_text_rebuild_sidecar",
        "metric_scope": "sidecar_only_proxy",
        "status": "missing_registry",
        "source_count": 0,
        "target_count": 5,
        "mapped_count": 0,
        "unmapped_source_count": 0,
        "unmapped_target_count": 5,
        "unmapped_source_ids": [],
        "unmapped_target_indexes": [0, 1, 2, 3, 4],
    }


def test_run_document_processing_fails_on_empty_lazy_docx_after_reader_cleanup_noop():
    runtime = _build_runtime_capture()

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0, "narration_include": True}],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={"reader_cleanup_enabled": True, "reader_cleanup_policy": "advisory"},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{code}:{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: json.dumps({"cleanup_operations": [], "warnings": []}, ensure_ascii=False)
        if kwargs["system_prompt"].startswith("You are cleaning translated book Markdown")
        else "block",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: b"",
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
    )

    assert result == "failed"
    assert "empty_docx_bytes" in runtime["state"]["last_error"]
    assert runtime["state"]["latest_docx_bytes"] is None


def test_run_document_processing_fails_on_empty_lazy_docx_before_cleanup_quality_gate():
    runtime = _build_runtime_capture()

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0, "narration_include": True}],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={
            "reader_cleanup_enabled": True,
            "reader_cleanup_policy": "advisory",
            "translation_output_quality_gate_policy": "strict",
        },
        model="gpt-5.4",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{code}:{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Заключение……29 Введение",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: b"",
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
    )

    assert result == "failed"
    assert "empty_docx_bytes" in runtime["state"]["last_error"]
    assert "translation_quality_gate_failed" not in runtime["state"]["last_error"]


def test_run_document_processing_reader_cleanup_uses_exact_raw_markdown_for_sidecar_and_hashes(tmp_path: Path):
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()
    artifact_calls = {}
    cleanup_payloads = []
    noisy_block = "This page intentionally left blank Chapter Nine STRATEGIES FOR NGO S"
    raw_markdown = f"Intro\n\n{noisy_block}\n\n{noisy_block}\n\nOutro"
    display_markdown = document_pipeline_late_phases._normalize_final_markdown_for_runtime_display(raw_markdown)

    assert display_markdown != raw_markdown
    assert len(build_cleanup_blocks(raw_markdown)) == 4
    assert len(build_cleanup_blocks(display_markdown)) == 6

    def generate_markdown_block(**kwargs):
        if kwargs["system_prompt"].startswith("You are cleaning translated book Markdown"):
            payload = json.loads(kwargs["target_text"])
            cleanup_payloads.append(payload)
            cleanup_operations = []
            for block in payload["blocks"]:
                if block["text"] == noisy_block:
                    cleanup_operations.append(
                        {
                            "id": block["id"],
                            "text_hash": block["text_hash"],
                            "operation": "delete_block",
                            "reason": "repeated_running_header",
                            "confidence": "high",
                            "evidence_before": block["text"],
                            "expected_after_preview": "",
                            "safety_note": "Test fixture deletes only exact repeated running-header blocks.",
                        }
                    )
            return json.dumps({"cleanup_operations": cleanup_operations, "warnings": []}, ensure_ascii=False)
        return kwargs["target_text"]

    def write_ui_result_artifacts(**kwargs):
        artifact_calls["kwargs"] = dict(kwargs)
        return {
            "markdown_path": str(tmp_path / "final.result.md"),
            "docx_path": str(tmp_path / "final.result.docx"),
        }

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {"target_text": "Intro", "context_before": "", "context_after": noisy_block, "target_chars": 5, "context_chars": len(noisy_block), "narration_include": True},
            {"target_text": noisy_block, "context_before": "Intro", "context_after": noisy_block, "target_chars": len(noisy_block), "context_chars": 10, "narration_include": True},
            {"target_text": noisy_block, "context_before": noisy_block, "context_after": "Outro", "target_chars": len(noisy_block), "context_chars": 10, "narration_include": True},
            {"target_text": "Outro", "context_before": noisy_block, "context_after": "", "target_chars": 5, "context_chars": len(noisy_block), "narration_include": True},
        ],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={
            "reader_cleanup_enabled": True,
            "reader_cleanup_policy": "advisory",
            "reader_cleanup_chunk_size": 500,
            "reader_cleanup_keep_toc": True,
            "reader_cleanup_max_delete_block_ratio": 0.8,
            "reader_cleanup_max_delete_char_ratio": 0.95,
        },
        model="gpt-5.4-translate",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=log_event,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == "Intro\n\nOutro"
    assert artifact_calls["kwargs"]["markdown_text"] == "Intro\n\nOutro"
    assert cleanup_payloads
    assert [block["text"] for block in cleanup_payloads[0]["blocks"]].count(noisy_block) == 2

    raw_sidecar_path = tmp_path / "final.raw.result.md"
    report_path = tmp_path / "final.reader_cleanup_report.json"
    assert raw_sidecar_path.read_text(encoding="utf-8") == raw_markdown
    report_payload = json.loads(report_path.read_text(encoding="utf-8"))
    expected_hashes = {block.text_hash for block in build_cleanup_blocks(raw_markdown) if block.text == noisy_block}
    accepted_hashes = {entry["text_hash"] for entry in report_payload["accepted_delete_blocks"]}
    assert report_payload["stats"]["accepted_delete_block_count"] == 2
    assert accepted_hashes == expected_hashes

    info_events = [event for event in events if event["level"] == logging.INFO]
    cleanup_event = next(event for event in info_events if event["event_id"] == "reader_cleanup_applied")
    assert cleanup_event["context"]["proposed_delete_block_count"] == 2


def test_run_document_processing_reader_cleanup_applies_runtime_anchor_repair(tmp_path: Path):
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()
    artifact_calls = {}
    cleanup_payloads = []
    target = "КАК ЭТО РАБОТАЕТ: Местные органы власти могут помочь."
    raw_markdown = f"Intro\n\n{target}\n\nOutro"
    blocks = build_cleanup_blocks(raw_markdown)

    def generate_markdown_block(**kwargs):
        if kwargs["system_prompt"].startswith("You are cleaning translated book Markdown"):
            payload = json.loads(kwargs["target_text"])
            cleanup_payloads.append(payload)
            if payload.get("pass_name") == "anchor_repair":
                target_block = next(block for block in payload["blocks"] if block["id"] == blocks[1].block_id)
                return json.dumps(
                    {
                        "cleanup_operations": [
                            {
                                "id": target_block["id"],
                                "text_hash": target_block["text_hash"],
                                "operation": "normalize_heading_boundary",
                                "reason": "page_furniture_heading",
                                "confidence": "high",
                                "evidence_before": target_block["text"],
                                "expected_after_preview": "КАК ЭТО РАБОТАЕТ:\n\nМестные органы власти могут помочь.",
                                "safety_note": "Split only the exact heading/body boundary from a verifier anchor.",
                                "heading_substring": "КАК ЭТО РАБОТАЕТ:",
                                "body_substring": "Местные органы власти могут помочь.",
                            }
                        ],
                        "warnings": [],
                    },
                    ensure_ascii=False,
                )
            return json.dumps({"cleanup_operations": [], "warnings": []}, ensure_ascii=False)
        return kwargs["target_text"]

    def write_ui_result_artifacts(**kwargs):
        artifact_calls["kwargs"] = dict(kwargs)
        return {
            "markdown_path": str(tmp_path / "final.result.md"),
            "docx_path": str(tmp_path / "final.result.docx"),
        }

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {"target_text": "Intro", "context_before": "", "context_after": target, "target_chars": 5, "context_chars": len(target), "narration_include": True},
            {"target_text": target, "context_before": "Intro", "context_after": "Outro", "target_chars": len(target), "context_chars": 10, "narration_include": True},
            {"target_text": "Outro", "context_before": target, "context_after": "", "target_chars": 5, "context_chars": len(target), "narration_include": True},
        ],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={
            "reader_cleanup_enabled": True,
            "reader_cleanup_policy": "advisory",
            "reader_cleanup_chunk_size": 500,
            "reader_cleanup_keep_toc": True,
            "reader_cleanup_anchor_repair_enabled": True,
            "reader_cleanup_anchor_targets": [
                {
                    "anchor_id": "runtime-anchor-1",
                    "category": "heading_fused_with_body",
                    "block_id": blocks[1].block_id,
                    "line_ref": "cleaned_markdown:3",
                    "snippet": target,
                }
            ],
        },
        model="gpt-5.4-translate",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=log_event,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == "Intro\n\nКАК ЭТО РАБОТАЕТ:\n\nМестные органы власти могут помочь.\n\nOutro"
    assert artifact_calls["kwargs"]["markdown_text"] == runtime["state"]["latest_markdown"]
    assert any(payload.get("pass_name") == "anchor_repair" for payload in cleanup_payloads)

    report_path = tmp_path / "final.reader_cleanup_report.json"
    report_payload = json.loads(report_path.read_text(encoding="utf-8"))
    anchor_pass = report_payload["passes"]["anchor_repair_pass"]
    assert anchor_pass["selected_anchor_count"] == 1
    assert anchor_pass["stats"]["accepted_cleanup_operation_count"] == 1
    assert report_payload["stats"]["accepted_cleanup_operation_count"] == 1

    info_events = [event for event in events if event["level"] == logging.INFO]
    anchor_events = [
        event for event in info_events if event["context"].get("pass") == "anchor_repair"
    ]
    assert any(event["event_id"] == "reader_cleanup_chunk_started" for event in anchor_events)
    cleanup_event = next(event for event in info_events if event["event_id"] == "reader_cleanup_applied")
    assert cleanup_event["context"]["accepted_cleanup_operation_count"] == 1


def test_run_document_processing_preserves_base_result_when_reader_cleanup_fails():
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()
    artifact_calls = {}
    converted_markdown_inputs = []

    def generate_markdown_block(**kwargs):
        if kwargs["system_prompt"].startswith("You are cleaning translated book Markdown"):
            raise RuntimeError("cleanup exploded")
        return kwargs["target_text"]

    def write_ui_result_artifacts(**kwargs):
        artifact_calls["kwargs"] = dict(kwargs)
        return {"markdown_path": "/tmp/report.result.md", "docx_path": "/tmp/report.result.docx"}

    def convert_markdown_to_docx_bytes(markdown_text):
        converted_markdown_inputs.append(markdown_text)
        return markdown_text.encode("utf-8")

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0, "narration_include": True}],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={"reader_cleanup_enabled": True, "reader_cleanup_policy": "advisory"},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=log_event,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == "block"
    assert runtime["state"]["latest_docx_bytes"] == b"block"
    assert converted_markdown_inputs == ["block"]
    assert artifact_calls["kwargs"]["markdown_text"] == "block"
    info_events = [event for event in events if event["level"] == logging.INFO]
    noop_event = next(event for event in info_events if event["event_id"] == "reader_cleanup_noop")
    assert any("reader_cleanup_chunk_failed" in warning for warning in noop_event["context"]["warnings"])


def test_run_document_processing_reader_cleanup_rebuild_preserves_images():
    reinsert_calls = []

    class FakeImageAsset:
        image_id = "img-1"

        def update_pipeline_metadata(self, **values):
            return None

    def reinsert_inline_images(docx_bytes, image_assets):
        reinsert_calls.append([asset.image_id for asset in image_assets])
        return docx_bytes + f"|images={len(image_assets)}".encode("utf-8")

    rebuilt = document_pipeline_late_phases._rebuild_docx_for_markdown(
        markdown_text="Intro\n\nBody\n\nOutro",
        context=SimpleNamespace(source_paragraphs=[]),
        dependencies=SimpleNamespace(
            convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
            preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
            reinsert_inline_images=reinsert_inline_images,
        ),
        state=SimpleNamespace(generated_paragraph_registry=[]),
        processed_image_assets=[FakeImageAsset()],
    )

    assert rebuilt == b"Intro\n\nBody\n\nOutro|images=1"
    assert reinsert_calls
    assert reinsert_calls[-1] == ["img-1"]


def test_reader_cleanup_docx_rebuild_markdown_restores_missing_image_placeholder_without_display_regression():
    rebuilt_markdown = document_pipeline_late_phases._build_docx_rebuild_markdown_after_reader_cleanup(
        raw_markdown="Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph\n\nOutro",
        cleaned_markdown="Intro\n\nBody paragraph\n\nOutro",
        accepted_delete_block_ids=[],
    )

    assert rebuilt_markdown == "Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph\n\nOutro"


def test_reader_cleanup_docx_rebuild_markdown_does_not_duplicate_existing_image_placeholder():
    rebuilt_markdown = document_pipeline_late_phases._build_docx_rebuild_markdown_after_reader_cleanup(
        raw_markdown="Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph",
        cleaned_markdown="Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph",
        accepted_delete_block_ids=[],
    )

    assert rebuilt_markdown.count("[[DOCX_IMAGE_img_001]]") == 1
    assert rebuilt_markdown == "Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph"


def test_reader_cleanup_docx_rebuild_markdown_preserves_consecutive_image_placeholder_order():
    rebuilt_markdown = document_pipeline_late_phases._build_docx_rebuild_markdown_after_reader_cleanup(
        raw_markdown="Intro\n\n[[DOCX_IMAGE_img_001]]\n\n[[DOCX_IMAGE_img_002]]\n\nBody paragraph",
        cleaned_markdown="Intro\n\nBody paragraph",
        accepted_delete_block_ids=[],
    )

    assert rebuilt_markdown == "Intro\n\n[[DOCX_IMAGE_img_001]]\n\n[[DOCX_IMAGE_img_002]]\n\nBody paragraph"


def test_reader_cleanup_docx_rebuild_markdown_anchors_missing_image_placeholder_by_paragraph_id():
    rebuilt_markdown = document_pipeline_late_phases._build_docx_rebuild_markdown_after_reader_cleanup(
        raw_markdown="Intro before cleanup\n\n[[DOCX_IMAGE_img_001]]\n\nBody before cleanup",
        cleaned_markdown="Intro after cleanup\n\nBody after cleanup",
        accepted_delete_block_ids=[],
        cleanup_block_metadata_by_index={
            0: {"paragraph_id": "p0001"},
            2: {"paragraph_id": "p0002"},
        },
        generated_paragraph_registry=[
            {"paragraph_id": "p0001", "text": "Intro after cleanup"},
            {"paragraph_id": "p0002", "text": "Body after cleanup"},
        ],
    )

    assert rebuilt_markdown == "Intro after cleanup\n\n[[DOCX_IMAGE_img_001]]\n\nBody after cleanup"


def test_run_document_processing_reader_cleanup_preserves_docx_image_anchor_when_markdown_cleanup_deletes_placeholder():
    runtime = _build_runtime_capture()
    artifact_calls = {}
    converted_markdown_inputs = []

    class FakeImageAsset:
        image_id = "img-1"

        def update_pipeline_metadata(self, **values):
            return None

    def generate_markdown_block(**kwargs):
        if kwargs["system_prompt"].startswith("You are cleaning translated book Markdown"):
            payload = json.loads(kwargs["target_text"])
            cleanup_operations = [
                {
                    "id": block["id"],
                    "text_hash": block["text_hash"],
                    "operation": "delete_block",
                    "reason": "extraction_artifact",
                    "confidence": "high",
                    "evidence_before": block["text"],
                    "expected_after_preview": "",
                    "safety_note": "Test fixture deletes only the exact placeholder block from display Markdown.",
                }
                for block in payload["blocks"]
                if block["text"] == "[[DOCX_IMAGE_img_001]]"
            ]
            return json.dumps({"cleanup_operations": cleanup_operations, "warnings": []}, ensure_ascii=False)
        return kwargs["target_text"]

    def convert_markdown_to_docx_bytes(markdown_text):
        converted_markdown_inputs.append(markdown_text)
        return markdown_text.encode("utf-8")

    def write_ui_result_artifacts(**kwargs):
        artifact_calls["kwargs"] = dict(kwargs)
        return {"markdown_path": "/tmp/report.result.md", "docx_path": "/tmp/report.result.docx"}

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {"target_text": "Intro", "context_before": "", "context_after": "[[DOCX_IMAGE_img_001]]", "target_chars": 5, "context_chars": 22, "narration_include": True},
            {"target_text": "[[DOCX_IMAGE_img_001]]", "context_before": "Intro", "context_after": "Body paragraph", "target_chars": 22, "context_chars": 19, "narration_include": True},
            {"target_text": "Body paragraph", "context_before": "[[DOCX_IMAGE_img_001]]", "context_after": "Outro", "target_chars": 14, "context_chars": 27, "narration_include": True},
            {"target_text": "Outro", "context_before": "Body paragraph", "context_after": "", "target_chars": 5, "context_chars": 14, "narration_include": True},
        ],
        source_paragraphs=[],
        image_assets=[FakeImageAsset()],
        image_mode="safe",
        app_config={
            "reader_cleanup_enabled": True,
            "reader_cleanup_policy": "advisory",
            "reader_cleanup_chunk_size": 500,
            "reader_cleanup_max_delete_block_ratio": 0.8,
            "reader_cleanup_max_delete_char_ratio": 0.8,
        },
        model="gpt-5.4-translate",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [FakeImageAsset()],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes + f"|images={len(image_assets)}".encode("utf-8"),
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == "Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph\n\nOutro"
    assert artifact_calls["kwargs"]["markdown_text"] == "Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph\n\nOutro"
    assert converted_markdown_inputs[-1] == "Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph\n\nOutro"
    assert runtime["state"]["latest_docx_bytes"] == b"Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph\n\nOutro|images=1"


def test_run_document_processing_reader_cleanup_strict_failure_preserves_base_result(tmp_path: Path):
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()
    artifact_calls = {}

    def generate_markdown_block(**kwargs):
        if kwargs["system_prompt"].startswith("You are cleaning translated book Markdown"):
            raise RuntimeError("cleanup exploded")
        return kwargs["target_text"]

    def write_ui_result_artifacts(**kwargs):
        artifact_calls["kwargs"] = dict(kwargs)
        return {
            "markdown_path": str(tmp_path / "report.result.md"),
            "docx_path": str(tmp_path / "report.result.docx"),
        }

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0, "narration_include": True}],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={"reader_cleanup_enabled": True, "reader_cleanup_policy": "strict"},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs['operation']}",
        log_event=log_event,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
        write_ui_result_artifacts=write_ui_result_artifacts,
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == "block"
    assert artifact_calls["kwargs"]["markdown_text"] == "block"
    assert runtime["state"]["latest_result_notice"] == {
        "level": "warning",
        "message": "Reader cleanup strict stage failed; preserved the raw translated result without cleanup.",
    }
    warning_events = [event for event in events if event["level"] == logging.WARNING]
    assert any(event["event_id"] == "reader_cleanup_strict_failed_base_result_preserved" for event in warning_events)
    raw_sidecar_path = tmp_path / "report.raw.result.md"
    report_path = tmp_path / "report.reader_cleanup_report.json"
    assert raw_sidecar_path.read_text(encoding="utf-8") == "block"
    report_payload = json.loads(report_path.read_text(encoding="utf-8"))
    assert report_payload["stage_status"] == "failed"
    assert report_payload["failure"]["kind"] == "chunk_failed"


def test_run_document_processing_fails_standalone_audiobook_with_invalid_narration_artifact():
    runtime = _build_runtime_capture()

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0, "narration_include": True}],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="audiobook",
        source_language="auto",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "[angry] text with DOI:10.1000/xyz",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
    )

    assert result == "failed"
    assert runtime["state"]["latest_docx_bytes"] is None
    assert runtime["state"]["latest_narration_text"] is None
    assert runtime["finalize"][-1][0] == "Ошибка проверки narration"


def test_run_document_processing_clears_stale_narration_on_docx_build_failure():
    runtime = _build_runtime_capture()
    runtime["state"]["latest_narration_text"] = "stale narration"

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="edit",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "edited",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: (_ for _ in ()).throw(RuntimeError("docx exploded")),
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
    )

    assert result == "failed"
    assert runtime["state"]["latest_narration_text"] is None


def test_run_document_processing_uses_base_model_for_audiobook_postprocess_when_configured_model_blank():
    runtime = _build_runtime_capture()
    generated_calls = []

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "target_text": "block",
            "context_before": "",
            "context_after": "",
            "target_chars": 5,
            "context_chars": 0,
            "narration_include": True,
        }],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={
            "audiobook_postprocess_enabled": True,
            "audiobook_model": "   ",
            "chunk_size": 6000,
        },
        model="gpt-5.4-base",
        max_retries=1,
        processing_operation="edit",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: "system:audiobook" if kwargs["operation"] == "audiobook" else "system:edit",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: generated_calls.append(dict(kwargs)) or ("edited" if len(generated_calls) == 1 else "[thoughtful] edited"),
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "succeeded"
    assert len(generated_calls) == 2
    assert generated_calls[1]["model"] == "gpt-5.4-base"
    assert runtime["state"]["latest_narration_text"] == "[thoughtful] edited"


def test_resolve_audiobook_postprocess_chunk_size_clamps_to_config_minimum():
    context = type("Context", (), {"app_config": {"chunk_size": 1200}})()

    assert document_pipeline_late_phases._resolve_audiobook_postprocess_chunk_size(context=context) == 3000


def test_run_document_processing_does_not_fail_when_ui_result_artifact_save_fails():
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()

    result = _run_processing(
        runtime,
        log_event=log_event,
        write_ui_result_artifacts=lambda **kwargs: (_ for _ in ()).throw(OSError("disk full")),
    )

    assert result == "succeeded"
    warning_events = [event for event in events if event["level"] == logging.WARNING]
    failed_event = next(event for event in warning_events if event["event_id"] == "ui_result_artifacts_save_failed")
    assert failed_event["context"]["error_message"] == "disk full"


def test_resolve_system_prompt_does_not_mask_internal_type_errors():
    def broken_loader(
        *,
        operation: str = "edit",
        source_language: str = "en",
        target_language: str = "ru",
        editorial_intensity: str = "literary",
        prompt_variant: str = "default",
        translation_domain: str = "general",
        source_text: str = "",
    ):
        raise TypeError("broken template")

    with pytest.raises(TypeError, match="broken template"):
        document_pipeline._resolve_system_prompt(
            broken_loader,
            operation="translate",
            source_language="en",
            target_language="de",
            editorial_intensity="literary",
        )


def test_resolve_system_prompt_falls_back_for_legacy_loader_without_editorial_intensity():
    captured = {}

    def legacy_loader(*, operation: str = "edit", source_language: str = "en", target_language: str = "ru"):
        captured["prompt"] = {
            "operation": operation,
            "source_language": source_language,
            "target_language": target_language,
        }
        return "system"

    resolved = document_pipeline._resolve_system_prompt(
        legacy_loader,  # type: ignore[arg-type]  # intentional: testing backward-compat fallback for pre-prompt_variant loaders
        operation="translate",
        source_language="en",
        target_language="de",
        editorial_intensity="conservative",
        translation_domain="theology",
        source_text="Great Tribulation",
    )

    assert resolved == "system"
    assert captured["prompt"] == {
        "operation": "translate",
        "source_language": "en",
        "target_language": "de",
    }


def test_run_document_processing_merges_document_context_prompt_into_system_prompt_source():
    runtime = _build_runtime_capture()
    prompts = []

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={
            "translation_domain_default": "theology",
            "translation_domain_instructions": "ДОМЕН ПЕРЕВОДА: богословие / эсхатология.",
        },
        document_context_prompt="КОНТЕКСТ ДОКУМЕНТА: Chapter 1; Great Tribulation -> die grosse Trubsal",
        model="gpt-5.4-mini",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="de",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: prompts.append(dict(kwargs)) or "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "translated",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "succeeded"
    assert "ДОМЕН ПЕРЕВОДА" in prompts[0]["source_text"]
    assert "КОНТЕКСТ ДОКУМЕНТА" in prompts[0]["source_text"]
    assert "Great Tribulation -> die grosse Trubsal" in prompts[0]["source_text"]


def test_run_document_processing_appends_block_segment_focus_to_generation_prompt():
    runtime = _build_runtime_capture()
    generated_calls = []

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {
                "segment_id": "seg_0002",
                "target_text": "block",
                "context_before": "",
                "context_after": "",
                "target_chars": 5,
                "context_chars": 0,
            }
        ],
        document_segments=[
            SimpleNamespace(segment_id="seg_0001", ordinal=1, level=1, structural_role="chapter", title="Chapter 1"),
            SimpleNamespace(segment_id="seg_0002", ordinal=2, level=1, structural_role="chapter", title="Chapter 2"),
            SimpleNamespace(segment_id="seg_0003", ordinal=3, level=1, structural_role="chapter", title="Chapter 3"),
        ],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4-mini",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="de",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: generated_calls.append(dict(kwargs)) or "translated",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "succeeded"
    assert "ТЕКУЩИЙ БЛОК ДОКУМЕНТА" in generated_calls[0]["system_prompt"]
    assert "Текущий сегмент: #2 | L1 | chapter | Chapter 2" in generated_calls[0]["system_prompt"]
    assert "Предыдущий сегмент: Chapter 1" in generated_calls[0]["system_prompt"]
    assert "Следующий сегмент: Chapter 3" in generated_calls[0]["system_prompt"]


def test_run_document_processing_appends_previous_completed_segment_summary_to_next_segment_prompt():
    runtime = _build_runtime_capture()
    generated_calls = []

    def generate_markdown_block(**kwargs):
        generated_calls.append(dict(kwargs))
        if kwargs["target_text"] == "chapter-one-block":
            return "Translated chapter one summary sentence. Additional translated detail for continuity."
        return "Translated chapter two body."

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {
                "segment_id": "seg_0001",
                "target_text": "chapter-one-block",
                "context_before": "",
                "context_after": "",
                "target_chars": 17,
                "context_chars": 0,
            },
            {
                "segment_id": "seg_0002",
                "target_text": "chapter-two-block",
                "context_before": "",
                "context_after": "",
                "target_chars": 17,
                "context_chars": 0,
            },
        ],
        document_segments=[
            SimpleNamespace(segment_id="seg_0001", ordinal=1, level=1, structural_role="chapter", title="Chapter 1"),
            SimpleNamespace(segment_id="seg_0002", ordinal=2, level=1, structural_role="chapter", title="Chapter 2"),
        ],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4-mini",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="de",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "succeeded"
    assert len(generated_calls) == 2
    assert "СВОДКА ПРЕДЫДУЩЕГО ЗАВЕРШЁННОГО СЕГМЕНТА" in generated_calls[1]["system_prompt"]
    assert "Сегмент: Chapter 1" in generated_calls[1]["system_prompt"]
    assert "Translated chapter one summary sentence." in generated_calls[1]["system_prompt"]


def test_run_document_processing_routes_toc_dominant_translate_block_through_toc_prompt_variant_and_retries():
    runtime = _build_runtime_capture()
    prompts = []
    generated_calls = []
    events, log_event = _capture_log_events()

    def generate_markdown_block(**kwargs):
        generated_calls.append(dict(kwargs))
        if len(generated_calls) == 1:
            return "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9"
        return "Содержание\n\nВведение ........ 1\n\nЗаключение ........ 9"

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "job_kind": "passthrough",
            "target_text": "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9",
            "target_text_with_markers": "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9",
            "paragraph_ids": ["p0000", "p0001", "p0002"],
            "structural_roles": ["toc_header", "toc_entry", "toc_entry"],
            "toc_dominant": True,
            "toc_paragraph_count": 3,
            "paragraph_count": 3,
            "context_before": "",
            "context_after": "",
            "target_chars": 58,
            "context_chars": 0,
        }],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={"translation_domain_default": "theology", "translation_domain_instructions": "ДОМЕН ПЕРЕВОДА: богословие / эсхатология."},
        model="gpt-5.4-mini",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: prompts.append(dict(kwargs)) or f"system:{kwargs.get('prompt_variant', 'default')}",
        log_event=log_event,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "succeeded"
    assert prompts[0]["prompt_variant"] == "toc_translate"
    assert prompts[0]["translation_domain"] == "theology"
    assert len(generated_calls) == 2
    assert runtime["state"]["latest_markdown"].startswith("Содержание")
    info_events = [event for event in events if event["level"] == logging.INFO]
    assert any(event["event_id"] == "toc_prompt_routing_selected" for event in info_events)
    warning_events = [event for event in events if event["level"] == logging.WARNING]
    assert any(event["event_id"] == "toc_validation_rejected" for event in warning_events)


def test_run_document_processing_uses_hardened_toc_retry_prompt_on_final_attempt():
    runtime = _build_runtime_capture()
    generated_calls = []

    def generate_markdown_block(**kwargs):
        generated_calls.append(dict(kwargs))
        if len(generated_calls) < 3:
            return "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9"
        return "Содержание\n\nВведение ........ 1\n\nЗаключение ........ 9"

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "job_kind": "passthrough",
            "target_text": "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9",
            "target_text_with_markers": "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9",
            "paragraph_ids": ["p0000", "p0001", "p0002"],
            "structural_roles": ["toc_header", "toc_entry", "toc_entry"],
            "toc_dominant": True,
            "toc_paragraph_count": 3,
            "paragraph_count": 3,
            "context_before": "",
            "context_after": "",
            "target_chars": 58,
            "context_chars": 0,
        }],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4-mini",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: f"system:{kwargs.get('prompt_variant', 'default')}",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=generate_markdown_block,
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "succeeded"
    assert len(generated_calls) == 3
    assert "TOC retry hardening." in generated_calls[2]["system_prompt"]


def test_run_document_processing_routes_mixed_toc_dominant_translate_block_through_toc_prompt_variant():
    runtime = _build_runtime_capture()
    prompts = []

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "job_kind": "llm",
            "target_text": "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9\n\nNote on sources",
            "target_text_with_markers": "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9\n\nNote on sources",
            "paragraph_ids": ["p0000", "p0001", "p0002", "p0003"],
            "structural_roles": ["toc_header", "toc_entry", "toc_entry", "body"],
            "toc_dominant": True,
            "toc_paragraph_count": 3,
            "paragraph_count": 4,
            "context_before": "",
            "context_after": "",
            "target_chars": 75,
            "context_chars": 0,
        }],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4-mini",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: prompts.append(dict(kwargs)) or f"system:{kwargs.get('prompt_variant', 'default')}",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Содержание\n\nВведение ........ 1\n\nЗаключение ........ 9\n\nПримечание к источникам",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "succeeded"
    assert prompts[0]["prompt_variant"] == "toc_translate"
    assert "Примечание к источникам" in runtime["state"]["latest_markdown"]


def test_run_document_processing_fails_with_dedicated_toc_error_after_retry_budget_exhausted():
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "job_kind": "passthrough",
            "target_text": "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9",
            "target_text_with_markers": "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9",
            "paragraph_ids": ["p0000", "p0001", "p0002"],
            "structural_roles": ["toc_header", "toc_entry", "toc_entry"],
            "toc_dominant": True,
            "toc_paragraph_count": 3,
            "paragraph_count": 3,
            "context_before": "",
            "context_after": "",
            "target_chars": 58,
            "context_chars": 0,
        }],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4-mini",
        max_retries=1,
        processing_operation="translate",
        source_language="en",
        target_language="ru",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **kwargs: "system",
        log_event=log_event,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Contents\n\nIntroduction ........ 1\n\nConclusion ........ 9",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: b"final-docx",
    )

    assert result == "failed"
    assert "Ошибка обработки блока оглавления" in runtime["state"]["last_error"]
    assert runtime["activity"][-1] == "Блок 1: отклонён TOC validation после исчерпания retry budget."
    warning_events = [event for event in events if event["level"] == logging.WARNING]
    terminal_event = next(event for event in warning_events if event["event_id"] == "toc_validation_failed_terminal")
    assert terminal_event["context"]["retry_attempt"] == 2
    assert terminal_event["context"]["structural_roles"] == ["toc_header", "toc_entry", "toc_entry"]
    assert terminal_event["context"]["toc_paragraph_count"] == 3
    assert terminal_event["context"]["paragraph_count"] == 3


def test_run_document_processing_applies_semantic_output_normalization_before_image_reinsertion():
    runtime = _build_runtime_capture()
    call_order = []
    image_assets = [AssetStub("img_001")]

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[ParagraphStub()],
        image_assets=image_assets,
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Обработанный блок",
        process_document_images=lambda **kwargs: image_assets,
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: call_order.append("convert") or b"docx-bytes",
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: call_order.append("preserve") or docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: call_order.append("reinsert") or docx_bytes,
    )

    assert result == "succeeded"
    assert call_order == ["convert", "preserve", "reinsert"]


def test_run_document_processing_surfaces_formatting_diagnostics_artifacts(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    monkeypatch.setattr(document_pipeline, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)

    def preserve_with_artifact(docx_bytes, paragraphs, generated_paragraph_registry=None):
        diagnostics_dir.mkdir(parents=True, exist_ok=True)
        (diagnostics_dir / "preserve_001.json").write_text(
            json.dumps(
                {
                    "stage": "preserve",
                    "source_count": 5,
                    "target_count": 4,
                    "mapped_count": 4,
                    "unmapped_source_ids": ["p0004"],
                    "unmapped_target_indexes": [],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        return docx_bytes

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[ParagraphStub()],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Обработанный блок",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: b"docx-bytes",
        preserve_source_paragraph_properties=preserve_with_artifact,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
    )

    assert result == "succeeded"
    assert runtime["activity"][-2] == "Сборка DOCX завершена; сохранена служебная диагностика форматирования."
    assert runtime["state"]["latest_result_notice"] == {
        "level": "info",
        "message": (
            "DOCX собран. Дополнительное восстановление форматирования было частично пропущено, "
            "потому что точное сопоставление абзацев нашлось не везде. Это нормально, когда модель объединяет, "
            "делит или переформулирует абзацы. Совпадение найдено для 4 из 5 исходных абзацев; "
            "без точного соответствия осталось 1."
        ),
    }
    assert all(entry["status"] != "INFO" for entry in runtime["log"])


def test_run_document_processing_warns_user_only_for_conflicting_formatting_diagnostics(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    monkeypatch.setattr(document_pipeline, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)

    def preserve_with_conflict_artifact(docx_bytes, paragraphs, generated_paragraph_registry=None):
        diagnostics_dir.mkdir(parents=True, exist_ok=True)
        (diagnostics_dir / "preserve_001.json").write_text(
            json.dumps(
                {
                    "stage": "preserve",
                    "source_count": 5,
                    "target_count": 5,
                    "mapped_count": 5,
                    "unmapped_source_ids": [],
                    "unmapped_target_indexes": [],
                    "caption_heading_conflicts": [
                        {"paragraph_id": "p0002", "target_heading_level": 2}
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        return docx_bytes

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[ParagraphStub()],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Обработанный блок",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: b"docx-bytes",
        preserve_source_paragraph_properties=preserve_with_conflict_artifact,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
    )

    assert result == "succeeded"
    assert runtime["activity"][-2] == "Сборка DOCX завершена; найдены места, где форматирование стоит проверить вручную."
    assert runtime["log"][-2]["status"] == "WARN"
    assert "спорные места форматирования" in runtime["log"][-2]["details"]
    assert "Конфликтов подписи/заголовка: 1." in runtime["log"][-2]["details"]
    assert runtime["state"].get("latest_result_notice") is None


def test_run_document_processing_warns_and_delivers_on_strict_unmapped_source_quality_gate(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [str(diagnostics_dir / "preserve_001.json")])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    def preserve_with_unmapped_artifact(docx_bytes, paragraphs, generated_paragraph_registry=None):
        diagnostics_dir.mkdir(parents=True, exist_ok=True)
        (diagnostics_dir / "preserve_001.json").write_text(
            json.dumps(
                {
                    "stage": "restore",
                    "source_count": 5,
                    "target_count": 4,
                    "mapped_count": 4,
                    "unmapped_source_ids": ["p0004"],
                    "unmapped_target_indexes": [],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        return docx_bytes

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[ParagraphStub()],
        image_assets=[],
        image_mode="safe",
        app_config={"translation_output_quality_gate_policy": "strict"},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="translate",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Обработанный блок",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: b"docx-bytes",
        preserve_source_paragraph_properties=preserve_with_unmapped_artifact,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
    )

    # spec 018: strict unmapped-source residue is review-grade — the document is usable,
    # so it is DELIVERED with a yellow warn notice, not blocked at a red failure.
    assert result == "succeeded"
    assert runtime["state"]["latest_result_notice"]["level"] == "warning"
    notice_message = runtime["state"]["latest_result_notice"]["message"]
    assert "translation_quality_gate_failed" not in notice_message
    assert "заблокирован" not in notice_message
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "warn"
    assert payload["gate_reasons"] == ["unmapped_source_paragraphs_present"]
    assert payload["bullet_heading_count"] == 0
    assert payload["toc_body_concat_detected"] is False


def test_run_document_processing_surfaces_advisory_quality_notice_on_mapping_drift(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [str(diagnostics_dir / "preserve_001.json")])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    def preserve_with_unmapped_artifact(docx_bytes, paragraphs, generated_paragraph_registry=None):
        diagnostics_dir.mkdir(parents=True, exist_ok=True)
        (diagnostics_dir / "preserve_001.json").write_text(
            json.dumps(
                {
                    "stage": "restore",
                    "source_count": 50,
                    "target_count": 48,
                    "mapped_count": 48,
                    "unmapped_source_ids": ["p0048", "p0049"],
                    "unmapped_target_indexes": [],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        return docx_bytes

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{"target_text": "block", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0}],
        source_paragraphs=[ParagraphStub()],
        image_assets=[],
        image_mode="safe",
        app_config={"translation_output_quality_gate_policy": "advisory"},
        model="gpt-5.4",
        max_retries=1,
        processing_operation="translate",
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Обработанный блок",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=lambda markdown_text: b"docx-bytes",
        preserve_source_paragraph_properties=preserve_with_unmapped_artifact,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
    )

    assert result == "succeeded"
    # Policy-independent discrepancy emission (GATE_TRUSTWORTHINESS Task B):
    # advisory now surfaces the unmapped review-item DATA (was 0), so the notice
    # switches to the actionable review message while the verdict severity
    # (warn + advisory-threshold reason) stays policy-scaled.
    assert runtime["state"]["latest_result_notice"] == {
        "level": "warning",
        "message": (
            "Перевод завершён. Документ готов к использованию, но требует ручной "
            "проверки оформления: 2 абзаца с замечаниями. "
            "Подробности — в отчёте проверки (formatting_review.txt)."
        ),
    }
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "warn"
    assert payload["gate_reasons"] == ["unmapped_source_paragraphs_above_advisory_threshold"]
    assert payload["formatting_review_required_count"] == 2
    assert payload["formatting_review_items"] == [
        {
            "reason": "unmapped_source_paragraphs_review_required",
            "label": "Абзацы без явного соответствия оригиналу",
            "count": 2,
            "severity": "review",
        }
    ]


def test_run_document_processing_keeps_false_fragment_cleanup_display_only_after_quality_gate_decoupling(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "strict", "enable_paragraph_markers": True},
        processing_operation="translate",
        generate_markdown_block=lambda **kwargs: (
            "Христос предупреждает нас\n\n"
            "## (Матфея 24:36)\n\n"
            "что день неизвестен.\n\n"
            "Это обсуждение подводит к вопросу\n\n"
            "## Спутники? Ракеты?)\n\n"
            "который дальше раскрывается в тексте."
        ),
    )

    assert result == "succeeded"
    assert "## (Матфея 24:36)" not in runtime["state"]["latest_markdown"]
    assert "## Спутники? Ракеты?)" not in runtime["state"]["latest_markdown"]
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "warn"
    assert payload["gate_reasons"] == ["false_fragment_headings_review_required"]
    assert payload["false_fragment_heading_count"] == 2
    assert payload["scripture_reference_heading_count"] == 1


def test_run_document_processing_quality_report_uses_pre_display_gate_input_for_sentence_split_case(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "strict"},
        processing_operation="translate",
        jobs=[{
            "target_text": "Исходный блок",
            "target_text_with_markers": "[[DOCX_PARA_p1]]\n[[DOCX_PARA_p2]]\n[[DOCX_PARA_p3]]\n[[DOCX_PARA_p4]]",
            "paragraph_ids": ["p1", "p2", "p3", "p4"],
            "context_before": "",
            "context_after": "",
            "target_chars": 13,
            "context_chars": 0,
        }],
        source_paragraphs=[
            SimpleNamespace(paragraph_id="p1", source_index=0, role="body", structural_role="body", heading_level=None, list_kind=None, boundary_source="raw", boundary_confidence="explicit"),
            SimpleNamespace(paragraph_id="p2", source_index=1, role="heading", structural_role="heading", heading_level=2, list_kind=None, boundary_source="raw", boundary_confidence="explicit"),
            SimpleNamespace(paragraph_id="p3", source_index=2, role="body", structural_role="body", heading_level=None, list_kind=None, boundary_source="raw", boundary_confidence="explicit"),
            SimpleNamespace(paragraph_id="p4", source_index=3, role="body", structural_role="body", heading_level=None, list_kind=None, boundary_source="raw", boundary_confidence="explicit"),
        ],
        generate_markdown_block=lambda **kwargs: (
            "Пожалуй, главный вывод состоит в том, что каждое поколение христиан должно приготовиться к возможности пережить\n\n"
            "## Великая скорбь\n\n"
            ".\n\n"
            "Практические шаги начинаются здесь."
        ),
    )

    assert result == "succeeded"
    assert "## Великая скорбь\n\n." not in runtime["state"]["latest_markdown"]
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "warn"
    assert payload["gate_reasons"] == ["false_fragment_headings_review_required"]
    assert payload["false_fragment_heading_count"] == 1


def test_run_document_processing_applies_residual_bullet_cleanup_before_display_hygiene_gating(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "strict"},
        processing_operation="translate",
        generate_markdown_block=lambda **kwargs: (
            "Посттрибулационисты считают, что Иисус придёт в конце ● скорби.\n\n"
            "● собирают армию в 200 миллионов солдат.\n\n"
            "- Соединённые Штаты формируют мировую ● культуру и политику?"
        ),
    )

    assert result == "succeeded"
    assert "●" not in runtime["state"]["latest_markdown"]
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "pass"
    assert payload["gate_reasons"] == []
    assert payload["residual_bullet_glyph_count"] == 0
    assert payload["residual_bullet_glyph_gate_source"] == "legacy_markdown"
    assert payload["raw_residual_bullet_glyph_count"] == 3


def test_run_document_processing_normalizes_list_fragment_regressions_before_quality_gate(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "strict"},
        processing_operation="translate",
        generate_markdown_block=lambda **kwargs: (
            "Поразительно, но все петли следуют одной и той же схеме: 1.\n\n"
            "Духовные существа восстают против Бога.\n\n"
            "2. Бог судит их за грех.\n\n"
            "3. Бог спасает остаток верных."
        ),
    )

    report_files = list(quality_dir.glob("*.json"))
    if report_files:
        payload = json.loads(report_files[0].read_text(encoding="utf-8"))
        assert payload["quality_status"] == "pass", json.dumps(payload, ensure_ascii=False, indent=2)

    assert result == "succeeded"
    assert "схеме: 1." not in runtime["state"]["latest_markdown"]
    assert "1. Духовные существа восстают против Бога." in runtime["state"]["latest_markdown"]
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "pass"
    assert payload["list_fragment_regression_count"] == 0


def test_run_document_processing_uses_runtime_normalized_markdown_for_docx_build(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    captured = {}
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "strict"},
        processing_operation="translate",
        generate_markdown_block=lambda **kwargs: (
            "Поразительно, но все петли следуют одной и той же схеме: 1.\n\n"
            "Духовные существа восстают против Бога.\n\n"
            "2. Бог судит их за грех."
        ),
        convert_markdown_to_docx_bytes=lambda markdown_text: captured.setdefault("markdown", markdown_text) or markdown_text.encode("utf-8"),
    )

    assert result == "succeeded"
    assert captured["markdown"] == runtime["state"]["latest_markdown"]
    assert "схеме: 1." not in captured["markdown"]
    assert "1. Духовные существа восстают против Бога." in captured["markdown"]


def test_run_document_processing_normalizes_mixed_script_before_quality_gate(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "strict"},
        processing_operation="translate",
        generate_markdown_block=lambda **kwargs: (
            "Прежде чем суперразумa догонит квантовый скачок.\n\n"
            "Это просто тестовая cтрока с латинскими символами."
        ),
    )

    assert result == "succeeded"
    assert "суперразумa" not in runtime["state"]["latest_markdown"]
    assert "суперразума" in runtime["state"]["latest_markdown"]
    assert "cтрока" not in runtime["state"]["latest_markdown"]
    assert "строка" in runtime["state"]["latest_markdown"]
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "pass"
    assert payload["mixed_script_term_count"] == 0


def test_run_document_processing_flags_untranslated_structural_heading_for_review(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "strict", "enable_paragraph_markers": True},
        processing_operation="translate",
        source_paragraphs=[
            ParagraphUnit(
                text="The Competitive Society",
                role="heading",
                structural_role="heading",
                heading_level=2,
                paragraph_id="p0001",
            )
        ],
        jobs=[
            {
                "target_text": "The Competitive Society",
                "target_text_with_markers": "[[DOCX_PARA_p0001]]\nThe Competitive Society",
                "context_before": "",
                "context_after": "",
                "target_chars": 23,
                "context_chars": 0,
                "paragraph_ids": ["p0001"],
            }
        ],
        generate_markdown_block=lambda **kwargs: "## The Competitive Society",
        write_ui_result_artifacts=lambda **kwargs: {
            "markdown_path": "/tmp/final.result.md",
            "docx_path": "/tmp/final.result.docx",
            "formatting_review_path": "/tmp/final.result.formatting_review.txt",
        },
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_result_notice"] == {
        "level": "warning",
        "message": (
            "Перевод завершён. Документ готов к использованию, но требует ручной "
            "проверки оформления: 1 абзац с замечаниями. "
            "Подробности — в отчёте проверки (formatting_review.txt)."
        ),
    }
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "warn"
    assert payload["gate_reasons"] == ["untranslated_structural_text_review_required"]
    assert payload["untranslated_structural_text_count"] == 1
    assert payload["formatting_review_required_count"] == 1
    assert payload["formatting_review_items"] == [
        {
            "reason": "untranslated_structural_text_review_required",
            "label": "Структурный элемент остался на исходном языке",
            "count": 1,
            "severity": "review",
            "sample": {
                "line": 1,
                # FR-004: the leaked markdown heading marker is stripped from the anchor.
                "text": "The Competitive Society",
                "reason": "untranslated_structural_text",
                "role": "heading",
                "structural_role": "heading",
                "paragraph_id": "p0001",
                "char_count": 23,
            },
        }
    ]


def test_run_document_processing_fails_large_untranslated_body_text(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)
    untranslated_body = (
        "This framework has been tested using years of quantitative data collected about how biomass "
        "flows through natural ecosystems. Natural ecosystems are large complex flow networks that "
        "show how resilience and efficiency interact across many scales. "
    ) * 12

    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "strict", "enable_paragraph_markers": True},
        processing_operation="translate",
        source_paragraphs=[
            ParagraphUnit(
                text="Исходный русский абзац для проверки.",
                role="body",
                structural_role="body",
                paragraph_id="p0001",
            )
        ],
        jobs=[
            {
                "target_text": "Исходный русский абзац для проверки.",
                "target_text_with_markers": "[[DOCX_PARA_p0001]]\nИсходный русский абзац для проверки.",
                "context_before": "",
                "context_after": "",
                "target_chars": 36,
                "context_chars": 0,
                "paragraph_ids": ["p0001"],
            }
        ],
        generate_markdown_block=lambda **kwargs: untranslated_body,
        write_ui_result_artifacts=lambda **kwargs: {
            "markdown_path": "/tmp/final.result.md",
            "docx_path": "/tmp/final.result.docx",
            "formatting_review_path": "/tmp/final.result.formatting_review.txt",
        },
    )

    assert result == "failed"
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "fail"
    assert payload["gate_reasons"] == ["untranslated_body_text_above_threshold"]
    assert payload["untranslated_body_text_count"] == 1
    assert payload["untranslated_body_text_chars"] > 2000
    assert payload["formatting_review_items"][0]["severity"] == "fix"


def test_build_translation_quality_report_flags_bullet_marker_headings_in_strict_translate_mode():
    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown="## ●\n\nПереведённый абзац",
        formatting_diagnostics_artifacts=[],
    )

    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == ["bullet_marker_headings_review_required"]
    assert report["bullet_heading_count"] == 1
    assert report["bullet_heading_gate_source"] == "legacy_markdown"
    assert report["bullet_heading_classification"] == "markdown_gate"
    assert report["raw_bullet_heading_count"] == 1
    assert report["formatting_review_items"] == [
        {
            "reason": "bullet_marker_headings_review_required",
            "label": "Маркер списка попал в заголовок",
            "count": 1,
            "severity": "fix",
            "sample": {
                "line": 1,
                # FR-004/006: "## ●" strips to a lone bullet glyph — not a locatable
                # anchor, so it is flagged for aggregation instead of shown.
                "text": "●",
                "reason": "bullet_marker_heading",
                "anchor_usable": False,
            },
        }
    ]


def test_build_translation_quality_report_counts_capped_legacy_hygiene_samples(monkeypatch):
    samples = [
        SimpleNamespace(line=index + 1, text=f"## ● {index}", reason="bullet_marker_heading")
        for index in range(10)
    ]
    monkeypatch.setattr(
        document_pipeline_late_phases,
        "collect_bullet_heading_samples",
        lambda markdown_text: list(samples),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
            paragraph_count=2000,
        ),
        final_markdown="Bullet headings",
        formatting_diagnostics_artifacts=[],
    )

    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == ["bullet_marker_headings_review_required"]
    assert report["bullet_heading_count"] == 10
    assert report["formatting_review_required_count"] == 10
    review_items = cast(list[dict[str, object]], report["formatting_review_items"])
    assert len(review_items) == 8
    assert review_items[0]["aggregate_count"] == 10
    assert [item["count"] for item in review_items] == [0] * 8


def test_build_translation_quality_report_flags_untranslated_structural_heading_but_not_proper_name():
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown="## The Competitive Society\n\n## Terra\n\nПереведённый текст.",
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="## The Competitive Society",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="heading",
                structural_role="heading",
                heading_level=2,
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="## Terra",
                block_index=1,
                paragraph_id="p2",
                source_index=1,
                role="heading",
                structural_role="heading",
                heading_level=2,
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="Переведённый текст.",
                block_index=1,
                paragraph_id="p3",
                source_index=2,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=assembly_result.final_markdown,
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == ["untranslated_structural_text_review_required"]
    assert report["untranslated_structural_text_count"] == 1
    assert report["untranslated_structural_text_samples"] == [
        {
            "line": 1,
            "text": "## The Competitive Society",
            "reason": "untranslated_structural_text",
            "role": "heading",
            "structural_role": "heading",
            "paragraph_id": "p1",
            "char_count": 23,
        }
    ]
    assert report["formatting_review_items"] == [
        {
            "reason": "untranslated_structural_text_review_required",
            "label": "Структурный элемент остался на исходном языке",
            "count": 1,
            "severity": "review",
            "sample": {
                "line": 1,
                # FR-004: the leaked markdown heading marker is stripped from the anchor.
                "text": "The Competitive Society",
                "reason": "untranslated_structural_text",
                "role": "heading",
                "structural_role": "heading",
                "paragraph_id": "p1",
                "char_count": 23,
            },
        }
    ]


def test_build_translation_quality_report_fails_large_untranslated_body_text():
    untranslated_body = (
        "This framework has been tested using years of quantitative data collected about how biomass "
        "flows through natural ecosystems. Natural ecosystems are large complex flow networks that "
        "show how resilience and efficiency interact across many scales. "
    ) * 12
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown=untranslated_body,
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text=untranslated_body,
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=assembly_result.final_markdown,
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    # spec 018 counter-proof: wholesale-untranslated BODY above the catastrophic
    # threshold is genuinely non-deliverable and STAYS a hard document-level fail.
    assert report["quality_status"] == "fail"
    assert report["gate_reasons"] == ["untranslated_body_text_above_threshold"]
    assert report["untranslated_body_text_count"] == 1
    assert report["untranslated_body_text_chars"] > 2000
    assert report["untranslated_body_text_ratio"] == 1.0
    assert report["formatting_review_items"][0]["severity"] == "fix"


def test_resolve_document_delivery_verdict_downgrades_review_grade_fail_to_warn():
    # spec 018 Level 1: every review-grade fail-driver (role_loss / heading_demotion /
    # false_fragment / list_fragment / unmapped-source / toc_body_concat / mixed_script)
    # resolves to a delivered ``warn`` — NOT a blocking document-level fail.
    review_grade_reasons = [
        "role_loss_above_manual_review_threshold",
        "heading_demotion_above_manual_review_threshold",
        "false_fragment_headings_present",
        "list_fragment_regressions_present",
        "unmapped_source_paragraphs_present",
        "toc_body_concatenation_detected",
        "mixed_script_terms_present",
    ]
    for reason in review_grade_reasons:
        assert (
            document_pipeline_late_phases._resolve_document_delivery_verdict(
                quality_status="fail",
                gate_reasons=[reason],
            )
            == "warn"
        ), reason


def test_resolve_document_delivery_verdict_keeps_untranslated_body_fatal():
    # spec 018 Level 1: the sole genuinely-fatal reason stays a hard fail, even when it
    # is mixed with review-grade reasons.
    assert (
        document_pipeline_late_phases._resolve_document_delivery_verdict(
            quality_status="fail",
            gate_reasons=["untranslated_body_text_above_threshold"],
        )
        == "fail"
    )
    assert (
        document_pipeline_late_phases._resolve_document_delivery_verdict(
            quality_status="fail",
            gate_reasons=[
                "role_loss_above_manual_review_threshold",
                "untranslated_body_text_above_threshold",
            ],
        )
        == "fail"
    )
    # pass / warn verdicts are returned unchanged.
    assert (
        document_pipeline_late_phases._resolve_document_delivery_verdict(
            quality_status="pass", gate_reasons=[]
        )
        == "pass"
    )
    assert (
        document_pipeline_late_phases._resolve_document_delivery_verdict(
            quality_status="warn", gate_reasons=["false_fragment_headings_review_required"]
        )
        == "warn"
    )


def test_build_quality_warn_notice_message_has_no_internal_tokens():
    # spec 018 Level 2: the yellow warn notice is human-readable Russian with NO internal
    # tokens and NONE of the fatal-path wording.
    message = document_pipeline_late_phases._build_quality_warn_notice_message(
        {"formatting_review_required_count": 216}
    )
    assert message == (
        "Перевод завершён. Документ готов к использованию, но требует ручной "
        "проверки оформления: 216 абзацев с замечаниями. "
        "Подробности — в отчёте проверки (formatting_review.txt)."
    )
    for forbidden in (
        "translation_quality_gate_failed",
        "quality_status",
        "gate_reasons",
        "заблокирован",
        "критическая",
        "role_loss",
    ):
        assert forbidden not in message.lower()
    # Zero-count degrade path still carries the deliverable framing.
    zero = document_pipeline_late_phases._build_quality_warn_notice_message(
        {"formatting_review_required_count": 0}
    )
    assert zero.startswith("Перевод завершён. Документ готов к использованию")
    assert "formatting_review.txt" in zero


def test_resistance_report_offline_replay_reclassifies_fail_to_warn():
    # spec 018 offline replay: the real EN->RU RESISTANCE run hard-FAILED with 5
    # review-grade reasons (0 genuine defects, DOCX fully usable). Re-running the
    # document-level verdict over that saved report must now yield ``warn`` with all
    # 5 reasons still present as review items. Loads the saved artifact when it is
    # available; otherwise replays its verdict-relevant shape inline (CI has no .run/).
    from docxaicorrector.core.constants import RUN_DIR

    resistance_reasons = [
        "role_loss_above_manual_review_threshold",
        "heading_demotion_above_manual_review_threshold",
        "false_fragment_headings_review_required",
        "list_fragment_regressions_review_required",
        "untranslated_structural_text_review_required",
    ]
    saved_reports = sorted(
        Path(RUN_DIR).glob(
            "quality_reports/RESISTANCE_FACTORS_AND_SPECIAL_FORCES_AREAS_UKRAINE.docx_*.json"
        )
    )
    if saved_reports:
        saved = json.loads(saved_reports[-1].read_text(encoding="utf-8"))
        assert saved["quality_status"] == "fail"
        original_status = str(saved["quality_status"])
        gate_reasons = list(saved["gate_reasons"])
        assert "untranslated_body_text_above_threshold" not in gate_reasons
    else:
        original_status = "fail"
        gate_reasons = list(resistance_reasons)

    reclassified = document_pipeline_late_phases._resolve_document_delivery_verdict(
        quality_status=original_status,
        gate_reasons=gate_reasons,
    )
    assert reclassified == "warn"
    # All review reasons are preserved verbatim — only the verdict severity moved.
    for reason in resistance_reasons:
        assert reason in gate_reasons


@pytest.mark.parametrize(
    "markdown_text",
    [
        "Заключение ........ 29 Введение",
        "Заключение……29 Введение",
        "Заключение··29 Введение",
        "Заключение  29 Введение",
    ],
)
def test_build_translation_quality_report_detects_toc_body_concat_across_leader_variants(markdown_text):
    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=markdown_text,
        formatting_diagnostics_artifacts=[],
    )

    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == ["toc_body_concatenation_review_required"]
    assert report["toc_body_concat_detected"] is True
    assert report["formatting_review_items"][0]["reason"] == "toc_body_concatenation_review_required"
    assert report["formatting_review_items"][0]["severity"] == "fix"


def test_build_translation_quality_report_exposes_quality_gate_audit_classifications():
    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown="## ●\n\nЗаключение ........ 29 Введение",
        formatting_diagnostics_artifacts=[],
    )

    audit = report["quality_gate_audit_classifications"]
    assert audit["bullet_heading"]["verdict"] == "unit_aware"
    assert audit["bullet_heading"]["severity_model"] == "legacy_hygiene_fix_review_threshold"
    assert audit["toc_body_concat"]["verdict"] == "tolerant"
    assert audit["toc_body_concat"]["severity_model"] == "structure_evidence_required_else_review"
    assert audit["mixed_script_term"]["verdict"] == "tolerant"
    assert audit["heading_body_concat_detected"]["verdict"] == "tolerant"
    assert audit["inline_page_furniture_leakage"]["verdict"] == "unit_aware_after_structural_label_exemption"


def test_build_translation_quality_report_keeps_source_backed_scripture_heading_out_of_false_fragment_gate():
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown="## (Матфея 24:36)\n\nХристос вернётся как вор в ночи.",
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="## (Матфея 24:36)",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="heading",
                structural_role="heading",
                heading_level=2,
                boundary_source="source_style",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="Христос вернётся как вор в ночи.",
                block_index=1,
                paragraph_id="p2",
                source_index=1,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=assembly_result.final_markdown,
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    assert report["quality_status"] == "pass"
    assert report["false_fragment_heading_count"] == 0
    assert report["scripture_reference_heading_count"] == 0
    assert report["raw_false_fragment_heading_count"] == 1
    assert report["quality_gate_audit_classifications"]["scripture_reference_heading"]["verdict"] == "tolerant"


def test_normalize_final_markdown_for_runtime_display_splits_placeholder_from_chapter_heading():
    normalized = document_pipeline_late_phases._normalize_final_markdown_for_runtime_display(
        "This page intentionally left blank Chapter Nine STRATEGIES FOR NGO S"
    )

    assert normalized == "This page intentionally left blank\n\nChapter Nine STRATEGIES FOR NGO S"


def test_restore_image_heading_lines_from_registry_splits_placeholder_and_heading_markers():
    restored = document_pipeline_late_phases._restore_image_heading_lines_from_registry(
        "[[DOCX_IMAGE_img_009]] Глава десятая ИСТИНА И ПОСЛЕДСТВИЯ\n\nBody",
        [
            {"paragraph_id": "p0076", "text": "## Глава десятая"},
            {"paragraph_id": "p0077", "text": "# ИСТИНА И ПОСЛЕДСТВИЯ"},
        ],
    )

    assert restored == "[[DOCX_IMAGE_img_009]]\n\n## Глава десятая\n# ИСТИНА И ПОСЛЕДСТВИЯ\n\nBody"


def test_restore_image_heading_lines_from_registry_leaves_unbacked_concat_unchanged():
    markdown = "[[DOCX_IMAGE_img_009]] Глава десятая ИСТИНА И ПОСЛЕДСТВИЯ"

    assert document_pipeline_late_phases._restore_image_heading_lines_from_registry(markdown, []) == markdown


def test_build_translation_quality_report_classifies_placeholder_heading_concat_as_display_hygiene_with_raw_observability():
    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown="This page intentionally left blank Chapter Nine STRATEGIES FOR NGO S",
        formatting_diagnostics_artifacts=[],
    )

    assert report["quality_status"] == "pass"
    assert report["gate_reasons"] == []
    assert report["page_placeholder_heading_concat_count"] == 0
    assert report["page_placeholder_heading_concat_source"] == "legacy_markdown"
    assert report["page_placeholder_heading_concat_classification"] == "display_hygiene"
    assert report["raw_page_placeholder_heading_concat_count"] == 1
    assert report["page_placeholder_heading_concat_samples"] == []
    raw_page_placeholder_heading_concat_samples = cast(
        list[dict[str, object]],
        report["raw_page_placeholder_heading_concat_samples"],
    )
    assert raw_page_placeholder_heading_concat_samples[0]["reason"] == "page_placeholder_heading_concat_markdown_present"


def test_normalize_final_markdown_for_quality_gate_preserves_placeholder_cleanup_as_non_gate_input():
    normalized = document_pipeline_late_phases._normalize_final_markdown_for_quality_gate(
        "This page intentionally left blank Chapter Nine STRATEGIES FOR NGO S"
    )

    assert normalized == "This page intentionally left blank Chapter Nine STRATEGIES FOR NGO S"


def test_normalize_final_markdown_for_display_hygiene_reporting_applies_residual_bullet_cleanup():
    normalized = document_pipeline_late_phases._normalize_final_markdown_for_display_hygiene_reporting(
        "Посттрибулационисты считают, что Иисус придёт в конце ● скорби."
    )

    assert normalized == "Посттрибулационисты считают, что Иисус придёт в конце скорби."


def test_normalize_final_markdown_for_runtime_display_applies_structure_compatibility_cleanup_only():
    normalized = document_pipeline_late_phases._normalize_final_markdown_for_runtime_display(
        "Наблюдайте внимательно.\n\n"
        "## (Матфея 24:36)\n\n"
        "Это продолжение абзаца.\n\n"
        "Поразительно, но все петли следуют одной и той же схеме: 1.\n\n"
        "Духовные существа восстают против Бога.\n\n"
        "2. Бог судит их за грех."
    )

    assert normalized == (
        "Наблюдайте внимательно.\n\n"
        "(Матфея 24:36)\n\n"
        "Это продолжение абзаца.\n\n"
        "Поразительно, но все петли следуют одной и той же схеме:\n\n"
        "1. Духовные существа восстают против Бога.\n\n"
        "2. Бог судит их за грех."
    )


def test_resolve_runtime_display_markdown_prefers_explicit_runtime_display_payload():
    resolved = document_pipeline_late_phases._resolve_runtime_display_markdown(
        docx_phase={
            "runtime_display_markdown": "display-cleaned",
        },
        fallback_markdown="raw-gate-input",
    )

    assert resolved == "display-cleaned"


def test_resolve_runtime_display_markdown_ignores_legacy_final_markdown_alias():
    resolved = document_pipeline_late_phases._resolve_runtime_display_markdown(
        docx_phase={"final_markdown": "legacy-display-cleaned"},
        fallback_markdown="This page intentionally left blank Chapter Nine STRATEGIES FOR NGO S",
    )

    assert resolved == "This page intentionally left blank\n\nChapter Nine STRATEGIES FOR NGO S"


def test_build_translation_quality_report_exposes_new_residual_quality_metrics_and_gate_reasons():
    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict", "translation_domain": "theology"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="theology",
        ),
        final_markdown=(
            "Является ли\n\n"
            "## начертание зверя\n\n"
            "на самом деле - квантовая технология?\n\n"
            "## (Матфея 24:36)\n\n"
            "- Сторонники мидтрибулационного взгляда считают, что христиане будут восхищены в середине\n"
            "- Великой скорби.\n\n"
            "Китай ... технологическими ● достижениями?\n\n"
            "## Суд над пятым печатью\n\n"
            "Создавайте кoinonia-сообщества и богословие imago Dei."
        ),
        formatting_diagnostics_artifacts=[],
    )

    # spec 018: review-grade fail-drivers (false_fragment / list_fragment / mixed_script)
    # are delivered as ``warn``, not a blocking ``fail`` — all gate_reasons preserved.
    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == [
        "false_fragment_headings_review_required",
        "list_fragment_regressions_present",
        "mixed_script_terms_review_required",
    ]
    assert report["bullet_heading_count"] == 0
    assert report["false_fragment_heading_count"] == 2
    assert report["scripture_reference_heading_count"] == 1
    assert report["residual_bullet_glyph_count"] == 0
    assert report["residual_bullet_glyph_classification"] == "display_hygiene"
    assert report["raw_residual_bullet_glyph_count"] == 1
    assert report["list_fragment_regression_count"] == 1
    mixed_script_term_count = report["mixed_script_term_count"]
    theology_style_issue_count = report["theology_style_deterministic_issue_count"]
    assert isinstance(mixed_script_term_count, int)
    assert isinstance(theology_style_issue_count, int)
    assert mixed_script_term_count >= 1
    assert theology_style_issue_count >= 2
    assert report["mixed_script_term_gate_source"] == "legacy_markdown"
    assert report["mixed_script_term_classification"] == "non_structural_hygiene"
    assert report["raw_mixed_script_term_count"] == mixed_script_term_count
    assert report["theology_style_deterministic_issue_source"] == "legacy_markdown"
    assert report["theology_style_deterministic_issue_classification"] == "domain_style_advisory"
    assert report["raw_theology_style_deterministic_issue_count"] == theology_style_issue_count
    assert report["translation_domain"] == "theology"
    assert report["worst_unmapped_source_count"] == 0
    assert report["suspicious_heading_repetition_count"] == 0


def test_build_translation_quality_report_measures_delivered_markdown_for_hygiene_metrics():
    # Spec 006 FR-002/003: the hygiene metrics describe the delivered artifact
    # (runtime_display_markdown), so a delivered markdown whose passes cleaned the
    # mixed-script / residual-glyph / page-placeholder residue reports the CLEANED
    # counts — mixed_script is NO LONGER set equal to raw.
    final_markdown = (
        "This page intentionally left blank Chapter Nine STRATEGIES FOR NGO S\n\n"
        "Создавайте кoinonia-сообщества и богословие imago Dei.\n\n"
        "Китай ... технологическими ● достижениями?"
    )
    runtime_display_markdown = (
        "This page intentionally left blank\n\n"
        "Chapter Nine STRATEGIES FOR NGO S\n\n"
        "Создавайте сообщества.\n\n"
        "Китай ... технологическими достижениями?"
    )
    context = SimpleNamespace(
        app_config={"translation_output_quality_gate_policy": "strict"},
        processing_operation="translate",
        uploaded_filename="report.docx",
        translation_domain="general",
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=context,
        final_markdown=final_markdown,
        formatting_diagnostics_artifacts=[],
        runtime_display_markdown=runtime_display_markdown,
    )

    # Delivered markdown is clean; raw (pre-display) baseline still counts the residue.
    assert report["mixed_script_term_count"] == 0
    assert report["raw_mixed_script_term_count"] == 1
    assert cast(int, report["mixed_script_term_count"]) < cast(int, report["raw_mixed_script_term_count"])
    assert report["page_placeholder_heading_concat_count"] == 0
    assert report["raw_page_placeholder_heading_concat_count"] == 1
    assert report["residual_bullet_glyph_count"] == 0
    assert report["raw_residual_bullet_glyph_count"] == 1

    # FR-006 degrade-safe: with no delivered markdown, behaviour is byte-identical to
    # today's — mixed_script falls back to the raw count (gated == raw).
    degraded = document_pipeline_late_phases._build_translation_quality_report(
        context=context,
        final_markdown=final_markdown,
        formatting_diagnostics_artifacts=[],
    )
    assert degraded["mixed_script_term_count"] == degraded["raw_mixed_script_term_count"] == 1
    assert degraded["page_placeholder_heading_concat_count"] == 0
    assert degraded["residual_bullet_glyph_count"] == 0


def test_build_translation_quality_report_delivered_markdown_leaves_structural_gates_source_aware():
    # Spec 006 FR-004: threading runtime_display_markdown must NOT repoint the
    # entry-based structural gates (specs 001/003). Passing the delivered markdown
    # leaves false_fragment / list_fragment gate sources and counts unchanged.
    final_markdown = "2. Goldman Sachs Annual Report, 2010.\n\n14. Forbes, 2017."
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown=final_markdown,
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="2. Goldman Sachs Annual Report, 2010.",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="14. Forbes, 2017.",
                block_index=1,
                paragraph_id="p2",
                source_index=1,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )
    context = SimpleNamespace(
        app_config={"translation_output_quality_gate_policy": "strict"},
        processing_operation="translate",
        uploaded_filename="report.docx",
        translation_domain="general",
    )

    baseline = document_pipeline_late_phases._build_translation_quality_report(
        context=context,
        final_markdown=final_markdown,
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )
    delivered = document_pipeline_late_phases._build_translation_quality_report(
        context=context,
        final_markdown=final_markdown,
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
        runtime_display_markdown=final_markdown,
    )

    assert delivered["list_fragment_regression_gate_source"] == "entry_assembly"
    assert delivered["false_fragment_heading_gate_source"] == baseline["false_fragment_heading_gate_source"]
    assert delivered["list_fragment_regression_gate_source"] == baseline["list_fragment_regression_gate_source"]
    assert delivered["list_fragment_regression_count"] == baseline["list_fragment_regression_count"]
    assert delivered["false_fragment_heading_count"] == baseline["false_fragment_heading_count"]


def test_build_translation_quality_report_prefers_entry_aware_false_heading_detection():
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown="## Введение\n\nТекст раздела.",
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="## Введение",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="heading",
                structural_role="heading",
                heading_level=2,
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="Текст раздела.",
                block_index=1,
                paragraph_id="p2",
                source_index=1,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown="## Введение\n\nТекст раздела.",
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    assert report["quality_status"] == "pass"
    assert report["false_fragment_heading_count"] == 0
    assert report["gate_reasons"] == []
    boundary_recovery = cast(dict[str, Any], report["boundary_recovery"])

    assert boundary_recovery["recovered_heading_entries"] == [
        {
            "paragraph_id": "p1",
            "source_index": 0,
            "role": "heading",
            "structural_role": "heading",
            "generated_heading_kind": None,
            "text": "## Введение",
        }
    ]


def test_build_translation_quality_report_reports_entry_aware_false_heading_sample():
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown="На людей, получивших начертание зверя и поклонявшихся его образу.",
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="На людей, получивших",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="## начертание зверя",
                block_index=2,
                paragraph_id="p2",
                source_index=1,
                role="heading",
                structural_role="heading",
                heading_level=2,
                from_registry=True,
                generated_heading_kind="false_fragment_heading",
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="и поклонявшихся его образу.",
                block_index=3,
                paragraph_id="p3",
                source_index=2,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(
            demoted_false_headings=1,
        ),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown="На людей, получивших начертание зверя и поклонявшихся его образу.",
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    assert report["quality_status"] == "pass"
    assert report["gate_reasons"] == []
    assert report["false_fragment_heading_count"] == 0
    assert report["false_fragment_heading_samples"] == []


def test_build_translation_quality_report_flags_raw_false_fragment_without_entry_authority():
    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=(
            "Иисус постоянно говорит о том, как важно распознавать знамения, чтобы, если им будет даровано пережить\n\n"
            "## Великую скорбь\n\n"
            "они могли устоять до конца."
        ),
        formatting_diagnostics_artifacts=[],
    )

    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == ["false_fragment_headings_review_required"]
    assert report["false_fragment_heading_count"] == 1
    assert report["false_fragment_heading_gate_source"] == "legacy_markdown"
    assert report["raw_false_fragment_heading_count"] == 1
    assert report["false_fragment_heading_samples"] == [
        {
            "line": 3,
            "text": "## Великую скорбь",
            "reason": "inline_term_heading_present",
        }
    ]
    assert report["formatting_review_items"][0]["reason"] == "false_fragment_headings_review_required"
    assert report["formatting_review_items"][0]["severity"] == "fix"


def test_build_translation_quality_report_fails_large_false_fragment_heading_set(monkeypatch):
    samples = [
        SimpleNamespace(
            line=index + 1,
            text=f"## Fragment {index}",
            reason="inline_term_heading_present",
        )
        for index in range(11)
    ]
    monkeypatch.setattr(
        document_pipeline_late_phases,
        "collect_false_fragment_heading_samples",
        lambda markdown_text: list(samples),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
            paragraph_count=1000,
        ),
        final_markdown="Fragments",
        formatting_diagnostics_artifacts=[],
    )

    # spec 018: a large false-fragment set is review-grade (fix severity) — delivered
    # as ``warn`` with the reason + review items preserved, not a blocking ``fail``.
    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == ["false_fragment_headings_present"]
    assert report["false_fragment_heading_count"] == 11
    assert report["formatting_review_required_count"] == 11


def test_build_translation_quality_report_allows_opening_chapter_marker_followed_by_title_heading():
    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=(
            "ГЛАВА 1\n\n"
            "## Что такое богатство?\n\n"
            "Первый абзац главы начинается здесь."
        ),
        formatting_diagnostics_artifacts=[],
    )

    assert report["quality_status"] == "pass"
    assert report["gate_reasons"] == []
    assert report["false_fragment_heading_count"] == 0
    assert report["raw_false_fragment_heading_count"] == 0
    assert report["false_fragment_heading_samples"] == []


def test_build_translation_quality_report_flags_scripture_reference_false_heading_without_normalizer_override():
    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown="Наблюдайте внимательно.\n\n## (Матфея 24:36)\n\nЭто продолжение абзаца.",
        formatting_diagnostics_artifacts=[],
    )

    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == ["false_fragment_headings_review_required"]
    assert report["false_fragment_heading_count"] == 1
    assert report["scripture_reference_heading_count"] == 1
    assert report["scripture_reference_heading_samples"] == [
        {
            "line": 3,
            "text": "## (Матфея 24:36)",
            "reason": "scripture_reference_heading_present",
        }
    ]


def test_build_translation_quality_report_prefers_entry_authority_over_raw_false_fragment_markdown():
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown="Иисус постоянно говорит о том, как важно распознавать знамения, чтобы, если им будет даровано пережить Великую скорбь они могли устоять до конца.",
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="Иисус постоянно говорит о том, как важно распознавать знамения, чтобы, если им будет даровано пережить",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="Великую скорбь",
                block_index=1,
                paragraph_id="p2",
                source_index=1,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="они могли устоять до конца.",
                block_index=1,
                paragraph_id="p3",
                source_index=2,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=(
            "Иисус постоянно говорит о том, как важно распознавать знамения, чтобы, если им будет даровано пережить\n\n"
            "## Великую скорбь\n\n"
            "они могли устоять до конца."
        ),
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    assert report["quality_status"] == "pass"
    assert report["gate_reasons"] == []
    assert report["false_fragment_heading_count"] == 0
    assert report["false_fragment_heading_samples"] == []
    assert report["false_fragment_heading_gate_source"] == "entry_assembly"
    assert report["raw_false_fragment_heading_count"] == 1


def test_build_translation_quality_report_keeps_entry_authority_with_mixed_fallback_entries():
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown=(
            "Иисус постоянно говорит о том, как важно распознавать знамения, чтобы, если им будет даровано пережить\n\n"
            "## Великую скорбь\n\n"
            "они могли устоять до конца."
        ),
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="Иисус постоянно говорит о том, как важно распознавать знамения, чтобы, если им будет даровано пережить",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="## Великую скорбь",
                block_index=2,
                paragraph_id="p2",
                source_index=1,
                role="heading",
                structural_role="heading",
                heading_level=2,
                from_registry=True,
                generated_heading_kind="real_heading",
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="они могли устоять до конца.",
                block_index=3,
                used_fallback=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=(
            "Иисус постоянно говорит о том, как важно распознавать знамения, чтобы, если им будет даровано пережить\n\n"
            "## Великую скорбь\n\n"
            "они могли устоять до конца."
        ),
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    assert report["quality_status"] == "pass"
    assert report["gate_reasons"] == []
    assert report["false_fragment_heading_count"] == 0
    assert report["false_fragment_heading_gate_source"] == "entry_assembly"
    assert report["raw_false_fragment_heading_count"] == 1


def test_collect_false_fragment_heading_samples_from_entries_preserves_source_backed_real_heading():
    entries = (
        document_pipeline_output_validation.FinalAssemblyEntry(
            text="Незавершённая вводная фраза о разделе",
            block_index=1,
            paragraph_id="p1",
            source_index=0,
            role="body",
            structural_role="body",
            from_registry=True,
        ),
        document_pipeline_output_validation.FinalAssemblyEntry(
            text="## Размышление о верности",
            block_index=2,
            paragraph_id="p2",
            source_index=1,
            role="heading",
            structural_role="heading",
            heading_level=2,
            from_registry=True,
            generated_heading_kind="real_heading",
        ),
        document_pipeline_output_validation.FinalAssemblyEntry(
            text="Этот раздел открывается полноценным абзацем.",
            block_index=3,
            paragraph_id="p3",
            source_index=2,
            role="body",
            structural_role="body",
            from_registry=True,
        ),
    )

    samples = document_pipeline_output_validation.collect_false_fragment_heading_samples_from_entries(entries)

    assert samples == []


def test_collect_false_fragment_heading_samples_from_entries_allows_opening_chapter_marker_pair():
    entries = (
        document_pipeline_output_validation.FinalAssemblyEntry(
            text="ГЛАВА 1",
            block_index=1,
            paragraph_id="p1",
            source_index=0,
            role="body",
            structural_role="body",
            from_registry=True,
        ),
        document_pipeline_output_validation.FinalAssemblyEntry(
            text="## Что такое богатство?",
            block_index=2,
            paragraph_id="p2",
            source_index=1,
            role="heading",
            structural_role="heading",
            heading_level=2,
            from_registry=True,
        ),
        document_pipeline_output_validation.FinalAssemblyEntry(
            text="Первый абзац главы начинается здесь.",
            block_index=3,
            paragraph_id="p3",
            source_index=2,
            role="body",
            structural_role="body",
            from_registry=True,
        ),
    )

    samples = document_pipeline_output_validation.collect_false_fragment_heading_samples_from_entries(entries)

    assert samples == []


def test_collect_false_fragment_heading_samples_from_entries_overrides_source_heading_for_parenthetical_question_tail():
    entries = (
        document_pipeline_output_validation.FinalAssemblyEntry(
            text="Кометы? Астероиды?",
            block_index=1,
            paragraph_id="p1",
            source_index=0,
            role="body",
            structural_role="body",
            from_registry=True,
        ),
        document_pipeline_output_validation.FinalAssemblyEntry(
            text="## Спутники? Ракеты?)",
            block_index=2,
            paragraph_id="p2",
            source_index=1,
            role="heading",
            structural_role="heading",
            heading_level=2,
            from_registry=True,
            generated_heading_kind=None,
        ),
        document_pipeline_output_validation.FinalAssemblyEntry(
            text="Треть земли будет сожжена.",
            block_index=3,
            paragraph_id="p3",
            source_index=2,
            role="body",
            structural_role="body",
            from_registry=True,
        ),
    )

    samples = document_pipeline_output_validation.collect_false_fragment_heading_samples_from_entries(entries)

    assert samples == [
        document_pipeline_output_validation.QualityIssueSample(
            line=3,
            text="## Спутники? Ракеты?)",
            reason="sentence_split_heading_present",
        )
    ]


def test_run_document_processing_quality_report_uses_same_final_markdown_as_runtime_state(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    final_markdown = "На людей, получивших начертание зверя и поклонявшихся его образу, приходят язвы."
    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "strict", "enable_paragraph_markers": True},
        processing_operation="translate",
        jobs=[{
            "target_text": "Исходный блок",
            "target_text_with_markers": "[[DOCX_PARA_p1]]\n[[DOCX_PARA_p2]]\n[[DOCX_PARA_p3]]",
            "paragraph_ids": ["p1", "p2", "p3"],
            "context_before": "",
            "context_after": "",
            "target_chars": 13,
            "context_chars": 0,
        }],
        source_paragraphs=[
            SimpleNamespace(paragraph_id="p1", source_index=0, role="body", structural_role="body", heading_level=None, list_kind=None, boundary_source="raw", boundary_confidence="explicit"),
            SimpleNamespace(paragraph_id="p2", source_index=1, role="body", structural_role="body", heading_level=None, list_kind=None, boundary_source="raw", boundary_confidence="explicit"),
            SimpleNamespace(paragraph_id="p3", source_index=2, role="body", structural_role="body", heading_level=None, list_kind=None, boundary_source="raw", boundary_confidence="explicit"),
        ],
        generate_markdown_block=lambda **kwargs: "На людей, получивших\n\n## начертание зверя\n\nи поклонявшихся его образу, приходят язвы.",
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == final_markdown
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "pass"
    assert payload["final_markdown_chars"] == len(final_markdown)
    assert payload["false_fragment_heading_count"] == 0

def test_run_document_processing_warns_on_legacy_markdown_toc_concat_quality_gate(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "strict"},
        processing_operation="translate",
        generate_markdown_block=lambda **kwargs: "Заключение........ 29 Введение",
    )

    assert result == "succeeded"
    assert runtime["state"]["last_error"] == ""
    assert runtime["state"]["latest_docx_bytes"] == b"docx-bytes"
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "warn"
    assert payload["gate_reasons"] == ["toc_body_concatenation_review_required"]
    assert payload["bullet_heading_count"] == 0
    assert payload["toc_body_concat_detected"] is True
    assert runtime["state"]["latest_result_notice"] == {
        "level": "warning",
        "message": (
            "Перевод завершён. Документ готов к использованию, но требует ручной "
            "проверки оформления: 1 абзац с замечаниями. "
            "Подробности — в отчёте проверки (formatting_review.txt)."
        ),
    }


def test_build_translation_quality_report_flags_suspicious_heading_repetition_without_intervening_body():
    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=(
            "## Начертание зверя\n\n"
            "## Начертание зверя\n\n"
            "Новый текст после подозрительного дубля."
        ),
        formatting_diagnostics_artifacts=[],
    )

    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == ["false_fragment_headings_review_required"]
    assert report["suspicious_heading_repetition_count"] == 1
    assert report["suspicious_heading_repetition_samples"] == [
        {
            "line": 3,
            "text": "Начертание зверя",
            "reason": "suspicious_heading_repetition_present",
        }
    ]


def test_build_translation_quality_report_keeps_list_fragment_runtime_cleanup_display_only():
    final_markdown = (
        "Поразительно, но все петли следуют одной и той же схеме: 1.\n"
        "Духовные существа восстают против Бога.\n"
        "2. Бог судит их за грех.\n"
        "3. Бог спасает остаток верных."
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=final_markdown,
        formatting_diagnostics_artifacts=[],
    )

    display_markdown = document_pipeline_late_phases._normalize_final_markdown_for_runtime_display(final_markdown)

    # spec 018: body-list fragmentation is review-grade — delivered as ``warn`` with the
    # reason preserved (the discrepancy still gates the acceptance verdict as review-DATA).
    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == ["list_fragment_regressions_present"]
    assert report["list_fragment_regression_count"] == 1
    assert report["list_fragment_regression_gate_source"] == "legacy_markdown"
    assert report["raw_list_fragment_regression_count"] == 1
    assert "схеме: 1." in final_markdown
    assert "схеме: 1." not in display_markdown
    assert "1. Духовные существа восстают против Бога." in display_markdown


def test_build_translation_quality_report_credits_source_backed_numbered_references_without_topology():
    final_markdown = (
        "2. Goldman Sachs Annual Report, 2010.\n\n"
        "14. Forbes, 2017."
    )
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown=final_markdown,
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="2. Goldman Sachs Annual Report, 2010.",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="14. Forbes, 2017.",
                block_index=1,
                paragraph_id="p2",
                source_index=1,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=final_markdown,
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    assert report["quality_status"] == "pass"
    assert report["gate_reasons"] == []
    assert report["list_fragment_regression_count"] == 0
    assert report["list_fragment_regression_samples"] == []
    assert report["list_fragment_regression_gate_source"] == "entry_assembly"
    assert report["raw_list_fragment_regression_count"] == 2


def test_build_translation_quality_report_keeps_standalone_numeric_continuation_after_reference_credit():
    final_markdown = (
        "26. Барба и де Виво, «Взгляд на финансы как на непроизводительный труд», с.\n\n"
        "1491.\n\n"
        "27. Хаттон и Кент, «Рынок валютных деривативов», с. 225."
    )
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown=final_markdown,
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="26. Барба и де Виво, «Взгляд на финансы как на непроизводительный труд», с.",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="1491.",
                block_index=1,
                paragraph_id="p2",
                source_index=1,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="27. Хаттон и Кент, «Рынок валютных деривативов», с. 225.",
                block_index=1,
                paragraph_id="p3",
                source_index=2,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=final_markdown,
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    assert report["quality_status"] == "warn"
    assert report["gate_reasons"] == ["list_fragment_regressions_review_required"]
    assert report["list_fragment_regression_count"] == 1
    assert report["list_fragment_regression_samples"] == [
        {
            "line": 3,
            "text": "1491.",
            "reason": "list_fragment_regressions_present",
        }
    ]
    assert report["list_fragment_regression_gate_source"] == "entry_assembly"
    assert report["raw_list_fragment_regression_count"] == 2
    assert report["formatting_review_required_count"] == 1
    assert report["formatting_review_items"] == [
        {
            "reason": "list_fragment_regressions_review_required",
            "label": "Одиночный номер в сносках или библиографии",
            "count": 1,
            "severity": "review",
            "sample": {
                "line": 3,
                "text": "1491.",
                "reason": "list_fragment_regressions_present",
            },
        }
    ]


def test_build_translation_quality_report_routes_four_standalone_numeric_backmatter_to_review():
    # Regression (GLOBAL_PLAN 1-B): four standalone-numeric back-matter residues
    # (footnote / page numbers like mazzucato's "18." / "1491." / "1489." / "249.")
    # must route to soft review, not the acceptance hard-fail. The old count cap
    # (<= 3) wrongly tipped an otherwise-good book into failure at 4 such numbers.
    final_markdown = (
        "2. Goldman Sachs Annual Report, 2010.\n\n"
        "18.\n\n"
        "14. Forbes, 2017.\n\n"
        "1491.\n\n"
        "24. Kaldor, Essays on Value and Distribution, 1960.\n\n"
        "1489.\n\n"
        "31. Hutton and Kent, Currency Derivatives, 2018.\n\n"
        "249."
    )
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown=final_markdown,
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="2. Goldman Sachs Annual Report, 2010.",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="18.",
                block_index=1,
                paragraph_id="p2",
                source_index=1,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="14. Forbes, 2017.",
                block_index=1,
                paragraph_id="p3",
                source_index=2,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="1491.",
                block_index=1,
                paragraph_id="p4",
                source_index=3,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="24. Kaldor, Essays on Value and Distribution, 1960.",
                block_index=1,
                paragraph_id="p5",
                source_index=4,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="1489.",
                block_index=1,
                paragraph_id="p6",
                source_index=5,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="31. Hutton and Kent, Currency Derivatives, 2018.",
                block_index=1,
                paragraph_id="p7",
                source_index=6,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="249.",
                block_index=1,
                paragraph_id="p8",
                source_index=7,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=final_markdown,
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    # Detector narrowed raw 8 -> effective 4, all standalone-numeric back-matter.
    assert report["list_fragment_regression_gate_source"] == "entry_assembly"
    assert report["raw_list_fragment_regression_count"] == 8
    assert report["list_fragment_regression_count"] == 4
    assert [sample["text"] for sample in report["list_fragment_regression_samples"]] == [
        "18.",
        "1491.",
        "1489.",
        "249.",
    ]
    # Softened path: review, NOT the acceptance hard-fail reason.
    assert "list_fragment_regressions_present" not in report["gate_reasons"]
    assert report["gate_reasons"] == ["list_fragment_regressions_review_required"]
    assert report["quality_status"] == "warn"
    assert report["formatting_review_required_count"] == 4


def test_build_translation_quality_report_hard_fails_non_numeric_body_list_fragment_residue():
    # Counter-check (GLOBAL_PLAN 1-B): a real body-text list fragment (a broken
    # bullet / item that is NOT standalone-numeric back-matter) must still be an
    # acceptance hard-fail, regardless of count. Softening the cap for numeric
    # back-matter must never silence genuine body-list fragmentation.
    final_markdown = (
        "2. Goldman Sachs Annual Report, 2010.\n\n"
        "- разорванный\n"
        "- пункт"
    )
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown=final_markdown,
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="2. Goldman Sachs Annual Report, 2010.",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="- разорванный",
                block_index=1,
                paragraph_id="p2",
                source_index=1,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="- пункт",
                block_index=1,
                paragraph_id="p3",
                source_index=2,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=final_markdown,
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    assert report["list_fragment_regression_gate_source"] == "entry_assembly"
    assert report["list_fragment_regression_count"] == 1
    non_numeric_sample = report["list_fragment_regression_samples"][0]
    assert not document_pipeline_late_phases._is_standalone_numeric_continuation_sample(
        SimpleNamespace(text=non_numeric_sample["text"])
    )
    # Body-list fragmentation is a hard (non-review) gate reason; spec 018 delivers it as
    # ``warn`` (review-DATA) rather than blocking, but the reason token is still recorded.
    assert "list_fragment_regressions_present" in report["gate_reasons"]
    assert "list_fragment_regressions_review_required" not in report["gate_reasons"]
    assert report["quality_status"] == "warn"


def test_build_translation_quality_report_credits_source_backed_reference_by_exact_text_when_line_offsets_drift():
    final_markdown = (
        "Intro paragraph that shifts assembly offsets.\n\n"
        "2. Goldman Sachs Annual Report, 2010.\n\n"
        "1491."
    )
    assembly_result = document_pipeline_output_validation.FinalMarkdownAssemblyResult(
        final_markdown=final_markdown,
        entries=(
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="2. Goldman Sachs Annual Report, 2010.",
                block_index=1,
                paragraph_id="p1",
                source_index=0,
                role="list",
                structural_role="list",
                list_kind="ordered",
                from_registry=True,
            ),
            document_pipeline_output_validation.FinalAssemblyEntry(
                text="1490–1491.",
                block_index=1,
                paragraph_id="p2",
                source_index=1,
                role="body",
                structural_role="body",
                from_registry=True,
            ),
        ),
        diagnostics=document_pipeline_output_validation.FinalAssemblyDiagnostics(),
    )

    report = document_pipeline_late_phases._build_translation_quality_report(
        context=SimpleNamespace(
            app_config={"translation_output_quality_gate_policy": "strict"},
            processing_operation="translate",
            uploaded_filename="report.docx",
            translation_domain="general",
        ),
        final_markdown=final_markdown,
        formatting_diagnostics_artifacts=[],
        assembly_result=assembly_result,
    )

    # The "2. Goldman Sachs…" reference is still credited by EXACT TEXT even though the
    # intro paragraph drifts the line→entry offsets (that is the original point of this test).
    # The bare footnote number "1491." now drops too: with only two entries its line 5 falls
    # off the entry sequence, so it has no resolvable list context (003 FR-002, User Story 1
    # scenario 3 — standalone footnote numbers with no list context are out of scope). Both
    # raw samples collapse; nothing reaches the gate.
    assert report["quality_status"] == "pass"
    assert report["gate_reasons"] == []
    assert report["list_fragment_regression_count"] == 0
    assert report["list_fragment_regression_samples"] == []
    assert report["list_fragment_regression_gate_source"] == "entry_assembly"
    assert report["raw_list_fragment_regression_count"] == 2
    assert report["formatting_review_required_count"] == 0


def test_run_document_processing_warns_on_advisory_structural_markdown_quality_gate(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    quality_dir = tmp_path / "quality_reports"
    artifact_calls = {}
    monkeypatch.setattr(document_pipeline_late_phases, "collect_recent_formatting_diagnostics_artifacts", lambda since_epoch_seconds, diagnostics_dir: [])
    monkeypatch.setattr(document_pipeline_late_phases, "QUALITY_REPORTS_DIR", quality_dir)

    result = _run_processing(
        runtime,
        app_config={"translation_output_quality_gate_policy": "advisory"},
        processing_operation="translate",
        generate_markdown_block=lambda **kwargs: "Заключение……29 Введение",
        write_ui_result_artifacts=lambda **kwargs: artifact_calls.setdefault("kwargs", dict(kwargs)) or {"markdown_path": "/tmp/report.result.md", "docx_path": "/tmp/report.result.docx", "metadata_path": "/tmp/report.result.meta.json"},
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_result_notice"] == {
        "level": "warning",
        "message": (
            "Перевод завершён. Документ готов к использованию, но требует ручной "
            "проверки оформления: 1 абзац с замечаниями. "
            "Подробности — в отчёте проверки (formatting_review.txt)."
        ),
    }
    assert artifact_calls["kwargs"]["quality_warning"]["gate_reasons"] == ["toc_body_concatenation_review_required"]
    report_files = list(quality_dir.glob("*.json"))
    assert len(report_files) == 1
    payload = json.loads(report_files[0].read_text(encoding="utf-8"))
    assert payload["quality_status"] == "warn"
    assert payload["gate_reasons"] == ["toc_body_concatenation_review_required"]
    assert payload["toc_body_concat_detected"] is True


def test_run_document_processing_logs_compact_block_plan_summary_at_info() -> None:
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()

    result = _run_processing(
        runtime,
        jobs=[
            {"target_text": "alpha", "context_before": "", "context_after": "", "target_chars": 5, "context_chars": 0},
            {"target_text": "beta", "context_before": "", "context_after": "", "target_chars": 8, "context_chars": 0, "job_kind": "passthrough"},
        ],
        log_event=log_event,
    )

    assert result == "succeeded"
    info_events = [event for event in events if event["level"] == logging.INFO]
    summary_event = next(event for event in info_events if event["event_id"] == "block_plan_summary")
    assert summary_event["context"]["block_count"] == 2
    assert summary_event["context"]["llm_block_count"] == 1
    assert summary_event["context"]["passthrough_block_count"] == 1
    assert summary_event["context"]["total_target_chars"] == 13
    assert summary_event["context"]["first_block_target_chars"] == [5, 8]
    assert "blocks" not in summary_event["context"]
    assert all(event["event_id"] != "block_map" for event in info_events)


def test_run_document_processing_demotes_block_chatter_to_debug() -> None:
    runtime = _build_runtime_capture()
    events, log_event = _capture_log_events()

    result = _run_processing(
        runtime,
        app_config={"enable_paragraph_markers": True},
        jobs=[{
            "target_text": "Исходный блок",
            "target_text_with_markers": "[[DOCX_PARA_p0001]]\nИсходный блок",
            "paragraph_ids": ["p0001"],
            "context_before": "",
            "context_after": "",
            "target_chars": 13,
            "context_chars": 0,
        }],
        generate_markdown_block=lambda **kwargs: "Очищенный блок",
        log_event=log_event,
    )

    assert result == "succeeded"
    info_event_ids = {event["event_id"] for event in events if event["level"] == logging.INFO}
    debug_event_ids = {event["event_id"] for event in events if event["level"] == logging.DEBUG}
    assert "processing_started" in info_event_ids
    assert "processing_completed" in info_event_ids
    assert "block_started" not in info_event_ids
    assert "block_completed" not in info_event_ids
    assert "block_marker_registry_built" not in info_event_ids
    assert {"block_started", "block_completed", "block_marker_registry_built"}.issubset(debug_event_ids)


def test_run_document_processing_passes_marker_wrapped_text_only_when_marker_mode_enabled():
    runtime = _build_runtime_capture()
    generate_calls = []

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "target_text": "Исходный блок",
            "target_text_with_markers": "[[DOCX_PARA_p0001]]\nИсходный блок",
            "paragraph_ids": ["p0001"],
            "context_before": "",
            "context_after": "",
            "target_chars": 13,
            "context_chars": 0,
        }],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={"enable_paragraph_markers": True},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: generate_calls.append(kwargs) or "Очищенный блок",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=_reinsert_inline_images,
    )

    assert result == "succeeded"
    assert generate_calls[0]["target_text"] == "[[DOCX_PARA_p0001]]\nИсходный блок"
    assert generate_calls[0]["expected_paragraph_ids"] == ["p0001"]
    assert generate_calls[0]["marker_mode"] is True
    assert runtime["state"]["latest_markdown"] == "Очищенный блок"
    assert runtime["state"]["processed_paragraph_registry"] == [
        {"block_index": 1, "paragraph_id": "p0001", "text": "Очищенный блок"}
    ]


def test_run_document_processing_passes_generated_paragraph_registry_into_docx_restoration():
    runtime = _build_runtime_capture()
    preserve_calls = []

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "target_text": "Исходный блок",
            "target_text_with_markers": "[[DOCX_PARA_p0001]]\nИсходный блок",
            "paragraph_ids": ["p0001"],
            "context_before": "",
            "context_after": "",
            "target_chars": 13,
            "context_chars": 0,
        }],
        source_paragraphs=[ParagraphStub()],
        image_assets=[],
        image_mode="safe",
        app_config={"enable_paragraph_markers": True},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Очищенный блок",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: preserve_calls.append(generated_paragraph_registry) or docx_bytes,
        reinsert_inline_images=_reinsert_inline_images,
    )

    assert result == "succeeded"
    expected_registry = [{"block_index": 1, "paragraph_id": "p0001", "text": "Очищенный блок"}]
    assert preserve_calls == [expected_registry]


def test_run_document_processing_registry_uses_logical_marker_for_merged_source_paragraph():
    runtime = _build_runtime_capture()
    generate_calls = []
    merged_paragraph = ParagraphUnit(
        text="Слитый логический абзац",
        role="body",
        paragraph_id="p0007",
        origin_raw_indexes=[0, 1],
        origin_raw_texts=["Слитый", "логический абзац"],
        boundary_source="normalized_merge",
        boundary_confidence="high",
    )

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "target_text": "Слитый логический абзац",
            "target_text_with_markers": "[[DOCX_PARA_p0007]]\nСлитый логический абзац",
            "paragraph_ids": ["p0007"],
            "context_before": "",
            "context_after": "",
            "target_chars": 24,
            "context_chars": 0,
        }],
        source_paragraphs=[merged_paragraph],
        image_assets=[],
        image_mode="safe",
        app_config={"enable_paragraph_markers": True},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: generate_calls.append(kwargs) or "Слитый логический абзац",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=_reinsert_inline_images,
    )

    assert result == "succeeded"
    assert generate_calls[0]["expected_paragraph_ids"] == ["p0007"]
    assert runtime["state"]["processed_paragraph_registry"] == [
        {"block_index": 1, "paragraph_id": "p0007", "text": "Слитый логический абзац"}
    ]


def test_run_document_processing_passes_assembly_aware_registry_into_docx_restoration():
    runtime = _build_runtime_capture()
    preserve_calls = []
    source_paragraphs = [
        ParagraphUnit(text="Первый фрагмент", role="body", paragraph_id="p0001"),
        ParagraphUnit(text="Второй фрагмент", role="body", paragraph_id="p0002"),
    ]

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "target_text": "Исходный блок",
            "target_text_with_markers": "[[DOCX_PARA_p0001]]\n[[DOCX_PARA_p0002]]",
            "paragraph_ids": ["p0001", "p0002"],
            "context_before": "",
            "context_after": "",
            "target_chars": 13,
            "context_chars": 0,
        }],
        source_paragraphs=source_paragraphs,
        image_assets=[],
        image_mode="safe",
        app_config={"enable_paragraph_markers": True},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Первый фрагмент\n\nВторой фрагмент",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: preserve_calls.append(generated_paragraph_registry) or docx_bytes,
        reinsert_inline_images=_reinsert_inline_images,
    )

    assert result == "succeeded"
    assert preserve_calls == [[
        {"block_index": 1, "paragraph_id": "p0001", "text": "Первый фрагмент Второй фрагмент", "merged_paragraph_ids": ["p0001", "p0002"]}
    ]]


def test_reader_cleanup_formatting_lineage_removes_deleted_registry_entries():
    raw_markdown = "Intro\n\nHeader\n\nBody paragraph"
    registry = [
        {"block_index": 1, "paragraph_id": "p0001", "text": "Intro"},
        {"block_index": 2, "paragraph_id": "p0002", "text": "Header"},
        {"block_index": 3, "paragraph_id": "p0003", "text": "Body paragraph"},
    ]
    cleanup_report = {
        "accepted_cleanup_operations": [
            {
                "id": "b_000001",
                "operation": "delete_block",
                "expected_after_preview": "",
            }
        ]
    }

    derived_registry, diagnostics = document_pipeline_late_phases._derive_reader_cleanup_generated_paragraph_registry(
        generated_paragraph_registry=registry,
        cleanup_report=cleanup_report,
        raw_markdown=raw_markdown,
    )

    assert diagnostics["status"] == "derived"
    assert diagnostics["deleted_registry_entry_count"] == 1
    assert derived_registry == [
        {"block_index": 1, "paragraph_id": "p0001", "text": "Intro"},
        {"block_index": 3, "paragraph_id": "p0003", "text": "Body paragraph"},
    ]


def test_reader_cleanup_formatting_lineage_updates_split_and_heading_boundary_text():
    raw_markdown = "HEADING Body paragraph"
    registry = [{"block_index": 1, "paragraph_id": "p0001", "text": "HEADING Body paragraph"}]
    cleanup_report = {
        "accepted_cleanup_operations": [
            {
                "id": "b_000000",
                "operation": "normalize_heading_boundary",
                "expected_after_preview": "### HEADING\n\nBody paragraph",
            }
        ]
    }

    derived_registry, diagnostics = document_pipeline_late_phases._derive_reader_cleanup_generated_paragraph_registry(
        generated_paragraph_registry=registry,
        cleanup_report=cleanup_report,
        raw_markdown=raw_markdown,
    )

    assert diagnostics["updated_registry_entry_count"] == 1
    assert derived_registry == [
        {
            "block_index": 1,
            "paragraph_id": "p0001",
            "text": "### HEADING\n\nBody paragraph",
            "reader_cleanup_operations": ["normalize_heading_boundary"],
        }
    ]


def test_reader_cleanup_formatting_lineage_merges_joined_registry_entries():
    raw_markdown = "First fragment\n\nsecond fragment"
    registry = [
        {"block_index": 1, "paragraph_id": "p0001", "text": "First fragment"},
        {"block_index": 2, "paragraph_id": "p0002", "text": "second fragment"},
    ]
    cleanup_report = {
        "accepted_cleanup_operations": [
            {
                "id": "b_000000",
                "operation": "join_fragmented_paragraph",
                "next_id": "b_000001",
                "expected_after_preview": "First fragment second fragment",
            }
        ]
    }

    derived_registry, diagnostics = document_pipeline_late_phases._derive_reader_cleanup_generated_paragraph_registry(
        generated_paragraph_registry=registry,
        cleanup_report=cleanup_report,
        raw_markdown=raw_markdown,
    )

    assert diagnostics["joined_registry_entry_count"] == 1
    assert derived_registry == [
        {
            "block_index": 1,
            "paragraph_id": "p0001",
            "text": "First fragment second fragment",
            "merged_paragraph_ids": ["p0001", "p0002"],
            "reader_cleanup_operations": ["join_fragmented_paragraph"],
        }
    ]


def test_reader_cleanup_formatting_lineage_skips_anchor_repair_operations():
    raw_markdown = "Intro\n\nHEADING Body"
    registry = [
        {"block_index": 1, "paragraph_id": "p0001", "text": "Intro"},
        {"block_index": 2, "paragraph_id": "p0002", "text": "HEADING Body"},
    ]
    cleanup_report = {
        "accepted_cleanup_operations": [
            {
                "id": "b_000001",
                "operation": "normalize_heading_boundary",
                "pass_name": "anchor_repair",
                "expected_after_preview": "### HEADING\n\nBody",
            }
        ]
    }

    derived_registry, diagnostics = document_pipeline_late_phases._derive_reader_cleanup_generated_paragraph_registry(
        generated_paragraph_registry=registry,
        cleanup_report=cleanup_report,
        raw_markdown=raw_markdown,
    )

    assert diagnostics["status"] == "derived"
    assert diagnostics["applied_operation_count"] == 0
    assert diagnostics["skipped_operation_count"] == 1
    assert derived_registry == registry


def test_reader_cleanup_formatting_lineage_skips_when_block_count_is_ambiguous():
    raw_markdown = "Intro\n\nBody"
    registry = [{"block_index": 1, "paragraph_id": "p0001", "text": "Intro"}]

    derived_registry, diagnostics = document_pipeline_late_phases._derive_reader_cleanup_generated_paragraph_registry(
        generated_paragraph_registry=registry,
        cleanup_report={"accepted_cleanup_operations": []},
        raw_markdown=raw_markdown,
    )

    assert diagnostics == {
        "status": "skipped",
        "reason": "cleanup_block_registry_count_mismatch",
        "sparse_alignment_failure_reason": "non_image_placeholder_registry_gaps",
        "alignment_gap_count": 1,
        "raw_cleanup_block_count": 2,
        "generated_registry_count": 1,
    }
    assert derived_registry == registry


def test_reader_cleanup_formatting_lineage_allows_sparse_image_placeholder_gap():
    raw_markdown = "Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph"
    registry = [
        {"block_index": 7, "paragraph_id": "p0001", "text": "Intro"},
        {"block_index": 7, "paragraph_id": "p0003", "text": "Body paragraph"},
    ]
    cleanup_report = {
        "accepted_cleanup_operations": [
            {
                "id": "b_000001",
                "operation": "delete_block",
                "expected_after_preview": "",
            }
        ]
    }

    derived_registry, diagnostics = document_pipeline_late_phases._derive_reader_cleanup_generated_paragraph_registry(
        generated_paragraph_registry=registry,
        cleanup_report=cleanup_report,
        raw_markdown=raw_markdown,
    )

    assert diagnostics["status"] == "derived"
    assert diagnostics["alignment_mode"] == "sparse_image_placeholders"
    assert diagnostics["alignment_gap_count"] == 1
    assert diagnostics["applied_operation_count"] == 0
    assert diagnostics["skipped_operation_count"] == 1
    assert derived_registry == registry


def test_reader_cleanup_formatting_lineage_uses_paragraph_id_when_cleanup_text_drifted():
    raw_markdown = "Intro before cleanup\n\n[[DOCX_IMAGE_img_001]]\n\nBody before cleanup"
    registry = [
        {"block_index": 7, "paragraph_id": "p0001", "text": "Intro after cleanup"},
        {"block_index": 7, "paragraph_id": "p0003", "text": "Body after cleanup"},
    ]

    derived_registry, diagnostics = document_pipeline_late_phases._derive_reader_cleanup_generated_paragraph_registry(
        generated_paragraph_registry=registry,
        cleanup_report={"accepted_cleanup_operations": []},
        raw_markdown=raw_markdown,
        cleanup_block_metadata_by_index={
            0: {"paragraph_id": "p0001"},
            2: {"paragraph_id": "p0003"},
        },
    )

    assert diagnostics["status"] == "derived"
    assert diagnostics["alignment_mode"] == "identity_sparse_image_placeholders"
    assert diagnostics["alignment_gap_count"] == 1
    assert diagnostics["applied_operation_count"] == 0
    assert derived_registry == registry


def test_reader_cleanup_formatting_lineage_rejects_sparse_non_image_gap():
    raw_markdown = "Intro\n\nDropped semantic text\n\nBody paragraph"
    registry = [
        {"block_index": 7, "paragraph_id": "p0001", "text": "Intro"},
        {"block_index": 7, "paragraph_id": "p0003", "text": "Body paragraph"},
    ]

    derived_registry, diagnostics = document_pipeline_late_phases._derive_reader_cleanup_generated_paragraph_registry(
        generated_paragraph_registry=registry,
        cleanup_report={"accepted_cleanup_operations": []},
        raw_markdown=raw_markdown,
    )

    assert diagnostics == {
        "status": "skipped",
        "reason": "cleanup_block_registry_count_mismatch",
        "sparse_alignment_failure_reason": "non_image_placeholder_registry_gaps",
        "alignment_gap_count": 1,
        "raw_cleanup_block_count": 3,
        "generated_registry_count": 2,
    }
    assert derived_registry == registry


def test_reader_cleanup_block_identity_metadata_is_serialized_to_model_payload():
    blocks = build_cleanup_blocks(
        "Intro",
        block_metadata_by_index={
            0: {
                "paragraph_id": "p0001",
                "merged_paragraph_ids": ["p0001", "p0002"],
                "layout_signals": {"font_size": 14.0, "centered": True},
            }
        },
    )

    assert blocks[0].paragraph_id == "p0001"
    assert blocks[0].merged_paragraph_ids == ("p0001", "p0002")
    payload = blocks[0].to_payload()
    assert payload["paragraph_id"] == "p0001"
    assert payload["merged_paragraph_ids"] == ["p0001", "p0002"]
    assert cast(dict[str, object], payload["layout_signals"])["font_size"] == 14.0


def test_reader_cleanup_block_identity_metadata_reports_id_match_and_gaps():
    raw_markdown = "Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph"
    registry = [
        {"block_index": 7, "paragraph_id": "p0001", "text": "Intro"},
        {"block_index": 7, "paragraph_id": "p0003", "text": "Body paragraph"},
    ]

    metadata, diagnostics = document_pipeline_late_phases._build_reader_cleanup_block_identity_metadata(
        raw_markdown=raw_markdown,
        generated_paragraph_registry=registry,
    )

    assert metadata[0]["paragraph_id"] == "p0001"
    assert metadata[0]["layout_signals"] == {
        "standalone_short_line": True,
        "looks_like_superscript_marker": False,
    }
    assert metadata[2]["paragraph_id"] == "p0003"
    assert metadata[2]["layout_signals"] == {
        "standalone_short_line": True,
        "looks_like_superscript_marker": False,
    }
    assert diagnostics == {
        "status": "available",
        "reason": None,
        "raw_cleanup_block_count": 3,
        "generated_registry_count": 2,
        "id_matched_block_count": 2,
        "missing_id_registry_entry_count": 0,
        "gap_count": 1,
        "image_gap_count": 1,
        "text_gap_count": 0,
    }


def test_reader_cleanup_postprocess_prefers_assembly_formatting_registry_over_stale_state_registry(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setattr(document_pipeline_late_phases, "READER_CLEANUP_LINEAGE_DIR", tmp_path / "reader_cleanup_lineage")
    runtime = _build_runtime_capture()
    preserve_calls = []
    log_events = []
    raw_markdown = "Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph"
    assembly_registry = [
        {"block_index": 1, "paragraph_id": "p0001", "text": "Intro"},
        {"block_index": 3, "paragraph_id": "p0003", "text": "Body paragraph"},
    ]
    stale_state_registry = [{"block_index": 1, "paragraph_id": "stale", "text": raw_markdown}]

    def generate_markdown_block(**kwargs):
        payload = json.loads(kwargs["target_text"])
        delete_target = next(block for block in payload["blocks"] if block["text"] == "[[DOCX_IMAGE_img_001]]")
        return json.dumps(
            {
                "cleanup_operations": [
                    {
                        "id": delete_target["id"],
                        "text_hash": delete_target["text_hash"],
                        "operation": "delete_block",
                        "reason": "extraction_artifact",
                        "confidence": "high",
                        "evidence_before": "[[DOCX_IMAGE_img_001]]",
                        "expected_after_preview": "",
                        "safety_note": "test fixture",
                    }
                ],
                "warnings": [],
            },
            ensure_ascii=False,
        )

    result = document_pipeline_late_phases._run_reader_cleanup_postprocess(
        context=SimpleNamespace(
            processing_operation="translate",
            app_config={
                "reader_cleanup_enabled": True,
                "reader_cleanup_policy": "advisory",
                "reader_cleanup_chunk_size": 500,
                "reader_cleanup_max_delete_block_ratio": 0.8,
                "reader_cleanup_max_delete_char_ratio": 0.8,
            },
            model="anthropic:claude-sonnet-4-6",
            max_retries=1,
            uploaded_filename="report.docx",
            runtime=runtime,
            source_paragraphs=[ParagraphUnit(text="Intro", role="body", paragraph_id="p0001")],
        ),
        dependencies=SimpleNamespace(
            get_client=lambda: object(),
            generate_markdown_block=generate_markdown_block,
            convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
            preserve_source_paragraph_properties=(
                lambda docx_bytes, paragraphs, generated_paragraph_registry=None: preserve_calls.append(
                    generated_paragraph_registry
                )
                or docx_bytes
            ),
            reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
            log_event=lambda level, event_id, message, **context: log_events.append(
                {"level": level, "event_id": event_id, "message": message, "context": context}
            ),
            present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        ),
        emitters=SimpleNamespace(
            emit_activity=lambda *args, **kwargs: None,
            emit_state=lambda *args, **kwargs: None,
        ),
        state=SimpleNamespace(generated_paragraph_registry=stale_state_registry),
        cleanup_input_markdown=raw_markdown,
        runtime_display_markdown=raw_markdown,
        base_docx_bytes=b"base-docx",
        job_count=1,
        processed_image_assets=[],
        formatting_registry=assembly_registry,
    )
    cleaned_markdown = result.markdown
    cleaned_docx_bytes = result.docx_bytes
    report = result.report
    final_registry = result.final_generated_paragraph_registry

    assert cleaned_markdown == raw_markdown
    assert cleaned_docx_bytes == b"base-docx"
    assert report is not None
    assert report["stats"]["accepted_delete_block_count"] == 0
    assert report["image_reconciliation"]["before_image_id_count"] == 1
    assert report["image_reconciliation"]["after_image_id_count"] == 1
    assert final_registry == [
        {"block_index": 1, "paragraph_id": "p0001", "text": "Intro", "target_paragraph_indexes": [0]},
        {"block_index": 3, "paragraph_id": "p0003", "text": "Body paragraph", "target_paragraph_indexes": [2]},
    ]
    assert preserve_calls == []
    noop_event = next(event for event in log_events if event["event_id"] == "reader_cleanup_noop")
    assert noop_event["context"]["cleanup_identity_status"] == "available"
    assert noop_event["context"]["cleanup_identity_id_matched_block_count"] == 2
    assert noop_event["context"]["cleanup_identity_image_gap_count"] == 1
    assert noop_event["context"]["cleanup_identity_text_gap_count"] == 0


def test_reader_cleanup_noop_restores_image_heading_concats_from_registry():
    runtime = _build_runtime_capture()
    raw_markdown = "[[DOCX_IMAGE_img_001]] Глава восьмая\n\nBody paragraph"
    assembly_registry = [
        {"block_index": 1, "paragraph_id": "p0026", "text": "## Глава восьмая"},
        {"block_index": 2, "paragraph_id": "p0027", "text": "Body paragraph"},
    ]

    result = document_pipeline_late_phases._run_reader_cleanup_postprocess(
        context=SimpleNamespace(
            processing_operation="translate",
            app_config={
                "reader_cleanup_enabled": True,
                "reader_cleanup_policy": "advisory",
                "reader_cleanup_chunk_size": 500,
                "reader_cleanup_max_delete_block_ratio": 0.8,
                "reader_cleanup_max_delete_char_ratio": 0.8,
            },
            model="anthropic:claude-sonnet-4-6",
            max_retries=1,
            uploaded_filename="report.docx",
            runtime=runtime,
            source_paragraphs=[ParagraphUnit(text="Chapter Eight", role="heading", paragraph_id="p0026")],
        ),
        dependencies=SimpleNamespace(
            get_client=lambda: object(),
            generate_markdown_block=lambda **kwargs: json.dumps({"cleanup_operations": [], "warnings": []}),
            convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
            preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
            reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
            log_event=lambda *args, **kwargs: None,
            present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        ),
        emitters=SimpleNamespace(
            emit_activity=lambda *args, **kwargs: None,
            emit_state=lambda *args, **kwargs: None,
        ),
        state=SimpleNamespace(generated_paragraph_registry=[]),
        cleanup_input_markdown=raw_markdown,
        runtime_display_markdown=raw_markdown,
        base_docx_bytes=None,
        base_docx_builder=None,
        job_count=1,
        processed_image_assets=[],
        formatting_registry=assembly_registry,
    )

    assert result.markdown == "[[DOCX_IMAGE_img_001]]\n\n## Глава восьмая\n\nBody paragraph"
    assert result.final_generated_paragraph_registry == [
        {"block_index": 1, "paragraph_id": "p0026", "text": "## Глава восьмая", "target_paragraph_indexes": [1]},
        {"block_index": 2, "paragraph_id": "p0027", "text": "Body paragraph", "target_paragraph_indexes": [2]},
    ]


def test_reader_cleanup_postprocess_persists_final_generated_registry_in_runtime_state(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setattr(document_pipeline_late_phases, "READER_CLEANUP_LINEAGE_DIR", tmp_path / "reader_cleanup_lineage")
    runtime = _build_runtime_capture()
    raw_markdown = "Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody paragraph"
    assembly_registry = [
        {"block_index": 1, "paragraph_id": "p0001", "text": "Intro"},
        {"block_index": 3, "paragraph_id": "p0003", "text": "Body paragraph"},
    ]

    def generate_markdown_block(**kwargs):
        payload = json.loads(kwargs["target_text"])
        delete_target = next(block for block in payload["blocks"] if block["text"] == "[[DOCX_IMAGE_img_001]]")
        return json.dumps(
            {
                "cleanup_operations": [
                    {
                        "id": delete_target["id"],
                        "text_hash": delete_target["text_hash"],
                        "operation": "delete_block",
                        "reason": "extraction_artifact",
                        "confidence": "high",
                        "evidence_before": "[[DOCX_IMAGE_img_001]]",
                        "expected_after_preview": "",
                        "safety_note": "test fixture",
                    }
                ],
                "warnings": [],
            },
            ensure_ascii=False,
        )

    result = document_pipeline_late_phases._run_reader_cleanup_postprocess(
        context=SimpleNamespace(
            processing_operation="translate",
            app_config={
                "reader_cleanup_enabled": True,
                "reader_cleanup_policy": "advisory",
                "reader_cleanup_chunk_size": 500,
                "reader_cleanup_max_delete_block_ratio": 0.8,
                "reader_cleanup_max_delete_char_ratio": 0.8,
            },
            model="anthropic:claude-sonnet-4-6",
            max_retries=1,
            uploaded_filename="report.docx",
            runtime=runtime,
            source_paragraphs=[ParagraphUnit(text="Intro", role="body", paragraph_id="p0001")],
        ),
        dependencies=SimpleNamespace(
            get_client=lambda: object(),
            generate_markdown_block=generate_markdown_block,
            convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
            preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
            reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
            log_event=lambda *args, **kwargs: None,
            present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        ),
        emitters=SimpleNamespace(
            emit_activity=lambda *args, **kwargs: None,
            emit_state=_emit_state,
        ),
        state=SimpleNamespace(generated_paragraph_registry=None),
        cleanup_input_markdown=raw_markdown,
        runtime_display_markdown=raw_markdown,
        base_docx_bytes=b"base-docx",
        job_count=1,
        processed_image_assets=[],
        formatting_registry=assembly_registry,
    )

    assert result.final_generated_paragraph_registry == [
        {"block_index": 1, "paragraph_id": "p0001", "text": "Intro", "target_paragraph_indexes": [0]},
        {"block_index": 3, "paragraph_id": "p0003", "text": "Body paragraph", "target_paragraph_indexes": [2]},
    ]
    assert "final_generated_paragraph_registry" not in runtime["state"]


def test_rebuild_docx_for_markdown_prefers_cleanup_formatting_registry_override():
    preserve_calls = []
    override_registry = [{"block_index": 1, "paragraph_id": "p0001", "text": "Cleaned"}]

    document_pipeline_late_phases._rebuild_docx_for_markdown(
        markdown_text="Cleaned",
        context=SimpleNamespace(source_paragraphs=[ParagraphUnit(text="Source", role="body", paragraph_id="p0001")]),
        dependencies=SimpleNamespace(
            convert_markdown_to_docx_bytes=lambda markdown_text: markdown_text.encode("utf-8"),
            preserve_source_paragraph_properties=(
                lambda docx_bytes, paragraphs, generated_paragraph_registry=None: preserve_calls.append(
                    generated_paragraph_registry
                )
                or docx_bytes
            ),
            reinsert_inline_images=lambda docx_bytes, processed_assets: docx_bytes,
        ),
        state=SimpleNamespace(generated_paragraph_registry=[{"block_index": 1, "paragraph_id": "stale", "text": "Stale"}]),
        processed_image_assets=[],
        generated_paragraph_registry=override_registry,
    )

    assert preserve_calls == [
        [{"block_index": 1, "paragraph_id": "p0001", "text": "Cleaned", "target_paragraph_indexes": [0]}]
    ]


def test_rebuild_identity_formatting_registry_attaches_target_indexes_without_visible_markers():
    registry = document_pipeline_late_phases._build_rebuild_identity_formatting_registry(
        markdown_text="# Translated Heading\n\nTranslated body",
        generated_paragraph_registry=[
            {"paragraph_id": "p0001", "text": "# Translated Heading"},
            {"paragraph_id": "p0002", "text": "Translated body"},
        ],
    )

    assert registry == [
        {"paragraph_id": "p0001", "text": "# Translated Heading", "target_paragraph_indexes": [0]},
        {"paragraph_id": "p0002", "text": "Translated body", "target_paragraph_indexes": [1]},
    ]


def test_rebuild_identity_formatting_registry_rolls_back_partial_multiblock_match():
    registry = document_pipeline_late_phases._build_rebuild_identity_formatting_registry(
        markdown_text="First block\n\nNext stable block",
        generated_paragraph_registry=[
            {"paragraph_id": "p0001", "text": "First block\n\nMissing continuation"},
            {"paragraph_id": "p0002", "text": "Next stable block"},
        ],
    )

    assert registry == [
        {"paragraph_id": "p0001", "text": "First block\n\nMissing continuation"},
        {"paragraph_id": "p0002", "text": "Next stable block", "target_paragraph_indexes": [1]},
    ]


def test_reader_cleanup_lineage_rebuild_harness_accepts_runtime_lineage_artifact(tmp_path):
    project_root = Path(__file__).resolve().parents[1]
    artifact_path = tmp_path / "reader_cleanup_lineage.json"
    output_path = tmp_path / "harness_result.json"
    artifact_path.write_text(
        json.dumps(
            {
                "stage": "reader_cleanup_lineage",
                "raw_markdown": "Intro\n\n[[DOCX_IMAGE_img_001]]\n\nBody",
                "cleaned_markdown": "Intro\n\nBody",
                "cleanup_report": {"accepted_cleanup_operations": []},
                "active_formatting_registry": [
                    {"block_index": 1, "paragraph_id": "p0001", "text": "Intro"},
                    {"block_index": 3, "paragraph_id": "p0003", "text": "Body"},
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    completed = subprocess.run(
        [
            sys.executable,
            str(project_root / "scripts/run-reader-cleanup-lineage-rebuild-harness.py"),
            "--lineage-artifact",
            str(artifact_path),
            "--output",
            str(output_path),
        ],
        cwd=project_root,
        check=False,
        text=True,
        capture_output=True,
    )

    assert completed.returncode == 0, completed.stdout + completed.stderr
    payload = json.loads(output_path.read_text(encoding="utf-8"))
    assert payload["status"] == "passed"
    attempt = payload["candidate_attempts"][0]
    assert attempt["lineage_diagnostics"]["status"] == "derived"
    assert attempt["lineage_diagnostics"]["alignment_mode"] == "identity_sparse_image_placeholders"
    assert attempt["artifact_shape"]["raw_image_placeholder_count"] == 1
    assert attempt["artifact_shape"]["cleaned_image_placeholder_count"] == 0
    assert attempt["artifact_shape"]["rebuilt_image_placeholder_count"] == 1


def test_run_document_processing_writes_marker_generation_diagnostics_artifact_on_marker_validation_failure(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    monkeypatch.setattr(document_pipeline, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "target_text": "Исходный блок",
            "target_text_with_markers": "[[DOCX_PARA_p0001]]\nИсходный блок",
            "paragraph_ids": ["p0001"],
            "context_before": "prev",
            "context_after": "next",
            "target_chars": 13,
            "context_chars": 8,
        }],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={"enable_paragraph_markers": True},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: (_ for _ in ()).throw(RuntimeError("paragraph_marker_validation_failed:markers_missing")),
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=_reinsert_inline_images,
    )

    assert result == "failed"
    artifact_path = Path(runtime["state"]["latest_marker_diagnostics_artifact"])
    assert artifact_path.exists()
    payload = json.loads(artifact_path.read_text(encoding="utf-8"))
    assert payload["stage"] == "generation"
    assert payload["error_code"] == "markers_missing"
    assert payload["paragraph_ids"] == ["p0001"]
    assert "marker diagnostics:" in runtime["log"][-1]["details"]


def test_run_document_processing_marker_generation_artifact_includes_found_ids_and_raw_response_preview(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    monkeypatch.setattr(document_pipeline, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "target_text": "Исходный блок",
            "target_text_with_markers": "[[DOCX_PARA_p0001]]\nИсходный блок",
            "paragraph_ids": ["p0001"],
            "context_before": "prev",
            "context_after": "next",
            "target_chars": 13,
            "context_chars": 8,
        }],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={"enable_paragraph_markers": True},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: (_ for _ in ()).throw(
            generation.MarkerValidationError(
                "marker_order_or_identity",
                raw_markdown="[[DOCX_PARA_p9999]]\nЧужой маркер",
                expected_paragraph_ids=["p0001"],
                found_paragraph_ids=["p9999"],
            )
        ),
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=_reinsert_inline_images,
    )

    assert result == "failed"
    artifact_path = Path(runtime["state"]["latest_marker_diagnostics_artifact"])
    payload = json.loads(artifact_path.read_text(encoding="utf-8"))
    assert payload["error_code"] == "marker_order_or_identity"
    assert payload["expected_paragraph_ids"] == ["p0001"]
    assert payload["found_paragraph_ids"] == ["p9999"]
    assert payload["raw_response_preview"] == "[[DOCX_PARA_p9999]]\nЧужой маркер"


def test_run_document_processing_writes_marker_registry_diagnostics_artifact_on_registry_mismatch(tmp_path, monkeypatch):
    runtime = _build_runtime_capture()
    diagnostics_dir = tmp_path / "formatting_diagnostics"
    monkeypatch.setattr(document_pipeline, "FORMATTING_DIAGNOSTICS_DIR", diagnostics_dir)

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[{
            "target_text": "Исходный блок",
            "target_text_with_markers": "[[DOCX_PARA_p0001]]\nИсходный блок",
            "paragraph_ids": ["p0001", "p0002"],
            "context_before": "prev",
            "context_after": "next",
            "target_chars": 13,
            "context_chars": 8,
        }],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={"enable_paragraph_markers": True},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: "Только один абзац",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=_reinsert_inline_images,
    )

    assert result == "failed"
    artifact_path = Path(runtime["state"]["latest_marker_diagnostics_artifact"])
    assert artifact_path.exists()
    payload = json.loads(artifact_path.read_text(encoding="utf-8"))
    assert payload["stage"] == "registry"
    assert payload["error_code"].startswith("block=1:expected=2:actual=1")
    assert payload["processed_chunk_preview"] == "Только один абзац"
    assert "marker diagnostics:" in runtime["log"][-1]["details"]


def test_run_document_processing_stops_before_second_block():
    runtime = _build_runtime_capture()
    stop_checks = {"count": 0}

    def should_stop(runtime):
        stop_checks["count"] += 1
        return stop_checks["count"] >= 2

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {"target_text": "block-1", "context_before": "", "context_after": "", "target_chars": 7, "context_chars": 0},
            {"target_text": "block-2", "context_before": "", "context_after": "", "target_chars": 7, "context_chars": 0},
        ],
        source_paragraphs=[],
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=should_stop,
        generate_markdown_block=lambda **kwargs: "ok",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=_reinsert_inline_images,
    )

    assert result == "stopped"
    assert runtime["finalize"][-1][0] == "Остановлено пользователем"
    assert runtime["finalize"][-1][3] == "stopped"
    assert runtime["log"][-1]["status"] == "STOP"


def test_run_document_processing_emits_segment_statuses_during_segmented_run_before_stopped():
    runtime = _build_runtime_capture()
    stop_checks = {"count": 0}

    def should_stop(runtime):
        stop_checks["count"] += 1
        return stop_checks["count"] >= 2

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {
                "target_text": "block-1",
                "context_before": "",
                "context_after": "",
                "target_chars": 7,
                "context_chars": 0,
                "segment_id": "seg_0001",
            },
            {
                "target_text": "block-2",
                "context_before": "",
                "context_after": "",
                "target_chars": 7,
                "context_chars": 0,
                "segment_id": "seg_0002",
            },
        ],
        source_paragraphs=cast(Any, [
            SimpleNamespace(segment_id="seg_0001", text="Chapter 1"),
            SimpleNamespace(segment_id="seg_0002", text="Chapter 2"),
        ]),
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=should_stop,
        generate_markdown_block=lambda **kwargs: "ok",
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=_reinsert_inline_images,
    )

    segment_status_events = [payload for payload in runtime["status"] if "segment_status_by_id" in payload]

    assert result == "stopped"
    assert segment_status_events[0]["segment_status_by_id"] == {"seg_0001": "processing", "seg_0002": "pending"}
    assert segment_status_events[1]["segment_status_by_id"] == {"seg_0001": "completed", "seg_0002": "pending"}


def test_run_document_processing_stops_after_image_phase_before_placeholder_validation(monkeypatch):
    runtime = _build_runtime_capture()
    stop_checks = {"count": 0}
    calls = {
        "image": 0,
        "image_integrity": 0,
        "late_validate": 0,
        "docx": 0,
        "artifacts": 0,
    }

    monkeypatch.setattr(
        document_pipeline,
        "_validate_placeholder_integrity_phase",
        lambda **kwargs: calls.__setitem__("late_validate", calls["late_validate"] + 1) or True,
    )

    def should_stop(runtime):
        stop_checks["count"] += 1
        return stop_checks["count"] >= 2

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {"target_text": "block-1", "context_before": "", "context_after": "", "target_chars": 7, "context_chars": 0},
        ],
        source_paragraphs=[],
        image_assets=[AssetStub("img_001")],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=should_stop,
        generate_markdown_block=lambda **kwargs: "ok",
        process_document_images=lambda **kwargs: calls.__setitem__("image", calls["image"] + 1) or kwargs["image_assets"],
        inspect_placeholder_integrity=lambda markdown_text, image_assets: calls.__setitem__("image_integrity", calls["image_integrity"] + 1) or {},
        convert_markdown_to_docx_bytes=lambda markdown_text: calls.__setitem__("docx", calls["docx"] + 1) or b"docx-bytes",
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=lambda docx_bytes, image_assets: docx_bytes,
        write_ui_result_artifacts=lambda **kwargs: calls.__setitem__("artifacts", calls["artifacts"] + 1) or {"markdown_path": "/tmp/final.result.md", "docx_path": "/tmp/final.result.docx"},
    )

    assert result == "stopped"
    assert calls == {
        "image": 1,
        "image_integrity": 1,
        "late_validate": 0,
        "docx": 0,
        "artifacts": 0,
    }
    assert runtime["finalize"][-1][0] == "Остановлено пользователем"
    assert runtime["finalize"][-1][3] == "stopped"
    assert runtime["log"][-1]["status"] == "STOP"


def test_run_document_processing_emits_active_segment_before_failed_terminal_result():
    runtime = _build_runtime_capture()

    result = document_pipeline.run_document_processing(
        uploaded_file="report.docx",
        jobs=[
            {
                "target_text": "block-1",
                "context_before": "",
                "context_after": "",
                "target_chars": 7,
                "context_chars": 0,
                "segment_id": "seg_0001",
            },
        ],
        source_paragraphs=cast(Any, [SimpleNamespace(segment_id="seg_0001", text="Chapter 1")]),
        image_assets=[],
        image_mode="safe",
        app_config={},
        model="gpt-5.4",
        max_retries=1,
        on_progress=lambda **kwargs: None,
        runtime=runtime,
        resolve_uploaded_filename=lambda uploaded_file: str(uploaded_file),
        get_client=lambda: object(),
        ensure_pandoc_available=lambda: None,
        load_system_prompt=lambda **_kw: "system",
        log_event=lambda *args, **kwargs: None,
        present_error=lambda code, exc, title, **kwargs: f"{title}: {exc}",
        emit_state=_emit_state,
        emit_finalize=_emit_finalize,
        emit_activity=_emit_activity,
        emit_log=_emit_log,
        emit_status=_emit_status,
        should_stop_processing=lambda runtime: False,
        generate_markdown_block=lambda **kwargs: (_ for _ in ()).throw(RuntimeError("boom")),
        process_document_images=lambda **kwargs: [],
        inspect_placeholder_integrity=_inspect_placeholder_integrity,
        convert_markdown_to_docx_bytes=_convert_markdown_to_docx_bytes,
        preserve_source_paragraph_properties=lambda docx_bytes, paragraphs, generated_paragraph_registry=None: docx_bytes,
        reinsert_inline_images=_reinsert_inline_images,
    )

    assert result == "failed"
    assert runtime["status"][0]["active_segment_id"] == "seg_0001"
    assert runtime["status"][0]["segment_status_by_id"] == {"seg_0001": "processing"}
    assert runtime["finalize"][-1][3] == "error"


def test_run_document_processing_persists_failed_job_result_records(monkeypatch):
    runtime = _build_runtime_capture()
    captured = {}

    def write_job_result_registry(*, records):
        captured.setdefault("job_records", []).extend(list(records))
        return {"job_0000": "/tmp/job-results/job_0000.job-result.json"}

    monkeypatch.setattr(document_pipeline, "write_job_result_registry_impl", write_job_result_registry)

    result = _run_processing(
        runtime,
        prepared_source_key="prep:report:1234",
        structure_fingerprint="struct-abc",
        jobs=[
            {
                "job_id": "job_0000",
                "segment_id": "seg_0001",
                "target_text": "block",
                "context_before": "",
                "context_after": "",
                "target_chars": 5,
                "context_chars": 0,
            }
        ],
        source_paragraphs=cast(Any, [SimpleNamespace(segment_id="seg_0001", text="Chapter 1")]),
        generate_markdown_block=lambda **kwargs: (_ for _ in ()).throw(RuntimeError("boom")),
    )

    assert result == "failed"
    assert len(captured["job_records"]) == 1
    failed_record = dict(captured["job_records"][0])
    assert isinstance(failed_record.pop("updated_at", ""), str)
    assert failed_record == {
        "schema_version": 1,
        "prepared_source_key": "prep:report:1234",
        "structure_fingerprint": "struct-abc",
        "job_id": "job_0000",
        "segment_id": "seg_0001",
        "status": "failed",
        "block_index": 1,
        "target_chars": 5,
        "context_chars": 0,
        "error_code": "block_failed",
        "error_message": "Ошибка обработки блока: boom",
    }


def test_run_document_processing_end_to_end_produces_openable_docx_artifact(tmp_path):
    image_path = tmp_path / "pipeline-image.png"
    image_path.write_bytes(PNG_BYTES)

    source_doc = Document()
    source_doc.add_heading("Глава", level=1)
    source_doc.add_paragraph().add_run().add_picture(str(image_path))
    source_doc.add_paragraph("Рисунок 1. Подпись")
    source_doc.add_paragraph("Исходный абзац")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)
    source_buffer.seek(0)

    source_paragraphs, image_assets = extract_document_content_from_docx(source_buffer)
    runtime = _build_runtime_capture()
    final_markdown = "Глава\n\n[[DOCX_IMAGE_img_001]]\n\nРисунок 1. Подпись\n\nОбновленный абзац"

    def build_docx_from_markdown(markdown_text):
        doc = Document()
        for block in markdown_text.split("\n\n"):
            doc.add_paragraph(block)
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    result = _run_processing(
        runtime,
        jobs=[{"target_text": "Исходный блок", "context_before": "", "context_after": "", "target_chars": 13, "context_chars": 0}],
        source_paragraphs=source_paragraphs,
        image_assets=image_assets,
        generate_markdown_block=lambda **kwargs: final_markdown,
        process_document_images=lambda **kwargs: image_assets,
        convert_markdown_to_docx_bytes=build_docx_from_markdown,
        preserve_source_paragraph_properties=__import__(
            "docxaicorrector.generation.formatting_transfer", fromlist=["preserve_source_paragraph_properties"]
        ).preserve_source_paragraph_properties,
        reinsert_inline_images=__import__(
            "docxaicorrector.image.reinsertion", fromlist=["reinsert_inline_images"]
        ).reinsert_inline_images,
    )

    assert result == "succeeded"
    assert runtime["state"]["latest_markdown"] == final_markdown
    assert runtime["state"]["latest_docx_bytes"]

    output_doc = Document(BytesIO(runtime["state"]["latest_docx_bytes"]))
    visible_text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)

    assert output_doc.paragraphs[0].style is not None
    assert output_doc.paragraphs[0].style.name == "Normal"
    assert output_doc.paragraphs[2].style is not None
    assert output_doc.paragraphs[2].style.name == "Caption"
    assert "Обновленный абзац" in visible_text
    assert "[[DOCX_IMAGE_img_001]]" not in output_doc._element.xml
    assert len(output_doc.inline_shapes) == 1
    assert runtime["finalize"][-1][0] == "Обработка завершена"
    assert runtime["log"][-1]["status"] == "DONE"


def test_emit_mapping_text_quality_defect_items_surfaces_bad_pairs():
    items: list[dict] = []
    document_pipeline_late_phases._emit_mapping_text_quality_defect_items(
        formatting_review_items=items,
        mapping_text_quality={
            "bad_pair_count": 3,
            "samples": [
                {
                    "source_text_preview": "Original source paragraph",
                    "target_text_preview": "Совсем другой перевод",
                },
                {
                    "source_text_preview": "Second source",
                    "target_text_preview": "Второй мимо",
                },
            ],
        },
    )

    assert len(items) == 2
    assert all(item["severity"] == "defect" for item in items)
    assert all(item["reason"] == "mapping_text_quality_bad_pair" for item in items)
    # 3 bad pairs but only 2 samples rendered -> aggregate rides on the first item.
    assert items[0]["aggregate_count"] == 3
    assert items[0]["sample"]["source_text"] == "Original source paragraph"
    assert items[0]["sample"]["text"] == "Совсем другой перевод"


def test_emit_mapping_text_quality_defect_items_noop_without_bad_pairs():
    items: list[dict] = []
    document_pipeline_late_phases._emit_mapping_text_quality_defect_items(
        formatting_review_items=items,
        mapping_text_quality={"bad_pair_count": 0, "samples": []},
    )
    document_pipeline_late_phases._emit_mapping_text_quality_defect_items(
        formatting_review_items=items,
        mapping_text_quality=None,
    )
    assert items == []
