import io
import logging
import threading
import zipfile
from types import SimpleNamespace
from typing import Any, cast

import lxml.etree as etree
import pytest
from docx import Document
from docx.document import Document as DocxDocument
from docx.styles.style import ParagraphStyle, _TableStyle

from PIL import Image

import docxaicorrector.core.model_accounting as model_accounting
import docxaicorrector.generation._generation as generation
import docxaicorrector.image.shared as image_shared
from docxaicorrector.image.generation import _normalize_generated_document_background

_THEME_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"
_DRAWINGML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def _as_openai_client(client: object) -> Any:
    return cast(Any, client)


def _batched_answer_only_client(marker_answer: str, *, requests: list[dict[str, Any]] | None = None) -> Any:
    """A model that answers the BATCHED marker request — badly — and nothing else.

    Needed by every test whose subject is the block-level source fallback. Since the
    degradation ladder re-asks each paragraph on its own with ``marker_mode=False``, a stub
    that returns the same string for every request makes the ladder SUCCEED, and the test
    stops describing the fallback it exists for. Answering nothing to the per-paragraph
    requests holds the ladder at its floor instead: it runs, it rescues nothing, and the
    block ends exactly where it ended before the ladder existed — which is the invariant
    those tests should be pinning, because the ladder is only allowed to ADD outcomes.

    The two request shapes are told apart by the markers in the prompt, not by a call
    counter: the ladder's units are marker-free by construction.
    """

    def create_response(**kwargs: Any) -> SimpleNamespace:
        if requests is not None:
            requests.append(dict(kwargs))
        prompt = kwargs["input"][1]["content"][0]["text"]
        if "[[DOCX_PARA_" in prompt:
            return SimpleNamespace(output_text=marker_answer)
        return SimpleNamespace(output_text="")

    return _as_openai_client(SimpleNamespace(responses=SimpleNamespace(create=create_response)))


def _as_paragraph_style(style: object) -> ParagraphStyle:
    return cast(ParagraphStyle, style)


def _as_table_style(style: object) -> _TableStyle:
    return cast(_TableStyle, style)


def _pt(value: object) -> float:
    return cast(Any, value).pt


def _numbering_level_signature(level: Any) -> dict[str, str | None]:
    level_xml = cast(Any, level)
    return {
        "num_fmt": _first_xpath_value(level_xml, './*[local-name()="numFmt"]/@*[local-name()="val"]'),
        "lvl_text": _first_xpath_value(level_xml, './*[local-name()="lvlText"]/@*[local-name()="val"]'),
        "left": _first_xpath_value(level_xml, './*[local-name()="pPr"]/*[local-name()="ind"]/@*[local-name()="left"]'),
        "hanging": _first_xpath_value(level_xml, './*[local-name()="pPr"]/*[local-name()="ind"]/@*[local-name()="hanging"]'),
        "after": _first_xpath_value(level_xml, './*[local-name()="pPr"]/*[local-name()="spacing"]/@*[local-name()="after"]'),
        "line": _first_xpath_value(level_xml, './*[local-name()="pPr"]/*[local-name()="spacing"]/@*[local-name()="line"]'),
        "line_rule": _first_xpath_value(level_xml, './*[local-name()="pPr"]/*[local-name()="spacing"]/@*[local-name()="lineRule"]'),
        "ascii_font": _first_xpath_value(level_xml, './*[local-name()="rPr"]/*[local-name()="rFonts"]/@*[local-name()="ascii"]'),
        "hansi_font": _first_xpath_value(level_xml, './*[local-name()="rPr"]/*[local-name()="rFonts"]/@*[local-name()="hAnsi"]'),
        "cs_font": _first_xpath_value(level_xml, './*[local-name()="rPr"]/*[local-name()="rFonts"]/@*[local-name()="cs"]'),
    }


def _first_xpath_value(node: Any, expression: str) -> str | None:
    values = cast(list[str], node.xpath(expression))
    return values[0] if values else None


def _find_matching_abstract_numbers(
    document: Any,
    *,
    num_fmt: str,
    level_texts: tuple[str, ...],
    body_font: str | None = None,
) -> list[Any]:
    numbering = document.part.numbering_part.element
    matches = []
    for abstract_num in numbering.xpath('./*[local-name()="abstractNum"]'):
        levels = cast(list[Any], abstract_num.xpath('./*[local-name()="lvl"]'))
        if len(levels) != len(level_texts):
            continue
        signatures = [_numbering_level_signature(level) for level in levels]
        expected_signatures = [
            {
                "num_fmt": num_fmt,
                "lvl_text": level_text,
                "left": str(720 + (index * 360)),
                "hanging": "360",
                "after": "80",
                "line": "264",
                "line_rule": "auto",
                "ascii_font": body_font,
                "hansi_font": body_font,
                "cs_font": body_font,
            }
            for index, level_text in enumerate(level_texts)
        ]
        if signatures == expected_signatures:
            matches.append(abstract_num)
    return matches


def _has_num_instance_for_abstract_num(document: Any, abstract_num: Any) -> bool:
    numbering = document.part.numbering_part.element
    abstract_num_id = cast(str | None, abstract_num.get(generation.qn("w:abstractNumId")))
    if abstract_num_id is None:
        return False
    return bool(
        numbering.xpath(
            f'./*[local-name()="num"]/*[local-name()="abstractNumId" and @*[local-name()="val"]="{abstract_num_id}"]'
        )
    )


def _paragraph_num_id(paragraph: Any) -> str | None:
    return _first_xpath_value(
        paragraph._p,
        './*[local-name()="pPr"]/*[local-name()="numPr"]/*[local-name()="numId"]/@*[local-name()="val"]',
    )


def _paragraph_ilvl(paragraph: Any) -> str | None:
    return _first_xpath_value(
        paragraph._p,
        './*[local-name()="pPr"]/*[local-name()="numPr"]/*[local-name()="ilvl"]/@*[local-name()="val"]',
    )


def _find_abstract_num_for_num_id(document: Any, num_id: str) -> Any | None:
    numbering = document.part.numbering_part.element
    abstract_num_ids = cast(
        list[Any],
        numbering.xpath(
            f'./*[local-name()="num" and @*[local-name()="numId"]="{num_id}"]/*[local-name()="abstractNumId"]/@*[local-name()="val"]'
        ),
    )
    if not abstract_num_ids:
        return None

    abstract_num_id = cast(str, abstract_num_ids[0])
    abstract_nums = cast(
        list[Any],
        numbering.xpath(f'./*[local-name()="abstractNum" and @*[local-name()="abstractNumId"]="{abstract_num_id}"]'),
    )
    return abstract_nums[0] if abstract_nums else None


def _pandoc_available() -> bool:
    generation.ensure_pandoc_available.cache_clear()
    try:
        generation.ensure_pandoc_available()
    except RuntimeError:
        return False
    finally:
        generation.ensure_pandoc_available.cache_clear()
    return True


class RetryableError(Exception):
    status_code = 429


def test_generate_markdown_block_retries_once_then_returns(monkeypatch):
    attempts = []
    sleep_calls = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        if len(attempts) == 1:
            raise RetryableError("rate limited")
        return SimpleNamespace(output_text="```markdown\nИсправленный текст\n```")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="target",
        context_before="   ",
        context_after="\n\t",
        max_retries=2,
    )

    assert result == "Исправленный текст"
    assert len(attempts) == 2
    assert sleep_calls == [1]
    user_payload = attempts[0]["input"][1]["content"][0]["text"]
    assert "[CONTEXT BEFORE]\n[no context]" in user_payload
    assert "[TARGET BLOCK]\ntarget" in user_payload
    assert "[CONTEXT AFTER]\n[no context]" in user_payload
    assert attempts[0]["max_output_tokens"] >= 512


def test_generate_markdown_block_uses_anthropic_messages_api():
    attempts = []

    def create_message(**kwargs):
        attempts.append(dict(kwargs))
        return SimpleNamespace(content=[SimpleNamespace(text="```markdown\nИсправленный текст\n```")])

    client = SimpleNamespace(messages=SimpleNamespace(create=create_message))

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="anthropic:claude-sonnet-4.6",
        system_prompt="system",
        target_text="target",
        context_before="before",
        context_after="after",
        max_retries=1,
    )

    assert result == "Исправленный текст"
    assert attempts == [
        {
            "model": "claude-sonnet-4-6",
            "messages": [
                {
                    "role": "user",
                    "content": "Below is a target document block and surrounding context.\n"
                    "Use the surrounding context only to understand meaning, terminology, and continuity.\n"
                    "Process only the target block according to the system instructions and return only its final text.\n\n"
                    "[CONTEXT BEFORE]\nbefore\n\n[TARGET BLOCK]\ntarget\n\n[CONTEXT AFTER]\nafter",
                }
            ],
            "max_tokens": 512,
            "system": "system",
            "temperature": 0.4,
        }
    ]


def test_generate_markdown_block_retries_on_empty_response(monkeypatch):
    attempts = []
    sleep_calls = []
    logged_events = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        if len(attempts) == 1:
            return SimpleNamespace(output_text="")
        return SimpleNamespace(output_text="Исправленный текст")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-empty",
    )

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="target",
        context_before="before",
        context_after="after",
        max_retries=2,
    )

    assert result == "Исправленный текст"
    assert len(attempts) == 2
    assert sleep_calls == [1]
    assert len(logged_events) == 1
    assert logged_events[0][0][1] == "model_empty_response_shape"
    assert logged_events[0][1]["error_code"] == "empty_response"


def test_generate_markdown_block_retries_on_incomplete_response(monkeypatch):
    attempts = []
    sleep_calls = []
    logged_events = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        if len(attempts) == 1:
            return SimpleNamespace(status="incomplete", output=[SimpleNamespace(type="reasoning", status="incomplete")])
        return SimpleNamespace(status="completed", output_text="Исправленный текст")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-incomplete",
    )

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="target",
        context_before="before",
        context_after="after",
        max_retries=2,
    )

    assert result == "Исправленный текст"
    assert len(attempts) == 2
    assert sleep_calls == [1]
    assert attempts[0]["max_output_tokens"] == 512
    assert attempts[1]["max_output_tokens"] == 1024
    assert logged_events[0][0][1] == "model_empty_response_shape"
    assert logged_events[0][1]["error_code"] == "incomplete_response"


def test_generate_markdown_block_uses_degraded_prompt_after_persistent_empty_response(monkeypatch):
    attempts = []
    sleep_calls = []
    logged_events = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        if len(attempts) <= 2:
            return SimpleNamespace(output_text="")
        return SimpleNamespace(output_text="Восстановленный текст")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-empty-recovery",
    )

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="target",
        context_before="before",
        context_after="after",
        max_retries=2,
    )

    assert result == "Восстановленный текст"
    assert len(attempts) == 3
    assert sleep_calls == [1]
    assert "[TARGET BLOCK ONLY]\ntarget" in attempts[-1]["input"][1]["content"][0]["text"]
    assert "[CONTEXT BEFORE]" not in attempts[-1]["input"][1]["content"][0]["text"]
    assert logged_events[-1][0][1] == "markdown_empty_response_recovery_started"


def test_generate_markdown_block_raises_on_empty_model_output(monkeypatch):
    logged_events = []
    client = SimpleNamespace(
        responses=SimpleNamespace(create=lambda **_: SimpleNamespace(output_text="```\n\n```"))
    )
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-collapsed",
    )

    try:
        generation.generate_markdown_block(
            client=_as_openai_client(client),
            model="gpt-5.4",
            system_prompt="system",
            target_text="target",
            context_before="before",
            context_after="after",
            max_retries=1,
        )
    except RuntimeError as exc:
        assert "collapsed_output" in str(exc)
        assert logged_events[0][1]["error_code"] == "collapsed_output"
    else:
        raise AssertionError("Expected RuntimeError for a collapsed model response")


def test_generate_markdown_block_retries_on_collapsed_output(monkeypatch):
    attempts = []
    sleep_calls = []
    logged_events = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        if len(attempts) == 1:
            return SimpleNamespace(output_text="```markdown\n   \n```")
        return SimpleNamespace(output_text="Итоговый текст")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-collapsed-retry",
    )

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="target",
        context_before="before",
        context_after="after",
        max_retries=2,
    )

    assert result == "Итоговый текст"
    assert len(attempts) == 2
    assert sleep_calls == [1]
    assert len(logged_events) == 1
    assert logged_events[0][1]["error_code"] == "collapsed_output"


def test_generate_markdown_block_retries_without_max_output_tokens_when_sdk_rejects_it():
    calls = []

    def create_response(**kwargs):
        calls.append(dict(kwargs))
        if len(calls) == 1:
            raise TypeError("unexpected keyword argument 'max_output_tokens'")
        return SimpleNamespace(output_text="ok")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="target",
        context_before="before",
        context_after="after",
        max_retries=1,
    )

    assert result == "ok"
    assert "max_output_tokens" in calls[0]
    assert "max_output_tokens" not in calls[1]


def test_generate_markdown_block_retries_without_temperature_when_model_rejects_it():
    calls = []

    class UnsupportedTemperatureError(Exception):
        status_code = 400

    def create_response(**kwargs):
        calls.append(dict(kwargs))
        if len(calls) == 1:
            raise UnsupportedTemperatureError("Unsupported parameter: 'temperature' is not supported with this model.")
        return SimpleNamespace(output_text="ok")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5-mini",
        system_prompt="system",
        target_text="target",
        context_before="before",
        context_after="after",
        max_retries=1,
    )

    assert result == "ok"
    assert "temperature" in calls[0]
    assert "temperature" not in calls[1]


def test_generate_markdown_block_falls_back_to_chat_completions_for_openrouter_compatibility(monkeypatch):
    responses_calls = []
    chat_calls = []
    logged_events = []

    class UnsupportedResponsesError(Exception):
        status_code = 400

    def responses_create(**kwargs):
        responses_calls.append(dict(kwargs))
        raise UnsupportedResponsesError("Responses API unsupported for this model; use chat.completions instead.")

    def chat_create(**kwargs):
        chat_calls.append(dict(kwargs))
        return SimpleNamespace(
            choices=[
                SimpleNamespace(
                    message=SimpleNamespace(content="fallback-ok"),
                )
            ]
        )

    client = SimpleNamespace(
        base_url="https://openrouter.ai/api/v1",
        responses=SimpleNamespace(create=responses_create),
        chat=SimpleNamespace(completions=SimpleNamespace(create=chat_create)),
    )
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-openrouter-fallback",
    )

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="google/gemini-3.1-flash-lite-preview",
        system_prompt="system",
        target_text="target",
        context_before="before",
        context_after="after",
        max_retries=1,
    )

    assert result == "fallback-ok"
    assert len(responses_calls) == 1
    assert len(chat_calls) == 1
    assert chat_calls[0]["model"] == "google/gemini-3.1-flash-lite-preview"
    assert chat_calls[0]["messages"] == [
        {"role": "system", "content": "system"},
        {"role": "user", "content": "Below is a target document block and surrounding context.\nUse the surrounding context only to understand meaning, terminology, and continuity.\nProcess only the target block according to the system instructions and return only its final text.\n\n[CONTEXT BEFORE]\nbefore\n\n[TARGET BLOCK]\ntarget\n\n[CONTEXT AFTER]\nafter"},
    ]
    assert logged_events[-1][0][1] == "provider_text_api_fallback_engaged"
    assert logged_events[-1][1]["canonical_model_selector"] == "openrouter:google/gemini-3.1-flash-lite-preview"


def test_generate_markdown_block_does_not_fallback_on_non_compatibility_openrouter_error():
    class UnauthorizedError(Exception):
        status_code = 401

    def responses_create(**kwargs):
        raise UnauthorizedError("Unauthorized provider call")

    client = SimpleNamespace(
        base_url="https://openrouter.ai/api/v1",
        responses=SimpleNamespace(create=responses_create),
        chat=SimpleNamespace(completions=SimpleNamespace(create=lambda **kwargs: (_ for _ in ()).throw(AssertionError("chat fallback must not run")))),
    )

    with pytest.raises(UnauthorizedError):
        generation.generate_markdown_block(
            client=_as_openai_client(client),
            model="google/gemini-3.1-flash-lite-preview",
            system_prompt="system",
            target_text="target",
            context_before="before",
            context_after="after",
            max_retries=1,
        )


def test_generate_markdown_block_does_not_fallback_on_generic_invalid_input_error():
    class InvalidInputError(Exception):
        status_code = 400

    def responses_create(**kwargs):
        raise InvalidInputError("Invalid input payload")

    client = SimpleNamespace(
        base_url="https://openrouter.ai/api/v1",
        responses=SimpleNamespace(create=responses_create),
        chat=SimpleNamespace(completions=SimpleNamespace(create=lambda **kwargs: (_ for _ in ()).throw(AssertionError("chat fallback must not run")))),
    )

    with pytest.raises(InvalidInputError):
        generation.generate_markdown_block(
            client=_as_openai_client(client),
            model="google/gemini-3.1-flash-lite-preview",
            system_prompt="system",
            target_text="target",
            context_before="before",
            context_after="after",
            max_retries=1,
        )


def test_generate_markdown_block_falls_back_to_source_after_persistent_empty_response(monkeypatch):
    attempts = []
    sleep_calls = []
    logged_events = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        return SimpleNamespace(output_text="")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-persistent-empty",
    )

    target_text = "Короткий исходный абзац, который должен сохраниться без падения пайплайна."
    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=target_text,
        context_before="before",
        context_after="after",
        max_retries=3,
    )

    assert result == target_text
    assert len(attempts) == 4
    assert sleep_calls == [1, 2]
    assert any(args[1] == "markdown_empty_response_recovery_started" for args, _ in logged_events)
    assert logged_events[-1][0][1] == "markdown_empty_response_source_fallback"


def test_generate_markdown_block_falls_back_when_recovery_returns_blank_markdown(monkeypatch):
    attempts = []
    sleep_calls = []
    logged_events = []

    def call_markdown_request_with_sdk_fallback(client, request_kwargs):
        attempts.append(dict(request_kwargs))
        if len(attempts) <= 2:
            raise RuntimeError("Модель вернула пустой ответ (empty_response).")
        return "", False

    client = SimpleNamespace(responses=SimpleNamespace(create=lambda **kwargs: SimpleNamespace(output_text="unused")))
    monkeypatch.setattr(generation, "_call_markdown_request_with_sdk_fallback", call_markdown_request_with_sdk_fallback)
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-blank-recovery",
    )

    target_text = "Исходный абзац должен сохраниться, если recovery вернул пустой markdown."
    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=target_text,
        context_before="before",
        context_after="after",
        max_retries=2,
    )

    assert result == target_text
    assert len(attempts) == 3
    assert sleep_calls == [1]
    assert any(args[1] == "markdown_empty_response_recovery_started" for args, _ in logged_events)
    assert logged_events[-1][0][1] == "markdown_empty_response_source_fallback"


def test_generate_markdown_block_falls_back_to_source_after_persistent_incomplete_response(monkeypatch):
    attempts = []
    sleep_calls = []
    logged_events = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        return SimpleNamespace(status="incomplete", output=[SimpleNamespace(type="reasoning", status="incomplete")])

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-persistent-incomplete",
    )

    target_text = "Короткий исходный абзац, который должен сохраниться без падения пайплайна."
    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=target_text,
        context_before="before",
        context_after="after",
        max_retries=2,
    )

    assert result == target_text
    assert len(attempts) == 3
    assert sleep_calls == [1]
    assert attempts[0]["max_output_tokens"] == 512
    assert attempts[1]["max_output_tokens"] == 1024
    assert attempts[2]["max_output_tokens"] == 1536
    assert any(args[1] == "markdown_empty_response_recovery_started" for args, _ in logged_events)
    assert logged_events[-1][0][1] == "markdown_incomplete_response_source_fallback"


def test_generate_markdown_block_falls_back_to_source_after_persistent_incomplete_response_for_long_block(monkeypatch):
    attempts = []
    logged_events = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        return SimpleNamespace(status="incomplete", output=[SimpleNamespace(type="reasoning", status="incomplete")])

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", lambda *_: None)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-persistent-incomplete-long",
    )

    target_text = "Длинный блок. " * 150
    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=target_text,
        context_before="before",
        context_after="after",
        max_retries=2,
    )

    assert result == target_text
    # The block is one paragraph of 150 sentences, so ``incomplete_response`` now buys a
    # sentence split before the source is substituted, and the calls are 3 + 3 + 3: the
    # block itself, the paragraph re-asked whole, and the FIRST sentence group. The ladder
    # abandons the paragraph as soon as one of its groups fails, which is why the second
    # group is never bought — an upper bound that is arithmetic, not a hope.
    assert len(attempts) == 9
    assert logged_events[-1][0][1] == "markdown_incomplete_response_source_fallback"
    # The ladder rescued nothing, so the delivered text, the event and the counter are
    # byte-for-byte what they were before it existed.
    assert any(args[1] == "degradation_ladder_completed" for args, _ in logged_events)


def test_generate_markdown_block_passthrough_for_image_only_target(monkeypatch):
    logged_events = []
    client = SimpleNamespace(responses=SimpleNamespace(create=lambda **_: (_ for _ in ()).throw(AssertionError("must not call API"))))
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-image-only",
    )

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_IMAGE_img_001]]",
        context_before="before",
        context_after="after",
        max_retries=1,
    )

    assert result == "[[DOCX_IMAGE_img_001]]"
    assert logged_events[0][0][1] == "image_only_target_passthrough"


def test_generate_markdown_block_passthrough_for_placeholder_only_marker_target(monkeypatch):
    logged_events = []
    client = SimpleNamespace(responses=SimpleNamespace(create=lambda **_: (_ for _ in ()).throw(AssertionError("must not call API"))))
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-placeholder-only",
    )

    target_text = "[[DOCX_PARA_p0001]]\n[[DOCX_IMAGE_img_001]]"
    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=target_text,
        context_before="before",
        context_after="after",
        max_retries=1,
        expected_paragraph_ids=["p0001"],
        marker_mode=True,
    )

    assert result == target_text
    assert logged_events[0][0][1] == "image_only_target_passthrough"


def test_generate_markdown_block_falls_back_to_source_on_missing_output_text(monkeypatch):
    logged_events = []
    client = SimpleNamespace(
        responses=SimpleNamespace(create=lambda **_: SimpleNamespace())
    )
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-missing-output",
    )

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="target",
        context_before="before",
        context_after="after",
        max_retries=1,
    )

    assert result == "target"
    assert logged_events[0][1]["error_code"] == "empty_response"
    assert any(args[1] == "markdown_empty_response_recovery_started" for args, _ in logged_events)
    assert logged_events[-1][0][1] == "markdown_empty_response_source_fallback"


def test_extract_response_output_text_falls_back_to_supported_response_output_items():
    response = SimpleNamespace(
        output=[
            SimpleNamespace(
                content=[SimpleNamespace(type="output_text", text="Структурированный ответ")]
            )
        ]
    )

    assert generation._extract_response_output_text(response) == "Структурированный ответ"


def test_extract_response_output_text_reads_supported_nested_text_value_from_response_output():
    response = SimpleNamespace(
        output=[
            SimpleNamespace(
                content=[
                    SimpleNamespace(
                        type="output_text",
                        text=SimpleNamespace(value="Ответ из value-поля"),
                    )
                ]
            )
        ]
    )

    assert generation._extract_response_output_text(response) == "Ответ из value-поля"


def test_generate_markdown_block_raises_on_unsupported_response_shape_in_output_items(monkeypatch):
    sleep_calls = []
    attempts = []

    def create_response(**_):
        attempts.append("call")
        return SimpleNamespace(
            output=[SimpleNamespace(content=[SimpleNamespace(type="refusal", text="not supported")])]
        )

    client = SimpleNamespace(
        responses=SimpleNamespace(create=create_response)
    )
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)

    try:
        generation.generate_markdown_block(
            client=_as_openai_client(client),
            model="gpt-5.4",
            system_prompt="system",
            target_text="target",
            context_before="before",
            context_after="after",
            max_retries=1,
        )
    except RuntimeError as exc:
        assert "unsupported_response_shape" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for unsupported response output shape")

    assert attempts == ["call"]
    assert sleep_calls == []


def test_generate_markdown_block_raises_when_supported_response_output_collapses_after_normalization(monkeypatch):
    logged_events = []
    client = SimpleNamespace(
        responses=SimpleNamespace(
            create=lambda **_: SimpleNamespace(
                output=[SimpleNamespace(content=[SimpleNamespace(type="output_text", text="```markdown\n   \n```")])]
            )
        )
    )
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-supported-collapse",
    )

    try:
        generation.generate_markdown_block(
            client=_as_openai_client(client),
            model="gpt-5.4",
            system_prompt="system",
            target_text="target",
            context_before="before",
            context_after="after",
            max_retries=1,
        )
    except RuntimeError as exc:
        assert "collapsed_output" in str(exc)
        assert logged_events[0][1]["error_code"] == "collapsed_output"
    else:
        raise AssertionError("Expected RuntimeError when normalized fallback output collapses")


def test_generate_markdown_block_raises_on_non_string_output_text():
    client = SimpleNamespace(
        responses=SimpleNamespace(create=lambda **_: SimpleNamespace(output_text=["invalid"]))
    )

    try:
        generation.generate_markdown_block(
            client=_as_openai_client(client),
            model="gpt-5.4",
            system_prompt="system",
            target_text="target",
            context_before="before",
            context_after="after",
            max_retries=1,
        )
    except RuntimeError as exc:
        assert "неподдерживаемом формате" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError when output_text is not a string")


def test_generate_markdown_block_rejects_max_retries_less_than_one():
    client = SimpleNamespace(
        responses=SimpleNamespace(create=lambda **_: SimpleNamespace(output_text="unused"))
    )

    try:
        generation.generate_markdown_block(
            client=_as_openai_client(client),
            model="gpt-5.4",
            system_prompt="system",
            target_text="target",
            context_before="before",
            context_after="after",
            max_retries=0,
        )
    except ValueError as exc:
        assert "max_retries" in str(exc)
    else:
        raise AssertionError("Expected ValueError when max_retries is less than 1")


def test_generate_markdown_block_rejects_non_integer_max_retries():
    client = SimpleNamespace(
        responses=SimpleNamespace(create=lambda **_: SimpleNamespace(output_text="unused"))
    )

    try:
        generation.generate_markdown_block(
            client=_as_openai_client(client),
            model="gpt-5.4",
            system_prompt="system",
            target_text="target",
            context_before="before",
            context_after="after",
            max_retries=cast(int, 1.5),
        )
    except TypeError as exc:
        assert "max_retries" in str(exc)
    else:
        raise AssertionError("Expected TypeError when max_retries is not an integer")


def test_extract_normalized_markdown_logs_empty_response_shape(monkeypatch):
    logged_events = []
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-empty-shape",
    )

    try:
        generation._extract_normalized_markdown(SimpleNamespace())
    except RuntimeError as exc:
        assert "empty_response" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for empty response shape")

    assert len(logged_events) == 1
    args, kwargs = logged_events[0]
    assert args[1] == "model_empty_response_shape"
    assert kwargs["error_code"] == "empty_response"
    assert kwargs["raw_output_len"] == 0


def test_extract_normalized_markdown_raises_on_incomplete_response(monkeypatch):
    logged_events = []
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-incomplete-shape",
    )

    try:
        generation._extract_normalized_markdown(
            SimpleNamespace(status="incomplete", output=[SimpleNamespace(type="reasoning", status="incomplete")])
        )
    except RuntimeError as exc:
        assert "incomplete_response" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for incomplete response")

    assert len(logged_events) == 1
    args, kwargs = logged_events[0]
    assert args[1] == "model_empty_response_shape"
    assert kwargs["error_code"] == "incomplete_response"


def test_extract_normalized_markdown_raises_hard_on_non_completed_response(monkeypatch):
    logged_events = []
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-non-completed-shape",
    )

    try:
        generation._extract_normalized_markdown(SimpleNamespace(status="failed"))
    except RuntimeError as exc:
        assert "non_completed_response" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for non-completed response")

    assert len(logged_events) == 1
    args, kwargs = logged_events[0]
    assert args[1] == "model_empty_response_shape"
    assert kwargs["error_code"] == "non_completed_response"


def test_incomplete_response_is_retryable():
    assert generation._is_retryable_empty_generation_error(RuntimeError("incomplete_response")) is True


def test_non_completed_response_is_not_retryable():
    assert generation._is_retryable_empty_generation_error(RuntimeError("non_completed_response")) is False


def test_generate_markdown_block_falls_back_to_source_after_persistent_non_completed_response(monkeypatch):
    # SC-001: persistent non_completed_response with non-empty target_text returns the source text.
    attempts = []
    sleep_calls = []
    logged_events = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        return SimpleNamespace(status="failed")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-persistent-non-completed",
    )

    target_text = "Короткий исходный абзац, который должен сохраниться без падения пайплайна."
    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=target_text,
        context_before="before",
        context_after="after",
        max_retries=3,
    )

    assert result == target_text
    assert len(attempts) == 3  # bounded by max_retries, no extra recovery call
    assert sleep_calls == [1, 2]
    # non_completed does NOT route through the empty-response recovery re-attempt
    assert not any(args[1] == "markdown_empty_response_recovery_started" for args, _ in logged_events)
    assert logged_events[-1][0][1] == "markdown_non_completed_response_source_fallback"


def test_generate_markdown_block_retries_then_succeeds_on_non_completed_response(monkeypatch):
    # SC-002: transient non_completed_response is retried within the bounded loop, then succeeds.
    attempts = []
    sleep_calls = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        if len(attempts) < 3:
            return SimpleNamespace(status="failed")
        return SimpleNamespace(status="completed", output_text="Исправленный текст")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(generation, "log_event", lambda *args, **kwargs: "evt-transient-non-completed")

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="target",
        context_before="before",
        context_after="after",
        max_retries=3,
    )

    assert result == "Исправленный текст"
    assert len(attempts) == 3  # proves it retried, did not fall back immediately
    assert sleep_calls == [1, 2]


def test_generate_markdown_block_raises_on_non_completed_response_with_empty_target(monkeypatch):
    # SC-003a (anti-vacuum): non_completed_response with empty substrate still hard-fails.
    # Bypass the empty-target passthrough so the error path is actually exercised.
    monkeypatch.setattr(generation, "_should_passthrough_target", lambda _target: False)
    monkeypatch.setattr(generation.time, "sleep", lambda *_: None)
    monkeypatch.setattr(generation, "log_event", lambda *args, **kwargs: "evt-empty-substrate")

    client = SimpleNamespace(responses=SimpleNamespace(create=lambda **_: SimpleNamespace(status="failed")))

    with pytest.raises(RuntimeError) as excinfo:
        generation.generate_markdown_block(
            client=_as_openai_client(client),
            model="gpt-5.4",
            system_prompt="system",
            target_text="   ",
            context_before="before",
            context_after="after",
            max_retries=2,
        )
    assert "non_completed_response" in str(excinfo.value)


def test_non_completed_response_empty_substrate_guard_rejects_fallback():
    # SC-003a (unit): the fallback guard refuses empty/whitespace substrate.
    assert generation._can_fallback_to_source_text_after_non_completed_response("") is False
    assert generation._can_fallback_to_source_text_after_non_completed_response("   ") is False
    assert generation._can_fallback_to_source_text_after_non_completed_response("target") is True


def test_generate_markdown_block_still_raises_on_auth_error_path():
    # SC-003b (anti-vacuum): an auth/SDK error never yields a status-bearing response, so the
    # non_completed_response fallback must NOT swallow it — the run still hard-fails.
    class UnauthorizedError(Exception):
        status_code = 401

    def responses_create(**kwargs):
        raise UnauthorizedError("Unauthorized provider call")

    client = SimpleNamespace(
        base_url="https://openrouter.ai/api/v1",
        responses=SimpleNamespace(create=responses_create),
        chat=SimpleNamespace(completions=SimpleNamespace(create=lambda **kwargs: (_ for _ in ()).throw(AssertionError("chat fallback must not run")))),
    )

    with pytest.raises(UnauthorizedError):
        generation.generate_markdown_block(
            client=_as_openai_client(client),
            model="google/gemini-3.1-flash-lite-preview",
            system_prompt="system",
            target_text="target",
            context_before="before",
            context_after="after",
            max_retries=1,
        )


def test_extract_normalized_markdown_logs_collapsed_output_shape(monkeypatch):
    logged_events = []
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-collapsed-shape",
    )

    try:
        generation._extract_normalized_markdown(SimpleNamespace(output_text="```markdown\n\n```") )
    except RuntimeError as exc:
        assert "collapsed_output" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for collapsed response shape")

    assert len(logged_events) == 1
    args, kwargs = logged_events[0]
    assert args[1] == "model_empty_response_shape"
    assert kwargs["error_code"] == "collapsed_output"
    assert kwargs["raw_output_len"] > 0


def test_ensure_pandoc_available_converts_os_error(monkeypatch):
    def raise_os_error():
        raise OSError("pandoc missing")

    generation.ensure_pandoc_available.cache_clear()
    monkeypatch.setattr(generation.pypandoc, "get_pandoc_version", raise_os_error)

    try:
        generation.ensure_pandoc_available()
    except RuntimeError as exc:
        assert "WSL runtime" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError when pandoc is unavailable")
    finally:
        generation.ensure_pandoc_available.cache_clear()


def test_convert_markdown_to_docx_bytes_calls_pandoc_and_reads_output(monkeypatch, tmp_path):
    monkeypatch.setattr(generation, "ensure_pandoc_available", lambda: None)

    def fake_convert_file(source_path, *, to, format, outputfile, extra_args):
        assert source_path.endswith("result.md")
        assert to == "docx"
        assert format == "markdown+raw_html+superscript+subscript"
        assert any(str(argument).startswith("--reference-doc=") for argument in extra_args)
        with open(outputfile, "wb") as file_handle:
            file_handle.write(b"docx-bytes")

    monkeypatch.setattr(generation.pypandoc, "convert_file", fake_convert_file)

    result = generation.convert_markdown_to_docx_bytes("# Title")

    assert result == b"docx-bytes"


def test_normalize_model_output_strips_any_code_fence_language_tag():
    assert generation.normalize_model_output("```python\nprint(1)\n```") == "print(1)"


def test_normalize_model_output_returns_empty_for_whitespace_only_fenced_block():
    assert generation.normalize_model_output("```markdown\n   \n\t\n```") == ""


def test_parse_json_object_with_backtick_in_content():
    result = image_shared.parse_json_object(
        '```json\n{"key": "val`ue"}\n```',
        empty_message="empty",
        no_json_message="nojson",
    )

    assert result == {"key": "val`ue"}


def test_parse_json_object_fence_without_newline():
    result = image_shared.parse_json_object(
        '```{"key": 1}```',
        empty_message="empty",
        no_json_message="nojson",
    )

    assert result == {"key": 1}


def test_call_responses_create_with_retry_retries_without_timeout_on_final_attempt():
    calls = []

    class Client:
        class Responses:
            def create(self, **kwargs):
                calls.append(dict(kwargs))
                if len(calls) == 1:
                    raise TypeError("unexpected keyword argument 'timeout'")
                return SimpleNamespace(output_text="ok")

        responses = Responses()

    result = image_shared.call_responses_create_with_retry(
        Client(),
        {"model": "gpt-5.4", "input": [], "timeout": 1},
        max_retries=1,
        retryable_error_predicate=lambda exc: False,
    )

    assert result.output_text == "ok"
    assert calls == [
        {"model": "gpt-5.4", "input": [], "timeout": 1},
        {"model": "gpt-5.4", "input": []},
    ]


def test_call_responses_create_with_retry_does_not_double_consume_budget_after_timeout_removal():
    class BudgetExceeded(RuntimeError):
        pass

    class Budget:
        def __init__(self):
            self.used_calls = 0

        def consume(self, operation_name):
            if self.used_calls >= 1:
                raise BudgetExceeded("exhausted")
            self.used_calls += 1

    calls = []
    budget = Budget()

    class Client:
        class Responses:
            def create(self, **kwargs):
                calls.append(dict(kwargs))
                if len(calls) == 1:
                    raise TypeError("unexpected keyword argument 'timeout'")
                return SimpleNamespace(output_text="ok")

        responses = Responses()

    result = image_shared.call_responses_create_with_retry(
        Client(),
        {"model": "gpt-5.4", "input": [], "timeout": 1},
        max_retries=1,
        retryable_error_predicate=lambda exc: False,
        budget=budget,
    )

    assert result.output_text == "ok"
    assert budget.used_calls == 1
    assert calls == [
        {"model": "gpt-5.4", "input": [], "timeout": 1},
        {"model": "gpt-5.4", "input": []},
    ]


def test_call_responses_create_with_retry_consumes_budget_only_after_retryable_success():
    class Budget:
        def __init__(self):
            self.used_calls = 0

        def ensure_available(self, operation_name):
            return None

        def consume(self, operation_name):
            self.used_calls += 1

    calls = []
    budget = Budget()

    class Client:
        class Responses:
            def create(self, **kwargs):
                calls.append(dict(kwargs))
                if len(calls) == 1:
                    raise RetryableError("rate limited")
                return SimpleNamespace(output_text="ok")

        responses = Responses()

    result = image_shared.call_responses_create_with_retry(
        Client(),
        {"model": "gpt-5.4", "input": []},
        max_retries=2,
        retryable_error_predicate=lambda exc: isinstance(exc, RetryableError),
        budget=budget,
    )

    assert result.output_text == "ok"
    assert len(calls) == 2
    assert budget.used_calls == 1


def test_call_responses_create_with_retry_retries_without_temperature_on_bad_request():
    calls = []

    class UnsupportedTemperatureError(Exception):
        status_code = 400

    class Client:
        class Responses:
            def create(self, **kwargs):
                calls.append(dict(kwargs))
                if len(calls) == 1:
                    raise UnsupportedTemperatureError("Unsupported parameter: 'temperature' is not supported with this model.")
                return SimpleNamespace(output_text="ok")

        responses = Responses()

    result = image_shared.call_responses_create_with_retry(
        Client(),
        {"model": "gpt-5-mini", "input": [], "temperature": 0.4},
        max_retries=1,
        retryable_error_predicate=lambda exc: False,
    )

    assert result.output_text == "ok"
    assert calls == [
        {"model": "gpt-5-mini", "input": [], "temperature": 0.4},
        {"model": "gpt-5-mini", "input": []},
    ]


def test_normalize_generated_document_background_whitens_dark_border_only():
    image = Image.new("RGBA", (12, 12), (0, 0, 0, 255))
    for x_coord in range(3, 9):
        for y_coord in range(3, 9):
            image.putpixel((x_coord, y_coord), (200, 0, 0, 255))

    normalized = _normalize_generated_document_background(image)

    assert normalized.getpixel((0, 0)) == (255, 255, 255, 255)
    assert normalized.getpixel((5, 5)) == (200, 0, 0, 255)


def test_generate_markdown_block_strips_image_placeholders_from_context(monkeypatch):
    captured_inputs = []

    def create_response(**kwargs):
        captured_inputs.append(kwargs.get("input", []))
        return SimpleNamespace(output_text="Исправленный текст")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="Основной текст без placeholder'а",
        context_before="Предшествующий блок\n\n[[DOCX_IMAGE_img_001]]\n\nДополнительный текст",
        context_after="[[DOCX_IMAGE_img_002]] Следующий блок",
        max_retries=1,
    )

    assert result == "Исправленный текст"
    assert len(captured_inputs) == 1
    all_prompt_text = " ".join(
        item.get("text", "") if isinstance(item, dict) else getattr(item, "text", "")
        for message in captured_inputs[0]
        for content_item in (
            message.get("content", []) if isinstance(message, dict) else getattr(message, "content", [])
        )
        for item in ([content_item] if isinstance(content_item, dict) else [])
    )
    assert "[[DOCX_IMAGE_img_" not in all_prompt_text


def test_generate_markdown_block_marker_mode_preserves_markers_and_returns_clean_markdown(monkeypatch):
    captured_inputs = []

    def create_response(**kwargs):
        captured_inputs.append(kwargs.get("input", []))
        return SimpleNamespace(
            output_text="[[DOCX_PARA_p0001]]\nИсправленный заголовок\n\n[[DOCX_PARA_p0002]]\nИсправленный абзац"
        )

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p0001]]\n# Заголовок\n\n[[DOCX_PARA_p0002]]\nАбзац",
        context_before="before",
        context_after="after",
        max_retries=1,
        expected_paragraph_ids=["p0001", "p0002"],
        marker_mode=True,
    )

    assert result == "Исправленный заголовок\n\nИсправленный абзац"
    prompt_text = captured_inputs[0][1]["content"][0]["text"]
    assert "[TARGET BLOCK WITH MARKERS]" in prompt_text
    assert "Preserve every marker exactly" in prompt_text


def test_generate_markdown_block_marker_mode_retries_and_recovers_when_markers_are_lost(monkeypatch):
    attempts = []
    sleep_calls = []
    logged_events = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        if len(attempts) <= 2:
            return SimpleNamespace(output_text="Маркеры потеряны")
        return SimpleNamespace(output_text="[[DOCX_PARA_p0001]]\nВосстановленный абзац")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-marker-recovery",
    )

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p0001]]\nАбзац",
        context_before="before",
        context_after="after",
        max_retries=2,
        expected_paragraph_ids=["p0001"],
        marker_mode=True,
    )

    assert result == "Восстановленный абзац"
    assert len(attempts) == 3
    assert sleep_calls == [1]
    assert logged_events[-1][0][1] == "markdown_empty_response_recovery_started"
    assert "[TARGET BLOCK WITH MARKERS ONLY]" in attempts[-1]["input"][1]["content"][0]["text"]
    assert "Required marker sequence:\n[[DOCX_PARA_p0001]]" in attempts[-1]["input"][1]["content"][0]["text"]
    assert "Previous invalid output preview:\nМаркеры потеряны" in attempts[-1]["input"][1]["content"][0]["text"]


def test_generate_markdown_block_marker_mode_falls_back_to_source_after_persistent_marker_validation_failure(monkeypatch):
    """A marker the model INVENTED is still block-fatal, and still ends at the source text.

    Spec 056 E changed the example this test used to carry. It used to send an EMPTY chunk
    (``[[DOCX_PARA_p0001]]\\n``) and assert three calls ending in
    ``markdown_marker_validation_source_fallback``. That assertion was a description of the
    defect, not of correct behaviour: emptiness is what the audiobook prompt ORDERS for a
    paragraph of pure reference apparatus while the user prompt forbids a stub in its
    place, so the block was billed for a retry and a recovery call it could never pass, and
    then thrown away whole. Under E an empty chunk is a paragraph status, not a block
    failure — see ``test_empty_chunk_no_longer_discards_the_paragraphs_around_it``.

    A wrong marker id is a different matter and stays exactly as it was: it is one of the
    three checks that detect real loss, so the retries, the recovery call and the
    block-level source fallback all still happen.

    What the degradation ladder adds is a step BETWEEN the recovery call and that fallback,
    so the model here has to fail the per-paragraph re-ask too — otherwise the block would
    (correctly) be translated and there would be no fallback left to describe. The subject
    of the test is therefore now the ladder's FLOOR: when the divided calls fail as well,
    the delivered text, the event and the record are exactly what they were before.
    """
    attempts = []
    sleep_calls = []
    logged_events = []

    client = _batched_answer_only_client("[[DOCX_PARA_p9999]]\nЧужой маркер", requests=attempts)
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-marker-source-fallback",
    )

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p0001]]\nАбзац-источник",
        context_before="before",
        context_after="after",
        max_retries=2,
        expected_paragraph_ids=["p0001"],
        marker_mode=True,
    )

    assert result == "Абзац-источник"
    # 3 for the block (2 attempts + the informed recovery) + 3 for the single paragraph
    # re-asked without markers, which the stub answers with nothing.
    assert len(attempts) == 6
    assert sleep_calls == [1, 1]
    assert any(args[1] == "markdown_empty_response_recovery_started" for args, _ in logged_events)
    assert any(args[1] == "degradation_ladder_started" for args, _ in logged_events)
    assert logged_events[-1][0][1] == "markdown_marker_validation_source_fallback"


@pytest.mark.parametrize(
    ("response_factory", "expected_event"),
    [
        (
            lambda: SimpleNamespace(
                status="incomplete",
                output=[SimpleNamespace(type="reasoning", status="incomplete")],
            ),
            "markdown_incomplete_response_source_fallback",
        ),
        (
            lambda: SimpleNamespace(output_text=""),
            "markdown_empty_response_source_fallback",
        ),
        (
            lambda: SimpleNamespace(status="failed"),
            "markdown_non_completed_response_source_fallback",
        ),
    ],
)
def test_generate_markdown_block_marker_mode_sanitizes_all_controlled_source_fallbacks(
    monkeypatch,
    response_factory,
    expected_event,
):
    logged_events = []
    client = SimpleNamespace(
        responses=SimpleNamespace(create=lambda **kwargs: response_factory())
    )
    monkeypatch.setattr(generation.time, "sleep", lambda *_: None)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-marker-fallback",
    )

    source_text = "**Текст [обычный]**\n[[DOCX_IMAGE_img_001]]"
    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=f"[[DOCX_PARA_p0001]]\n{source_text}",
        context_before="before",
        context_after="after",
        max_retries=2,
        expected_paragraph_ids=["p0001"],
        marker_mode=True,
    )

    assert result == source_text
    assert "[[DOCX_PARA_" not in result
    fallback_event = next(
        (args, kwargs)
        for args, kwargs in reversed(logged_events)
        if len(args) > 1 and args[1] == expected_event
    )
    assert fallback_event[1]["target_chars"] == len(source_text)


def test_generate_markdown_block_non_marker_fallback_preserves_bracketed_unicode_markdown(monkeypatch):
    logged_events = []
    client = SimpleNamespace(
        responses=SimpleNamespace(create=lambda **kwargs: SimpleNamespace(status="failed"))
    )
    monkeypatch.setattr(generation.time, "sleep", lambda *_: None)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-standard-fallback",
    )

    source_text = "**Текст [обычный]** [[DOCX_IMAGE_img_001]]"
    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=source_text,
        context_before="before",
        context_after="after",
        max_retries=2,
        marker_mode=False,
    )

    assert result == source_text
    assert logged_events[-1][0][1] == "markdown_non_completed_response_source_fallback"
    assert logged_events[-1][1]["target_chars"] == len(source_text)


def test_split_marker_preserved_markdown_raises_structured_marker_diagnostics():
    with pytest.raises(generation.MarkerValidationError) as exc_info:
        generation._split_marker_preserved_markdown(
            "[[DOCX_PARA_p9999]]\nНеверный маркер",
            ["p0001"],
        )

    exc = exc_info.value
    assert exc.error_code == "marker_order_or_identity"
    assert exc.expected_paragraph_ids == ("p0001",)
    assert exc.found_paragraph_ids == ("p9999",)
    assert exc.raw_markdown_preview == "[[DOCX_PARA_p9999]]\nНеверный маркер"


# --- one marker — one paragraph: stub answers never reach the document -------------
# 2026-08-03 literary-edit first live run: the import tore endnotes in two, the model
# merged the tail into the previous marker, and — forbidden to delete the emptied marker —
# filled it with the invented word "(Пусто)", which shipped to the reader 7 times.

_STUB_SOURCE_ABSORBED = (
    "4. Подробнее см.: Bernard Lietaer, Robert Ulanowicz, and Sally Goerner, «White Paper»"
)
_STUB_SOURCE_COLLAPSED = (
    "«О всех вариантах управления системным банковским кризисом»: доклад, в котором авторы "
    "разбирают устойчивость денежных систем и показывают, почему монокультура валют делает "
    "экономику хрупкой, а разнообразие — устойчивой."
)


def test_restore_collapsed_marker_paragraphs_replaces_invented_stub_with_source():
    returned = [
        f"{_STUB_SOURCE_ABSORBED} {_STUB_SOURCE_COLLAPSED}",
        "(Пусто)",
    ]

    restored = generation.restore_collapsed_marker_paragraphs(
        returned,
        [_STUB_SOURCE_ABSORBED, _STUB_SOURCE_COLLAPSED],
        expected_paragraph_ids=["p1344", "p1345"],
    )

    assert "(Пусто)" not in restored
    # The neighbour that swallowed the paragraph is restored too, otherwise the swallowed
    # text would ship twice.
    assert restored == [_STUB_SOURCE_ABSORBED, _STUB_SOURCE_COLLAPSED]


def test_restore_collapsed_marker_paragraphs_restores_the_pair_when_the_merge_went_forward():
    """The model merges forward as readily as back, and then the pair must move together.

    Looking for the absorber only at ``index - 1`` left the FIRST paragraph of a block
    unpaired: its source was re-instated while the next chunk kept holding the same text,
    so the reader got that paragraph twice.
    """
    returned = [
        "(Пусто)",
        f"{_STUB_SOURCE_COLLAPSED} {_STUB_SOURCE_ABSORBED}",
    ]

    restored = generation.restore_collapsed_marker_paragraphs(
        returned,
        [_STUB_SOURCE_COLLAPSED, _STUB_SOURCE_ABSORBED],
        expected_paragraph_ids=["p1344", "p1345"],
    )

    assert "(Пусто)" not in restored
    assert restored == [_STUB_SOURCE_COLLAPSED, _STUB_SOURCE_ABSORBED]
    # The decisive assertion: the swallowed paragraph is delivered ONCE.
    assert sum(_STUB_SOURCE_COLLAPSED in chunk for chunk in restored) == 1


def test_restore_collapsed_marker_paragraphs_keeps_a_short_answer_nobody_absorbed():
    """A shrunken paragraph is not evidence of a merge, and a revert costs a real edit.

    Measured on the 2026-08-03 run: the single collapse with no absorbing neighbour
    (p1533) was not a merge at all — the model had shifted a whole endnote region, and its
    text was already shipping five markers earlier. Restoring the source there delivered
    that 655-character quote TWICE.
    """
    wordy_source = (
        "В связи с вышеизложенным следует констатировать, что осуществление данного "
        "мероприятия в текущих обстоятельствах представляется в высшей степени "
        "нецелесообразным и не может быть рекомендовано к исполнению."
    )
    tightened = "Делать этого не стоит."
    neighbour_source = "1. Там же, с. 131, со ссылкой на архивные материалы Бундесбанка."
    neighbour_edited = "1. Там же, с. 131."
    returned = [neighbour_edited, tightened]

    restored = generation.restore_collapsed_marker_paragraphs(
        returned,
        [neighbour_source, wordy_source],
        expected_paragraph_ids=["p1532", "p1533"],
    )

    assert restored == [neighbour_edited, tightened]


def test_restore_collapsed_marker_paragraphs_keeps_an_audiobook_reference_paragraph_drop():
    """Audiobook prompt rule 1 ORDERS reference paragraphs out of the narration.

    Marker mode is enabled by configuration, not by operation, so this function also runs
    on ``processing_operation="audiobook"``. Restoring on shrinkage alone put the
    bibliography the model was told to delete straight back into the text read aloud.
    """
    reference_source = (
        "12. Bernard Lietaer, Robert Ulanowicz, Sally Goerner, «Options for Managing a "
        "Systemic Bank Crisis», S.A.P.I.EN.S 2, no. 1 (2009): 1–15, doi:10.5194/sapiens-2-1."
    )
    body_source = "Устойчивость денежной системы держится не на эффективности, а на разнообразии."
    body_narrated = "Устойчивость денежной системы держится на разнообразии, а не на эффективности."
    returned = [body_narrated, "."]

    restored = generation.restore_collapsed_marker_paragraphs(
        returned,
        [body_source, reference_source],
        expected_paragraph_ids=["p0101", "p0102"],
    )

    assert restored == returned


def test_restore_collapsed_marker_paragraphs_leaves_genuine_edits_untouched():
    source = [
        "Он поместил свои руки в свои карманы и начал осуществлять прогулку вдоль набережной.",
        "Она была освещена светом фонарей, которые горели вдоль всей её протяжённости.",
    ]
    returned = [
        "Он сунул руки в карманы и побрёл вдоль набережной.",
        "Её освещали фонари, горевшие по всей длине.",
    ]

    assert generation.restore_collapsed_marker_paragraphs(returned, source) == returned


def test_restore_collapsed_marker_paragraphs_leaves_a_translation_untouched():
    # The check is length-only, so it must stay quiet on `translate`, where the answer
    # shares no characters with its source but keeps its size.
    source = [
        "Деньги для людей — то же, что вода для рыб: среда, которую замечают, только когда её нет.",
        "Современные деньги были придуманы в совершенно иную эпоху, для совершенно иных задач.",
    ]
    returned = [
        "Money is to people what water is to fish: a medium noticed only in its absence.",
        "Modern money was invented in a wholly different era, for wholly different purposes.",
    ]

    assert generation.restore_collapsed_marker_paragraphs(returned, source) == returned


def test_generate_markdown_block_marker_mode_replaces_stub_answer_with_source_paragraph(monkeypatch):
    target_text = (
        f"[[DOCX_PARA_p1344]]\n{_STUB_SOURCE_ABSORBED}\n\n[[DOCX_PARA_p1345]]\n{_STUB_SOURCE_COLLAPSED}"
    )

    def create_response(**kwargs):
        return SimpleNamespace(
            output_text=(
                f"[[DOCX_PARA_p1344]]\n{_STUB_SOURCE_ABSORBED} {_STUB_SOURCE_COLLAPSED}\n\n"
                "[[DOCX_PARA_p1345]]\n(Пусто)"
            )
        )

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=target_text,
        context_before="before",
        context_after="after",
        max_retries=1,
        expected_paragraph_ids=["p1344", "p1345"],
        marker_mode=True,
    )

    assert "(Пусто)" not in result
    assert result == f"{_STUB_SOURCE_ABSORBED}\n\n{_STUB_SOURCE_COLLAPSED}"
    # The paragraph-per-marker count is the invariant that must not move.
    assert len(result.split("\n\n")) == 2


# --- spec 056 E: a typed disposition per paragraph ----------------------------------
# A block used to be all-or-nothing. On the 2026-08-04 audiobook run block 274 held ten
# paragraphs and all ten were discarded — replaced by the block's own English source in a
# Russian narration — because paragraph p1336, whose entire text is "14", came back empty.


def test_paragraph_disposition_status_names_match_the_reporting_vocabulary():
    """One source of truth: the generator's names are the buckets the run report publishes."""

    assert sorted(
        [
            generation.PARAGRAPH_STATUS_ACCEPTED,
            generation.PARAGRAPH_STATUS_OMITTED,
            generation.PARAGRAPH_STATUS_SOURCE_RESTORED,
            generation.PARAGRAPH_STATUS_RETRY_REQUIRED,
        ]
    ) == sorted(model_accounting.PARAGRAPH_DISPOSITION_STATUSES)


# --- the three checks that detect REAL loss stay block-fatal ------------------------


def test_missing_markers_still_fail_the_block():
    with pytest.raises(generation.MarkerValidationError) as exc_info:
        generation.split_marker_preserved_paragraph_dispositions(
            "Просто текст без единого маркера.",
            ["p0001", "p0002"],
        )
    assert exc_info.value.error_code == "markers_missing"


def test_a_dropped_marker_still_fails_the_block():
    """Calls 199/200/201 of the 2026-08-04 run: the model deleted a paragraph WITH its marker.

    The lost ``p0957`` is a heading (``## NGO Initiative s :``) and it is
    ``absent_from_artifact`` — this is the one mechanism of the three where content is
    genuinely lost rather than reverted to source, so it must never become a status.
    """
    with pytest.raises(generation.MarkerValidationError) as exc_info:
        generation.split_marker_preserved_paragraph_dispositions(
            "[[DOCX_PARA_p0958]]\nАльянс токенов здоровья.",
            ["p0957", "p0958"],
        )
    assert exc_info.value.error_code == "marker_order_or_identity"
    assert exc_info.value.expected_paragraph_ids == ("p0957", "p0958")
    assert exc_info.value.found_paragraph_ids == ("p0958",)


def test_a_duplicated_marker_still_fails_the_block():
    with pytest.raises(generation.MarkerValidationError) as exc_info:
        generation.split_marker_preserved_paragraph_dispositions(
            "[[DOCX_PARA_p0001]]\nПервый.\n\n[[DOCX_PARA_p0001]]\nОн же снова.",
            ["p0001", "p0002"],
        )
    assert exc_info.value.error_code == "marker_order_or_identity"


def test_reordered_markers_still_fail_the_block():
    with pytest.raises(generation.MarkerValidationError) as exc_info:
        generation.split_marker_preserved_paragraph_dispositions(
            "[[DOCX_PARA_p0002]]\nВторой.\n\n[[DOCX_PARA_p0001]]\nПервый.",
            ["p0001", "p0002"],
        )
    assert exc_info.value.error_code == "marker_order_or_identity"


def test_text_before_the_first_marker_still_fails_the_block():
    with pytest.raises(generation.MarkerValidationError) as exc_info:
        generation.split_marker_preserved_paragraph_dispositions(
            "Вот обработанный блок:\n\n[[DOCX_PARA_p0001]]\nПервый абзац.",
            ["p0001"],
        )
    assert exc_info.value.error_code == "unexpected_prefix"
    assert exc_info.value.leading_text_preview == "Вот обработанный блок:"


# --- the paragraph-break counterexample: p1 must not be able to take p2's heading ----


_COUNTEREXAMPLE_ANSWER = (
    "[[DOCX_PARA_p1]]\n"
    "Перевод первого абзаца.\n"
    "\n"
    "## Безопасность\n"
    "[[DOCX_PARA_p2]]\n"
    "[short pause]"
)


def test_a_paragraph_break_in_a_multi_marker_block_is_still_rejected():
    """The standing counterexample of spec 056 anti-regression 2.

    Marker identity and order are exact, both chunks are non-empty, and ``p2``'s source is
    below ``_COLLAPSED_MARKER_CHUNK_MIN_SOURCE_CHARS`` so the collapse-restore is skipped.
    If the break were collapsed, ``p1`` would take ``p2``'s heading and every remaining
    check would pass — a detected failure traded for a silent corruption. There is no
    signal in this answer that distinguishes "the model split my paragraph" from "the model
    placed the next paragraph's heading early", so the break stays block-fatal wherever a
    neighbour exists. This must fail loudly; it must never pass silently.
    """

    with pytest.raises(generation.MarkerValidationError) as exc_info:
        generation.split_marker_preserved_paragraph_dispositions(
            _COUNTEREXAMPLE_ANSWER,
            ["p1", "p2"],
        )
    assert exc_info.value.error_code == "paragraph_split_detected"


def test_the_counterexample_never_reaches_the_document_through_the_generator(monkeypatch):
    """End-to-end twin of the check above: the heading never migrates to ``p1``."""

    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    source = "[[DOCX_PARA_p1]]\nFirst source paragraph.\n\n[[DOCX_PARA_p2]]\n## Safety"

    client = SimpleNamespace(
        responses=SimpleNamespace(create=lambda **_kwargs: SimpleNamespace(output_text=_COUNTEREXAMPLE_ANSWER))
    )

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=source,
        context_before="",
        context_after="",
        max_retries=2,
        expected_paragraph_ids=["p1", "p2"],
        marker_mode=True,
    )

    assert result == "First source paragraph.\n\n## Safety"
    assert "Перевод первого абзаца. ## Безопасность" not in result


def test_a_paragraph_break_in_a_single_marker_block_is_collapsed():
    """Block 118 of the 2026-08-04 run: one 4 095-character welded quotation, seven
    spoken paragraphs back, whole block discarded.

    With one marker there is no neighbour to steal from, and ``unexpected_prefix`` already
    forbids a fragment before it, so every character after the marker provably belongs to
    that paragraph.
    """

    dispositions = generation.split_marker_preserved_paragraph_dispositions(
        "[[DOCX_PARA_p0554]]\nПервое предложение.\n\nВторое предложение.\n\nТретье.",
        ["p0554"],
    )

    assert [(item.paragraph_id, item.status) for item in dispositions] == [("p0554", "accepted")]
    assert dispositions[0].text == "Первое предложение. Второе предложение. Третье."


# --- an empty chunk is a paragraph status, not a block failure ----------------------


def test_empty_chunk_asks_for_a_retry_while_there_is_budget():
    """Two of the seven bare-number paragraphs on the 2026-08-04 run WERE rescued by a
    resend, so the retry is kept exactly as it was — same error code, same accounting."""

    dispositions = generation.split_marker_preserved_paragraph_dispositions(
        "[[DOCX_PARA_p1336]]\n\n[[DOCX_PARA_p1337]]\nПеревод.",
        ["p1336", "p1337"],
    )
    assert [item.status for item in dispositions] == ["retry_required", "accepted"]

    with pytest.raises(generation.MarkerValidationError) as exc_info:
        generation.resolve_marker_paragraph_dispositions(
            dispositions,
            source_paragraph_chunks=["14", "Why not mobilise the Fund?"],
            allow_unresolved_paragraphs=False,
        )
    assert exc_info.value.error_code == "empty_marker_chunk"


def test_empty_chunk_no_longer_discards_the_paragraphs_around_it():
    """Block 274, replayed from its real recorded answer shape: 9 kept, 1 typed."""

    source_chunks = [f"Source paragraph {index}." for index in range(10)]
    source_chunks[3] = "14"
    answer = "\n\n".join(
        f"[[DOCX_PARA_p{1333 + index}]]\n" + ("" if index == 3 else f"Перевод абзаца {index}.")
        for index in range(10)
    )
    paragraph_ids = [f"p{1333 + index}" for index in range(10)]

    dispositions = generation.resolve_marker_paragraph_dispositions(
        generation.split_marker_preserved_paragraph_dispositions(answer, paragraph_ids),
        source_paragraph_chunks=source_chunks,
        allow_unresolved_paragraphs=True,
    )

    statuses = [item.status for item in dispositions]
    assert statuses.count("accepted") == 9
    assert statuses.count("omitted") == 1
    assert dispositions[3].paragraph_id == "p1336"
    assert dispositions[3].status == "omitted"
    # The source stands in the DOCUMENT so the paragraph-per-marker mapping survives.
    assert dispositions[3].text == "14"
    assert dispositions[0].text == "Перевод абзаца 0."


def test_an_empty_chunk_is_an_omission_and_never_a_merge():
    """An EMPTY answer must not be read as evidence that a neighbour swallowed the text.

    Until spec 056 E an empty chunk raised ``empty_marker_chunk`` before the restorer was
    reached, so the restorer was calibrated on stubs and never saw one. Feeding empties to
    it was measurably wrong: the merge evidence is "a neighbour grew by at least half of
    what this paragraph holds", and a Russian translation of an English paragraph is
    routinely 1.5x its source — so an ordinary neighbour looks like an absorber. The block
    then reverted the emptied paragraph AND its good neighbour to the source, marked both
    ``source_restored`` (a status the narration keeps) and read the source aloud.
    """

    returned = ["", f"{_STUB_SOURCE_COLLAPSED} {_STUB_SOURCE_ABSORBED}"]
    dispositions = [
        generation.ParagraphDisposition(paragraph_id="p1344", text=returned[0], status="retry_required"),
        generation.ParagraphDisposition(paragraph_id="p1345", text=returned[1], status="accepted"),
    ]

    resolved = generation.resolve_marker_paragraph_dispositions(
        dispositions,
        source_paragraph_chunks=[_STUB_SOURCE_COLLAPSED, _STUB_SOURCE_ABSORBED],
        allow_unresolved_paragraphs=True,
    )

    assert [item.status for item in resolved] == ["omitted", "accepted"]
    # The emptied paragraph keeps its source in the DOCUMENT, so the paragraph-per-marker
    # mapping holds; the narration filter drops it by status.
    assert resolved[0].text == _STUB_SOURCE_COLLAPSED
    # ANTI-VACUUM: the neighbour's answer is NOT thrown away with it.
    assert resolved[1].text == returned[1]


def test_an_empty_chunk_does_not_stop_a_real_stub_from_being_restored_as_a_pair():
    """ANTI-VACUUM for the rule above: the merge the restorer exists for still fires.

    A merge leaves a STUB behind, not an empty chunk ("(Пусто)", 7 occurrences on the
    2026-08-03 literary-edit run). Excluding empties must not blunt that.
    """

    returned = ["(Пусто)", f"{_STUB_SOURCE_COLLAPSED} {_STUB_SOURCE_ABSORBED}"]
    dispositions = [
        generation.ParagraphDisposition(paragraph_id="p1344", text=returned[0], status="accepted"),
        generation.ParagraphDisposition(paragraph_id="p1345", text=returned[1], status="accepted"),
    ]

    resolved = generation.resolve_marker_paragraph_dispositions(
        dispositions,
        source_paragraph_chunks=[_STUB_SOURCE_COLLAPSED, _STUB_SOURCE_ABSORBED],
        allow_unresolved_paragraphs=True,
    )

    assert [item.status for item in resolved] == ["source_restored", "source_restored"]
    assert [item.text for item in resolved] == [_STUB_SOURCE_COLLAPSED, _STUB_SOURCE_ABSORBED]
    # The swallowed paragraph is delivered ONCE.
    assert sum(_STUB_SOURCE_COLLAPSED in item.text for item in resolved) == 1


def test_generated_block_carries_its_paragraph_dispositions(monkeypatch):
    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    source = "[[DOCX_PARA_p1336]]\n14\n\n[[DOCX_PARA_p1337]]\nWhy not mobilise the Fund?"
    answer = "[[DOCX_PARA_p1336]]\n\n[[DOCX_PARA_p1337]]\nПочему бы не задействовать фонд?"

    client = SimpleNamespace(responses=SimpleNamespace(create=lambda **_kwargs: SimpleNamespace(output_text=answer)))

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=source,
        context_before="",
        context_after="",
        max_retries=2,
        expected_paragraph_ids=["p1336", "p1337"],
        marker_mode=True,
        block_index=274,
    )

    dispositions = generation.marker_paragraph_dispositions(result)
    assert dispositions is not None
    assert [(item.paragraph_id, item.status) for item in dispositions] == [
        ("p1336", "omitted"),
        ("p1337", "accepted"),
    ]
    # Still one paragraph per marker in the joined text, so nothing downstream re-counts.
    assert result == "14\n\nПочему бы не задействовать фонд?"
    assert isinstance(result, str)


def test_paragraph_dispositions_are_counted_for_the_run_report():
    model_accounting.reset_run_model_accounting()
    generation._finalize_generated_markdown(
        "[[DOCX_PARA_p1]]\nПервый.\n\n[[DOCX_PARA_p2]]\n",
        target_text="First.\n\n14",
        context_before="",
        context_after="",
        expected_paragraph_ids=["p1", "p2"],
        marker_mode=True,
        allow_persistent_context_leakage=True,
        source_paragraph_chunks=["First.", "14"],
        allow_unresolved_paragraphs=True,
    )

    counts = model_accounting.snapshot_run_model_accounting()["paragraph_disposition_counts"]
    # A zero is an assertion here, not a missing field.
    assert counts == {"accepted": 1, "omitted": 1, "retry_required": 0, "source_restored": 0}


def test_a_plain_string_carries_no_dispositions():
    assert generation.marker_paragraph_dispositions("просто строка") is None


# --- spec 056 D': the rejected answer is written down, inside the attempt loop ------
# A $0.53 audiobook run left NO record of what the model answered for the six blocks it
# dropped, because the only writer that keeps a raw response is reachable from the path
# where a block RAISES — and a controlled fallback returns a plain string instead.


def _capture_marker_attempts(monkeypatch, tmp_path):
    """Point the attempt-capture family at ``tmp_path`` and return the calls it logged."""

    from docxaicorrector.generation import marker_attempt_capture

    monkeypatch.setattr(
        marker_attempt_capture,
        "MARKER_ATTEMPT_DIAGNOSTICS_DIR",
        tmp_path / "marker_attempts",
    )
    logged: list[tuple[tuple[Any, ...], dict[str, Any]]] = []
    monkeypatch.setattr(
        marker_attempt_capture,
        "log_event",
        lambda *args, **kwargs: logged.append((args, kwargs)),
    )
    return logged


def _read_marker_attempt_artifacts(tmp_path) -> list[dict[str, Any]]:
    import json

    directory = tmp_path / "marker_attempts"
    if not directory.exists():
        return []
    payloads = [json.loads(path.read_text(encoding="utf-8")) for path in sorted(directory.glob("*.json"))]
    return sorted(payloads, key=lambda payload: (payload["attempt"], payload["stage"]))


def test_controlled_source_fallback_still_writes_every_rejected_answer(monkeypatch, tmp_path):
    """The path that leaves NO evidence today: the block falls back and returns normally.

    ``generate_markdown_block`` swallows the recovery exception and hands back the block's
    own source text, so the call site holds neither the answer nor the exception. Every
    attempt must therefore have been captured before that happened.
    """

    logged = _capture_marker_attempts(monkeypatch, tmp_path)
    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    target_text = "[[DOCX_PARA_p0957]]\n## NGO Initiatives:\n\n[[DOCX_PARA_p0958]]\nThe Wellness Token Alliance."
    # Call 199/200/201 of the 2026-08-04 run: expected ['p0957','p0958'], got ['p0958'] —
    # the model deleted a paragraph TOGETHER with its marker, and the heading was lost.
    rejected_answer = "[[DOCX_PARA_p0958]]\nАльянс токенов здоровья."

    # The per-paragraph re-ask must fail too, or the block gets translated by the ladder and
    # there is no controlled fallback left whose evidence this test can check.
    client = _batched_answer_only_client(rejected_answer)

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text=target_text,
        context_before="before",
        context_after="after",
        max_retries=2,
        expected_paragraph_ids=["p0957", "p0958"],
        marker_mode=True,
        block_index=200,
    )

    # A dropped marker is real loss and stays block-fatal: the source stands.
    assert result == "## NGO Initiatives:\n\nThe Wellness Token Alliance."

    payloads = _read_marker_attempt_artifacts(tmp_path)
    # 2 loop attempts + the informed recovery call.
    assert [(payload["attempt"], payload["stage"]) for payload in payloads] == [
        (1, "attempt"),
        (2, "attempt"),
        (3, "recovery"),
    ]
    for payload in payloads:
        assert payload["schema_version"] == 1
        assert payload["block_index"] == 200
        assert payload["error_code"] == "marker_order_or_identity"
        assert payload["expected_paragraph_ids"] == ["p0957", "p0958"]
        assert payload["found_paragraph_ids"] == ["p0958"]
        # The FULL answer, not a preview: the point is that it can be replayed offline.
        assert payload["raw_response"] == rejected_answer
        assert payload["raw_response_chars"] == len(rejected_answer)
    assert [call[0][1] for call in logged] == ["marker_attempt_rejected"] * 3
    # The model payload never reaches the log line (LOGGING_AND_ARTIFACT_RETENTION §1.5).
    assert all("raw_response" not in call[1] for call in logged)
    assert all(call[1]["raw_response_chars"] == len(rejected_answer) for call in logged)


def test_rejected_answer_capture_keeps_the_full_response_past_the_preview_limit(monkeypatch, tmp_path):
    """``MarkerValidationError`` truncates its preview at 1000 chars; the artifact must not."""

    _capture_marker_attempts(monkeypatch, tmp_path)
    long_paragraph = "Длинный переведённый абзац. " * 120
    rejected_answer = f"[[DOCX_PARA_p0002]]\n{long_paragraph}"

    client = _batched_answer_only_client(rejected_answer)

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p0001]]\nSource paragraph.",
        context_before="",
        context_after="",
        max_retries=1,
        expected_paragraph_ids=["p0001"],
        marker_mode=True,
        block_index=7,
    )

    # Wrong marker id is block-fatal, so the source stands — and that is exactly the path
    # that used to leave nothing behind.
    assert result == "Source paragraph."
    payloads = _read_marker_attempt_artifacts(tmp_path)
    assert payloads
    assert len(rejected_answer) > 1000
    # The captured answer is the NORMALISED model output (fences stripped) — exactly what
    # the validator judged, which is what a replay has to re-judge.
    assert payloads[0]["raw_response"] == rejected_answer.strip()
    assert payloads[0]["error_code"] == "marker_order_or_identity"
    assert payloads[0]["found_paragraph_ids"] == ["p0002"]


def test_rejected_answer_capture_is_silent_for_failures_that_carry_no_answer(monkeypatch, tmp_path):
    """A transient API error has no model answer to keep — do not write an empty record."""

    _capture_marker_attempts(monkeypatch, tmp_path)
    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)

    def create_response(**kwargs):
        return SimpleNamespace(output_text="")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))

    generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="Обычный блок без маркеров.",
        context_before="",
        context_after="",
        max_retries=1,
        expected_paragraph_ids=None,
        marker_mode=False,
        block_index=3,
    )

    assert _read_marker_attempt_artifacts(tmp_path) == []


def test_rejected_answer_capture_failure_never_breaks_generation(monkeypatch, tmp_path):
    """Losing a diagnostic must not take down the generation it only observes."""

    from docxaicorrector.generation import marker_attempt_capture

    _capture_marker_attempts(monkeypatch, tmp_path)

    def explode(**_kwargs):
        raise OSError("disk full")

    monkeypatch.setattr(marker_attempt_capture, "write_marker_attempt_artifact", explode)
    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)

    answers = iter(
        [
            "[[DOCX_PARA_p0002]]\nWrong marker.",
            "[[DOCX_PARA_p0001]]\nПравильный ответ.",
        ]
    )

    def create_response(**kwargs):
        return SimpleNamespace(output_text=next(answers))

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p0001]]\nSource paragraph.",
        context_before="",
        context_after="",
        max_retries=2,
        expected_paragraph_ids=["p0001"],
        marker_mode=True,
        block_index=11,
    )

    assert result == "Правильный ответ."


def test_marker_preserving_user_prompt_tells_the_model_what_to_do_instead_of_a_stub():
    prompt = generation._build_marker_preserving_user_prompt(
        target_text="[[DOCX_PARA_p0001]]\nАбзац",
        context_before="before",
        context_after="after",
    )

    assert "Never answer with a placeholder" in prompt
    assert "(Пусто)" in prompt


def test_detect_context_leakage_finds_verbatim_fragment_absent_from_target():
    leaked_fragment = generation._detect_context_leakage(
        response_text=(
            "Исправленный блок. Возможно, вы взяли эту книгу, думая, что она подскажет, как увеличить личное состояние."
        ),
        target_text="Исправленный блок.",
        context_before=(
            "Возможно, вы взяли эту книгу, думая, что она подскажет, как увеличить личное состояние."
        ),
        context_after="Следующий абзац без совпадений.",
    )

    assert leaked_fragment == "Возможно, вы взяли эту книгу, думая"


def test_generate_markdown_block_retries_on_context_leakage_and_reinforces_prompt(monkeypatch):
    attempts = []
    sleep_calls = []

    def create_response(**kwargs):
        attempts.append(dict(kwargs))
        if len(attempts) == 1:
            return SimpleNamespace(
                output_text=(
                    "Исправленный блок. Возможно, вы взяли эту книгу, думая, что она подскажет, как увеличить личное состояние."
                )
            )
        return SimpleNamespace(output_text="Исправленный блок.")

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", sleep_calls.append)

    result = generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="Исправленный блок.",
        context_before="Возможно, вы взяли эту книгу, думая, что она подскажет, как увеличить личное состояние.",
        context_after="Следующий абзац без совпадений.",
        max_retries=2,
    )

    assert result == "Исправленный блок."
    assert len(attempts) == 2
    assert sleep_calls == [1]
    assert generation._CONTEXT_LEAKAGE_RETRY_WARNING in attempts[1]["input"][1]["content"][0]["text"]


def test_strip_image_placeholders_removes_only_placeholder_tokens():
    result = generation._strip_image_placeholders(
        "Текст перед\n\n[[DOCX_IMAGE_img_001]]\n\nТекст после"
    )
    assert "[[DOCX_IMAGE_img_" not in result
    assert "Текст перед" in result
    assert "Текст после" in result


def test_build_reference_docx_configures_body_and_heading_baselines(tmp_path):
    reference_docx_path = tmp_path / "reference.docx"

    generation._build_reference_docx(reference_docx_path)

    document = Document(str(reference_docx_path))
    styles = document.styles

    normal_style = _as_paragraph_style(styles["Normal"])
    body_text_style = _as_paragraph_style(styles["Body Text"])
    list_paragraph_style = _as_paragraph_style(styles["List Paragraph"])
    caption_style = _as_paragraph_style(styles["Caption"])
    table_grid_style = _as_table_style(styles["Table Grid"])
    normal_attrs = _style_rfonts_attrs(document, "Normal")
    heading_attrs = _style_rfonts_attrs(document, "Heading 1")
    list_attrs = _style_rfonts_attrs(document, "List Paragraph")
    caption_attrs = _style_rfonts_attrs(document, "Caption")

    assert _pt(normal_style.font.size) == 11
    assert _pt(normal_style.paragraph_format.space_after) == 8
    assert normal_style.paragraph_format.line_spacing == 1.15
    assert "Aptos" not in normal_attrs.values()

    assert _pt(body_text_style.font.size) == 11
    assert _pt(body_text_style.paragraph_format.space_after) == 8
    assert body_text_style.paragraph_format.line_spacing == 1.15

    heading_sizes = []
    heading_space_before = []
    heading_space_after = []
    for level in range(1, 7):
        style = _as_paragraph_style(styles[f"Heading {level}"])
        heading_sizes.append(_pt(style.font.size))
        heading_space_before.append(_pt(style.paragraph_format.space_before))
        heading_space_after.append(_pt(style.paragraph_format.space_after))
        assert style.font.bold is True
        assert style.paragraph_format.keep_with_next is True
        assert style.paragraph_format.line_spacing == 1.1

    assert heading_sizes == sorted(heading_sizes, reverse=True)
    assert heading_space_before == sorted(heading_space_before, reverse=True)
    assert heading_space_after == sorted(heading_space_after, reverse=True)
    assert "Aptos Display" not in heading_attrs.values()

    assert _pt(list_paragraph_style.font.size) == 11
    assert _pt(list_paragraph_style.paragraph_format.space_before) == 0
    assert _pt(list_paragraph_style.paragraph_format.space_after) == 4
    assert list_paragraph_style.paragraph_format.line_spacing == 1.1
    assert "Aptos" not in list_attrs.values()

    assert _pt(caption_style.font.size) == 10
    assert caption_style.font.italic is True
    assert _pt(caption_style.paragraph_format.space_before) == 4
    assert _pt(caption_style.paragraph_format.space_after) == 10
    assert "Aptos" not in caption_attrs.values()

    assert _pt(table_grid_style.font.size) == 10
    assert table_grid_style.font.name is None


def test_build_reference_docx_ensures_decimal_and_bullet_numbering_definitions(tmp_path):
    reference_docx_path = tmp_path / "reference.docx"

    generation._build_reference_docx(reference_docx_path)

    document = Document(str(reference_docx_path))
    decimal_matches = _find_matching_abstract_numbers(
        document,
        num_fmt="decimal",
        level_texts=("%1.", "%1.%2.", "%1.%2.%3."),
    )
    bullet_matches = _find_matching_abstract_numbers(
        document,
        num_fmt="bullet",
        level_texts=(chr(0x2022), chr(0x25E6), chr(0x25AA)),
    )

    assert len(decimal_matches) == 1
    assert len(bullet_matches) == 1
    assert _has_num_instance_for_abstract_num(document, decimal_matches[0])
    assert _has_num_instance_for_abstract_num(document, bullet_matches[0])


def test_ensure_reference_numbering_definitions_is_idempotent_for_baseline_definitions():
    document = Document()

    generation._ensure_reference_numbering_definitions(document)
    generation._ensure_reference_numbering_definitions(document)

    decimal_matches = _find_matching_abstract_numbers(
        document,
        num_fmt="decimal",
        level_texts=("%1.", "%1.%2.", "%1.%2.%3."),
    )
    bullet_matches = _find_matching_abstract_numbers(
        document,
        num_fmt="bullet",
        level_texts=(chr(0x2022), chr(0x25E6), chr(0x25AA)),
    )

    assert len(decimal_matches) == 1
    assert len(bullet_matches) == 1
    assert _has_num_instance_for_abstract_num(document, decimal_matches[0])
    assert _has_num_instance_for_abstract_num(document, bullet_matches[0])


def test_build_reference_docx_without_font_config_does_not_write_aptos_to_numbering(tmp_path):
    reference_docx_path = tmp_path / "reference.docx"

    generation._build_reference_docx(reference_docx_path)

    with zipfile.ZipFile(reference_docx_path) as docx_archive:
        numbering_xml = docx_archive.read("word/numbering.xml").decode("utf-8")

    assert "Aptos" not in numbering_xml


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_applies_reference_doc_heading_baseline():
    result = generation.convert_markdown_to_docx_bytes("# Заголовок\n\nАбзац")

    document = Document(io.BytesIO(result))
    heading = document.paragraphs[0]
    heading_style = _as_paragraph_style(heading.style)

    assert heading.text == "Заголовок"
    assert heading_style.name == "Heading 1"
    assert heading_style.font.name is None
    assert _pt(heading_style.font.size) == 18
    assert _pt(heading_style.paragraph_format.space_before) == 18
    assert _pt(heading_style.paragraph_format.space_after) == 8
    assert heading_style.paragraph_format.keep_with_next is True
    assert heading_style.paragraph_format.line_spacing == 1.1

    with zipfile.ZipFile(io.BytesIO(result)) as docx_archive:
        document_xml = docx_archive.read("word/document.xml").decode("utf-8")
    assert "Заголовок" in document_xml


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_preserves_ordered_list_word_numbering_semantics():
    result = generation.convert_markdown_to_docx_bytes("1. Первый пункт\n2. Второй пункт\n3. Третий пункт")

    document = Document(io.BytesIO(result))
    list_paragraphs = document.paragraphs[:3]

    assert [paragraph.text for paragraph in list_paragraphs] == [
        "Первый пункт",
        "Второй пункт",
        "Третий пункт",
    ]

    num_ids = [_paragraph_num_id(paragraph) for paragraph in list_paragraphs]
    ilvls = [_paragraph_ilvl(paragraph) for paragraph in list_paragraphs]

    assert all(num_id is not None for num_id in num_ids)
    assert ilvls == ["0", "0", "0"]
    assert len(set(cast(list[str], num_ids))) == 1

    abstract_num = _find_abstract_num_for_num_id(document, cast(str, num_ids[0]))
    assert abstract_num is not None

    # Pandoc may choose any concrete numId and may materialize a full 9-level
    # decimal definition, so the stable contract here is real Word numbering
    # semantics for the emitted list paragraphs rather than an exact custom
    # reference-doc baseline signature.
    levels = cast(list[Any], abstract_num.xpath('./*[local-name()="lvl"]'))
    assert levels

    signatures = [_numbering_level_signature(level) for level in levels]
    assert all(signature["num_fmt"] == "decimal" for signature in signatures)
    assert signatures[0]["lvl_text"] == "%1."
    assert signatures[0]["left"] is not None
    assert signatures[0]["hanging"] is not None


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_preserves_extractor_underline_tag():
    """<u> is raw HTML that Pandoc's DOCX writer drops, so it must be translated."""
    result = generation.convert_markdown_to_docx_bytes("Обычный <u>подчёркнутый</u> текст")

    document = Document(io.BytesIO(result))
    paragraph = document.paragraphs[0]

    assert paragraph.text == "Обычный подчёркнутый текст"
    underlined = [run.text for run in paragraph.runs if run.underline]
    assert underlined == ["подчёркнутый"]


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
@pytest.mark.parametrize(
    ("markdown_text", "expected_text", "expected_underlined"),
    [
        # DOCX splits runs anywhere, so an underlined footnote/reference tail such as
        # "1]" or "см. п. 5]" arrives as its own <u> run. Unescaped, its bracket closed
        # Pandoc's span early and the reader saw literal "[1]]{.underline}" in the body.
        ("Text <u>a] b</u> more", "Text a] b more", ["a] b"]),
        ("Text <u>1]</u> more", "Text 1] more", ["1]"]),
        ("Text <u>note]</u> more", "Text note] more", ["note]"]),
        ("Text <u>see [1</u> more", "Text see [1 more", ["see [1"]),
        ("Text <u>]a[</u> more", "Text ]a[ more", ["]a["]),
        ("Текст <u>см. п. 5]</u> дальше", "Текст см. п. 5] дальше", ["см. п. 5]"]),
        ("Text <u>C:\\</u> more", "Text C:\\ more", ["C:\\"]),
        # A "!" in front of the span turned it into Pandoc image syntax.
        ("Text!<u>x</u> more", "Text!x more", ["x"]),
        # Balanced brackets must keep working: escaping may not cost the underline.
        ("Text <u>see [1]</u> more", "Text see [1] more", ["see [1]"]),
        ("Текст <u>[прим. 3]</u> дальше", "Текст [прим. 3] дальше", ["[прим. 3]"]),
    ],
)
def test_convert_markdown_to_docx_bytes_keeps_underline_span_markup_out_of_body_text(
    markdown_text: str, expected_text: str, expected_underlined: list[str]
):
    """Underlined content may not leak Pandoc span markup into the delivered text."""
    result = generation.convert_markdown_to_docx_bytes(markdown_text)

    document = Document(io.BytesIO(result))
    paragraph = document.paragraphs[0]

    assert "{.underline}" not in paragraph.text
    assert paragraph.text == expected_text
    assert [run.text for run in paragraph.runs if run.underline] == expected_underlined


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_keeps_underlined_superscript_role_with_bracket_tail():
    """Escaping the span must not disturb the superscript/subscript escaping around it."""
    non_breaking_space = chr(0x00A0)

    result = generation.convert_markdown_to_docx_bytes(
        "Text <u>see<sup>note 1</sup></u> more\n\n"
        "Text <sup><u>a]</u></sup> more\n\n"
        "Text <u>x<sub>2]</sub></u> more"
    )

    document = Document(io.BytesIO(result))
    nested, wrapping, subscripted = document.paragraphs[:3]

    assert "{.underline}" not in nested.text
    assert nested.text == f"Text seenote{non_breaking_space}1 more"
    assert [run.text for run in nested.runs if run.underline] == ["see", f"note{non_breaking_space}1"]
    assert [run.text for run in nested.runs if run.font.superscript] == [f"note{non_breaking_space}1"]

    assert "{.underline}" not in wrapping.text
    assert wrapping.text == "Text a] more"
    assert [run.text for run in wrapping.runs if run.underline] == ["a]"]
    assert [run.text for run in wrapping.runs if run.font.superscript] == ["a]"]

    assert "{.underline}" not in subscripted.text
    assert subscripted.text == "Text x2] more"
    assert [run.text for run in subscripted.runs if run.underline] == ["x", "2]"]
    assert [run.text for run in subscripted.runs if run.font.subscript] == ["2]"]


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_keeps_emphasis_inside_underline_with_bracket_tail():
    """Bold/italic nested in an underlined run keep both roles once the span is escaped."""
    result = generation.convert_markdown_to_docx_bytes("Text <u>**bold]** x</u> more")

    document = Document(io.BytesIO(result))
    paragraph = document.paragraphs[0]

    assert "{.underline}" not in paragraph.text
    assert "*" not in paragraph.text
    assert paragraph.text == "Text bold] x more"
    assert [run.text for run in paragraph.runs if run.bold] == ["bold]"]
    assert [run.text for run in paragraph.runs if run.underline] == ["bold]", " ", "x"]


def test_preprocess_markdown_for_docx_escapes_span_metacharacters_before_script_escaping():
    """Ordering guard: span escaping may not double the backslashes ^…^ needs for spaces."""
    assert generation._preprocess_markdown_for_docx("Text <u>1]</u>") == "Text [1\\]]{.underline}"
    assert (
        generation._preprocess_markdown_for_docx("Text <u>see<sup>note 1</sup></u>")
        == "Text [see^note\\ 1^]{.underline}"
    )


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_keeps_trailing_space_emphasis_out_of_body_text():
    """A bold DOCX run carrying its trailing space must not leak literal asterisks."""
    import docxaicorrector.document.extraction as document_extraction

    source = Document()
    source_run = source.add_paragraph().add_run("жирный ")
    source_run.bold = True
    emphasized = document_extraction._apply_run_markdown("жирный ", source_run._element)

    result = generation.convert_markdown_to_docx_bytes(f"{emphasized}хвост")

    document = Document(io.BytesIO(result))
    paragraph = document.paragraphs[0]

    assert "*" not in paragraph.text
    assert paragraph.text == "жирный хвост"
    assert [run.text for run in paragraph.runs if run.bold] == ["жирный"]


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_preserves_superscript_containing_a_space():
    """Pandoc superscript may not hold unescaped spaces, or the carets stay literal."""
    result = generation.convert_markdown_to_docx_bytes("Ссылка<sup>прим. 1</sup>.")

    document = Document(io.BytesIO(result))
    paragraph = document.paragraphs[0]

    # Pandoc renders the escaped space as a non-breaking space: the accepted cost of
    # keeping the superscript role instead of demoting the marker to ordinary body text.
    non_breaking_space = chr(0x00A0)
    assert "^" not in paragraph.text
    assert paragraph.text == f"Ссылкаприм.{non_breaking_space}1."
    assert [run.text for run in paragraph.runs if run.font.superscript] == [
        f"прим.{non_breaking_space}1"
    ]


def test_preprocess_markdown_for_docx_leaves_ordinary_body_punctuation_unchanged():
    """Anti-vacuum guard: legitimate asterisks/carets/brackets are not rewritten."""
    body_text = "Формула a^2 и звёздочка * в тексте, а также [ссылка](url) и 2 ~ 3."

    assert generation._preprocess_markdown_for_docx(body_text) == body_text


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_preserves_superscript_and_subscript_inline_tags():
    result = generation.convert_markdown_to_docx_bytes("Alpha<sup>13</sup> beta\n\nH<sub>2</sub>O")

    document = Document(io.BytesIO(result))

    first_runs = document.paragraphs[0].runs
    second_runs = document.paragraphs[1].runs

    assert document.paragraphs[0].text == "Alpha13 beta"
    assert first_runs[1].text == "13"
    assert first_runs[1].font.superscript is True

    assert document.paragraphs[1].text == "H2O"
    assert second_runs[1].text == "2"
    assert second_runs[1].font.subscript is True


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_preserves_inline_html_line_breaks():
    result = generation.convert_markdown_to_docx_bytes("Line one<br/>Line two")

    document = Document(io.BytesIO(result))

    assert len(document.paragraphs) == 1
    assert "w:br" in document.paragraphs[0]._p.xml


# ---------------------------------------------------------------------------
# _patch_reference_theme_fonts
# ---------------------------------------------------------------------------

def _theme_font(doc, slot: str) -> str | None:
    """Return the Latin typeface for *slot* ('major' or 'minor') from theme XML."""
    try:
        theme_part = doc.part.part_related_by(_THEME_REL)
    except KeyError:
        return None
    root = etree.fromstring(theme_part.blob)
    elements = root.findall(f".//{{{_DRAWINGML_NS}}}{slot}Font/{{{_DRAWINGML_NS}}}latin")
    return elements[0].get("typeface") if elements else None


def test_patch_reference_theme_fonts_sets_both_slots():
    doc = Document()
    generation._patch_reference_theme_fonts(doc, body_font="Times New Roman", heading_font="Georgia")

    assert _theme_font(doc, "minor") == "Times New Roman"
    assert _theme_font(doc, "major") == "Georgia"


def test_patch_reference_theme_fonts_only_heading():
    doc = Document()
    original_minor = _theme_font(doc, "minor")
    generation._patch_reference_theme_fonts(doc, body_font=None, heading_font="Georgia")

    assert _theme_font(doc, "major") == "Georgia"
    assert _theme_font(doc, "minor") == original_minor  # body slot unchanged


def test_patch_reference_theme_fonts_only_body():
    doc = Document()
    original_major = _theme_font(doc, "major")
    generation._patch_reference_theme_fonts(doc, body_font="Arial", heading_font=None)

    assert _theme_font(doc, "minor") == "Arial"
    assert _theme_font(doc, "major") == original_major  # heading slot unchanged


def test_patch_reference_theme_fonts_does_not_touch_style_rfonts_ascii():
    """Patching the theme must not alter w:ascii on individual heading styles.

    The OOXML contract is that w:asciiTheme resolves via the theme, so
    w:ascii should remain unset unless an explicit style font override writes it.
    """
    doc = Document()
    from docx.oxml.ns import qn as _qn

    generation._patch_reference_theme_fonts(doc, body_font="Arial", heading_font="Georgia")

    h1 = doc.styles["Heading 1"]
    rpr = h1.element.find(_qn("w:rPr"))
    if rpr is not None:
        rfonts = rpr.find(_qn("w:rFonts"))
        if rfonts is not None:
            # w:ascii was NOT set by _patch_reference_theme_fonts — only the theme blob changed.
            assert rfonts.get(_qn("w:ascii")) is None


def _style_rfonts_attrs(doc: DocxDocument, style_name: str) -> dict[str, str]:
    from docx.oxml.ns import qn as _qn

    style = doc.styles[style_name]
    rpr = style.element.find(_qn("w:rPr"))
    if rpr is None:
        return {}
    rfonts = rpr.find(_qn("w:rFonts"))
    if rfonts is None:
        return {}
    return {key: value for key, value in rfonts.attrib.items()}


def test_build_reference_docx_applies_configured_fonts_to_effective_styles(tmp_path):
    from docx.oxml.ns import qn as _qn

    reference_docx_path = tmp_path / "reference.docx"

    generation._build_reference_docx(
        reference_docx_path,
        body_font="Times New Roman",
        heading_font="Georgia",
    )

    reference_doc = Document(reference_docx_path)
    body_attrs = _style_rfonts_attrs(reference_doc, "Normal")
    heading_attrs = _style_rfonts_attrs(reference_doc, "Heading 1")
    caption_attrs = _style_rfonts_attrs(reference_doc, "Caption")

    with zipfile.ZipFile(reference_docx_path) as docx_archive:
        numbering_xml = docx_archive.read("word/numbering.xml").decode("utf-8")

    assert body_attrs[_qn("w:ascii")] == "Times New Roman"
    assert body_attrs[_qn("w:hAnsi")] == "Times New Roman"
    assert heading_attrs[_qn("w:ascii")] == "Georgia"
    assert heading_attrs[_qn("w:hAnsi")] == "Georgia"
    assert caption_attrs[_qn("w:ascii")] == "Times New Roman"
    assert "Times New Roman" in numbering_xml
    assert _theme_font(reference_doc, "minor") == "Times New Roman"
    assert _theme_font(reference_doc, "major") == "Georgia"


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_theme_fonts_applied_when_configured():
    from docx.oxml.ns import qn as _qn

    result = generation.convert_markdown_to_docx_bytes(
        "# Заголовок\n\nАбзац",
        body_font="Times New Roman",
        heading_font="Georgia",
    )

    with zipfile.ZipFile(io.BytesIO(result)) as z:
        theme_xml = z.read("word/theme/theme1.xml").decode("utf-8")

    result_doc = Document(io.BytesIO(result))
    body_attrs = _style_rfonts_attrs(result_doc, "Normal")
    heading_attrs = _style_rfonts_attrs(result_doc, "Heading 1")

    assert "Georgia" in theme_xml
    assert "Times New Roman" in theme_xml
    assert body_attrs[_qn("w:ascii")] == "Times New Roman"
    assert heading_attrs[_qn("w:ascii")] == "Georgia"


@pytest.mark.skipif(not _pandoc_available(), reason="pandoc is unavailable in current runtime")
def test_convert_markdown_to_docx_bytes_no_font_args_leaves_theme_unchanged():
    """When no font args are passed the theme in the output must not contain
    any font name that was not already present in the python-docx default template.
    'Aptos' must NOT appear in the theme — it should only appear via w:ascii on styles.
    """
    result = generation.convert_markdown_to_docx_bytes("# Заголовок\n\nАбзац")

    with zipfile.ZipFile(io.BytesIO(result)) as z:
        theme_xml = z.read("word/theme/theme1.xml").decode("utf-8")
        numbering_xml = z.read("word/numbering.xml").decode("utf-8")
        styles_xml = z.read("word/styles.xml").decode("utf-8")

    # The default template uses Calibri/Cambria, not Aptos, in its theme slots.
    assert "Aptos" not in theme_xml
    assert "Aptos" not in numbering_xml
    assert "Aptos" not in styles_xml


# --- Token / cost accounting -------------------------------------------------------
#
# The product spent real money on every run while counting neither tokens nor dollars:
# before this block there was not one reference to ``usage``/``prompt_tokens``/``cost``
# in ``src/``. These tests pin the two halves of the contract that matter — the numbers
# a provider DOES report must arrive intact, and the numbers it does NOT report must stay
# visibly unknown rather than collapse into a free-looking zero.


def _reset_accounting():
    model_accounting.reset_run_model_accounting()


def _usage_response(markdown: str, **usage_fields: object) -> SimpleNamespace:
    return SimpleNamespace(output_text=markdown, usage=SimpleNamespace(**usage_fields))


def test_generate_markdown_block_records_provider_reported_tokens_and_cost():
    _reset_accounting()
    client = SimpleNamespace(
        responses=SimpleNamespace(
            create=lambda **_: _usage_response(
                "Исправленный текст",
                prompt_tokens=2315,
                completion_tokens=693,
                total_tokens=3008,
                cost=0.00161825,
            )
        )
    )

    generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="target",
        context_before="",
        context_after="",
        max_retries=1,
    )

    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["model_call_count"] == 1
    assert snapshot["model_calls_with_usage"] == 1
    assert snapshot["model_calls_without_usage"] == 0
    assert snapshot["prompt_tokens"] == 2315
    assert snapshot["completion_tokens"] == 693
    assert snapshot["total_tokens"] == 3008
    assert snapshot["cost_usd_reported_by_provider"] == 0.001618
    assert snapshot["token_accounting_complete"] is True
    assert snapshot["cost_accounting_complete"] is True
    stages = cast(dict[str, Any], snapshot["stages"])
    assert stages["text_generation"]["total_tokens"] == 3008


def test_generate_markdown_block_without_usage_counts_the_call_as_unaccounted():
    """A silent provider must never read as a free call.

    Zero tokens plus ``token_accounting_complete=False`` is the honest shape; inferring a
    token count from text length would be exactly the invented number the accounting
    module refuses to produce.
    """

    _reset_accounting()
    client = SimpleNamespace(
        responses=SimpleNamespace(create=lambda **_: SimpleNamespace(output_text="Исправленный текст"))
    )

    generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="target",
        context_before="",
        context_after="",
        max_retries=1,
    )

    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["model_call_count"] == 1
    assert snapshot["model_calls_without_usage"] == 1
    assert snapshot["model_calls_with_usage"] == 0
    assert snapshot["total_tokens"] == 0
    assert snapshot["token_accounting_complete"] is False
    assert snapshot["cost_usd_reported_by_provider"] == 0.0
    assert snapshot["cost_accounting_complete"] is False
    assert snapshot["model_calls_without_cost"] == 1


def test_model_call_usage_without_cost_leaves_cost_unknown_but_keeps_tokens():
    """Anthropic reports tokens and no cost. Tokens land; no price list fills the gap."""

    usage = model_accounting.extract_model_call_usage(
        SimpleNamespace(usage=SimpleNamespace(input_tokens=1200, output_tokens=340))
    )

    assert usage.usage_reported is True
    assert usage.prompt_tokens == 1200
    assert usage.completion_tokens == 340
    assert usage.total_tokens == 1540  # derived from the two reported halves, not guessed
    assert usage.cost_reported is False
    assert usage.cost_usd == 0.0


def test_model_call_usage_keeps_a_reported_zero_cost_distinct_from_unknown():
    reported_zero = model_accounting.extract_model_call_usage(
        SimpleNamespace(usage=SimpleNamespace(prompt_tokens=5, completion_tokens=5, cost=0.0))
    )
    unknown = model_accounting.extract_model_call_usage(
        SimpleNamespace(usage=SimpleNamespace(prompt_tokens=5, completion_tokens=5))
    )

    assert (reported_zero.cost_reported, reported_zero.cost_usd) == (True, 0.0)
    assert unknown.cost_reported is False


def test_model_call_usage_ignores_a_usage_container_with_no_token_counts():
    usage = model_accounting.extract_model_call_usage(SimpleNamespace(usage=SimpleNamespace(foo="bar")))

    assert usage.usage_reported is False
    assert usage.total_tokens == 0


def test_generate_markdown_block_counts_retries_and_the_paragraphs_they_covered(monkeypatch):
    """107 paragraphs went through retries on the 2026-08-03 run and the report said nothing."""

    _reset_accounting()
    attempts: list[int] = []

    def create_response(**_kwargs):
        attempts.append(1)
        if len(attempts) == 1:
            return SimpleNamespace(output_text="")
        return SimpleNamespace(
            output_text="[[DOCX_PARA_p1]]\nПервый абзац\n\n[[DOCX_PARA_p2]]\nВторой абзац"
        )

    client = SimpleNamespace(responses=SimpleNamespace(create=create_response))
    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)

    generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p1]]\nПервый абзац\n\n[[DOCX_PARA_p2]]\nВторой абзац",
        context_before="",
        context_after="",
        max_retries=2,
        expected_paragraph_ids=["p1", "p2"],
        marker_mode=True,
    )

    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["retry_attempt_count"] == 1
    assert snapshot["retried_block_count"] == 1
    assert snapshot["retried_paragraph_count"] == 2
    assert cast(dict[str, Any], snapshot["retry_reason_counts"])["empty_generation"] == 1
    # Both attempts are still charged for: a retry is a second paid call.
    assert snapshot["model_call_count"] == 2


def test_generate_markdown_block_counts_paragraphs_whose_answer_was_discarded():
    """2 paragraphs kept their source while the block still logged OK. Make it countable."""

    _reset_accounting()
    target_text = (
        f"[[DOCX_PARA_p1344]]\n{_STUB_SOURCE_ABSORBED}\n\n[[DOCX_PARA_p1345]]\n{_STUB_SOURCE_COLLAPSED}"
    )
    client = SimpleNamespace(
        responses=SimpleNamespace(
            create=lambda **_: SimpleNamespace(
                output_text=(
                    f"[[DOCX_PARA_p1344]]\n{_STUB_SOURCE_ABSORBED} {_STUB_SOURCE_COLLAPSED}\n\n"
                    "[[DOCX_PARA_p1345]]\n(Пусто)"
                )
            )
        )
    )

    generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text=target_text,
        context_before="before",
        context_after="after",
        max_retries=1,
        expected_paragraph_ids=["p1344", "p1345"],
        marker_mode=True,
    )

    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["model_output_discarded_paragraph_count"] == 2
    assert cast(dict[str, Any], snapshot["model_output_discarded_reason_counts"]) == {
        "marker_chunk_collapse": 1
    }


def test_chat_completions_fallback_path_is_accounted_too():
    """The OpenRouter Chat Completions fallback is a separate SDK surface, and therefore a
    separate place where accounting could silently go missing."""

    _reset_accounting()
    response = SimpleNamespace(
        choices=[SimpleNamespace(message=SimpleNamespace(content="Исправленный текст"))],
        usage=SimpleNamespace(prompt_tokens=101, completion_tokens=7, total_tokens=108, cost=0.0004),
    )
    client = SimpleNamespace(chat=SimpleNamespace(completions=SimpleNamespace(create=lambda **_: response)))

    generation._call_chat_completions_create(
        _as_openai_client(client),
        {"model": "openrouter/model", "input": [{"role": "user", "content": "hi"}], "temperature": 0.4},
    )

    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["prompt_tokens"] == 101
    assert snapshot["total_tokens"] == 108
    assert snapshot["cost_usd_reported_by_provider"] == 0.0004


def test_anthropic_messages_path_is_accounted_too():
    _reset_accounting()
    response = SimpleNamespace(
        content=[SimpleNamespace(text="Исправленный текст")],
        usage=SimpleNamespace(input_tokens=55, output_tokens=11),
    )
    client = SimpleNamespace(messages=SimpleNamespace(create=lambda **_: response))

    generation._call_anthropic_messages_create(
        client,
        {
            "model": "claude-sonnet-4.6",
            "input": [{"role": "user", "content": "hi"}],
            "max_output_tokens": 512,
        },
    )

    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["total_tokens"] == 66
    assert snapshot["model_calls_without_cost"] == 1
    assert snapshot["cost_accounting_complete"] is False


def test_image_and_text_calls_are_separated_by_stage():
    """Stage breakdown comes free from the call site, so text-vs-images is answerable."""

    _reset_accounting()
    model_accounting.record_model_call_usage(
        stage=model_accounting.STAGE_TEXT_GENERATION,
        response=SimpleNamespace(usage=SimpleNamespace(prompt_tokens=100, completion_tokens=10, cost=0.01)),
    )
    model_accounting.record_model_call_usage(
        stage=model_accounting.STAGE_IMAGE_ANALYSIS,
        response=SimpleNamespace(usage=SimpleNamespace(prompt_tokens=900, completion_tokens=20, cost=0.05)),
    )

    stages = cast(dict[str, Any], model_accounting.snapshot_run_model_accounting()["stages"])
    assert stages["text_generation"]["total_tokens"] == 110
    assert stages["image_analysis"]["total_tokens"] == 920
    assert stages["image_analysis"]["cost_usd_reported_by_provider"] == 0.05


def test_call_responses_create_with_retry_records_usage_for_every_caller():
    """One recording point covers translate, literary edit, proofreading and images."""

    _reset_accounting()

    class Responses:
        def create(self, **_kwargs):
            return SimpleNamespace(
                output_text="ok",
                usage=SimpleNamespace(input_tokens=7, output_tokens=3, total_tokens=10),
            )

    class Client:
        responses = Responses()

    image_shared.call_responses_create_with_retry(
        Client(),
        {"model": "m"},
        max_retries=1,
        retryable_error_predicate=lambda _exc: False,
        usage_stage=model_accounting.STAGE_IMAGE_VALIDATION,
    )

    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["model_call_count"] == 1
    assert snapshot["total_tokens"] == 10
    assert cast(dict[str, Any], snapshot["stages"])["image_validation"]["model_call_count"] == 1


# --- Concurrent runs (Codex round 3, P1-A) -----------------------------------------
#
# Two runs are admitted at once by default
# (processing_runtime._DEFAULT_PROCESSING_ADMISSION_LIMIT == 2). The accounting these
# tests defend exists to answer "what did THIS run cost"; an answer that silently belongs
# to another document is worse than no answer, because it looks authoritative.


def _record_run(*, run_id: str, source_token: str, tokens: int, cost: float, calls: int, barrier, out: dict) -> None:
    with model_accounting.run_model_accounting_scope(run_id=run_id, source_token=source_token):
        for _ in range(calls):
            model_accounting.record_model_call_usage(
                stage=model_accounting.STAGE_TEXT_GENERATION,
                response=SimpleNamespace(
                    usage=SimpleNamespace(prompt_tokens=tokens, completion_tokens=0, cost=cost)
                ),
            )
            # Force the two runs to interleave: neither can finish its calls before the
            # other has started recording, which is exactly the window the global ledger
            # lost spend in.
            barrier.wait(timeout=10)
        out[run_id] = model_accounting.snapshot_run_model_accounting()


def test_two_concurrent_runs_each_get_their_own_spend():
    """Interleaved runs must not reset, share or absorb each other's counters."""

    barrier = threading.Barrier(2)
    out: dict[str, Any] = {}
    threads = [
        threading.Thread(
            target=_record_run,
            kwargs={
                "run_id": "run-a",
                "source_token": "token-a",
                "tokens": 100,
                "cost": 0.01,
                "calls": 5,
                "barrier": barrier,
                "out": out,
            },
        ),
        threading.Thread(
            target=_record_run,
            kwargs={
                "run_id": "run-b",
                "source_token": "token-b",
                "tokens": 7,
                "cost": 0.002,
                "calls": 5,
                "barrier": barrier,
                "out": out,
            },
        ),
    ]
    for thread in threads:
        thread.start()
    for thread in threads:
        thread.join(timeout=30)

    assert set(out) == {"run-a", "run-b"}
    run_a, run_b = out["run-a"], out["run-b"]

    # Each run reports its OWN spend: not the other's, not the sum of both.
    assert run_a["model_call_count"] == 5
    assert run_a["total_tokens"] == 500
    assert run_a["cost_usd_reported_by_provider"] == 0.05
    assert run_b["model_call_count"] == 5
    assert run_b["total_tokens"] == 35
    assert run_b["cost_usd_reported_by_provider"] == 0.01

    # And each snapshot NAMES the run it describes, so attribution is checkable in data.
    assert run_a["run_id"] == "run-a"
    assert run_a["source_token"] == "token-a"
    assert run_a["run_identity_complete"] is True
    assert run_b["run_id"] == "run-b"
    assert run_b["source_token"] == "token-b"


def test_a_second_run_starting_does_not_erase_the_first_runs_spend():
    """The precise loss: run B's start used to reset the ledger run A was still filling."""

    with model_accounting.run_model_accounting_scope(run_id="run-a", source_token="token-a") as ledger_a:
        model_accounting.record_model_call_usage(
            stage=model_accounting.STAGE_TEXT_GENERATION,
            response=SimpleNamespace(usage=SimpleNamespace(prompt_tokens=1000, completion_tokens=0, cost=0.5)),
        )

        started: dict[str, Any] = {}

        def _second_run() -> None:
            with model_accounting.run_model_accounting_scope(run_id="run-b", source_token="token-b"):
                model_accounting.record_model_call_usage(
                    stage=model_accounting.STAGE_TEXT_GENERATION,
                    response=SimpleNamespace(usage=SimpleNamespace(prompt_tokens=3, completion_tokens=0, cost=0.001)),
                )
                started["b"] = model_accounting.snapshot_run_model_accounting()

        thread = threading.Thread(target=_second_run)
        thread.start()
        thread.join(timeout=30)

        # Run A continues after B has come and gone.
        model_accounting.record_model_call_usage(
            stage=model_accounting.STAGE_TEXT_GENERATION,
            response=SimpleNamespace(usage=SimpleNamespace(prompt_tokens=1000, completion_tokens=0, cost=0.5)),
        )
        snapshot_a = model_accounting.snapshot_run_model_accounting()

    assert started["b"]["total_tokens"] == 3
    assert started["b"]["model_call_count"] == 1
    # Both of A's calls survived B's lifetime, and B's call did not join them.
    assert snapshot_a["total_tokens"] == 2000
    assert snapshot_a["model_call_count"] == 2
    assert snapshot_a["cost_usd_reported_by_provider"] == 1.0
    assert ledger_a.run_id == "run-a"


def test_preparation_spend_is_attributed_to_its_source_and_reported_beside_the_run():
    """Preparation runs in another worker before the run exists.

    It is neither charged to whichever run happens to be in flight nor left implied: the
    run's snapshot carries it, marked as NOT part of the run totals, because one
    preparation can feed several runs of the same source.
    """

    with model_accounting.preparation_model_accounting_scope(source_token="token-prep"):
        model_accounting.record_model_call_usage(
            stage=model_accounting.STAGE_BOUNDARY_REVIEW,
            response=SimpleNamespace(usage=SimpleNamespace(prompt_tokens=400, completion_tokens=40, cost=0.02)),
        )

    with model_accounting.run_model_accounting_scope(run_id="run-p", source_token="token-prep"):
        model_accounting.record_model_call_usage(
            stage=model_accounting.STAGE_TEXT_GENERATION,
            response=SimpleNamespace(usage=SimpleNamespace(prompt_tokens=10, completion_tokens=0, cost=0.001)),
        )
        snapshot = model_accounting.snapshot_run_model_accounting()

    # Preparation is NOT inside the run totals.
    assert snapshot["total_tokens"] == 10
    assert snapshot["model_call_count"] == 1
    # But it is visible in the same payload, attributed to the source it prepared.
    preparation = cast(dict[str, Any], snapshot["preparation_accounting"])
    assert preparation["source_token"] == "token-prep"
    assert preparation["included_in_run_totals"] is False
    assert preparation["total_tokens"] == 440
    assert preparation["model_call_count"] == 1

    # A different source's run does not inherit that preparation.
    with model_accounting.run_model_accounting_scope(run_id="run-q", source_token="token-other"):
        other = model_accounting.snapshot_run_model_accounting()
    assert other["preparation_accounting"] is None


def test_preparation_calls_do_not_land_in_a_concurrent_unrelated_run():
    """The mis-attribution half of P1-A: preparation for source B, run in flight for A."""

    with model_accounting.run_model_accounting_scope(run_id="run-a", source_token="token-a"):
        model_accounting.record_model_call_usage(
            stage=model_accounting.STAGE_TEXT_GENERATION,
            response=SimpleNamespace(usage=SimpleNamespace(prompt_tokens=50, completion_tokens=0, cost=0.005)),
        )

        def _prepare_other_source() -> None:
            with model_accounting.preparation_model_accounting_scope(source_token="token-b"):
                model_accounting.record_model_call_usage(
                    stage=model_accounting.STAGE_BOUNDARY_REVIEW,
                    response=SimpleNamespace(usage=SimpleNamespace(prompt_tokens=9999, completion_tokens=0, cost=1.0)),
                )

        thread = threading.Thread(target=_prepare_other_source)
        thread.start()
        thread.join(timeout=30)

        snapshot = model_accounting.snapshot_run_model_accounting()

    assert snapshot["total_tokens"] == 50
    assert snapshot["cost_usd_reported_by_provider"] == 0.005
    assert "boundary_review" not in cast(dict[str, Any], snapshot["stages"])
    # Run A's source has no preparation of its own, and it does not pick up B's.
    assert snapshot["preparation_accounting"] is None


def test_block_level_source_fallback_counts_its_paragraphs_as_source_restored(monkeypatch):
    """A whole block reverting to its own source is ``source_restored`` for every paragraph.

    Without this the run report would say ``source_restored: 0`` on exactly the run where a
    block was thrown away and replaced by its untranslated source. The degradation ladder
    has to fail on the per-paragraph re-ask for the block-level fallback to be reached at
    all — see ``_batched_answer_only_client``.
    """
    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    model_accounting.reset_run_model_accounting()

    result = generation.generate_markdown_block(
        client=_batched_answer_only_client("[[DOCX_PARA_p9999]]\nЧужой маркер"),
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p0001]]\nFirst.\n\n[[DOCX_PARA_p0002]]\nSecond.",
        context_before="",
        context_after="",
        max_retries=1,
        expected_paragraph_ids=["p0001", "p0002"],
        marker_mode=True,
    )

    assert result == "First.\n\nSecond."
    counts = model_accounting.snapshot_run_model_accounting()["paragraph_disposition_counts"]
    assert counts == {"accepted": 0, "omitted": 0, "retry_required": 0, "source_restored": 2}


def test_the_rejected_answer_capture_still_holds_the_model_text_for_an_empty_chunk(monkeypatch, tmp_path):
    """Spec 056 E must not degrade D': ``empty_marker_chunk`` now originates in the
    resolver, so the exception has to keep carrying the model's ACTUAL answer, not a
    reconstruction of it."""

    _capture_marker_attempts(monkeypatch, tmp_path)
    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    rejected_answer = "[[DOCX_PARA_p1336]]\n   \n[[DOCX_PARA_p1337]]\nПочему бы не задействовать фонд?"

    client = SimpleNamespace(
        responses=SimpleNamespace(create=lambda **_kwargs: SimpleNamespace(output_text=rejected_answer))
    )

    generation.generate_markdown_block(
        client=_as_openai_client(client),
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p1336]]\n14\n\n[[DOCX_PARA_p1337]]\nWhy not mobilise the Fund?",
        context_before="",
        context_after="",
        max_retries=2,
        expected_paragraph_ids=["p1336", "p1337"],
        marker_mode=True,
        block_index=274,
    )

    payloads = _read_marker_attempt_artifacts(tmp_path)
    # One rejected attempt: the second one resolves per paragraph instead of failing.
    assert [(payload["attempt"], payload["error_code"]) for payload in payloads] == [(1, "empty_marker_chunk")]
    assert payloads[0]["raw_response"] == rejected_answer.strip()


# --- rev41 P0-2: the transport must not degrade into a plain string in silence -------
# The per-paragraph record rode on a ``str`` subclass, and ``_trim_boundary_context_leakage``
# built its result with ``.strip()`` and slices. That returned a plain ``str`` INSIDE the
# generator, before the value reached any caller, and the paragraph COUNT still matched
# afterwards - so the registry rebuilt itself by re-splitting on a blank line, every status
# was lost, and the source text standing in an ``omitted`` slot was read aloud with every
# check green. Reproduced by execution in ``.run/rev41_attack6.py``.


def test_context_leakage_trim_keeps_the_per_paragraph_record():
    """`.run/rev41_attack6.py`, as a test. Before the fix the record was LOST here."""

    leak = "alpha beta gamma delta epsilon zeta eta theta"
    sources = ["First source paragraph of the book.", "Notes 7 Ibid page 214."]
    answer = f"[[DOCX_PARA_p1]]\n{leak} Переведённый текст первого абзаца.\n[[DOCX_PARA_p2]]\n"

    result = generation._finalize_generated_markdown(
        answer,
        target_text="\n\n".join(sources),
        context_before=leak,
        context_after="",
        expected_paragraph_ids=["p1", "p2"],
        marker_mode=True,
        allow_persistent_context_leakage=True,
        source_paragraph_chunks=sources,
        allow_unresolved_paragraphs=True,
    )

    # The trim really fired - otherwise the test would pass for the wrong reason.
    assert str(result).startswith("eta theta ")
    dispositions = generation.marker_paragraph_dispositions(result)
    assert dispositions is not None
    assert [(item.paragraph_id, item.status) for item in dispositions] == [
        ("p1", "accepted"),
        ("p2", "omitted"),
    ]
    # ANTI-VACUUM: the trim removed only the leaked words, not the paragraph's own text.
    assert dispositions[0].text == "eta theta Переведённый текст первого абзаца."
    assert str(result) == "\n\n".join(item.text for item in dispositions)


# --- the boundary trim deletes text: it must leave a number behind -------------------
# The trim ran on the per-paragraph record, deleted characters and recorded nothing at all
# - no log line, no counter - while the delivered block still had the paragraph count the
# checks downstream compare. Behaviour is untouched below; only the counters are new, and
# each is proved to move. The shape the review feared (an EMPTY paragraph still recorded
# ``accepted``) is measured here too, and it turns out the attribution check already stops
# it - so that is asserted as a guard rather than assumed.


def _leak_trim_record(*texts: str):
    return generation._marker_preserved_block_text(
        [
            generation.ParagraphDisposition(paragraph_id=f"p{index + 1}", text=text, status="accepted")
            for index, text in enumerate(texts)
        ]
    )


def test_the_boundary_leakage_trim_counts_the_characters_it_removed():
    leak = "Дословный кусок соседнего блока."
    value = _leak_trim_record(f"{leak} Собственный текст абзаца.", "Второй абзац.")

    with model_accounting.run_model_accounting_scope(
        run_id="run-leak-trim", source_token="token-leak-trim"
    ) as ledger:
        trimmed, was_trimmed = generation._trim_marker_preserved_boundary_leakage(value, leak)
        snapshot = ledger.snapshot()

    # Behaviour unchanged: the same characters still go.
    assert was_trimmed
    dispositions = generation.marker_paragraph_dispositions(trimmed)
    assert dispositions is not None
    assert dispositions[0].text == "Собственный текст абзаца."
    assert dispositions[1].text == "Второй абзац."

    removed_chars = len(f"{leak} Собственный текст абзаца.") - len("Собственный текст абзаца.")
    assert snapshot["text_removal_event_count"] == 1
    assert snapshot["text_removal_chars"] == removed_chars
    assert snapshot["text_removal_emptied_unit_count"] == 0
    assert snapshot["text_removal_site_counts"] == {"context_leakage_boundary_trim": 1}
    assert snapshot["text_removal_site_chars"] == {"context_leakage_boundary_trim": removed_chars}


@pytest.mark.parametrize(
    ("first_text", "second_text"),
    [
        ("Дословный кусок соседнего блока.", "Второй абзац."),
        ("Первый абзац.", "Дословный кусок соседнего блока."),
    ],
)
def test_a_trim_that_would_empty_a_paragraph_is_refused_instead_of_shipped(first_text, second_text):
    """The shape the review called the worst of the three does not get out.

    A paragraph emptied by the trim would keep its ``accepted`` status while holding no
    text, and the paragraph COUNT would still match, so nothing downstream would fire. It
    cannot ship: the string-level trim also strips the boundary whitespace it removed with
    the leak, so the rebuilt record no longer matches character for character and the
    attribution check raises - the block is retried instead of delivered. Nothing left the
    document, so nothing is counted either.
    """

    leak = "Дословный кусок соседнего блока."
    value = _leak_trim_record(first_text, second_text)

    with model_accounting.run_model_accounting_scope(
        run_id="run-leak-empty", source_token="token-leak-empty"
    ) as ledger:
        with pytest.raises(generation.MarkerValidationError) as exc_info:
            generation._trim_marker_preserved_boundary_leakage(value, leak)
        snapshot = ledger.snapshot()

    assert "context_leakage_trim_unattributable" in str(exc_info.value)
    assert snapshot["text_removal_event_count"] == 0
    assert snapshot["text_removal_emptied_unit_count"] == 0


def test_an_answer_with_no_boundary_leak_leaves_the_removal_counters_at_zero():
    """Anti-vacuum: the counter must stay at zero when nothing is trimmed, or a non-zero
    total would say nothing."""

    value = _leak_trim_record("Первый абзац книги.", "Второй абзац книги.")

    with model_accounting.run_model_accounting_scope(
        run_id="run-leak-none", source_token="token-leak-none"
    ) as ledger:
        trimmed, was_trimmed = generation._trim_marker_preserved_boundary_leakage(
            value, "Фрагмент, которого в ответе нет."
        )
        snapshot = ledger.snapshot()

    assert not was_trimmed
    assert str(trimmed) == str(value)
    assert snapshot["text_removal_event_count"] == 0
    assert snapshot["text_removal_chars"] == 0
    assert snapshot["text_removal_emptied_unit_count"] == 0
    assert snapshot["text_removal_site_counts"] == {}


def test_a_marker_mode_answer_without_its_record_is_refused_loudly(monkeypatch):
    """The class of defect, not the one instance of it.

    Any future string operation inside the generator that drops the subclass must stop the
    block with a named cause instead of shipping a block whose statuses silently vanished.
    """

    monkeypatch.setattr(
        generation,
        "_strip_and_validate_paragraph_markers",
        lambda *args, **kwargs: "Первый.\n\nВторой.",
    )

    with pytest.raises(generation.MarkerParagraphRecordLost) as exc_info:
        generation._finalize_generated_markdown(
            "[[DOCX_PARA_p1]]\nПервый.\n[[DOCX_PARA_p2]]\nВторой.",
            target_text="First.\n\nSecond.",
            context_before="",
            context_after="",
            expected_paragraph_ids=["p1", "p2"],
            marker_mode=True,
            allow_persistent_context_leakage=True,
            source_paragraph_chunks=["First.", "Second."],
            allow_unresolved_paragraphs=True,
        )

    assert exc_info.value.stage == "finalize_clean"
    # NOT a marker validation error: the model did nothing wrong, so this must not be
    # retried as though a resend could put the record back.
    assert not generation._is_retryable_marker_validation_error(exc_info.value)


def test_a_block_level_source_fallback_still_carries_a_record(monkeypatch):
    """The four block-level fallbacks used to return a bare string, which made "no record"
    an ambiguous state downstream. Every marker-mode exit now says what happened."""

    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    model_accounting.reset_run_model_accounting()

    result = generation.generate_markdown_block(
        client=_batched_answer_only_client("[[DOCX_PARA_p9999]]\nЧужой маркер"),
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p0001]]\nFirst.\n\n[[DOCX_PARA_p0002]]\nSecond.",
        context_before="",
        context_after="",
        max_retries=1,
        expected_paragraph_ids=["p0001", "p0002"],
        marker_mode=True,
    )

    assert result == "First.\n\nSecond."
    dispositions = generation.marker_paragraph_dispositions(result)
    assert dispositions is not None
    assert [(item.paragraph_id, item.status) for item in dispositions] == [
        ("p0001", "source_restored"),
        ("p0002", "source_restored"),
    ]


# --- rev41 P1-1: the counters counted ATTEMPTS, not delivered paragraphs -------------


def test_rejected_attempts_do_not_reach_the_paragraph_disposition_counters():
    """A two-paragraph block that succeeded once after two rejected attempts reported
    ``accepted: 6``, because the statuses were recorded where they are RESOLVED instead of
    where the block is delivered."""

    model_accounting.reset_run_model_accounting()
    leak = "Это дословный кусок соседнего блока, который модель повторила целиком в ответе."
    sources = ["First source paragraph, long enough to matter here.", "Second source paragraph."]
    answer = (
        f"[[DOCX_PARA_p1]]\nПервый переведён. {leak} Ещё немного текста.\n"
        "[[DOCX_PARA_p2]]\nВторой переведён."
    )

    for attempt in (1, 2, 3):
        try:
            generation._finalize_generated_markdown(
                answer,
                target_text="\n\n".join(sources),
                context_before=leak,
                context_after="",
                expected_paragraph_ids=["p1", "p2"],
                marker_mode=True,
                allow_persistent_context_leakage=attempt >= 3,
                source_paragraph_chunks=sources,
                allow_unresolved_paragraphs=attempt >= 3,
            )
        except generation.ContextLeakageError:
            pass

    counts = model_accounting.snapshot_run_model_accounting()["paragraph_disposition_counts"]
    assert counts == {"accepted": 2, "omitted": 0, "retry_required": 0, "source_restored": 0}


# --- rev41 P1-3: the omission WARNING has to carry characters, not only a count -------


def test_the_omission_warning_reports_the_characters_it_withholds(monkeypatch):
    """A real paragraph of prose, not the bare footnote number the first test used.

    Spec 054's metric is the share of source-language CHARACTERS in the artifact, so a
    WARNING reading "1 paragraph" cannot be compared against it: here that one line stands
    for three thousand characters of prose that stop being spoken.
    """

    prose = (
        "The central bank acts as lender of last resort for commercial banks but never for "
        "sovereign states, and the whole architecture of the euro crisis follows from that "
        "single asymmetry, which no treaty ever debated in public. " * 14
    ).strip()
    assert len(prose) > 3000
    answer = "[[DOCX_PARA_p1]]\n\n[[DOCX_PARA_p2]]\nВторой абзац переведён."
    logged_events: list[tuple[tuple[object, ...], dict[str, object]]] = []
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt",
    )

    resolved = generation.resolve_marker_paragraph_dispositions(
        generation.split_marker_preserved_paragraph_dispositions(answer, ["p1", "p2"]),
        source_paragraph_chunks=[prose, "Second paragraph."],
        allow_unresolved_paragraphs=True,
    )

    assert [item.status for item in resolved] == ["omitted", "accepted"]
    omissions = [payload for args, payload in logged_events if args[1] == "marker_paragraph_omitted"]
    assert len(omissions) == 1
    assert omissions[0]["omitted_paragraph_count"] == 1
    assert omissions[0]["omitted_source_chars"] == len(prose)
    assert logging.WARNING in {args[0] for args, _payload in logged_events}
    # ANTI-VACUUM: the prose paragraph next to it is still delivered.
    assert resolved[1].text == "Второй абзац переведён."


# --- rev41 P1-4: the single-marker collapse must not weld a structural line ----------
# Measured over every recorded model answer the repository holds (two books, 488 answers
# carrying a readable marker block): the collapse rescues 6 answers and NONE of them holds
# a heading, a list item or a code fence, so refusing those costs nothing that was being
# gained. Refusing returns the block to ``paragraph_split_detected``, i.e. exactly the
# behaviour before spec 056.


@pytest.mark.parametrize(
    "second_part",
    [
        "## Безопасность денежной системы",
        "- первый пункт списка",
        "1. первый пункт нумерованного списка",
        "```python",
    ],
)
def test_a_single_marker_break_around_a_structural_line_is_not_collapsed(second_part):
    answer = f"[[DOCX_PARA_p1]]\nПервая часть абзаца.\n\n{second_part}"

    with pytest.raises(generation.MarkerValidationError) as exc_info:
        generation.split_marker_preserved_paragraph_dispositions(answer, ["p1"])

    assert exc_info.value.error_code == "paragraph_split_detected"


def test_a_single_marker_prose_break_is_still_collapsed():
    """ANTI-VACUUM for the guard above: the 6 real rescues are prose and stay rescued."""

    answer = (
        "[[DOCX_PARA_p1]]\nДлинная цитата, которую импортёр сварил в один абзац.\n\n"
        "Модель разбила её на две произносимые части."
    )

    dispositions = generation.split_marker_preserved_paragraph_dispositions(answer, ["p1"])

    assert [item.status for item in dispositions] == ["accepted"]
    assert dispositions[0].text == (
        "Длинная цитата, которую импортёр сварил в один абзац. "
        "Модель разбила её на две произносимые части."
    )


# --- The degradation ladder: divide instead of substituting the source ----------------
#
# Spec about prose being lost. Step 2 answers a marker-contract rejection by re-asking each
# paragraph on its own without markers; step 3 answers ``incomplete_response`` by dividing.
# Every test below fails on ``origin/main``, where the four controlled fallbacks are the
# FIRST answer to a rejection rather than the last.


def _ladder_client(
    answers_by_target: dict[str, str],
    *,
    requests: list[dict[str, Any]] | None = None,
    default_answer: str | None = None,
) -> Any:
    """A model keyed by what it was ASKED, so a per-paragraph re-ask can answer that paragraph.

    Matching is by containment of the target text in the prompt: the ladder builds the
    ordinary standard prompt around the unit, so the unit's own text is in there verbatim.
    A target with no entry and no ``default_answer`` returns nothing, which is how a test
    makes one paragraph fail while its neighbours succeed.
    """

    def create_response(**kwargs: Any) -> SimpleNamespace:
        if requests is not None:
            requests.append(dict(kwargs))
        prompt = kwargs["input"][1]["content"][0]["text"]
        target = prompt.split("[TARGET BLOCK", 1)[-1]
        for source, answer in answers_by_target.items():
            if source in target:
                return SimpleNamespace(output_text=answer)
        if default_answer is not None:
            return SimpleNamespace(output_text=default_answer)
        return SimpleNamespace(output_text="")

    return _as_openai_client(SimpleNamespace(responses=SimpleNamespace(create=create_response)))


def _incomplete_response() -> SimpleNamespace:
    return SimpleNamespace(status="incomplete", output=[SimpleNamespace(type="reasoning", status="incomplete")])


def test_step_two_translates_every_paragraph_instead_of_delivering_the_english_block(monkeypatch):
    """The core of the change: a marker-contract rejection stops costing prose.

    The batched answer violates marker identity, which is one of the three checks that
    detect real loss and therefore stays block-fatal. Before the ladder that verdict was the
    end of the block: its own English source was delivered and, in audiobook mode, the block
    was dropped from the narration entirely. Now each paragraph is asked for on its own with
    ``marker_mode=False``, where the error class is unreachable by construction, and the
    block ships translated.
    """

    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    model_accounting.reset_run_model_accounting()
    requests: list[dict[str, Any]] = []
    client = _ladder_client(
        {
            "The current monetary system is taken for granted.": "Нынешняя денежная система воспринимается как данность.",
            "Money is the last taboo.": "Деньги — последнее табу.",
        },
        requests=requests,
    )

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text=(
            "[[DOCX_PARA_p0001]]\nThe current monetary system is taken for granted.\n\n"
            "[[DOCX_PARA_p0002]]\nMoney is the last taboo."
        ),
        context_before="предыдущий блок",
        context_after="следующий блок",
        max_retries=2,
        expected_paragraph_ids=["p0001", "p0002"],
        marker_mode=True,
    )

    assert result == (
        "Нынешняя денежная система воспринимается как данность.\n\nДеньги — последнее табу."
    )
    # The paragraph count is the invariant the whole ladder rests on: two markers in, two
    # paragraphs out, each with its own translation and its own accepted status.
    dispositions = generation.marker_paragraph_dispositions(result)
    assert dispositions is not None
    assert [(item.paragraph_id, item.status) for item in dispositions] == [
        ("p0001", "accepted"),
        ("p0002", "accepted"),
    ]
    # And the volume of delivered text did not drop: a rule that stops substituting without
    # translating would be loss, not a repair.
    assert len(result) >= len("The current monetary system is taken for granted.\n\nMoney is the last taboo.")

    # The re-asks carry NO markers and NOT the marker prompt — that is what makes the error
    # class unreachable rather than merely unlikely.
    ladder_prompts = [
        request["input"][1]["content"][0]["text"]
        for request in requests
        if "[[DOCX_PARA_" not in request["input"][1]["content"][0]["text"]
    ]
    assert len(ladder_prompts) == 2
    assert all("[TARGET BLOCK WITH MARKERS" not in prompt for prompt in ladder_prompts)
    # Neighbouring context is handed over, not thrown away.
    assert "предыдущий блок" in ladder_prompts[0]
    assert "Money is the last taboo." in ladder_prompts[0]
    assert "The current monetary system is taken for granted." in ladder_prompts[1]
    assert "следующий блок" in ladder_prompts[1]

    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["degradation_ladder_block_count"] == 1
    assert snapshot["degradation_ladder_trigger_counts"] == {"marker_contract": 1}
    assert snapshot["degradation_ladder_translated_paragraph_count"] == 2
    assert snapshot["degradation_ladder_unrescued_paragraph_count"] == 0
    # 3 block calls (2 attempts + the informed recovery) preceded the ladder; the ladder
    # itself bought exactly one call per paragraph.
    assert snapshot["degradation_ladder_model_call_count"] == 2
    assert snapshot["model_call_count"] == 5
    assert snapshot["paragraph_disposition_counts"] == {
        "accepted": 2,
        "omitted": 0,
        "retry_required": 0,
        "source_restored": 0,
    }


def test_step_two_keeps_the_paragraph_it_could_not_translate_and_only_that_one(monkeypatch):
    """Partial rescue is the point: one bad paragraph no longer costs its neighbours.

    ``origin/main`` reverts the WHOLE block, so a five-paragraph block loses five
    translations to one unanswerable paragraph. The unanswered one keeps its own source —
    nothing is lost — and it is the only one that does.
    """

    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    model_accounting.reset_run_model_accounting()
    client = _ladder_client(
        {"First source paragraph.": "Первый абзац переведён."},
        # "Second source paragraph." has no entry, so the model answers nothing for it.
    )

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text=(
            "[[DOCX_PARA_p0001]]\nFirst source paragraph.\n\n"
            "[[DOCX_PARA_p0002]]\nSecond source paragraph."
        ),
        context_before="",
        context_after="",
        max_retries=1,
        expected_paragraph_ids=["p0001", "p0002"],
        marker_mode=True,
    )

    assert result == "Первый абзац переведён.\n\nSecond source paragraph."
    dispositions = generation.marker_paragraph_dispositions(result)
    assert dispositions is not None
    # ``omitted``, not ``source_restored``: the source stands in the DOCX (nothing lost) and
    # the narration skips it. The block-level fallback this ladder replaces is dropped from
    # the narration WHOLE, so a partial rescue must not become the one paragraph a listener
    # hears in English.
    assert [(item.paragraph_id, item.status) for item in dispositions] == [
        ("p0001", "accepted"),
        ("p0002", "omitted"),
    ]
    from docxaicorrector.pipeline.block_execution import narration_projection_for_processed_block

    spoken, withheld_paragraphs, withheld_chars = narration_projection_for_processed_block(result)
    assert spoken == "Первый абзац переведён."
    assert (withheld_paragraphs, withheld_chars) == (1, len("Second source paragraph."))
    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["degradation_ladder_translated_paragraph_count"] == 1
    assert snapshot["degradation_ladder_unrescued_paragraph_count"] == 1
    # The paragraph left in the source language is still counted as a discarded answer, so
    # the run report cannot report a rescue it did not fully achieve.
    assert snapshot["model_output_discarded_reason_counts"] == {
        "degradation_ladder_paragraph_source_fallback": 1
    }


def test_step_two_refuses_an_answer_that_would_turn_one_paragraph_into_two(monkeypatch):
    """A one-paragraph request answered with a heading plus prose is REFUSED, not shipped.

    The ladder's whole safety argument is the paragraph count, and a two-paragraph answer
    dropped into one paragraph's slot would deliver three paragraphs under two markers with
    every remaining check green. The guard is the one
    ``split_marker_preserved_paragraph_dispositions`` already uses for a one-marker block.
    """

    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    model_accounting.reset_run_model_accounting()
    client = _ladder_client(
        {"First source paragraph.": "Перевод абзаца.\n\n## Заголовок из соседнего абзаца"}
    )

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p1]]\nFirst source paragraph.\n\n[[DOCX_PARA_p2]]\n## Safety",
        context_before="",
        context_after="",
        max_retries=1,
        expected_paragraph_ids=["p1", "p2"],
        marker_mode=True,
    )

    assert result == "First source paragraph.\n\n## Safety"
    assert result.count("\n\n") == 1
    dispositions = generation.marker_paragraph_dispositions(result)
    assert dispositions is not None
    assert [item.status for item in dispositions] == ["source_restored", "source_restored"]
    # That outcome is ``origin/main``'s outcome too, so without the next two lines the test
    # would pass vacuously. The ladder DID run, refused both answers, and handed the block
    # back to the block-level fallback instead of shipping three paragraphs under two
    # markers — which is what it would have done had the answer been accepted.
    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["degradation_ladder_block_count"] == 1
    assert snapshot["degradation_ladder_translated_paragraph_count"] == 0


def test_step_three_divides_an_incomplete_block_into_its_paragraphs(monkeypatch):
    """``incomplete_response`` is a failure of SIZE, so the remedy is a smaller request.

    ``_boost_request_output_budget`` doubles the budget but ``min``s it with the ceiling, so
    the block is never answered whole however often it is retried. The model here answers
    the block with ``incomplete`` and each paragraph normally.
    """

    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    model_accounting.reset_run_model_accounting()
    first = "First source paragraph, long enough to matter."
    second = "Second source paragraph, also long enough."

    def create_response(**kwargs):
        prompt = kwargs["input"][1]["content"][0]["text"]
        target = prompt.split("[TARGET BLOCK", 1)[-1]
        if "[[DOCX_PARA_" in target:
            return _incomplete_response()
        if first in target.split("[CONTEXT AFTER]")[0]:
            return SimpleNamespace(output_text="Первый абзац.")
        return SimpleNamespace(output_text="Второй абзац.")

    client = _as_openai_client(SimpleNamespace(responses=SimpleNamespace(create=create_response)))

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text=f"[[DOCX_PARA_p0001]]\n{first}\n\n[[DOCX_PARA_p0002]]\n{second}",
        context_before="",
        context_after="",
        max_retries=2,
        expected_paragraph_ids=["p0001", "p0002"],
        marker_mode=True,
    )

    assert result == "Первый абзац.\n\nВторой абзац."
    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["degradation_ladder_trigger_counts"] == {"incomplete_response": 1}
    assert snapshot["degradation_ladder_translated_paragraph_count"] == 2
    assert snapshot["degradation_ladder_sentence_split_paragraph_count"] == 0


def test_step_three_divides_one_paragraph_by_sentences_and_joins_them_back_into_it(monkeypatch):
    """Below one paragraph there is no marker, so the pieces come back into the SAME paragraph.

    The model answers ``incomplete`` for anything as long as the whole paragraph and
    normally for a half of it, which is exactly the shape the output ceiling produces.
    """

    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    model_accounting.reset_run_model_accounting()
    sentences = [f"Sentence number {index} of the paragraph." for index in range(1, 9)]
    paragraph = " ".join(sentences)

    def create_response(**kwargs):
        prompt = kwargs["input"][1]["content"][0]["text"]
        target = prompt.split("[TARGET BLOCK", 1)[-1].split("[CONTEXT AFTER]")[0]
        if len(target) >= len(paragraph):
            return _incomplete_response()
        return SimpleNamespace(output_text=f"Перевод {target.count('Sentence')} предложений.")

    client = _as_openai_client(SimpleNamespace(responses=SimpleNamespace(create=create_response)))

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text=f"[[DOCX_PARA_p0001]]\n{paragraph}",
        context_before="",
        context_after="",
        max_retries=1,
        expected_paragraph_ids=["p0001"],
        marker_mode=True,
    )

    # ONE paragraph in, ONE paragraph out — the marker is neither created nor destroyed.
    assert "\n\n" not in result
    assert result == "Перевод 4 предложений. Перевод 4 предложений."
    dispositions = generation.marker_paragraph_dispositions(result)
    assert dispositions is not None
    assert [(item.paragraph_id, item.status) for item in dispositions] == [("p0001", "accepted")]
    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["degradation_ladder_sentence_split_paragraph_count"] == 1
    assert snapshot["degradation_ladder_translated_paragraph_count"] == 1
    assert snapshot["degradation_ladder_oversized_sentence_count"] == 0


def test_a_sentence_over_the_output_ceiling_is_named_and_never_substituted_silently(monkeypatch):
    """The honest edge of step 3, said out loud rather than swallowed.

    One sentence longer than what the provider's output ceiling can budget for does not
    divide further — there is nothing below a sentence that keeps the paragraph intact. The
    outcome is still today's outcome (the paragraph's own source stands), but it is a WARNING
    with the two numbers in it and a counter of its own, so a run report can say how often
    the ladder met its own limit.
    """

    logged_events = []
    monkeypatch.setattr(generation.time, "sleep", lambda *_: None)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-ceiling",
    )
    model_accounting.reset_run_model_accounting()

    # A single sentence: no terminal punctuation inside it at all.
    oversized = "word " * (generation._LADDER_BUDGETABLE_SOURCE_CHARS // 4)
    assert len(oversized) > generation._LADDER_BUDGETABLE_SOURCE_CHARS
    client = _as_openai_client(
        SimpleNamespace(responses=SimpleNamespace(create=lambda **_kwargs: _incomplete_response()))
    )

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text=f"[[DOCX_PARA_p0001]]\n{oversized.strip()}\n\n[[DOCX_PARA_p0002]]\nSecond.",
        context_before="",
        context_after="",
        max_retries=1,
        expected_paragraph_ids=["p0001", "p0002"],
        marker_mode=True,
    )

    assert result == f"{oversized.strip()}\n\nSecond."
    ceiling_events = [
        kwargs
        for args, kwargs in logged_events
        if len(args) > 1 and args[1] == "degradation_ladder_sentence_exceeds_output_ceiling"
    ]
    assert ceiling_events
    assert ceiling_events[0]["paragraph_chars"] == len(oversized.strip())
    assert ceiling_events[0]["budgetable_chars"] == generation._LADDER_BUDGETABLE_SOURCE_CHARS
    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["degradation_ladder_oversized_sentence_count"] == 1
    assert snapshot["degradation_ladder_translated_paragraph_count"] == 0


def test_the_ladder_never_runs_on_a_block_that_passes(monkeypatch):
    """ANTI-VACUUM. A block the model answers on the first try costs exactly one call.

    A remedy that fires on healthy work is not a remedy, it is a tax. Every ladder counter
    stays at zero and the mapping stays empty, so "the ladder did not run" is a number a
    reader can check rather than an absence he has to trust.
    """

    monkeypatch.setattr(generation.time, "sleep", lambda _seconds: None)
    model_accounting.reset_run_model_accounting()
    requests: list[dict[str, Any]] = []

    def create_response(**kwargs):
        requests.append(dict(kwargs))
        return SimpleNamespace(output_text="[[DOCX_PARA_p0001]]\nПеревод абзаца.")

    client = _as_openai_client(SimpleNamespace(responses=SimpleNamespace(create=create_response)))

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p0001]]\nThe source paragraph.",
        context_before="",
        context_after="",
        max_retries=3,
        expected_paragraph_ids=["p0001"],
        marker_mode=True,
    )

    assert result == "Перевод абзаца."
    assert len(requests) == 1
    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["degradation_ladder_block_count"] == 0
    assert snapshot["degradation_ladder_model_call_count"] == 0
    assert snapshot["degradation_ladder_trigger_counts"] == {}
    assert snapshot["degradation_ladder_translated_paragraph_count"] == 0


def test_the_ladder_declines_a_block_it_cannot_divide(monkeypatch):
    """ANTI-VACUUM, the second half: an indivisible block buys no extra call.

    One paragraph, one sentence, no markers — dividing it would re-ask the identical text
    and change nothing but the bill. The three calls are the two attempts and the recovery,
    exactly as before, and no ladder counter moves.
    """

    monkeypatch.setattr(generation.time, "sleep", lambda *_: None)
    model_accounting.reset_run_model_accounting()
    requests: list[dict[str, Any]] = []

    def create_response(**kwargs):
        requests.append(dict(kwargs))
        return _incomplete_response()

    client = _as_openai_client(SimpleNamespace(responses=SimpleNamespace(create=create_response)))

    target_text = "Единственное предложение блока без внутренних границ"
    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text=target_text,
        context_before="before",
        context_after="after",
        max_retries=2,
    )

    assert result == target_text
    assert len(requests) == 3
    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["degradation_ladder_block_count"] == 0
    assert snapshot["model_output_discarded_reason_counts"] == {"incomplete_response_source_fallback": 1}


def test_the_ladder_terminates_and_is_bounded_on_a_client_that_always_fails(monkeypatch):
    """The ladder cannot call itself, and one block's calls have an arithmetic ceiling.

    A client that rejects everything is the adversarial case: if a divided call could open a
    ladder of its own, this would not return. It does, ``degradation_ladder_started`` is
    logged exactly ONCE for the block, and the calls stay inside
    ``(paragraphs + 1) * (max_retries + 1)`` — 3 paragraphs and ``max_retries=2`` give
    3 for the block plus 3 per paragraph, i.e. 12.
    """

    logged_events = []
    requests: list[dict[str, Any]] = []
    monkeypatch.setattr(generation.time, "sleep", lambda *_: None)
    monkeypatch.setattr(
        generation,
        "log_event",
        lambda *args, **kwargs: logged_events.append((args, kwargs)) or "evt-bounded",
    )
    model_accounting.reset_run_model_accounting()

    client = _batched_answer_only_client(
        "[[DOCX_PARA_p9999]]\nОтвет с чужим маркером", requests=requests
    )

    result = generation.generate_markdown_block(
        client=client,
        model="gpt-5.4",
        system_prompt="system",
        target_text="[[DOCX_PARA_p1]]\nOne.\n\n[[DOCX_PARA_p2]]\nTwo.\n\n[[DOCX_PARA_p3]]\nThree.",
        context_before="",
        context_after="",
        max_retries=2,
        expected_paragraph_ids=["p1", "p2", "p3"],
        marker_mode=True,
    )

    # It returned at all — that is the first half of the assertion. Nothing was lost either:
    # every paragraph kept its own source, exactly today's outcome.
    assert result == "One.\n\nTwo.\n\nThree."
    assert len(requests) == 12
    started = [args for args, _ in logged_events if len(args) > 1 and args[1] == "degradation_ladder_started"]
    assert len(started) == 1
    snapshot = model_accounting.snapshot_run_model_accounting()
    assert snapshot["degradation_ladder_block_count"] == 1


def test_a_divided_call_never_substitutes_the_source_behind_the_ladders_back(monkeypatch):
    """``allow_controlled_source_fallback=False`` is what keeps the ladder honest.

    A nested call that answered a rejection with the source text would be indistinguishable
    from a translation, and the ladder would report a rescue that shipped English. It raises
    instead — and the same flag is what makes the ladder unable to re-enter itself.
    """

    monkeypatch.setattr(generation.time, "sleep", lambda *_: None)
    client = _as_openai_client(
        SimpleNamespace(responses=SimpleNamespace(create=lambda **_kwargs: _incomplete_response()))
    )

    with pytest.raises(RuntimeError) as exc_info:
        generation.generate_markdown_block(
            client=client,
            model="gpt-5.4",
            system_prompt="system",
            target_text="Один абзац, который модель не осилила",
            context_before="",
            context_after="",
            max_retries=1,
            allow_controlled_source_fallback=False,
        )

    assert "incomplete_response" in str(exc_info.value)


def test_the_run_report_prints_what_the_ladder_did_and_what_it_cost() -> None:
    """The counters have to reach a human, and ``run_summary.txt`` is where he reads them.

    Constitution VIII: without these lines "the ladder works" cannot be checked against a
    run, only believed.
    """

    import importlib.util
    from pathlib import Path

    harness_path = (
        Path(__file__).resolve().parent / "artifacts" / "real_document_pipeline" / "run_lietaer_validation.py"
    )
    spec = importlib.util.spec_from_file_location("run_lietaer_validation_ladder_probe", harness_path)
    assert spec is not None and spec.loader is not None
    harness = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(harness)

    lines = harness._build_model_accounting_summary_lines(
        {
            "degradation_ladder_block_count": 2,
            "degradation_ladder_model_call_count": 11,
            "degradation_ladder_translated_paragraph_count": 11,
            "degradation_ladder_unrescued_paragraph_count": 0,
            "degradation_ladder_sentence_split_paragraph_count": 1,
            "degradation_ladder_oversized_sentence_count": 0,
            "degradation_ladder_trigger_counts": {"marker_contract": 2},
        }
    )

    assert "model_accounting_degradation_ladder_block_count=2" in lines
    assert "model_accounting_degradation_ladder_model_call_count=11" in lines
    assert "model_accounting_degradation_ladder_translated_paragraph_count=11" in lines
    assert "model_accounting_degradation_ladder_unrescued_paragraph_count=0" in lines
    assert "model_accounting_degradation_ladder_sentence_split_paragraph_count=1" in lines
    assert "model_accounting_degradation_ladder_oversized_sentence_count=0" in lines
    assert 'model_accounting_degradation_ladder_trigger_counts={"marker_contract": 2}' in lines


def test_every_ladder_counter_is_seeded_so_a_zero_is_an_assertion() -> None:
    """A missing field is a guess; a zero is a statement. All seven are always published."""

    model_accounting.reset_run_model_accounting()
    snapshot = model_accounting.snapshot_run_model_accounting()

    assert snapshot["degradation_ladder_block_count"] == 0
    assert snapshot["degradation_ladder_model_call_count"] == 0
    assert snapshot["degradation_ladder_translated_paragraph_count"] == 0
    assert snapshot["degradation_ladder_unrescued_paragraph_count"] == 0
    assert snapshot["degradation_ladder_sentence_split_paragraph_count"] == 0
    assert snapshot["degradation_ladder_oversized_sentence_count"] == 0
    assert snapshot["degradation_ladder_trigger_counts"] == {}


def test_the_sentence_split_refuses_to_cut_a_structural_line():
    """Constitution VII rule 7: repairing size must not destroy a role.

    A heading welded onto the sentence after it stops being a heading, so a unit whose own
    text carries a structural line is not divided at all — the same guard
    ``split_marker_preserved_paragraph_dispositions`` uses, not a second one.
    """

    prose = "Первое предложение. Второе предложение. Третье предложение. Четвёртое."
    assert len(generation._ladder_sentence_groups(prose)) >= 2
    assert generation._ladder_sentence_groups(f"## Заголовок\n{prose}") == []
    # A single sentence declines too: dividing it would re-ask the identical text.
    assert generation._ladder_sentence_groups("Одно предложение без границ.") == []
