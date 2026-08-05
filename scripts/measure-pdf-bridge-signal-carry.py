"""Census the PDF->DOCX bridge: what the importer measured, and what survives into the DOCX.

Spec 055 asserts that `pdf_import/logical_import.py` measures real layout evidence and that
`_append_pdf_text_paragraph_to_docx` writes only a style name plus bold/italic, so
`document/extraction.py` re-guesses structure from an impoverished document. This script is the
measurement that has to exist before any of that is changed, and the same script re-run after.

Three sides are counted for every book:

* IMPORTER  -- the `ParagraphUnit` list `build_paragraph_units_from_text_spans` returns, plus the
  span geometry and the layout profile it derived, plus how many paragraphs a purely geometric
  alignment / space-before derivation could reach. That last block is a FEASIBILITY count, not a
  claim: it answers "could the bridge write this", which is exactly the census question.
* BRIDGE    -- the intermediate DOCX that `normalize_uploaded_document` produces for the PDF, read
  back through the very resolvers `extraction.py` uses (`resolve_paragraph_alignment`,
  `resolve_effective_paragraph_font_size`, the `w:spacing w:before` regex).
* NATIVE    -- the same book's hand-made `.docx` twin, read through the same resolvers. This is the
  control group: the difference between BRIDGE and NATIVE is the bridge and nothing else.

Everything here is offline. No LLM client is constructed and none is needed.

Run (WSL, `.venv` has pdfminer):
    . .venv/bin/activate && export PYTHONPATH="$PWD/src:$PWD" \
        && python scripts/measure-pdf-bridge-signal-carry.py tests/sources/book/*.pdf \
            --json-out .run/brg_census.json
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import Any

PROJECT_ROOT = Path(__file__).resolve().parents[1]
SRC_ROOT = PROJECT_ROOT / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

import docxaicorrector.pdf_import.logical_import as logical_import
import docxaicorrector.processing.processing_runtime as processing_runtime
from docxaicorrector.document.roles import (
    resolve_effective_paragraph_font_size,
    resolve_paragraph_alignment,
)
from docxaicorrector.pdf_import.text_layer_quality import extract_pdf_text_spans_with_pdfminer

_SPACING_BEFORE_PATTERN = re.compile(r"w:before=\"(?P<before>-?\d+)\"")

# Signals the importer holds per paragraph that the bridge could in principle express in a DOCX.
_PARAGRAPH_SIGNALS = (
    "font_size_pt",
    "paragraph_alignment",
    "vertical_gap_before_pt",
    "is_bold",
    "is_italic",
    "heading_level",
    "list_kind",
    "style_name",
)


def _populated(value: object) -> bool:
    if value is None:
        return False
    if isinstance(value, bool):
        return bool(value)
    if isinstance(value, str):
        return bool(value.strip())
    return True


# --------------------------------------------------------------------------------------
# Feasibility probes. These mirror what a bridge COULD derive from geometry the importer
# already computes; they classify on coordinates only, never on the shape of the text.
# --------------------------------------------------------------------------------------


def _span_is_centered(span: Any, profile: Any) -> bool:
    left = profile.body_left_x0
    right = profile.body_right_x1
    if left is None or right is None:
        return False
    leading = float(profile.body_leading or 0.0) or 12.0
    column_center = (float(left) + float(right)) / 2.0
    span_center = (float(span.x0) + float(span.x1)) / 2.0
    indented = float(span.x0) - float(left) > leading * 0.5
    balanced = abs(span_center - column_center) <= leading * 0.5
    return indented and balanced


def _feasibility(spans: list[Any], paragraphs: list[Any]) -> dict[str, Any]:
    """How many paragraphs a geometric derivation could reach, and with what geometry."""

    if not spans:
        return {}
    ordered = sorted(spans, key=lambda item: (item.page_number, item.top, item.x0))
    furniture = logical_import._detect_repeated_page_furniture_keys(tuple(ordered))
    profile = logical_import._build_heading_layout_profile(ordered, repeated_furniture_keys=furniture)

    by_origin: dict[int, list[Any]] = {}
    for span in ordered:
        by_origin.setdefault(logical_import._span_origin_index(span), []).append(span)

    # Vertical gap: distance from the bottom of the previous content line on the SAME page.
    previous_bottom: dict[int, float] = {}
    gap_by_origin: dict[int, float] = {}
    last_page: int | None = None
    last_bottom: float | None = None
    for span in ordered:
        key = logical_import._span_origin_index(span)
        if last_page == span.page_number and last_bottom is not None:
            gap_by_origin.setdefault(key, max(0.0, float(span.top) - last_bottom))
        last_page = span.page_number
        last_bottom = float(span.bottom)
    previous_bottom.clear()

    centered_paragraphs = 0
    gap_paragraphs = 0
    font_size_paragraphs = 0
    centered_samples: list[str] = []
    for paragraph in paragraphs:
        origins = list(getattr(paragraph, "origin_raw_indexes", []) or [])
        paragraph_spans = [span for origin in origins for span in by_origin.get(int(origin), [])]
        if paragraph.font_size_pt is not None:
            font_size_paragraphs += 1
        if paragraph_spans and all(_span_is_centered(span, profile) for span in paragraph_spans):
            centered_paragraphs += 1
            if len(centered_samples) < 12:
                centered_samples.append(str(paragraph.text)[:120])
        if origins and int(origins[0]) in gap_by_origin:
            gap_paragraphs += 1

    return {
        "layout_profile": {
            "body_font_size": profile.body_font_size,
            "body_left_x0": profile.body_left_x0,
            "body_right_x1": profile.body_right_x1,
            "body_leading": profile.body_leading,
        },
        "spans_with_x_geometry": sum(
            1 for span in ordered if span.x0 is not None and span.x1 is not None
        ),
        "spans_with_page_height": sum(1 for span in ordered if span.page_height is not None),
        "spans_with_font_size": sum(1 for span in ordered if span.font_size),
        "paragraphs_font_size_derivable": font_size_paragraphs,
        "paragraphs_center_derivable": centered_paragraphs,
        "paragraphs_gap_before_derivable": gap_paragraphs,
        "center_samples": centered_samples,
    }


# --------------------------------------------------------------------------------------
# Censuses
# --------------------------------------------------------------------------------------


def census_importer(pdf_path: Path) -> dict[str, Any]:
    spans = list(extract_pdf_text_spans_with_pdfminer(pdf_path))
    result = logical_import.build_paragraph_units_from_text_spans(spans)
    paragraphs = list(result.paragraphs)

    populated = {name: 0 for name in _PARAGRAPH_SIGNALS}
    font_sizes: Counter[float] = Counter()
    styles: Counter[str] = Counter()
    roles: Counter[str] = Counter()
    structural_roles: Counter[str] = Counter()
    role_chars: Counter[str] = Counter()
    for paragraph in paragraphs:
        for name in _PARAGRAPH_SIGNALS:
            if _populated(getattr(paragraph, name, None)):
                populated[name] += 1
        if paragraph.font_size_pt is not None:
            font_sizes[round(float(paragraph.font_size_pt), 2)] += 1
        styles[str(paragraph.style_name or "")] += 1
        roles[str(paragraph.role or "")] += 1
        structural_roles[str(paragraph.structural_role or "")] += 1
        role_chars[str(paragraph.structural_role or "")] += len(str(paragraph.text or ""))

    return {
        "span_count": len(spans),
        "paragraph_count": len(paragraphs),
        "populated_on_paragraph_unit": populated,
        "distinct_font_size_pt": len(font_sizes),
        "font_size_pt_top": dict(font_sizes.most_common(8)),
        "style_name_counts": dict(styles.most_common()),
        "role_counts": dict(roles.most_common()),
        "structural_role_counts": dict(structural_roles.most_common()),
        "structural_role_chars": dict(role_chars.most_common()),
        "feasibility": _feasibility(spans, paragraphs),
    }


def census_docx_bytes(docx_bytes: bytes) -> dict[str, Any]:
    """Read a DOCX exactly the way `extraction.py` reads it."""

    from docx import Document

    document = Document(BytesIO(docx_bytes))
    paragraphs = list(document.paragraphs)

    alignment_populated = 0
    alignment_values: Counter[str] = Counter()
    font_size_populated = 0
    font_sizes: Counter[float] = Counter()
    spacing_before_populated = 0
    styles: Counter[str] = Counter()
    ppr_xml: Counter[str] = Counter()
    centered_samples: list[str] = []

    for paragraph in paragraphs:
        alignment = resolve_paragraph_alignment(paragraph)
        if alignment is not None:
            alignment_populated += 1
            alignment_values[str(alignment)] += 1
            if str(alignment) == "center" and len(centered_samples) < 12:
                centered_samples.append(paragraph.text[:120])
        size = resolve_effective_paragraph_font_size(paragraph)
        if size is not None:
            font_size_populated += 1
            font_sizes[round(float(size), 2)] += 1
        properties = getattr(paragraph._element, "pPr", None)
        properties_xml = str(properties.xml) if properties is not None else ""
        ppr_xml[properties_xml] += 1
        if _SPACING_BEFORE_PATTERN.search(properties_xml):
            spacing_before_populated += 1
        styles[str(getattr(getattr(paragraph, "style", None), "name", "") or "")] += 1

    return {
        "paragraph_count": len(paragraphs),
        "paragraph_alignment_populated": alignment_populated,
        "paragraph_alignment_values": dict(alignment_values.most_common()),
        "font_size_pt_populated": font_size_populated,
        "distinct_font_size_pt": len(font_sizes),
        "font_size_pt_top": dict(font_sizes.most_common(8)),
        "spacing_before_populated": spacing_before_populated,
        "distinct_paragraph_properties_xml": len(ppr_xml),
        "style_name_counts": dict(styles.most_common(12)),
        "center_samples": centered_samples,
    }


def census_bridge(pdf_path: Path) -> dict[str, Any]:
    normalized = processing_runtime.normalize_uploaded_document(
        filename=pdf_path.name,
        source_bytes=pdf_path.read_bytes(),
    )
    census = census_docx_bytes(normalized.content_bytes)
    census["conversion_backend"] = normalized.conversion_backend
    return census


def _native_twin(pdf_path: Path) -> Path | None:
    candidate = pdf_path.with_suffix(".docx")
    return candidate if candidate.exists() else None


def measure(pdf_path: Path) -> dict[str, Any]:
    payload: dict[str, Any] = {"source": pdf_path.name}
    payload["importer"] = census_importer(pdf_path)
    payload["bridge_docx"] = census_bridge(pdf_path)
    twin = _native_twin(pdf_path)
    if twin is not None:
        payload["native_docx_twin"] = {
            "source": twin.name,
            **census_docx_bytes(twin.read_bytes()),
        }
    return payload


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument("sources", nargs="+", type=Path, help="PDF books to census")
    parser.add_argument("--json-out", type=Path, default=None)
    args = parser.parse_args(argv)

    results: list[dict[str, Any]] = []
    for source_path in args.sources:
        if not source_path.exists():
            print(f"!! missing: {source_path}", flush=True)
            continue
        print(f"== {source_path.name}", flush=True)
        try:
            result = measure(source_path)
        except Exception as exc:  # a failed book must not hide the others
            print(f"   FAILED: {type(exc).__name__}: {exc}", flush=True)
            results.append({"source": source_path.name, "error": f"{type(exc).__name__}: {exc}"})
            continue
        results.append(result)
        importer = result["importer"]
        bridge = result["bridge_docx"]
        native = result.get("native_docx_twin") or {}
        print(
            f"   importer: paragraphs={importer['paragraph_count']}"
            f" font_size={importer['populated_on_paragraph_unit']['font_size_pt']}"
            f" align={importer['populated_on_paragraph_unit']['paragraph_alignment']}"
            f" gap={importer['populated_on_paragraph_unit']['vertical_gap_before_pt']}",
            flush=True,
        )
        feasibility = importer.get("feasibility") or {}
        print(
            f"   derivable: font_size={feasibility.get('paragraphs_font_size_derivable')}"
            f" center={feasibility.get('paragraphs_center_derivable')}"
            f" gap={feasibility.get('paragraphs_gap_before_derivable')}",
            flush=True,
        )
        print(
            f"   bridge   : paragraphs={bridge['paragraph_count']}"
            f" align={bridge['paragraph_alignment_populated']}"
            f" font_size={bridge['font_size_pt_populated']} ({bridge['distinct_font_size_pt']} distinct)"
            f" spacing_before={bridge['spacing_before_populated']}"
            f" pPr={bridge['distinct_paragraph_properties_xml']}",
            flush=True,
        )
        if native:
            print(
                f"   native   : paragraphs={native['paragraph_count']}"
                f" align={native['paragraph_alignment_populated']}"
                f" font_size={native['font_size_pt_populated']} ({native['distinct_font_size_pt']} distinct)"
                f" spacing_before={native['spacing_before_populated']}"
                f" pPr={native['distinct_paragraph_properties_xml']}",
                flush=True,
            )

    if args.json_out:
        args.json_out.parent.mkdir(parents=True, exist_ok=True)
        args.json_out.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"\nwrote {args.json_out}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
