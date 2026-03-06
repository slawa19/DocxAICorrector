import os
import re
import tempfile
import time
import tomllib
from dataclasses import dataclass
from logging.handlers import RotatingFileHandler
from pathlib import Path
import json
import logging

import pypandoc
import streamlit as st
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
PROMPTS_DIR = BASE_DIR / "prompts"
CONFIG_PATH = BASE_DIR / "config.toml"
SYSTEM_PROMPT_PATH = PROMPTS_DIR / "system_prompt.txt"
RUN_DIR = BASE_DIR / ".run"
APP_LOG_PATH = RUN_DIR / "app.log"

st.set_page_config(
    page_title="AI DOCX Editor",
    layout="wide",
)

DEFAULT_MODEL = "gpt-5-mini"
DEFAULT_MODEL_OPTIONS = [
    "gpt-5.4",
    "gpt-5.4-pro",
    "gpt-5.2",
    "gpt-5.1",
    "gpt-5-mini",
]
DEFAULT_CHUNK_SIZE = 6000
DEFAULT_MAX_RETRIES = 3


@dataclass
class ParagraphUnit:
    text: str
    role: str


@dataclass
class DocumentBlock:
    paragraphs: list[ParagraphUnit]

    @property
    def text(self) -> str:
        return "\n\n".join(paragraph.text for paragraph in self.paragraphs)


def setup_logger() -> logging.Logger:
    RUN_DIR.mkdir(parents=True, exist_ok=True)
    logger = logging.getLogger("docxaicorrector")
    if logger.handlers:
        return logger

    logger.setLevel(logging.INFO)
    handler = RotatingFileHandler(APP_LOG_PATH, maxBytes=1_000_000, backupCount=3, encoding="utf-8")
    formatter = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s")
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    logger.propagate = False
    return logger


LOGGER = setup_logger()


def make_event_id(prefix: str = "evt") -> str:
    return f"{prefix}-{int(time.time() * 1000)}"


def sanitize_log_context(value: object) -> object:
    if isinstance(value, (str, int, float, bool)) or value is None:
        return value
    if isinstance(value, Path):
        return str(value)
    if isinstance(value, dict):
        return {str(key): sanitize_log_context(item) for key, item in value.items()}
    if isinstance(value, (list, tuple, set)):
        return [sanitize_log_context(item) for item in value]
    return str(value)


def log_event(level: int, event: str, message: str, **context: object) -> str:
    event_id = make_event_id(event)
    payload = {
        "event_id": event_id,
        "event": event,
        "message": message,
        "context": sanitize_log_context(context),
    }
    LOGGER.log(level, json.dumps(payload, ensure_ascii=False))
    return event_id


def extract_exception_message(exc: Exception) -> str:
    body = getattr(exc, "body", None)
    if isinstance(body, dict):
        error_data = body.get("error")
        if isinstance(error_data, dict):
            message = error_data.get("message")
            if isinstance(message, str) and message.strip():
                return message.strip()

    response = getattr(exc, "response", None)
    if response is not None:
        response_json = getattr(response, "json", None)
        if callable(response_json):
            try:
                response_data = response_json()
                if isinstance(response_data, dict):
                    error_data = response_data.get("error")
                    if isinstance(error_data, dict):
                        message = error_data.get("message")
                        if isinstance(message, str) and message.strip():
                            return message.strip()
            except Exception:
                pass

    return str(exc).strip() or exc.__class__.__name__


def format_user_error(exc: Exception) -> str:
    message = extract_exception_message(exc)
    status_code = getattr(exc, "status_code", None)
    lowered = message.lower()

    if "unsupported parameter" in lowered and "temperature" in lowered:
        return "Выбранная модель не поддерживает параметр temperature. Запрос отправлен с неподдерживаемой настройкой модели."
    if status_code == 400:
        return f"OpenAI отклонил запрос: {message}"
    if status_code == 401:
        return "OpenAI отклонил запрос из-за неверного или отсутствующего API-ключа."
    if status_code == 403:
        return "OpenAI отклонил запрос из-за ограничений доступа к модели или аккаунту."
    if status_code == 404:
        return f"Запрошенная модель или ресурс не найдены: {message}"
    if status_code == 408:
        return "Запрос к OpenAI превысил время ожидания. Попробуйте повторить запуск."
    if status_code == 409:
        return f"OpenAI вернул конфликт состояния запроса: {message}"
    if status_code == 429:
        return "OpenAI временно ограничил запросы. Попробуйте позже или уменьшите нагрузку."
    if isinstance(status_code, int) and status_code >= 500:
        return "OpenAI временно недоступен. Попробуйте повторить запуск позже."
    if exc.__class__.__name__ == "APIConnectionError":
        return "Не удалось подключиться к OpenAI. Проверьте интернет или сетевые ограничения."
    if exc.__class__.__name__ == "APITimeoutError":
        return "OpenAI не ответил вовремя. Попробуйте повторить запуск."
    if isinstance(exc, RuntimeError) or isinstance(exc, ValueError):
        return message
    return f"Непредвиденная ошибка: {message}"


def log_exception(event: str, exc: Exception, message: str, **context: object) -> str:
    event_id = make_event_id(event)
    payload = {
        "event_id": event_id,
        "event": event,
        "message": message,
        "error_type": exc.__class__.__name__,
        "error_message": extract_exception_message(exc),
        "status_code": getattr(exc, "status_code", None),
        "context": sanitize_log_context(context),
    }
    LOGGER.exception(json.dumps(payload, ensure_ascii=False))
    return event_id


def present_error(event: str, exc: Exception, message: str, **context: object) -> str:
    event_id = log_exception(event, exc, message, **context)
    return f"{format_user_error(exc)} [log: {event_id}]"


def parse_int_env(name: str, default: int) -> int:
    raw_value = os.getenv(name, "").strip()
    if not raw_value:
        return default

    try:
        return int(raw_value)
    except ValueError as exc:
        raise RuntimeError(f"Некорректное целое значение в {name}: {raw_value}") from exc


def parse_csv_env(name: str) -> list[str] | None:
    raw_value = os.getenv(name, "").strip()
    if not raw_value:
        return None

    items = [item.strip() for item in raw_value.split(",") if item.strip()]
    if not items:
        raise RuntimeError(f"Переменная {name} задана, но список моделей пуст.")
    return items


def load_app_config() -> dict[str, object]:
    config_data: dict[str, object] = {}
    if CONFIG_PATH.exists():
        with CONFIG_PATH.open("rb") as file_handle:
            config_data = tomllib.load(file_handle)

    model_options = config_data.get("model_options", DEFAULT_MODEL_OPTIONS)
    if not isinstance(model_options, list) or not all(isinstance(item, str) and item.strip() for item in model_options):
        raise RuntimeError(f"Некорректное поле model_options в {CONFIG_PATH}")

    default_model = config_data.get("default_model", DEFAULT_MODEL)
    if not isinstance(default_model, str) or not default_model.strip():
        raise RuntimeError(f"Некорректное поле default_model в {CONFIG_PATH}")

    chunk_size = config_data.get("chunk_size", DEFAULT_CHUNK_SIZE)
    if not isinstance(chunk_size, int):
        raise RuntimeError(f"Некорректное поле chunk_size в {CONFIG_PATH}")

    max_retries = config_data.get("max_retries", DEFAULT_MAX_RETRIES)
    if not isinstance(max_retries, int):
        raise RuntimeError(f"Некорректное поле max_retries в {CONFIG_PATH}")

    env_model_options = parse_csv_env("DOCX_AI_MODEL_OPTIONS")
    if env_model_options is not None:
        model_options = env_model_options

    default_model = os.getenv("DOCX_AI_DEFAULT_MODEL", default_model).strip() or default_model
    chunk_size = parse_int_env("DOCX_AI_CHUNK_SIZE", chunk_size)
    max_retries = parse_int_env("DOCX_AI_MAX_RETRIES", max_retries)

    if default_model not in model_options:
        model_options = [default_model, *[item for item in model_options if item != default_model]]

    return {
        "default_model": default_model,
        "model_options": model_options,
        "chunk_size": max(3000, min(chunk_size, 12000)),
        "max_retries": max(1, min(max_retries, 5)),
    }


def load_system_prompt() -> str:
    try:
        prompt_text = SYSTEM_PROMPT_PATH.read_text(encoding="utf-8").strip()
    except FileNotFoundError as exc:
        raise RuntimeError(f"Не найден файл системного промпта: {SYSTEM_PROMPT_PATH}") from exc

    if not prompt_text:
        raise RuntimeError(f"Файл системного промпта пуст: {SYSTEM_PROMPT_PATH}")

    return prompt_text


def get_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("Не найден OPENAI_API_KEY. Добавьте его в .env или переменные окружения.")
    return OpenAI(api_key=api_key)


def ensure_pandoc_available() -> None:
    try:
        pypandoc.get_pandoc_version()
    except OSError as exc:
        raise RuntimeError(
            "Pandoc не найден. Для Windows PowerShell установите его командой: "
            "winget install --id JohnMacFarlane.Pandoc -e"
        ) from exc


def classify_paragraph_role(text: str, style_name: str) -> str:
    normalized_style = style_name.strip().lower()
    stripped_text = text.lstrip()

    if normalized_style.startswith("heading") or normalized_style.startswith("заголовок"):
        return "heading"

    if "list" in normalized_style or "спис" in normalized_style:
        return "list"

    if stripped_text.startswith(("- ", "* ", "• ", "— ")):
        return "list"

    if re.match(r"^\d+[\.)]\s+", stripped_text):
        return "list"

    return "body"


def extract_paragraph_units_from_docx(uploaded_file) -> list[ParagraphUnit]:
    uploaded_file.seek(0)
    document = Document(uploaded_file)
    paragraphs: list[ParagraphUnit] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
        paragraphs.append(ParagraphUnit(text=text, role=classify_paragraph_role(text, style_name)))

    if not paragraphs:
        raise ValueError("В документе не найден текст для обработки.")
    return paragraphs


def build_document_text(paragraphs: list[ParagraphUnit]) -> str:
    return "\n\n".join(paragraph.text for paragraph in paragraphs).strip()


def build_semantic_blocks(paragraphs: list[ParagraphUnit], max_chars: int = 6000) -> list[DocumentBlock]:
    if not paragraphs:
        return []

    soft_limit = max(1200, min(max_chars, int(max_chars * 0.7)))
    blocks: list[DocumentBlock] = []
    current: list[ParagraphUnit] = []
    current_size = 0

    def flush_current() -> None:
        nonlocal current, current_size
        if current:
            blocks.append(DocumentBlock(paragraphs=current))
            current = []
            current_size = 0

    def append_paragraph(paragraph: ParagraphUnit) -> None:
        nonlocal current_size
        separator_size = 2 if current else 0
        current.append(paragraph)
        current_size += separator_size + len(paragraph.text)

    for paragraph in paragraphs:
        if not current:
            append_paragraph(paragraph)
            continue

        projected_size = current_size + 2 + len(paragraph.text)
        current_only_heading = len(current) == 1 and current[0].role == "heading"
        current_is_list = all(item.role == "list" for item in current)

        if paragraph.role == "heading":
            flush_current()
            append_paragraph(paragraph)
            continue

        if current_only_heading:
            append_paragraph(paragraph)
            continue

        if current_is_list and paragraph.role == "list":
            if projected_size <= max_chars or current_size < soft_limit:
                append_paragraph(paragraph)
            else:
                flush_current()
                append_paragraph(paragraph)
            continue

        if current_is_list and paragraph.role != "list":
            if current_size >= max(600, soft_limit // 2) or len(current) > 1:
                flush_current()
                append_paragraph(paragraph)
                continue

        if projected_size <= max_chars and current_size < soft_limit:
            append_paragraph(paragraph)
            continue

        if projected_size <= max_chars and len(paragraph.text) <= max(500, max_chars // 4) and current_size < int(max_chars * 0.9):
            append_paragraph(paragraph)
            continue

        flush_current()
        append_paragraph(paragraph)

    flush_current()
    return blocks


def build_context_excerpt(blocks: list[DocumentBlock], block_index: int, limit_chars: int, *, reverse: bool) -> str:
    if limit_chars <= 0:
        return ""

    indexes = range(block_index - 1, -1, -1) if reverse else range(block_index + 1, len(blocks))
    collected: list[str] = []
    total_size = 0

    for index in indexes:
        block_text = blocks[index].text.strip()
        if not block_text:
            continue

        separator_size = 2 if collected else 0
        projected_size = total_size + separator_size + len(block_text)
        if projected_size <= limit_chars:
            collected.append(block_text)
            total_size = projected_size
            continue

        remaining = limit_chars - total_size - separator_size
        if remaining > 0:
            excerpt = block_text[-remaining:] if reverse else block_text[:remaining]
            if excerpt.strip():
                collected.append(excerpt.strip())
        break

    if reverse:
        collected.reverse()

    return "\n\n".join(collected).strip()


def build_editing_jobs(blocks: list[DocumentBlock], max_chars: int) -> list[dict[str, str | int]]:
    context_before_chars = max(600, min(1400, int(max_chars * 0.2)))
    context_after_chars = max(300, min(800, int(max_chars * 0.12)))
    jobs: list[dict[str, str | int]] = []

    for index, block in enumerate(blocks):
        context_before = build_context_excerpt(blocks, index, context_before_chars, reverse=True)
        context_after = build_context_excerpt(blocks, index, context_after_chars, reverse=False)
        jobs.append(
            {
                "target_text": block.text,
                "context_before": context_before,
                "context_after": context_after,
                "target_chars": len(block.text),
                "context_chars": len(context_before) + len(context_after),
            }
        )

    return jobs


def normalize_model_output(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```markdown"):
        cleaned = cleaned[len("```markdown") :].strip()
    elif cleaned.startswith("```"):
        cleaned = cleaned[len("```") :].strip()
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3].strip()
    return cleaned


def is_retryable_error(exc: Exception) -> bool:
    status_code = getattr(exc, "status_code", None)
    if status_code in {408, 409, 429}:
        return True
    if isinstance(status_code, int) and status_code >= 500:
        return True
    return exc.__class__.__name__ in {"APIConnectionError", "APITimeoutError"}


def generate_markdown_block(
    client: OpenAI,
    model: str,
    system_prompt: str,
    target_text: str,
    context_before: str,
    context_after: str,
    max_retries: int,
) -> str:
    context_before_text = context_before or "[контекст отсутствует]"
    context_after_text = context_after or "[контекст отсутствует]"
    user_prompt = (
        "Ниже передан целевой блок документа и соседний контекст.\n"
        "Используй соседний контекст только для понимания смысла, терминологии и связности.\n"
        "Редактируй только целевой блок и верни только его итоговый текст.\n\n"
        f"КОНТЕКСТ ДО (только для понимания, не переписывать):\n{context_before_text}\n\n"
        f"ЦЕЛЕВОЙ БЛОК (отредактируй и верни только его):\n{target_text}\n\n"
        f"КОНТЕКСТ ПОСЛЕ (только для понимания, не переписывать):\n{context_after_text}"
    )
    payload = [
        {
            "role": "system",
            "content": [{"type": "input_text", "text": system_prompt}],
        },
        {
            "role": "user",
            "content": [{"type": "input_text", "text": user_prompt}],
        },
    ]

    for attempt in range(1, max_retries + 1):
        try:
            response = client.responses.create(model=model, input=payload)
            markdown = normalize_model_output(response.output_text)
            if not markdown:
                raise RuntimeError("Модель вернула пустой ответ.")
            return markdown
        except Exception as exc:
            should_retry = attempt < max_retries and is_retryable_error(exc)
            if not should_retry:
                raise
            time.sleep(min(2 ** (attempt - 1), 8))

    raise RuntimeError("Не удалось получить ответ модели.")


def convert_markdown_to_docx_bytes(markdown_text: str) -> bytes:
    ensure_pandoc_available()
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            markdown_path = temp_path / "result.md"
            docx_path = temp_path / "result.docx"
            markdown_path.write_text(markdown_text, encoding="utf-8")
            pypandoc.convert_file(
                str(markdown_path),
                to="docx",
                format="md",
                outputfile=str(docx_path),
            )
            return docx_path.read_bytes()
    except Exception as exc:
        raise RuntimeError(f"Ошибка при сборке DOCX: {exc}") from exc


def build_output_filename(filename: str) -> str:
    return f"{Path(filename).stem}_edited.docx"


def build_markdown_filename(filename: str) -> str:
    return f"{Path(filename).stem}_edited.md"


def init_session_state() -> None:
    st.session_state.setdefault("run_log", [])
    st.session_state.setdefault("latest_markdown", "")
    st.session_state.setdefault("latest_docx_bytes", None)
    st.session_state.setdefault("latest_source_name", "")
    st.session_state.setdefault("last_error", "")
    st.session_state.setdefault("last_log_hint", f"Подробный лог приложения: {APP_LOG_PATH}")


def append_log(
    status: str,
    block_index: int,
    block_count: int,
    target_chars: int,
    context_chars: int,
    details: str,
) -> None:
    st.session_state.run_log.append(
        {
            "status": status,
            "block_index": block_index,
            "block_count": block_count,
            "target_chars": target_chars,
            "context_chars": context_chars,
            "details": details,
        }
    )


def reset_run_state() -> None:
    st.session_state.run_log = []
    st.session_state.latest_markdown = ""
    st.session_state.latest_docx_bytes = None
    st.session_state.latest_source_name = ""
    st.session_state.last_error = ""


def render_run_log() -> None:
    if not st.session_state.run_log:
        return

    with st.expander("Журнал обработки", expanded=True):
        for entry in st.session_state.run_log:
            st.write(
                f"[{entry['status']}] Блок {entry['block_index']}/{entry['block_count']} | "
                f"цель: {entry['target_chars']} симв. | контекст: {entry['context_chars']} симв. | {entry['details']}"
            )
        st.caption(st.session_state.last_log_hint)


def render_sidebar(config: dict[str, object]) -> tuple[str, int, int]:
    st.sidebar.header("Настройки")
    model_options = [*config["model_options"], "custom"]
    default_model = str(config["default_model"])
    default_index = model_options.index(default_model) if default_model in model_options else 0
    selected_model = st.sidebar.selectbox("Модель", model_options, index=default_index)
    custom_model = ""
    if selected_model == "custom":
        custom_model = st.sidebar.text_input("Имя модели", value=default_model).strip()

    model = custom_model or selected_model
    chunk_size = st.sidebar.slider(
        "Размер целевого блока, символов",
        min_value=3000,
        max_value=12000,
        value=int(config["chunk_size"]),
        step=500,
    )
    max_retries = st.sidebar.slider(
        "Количество retry",
        min_value=1,
        max_value=5,
        value=int(config["max_retries"]),
    )
    return model, chunk_size, max_retries


def render_result(docx_bytes: bytes, markdown_text: str, original_filename: str) -> None:
    st.success("Документ обработан.")
    st.download_button(
        label="Скачать итоговый DOCX",
        data=docx_bytes,
        file_name=build_output_filename(original_filename),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
    st.download_button(
        label="Скачать итоговый Markdown",
        data=markdown_text.encode("utf-8"),
        file_name=build_markdown_filename(original_filename),
        mime="text/markdown",
        use_container_width=True,
    )

    with st.expander("Предпросмотр Markdown", expanded=True):
        st.text_area("Результат", value=markdown_text, height=420)


def render_partial_result() -> None:
    markdown_text = st.session_state.latest_markdown
    if not markdown_text or st.session_state.latest_docx_bytes is not None:
        return

    source_name = st.session_state.latest_source_name or "result.docx"
    st.warning("Доступен промежуточный Markdown-результат последнего запуска.")
    st.download_button(
        label="Скачать текущий Markdown",
        data=markdown_text.encode("utf-8"),
        file_name=build_markdown_filename(source_name),
        mime="text/markdown",
        use_container_width=True,
    )

    with st.expander("Текущий Markdown", expanded=False):
        st.text_area("Промежуточный результат", value=markdown_text, height=320)


def main() -> None:
    init_session_state()
    log_event(logging.INFO, "app_start", "Приложение инициализировано")
    st.title("AI-редактор DOCX через Markdown")
    st.write(
        "Загрузите DOCX, приложение соберет смысловые блоки из нескольких абзацев, "
        "добавит соседний контекст для модели и соберет новый DOCX."
    )

    try:
        app_config = load_app_config()
    except Exception as exc:
        user_message = present_error("config_load_failed", exc, "Ошибка загрузки конфигурации")
        st.error(f"Ошибка загрузки конфигурации: {user_message}")
        return

    model, chunk_size, max_retries = render_sidebar(app_config)
    uploaded_file = st.file_uploader("Загрузите DOCX-файл", type=["docx"])

    if st.button("Сбросить результаты", use_container_width=True):
        reset_run_state()
        st.rerun()

    if not uploaded_file:
        st.info("Ожидается файл .docx")
        render_run_log()
        render_partial_result()
        return

    try:
        paragraphs = extract_paragraph_units_from_docx(uploaded_file)
        source_text = build_document_text(paragraphs)
        blocks = build_semantic_blocks(paragraphs, max_chars=chunk_size)
        jobs = build_editing_jobs(blocks, max_chars=chunk_size)
        log_event(
            logging.INFO,
            "document_prepared",
            "Документ подготовлен к обработке",
            filename=uploaded_file.name,
            paragraph_count=len(paragraphs),
            block_count=len(jobs),
            source_chars=len(source_text),
            chunk_size=chunk_size,
        )
    except Exception as exc:
        user_message = present_error(
            "document_read_failed",
            exc,
            "Ошибка чтения документа",
            filename=uploaded_file.name,
        )
        st.error(f"Ошибка чтения документа: {user_message}")
        return

    st.caption(
        f"Символов: {len(source_text)} | Абзацев: {len(paragraphs)} | Блоков: {len(jobs)}"
    )

    if len(jobs) == 1:
        st.info("Документ помещается в один блок. Для длинных файлов обработка пойдет по блокам с соседним контекстом.")

    if st.session_state.last_error:
        st.error(st.session_state.last_error)
        st.caption(st.session_state.last_log_hint)

    render_run_log()
    render_partial_result()

    if st.session_state.latest_docx_bytes and st.session_state.latest_source_name == uploaded_file.name:
        render_result(st.session_state.latest_docx_bytes, st.session_state.latest_markdown, uploaded_file.name)

    if st.button("Начать обработку", type="primary", use_container_width=True):
        reset_run_state()
        st.session_state.latest_source_name = uploaded_file.name

        try:
            client = get_client()
            ensure_pandoc_available()
            system_prompt = load_system_prompt()
            log_event(
                logging.INFO,
                "processing_started",
                "Запуск обработки документа",
                filename=uploaded_file.name,
                model=model,
                block_count=len(jobs),
                max_retries=max_retries,
            )
        except Exception as exc:
            st.session_state.last_error = present_error(
                "processing_init_failed",
                exc,
                "Ошибка инициализации обработки",
                filename=uploaded_file.name,
                model=model,
            )
            st.error(st.session_state.last_error)
            st.caption(st.session_state.last_log_hint)
            return

        processed_chunks: list[str] = []
        progress_bar = st.progress(0.0)
        status_text = st.empty()
        started_at = time.perf_counter()

        for index, job in enumerate(jobs, start=1):
            status_text.text(f"Обработка блока {index} из {len(jobs)}")
            log_event(
                logging.INFO,
                "block_started",
                "Начата обработка блока",
                filename=uploaded_file.name,
                block_index=index,
                block_count=len(jobs),
                target_chars=int(job["target_chars"]),
                context_chars=int(job["context_chars"]),
                model=model,
            )
            try:
                processed_chunk = generate_markdown_block(
                    client=client,
                    model=model,
                    system_prompt=system_prompt,
                    target_text=str(job["target_text"]),
                    context_before=str(job["context_before"]),
                    context_after=str(job["context_after"]),
                    max_retries=max_retries,
                )
            except Exception as exc:
                st.session_state.latest_markdown = "\n\n".join(processed_chunks).strip()
                error_message = present_error(
                    "block_failed",
                    exc,
                    "Ошибка обработки блока",
                    filename=uploaded_file.name,
                    block_index=index,
                    block_count=len(jobs),
                    target_chars=int(job["target_chars"]),
                    context_chars=int(job["context_chars"]),
                    model=model,
                )
                st.session_state.last_error = f"Ошибка на блоке {index}: {error_message}"
                append_log(
                    "ERROR",
                    index,
                    len(jobs),
                    int(job["target_chars"]),
                    int(job["context_chars"]),
                    error_message,
                )
                st.error(st.session_state.last_error)
                st.caption(st.session_state.last_log_hint)
                render_run_log()
                render_partial_result()
                return

            processed_chunks.append(processed_chunk)
            st.session_state.latest_markdown = "\n\n".join(processed_chunks).strip()
            append_log(
                "OK",
                index,
                len(jobs),
                int(job["target_chars"]),
                int(job["context_chars"]),
                f"готово за {time.perf_counter() - started_at:.1f} сек. с начала запуска",
            )
            log_event(
                logging.INFO,
                "block_completed",
                "Блок обработан успешно",
                filename=uploaded_file.name,
                block_index=index,
                block_count=len(jobs),
                target_chars=int(job["target_chars"]),
                context_chars=int(job["context_chars"]),
                output_chars=len(processed_chunk),
            )
            progress_bar.progress(index / len(jobs))

        final_markdown = "\n\n".join(processed_chunks).strip()
        st.session_state.latest_markdown = final_markdown
        status_text.text("Сборка итогового DOCX")

        try:
            docx_bytes = convert_markdown_to_docx_bytes(final_markdown)
        except Exception as exc:
            error_message = present_error(
                "docx_build_failed",
                exc,
                "Ошибка сборки DOCX",
                filename=uploaded_file.name,
                final_markdown_chars=len(final_markdown),
            )
            st.session_state.last_error = error_message
            append_log("ERROR", len(jobs), len(jobs), len(final_markdown), 0, error_message)
            st.error(st.session_state.last_error)
            st.caption(st.session_state.last_log_hint)
            render_run_log()
            render_partial_result()
            return

        st.session_state.latest_docx_bytes = docx_bytes
        st.session_state.last_error = ""
        log_event(
            logging.INFO,
            "processing_completed",
            "Документ обработан полностью",
            filename=uploaded_file.name,
            block_count=len(jobs),
            final_markdown_chars=len(final_markdown),
            elapsed_seconds=round(time.perf_counter() - started_at, 2),
        )
        append_log(
            "DONE",
            len(jobs),
            len(jobs),
            len(final_markdown),
            0,
            f"весь документ обработан за {time.perf_counter() - started_at:.1f} сек.",
        )
        render_run_log()
        render_result(docx_bytes, final_markdown, uploaded_file.name)


if __name__ == "__main__":
    main()