import html
import hashlib
import json
import re
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any, cast

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import lxml.etree as etree

from constants import (
    MAX_DOCX_ARCHIVE_SIZE_BYTES,
    MAX_DOCX_COMPRESSION_RATIO,
    MAX_DOCX_ENTRY_COUNT,
    MAX_DOCX_UNCOMPRESSED_SIZE_BYTES,
)
from image_shared import (
    call_responses_create_with_retry,
    extract_model_response_error_code,
    extract_response_text,
    is_retryable_error,
    parse_json_object,
)
from models import (
    PARAGRAPH_BOUNDARY_AI_REVIEW_MODE_VALUES,
    PARAGRAPH_BOUNDARY_NORMALIZATION_MODE_VALUES,
    RELATION_NORMALIZATION_KIND_VALUES,
    DocumentBlock,
    ImageAsset,
    ParagraphBoundaryDecision,
    ParagraphBoundaryNormalizationReport,
    ParagraphRelation,
    ParagraphRelationDecision,
    ParagraphUnit,
    RawBlock,
    RawParagraph,
    RawTable,
    RelationNormalizationReport,
)
from processing_runtime import normalize_uploaded_document, read_uploaded_file_bytes, resolve_uploaded_filename

IMAGE_PLACEHOLDER_PATTERN = re.compile(r"\[\[DOCX_IMAGE_img_\d+\]\]")
PARAGRAPH_MARKER_PATTERN = re.compile(r"\[\[DOCX_PARA_([A-Za-z0-9_]+)\]\]")
IMAGE_ONLY_PATTERN = re.compile(r"^(?:\s*\[\[DOCX_IMAGE_img_\d+\]\]\s*)+$")
CAPTION_PREFIX_PATTERN = re.compile(r"^(?:рис\.?|рисунок|figure|fig\.?|табл\.?|таблица|table)\b", re.IGNORECASE)
HEADING_STYLE_PATTERN = re.compile(r"^(?:heading|заголовок)\s*(\d+)?$", re.IGNORECASE)
COMPARE_ALL_VARIANT_LABELS = {
    "safe": "Вариант 1: Просто улучшить",
    "semantic_redraw_direct": "Вариант 2: Креативная AI-перерисовка",
    "semantic_redraw_structured": "Вариант 3: Структурная AI-перерисовка",
}
MANUAL_REVIEW_SAFE_LABEL = "safe"
RELATIONSHIP_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
ORDERED_LIST_FORMATS = {
    "aiueo",
    "cardinalText",
    "chicago",
    "decimal",
    "decimalEnclosedCircle",
    "decimalEnclosedFullstop",
    "decimalEnclosedParen",
    "decimalFullWidth",
    "decimalFullWidth2",
    "decimalHalfWidth",
    "ganada",
    "hebrew1",
    "hebrew2",
    "hindiConsonants",
    "hindiCounting",
    "hindiNumbers",
    "hindiVowels",
    "ideographDigital",
    "ideographEnclosedCircle",
    "ideographLegalTraditional",
    "ideographTraditional",
    "iroha",
    "japaneseCounting",
    "japaneseDigitalTenThousand",
    "japaneseLegal",
    "koreanCounting",
    "koreanDigital",
    "koreanDigital2",
    "koreanLegal",
    "lowerLetter",
    "lowerRoman",
    "numberInDash",
    "ordinal",
    "ordinalText",
    "russianLower",
    "russianUpper",
    "taiwaneseCounting",
    "taiwaneseCountingThousand",
    "taiwaneseDigital",
    "thaiCounting",
    "thaiLetters",
    "thaiNumbers",
    "upperLetter",
    "upperRoman",
    "vietnameseCounting",
}
UNORDERED_LIST_FORMATS = {"bullet", "none"}
INLINE_HTML_TAG_PATTERN = re.compile(r"</?(?:u|sup|sub)>", re.IGNORECASE)
MARKDOWN_LINK_PATTERN = re.compile(r"\[([^\]]+)\]\([^\)]+\)")
STRONG_PARAGRAPH_TERMINATOR_PATTERN = re.compile(r"[.!?…]\s*$")
TOC_ENTRY_PATTERN = re.compile(r"^.{1,120}(?:\.{2,}|\s{2,})\d+\s*$")
PARAGRAPH_BOUNDARY_REPORTS_DIR = Path(".run") / "paragraph_boundary_reports"
RELATION_NORMALIZATION_REPORTS_DIR = Path(".run") / "relation_normalization_reports"
PARAGRAPH_BOUNDARY_AI_REVIEW_DIR = Path(".run") / "paragraph_boundary_ai_review"


def classify_paragraph_role(text: str, style_name: str, *, heading_level: int | None = None) -> str:
    normalized_style = style_name.strip().lower()
    stripped_text = text.lstrip()

    if _is_image_only_text(text):
        return "image"

    if heading_level is not None:
        return "heading"

    if _is_caption_style(normalized_style):
        return "caption"

    if "list" in normalized_style or "спис" in normalized_style:
        return "list"

    if stripped_text.startswith(("- ", "* ", "• ")):
        return "list"

    if re.match(r"^\d+[\.)]\s+", stripped_text):
        return "list"

    return "body"


def _infer_role_confidence(
    *,
    role: str,
    text: str,
    normalized_style: str,
    explicit_heading_level: int | None,
    heading_source: str | None,
) -> str:
    if role in {"image", "table"}:
        return "explicit"
    if role == "heading":
        return "explicit" if explicit_heading_level is not None or heading_source == "explicit" else "heuristic"
    if role == "caption":
        return "explicit" if _is_caption_style(normalized_style) else "heuristic"
    if role == "list":
        if "list" in normalized_style or "спис" in normalized_style or _detect_explicit_list_kind(text) is not None:
            return "explicit"
    return "heuristic"


def _assign_paragraph_identity(paragraph: ParagraphUnit, source_index: int) -> None:
    paragraph.source_index = source_index
    paragraph.paragraph_id = f"p{source_index:04d}"
    if not paragraph.structural_role or paragraph.structural_role == "body":
        paragraph.structural_role = paragraph.role
    if not paragraph.origin_raw_indexes:
        paragraph.origin_raw_indexes = [source_index]
    if not paragraph.origin_raw_texts:
        paragraph.origin_raw_texts = [paragraph.text]


def extract_paragraph_units_from_docx(uploaded_file) -> list[ParagraphUnit]:
    paragraphs, _ = extract_document_content_from_docx(uploaded_file)
    return paragraphs


def extract_inline_images(uploaded_file) -> list[ImageAsset]:
    _, image_assets = extract_document_content_from_docx(uploaded_file)
    return image_assets


def extract_document_content_from_docx(uploaded_file) -> tuple[list[ParagraphUnit], list[ImageAsset]]:
    paragraphs, image_assets, _, _, _ = extract_document_content_with_normalization_reports(uploaded_file)
    return paragraphs, image_assets


def extract_document_content_with_normalization_reports(
    uploaded_file,
) -> tuple[
    list[ParagraphUnit],
    list[ImageAsset],
    ParagraphBoundaryNormalizationReport,
    list[ParagraphRelation],
    RelationNormalizationReport,
]:
    source_bytes = _read_uploaded_docx_bytes(uploaded_file)
    validate_docx_source_bytes(source_bytes)
    document = Document(BytesIO(source_bytes))
    raw_blocks, image_assets = _build_raw_document_blocks(document)
    normalization_mode, save_boundary_debug_artifacts = _resolve_paragraph_boundary_normalization_settings()
    normalized_blocks, boundary_report = _normalize_paragraph_boundaries(raw_blocks, mode=normalization_mode)
    paragraphs = _build_logical_paragraph_units(normalized_blocks)
    _promote_short_standalone_headings(paragraphs)
    (
        relation_enabled,
        relation_profile,
        enabled_relation_kinds,
        save_relation_debug_artifacts,
    ) = _resolve_relation_normalization_settings()
    (
        ai_review_enabled,
        ai_review_mode,
        ai_review_candidate_limit,
        ai_review_timeout_seconds,
        ai_review_max_tokens_per_candidate,
        ai_review_model,
    ) = _resolve_paragraph_boundary_ai_review_settings()
    relations, relation_report = build_paragraph_relations(
        paragraphs,
        enabled_relation_kinds=enabled_relation_kinds if relation_enabled else (),
    )
    _apply_relation_side_effects(paragraphs, relations)
    _reclassify_adjacent_captions(paragraphs)

    if ai_review_enabled and ai_review_mode != "off":
        _run_paragraph_boundary_ai_review(
            source_name=resolve_uploaded_filename(uploaded_file),
            source_bytes=source_bytes,
            mode=ai_review_mode,
            model=ai_review_model,
            raw_blocks=raw_blocks,
            paragraphs=paragraphs,
            boundary_report=boundary_report,
            relation_report=relation_report,
            candidate_limit=ai_review_candidate_limit,
            timeout_seconds=ai_review_timeout_seconds,
            max_tokens_per_candidate=ai_review_max_tokens_per_candidate,
        )

    if save_boundary_debug_artifacts:
        _write_paragraph_boundary_report_artifact(
            source_name=resolve_uploaded_filename(uploaded_file),
            source_bytes=source_bytes,
            mode=normalization_mode,
            report=boundary_report,
        )
    if relation_enabled and save_relation_debug_artifacts:
        _write_relation_normalization_report_artifact(
            source_name=resolve_uploaded_filename(uploaded_file),
            source_bytes=source_bytes,
            profile=relation_profile,
            enabled_relation_kinds=enabled_relation_kinds,
            report=relation_report,
        )

    if not paragraphs:
        raise ValueError("В документе не найден текст для обработки.")
    return paragraphs, image_assets, boundary_report, relations, relation_report


def extract_document_content_with_boundary_report(
    uploaded_file,
) -> tuple[list[ParagraphUnit], list[ImageAsset], ParagraphBoundaryNormalizationReport]:
    paragraphs, image_assets, boundary_report, _, _ = extract_document_content_with_normalization_reports(uploaded_file)
    return paragraphs, image_assets, boundary_report


def build_document_text(paragraphs: list[ParagraphUnit]) -> str:
    return "\n\n".join(paragraph.rendered_text for paragraph in paragraphs).strip()


def _resolve_marker_paragraph_id(paragraph: ParagraphUnit, fallback_index: int) -> str:
    if paragraph.paragraph_id:
        return paragraph.paragraph_id
    if paragraph.source_index >= 0:
        return f"p{paragraph.source_index:04d}"
    return f"p{fallback_index:04d}"


def build_marker_wrapped_block_text(paragraphs: list[ParagraphUnit], *, paragraph_ids: list[str] | None = None) -> str:
    parts: list[str] = []
    for index, paragraph in enumerate(paragraphs):
        paragraph_id = paragraph_ids[index] if paragraph_ids is not None else _resolve_marker_paragraph_id(paragraph, index)
        parts.append(f"[[DOCX_PARA_{paragraph_id}]]\n{paragraph.rendered_text}")
    return "\n\n".join(parts).strip()


def _paragraph_structural_kind(paragraph: ParagraphUnit) -> str:
    return str(getattr(paragraph, "structural_role", None) or getattr(paragraph, "role", None) or "").strip().lower()


def _is_quote_structural_role(paragraph: ParagraphUnit) -> bool:
    return _paragraph_structural_kind(paragraph) in {"epigraph", "attribution", "dedication"}


def _is_toc_structural_role(paragraph: ParagraphUnit) -> bool:
    return _paragraph_structural_kind(paragraph) in {"toc_header", "toc_entry"}


def inspect_placeholder_integrity(markdown_text: str, image_assets: list[ImageAsset]) -> dict[str, str]:
    status_map: dict[str, str] = {}
    expected_placeholders = {asset.placeholder for asset in image_assets}
    for asset in image_assets:
        occurrence_count = markdown_text.count(asset.placeholder)
        if occurrence_count == 1:
            status_map[asset.image_id] = "ok"
        elif occurrence_count == 0:
            status_map[asset.image_id] = "lost"
        else:
            status_map[asset.image_id] = "duplicated"
    for unexpected_placeholder in sorted(set(IMAGE_PLACEHOLDER_PATTERN.findall(markdown_text)) - expected_placeholders):
        status_map[f"unexpected:{unexpected_placeholder}"] = "unexpected"
    return status_map


def build_semantic_blocks(
    paragraphs: list[ParagraphUnit],
    max_chars: int = 6000,
    *,
    relations: list[ParagraphRelation] | None = None,
) -> list[DocumentBlock]:
    if not paragraphs:
        return []

    resolved_relations = relations
    if resolved_relations is None:
        resolved_relations, _ = build_paragraph_relations(
            paragraphs,
            enabled_relation_kinds=resolve_effective_relation_kinds(),
        )
    paragraph_units = _build_semantic_block_units(paragraphs, resolved_relations)
    soft_limit = max(1200, min(max_chars, int(max_chars * 0.7)))
    blocks: list[DocumentBlock] = []
    current: list[ParagraphUnit] = []
    current_size = 0

    def flush_current() -> None:
        nonlocal current, current_size
        if current:
            blocks.append(DocumentBlock(paragraphs=current))
            current = []
            current_size = 0

    def append_unit(unit_paragraphs: list[ParagraphUnit]) -> None:
        nonlocal current_size
        separator_size = 2 if current else 0
        current.extend(unit_paragraphs)
        unit_text = "\n\n".join(paragraph.rendered_text for paragraph in unit_paragraphs)
        current_size += separator_size + len(unit_text)

    for unit_paragraphs in paragraph_units:
        unit_text = "\n\n".join(paragraph.rendered_text for paragraph in unit_paragraphs)
        unit_contains_atomic_block = any(paragraph.role in {"image", "table"} for paragraph in unit_paragraphs)
        unit_all_headings = all(paragraph.role == "heading" for paragraph in unit_paragraphs)
        unit_is_list = all(paragraph.role == "list" for paragraph in unit_paragraphs)
        unit_is_quote_cluster = bool(unit_paragraphs) and all(_is_quote_structural_role(paragraph) for paragraph in unit_paragraphs)
        unit_is_toc_cluster = bool(unit_paragraphs) and all(_is_toc_structural_role(paragraph) for paragraph in unit_paragraphs)
        if not current:
            append_unit(unit_paragraphs)
            continue

        current_contains_atomic_block = any(item.role in {"image", "table"} for item in current)
        if current_contains_atomic_block:
            flush_current()
            append_unit(unit_paragraphs)
            continue

        if unit_contains_atomic_block:
            flush_current()
            append_unit(unit_paragraphs)
            continue

        projected_size = current_size + 2 + len(unit_text)
        current_all_headings = all(item.role == "heading" for item in current)
        current_is_list = all(item.role == "list" for item in current)
        current_is_toc_cluster = bool(current) and all(_is_toc_structural_role(item) for item in current)

        if unit_is_toc_cluster and not current_is_toc_cluster:
            flush_current()
            append_unit(unit_paragraphs)
            continue

        if current_is_toc_cluster and not unit_is_toc_cluster:
            flush_current()
            append_unit(unit_paragraphs)
            continue

        if unit_all_headings:
            if current_all_headings:
                append_unit(unit_paragraphs)
                continue
            flush_current()
            append_unit(unit_paragraphs)
            continue

        if current_all_headings:
            append_unit(unit_paragraphs)
            continue

        if current[-1].role == "heading" and unit_is_quote_cluster:
            append_unit(unit_paragraphs)
            continue

        if current[-1].role == "heading" and all(paragraph.role == "caption" for paragraph in unit_paragraphs):
            append_unit(unit_paragraphs)
            continue

        if current_is_list and unit_is_list:
            if projected_size <= max_chars or current_size < soft_limit:
                append_unit(unit_paragraphs)
            else:
                flush_current()
                append_unit(unit_paragraphs)
            continue

        if current_is_list and not unit_is_list:
            if current_size >= max(600, soft_limit // 2) or len(current) > 1:
                flush_current()
                append_unit(unit_paragraphs)
                continue

        if projected_size <= max_chars and current_size < soft_limit:
            append_unit(unit_paragraphs)
            continue

        if projected_size <= max_chars and len(unit_text) <= max(500, max_chars // 4) and current_size < int(max_chars * 0.9):
            append_unit(unit_paragraphs)
            continue

        flush_current()
        append_unit(unit_paragraphs)

    flush_current()
    return blocks


def build_paragraph_relations(
    paragraphs: list[ParagraphUnit],
    *,
    enabled_relation_kinds: tuple[str, ...] | list[str] | set[str] | None = None,
) -> tuple[list[ParagraphRelation], RelationNormalizationReport]:
    relations: list[ParagraphRelation] = []
    decisions: list[ParagraphRelationDecision] = []
    relation_counts: dict[str, int] = {}
    rejected_candidate_count = 0
    next_relation_id = 1
    enabled_kinds = set(enabled_relation_kinds or RELATION_NORMALIZATION_KIND_VALUES)

    def append_relation(
        *,
        relation_kind: str,
        member_paragraph_ids: tuple[str, ...],
        anchor_asset_id: str | None = None,
        rationale: tuple[str, ...] = (),
    ) -> None:
        nonlocal next_relation_id
        relation_id = f"rel_{next_relation_id:04d}"
        next_relation_id += 1
        relations.append(
            ParagraphRelation(
                relation_id=relation_id,
                relation_kind=relation_kind,
                member_paragraph_ids=member_paragraph_ids,
                anchor_asset_id=anchor_asset_id,
                confidence="high",
                rationale=rationale,
            )
        )
        relation_counts[relation_kind] = relation_counts.get(relation_kind, 0) + 1
        decisions.append(
            ParagraphRelationDecision(
                relation_kind=relation_kind,
                decision="accept",
                member_paragraph_ids=member_paragraph_ids,
                anchor_asset_id=anchor_asset_id,
                reasons=rationale,
            )
        )

    def append_rejection(
        *,
        relation_kind: str,
        member_paragraph_ids: tuple[str, ...],
        reasons: tuple[str, ...],
        anchor_asset_id: str | None = None,
    ) -> None:
        nonlocal rejected_candidate_count
        rejected_candidate_count += 1
        decisions.append(
            ParagraphRelationDecision(
                relation_kind=relation_kind,
                decision="reject",
                member_paragraph_ids=member_paragraph_ids,
                anchor_asset_id=anchor_asset_id,
                reasons=reasons,
            )
        )

    for index, paragraph in enumerate(paragraphs):
        paragraph_role = getattr(paragraph, "role", None)
        paragraph_id = getattr(paragraph, "paragraph_id", None)
        is_caption_candidate = paragraph_role == "caption"
        if not is_caption_candidate and index > 0:
            previous_paragraph = paragraphs[index - 1]
            if previous_paragraph.role in {"image", "table"} and _is_likely_caption_candidate_for_relation(paragraph):
                is_caption_candidate = True
        if not is_caption_candidate:
            continue
        if index == 0:
            append_rejection(
                relation_kind="caption_attachment",
                member_paragraph_ids=((paragraph_id or f"p{index:04d}"),),
                reasons=("caption_without_preceding_asset",),
            )
            continue
        previous_paragraph = paragraphs[index - 1]
        previous_role = getattr(previous_paragraph, "role", None)
        previous_paragraph_id = getattr(previous_paragraph, "paragraph_id", None)
        relation_kind = f"{previous_role}_caption" if previous_role in {"image", "table"} else "caption_attachment"
        if previous_role not in {"image", "table"}:
            append_rejection(
                relation_kind="caption_attachment",
                member_paragraph_ids=((paragraph_id or f"p{index:04d}"),),
                reasons=("caption_not_adjacent_to_asset",),
            )
            continue
        if relation_kind not in enabled_kinds:
            continue
        if not previous_paragraph_id or not paragraph_id or getattr(previous_paragraph, "asset_id", None) is None:
            append_rejection(
                relation_kind=relation_kind,
                member_paragraph_ids=tuple(
                    paragraph_key
                    for paragraph_key in (previous_paragraph_id, paragraph_id)
                    if paragraph_key
                ) or ((paragraph_id or f"p{index:04d}"),),
                reasons=("missing_caption_anchor_identity",),
                anchor_asset_id=getattr(previous_paragraph, "asset_id", None),
            )
            continue
        append_relation(
            relation_kind=relation_kind,
            member_paragraph_ids=(previous_paragraph_id, paragraph_id),
            anchor_asset_id=getattr(previous_paragraph, "asset_id", None),
            rationale=("adjacent_asset_caption",),
        )

    if "epigraph_attribution" in enabled_kinds:
        for index in range(len(paragraphs) - 1):
            left = paragraphs[index]
            right = paragraphs[index + 1]
            left_paragraph_id = getattr(left, "paragraph_id", None)
            right_paragraph_id = getattr(right, "paragraph_id", None)
            if not left_paragraph_id or not right_paragraph_id:
                continue
            if not _is_epigraph_relation_candidate(left, right):
                rejection_reasons = _epigraph_relation_rejection_reasons(left, right)
                if rejection_reasons:
                    append_rejection(
                        relation_kind="epigraph_attribution",
                        member_paragraph_ids=(left_paragraph_id, right_paragraph_id),
                        reasons=rejection_reasons,
                    )
                continue
            append_relation(
                relation_kind="epigraph_attribution",
                member_paragraph_ids=(left_paragraph_id, right_paragraph_id),
                rationale=("adjacent_epigraph_attribution",),
            )

    index = 0
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        if _is_toc_header_paragraph(paragraph):
            member_indexes = [index]
            look_ahead = index + 1
            while look_ahead < len(paragraphs) and _is_toc_entry_paragraph(paragraphs[look_ahead]):
                member_indexes.append(look_ahead)
                look_ahead += 1
            if len(member_indexes) >= 2:
                if "toc_region" in enabled_kinds:
                    append_relation(
                        relation_kind="toc_region",
                        member_paragraph_ids=tuple(
                            paragraphs[member_index].paragraph_id or f"p{member_index:04d}" for member_index in member_indexes
                        ),
                        rationale=("toc_header_with_entries",),
                    )
                index = look_ahead
                continue
            append_rejection(
                relation_kind="toc_region",
                member_paragraph_ids=((paragraph.paragraph_id or f"p{index:04d}"),),
                reasons=("toc_header_without_entries",),
            )

        if _is_toc_entry_paragraph(paragraph):
            member_indexes = [index]
            look_ahead = index + 1
            while look_ahead < len(paragraphs) and _is_toc_entry_paragraph(paragraphs[look_ahead]):
                member_indexes.append(look_ahead)
                look_ahead += 1
            if len(member_indexes) >= 2:
                if "toc_region" in enabled_kinds:
                    append_relation(
                        relation_kind="toc_region",
                        member_paragraph_ids=tuple(
                            paragraphs[member_index].paragraph_id or f"p{member_index:04d}" for member_index in member_indexes
                        ),
                        rationale=("contiguous_toc_entries",),
                    )
                index = look_ahead
                continue
            append_rejection(
                relation_kind="toc_region",
                member_paragraph_ids=((paragraph.paragraph_id or f"p{index:04d}"),),
                reasons=("isolated_toc_entry",),
            )
        index += 1

    report = RelationNormalizationReport(
        total_relations=len(relations),
        relation_counts=relation_counts,
        rejected_candidate_count=rejected_candidate_count,
        decisions=decisions,
    )
    return relations, report


def _is_likely_caption_candidate_for_relation(paragraph: ParagraphUnit) -> bool:
    if paragraph.role == "heading" and paragraph.heading_source != "heuristic":
        return False
    return _is_likely_caption_text(paragraph.text)


def _apply_relation_side_effects(paragraphs: list[ParagraphUnit], relations: list[ParagraphRelation]) -> None:
    paragraph_by_id = {paragraph.paragraph_id: paragraph for paragraph in paragraphs if paragraph.paragraph_id}
    for paragraph in paragraphs:
        if paragraph.role == "caption":
            paragraph.attached_to_asset_id = None

    for relation in relations:
        if relation.relation_kind not in {"image_caption", "table_caption"}:
            continue
        if len(relation.member_paragraph_ids) < 2:
            continue
        caption_paragraph = paragraph_by_id.get(relation.member_paragraph_ids[-1])
        if caption_paragraph is not None:
            caption_paragraph.attached_to_asset_id = relation.anchor_asset_id


def _build_semantic_block_units(
    paragraphs: list[ParagraphUnit],
    relations: list[ParagraphRelation],
) -> list[list[ParagraphUnit]]:
    index_by_paragraph_id = {
        paragraph.paragraph_id: index for index, paragraph in enumerate(paragraphs) if paragraph.paragraph_id
    }
    parent = list(range(len(paragraphs)))

    def find(index: int) -> int:
        while parent[index] != index:
            parent[index] = parent[parent[index]]
            index = parent[index]
        return index

    def union(left_index: int, right_index: int) -> None:
        left_root = find(left_index)
        right_root = find(right_index)
        if left_root != right_root:
            parent[right_root] = left_root

    for relation in relations:
        member_indexes = [index_by_paragraph_id[paragraph_id] for paragraph_id in relation.member_paragraph_ids if paragraph_id in index_by_paragraph_id]
        if len(member_indexes) < 2:
            continue
        for member_index in member_indexes[1:]:
            union(member_indexes[0], member_index)

    for index in range(len(paragraphs) - 1):
        left = paragraphs[index]
        right = paragraphs[index + 1]
        if _is_quote_structural_role(left) and _is_quote_structural_role(right):
            union(index, index + 1)
            continue
        if _is_toc_structural_role(left) and _is_toc_structural_role(right):
            union(index, index + 1)

    grouped_indexes: dict[int, list[int]] = {}
    for index in range(len(paragraphs)):
        grouped_indexes.setdefault(find(index), []).append(index)

    clusters = sorted((sorted(indexes) for indexes in grouped_indexes.values()), key=lambda indexes: indexes[0])
    units: list[list[ParagraphUnit]] = []
    for indexes in clusters:
        if indexes != list(range(indexes[0], indexes[-1] + 1)):
            for index in indexes:
                units.append([paragraphs[index]])
            continue
        units.append([paragraphs[index] for index in indexes])
    return units


def _is_epigraph_relation_candidate(left: ParagraphUnit, right: ParagraphUnit) -> bool:
    left_role = getattr(left, "role", None)
    right_role = getattr(right, "role", None)
    if left_role in {"image", "table", "caption", "list"} or right_role in {"image", "table", "caption", "list"}:
        return False
    left_structural = str(getattr(left, "structural_role", None) or left_role or "").strip().lower()
    right_structural = str(getattr(right, "structural_role", None) or right_role or "").strip().lower()
    if left_structural == "epigraph" and right_structural == "attribution":
        return True
    if left_structural == "epigraph" and _is_likely_attribution_text(str(getattr(right, "text", ""))):
        return True
    if right_structural == "attribution" and getattr(left, "paragraph_alignment", None) == "center":
        return True
    return False


def _epigraph_relation_rejection_reasons(left: ParagraphUnit, right: ParagraphUnit) -> tuple[str, ...]:
    left_role = getattr(left, "role", None)
    right_role = getattr(right, "role", None)
    left_structural = str(getattr(left, "structural_role", None) or left_role or "").strip().lower()
    right_structural = str(getattr(right, "structural_role", None) or right_role or "").strip().lower()
    right_text = str(getattr(right, "text", ""))
    reasons: list[str] = []

    if left_structural == "epigraph" and right_structural != "attribution":
        reasons.append("epigraph_without_attribution")
    elif right_structural == "attribution" and left_structural != "epigraph":
        reasons.append("attribution_without_epigraph")
    elif left_structural == "epigraph" and _is_likely_attribution_text(right_text):
        reasons.append("epigraph_candidate_rejected")
    elif right_structural == "attribution" and getattr(left, "paragraph_alignment", None) != "center":
        reasons.append("attribution_alignment_mismatch")

    return tuple(reasons)


def _is_toc_header_paragraph(paragraph: ParagraphUnit) -> bool:
    structural = str(getattr(paragraph, "structural_role", None) or getattr(paragraph, "role", None) or "").strip().lower()
    if structural == "toc_header":
        return True
    return str(getattr(paragraph, "text", "")).strip().lower() in {"содержание", "contents"}


def _is_toc_entry_paragraph(paragraph: ParagraphUnit) -> bool:
    structural = str(getattr(paragraph, "structural_role", None) or getattr(paragraph, "role", None) or "").strip().lower()
    if structural == "toc_entry":
        return True
    return _is_likely_toc_entry_text(str(getattr(paragraph, "text", "")))


def build_context_excerpt(blocks: list[DocumentBlock], block_index: int, limit_chars: int, *, reverse: bool) -> str:
    if limit_chars <= 0:
        return ""

    indexes = range(block_index - 1, -1, -1) if reverse else range(block_index + 1, len(blocks))
    collected: list[str] = []
    total_size = 0

    for index in indexes:
        block_text = blocks[index].text.strip()
        if not block_text:
            continue

        separator_size = 2 if collected else 0
        projected_size = total_size + separator_size + len(block_text)
        if projected_size <= limit_chars:
            collected.append(block_text)
            total_size = projected_size
            continue

        remaining = limit_chars - total_size - separator_size
        if remaining > 0:
            excerpt = block_text[-remaining:] if reverse else block_text[:remaining]
            if excerpt.strip():
                collected.append(excerpt.strip())
        break

    if reverse:
        collected.reverse()

    return "\n\n".join(collected).strip()


def build_editing_jobs(blocks: list[DocumentBlock], max_chars: int) -> list[dict[str, object]]:
    context_before_chars = max(600, min(1400, int(max_chars * 0.2)))
    context_after_chars = max(300, min(800, int(max_chars * 0.12)))
    jobs: list[dict[str, object]] = []
    fallback_paragraph_index = 0

    for index, block in enumerate(blocks):
        context_before = build_context_excerpt(blocks, index, context_before_chars, reverse=True)
        context_after = build_context_excerpt(blocks, index, context_after_chars, reverse=False)
        job_kind = (
            "passthrough"
            if block.paragraphs
            and (
                all(paragraph.role == "image" for paragraph in block.paragraphs)
                or all(_is_toc_structural_role(paragraph) for paragraph in block.paragraphs)
            )
            else "llm"
        )
        paragraph_ids = [
            _resolve_marker_paragraph_id(paragraph, fallback_paragraph_index + paragraph_index)
            for paragraph_index, paragraph in enumerate(block.paragraphs)
        ]
        jobs.append(
            {
                "job_kind": job_kind,
                "target_text": block.text,
                "target_text_with_markers": build_marker_wrapped_block_text(block.paragraphs, paragraph_ids=paragraph_ids),
                "paragraph_ids": paragraph_ids,
                "context_before": context_before,
                "context_after": context_after,
                "target_chars": len(block.text),
                "context_chars": len(context_before) + len(context_after),
            }
        )
        fallback_paragraph_index += len(block.paragraphs)

    return jobs


def _build_paragraph_text_with_placeholders(paragraph, image_assets: list[ImageAsset]) -> str:
    parts: list[str] = []
    for child in paragraph._element:
        local_name = _xml_local_name(child.tag)
        if local_name == "r":
            parts.append(_render_run_element(child, paragraph.part, image_assets))
            continue
        if local_name == "hyperlink":
            parts.append(_render_hyperlink_element(child, paragraph, image_assets))
    return "".join(parts)


def _build_raw_document_blocks(document) -> tuple[list[RawBlock], list[ImageAsset]]:
    raw_blocks: list[RawBlock] = []
    image_assets: list[ImageAsset] = []
    table_count = 0

    for block_kind, block in _iter_document_block_items(document):
        raw_index = len(raw_blocks)
        if block_kind == "paragraph":
            raw_block = _build_raw_paragraph(cast(Paragraph, block), image_assets, raw_index=raw_index)
        else:
            table_count += 1
            raw_block = _build_raw_table(
                cast(Table, block),
                image_assets,
                raw_index=raw_index,
                asset_id=f"table_{table_count:03d}",
            )
        if raw_block is not None:
            raw_blocks.append(raw_block)

    return raw_blocks, image_assets


def _build_raw_paragraph(paragraph, image_assets: list[ImageAsset], *, raw_index: int) -> RawParagraph | None:
    text = _build_paragraph_text_with_placeholders(paragraph, image_assets).strip()
    if not text:
        return None

    style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
    normalized_style = style_name.strip().lower()
    explicit_heading_level = _extract_explicit_heading_level(paragraph, style_name)
    heading_level = explicit_heading_level
    heading_source = "explicit" if explicit_heading_level is not None else None
    if heading_level is None and not _is_caption_style(normalized_style):
        if _is_probable_heading(paragraph, text, normalized_style):
            heading_level = _infer_heuristic_heading_level(text)
            heading_source = "heuristic"
    role = classify_paragraph_role(text, style_name, heading_level=heading_level)
    list_metadata = _extract_paragraph_list_metadata(paragraph, text, style_name, role)
    if role != "list" and list_metadata["list_kind"] is not None:
        role = "list"
    if role == "list" and list_metadata.get("_is_typographic_emdash_bullet"):
        role = "body"
        list_metadata = _empty_list_metadata()
    asset_id = _extract_paragraph_asset_id(text, role=role)
    role_confidence = _infer_role_confidence(
        role=role,
        text=text,
        normalized_style=normalized_style,
        explicit_heading_level=explicit_heading_level,
        heading_source=heading_source,
    )
    return RawParagraph(
        raw_index=raw_index,
        text=text,
        style_name=style_name,
        paragraph_alignment=_resolve_paragraph_alignment(paragraph),
        is_bold=_paragraph_is_effectively_bold(paragraph),
        is_italic=_paragraph_is_effectively_italic(paragraph),
        font_size_pt=_resolve_effective_paragraph_font_size(paragraph),
        explicit_heading_level=explicit_heading_level,
        heading_level=heading_level,
        heading_source=heading_source,
        list_kind=cast(str | None, list_metadata["list_kind"]),
        list_level=cast(int, list_metadata["list_level"]),
        list_numbering_format=cast(str | None, list_metadata["list_numbering_format"]),
        list_num_id=cast(str | None, list_metadata["list_num_id"]),
        list_abstract_num_id=cast(str | None, list_metadata["list_abstract_num_id"]),
        list_num_xml=cast(str | None, list_metadata["list_num_xml"]),
        list_abstract_num_xml=cast(str | None, list_metadata["list_abstract_num_xml"]),
        role_hint=role,
        source_xml_fingerprint=_build_source_xml_fingerprint(paragraph),
        origin_raw_indexes=(raw_index,),
        origin_raw_texts=(text,),
        boundary_source="raw",
        boundary_confidence="explicit" if role_confidence == "explicit" else "high",
    )


def _build_raw_table(table: Table, image_assets: list[ImageAsset], *, raw_index: int, asset_id: str) -> RawTable | None:
    html_table = _render_table_html(table, image_assets)
    if not html_table.strip():
        return None
    return RawTable(raw_index=raw_index, html_text=html_table, asset_id=asset_id)


def _build_logical_paragraph_units(raw_blocks: list[RawBlock]) -> list[ParagraphUnit]:
    paragraphs: list[ParagraphUnit] = []
    for block in raw_blocks:
        if isinstance(block, RawParagraph):
            paragraph = ParagraphUnit(
                text=block.text,
                role=block.role_hint,
                asset_id=_extract_paragraph_asset_id(block.text, role=block.role_hint),
                paragraph_alignment=block.paragraph_alignment,
                heading_level=block.heading_level,
                heading_source=block.heading_source,
                list_kind=block.list_kind,
                list_level=block.list_level,
                list_numbering_format=block.list_numbering_format,
                list_num_id=block.list_num_id,
                list_abstract_num_id=block.list_abstract_num_id,
                list_num_xml=block.list_num_xml,
                list_abstract_num_xml=block.list_abstract_num_xml,
                structural_role=block.role_hint,
                role_confidence=_infer_role_confidence(
                    role=block.role_hint,
                    text=block.text,
                    normalized_style=block.style_name.strip().lower(),
                    explicit_heading_level=block.explicit_heading_level,
                    heading_source=block.heading_source,
                ),
                style_name=block.style_name,
                is_bold=block.is_bold,
                is_italic=block.is_italic,
                font_size_pt=block.font_size_pt,
                origin_raw_indexes=list(block.origin_raw_indexes or (block.raw_index,)),
                origin_raw_texts=list(block.origin_raw_texts or (block.text,)),
                boundary_source=block.boundary_source,
                boundary_confidence=block.boundary_confidence,
                boundary_rationale=block.boundary_rationale,
            )
        else:
            paragraph = ParagraphUnit(
                text=block.html_text,
                role="table",
                asset_id=block.asset_id,
                structural_role="table",
                role_confidence="explicit",
                origin_raw_indexes=[block.raw_index],
                origin_raw_texts=[block.html_text],
            )
        _assign_paragraph_identity(paragraph, len(paragraphs))
        paragraphs.append(paragraph)
    return paragraphs


def _resolve_paragraph_boundary_normalization_settings() -> tuple[str, bool]:
    from config import load_app_config

    app_config = load_app_config()
    enabled = bool(app_config.get("paragraph_boundary_normalization_enabled", True))
    mode = str(app_config.get("paragraph_boundary_normalization_mode", "high_only"))
    if mode not in PARAGRAPH_BOUNDARY_NORMALIZATION_MODE_VALUES:
        mode = "high_only"
    if not enabled:
        mode = "off"
    return mode, bool(app_config.get("paragraph_boundary_normalization_save_debug_artifacts", True))


def _resolve_relation_normalization_settings() -> tuple[bool, str, tuple[str, ...], bool]:
    from config import load_app_config

    app_config = load_app_config()
    enabled = bool(app_config.get("relation_normalization_enabled", True))
    profile = str(app_config.get("relation_normalization_profile", "phase2_default") or "phase2_default")
    configured_relation_kinds = app_config.get(
        "relation_normalization_enabled_relation_kinds",
        RELATION_NORMALIZATION_KIND_VALUES,
    )
    if not isinstance(configured_relation_kinds, (list, tuple, set)):
        configured_relation_kinds = RELATION_NORMALIZATION_KIND_VALUES
    enabled_relation_kinds = tuple(
        kind
        for kind in configured_relation_kinds
        if kind in RELATION_NORMALIZATION_KIND_VALUES
    )
    if not enabled:
        enabled_relation_kinds = ()
    return (
        enabled,
        profile,
        enabled_relation_kinds,
        bool(app_config.get("relation_normalization_save_debug_artifacts", True)),
    )


def _resolve_paragraph_boundary_ai_review_settings() -> tuple[bool, str, int, int, int, str]:
    from config import load_app_config

    app_config = load_app_config()
    enabled = bool(app_config.get("paragraph_boundary_ai_review_enabled", False))
    mode = str(app_config.get("paragraph_boundary_ai_review_mode", "off") or "off")
    if mode not in PARAGRAPH_BOUNDARY_AI_REVIEW_MODE_VALUES:
        mode = "off"
    if not enabled:
        mode = "off"
    return (
        enabled and mode != "off",
        mode,
        _coerce_int_config_value(app_config.get("paragraph_boundary_ai_review_candidate_limit"), 200),
        _coerce_int_config_value(app_config.get("paragraph_boundary_ai_review_timeout_seconds"), 30),
        _coerce_int_config_value(app_config.get("paragraph_boundary_ai_review_max_tokens_per_candidate"), 120),
        str(app_config.get("default_model", "gpt-5-mini") or "gpt-5-mini"),
    )


def _coerce_int_config_value(value: object, default: int) -> int:
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return int(value)
    if isinstance(value, str) and value.strip():
        try:
            return int(value.strip())
        except ValueError:
            return default
    return default


def _build_ai_review_candidates(
    *,
    raw_blocks: list[RawBlock],
    paragraphs: list[ParagraphUnit],
    boundary_report: ParagraphBoundaryNormalizationReport,
    relation_report: RelationNormalizationReport,
    candidate_limit: int,
) -> list[dict[str, object]]:
    raw_paragraphs_by_index = {
        block.raw_index: block
        for block in raw_blocks
        if isinstance(block, RawParagraph)
    }
    paragraph_text_by_id = {
        paragraph.paragraph_id: paragraph.text
        for paragraph in paragraphs
        if paragraph.paragraph_id
    }

    candidates: list[dict[str, object]] = []
    for decision in boundary_report.decisions:
        if decision.confidence != "medium":
            continue
        left = raw_paragraphs_by_index.get(decision.left_raw_index)
        right = raw_paragraphs_by_index.get(decision.right_raw_index)
        candidates.append(
            {
                "candidate_kind": "boundary_medium",
                "candidate_id": f"{decision.left_raw_index}:{decision.right_raw_index}",
                "deterministic_decision": decision.decision,
                "confidence": decision.confidence,
                "left_raw_index": decision.left_raw_index,
                "right_raw_index": decision.right_raw_index,
                "left_text": None if left is None else left.text,
                "right_text": None if right is None else right.text,
                "reasons": list(decision.reasons),
            }
        )

    for decision in relation_report.decisions:
        if decision.decision == "accept":
            continue
        candidates.append(
            {
                "candidate_kind": "relation_rejected",
                "candidate_id": f"{decision.relation_kind}:{'|'.join(decision.member_paragraph_ids)}",
                "deterministic_decision": decision.decision,
                "relation_kind": decision.relation_kind,
                "member_paragraph_ids": list(decision.member_paragraph_ids),
                "member_texts": [paragraph_text_by_id.get(paragraph_id, "") for paragraph_id in decision.member_paragraph_ids],
                "anchor_asset_id": decision.anchor_asset_id,
                "reasons": list(decision.reasons),
            }
        )

    return candidates[: max(0, candidate_limit)]


def _build_ai_review_request_payload(
    *,
    model: str,
    candidates: list[dict[str, object]],
    timeout_seconds: int,
    max_tokens_per_candidate: int,
) -> dict[str, object]:
    system_prompt = (
        "You review ambiguous DOCX paragraph-boundary and grouping candidates. "
        "Return only JSON with a top-level recommendations array. "
        "For boundary candidates use recommendation merge or keep. "
        "For relation candidates use recommendation accept or reject."
    )
    user_prompt = json.dumps(
        {
            "instructions": {
                "review_scope": "ambiguous normalization candidates",
                "required_output_shape": {
                    "recommendations": [
                        {
                            "candidate_id": "string",
                            "recommendation": "merge|keep|accept|reject",
                            "reasons": ["string"],
                        }
                    ]
                },
            },
            "candidates": candidates,
        },
        ensure_ascii=False,
    )
    return {
        "model": model,
        "input": [
            {
                "role": "system",
                "content": [{"type": "input_text", "text": system_prompt}],
            },
            {
                "role": "user",
                "content": [{"type": "input_text", "text": user_prompt}],
            },
        ],
        "max_output_tokens": max(256, min(len(candidates) * max_tokens_per_candidate, 8192)),
        "timeout": timeout_seconds,
    }


def _request_ai_review_recommendations(
    *,
    model: str,
    candidates: list[dict[str, object]],
    timeout_seconds: int,
    max_tokens_per_candidate: int,
) -> dict[str, dict[str, object]]:
    from config import get_client

    response = call_responses_create_with_retry(
        get_client(),
        _build_ai_review_request_payload(
            model=model,
            candidates=candidates,
            timeout_seconds=timeout_seconds,
            max_tokens_per_candidate=max_tokens_per_candidate,
        ),
        max_retries=2,
        retryable_error_predicate=is_retryable_error,
    )
    raw_text = extract_response_text(
        response,
        empty_message="AI review did not return text.",
        incomplete_message="AI review returned incomplete response.",
        unsupported_message="AI review returned unsupported response shape.",
    )
    payload = parse_json_object(
        raw_text,
        empty_message="AI review returned empty output.",
        no_json_message="AI review did not return JSON.",
    )
    recommendations = payload.get("recommendations", [])
    if not isinstance(recommendations, list):
        raise RuntimeError("AI review returned invalid recommendations payload.")

    result: dict[str, dict[str, object]] = {}
    for entry in recommendations:
        if not isinstance(entry, dict):
            continue
        candidate_id = entry.get("candidate_id")
        recommendation = entry.get("recommendation")
        reasons = entry.get("reasons", [])
        if not isinstance(candidate_id, str) or not candidate_id.strip():
            continue
        if not isinstance(recommendation, str) or not recommendation.strip():
            continue
        if not isinstance(reasons, list):
            reasons = []
        result[candidate_id] = {
            "recommendation": recommendation.strip().lower(),
            "reasons": [str(reason) for reason in reasons if str(reason).strip()],
        }
    return result


def _build_ai_review_decision_records(
    *,
    candidates: list[dict[str, object]],
    recommendations: dict[str, dict[str, object]],
    mode: str,
    error_code: str | None,
) -> list[dict[str, object]]:
    decisions: list[dict[str, object]] = []
    for candidate in candidates:
        candidate_id = str(candidate.get("candidate_id") or "")
        deterministic_decision = str(candidate.get("deterministic_decision") or "keep")
        recommendation_payload = recommendations.get(candidate_id, {})
        ai_recommendation = recommendation_payload.get("recommendation")
        reasons = [f"{mode}_mode"]
        if error_code is not None:
            reasons.append(f"ai_review_unavailable:{error_code}")
        elif ai_recommendation is None:
            reasons.append("ai_review_no_recommendation")
        elif ai_recommendation != deterministic_decision:
            reasons.append("deterministic_decision_retained")
        else:
            reasons.append("ai_agreed_with_deterministic")

        decisions.append(
            {
                "candidate_kind": candidate.get("candidate_kind"),
                "candidate_id": candidate_id,
                "deterministic_decision": deterministic_decision,
                "ai_recommendation": ai_recommendation,
                "final_decision": deterministic_decision,
                "reasons": reasons,
            }
        )
    return decisions


def _write_paragraph_boundary_ai_review_artifact(
    *,
    source_name: str,
    source_bytes: bytes,
    mode: str,
    decisions: list[dict[str, object]],
    error_code: str | None = None,
) -> str | None:
    try:
        PARAGRAPH_BOUNDARY_AI_REVIEW_DIR.mkdir(parents=True, exist_ok=True)
        source_hash = hashlib.sha1(source_bytes).hexdigest()[:8]
        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", source_name or "document.docx").strip("_") or "document.docx"
        artifact_path = PARAGRAPH_BOUNDARY_AI_REVIEW_DIR / f"{safe_name}_{source_hash}.json"
        payload: dict[str, object] = {
            "version": 1,
            "source_file": source_name,
            "source_hash": source_hash,
            "mode": mode,
            "reviewed_candidate_count": len(decisions),
            "accepted_candidate_count": sum(
                1 for decision in decisions if decision.get("final_decision") in {"merge", "accept", "group"}
            ),
            "decisions": decisions,
        }
        if error_code is not None:
            payload["error_code"] = error_code
        artifact_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(artifact_path)
    except Exception:
        return None


def _run_paragraph_boundary_ai_review(
    *,
    source_name: str,
    source_bytes: bytes,
    mode: str,
    model: str,
    raw_blocks: list[RawBlock],
    paragraphs: list[ParagraphUnit],
    boundary_report: ParagraphBoundaryNormalizationReport,
    relation_report: RelationNormalizationReport,
    candidate_limit: int,
    timeout_seconds: int,
    max_tokens_per_candidate: int,
) -> str | None:
    candidates = _build_ai_review_candidates(
        raw_blocks=raw_blocks,
        paragraphs=paragraphs,
        boundary_report=boundary_report,
        relation_report=relation_report,
        candidate_limit=candidate_limit,
    )
    if not candidates:
        return None

    recommendations: dict[str, dict[str, object]] = {}
    error_code: str | None = None
    try:
        recommendations = _request_ai_review_recommendations(
            model=model,
            candidates=candidates,
            timeout_seconds=timeout_seconds,
            max_tokens_per_candidate=max_tokens_per_candidate,
        )
    except Exception as exc:
        error_code = extract_model_response_error_code(exc)
        if error_code is None and "timeout" in str(exc).lower():
            error_code = "timeout"
        if error_code is None:
            error_code = "review_failed"

    decisions = _build_ai_review_decision_records(
        candidates=candidates,
        recommendations=recommendations,
        mode=mode,
        error_code=error_code,
    )
    return _write_paragraph_boundary_ai_review_artifact(
        source_name=source_name,
        source_bytes=source_bytes,
        mode=mode,
        decisions=decisions,
        error_code=error_code,
    )


def resolve_effective_relation_kinds() -> tuple[str, ...]:
    enabled, _, enabled_relation_kinds, _ = _resolve_relation_normalization_settings()
    if not enabled:
        return ()
    return enabled_relation_kinds


def summarize_boundary_normalization_metrics(
    report: ParagraphBoundaryNormalizationReport | None,
) -> dict[str, int]:
    if report is None:
        return {}
    high_confidence_merge_count = 0
    medium_accepted_merge_count = 0
    medium_rejected_candidate_count = 0
    decisions = getattr(report, "decisions", ()) or ()
    for decision in decisions:
        if decision.decision == "merge" and decision.confidence == "high":
            high_confidence_merge_count += 1
        elif decision.decision == "merge" and decision.confidence == "medium":
            medium_accepted_merge_count += 1
        elif decision.decision == "keep" and decision.confidence == "medium":
            medium_rejected_candidate_count += 1
    return {
        "high_confidence_merge_count": high_confidence_merge_count,
        "medium_accepted_merge_count": medium_accepted_merge_count,
        "medium_rejected_candidate_count": medium_rejected_candidate_count,
    }


def _normalize_paragraph_boundaries(
    raw_blocks: list[RawBlock],
    *,
    mode: str,
) -> tuple[list[RawBlock], ParagraphBoundaryNormalizationReport]:
    total_raw_paragraphs = sum(1 for block in raw_blocks if isinstance(block, RawParagraph))
    if mode == "off":
        report = ParagraphBoundaryNormalizationReport(
            total_raw_paragraphs=total_raw_paragraphs,
            total_logical_paragraphs=total_raw_paragraphs,
            merged_group_count=0,
            merged_raw_paragraph_count=0,
            decisions=[],
        )
        return list(raw_blocks), report

    normalized_blocks: list[RawBlock] = []
    decisions: list[ParagraphBoundaryDecision] = []
    merged_group_count = 0
    merged_raw_paragraph_count = 0
    index = 0

    while index < len(raw_blocks):
        block = raw_blocks[index]
        if not isinstance(block, RawParagraph):
            normalized_blocks.append(block)
            index += 1
            continue

        group = [block]
        group_reasons: list[str] = []
        group_confidences: list[str] = []
        look_ahead = index
        while look_ahead + 1 < len(raw_blocks) and isinstance(raw_blocks[look_ahead + 1], RawParagraph):
            next_block = cast(RawParagraph, raw_blocks[look_ahead + 1])
            decision = _evaluate_paragraph_boundary(group[-1], next_block)
            effective_decision = decision
            if decision.decision == "merge" and decision.confidence == "medium" and mode != "high_and_medium":
                effective_decision = ParagraphBoundaryDecision(
                    left_raw_index=decision.left_raw_index,
                    right_raw_index=decision.right_raw_index,
                    decision="keep",
                    confidence="medium",
                    reasons=tuple((*decision.reasons, "medium_mode_disabled")),
                )
            decisions.append(effective_decision)
            if effective_decision.decision != "merge":
                break
            group.append(next_block)
            group_reasons.extend(effective_decision.reasons)
            group_confidences.append(effective_decision.confidence)
            look_ahead += 1

        if len(group) == 1:
            normalized_blocks.append(block)
            index += 1
            continue

        merged_group_count += 1
        merged_raw_paragraph_count += len(group)
        normalized_blocks.append(_merge_raw_paragraph_group(group, group_reasons, group_confidences))
        index += len(group)

    report = ParagraphBoundaryNormalizationReport(
        total_raw_paragraphs=total_raw_paragraphs,
        total_logical_paragraphs=sum(1 for block in normalized_blocks if isinstance(block, RawParagraph)),
        merged_group_count=merged_group_count,
        merged_raw_paragraph_count=merged_raw_paragraph_count,
        decisions=decisions,
    )
    return normalized_blocks, report


def _evaluate_paragraph_boundary(left: RawParagraph, right: RawParagraph) -> ParagraphBoundaryDecision:
    blocked_reasons: list[str] = []
    positive_reasons: list[str] = []

    if left.heading_level is not None or right.heading_level is not None:
        blocked_reasons.append("heading_boundary")
    if left.role_hint != "body" or right.role_hint != "body":
        blocked_reasons.append("non_body_role")
    if left.list_kind is not None or right.list_kind is not None:
        blocked_reasons.append("list_metadata")
    if _detect_explicit_list_kind(right.text) is not None:
        blocked_reasons.append("right_explicit_list_marker")
    if _is_likely_caption_text(left.text) or _is_likely_caption_text(right.text):
        blocked_reasons.append("caption_like_boundary")
    if _is_likely_attribution_text(right.text):
        blocked_reasons.append("right_attribution_like")
    if _is_likely_toc_entry_text(right.text):
        blocked_reasons.append("right_toc_like")
    if _style_transition_implies_structure(left, right):
        blocked_reasons.append("style_transition")
    if _alignment_transition_implies_structure(left, right):
        blocked_reasons.append("alignment_transition")
    if _ends_with_strong_paragraph_terminator(left.text) and _starts_with_new_sentence_signal(right.text):
        blocked_reasons.append("terminal_punctuation_sentence_reset")

    if blocked_reasons:
        return ParagraphBoundaryDecision(
            left_raw_index=left.raw_index,
            right_raw_index=right.raw_index,
            decision="keep",
            confidence="blocked",
            reasons=tuple(blocked_reasons),
        )

    if _styles_are_compatible(left, right):
        positive_reasons.append("same_body_style")
    if _alignments_are_compatible(left, right):
        positive_reasons.append("compatible_alignment")
    if not _ends_with_strong_paragraph_terminator(left.text):
        positive_reasons.append("left_not_terminal")
    if _starts_with_continuation_signal(right.text):
        positive_reasons.append("right_starts_continuation")
    if _left_paragraph_looks_incomplete(left.text):
        positive_reasons.append("left_incomplete")
    if _combined_text_reads_as_continuation(left.text, right.text):
        positive_reasons.append("combined_sentence_plausible")

    if {"same_body_style", "left_not_terminal", "right_starts_continuation"}.issubset(set(positive_reasons)):
        return ParagraphBoundaryDecision(
            left_raw_index=left.raw_index,
            right_raw_index=right.raw_index,
            decision="merge",
            confidence="high",
            reasons=tuple(positive_reasons),
        )

    if _should_promote_medium_merge(positive_reasons):
        return ParagraphBoundaryDecision(
            left_raw_index=left.raw_index,
            right_raw_index=right.raw_index,
            decision="merge",
            confidence="medium",
            reasons=tuple(positive_reasons),
        )

    return ParagraphBoundaryDecision(
        left_raw_index=left.raw_index,
        right_raw_index=right.raw_index,
        decision="keep",
        confidence="medium",
        reasons=tuple(positive_reasons or ("insufficient_merge_signals",)),
    )


def _merge_raw_paragraph_group(group: list[RawParagraph], reasons: list[str], confidences: list[str]) -> RawParagraph:
    dominant = group[0]
    merged_text = _join_merged_paragraph_text(group)
    merged_indexes = tuple(index for paragraph in group for index in paragraph.origin_raw_indexes)
    merged_texts = tuple(text for paragraph in group for text in paragraph.origin_raw_texts)
    rationale = ", ".join(dict.fromkeys(reasons)) or None
    boundary_confidence = "medium" if "medium" in confidences else "high"
    return RawParagraph(
        raw_index=dominant.raw_index,
        text=merged_text,
        style_name=dominant.style_name,
        paragraph_alignment=dominant.paragraph_alignment,
        is_bold=dominant.is_bold,
        is_italic=dominant.is_italic,
        font_size_pt=dominant.font_size_pt,
        explicit_heading_level=dominant.explicit_heading_level,
        heading_level=dominant.heading_level,
        heading_source=dominant.heading_source,
        list_kind=dominant.list_kind,
        list_level=dominant.list_level,
        list_numbering_format=dominant.list_numbering_format,
        list_num_id=dominant.list_num_id,
        list_abstract_num_id=dominant.list_abstract_num_id,
        list_num_xml=dominant.list_num_xml,
        list_abstract_num_xml=dominant.list_abstract_num_xml,
        role_hint=dominant.role_hint,
        source_xml_fingerprint=dominant.source_xml_fingerprint,
        origin_raw_indexes=merged_indexes,
        origin_raw_texts=merged_texts,
        boundary_source="normalized_merge",
        boundary_confidence=boundary_confidence,
        boundary_rationale=rationale,
    )


def _join_merged_paragraph_text(group: list[RawParagraph]) -> str:
    merged_text = " ".join(paragraph.text.strip() for paragraph in group if paragraph.text.strip())
    merged_text = re.sub(r"\s+([,.;:!?…])", r"\1", merged_text)
    merged_text = re.sub(r"\s+", " ", merged_text)
    return merged_text.strip()


def _styles_are_compatible(left: RawParagraph, right: RawParagraph) -> bool:
    left_style = left.style_name.strip().lower()
    right_style = right.style_name.strip().lower()
    if left_style == right_style:
        return True
    body_aliases = {"", "normal", "body text", "текст", "обычный"}
    return left_style in body_aliases and right_style in body_aliases


def _alignments_are_compatible(left: RawParagraph, right: RawParagraph) -> bool:
    compatible = {None, "left", "start", "both"}
    if left.paragraph_alignment == right.paragraph_alignment:
        return True
    return left.paragraph_alignment in compatible and right.paragraph_alignment in compatible


def _alignment_transition_implies_structure(left: RawParagraph, right: RawParagraph) -> bool:
    if _alignments_are_compatible(left, right):
        return False
    structured_alignments = {"center", "right", "end"}
    return left.paragraph_alignment in structured_alignments or right.paragraph_alignment in structured_alignments


def _style_transition_implies_structure(left: RawParagraph, right: RawParagraph) -> bool:
    for style_name in (left.style_name, right.style_name):
        normalized_style = style_name.strip().lower()
        if _is_caption_style(normalized_style):
            return True
        if HEADING_STYLE_PATTERN.match(normalized_style) is not None:
            return True
        if "list" in normalized_style or "спис" in normalized_style:
            return True
    return False


def _ends_with_strong_paragraph_terminator(text: str) -> bool:
    return STRONG_PARAGRAPH_TERMINATOR_PATTERN.search(text.strip()) is not None


def _starts_with_continuation_signal(text: str) -> bool:
    stripped = text.lstrip()
    if not stripped:
        return False
    for char in stripped:
        if char in {'"', "'", "«", "(", "["}:
            continue
        if char.islower() or char.isdigit():
            return True
        break
    first_word = stripped.split()[0].strip("\"'«»()[]").lower() if stripped.split() else ""
    return first_word in {"и", "а", "но", "или", "что", "как", "поэтому", "and", "but", "or", "that", "which"}


def _starts_with_new_sentence_signal(text: str) -> bool:
    stripped = text.lstrip()
    if not stripped:
        return False
    for char in stripped:
        if char in {'"', "'", "«", "(", "["}:
            continue
        if char.isupper():
            return True
        break
    return _has_heading_text_signal(stripped)


def _left_paragraph_looks_incomplete(text: str) -> bool:
    stripped = text.rstrip()
    if not stripped:
        return False
    if _ends_with_strong_paragraph_terminator(stripped):
        return False
    return stripped[-1].isalnum() or stripped.endswith((",", ";", ":", "-", "(", "["))


def _combined_text_reads_as_continuation(left_text: str, right_text: str) -> bool:
    if not left_text.strip() or not right_text.strip():
        return False
    return _left_paragraph_looks_incomplete(left_text) and _starts_with_continuation_signal(right_text)


def _should_promote_medium_merge(positive_reasons: list[str]) -> bool:
    positive_reason_set = set(positive_reasons)
    if not {"same_body_style", "compatible_alignment"}.issubset(positive_reason_set):
        return False
    supporting_signals = {
        "left_not_terminal",
        "left_incomplete",
        "right_starts_continuation",
        "combined_sentence_plausible",
    }
    return len(positive_reason_set & supporting_signals) >= 2


def _is_likely_attribution_text(text: str) -> bool:
    stripped = text.strip()
    return bool(stripped) and len(stripped) <= 120 and stripped.startswith(("-", "—", "–"))


def _is_likely_toc_entry_text(text: str) -> bool:
    return TOC_ENTRY_PATTERN.match(text.strip()) is not None


def _build_source_xml_fingerprint(paragraph) -> str | None:
    try:
        xml_text = etree.tostring(paragraph._element, encoding="utf-8")
    except Exception:
        return None
    return hashlib.sha1(xml_text).hexdigest()[:12]


def _write_paragraph_boundary_report_artifact(
    *,
    source_name: str,
    source_bytes: bytes,
    mode: str,
    report: ParagraphBoundaryNormalizationReport,
) -> str | None:
    try:
        PARAGRAPH_BOUNDARY_REPORTS_DIR.mkdir(parents=True, exist_ok=True)
        source_hash = hashlib.sha1(source_bytes).hexdigest()[:8]
        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", source_name or "document.docx").strip("_") or "document.docx"
        artifact_path = PARAGRAPH_BOUNDARY_REPORTS_DIR / f"{safe_name}_{source_hash}.json"
        payload = {
            "version": 1,
            "source_file": source_name,
            "source_hash": source_hash,
            "mode": mode,
            "total_raw_paragraphs": report.total_raw_paragraphs,
            "total_logical_paragraphs": report.total_logical_paragraphs,
            "merged_group_count": report.merged_group_count,
            "merged_raw_paragraph_count": report.merged_raw_paragraph_count,
            **summarize_boundary_normalization_metrics(report),
            "decisions": [
                {
                    "left_raw_index": decision.left_raw_index,
                    "right_raw_index": decision.right_raw_index,
                    "decision": decision.decision,
                    "confidence": decision.confidence,
                    "reasons": list(decision.reasons),
                }
                for decision in report.decisions
            ],
        }
        artifact_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(artifact_path)
    except Exception:
        return None


def _write_relation_normalization_report_artifact(
    *,
    source_name: str,
    source_bytes: bytes,
    profile: str,
    enabled_relation_kinds: tuple[str, ...],
    report: RelationNormalizationReport,
) -> str | None:
    try:
        RELATION_NORMALIZATION_REPORTS_DIR.mkdir(parents=True, exist_ok=True)
        source_hash = hashlib.sha1(source_bytes).hexdigest()[:8]
        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", source_name or "document.docx").strip("_") or "document.docx"
        artifact_path = RELATION_NORMALIZATION_REPORTS_DIR / f"{safe_name}_{source_hash}.json"
        payload = {
            "version": 1,
            "source_file": source_name,
            "source_hash": source_hash,
            "profile": profile,
            "enabled_relation_kinds": list(enabled_relation_kinds),
            "total_relations": report.total_relations,
            "relation_counts": dict(report.relation_counts),
            "rejected_candidate_count": report.rejected_candidate_count,
            "decisions": [
                {
                    "relation_kind": decision.relation_kind,
                    "decision": decision.decision,
                    "member_paragraph_ids": list(decision.member_paragraph_ids),
                    "anchor_asset_id": decision.anchor_asset_id,
                    "reasons": list(decision.reasons),
                }
                for decision in report.decisions
            ],
        }
        artifact_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(artifact_path)
    except Exception:
        return None


def _iter_document_block_items(document):
    for child in document.element.body.iterchildren():
        local_name = _xml_local_name(child.tag)
        if local_name == "p":
            yield "paragraph", Paragraph(child, document)
        elif local_name == "tbl":
            yield "table", Table(child, document)


def _render_table_html(table: Table, image_assets: list[ImageAsset]) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rendered_row = [_render_table_cell(cell, image_assets) for cell in row.cells]
        rows.append(rendered_row)

    if not any(any(cell.strip() for cell in row) for row in rows):
        return ""

    has_header = len(rows) > 1 and all(cell.strip() for cell in rows[0])
    lines = ["<table>"]
    if has_header:
        lines.append("<thead>")
        lines.append(_render_table_html_row(rows[0], cell_tag="th"))
        lines.append("</thead>")
        body_rows = rows[1:]
    else:
        body_rows = rows

    lines.append("<tbody>")
    for row in body_rows:
        lines.append(_render_table_html_row(row, cell_tag="td"))
    lines.append("</tbody>")
    lines.append("</table>")
    return "\n".join(lines)


def _render_table_cell(cell, image_assets: list[ImageAsset]) -> str:
    cell_parts: list[str] = []
    for paragraph in cell.paragraphs:
        text = _build_paragraph_text_with_placeholders(paragraph, image_assets).strip()
        if text:
            cell_parts.append(_escape_html_preserving_breaks(text))
    return "<br/>".join(cell_parts)


def _render_table_html_row(cells: list[str], *, cell_tag: str) -> str:
    rendered_cells = "".join(f"<{cell_tag}>{cell or '&nbsp;'}</{cell_tag}>" for cell in cells)
    return f"<tr>{rendered_cells}</tr>"


def _escape_html_preserving_breaks(text: str) -> str:
    return "<br/>".join(html.escape(part, quote=False) for part in text.split("<br/>"))


def _extract_run_images(run) -> list[tuple[bytes, str | None, int | None, int | None, dict[str, object]]]:
    return _extract_run_element_images(run._element, run.part)


def _extract_run_element_images(run_element, part) -> list[tuple[bytes, str | None, int | None, int | None, dict[str, object]]]:
    images: list[tuple[bytes, str | None, int | None, int | None, dict[str, object]]] = []
    for drawing in run_element.xpath(".//w:drawing"):
        blips = drawing.xpath(".//a:blip")
        width_emu, height_emu = _resolve_drawing_extent_emu(drawing)
        for blip in blips:
            embed_id = blip.get(f"{{{RELATIONSHIP_NAMESPACE}}}embed")
            if not embed_id:
                continue
            image_part = part.related_parts.get(embed_id)
            if image_part is None:
                continue
            images.append(
                (
                    image_part.blob,
                    getattr(image_part, "content_type", None),
                    width_emu,
                    height_emu,
                    _build_drawing_forensics(drawing, embed_id=embed_id),
                )
            )
    return images


def _resolve_drawing_extent_emu(drawing) -> tuple[int | None, int | None]:
    extents = drawing.xpath(".//wp:extent")
    if not extents:
        return None, None

    extent = extents[0]
    try:
        width_emu = int(extent.get("cx"))
        height_emu = int(extent.get("cy"))
    except (TypeError, ValueError):
        return None, None

    if width_emu <= 0 or height_emu <= 0:
        return None, None
    return width_emu, height_emu


def _build_drawing_forensics(drawing, *, embed_id: str) -> dict[str, object]:
    doc_properties = _resolve_drawing_doc_properties(drawing)
    return {
        "relationship_id": embed_id,
        "drawing_container": _resolve_drawing_container_kind(drawing),
        "drawing_container_xml": _resolve_drawing_container_xml(drawing),
        "source_rect": _resolve_drawing_source_rect(drawing),
        "doc_properties": doc_properties,
    }


def _resolve_drawing_container_kind(drawing) -> str | None:
    if drawing.xpath("./wp:inline"):
        return "inline"
    if drawing.xpath("./wp:anchor"):
        return "anchor"
    return None


def _resolve_drawing_container_xml(drawing) -> str | None:
    containers = drawing.xpath("./wp:inline | ./wp:anchor")
    if not containers:
        return None
    return etree.tostring(containers[0], encoding="unicode")


def _resolve_drawing_source_rect(drawing) -> dict[str, int] | None:
    source_rects = drawing.xpath(".//a:srcRect")
    if not source_rects:
        return None
    source_rect = source_rects[0]
    resolved: dict[str, int] = {}
    for key in ("l", "t", "r", "b"):
        raw_value = source_rect.get(key)
        if raw_value is None:
            continue
        try:
            resolved[key] = int(raw_value)
        except (TypeError, ValueError):
            continue
    return resolved or None


def _resolve_drawing_doc_properties(drawing) -> dict[str, object] | None:
    properties = drawing.xpath(".//wp:docPr")
    if not properties:
        return None
    doc_pr = properties[0]
    payload = {
        "id": doc_pr.get("id"),
        "name": doc_pr.get("name"),
        "descr": doc_pr.get("descr"),
        "title": doc_pr.get("title"),
    }
    return {key: value for key, value in payload.items() if value not in {None, ""}}


def _xml_local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _read_uploaded_docx_bytes(uploaded_file) -> bytes:
    try:
        source_bytes = read_uploaded_file_bytes(uploaded_file)
    except ValueError as exc:
        raise ValueError("Не удалось прочитать содержимое DOCX-файла.") from exc
    if zipfile.is_zipfile(BytesIO(source_bytes)):
        return source_bytes
    normalized_document = normalize_uploaded_document(
        filename=resolve_uploaded_filename(uploaded_file),
        source_bytes=source_bytes,
    )
    return normalized_document.content_bytes


def _extract_explicit_heading_level(paragraph, style_name: str) -> int | None:
    normalized_style = style_name.strip().lower()
    if normalized_style == "title":
        return 1
    if normalized_style == "subtitle":
        return 2

    style_match = HEADING_STYLE_PATTERN.match(normalized_style)
    if style_match is not None:
        level_text = style_match.group(1)
        if level_text:
            try:
                return max(1, min(int(level_text), 6))
            except ValueError:
                return 1
        return 1

    outline_level = _resolve_paragraph_outline_level(paragraph)
    if outline_level is not None:
        return outline_level
    return None


def _resolve_paragraph_outline_level(paragraph) -> int | None:
    outline_element = _find_paragraph_property_element(paragraph, "outlineLvl")
    outline_value = _get_xml_attribute(outline_element, "val") if outline_element is not None else None
    try:
        if outline_value is None:
            return None
        return max(1, min(int(outline_value) + 1, 6))
    except (TypeError, ValueError):
        return None


def _find_paragraph_property_element(paragraph, local_name: str):
    paragraph_properties = _find_child_element(paragraph._element, "pPr")
    element = _find_child_element(paragraph_properties, local_name)
    if element is not None:
        return element

    style = getattr(paragraph, "style", None)
    while style is not None:
        style_properties = _find_child_element(getattr(style, "_element", None), "pPr")
        element = _find_child_element(style_properties, local_name)
        if element is not None:
            return element
        style = getattr(style, "base_style", None)
    return None


def _resolve_paragraph_alignment(paragraph) -> str | None:
    alignment = _find_paragraph_property_element(paragraph, "jc")
    return _get_xml_attribute(alignment, "val") if alignment is not None else None


def _normalize_text_for_heading_heuristics(text: str) -> str:
    normalized = MARKDOWN_LINK_PATTERN.sub(r"\1", text)
    normalized = INLINE_HTML_TAG_PATTERN.sub("", normalized)
    normalized = normalized.replace("**", "").replace("*", "")
    return normalized.strip()


def _infer_heuristic_heading_level(text: str) -> int:
    normalized_text = _normalize_text_for_heading_heuristics(text)
    lower_text = normalized_text.lower()

    if re.match(r"^(?:глава|часть|chapter|part|appendix|приложение)\b", lower_text):
        return 1
    if re.match(r"^(?:раздел|section)\b", lower_text):
        return 2

    numeric_match = re.match(r"^(\d+(?:\.\d+){0,4})(?:[\):]|\s)", normalized_text)
    if numeric_match is not None:
        return min(numeric_match.group(1).count(".") + 2, 6)

    return 2


def _is_probable_heading(paragraph, text: str, normalized_style: str) -> bool:
    stripped_text = _normalize_text_for_heading_heuristics(text)
    if not stripped_text or len(stripped_text) > 140:
        return False
    word_count = len(stripped_text.split())
    if word_count > 18:
        return False
    if stripped_text.endswith(".") and word_count > 4:
        return False
    if stripped_text.count(".") > 1:
        return False
    has_strong_format = _paragraph_has_strong_heading_format(paragraph)
    if normalized_style in {"body text", "normal"} and not has_strong_format:
        return False
    if not has_strong_format:
        return False
    if _is_caption_style(normalized_style):
        return False
    resolved_alignment = _resolve_paragraph_alignment(paragraph)
    if word_count <= 8 and len(stripped_text) <= 100:
        if resolved_alignment == "center" and word_count > 2 and not _has_heading_text_signal(stripped_text):
            return False
        return _has_heading_text_signal(stripped_text) or resolved_alignment == "center"
    return _has_heading_text_signal(stripped_text)


def _has_heading_text_signal(text: str) -> bool:
    stripped_text = _normalize_text_for_heading_heuristics(text)
    word_count = len(stripped_text.split())
    lower_text = stripped_text.lower()

    if re.match(r"^(?:глава|раздел|часть|приложение|chapter|section|appendix)\b", lower_text):
        return True
    if re.match(r"^\d+(?:\.\d+){0,4}(?:[\):]|\s)", stripped_text):
        return True
    if ":" in stripped_text and word_count <= 12 and not stripped_text.endswith("."):
        return True
    return False


def _paragraph_is_effectively_bold(paragraph) -> bool:
    alignment_value = _resolve_paragraph_alignment(paragraph)
    visible_runs = [run for run in paragraph.runs if run.text and run.text.strip()]
    if not visible_runs:
        return False

    bold_runs = [run for run in visible_runs if bool(run.bold)]
    if len(bold_runs) == len(visible_runs):
        return True

    visible_chars = sum(len(run.text.strip()) for run in visible_runs)
    bold_chars = sum(len(run.text.strip()) for run in bold_runs)
    return bool(bold_runs) and visible_chars > 0 and (bold_chars / visible_chars) >= 0.5


def paragraph_has_strong_heading_format(paragraph) -> bool:
    alignment_value = _resolve_paragraph_alignment(paragraph)
    if alignment_value == "center":
        return True
    return _paragraph_is_effectively_bold(paragraph)


def _paragraph_has_strong_heading_format(paragraph) -> bool:
    return paragraph_has_strong_heading_format(paragraph)


def _paragraph_unit_has_strong_heading_format(paragraph: ParagraphUnit) -> bool:
    return paragraph.paragraph_alignment == "center" or paragraph.is_bold


def _is_image_only_text(text: str) -> bool:
    return IMAGE_ONLY_PATTERN.fullmatch(text.strip()) is not None


def _extract_paragraph_asset_id(text: str, *, role: str) -> str | None:
    if role != "image":
        return None
    placeholders = IMAGE_PLACEHOLDER_PATTERN.findall(text)
    if len(placeholders) != 1:
        return None
    placeholder = placeholders[0]
    match = re.match(r"\[\[DOCX_IMAGE_(img_\d+)\]\]", placeholder)
    if match is None:
        return None
    return match.group(1)


def _is_caption_style(normalized_style: str) -> bool:
    return normalized_style in {"caption", "подпись"} or "caption" in normalized_style or "подпись" in normalized_style


def _is_likely_caption_text(text: str) -> bool:
    stripped_text = text.strip()
    if not stripped_text or len(stripped_text) > 140:
        return False
    return CAPTION_PREFIX_PATTERN.match(stripped_text) is not None


def _is_short_standalone_heading_text(text: str) -> bool:
    stripped_text = _normalize_text_for_heading_heuristics(text)
    if not stripped_text or len(stripped_text) > 80:
        return False
    if _is_likely_caption_text(stripped_text):
        return False
    word_count = len(stripped_text.split())
    if word_count == 0 or word_count > 6:
        return False
    if stripped_text.endswith((".", "?", "!", ";")):
        return False
    if stripped_text.count(".") > 0:
        return False
    return True


def _is_very_short_standalone_heading_text(text: str) -> bool:
    stripped_text = _normalize_text_for_heading_heuristics(text)
    if not _is_short_standalone_heading_text(stripped_text):
        return False
    word_count = len(stripped_text.split())
    return word_count <= 4 and len(stripped_text) <= 48


def _has_body_context_signal(text: str) -> bool:
    stripped_text = _normalize_text_for_heading_heuristics(text)
    word_count = len(stripped_text.split())
    if word_count >= 8:
        return True
    if len(stripped_text) >= 60:
        return True
    return any(marker in stripped_text for marker in (",", ":")) and word_count >= 5


def _length_to_points(length) -> float | None:
    if length is None:
        return None
    points = getattr(length, "pt", None)
    if points is None:
        return None
    try:
        return float(points)
    except (TypeError, ValueError):
        return None


def _resolve_style_font_size(style) -> float | None:
    while style is not None:
        font = getattr(style, "font", None)
        points = _length_to_points(getattr(font, "size", None))
        if points is not None:
            return points
        style = getattr(style, "base_style", None)
    return None


def resolve_effective_paragraph_font_size(paragraph) -> float | None:
    weighted_sizes: dict[float, int] = {}
    for run in paragraph.runs:
        text = run.text.strip()
        if not text:
            continue
        points = _length_to_points(getattr(getattr(run, "font", None), "size", None))
        if points is None:
            points = _resolve_style_font_size(getattr(run, "style", None))
        if points is None:
            continue
        normalized_points = round(points, 2)
        weighted_sizes[normalized_points] = weighted_sizes.get(normalized_points, 0) + len(text)

    if weighted_sizes:
        return max(weighted_sizes.items(), key=lambda item: (item[1], item[0]))[0]
    return _resolve_style_font_size(getattr(paragraph, "style", None))


def _resolve_effective_paragraph_font_size(paragraph) -> float | None:
    return resolve_effective_paragraph_font_size(paragraph)


def _paragraph_is_effectively_italic(paragraph) -> bool:
    visible_runs = [run for run in paragraph.runs if run.text and run.text.strip()]
    if not visible_runs:
        return False

    italic_runs = [run for run in visible_runs if bool(run.italic)]
    if len(italic_runs) == len(visible_runs):
        return True

    visible_chars = sum(len(run.text.strip()) for run in visible_runs)
    italic_chars = sum(len(run.text.strip()) for run in italic_runs)
    return bool(italic_runs) and visible_chars > 0 and (italic_chars / visible_chars) >= 0.5


def _infer_contextual_heading_level(paragraphs: list[ParagraphUnit], index: int) -> int:
    for previous_index in range(index - 1, -1, -1):
        previous_paragraph = paragraphs[previous_index]
        if previous_paragraph.role != "heading" or previous_paragraph.heading_level is None:
            continue
        if previous_paragraph.heading_level <= 1:
            return 2
        return previous_paragraph.heading_level
    return 2


def _promote_short_standalone_headings(paragraphs: list[ParagraphUnit]) -> None:
    if len(paragraphs) < 3:
        return

    for index in range(1, len(paragraphs) - 1):
        paragraph = paragraphs[index]
        if paragraph.role != "body":
            continue
        if paragraph.role_confidence == "ai":
            continue
        if not _is_short_standalone_heading_text(paragraph.text):
            continue

        previous_paragraph = paragraphs[index - 1]
        next_paragraph = paragraphs[index + 1]
        if previous_paragraph.role != "body" or next_paragraph.role != "body":
            continue
        if not _has_body_context_signal(previous_paragraph.text) or not _has_body_context_signal(next_paragraph.text):
            continue

        if _is_very_short_standalone_heading_text(paragraph.text):
            paragraph.role = "heading"
            paragraph.structural_role = "heading"
            paragraph.role_confidence = "heuristic"
            paragraph.heading_source = "heuristic"
            paragraph.heading_level = _infer_contextual_heading_level(paragraphs, index)
            continue

        candidate_font_size = paragraph.font_size_pt
        if candidate_font_size is None:
            continue

        context_font_sizes: list[float] = []
        previous_font_size = paragraphs[index - 1].font_size_pt
        if previous_font_size is not None:
            context_font_sizes.append(previous_font_size)
        next_font_size = paragraphs[index + 1].font_size_pt
        if next_font_size is not None:
            context_font_sizes.append(next_font_size)
        if not context_font_sizes:
            continue

        required_delta = 1.0 if _paragraph_unit_has_strong_heading_format(paragraph) else 1.5
        if candidate_font_size < max(context_font_sizes) + required_delta:
            continue

        paragraph.role = "heading"
        paragraph.structural_role = "heading"
        paragraph.role_confidence = "heuristic"
        paragraph.heading_source = "heuristic"
        paragraph.heading_level = _infer_contextual_heading_level(paragraphs, index)


def _reclassify_adjacent_captions(paragraphs: list[ParagraphUnit]) -> None:
    for index, paragraph in enumerate(paragraphs):
        if index == 0:
            continue
        previous_paragraph = paragraphs[index - 1]
        if previous_paragraph.role not in {"image", "table"}:
            continue
        if paragraph.role == "caption":
            continue
        if _is_likely_caption_text(paragraph.text):
            if paragraph.role == "heading" and paragraph.heading_source != "heuristic":
                continue
            paragraph.role = "caption"
            paragraph.structural_role = "caption"
            paragraph.role_confidence = "adjacent"
            paragraph.heading_level = None
            paragraph.heading_source = None


def _render_hyperlink_element(hyperlink_element, paragraph, image_assets: list[ImageAsset]) -> str:
    text_parts: list[str] = []
    for child in hyperlink_element:
        if _xml_local_name(child.tag) != "r":
            continue
        text_parts.append(_render_run_element(child, paragraph.part, image_assets, allow_hyperlink_markdown=False))

    text = "".join(text_parts)
    if not text.strip():
        return text

    relationship_id = hyperlink_element.get(f"{{{RELATIONSHIP_NAMESPACE}}}id")
    if not relationship_id:
        return text

    relationship = paragraph.part.rels.get(relationship_id)
    url = getattr(relationship, "target_ref", None)
    if not url:
        return text
    return f"[{text}]({url})"


def _render_run_element(run_element, part, image_assets: list[ImageAsset], *, allow_hyperlink_markdown: bool = True) -> str:
    text = _extract_run_text(run_element)
    formatted_text = _apply_run_markdown(text, run_element) if allow_hyperlink_markdown else text
    image_placeholders = _extract_run_image_placeholders(run_element, part, image_assets)
    return formatted_text + "".join(image_placeholders)


def _extract_run_text(run_element) -> str:
    text_parts: list[str] = []
    for child in run_element:
        local_name = _xml_local_name(child.tag)
        if local_name in {"t", "delText", "instrText"}:
            text_parts.append(child.text or "")
            continue
        if local_name == "tab":
            text_parts.append("\t")
            continue
        if local_name in {"br", "cr"}:
            text_parts.append("<br/>")
    return "".join(text_parts)


def _apply_run_markdown(text: str, run_element) -> str:
    if not text:
        return text

    run_properties = _find_child_element(run_element, "rPr")
    if run_properties is None:
        return text

    is_bold = _find_child_element(run_properties, "b") is not None
    is_italic = _find_child_element(run_properties, "i") is not None
    is_underline = _find_child_element(run_properties, "u") is not None
    vertical_align = _extract_vertical_align(run_properties)

    formatted = text
    if is_bold and is_italic:
        formatted = f"***{formatted}***"
    elif is_bold:
        formatted = f"**{formatted}**"
    elif is_italic:
        formatted = f"*{formatted}*"

    if is_underline:
        formatted = f"<u>{formatted}</u>"
    if vertical_align == "superscript":
        formatted = f"<sup>{formatted}</sup>"
    elif vertical_align == "subscript":
        formatted = f"<sub>{formatted}</sub>"
    return formatted


def _extract_vertical_align(run_properties) -> str | None:
    vertical_align = _find_child_element(run_properties, "vertAlign")
    return _get_xml_attribute(vertical_align, "val") if vertical_align is not None else None


def _extract_run_image_placeholders(run_element, part, image_assets: list[ImageAsset]) -> list[str]:
    placeholders: list[str] = []
    for image_blob, mime_type, width_emu, height_emu, source_forensics in _extract_run_element_images(run_element, part):
        image_index = len(image_assets) + 1
        placeholder = f"[[DOCX_IMAGE_img_{image_index:03d}]]"
        image_assets.append(
            ImageAsset(
                image_id=f"img_{image_index:03d}",
                placeholder=placeholder,
                original_bytes=image_blob,
                mime_type=mime_type,
                position_index=image_index - 1,
                width_emu=width_emu,
                height_emu=height_emu,
                source_forensics=source_forensics,
            )
        )
        placeholders.append(placeholder)
    return placeholders


_TYPOGRAPHIC_BULLET_CHARS = {"\u2014", "\u2013"}  # em-dash, en-dash


def _empty_list_metadata() -> dict[str, object]:
    return {
        "list_kind": None,
        "list_level": 0,
        "list_numbering_format": None,
        "list_num_id": None,
        "list_abstract_num_id": None,
        "list_num_xml": None,
        "list_abstract_num_xml": None,
    }


def _is_typographic_emdash_bullet(numbering_details: dict[str, str | None]) -> bool:
    return (
        numbering_details.get("num_format") == "bullet"
        and (numbering_details.get("lvl_text") or "") in _TYPOGRAPHIC_BULLET_CHARS
    )


def _extract_paragraph_list_metadata(paragraph, text: str, style_name: str, role: str) -> dict[str, object]:
    metadata: dict[str, object] = _empty_list_metadata()

    explicit_kind = _detect_explicit_list_kind(text)
    style_level = _extract_style_list_level(style_name)
    num_pr = _resolve_paragraph_num_pr(paragraph)

    if role != "list" and explicit_kind is None and num_pr is None:
        return metadata

    if explicit_kind is not None:
        metadata["list_kind"] = explicit_kind
        # Still try to capture Word numbering XML so DOCX list restoration works
        # even when the source paragraph already has visible text markers in its text.
        # Per spec: numbered lists must be restored as real Word lists even if visible
        # markdown markers are present.
        if num_pr is not None:
            numbering_details = _resolve_num_pr_details(paragraph, num_pr)
            if _is_typographic_emdash_bullet(numbering_details):
                return _empty_list_metadata()
            numbering_format = numbering_details["num_format"]
            metadata["list_level"] = _extract_num_pr_level(num_pr)
            metadata["list_numbering_format"] = numbering_format
            metadata["list_num_id"] = numbering_details["num_id"]
            metadata["list_abstract_num_id"] = numbering_details["abstract_num_id"]
            metadata["list_num_xml"] = numbering_details["num_xml"]
            metadata["list_abstract_num_xml"] = numbering_details["abstract_num_xml"]
            if numbering_format in ORDERED_LIST_FORMATS:
                metadata["list_kind"] = "ordered"
            elif numbering_format in UNORDERED_LIST_FORMATS:
                metadata["list_kind"] = "unordered"
        return metadata

    if num_pr is not None:
        list_level = max(_extract_num_pr_level(num_pr), style_level)
        numbering_details = _resolve_num_pr_details(paragraph, num_pr)
        if _is_typographic_emdash_bullet(numbering_details):
            metadata["_is_typographic_emdash_bullet"] = True
            return metadata
        numbering_format = numbering_details["num_format"]
        metadata["list_level"] = list_level
        metadata["list_numbering_format"] = numbering_format
        metadata["list_num_id"] = numbering_details["num_id"]
        metadata["list_abstract_num_id"] = numbering_details["abstract_num_id"]
        metadata["list_num_xml"] = numbering_details["num_xml"]
        metadata["list_abstract_num_xml"] = numbering_details["abstract_num_xml"]
        if numbering_format in ORDERED_LIST_FORMATS:
            metadata["list_kind"] = "ordered"
            return metadata
        if numbering_format in UNORDERED_LIST_FORMATS:
            metadata["list_kind"] = "unordered"
            return metadata
        if role != "list":
            return metadata

    if role != "list":
        return metadata

    normalized_style = style_name.strip().lower()
    if any(token in normalized_style for token in ("number", "num", "нумер", "числ")):
        metadata["list_kind"] = "ordered"
        metadata["list_level"] = style_level
        return metadata
    if any(token in normalized_style for token in ("bullet", "bulleted", "маркир", "маркер")):
        metadata["list_kind"] = "unordered"
        metadata["list_level"] = style_level
        return metadata
    metadata["list_kind"] = "unordered"
    metadata["list_level"] = style_level
    return metadata


def _detect_explicit_list_kind(text: str) -> str | None:
    stripped_text = text.lstrip()
    if stripped_text.startswith(("- ", "* ", "• ")):
        return "unordered"
    if re.match(r"^\d+[\.)]\s+", stripped_text):
        return "ordered"
    return None


def _extract_style_list_level(style_name: str) -> int:
    match = re.search(r"(\d+)\s*$", style_name.strip())
    if match is None:
        return 0
    try:
        return max(0, int(match.group(1)) - 1)
    except ValueError:
        return 0


def _resolve_paragraph_num_pr(paragraph):
    paragraph_properties = _find_child_element(paragraph._element, "pPr")
    num_pr = _find_child_element(paragraph_properties, "numPr")
    if num_pr is not None:
        return num_pr

    style = getattr(paragraph, "style", None)
    while style is not None:
        style_properties = _find_child_element(getattr(style, "_element", None), "pPr")
        num_pr = _find_child_element(style_properties, "numPr")
        if num_pr is not None:
            return num_pr
        style = getattr(style, "base_style", None)
    return None


def _extract_num_pr_level(num_pr) -> int:
    ilvl = _find_child_element(num_pr, "ilvl")
    level_value = _get_xml_attribute(ilvl, "val") if ilvl is not None else None
    if level_value is None:
        return 0
    try:
        return max(0, int(level_value))
    except (TypeError, ValueError):
        return 0


def _resolve_num_pr_details(paragraph, num_pr) -> dict[str, str | None]:
    num_id_element = _find_child_element(num_pr, "numId")
    ilvl_element = _find_child_element(num_pr, "ilvl")
    num_id = _get_xml_attribute(num_id_element, "val") if num_id_element is not None else None
    ilvl = _get_xml_attribute(ilvl_element, "val") if ilvl_element is not None else "0"
    if num_id is None:
        return {
            "num_id": None,
            "abstract_num_id": None,
            "num_format": None,
            "num_xml": None,
            "abstract_num_xml": None,
        }

    numbering_part = getattr(paragraph.part, "numbering_part", None)
    numbering_root = getattr(numbering_part, "element", None)
    if numbering_root is None:
        return {
            "num_id": num_id,
            "abstract_num_id": None,
            "num_format": None,
            "num_xml": None,
            "abstract_num_xml": None,
        }

    abstract_num_id = None
    num_xml = None
    for child in numbering_root:
        if _xml_local_name(child.tag) != "num":
            continue
        if _get_xml_attribute(child, "numId") != num_id:
            continue
        abstract_num = _find_child_element(child, "abstractNumId")
        abstract_num_id = _get_xml_attribute(abstract_num, "val") if abstract_num is not None else None
        num_xml = etree.tostring(child, encoding="unicode")
        break

    if abstract_num_id is None:
        return {
            "num_id": num_id,
            "abstract_num_id": None,
            "num_format": None,
            "num_xml": num_xml,
            "abstract_num_xml": None,
        }

    for child in numbering_root:
        if _xml_local_name(child.tag) != "abstractNum":
            continue
        if _get_xml_attribute(child, "abstractNumId") != abstract_num_id:
            continue
        abstract_num_xml = etree.tostring(child, encoding="unicode")
        for level in child:
            if _xml_local_name(level.tag) != "lvl":
                continue
            if _get_xml_attribute(level, "ilvl") != ilvl:
                continue
            num_format = _find_child_element(level, "numFmt")
            lvl_text = _find_child_element(level, "lvlText")
            return {
                "num_id": num_id,
                "abstract_num_id": abstract_num_id,
                "num_format": _get_xml_attribute(num_format, "val") if num_format is not None else None,
                "lvl_text": _get_xml_attribute(lvl_text, "val") if lvl_text is not None else None,
                "num_xml": num_xml,
                "abstract_num_xml": abstract_num_xml,
            }
        return {
            "num_id": num_id,
            "abstract_num_id": abstract_num_id,
            "num_format": None,
            "num_xml": num_xml,
            "abstract_num_xml": abstract_num_xml,
        }
    return {
        "num_id": num_id,
        "abstract_num_id": abstract_num_id,
        "num_format": None,
        "num_xml": num_xml,
        "abstract_num_xml": None,
    }


def _find_child_element(parent, local_name: str):
    if parent is None:
        return None
    for child in parent:
        if _xml_local_name(child.tag) == local_name:
            return child
    return None


def _get_xml_attribute(element, attribute_name: str) -> str | None:
    if element is None:
        return None
    for key, value in element.attrib.items():
        if _xml_local_name(key) == attribute_name:
            return value
    return None


xml_local_name = _xml_local_name
resolve_paragraph_outline_level = _resolve_paragraph_outline_level
infer_heuristic_heading_level = _infer_heuristic_heading_level
is_image_only_text = _is_image_only_text
is_likely_caption_text = _is_likely_caption_text
detect_explicit_list_kind = _detect_explicit_list_kind
find_child_element = _find_child_element
get_xml_attribute = _get_xml_attribute


def _validate_docx_archive(source_bytes: bytes) -> None:
    if len(source_bytes) > MAX_DOCX_ARCHIVE_SIZE_BYTES:
        raise RuntimeError("DOCX-файл превышает допустимый размер архива.")

    try:
        with zipfile.ZipFile(BytesIO(source_bytes)) as archive:
            entries = archive.infolist()
    except zipfile.BadZipFile as exc:
        raise RuntimeError("Передан поврежденный или неподдерживаемый DOCX-архив.") from exc

    if not entries:
        raise RuntimeError("Передан пустой DOCX-архив.")
    if len(entries) > MAX_DOCX_ENTRY_COUNT:
        raise RuntimeError("DOCX-архив содержит слишком много файлов и отклонен из соображений безопасности.")

    total_uncompressed_size = sum(max(0, entry.file_size) for entry in entries)
    total_compressed_size = sum(max(0, entry.compress_size) for entry in entries)
    if total_uncompressed_size > MAX_DOCX_UNCOMPRESSED_SIZE_BYTES:
        raise RuntimeError("DOCX-архив слишком велик после распаковки и отклонен из соображений безопасности.")
    if total_compressed_size > 0 and (total_uncompressed_size / total_compressed_size) > MAX_DOCX_COMPRESSION_RATIO:
        raise RuntimeError("DOCX-архив имеет подозрительно высокий коэффициент сжатия и отклонен из соображений безопасности.")

    for entry in entries:
        entry_name = entry.filename
        parts = entry_name.replace("\\", "/").split("/")
        if any(part == ".." for part in parts):
            raise RuntimeError("DOCX-архив содержит подозрительные пути и отклонён из соображений безопасности.")
        if entry_name.startswith("/"):
            raise RuntimeError("DOCX-архив содержит абсолютные пути и отклонён из соображений безопасности.")

    filenames = {entry.filename for entry in entries}
    if "[Content_Types].xml" not in filenames:
        raise RuntimeError("Передан невалидный DOCX-архив: отсутствует [Content_Types].xml.")


def validate_docx_source_bytes(source_bytes: bytes) -> None:
    _validate_docx_archive(source_bytes)
